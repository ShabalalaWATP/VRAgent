"""
Combined Analysis Router

API endpoints for the Combined Results Analysis Report feature.
"""

import io
import re
from typing import List, Optional
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from backend.core.database import get_db
from backend.core.auth import get_current_user
from backend.core.logging import get_logger
from backend.models.models import User, Conversation, ConversationParticipant, Message, Project
from backend.schemas.combined_analysis import (
    CombinedAnalysisRequest,
    AvailableScansResponse,
    CombinedAnalysisReportResponse,
    CombinedAnalysisListItem,
    CombinedAnalysisListResponse,
    ReportSection,
    CrossAnalysisFinding,
    ExploitDevelopmentArea,
    SelectedScan,
    CombinedAnalysisChatRequest,
    CombinedAnalysisChatResponse,
)
from backend.services import combined_analysis_service


def parse_markdown_for_word(content: str, doc, Pt):
    """
    Parse markdown content and add to Word document with proper formatting.
    Handles code blocks, headings, bold, lists, etc.
    """
    # Split content by code blocks
    code_block_pattern = r'```(\w*)\n(.*?)```'
    parts = re.split(code_block_pattern, content, flags=re.DOTALL)
    
    i = 0
    while i < len(parts):
        if i + 2 < len(parts) and parts[i+1] and parts[i+2]:
            # Text before code block
            text_before = parts[i].strip()
            if text_before:
                _add_markdown_text_to_word(text_before, doc)
            
            # Code block
            lang = parts[i + 1]
            code = parts[i + 2]
            _add_code_block_to_word(code, lang, doc, Pt)
            i += 3
        else:
            # Just text, no code block
            text = parts[i].strip()
            if text:
                _add_markdown_text_to_word(text, doc)
            i += 1


def _add_markdown_text_to_word(text: str, doc):
    """Add markdown-formatted text to Word document."""
    lines = text.split('\n')
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue
        
        # Handle headers - these become Word headings
        if stripped_line.startswith('#### '):
            doc.add_heading(stripped_line[5:], level=4)
        elif stripped_line.startswith('### '):
            doc.add_heading(stripped_line[4:], level=3)
        elif stripped_line.startswith('## '):
            doc.add_heading(stripped_line[3:], level=2)
        elif stripped_line.startswith('# '):
            doc.add_heading(stripped_line[2:], level=1)
        # Handle bullet points
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
            # Remove markdown bold from bullet content
            content = stripped_line[2:]
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            para = doc.add_paragraph(content, style='List Bullet')
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', stripped_line):
            content = re.sub(r'^\d+\.\s', '', stripped_line)
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            para = doc.add_paragraph(content, style='List Number')
        else:
            # Regular paragraph - handle bold and inline code
            para = doc.add_paragraph()
            # Process inline formatting
            remaining = stripped_line
            while remaining:
                # Look for bold text (**text**)
                bold_match = re.search(r'\*\*(.*?)\*\*', remaining)
                # Look for inline code (`code`)
                code_match = re.search(r'`([^`]+)`', remaining)
                
                # Find earliest match
                if bold_match and (not code_match or bold_match.start() < code_match.start()):
                    # Add text before bold
                    if bold_match.start() > 0:
                        para.add_run(remaining[:bold_match.start()])
                    # Add bold text
                    para.add_run(bold_match.group(1)).bold = True
                    remaining = remaining[bold_match.end():]
                elif code_match and (not bold_match or code_match.start() < bold_match.start()):
                    # Add text before code
                    if code_match.start() > 0:
                        para.add_run(remaining[:code_match.start()])
                    # Add code text with monospace font
                    from docx.shared import Pt
                    code_run = para.add_run(code_match.group(1))
                    code_run.font.name = "Courier New"
                    code_run.font.size = Pt(9)
                    remaining = remaining[code_match.end():]
                else:
                    # No more formatting, add rest as plain text
                    para.add_run(remaining)
                    remaining = ""


def _add_code_block_to_word(code: str, lang: str, doc, Pt):
    """Add a formatted code block to Word document."""
    from docx.shared import RGBColor
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    # Add language label if present
    if lang:
        label_para = doc.add_paragraph()
        label_run = label_para.add_run(f"[{lang.upper()}]")
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add code with monospace font and background
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Pt(20)
    
    # Add shading to paragraph
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E1E"/>')
    code_para._p.get_or_add_pPr().append(shading_elm)
    
    code_run = code_para.add_run(code.strip())
    code_run.font.name = "Courier New"
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = RGBColor(212, 212, 212)  # Light gray text


def parse_markdown_for_pdf(content: str, story, styles):
    """
    Parse markdown content and add to PDF with proper formatting.
    """
    # Split content by code blocks
    code_block_pattern = r'```(\w*)\n(.*?)```'
    parts = re.split(code_block_pattern, content, flags=re.DOTALL)
    
    i = 0
    while i < len(parts):
        if i + 2 < len(parts) and parts[i+1] is not None and parts[i+2] is not None:
            # Text before code block
            text_before = parts[i].strip()
            if text_before:
                _add_markdown_text_to_pdf(text_before, story, styles)
            
            # Code block
            lang = parts[i + 1]
            code = parts[i + 2]
            _add_code_block_to_pdf(code, lang, story, styles)
            i += 3
        else:
            # Just text, no code block
            text = parts[i].strip()
            if text:
                _add_markdown_text_to_pdf(text, story, styles)
            i += 1


def _add_markdown_text_to_pdf(text: str, story, styles):
    """Add markdown-formatted text to PDF."""
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.units import inch
    
    lines = text.split('\n')
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            story.append(Spacer(1, 0.05 * inch))
            continue
        
        # Escape HTML special characters first
        line_escaped = stripped_line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Handle headers - strip markdown and use heading styles
        if stripped_line.startswith('#### '):
            clean_text = line_escaped[5:]
            # Remove any remaining markdown bold
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            story.append(Paragraph(clean_text, styles['Heading3Custom']))
        elif stripped_line.startswith('### '):
            clean_text = line_escaped[4:]
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            story.append(Paragraph(clean_text, styles['Heading3Custom']))
        elif stripped_line.startswith('## '):
            clean_text = line_escaped[3:]
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            story.append(Paragraph(clean_text, styles['Heading2Custom']))
        elif stripped_line.startswith('# '):
            clean_text = line_escaped[2:]
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            story.append(Paragraph(clean_text, styles['Heading1Custom']))
        # Handle bullet points
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
            content = line_escaped[2:]
            # Convert markdown bold to HTML bold
            formatted = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
            # Convert inline code to styled text
            formatted = re.sub(r'`([^`]+)`', r'<font face="Courier" size="9">\1</font>', formatted)
            story.append(Paragraph(f"• {formatted}", styles['BodyCustom']))
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', stripped_line):
            # Extract number and content
            match = re.match(r'^(\d+)\.\s(.*)$', line_escaped)
            if match:
                num = match.group(1)
                content = match.group(2)
                formatted = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
                formatted = re.sub(r'`([^`]+)`', r'<font face="Courier" size="9">\1</font>', formatted)
                story.append(Paragraph(f"{num}. {formatted}", styles['BodyCustom']))
        else:
            # Regular paragraph - convert markdown to HTML
            formatted = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line_escaped)
            # Convert inline code to styled text
            formatted = re.sub(r'`([^`]+)`', r'<font face="Courier" size="9">\1</font>', formatted)
            story.append(Paragraph(formatted, styles['BodyCustom']))


def _add_code_block_to_pdf(code: str, lang: str, story, styles):
    """Add a formatted code block to PDF."""
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.units import inch
    
    # Add language label
    if lang:
        story.append(Paragraph(f"<font color='#666666' size='8'>[{lang.upper()}]</font>", styles['MetaStyle']))
    
    # Format code for PDF - escape special characters
    code_escaped = code.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
    story.append(Paragraph(code_escaped, styles['CodeStyle']))
    story.append(Spacer(1, 0.1 * inch))

logger = get_logger(__name__)

router = APIRouter()


@router.get("/projects/{project_id}/available-scans", response_model=AvailableScansResponse)
async def get_available_scans(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    Get all available scans/reports for a project that can be included in combined analysis.
    
    Returns:
        AvailableScansResponse: Lists of available security scans, network reports,
                                RE reports, and fuzzing sessions.
    """
    try:
        return combined_analysis_service.get_available_scans(db, project_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error getting available scans: {e}")
        raise HTTPException(status_code=500, detail=f"Error getting available scans: {str(e)}")


@router.post("/projects/{project_id}/generate", response_model=CombinedAnalysisReportResponse)
async def generate_combined_analysis(
    project_id: int,
    request: CombinedAnalysisRequest,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    Generate a comprehensive combined analysis report.
    
    This endpoint aggregates data from selected security scans, network reports,
    RE reports, and fuzzing sessions to generate a cross-domain security analysis.
    
    Args:
        project_id: The project ID
        request: CombinedAnalysisRequest with selected scans and options
        
    Returns:
        CombinedAnalysisReportResponse: The generated comprehensive report
    """
    if request.project_id != project_id:
        raise HTTPException(status_code=400, detail="Project ID mismatch")
    
    if not request.selected_scans:
        raise HTTPException(status_code=400, detail="At least one scan must be selected")
    
    try:
        user_id = current_user.id if current_user else None
        report = await combined_analysis_service.generate_combined_analysis(
            db, request, user_id
        )
        
        # Convert to response model
        return _report_to_response(report)
        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error generating combined analysis: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating combined analysis: {str(e)}")


@router.get("/projects/{project_id}/reports", response_model=CombinedAnalysisListResponse)
async def list_combined_analysis_reports(
    project_id: int,
    limit: int = 50,
    offset: int = 0,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    List all combined analysis reports for a project.
    
    Returns:
        CombinedAnalysisListResponse: List of report summaries
    """
    try:
        reports, total = combined_analysis_service.list_combined_analysis_reports(
            db, project_id, limit, offset
        )
        
        items = []
        for r in reports:
            items.append(CombinedAnalysisListItem(
                id=r.id,
                project_id=r.project_id,
                title=r.title,
                created_at=r.created_at,
                overall_risk_level=r.overall_risk_level or "Unknown",
                overall_risk_score=r.overall_risk_score or 0,
                total_findings_analyzed=r.total_findings_analyzed or 0,
                scans_included=r.scans_included or 0,
            ))
        
        return CombinedAnalysisListResponse(reports=items, total=total)
        
    except Exception as e:
        logger.error(f"Error listing combined analysis reports: {e}")
        raise HTTPException(status_code=500, detail=f"Error listing reports: {str(e)}")


@router.get("/reports/{report_id}", response_model=CombinedAnalysisReportResponse)
async def get_combined_analysis_report(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    Get a specific combined analysis report.
    
    Returns:
        CombinedAnalysisReportResponse: The full report
    """
    report = combined_analysis_service.get_combined_analysis_report(db, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return _report_to_response(report)


@router.delete("/reports/{report_id}")
async def delete_combined_analysis_report(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    Delete a combined analysis report.
    
    Returns:
        dict: Status message
    """
    success = combined_analysis_service.delete_combined_analysis_report(db, report_id)
    if not success:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return {"status": "deleted", "id": report_id}


def _report_to_response(report) -> CombinedAnalysisReportResponse:
    """Convert database model to response schema."""
    
    # Convert sections
    sections = []
    if report.report_sections:
        for s in report.report_sections:
            sections.append(ReportSection(
                title=s.get("title", ""),
                content=s.get("content", ""),
                section_type=s.get("section_type", "text"),
                severity=s.get("severity"),
                metadata=s.get("metadata"),
            ))
    
    # Convert cross-analysis findings
    cross_findings = []
    if report.cross_analysis_findings:
        for cf in report.cross_analysis_findings:
            cross_findings.append(CrossAnalysisFinding(
                title=cf.get("title", ""),
                description=cf.get("description", ""),
                severity=cf.get("severity", "Medium"),
                sources=cf.get("sources", []),
                source_details=cf.get("source_details"),
                exploitability_score=cf.get("exploitability_score"),
                exploit_guidance=cf.get("exploit_guidance"),
                remediation=cf.get("remediation"),
            ))
    
    # Convert exploit areas
    exploit_areas = []
    if report.exploit_development_areas:
        for ea in report.exploit_development_areas:
            exploit_areas.append(ExploitDevelopmentArea(
                title=ea.get("title", ""),
                description=ea.get("description", ""),
                vulnerability_chain=ea.get("vulnerability_chain", []),
                attack_vector=ea.get("attack_vector", "Network"),
                complexity=ea.get("complexity", "Medium"),
                impact=ea.get("impact", ""),
                prerequisites=ea.get("prerequisites"),
                poc_guidance=ea.get("poc_guidance"),
            ))
    
    # Convert selected scans
    selected_scans = []
    if report.selected_scans:
        for ss in report.selected_scans:
            selected_scans.append(SelectedScan(
                scan_type=ss.get("scan_type", ""),
                scan_id=ss.get("scan_id", 0),
                title=ss.get("title"),
            ))
    
    return CombinedAnalysisReportResponse(
        id=report.id,
        project_id=report.project_id,
        title=report.title,
        created_at=report.created_at,
        executive_summary=report.executive_summary or "",
        overall_risk_level=report.overall_risk_level or "Unknown",
        overall_risk_score=report.overall_risk_score or 0,
        risk_justification=report.risk_justification or "",
        total_findings_analyzed=report.total_findings_analyzed or 0,
        scans_included=report.scans_included or 0,
        scan_types_breakdown=report.scan_types_breakdown or {},
        sections=sections,
        cross_analysis_findings=cross_findings,
        attack_surface_diagram=report.attack_surface_diagram,
        exploit_development_areas=exploit_areas,
        prioritized_vulnerabilities=report.prioritized_vulnerabilities,
        included_scans=selected_scans,
    )


# ============================================================================
# Export Endpoints
# ============================================================================

def _generate_markdown_report(report) -> str:
    """Generate a formatted Markdown report from the combined analysis."""
    md = []
    
    # Title and metadata
    md.append(f"# {report.title}")
    md.append("")
    md.append(f"**Generated:** {report.created_at.strftime('%Y-%m-%d %H:%M:%S') if hasattr(report.created_at, 'strftime') else report.created_at}")
    md.append(f"**Overall Risk Level:** {report.overall_risk_level or 'Unknown'}")
    md.append(f"**Risk Score:** {report.overall_risk_score or 0}/100")
    md.append(f"**Total Findings Analyzed:** {report.total_findings_analyzed or 0}")
    md.append(f"**Scans Included:** {report.scans_included or 0}")
    md.append("")
    md.append("---")
    md.append("")
    
    # Executive Summary
    md.append("## Executive Summary")
    md.append("")
    md.append(report.executive_summary or "_No executive summary available._")
    md.append("")
    
    # Risk Justification
    if report.risk_justification:
        md.append("## Risk Justification")
        md.append("")
        md.append(report.risk_justification)
        md.append("")
    
    # Report Sections
    if report.report_sections:
        md.append("---")
        md.append("")
        md.append("# Detailed Analysis")
        md.append("")
        for section in report.report_sections:
            severity = section.get("severity", "")
            severity_badge = f" [{severity.upper()}]" if severity else ""
            md.append(f"## {section.get('title', 'Section')}{severity_badge}")
            md.append("")
            content = section.get("content", "")
            if section.get("section_type") == "code":
                md.append(f"```\n{content}\n```")
            else:
                md.append(content)
            md.append("")
    
    # Cross-Analysis Findings
    if report.cross_analysis_findings:
        md.append("---")
        md.append("")
        md.append("# Cross-Analysis Findings")
        md.append("")
        for i, finding in enumerate(report.cross_analysis_findings, 1):
            severity = finding.get("severity", "Medium")
            md.append(f"## {i}. {finding.get('title', 'Finding')} [{severity.upper()}]")
            md.append("")
            md.append(f"**Severity:** {severity}")
            md.append(f"**Sources:** {', '.join(finding.get('sources', []))}")
            if finding.get("exploitability_score"):
                md.append(f"**Exploitability Score:** {finding.get('exploitability_score')}")
            md.append("")
            md.append("### Description")
            md.append("")
            md.append(finding.get("description", "_No description._"))
            md.append("")
            if finding.get("exploit_guidance"):
                md.append("### Exploitation Guidance")
                md.append("")
                md.append(finding.get("exploit_guidance"))
                md.append("")
            if finding.get("remediation"):
                md.append("### Remediation")
                md.append("")
                md.append(finding.get("remediation"))
                md.append("")
    
    # PoC Scripts
    if report.poc_scripts:
        md.append("---")
        md.append("")
        md.append("# Proof-of-Concept Scripts")
        md.append("")
        for i, poc in enumerate(report.poc_scripts, 1):
            md.append(f"## PoC {i}: {poc.get('vulnerability_name', 'Unknown Vulnerability')}")
            md.append("")
            md.append(f"**Language:** {poc.get('language', 'Unknown')}")
            md.append("")
            if poc.get("description"):
                md.append("### Description")
                md.append("")
                md.append(poc.get("description"))
                md.append("")
            if poc.get("usage_instructions"):
                md.append("### Usage Instructions")
                md.append("")
                md.append(poc.get("usage_instructions"))
                md.append("")
            if poc.get("script_code"):
                lang = poc.get("language", "").lower()
                md.append("### Script Code")
                md.append("")
                md.append(f"```{lang}")
                md.append(poc.get("script_code"))
                md.append("```")
                md.append("")
            if poc.get("expected_output"):
                md.append("### Expected Output")
                md.append("")
                md.append(f"```\n{poc.get('expected_output')}\n```")
                md.append("")
    
    # Beginner Attack Guides
    if report.beginner_attack_guide:
        md.append("---")
        md.append("")
        md.append("# Beginner Attack Guides")
        md.append("")
        for guide in report.beginner_attack_guide:
            md.append(f"## {guide.get('attack_name', 'Attack Guide')}")
            md.append("")
            md.append(f"**Difficulty:** {guide.get('difficulty_level', 'Unknown')}")
            md.append(f"**Estimated Time:** {guide.get('estimated_time', 'Unknown')}")
            md.append("")
            if guide.get("prerequisites"):
                md.append("### Prerequisites")
                md.append("")
                for prereq in guide.get("prerequisites", []):
                    md.append(f"- {prereq}")
                md.append("")
            if guide.get("tools_needed"):
                md.append("### Tools Needed")
                md.append("")
                for tool in guide.get("tools_needed", []):
                    md.append(f"- **{tool.get('tool', 'Tool')}**: {tool.get('purpose', '')}")
                    if tool.get("installation"):
                        md.append(f"  - Installation: `{tool.get('installation')}`")
                md.append("")
            if guide.get("step_by_step_guide"):
                md.append("### Step-by-Step Guide")
                md.append("")
                for step in guide.get("step_by_step_guide", []):
                    md.append(f"#### Step {step.get('step_number', '?')}: {step.get('title', 'Step')}")
                    md.append("")
                    if step.get("explanation"):
                        md.append(step.get("explanation"))
                        md.append("")
                    if step.get("command_or_action"):
                        md.append("**Command/Action:**")
                        md.append(f"```bash\n{step.get('command_or_action')}\n```")
                        md.append("")
                    if step.get("expected_output"):
                        md.append(f"**Expected Output:** {step.get('expected_output')}")
                        md.append("")
                    if step.get("troubleshooting"):
                        md.append(f"**Troubleshooting:** {step.get('troubleshooting')}")
                        md.append("")
            if guide.get("success_indicators"):
                md.append("### Success Indicators")
                md.append("")
                for indicator in guide.get("success_indicators", []):
                    md.append(f"- {indicator}")
                md.append("")
    
    # Attack Chains
    if report.attack_chains:
        md.append("---")
        md.append("")
        md.append("# Attack Chains")
        md.append("")
        for chain in report.attack_chains:
            md.append(f"## {chain.get('chain_name', 'Attack Chain')}")
            md.append("")
            md.append(f"**Entry Point:** {chain.get('entry_point', 'Unknown')}")
            md.append(f"**Likelihood:** {chain.get('likelihood', 'Unknown')}")
            md.append(f"**Final Impact:** {chain.get('final_impact', 'Unknown')}")
            md.append("")
            if chain.get("steps"):
                md.append("### Attack Steps")
                md.append("")
                for step in chain.get("steps", []):
                    md.append(f"**Step {step.get('step', '?')}:** {step.get('action', 'Action')}")
                    md.append(f"- Vulnerability Used: {step.get('vulnerability_used', 'N/A')}")
                    md.append(f"- Outcome: {step.get('outcome', 'N/A')}")
                    md.append("")
    
    # Exploit Development Areas
    if report.exploit_development_areas:
        md.append("---")
        md.append("")
        md.append("# Exploit Development Areas")
        md.append("")
        for area in report.exploit_development_areas:
            md.append(f"## {area.get('title', 'Exploit Area')}")
            md.append("")
            md.append(f"**Attack Vector:** {area.get('attack_vector', 'Unknown')}")
            md.append(f"**Complexity:** {area.get('complexity', 'Unknown')}")
            md.append("")
            if area.get("description"):
                md.append(area.get("description"))
                md.append("")
            if area.get("impact"):
                md.append("### Impact")
                md.append("")
                md.append(area.get("impact"))
                md.append("")
            if area.get("poc_guidance"):
                md.append("### PoC Guidance")
                md.append("")
                md.append(area.get("poc_guidance"))
                md.append("")
    
    # Prioritized Vulnerabilities
    if report.prioritized_vulnerabilities:
        md.append("---")
        md.append("")
        md.append("# Prioritized Vulnerabilities")
        md.append("")
        md.append("_Ranked by exploitability, impact, and remediation complexity_")
        md.append("")
        for vuln in report.prioritized_vulnerabilities:
            rank = vuln.get('rank', '?')
            severity = vuln.get('severity', 'Medium').upper()
            severity_emoji = {"CRITICAL": "🔴", "HIGH": "🟠", "MEDIUM": "🟡", "LOW": "🟢"}.get(severity, "⚪")
            md.append(f"## {severity_emoji} #{rank}: {vuln.get('title', 'Vulnerability')} [{severity}]")
            md.append("")
            md.append(f"**CVSS Estimate:** {vuln.get('cvss_estimate', 'N/A')} | **Exploitability:** {vuln.get('exploitability', 'Unknown')} | **Remediation Priority:** {vuln.get('remediation_priority', 'N/A')}")
            md.append("")
            md.append(f"**Affected Component:** `{vuln.get('affected_component', 'Unknown')}`")
            md.append("")
            if vuln.get('impact'):
                md.append("### Impact")
                md.append("")
                md.append(vuln.get('impact'))
                md.append("")
            if vuln.get('exploitation_steps'):
                md.append("### Exploitation Steps")
                md.append("")
                for i, step in enumerate(vuln.get('exploitation_steps', []), 1):
                    md.append(f"{i}. {step}")
                md.append("")
            if vuln.get('remediation_steps'):
                md.append("### Remediation Steps")
                md.append("")
                for step in vuln.get('remediation_steps', []):
                    md.append(f"- {step}")
                md.append("")
            if vuln.get('references'):
                md.append("### References")
                md.append("")
                for ref in vuln.get('references', []):
                    md.append(f"- {ref}")
                md.append("")
    
    # Source Code Findings
    if report.source_code_findings:
        md.append("---")
        md.append("")
        md.append("# Source Code Findings")
        md.append("")
        for finding in report.source_code_findings:
            md.append(f"## {finding.get('title', 'Code Finding')}")
            md.append("")
            md.append(f"**File:** `{finding.get('file_path', 'Unknown')}`")
            if finding.get("line_numbers"):
                md.append(f"**Lines:** {finding.get('line_numbers')}")
            md.append(f"**Severity:** {finding.get('severity', 'Unknown')}")
            md.append("")
            if finding.get("analysis"):
                md.append(finding.get("analysis"))
                md.append("")
            if finding.get("code_snippet"):
                md.append("### Vulnerable Code")
                md.append("")
                md.append(f"```\n{finding.get('code_snippet')}\n```")
                md.append("")
    
    # Attack Surface Diagram
    if report.attack_surface_diagram:
        md.append("---")
        md.append("")
        md.append("# Attack Surface Diagram")
        md.append("")
        md.append("```mermaid")
        md.append(report.attack_surface_diagram)
        md.append("```")
        md.append("")
    
    # Documentation Analysis
    if report.documentation_analysis:
        md.append("---")
        md.append("")
        md.append("# Documentation Analysis")
        md.append("")
        md.append(report.documentation_analysis)
        md.append("")
    
    md.append("---")
    md.append("")
    md.append("*Report generated by VRAgent Combined Analysis*")
    
    return "\n".join(md)


@router.get("/reports/{report_id}/export/markdown")
async def export_report_markdown(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    Export combined analysis report as Markdown.
    """
    report = combined_analysis_service.get_combined_analysis_report(db, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    markdown_content = _generate_markdown_report(report)
    
    # Create a safe filename
    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in report.title)[:50]
    filename = f"{safe_title}_report.md"
    
    return StreamingResponse(
        io.BytesIO(markdown_content.encode("utf-8")),
        media_type="text/markdown",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
    )


@router.get("/reports/{report_id}/export/word")
async def export_report_word(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    Export combined analysis report as Word document (.docx).
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Please install it with: pip install python-docx")
    
    report = combined_analysis_service.get_combined_analysis_report(db, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading(report.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Generated: ").bold = True
    meta_para.add_run(f"{report.created_at.strftime('%Y-%m-%d %H:%M:%S') if hasattr(report.created_at, 'strftime') else report.created_at}\n")
    meta_para.add_run(f"Overall Risk Level: ").bold = True
    meta_para.add_run(f"{report.overall_risk_level or 'Unknown'}\n")
    meta_para.add_run(f"Risk Score: ").bold = True
    meta_para.add_run(f"{report.overall_risk_score or 0}/100\n")
    meta_para.add_run(f"Total Findings: ").bold = True
    meta_para.add_run(f"{report.total_findings_analyzed or 0}\n")
    meta_para.add_run(f"Scans Included: ").bold = True
    meta_para.add_run(f"{report.scans_included or 0}")
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    parse_markdown_for_word(report.executive_summary or "No executive summary available.", doc, Pt)
    
    # Risk Justification
    if report.risk_justification:
        doc.add_heading("Risk Justification", level=1)
        parse_markdown_for_word(report.risk_justification, doc, Pt)
    
    # Report Sections (includes Detailed Exploit Scenarios with code blocks)
    if report.report_sections:
        doc.add_heading("Detailed Analysis", level=1)
        for section in report.report_sections:
            severity = section.get("severity", "")
            title_text = section.get("title", "Section")
            if severity:
                title_text += f" [{severity.upper()}]"
            doc.add_heading(title_text, level=2)
            # Parse markdown content including code blocks
            parse_markdown_for_word(section.get("content", ""), doc, Pt)
    
    # Cross-Analysis Findings
    if report.cross_analysis_findings:
        doc.add_heading("Cross-Analysis Findings", level=1)
        for i, finding in enumerate(report.cross_analysis_findings, 1):
            doc.add_heading(f"{i}. {finding.get('title', 'Finding')} [{finding.get('severity', 'Medium').upper()}]", level=2)
            para = doc.add_paragraph()
            para.add_run("Severity: ").bold = True
            para.add_run(f"{finding.get('severity', 'Medium')}\n")
            para.add_run("Sources: ").bold = True
            para.add_run(f"{', '.join(finding.get('sources', []))}\n")
            parse_markdown_for_word(finding.get("description", ""), doc, Pt)
            if finding.get("exploit_guidance"):
                doc.add_heading("Exploitation Guidance", level=3)
                parse_markdown_for_word(finding.get("exploit_guidance"), doc, Pt)
            if finding.get("remediation"):
                doc.add_heading("Remediation", level=3)
                parse_markdown_for_word(finding.get("remediation"), doc, Pt)
    
    # PoC Scripts
    if report.poc_scripts:
        doc.add_heading("Proof-of-Concept Scripts", level=1)
        for i, poc in enumerate(report.poc_scripts, 1):
            doc.add_heading(f"PoC {i}: {poc.get('vulnerability_name', 'Unknown')}", level=2)
            para = doc.add_paragraph()
            para.add_run("Language: ").bold = True
            para.add_run(f"{poc.get('language', 'Unknown')}\n")
            if poc.get("description"):
                doc.add_paragraph(poc.get("description"))
            if poc.get("usage_instructions"):
                doc.add_heading("Usage Instructions", level=3)
                doc.add_paragraph(poc.get("usage_instructions"))
            if poc.get("script_code"):
                doc.add_heading("Script Code", level=3)
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(poc.get("script_code"))
                code_run.font.name = "Courier New"
                code_run.font.size = Pt(9)
    
    # Beginner Attack Guides
    if report.beginner_attack_guide:
        doc.add_heading("Beginner Attack Guides", level=1)
        for guide in report.beginner_attack_guide:
            doc.add_heading(guide.get("attack_name", "Attack Guide"), level=2)
            para = doc.add_paragraph()
            para.add_run("Difficulty: ").bold = True
            para.add_run(f"{guide.get('difficulty_level', 'Unknown')}\n")
            para.add_run("Estimated Time: ").bold = True
            para.add_run(f"{guide.get('estimated_time', 'Unknown')}\n")
            
            if guide.get("step_by_step_guide"):
                doc.add_heading("Step-by-Step Guide", level=3)
                for step in guide.get("step_by_step_guide", []):
                    step_para = doc.add_paragraph()
                    step_para.add_run(f"Step {step.get('step_number', '?')}: {step.get('title', 'Step')}").bold = True
                    if step.get("explanation"):
                        doc.add_paragraph(step.get("explanation"))
                    if step.get("command_or_action"):
                        cmd_para = doc.add_paragraph()
                        cmd_run = cmd_para.add_run(step.get("command_or_action"))
                        cmd_run.font.name = "Courier New"
                        cmd_run.font.size = Pt(9)
    
    # Exploit Development Areas
    if report.exploit_development_areas:
        doc.add_heading("Exploit Development Areas", level=1)
        for area in report.exploit_development_areas:
            doc.add_heading(area.get("title", "Exploit Area"), level=2)
            para = doc.add_paragraph()
            para.add_run("Attack Vector: ").bold = True
            para.add_run(f"{area.get('attack_vector', 'Unknown')}\n")
            para.add_run("Complexity: ").bold = True
            para.add_run(f"{area.get('complexity', 'Unknown')}\n")
            if area.get("description"):
                doc.add_paragraph(area.get("description"))
            if area.get("impact"):
                doc.add_heading("Impact", level=3)
                doc.add_paragraph(area.get("impact"))
    
    # Prioritized Vulnerabilities
    if report.prioritized_vulnerabilities:
        doc.add_heading("Prioritized Vulnerabilities", level=1)
        intro = doc.add_paragraph()
        intro.add_run("Ranked by exploitability, impact, and remediation complexity").italic = True
        
        for vuln in report.prioritized_vulnerabilities:
            rank = vuln.get('rank', '?')
            severity = vuln.get('severity', 'Medium').upper()
            doc.add_heading(f"#{rank}: {vuln.get('title', 'Vulnerability')} [{severity}]", level=2)
            
            meta_para = doc.add_paragraph()
            meta_para.add_run("CVSS Estimate: ").bold = True
            meta_para.add_run(f"{vuln.get('cvss_estimate', 'N/A')} | ")
            meta_para.add_run("Exploitability: ").bold = True
            meta_para.add_run(f"{vuln.get('exploitability', 'Unknown')} | ")
            meta_para.add_run("Remediation Priority: ").bold = True
            meta_para.add_run(f"{vuln.get('remediation_priority', 'N/A')}")
            
            comp_para = doc.add_paragraph()
            comp_para.add_run("Affected Component: ").bold = True
            comp_run = comp_para.add_run(vuln.get('affected_component', 'Unknown'))
            comp_run.font.name = "Courier New"
            comp_run.font.size = Pt(9)
            
            if vuln.get('impact'):
                doc.add_heading("Impact", level=3)
                doc.add_paragraph(vuln.get('impact'))
            
            if vuln.get('exploitation_steps'):
                doc.add_heading("Exploitation Steps", level=3)
                for i, step in enumerate(vuln.get('exploitation_steps', []), 1):
                    doc.add_paragraph(f"{i}. {step}")
            
            if vuln.get('remediation_steps'):
                doc.add_heading("Remediation Steps", level=3)
                for step in vuln.get('remediation_steps', []):
                    doc.add_paragraph(f"• {step}")
            
            if vuln.get('references'):
                doc.add_heading("References", level=3)
                for ref in vuln.get('references', []):
                    doc.add_paragraph(f"• {ref}")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create a safe filename
    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in report.title)[:50]
    filename = f"{safe_title}_report.docx"
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
    )


@router.get("/reports/{report_id}/export/pdf")
async def export_report_pdf(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    Export combined analysis report as PDF.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab not installed. Please install it with: pip install reportlab")
    
    report = combined_analysis_service.get_combined_analysis_report(db, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    
    # Styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='MainTitle', fontSize=24, alignment=TA_CENTER, spaceAfter=30, textColor=colors.darkblue))
    styles.add(ParagraphStyle(name='Heading1Custom', fontSize=18, spaceAfter=12, spaceBefore=20, textColor=colors.darkblue, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='Heading2Custom', fontSize=14, spaceAfter=8, spaceBefore=15, textColor=colors.darkblue, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='Heading3Custom', fontSize=12, spaceAfter=6, spaceBefore=10, fontName='Helvetica-Bold'))
    styles.add(ParagraphStyle(name='BodyCustom', fontSize=10, spaceAfter=8, alignment=TA_JUSTIFY, leading=14))
    styles.add(ParagraphStyle(name='CodeStyle', fontSize=8, fontName='Courier', spaceAfter=8, leftIndent=20, backColor=colors.lightgrey))
    styles.add(ParagraphStyle(name='MetaStyle', fontSize=10, spaceAfter=4))
    
    story = []
    
    # Title
    story.append(Paragraph(report.title, styles['MainTitle']))
    story.append(Spacer(1, 20))
    
    # Metadata
    meta_text = f"""
    <b>Generated:</b> {report.created_at.strftime('%Y-%m-%d %H:%M:%S') if hasattr(report.created_at, 'strftime') else report.created_at}<br/>
    <b>Overall Risk Level:</b> {report.overall_risk_level or 'Unknown'}<br/>
    <b>Risk Score:</b> {report.overall_risk_score or 0}/100<br/>
    <b>Total Findings:</b> {report.total_findings_analyzed or 0}<br/>
    <b>Scans Included:</b> {report.scans_included or 0}
    """
    story.append(Paragraph(meta_text, styles['MetaStyle']))
    story.append(Spacer(1, 20))
    
    # Executive Summary
    story.append(Paragraph("Executive Summary", styles['Heading1Custom']))
    parse_markdown_for_pdf(report.executive_summary or "No executive summary available.", story, styles)
    
    # Risk Justification
    if report.risk_justification:
        story.append(Paragraph("Risk Justification", styles['Heading1Custom']))
        parse_markdown_for_pdf(report.risk_justification, story, styles)
    
    # Report Sections (includes Detailed Exploit Scenarios with code blocks)
    if report.report_sections:
        story.append(PageBreak())
        story.append(Paragraph("Detailed Analysis", styles['Heading1Custom']))
        for section in report.report_sections:
            severity = section.get("severity", "")
            title_text = section.get("title", "Section")
            if severity:
                title_text += f" [{severity.upper()}]"
            story.append(Paragraph(title_text, styles['Heading2Custom']))
            # Parse markdown content including code blocks
            parse_markdown_for_pdf(section.get("content", ""), story, styles)
    
    # Cross-Analysis Findings
    if report.cross_analysis_findings:
        story.append(PageBreak())
        story.append(Paragraph("Cross-Analysis Findings", styles['Heading1Custom']))
        for i, finding in enumerate(report.cross_analysis_findings, 1):
            story.append(Paragraph(f"{i}. {finding.get('title', 'Finding')} [{finding.get('severity', 'Medium').upper()}]", styles['Heading2Custom']))
            meta = f"<b>Severity:</b> {finding.get('severity', 'Medium')} | <b>Sources:</b> {', '.join(finding.get('sources', []))}"
            story.append(Paragraph(meta, styles['MetaStyle']))
            parse_markdown_for_pdf(finding.get("description", ""), story, styles)
            if finding.get("exploit_guidance"):
                story.append(Paragraph("Exploitation Guidance", styles['Heading3Custom']))
                parse_markdown_for_pdf(finding.get("exploit_guidance"), story, styles)
            if finding.get("remediation"):
                story.append(Paragraph("Remediation", styles['Heading3Custom']))
                parse_markdown_for_pdf(finding.get("remediation"), story, styles)
    
    # PoC Scripts
    if report.poc_scripts:
        story.append(PageBreak())
        story.append(Paragraph("Proof-of-Concept Scripts", styles['Heading1Custom']))
        for i, poc in enumerate(report.poc_scripts, 1):
            story.append(Paragraph(f"PoC {i}: {poc.get('vulnerability_name', 'Unknown')}", styles['Heading2Custom']))
            story.append(Paragraph(f"<b>Language:</b> {poc.get('language', 'Unknown')}", styles['MetaStyle']))
            if poc.get("description"):
                story.append(Paragraph(poc.get("description").replace('\n', '<br/>'), styles['BodyCustom']))
            if poc.get("usage_instructions"):
                story.append(Paragraph("Usage Instructions", styles['Heading3Custom']))
                story.append(Paragraph(poc.get("usage_instructions").replace('\n', '<br/>'), styles['BodyCustom']))
            if poc.get("script_code"):
                story.append(Paragraph("Script Code", styles['Heading3Custom']))
                # Escape special characters for PDF
                code = poc.get("script_code", "").replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
                story.append(Paragraph(code, styles['CodeStyle']))
    
    # Beginner Attack Guides
    if report.beginner_attack_guide:
        story.append(PageBreak())
        story.append(Paragraph("Beginner Attack Guides", styles['Heading1Custom']))
        for guide in report.beginner_attack_guide:
            story.append(Paragraph(guide.get("attack_name", "Attack Guide"), styles['Heading2Custom']))
            meta = f"<b>Difficulty:</b> {guide.get('difficulty_level', 'Unknown')} | <b>Time:</b> {guide.get('estimated_time', 'Unknown')}"
            story.append(Paragraph(meta, styles['MetaStyle']))
            
            if guide.get("step_by_step_guide"):
                story.append(Paragraph("Step-by-Step Guide", styles['Heading3Custom']))
                for step in guide.get("step_by_step_guide", []):
                    story.append(Paragraph(f"<b>Step {step.get('step_number', '?')}: {step.get('title', 'Step')}</b>", styles['BodyCustom']))
                    if step.get("explanation"):
                        story.append(Paragraph(step.get("explanation").replace('\n', '<br/>'), styles['BodyCustom']))
                    if step.get("command_or_action"):
                        cmd = step.get("command_or_action", "").replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(cmd, styles['CodeStyle']))
    
    # Prioritized Vulnerabilities
    if report.prioritized_vulnerabilities:
        story.append(PageBreak())
        story.append(Paragraph("Prioritized Vulnerabilities", styles['Heading1Custom']))
        story.append(Paragraph("<i>Ranked by exploitability, impact, and remediation complexity</i>", styles['BodyCustom']))
        
        for vuln in report.prioritized_vulnerabilities:
            rank = vuln.get('rank', '?')
            severity = vuln.get('severity', 'Medium').upper()
            severity_color = {"CRITICAL": colors.red, "HIGH": colors.orangered, "MEDIUM": colors.orange, "LOW": colors.green}.get(severity, colors.grey)
            
            story.append(Paragraph(f"#{rank}: {vuln.get('title', 'Vulnerability')} [{severity}]", styles['Heading2Custom']))
            
            meta = f"<b>CVSS Estimate:</b> {vuln.get('cvss_estimate', 'N/A')} | <b>Exploitability:</b> {vuln.get('exploitability', 'Unknown')} | <b>Remediation Priority:</b> {vuln.get('remediation_priority', 'N/A')}"
            story.append(Paragraph(meta, styles['MetaStyle']))
            
            comp = f"<b>Affected Component:</b> <font face='Courier'>{vuln.get('affected_component', 'Unknown')}</font>"
            story.append(Paragraph(comp, styles['MetaStyle']))
            
            if vuln.get('impact'):
                story.append(Paragraph("Impact", styles['Heading3Custom']))
                story.append(Paragraph(vuln.get('impact').replace('\n', '<br/>'), styles['BodyCustom']))
            
            if vuln.get('exploitation_steps'):
                story.append(Paragraph("Exploitation Steps", styles['Heading3Custom']))
                for i, step in enumerate(vuln.get('exploitation_steps', []), 1):
                    story.append(Paragraph(f"{i}. {step}", styles['BodyCustom']))
            
            if vuln.get('remediation_steps'):
                story.append(Paragraph("Remediation Steps", styles['Heading3Custom']))
                for step in vuln.get('remediation_steps', []):
                    story.append(Paragraph(f"• {step}", styles['BodyCustom']))
            
            if vuln.get('references'):
                story.append(Paragraph("References", styles['Heading3Custom']))
                for ref in vuln.get('references', []):
                    ref_escaped = ref.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(f"• {ref_escaped}", styles['BodyCustom']))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    # Create a safe filename
    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in report.title)[:50]
    filename = f"{safe_title}_report.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
    )


# ============================================================================
# Send to Team Chat
# ============================================================================

@router.post("/reports/{report_id}/send-to-team-chat")
async def send_report_to_team_chat(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Send a summary of the combined analysis report to the project's team chat.
    """
    if not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    report = combined_analysis_service.get_combined_analysis_report(db, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get the project's team chat conversation
    project = db.query(Project).filter(Project.id == report.project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Find team chat conversation
    team_chat = db.query(Conversation).filter(
        Conversation.project_id == report.project_id
    ).first()
    
    if not team_chat:
        raise HTTPException(status_code=404, detail="Team chat not found for this project. Please set up a team chat first.")
    
    # Check if user is participant
    is_participant = db.query(ConversationParticipant).filter(
        ConversationParticipant.conversation_id == team_chat.id,
        ConversationParticipant.user_id == current_user.id
    ).first()
    
    if not is_participant:
        raise HTTPException(status_code=403, detail="You are not a member of this team chat")
    
    # Create report summary message
    risk_emoji = {
        "Critical": "🔴",
        "High": "🟠",
        "Medium": "🟡",
        "Low": "🟢",
        "Clean": "✅",
    }.get(report.overall_risk_level or "Unknown", "⚪")
    
    summary_message = f"""📊 **Combined Analysis Report Shared**

**{report.title}**

{risk_emoji} **Risk Level:** {report.overall_risk_level or 'Unknown'}
📈 **Risk Score:** {report.overall_risk_score or 0}/100
🔍 **Findings Analyzed:** {report.total_findings_analyzed or 0}
📋 **Scans Included:** {report.scans_included or 0}

---

**Executive Summary:**
{(report.executive_summary or 'No summary available.')[:500]}{'...' if len(report.executive_summary or '') > 500 else ''}

---

📎 [View Full Report](/projects/{report.project_id}/combined-analysis?report={report_id})

_Shared by {current_user.username}_"""

    # Create the message
    message = Message(
        conversation_id=team_chat.id,
        sender_id=current_user.id,
        content=summary_message,
        message_type="report_share",
        attachment_data={
            "type": "combined_analysis_report",
            "report_id": report_id,
            "report_title": report.title,
            "risk_level": report.overall_risk_level,
            "risk_score": report.overall_risk_score,
        }
    )
    db.add(message)
    
    # Update conversation last_message_at
    team_chat.last_message_at = datetime.utcnow()
    
    db.commit()
    db.refresh(message)
    
    return {
        "status": "success",
        "message": "Report summary sent to team chat",
        "conversation_id": team_chat.id,
        "message_id": message.id
    }


# ============================================================================
# AI Chat
# ============================================================================

@router.post("/reports/{report_id}/chat", response_model=CombinedAnalysisChatResponse)
async def chat_about_report(
    report_id: int,
    request: CombinedAnalysisChatRequest,
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_current_user),
):
    """
    Chat with AI about a combined analysis report.
    The AI has full access to the report content to answer questions.
    """
    import os
    
    report = combined_analysis_service.get_combined_analysis_report(db, report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Check for API key
    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if not api_key:
        return CombinedAnalysisChatResponse(
            response="AI chat is not available - API key not configured. Please contact the administrator.",
            suggestions=["What are the most critical vulnerabilities?", "Summarize the key findings"]
        )
    
    # Build report context for the AI
    report_context = _build_report_context(report)
    
    # Build conversation history
    history_text = ""
    if request.history:
        for msg in request.history[-10:]:  # Last 10 messages
            role = "User" if msg.role == "user" else "Assistant"
            history_text += f"\n{role}: {msg.content}\n"
    
    # Create prompt for the AI
    system_prompt = f"""You are an expert security analyst assistant helping to understand and analyze a Combined Security Analysis Report.

You have FULL ACCESS to the following report content:

=== REPORT: {report.title} ===

**Overall Risk Level:** {report.overall_risk_level}
**Risk Score:** {report.overall_risk_score}/100
**Total Findings:** {report.total_findings_analyzed}
**Scans Included:** {report.scans_included}

**Executive Summary:**
{report.executive_summary or 'No summary available'}

**Risk Justification:**
{report.risk_justification or 'No justification available'}

{report_context}

=== END OF REPORT ===

IMPORTANT GUIDELINES:
- Answer questions specifically about this report's content
- Reference specific findings, vulnerabilities, or sections when relevant
- Provide actionable insights and explanations
- If asked about something not in the report, say so clearly
- Be concise but thorough
- Use markdown formatting for better readability
- When showing code examples, ALWAYS use proper markdown code blocks with language specification (```python, ```javascript, etc.)
- When discussing vulnerabilities, mention their severity and potential impact

{history_text}
"""
    
    user_prompt = request.message
    
    try:
        from google import genai
        from backend.core.config import settings
        
        client = genai.Client(api_key=api_key)
        
        # Build conversation messages
        messages = []
        
        # Add conversation history
        if request.history:
            for msg in request.history[-10:]:
                messages.append({
                    "role": "user" if msg.role == "user" else "model",
                    "parts": [{"text": msg.content}]
                })
        
        # Add current message
        messages.append({
            "role": "user",
            "parts": [{"text": user_prompt}]
        })
        
        # Generate response
        response = client.models.generate_content(
            model=settings.gemini_model_id,
            contents=messages,
            config={
                "system_instruction": system_prompt,
                "temperature": 0.7,
                "max_output_tokens": 2000,
            }
        )
        
        if response.text:
            # Generate follow-up suggestions based on the context
            suggestions = _generate_chat_suggestions(request.message, report)
            
            return CombinedAnalysisChatResponse(
                response=response.text,
                suggestions=suggestions
            )
        else:
            return CombinedAnalysisChatResponse(
                response="I couldn't generate a response. Please try rephrasing your question.",
                suggestions=["What are the most critical vulnerabilities?", "Summarize the key findings"]
            )
            
    except ImportError:
        logger.error("google-genai not installed")
        return CombinedAnalysisChatResponse(
            response="AI module not available. Please ensure google-genai is installed.",
            suggestions=["What are the most critical vulnerabilities?"]
        )
    except Exception as e:
        logger.error(f"Error in AI chat: {e}")
        import traceback
        traceback.print_exc()
        # Provide a helpful fallback response
        return CombinedAnalysisChatResponse(
            response=f"I apologize, but I encountered an error processing your request. Here's what I can tell you about the report:\n\n"
                     f"**{report.title}**\n"
                     f"- Risk Level: {report.overall_risk_level}\n"
                     f"- Risk Score: {report.overall_risk_score}/100\n"
                     f"- Total Findings: {report.total_findings_analyzed}\n\n"
                     f"Please try rephrasing your question or ask about specific sections of the report.",
            suggestions=[
                "What are the most critical vulnerabilities?",
                "Summarize the key findings",
                "What remediation steps are recommended?",
            ]
        )


def _build_report_context(report) -> str:
    """Build detailed context from the report for AI consumption."""
    context_parts = []
    
    # Report sections (includes Detailed Exploit Scenarios)
    if report.report_sections:
        context_parts.append("\n**DETAILED ANALYSIS SECTIONS:**")
        for section in report.report_sections:
            severity = section.get("severity", "")
            title = section.get("title", "Section")
            content = section.get("content", "")[:3000]  # Limit to avoid token overflow
            if severity:
                context_parts.append(f"\n### {title} [{severity.upper()}]\n{content}")
            else:
                context_parts.append(f"\n### {title}\n{content}")
    
    # Cross-Analysis Findings
    if report.cross_analysis_findings:
        context_parts.append("\n**CROSS-ANALYSIS FINDINGS:**")
        for i, finding in enumerate(report.cross_analysis_findings[:10], 1):  # Top 10
            context_parts.append(
                f"\n{i}. **{finding.get('title', 'Finding')}** [{finding.get('severity', 'Medium')}]\n"
                f"   Sources: {', '.join(finding.get('sources', []))}\n"
                f"   {finding.get('description', '')[:500]}"
            )
    
    # Prioritized Vulnerabilities
    if report.prioritized_vulnerabilities:
        context_parts.append("\n**PRIORITIZED VULNERABILITIES:**")
        for vuln in report.prioritized_vulnerabilities[:10]:  # Top 10
            context_parts.append(
                f"\n- **#{vuln.get('rank', '?')}: {vuln.get('title', 'Vulnerability')}** [{vuln.get('severity', 'Medium')}]\n"
                f"  CVSS: {vuln.get('cvss_estimate', 'N/A')} | Exploitability: {vuln.get('exploitability', 'Unknown')}\n"
                f"  Component: {vuln.get('affected_component', 'Unknown')}\n"
                f"  Impact: {vuln.get('impact', 'Unknown')[:300]}"
            )
    
    # PoC Scripts
    if report.poc_scripts:
        context_parts.append("\n**PROOF-OF-CONCEPT SCRIPTS:**")
        for poc in report.poc_scripts[:5]:  # Top 5
            context_parts.append(
                f"\n- **{poc.get('vulnerability_name', 'PoC')}** ({poc.get('language', 'Unknown')})\n"
                f"  {poc.get('description', '')[:200]}"
            )
    
    # Attack Chains
    if report.attack_chains:
        context_parts.append("\n**ATTACK CHAINS:**")
        for chain in report.attack_chains[:3]:  # Top 3
            context_parts.append(
                f"\n- **{chain.get('name', 'Attack Chain')}** [{chain.get('risk_level', 'Unknown')}]\n"
                f"  {chain.get('description', '')[:300]}"
            )
    
    # Exploit Development Areas
    if report.exploit_development_areas:
        context_parts.append("\n**EXPLOIT DEVELOPMENT AREAS:**")
        for area in report.exploit_development_areas[:5]:
            context_parts.append(
                f"\n- **{area.get('title', 'Area')}**\n"
                f"  Vector: {area.get('attack_vector', 'Unknown')} | Complexity: {area.get('complexity', 'Unknown')}\n"
                f"  {area.get('description', '')[:200]}"
            )
    
    # Attack Surface Diagram
    if report.attack_surface_diagram:
        context_parts.append("\n**ATTACK SURFACE DIAGRAM:**")
        context_parts.append(f"\n{report.attack_surface_diagram[:1500]}")
    
    return "\n".join(context_parts)


def _generate_chat_suggestions(user_message: str, report) -> list:
    """Generate contextual follow-up suggestions."""
    msg_lower = user_message.lower()
    suggestions = []
    
    # Context-aware suggestions
    if "critical" in msg_lower or "severe" in msg_lower:
        suggestions.append("What are the exploitation steps for the top vulnerability?")
        suggestions.append("Are there any PoC scripts for critical findings?")
    elif "fix" in msg_lower or "remediat" in msg_lower:
        suggestions.append("What's the priority order for fixes?")
        suggestions.append("Are there any quick wins I can implement?")
    elif "exploit" in msg_lower or "attack" in msg_lower:
        suggestions.append("What attack chains were identified?")
        suggestions.append("Which vulnerabilities have the highest exploitability?")
    elif "summary" in msg_lower or "overview" in msg_lower:
        suggestions.append("What are the most critical findings?")
        suggestions.append("What's the overall risk posture?")
    else:
        # Generic suggestions based on report content
        if report.prioritized_vulnerabilities:
            suggestions.append("Explain the top vulnerability in detail")
        if report.poc_scripts:
            suggestions.append("Tell me about the PoC scripts")
        if report.attack_chains:
            suggestions.append("Describe the attack chains")
        if not suggestions:
            suggestions = [
                "What are the key findings?",
                "What should I fix first?",
                "Are there any critical vulnerabilities?",
            ]
    
    return suggestions[:4]  # Max 4 suggestions