"""
Reverse Engineering Router for VRAgent.

Provides endpoints for analyzing:
- Binary files (EXE, ELF, DLL, SO)
- Android APK files
- Docker image layers
"""

import os
import time
import shutil
import tempfile
from pathlib import Path
from typing import Optional, List, Dict, Any, Set, Union
from datetime import datetime
import asyncio
import json
import uuid
from types import SimpleNamespace

from fastapi import APIRouter, Body, File, HTTPException, UploadFile, Query, Depends, Form
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field
from sqlalchemy import and_, or_
from sqlalchemy.orm import Session

from backend.core.logging import get_logger
from backend.core.database import get_db
from backend.core.config import settings
from backend.core.file_validator import sanitize_filename
from backend.core.auth import get_current_active_user
from backend.models.models import (
    Project,
    ProjectCollaborator,
    ProjectDockerImage,
    ReverseEngineeringReport,
    User,
    UserDockerImage,
)
from backend.services import project_service
from backend.services import reverse_engineering_service as re_service
from backend.services import deduplication_service

router = APIRouter(prefix="/reverse", tags=["reverse-engineering"])
logger = get_logger(__name__)


class ExportSafeNamespace(SimpleNamespace):
    """Namespace that returns None for missing attributes during export rendering."""

    def __getattr__(self, name: str):
        return None

# Constants
MAX_FILE_SIZE = 5 * 1024 * 1024 * 1024  # 5GB - supports large APKs, firmware images, binaries (was 500MB)
MAX_SCAN_TIMEOUT = 3600  # 60 minutes global timeout for unified scans
ALLOWED_BINARY_EXTENSIONS = {".exe", ".dll", ".so", ".elf", ".bin", ".o", ".dylib", ".mach"}
ALLOWED_APK_EXTENSIONS = {".apk", ".aab"}

# Store active unified scans for cancellation (must be defined before endpoint functions)
_unified_scan_sessions: Dict[str, Dict[str, Any]] = {}
_unified_binary_scan_sessions: Dict[str, Dict[str, Any]] = {}


def _serialize_symbolic_result(sym_result) -> Optional[Dict[str, Any]]:
    """Serialize a SymbolicExecutionResult dataclass to a JSON-compatible dict."""
    if sym_result is None:
        return None
    try:
        return {
            "paths_explored": sym_result.paths_explored,
            "paths_deadended": sym_result.paths_deadended,
            "paths_errored": sym_result.paths_errored,
            "max_depth_reached": sym_result.max_depth_reached,
            "symbolic_inputs": [
                {
                    "name": si.name,
                    "type": si.type,
                    "size_bits": si.size_bits,
                    "constraints": si.constraints[:10],
                    "concrete_examples": si.concrete_examples[:5],
                }
                for si in (sym_result.symbolic_inputs or [])
            ],
            "crash_inputs": [
                {
                    "input_type": ci.input_type,
                    "input_value": ci.input_value.hex() if isinstance(ci.input_value, bytes) else str(ci.input_value),
                    "crash_address": hex(ci.crash_address) if ci.crash_address else "0x0",
                    "crash_type": ci.crash_type,
                    "vulnerability_type": ci.vulnerability_type,
                    "cwe_id": ci.cwe_id,
                    "exploitability": ci.exploitability,
                }
                for ci in (sym_result.crash_inputs or [])[:20]
            ],
            "target_reaches": [
                {
                    "target_address": hex(tr.target_address) if tr.target_address else "0x0",
                    "target_name": tr.target_name,
                    "reached": tr.reached,
                    "input_to_reach": tr.input_to_reach.hex() if isinstance(tr.input_to_reach, bytes) else str(tr.input_to_reach) if tr.input_to_reach else None,
                    "path_length": tr.path_length,
                    "constraints_solved": tr.constraints_solved,
                }
                for tr in (sym_result.target_reaches or [])[:20]
            ],
            "interesting_paths": [
                {
                    "path_id": ip.path_id,
                    "depth": ip.depth,
                    "constraints_count": ip.constraints_count,
                    "is_feasible": ip.is_feasible,
                    "termination_reason": ip.termination_reason,
                    "addresses_visited": [hex(a) for a in (ip.addresses_visited or [])[:50]],
                }
                for ip in (sym_result.interesting_paths or [])[:20]
            ],
            "vulnerabilities_found": sym_result.vulnerabilities_found or [],
            "execution_time_seconds": sym_result.execution_time_seconds,
            "memory_used_mb": sym_result.memory_used_mb,
            "timeout_reached": sym_result.timeout_reached,
            "error": sym_result.error,
        }
    except Exception as e:
        logger.error(f"Failed to serialize symbolic result: {e}")
        return {"error": str(e), "paths_explored": 0, "vulnerabilities_found": []}


# ============================================================================
# Response Models
# ============================================================================

class BinaryStringResponse(BaseModel):
    """A string extracted from a binary."""
    value: str
    offset: int
    encoding: str
    category: Optional[str] = None


class ImportedFunctionResponse(BaseModel):
    """An imported function from a binary."""
    name: str
    library: str
    ordinal: Optional[int] = None
    is_suspicious: bool = False
    reason: Optional[str] = None


class RichHeaderEntryResponse(BaseModel):
    """An entry in the PE Rich header."""
    product_id: int
    build_id: int
    count: int
    product_name: Optional[str] = None
    vs_version: Optional[str] = None


class RichHeaderResponse(BaseModel):
    """PE Rich header information for compiler/linker identification."""
    entries: List[RichHeaderEntryResponse]
    rich_hash: str  # MD5 hash for malware identification
    checksum: int
    raw_data: str
    clear_data: str


class BinaryMetadataResponse(BaseModel):
    """Metadata from a binary file."""
    file_type: str
    architecture: str
    file_size: int
    entry_point: Optional[int] = None
    is_packed: bool = False
    packer_name: Optional[str] = None
    compile_time: Optional[str] = None
    sections: List[Dict[str, Any]] = []
    headers: Dict[str, Any] = {}
    # PE-specific
    rich_header: Optional[RichHeaderResponse] = None
    imphash: Optional[str] = None
    tls_callbacks: List[int] = []
    mitigations: Dict[str, Any] = {}
    resource_summary: Dict[str, Any] = {}
    version_info: Dict[str, Any] = {}
    authenticode: Optional[Dict[str, Any]] = None
    overlay: Optional[Dict[str, Any]] = None
    pe_delay_imports: List[Dict[str, Any]] = []
    pe_relocations: Dict[str, Any] = {}
    pe_debug: Dict[str, Any] = {}
    pe_data_directories: List[Dict[str, Any]] = []
    pe_manifest: Optional[str] = None
    # ELF-specific
    relro: Optional[str] = None
    stack_canary: bool = False
    nx_enabled: bool = False
    pie_enabled: bool = False
    interpreter: Optional[str] = None
    linked_libraries: List[str] = []
    elf_dynamic: Dict[str, Any] = {}
    elf_relocations: Dict[str, Any] = {}
    elf_version_info: Dict[str, Any] = {}
    elf_build_id: Optional[str] = None
    elf_program_headers: List[Dict[str, Any]] = []


class SecretResponse(BaseModel):
    """A potential secret found."""
    type: str
    value: str
    masked_value: str
    severity: str
    context: Optional[str] = None
    offset: Optional[int] = None


class SuspiciousIndicatorResponse(BaseModel):
    """A suspicious indicator found in analysis."""
    category: str
    severity: str
    description: str
    details: Optional[Any] = None


class BinaryAnalysisResponse(BaseModel):
    """Complete binary analysis response."""
    filename: str
    metadata: BinaryMetadataResponse
    strings_count: int
    strings_sample: List[BinaryStringResponse]
    imports: List[ImportedFunctionResponse]
    exports: List[str]
    secrets: List[SecretResponse]
    suspicious_indicators: List[SuspiciousIndicatorResponse]
    fuzzy_hashes: Dict[str, Optional[str]] = {}
    yara_matches: List[Dict[str, Any]] = []
    capa_summary: Optional[Dict[str, Any]] = None
    deobfuscated_strings: List[Dict[str, Any]] = []
    ai_analysis: Optional[str] = None
    ghidra_analysis: Optional[Dict[str, Any]] = None
    ghidra_ai_summaries: Optional[List[Dict[str, Any]]] = None
    vuln_hunt_result: Optional[Dict[str, Any]] = None  # Multi-pass AI vulnerability hunt results
    # Enhanced analysis results (matching APK analyzer)
    pattern_scan_result: Optional[Dict[str, Any]] = None  # Pattern-based vulnerability scan
    cve_lookup_result: Optional[Dict[str, Any]] = None  # CVE lookup results
    sensitive_scan_result: Optional[Dict[str, Any]] = None  # Sensitive data discovery
    verification_result: Optional[Dict[str, Any]] = None  # Unified AI verification (includes risk, FP filtering, attack chains)
    attack_chains: Optional[List[Dict[str, Any]]] = None  # Detected attack chains from correlated findings
    prioritized_actions: Optional[List[Dict[str, Any]]] = None  # AI-prioritized remediation actions
    # NEW: APK-matching features
    obfuscation_analysis: Optional[Dict[str, Any]] = None  # Obfuscation/packing detection
    attack_surface: Optional[Dict[str, Any]] = None  # Attack surface mapping
    dynamic_analysis: Optional[Dict[str, Any]] = None  # Frida scripts for dynamic analysis
    emulation_analysis: Optional[Dict[str, Any]] = None  # Unicorn emulation results
    ai_functionality_report: Optional[str] = None  # AI report: What does this binary do?
    ai_security_report: Optional[str] = None  # AI report: Security assessment
    ai_architecture_diagram: Optional[str] = None  # AI-generated Mermaid architecture diagram
    ai_attack_surface_map: Optional[str] = None  # AI-generated Mermaid attack tree
    # Entropy analysis (per-section heatmap visualization)
    entropy_analysis: Optional[Dict[str, Any]] = None
    # Symbolic execution results
    symbolic_execution: Optional[Dict[str, Any]] = None
    # Legitimacy detection (reduces false positives)
    is_legitimate_software: Optional[bool] = None  # Whether binary appears to be from known publisher
    legitimacy_indicators: Optional[List[str]] = None  # Reasons why it appears legitimate
    error: Optional[str] = None


class UnifiedBinaryScanPhase(BaseModel):
    """A phase in the unified binary scan."""
    id: str
    label: str
    description: str
    status: str  # "pending", "in_progress", "completed", "error"
    progress: int = 0  # 0-100
    details: Optional[str] = None
    started_at: Optional[str] = None
    completed_at: Optional[str] = None


class UnifiedBinaryScanProgress(BaseModel):
    """Progress update for unified binary scan."""
    scan_id: str
    current_phase: str
    overall_progress: int  # 0-100
    phases: List[UnifiedBinaryScanPhase]
    message: str
    error: Optional[str] = None
    # Time tracking
    elapsed_seconds: Optional[float] = None
    estimated_remaining_seconds: Optional[float] = None
    estimated_total_seconds: Optional[float] = None


class ApkPermissionResponse(BaseModel):
    """An Android permission."""
    name: str
    is_dangerous: bool
    description: Optional[str] = None


class ApkComponentResponse(BaseModel):
    """An Android app component."""
    name: str
    component_type: str
    is_exported: bool
    intent_filters: List[str] = []


class ApkSecurityIssueResponse(BaseModel):
    """A security issue found in APK."""
    category: str
    severity: str
    description: str
    details: Optional[Any] = None


class ApkAnalysisResponse(BaseModel):
    """Complete APK analysis response."""
    filename: str
    package_name: str
    version_name: Optional[str] = None
    version_code: Optional[int] = None
    min_sdk: Optional[int] = None
    target_sdk: Optional[int] = None
    permissions: List[ApkPermissionResponse]
    dangerous_permissions_count: int
    components: List[ApkComponentResponse]
    strings_count: int
    secrets: List[SecretResponse]
    urls: List[str]
    native_libraries: List[str]
    security_issues: List[ApkSecurityIssueResponse]
    ai_analysis: Optional[str] = None
    ai_report_functionality: Optional[str] = None  # "What does this APK do" report
    ai_report_security: Optional[str] = None  # Security findings report
    ai_architecture_diagram: Optional[str] = None  # AI-generated Mermaid architecture diagram
    ai_data_flow_diagram: Optional[str] = None  # AI-generated Mermaid data flow diagram
    error: Optional[str] = None


class DockerLayerResponse(BaseModel):
    """A Docker image layer."""
    id: str
    command: str
    size: int


class DockerSecretResponse(BaseModel):
    """A secret found in Docker layer."""
    layer_id: str
    layer_command: str
    secret_type: str
    value: str
    masked_value: str
    context: str
    severity: str


class DockerSecurityIssueResponse(BaseModel):
    """A security issue in Docker image."""
    category: str
    severity: str
    description: str
    command: Optional[str] = None
    attack_vector: Optional[str] = None  # Offensive security context
    rule_id: Optional[str] = None


class AdjudicationResult(BaseModel):
    """Result of AI false positive adjudication."""
    verdict: str  # "confirmed" or "false_positive"
    reason: str


class BaseImageIntelligenceResponse(BaseModel):
    """Base image intelligence finding."""
    image: str
    category: str  # eol, compromised, vulnerable, discouraged, typosquatting, untrusted
    severity: str
    message: str
    attack_vector: str = ""
    recommendation: str = ""


class LayerSecretResponse(BaseModel):
    """Secret found in image layer (deep scan)."""
    layer_id: str
    layer_index: int
    file_path: str
    file_type: str
    severity: str
    size_bytes: int
    is_deleted: bool  # True = file was "deleted" but still recoverable!
    content_preview: Optional[str] = None
    entropy: Optional[float] = None
    attack_vector: str = ""


class DockerCVEVulnerabilityResponse(BaseModel):
    """A CVE vulnerability found in a Docker image package."""
    external_id: str  # CVE ID (e.g., CVE-2021-44228)
    title: Optional[str] = None
    description: Optional[str] = None
    severity: Optional[str] = None
    cvss_score: Optional[float] = None
    cvss_vector: Optional[str] = None
    source: str = "osv"  # osv, nvd, etc.
    # Package context
    package_name: Optional[str] = None
    package_version: Optional[str] = None
    package_ecosystem: Optional[str] = None
    package_type: Optional[str] = None
    # Enrichment data
    in_kev: bool = False  # In CISA Known Exploited Vulnerabilities
    epss_score: Optional[float] = None  # Exploit Prediction Scoring System
    epss_percentile: Optional[float] = None
    epss_priority: Optional[str] = None
    combined_priority: Optional[float] = None
    priority_label: Optional[str] = None
    # NVD enrichment
    nvd_enrichment: Optional[Dict[str, Any]] = None


class DockerCVEScanResponse(BaseModel):
    """CVE scan results for a Docker image."""
    image_name: str
    packages_scanned: int
    packages_with_vulns: int
    vulnerabilities: List[DockerCVEVulnerabilityResponse]
    total_vulnerabilities: int = 0
    # Severity counts
    critical_count: int = 0
    high_count: int = 0
    medium_count: int = 0
    low_count: int = 0
    # Special indicators
    kev_count: int = 0  # CVEs in CISA KEV catalog
    high_epss_count: int = 0  # CVEs with EPSS > 0.5
    # Metadata
    scan_duration_seconds: float = 0.0
    trivy_available: bool = False
    enrichment_applied: bool = False
    epss_enabled: bool = False
    epss_available: bool = True
    epss_source: Optional[str] = None
    error: Optional[str] = None


class DockerTrivyNativeScanResponse(BaseModel):
    """Native Trivy scanner results for a Docker image."""
    image_name: str
    enabled_scanners: List[str] = []
    trivy_available: bool = False
    artifact_name: Optional[str] = None
    artifact_type: Optional[str] = None
    os_family: Optional[str] = None
    os_name: Optional[str] = None
    trivy_version: Optional[str] = None
    schema_version: Optional[int] = None
    package_count: int = 0
    vulnerabilities: List[Dict[str, Any]] = []
    misconfigurations: List[Dict[str, Any]] = []
    secrets: List[Dict[str, Any]] = []
    licenses: List[Dict[str, Any]] = []
    vulnerability_severity_counts: Dict[str, int] = {}
    misconfiguration_severity_counts: Dict[str, int] = {}
    secret_severity_counts: Dict[str, int] = {}
    license_severity_counts: Dict[str, int] = {}
    scan_duration_seconds: float = 0.0
    error: Optional[str] = None


class DockerAnalysisResponse(BaseModel):
    """Complete Docker image analysis response."""
    image_name: str
    image_id: str
    total_layers: int
    total_size: int
    total_size_human: str
    base_image: Optional[str] = None
    layers: List[DockerLayerResponse]
    secrets: List[DockerSecretResponse]
    deleted_files: List[Dict[str, Any]]
    security_issues: List[DockerSecurityIssueResponse]
    ai_analysis: Optional[str] = None
    error: Optional[str] = None
    # AI False Positive Adjudication
    adjudication_enabled: bool = False
    adjudication_summary: Optional[str] = None
    rejected_findings: List[Dict[str, Any]] = []
    adjudication_stats: Optional[Dict[str, int]] = None
    # Base Image Intelligence
    base_image_intel: List[BaseImageIntelligenceResponse] = []
    # Layer Deep Scan (recoverable secrets)
    layer_secrets: List[LayerSecretResponse] = []
    layer_scan_metadata: Optional[Dict[str, Any]] = None
    deleted_secrets_count: int = 0
    # CVE Scanning (NEW)
    cve_scan: Optional[DockerCVEScanResponse] = None
    # Native Trivy scanners
    trivy_scan: Optional[DockerTrivyNativeScanResponse] = None
    # Scan coverage summary
    scan_summary: Optional[Dict[str, Any]] = None


class StatusResponse(BaseModel):
    """Status of reverse engineering capabilities."""
    binary_analysis: bool
    apk_analysis: bool
    docker_analysis: bool
    jadx_available: bool
    ghidra_available: bool
    docker_available: bool
    message: str
    # Ghidra diagnostics (optional, included when Ghidra is not available)
    ghidra_path: Optional[str] = None
    ghidra_setup_instructions: Optional[str] = None
    # AI availability
    ai_available: bool = False
    ai_status: Optional[str] = None
    # Symbolic execution availability
    symbolic_execution_available: bool = False
    symbolic_execution_mode: Optional[str] = None  # "angr" or "lightweight"
    # Offline/air-gapped capability status
    offline_mode: bool = False
    osv_db_available: bool = False
    nvd_db_available: bool = False
    epss_available: bool = False
    epss_source: Optional[str] = None
    epss_last_sync: Optional[str] = None


# ============================================================================
# Helper Functions
# ============================================================================

def format_size(size_bytes: int) -> str:
    """Format bytes to human readable string."""
    if size_bytes >= 1024 * 1024 * 1024:
        return f"{size_bytes / (1024 * 1024 * 1024):.1f} GB"
    elif size_bytes >= 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.1f} MB"
    elif size_bytes >= 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes} B"


def check_docker_available() -> bool:
    """Check if Docker CLI is available."""
    import subprocess
    try:
        result = subprocess.run(["docker", "--version"], capture_output=True, timeout=5)
        return result.returncode == 0
    except (OSError, subprocess.TimeoutExpired, FileNotFoundError) as e:
        logger.debug(f"Docker check failed: {e}")
        return False


def _normalize_docker_image_name(image_name: str) -> str:
    normalized = image_name.strip()
    if not normalized:
        raise HTTPException(status_code=400, detail="Docker image name is required")
    return normalized


def _ensure_project_access(db: Session, project_id: int, user_id: int, require_edit: bool = False) -> None:
    project = db.get(Project, project_id)
    if not project:
        raise HTTPException(status_code=404, detail=f"Project {project_id} not found")

    allowed = (
        project_service.can_edit_project(db, project_id, user_id)
        if require_edit
        else project_service.can_access_project(db, project_id, user_id)[0]
    )
    if not allowed:
        raise HTTPException(status_code=403, detail="Access denied")


def _get_accessible_project_ids(db: Session, user_id: int) -> Set[int]:
    rows = (
        db.query(Project.id)
        .outerjoin(ProjectCollaborator, Project.id == ProjectCollaborator.project_id)
        .filter(
            or_(
                Project.owner_id == user_id,
                ProjectCollaborator.user_id == user_id,
            )
        )
        .distinct()
        .all()
    )
    return {row[0] for row in rows if row and row[0] is not None}


def _get_docker_allowlist_scope_map(
    db: Session,
    user_id: int,
    project_id: Optional[int] = None,
) -> Dict[str, str]:
    scope_map: Dict[str, str] = {}

    for row in db.query(UserDockerImage.image_name).filter(UserDockerImage.user_id == user_id).all():
        if not row or not row[0]:
            continue
        scope_map[_normalize_docker_image_name(row[0])] = "user"

    if project_id is not None:
        _ensure_project_access(db, project_id, user_id)
        for row in (
            db.query(ProjectDockerImage.image_name)
            .filter(ProjectDockerImage.project_id == project_id)
            .all()
        ):
            if not row or not row[0]:
                continue
            normalized_name = _normalize_docker_image_name(row[0])
            if normalized_name in scope_map:
                scope_map[normalized_name] = "user+project"
            else:
                scope_map[normalized_name] = "project"

    return scope_map


def _get_allowed_docker_image_names(db: Session, user_id: int, project_id: Optional[int] = None) -> Set[str]:
    return set(_get_docker_allowlist_scope_map(db, user_id, project_id).keys())


def _ensure_docker_image_allowlisted(
    db: Session,
    user_id: int,
    image_name: str,
    project_id: Optional[int] = None,
) -> str:
    normalized_name = _normalize_docker_image_name(image_name)
    allowed_names = _get_allowed_docker_image_names(db, user_id, project_id)
    if normalized_name not in allowed_names:
        raise HTTPException(
            status_code=403,
            detail="Docker image is not allowlisted for this user or project. Allow it before analysis.",
        )
    return normalized_name


def _list_local_docker_images_raw() -> List[Dict[str, str]]:
    import subprocess

    try:
        result = subprocess.run(
            ["docker", "images", "--format", "{{.Repository}}:{{.Tag}}|||{{.ID}}|||{{.Size}}|||{{.CreatedAt}}"],
            capture_output=True,
            text=True,
            timeout=30,
        )
    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=504, detail="Docker command timed out")
    except OSError as exc:
        logger.error(f"Failed to invoke Docker CLI: {exc}")
        raise HTTPException(status_code=503, detail="Docker is not available")

    if result.returncode != 0:
        raise HTTPException(status_code=500, detail="Failed to list Docker images")

    images: List[Dict[str, str]] = []
    for line in result.stdout.strip().split("\n"):
        if not line.strip():
            continue
        parts = line.split("|||")
        if len(parts) < 4:
            continue
        image_name = parts[0].strip()
        if not image_name:
            continue
        images.append(
            {
                "name": image_name,
                "id": parts[1].strip()[:12],
                "size": parts[2].strip(),
                "created": parts[3].strip(),
            }
        )
    return images


def _docker_image_exists(image_name: str) -> bool:
    import subprocess

    normalized_name = _normalize_docker_image_name(image_name)
    try:
        result = subprocess.run(
            ["docker", "image", "inspect", normalized_name],
            capture_output=True,
            text=True,
            timeout=20,
        )
        return result.returncode == 0
    except (OSError, subprocess.TimeoutExpired):
        return False


def _get_reverse_engineering_report_or_404(db: Session, report_id: int) -> ReverseEngineeringReport:
    report = db.query(ReverseEngineeringReport).filter(ReverseEngineeringReport.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    return report


def _ensure_reverse_engineering_report_access(
    db: Session,
    report: ReverseEngineeringReport,
    user_id: int,
    require_edit: bool = False,
) -> ReverseEngineeringReport:
    if report.project_id is not None:
        _ensure_project_access(db, report.project_id, user_id, require_edit=require_edit)
        return report
    if report.user_id != user_id:
        raise HTTPException(status_code=403, detail="Access denied")
    return report


def check_jadx_available() -> bool:
    """Check if jadx is available for APK decompilation."""
    import subprocess
    try:
        result = subprocess.run(["jadx", "--version"], capture_output=True, timeout=5)
        return result.returncode == 0
    except (OSError, subprocess.TimeoutExpired, FileNotFoundError) as e:
        logger.debug(f"JADX check failed: {e}")
        return False


# ============================================================================
# Endpoints
# ============================================================================

@router.get("/status", response_model=StatusResponse)
def get_status(
    current_user: User = Depends(get_current_active_user),
):
    """
    Check status of reverse engineering capabilities.
    """
    docker_available = check_docker_available()
    jadx_available = check_jadx_available()

    # Ghidra detection with diagnostics
    ghidra_ok = False
    ghidra_path = None
    ghidra_setup = None
    try:
        from backend.services.ghidra_service import ghidra_status
        gstatus = ghidra_status()
        ghidra_ok = gstatus["available"]
        ghidra_path = gstatus.get("path")
        ghidra_setup = gstatus.get("setup_instructions")
    except Exception:
        ghidra_ok = False

    # AI availability check
    ai_ok = bool(settings.gemini_api_key)
    ai_status_msg = None
    if not ai_ok:
        ai_status_msg = "GEMINI_API_KEY not configured. AI analysis features will be skipped."

    # Symbolic execution availability
    angr_available = getattr(re_service, 'ANGR_AVAILABLE', False)
    sym_mode = "angr" if angr_available else "lightweight"

    from backend.services import cve_service, nvd_service, epss_service
    osv_stats = cve_service.get_local_osv_db_stats()
    nvd_stats = nvd_service.get_local_db_stats()
    epss_stats = epss_service.get_local_epss_db_stats()

    return StatusResponse(
        binary_analysis=True,  # Always available (pure Python)
        apk_analysis=True,     # Basic analysis always available
        docker_analysis=docker_available,
        jadx_available=jadx_available,
        ghidra_available=ghidra_ok,
        docker_available=docker_available,
        message="Reverse engineering tools ready" if docker_available else "Docker not available - Docker analysis disabled",
        ghidra_path=ghidra_path,
        ghidra_setup_instructions=ghidra_setup,
        ai_available=ai_ok,
        ai_status=ai_status_msg,
        symbolic_execution_available=True,  # Always available (lightweight fallback)
        symbolic_execution_mode=sym_mode,
        offline_mode=settings.app_offline_mode,
        osv_db_available=bool(osv_stats.get("available")),
        nvd_db_available=bool(nvd_stats.get("available")),
        epss_available=bool(epss_stats.get("available")),
        epss_source="local_db" if epss_stats.get("available") else None,
        epss_last_sync=(
            epss_stats.get("score_date")
            or epss_stats.get("last_sync")
        ),
    )


@router.post("/analyze-binary", response_model=BinaryAnalysisResponse)
async def analyze_binary(
    file: UploadFile = File(..., description="Binary file to analyze (EXE, ELF, DLL, SO)"),
    include_ai: bool = Query(True, description="Include AI-powered analysis"),
    include_ghidra: bool = Query(True, description="Include Ghidra headless decompilation"),
    extended_ghidra: bool = Query(False, description="Extended Ghidra scan (2x functions and decompilation limit)"),
    ghidra_selection_mode: str = Query(
        "smart",
        description="Ghidra function selection mode: smart, security, size, or all",
    ),
    ghidra_max_functions: int = Query(500, ge=1, le=5000, description="Max functions to export from Ghidra"),
    ghidra_decomp_limit: int = Query(10000, ge=200, le=50000, description="Max decompilation chars per function"),
    include_ghidra_ai: bool = Query(True, description="Include Gemini summaries for decompiled functions"),
    ghidra_ai_max_functions: int = Query(30, ge=1, le=200, description="Max functions to summarize with Gemini"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Analyze a binary executable file.
    
    Extracts:
    - File metadata (type, architecture, entry point)
    - Strings (ASCII and UTF-16)
    - Imported functions (with suspicious API detection)
    - Potential secrets and credentials
    - Packer/obfuscation detection
    
    Supported formats: EXE, DLL, ELF, SO, Mach-O
    """
    # Validate file extension
    filename = file.filename or "unknown"
    suffix = Path(filename).suffix.lower()
    
    # Allow any binary file (we'll detect type from content)
    if suffix not in ALLOWED_BINARY_EXTENSIONS and suffix not in {".bin", ""}:
        # Still allow if no extension - could be ELF
        pass
    
    tmp_dir = None
    try:
        # Save file to temp location with sanitized filename to prevent path traversal
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_binary_"))
        safe_filename = sanitize_filename(filename)
        tmp_path = tmp_dir / safe_filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Analyzing binary: {filename} ({file_size:,} bytes)")
        
        # Perform analysis
        result = re_service.analyze_binary(tmp_path)

        # Run Ghidra decompilation if requested
        if include_ghidra:
            # Apply 2x multiplier for extended scan
            effective_max_functions = ghidra_max_functions * 2 if extended_ghidra else ghidra_max_functions
            effective_decomp_limit = ghidra_decomp_limit * 2 if extended_ghidra else ghidra_decomp_limit
            
            result.ghidra_analysis = re_service.analyze_binary_with_ghidra(
                tmp_path,
                max_functions=effective_max_functions,
                decomp_limit=effective_decomp_limit,
                selection_mode=ghidra_selection_mode,
            )

            if include_ghidra_ai and result.ghidra_analysis and "error" not in result.ghidra_analysis:
                result.ghidra_ai_summaries = await re_service.analyze_ghidra_functions_with_ai(
                    result.ghidra_analysis,
                    max_functions=ghidra_ai_max_functions,
                )

        # Run AI analysis if requested
        if include_ai and not result.error:
            result.ai_analysis = await re_service.analyze_binary_with_ai(result)
        
        # Convert to response model
        # Build rich header response if present
        rich_header_response = None
        if result.metadata.rich_header:
            rich_header_response = RichHeaderResponse(
                entries=[
                    RichHeaderEntryResponse(
                        product_id=e.product_id,
                        build_id=e.build_id,
                        count=e.count,
                        product_name=e.product_name,
                        vs_version=e.vs_version,
                    )
                    for e in result.metadata.rich_header.entries
                ],
                rich_hash=result.metadata.rich_header.rich_hash,
                checksum=result.metadata.rich_header.checksum,
                raw_data=result.metadata.rich_header.raw_data,
                clear_data=result.metadata.rich_header.clear_data,
            )
        
        return BinaryAnalysisResponse(
            filename=result.filename,
            metadata=BinaryMetadataResponse(
                file_type=result.metadata.file_type,
                architecture=result.metadata.architecture,
                file_size=result.metadata.file_size,
                entry_point=result.metadata.entry_point,
                is_packed=result.metadata.is_packed,
                packer_name=result.metadata.packer_name,
                compile_time=result.metadata.compile_time,
                sections=result.metadata.sections,
                headers=result.metadata.headers,
                # PE-specific
                rich_header=rich_header_response,
                imphash=result.metadata.imphash,
                tls_callbacks=result.metadata.tls_callbacks,
                mitigations=result.metadata.mitigations,
                resource_summary=result.metadata.resource_summary,
                version_info=result.metadata.version_info,
                authenticode=result.metadata.authenticode,
                overlay=result.metadata.overlay,
                pe_delay_imports=result.metadata.pe_delay_imports,
                pe_relocations=result.metadata.pe_relocations,
                pe_debug=result.metadata.pe_debug,
                pe_data_directories=result.metadata.pe_data_directories,
                pe_manifest=result.metadata.pe_manifest,
                # ELF-specific
                relro=result.metadata.relro,
                stack_canary=result.metadata.stack_canary,
                nx_enabled=result.metadata.nx_enabled,
                pie_enabled=result.metadata.pie_enabled,
                interpreter=result.metadata.interpreter,
                linked_libraries=result.metadata.linked_libraries,
                elf_dynamic=result.metadata.elf_dynamic,
                elf_relocations=result.metadata.elf_relocations,
                elf_version_info=result.metadata.elf_version_info,
                elf_build_id=result.metadata.elf_build_id,
                elf_program_headers=result.metadata.elf_program_headers,
            ),
            strings_count=len(result.strings),
            strings_sample=[
                BinaryStringResponse(
                    value=s.value[:500],
                    offset=s.offset,
                    encoding=s.encoding,
                    category=s.category,
                )
                for s in result.strings[:200]
            ],
            imports=[
                ImportedFunctionResponse(
                    name=imp.name,
                    library=imp.library,
                    ordinal=imp.ordinal,
                    is_suspicious=imp.is_suspicious,
                    reason=imp.reason,
                )
                for imp in result.imports
            ],
            exports=result.exports,
            secrets=[
                SecretResponse(
                    type=s["type"],
                    value=s["value"],
                    masked_value=s["masked_value"],
                    severity=s["severity"],
                    context=s.get("context"),
                    offset=s.get("offset"),
                )
                for s in result.secrets
            ],
            suspicious_indicators=[
                SuspiciousIndicatorResponse(
                    category=ind["category"],
                    severity=ind["severity"],
                    description=ind["description"],
                    details=ind.get("details"),
                )
                for ind in result.suspicious_indicators
            ],
            fuzzy_hashes=result.fuzzy_hashes,
            yara_matches=result.yara_matches,
            capa_summary=result.capa_summary,
            deobfuscated_strings=result.deobfuscated_strings,
            ai_analysis=result.ai_analysis,
            ghidra_analysis=result.ghidra_analysis,
            ghidra_ai_summaries=result.ghidra_ai_summaries,
            entropy_analysis=result.entropy_analysis,
            error=result.error,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Binary analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


@router.post("/binary/unified-scan")
async def unified_binary_scan(
    file: UploadFile = File(..., description="Binary file to analyze (EXE, ELF, DLL, SO)"),
    include_ai: bool = Query(True, description="Include AI-powered analysis"),
    include_ghidra: bool = Query(True, description="Include Ghidra headless decompilation"),
    extended_ghidra: bool = Query(False, description="Extended Ghidra scan (2x functions and decompilation limit)"),
    ghidra_selection_mode: str = Query(
        "smart",
        description="Ghidra function selection mode: smart, security, size, or all",
    ),
    ghidra_max_functions: int = Query(500, ge=1, le=5000, description="Max functions to export from Ghidra"),
    ghidra_decomp_limit: int = Query(10000, ge=200, le=50000, description="Max decompilation chars per function"),
    include_ghidra_ai: bool = Query(True, description="Include Gemini summaries for decompiled functions"),
    ghidra_ai_max_functions: int = Query(30, ge=1, le=200, description="Max functions to summarize with Gemini"),
    include_vuln_hunt: bool = Query(True, description="Include multi-pass AI vulnerability hunting (enabled by default)"),
    vuln_hunt_max_passes: int = Query(4, ge=1, le=6, description="Max vulnerability hunting passes"),
    vuln_hunt_max_targets: int = Query(50, ge=5, le=80, description="Max targets per hunting pass"),
    include_pattern_scan: bool = Query(True, description="Include pattern-based vulnerability scanning"),
    include_cve_lookup: bool = Query(True, description="Include CVE lookup for libraries"),
    include_sensitive_scan: bool = Query(True, description="Include sensitive data discovery"),
    include_unified_verification: bool = Query(True, description="Include unified AI verification of all findings"),
    include_symbolic: bool = Query(True, description="Include symbolic execution analysis (uses angr if available, otherwise lightweight pattern analysis)"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Perform a complete binary analysis with streaming progress updates.
    This unified scan combines up to 12 phases (matching APK analyzer capabilities):

    Phase 1: Static metadata, strings, imports, secrets
    Phase 2: Optional Ghidra decompilation
    Phase 3: Optional Gemini function summaries
    Phase 4: Optional Gemini overall analysis
    Phase 5: Pattern-based vulnerability scanning (80+ patterns)
    Phase 6: CVE lookup for libraries (OSV.dev + NVD)
    Phase 7: Sensitive data discovery (40+ patterns)
    Phase 8: Unified AI verification of all findings
    Phase 9: Multi-pass AI vulnerability hunting
    Phase 10: Symbolic execution / path analysis
    Phase 11: Final risk assessment
    Phase 12: Report generation
    """
    filename = file.filename or "unknown"
    suffix = Path(filename).suffix.lower()

    if suffix not in ALLOWED_BINARY_EXTENSIONS and suffix not in {".bin", ""}:
        pass

    tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_binary_unified_"))
    safe_filename = sanitize_filename(filename)
    tmp_path = tmp_dir / safe_filename
    file_size = 0

    try:
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)

        scan_id = str(uuid.uuid4())
        _unified_binary_scan_sessions[scan_id] = {"cancelled": False, "tmp_dir": str(tmp_dir)}
        _unified_binary_scan_timestamps[scan_id] = datetime.now()

        # Build phases list based on enabled options (up to 11 phases)
        phases: List[UnifiedBinaryScanPhase] = [
            UnifiedBinaryScanPhase(
                id="static",
                label="Static Analysis",
                description="Extract metadata, strings, imports, and secrets",
                status="pending",
            ),
        ]
        if include_ghidra:
            phases.append(UnifiedBinaryScanPhase(
                id="ghidra",
                label="Ghidra Decompilation",
                description="Run headless decompiler and export functions",
                status="pending",
            ))
        if include_ghidra and include_ghidra_ai:
            phases.append(UnifiedBinaryScanPhase(
                id="ghidra_ai",
                label="Ghidra AI Summaries",
                description="Summarize decompiled functions with Gemini",
                status="pending",
            ))
        if include_ai:
            phases.append(UnifiedBinaryScanPhase(
                id="ai_summary",
                label="AI Security Summary",
                description="Generate overall Gemini analysis",
                status="pending",
            ))
        # NEW PHASES - Pattern scan, CVE lookup, Sensitive data, Verification
        if include_pattern_scan and include_ghidra:
            phases.append(UnifiedBinaryScanPhase(
                id="pattern_scan",
                label="Pattern Vulnerability Scan",
                description="Scan decompiled code with 80+ vulnerability patterns",
                status="pending",
            ))
        if include_cve_lookup:
            phases.append(UnifiedBinaryScanPhase(
                id="cve_lookup",
                label="CVE Lookup",
                description="Query OSV.dev and NVD for library CVEs",
                status="pending",
            ))
        if include_sensitive_scan:
            phases.append(UnifiedBinaryScanPhase(
                id="sensitive_scan",
                label="Sensitive Data Discovery",
                description="Scan for secrets, credentials, and API keys (40+ patterns)",
                status="pending",
            ))
        # Vuln hunt runs BEFORE verification so its findings get verified
        if include_vuln_hunt and include_ghidra:
            phases.append(UnifiedBinaryScanPhase(
                id="vuln_hunt",
                label="AI Vulnerability Hunt",
                description=f"Multi-pass AI vulnerability hunting ({vuln_hunt_max_passes} passes)",
                status="pending",
            ))
        # Verification runs AFTER vuln_hunt to verify ALL findings (pattern + CVE + sensitive + vuln_hunt)
        if include_unified_verification and (include_pattern_scan or include_cve_lookup or include_sensitive_scan or include_vuln_hunt):
            phases.append(UnifiedBinaryScanPhase(
                id="ai_verification",
                label="AI Findings Verification",
                description="Unified AI verification to eliminate false positives and detect attack chains",
                status="pending",
            ))
        # APK-matching phases - these use verified findings
        phases.append(UnifiedBinaryScanPhase(
            id="advanced_analysis",
            label="Advanced Analysis",
            description="Obfuscation detection and packing analysis",
            status="pending",
        ))
        phases.append(UnifiedBinaryScanPhase(
            id="attack_surface",
            label="Attack Surface Mapping",
            description="Map entry points, exports, and attack vectors",
            status="pending",
        ))
        phases.append(UnifiedBinaryScanPhase(
            id="dynamic_scripts",
            label="Dynamic Analysis Scripts",
            description="Generate Frida hooks for runtime analysis",
            status="pending",
        ))
        phases.append(UnifiedBinaryScanPhase(
            id="emulation",
            label="Emulation Analysis",
            description="Emulate code with Unicorn for runtime behavior analysis",
            status="pending",
        ))
        if include_symbolic:
            phases.append(UnifiedBinaryScanPhase(
                id="symbolic_execution",
                label="Symbolic Execution",
                description="Discover vulnerability-triggering inputs via symbolic/pattern analysis",
                status="pending",
            ))
        phases.append(UnifiedBinaryScanPhase(
            id="ai_reports",
            label="AI Report Generation",
            description="Generate functionality, security, and architecture reports",
            status="pending",
        ))

        current_phase_idx = 0
        vuln_hunt_result = None  # Store vulnerability hunt results
        # NEW: Store results for new phases
        pattern_scan_result = None
        cve_lookup_result = None
        sensitive_scan_result = None
        verification_result = None
        # NEW: APK-matching results
        obfuscation_analysis = None
        attack_surface_result = None
        dynamic_analysis_result = None
        emulation_result = None  # Unicorn emulation results
        symbolic_execution_result = None  # Symbolic execution results
        ai_reports_result = None
        
        # Time tracking for estimates
        scan_start_time = datetime.utcnow()
        phase_times: Dict[str, float] = {}  # Track time per phase for estimates
        
        # Average phase durations (in seconds) based on typical scans
        PHASE_TIME_ESTIMATES = {
            "static": 5,
            "ghidra": 30,
            "ghidra_ai": 20,
            "ai_summary": 10,
            "pattern_scan": 8,
            "cve_lookup": 12,
            "sensitive_scan": 10,
            "vuln_hunt": 60,
            "ai_verification": 15,
            "advanced_analysis": 8,
            "attack_surface": 10,
            "dynamic_scripts": 5,
            "emulation": 25,
            "symbolic_execution": 15,
            "ai_reports": 30,
        }

        def make_progress(message: str, phase_progress: int = 0) -> str:
            nonlocal phases, current_phase_idx, scan_start_time
            overall = (current_phase_idx * 100 // max(len(phases), 1)) + (phase_progress // max(len(phases), 1))
            
            # Calculate time estimates
            elapsed = (datetime.utcnow() - scan_start_time).total_seconds()
            
            # Estimate remaining time based on phase estimates
            remaining_phases = [p for p in phases[current_phase_idx:] if p.status != "completed"]
            estimated_remaining = sum(
                PHASE_TIME_ESTIMATES.get(p.id, 10) * (1 - (p.progress / 100 if p.progress else 0))
                for p in remaining_phases
            )
            
            # Adjust based on actual elapsed time if we have data
            if current_phase_idx > 0 and elapsed > 0:
                # Calculate actual average phase time
                completed_phases = [p for p in phases if p.status == "completed"]
                if completed_phases:
                    estimated_total = elapsed + estimated_remaining
                else:
                    estimated_total = sum(PHASE_TIME_ESTIMATES.get(p.id, 10) for p in phases)
            else:
                estimated_total = sum(PHASE_TIME_ESTIMATES.get(p.id, 10) for p in phases)
            
            progress = UnifiedBinaryScanProgress(
                scan_id=scan_id,
                current_phase=phases[current_phase_idx].id,
                overall_progress=min(overall, 100),
                phases=phases,
                message=message,
                elapsed_seconds=round(elapsed, 1),
                estimated_remaining_seconds=round(max(0, estimated_remaining), 1),
                estimated_total_seconds=round(estimated_total, 1),
            )
            return f"data: {json.dumps({'type': 'progress', 'data': progress.model_dump()})}\n\n"

        def update_phase(phase_id: str, status: str, details: str = None, progress: int = 0):
            for p in phases:
                if p.id == phase_id:
                    p.status = status
                    p.progress = progress
                    if details:
                        p.details = details
                    if status == "in_progress" and not p.started_at:
                        p.started_at = datetime.utcnow().isoformat()
                    if status in ("completed", "error"):
                        p.completed_at = datetime.utcnow().isoformat()
                        p.progress = 100

        def get_phase_idx(phase_id: str) -> int:
            """Get the index of a phase by ID."""
            for i, p in enumerate(phases):
                if p.id == phase_id:
                    return i
            return len(phases) - 1

        async def run_unified_scan():
            nonlocal current_phase_idx, pattern_scan_result, cve_lookup_result, sensitive_scan_result, verification_result, obfuscation_analysis, attack_surface_result, dynamic_analysis_result, emulation_result, symbolic_execution_result, ai_reports_result
            result = None
            is_legitimate_software = False  # Will be detected from static analysis
            legitimacy_indicators = []
            
            try:
                if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                current_phase_idx = 0
                update_phase("static", "in_progress", progress=10)
                yield make_progress("Analyzing binary...", 10)
                result = re_service.analyze_binary(tmp_path)
                
                # ================================================================
                # LEGITIMACY DETECTION - centralized helper with PE/ELF/macOS signals
                # ================================================================
                is_legitimate_software, legitimacy_indicators = re_service.assess_binary_legitimacy(
                    result,
                    filename=filename,
                )

                if is_legitimate_software:
                    logger.info(f"Binary appears legitimate: {legitimacy_indicators[:5]}")
                elif legitimacy_indicators:
                    logger.info(f"Some legitimacy indicators found but not conclusive: {legitimacy_indicators}")
                
                update_phase(
                    "static",
                    "completed",
                    f"{len(result.strings)} strings, {len(result.imports)} imports" + 
                    (" (legitimate software detected)" if is_legitimate_software else ""),
                    100,
                )
                yield make_progress("Static analysis complete", 100)

                if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                # Apply 2x multiplier for extended Ghidra scan
                effective_max_functions = ghidra_max_functions * 2 if extended_ghidra else ghidra_max_functions
                effective_decomp_limit = ghidra_decomp_limit * 2 if extended_ghidra else ghidra_decomp_limit

                if include_ghidra:
                    current_phase_idx = 1
                    update_phase("ghidra", "in_progress", progress=10)
                    extended_label = " (extended)" if extended_ghidra else ""
                    yield make_progress(f"Running Ghidra decompilation{extended_label}...", 10)
                    result.ghidra_analysis = re_service.analyze_binary_with_ghidra(
                        tmp_path,
                        max_functions=effective_max_functions,
                        decomp_limit=effective_decomp_limit,
                        selection_mode=ghidra_selection_mode,
                    )
                    if result.ghidra_analysis and "error" in result.ghidra_analysis:
                        update_phase("ghidra", "error", result.ghidra_analysis.get("error"), 100)
                    else:
                        fn_total = (result.ghidra_analysis or {}).get("functions_total", 0)
                        fn_exported = (result.ghidra_analysis or {}).get("functions_exported", fn_total)
                        mode_used = (
                            (result.ghidra_analysis or {}).get("selection_mode_effective")
                            or (result.ghidra_analysis or {}).get("selection_mode")
                            or "unknown"
                        )
                        update_phase(
                            "ghidra",
                            "completed",
                            f"Exported {fn_exported}/{fn_total} functions (mode: {mode_used})",
                            100,
                        )
                    yield make_progress("Ghidra decompilation complete", 100)

                    if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                        yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                        yield "data: {\"type\":\"done\"}\n\n"
                        return

                if include_ghidra and include_ghidra_ai:
                    current_phase_idx = 2
                    update_phase("ghidra_ai", "in_progress", progress=10)
                    yield make_progress("Summarizing functions with Gemini...", 10)
                    if result and result.ghidra_analysis and "error" not in result.ghidra_analysis:
                        result.ghidra_ai_summaries = await re_service.analyze_ghidra_functions_with_ai(
                            result.ghidra_analysis,
                            max_functions=ghidra_ai_max_functions,
                        )
                    update_phase("ghidra_ai", "completed", "Function summaries generated", 100)
                    yield make_progress("Ghidra AI summaries complete", 100)

                    if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                        yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                        yield "data: {\"type\":\"done\"}\n\n"
                        return

                if include_ai and result and not result.error:
                    # Find the phase index for ai_summary
                    ai_phase_idx = next((i for i, p in enumerate(phases) if p.id == "ai_summary"), len(phases) - 1)
                    current_phase_idx = ai_phase_idx
                    update_phase("ai_summary", "in_progress", progress=10)
                    yield make_progress("Generating AI security summary...", 10)
                    result.ai_analysis = await re_service.analyze_binary_with_ai(result)
                    update_phase("ai_summary", "completed", "AI summary generated", 100)
                    yield make_progress("AI summary complete", 100)

                if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                # ================================================================
                # PARALLEL PHASES: Pattern Scan + CVE Lookup + Sensitive Data
                # These phases don't depend on each other, only on static/ghidra
                # Running them in parallel saves ~30% time
                # ================================================================
                parallel_phases_enabled = any([
                    include_pattern_scan and include_ghidra and result and result.ghidra_analysis,
                    include_cve_lookup and result,
                    include_sensitive_scan and result
                ])
                
                if parallel_phases_enabled:
                    yield make_progress("Running parallel scans (pattern + CVE + sensitive)...", 10)
                    
                    # Mark all parallel phases as in_progress
                    if include_pattern_scan and include_ghidra and result and result.ghidra_analysis:
                        update_phase("pattern_scan", "in_progress", "Running in parallel...", 20)
                    if include_cve_lookup and result:
                        update_phase("cve_lookup", "in_progress", "Running in parallel...", 20)
                    if include_sensitive_scan and result:
                        update_phase("sensitive_scan", "in_progress", "Running in parallel...", 20)
                    
                    # Define async tasks for parallel execution
                    async def run_pattern_scan():
                        if not (include_pattern_scan and include_ghidra and result and result.ghidra_analysis):
                            return None
                        try:
                            res = re_service.scan_decompiled_binary_comprehensive(
                                result.ghidra_analysis, 
                                is_legitimate_software=is_legitimate_software
                            )
                            findings_count = len(res.get("findings", []))
                            severity_label = " (filtered for legitimate software)" if is_legitimate_software else ""
                            update_phase("pattern_scan", "completed", f"Found {findings_count} potential vulnerabilities{severity_label}", 100)
                            return res
                        except Exception as e:
                            logger.error(f"Pattern scan failed: {e}")
                            update_phase("pattern_scan", "error", str(e), 100)
                            return {"findings": [], "error": str(e)}
                    
                    async def run_cve_lookup():
                        if not (include_cve_lookup and result):
                            return None
                        try:
                            binary_metadata = {
                                "file_type": result.metadata.file_type,
                                "architecture": result.metadata.architecture,
                                "is_packed": result.metadata.is_packed,
                                "linked_libraries": getattr(result.metadata, 'linked_libraries', []),
                                "imports": [
                                    {"name": imp.name, "library": imp.library}
                                    for imp in result.imports
                                ] if result.imports else []
                            }
                            strings_list = [
                                {"value": s.value, "category": s.category}
                                for s in result.strings[:1000]
                            ] if result.strings else []
                            
                            res = await re_service.comprehensive_binary_cve_scan(binary_metadata, strings_list)
                            cve_count = len(res.get("findings", []))
                            update_phase("cve_lookup", "completed", f"Found {cve_count} CVEs", 100)
                            return res
                        except Exception as e:
                            logger.error(f"CVE lookup failed: {e}")
                            update_phase("cve_lookup", "error", str(e), 100)
                            return {"findings": [], "error": str(e)}
                    
                    async def run_sensitive_scan():
                        if not (include_sensitive_scan and result):
                            return None
                        try:
                            strings_list = [
                                {"value": s.value, "category": s.category, "offset": s.offset}
                                for s in result.strings
                            ] if result.strings else []
                            
                            decompiled_code = result.ghidra_analysis if include_ghidra and result.ghidra_analysis else None
                            
                            # Enable AI verification to match APK scanner behavior
                            res = await re_service.comprehensive_binary_sensitive_scan(
                                strings=strings_list,
                                decompiled_code=decompiled_code,
                                verify_with_ai=include_ai  # Now uses AI like APK scanner!
                            )
                            secrets_count = len(res.get("findings", []))
                            ai_note = " (AI verified)" if include_ai else ""
                            update_phase("sensitive_scan", "completed", f"Found {secrets_count} secrets{ai_note}", 100)
                            return res
                        except Exception as e:
                            logger.error(f"Sensitive data scan failed: {e}")
                            update_phase("sensitive_scan", "error", str(e), 100)
                            return {"findings": [], "error": str(e)}
                    
                    # Run all three in parallel
                    parallel_results = await asyncio.gather(
                        run_pattern_scan(),
                        run_cve_lookup(),
                        run_sensitive_scan(),
                        return_exceptions=True
                    )
                    
                    # Unpack results
                    pattern_scan_result = parallel_results[0] if not isinstance(parallel_results[0], Exception) else None
                    cve_lookup_result = parallel_results[1] if not isinstance(parallel_results[1], Exception) else None
                    sensitive_scan_result = parallel_results[2] if not isinstance(parallel_results[2], Exception) else None
                    
                    # Log any exceptions
                    for i, r in enumerate(parallel_results):
                        if isinstance(r, Exception):
                            logger.error(f"Parallel phase {i} failed with exception: {r}")
                    
                    # Update phase index to after parallel phases
                    current_phase_idx = get_phase_idx("sensitive_scan")
                    
                    # Summary of parallel phase results
                    pattern_count = len(pattern_scan_result.get("findings", [])) if pattern_scan_result else 0
                    cve_count = len(cve_lookup_result.get("findings", [])) if cve_lookup_result else 0
                    secrets_count = len(sensitive_scan_result.get("findings", [])) if sensitive_scan_result else 0
                    
                    yield make_progress(
                        f"Parallel scans complete: {pattern_count} patterns, {cve_count} CVEs, {secrets_count} secrets",
                        100
                    )

                if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                # ================================================================
                # Multi-pass AI Vulnerability Hunt (runs BEFORE verification!)
                # OPTIMIZATION: Skip for legitimate software to save time
                # ================================================================
                if include_vuln_hunt and include_ghidra and result and not result.error:
                    vuln_phase_idx = next((i for i, p in enumerate(phases) if p.id == "vuln_hunt"), len(phases) - 1)
                    current_phase_idx = vuln_phase_idx
                    
                    # SKIP VulnHuntr for legitimate software - saves significant time
                    if is_legitimate_software:
                        skip_reason = "Skipped for legitimate software (false positive reduction)"
                        update_phase("vuln_hunt", "completed", skip_reason, 100)
                        yield make_progress(skip_reason, 100)
                        logger.info(f"Skipping VulnHuntr for legitimate software: {legitimacy_indicators[:3]}")
                    else:
                        update_phase("vuln_hunt", "in_progress", progress=5)
                        yield make_progress("Starting multi-pass AI vulnerability hunt...", 5)
                        
                        try:
                            # Progress queue for streaming updates to client
                            progress_queue = asyncio.Queue()
                            
                            async def vuln_progress(phase: str, progress: int, message: str):
                                update_phase("vuln_hunt", "in_progress", message, progress)
                                # Also put in queue for streaming to client
                                await progress_queue.put((phase, progress, message))
                            
                            # Create background task for progress streaming
                            async def progress_streamer():
                                """Stream progress updates to client."""
                                while True:
                                    try:
                                        phase, prog, msg = await asyncio.wait_for(progress_queue.get(), timeout=0.5)
                                        # Note: We can't yield from here, but updates are via update_phase
                                    except asyncio.TimeoutError:
                                        continue
                                    except Exception:
                                        break
                            
                            nonlocal vuln_hunt_result
                            vuln_hunt_result = await re_service.ai_vulnerability_hunt(
                                tmp_path,
                                focus_categories=None,
                                max_passes=vuln_hunt_max_passes,
                                max_targets_per_pass=vuln_hunt_max_targets,
                                ghidra_max_functions=ghidra_max_functions,
                                ghidra_decomp_limit=ghidra_decomp_limit,
                                on_progress=vuln_progress,
                            )
                            vulns_found = len(vuln_hunt_result.vulnerabilities) if vuln_hunt_result else 0
                            update_phase("vuln_hunt", "completed", f"Found {vulns_found} vulnerabilities", 100)
                            yield make_progress(f"Vulnerability hunt complete: {vulns_found} found", 100)
                        except Exception as vuln_err:
                            logger.error(f"Vulnerability hunt failed: {vuln_err}")
                            update_phase("vuln_hunt", "error", str(vuln_err), 100)
                            yield make_progress(f"Vulnerability hunt failed: {vuln_err}", 100)

                    if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                        yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                        yield "data: {\"type\":\"done\"}\n\n"
                        return

                # ================================================================
                # Unified AI Verification (runs AFTER vuln_hunt to verify ALL findings)
                # ================================================================
                has_findings = (
                    (pattern_scan_result and pattern_scan_result.get("findings")) or
                    (cve_lookup_result and cve_lookup_result.get("findings")) or
                    (sensitive_scan_result and sensitive_scan_result.get("findings")) or
                    (vuln_hunt_result and vuln_hunt_result.vulnerabilities)
                )
                
                if include_unified_verification and has_findings:
                    current_phase_idx = get_phase_idx("ai_verification")
                    update_phase("ai_verification", "in_progress", progress=10)
                    yield make_progress("Running unified AI verification (includes vuln_hunt findings)...", 10)
                    
                    try:
                        pattern_findings = pattern_scan_result.get("findings", []) if pattern_scan_result else []
                        cve_findings = cve_lookup_result.get("findings", []) if cve_lookup_result else []
                        sensitive_findings = sensitive_scan_result.get("findings", []) if sensitive_scan_result else []
                        
                        # Convert vuln_hunt findings to dict format for unified verification
                        vuln_hunt_findings = []
                        if vuln_hunt_result and vuln_hunt_result.vulnerabilities:
                            for v in vuln_hunt_result.vulnerabilities:
                                vuln_hunt_findings.append({
                                    "id": v.id,
                                    "title": v.title,
                                    "severity": v.severity,
                                    "category": v.category,
                                    "cwe_id": v.cwe_id,
                                    "cvss_estimate": v.cvss_estimate,
                                    "function_name": v.function_name,
                                    "entry_address": v.entry_address,
                                    "description": v.description,
                                    "technical_details": v.technical_details,
                                    "proof_of_concept": v.proof_of_concept,
                                    "exploitation_steps": v.exploitation_steps,
                                    "remediation": v.remediation,
                                    "confidence": v.confidence,
                                    "ai_reasoning": v.ai_reasoning,
                                    "code_snippet": v.code_snippet,
                                    "source": "vuln_hunt",  # Tag source for filtering
                                })
                        
                        binary_meta = {
                            "file_type": result.metadata.file_type,
                            "architecture": result.metadata.architecture,
                            "is_packed": result.metadata.is_packed,
                            "mitigations": result.metadata.mitigations,
                        } if result else {}
                        
                        decompiled_code = result.ghidra_analysis if include_ghidra and result.ghidra_analysis else None
                        
                        verification_result = await re_service.verify_binary_findings_unified(
                            pattern_findings=pattern_findings,
                            cve_findings=cve_findings,
                            sensitive_findings=sensitive_findings,
                            vuln_hunt_findings=vuln_hunt_findings,  # NEW: Include vuln_hunt
                            decompiled_code=decompiled_code,
                            binary_metadata=binary_meta,
                            is_legitimate_software=is_legitimate_software,
                            legitimacy_indicators=legitimacy_indicators
                        )
                        
                        verified_total = verification_result.get("summary", {}).get("verified_total", 0)
                        filtered_total = verification_result.get("summary", {}).get("filtered_total", 0)
                        attack_chains = len(verification_result.get("attack_chains", []))
                        
                        details = f"Verified {verified_total}, filtered {filtered_total} FPs"
                        if attack_chains > 0:
                            details += f", {attack_chains} attack chains detected"
                        
                        update_phase("ai_verification", "completed", details, 100)
                        yield make_progress(f"AI verification complete: {details}", 100)
                    except Exception as verify_err:
                        logger.error(f"AI verification failed: {verify_err}")
                        update_phase("ai_verification", "error", str(verify_err), 100)
                        yield make_progress(f"AI verification failed: {verify_err}", 100)

                    if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                        yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                        yield "data: {\"type\":\"done\"}\n\n"
                        return

                # ================================================================
                # NEW PHASE: Advanced Analysis (Obfuscation/Packing Detection)
                # ================================================================
                current_phase_idx = get_phase_idx("advanced_analysis")
                update_phase("advanced_analysis", "in_progress", progress=10)
                yield make_progress("Analyzing obfuscation and packing...", 10)
                
                try:
                    obfuscation_analysis = re_service.analyze_binary_obfuscation(
                        str(tmp_path),
                        static_result=result,
                        ghidra_result=result.ghidra_analysis if include_ghidra else None
                    )
                    obfusc_level = obfuscation_analysis.get("overall_obfuscation_level", "none")
                    obfusc_score = obfuscation_analysis.get("obfuscation_score", 0)
                    packers = [p["name"] for p in obfuscation_analysis.get("detected_packers", [])]
                    
                    details = f"Level: {obfusc_level}, Score: {obfusc_score}/100"
                    if packers:
                        details += f", Packers: {', '.join(packers)}"
                    
                    update_phase("advanced_analysis", "completed", details, 100)
                    yield make_progress(f"Obfuscation analysis complete: {obfusc_level}", 100)
                except Exception as obfusc_err:
                    logger.error(f"Obfuscation analysis failed: {obfusc_err}")
                    update_phase("advanced_analysis", "error", str(obfusc_err), 100)

                if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                # ================================================================
                # NEW PHASE: Attack Surface Mapping
                # ================================================================
                current_phase_idx = get_phase_idx("attack_surface")
                update_phase("attack_surface", "in_progress", progress=10)
                yield make_progress("Mapping attack surface...", 10)
                
                try:
                    attack_surface_result = re_service.generate_binary_attack_surface(
                        static_result=result,
                        ghidra_result=result.ghidra_analysis if include_ghidra else None,
                        binary_path=str(tmp_path)
                    )
                    summary = attack_surface_result.get("summary", {})
                    entry_points = summary.get("total_entry_points", 0)
                    vectors = summary.get("total_attack_vectors", 0)
                    overall_risk = summary.get("overall_risk", "unknown")
                    
                    update_phase("attack_surface", "completed", 
                        f"{entry_points} entry points, {vectors} attack vectors ({overall_risk} risk)", 100)
                    yield make_progress(f"Attack surface mapped: {entry_points} entry points", 100)
                except Exception as attack_err:
                    logger.error(f"Attack surface mapping failed: {attack_err}")
                    update_phase("attack_surface", "error", str(attack_err), 100)

                if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                # ================================================================
                # NEW PHASE: Dynamic Analysis Scripts (Frida)
                # ================================================================
                current_phase_idx = get_phase_idx("dynamic_scripts")
                update_phase("dynamic_scripts", "in_progress", progress=10)
                yield make_progress("Generating Frida scripts...", 10)
                
                try:
                    # Get verified findings for vulnerability-specific hooks
                    verified_for_frida = []
                    if verification_result:
                        verified_for_frida = verification_result.get("verified_vulnerabilities", [])
                    elif pattern_scan_result:
                        verified_for_frida = pattern_scan_result.get("findings", [])
                    
                    # Include vuln_hunt findings for targeted hooks
                    vuln_hunt_findings = []
                    if vuln_hunt_result and vuln_hunt_result.vulnerabilities:
                        vuln_hunt_findings = [
                            {
                                "title": v.title,
                                "severity": v.severity,
                                "function_name": v.function_name,
                                "category": v.vuln_type,
                                "description": v.description,
                            }
                            for v in vuln_hunt_result.vulnerabilities
                        ]
                    
                    dynamic_analysis_result = re_service.generate_binary_frida_scripts(
                        binary_name=filename,
                        static_result=result,
                        ghidra_result=result.ghidra_analysis if include_ghidra else None,
                        obfuscation_result=obfuscation_analysis,
                        verified_findings=verified_for_frida,
                        vuln_hunt_findings=vuln_hunt_findings,
                        attack_surface_result=attack_surface_result
                    )
                    scripts_count = dynamic_analysis_result.get("total_scripts", 0)
                    categories = list(dynamic_analysis_result.get("categories", {}).keys())
                    
                    # Build protection detection summary
                    protections_found = []
                    if dynamic_analysis_result.get("anti_debug_detected"):
                        protections_found.append("Anti-Debug")
                    if dynamic_analysis_result.get("anti_vm_detected"):
                        protections_found.append("Anti-VM")
                    if dynamic_analysis_result.get("anti_tampering_detected"):
                        protections_found.append("Anti-Tampering")
                    if dynamic_analysis_result.get("packing_detected"):
                        protections_found.append("Packed")
                    
                    protection_info = ""
                    if protections_found:
                        protection_info = f" | Protections: {', '.join(protections_found)}"
                    
                    update_phase("dynamic_scripts", "completed", 
                        f"Generated {scripts_count} scripts ({', '.join(categories)}){protection_info}", 100)
                    yield make_progress(f"Generated {scripts_count} Frida scripts{protection_info}", 100)
                except Exception as frida_err:
                    logger.error(f"Frida script generation failed: {frida_err}")
                    update_phase("dynamic_scripts", "error", str(frida_err), 100)

                if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                # ================================================================
                # NEW PHASE: Emulation Analysis (Unicorn-based)
                # ================================================================
                current_phase_idx = get_phase_idx("emulation")
                update_phase("emulation", "in_progress", progress=5)
                yield make_progress("Running enhanced emulation analysis...", 5)
                
                try:
                    # Read binary data for emulation
                    with open(tmp_path, 'rb') as f:
                        binary_data = f.read()
                    
                    yield make_progress("Parsing PE/ELF sections for proper mapping...", 10)
                    
                    # Get entry point from static analysis
                    entry_point = None
                    if result and result.metadata:
                        ep = getattr(result.metadata, 'entry_point', None)
                        if ep:
                            entry_point = ep
                    
                    yield make_progress("Running emulation with API hooks and evasion detection...", 25)
                    
                    # Extract additional entry points from attack surface for multi-path emulation
                    additional_entry_points = []
                    if attack_surface_result:
                        entry_points_data = attack_surface_result.get("entry_points", [])
                        for ep in entry_points_data[:10]:  # Limit to top 10 entry points
                            if not isinstance(ep, dict):
                                continue
                            address = ep.get("address")
                            if address and address != entry_point:
                                additional_entry_points.append(address)
                    
                    # Run ENHANCED emulation with AI verification
                    emulation_result = await re_service.run_enhanced_emulation_with_verification(
                        binary_data=binary_data,
                        architecture=getattr(result.metadata, 'architecture', 'x86') if result and result.metadata else 'x86',
                        ghidra_result=result.ghidra_analysis if include_ghidra else None,
                        static_result=result,
                        binary_name=filename,
                        base_address=0x400000,
                        entry_point=entry_point,
                        additional_entry_points=additional_entry_points,
                        verify_with_ai=include_ai  # Use AI verification if AI is enabled
                    )

                    # Some emulation paths may return None; normalize to a failed result object.
                    if not isinstance(emulation_result, dict):
                        emulation_result = {
                            "success": False,
                            "errors": ["No emulation result returned"],
                        }
                    
                    yield make_progress("Detecting evasion techniques and malicious patterns...", 60)
                    
                    if emulation_result.get("ai_verification"):
                        yield make_progress("AI verifying emulation findings...", 80)
                    
                    yield make_progress("Compiling emulation results...", 95)
                    
                    if emulation_result.get("success"):
                        summary = emulation_result.get("summary") or {}
                        verdict = emulation_result.get("final_verdict") or {}
                        if not isinstance(summary, dict):
                            summary = {}
                        if not isinstance(verdict, dict):
                            verdict = {}
                        
                        detail_parts = []
                        if summary.get("strings_recovered"):
                            detail_parts.append(f"{summary['strings_recovered']} strings")
                        if summary.get("evasion_techniques"):
                            detail_parts.append(f"{summary['evasion_techniques']} evasion techniques")
                        if summary.get("malicious_patterns"):
                            detail_parts.append(f"{summary['malicious_patterns']} malicious patterns")
                        if summary.get("is_packed"):
                            detail_parts.append("PACKED")
                        if verdict.get("verdict"):
                            v = verdict["verdict"].upper()
                            conf = verdict.get("confidence", 0)
                            detail_parts.append(f"Verdict: {v} ({conf}%)")
                        if summary.get("ai_verified"):
                            detail_parts.append("✓ AI Verified")
                        
                        update_phase("emulation", "completed", 
                            " | ".join(detail_parts) if detail_parts else "Emulation complete", 100)
                        yield make_progress(f"Emulation complete: {' | '.join(detail_parts)}", 100)
                    else:
                        errors = emulation_result.get("errors") or []
                        if not isinstance(errors, list):
                            errors = [str(errors)]
                        error_msg = errors[0] if errors else "Emulation failed"
                        update_phase("emulation", "completed", error_msg, 100)
                        yield make_progress(f"Emulation: {error_msg}", 100)
                except Exception as emu_err:
                    logger.error(f"Emulation analysis failed: {emu_err}")
                    update_phase("emulation", "error", str(emu_err), 100)

                if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                # ================================================================
                # NEW PHASE: Symbolic Execution / Path Analysis
                # ================================================================
                if include_symbolic:
                    current_phase_idx = get_phase_idx("symbolic_execution")
                    update_phase("symbolic_execution", "in_progress", progress=5)

                    angr_available = getattr(re_service, 'ANGR_AVAILABLE', False)
                    mode_label = "angr symbolic execution" if angr_available else "lightweight path analysis"
                    yield make_progress(f"Running {mode_label}...", 10)

                    try:
                        if angr_available:
                            yield make_progress("Exploring paths with angr constraint solver...", 20)
                            symbolic_execution_result = re_service.perform_symbolic_execution(
                                tmp_path, max_time_seconds=60, max_paths=500
                            )
                        else:
                            yield make_progress("Analyzing taint flows and dangerous sink reachability...", 20)
                            symbolic_execution_result = re_service.perform_lightweight_symbolic_analysis(
                                tmp_path,
                                result=result.disassembly if result else None,
                                imports=list(result.imports) if result else None,
                                is_legitimate_software=is_legitimate_software,
                            )

                        yield make_progress("Analyzing symbolic execution results...", 70)

                        if symbolic_execution_result:
                            vulns = symbolic_execution_result.vulnerabilities_found
                            reaches = symbolic_execution_result.target_reaches
                            crashes = symbolic_execution_result.crash_inputs
                            paths = symbolic_execution_result.paths_explored

                            detail_parts = [f"{paths} paths explored"]
                            if vulns:
                                detail_parts.append(f"{len(vulns)} vulnerabilities")
                            if reaches:
                                detail_parts.append(f"{len(reaches)} targets reached")
                            if crashes:
                                detail_parts.append(f"{len(crashes)} crash inputs")
                            if not angr_available:
                                detail_parts.append("(lightweight mode)")

                            update_phase("symbolic_execution", "completed",
                                " | ".join(detail_parts), 100)
                            yield make_progress(f"Symbolic execution complete: {' | '.join(detail_parts)}", 100)
                        else:
                            update_phase("symbolic_execution", "completed",
                                "No results (binary may not be suitable for symbolic execution)", 100)
                            yield make_progress("Symbolic execution: no results", 100)

                    except Exception as sym_err:
                        logger.error(f"Symbolic execution failed: {sym_err}")
                        update_phase("symbolic_execution", "error", str(sym_err), 100)

                    if _unified_binary_scan_sessions.get(scan_id, {}).get("cancelled"):
                        yield f"data: {json.dumps({'type': 'error', 'error': 'Scan cancelled'})}\n\n"
                        yield "data: {\"type\":\"done\"}\n\n"
                        return

                # ================================================================
                # NEW PHASE: AI Report Generation (4 reports like APK)
                # ================================================================
                current_phase_idx = get_phase_idx("ai_reports")
                update_phase("ai_reports", "in_progress", progress=10)
                yield make_progress("Generating AI reports...", 10)
                
                try:
                    # Get verified findings and CVEs for reports
                    verified_for_reports = []
                    cve_for_reports = []
                    
                    if verification_result:
                        verified_for_reports = verification_result.get("verified_vulnerabilities", [])
                        cve_for_reports = verification_result.get("verified_cves", [])
                    else:
                        if pattern_scan_result:
                            verified_for_reports = pattern_scan_result.get("findings", [])
                        if cve_lookup_result:
                            cve_for_reports = cve_lookup_result.get("findings", [])
                    
                    yield make_progress("Generating functionality report...", 25)
                    ai_reports_result = await re_service.generate_binary_ai_reports(
                        binary_name=filename,
                        static_result=result,
                        ghidra_result=result.ghidra_analysis if include_ghidra else None,
                        attack_surface=attack_surface_result,
                        obfuscation_result=obfuscation_analysis,
                        verified_findings=verified_for_reports,
                        cve_findings=cve_for_reports,
                        emulation_result=emulation_result,
                        is_legitimate_software=is_legitimate_software,
                        legitimacy_indicators=legitimacy_indicators
                    )
                    
                    reports_generated = sum([
                        1 if ai_reports_result.get("functionality_report") else 0,
                        1 if ai_reports_result.get("security_report") else 0,
                        1 if ai_reports_result.get("architecture_diagram") else 0,
                        1 if ai_reports_result.get("attack_surface_map") else 0,
                    ])
                    
                    update_phase("ai_reports", "completed", f"Generated {reports_generated}/4 AI reports", 100)
                    yield make_progress(f"AI reports complete: {reports_generated}/4 generated", 100)
                except Exception as report_err:
                    logger.error(f"AI report generation failed: {report_err}")
                    update_phase("ai_reports", "error", str(report_err), 100)

                if not result:
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Analysis failed'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return

                rich_header_response = None
                if result.metadata.rich_header:
                    rich_header_response = RichHeaderResponse(
                        entries=[
                            RichHeaderEntryResponse(
                                product_id=e.product_id,
                                build_id=e.build_id,
                                count=e.count,
                                product_name=e.product_name,
                                vs_version=e.vs_version,
                            )
                            for e in result.metadata.rich_header.entries
                        ],
                        rich_hash=result.metadata.rich_header.rich_hash,
                        checksum=result.metadata.rich_header.checksum,
                        raw_data=result.metadata.rich_header.raw_data,
                        clear_data=result.metadata.rich_header.clear_data,
                    )

                response = BinaryAnalysisResponse(
                    filename=result.filename,
                    metadata=BinaryMetadataResponse(
                        file_type=result.metadata.file_type,
                        architecture=result.metadata.architecture,
                        file_size=result.metadata.file_size,
                        entry_point=result.metadata.entry_point,
                        is_packed=result.metadata.is_packed,
                        packer_name=result.metadata.packer_name,
                        compile_time=result.metadata.compile_time,
                        sections=result.metadata.sections,
                        headers=result.metadata.headers,
                        rich_header=rich_header_response,
                        imphash=result.metadata.imphash,
                        tls_callbacks=result.metadata.tls_callbacks,
                        mitigations=result.metadata.mitigations,
                        resource_summary=result.metadata.resource_summary,
                        version_info=result.metadata.version_info,
                        authenticode=result.metadata.authenticode,
                        overlay=result.metadata.overlay,
                        pe_delay_imports=result.metadata.pe_delay_imports,
                        pe_relocations=result.metadata.pe_relocations,
                        pe_debug=result.metadata.pe_debug,
                        pe_data_directories=result.metadata.pe_data_directories,
                        pe_manifest=result.metadata.pe_manifest,
                        relro=result.metadata.relro,
                        stack_canary=result.metadata.stack_canary,
                        nx_enabled=result.metadata.nx_enabled,
                        pie_enabled=result.metadata.pie_enabled,
                        interpreter=result.metadata.interpreter,
                        linked_libraries=result.metadata.linked_libraries,
                        elf_dynamic=result.metadata.elf_dynamic,
                        elf_relocations=result.metadata.elf_relocations,
                        elf_version_info=result.metadata.elf_version_info,
                        elf_build_id=result.metadata.elf_build_id,
                        elf_program_headers=result.metadata.elf_program_headers,
                    ),
                    strings_count=len(result.strings),
                    strings_sample=[
                        BinaryStringResponse(
                            value=s.value[:500],
                            offset=s.offset,
                            encoding=s.encoding,
                            category=s.category,
                        )
                        for s in result.strings[:200]
                    ],
                    imports=[
                        ImportedFunctionResponse(
                            name=imp.name,
                            library=imp.library,
                            ordinal=imp.ordinal,
                            is_suspicious=imp.is_suspicious,
                            reason=imp.reason,
                        )
                        for imp in result.imports
                    ],
                    exports=result.exports,
                    secrets=[
                        SecretResponse(
                            type=s["type"],
                            value=s["value"],
                            masked_value=s["masked_value"],
                            severity=s["severity"],
                            context=s.get("context"),
                            offset=s.get("offset"),
                        )
                        for s in result.secrets
                    ],
                    # Filter suspicious indicators for legitimate software
                    suspicious_indicators=[
                        SuspiciousIndicatorResponse(
                            category=ind["category"],
                            severity=ind["severity"],
                            description=ind["description"],
                            details=ind.get("details"),
                        )
                        for ind in result.suspicious_indicators
                        # For legitimate software, only show info-level and legitimacy indicators
                        # Filter out alarmist "suspicious API" warnings that are false positives
                        if not is_legitimate_software or (
                            ind.get("severity") == "info" or 
                            "Legitimate" in ind.get("category", "") or
                            "Security Features" in ind.get("category", "") or
                            "Cryptographic" in ind.get("category", "") or
                            ind.get("severity") == "high"  # Only show genuinely high severity
                        )
                    ],
                    fuzzy_hashes=result.fuzzy_hashes,
                    yara_matches=result.yara_matches,
                    capa_summary=result.capa_summary,
                    deobfuscated_strings=result.deobfuscated_strings,
                    ai_analysis=result.ai_analysis,
                    ghidra_analysis=result.ghidra_analysis,
                    ghidra_ai_summaries=result.ghidra_ai_summaries,
                    vuln_hunt_result=re_service.vulnerability_hunt_result_to_dict(vuln_hunt_result) if vuln_hunt_result else None,
                    # NEW: Include enhanced analysis results
                    pattern_scan_result=pattern_scan_result,
                    cve_lookup_result=cve_lookup_result,
                    sensitive_scan_result=sensitive_scan_result,
                    verification_result=verification_result,
                    attack_chains=verification_result.get("attack_chains", []) if verification_result else None,
                    prioritized_actions=verification_result.get("prioritized_actions", []) if verification_result else None,
                    # NEW: APK-matching features
                    obfuscation_analysis=obfuscation_analysis,
                    attack_surface=attack_surface_result,
                    dynamic_analysis=dynamic_analysis_result,
                    emulation_analysis=emulation_result,
                    symbolic_execution=_serialize_symbolic_result(symbolic_execution_result) if symbolic_execution_result else None,
                    ai_functionality_report=ai_reports_result.get("functionality_report") if ai_reports_result else None,
                    ai_security_report=ai_reports_result.get("security_report") if ai_reports_result else None,
                    ai_architecture_diagram=ai_reports_result.get("architecture_diagram") if ai_reports_result else None,
                    ai_attack_surface_map=ai_reports_result.get("attack_surface_map") if ai_reports_result else None,
                    # Legitimacy detection
                    is_legitimate_software=is_legitimate_software,
                    legitimacy_indicators=legitimacy_indicators if is_legitimate_software else None,
                    error=result.error,
                )

                yield f"data: {json.dumps({'type': 'result', 'data': response.model_dump()})}\n\n"
                yield "data: {\"type\":\"done\"}\n\n"

            except Exception as exc:
                logger.error(f"Unified binary scan failed: {exc}")
                yield f"data: {json.dumps({'type': 'error', 'error': str(exc)})}\n\n"
                yield "data: {\"type\":\"done\"}\n\n"
            finally:
                if scan_id in _unified_binary_scan_sessions:
                    del _unified_binary_scan_sessions[scan_id]
                if tmp_dir:
                    shutil.rmtree(tmp_dir, ignore_errors=True)

        return StreamingResponse(
            run_unified_scan(),
            media_type="text/event-stream",
        )

    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"Unified binary scan failed: {exc}")
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(exc)}")


@router.post("/binary/unified-scan/{scan_id}/cancel")
async def cancel_unified_binary_scan(
    scan_id: str,
    current_user: User = Depends(get_current_active_user),
):
    """Cancel an in-progress unified binary scan."""
    if scan_id in _unified_binary_scan_sessions:
        _unified_binary_scan_sessions[scan_id]["cancelled"] = True
        return {"message": "Binary scan cancelled"}
    raise HTTPException(status_code=404, detail="Scan not found")


# ============================================================================
# AI Vulnerability Hunter
# ============================================================================

# Store active vulnerability hunts for cancellation
_vulnerability_hunt_sessions: Dict[str, Dict[str, Any]] = {}


class VulnerabilityHuntPhase(BaseModel):
    """A phase in the vulnerability hunt."""
    id: str
    label: str
    description: str
    status: str  # "pending", "in_progress", "completed", "error"
    progress: int = 0
    details: Optional[str] = None
    started_at: Optional[str] = None
    completed_at: Optional[str] = None


class VulnerabilityHuntProgress(BaseModel):
    """Progress update for vulnerability hunt."""
    scan_id: str
    current_phase: str
    overall_progress: int
    phases: List[VulnerabilityHuntPhase]
    message: str
    targets_identified: int = 0
    vulnerabilities_found: int = 0


class VulnerabilityFindingResponse(BaseModel):
    """A vulnerability finding from the hunt."""
    id: str
    title: str
    severity: str
    category: str
    cwe_id: Optional[str] = None
    cvss_estimate: float
    function_name: str
    entry_address: str
    description: str
    technical_details: str
    proof_of_concept: str
    exploitation_steps: List[str]
    remediation: str
    confidence: float
    ai_reasoning: str
    code_snippet: str


class VulnerabilityHuntResultResponse(BaseModel):
    """Complete result of a vulnerability hunt."""
    scan_id: str
    filename: str
    passes_completed: int
    total_functions_analyzed: int
    targets_identified: int
    vulnerabilities: List[VulnerabilityFindingResponse]
    attack_surface_summary: Dict[str, Any]
    hunting_log: List[Dict[str, Any]]
    executive_summary: str
    risk_score: int
    recommended_focus_areas: List[str]


@router.post("/binary/vulnerability-hunt")
async def vulnerability_hunt(
    file: UploadFile = File(..., description="Binary file to analyze (EXE, ELF, DLL, SO)"),
    focus_categories: Optional[str] = Query(
        None,
        description="Comma-separated vulnerability categories to focus on (buffer_overflow,format_string,integer_overflow,use_after_free,command_injection,path_traversal,race_condition,crypto_weakness)"
    ),
    max_passes: int = Query(3, ge=1, le=5, description="Maximum analysis passes"),
    max_targets_per_pass: int = Query(20, ge=5, le=50, description="Max targets to analyze per pass"),
    ghidra_max_functions: int = Query(500, ge=100, le=2000, description="Max functions for Ghidra to export"),
    ghidra_decomp_limit: int = Query(8000, ge=2000, le=20000, description="Max decompilation chars per function"),
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-powered autonomous vulnerability hunting.
    
    Performs multi-pass analysis:
    1. **Pass 1 - Reconnaissance**: Ghidra decompilation + identify dangerous function calls
    2. **Pass 2 - AI Triage**: AI identifies highest-priority targets for deep analysis
    3. **Pass 3+ - Deep Analysis**: AI performs thorough vulnerability analysis on each target
    
    Streams progress updates via Server-Sent Events.
    
    Returns findings with:
    - Vulnerability details and severity
    - CWE classification and CVSS estimate
    - Proof of concept code
    - Exploitation steps
    - Remediation recommendations
    
    Categories: buffer_overflow, format_string, integer_overflow, use_after_free,
    command_injection, path_traversal, race_condition, crypto_weakness
    """
    filename = file.filename or "unknown"
    suffix = Path(filename).suffix.lower()
    
    if suffix not in ALLOWED_BINARY_EXTENSIONS and suffix not in {".bin", ""}:
        pass  # Allow any file for analysis
    
    tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_vuln_hunt_"))
    safe_filename = sanitize_filename(filename)
    tmp_path = tmp_dir / safe_filename
    file_size = 0
    
    try:
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        scan_id = str(uuid.uuid4())
        _vulnerability_hunt_sessions[scan_id] = {"cancelled": False, "tmp_dir": str(tmp_dir)}
        
        # Parse focus categories
        categories = None
        if focus_categories:
            categories = [c.strip() for c in focus_categories.split(",")]
        
        # Build phases
        phases: List[VulnerabilityHuntPhase] = [
            VulnerabilityHuntPhase(
                id="pass1",
                label="Reconnaissance",
                description="Ghidra decompilation and attack surface mapping",
                status="pending",
            ),
            VulnerabilityHuntPhase(
                id="pass2",
                label="AI Triage",
                description="AI identifies high-priority targets",
                status="pending",
            ),
        ]
        for i in range(3, max_passes + 1):
            phases.append(VulnerabilityHuntPhase(
                id=f"pass{i}",
                label=f"Deep Analysis Pass {i-2}",
                description="AI deep vulnerability analysis",
                status="pending",
            ))
        phases.append(VulnerabilityHuntPhase(
            id="summary",
            label="Summary",
            description="Generate executive summary",
            status="pending",
        ))
        
        current_phase_idx = 0
        targets_identified = 0
        vulns_found = 0
        
        def make_progress(message: str, phase_progress: int = 0) -> str:
            nonlocal phases, current_phase_idx, targets_identified, vulns_found
            overall = (current_phase_idx * 100 // max(len(phases), 1)) + (phase_progress // max(len(phases), 1))
            progress = VulnerabilityHuntProgress(
                scan_id=scan_id,
                current_phase=phases[current_phase_idx].id,
                overall_progress=min(overall, 100),
                phases=phases,
                message=message,
                targets_identified=targets_identified,
                vulnerabilities_found=vulns_found,
            )
            return f"data: {json.dumps({'type': 'progress', 'data': progress.model_dump()})}\n\n"
        
        def update_phase(phase_id: str, status: str, details: str = None, progress: int = 0):
            for p in phases:
                if p.id == phase_id:
                    p.status = status
                    p.progress = progress
                    if details:
                        p.details = details
                    if status == "in_progress" and not p.started_at:
                        p.started_at = datetime.utcnow().isoformat()
                    if status in ("completed", "error"):
                        p.completed_at = datetime.utcnow().isoformat()
                        p.progress = 100
        
        async def run_vulnerability_hunt():
            nonlocal current_phase_idx, targets_identified, vulns_found
            
            try:
                if _vulnerability_hunt_sessions.get(scan_id, {}).get("cancelled"):
                    yield f"data: {json.dumps({'type': 'error', 'error': 'Hunt cancelled'})}\n\n"
                    yield "data: {\"type\":\"done\"}\n\n"
                    return
                
                # Progress callback for the hunt
                async def on_progress(phase: str, progress: int, message: str):
                    nonlocal current_phase_idx, targets_identified, vulns_found
                    
                    # Map service phases to router phases
                    phase_map = {
                        "pass1": 0,
                        "pass2": 1,
                        "pass3": 2,
                        "pass4": 3,
                        "pass5": 4,
                        "summary": len(phases) - 1,
                        "complete": len(phases) - 1,
                    }
                    if phase in phase_map:
                        current_phase_idx = min(phase_map[phase], len(phases) - 1)
                    
                    # Update phase status
                    current_id = phases[current_phase_idx].id
                    if progress == 0:
                        update_phase(current_id, "in_progress", message, 0)
                    elif progress >= 100:
                        update_phase(current_id, "completed", message, 100)
                    else:
                        update_phase(current_id, "in_progress", message, progress)
                
                # Yield initial progress
                yield make_progress("Starting AI vulnerability hunt...", 0)
                
                # Run the hunt
                result = await re_service.ai_vulnerability_hunt(
                    tmp_path,
                    focus_categories=categories,
                    max_passes=max_passes,
                    max_targets_per_pass=max_targets_per_pass,
                    ghidra_max_functions=ghidra_max_functions,
                    ghidra_decomp_limit=ghidra_decomp_limit,
                    on_progress=on_progress,
                )
                
                targets_identified = result.targets_identified
                vulns_found = len(result.vulnerabilities)
                
                # Convert to response
                response = re_service.vulnerability_hunt_result_to_dict(result)
                
                # Final progress
                update_phase("summary", "completed", f"Found {vulns_found} vulnerabilities", 100)
                yield make_progress(f"Hunt complete: {vulns_found} vulnerabilities found", 100)
                
                # Yield final result
                yield f"data: {json.dumps({'type': 'result', 'data': response})}\n\n"
                yield "data: {\"type\":\"done\"}\n\n"
                
            except Exception as exc:
                logger.exception(f"Vulnerability hunt failed: {exc}")
                yield f"data: {json.dumps({'type': 'error', 'error': str(exc)})}\n\n"
                yield "data: {\"type\":\"done\"}\n\n"
            finally:
                # Cleanup in background
                try:
                    if scan_id in _vulnerability_hunt_sessions:
                        del _vulnerability_hunt_sessions[scan_id]
                    shutil.rmtree(tmp_dir, ignore_errors=True)
                except Exception as e:
                    logger.debug(f"Cleanup failed (non-critical): {e}")
        
        return StreamingResponse(
            run_vulnerability_hunt(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            }
        )
        
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"Vulnerability hunt failed: {exc}")
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Hunt failed: {str(exc)}")


@router.post("/binary/vulnerability-hunt/{scan_id}/cancel")
async def cancel_vulnerability_hunt(
    scan_id: str,
    current_user: User = Depends(get_current_active_user),
):
    """Cancel an in-progress vulnerability hunt."""
    if scan_id in _vulnerability_hunt_sessions:
        _vulnerability_hunt_sessions[scan_id]["cancelled"] = True
        return {"message": "Vulnerability hunt cancelled"}
    raise HTTPException(status_code=404, detail="Hunt not found")


# ============================================================================
# Binary Purpose Analysis ("What does this Binary do?")
# ============================================================================

class BinaryPurposeProgress(BaseModel):
    """Progress update for binary purpose analysis."""
    phase: str
    progress: int
    message: str


class SuspiciousBehavior(BaseModel):
    """A suspicious behavior identified in the binary."""
    behavior: str
    severity: str  # low, medium, high, critical
    indicator: str


class BinaryPurposeResponse(BaseModel):
    """Response from binary purpose analysis."""
    filename: str
    purpose_summary: str
    detailed_description: str
    category: str
    functionality: List[str]
    capabilities: Dict[str, List[str]]
    api_usage: Dict[str, List[str]]
    suspicious_behaviors: List[SuspiciousBehavior]
    data_handling: Dict[str, Any]
    network_activity: Dict[str, Any]
    file_operations: Dict[str, Any]
    process_operations: Dict[str, Any]
    crypto_usage: Dict[str, Any]
    ui_type: str
    confidence: float
    analysis_notes: List[str]


@router.post("/binary/analyze-purpose")
async def analyze_binary_purpose(
    file: UploadFile = File(..., description="Binary file to analyze"),
    use_ghidra: bool = Query(True, description="Use Ghidra for deeper analysis"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Analyze a binary to understand what it does.
    
    Provides:
    - Purpose summary and detailed description
    - Category classification (utility, malware, game, server, etc.)
    - Functionality list
    - API usage by category (file, network, process, crypto, etc.)
    - Suspicious behavior detection
    - Data handling analysis
    - Network activity patterns
    - UI type detection
    
    Streams progress updates via Server-Sent Events.
    """
    filename = file.filename or "unknown"
    suffix = Path(filename).suffix.lower()
    
    tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_purpose_"))
    safe_filename = sanitize_filename(filename)
    tmp_path = tmp_dir / safe_filename
    
    try:
        # Save file
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(status_code=400, detail=f"File too large (max {MAX_FILE_SIZE // (1024*1024)}MB)")
                f.write(chunk)
        
        logger.info(f"Analyzing binary purpose: {filename} ({file_size:,} bytes)")
        
        # Track progress
        progress_updates = []
        
        async def progress_callback(phase: str, pct: int, msg: str):
            progress_updates.append({"phase": phase, "progress": pct, "message": msg})
        
        # SSE generator
        async def generate_events():
            try:
                # Run analysis in background
                analysis_task = asyncio.create_task(
                    re_service.analyze_binary_purpose(
                        tmp_path,
                        use_ghidra=use_ghidra,
                        progress_callback=progress_callback,
                    )
                )
                
                last_idx = 0
                while not analysis_task.done():
                    # Send any new progress updates
                    while last_idx < len(progress_updates):
                        update = progress_updates[last_idx]
                        yield f"data: {json.dumps({'type': 'progress', 'data': update})}\n\n"
                        last_idx += 1
                    await asyncio.sleep(0.3)
                
                # Get result
                result = await analysis_task
                result_dict = re_service.binary_purpose_to_dict(result)
                
                yield f"data: {json.dumps({'type': 'result', 'data': result_dict})}\n\n"
                yield f"data: {json.dumps({'type': 'done'})}\n\n"
                
            except Exception as e:
                logger.error(f"Binary purpose analysis failed: {e}")
                yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
            finally:
                shutil.rmtree(tmp_dir, ignore_errors=True)
        
        return StreamingResponse(
            generate_events(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            }
        )
        
    except Exception as exc:
        logger.error(f"Binary purpose analysis failed: {exc}")
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(exc)}")


# ============================================================================
# Proof-of-Concept Exploit Generation
# ============================================================================

class PoCExploitRequest(BaseModel):
    """Request to generate a PoC exploit."""
    vulnerability: Dict[str, Any] = Field(..., description="The vulnerability finding to generate exploit for")
    target_platform: str = Field("linux", description="Target platform: linux, windows, both")
    exploit_style: str = Field("python", description="Exploit language: python, c, shellcode")
    include_shellcode: bool = Field(False, description="Include raw shellcode in exploit")


class PoCExploitResponse(BaseModel):
    """A proof-of-concept exploit."""
    vuln_id: str
    vuln_title: str
    exploit_type: str
    language: str
    code: str
    description: str
    prerequisites: List[str]
    usage_instructions: str
    expected_outcome: str
    limitations: List[str]
    safety_notes: List[str]
    tested_on: str
    reliability: str
    evasion_notes: List[str]


class MultiplePoCRequest(BaseModel):
    """Request to generate multiple PoC exploits."""
    vulnerabilities: List[Dict[str, Any]] = Field(..., description="List of vulnerability findings")
    target_platform: str = Field("linux", description="Target platform: linux, windows, both")
    exploit_style: str = Field("python", description="Exploit language: python, c, shellcode")


class MultiplePoCResponse(BaseModel):
    """Response containing multiple PoC exploits."""
    success: bool
    exploits: List[PoCExploitResponse]
    generation_log: List[str]
    warnings: List[str]
    disclaimer: str


@router.post("/binary/generate-poc", response_model=PoCExploitResponse)
async def generate_poc_exploit(
    request: PoCExploitRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate a proof-of-concept exploit for a vulnerability.
    
    Supports various vulnerability types:
    - Buffer overflows
    - Format string bugs
    - Use-after-free
    - Command injection
    - Integer overflows
    - Path traversal
    
    Returns working exploit code with:
    - Full commented code
    - Prerequisites and setup instructions
    - Usage instructions
    - Expected outcome
    - Safety notes and legal warnings
    """
    logger.info(f"Generating PoC for: {request.vulnerability.get('title', 'unknown')}")
    
    try:
        result = await re_service.generate_poc_exploit(
            vulnerability=request.vulnerability,
            target_platform=request.target_platform,
            exploit_style=request.exploit_style,
            include_shellcode=request.include_shellcode,
        )
        
        return PoCExploitResponse(
            vuln_id=result.vuln_id,
            vuln_title=result.vuln_title,
            exploit_type=result.exploit_type,
            language=result.language,
            code=result.code,
            description=result.description,
            prerequisites=result.prerequisites,
            usage_instructions=result.usage_instructions,
            expected_outcome=result.expected_outcome,
            limitations=result.limitations,
            safety_notes=result.safety_notes,
            tested_on=result.tested_on,
            reliability=result.reliability,
            evasion_notes=result.evasion_notes,
        )
        
    except Exception as e:
        logger.error(f"PoC generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PoC: {str(e)}")


@router.post("/binary/generate-pocs", response_model=MultiplePoCResponse)
async def generate_multiple_pocs(
    request: MultiplePoCRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate proof-of-concept exploits for multiple vulnerabilities.
    
    Batch generates exploits for all provided vulnerabilities.
    Returns a summary with all generated exploits and any failures.
    """
    logger.info(f"Generating PoCs for {len(request.vulnerabilities)} vulnerabilities")
    
    try:
        result = await re_service.generate_multiple_pocs(
            vulnerabilities=request.vulnerabilities,
            target_platform=request.target_platform,
            exploit_style=request.exploit_style,
        )
        
        return MultiplePoCResponse(
            success=result.success,
            exploits=[
                PoCExploitResponse(
                    vuln_id=e.vuln_id,
                    vuln_title=e.vuln_title,
                    exploit_type=e.exploit_type,
                    language=e.language,
                    code=e.code,
                    description=e.description,
                    prerequisites=e.prerequisites,
                    usage_instructions=e.usage_instructions,
                    expected_outcome=e.expected_outcome,
                    limitations=e.limitations,
                    safety_notes=e.safety_notes,
                    tested_on=e.tested_on,
                    reliability=e.reliability,
                    evasion_notes=e.evasion_notes,
                )
                for e in result.exploits
            ],
            generation_log=result.generation_log,
            warnings=result.warnings,
            disclaimer=result.disclaimer,
        )
        
    except Exception as e:
        logger.error(f"Multiple PoC generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PoCs: {str(e)}")


# ============================================================================
# AI Chat for Analysis Context
# ============================================================================

class ChatMessage(BaseModel):
    """A single chat message."""
    role: str  # "user" or "assistant"
    content: str

class AnalysisChatRequest(BaseModel):
    """Request for chatting about analysis results."""
    message: str = Field(..., description="User's question")
    conversation_history: List[ChatMessage] = Field(default_factory=list, description="Previous messages")
    analysis_context: Dict[str, Any] = Field(..., description="Analysis context (vulnerabilities, purpose, etc.)")

class AnalysisChatResponse(BaseModel):
    """Response from the chat endpoint."""
    response: str
    error: Optional[str] = None

@router.post("/chat", response_model=AnalysisChatResponse)
async def chat_about_analysis(
    request: AnalysisChatRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Chat with AI about vulnerability analysis results.
    
    Allows users to ask follow-up questions about:
    - Discovered vulnerabilities
    - Binary purpose analysis
    - PoC exploits
    - Remediation strategies
    - Technical details
    """
    from backend.core.config import settings
    
    if not settings.gemini_api_key:
        return AnalysisChatResponse(
            response="",
            error="Chat unavailable: GEMINI_API_KEY not configured"
        )
    
    try:
        from google import genai
        from google.genai import types
        import json
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # Extract context components
        binary_info = request.analysis_context.get("binary_info", {})
        purpose_analysis = request.analysis_context.get("purpose_analysis", {})
        vulnerabilities = request.analysis_context.get("vulnerabilities", [])
        poc_exploits = request.analysis_context.get("poc_exploits", [])
        attack_surface = request.analysis_context.get("attack_surface", {})
        hunt_result = request.analysis_context.get("hunt_result", {})
        
        # Build the system context
        context = f"""You are an expert binary security analyst and reverse engineering assistant. You have analyzed a binary file and have detailed context about its vulnerabilities, purpose, and security characteristics.

## BINARY ANALYSIS CONTEXT

### Binary Information
{json.dumps(binary_info, indent=2) if binary_info else "No binary info available."}

### Binary Purpose Analysis
{json.dumps(purpose_analysis, indent=2) if purpose_analysis else "Not yet analyzed for purpose."}

### Discovered Vulnerabilities ({len(vulnerabilities)} total)
{json.dumps(vulnerabilities[:10], indent=2) if vulnerabilities else "No vulnerabilities discovered yet."}
{f"... and {len(vulnerabilities) - 10} more" if len(vulnerabilities) > 10 else ""}

### Generated PoC Exploits ({len(poc_exploits)} total)
{json.dumps(poc_exploits[:5], indent=2) if poc_exploits else "No PoC exploits generated."}

### Attack Surface Summary
{json.dumps(attack_surface, indent=2) if attack_surface else "Not yet mapped."}

### Hunt Result Summary
- Risk Score: {hunt_result.get('risk_score', 'N/A')}
- Total Functions Analyzed: {hunt_result.get('total_functions_analyzed', 'N/A')}
- Executive Summary: {hunt_result.get('executive_summary', 'N/A')}
- Recommended Focus Areas: {json.dumps(hunt_result.get('recommended_focus_areas', []), indent=2)}

---

## YOUR ROLE
- Answer questions about the vulnerabilities, their severity, and exploitation
- Explain technical details in a clear way
- Provide remediation guidance and security best practices
- Help interpret the analysis results
- Suggest additional testing or analysis if relevant
- If asked about exploitation, always remind users that testing should only be done on authorized systems

Be precise, technical when needed, and always reference specific findings from the analysis when relevant. If a question is outside the scope of the analysis, let the user know what information is available."""

        # Build conversation history for multi-turn chat
        contents = []
        contents.append(types.Content(
            role="user",
            parts=[types.Part(text=context + "\n\nThe user's first question is below.")]
        ))
        
        # Add conversation history
        for msg in request.conversation_history:
            contents.append(types.Content(
                role="user" if msg.role == "user" else "model",
                parts=[types.Part(text=msg.content)]
            ))
        
        # Add current message
        contents.append(types.Content(
            role="user",
            parts=[types.Part(text=request.message)]
        ))
        
        # Generate response
        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=contents
        )
        
        return AnalysisChatResponse(response=response.text)
        
    except Exception as e:
        logger.error(f"Analysis chat error: {e}")
        return AnalysisChatResponse(
            response="",
            error=f"Failed to generate response: {str(e)}"
        )


# ============================================================================
# Notes Management
# ============================================================================

class AnalysisNote(BaseModel):
    """A note associated with analysis."""
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    content: str
    vulnerability_id: Optional[str] = None
    category: str = "general"  # general, vulnerability, poc, remediation
    created_at: str = Field(default_factory=lambda: datetime.utcnow().isoformat())
    updated_at: Optional[str] = None

class NotesStore(BaseModel):
    """Collection of notes for an analysis session."""
    session_id: str
    binary_name: str
    notes: List[AnalysisNote] = Field(default_factory=list)
    created_at: str = Field(default_factory=lambda: datetime.utcnow().isoformat())

class CreateNoteRequest(BaseModel):
    """Request to create a new note."""
    session_id: str
    binary_name: str
    content: str
    vulnerability_id: Optional[str] = None
    category: str = "general"

class UpdateNoteRequest(BaseModel):
    """Request to update an existing note."""
    session_id: str
    note_id: str
    content: str

class DeleteNoteRequest(BaseModel):
    """Request to delete a note."""
    session_id: str
    note_id: str

class ExportNotesRequest(BaseModel):
    """Request to export notes."""
    session_id: str
    format: str = "markdown"  # markdown, json, txt

# In-memory storage for notes (in production, use database)
_notes_storage: Dict[str, NotesStore] = {}
_notes_storage_timestamps: Dict[str, datetime] = {}
NOTES_STORAGE_TTL_HOURS = 24  # Clean up notes after 24 hours of inactivity


async def cleanup_old_notes_storage():
    """Background task to clean up old notes storage entries."""
    while True:
        try:
            await asyncio.sleep(3600)  # Run every hour
            now = datetime.now()
            expired_sessions = []
            
            for session_id, timestamp in list(_notes_storage_timestamps.items()):
                age_hours = (now - timestamp).total_seconds() / 3600
                if age_hours > NOTES_STORAGE_TTL_HOURS:
                    expired_sessions.append(session_id)
            
            for session_id in expired_sessions:
                if session_id in _notes_storage:
                    del _notes_storage[session_id]
                    logger.info(f"Cleaned up notes storage: {session_id}")
                if session_id in _notes_storage_timestamps:
                    del _notes_storage_timestamps[session_id]
            
            if expired_sessions:
                logger.info(f"Cleaned up {len(expired_sessions)} old notes storage entries")
        except Exception as e:
            logger.error(f"Notes storage cleanup error: {e}")


# Session timestamps for scan session cleanup
_unified_scan_timestamps: Dict[str, datetime] = {}
_unified_binary_scan_timestamps: Dict[str, datetime] = {}
SCAN_SESSION_TTL_SECONDS = MAX_SCAN_TIMEOUT + 600  # Scan timeout + 10 minute buffer


async def cleanup_stale_scan_sessions():
    """Background task to clean up stale scan sessions that timed out or had disconnected clients."""
    while True:
        try:
            await asyncio.sleep(300)  # Run every 5 minutes
            now = datetime.now()
            
            # Clean unified APK scan sessions
            expired_apk_sessions = []
            for session_id, timestamp in list(_unified_scan_timestamps.items()):
                age_seconds = (now - timestamp).total_seconds()
                if age_seconds > SCAN_SESSION_TTL_SECONDS:
                    expired_apk_sessions.append(session_id)
            
            for session_id in expired_apk_sessions:
                if session_id in _unified_scan_sessions:
                    session = _unified_scan_sessions[session_id]
                    # Clean up temp directory if exists
                    if "temp_dir" in session and session["temp_dir"]:
                        try:
                            temp_path = Path(session["temp_dir"])
                            if temp_path.exists():
                                shutil.rmtree(temp_path, ignore_errors=True)
                                logger.info(f"Cleaned up temp dir for session {session_id}")
                        except Exception as e:
                            logger.warning(f"Failed to cleanup temp dir for {session_id}: {e}")
                    # Cancel any running task
                    if "task" in session and session["task"] and not session["task"].done():
                        session["task"].cancel()
                    del _unified_scan_sessions[session_id]
                    logger.info(f"Cleaned up stale APK scan session: {session_id}")
                if session_id in _unified_scan_timestamps:
                    del _unified_scan_timestamps[session_id]
            
            # Clean unified binary scan sessions
            expired_binary_sessions = []
            for session_id, timestamp in list(_unified_binary_scan_timestamps.items()):
                age_seconds = (now - timestamp).total_seconds()
                if age_seconds > SCAN_SESSION_TTL_SECONDS:
                    expired_binary_sessions.append(session_id)
            
            for session_id in expired_binary_sessions:
                if session_id in _unified_binary_scan_sessions:
                    session = _unified_binary_scan_sessions[session_id]
                    # Clean up temp directory if exists
                    if "temp_dir" in session and session["temp_dir"]:
                        try:
                            temp_path = Path(session["temp_dir"])
                            if temp_path.exists():
                                shutil.rmtree(temp_path, ignore_errors=True)
                                logger.info(f"Cleaned up temp dir for binary session {session_id}")
                        except Exception as e:
                            logger.warning(f"Failed to cleanup temp dir for binary {session_id}: {e}")
                    # Cancel any running task
                    if "task" in session and session["task"] and not session["task"].done():
                        session["task"].cancel()
                    del _unified_binary_scan_sessions[session_id]
                    logger.info(f"Cleaned up stale binary scan session: {session_id}")
                if session_id in _unified_binary_scan_timestamps:
                    del _unified_binary_scan_timestamps[session_id]
            
            total_cleaned = len(expired_apk_sessions) + len(expired_binary_sessions)
            if total_cleaned:
                logger.info(f"Cleaned up {total_cleaned} stale scan sessions")
        except Exception as e:
            logger.error(f"Scan session cleanup error: {e}")


async def cleanup_orphaned_temp_dirs():
    """Startup task to clean up orphaned temp directories from crashed/restarted sessions."""
    try:
        import glob
        temp_dir = tempfile.gettempdir()
        patterns = [
            "vragent_binary_*",
            "vragent_vuln_hunt_*", 
            "vragent_binary_unified_*",
            "vragent_jadx_*",
            "jadx_*",
            "vragent_apk_*",
        ]
        
        total_cleaned = 0
        for pattern in patterns:
            full_pattern = os.path.join(temp_dir, pattern)
            for dir_path in glob.glob(full_pattern):
                try:
                    if os.path.isdir(dir_path):
                        # Only clean directories older than 2 hours to avoid race conditions
                        dir_age = time.time() - os.path.getmtime(dir_path)
                        if dir_age > 7200:  # 2 hours
                            shutil.rmtree(dir_path, ignore_errors=True)
                            total_cleaned += 1
                            logger.debug(f"Cleaned orphaned temp dir: {dir_path}")
                except Exception as e:
                    logger.warning(f"Failed to clean temp dir {dir_path}: {e}")
        
        if total_cleaned:
            logger.info(f"Cleaned up {total_cleaned} orphaned temp directories on startup")
    except Exception as e:
        logger.error(f"Orphaned temp dir cleanup error: {e}")


@router.post("/notes/create", response_model=AnalysisNote)
async def create_note(
    request: CreateNoteRequest,
    current_user: User = Depends(get_current_active_user),
):
    """Create a new note for the analysis session."""
    session_id = request.session_id
    
    # Initialize session if needed
    if session_id not in _notes_storage:
        _notes_storage[session_id] = NotesStore(
            session_id=session_id,
            binary_name=request.binary_name
        )
    
    # Update timestamp on any note activity
    _notes_storage_timestamps[session_id] = datetime.now()
    
    # Create new note
    note = AnalysisNote(
        content=request.content,
        vulnerability_id=request.vulnerability_id,
        category=request.category
    )
    
    _notes_storage[session_id].notes.append(note)
    
    logger.info(f"Created note {note.id} for session {session_id}")
    return note


@router.post("/notes/list", response_model=NotesStore)
async def list_notes(
    session_id: str,
    current_user: User = Depends(get_current_active_user),
):
    """List all notes for an analysis session."""
    if session_id not in _notes_storage:
        return NotesStore(session_id=session_id, binary_name="Unknown")
    
    # Update timestamp on access
    _notes_storage_timestamps[session_id] = datetime.now()
    
    return _notes_storage[session_id]


@router.post("/notes/update", response_model=AnalysisNote)
async def update_note(
    request: UpdateNoteRequest,
    current_user: User = Depends(get_current_active_user),
):
    """Update an existing note."""
    if request.session_id not in _notes_storage:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Update timestamp on any note activity
    _notes_storage_timestamps[request.session_id] = datetime.now()
    
    store = _notes_storage[request.session_id]
    for note in store.notes:
        if note.id == request.note_id:
            note.content = request.content
            note.updated_at = datetime.utcnow().isoformat()
            logger.info(f"Updated note {note.id}")
            return note
    
    raise HTTPException(status_code=404, detail="Note not found")


@router.delete("/notes/delete")
async def delete_note(
    request: DeleteNoteRequest,
    current_user: User = Depends(get_current_active_user),
):
    """Delete a note."""
    if request.session_id not in _notes_storage:
        raise HTTPException(status_code=404, detail="Session not found")
    
    store = _notes_storage[request.session_id]
    original_length = len(store.notes)
    store.notes = [n for n in store.notes if n.id != request.note_id]
    
    if len(store.notes) == original_length:
        raise HTTPException(status_code=404, detail="Note not found")
    
    logger.info(f"Deleted note {request.note_id}")
    return {"status": "deleted", "note_id": request.note_id}


@router.post("/notes/export")
async def export_notes(
    request: ExportNotesRequest,
    current_user: User = Depends(get_current_active_user),
):
    """Export notes in the specified format."""
    if request.session_id not in _notes_storage:
        raise HTTPException(status_code=404, detail="Session not found")
    
    store = _notes_storage[request.session_id]
    
    if request.format == "json":
        return {
            "format": "json",
            "content": store.model_dump(),
            "filename": f"{store.binary_name}_notes.json"
        }
    
    elif request.format == "markdown":
        md_content = f"# Analysis Notes: {store.binary_name}\n\n"
        md_content += f"*Session: {store.session_id}*\n"
        md_content += f"*Created: {store.created_at}*\n\n"
        
        # Group by category
        categories = {}
        for note in store.notes:
            cat = note.category
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(note)
        
        for category, notes in categories.items():
            md_content += f"## {category.title()} Notes\n\n"
            for note in notes:
                md_content += f"### Note ({note.created_at})\n"
                if note.vulnerability_id:
                    md_content += f"*Linked to vulnerability: {note.vulnerability_id}*\n\n"
                md_content += f"{note.content}\n\n"
                md_content += "---\n\n"
        
        return {
            "format": "markdown",
            "content": md_content,
            "filename": f"{store.binary_name}_notes.md"
        }
    
    else:  # txt
        txt_content = f"Analysis Notes: {store.binary_name}\n"
        txt_content += f"Session: {store.session_id}\n"
        txt_content += f"Created: {store.created_at}\n"
        txt_content += "=" * 50 + "\n\n"
        
        for note in store.notes:
            txt_content += f"[{note.category.upper()}] {note.created_at}\n"
            if note.vulnerability_id:
                txt_content += f"Vulnerability: {note.vulnerability_id}\n"
            txt_content += "-" * 30 + "\n"
            txt_content += f"{note.content}\n\n"
        
        return {
            "format": "txt",
            "content": txt_content,
            "filename": f"{store.binary_name}_notes.txt"
        }


# ============================================================================
# AI Decompiler Enhancement
# ============================================================================

class EnhanceCodeRequest(BaseModel):
    """Request to enhance decompiled code with AI."""
    code: str = Field(..., description="Decompiled code to enhance")
    function_name: str = Field(default="unknown", description="Original function name")
    binary_context: Optional[Dict[str, Any]] = Field(default=None, description="Additional binary context")
    enhancement_level: str = Field(default="full", description="Enhancement level: basic, standard, full")
    include_security_analysis: bool = Field(default=True, description="Include security vulnerability annotations")


class EnhancedVariable(BaseModel):
    """An enhanced/renamed variable."""
    original_name: str
    suggested_name: str
    inferred_type: str
    confidence: float
    reasoning: str


class EnhancedCodeBlock(BaseModel):
    """A code block with explanation."""
    start_line: int
    end_line: int
    purpose: str
    security_notes: Optional[str] = None


class DataStructure(BaseModel):
    """Reconstructed data structure."""
    name: str
    inferred_type: str
    fields: List[Dict[str, str]]
    usage_context: str


class SecurityAnnotation(BaseModel):
    """Security-related annotation in code."""
    line: int
    severity: str  # info, low, medium, high, critical
    issue_type: str
    description: str
    cwe_id: Optional[str] = None


class EnhanceCodeResponse(BaseModel):
    """Response with enhanced decompiled code."""
    original_code: str
    enhanced_code: str
    suggested_function_name: str
    function_purpose: str
    variables: List[EnhancedVariable]
    code_blocks: List[EnhancedCodeBlock]
    data_structures: List[DataStructure]
    security_annotations: List[SecurityAnnotation]
    complexity_score: int  # 1-10
    readability_improvement: int  # percentage
    inline_comments_added: int
    api_calls_identified: List[Dict[str, str]]


@router.post("/binary/enhance-code", response_model=EnhanceCodeResponse)
async def enhance_decompiled_code(
    request: EnhanceCodeRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-powered decompiled code enhancement.
    
    Takes Ghidra/IDA decompiled code and makes it human-readable:
    - Renames variables intelligently (var_14 → encryptionKey)
    - Adds inline comments explaining code blocks
    - Identifies and documents data structures
    - Suggests original function names based on behavior
    - Annotates security vulnerabilities inline
    - Identifies API calls and their purposes
    
    Enhancement levels:
    - basic: Variable renaming only
    - standard: Variables + inline comments
    - full: Everything including security analysis
    """
    from backend.core.config import settings
    from google import genai
    from google.genai import types

    if not settings.GEMINI_API_KEY:
        raise HTTPException(status_code=503, detail="Gemini API key not configured")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    model_id = "gemini-3-flash-preview"
    
    # Build context
    context_info = ""
    if request.binary_context:
        if "imports" in request.binary_context:
            context_info += f"\nImported functions: {', '.join(request.binary_context['imports'][:20])}"
        if "strings" in request.binary_context:
            context_info += f"\nInteresting strings: {', '.join(request.binary_context['strings'][:10])}"
        if "architecture" in request.binary_context:
            context_info += f"\nArchitecture: {request.binary_context['architecture']}"
    
    # Determine what to analyze based on enhancement level
    analyze_security = request.include_security_analysis and request.enhancement_level == "full"
    analyze_structures = request.enhancement_level in ("standard", "full")
    
    prompt = f"""You are an expert reverse engineer and code analyst. Analyze this decompiled code and enhance it for readability.

**Original Function Name:** {request.function_name}
**Enhancement Level:** {request.enhancement_level}
{context_info}

**Decompiled Code:**
```c
{request.code}
```

Provide a comprehensive analysis in the following JSON format:

{{
    "suggested_function_name": "descriptive_name_based_on_behavior",
    "function_purpose": "Clear description of what this function does",
    "enhanced_code": "The complete enhanced code with renamed variables and inline comments",
    "variables": [
        {{
            "original_name": "var_14",
            "suggested_name": "buffer_size",
            "inferred_type": "size_t",
            "confidence": 0.85,
            "reasoning": "Used as size parameter in memcpy"
        }}
    ],
    "code_blocks": [
        {{
            "start_line": 1,
            "end_line": 5,
            "purpose": "Initialize encryption context",
            "security_notes": "Uses hardcoded IV - potential weakness"
        }}
    ],
    "data_structures": [
        {{
            "name": "connection_info",
            "inferred_type": "struct",
            "fields": [{{"name": "socket_fd", "type": "int"}}, {{"name": "buffer", "type": "char*"}}],
            "usage_context": "Network connection state"
        }}
    ],
    "security_annotations": [
        {{
            "line": 12,
            "severity": "high",
            "issue_type": "buffer_overflow",
            "description": "strcpy without bounds checking",
            "cwe_id": "CWE-120"
        }}
    ],
    "api_calls_identified": [
        {{
            "name": "memcpy",
            "purpose": "Copy decrypted data to output buffer",
            "line": 15
        }}
    ],
    "complexity_score": 6,
    "readability_improvement": 75
}}

Guidelines for enhanced_code:
1. Replace generic variable names (var_X, param_X) with meaningful names
2. Add // comments explaining non-obvious operations
3. Group related operations with block comments
4. Keep the code syntactically valid C
5. Preserve the original logic exactly
{"6. Add security warning comments for dangerous patterns" if analyze_security else ""}

Return ONLY valid JSON, no markdown formatting."""

    try:
        response = await asyncio.to_thread(
            lambda: client.models.generate_content(
                model=model_id,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(thinking_level="medium"),
                    max_output_tokens=8000,
                )
            )
        )
        
        result_text = response.text.strip()
        
        # Clean up response
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        result_text = result_text.strip()
        
        result = json.loads(result_text)
        
        return EnhanceCodeResponse(
            original_code=request.code,
            enhanced_code=result.get("enhanced_code", request.code),
            suggested_function_name=result.get("suggested_function_name", request.function_name),
            function_purpose=result.get("function_purpose", "Unknown"),
            variables=[EnhancedVariable(**v) for v in result.get("variables", [])],
            code_blocks=[EnhancedCodeBlock(**b) for b in result.get("code_blocks", [])],
            data_structures=[DataStructure(**d) for d in result.get("data_structures", [])],
            security_annotations=[SecurityAnnotation(**s) for s in result.get("security_annotations", [])] if analyze_security else [],
            complexity_score=result.get("complexity_score", 5),
            readability_improvement=result.get("readability_improvement", 50),
            inline_comments_added=result.get("enhanced_code", "").count("//") - request.code.count("//"),
            api_calls_identified=result.get("api_calls_identified", []),
        )
        
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse AI response: {e}")
        # Return basic response
        return EnhanceCodeResponse(
            original_code=request.code,
            enhanced_code=request.code,
            suggested_function_name=request.function_name,
            function_purpose="Analysis failed - could not parse AI response",
            variables=[],
            code_blocks=[],
            data_structures=[],
            security_annotations=[],
            complexity_score=5,
            readability_improvement=0,
            inline_comments_added=0,
            api_calls_identified=[],
        )
    except Exception as e:
        logger.error(f"Code enhancement failed: {e}")
        raise HTTPException(status_code=500, detail=f"Enhancement failed: {str(e)}")


# ============================================================================
# Natural Language Binary Search
# ============================================================================

class NaturalLanguageSearchRequest(BaseModel):
    """Request for natural language search across binary."""
    query: str = Field(..., description="Natural language search query")
    functions: List[Dict[str, Any]] = Field(..., description="List of functions with their decompiled code")
    max_results: int = Field(default=10, ge=1, le=50, description="Maximum results to return")
    include_explanations: bool = Field(default=True, description="Include AI explanations for matches")
    search_scope: str = Field(default="all", description="Search scope: all, security, network, crypto, file_io")


class SearchMatch(BaseModel):
    """A matching function from natural language search."""
    function_name: str
    address: Optional[str] = None
    relevance_score: float  # 0-1
    match_reason: str
    code_snippet: str
    key_indicators: List[str]
    security_relevant: bool
    category: str  # network, crypto, file_io, memory, string, etc.


class NaturalLanguageSearchResponse(BaseModel):
    """Response from natural language binary search."""
    query: str
    interpreted_query: str  # How AI understood the query
    total_functions_searched: int
    matches: List[SearchMatch]
    search_suggestions: List[str]  # Related queries user might want
    categories_found: Dict[str, int]  # Category distribution


@router.post("/binary/nl-search", response_model=NaturalLanguageSearchResponse)
async def natural_language_search(
    request: NaturalLanguageSearchRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Natural language search across decompiled binary functions.
    
    Search by description instead of exact strings:
    - "Find the function that handles network authentication"
    - "Show me code that writes to files"
    - "Where is the encryption key derived?"
    - "Functions that parse user input"
    - "Code vulnerable to buffer overflow"
    
    The AI semantically searches across all decompiled functions
    and returns the most relevant matches with explanations.
    """
    from backend.core.config import settings
    from google import genai
    from google.genai import types

    if not settings.GEMINI_API_KEY:
        raise HTTPException(status_code=503, detail="Gemini API key not configured")

    if not request.functions:
        raise HTTPException(status_code=400, detail="No functions provided to search")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    model_id = "gemini-3-flash-preview"
    
    # Build function summaries for the AI to search
    # Limit to prevent token overflow
    max_funcs = min(len(request.functions), 100)
    function_summaries = []
    
    for i, func in enumerate(request.functions[:max_funcs]):
        code = func.get("code", func.get("decompiled", ""))[:2000]  # Limit code length
        name = func.get("name", func.get("function_name", f"func_{i}"))
        addr = func.get("address", func.get("entry_point", "unknown"))
        
        function_summaries.append({
            "index": i,
            "name": name,
            "address": str(addr),
            "code_preview": code
        })
    
    # Build the search prompt
    scope_hint = ""
    if request.search_scope != "all":
        scope_hints = {
            "security": "Focus on security-sensitive functions (input validation, auth, crypto, memory operations)",
            "network": "Focus on network-related functions (sockets, HTTP, DNS, protocols)",
            "crypto": "Focus on cryptographic functions (encryption, hashing, key management)",
            "file_io": "Focus on file operations (read, write, open, delete, permissions)",
        }
        scope_hint = scope_hints.get(request.search_scope, "")
    
    prompt = f"""You are an expert reverse engineer searching through decompiled binary code.

**User's Search Query:** "{request.query}"
{f"**Search Scope:** {scope_hint}" if scope_hint else ""}

**Functions to Search ({len(function_summaries)} total):**

{json.dumps(function_summaries, indent=2)}

Analyze each function and find the ones that best match the user's query. Consider:
1. Function behavior and purpose
2. API calls and system interactions
3. Data flow and transformations
4. Security implications
5. Semantic meaning, not just keyword matching

Return your analysis as JSON:

{{
    "interpreted_query": "How you understood the user's intent",
    "matches": [
        {{
            "function_index": 0,
            "function_name": "actual_name",
            "address": "0x401000",
            "relevance_score": 0.95,
            "match_reason": "This function implements network authentication by...",
            "code_snippet": "Key lines that match the query",
            "key_indicators": ["calls socket()", "compares credentials", "returns auth token"],
            "security_relevant": true,
            "category": "network"
        }}
    ],
    "search_suggestions": [
        "Related query 1 the user might want",
        "Related query 2"
    ],
    "categories_found": {{
        "network": 2,
        "crypto": 1,
        "file_io": 0
    }}
}}

Return up to {request.max_results} most relevant matches, sorted by relevance_score (highest first).
Only include functions with relevance_score >= 0.3.
Return ONLY valid JSON, no markdown."""

    try:
        response = await asyncio.to_thread(
            lambda: client.models.generate_content(
                model=model_id,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(thinking_level="medium"),
                    max_output_tokens=4000,
                )
            )
        )
        
        result_text = response.text.strip()
        
        # Clean up response
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        result_text = result_text.strip()
        
        result = json.loads(result_text)
        
        # Build response matches
        matches = []
        for match in result.get("matches", []):
            idx = match.get("function_index", -1)
            if 0 <= idx < len(function_summaries):
                orig_func = function_summaries[idx]
                matches.append(SearchMatch(
                    function_name=match.get("function_name", orig_func["name"]),
                    address=match.get("address", orig_func["address"]),
                    relevance_score=min(1.0, max(0.0, match.get("relevance_score", 0.5))),
                    match_reason=match.get("match_reason", "Matched search criteria"),
                    code_snippet=match.get("code_snippet", orig_func["code_preview"][:500]),
                    key_indicators=match.get("key_indicators", []),
                    security_relevant=match.get("security_relevant", False),
                    category=match.get("category", "unknown"),
                ))
        
        return NaturalLanguageSearchResponse(
            query=request.query,
            interpreted_query=result.get("interpreted_query", request.query),
            total_functions_searched=len(function_summaries),
            matches=matches[:request.max_results],
            search_suggestions=result.get("search_suggestions", [])[:5],
            categories_found=result.get("categories_found", {}),
        )
        
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse search response: {e}")
        return NaturalLanguageSearchResponse(
            query=request.query,
            interpreted_query=request.query,
            total_functions_searched=len(function_summaries),
            matches=[],
            search_suggestions=["Try a more specific query", "Use keywords like 'network', 'file', 'crypto'"],
            categories_found={},
        )
    except Exception as e:
        logger.error(f"Natural language search failed: {e}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")


@router.post("/binary/smart-rename")
async def smart_rename_function(
    function_code: str = Query(..., description="Decompiled function code"),
    current_name: str = Query(default="FUN_00401000", description="Current function name"),
    context_hints: Optional[str] = Query(default=None, description="Additional context about the binary"),
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-powered function renaming suggestion.
    
    Analyzes decompiled code and suggests a meaningful function name
    based on the function's behavior, API calls, and data flow.
    """
    from backend.core.config import settings
    from google import genai
    from google.genai import types

    if not settings.GEMINI_API_KEY:
        raise HTTPException(status_code=503, detail="Gemini API key not configured")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    model_id = "gemini-3-flash-preview"
    
    prompt = f"""Analyze this decompiled function and suggest a meaningful name.

**Current Name:** {current_name}
{f"**Context:** {context_hints}" if context_hints else ""}

**Decompiled Code:**
```c
{function_code[:4000]}
```

Respond with JSON:
{{
    "suggested_name": "descriptive_function_name",
    "confidence": 0.85,
    "reasoning": "Why this name fits",
    "alternative_names": ["other_option_1", "other_option_2"],
    "detected_purpose": "Brief description of what the function does",
    "naming_convention": "snake_case"
}}

Use snake_case. Be specific but concise. Examples:
- "validate_user_credentials" not "check_stuff"
- "decrypt_aes_buffer" not "crypto_function"
- "parse_http_headers" not "process_data"

Return ONLY valid JSON."""

    try:
        response = await asyncio.to_thread(
            lambda: client.models.generate_content(
                model=model_id,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(thinking_level="medium"),
                    max_output_tokens=500,
                )
            )
        )
        
        result_text = response.text.strip()
        if result_text.startswith("```"):
            result_text = result_text.split("```")[1]
            if result_text.startswith("json"):
                result_text = result_text[4:]
        result_text = result_text.strip()
        
        result = json.loads(result_text)
        return result
        
    except Exception as e:
        logger.error(f"Smart rename failed: {e}")
        return {
            "suggested_name": current_name,
            "confidence": 0.0,
            "reasoning": f"Analysis failed: {str(e)}",
            "alternative_names": [],
            "detected_purpose": "Unknown",
            "naming_convention": "unknown"
        }


# ============================================================================
# Live Attack Simulation Mode
# ============================================================================

class AttackSimulationRequest(BaseModel):
    """Request for attack simulation."""
    vulnerability: Dict[str, Any] = Field(..., description="The vulnerability finding to simulate")
    target_code: Optional[str] = Field(None, description="Relevant decompiled code")
    binary_context: Optional[Dict[str, Any]] = Field(None, description="Binary context (architecture, imports)")
    simulation_depth: str = Field("standard", description="Depth: quick, standard, or comprehensive")


class RegisterState(BaseModel):
    """CPU register state at a simulation step."""
    name: str
    value: str
    description: Optional[str] = None
    is_controlled: bool = False  # Whether attacker controls this value


class MemoryRegion(BaseModel):
    """Memory region state at a simulation step."""
    address: str
    size: int
    label: str  # e.g., "stack buffer", "heap chunk", "return address"
    state: str  # "normal", "corrupted", "overwritten", "controlled"
    content_preview: Optional[str] = None


class AttackStep(BaseModel):
    """A single step in the attack simulation."""
    step_number: int
    phase: str  # "setup", "trigger", "corruption", "control", "payload", "execution"
    title: str
    description: str
    technical_detail: str
    code_location: Optional[str] = None
    registers: List[RegisterState] = []
    memory_regions: List[MemoryRegion] = []
    attacker_input: Optional[str] = None
    visual_indicator: str  # "info", "warning", "danger", "success"


class ExploitPrimitive(BaseModel):
    """An exploit primitive that can be achieved."""
    name: str  # e.g., "arbitrary write", "code execution", "info leak"
    description: str
    prerequisites: List[str]
    achieved_at_step: int


class Mitigation(BaseModel):
    """A mitigation that could prevent the attack."""
    name: str
    description: str
    effectiveness: str  # "blocks", "hinders", "ineffective"
    bypass_possible: bool
    bypass_technique: Optional[str] = None


class AttackSimulationResponse(BaseModel):
    """Response containing the full attack simulation."""
    vulnerability_title: str
    vulnerability_type: str
    severity: str
    
    # Attack overview
    attack_summary: str
    success_probability: float  # 0-1
    required_conditions: List[str]
    
    # Step-by-step simulation
    total_steps: int
    steps: List[AttackStep]
    
    # What the attacker achieves
    exploit_primitives: List[ExploitPrimitive]
    final_impact: str
    
    # Defense analysis
    mitigations: List[Mitigation]
    detection_opportunities: List[str]
    
    # Additional context
    real_world_examples: List[str]
    cve_references: List[str]
    difficulty_rating: str  # "trivial", "easy", "moderate", "hard", "expert"


@router.post("/binary/simulate-attack", response_model=AttackSimulationResponse)
async def simulate_attack(
    request: AttackSimulationRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Live Attack Simulation Mode.
    
    Takes a discovered vulnerability and generates a step-by-step simulation
    of how an attacker would exploit it. Shows:
    - Each phase of the attack (setup, trigger, corruption, control, execution)
    - Register and memory state at each step
    - What attacker input causes each state change
    - Exploit primitives achieved
    - Mitigations and their effectiveness
    - Real-world examples of similar attacks
    
    This helps security teams understand the practical impact of vulnerabilities
    and prioritize fixes based on exploitability.
    """
    from backend.core.config import settings
    from google import genai
    from google.genai import types

    if not settings.GEMINI_API_KEY:
        raise HTTPException(status_code=503, detail="Gemini API key not configured")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    model_id = "gemini-3-flash-preview"
    
    # Extract vulnerability details
    vuln = request.vulnerability
    vuln_type = vuln.get("vulnerability_type", vuln.get("type", vuln.get("title", "Unknown")))
    vuln_title = vuln.get("title", vuln.get("function_name", f"{vuln_type} Vulnerability"))
    severity = vuln.get("severity", "medium")
    technical_details = vuln.get("technical_details", vuln.get("description", ""))
    poc = vuln.get("proof_of_concept", "")
    
    # Build context
    context_parts = []
    if request.target_code:
        context_parts.append(f"**Vulnerable Code:**\n```c\n{request.target_code[:3000]}\n```")
    if request.binary_context:
        if request.binary_context.get("architecture"):
            context_parts.append(f"**Architecture:** {request.binary_context['architecture']}")
        if request.binary_context.get("imports"):
            context_parts.append(f"**Key Imports:** {', '.join(request.binary_context['imports'][:15])}")
    
    context = "\n".join(context_parts) if context_parts else "No additional context provided."
    
    # Determine simulation depth
    step_count = {"quick": 4, "standard": 6, "comprehensive": 10}.get(request.simulation_depth, 6)
    
    prompt = f"""You are an expert exploit developer and security researcher. Generate a detailed, step-by-step attack simulation for this vulnerability.

**Vulnerability Type:** {vuln_type}
**Title:** {vuln_title}
**Severity:** {severity}

**Technical Details:**
{technical_details}

{f"**Proof of Concept:**{chr(10)}{poc}" if poc else ""}

{context}

Generate a realistic attack simulation showing exactly how an attacker would exploit this vulnerability. Be technical and specific.

Respond with JSON (and ONLY JSON, no markdown):

{{
    "vulnerability_title": "{vuln_title}",
    "vulnerability_type": "{vuln_type}",
    "severity": "{severity}",
    
    "attack_summary": "A 2-3 sentence summary of the attack",
    "success_probability": 0.75,
    "required_conditions": [
        "Condition 1 needed for exploit",
        "Condition 2"
    ],
    
    "total_steps": {step_count},
    "steps": [
        {{
            "step_number": 1,
            "phase": "setup",
            "title": "Prepare Malicious Input",
            "description": "The attacker crafts a specially formatted input...",
            "technical_detail": "Create a buffer of 0x108 bytes followed by the target return address",
            "code_location": "Line 42: memcpy(local_buf, user_input, input_len)",
            "registers": [
                {{"name": "RSP", "value": "0x7fffffffe000", "description": "Stack pointer", "is_controlled": false}},
                {{"name": "RDI", "value": "0x7fffffffe010", "description": "Destination buffer", "is_controlled": false}}
            ],
            "memory_regions": [
                {{
                    "address": "0x7fffffffe010",
                    "size": 256,
                    "label": "local_buf (stack buffer)",
                    "state": "normal",
                    "content_preview": "00 00 00 00 00 00 00 00..."
                }},
                {{
                    "address": "0x7fffffffe110",
                    "size": 8,
                    "label": "saved return address",
                    "state": "normal",
                    "content_preview": "Address of caller"
                }}
            ],
            "attacker_input": "A * 264 + pack('<Q', 0x401234)",
            "visual_indicator": "info"
        }},
        {{
            "step_number": 2,
            "phase": "trigger",
            "title": "Trigger the Vulnerability",
            "description": "The vulnerable function is called with oversized input...",
            "technical_detail": "memcpy copies 280 bytes into 256-byte buffer",
            "code_location": "Line 42: memcpy(local_buf, user_input, input_len)",
            "registers": [
                {{"name": "RDX", "value": "0x118", "description": "Copy size (280 bytes)", "is_controlled": true}}
            ],
            "memory_regions": [
                {{
                    "address": "0x7fffffffe010",
                    "size": 256,
                    "label": "local_buf",
                    "state": "corrupted",
                    "content_preview": "41 41 41 41 41 41 41 41..."
                }}
            ],
            "attacker_input": null,
            "visual_indicator": "warning"
        }},
        {{
            "step_number": 3,
            "phase": "corruption",
            "title": "Stack Corruption Occurs",
            "description": "The overflow overwrites the saved return address...",
            "technical_detail": "Bytes 256-263 overwrite saved RBP, bytes 264-271 overwrite return address",
            "code_location": null,
            "registers": [],
            "memory_regions": [
                {{
                    "address": "0x7fffffffe110",
                    "size": 8,
                    "label": "saved return address",
                    "state": "overwritten",
                    "content_preview": "34 12 40 00 00 00 00 00"
                }}
            ],
            "attacker_input": null,
            "visual_indicator": "danger"
        }},
        {{
            "step_number": 4,
            "phase": "control",
            "title": "Attacker Gains Control",
            "description": "When the function returns, execution jumps to attacker-controlled address...",
            "technical_detail": "RET instruction pops 0x401234 into RIP",
            "code_location": "Function epilogue: ret",
            "registers": [
                {{"name": "RIP", "value": "0x401234", "description": "Instruction pointer", "is_controlled": true}}
            ],
            "memory_regions": [],
            "attacker_input": null,
            "visual_indicator": "danger"
        }},
        {{
            "step_number": 5,
            "phase": "payload",
            "title": "ROP Chain Execution",
            "description": "The attacker's ROP chain begins executing...",
            "technical_detail": "First gadget: pop rdi; ret - sets up argument for system()",
            "code_location": "0x401234 (gadget)",
            "registers": [
                {{"name": "RDI", "value": "0x402000", "description": "Pointer to '/bin/sh'", "is_controlled": true}}
            ],
            "memory_regions": [],
            "attacker_input": null,
            "visual_indicator": "danger"
        }},
        {{
            "step_number": 6,
            "phase": "execution",
            "title": "Code Execution Achieved",
            "description": "The attacker achieves arbitrary code execution...",
            "technical_detail": "system('/bin/sh') spawns a shell with process privileges",
            "code_location": "system() in libc",
            "registers": [
                {{"name": "RAX", "value": "0x0", "description": "system() return value", "is_controlled": false}}
            ],
            "memory_regions": [],
            "attacker_input": null,
            "visual_indicator": "success"
        }}
    ],
    
    "exploit_primitives": [
        {{
            "name": "Stack Buffer Overflow",
            "description": "Overwrite stack data beyond buffer bounds",
            "prerequisites": ["No stack canaries or canary bypass"],
            "achieved_at_step": 2
        }},
        {{
            "name": "Return Address Control",
            "description": "Redirect execution to arbitrary address",
            "prerequisites": ["Stack overflow reaching return address"],
            "achieved_at_step": 3
        }},
        {{
            "name": "Arbitrary Code Execution",
            "description": "Execute attacker-controlled code/ROP chain",
            "prerequisites": ["Return address control", "Known gadget addresses"],
            "achieved_at_step": 6
        }}
    ],
    
    "final_impact": "The attacker achieves remote code execution with the privileges of the vulnerable process, potentially leading to full system compromise.",
    
    "mitigations": [
        {{
            "name": "Stack Canaries",
            "description": "Random value placed before return address, checked before function returns",
            "effectiveness": "blocks",
            "bypass_possible": true,
            "bypass_technique": "Information leak to read canary value, or format string to overwrite"
        }},
        {{
            "name": "ASLR",
            "description": "Randomize memory layout including stack and library addresses",
            "effectiveness": "hinders",
            "bypass_possible": true,
            "bypass_technique": "Information leak or brute force (32-bit)"
        }},
        {{
            "name": "DEP/NX",
            "description": "Mark stack as non-executable",
            "effectiveness": "hinders",
            "bypass_possible": true,
            "bypass_technique": "Return-Oriented Programming (ROP)"
        }},
        {{
            "name": "Safe Functions",
            "description": "Use strncpy/memcpy_s with size limits",
            "effectiveness": "blocks",
            "bypass_possible": false,
            "bypass_technique": null
        }}
    ],
    
    "detection_opportunities": [
        "Monitor for crashes/segfaults indicating exploitation attempts",
        "Detect anomalous memory access patterns",
        "Intrusion detection signatures for known exploit payloads",
        "Runtime canary violation alerts"
    ],
    
    "real_world_examples": [
        "CVE-2021-3156 (sudo heap overflow)",
        "CVE-2017-5638 (Apache Struts RCE)"
    ],
    
    "cve_references": [],
    
    "difficulty_rating": "moderate"
}}

IMPORTANT:
- Generate realistic register values (use x86-64 conventions)
- Show actual memory corruption with hex previews
- Make the attack technically accurate for the vulnerability type
- Adjust the simulation for the specific vulnerability (buffer overflow, format string, use-after-free, etc.)
- Include {step_count} detailed steps covering setup through execution
- Return ONLY valid JSON"""

    try:
        response = await asyncio.to_thread(
            lambda: client.models.generate_content(
                model=model_id,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(thinking_level="medium"),
                    max_output_tokens=8000,
                )
            )
        )
        
        result_text = response.text.strip()
        
        # Clean up response
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        result_text = result_text.strip()
        
        result = json.loads(result_text)
        
        # Build response with proper validation
        steps = []
        for step_data in result.get("steps", []):
            registers = [RegisterState(**r) for r in step_data.get("registers", [])]
            memory_regions = [MemoryRegion(**m) for m in step_data.get("memory_regions", [])]
            steps.append(AttackStep(
                step_number=step_data.get("step_number", 0),
                phase=step_data.get("phase", "unknown"),
                title=step_data.get("title", "Unknown Step"),
                description=step_data.get("description", ""),
                technical_detail=step_data.get("technical_detail", ""),
                code_location=step_data.get("code_location"),
                registers=registers,
                memory_regions=memory_regions,
                attacker_input=step_data.get("attacker_input"),
                visual_indicator=step_data.get("visual_indicator", "info"),
            ))
        
        exploit_primitives = [
            ExploitPrimitive(**p) for p in result.get("exploit_primitives", [])
        ]
        
        mitigations = [
            Mitigation(**m) for m in result.get("mitigations", [])
        ]
        
        return AttackSimulationResponse(
            vulnerability_title=result.get("vulnerability_title", vuln_title),
            vulnerability_type=result.get("vulnerability_type", vuln_type),
            severity=result.get("severity", severity),
            attack_summary=result.get("attack_summary", "Attack simulation generated"),
            success_probability=min(1.0, max(0.0, result.get("success_probability", 0.5))),
            required_conditions=result.get("required_conditions", []),
            total_steps=len(steps),
            steps=steps,
            exploit_primitives=exploit_primitives,
            final_impact=result.get("final_impact", "Unknown impact"),
            mitigations=mitigations,
            detection_opportunities=result.get("detection_opportunities", []),
            real_world_examples=result.get("real_world_examples", []),
            cve_references=result.get("cve_references", []),
            difficulty_rating=result.get("difficulty_rating", "unknown"),
        )
        
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse attack simulation response: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate attack simulation")
    except Exception as e:
        logger.error(f"Attack simulation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Simulation failed: {str(e)}")


# ============================================================================
# Symbolic Execution Traces
# ============================================================================

class PathConstraint(BaseModel):
    """A constraint on a path (e.g., input must be > 0)."""
    variable: str
    operator: str  # "==", "!=", ">", "<", ">=", "<=", "contains", "matches"
    value: str
    description: str
    is_satisfiable: bool = True


class TaintedVariable(BaseModel):
    """A variable that is tainted (derived from user input)."""
    name: str
    source: str  # Where the taint originated (e.g., "argv[1]", "read()", "recv()")
    taint_type: str  # "direct", "derived", "partial"
    propagation_chain: List[str]  # How taint spread: ["argv[1]", "buffer", "ptr"]
    reaches_sink: bool = False
    sink_name: Optional[str] = None  # e.g., "strcpy", "system", "eval"


class BasicBlock(BaseModel):
    """A basic block in the control flow graph."""
    id: str
    start_address: str
    end_address: str
    instructions_count: int
    code_preview: str
    is_entry: bool = False
    is_exit: bool = False
    is_reachable: bool = True
    predecessors: List[str] = []
    successors: List[str] = []
    dominates: List[str] = []
    tainted_at_entry: List[str] = []  # Which variables are tainted when entering this block


class ExecutionPath(BaseModel):
    """A single execution path through the function."""
    path_id: str
    blocks: List[str]  # Block IDs in order
    constraints: List[PathConstraint]  # Constraints to reach this path
    probability: float  # Estimated likelihood (0-1)
    is_feasible: bool = True
    leads_to_vulnerability: bool = False
    vulnerability_type: Optional[str] = None
    tainted_at_end: List[TaintedVariable] = []
    path_description: str
    interesting_operations: List[str] = []  # Notable things that happen on this path


class SymbolicState(BaseModel):
    """Symbolic state at a specific point in execution."""
    location: str  # Address or line
    variables: Dict[str, str]  # Variable name -> symbolic expression
    memory_regions: Dict[str, str]  # Address range -> description
    constraints: List[PathConstraint]
    tainted: List[str]  # Variable names that are tainted


class DangerousSink(BaseModel):
    """A dangerous function that tainted data might reach."""
    function_name: str
    address: str
    sink_type: str  # "memory", "command", "format", "file", "network"
    severity: str  # "critical", "high", "medium", "low"
    tainted_arguments: List[int]  # Which arguments are tainted (0-indexed)
    reachable_from: List[str]  # Path IDs that can reach this sink
    description: str
    cwe_id: Optional[str] = None


class SymbolicTraceRequest(BaseModel):
    """Request for symbolic execution trace."""
    code: str = Field(..., description="Decompiled function code")
    function_name: str = Field(default="unknown", description="Function name")
    entry_point: Optional[str] = Field(None, description="Entry point address")
    input_sources: List[str] = Field(
        default=["argv", "stdin", "recv", "read", "fgets", "scanf", "getenv"],
        description="Functions/variables considered as input sources (taint sources)"
    )
    dangerous_sinks: List[str] = Field(
        default=["strcpy", "strcat", "sprintf", "gets", "system", "exec", "eval", "memcpy", "memmove"],
        description="Dangerous functions to track (taint sinks)"
    )
    max_paths: int = Field(default=20, ge=1, le=100, description="Maximum paths to analyze")
    max_depth: int = Field(default=50, ge=5, le=200, description="Maximum path depth")
    binary_context: Optional[Dict[str, Any]] = Field(None, description="Additional binary context")


class SymbolicTraceResponse(BaseModel):
    """Response containing symbolic execution trace analysis."""
    function_name: str
    analysis_summary: str
    
    # Control Flow
    total_basic_blocks: int
    basic_blocks: List[BasicBlock]
    entry_block: str
    exit_blocks: List[str]
    
    # Paths
    total_paths_analyzed: int
    feasible_paths: int
    vulnerable_paths: int
    paths: List[ExecutionPath]
    
    # Taint Analysis
    input_sources_found: List[str]
    tainted_variables: List[TaintedVariable]
    dangerous_sinks_reached: List[DangerousSink]
    taint_summary: str
    
    # Symbolic States (at key points)
    symbolic_states: List[SymbolicState]
    
    # Security Assessment
    reachability_score: float  # 0-1, how much code is reachable
    vulnerability_score: float  # 0-1, likelihood of exploitable bugs
    recommended_focus_areas: List[str]
    
    # Integration data for AI Decompiler
    path_annotations: Dict[str, str]  # line/address -> annotation
    taint_annotations: Dict[str, str]  # variable -> taint info


@router.post("/binary/symbolic-trace", response_model=SymbolicTraceResponse)
async def analyze_symbolic_trace(
    request: SymbolicTraceRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-powered Symbolic Execution Trace Analysis.
    
    Performs comprehensive path and taint analysis on decompiled code:
    
    **Control Flow Analysis:**
    - Identifies all basic blocks and their relationships
    - Maps control flow graph structure
    - Determines reachable vs unreachable code
    
    **Path Enumeration:**
    - Enumerates execution paths through the function
    - Calculates constraints needed to reach each path
    - Identifies which paths lead to dangerous operations
    
    **Taint Analysis:**
    - Tracks user input from sources (argv, stdin, recv, etc.)
    - Propagates taint through assignments and operations
    - Identifies when tainted data reaches dangerous sinks
    
    **Integration with AI Decompiler:**
    - Returns annotations that can be used to enhance code
    - Marks lines with taint information
    - Highlights vulnerable paths in the code
    
    This analysis helps answer:
    - "Can user input reach this vulnerable function?"
    - "What input triggers this specific code path?"
    - "Which variables are attacker-controlled?"
    """
    from backend.core.config import settings
    from google import genai
    from google.genai import types

    if not settings.GEMINI_API_KEY:
        raise HTTPException(status_code=503, detail="Gemini API key not configured")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    model_id = "gemini-3-flash-preview"
    
    # Build context
    context_parts = []
    if request.binary_context:
        if request.binary_context.get("architecture"):
            context_parts.append(f"Architecture: {request.binary_context['architecture']}")
        if request.binary_context.get("imports"):
            context_parts.append(f"Imports: {', '.join(request.binary_context['imports'][:20])}")
    context = "\n".join(context_parts) if context_parts else "No additional context."
    
    prompt = f"""You are an expert binary analyst performing symbolic execution and taint analysis on decompiled code.

**Function Name:** {request.function_name}
**Entry Point:** {request.entry_point or "Unknown"}

**Input Sources (Taint Sources):** {", ".join(request.input_sources)}
**Dangerous Sinks:** {", ".join(request.dangerous_sinks)}

**Binary Context:**
{context}

**Decompiled Code:**
```c
{request.code[:8000]}
```

Perform comprehensive symbolic execution analysis:

1. **Control Flow Analysis**: Identify basic blocks, their relationships, and the CFG structure
2. **Path Analysis**: Enumerate distinct execution paths with their constraints
3. **Taint Analysis**: Track data flow from input sources to dangerous sinks
4. **Vulnerability Assessment**: Identify paths that lead to exploitable conditions

Respond with JSON (ONLY valid JSON, no markdown):

{{
    "function_name": "{request.function_name}",
    "analysis_summary": "Brief summary of the analysis findings",
    
    "total_basic_blocks": 8,
    "basic_blocks": [
        {{
            "id": "BB0",
            "start_address": "0x401000",
            "end_address": "0x401020",
            "instructions_count": 12,
            "code_preview": "if (argc < 2) return -1;",
            "is_entry": true,
            "is_exit": false,
            "is_reachable": true,
            "predecessors": [],
            "successors": ["BB1", "BB2"],
            "dominates": ["BB1", "BB2", "BB3"],
            "tainted_at_entry": []
        }},
        {{
            "id": "BB1",
            "start_address": "0x401024",
            "end_address": "0x401050",
            "instructions_count": 15,
            "code_preview": "buffer = malloc(size); strcpy(buffer, argv[1]);",
            "is_entry": false,
            "is_exit": false,
            "is_reachable": true,
            "predecessors": ["BB0"],
            "successors": ["BB3"],
            "dominates": [],
            "tainted_at_entry": ["argv[1]", "size"]
        }}
    ],
    "entry_block": "BB0",
    "exit_blocks": ["BB7"],
    
    "total_paths_analyzed": 5,
    "feasible_paths": 4,
    "vulnerable_paths": 2,
    "paths": [
        {{
            "path_id": "P1",
            "blocks": ["BB0", "BB1", "BB3", "BB5", "BB7"],
            "constraints": [
                {{
                    "variable": "argc",
                    "operator": ">=",
                    "value": "2",
                    "description": "Program requires at least one argument",
                    "is_satisfiable": true
                }},
                {{
                    "variable": "strlen(argv[1])",
                    "operator": ">",
                    "value": "256",
                    "description": "Input must overflow the buffer",
                    "is_satisfiable": true
                }}
            ],
            "probability": 0.3,
            "is_feasible": true,
            "leads_to_vulnerability": true,
            "vulnerability_type": "buffer_overflow",
            "tainted_at_end": [
                {{
                    "name": "buffer",
                    "source": "argv[1]",
                    "taint_type": "direct",
                    "propagation_chain": ["argv[1]", "buffer"],
                    "reaches_sink": true,
                    "sink_name": "strcpy"
                }}
            ],
            "path_description": "Main execution path that processes user input and copies to fixed buffer",
            "interesting_operations": ["malloc allocation", "strcpy without bounds check", "buffer passed to printf"]
        }}
    ],
    
    "input_sources_found": ["argv[1]", "getenv('HOME')"],
    "tainted_variables": [
        {{
            "name": "buffer",
            "source": "argv[1]",
            "taint_type": "direct",
            "propagation_chain": ["argv[1]", "buffer"],
            "reaches_sink": true,
            "sink_name": "strcpy"
        }},
        {{
            "name": "size",
            "source": "argv[2]",
            "taint_type": "derived",
            "propagation_chain": ["argv[2]", "atoi()", "size"],
            "reaches_sink": true,
            "sink_name": "malloc"
        }}
    ],
    "dangerous_sinks_reached": [
        {{
            "function_name": "strcpy",
            "address": "0x401030",
            "sink_type": "memory",
            "severity": "critical",
            "tainted_arguments": [1],
            "reachable_from": ["P1", "P2"],
            "description": "strcpy called with tainted source from argv[1], destination is 256-byte stack buffer",
            "cwe_id": "CWE-120"
        }},
        {{
            "function_name": "printf",
            "address": "0x401080",
            "sink_type": "format",
            "severity": "high",
            "tainted_arguments": [0],
            "reachable_from": ["P1"],
            "description": "printf called with tainted format string",
            "cwe_id": "CWE-134"
        }}
    ],
    "taint_summary": "User input from argv[1] flows directly to strcpy without length validation, allowing buffer overflow. Additionally, the buffer is later used as a printf format string.",
    
    "symbolic_states": [
        {{
            "location": "0x401024",
            "variables": {{
                "argc": "concrete: user-provided",
                "argv[1]": "symbolic: user_input_0",
                "buffer": "uninitialized"
            }},
            "memory_regions": {{
                "stack[rbp-0x100]": "256-byte local buffer (uninitialized)"
            }},
            "constraints": [
                {{
                    "variable": "argc",
                    "operator": ">=",
                    "value": "2",
                    "description": "Passed argument check",
                    "is_satisfiable": true
                }}
            ],
            "tainted": ["argv[1]"]
        }}
    ],
    
    "reachability_score": 0.85,
    "vulnerability_score": 0.75,
    "recommended_focus_areas": [
        "The strcpy at 0x401030 is the primary vulnerability - user input copied without bounds checking",
        "The printf at 0x401080 may allow format string exploitation",
        "Consider the malloc size being derived from user input - potential integer overflow"
    ],
    
    "path_annotations": {{
        "0x401024": "⚠️ Taint: argv[1] enters here",
        "0x401030": "🔴 SINK: strcpy with tainted source, CWE-120",
        "0x401050": "Taint propagates: buffer now tainted",
        "0x401080": "🔴 SINK: printf with tainted format, CWE-134"
    }},
    "taint_annotations": {{
        "buffer": "🔴 TAINTED from argv[1] → reaches strcpy, printf",
        "size": "⚠️ TAINTED from argv[2] → reaches malloc (potential int overflow)",
        "result": "✅ CLEAN - derived from return value, not user input"
    }}
}}

Analysis Guidelines:
1. Be thorough in identifying ALL paths, not just obvious ones
2. Track taint precisely through all operations (assignments, arithmetic, casts)
3. Consider indirect taint (e.g., array index from tainted value)
4. Identify ALL dangerous sinks the tainted data reaches
5. Provide actionable annotations for code enhancement
6. Focus on {request.max_paths} most interesting/vulnerable paths
7. Generate realistic addresses based on typical binary layouts

Return ONLY valid JSON."""

    try:
        response = await asyncio.to_thread(
            lambda: client.models.generate_content(
                model=model_id,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(thinking_level="medium"),
                    max_output_tokens=12000,
                )
            )
        )
        
        result_text = response.text.strip()
        
        # Clean up response
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        result_text = result_text.strip()
        
        result = json.loads(result_text)
        
        # Build response with proper validation
        basic_blocks = [BasicBlock(**bb) for bb in result.get("basic_blocks", [])]
        
        paths = []
        for path_data in result.get("paths", []):
            constraints = [PathConstraint(**c) for c in path_data.get("constraints", [])]
            tainted = [TaintedVariable(**t) for t in path_data.get("tainted_at_end", [])]
            paths.append(ExecutionPath(
                path_id=path_data.get("path_id", ""),
                blocks=path_data.get("blocks", []),
                constraints=constraints,
                probability=path_data.get("probability", 0.5),
                is_feasible=path_data.get("is_feasible", True),
                leads_to_vulnerability=path_data.get("leads_to_vulnerability", False),
                vulnerability_type=path_data.get("vulnerability_type"),
                tainted_at_end=tainted,
                path_description=path_data.get("path_description", ""),
                interesting_operations=path_data.get("interesting_operations", []),
            ))
        
        tainted_vars = [TaintedVariable(**t) for t in result.get("tainted_variables", [])]
        
        dangerous_sinks = [DangerousSink(**s) for s in result.get("dangerous_sinks_reached", [])]
        
        symbolic_states = []
        for state_data in result.get("symbolic_states", []):
            constraints = [PathConstraint(**c) for c in state_data.get("constraints", [])]
            symbolic_states.append(SymbolicState(
                location=state_data.get("location", ""),
                variables=state_data.get("variables", {}),
                memory_regions=state_data.get("memory_regions", {}),
                constraints=constraints,
                tainted=state_data.get("tainted", []),
            ))
        
        return SymbolicTraceResponse(
            function_name=result.get("function_name", request.function_name),
            analysis_summary=result.get("analysis_summary", ""),
            total_basic_blocks=result.get("total_basic_blocks", len(basic_blocks)),
            basic_blocks=basic_blocks,
            entry_block=result.get("entry_block", ""),
            exit_blocks=result.get("exit_blocks", []),
            total_paths_analyzed=result.get("total_paths_analyzed", len(paths)),
            feasible_paths=result.get("feasible_paths", len(paths)),
            vulnerable_paths=result.get("vulnerable_paths", 0),
            paths=paths,
            input_sources_found=result.get("input_sources_found", []),
            tainted_variables=tainted_vars,
            dangerous_sinks_reached=dangerous_sinks,
            taint_summary=result.get("taint_summary", ""),
            symbolic_states=symbolic_states,
            reachability_score=min(1.0, max(0.0, result.get("reachability_score", 0.5))),
            vulnerability_score=min(1.0, max(0.0, result.get("vulnerability_score", 0.0))),
            recommended_focus_areas=result.get("recommended_focus_areas", []),
            path_annotations=result.get("path_annotations", {}),
            taint_annotations=result.get("taint_annotations", {}),
        )
        
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse symbolic trace response: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse symbolic analysis results")
    except Exception as e:
        logger.error(f"Symbolic trace analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


# ============================================================================
# Enhanced Code with Symbolic Data Integration
# ============================================================================

class EnhanceCodeWithSymbolicRequest(BaseModel):
    """Request to enhance code with symbolic execution data."""
    code: str = Field(..., description="Decompiled code to enhance")
    function_name: str = Field(default="unknown", description="Function name")
    symbolic_trace: Optional[Dict[str, Any]] = Field(None, description="Symbolic trace data from /binary/symbolic-trace")
    binary_context: Optional[Dict[str, Any]] = Field(None, description="Additional binary context")
    enhancement_level: str = Field(default="full", description="Enhancement level: basic, standard, full")


class SymbolicEnhancedCodeResponse(BaseModel):
    """Response with code enhanced using symbolic execution data."""
    original_code: str
    enhanced_code: str
    suggested_function_name: str
    function_purpose: str
    
    # Standard enhancements
    variables: List[EnhancedVariable]
    code_blocks: List[EnhancedCodeBlock]
    data_structures: List[DataStructure]
    security_annotations: List[SecurityAnnotation]
    
    # Symbolic execution enhancements
    taint_annotations: List[Dict[str, Any]]  # Line-level taint info
    path_annotations: List[Dict[str, Any]]  # Which paths reach each line
    reachability_info: Dict[str, bool]  # Line -> is reachable
    constraint_annotations: List[Dict[str, Any]]  # Input constraints per block
    
    # Summary
    complexity_score: int
    vulnerability_paths_highlighted: int
    tainted_sinks_marked: int
    integration_quality: str  # How well symbolic data enhanced the analysis


@router.post("/binary/enhance-code-symbolic", response_model=SymbolicEnhancedCodeResponse)
async def enhance_code_with_symbolic(
    request: EnhanceCodeWithSymbolicRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-powered code enhancement with symbolic execution integration.
    
    This endpoint combines:
    1. **Standard Enhancement**: Variable renaming, comments, structure detection
    2. **Symbolic Integration**: Taint annotations, path info, reachability
    
    When symbolic_trace data is provided (from /binary/symbolic-trace):
    - Annotates variables with taint information
    - Shows which paths reach each code block
    - Highlights lines where tainted data reaches dangerous sinks
    - Marks unreachable code
    - Shows input constraints needed to reach specific lines
    
    This creates a comprehensive view that shows:
    - WHAT the code does (standard enhancement)
    - HOW data flows through it (taint tracking)
    - WHICH inputs trigger which behaviors (path constraints)
    - WHERE vulnerabilities exist (sink annotations)
    """
    from backend.core.config import settings
    from google import genai
    from google.genai import types

    if not settings.GEMINI_API_KEY:
        raise HTTPException(status_code=503, detail="Gemini API key not configured")

    client = genai.Client(api_key=settings.GEMINI_API_KEY)
    model_id = "gemini-3-flash-preview"
    
    # Build symbolic context
    symbolic_context = ""
    if request.symbolic_trace:
        st = request.symbolic_trace
        symbolic_context = f"""
**SYMBOLIC EXECUTION DATA AVAILABLE:**

Taint Summary: {st.get('taint_summary', 'N/A')}

Tainted Variables:
{json.dumps(st.get('tainted_variables', [])[:10], indent=2)}

Dangerous Sinks Reached:
{json.dumps(st.get('dangerous_sinks_reached', [])[:5], indent=2)}

Path Annotations:
{json.dumps(st.get('path_annotations', {}), indent=2)}

Taint Annotations:
{json.dumps(st.get('taint_annotations', {}), indent=2)}

Vulnerable Paths: {st.get('vulnerable_paths', 0)}
Vulnerability Score: {st.get('vulnerability_score', 0)}

Recommended Focus Areas:
{json.dumps(st.get('recommended_focus_areas', []), indent=2)}
"""
    
    # Build binary context
    binary_context = ""
    if request.binary_context:
        if request.binary_context.get("architecture"):
            binary_context += f"Architecture: {request.binary_context['architecture']}\n"
        if request.binary_context.get("imports"):
            binary_context += f"Imports: {', '.join(request.binary_context['imports'][:15])}\n"
    
    prompt = f"""You are an expert reverse engineer enhancing decompiled code with symbolic execution insights.

**Function Name:** {request.function_name}
**Enhancement Level:** {request.enhancement_level}

{binary_context}

{symbolic_context}

**Decompiled Code:**
```c
{request.code[:6000]}
```

Enhance this code by:
1. Renaming variables based on their usage AND taint status
2. Adding comments that explain BOTH behavior AND data flow
3. Marking tainted variables with special annotations
4. Highlighting lines where tainted data reaches dangerous sinks
5. Showing path constraints as comments where relevant

Respond with JSON (ONLY valid JSON):

{{
    "suggested_function_name": "process_user_input_unsafe",
    "function_purpose": "Processes command-line argument and copies to buffer without validation, vulnerable to overflow",
    
    "enhanced_code": "/* FUNCTION: process_user_input_unsafe
 * PURPOSE: Process user input (VULNERABLE - buffer overflow)
 * TAINT: argv[1] -> buffer -> strcpy sink
 * PATHS: 2 paths lead to vulnerable strcpy
 */
int process_user_input_unsafe(int argc, char **argv) {{
    // [PATH CONSTRAINT: argc >= 2 to reach this code]
    if (argc < 2) {{
        return -1;  // Early exit path (safe)
    }}
    
    // [TAINT SOURCE] user_input receives tainted data from argv[1]
    char *user_input = argv[1];  // 🔴 TAINTED: direct from argv
    
    // [ALLOCATION] Fixed-size buffer on stack
    char local_buffer[256];  // ⚠️ Target of overflow
    
    // [VULNERABLE SINK] strcpy with tainted source
    // 🔴 CWE-120: Buffer overflow - no bounds checking
    // CONSTRAINT: strlen(argv[1]) > 256 triggers overflow
    strcpy(local_buffer, user_input);  // 🔴 SINK: tainted data
    
    // [SECOND SINK] Format string vulnerability
    // 🔴 CWE-134: Tainted data used as format string
    printf(local_buffer);  // 🔴 SINK: format string
    
    return 0;
}}",
    
    "variables": [
        {{
            "original_name": "param_1",
            "suggested_name": "argc",
            "inferred_type": "int",
            "confidence": 0.95,
            "reasoning": "Standard main() argument count parameter"
        }},
        {{
            "original_name": "param_2",
            "suggested_name": "argv",
            "inferred_type": "char**",
            "confidence": 0.95,
            "reasoning": "Standard main() argument vector, TAINT SOURCE"
        }},
        {{
            "original_name": "local_108",
            "suggested_name": "local_buffer",
            "inferred_type": "char[256]",
            "confidence": 0.9,
            "reasoning": "Stack buffer used as strcpy destination, becomes tainted"
        }}
    ],
    
    "code_blocks": [
        {{
            "start_line": 1,
            "end_line": 4,
            "purpose": "Argument validation - early exit if no input",
            "security_notes": "Safe path - exits before processing tainted data"
        }},
        {{
            "start_line": 6,
            "end_line": 10,
            "purpose": "Tainted input processing - copies to fixed buffer",
            "security_notes": "CRITICAL: Buffer overflow vulnerability, tainted data from argv[1] copied without bounds check"
        }}
    ],
    
    "data_structures": [],
    
    "security_annotations": [
        {{
            "line": 8,
            "severity": "critical",
            "issue_type": "buffer_overflow",
            "description": "strcpy copies tainted argv[1] to 256-byte buffer without length check",
            "cwe_id": "CWE-120"
        }},
        {{
            "line": 10,
            "severity": "high",
            "issue_type": "format_string",
            "description": "Tainted buffer used directly as printf format string",
            "cwe_id": "CWE-134"
        }}
    ],
    
    "taint_annotations": [
        {{
            "line": 6,
            "variable": "user_input",
            "taint_source": "argv[1]",
            "taint_type": "direct",
            "annotation": "🔴 TAINTED: Direct user input from command line"
        }},
        {{
            "line": 8,
            "variable": "local_buffer",
            "taint_source": "argv[1]",
            "taint_type": "propagated",
            "annotation": "🔴 TAINTED: Contains copy of tainted argv[1]"
        }}
    ],
    
    "path_annotations": [
        {{
            "line": 3,
            "paths": ["P1"],
            "constraint": "argc < 2",
            "annotation": "Early exit path - safe, no tainted data processed"
        }},
        {{
            "line": 8,
            "paths": ["P2", "P3"],
            "constraint": "argc >= 2",
            "annotation": "Vulnerable path - tainted data reaches strcpy"
        }}
    ],
    
    "reachability_info": {{
        "1": true,
        "2": true,
        "3": true,
        "6": true,
        "8": true,
        "10": true
    }},
    
    "constraint_annotations": [
        {{
            "block": "argument_check",
            "constraint": "argc >= 2",
            "satisfiable": true,
            "annotation": "Requires at least one command-line argument"
        }},
        {{
            "block": "overflow_trigger",
            "constraint": "strlen(argv[1]) > 256",
            "satisfiable": true,
            "annotation": "Triggers buffer overflow when input exceeds buffer size"
        }}
    ],
    
    "complexity_score": 3,
    "vulnerability_paths_highlighted": 2,
    "tainted_sinks_marked": 2,
    "integration_quality": "excellent"
}}

Guidelines:
1. In enhanced_code, add visual markers: 🔴 for sinks, ⚠️ for warnings, ✅ for safe
2. Include taint flow in function header comment
3. Add PATH CONSTRAINT comments showing what inputs reach each block
4. Mark EVERY tainted variable with its source
5. Highlight ALL dangerous sinks with CWE IDs
6. Make the code self-documenting for security review

Return ONLY valid JSON."""

    try:
        response = await asyncio.to_thread(
            lambda: client.models.generate_content(
                model=model_id,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(thinking_level="medium"),
                    max_output_tokens=10000,
                )
            )
        )
        
        result_text = response.text.strip()
        
        # Clean up response
        if result_text.startswith("```json"):
            result_text = result_text[7:]
        if result_text.startswith("```"):
            result_text = result_text[3:]
        if result_text.endswith("```"):
            result_text = result_text[:-3]
        result_text = result_text.strip()
        
        result = json.loads(result_text)
        
        return SymbolicEnhancedCodeResponse(
            original_code=request.code,
            enhanced_code=result.get("enhanced_code", request.code),
            suggested_function_name=result.get("suggested_function_name", request.function_name),
            function_purpose=result.get("function_purpose", "Unknown"),
            variables=[EnhancedVariable(**v) for v in result.get("variables", [])],
            code_blocks=[EnhancedCodeBlock(**b) for b in result.get("code_blocks", [])],
            data_structures=[DataStructure(**d) for d in result.get("data_structures", [])],
            security_annotations=[SecurityAnnotation(**s) for s in result.get("security_annotations", [])],
            taint_annotations=result.get("taint_annotations", []),
            path_annotations=result.get("path_annotations", []),
            reachability_info=result.get("reachability_info", {}),
            constraint_annotations=result.get("constraint_annotations", []),
            complexity_score=result.get("complexity_score", 5),
            vulnerability_paths_highlighted=result.get("vulnerability_paths_highlighted", 0),
            tainted_sinks_marked=result.get("tainted_sinks_marked", 0),
            integration_quality=result.get("integration_quality", "unknown"),
        )
        
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse symbolic enhancement response: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse enhancement results")
    except Exception as e:
        logger.error(f"Symbolic code enhancement failed: {e}")
        raise HTTPException(status_code=500, detail=f"Enhancement failed: {str(e)}")


# ============================================================================
# APK Analysis
# ============================================================================

async def analyze_apk(
    file: UploadFile = File(..., description="Android APK file to analyze"),
    include_ai: bool = Query(True, description="Include AI-powered analysis"),
):
    """
    Analyze an Android APK file.
    
    Extracts:
    - Package info (name, version, SDK levels)
    - Permissions (with dangerous permission detection)
    - App components (activities, services, receivers, providers)
    - Strings from DEX files
    - Hardcoded URLs and secrets
    - Native libraries
    - Security issues
    
    Supported formats: APK, AAB
    """
    filename = file.filename or "unknown.apk"
    suffix = Path(filename).suffix.lower()
    
    if suffix not in ALLOWED_APK_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed extensions: {', '.join(ALLOWED_APK_EXTENSIONS)}"
        )
    
    tmp_dir = None
    try:
        # Save file to temp location with sanitized filename to prevent path traversal
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_apk_"))
        safe_filename = sanitize_filename(filename)
        tmp_path = tmp_dir / safe_filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Analyzing APK: {filename} ({file_size:,} bytes)")
        
        # Perform analysis
        result = re_service.analyze_apk(tmp_path)
        
        # Run AI analysis if requested (text analysis only - diagrams are generated after JADX decompilation)
        if include_ai and not result.error:
            result.ai_analysis = await re_service.analyze_apk_with_ai(result)
            # Note: Architecture diagram is only generated after JADX decompilation for proper context
        
        # Count dangerous permissions
        dangerous_count = sum(1 for p in result.permissions if p.is_dangerous)
        
        # Convert to response model
        return ApkAnalysisResponse(
            filename=result.filename,
            package_name=result.package_name,
            version_name=result.version_name,
            version_code=result.version_code,
            min_sdk=result.min_sdk,
            target_sdk=result.target_sdk,
            permissions=[
                ApkPermissionResponse(
                    name=p.name,
                    is_dangerous=p.is_dangerous,
                    description=p.description,
                )
                for p in result.permissions
            ],
            dangerous_permissions_count=dangerous_count,
            components=[
                ApkComponentResponse(
                    name=c.name,
                    component_type=c.component_type,
                    is_exported=c.is_exported,
                    intent_filters=c.intent_filters,
                )
                for c in result.components
            ],
            strings_count=len(result.strings),
            secrets=[
                SecretResponse(
                    type=s["type"],
                    value=s["value"],
                    masked_value=s["masked_value"],
                    severity=s["severity"],
                    context=s.get("context"),
                )
                for s in result.secrets
            ],
            urls=result.urls[:100],
            native_libraries=result.native_libraries,
            security_issues=[
                ApkSecurityIssueResponse(
                    category=issue["category"],
                    severity=issue["severity"],
                    description=issue["description"],
                    details=issue.get("details"),
                )
                for issue in result.security_issues
            ],
            ai_analysis=result.ai_analysis,
            ai_report_functionality=result.ai_report_functionality,
            ai_report_security=result.ai_report_security,
            ai_architecture_diagram=result.ai_architecture_diagram,
            ai_data_flow_diagram=result.ai_data_flow_diagram,
            error=result.error,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"APK analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


@router.get("/analyze-docker/{image_name:path}", response_model=DockerAnalysisResponse)
async def analyze_docker_image(
    image_name: str,
    project_id: Optional[int] = Query(None, description="Optional project scope for shared Docker images"),
    include_ai: bool = Query(True, description="Include AI-powered analysis"),
    adjudicate_findings: bool = Query(True, description="Use AI to filter false positives (recommended)"),
    skepticism_level: str = Query("high", description="How aggressively to filter: 'high' (default), 'medium', or 'low'"),
    deep_layer_scan: bool = Query(False, description="Extract and scan image layers for recoverable secrets"),
    check_base_image: bool = Query(True, description="Check base image against intelligence database (EOL, compromised, etc.)"),
    # CVE Scanning parameters (NEW)
    scan_cves: bool = Query(True, description="Scan image packages for CVEs using OSV/NVD"),
    include_nvd_enrichment: bool = Query(True, description="Enrich CVEs with NVD data (CVSS vectors, CWEs)"),
    include_kev: bool = Query(True, description="Check CISA Known Exploited Vulnerabilities catalog"),
    include_epss: bool = Query(True, description="Include EPSS exploit probability scores"),
    include_trivy_vuln_scan: bool = Query(False, description="Run Trivy's native vulnerability scanner alongside the OSV/NVD pipeline"),
    include_trivy_misconfig_scan: bool = Query(True, description="Run Trivy's misconfiguration scanner against the image"),
    include_trivy_secret_scan: bool = Query(True, description="Run Trivy's native secret scanner against the image filesystem"),
    include_trivy_license_scan: bool = Query(False, description="Run Trivy's license scanner against packages in the image"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    Analyze Docker image layers for secrets and security issues.

    Examines:
    - Image history and layer commands
    - ENV/ARG secrets
    - Hardcoded credentials in RUN commands
    - Security misconfigurations (root user, chmod 777, etc.)
    - Suspicious operations (curl | sh, sensitive file access)

    Features:
    - Image history and layer command analysis
    - Regex-based secret detection for build commands and env vars
    - AI False Positive Adjudicator (filters out noise with high skepticism)
    - **Layer Deep Scan**: Extract layers and find secrets in "deleted" files
    - **Base Image Intelligence**: Check curated EOL, compromised, typosquatting, and registry rules
    - **CVE Scanning**: Scan packages for known vulnerabilities with optional NVD/EPSS/KEV enrichment
    - **Native Trivy Scanners**: Optional vuln, misconfiguration, secret, and license findings

    Note: Requires Docker to be installed and the image must be pulled locally.
    """
    if not check_docker_available():
        raise HTTPException(
            status_code=503,
            detail="Docker is not available. Please install Docker to use this feature."
        )

    image_name = _ensure_docker_image_allowlisted(db, current_user.id, image_name, project_id)
    logger.info(f"Analyzing Docker image: {image_name}")

    try:
        analysis_started = time.perf_counter()

        # Perform analysis
        history_started = time.perf_counter()
        result = re_service.analyze_docker_image(image_name)
        history_duration = time.perf_counter() - history_started

        if result.error:
            raise HTTPException(status_code=400, detail=result.error)

        history_detail = (
            f"Parsed {result.total_layers} layers, found {len(result.secrets)} image-history secrets "
            f"and {len(result.security_issues)} raw security issues."
        )

        # Prepare security issues for potential adjudication
        security_issues = result.security_issues.copy()
        adjudication_summary = None
        rejected_findings = []
        adjudication_stats = None
        adjudication_status = "skipped"
        adjudication_detail = "AI adjudication was disabled for this run."
        adjudication_duration = 0.0

        # Run AI False Positive Adjudication if enabled
        if adjudicate_findings and security_issues:
            adjudication_started = time.perf_counter()
            try:
                from backend.services.docker_scan_service import ai_false_positive_adjudicator

                # Convert security issues to dict format for adjudicator
                issues_as_dicts = [
                    {
                        "finding_index": idx,
                        "rule_id": issue.get("rule_id", "N/A"),
                        "severity": issue.get("severity", "unknown"),
                        "message": issue.get("description", ""),
                        "category": issue.get("category", ""),
                        "command": issue.get("command", ""),
                    }
                    for idx, issue in enumerate(security_issues)
                ]

                # Build context for adjudicator (reconstruct from layers)
                dockerfile_context = "\n".join([
                    f"# Layer: {layer.get('command', '')}"
                    for layer in result.layers
                ])

                confirmed, rejected, summary = await ai_false_positive_adjudicator(
                    findings=issues_as_dicts,
                    dockerfile_content=dockerfile_context,
                    skepticism_level=skepticism_level,
                )

                # Filter security_issues to only confirmed (by index when available)
                confirmed_indices = {
                    f.get("finding_index") for f in confirmed
                    if f.get("finding_index") is not None
                }
                if confirmed_indices:
                    security_issues = [
                        issue for idx, issue in enumerate(security_issues)
                        if idx in confirmed_indices
                    ]
                else:
                    confirmed_messages = {f.get("message") for f in confirmed}
                    security_issues = [
                        issue for issue in security_issues
                        if issue.get("description") in confirmed_messages
                    ]

                adjudication_summary = summary
                rejected_findings = rejected
                adjudication_stats = {
                    "confirmed": len(confirmed),
                    "rejected": len(rejected),
                    "total": len(issues_as_dicts),
                }
                adjudication_status = "completed"
                adjudication_detail = (
                    f"Confirmed {len(confirmed)} of {len(issues_as_dicts)} raw findings and "
                    f"rejected {len(rejected)} as likely false positives."
                )

                logger.info(f"Adjudication complete: {len(confirmed)} confirmed, {len(rejected)} rejected")

            except Exception as e:
                logger.warning(f"Adjudication failed, returning all findings: {e}")
                adjudication_summary = f"Adjudication failed: {e}"
                adjudication_status = "error"
                adjudication_detail = f"AI adjudication failed: {e}"
            finally:
                adjudication_duration = time.perf_counter() - adjudication_started
        elif adjudicate_findings:
            adjudication_status = "completed"
            adjudication_detail = "AI adjudication was enabled, but there were no initial security issues to review."

        # Keep result in sync for AI analysis
        result.security_issues = security_issues

        # Base Image Intelligence check
        base_image_intel = []
        base_image_status = "skipped"
        base_image_detail = "Base-image intelligence was disabled for this run."
        base_image_duration = 0.0
        if check_base_image and result.base_image:
            base_image_started = time.perf_counter()
            try:
                from backend.services.docker_scan_service import check_base_image_intelligence
                intel_findings = check_base_image_intelligence(result.base_image)
                base_image_intel = [
                    BaseImageIntelligenceResponse(
                        image=f.image,
                        category=f.category,
                        severity=f.severity,
                        message=f.message,
                        attack_vector=f.attack_vector,
                        recommendation=f.recommendation,
                    )
                    for f in intel_findings
                ]
                base_image_status = "completed"
                base_image_detail = (
                    f"Checked inferred base image {result.base_image}; "
                    f"{len(intel_findings)} intelligence findings returned."
                )
                if intel_findings:
                    logger.info(f"Base image intelligence: {len(intel_findings)} findings for {result.base_image}")
            except Exception as e:
                logger.warning(f"Base image intelligence check failed: {e}")
                base_image_status = "error"
                base_image_detail = f"Base-image intelligence failed: {e}"
            finally:
                base_image_duration = time.perf_counter() - base_image_started
        elif check_base_image:
            base_image_detail = "Base-image intelligence was requested, but the base image could not be inferred."

        # Layer Deep Scan for recoverable secrets
        layer_secrets = []
        layer_scan_metadata = None
        deleted_secrets_count = 0
        deleted_files = result.deleted_files or []
        deep_layer_status = "skipped"
        deep_layer_detail = "Deep layer extraction was disabled for this run."
        deep_layer_duration = 0.0
        if deep_layer_scan:
            deep_layer_started = time.perf_counter()
            try:
                from backend.services.docker_scan_service import extract_and_scan_layers
                secrets_found, scan_meta = extract_and_scan_layers(image_name, max_layers=20)
                layer_secrets = [
                    LayerSecretResponse(
                        layer_id=s.layer_id,
                        layer_index=s.layer_index,
                        file_path=s.file_path,
                        file_type=s.file_type,
                        severity=s.severity,
                        size_bytes=s.size_bytes,
                        is_deleted=s.is_deleted,
                        content_preview=s.content_preview[:100] if s.content_preview else None,
                        entropy=s.entropy,
                        attack_vector=s.attack_vector,
                    )
                    for s in secrets_found
                ]
                layer_scan_metadata = scan_meta
                deleted_secrets_count = scan_meta.get("deleted_secrets_found", 0)
                if isinstance(scan_meta, dict) and scan_meta.get("deleted_files") is not None:
                    deleted_files = scan_meta.get("deleted_files", deleted_files)
                deep_layer_status = "completed"
                deep_layer_detail = (
                    f"Scanned {scan_meta.get('layers_scanned', 'N/A')} layers and "
                    f"{scan_meta.get('files_scanned', 'N/A')} files; found {len(secrets_found)} layer secrets, "
                    f"including {deleted_secrets_count} recoverable deleted-file secrets."
                )
                if secrets_found:
                    logger.info(f"Layer deep scan: {len(secrets_found)} secrets found, {deleted_secrets_count} in 'deleted' files")
            except Exception as e:
                logger.warning(f"Layer deep scan failed: {e}")
                layer_scan_metadata = {"error": str(e)}
                deep_layer_status = "error"
                deep_layer_detail = f"Layer deep scan failed: {e}"
            finally:
                deep_layer_duration = time.perf_counter() - deep_layer_started

        # CVE Scanning - scan image packages for known vulnerabilities
        cve_scan_result = None
        cve_status = "skipped"
        cve_detail = "Package CVE scanning was disabled for this run."
        cve_duration = 0.0
        if scan_cves:
            cve_started = time.perf_counter()
            try:
                from backend.services.docker_scan_service import scan_image_packages_for_cves

                cve_result = await scan_image_packages_for_cves(
                    image_name,
                    include_nvd_enrichment=include_nvd_enrichment,
                    include_kev=include_kev,
                    include_epss=include_epss,
                )

                if cve_result.vulnerabilities or cve_result.packages_scanned > 0:
                    # Convert vulnerabilities to response format
                    vuln_responses = []
                    for vuln in cve_result.vulnerabilities:
                        vuln_responses.append(DockerCVEVulnerabilityResponse(
                            external_id=vuln.get("external_id", ""),
                            title=vuln.get("title"),
                            description=vuln.get("description"),
                            severity=vuln.get("severity"),
                            cvss_score=vuln.get("cvss_score"),
                            cvss_vector=vuln.get("cvss_vector"),
                            source=vuln.get("source", "osv"),
                            package_name=vuln.get("package_name"),
                            package_version=vuln.get("package_version"),
                            package_ecosystem=vuln.get("package_ecosystem"),
                            package_type=vuln.get("package_type"),
                            in_kev=vuln.get("in_kev", False),
                            epss_score=vuln.get("epss_score"),
                            epss_percentile=vuln.get("epss_percentile"),
                            epss_priority=vuln.get("epss_priority"),
                            combined_priority=vuln.get("combined_priority"),
                            priority_label=vuln.get("priority_label"),
                            nvd_enrichment=vuln.get("nvd_enrichment"),
                        ))

                cve_scan_result = DockerCVEScanResponse(
                    image_name=cve_result.image_name,
                    packages_scanned=cve_result.packages_scanned,
                    packages_with_vulns=cve_result.packages_with_vulns,
                    vulnerabilities=vuln_responses,
                    total_vulnerabilities=len(vuln_responses),
                    critical_count=cve_result.critical_count,
                    high_count=cve_result.high_count,
                    medium_count=cve_result.medium_count,
                    low_count=cve_result.low_count,
                    kev_count=cve_result.kev_count,
                    high_epss_count=cve_result.high_epss_count,
                    scan_duration_seconds=cve_result.scan_duration_seconds,
                    trivy_available=cve_result.trivy_available,
                    enrichment_applied=cve_result.enrichment_applied,
                    epss_enabled=cve_result.epss_enabled,
                    epss_available=cve_result.epss_available,
                    epss_source=cve_result.epss_source,
                    error=cve_result.error,
                )
                cve_status = "error" if cve_result.error else "completed"
                if cve_result.error:
                    cve_detail = f"Package CVE scan failed: {cve_result.error}"
                else:
                    cve_detail = (
                        f"Scanned {cve_result.packages_scanned} packages and found {len(vuln_responses)} vulnerabilities "
                        f"({cve_result.critical_count} critical, {cve_result.high_count} high, "
                        f"{cve_result.kev_count} in KEV)."
                    )

                logger.info(
                    f"CVE scan: {len(vuln_responses)} CVEs found in {cve_result.packages_scanned} packages "
                    f"({cve_result.critical_count} critical, {cve_result.high_count} high, "
                    f"{cve_result.kev_count} in KEV)"
                )

            except Exception as e:
                logger.warning(f"CVE scanning failed: {e}")
                cve_scan_result = DockerCVEScanResponse(
                    image_name=image_name,
                    packages_scanned=0,
                    packages_with_vulns=0,
                    vulnerabilities=[],
                    total_vulnerabilities=0,
                    epss_enabled=False,
                    error=str(e),
                )
                cve_status = "error"
                cve_detail = f"Package CVE scan failed: {e}"
            finally:
                cve_duration = time.perf_counter() - cve_started

        trivy_scan_result = None
        trivy_requested = any([include_trivy_vuln_scan, include_trivy_misconfig_scan, include_trivy_secret_scan, include_trivy_license_scan])
        trivy_status = "skipped"
        trivy_detail = "Native Trivy scanners were disabled for this run."
        trivy_duration = 0.0
        if trivy_requested:
            trivy_started = time.perf_counter()
            try:
                from backend.services.docker_scan_service import scan_image_with_trivy_native

                native_trivy = await scan_image_with_trivy_native(
                    image_name,
                    include_vulnerabilities=include_trivy_vuln_scan,
                    include_misconfigurations=include_trivy_misconfig_scan,
                    include_secrets=include_trivy_secret_scan,
                    include_licenses=include_trivy_license_scan,
                )
                trivy_scan_result = DockerTrivyNativeScanResponse(
                    image_name=native_trivy.image_name,
                    enabled_scanners=native_trivy.enabled_scanners,
                    trivy_available=native_trivy.trivy_available,
                    artifact_name=native_trivy.artifact_name,
                    artifact_type=native_trivy.artifact_type,
                    os_family=native_trivy.os_family,
                    os_name=native_trivy.os_name,
                    trivy_version=native_trivy.trivy_version,
                    schema_version=native_trivy.schema_version,
                    package_count=native_trivy.package_count,
                    vulnerabilities=native_trivy.vulnerabilities,
                    misconfigurations=native_trivy.misconfigurations,
                    secrets=native_trivy.secrets,
                    licenses=native_trivy.licenses,
                    vulnerability_severity_counts=native_trivy.vulnerability_severity_counts,
                    misconfiguration_severity_counts=native_trivy.misconfiguration_severity_counts,
                    secret_severity_counts=native_trivy.secret_severity_counts,
                    license_severity_counts=native_trivy.license_severity_counts,
                    scan_duration_seconds=native_trivy.scan_duration_seconds,
                    error=native_trivy.error,
                )
                total_trivy_findings = (
                    len(native_trivy.vulnerabilities)
                    + len(native_trivy.misconfigurations)
                    + len(native_trivy.secrets)
                    + len(native_trivy.licenses)
                )
                trivy_status = "error" if native_trivy.error else "completed"
                trivy_detail = (
                    f"Ran scanners: {', '.join(native_trivy.enabled_scanners) or 'none'}; "
                    f"observed {native_trivy.package_count} packages and collected {total_trivy_findings} findings."
                )
                if native_trivy.error:
                    trivy_detail = f"Native Trivy scan failed: {native_trivy.error}"
            except Exception as e:
                logger.warning(f"Native Trivy scanning failed: {e}")
                trivy_scan_result = DockerTrivyNativeScanResponse(
                    image_name=image_name,
                    enabled_scanners=[
                        scanner
                        for scanner, enabled in [
                            ("vuln", include_trivy_vuln_scan),
                            ("misconfig", include_trivy_misconfig_scan),
                            ("secret", include_trivy_secret_scan),
                            ("license", include_trivy_license_scan),
                        ]
                        if enabled
                    ],
                    trivy_available=False,
                    error=str(e),
                )
                trivy_status = "error"
                trivy_detail = f"Native Trivy scan failed: {e}"
            finally:
                trivy_duration = time.perf_counter() - trivy_started

        # Run AI analysis if requested
        ai_status = "skipped"
        ai_detail = "AI security analysis was disabled for this run."
        ai_duration = 0.0
        if include_ai:
            ai_started = time.perf_counter()
            base_image_intel_payload = [
                intel.dict() if hasattr(intel, "dict") else intel
                for intel in base_image_intel
            ]
            layer_secrets_payload = [
                secret.dict() if hasattr(secret, "dict") else secret
                for secret in layer_secrets
            ]
            cve_scan_payload = cve_scan_result.dict() if hasattr(cve_scan_result, "dict") and cve_scan_result else cve_scan_result
            trivy_scan_payload = trivy_scan_result.dict() if hasattr(trivy_scan_result, "dict") and trivy_scan_result else trivy_scan_result
            result.ai_analysis = await re_service.analyze_docker_with_ai(
                result,
                base_image_intel=base_image_intel_payload,
                layer_secrets=layer_secrets_payload,
                layer_scan_metadata=layer_scan_metadata,
                cve_scan=cve_scan_payload,
                trivy_scan=trivy_scan_payload,
            )
            ai_duration = time.perf_counter() - ai_started
            if result.ai_analysis and not str(result.ai_analysis).startswith("AI analysis unavailable"):
                ai_status = "completed"
                ai_detail = "Generated AI security analysis for the completed scan results."
            else:
                ai_status = "error"
                ai_detail = result.ai_analysis or "AI analysis returned no content."

        total_duration = time.perf_counter() - analysis_started
        scan_summary = {
            "total_duration_seconds": round(total_duration, 3),
            "checks": [
                {
                    "id": "history_analysis",
                    "label": "Image history and layer command analysis",
                    "status": "completed",
                    "duration_seconds": round(history_duration, 3),
                    "detail": history_detail,
                },
                {
                    "id": "ai_adjudication",
                    "label": "AI false-positive adjudication",
                    "status": adjudication_status,
                    "duration_seconds": round(adjudication_duration, 3),
                    "detail": adjudication_detail,
                },
                {
                    "id": "base_image_intelligence",
                    "label": "Base-image intelligence",
                    "status": base_image_status,
                    "duration_seconds": round(base_image_duration, 3),
                    "detail": base_image_detail,
                },
                {
                    "id": "deep_layer_scan",
                    "label": "Deep layer extraction and deleted-file forensics",
                    "status": deep_layer_status,
                    "duration_seconds": round(deep_layer_duration, 3),
                    "detail": deep_layer_detail,
                },
                {
                    "id": "package_cve_scan",
                    "label": "Package CVE scan",
                    "status": cve_status,
                    "duration_seconds": round(cve_duration, 3),
                    "detail": cve_detail,
                },
                {
                    "id": "native_trivy",
                    "label": "Native Trivy scanners",
                    "status": trivy_status,
                    "duration_seconds": round(trivy_duration, 3),
                    "detail": trivy_detail,
                },
                {
                    "id": "ai_summary",
                    "label": "AI security summary",
                    "status": ai_status,
                    "duration_seconds": round(ai_duration, 3),
                    "detail": ai_detail,
                },
            ],
        }

        # Convert to response model
        return DockerAnalysisResponse(
            image_name=result.image_name,
            image_id=result.image_id,
            total_layers=result.total_layers,
            total_size=result.total_size,
            total_size_human=format_size(result.total_size),
            base_image=result.base_image,
            layers=[
                DockerLayerResponse(
                    id=layer["id"],
                    command=layer["command"],
                    size=layer["size"],
                )
                for layer in result.layers
            ],
            secrets=[
                DockerSecretResponse(
                    layer_id=s.layer_id,
                    layer_command=s.layer_command,
                    secret_type=s.secret_type,
                    value=s.value,
                    masked_value=s.masked_value,
                    context=s.context,
                    severity=s.severity,
                )
                for s in result.secrets
            ],
            deleted_files=deleted_files,
            security_issues=[
                DockerSecurityIssueResponse(
                    category=issue["category"],
                    severity=issue["severity"],
                    description=issue["description"],
                    command=issue.get("command"),
                    attack_vector=issue.get("attack_vector"),
                    rule_id=issue.get("rule_id"),
                )
                for issue in security_issues
            ],
            ai_analysis=result.ai_analysis,
            error=result.error,
            adjudication_enabled=adjudicate_findings,
            adjudication_summary=adjudication_summary,
            rejected_findings=rejected_findings,
            adjudication_stats=adjudication_stats,
            # New fields
            base_image_intel=base_image_intel,
            layer_secrets=layer_secrets,
            layer_scan_metadata=layer_scan_metadata,
            deleted_secrets_count=deleted_secrets_count,
            # CVE scanning
            cve_scan=cve_scan_result,
            trivy_scan=trivy_scan_result,
            scan_summary=scan_summary,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Docker analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


# =============================================================================
# Docker Export Endpoint - Comprehensive formatted exports
# =============================================================================

class DockerExportRequest(BaseModel):
    """Request body for Docker export."""
    image_name: str
    image_id: str
    total_layers: int
    total_size: int
    total_size_human: str
    base_image: Optional[str] = None
    layers: List[Dict[str, Any]] = []
    secrets: List[Dict[str, Any]] = []
    deleted_files: List[Dict[str, Any]] = []
    security_issues: List[Dict[str, Any]] = []
    ai_analysis: Optional[str] = None
    adjudication_enabled: bool = False
    adjudication_summary: Optional[str] = None
    rejected_findings: List[Dict[str, Any]] = []
    adjudication_stats: Optional[Dict[str, int]] = None
    base_image_intel: List[Dict[str, Any]] = []
    layer_secrets: List[Dict[str, Any]] = []
    layer_scan_metadata: Optional[Dict[str, Any]] = None
    deleted_secrets_count: int = 0
    cve_scan: Optional[Dict[str, Any]] = None
    trivy_scan: Optional[Dict[str, Any]] = None
    scan_summary: Optional[Dict[str, Any]] = None
    risk_score: Optional[int] = None
    risk_level: Optional[str] = None


@router.post("/docker/export")
async def export_docker_analysis(
    request: DockerExportRequest,
    format: str = Query(..., description="Export format: markdown, pdf, docx"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Export Docker analysis results to Markdown, PDF, or Word format.

    Features properly formatted exports with:
    - Executive summary with risk assessment
    - Base image intelligence findings (EOL, typosquatting, etc.)
    - Security issues with attack vectors
    - Layer deep scan results (recoverable secrets)
    - AI analysis and false positive adjudication summary
    """
    from fastapi.responses import Response
    from io import BytesIO
    from datetime import datetime

    if format not in ["markdown", "pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Invalid format. Use: markdown, pdf, docx")

    try:
        # Generate comprehensive markdown content
        md_content = _generate_docker_export_markdown(request)

        # Return based on format
        base_name = request.image_name.replace('/', '_').replace(':', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        if format == "markdown":
            return Response(
                content=md_content.encode('utf-8'),
                media_type="text/markdown",
                headers={"Content-Disposition": f'attachment; filename="{base_name}_docker_analysis_{timestamp}.md"'}
            )
        elif format == "pdf":
            pdf_content = _generate_docker_pdf(md_content, request)
            return Response(
                content=pdf_content,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{base_name}_docker_analysis_{timestamp}.pdf"'}
            )
        elif format == "docx":
            docx_content = _generate_docker_docx(md_content, request)
            return Response(
                content=docx_content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{base_name}_docker_analysis_{timestamp}.docx"'}
            )

    except Exception as e:
        logger.error(f"Docker export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


@router.get("/docker/sbom/{image_name:path}")
async def export_docker_sbom(
    image_name: str,
    project_id: Optional[int] = Query(None, description="Optional project scope for shared Docker images"),
    format: str = Query("cyclonedx", description="SBOM format: cyclonedx or spdx-json"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """Export a Docker image SBOM using Trivy."""
    from fastapi.responses import Response
    from datetime import datetime
    from backend.services.docker_scan_service import generate_trivy_sbom

    image_name = _ensure_docker_image_allowlisted(db, current_user.id, image_name, project_id)

    sbom_content, media_type, error = generate_trivy_sbom(image_name, format=format)
    if error:
        raise HTTPException(status_code=400, detail=error)
    if sbom_content is None:
        raise HTTPException(status_code=500, detail="Failed to generate SBOM")

    normalized_format = str(format or "cyclonedx").strip().lower()
    ext = "json" if normalized_format in {"cyclonedx", "spdx-json"} else "txt"
    safe_name = image_name.replace("/", "_").replace(":", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return Response(
        content=sbom_content,
        media_type=media_type or "application/json",
        headers={
            "Content-Disposition": (
                f'attachment; filename="{safe_name}_sbom_{normalized_format}_{timestamp}.{ext}"'
            )
        },
    )


def _format_docker_export_duration(value: Any) -> str:
    """Format seconds for export-friendly display."""
    try:
        seconds = float(value or 0)
    except (TypeError, ValueError):
        seconds = 0.0
    if seconds >= 60:
        return f"{seconds / 60:.1f} min"
    return f"{seconds:.1f} sec"


def _normalize_docker_export_text(value: Any) -> str:
    """Normalize AI/report text for Markdown, PDF, and Word exports."""
    import re

    text = str(value or "").strip()
    if not text:
        return ""

    text = re.sub(r"^```(?:html|markdown|md)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    if "<" in text and ">" in text:
        text = _html_to_plain_text(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?m)^---+\s*$", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _collect_docker_export_findings(request: DockerExportRequest) -> Dict[str, Any]:
    """Build a shared summary model for Docker report exports."""
    security_issues = request.security_issues or []
    cve_scan = request.cve_scan if isinstance(request.cve_scan, dict) else {}
    cve_vulnerabilities = cve_scan.get("vulnerabilities") if isinstance(cve_scan.get("vulnerabilities"), list) else []
    trivy_scan = request.trivy_scan if isinstance(request.trivy_scan, dict) else {}
    trivy_vulns = trivy_scan.get("vulnerabilities") if isinstance(trivy_scan.get("vulnerabilities"), list) else []
    trivy_misconfigs = trivy_scan.get("misconfigurations") if isinstance(trivy_scan.get("misconfigurations"), list) else []
    trivy_secrets = trivy_scan.get("secrets") if isinstance(trivy_scan.get("secrets"), list) else []
    trivy_licenses = trivy_scan.get("licenses") if isinstance(trivy_scan.get("licenses"), list) else []

    severity_counts = {
        "critical": sum(1 for item in security_issues if item.get("severity", "").lower() == "critical"),
        "high": sum(1 for item in security_issues if item.get("severity", "").lower() == "high"),
        "medium": sum(1 for item in security_issues if item.get("severity", "").lower() == "medium"),
        "low": sum(1 for item in security_issues if item.get("severity", "").lower() == "low"),
    }

    total_cves = int(cve_scan.get("total_vulnerabilities") or len(cve_vulnerabilities) or 0)
    total_trivy_findings = len(trivy_vulns) + len(trivy_misconfigs) + len(trivy_secrets) + len(trivy_licenses)
    total_secrets = len(request.secrets or []) + len(request.layer_secrets or [])
    deleted_recoverable = int(request.deleted_secrets_count or sum(1 for item in (request.layer_secrets or []) if item.get("is_deleted")))

    risk_score = request.risk_score
    if risk_score is None:
        risk_score = (
            severity_counts["critical"] * 40
            + severity_counts["high"] * 25
            + severity_counts["medium"] * 10
            + severity_counts["low"] * 3
        )
        if cve_scan.get("kev_count"):
            risk_score += 25
        if deleted_recoverable:
            risk_score += 20
    risk_score = min(int(risk_score), 100)

    risk_level = request.risk_level or (
        "Critical" if risk_score >= 75 else
        "High" if risk_score >= 45 else
        "Medium" if risk_score >= 20 else
        "Low"
    )

    scan_summary = request.scan_summary if isinstance(request.scan_summary, dict) else {}
    scan_checks = scan_summary.get("checks") if isinstance(scan_summary.get("checks"), list) else []
    coverage_rows: List[Dict[str, str]] = []
    if scan_checks:
        for check in scan_checks:
            coverage_rows.append({
                "label": str(check.get("label") or check.get("id") or "Check"),
                "status": str(check.get("status") or "unknown").replace("_", " ").title(),
                "duration": _format_docker_export_duration(check.get("duration_seconds")),
                "detail": str(check.get("detail") or "").strip(),
            })
    else:
        inferred_checks = [
            ("Image history analysis", "Completed", "Included in every Docker report."),
            ("Base image intelligence", "Completed" if request.base_image_intel else "Not requested", ""),
            ("Layer deep scan", "Completed" if (request.layer_secrets or request.layer_scan_metadata) else "Not requested", ""),
            ("Package CVE scan", "Completed" if request.cve_scan else "Not requested", cve_scan.get("error", "")),
            ("Native Trivy scan", "Completed" if request.trivy_scan else "Not requested", trivy_scan.get("error", "")),
            ("AI security analysis", "Completed" if request.ai_analysis else "Not requested", ""),
        ]
        for label, status, detail in inferred_checks:
            coverage_rows.append({
                "label": label,
                "status": status,
                "duration": "N/A",
                "detail": detail,
            })

    findings: List[str] = []
    if deleted_recoverable:
        findings.append(
            f"{deleted_recoverable} deleted-but-recoverable secrets remain inside preserved image layers."
        )
    if request.secrets:
        findings.append(
            f"{len(request.secrets)} secrets were exposed directly in image history or build commands."
        )
    if cve_scan.get("kev_count"):
        findings.append(
            f"{cve_scan.get('kev_count')} package vulnerabilities are listed in CISA KEV."
        )
    if severity_counts["critical"] or severity_counts["high"]:
        findings.append(
            f"{severity_counts['critical']} critical and {severity_counts['high']} high heuristic findings were identified."
        )
    if request.base_image_intel:
        findings.append(
            f"{len(request.base_image_intel)} base image intelligence findings require supply-chain review."
        )
    if trivy_misconfigs:
        findings.append(
            f"Trivy reported {len(trivy_misconfigs)} image misconfigurations in addition to native history analysis."
        )
    if not findings:
        findings.append("No high-priority findings were highlighted in the exported Docker result.")

    actions: List[str] = []
    if total_secrets:
        actions.append("Rotate all exposed credentials, scrub them from build inputs, and rebuild the image from a clean context.")
    if deleted_recoverable:
        actions.append("Rebuild the image so sensitive files never enter intermediate layers; deleting them later is not sufficient.")
    if total_cves:
        actions.append("Patch vulnerable packages, prioritizing KEV and high-EPSS CVEs, then republish the image with a pinned base tag.")
    if request.base_image_intel:
        actions.append("Review the base image lineage and move to a supported, trusted base image where findings indicate risk.")
    if severity_counts["critical"] or severity_counts["high"]:
        actions.append("Harden the runtime posture by removing dangerous build steps, dropping privileges, and minimizing image contents.")
    if not actions:
        actions.append("Retain this report as the clean baseline and re-run after the next image update.")

    return {
        "severity_counts": severity_counts,
        "cve_scan": cve_scan,
        "cve_vulnerabilities": cve_vulnerabilities,
        "total_cves": total_cves,
        "trivy_scan": trivy_scan,
        "trivy_vulns": trivy_vulns,
        "trivy_misconfigs": trivy_misconfigs,
        "trivy_secrets": trivy_secrets,
        "trivy_licenses": trivy_licenses,
        "total_trivy_findings": total_trivy_findings,
        "total_secrets": total_secrets,
        "deleted_recoverable": deleted_recoverable,
        "risk_score": risk_score,
        "risk_level": risk_level,
        "coverage_rows": coverage_rows,
        "key_findings": findings[:5],
        "priority_actions": actions[:5],
        "ai_text": _normalize_docker_export_text(request.ai_analysis),
        "scan_summary": scan_summary,
    }


def _generate_docker_export_markdown(request: DockerExportRequest) -> str:
    """Generate comprehensive markdown export for Docker analysis."""
    from datetime import datetime

    def _inline_code(value: Any) -> str:
        return str(value or "N/A").replace("`", "'")

    overview = _collect_docker_export_findings(request)
    severity_counts = overview["severity_counts"]
    critical_count = severity_counts["critical"]
    high_count = severity_counts["high"]
    medium_count = severity_counts["medium"]
    low_count = severity_counts["low"]
    cve_scan = overview["cve_scan"]
    cve_vulnerabilities = overview["cve_vulnerabilities"]
    total_cves = overview["total_cves"]
    trivy_scan = overview["trivy_scan"]
    trivy_vulns = overview["trivy_vulns"]
    trivy_misconfigs = overview["trivy_misconfigs"]
    trivy_secrets = overview["trivy_secrets"]
    trivy_licenses = overview["trivy_licenses"]
    total_trivy_findings = overview["total_trivy_findings"]
    total_secrets = overview["total_secrets"]
    risk_score = overview["risk_score"]
    risk_level = overview["risk_level"]

    md = f"""# Docker Security Analysis Report

> Analyst-ready export from **VRAgent Docker Inspector**
> Generated **{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}**

---

## Report Header

| Field | Value |
|-------|-------|
| **Image** | `{request.image_name}` |
| **Image ID** | `{request.image_id[:12]}` |
| **Base Image** | `{request.base_image or 'Unknown'}` |
| **Image Size** | {request.total_size_human} |
| **Total Layers** | {request.total_layers} |
| **Risk Rating** | **{risk_level.upper()}** ({risk_score}/100) |

## At A Glance

"""
    for finding in overview["key_findings"]:
        md += f"- {finding}\n"

    md += f"""

## Priority Actions

"""
    for index, action in enumerate(overview["priority_actions"], start=1):
        md += f"{index}. {action}\n"

    md += f"""

## Executive Summary

| Metric | Value |
|--------|-------|
| **Risk Level** | **{risk_level}** ({risk_score}/100) |
| **Total Layers** | {request.total_layers} |
| **Image Size** | {request.total_size_human} |
| **Base Image** | `{request.base_image or 'Unknown'}` |
| **Security Issues** | {len(request.security_issues)} |
| **Secrets Found** | {total_secrets} |
| **CVEs Found** | {total_cves} |
| **Native Trivy Findings** | {total_trivy_findings} |
| **Deleted but Recoverable** | {overview["deleted_recoverable"]} |

### Risk Breakdown

| Severity | Count | Impact |
|----------|-------|--------|
| Critical | {critical_count} | Immediate exploitation possible |
| High | {high_count} | High-value attack targets |
| Medium | {medium_count} | Secondary targets |
| Low | {low_count} | Opportunistic |

"""

    if overview["coverage_rows"]:
        md += """---

## Scan Coverage

| Check | Status | Duration | Notes |
|-------|--------|----------|-------|
"""
        for row in overview["coverage_rows"]:
            detail = str(row.get("detail") or "").replace("\n", " ").strip() or " "
            md += (
                f"| {row.get('label', 'Check')} | {row.get('status', 'Unknown')} | "
                f"{row.get('duration', 'N/A')} | {detail} |\n"
            )
        md += "\n"

    if cve_scan:
        md += f"""---

## Package CVE Scan

*Known vulnerabilities detected in packages extracted from the image.*

| Metric | Value |
|--------|-------|
| **Packages Scanned** | {cve_scan.get('packages_scanned', 0)} |
| **Packages With Vulnerabilities** | {cve_scan.get('packages_with_vulns', 0)} |
| **Total CVEs** | {total_cves} |
| **Critical** | {cve_scan.get('critical_count', 0)} |
| **High** | {cve_scan.get('high_count', 0)} |
| **In CISA KEV** | {cve_scan.get('kev_count', 0)} |
| **High EPSS** | {cve_scan.get('high_epss_count', 0)} |

"""
        if cve_scan.get("error"):
            md += f"**Scan Error:** {cve_scan.get('error')}\n\n"
        elif cve_vulnerabilities:
            md += """### Highest Priority CVEs

| CVE | Severity | Package | Version | Signals |
|-----|----------|---------|---------|---------|
"""
            for vuln in cve_vulnerabilities[:20]:
                signals: List[str] = []
                if vuln.get("in_kev"):
                    signals.append("KEV")
                if vuln.get("priority_label"):
                    signals.append(str(vuln.get("priority_label")))
                elif vuln.get("epss_percentile") is not None:
                    signals.append(f"EPSS {vuln.get('epss_percentile')}")
                md += (
                    f"| {vuln.get('external_id', vuln.get('id', 'N/A'))} | "
                    f"{str(vuln.get('severity', 'unknown')).upper()} | "
                    f"{vuln.get('package_name', vuln.get('package', 'N/A'))} | "
                    f"{vuln.get('package_version', vuln.get('installed_version', 'N/A'))} | "
                    f"{', '.join(signals) or 'N/A'} |\n"
                )
            md += "\n"

    if trivy_scan:
        native_vulns = trivy_scan.get("vulnerabilities") if isinstance(trivy_scan.get("vulnerabilities"), list) else []
        native_misconfigs = trivy_scan.get("misconfigurations") if isinstance(trivy_scan.get("misconfigurations"), list) else []
        native_secrets = trivy_scan.get("secrets") if isinstance(trivy_scan.get("secrets"), list) else []
        native_licenses = trivy_scan.get("licenses") if isinstance(trivy_scan.get("licenses"), list) else []
        enabled_scanners = ", ".join(trivy_scan.get("enabled_scanners") or []) or "None"
        trivy_duration = trivy_scan.get("scan_duration_seconds") or 0
        try:
            trivy_duration = float(trivy_duration)
        except (TypeError, ValueError):
            trivy_duration = 0.0
        md += f"""---

## Native Trivy Scan

*Raw Trivy scanner output captured alongside the Docker Inspector's own analysis.*

| Metric | Value |
|--------|-------|
| **Enabled Scanners** | {enabled_scanners} |
| **Packages Observed** | {trivy_scan.get('package_count', 0)} |
| **Native Vulnerabilities** | {len(native_vulns)} |
| **Misconfigurations** | {len(native_misconfigs)} |
| **Secret Findings** | {len(native_secrets)} |
| **License Findings** | {len(native_licenses)} |
| **Duration** | {trivy_duration:.1f}s |

"""
        if trivy_scan.get("error"):
            md += f"**Scan Error:** {trivy_scan.get('error')}\n\n"

        if native_vulns:
            md += """### Native Trivy Vulnerabilities

| ID | Severity | Package | Installed | Fixed |
|----|----------|---------|-----------|-------|
"""
            for vuln in native_vulns[:20]:
                md += (
                    f"| {vuln.get('vulnerability_id', 'N/A')} | "
                    f"{str(vuln.get('severity', 'unknown')).upper()} | "
                    f"{vuln.get('package_name', 'N/A')} | "
                    f"{vuln.get('installed_version', 'N/A')} | "
                    f"{vuln.get('fixed_version', 'N/A') or '-'} |\n"
                )
            md += "\n"

        if native_misconfigs:
            md += """### Trivy Misconfigurations

| ID | Severity | Target | Message |
|----|----------|--------|---------|
"""
            for finding in native_misconfigs[:20]:
                md += (
                    f"| {finding.get('id', 'N/A')} | "
                    f"{str(finding.get('severity', 'unknown')).upper()} | "
                    f"{str(finding.get('target', 'N/A'))[:40]} | "
                    f"{str(finding.get('message') or finding.get('title') or 'N/A')[:80]} |\n"
                )
            md += "\n"

        if native_secrets:
            md += """### Trivy Native Secret Findings

*Operator note: raw Trivy secret matches are intentionally included for remediation workflows.*

| Rule | Severity | Target | Match |
|------|----------|--------|-------|
"""
            for finding in native_secrets[:20]:
                md += (
                    f"| {finding.get('rule_id', 'N/A')} | "
                    f"{str(finding.get('severity', 'unknown')).upper()} | "
                    f"{str(finding.get('target', 'N/A'))[:40]} | "
                    f"`{_inline_code(finding.get('match'))}` |\n"
                )
            md += "\n"

        if native_licenses:
            md += """### Trivy License Findings

| License | Severity | Package | Category |
|---------|----------|---------|----------|
"""
            for finding in native_licenses[:20]:
                md += (
                    f"| {finding.get('name', 'N/A')} | "
                    f"{str(finding.get('severity', 'unknown')).upper()} | "
                    f"{finding.get('package_name', 'N/A')} | "
                    f"{finding.get('category', 'N/A')} |\n"
                )
            md += "\n"

    # Base Image Intelligence Section
    if request.base_image_intel:
        md += """---

## Base Image Intelligence

*Findings from automated intelligence checks on the base image.*

"""
        for intel in request.base_image_intel:
            severity_icon = "🔴" if intel.get('severity') == 'critical' else "🟠" if intel.get('severity') == 'high' else "🟡" if intel.get('severity') == 'medium' else "🟢"
            md += f"""### {severity_icon} {intel.get('category', 'Unknown').upper()}: {intel.get('message', '')}

- **Image:** `{intel.get('image', 'N/A')}`
- **Severity:** {intel.get('severity', 'N/A').upper()}
- **Attack Vector:** {intel.get('attack_vector', 'N/A')}
- **Recommendation:** {intel.get('recommendation', 'N/A')}

"""

    # Security Issues Section
    if request.security_issues:
        md += """---

## Security Issues

*Identified security vulnerabilities and misconfigurations.*

"""
        # Group by severity
        for severity in ['critical', 'high', 'medium', 'low']:
            issues = [i for i in request.security_issues if i.get('severity', '').lower() == severity]
            if issues:
                severity_title = severity.upper()
                md += f"### {severity_title} Severity ({len(issues)})\n\n"
                for issue in issues:
                    md += f"""#### {issue.get('category', 'Unknown')}

- **Description:** {issue.get('description', 'N/A')}
- **Rule ID:** `{issue.get('rule_id', 'N/A')}`
"""
                    if issue.get('attack_vector'):
                        md += f"- **Attack Vector:** {issue.get('attack_vector')}\n"
                    if issue.get('command'):
                        md += f"- **Command:** `{issue.get('command')[:100]}...`\n"
                    md += "\n"

    # Layer Secrets Section (Deep Scan)
    if request.layer_secrets:
        md += """---

## Layer Deep Scan - Recoverable Secrets

*Sensitive files found in image layers, including "deleted" files that are still recoverable.*

"""
        # Highlight deleted but recoverable
        deleted_secrets = [s for s in request.layer_secrets if s.get('is_deleted')]
        if deleted_secrets:
            md += f"""### ⚠️ CRITICAL: {len(deleted_secrets)} Deleted But Recoverable Secrets

These files were "deleted" in later layers but **can still be extracted** from the image using `docker save`.

| Layer | File | Type | Severity |
|-------|------|------|----------|
"""
            for secret in deleted_secrets:
                md += f"| {secret.get('layer_index', '?')} | `{secret.get('file_path', 'N/A')[:50]}` | {secret.get('file_type', 'N/A')} | {secret.get('severity', 'N/A').upper()} |\n"
            md += "\n"

        # All layer secrets
        md += f"""### All Sensitive Files ({len(request.layer_secrets)})

| Layer | File Path | Type | Deleted? | Severity |
|-------|-----------|------|----------|----------|
"""
        for secret in request.layer_secrets[:30]:
            deleted_mark = "⚠️ YES" if secret.get('is_deleted') else "No"
            md += f"| {secret.get('layer_index', '?')} | `{secret.get('file_path', 'N/A')[:40]}` | {secret.get('file_type', 'N/A')} | {deleted_mark} | {secret.get('severity', 'N/A').upper()} |\n"

        if len(request.layer_secrets) > 30:
            md += f"\n*...and {len(request.layer_secrets) - 30} more secrets*\n"
        md += "\n"

    # Traditional Secrets Section
    if request.secrets:
        md += """---

## Secrets in Image History

*Secrets detected in Docker image build commands and environment variables.*

*Operator note: raw secret values are intentionally included in this Docker report for remediation workflows.*

"""
        for secret in request.secrets:
            md += f"""### {secret.get('secret_type', 'Unknown')} ({secret.get('severity', 'N/A').upper()})

- **Layer:** `{secret.get('layer_id', 'N/A')[:12]}`
- **Context:** {secret.get('context', 'N/A')}
- **Value (raw):** `{_inline_code(secret.get('value'))}`
- **Value (masked):** `{_inline_code(secret.get('masked_value'))}`
- **Layer Command:** `{_inline_code(secret.get('layer_command'))}`

"""

    # AI Analysis Section
    if overview["ai_text"]:
        md += f"""---

## AI Security Analysis

{overview["ai_text"]}

"""

    # Adjudication Summary
    if request.adjudication_enabled and request.adjudication_summary:
        md += f"""---

## False Positive Adjudication

*AI-powered filtering of false positives with high skepticism.*

**Summary:** {request.adjudication_summary}

"""
        if request.adjudication_stats:
            stats = request.adjudication_stats
            md += f"""| Metric | Count |
|--------|-------|
| Total Findings | {stats.get('total', 0)} |
| Confirmed Real | {stats.get('confirmed', 0)} |
| Rejected as False Positive | {stats.get('rejected', 0)} |

"""

        if request.rejected_findings:
            md += """### Rejected Findings (False Positives)

| Finding | Reason |
|---------|--------|
"""
            for finding in request.rejected_findings[:10]:
                reason = finding.get('_adjudication', {}).get('reason', 'N/A')[:60]
                message = finding.get('message', 'N/A')[:40]
                md += f"| {message} | {reason} |\n"
            md += "\n"

    # Layer History Section
    if request.layers:
        md += """---

## Image Layer History

*Build commands that created each layer.*

| Layer | Command | Size |
|-------|---------|------|
"""
        for i, layer in enumerate(request.layers[:20]):
            cmd = layer.get('command', 'N/A')[:60]
            size = layer.get('size', 0)
            md += f"| {i} | `{cmd}` | {size} |\n"

        if len(request.layers) > 20:
            md += f"\n*...and {len(request.layers) - 20} more layers*\n"
        md += "\n"

    # Footer
    md += """---

## Report Information

- **Tool:** VRAgent Docker Inspector
- **Features Used:** Image history analysis, regex secret detection, AI adjudication, layer deep scan, base image intelligence, package CVE scanning, native Trivy scanners, SBOM export
- **Document Structure:** Report header, at-a-glance findings, priority actions, scan coverage, technical sections, AI narrative
- **Offensive Focus:** Container escape, privilege escalation, secrets extraction, supply chain risks

*This report was generated with an offensive security perspective, emphasizing exploitable weaknesses and attack vectors.*
"""

    return md


def _generate_docker_pdf(md_content: str, request: DockerExportRequest) -> bytes:
    """Generate PDF from Docker analysis with proper formatting."""
    from io import BytesIO

    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem, KeepTogether
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        sanitize_pdf_text = globals().get("_sanitize_pdf_text")
        if callable(sanitize_pdf_text):
            _sanitize_pdf_text = sanitize_pdf_text
        else:
            def _sanitize_pdf_text(value: Any) -> str:
                text = str(value or "")
                return "".join(ch for ch in text if ch == "\n" or ch == "\t" or ord(ch) >= 32)

        overview = _collect_docker_export_findings(request)
        severity_counts = overview["severity_counts"]
        cve_scan = overview["cve_scan"]
        trivy_scan = overview["trivy_scan"]
        trivy_vulns = overview["trivy_vulns"]
        trivy_misconfigs = overview["trivy_misconfigs"]
        trivy_secrets = overview["trivy_secrets"]
        trivy_licenses = overview["trivy_licenses"]
        total_trivy_findings = overview["total_trivy_findings"]
        risk_score = overview["risk_score"]
        risk_level = overview["risk_level"]

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.55*inch, bottomMargin=0.55*inch)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=20, spaceAfter=6, textColor=colors.HexColor('#0f172a'), alignment=TA_CENTER)
        subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=9.5, textColor=colors.HexColor('#475569'), alignment=TA_CENTER, spaceAfter=14)
        h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=14, spaceBefore=16, spaceAfter=8, textColor=colors.HexColor('#16213e'))
        h3_style = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=11.5, spaceBefore=10, spaceAfter=6, textColor=colors.HexColor('#0f3460'))
        body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, leading=13, spaceAfter=6, textColor=colors.HexColor('#1f2937'))
        note_style = ParagraphStyle('Note', parent=body_style, textColor=colors.HexColor('#475569'))
        small_style = ParagraphStyle('Small', parent=body_style, fontSize=8.5, leading=11, textColor=colors.HexColor('#64748b'))

        def _build_table(data: List[List[str]], widths: List[float], header_color: str) -> Table:
            table = Table(data, colWidths=widths, hAlign="LEFT")
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(header_color)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8.5),
                ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#cbd5e1')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#f8fafc'), colors.white]),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ]))
            return table

        def _draw_page_chrome(canvas, pdf_doc):
            canvas.saveState()
            canvas.setStrokeColor(colors.HexColor('#cbd5e1'))
            canvas.setLineWidth(0.5)
            canvas.line(pdf_doc.leftMargin, 0.45 * inch, letter[0] - pdf_doc.rightMargin, 0.45 * inch)
            canvas.setFont('Helvetica', 8)
            canvas.setFillColor(colors.HexColor('#64748b'))
            canvas.drawString(pdf_doc.leftMargin, 0.28 * inch, f"VRAgent Docker Inspector • {request.image_name}")
            canvas.drawRightString(letter[0] - pdf_doc.rightMargin, 0.28 * inch, f"Page {canvas.getPageNumber()}")
            canvas.restoreState()

        elements = []

        title_block = [
            [Paragraph("<b>Docker Security Analysis Report</b>", title_style)],
            [Paragraph("Analyst-ready export covering native Docker Inspector findings, CVEs, Trivy, and AI narrative.", subtitle_style)],
        ]
        title_table = Table(title_block, colWidths=[6.7 * inch], hAlign="CENTER")
        title_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#e2e8f0')),
            ('BOX', (0, 0), (-1, -1), 0.6, colors.HexColor('#cbd5e1')),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ]))
        elements.append(title_table)
        elements.append(Spacer(1, 0.18 * inch))

        metadata_data = [
            ['Field', 'Value'],
            ['Image', _sanitize_pdf_text(request.image_name)],
            ['Image ID', _sanitize_pdf_text(request.image_id[:12])],
            ['Base Image', _sanitize_pdf_text(request.base_image or 'Unknown')],
            ['Generated', _sanitize_pdf_text(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))],
            ['Risk Rating', f"{risk_level.upper()} ({risk_score}/100)"],
        ]
        elements.append(_build_table(metadata_data, [1.6 * inch, 5.0 * inch], '#0f172a'))
        elements.append(Spacer(1, 0.16 * inch))

        elements.append(Paragraph("At A Glance", h2_style))
        elements.append(ListFlowable(
            [ListItem(Paragraph(_sanitize_pdf_text(item), body_style)) for item in overview["key_findings"]],
            bulletType='bullet',
            leftIndent=16,
        ))
        elements.append(Spacer(1, 0.08 * inch))

        elements.append(Paragraph("Priority Actions", h2_style))
        elements.append(ListFlowable(
            [ListItem(Paragraph(_sanitize_pdf_text(item), body_style), value=index) for index, item in enumerate(overview["priority_actions"], start=1)],
            bulletType='1',
            leftIndent=18,
        ))
        elements.append(Spacer(1, 0.12 * inch))

        elements.append(Paragraph("Executive Summary", h2_style))

        summary_data = [
            ['Metric', 'Value'],
            ['Risk Score', f"{risk_score}/100"],
            ['Total Layers', str(request.total_layers)],
            ['Image Size', request.total_size_human],
            ['Security Issues', str(len(request.security_issues))],
            ['Secrets Found', str(overview["total_secrets"])],
            ['CVEs Found', str(overview["total_cves"])],
            ['Native Trivy Findings', str(total_trivy_findings)],
            ['Deleted Secrets', str(overview["deleted_recoverable"])],
        ]
        elements.append(_build_table(summary_data, [2 * inch, 3.2 * inch], '#1a1a2e'))
        elements.append(Spacer(1, 0.12 * inch))

        risk_breakdown = [
            ['Severity', 'Count', 'Meaning'],
            ['Critical', str(severity_counts['critical']), 'Immediate exploitation or direct compromise'],
            ['High', str(severity_counts['high']), 'Material attacker advantage'],
            ['Medium', str(severity_counts['medium']), 'Important but not immediately critical'],
            ['Low', str(severity_counts['low']), 'Hardening and hygiene gaps'],
        ]
        elements.append(_build_table(risk_breakdown, [1.1 * inch, 0.7 * inch, 4.2 * inch], '#334155'))
        elements.append(Spacer(1, 0.16 * inch))

        if overview["coverage_rows"]:
            coverage_data = [['Check', 'Status', 'Duration', 'Notes']]
            for row in overview["coverage_rows"]:
                coverage_data.append([
                    _sanitize_pdf_text(row.get('label', 'Check')),
                    _sanitize_pdf_text(row.get('status', 'Unknown')),
                    _sanitize_pdf_text(row.get('duration', 'N/A')),
                    _sanitize_pdf_text((row.get('detail') or '').strip() or '-'),
                ])
            elements.append(Paragraph("Scan Coverage", h2_style))
            elements.append(_build_table(coverage_data, [1.8 * inch, 1.0 * inch, 0.8 * inch, 3.1 * inch], '#475569'))
            elements.append(Spacer(1, 0.15 * inch))

        # Base Image Intelligence
        if request.base_image_intel:
            elements.append(Paragraph("Base Image Intelligence", h2_style))
            for intel in request.base_image_intel:
                sev = intel.get('severity', 'N/A').upper()
                elements.append(Paragraph(f"<b>{sev}:</b> {intel.get('message', 'N/A')}", body_style))
                elements.append(Paragraph(f"<i>Recommendation:</i> {intel.get('recommendation', 'N/A')}", body_style))
            elements.append(Spacer(1, 0.1*inch))

        # Security Issues
        if request.security_issues:
            elements.append(Paragraph("Security Issues", h2_style))
            issues_data = [['Severity', 'Category', 'Description']]
            for issue in request.security_issues[:15]:
                issues_data.append([
                    issue.get('severity', 'N/A').upper(),
                    issue.get('category', 'N/A')[:20],
                    issue.get('description', 'N/A')[:50] + '...' if len(issue.get('description', '')) > 50 else issue.get('description', 'N/A')
                ])
            issues_table = Table(issues_data, colWidths=[1*inch, 1.5*inch, 4*inch])
            issues_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e94560')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(issues_table)
            elements.append(Spacer(1, 0.2*inch))

        if request.cve_scan:
            elements.append(Paragraph("Package CVE Scan", h2_style))
            cve_data = [
                ['Metric', 'Value'],
                ['Packages Scanned', str(request.cve_scan.get('packages_scanned', 0))],
                ['Packages With Vulnerabilities', str(request.cve_scan.get('packages_with_vulns', 0))],
                ['Total CVEs', str(request.cve_scan.get('total_vulnerabilities') or len(request.cve_scan.get('vulnerabilities', [])))],
                ['Critical', str(request.cve_scan.get('critical_count', 0))],
                ['High', str(request.cve_scan.get('high_count', 0))],
                ['KEV', str(request.cve_scan.get('kev_count', 0))],
            ]
            cve_table = Table(cve_data, colWidths=[2.2*inch, 2.8*inch])
            cve_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d6a4f')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(cve_table)
            elements.append(Spacer(1, 0.2*inch))

        if trivy_scan:
            elements.append(Paragraph("Native Trivy Scan", h2_style))
            enabled_scanners = ", ".join(trivy_scan.get("enabled_scanners") or []) or "None"
            trivy_duration = trivy_scan.get("scan_duration_seconds") or 0
            try:
                trivy_duration = float(trivy_duration)
            except (TypeError, ValueError):
                trivy_duration = 0.0

            trivy_data = [
                ['Metric', 'Value'],
                ['Enabled Scanners', _sanitize_pdf_text(enabled_scanners)],
                ['Packages Observed', str(trivy_scan.get('package_count', 0))],
                ['Native Vulnerabilities', str(len(trivy_vulns))],
                ['Misconfigurations', str(len(trivy_misconfigs))],
                ['Secret Findings', str(len(trivy_secrets))],
                ['License Findings', str(len(trivy_licenses))],
                ['Duration', f"{trivy_duration:.1f}s"],
            ]
            trivy_table = Table(trivy_data, colWidths=[2.2*inch, 3.8*inch])
            trivy_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1d4ed8')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(trivy_table)
            if trivy_scan.get("error"):
                elements.append(Spacer(1, 0.08*inch))
                elements.append(Paragraph(
                    f"<b>Scan Error:</b> {_sanitize_pdf_text(trivy_scan.get('error'))}",
                    body_style,
                ))
            elements.append(Spacer(1, 0.15*inch))

            if trivy_vulns:
                elements.append(Paragraph("Trivy Vulnerabilities", h3_style))
                vuln_rows = [['ID', 'Severity', 'Package', 'Installed', 'Fixed']]
                for finding in trivy_vulns[:12]:
                    vuln_rows.append([
                        _sanitize_pdf_text(str(finding.get('vulnerability_id', 'N/A'))[:24]),
                        _sanitize_pdf_text(str(finding.get('severity', 'N/A')).upper()),
                        _sanitize_pdf_text(str(finding.get('package_name', 'N/A'))[:18]),
                        _sanitize_pdf_text(str(finding.get('installed_version', 'N/A'))[:18]),
                        _sanitize_pdf_text(str(finding.get('fixed_version') or '-')[:18]),
                    ])
                vuln_table = Table(vuln_rows, colWidths=[1.4*inch, 0.8*inch, 1.4*inch, 1.2*inch, 1.2*inch])
                vuln_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                elements.append(vuln_table)
                elements.append(Spacer(1, 0.12*inch))

            if trivy_misconfigs:
                elements.append(Paragraph("Trivy Misconfigurations", h3_style))
                misconfig_rows = [['ID', 'Severity', 'Target', 'Message']]
                for finding in trivy_misconfigs[:12]:
                    misconfig_rows.append([
                        _sanitize_pdf_text(str(finding.get('id', 'N/A'))[:24]),
                        _sanitize_pdf_text(str(finding.get('severity', 'N/A')).upper()),
                        _sanitize_pdf_text(str(finding.get('target', 'N/A'))[:26]),
                        _sanitize_pdf_text(str(finding.get('message') or finding.get('title') or 'N/A')[:70]),
                    ])
                misconfig_table = Table(misconfig_rows, colWidths=[1.3*inch, 0.8*inch, 1.4*inch, 2.5*inch])
                misconfig_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f766e')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                elements.append(misconfig_table)
                elements.append(Spacer(1, 0.12*inch))

            if trivy_secrets:
                elements.append(Paragraph("Trivy Secret Findings", h3_style))
                elements.append(Paragraph(
                    "Operator note: raw Trivy secret matches are intentionally included for remediation workflows.",
                    body_style,
                ))
                for finding in trivy_secrets[:10]:
                    elements.append(Paragraph(
                        f"{_sanitize_pdf_text(finding.get('rule_id', 'Unknown'))} "
                        f"({_sanitize_pdf_text(str(finding.get('severity', 'N/A')).upper())})",
                        h3_style,
                    ))
                    elements.append(Paragraph(
                        f"<b>Target:</b> {_sanitize_pdf_text(finding.get('target', 'N/A'))}",
                        body_style,
                    ))
                    elements.append(Paragraph(
                        f"<b>Match:</b> {_sanitize_pdf_text(finding.get('match', 'N/A'))}",
                        body_style,
                    ))
                    elements.append(Spacer(1, 0.06*inch))
                elements.append(Spacer(1, 0.08*inch))

            if trivy_licenses:
                elements.append(Paragraph("Trivy License Findings", h3_style))
                license_rows = [['License', 'Severity', 'Package', 'Category']]
                for finding in trivy_licenses[:12]:
                    license_rows.append([
                        _sanitize_pdf_text(str(finding.get('name', 'N/A'))[:20]),
                        _sanitize_pdf_text(str(finding.get('severity', 'N/A')).upper()),
                        _sanitize_pdf_text(str(finding.get('package_name', 'N/A'))[:18]),
                        _sanitize_pdf_text(str(finding.get('category', 'N/A'))[:18]),
                    ])
                license_table = Table(license_rows, colWidths=[1.4*inch, 0.8*inch, 1.5*inch, 1.3*inch])
                license_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                elements.append(license_table)
                elements.append(Spacer(1, 0.12*inch))

        # Layer Secrets
        if request.layer_secrets:
            elements.append(Paragraph("Layer Deep Scan - Recoverable Secrets", h2_style))
            secrets_data = [['Layer', 'File Path', 'Type', 'Deleted?']]
            for secret in request.layer_secrets[:15]:
                secrets_data.append([
                    str(secret.get('layer_index', '?')),
                    secret.get('file_path', 'N/A')[:35],
                    secret.get('file_type', 'N/A'),
                    'YES' if secret.get('is_deleted') else 'No'
                ])
            secrets_table = Table(secrets_data, colWidths=[0.6*inch, 3.5*inch, 1.2*inch, 0.8*inch])
            secrets_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#ff6b6b')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            elements.append(secrets_table)
            elements.append(Spacer(1, 0.15*inch))

        if request.secrets:
            elements.append(Paragraph("Secrets in Image History", h2_style))
            elements.append(Paragraph(
                "Operator note: raw secret values are intentionally included in this Docker report for remediation workflows.",
                body_style,
            ))
            for secret in request.secrets[:20]:
                elements.append(Paragraph(
                    f"{_sanitize_pdf_text(secret.get('secret_type', 'Unknown'))} "
                    f"({_sanitize_pdf_text(str(secret.get('severity', 'N/A')).upper())})",
                    h3_style,
                ))
                elements.append(Paragraph(
                    f"<b>Layer:</b> {_sanitize_pdf_text(secret.get('layer_id', 'N/A')[:12])}",
                    body_style,
                ))
                elements.append(Paragraph(
                    f"<b>Raw Value:</b> {_sanitize_pdf_text(secret.get('value', 'N/A'))}",
                    body_style,
                ))
                elements.append(Paragraph(
                    f"<b>Context:</b> {_sanitize_pdf_text(secret.get('context', 'N/A'))}",
                    body_style,
                ))
                elements.append(Paragraph(
                    f"<b>Layer Command:</b> {_sanitize_pdf_text(secret.get('layer_command', 'N/A'))}",
                    body_style,
                ))
                elements.append(Spacer(1, 0.08*inch))

        if overview["ai_text"]:
            elements.append(Paragraph("AI Security Analysis", h2_style))
            elements.append(Paragraph(
                "Narrative summary from the analyst-facing AI layer, normalized for export formatting.",
                note_style,
            ))
            for block in overview["ai_text"].split("\n\n"):
                clean_block = _sanitize_pdf_text(block.strip())
                if not clean_block:
                    continue
                if clean_block.startswith("• "):
                    bullets = [line[2:].strip() for line in clean_block.splitlines() if line.strip().startswith("• ")]
                    elements.append(ListFlowable(
                        [ListItem(Paragraph(_html_to_reportlab(item), body_style)) for item in bullets],
                        bulletType='bullet',
                        leftIndent=16,
                    ))
                else:
                    elements.append(Paragraph(_html_to_reportlab(clean_block.replace("\n", "<br/>")), body_style))
                elements.append(Spacer(1, 0.04 * inch))

        elements.append(Spacer(1, 0.08 * inch))
        elements.append(Paragraph("Report Information", h2_style))
        elements.append(Paragraph(
            "Prepared by VRAgent Docker Inspector with consistent formatting across Markdown, PDF, and Word exports.",
            small_style,
        ))
        elements.append(Paragraph(
            f"Document sections included: {', '.join(['overview', 'actions', 'coverage', 'findings', 'AI narrative'])}.",
            small_style,
        ))

        doc.build(elements, onFirstPage=_draw_page_chrome, onLaterPages=_draw_page_chrome)
        return buffer.getvalue()

    except ImportError as e:
        logger.warning(f"PDF generation requires reportlab: {e}")
        # Fallback: return markdown as plain text
        return md_content.encode('utf-8')


def _generate_docker_docx(md_content: str, request: DockerExportRequest) -> bytes:
    """Generate Word document from Docker analysis with proper formatting."""
    from io import BytesIO

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        sanitize_word_text = globals().get("_sanitize_word_text")
        if callable(sanitize_word_text):
            _sanitize_word_text = sanitize_word_text
        else:
            def _sanitize_word_text(value: Any) -> str:
                text = str(value or "")
                return "".join(ch for ch in text if ch == "\n" or ch == "\t" or ord(ch) >= 32)

        overview = _collect_docker_export_findings(request)
        severity_counts = overview["severity_counts"]
        cve_scan = overview["cve_scan"]
        trivy_scan = overview["trivy_scan"]
        trivy_vulns = overview["trivy_vulns"]
        trivy_misconfigs = overview["trivy_misconfigs"]
        trivy_secrets = overview["trivy_secrets"]
        trivy_licenses = overview["trivy_licenses"]
        total_trivy_findings = overview["total_trivy_findings"]
        risk_score = overview["risk_score"]
        risk_level = overview["risk_level"]

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('VRAgent Docker Inspector')
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(100, 116, 139)

        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Aptos'
        normal_style.font.size = Pt(10)
        doc.styles['Heading 1'].font.name = 'Aptos'
        doc.styles['Heading 1'].font.size = Pt(16)
        doc.styles['Heading 1'].font.color.rgb = RGBColor(15, 23, 42)
        doc.styles['Heading 2'].font.name = 'Aptos'
        doc.styles['Heading 2'].font.size = Pt(12)
        doc.styles['Heading 2'].font.color.rgb = RGBColor(30, 41, 59)

        # Title
        title = doc.add_heading('Docker Security Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(
            'Analyst-ready export covering native Docker Inspector findings, CVEs, Trivy, and AI narrative.'
        )
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(9)
        subtitle_run.font.color.rgb = RGBColor(71, 85, 105)

        metadata_table = doc.add_table(rows=6, cols=2)
        metadata_table.style = 'Table Grid'
        metadata_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        metadata_rows = [
            ('Image', request.image_name),
            ('Image ID', request.image_id[:12]),
            ('Base Image', request.base_image or 'Unknown'),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Risk Rating', f'{risk_level.upper()} ({risk_score}/100)'),
            ('Image Size / Layers', f'{request.total_size_human} / {request.total_layers}'),
        ]
        for index, (key, value) in enumerate(metadata_rows):
            metadata_table.rows[index].cells[0].text = key
            metadata_table.rows[index].cells[1].text = _sanitize_word_text(value)
            metadata_table.rows[index].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_heading('At A Glance', level=1)
        for finding in overview["key_findings"]:
            doc.add_paragraph(_sanitize_word_text(finding), style='List Bullet')

        doc.add_heading('Priority Actions', level=1)
        for action in overview["priority_actions"]:
            doc.add_paragraph(_sanitize_word_text(action), style='List Number')

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary_table = doc.add_table(rows=9, cols=2)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        summary_data = [
            ('Metric', 'Value'),
            ('Risk Score', f'{risk_score}/100'),
            ('Total Layers', str(request.total_layers)),
            ('Security Issues', str(len(request.security_issues))),
            ('Secrets Found', str(overview["total_secrets"])),
            ('CVEs Found', str(overview["total_cves"])),
            ('Native Trivy Findings', str(total_trivy_findings)),
            ('Deleted but Recoverable', str(overview["deleted_recoverable"])),
            ('Risk Rating', risk_level.upper()),
        ]
        for i, (key, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = key
            summary_table.rows[i].cells[1].text = value
            if summary_table.rows[i].cells[0].paragraphs[0].runs:
                summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            if i == 0:
                for cell in summary_table.rows[i].cells:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True

        risk_table = doc.add_table(rows=5, cols=3)
        risk_table.style = 'Table Grid'
        risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        risk_rows = [
            ('Severity', 'Count', 'Meaning'),
            ('Critical', str(severity_counts['critical']), 'Immediate exploitation or direct compromise'),
            ('High', str(severity_counts['high']), 'Material attacker advantage'),
            ('Medium', str(severity_counts['medium']), 'Important but not immediately critical'),
            ('Low', str(severity_counts['low']), 'Hardening and hygiene gaps'),
        ]
        for row_index, row_data in enumerate(risk_rows):
            for col_index, value in enumerate(row_data):
                risk_table.rows[row_index].cells[col_index].text = value
                if row_index == 0 and risk_table.rows[row_index].cells[col_index].paragraphs[0].runs:
                    risk_table.rows[row_index].cells[col_index].paragraphs[0].runs[0].bold = True

        if overview["coverage_rows"]:
            doc.add_heading('Scan Coverage', level=1)
            coverage_table = doc.add_table(rows=1, cols=4)
            coverage_table.style = 'Table Grid'
            coverage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = coverage_table.rows[0].cells
            headers[0].text = 'Check'
            headers[1].text = 'Status'
            headers[2].text = 'Duration'
            headers[3].text = 'Notes'
            for cell in headers:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
            for row in overview["coverage_rows"]:
                cells = coverage_table.add_row().cells
                cells[0].text = _sanitize_word_text(row.get('label', 'Check'))
                cells[1].text = _sanitize_word_text(row.get('status', 'Unknown'))
                cells[2].text = _sanitize_word_text(row.get('duration', 'N/A'))
                cells[3].text = _sanitize_word_text((row.get('detail') or '').strip() or '-')

        # Base Image Intelligence
        if request.base_image_intel:
            doc.add_heading('Base Image Intelligence', level=1)
            for intel in request.base_image_intel:
                p = doc.add_paragraph()
                run = p.add_run(f"{intel.get('severity', 'N/A').upper()}: ")
                run.bold = True
                if intel.get('severity', '').lower() == 'critical':
                    run.font.color.rgb = RGBColor(220, 53, 69)
                p.add_run(intel.get('message', 'N/A'))
                doc.add_paragraph(f"Recommendation: {intel.get('recommendation', 'N/A')}", style='List Bullet')

        # Security Issues
        if request.security_issues:
            doc.add_heading('Security Issues', level=1)
            issues_table = doc.add_table(rows=1, cols=3)
            issues_table.style = 'Table Grid'
            hdr_cells = issues_table.rows[0].cells
            hdr_cells[0].text = 'Severity'
            hdr_cells[1].text = 'Category'
            hdr_cells[2].text = 'Description'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True

            for issue in request.security_issues[:20]:
                row_cells = issues_table.add_row().cells
                row_cells[0].text = issue.get('severity', 'N/A').upper()
                row_cells[1].text = issue.get('category', 'N/A')
                row_cells[2].text = issue.get('description', 'N/A')[:80]

        if request.cve_scan:
            doc.add_heading('Package CVE Scan', level=1)
            doc.add_paragraph(f"Packages scanned: {request.cve_scan.get('packages_scanned', 0)}")
            doc.add_paragraph(f"Packages with vulnerabilities: {request.cve_scan.get('packages_with_vulns', 0)}")
            doc.add_paragraph(
                f"Total CVEs: {request.cve_scan.get('total_vulnerabilities') or len(request.cve_scan.get('vulnerabilities', []))}"
            )
            doc.add_paragraph(
                f"Critical: {request.cve_scan.get('critical_count', 0)} | High: {request.cve_scan.get('high_count', 0)} | KEV: {request.cve_scan.get('kev_count', 0)}"
            )

        if trivy_scan:
            doc.add_heading('Native Trivy Scan', level=1)
            enabled_scanners = ", ".join(trivy_scan.get("enabled_scanners") or []) or "None"
            trivy_duration = trivy_scan.get("scan_duration_seconds") or 0
            try:
                trivy_duration = float(trivy_duration)
            except (TypeError, ValueError):
                trivy_duration = 0.0

            doc.add_paragraph(f"Enabled scanners: {enabled_scanners}")
            doc.add_paragraph(f"Packages observed: {trivy_scan.get('package_count', 0)}")
            doc.add_paragraph(
                f"Native vulnerabilities: {len(trivy_vulns)} | Misconfigurations: {len(trivy_misconfigs)} | "
                f"Secret findings: {len(trivy_secrets)} | License findings: {len(trivy_licenses)}"
            )
            doc.add_paragraph(f"Duration: {trivy_duration:.1f}s")
            if trivy_scan.get("error"):
                warn_p = doc.add_paragraph()
                warn_run = warn_p.add_run(f"Scan Error: {_sanitize_word_text(trivy_scan.get('error'))}")
                warn_run.bold = True
                warn_run.font.color.rgb = RGBColor(220, 53, 69)

            if trivy_vulns:
                doc.add_heading('Trivy Vulnerabilities', level=2)
                vuln_table = doc.add_table(rows=1, cols=5)
                vuln_table.style = 'Table Grid'
                hdr = vuln_table.rows[0].cells
                hdr[0].text = 'ID'
                hdr[1].text = 'Severity'
                hdr[2].text = 'Package'
                hdr[3].text = 'Installed'
                hdr[4].text = 'Fixed'
                for cell in hdr:
                    cell.paragraphs[0].runs[0].bold = True
                for finding in trivy_vulns[:15]:
                    row = vuln_table.add_row().cells
                    row[0].text = _sanitize_word_text(str(finding.get('vulnerability_id', 'N/A'))[:28])
                    row[1].text = _sanitize_word_text(str(finding.get('severity', 'N/A')).upper())
                    row[2].text = _sanitize_word_text(str(finding.get('package_name', 'N/A'))[:24])
                    row[3].text = _sanitize_word_text(str(finding.get('installed_version', 'N/A'))[:24])
                    row[4].text = _sanitize_word_text(str(finding.get('fixed_version') or '-')[:24])

            if trivy_misconfigs:
                doc.add_heading('Trivy Misconfigurations', level=2)
                misconfig_table = doc.add_table(rows=1, cols=4)
                misconfig_table.style = 'Table Grid'
                hdr = misconfig_table.rows[0].cells
                hdr[0].text = 'ID'
                hdr[1].text = 'Severity'
                hdr[2].text = 'Target'
                hdr[3].text = 'Message'
                for cell in hdr:
                    cell.paragraphs[0].runs[0].bold = True
                for finding in trivy_misconfigs[:15]:
                    row = misconfig_table.add_row().cells
                    row[0].text = _sanitize_word_text(str(finding.get('id', 'N/A'))[:28])
                    row[1].text = _sanitize_word_text(str(finding.get('severity', 'N/A')).upper())
                    row[2].text = _sanitize_word_text(str(finding.get('target', 'N/A'))[:32])
                    row[3].text = _sanitize_word_text(str(finding.get('message') or finding.get('title') or 'N/A')[:120])

            if trivy_secrets:
                doc.add_heading('Trivy Secret Findings', level=2)
                doc.add_paragraph('Operator note: raw Trivy secret matches are intentionally included for remediation workflows.')
                for finding in trivy_secrets[:15]:
                    doc.add_heading(
                        f"{_sanitize_word_text(finding.get('rule_id', 'Unknown'))} "
                        f"({_sanitize_word_text(str(finding.get('severity', 'N/A')).upper())})",
                        level=3,
                    )
                    doc.add_paragraph(f"Target: {_sanitize_word_text(finding.get('target', 'N/A'))}")
                    doc.add_paragraph(f"Match: {_sanitize_word_text(finding.get('match', 'N/A'))}")

            if trivy_licenses:
                doc.add_heading('Trivy License Findings', level=2)
                license_table = doc.add_table(rows=1, cols=4)
                license_table.style = 'Table Grid'
                hdr = license_table.rows[0].cells
                hdr[0].text = 'License'
                hdr[1].text = 'Severity'
                hdr[2].text = 'Package'
                hdr[3].text = 'Category'
                for cell in hdr:
                    cell.paragraphs[0].runs[0].bold = True
                for finding in trivy_licenses[:15]:
                    row = license_table.add_row().cells
                    row[0].text = _sanitize_word_text(str(finding.get('name', 'N/A'))[:24])
                    row[1].text = _sanitize_word_text(str(finding.get('severity', 'N/A')).upper())
                    row[2].text = _sanitize_word_text(str(finding.get('package_name', 'N/A'))[:24])
                    row[3].text = _sanitize_word_text(str(finding.get('category', 'N/A'))[:24])

        # Layer Secrets
        if request.layer_secrets:
            doc.add_heading('Layer Deep Scan - Recoverable Secrets', level=1)
            deleted_count = sum(1 for s in request.layer_secrets if s.get('is_deleted'))
            if deleted_count > 0:
                warn_p = doc.add_paragraph()
                warn_run = warn_p.add_run(f"WARNING: {deleted_count} files were deleted but are still recoverable!")
                warn_run.bold = True
                warn_run.font.color.rgb = RGBColor(220, 53, 69)

            secrets_table = doc.add_table(rows=1, cols=4)
            secrets_table.style = 'Table Grid'
            hdr = secrets_table.rows[0].cells
            hdr[0].text = 'Layer'
            hdr[1].text = 'File Path'
            hdr[2].text = 'Type'
            hdr[3].text = 'Deleted?'
            for cell in hdr:
                cell.paragraphs[0].runs[0].bold = True

            for secret in request.layer_secrets[:20]:
                row = secrets_table.add_row().cells
                row[0].text = str(secret.get('layer_index', '?'))
                row[1].text = secret.get('file_path', 'N/A')[:40]
                row[2].text = secret.get('file_type', 'N/A')
                row[3].text = 'YES' if secret.get('is_deleted') else 'No'

        if request.secrets:
            doc.add_heading('Secrets in Image History', level=1)
            doc.add_paragraph('Operator note: raw secret values are intentionally included in this Docker report for remediation workflows.')
            for secret in request.secrets[:20]:
                doc.add_heading(
                    f"{_sanitize_word_text(secret.get('secret_type', 'Unknown'))} "
                    f"({_sanitize_word_text(str(secret.get('severity', 'N/A')).upper())})",
                    level=2,
                )
                doc.add_paragraph(f"Layer: {_sanitize_word_text(secret.get('layer_id', 'N/A')[:12])}")
                doc.add_paragraph(f"Raw Value: {_sanitize_word_text(secret.get('value', 'N/A'))}")
                doc.add_paragraph(f"Context: {_sanitize_word_text(secret.get('context', 'N/A'))}")
                doc.add_paragraph(f"Layer Command: {_sanitize_word_text(secret.get('layer_command', 'N/A'))}")

        # AI Analysis
        if overview["ai_text"]:
            doc.add_heading('AI Security Analysis', level=1)
            intro = doc.add_paragraph()
            intro_run = intro.add_run('Narrative summary from the analyst-facing AI layer, normalized for export formatting.')
            intro_run.italic = True
            intro_run.font.color.rgb = RGBColor(71, 85, 105)
            for block in overview["ai_text"].split("\n\n"):
                clean_block = _sanitize_word_text(block.strip())
                if not clean_block:
                    continue
                if clean_block.startswith("• "):
                    for line in clean_block.splitlines():
                        if line.strip().startswith("• "):
                            doc.add_paragraph(line.strip()[2:], style='List Bullet')
                else:
                    paragraph = doc.add_paragraph()
                    _add_formatted_text(paragraph, clean_block)

        # Footer
        doc.add_heading('Report Information', level=2)
        doc.add_paragraph('Generated by VRAgent Docker Inspector', style='List Bullet')
        doc.add_paragraph('Features: Image history analysis, regex secret detection, AI adjudication, layer deep scan, base image intelligence, package CVE scanning, native Trivy scanners, SBOM export', style='List Bullet')
        doc.add_paragraph('Document structure: report header, at-a-glance findings, priority actions, scan coverage, technical sections, AI narrative', style='List Bullet')

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError as e:
        logger.warning(f"DOCX generation requires python-docx: {e}")
        return md_content.encode('utf-8')


class DockerImageAllowRequest(BaseModel):
    """Request to allow a local Docker image for analysis."""
    image_name: str
    project_id: Optional[int] = None


@router.get("/docker-images")
async def list_local_docker_images(
    project_id: Optional[int] = Query(None, description="Optional project scope for shared Docker images"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    List Docker images for the Docker Inspector.

    Non-admin users only see allowlisted images.
    Administrators see all local images with allowlist status so they can curate the list.
    """
    if not check_docker_available():
        raise HTTPException(
            status_code=503,
            detail="Docker is not available."
        )

    try:
        allowlist_scope_map = _get_docker_allowlist_scope_map(db, current_user.id, project_id)
        local_images = _list_local_docker_images_raw()
        if current_user.role == "admin":
            images = [
                {
                    **image,
                    "allowlisted": image["name"] in allowlist_scope_map,
                    "allowlist_scope": allowlist_scope_map.get(image["name"]),
                }
                for image in local_images
            ]
            allowlisted_total = sum(1 for image in images if image["allowlisted"])
            return {
                "images": images,
                "total": len(images),
                "allowlisted_total": allowlisted_total,
                "showing_all_local": True,
            }

        images = [
            {
                **image,
                "allowlisted": True,
                "allowlist_scope": allowlist_scope_map.get(image["name"]),
            }
            for image in local_images
            if image["name"] in allowlist_scope_map
        ]
        return {
            "images": images,
            "total": len(images),
            "allowlisted_total": len(images),
            "showing_all_local": False,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to list Docker images: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/docker-images/allow")
async def allow_docker_image(
    request: DockerImageAllowRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    Allow a local Docker image to appear in the Docker Inspector and be analyzed.
    """
    if not check_docker_available():
        raise HTTPException(status_code=503, detail="Docker is not available.")

    if current_user.role != "admin":
        raise HTTPException(
            status_code=403,
            detail="Only administrators can allow Docker images on shared infrastructure.",
        )

    image_name = _normalize_docker_image_name(request.image_name)
    if not _docker_image_exists(image_name):
        raise HTTPException(
            status_code=404,
            detail="Docker image was not found locally. Pull the image before allowlisting it.",
        )

    try:
        if request.project_id is not None:
            _ensure_project_access(db, request.project_id, current_user.id, require_edit=True)
            existing = (
                db.query(ProjectDockerImage)
                .filter(
                    ProjectDockerImage.project_id == request.project_id,
                    ProjectDockerImage.image_name == image_name,
                )
                .first()
            )
            if not existing:
                db.add(
                    ProjectDockerImage(
                        project_id=request.project_id,
                        image_name=image_name,
                        added_by=current_user.id,
                    )
                )
                db.commit()
            return {
                "image_name": image_name,
                "project_id": request.project_id,
                "scope": "project",
                "message": "Docker image allowlisted for the project",
            }

        existing = (
            db.query(UserDockerImage)
            .filter(
                UserDockerImage.user_id == current_user.id,
                UserDockerImage.image_name == image_name,
            )
            .first()
        )
        if not existing:
            db.add(UserDockerImage(user_id=current_user.id, image_name=image_name))
            db.commit()
        return {
            "image_name": image_name,
            "project_id": None,
            "scope": "user",
            "message": "Docker image allowlisted for the user",
        }
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Failed to allow Docker image {image_name}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to allow Docker image: {str(e)}")


# ============================================================================
# Report Management Endpoints
# ============================================================================

class SaveReportRequest(BaseModel):
    """Request to save a reverse engineering report."""
    analysis_type: str  # 'binary', 'apk', 'docker'
    title: str
    filename: Optional[str] = None
    project_id: Optional[int] = None
    
    # Risk assessment
    risk_level: Optional[str] = None
    risk_score: Optional[int] = None
    
    # Type-specific fields
    file_type: Optional[str] = None
    architecture: Optional[str] = None
    file_size: Optional[int] = None
    is_packed: Optional[bool] = None
    packer_name: Optional[str] = None
    
    package_name: Optional[str] = None
    version_name: Optional[str] = None
    min_sdk: Optional[int] = None
    target_sdk: Optional[int] = None
    
    image_name: Optional[str] = None
    image_id: Optional[str] = None
    total_layers: Optional[int] = None
    base_image: Optional[str] = None
    
    # Counts
    strings_count: Optional[int] = None
    imports_count: Optional[int] = None
    exports_count: Optional[int] = None
    secrets_count: Optional[int] = None
    
    # JSON data
    suspicious_indicators: Optional[List[Dict[str, Any]]] = None
    permissions: Optional[List[Dict[str, Any]]] = None
    security_issues: Optional[List[Dict[str, Any]]] = None
    full_analysis_data: Optional[Dict[str, Any]] = None
    
    # AI Analysis
    ai_analysis_raw: Optional[str] = None
    
    # JADX Full Scan Data
    jadx_total_classes: Optional[int] = None
    jadx_total_files: Optional[int] = None
    jadx_output_directory: Optional[str] = None
    jadx_classes_sample: Optional[List[Dict[str, Any]]] = None
    jadx_security_issues: Optional[List[Dict[str, Any]]] = None
    jadx_source_tree: Optional[Dict[str, Any]] = None  # Directory structure for source browser
    jadx_source_code_samples: Optional[List[Dict[str, Any]]] = None  # Actual source code for key classes
    
    # AI-Generated Reports (Deep Analysis)
    ai_functionality_report: Optional[str] = None
    ai_security_report: Optional[str] = None
    ai_privacy_report: Optional[str] = None
    ai_architecture_diagram: Optional[str] = None
    ai_attack_surface_map: Optional[str] = None  # Mermaid attack tree diagram
    ai_threat_model: Optional[Dict[str, Any]] = None
    ai_vuln_scan_result: Optional[Dict[str, Any]] = None
    ai_chat_history: Optional[List[Dict[str, Any]]] = None
    
    # Library CVE Analysis
    detected_libraries: Optional[List[Dict[str, Any]]] = None  # Libraries detected in APK
    library_cves: Optional[List[Dict[str, Any]]] = None  # CVEs found in libraries
    
    # Dynamic Analysis / Frida Scripts
    dynamic_analysis: Optional[Dict[str, Any]] = None  # Full Frida scripts data
    
    # Decompiled Code Analysis Results (Pattern-based scanners)
    decompiled_code_findings: Optional[List[Dict[str, Any]]] = None  # Security findings from decompiled code
    decompiled_code_summary: Optional[Dict[str, Any]] = None  # Summary by severity/scanner/category
    
    # CVE Scan Results
    cve_scan_results: Optional[Dict[str, Any]] = None  # CVE database lookup results
    
    # Vulnerability-specific Frida Hooks
    vulnerability_frida_hooks: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None  # Auto-generated hooks for discovered vulns
    
    # Unified APK vulnerability hunt results
    vuln_hunt_result: Optional[Dict[str, Any]] = None
    
    # Manifest Visualization (component graph, deep links, AI analysis)
    manifest_visualization: Optional[Dict[str, Any]] = None
    
    # Obfuscation Analysis (detection, deobfuscation strategies, Frida hooks)
    obfuscation_analysis: Optional[Dict[str, Any]] = None
    
    # AI Finding Verification Results (confidence scores, attack chains, FP filtering)
    verification_results: Optional[Dict[str, Any]] = None
    
    # Sensitive Data Discovery (AI-verified passwords, API keys, emails, phone numbers, PII)
    sensitive_data_findings: Optional[Dict[str, Any]] = None
    
    # Tags and notes
    tags: Optional[List[str]] = None
    notes: Optional[str] = None


class ReportSummaryResponse(BaseModel):
    """Summary of a saved report."""
    id: int
    analysis_type: str
    title: str
    filename: Optional[str] = None
    risk_level: Optional[str] = None
    risk_score: Optional[int] = None
    created_at: datetime
    tags: Optional[List[str]] = None


class ReportDetailResponse(BaseModel):
    """Full report detail response."""
    id: int
    analysis_type: str
    title: str
    filename: Optional[str] = None
    created_at: datetime
    updated_at: datetime
    project_id: Optional[int] = None
    
    risk_level: Optional[str] = None
    risk_score: Optional[int] = None
    
    file_type: Optional[str] = None
    architecture: Optional[str] = None
    file_size: Optional[int] = None
    is_packed: Optional[str] = None
    packer_name: Optional[str] = None
    
    package_name: Optional[str] = None
    version_name: Optional[str] = None
    min_sdk: Optional[int] = None
    target_sdk: Optional[int] = None
    
    image_name: Optional[str] = None
    image_id: Optional[str] = None
    total_layers: Optional[int] = None
    base_image: Optional[str] = None
    
    strings_count: Optional[int] = None
    imports_count: Optional[int] = None
    exports_count: Optional[int] = None
    secrets_count: Optional[int] = None
    
    suspicious_indicators: Optional[List[Dict[str, Any]]] = None
    permissions: Optional[List[Dict[str, Any]]] = None
    security_issues: Optional[List[Dict[str, Any]]] = None
    full_analysis_data: Optional[Dict[str, Any]] = None
    
    ai_analysis_raw: Optional[str] = None
    ai_analysis_structured: Optional[Dict[str, Any]] = None
    
    # JADX Full Scan Data
    jadx_total_classes: Optional[int] = None
    jadx_total_files: Optional[int] = None
    jadx_output_directory: Optional[str] = None
    jadx_classes_sample: Optional[List[Dict[str, Any]]] = None
    jadx_security_issues: Optional[List[Dict[str, Any]]] = None
    jadx_source_tree: Optional[Dict[str, Any]] = None  # Directory structure for source browser
    jadx_source_code_samples: Optional[List[Dict[str, Any]]] = None  # Actual source code for key classes
    jadx_data: Optional[Dict[str, Any]] = None
    
    # AI-Generated Reports
    ai_functionality_report: Optional[str] = None
    ai_security_report: Optional[str] = None
    ai_privacy_report: Optional[str] = None
    ai_architecture_diagram: Optional[str] = None
    ai_attack_surface_map: Optional[str] = None  # Mermaid attack tree diagram
    ai_threat_model: Optional[Dict[str, Any]] = None
    ai_vuln_scan_result: Optional[Dict[str, Any]] = None
    ai_chat_history: Optional[List[Dict[str, Any]]] = None
    
    # Library CVE Analysis
    detected_libraries: Optional[List[Dict[str, Any]]] = None
    library_cves: Optional[List[Dict[str, Any]]] = None
    
    # Dynamic Analysis / Frida Scripts
    dynamic_analysis: Optional[Dict[str, Any]] = None  # Full Frida scripts data
    
    # Decompiled Code Analysis Results (Pattern-based scanners)
    decompiled_code_findings: Optional[List[Dict[str, Any]]] = None  # Security findings from decompiled code
    decompiled_code_summary: Optional[Dict[str, Any]] = None  # Summary by severity/scanner/category
    
    # CVE Scan Results
    cve_scan_results: Optional[Dict[str, Any]] = None  # CVE database lookup results
    
    # Vulnerability-specific Frida Hooks
    vulnerability_frida_hooks: Optional[Dict[str, Any]] = None  # Auto-generated hooks for discovered vulns
    
    # Unified APK vulnerability hunt results
    vuln_hunt_result: Optional[Dict[str, Any]] = None
    
    # Manifest Visualization (component graph, deep links, AI analysis)
    manifest_visualization: Optional[Dict[str, Any]] = None
    
    # Obfuscation Analysis (detection, deobfuscation strategies, Frida hooks)
    obfuscation_analysis: Optional[Dict[str, Any]] = None
    
    # AI Finding Verification Results (confidence scores, attack chains, FP filtering)
    verification_results: Optional[Dict[str, Any]] = None
    
    # Sensitive Data Discovery (AI-verified passwords, API keys, emails, phone numbers, PII)
    sensitive_data_findings: Optional[Dict[str, Any]] = None
    
    tags: Optional[List[str]] = None
    notes: Optional[str] = None


@router.post("/reports", response_model=ReportSummaryResponse)
def save_report(
    request: SaveReportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    Save a reverse engineering analysis report.
    
    Stores the full analysis data for later review, comparison, or export.
    """
    try:
        def _normalize_vulnerability_hooks_payload(value: Any) -> Optional[Dict[str, Any]]:
            if isinstance(value, dict):
                return value
            if isinstance(value, list):
                return {"vulnerability_scripts": value}
            return None

        if request.project_id is not None:
            _ensure_project_access(db, request.project_id, current_user.id, require_edit=True)

        # Parse risk level from AI analysis if not provided
        risk_level = request.risk_level
        risk_score = request.risk_score
        
        if not risk_level and request.ai_analysis_raw:
            # Try to extract risk level from AI analysis
            ai_text = request.ai_analysis_raw.lower()
            if 'critical' in ai_text[:500]:
                risk_level = 'Critical'
                risk_score = risk_score or 90
            elif 'high' in ai_text[:500]:
                risk_level = 'High'
                risk_score = risk_score or 70
            elif 'medium' in ai_text[:500]:
                risk_level = 'Medium'
                risk_score = risk_score or 50
            elif 'low' in ai_text[:500]:
                risk_level = 'Low'
                risk_score = risk_score or 30
            elif 'clean' in ai_text[:500]:
                risk_level = 'Clean'
                risk_score = risk_score or 10

        normalized_full_analysis_data = dict(request.full_analysis_data or {})
        normalized_vulnerability_hooks = _normalize_vulnerability_hooks_payload(request.vulnerability_frida_hooks)

        if normalized_vulnerability_hooks and "vulnerability_frida_hooks" not in normalized_full_analysis_data:
            normalized_full_analysis_data["vulnerability_frida_hooks"] = normalized_vulnerability_hooks
        if request.vuln_hunt_result and "vuln_hunt_result" not in normalized_full_analysis_data:
            normalized_full_analysis_data["vuln_hunt_result"] = request.vuln_hunt_result
        if request.manifest_visualization and "manifest_visualization" not in normalized_full_analysis_data:
            normalized_full_analysis_data["manifest_visualization"] = request.manifest_visualization
        if request.obfuscation_analysis and "obfuscation_analysis" not in normalized_full_analysis_data:
            normalized_full_analysis_data["obfuscation_analysis"] = request.obfuscation_analysis
        if request.verification_results and "verification_results" not in normalized_full_analysis_data:
            normalized_full_analysis_data["verification_results"] = request.verification_results
        if request.sensitive_data_findings and "sensitive_data_findings" not in normalized_full_analysis_data:
            normalized_full_analysis_data["sensitive_data_findings"] = request.sensitive_data_findings
        
        report = ReverseEngineeringReport(
            analysis_type=request.analysis_type,
            title=request.title,
            filename=request.filename,
            project_id=request.project_id,
            user_id=current_user.id,
            
            risk_level=risk_level,
            risk_score=risk_score,
            
            file_type=request.file_type,
            architecture=request.architecture,
            file_size=request.file_size,
            is_packed=str(request.is_packed) if request.is_packed is not None else None,
            packer_name=request.packer_name,
            
            package_name=request.package_name,
            version_name=request.version_name,
            min_sdk=request.min_sdk,
            target_sdk=request.target_sdk,
            
            image_name=request.image_name,
            image_id=request.image_id,
            total_layers=request.total_layers,
            base_image=request.base_image,
            
            strings_count=request.strings_count,
            imports_count=request.imports_count,
            exports_count=request.exports_count,
            secrets_count=request.secrets_count,
            
            suspicious_indicators=request.suspicious_indicators,
            permissions=request.permissions,
            security_issues=request.security_issues,
            full_analysis_data=normalized_full_analysis_data or request.full_analysis_data,
            
            ai_analysis_raw=request.ai_analysis_raw,
            
            # JADX Full Scan Data
            jadx_total_classes=request.jadx_total_classes,
            jadx_total_files=request.jadx_total_files,
            jadx_data={
                "output_directory": request.jadx_output_directory,
                "classes_sample": request.jadx_classes_sample,
                "security_issues": request.jadx_security_issues,
                "source_tree": request.jadx_source_tree,
                "source_code_samples": request.jadx_source_code_samples,
            } if request.jadx_total_classes else None,
            
            # AI-Generated Reports
            ai_functionality_report=request.ai_functionality_report,
            ai_security_report=request.ai_security_report,
            ai_privacy_report=request.ai_privacy_report,
            ai_architecture_diagram=request.ai_architecture_diagram,
            ai_attack_surface_map=request.ai_attack_surface_map,
            ai_threat_model=request.ai_threat_model,
            ai_vuln_scan_result=request.ai_vuln_scan_result,
            ai_chat_history=request.ai_chat_history,
            
            # Library CVE Analysis
            detected_libraries=request.detected_libraries,
            library_cves=request.library_cves,
            
            # Dynamic Analysis / Frida Scripts
            dynamic_analysis=request.dynamic_analysis,
            
            # Decompiled Code Analysis Results
            decompiled_code_findings=request.decompiled_code_findings,
            decompiled_code_summary=request.decompiled_code_summary,
            
            # CVE Scan Results
            cve_scan_results=request.cve_scan_results,
            
            # Vulnerability-specific Frida Hooks
            vulnerability_frida_hooks=normalized_vulnerability_hooks,
            
            # Manifest Visualization (component graph, deep links, AI analysis)
            manifest_visualization=request.manifest_visualization,
            
            # Obfuscation Analysis (detection, deobfuscation strategies, Frida hooks)
            obfuscation_analysis=request.obfuscation_analysis,
            
            # AI Finding Verification Results (confidence scores, attack chains, FP filtering)
            verification_results=request.verification_results,
            
            # Sensitive Data Discovery (AI-verified passwords, API keys, emails, phone numbers, PII)
            sensitive_data_findings=request.sensitive_data_findings,
            
            tags=request.tags,
            notes=request.notes,
        )
        
        db.add(report)
        db.commit()
        db.refresh(report)
        
        logger.info(f"Saved RE report {report.id}: {request.title}")
        
        return ReportSummaryResponse(
            id=report.id,
            analysis_type=report.analysis_type,
            title=report.title,
            filename=report.filename,
            risk_level=report.risk_level,
            risk_score=report.risk_score,
            created_at=report.created_at,
            tags=report.tags,
        )
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Failed to save report: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save report: {str(e)}")


@router.get("/reports", response_model=List[ReportSummaryResponse])
def list_reports(
    analysis_type: Optional[str] = Query(None, description="Filter by analysis type (binary, apk, docker)"),
    project_id: Optional[int] = Query(None, description="Filter by project ID"),
    include_project_reports: bool = Query(
        False,
        description="When no project_id is provided, include project-scoped reports as well.",
    ),
    risk_level: Optional[str] = Query(None, description="Filter by risk level"),
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    List saved reverse engineering reports.
    
    Returns summaries of saved reports with optional filtering.
    """
    query = db.query(ReverseEngineeringReport)

    if analysis_type:
        query = query.filter(ReverseEngineeringReport.analysis_type == analysis_type)
    if project_id is not None:
        _ensure_project_access(db, project_id, current_user.id)
        query = query.filter(ReverseEngineeringReport.project_id == project_id)
    elif not include_project_reports:
        # Standalone Hub view: only show reports that are not attached to a project.
        query = query.filter(
            ReverseEngineeringReport.project_id.is_(None),
            ReverseEngineeringReport.user_id == current_user.id,
        )
    else:
        accessible_project_ids = _get_accessible_project_ids(db, current_user.id)
        query = query.filter(
            or_(
                and_(
                    ReverseEngineeringReport.project_id.is_(None),
                    ReverseEngineeringReport.user_id == current_user.id,
                ),
                ReverseEngineeringReport.project_id.in_(accessible_project_ids or [-1]),
            )
        )
    if risk_level:
        query = query.filter(ReverseEngineeringReport.risk_level == risk_level)
    
    reports = query.order_by(ReverseEngineeringReport.created_at.desc()).offset(offset).limit(limit).all()
    
    return [
        ReportSummaryResponse(
            id=r.id,
            analysis_type=r.analysis_type,
            title=r.title,
            filename=r.filename,
            risk_level=r.risk_level,
            risk_score=r.risk_score,
            created_at=r.created_at,
            tags=r.tags,
        )
        for r in reports
    ]


@router.get("/reports/{report_id}", response_model=ReportDetailResponse)
def get_report(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    Get full details of a saved reverse engineering report.
    """
    report = _ensure_reverse_engineering_report_access(
        db,
        _get_reverse_engineering_report_or_404(db, report_id),
        current_user.id,
    )
    
    def _coerce_dict(value: Any) -> Dict[str, Any]:
        if isinstance(value, dict):
            return value
        if isinstance(value, str):
            try:
                parsed = json.loads(value)
                if isinstance(parsed, dict):
                    return parsed
            except Exception:
                return {}
        return {}

    def _collect_dicts(root: Any, max_depth: int = 5) -> List[Dict[str, Any]]:
        collected: List[Dict[str, Any]] = []
        seen: set[int] = set()

        def _walk(node: Any, depth: int) -> None:
            if depth > max_depth or node is None:
                return
            if isinstance(node, dict):
                node_id = id(node)
                if node_id in seen:
                    return
                seen.add(node_id)
                collected.append(node)
                for value in node.values():
                    _walk(value, depth + 1)
                return
            if isinstance(node, list):
                for item in node:
                    _walk(item, depth + 1)

        _walk(root, 0)
        return collected

    def _clean_text(value: Any) -> Optional[str]:
        if not isinstance(value, str):
            return None
        trimmed = value.strip()
        if not trimmed:
            return None
        if trimmed.lower() in {"not available", "n/a", "none", "null", "undefined"}:
            return None
        return trimmed

    def _pick_text(dicts: List[Dict[str, Any]], keys: List[str]) -> Optional[str]:
        for data in dicts:
            for key in keys:
                cleaned = _clean_text(data.get(key))
                if cleaned:
                    return cleaned
        return None

    def _pick_text_by_hint(
        dicts: List[Dict[str, Any]],
        hints: List[str],
        minimum_length: int = 120,
    ) -> Optional[str]:
        for data in dicts:
            for key, raw_value in data.items():
                lowered_key = str(key).lower()
                if not any(hint in lowered_key for hint in hints):
                    continue
                cleaned = _clean_text(raw_value)
                if cleaned and len(cleaned) >= minimum_length:
                    return cleaned
        return None

    def _is_mermaid(text: Optional[str]) -> bool:
        if not isinstance(text, str):
            return False
        trimmed = text.strip()
        if not trimmed:
            return False
        lowered = trimmed.lower()
        if "```mermaid" in lowered:
            return True
        return lowered.startswith((
            "flowchart",
            "graph ",
            "sequencediagram",
            "classdiagram",
            "statediagram",
            "erdiagram",
            "journey",
            "gantt",
        ))

    def _pick_mermaid(dicts: List[Dict[str, Any]], keys: List[str]) -> Optional[str]:
        direct = _pick_text(dicts, keys)
        if _is_mermaid(direct):
            return direct
        for data in dicts:
            for value in data.values():
                if isinstance(value, str) and _is_mermaid(value):
                    return value.strip()
        return None

    full_data_dict = _coerce_dict(report.full_analysis_data)
    structured_dict = _coerce_dict(report.ai_analysis_structured)
    ai_raw = _clean_text(report.ai_analysis_raw)
    ai_raw_dict = _coerce_dict(ai_raw) if ai_raw else {}
    search_space = (
        _collect_dicts(full_data_dict)
        + _collect_dicts(structured_dict)
        + _collect_dicts(ai_raw_dict, 4)
    )

    functionality_report = _clean_text(report.ai_functionality_report) or _pick_text(search_space, [
        "ai_functionality_report",
        "ai_report_functionality",
        "functionality_report",
        "functionality",
        "report_functionality",
        "binary_functionality_report",
    ]) or _pick_text_by_hint(search_space, ["functionality", "purpose", "behavior", "what_does"])
    security_report = _clean_text(report.ai_security_report) or _pick_text(search_space, [
        "ai_security_report",
        "ai_report_security",
        "security_report",
        "security",
        "report_security",
        "binary_security_report",
    ]) or _pick_text_by_hint(search_space, ["security", "risk", "vulnerab", "threat"])
    architecture_diagram = _clean_text(report.ai_architecture_diagram) or _pick_mermaid(search_space, [
        "ai_architecture_diagram",
        "architecture_diagram",
        "ai_architecture",
        "architecture",
        "architecture_map",
        "system_architecture_diagram",
    ])
    attack_surface_map = _clean_text(report.ai_attack_surface_map) or _pick_mermaid(search_space, [
        "ai_attack_surface_map",
        "attack_surface_map",
        "attack_surface",
        "ai_attack_surface",
        "attack_tree",
        "attack_tree_diagram",
    ])

    if not functionality_report and ai_raw:
        functionality_report = ai_raw
    if not security_report and ai_raw:
        security_report = ai_raw

    def _normalize_vulnerability_hooks_payload(value: Any) -> Optional[Dict[str, Any]]:
        if isinstance(value, dict):
            return value
        if isinstance(value, list):
            return {"vulnerability_scripts": value}
        return None

    normalized_vulnerability_hooks = _normalize_vulnerability_hooks_payload(
        getattr(report, "vulnerability_frida_hooks", None)
    ) or _normalize_vulnerability_hooks_payload(full_data_dict.get("vulnerability_frida_hooks"))
    vuln_hunt_result = full_data_dict.get("vuln_hunt_result")
    if not isinstance(vuln_hunt_result, dict):
        vuln_hunt_result = None

    normalized_full_data = dict(full_data_dict) if full_data_dict else {}
    if functionality_report and not normalized_full_data.get("ai_functionality_report"):
        normalized_full_data["ai_functionality_report"] = functionality_report
    if security_report and not normalized_full_data.get("ai_security_report"):
        normalized_full_data["ai_security_report"] = security_report
    if architecture_diagram and not normalized_full_data.get("ai_architecture_diagram"):
        normalized_full_data["ai_architecture_diagram"] = architecture_diagram
    if attack_surface_map and not normalized_full_data.get("ai_attack_surface_map"):
        normalized_full_data["ai_attack_surface_map"] = attack_surface_map
    if normalized_vulnerability_hooks and not normalized_full_data.get("vulnerability_frida_hooks"):
        normalized_full_data["vulnerability_frida_hooks"] = normalized_vulnerability_hooks
    if report.manifest_visualization and not normalized_full_data.get("manifest_visualization"):
        normalized_full_data["manifest_visualization"] = report.manifest_visualization
    if report.obfuscation_analysis and not normalized_full_data.get("obfuscation_analysis"):
        normalized_full_data["obfuscation_analysis"] = report.obfuscation_analysis
    if report.verification_results and not normalized_full_data.get("verification_results"):
        normalized_full_data["verification_results"] = report.verification_results
    if report.sensitive_data_findings and not normalized_full_data.get("sensitive_data_findings"):
        normalized_full_data["sensitive_data_findings"] = report.sensitive_data_findings
    if vuln_hunt_result and not normalized_full_data.get("vuln_hunt_result"):
        normalized_full_data["vuln_hunt_result"] = vuln_hunt_result
    if report.cve_scan_results and not normalized_full_data.get("cve_scan"):
        normalized_full_data["cve_scan"] = report.cve_scan_results

    return ReportDetailResponse(
        id=report.id,
        analysis_type=report.analysis_type,
        title=report.title,
        filename=report.filename,
        created_at=report.created_at,
        updated_at=report.updated_at,
        project_id=report.project_id,
        
        risk_level=report.risk_level,
        risk_score=report.risk_score,
        
        file_type=report.file_type,
        architecture=report.architecture,
        file_size=report.file_size,
        is_packed=report.is_packed,
        packer_name=report.packer_name,
        
        package_name=report.package_name,
        version_name=report.version_name,
        min_sdk=report.min_sdk,
        target_sdk=report.target_sdk,
        
        image_name=report.image_name,
        image_id=report.image_id,
        total_layers=report.total_layers,
        base_image=report.base_image,
        
        strings_count=report.strings_count,
        imports_count=report.imports_count,
        exports_count=report.exports_count,
        secrets_count=report.secrets_count,
        
        suspicious_indicators=report.suspicious_indicators,
        permissions=report.permissions,
        security_issues=report.security_issues,
        full_analysis_data=normalized_full_data or full_data_dict or None,
        
        ai_analysis_raw=report.ai_analysis_raw,
        ai_analysis_structured=report.ai_analysis_structured,
        
        # JADX Full Scan Data
        jadx_total_classes=report.jadx_total_classes,
        jadx_total_files=report.jadx_total_files,
        jadx_data=report.jadx_data,
        
        # AI-Generated Reports
        ai_functionality_report=functionality_report,
        ai_security_report=security_report,
        ai_privacy_report=report.ai_privacy_report,
        ai_architecture_diagram=architecture_diagram,
        ai_attack_surface_map=attack_surface_map,
        ai_threat_model=report.ai_threat_model,
        ai_vuln_scan_result=report.ai_vuln_scan_result,
        ai_chat_history=report.ai_chat_history,
        
        # Library CVE Analysis
        detected_libraries=report.detected_libraries,
        library_cves=report.library_cves,
        
        # Dynamic Analysis / Frida Scripts
        dynamic_analysis=report.dynamic_analysis,
        
        # Decompiled Code Analysis Results
        decompiled_code_findings=report.decompiled_code_findings,
        decompiled_code_summary=report.decompiled_code_summary,
        
        # CVE Scan Results
        cve_scan_results=report.cve_scan_results,
        
        # Vulnerability-specific Frida Hooks
        vulnerability_frida_hooks=normalized_vulnerability_hooks,
        
        # Unified APK vulnerability hunt results
        vuln_hunt_result=vuln_hunt_result,
        
        # Manifest Visualization (component graph, deep links, AI analysis)
        manifest_visualization=report.manifest_visualization,
        
        # Obfuscation Analysis (detection, deobfuscation strategies, Frida hooks)
        obfuscation_analysis=report.obfuscation_analysis,
        
        # AI Finding Verification Results (confidence scores, attack chains, FP filtering)
        verification_results=report.verification_results,
        
        # Sensitive Data Discovery (AI-verified passwords, API keys, emails, phone numbers, PII)
        sensitive_data_findings=report.sensitive_data_findings,
        
        tags=report.tags,
        notes=report.notes,
    )


@router.get("/reports/{report_id}/export")
def export_saved_report(
    report_id: int,
    format: str = Query(..., description="Export format: markdown, pdf, docx"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    Export a saved reverse engineering report to Markdown, PDF, or Word format.
    
    Includes all analysis data, AI reports, and JADX findings if available.
    """
    if format not in ["markdown", "pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Invalid format. Use: markdown, pdf, docx")

    report = _ensure_reverse_engineering_report_access(
        db,
        _get_reverse_engineering_report_or_404(db, report_id),
        current_user.id,
    )
    
    try:
        # Generate comprehensive export
        content = _generate_full_report_export(report, format)
        
        # Prefer report title for exported filename so named scans keep their identity.
        base_name_source = report.title or report.package_name or report.image_name or report.filename or f"report_{report_id}"
        base_name = "".join(ch if (ch.isalnum() or ch in {"-", "_"}) else "_" for ch in str(base_name_source))
        base_name = base_name.strip("_")[:120] or f"report_{report_id}"
        
        if format == "markdown":
            return Response(
                content=content.encode('utf-8'),
                media_type="text/markdown",
                headers={"Content-Disposition": f'attachment; filename="{base_name}_full_report.md"'}
            )
        elif format == "pdf":
            return Response(
                content=content,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{base_name}_full_report.pdf"'}
            )
        elif format == "docx":
            return Response(
                content=content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{base_name}_full_report.docx"'}
            )
    except Exception as e:
        logger.error(f"Failed to export report {report_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


@router.post("/binary/export-from-result")
async def export_binary_report_from_result(
    result: BinaryAnalysisResponse,
    format: str = Query(..., description="Export format: markdown, pdf, docx"),
    report_title: Optional[str] = Query(None, description="Optional custom report title"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Export a binary analysis result to Markdown, PDF, or Word format.
    """
    if format not in ["markdown", "pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Invalid format. Use: markdown, pdf, docx")

    def _to_dict_list(items: Any) -> List[Dict[str, Any]]:
        output: List[Dict[str, Any]] = []
        if not isinstance(items, list):
            return output
        for item in items:
            if isinstance(item, dict):
                output.append(item)
            elif hasattr(item, "model_dump"):
                dumped = item.model_dump()
                if isinstance(dumped, dict):
                    output.append(dumped)
            elif hasattr(item, "dict"):
                dumped = item.dict()
                if isinstance(dumped, dict):
                    output.append(dumped)
        return output

    verification_result = result.verification_result if isinstance(result.verification_result, dict) else {}
    pattern_scan_result = result.pattern_scan_result if isinstance(result.pattern_scan_result, dict) else {}
    cve_lookup_result = result.cve_lookup_result if isinstance(result.cve_lookup_result, dict) else {}
    vuln_hunt_result = result.vuln_hunt_result if isinstance(result.vuln_hunt_result, dict) else {}

    export_security_issues = (
        _to_dict_list(verification_result.get("verified_vulnerabilities"))
        + _to_dict_list(verification_result.get("verified_findings"))
        + _to_dict_list(verification_result.get("findings"))
        + _to_dict_list(pattern_scan_result.get("findings"))
        + _to_dict_list(vuln_hunt_result.get("vulnerabilities"))
    )
    if not export_security_issues:
        export_security_issues = _to_dict_list(cve_lookup_result.get("findings"))

    deduped_export_findings: List[Dict[str, Any]] = []
    seen_export_keys: Set[str] = set()
    for finding in export_security_issues:
        title = str(finding.get("title") or finding.get("name") or finding.get("type") or finding.get("category") or "Security Finding")
        location = str(finding.get("file") or finding.get("location") or finding.get("affected_class") or "")
        description = str(finding.get("description") or finding.get("summary") or finding.get("details") or "")
        dedupe_key = f"{title.strip().lower()}::{location.strip().lower()}::{description[:120].strip().lower()}"
        if dedupe_key in seen_export_keys:
            continue
        seen_export_keys.add(dedupe_key)
        deduped_export_findings.append(finding)

    resolved_title = (report_title or "").strip()
    if resolved_title:
        resolved_title = resolved_title[:180]
    else:
        resolved_title = f"Binary Analysis: {result.filename}"

    report = ExportSafeNamespace(
        title=resolved_title,
        created_at=datetime.now(),
        updated_at=datetime.now(),
        analysis_type="binary",
        filename=result.filename,
        risk_level=None,
        risk_score=None,
        file_type=result.metadata.file_type,
        architecture=result.metadata.architecture,
        file_size=result.metadata.file_size,
        is_packed=str(result.metadata.is_packed),
        packer_name=result.metadata.packer_name,
        strings_count=result.strings_count,
        imports_count=len(result.imports),
        exports_count=len(result.exports),
        secrets_count=len(result.secrets),
        suspicious_indicators=[s.model_dump() for s in result.suspicious_indicators],
        permissions=None,
        security_issues=deduped_export_findings or None,
        full_analysis_data={
            "metadata": result.metadata.model_dump(),
            "strings_sample": [s.model_dump() for s in result.strings_sample],
            "imports": [i.model_dump() for i in result.imports],
            "exports": result.exports,
            "secrets": [s.model_dump() for s in result.secrets],
            "ghidra_analysis": result.ghidra_analysis,
            "ghidra_ai_summaries": result.ghidra_ai_summaries,
            "security_issues": deduped_export_findings or None,
            "verified_findings": _to_dict_list(verification_result.get("verified_findings")),
            "verified_vulnerabilities": _to_dict_list(verification_result.get("verified_vulnerabilities")),
            "verification_results": verification_result or None,
            "pattern_scan_result": pattern_scan_result or None,
            "cve_lookup_result": cve_lookup_result or None,
            "vuln_hunt_result": vuln_hunt_result or None,
            "attack_surface": result.attack_surface,
            "ai_functionality_report": result.ai_functionality_report,
            "ai_security_report": result.ai_security_report,
            "ai_architecture_diagram": result.ai_architecture_diagram,
            "ai_attack_surface_map": result.ai_attack_surface_map,
        },
        ai_analysis_raw=result.ai_analysis,
        ai_analysis_structured=None,
        tags=None,
        notes=None,
        ai_functionality_report=result.ai_functionality_report,
        ai_security_report=result.ai_security_report,
        ai_privacy_report=None,
        ai_architecture_diagram=result.ai_architecture_diagram,
        ai_attack_surface_map=result.ai_attack_surface_map,
        ai_threat_model=None,
        ai_vuln_scan_result=verification_result or None,
        ai_chat_history=None,
        detected_libraries=None,
        library_cves=_to_dict_list(cve_lookup_result.get("findings")) or None,
        # NEW: Include advanced analysis data
        obfuscation_analysis=result.obfuscation_analysis,
        attack_surface=result.attack_surface,
        pattern_scan_result=pattern_scan_result or None,
        cve_lookup_result=cve_lookup_result or None,
        vuln_hunt_result=vuln_hunt_result or None,
        is_legitimate_software=result.is_legitimate_software,
        legitimacy_indicators=result.legitimacy_indicators,
        verification_results=verification_result or None,  # Also expose as verification_results for export
    )

    try:
        content = _generate_full_report_export(report, format)
        base_name_source = (report_title or "").strip() or f"binary_analysis_{result.filename}"
        base_name = "".join(ch if (ch.isalnum() or ch in {"-", "_"}) else "_" for ch in base_name_source)
        base_name = base_name.strip("_")[:120] or "binary_analysis_report"

        if format == "markdown":
            return Response(
                content=content.encode("utf-8"),
                media_type="text/markdown",
                headers={"Content-Disposition": f'attachment; filename="{base_name}.md"'}
            )
        if format == "pdf":
            return Response(
                content=content,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{base_name}.pdf"'}
            )
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base_name}.docx"'}
        )
    except Exception as e:
        logger.error(f"Binary export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


def _generate_full_report_export(report, format: str):
    """Generate full export content for a saved report with clean, organized structure."""
    from io import BytesIO
    
    # Build comprehensive markdown content with clean structure
    # Structure: Executive Summary → Security Findings → Architecture → AI Analysis → Attack Surface → Secrets
    
    md_content = f"""# {report.title}

**Generated:** {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}
**Risk Level:** {report.risk_level or 'Not Assessed'} ({report.risk_score or 0}/100)

---

"""
    # Normalize report fields from nested/full payloads so saved and live exports behave the same.
    def _coerce_dict(value: Any) -> Dict[str, Any]:
        if isinstance(value, dict):
            return value
        if isinstance(value, str):
            try:
                parsed = json.loads(value)
                if isinstance(parsed, dict):
                    return parsed
            except Exception:
                return {}
        if hasattr(value, "model_dump"):
            dumped = value.model_dump()
            if isinstance(dumped, dict):
                return dumped
        if hasattr(value, "dict"):
            dumped = value.dict()
            if isinstance(dumped, dict):
                return dumped
        return {}

    def _coerce_list(value: Any) -> List[Any]:
        if isinstance(value, list):
            return value
        if isinstance(value, tuple):
            return list(value)
        return []

    def _clean_text(value: Any) -> Optional[str]:
        if not isinstance(value, str):
            return None
        trimmed = value.strip()
        if not trimmed:
            return None
        if trimmed.lower() in {"not available", "n/a", "none", "null", "undefined"}:
            return None
        return trimmed

    def _collect_dicts(data: Dict[str, Any], depth: int = 3) -> List[Dict[str, Any]]:
        collected: List[Dict[str, Any]] = []
        seen_ids: Set[int] = set()

        def _walk(node: Any, level: int):
            if level > depth:
                return
            if isinstance(node, dict):
                node_id = id(node)
                if node_id in seen_ids:
                    return
                seen_ids.add(node_id)
                collected.append(node)
                for child in node.values():
                    _walk(child, level + 1)
            elif isinstance(node, list):
                for child in node:
                    _walk(child, level + 1)

        _walk(data, 0)
        return collected

    def _pick_text(dicts: List[Dict[str, Any]], keys: List[str]) -> Optional[str]:
        for data in dicts:
            for key in keys:
                cleaned = _clean_text(data.get(key))
                if cleaned:
                    return cleaned
        return None

    def _pick_text_by_hint(dicts: List[Dict[str, Any]], hints: List[str], minimum_length: int = 120) -> Optional[str]:
        for data in dicts:
            for key, raw_value in data.items():
                lowered_key = str(key).lower()
                if not any(hint in lowered_key for hint in hints):
                    continue
                cleaned = _clean_text(raw_value)
                if cleaned and len(cleaned) >= minimum_length:
                    return cleaned
        return None

    def _is_mermaid(text: Optional[str]) -> bool:
        if not isinstance(text, str):
            return False
        trimmed = text.strip()
        if not trimmed:
            return False
        lowered = trimmed.lower()
        if "```mermaid" in lowered:
            return True
        return lowered.startswith((
            "flowchart",
            "graph ",
            "sequencediagram",
            "classdiagram",
            "statediagram",
            "erdiagram",
            "journey",
            "gantt",
        ))

    def _pick_mermaid(dicts: List[Dict[str, Any]], keys: List[str]) -> Optional[str]:
        direct = _pick_text(dicts, keys)
        if _is_mermaid(direct):
            return direct
        for data in dicts:
            for value in data.values():
                if isinstance(value, str) and _is_mermaid(value):
                    return value.strip()
        return None

    def _as_dict(value: Any) -> Dict[str, Any]:
        if isinstance(value, dict):
            return value
        if hasattr(value, "model_dump"):
            dumped = value.model_dump()
            if isinstance(dumped, dict):
                return dumped
        if hasattr(value, "dict"):
            dumped = value.dict()
            if isinstance(dumped, dict):
                return dumped
        return {}

    def _dict_items(value: Any) -> List[Dict[str, Any]]:
        items: List[Dict[str, Any]] = []
        for raw in _coerce_list(value):
            as_dict = _as_dict(raw)
            if as_dict:
                items.append(as_dict)
        return items

    def _severity_rank(severity: str) -> int:
        sev = str(severity or "").lower()
        if sev == "critical":
            return 0
        if sev == "high":
            return 1
        if sev == "medium":
            return 2
        if sev == "low":
            return 3
        return 4

    def _to_int(value: Any, default: int = 0) -> int:
        try:
            if value is None or value == "":
                return default
            return int(float(value))
        except (TypeError, ValueError):
            return default

    def _to_float(value: Any, default: float = 0.0) -> float:
        try:
            if value is None or value == "":
                return default
            return float(value)
        except (TypeError, ValueError):
            return default

    def _markdown_code(value: Any) -> str:
        return str(value or "N/A").replace("`", "'").replace("\r", " ").replace("\n", " ")

    full_data_dict = _coerce_dict(getattr(report, "full_analysis_data", None))
    structured_dict = _coerce_dict(getattr(report, "ai_analysis_structured", None))
    ai_raw = _clean_text(getattr(report, "ai_analysis_raw", None))
    ai_raw_dict = _coerce_dict(ai_raw) if ai_raw else {}
    verification_dict = _coerce_dict(getattr(report, "verification_results", None))
    ai_vuln_dict = _coerce_dict(getattr(report, "ai_vuln_scan_result", None))
    pattern_scan_dict = _coerce_dict(getattr(report, "pattern_scan_result", None)) or _coerce_dict(full_data_dict.get("pattern_scan_result"))
    cve_lookup_dict = _coerce_dict(getattr(report, "cve_lookup_result", None)) or _coerce_dict(full_data_dict.get("cve_lookup_result"))
    vuln_hunt_dict = _coerce_dict(getattr(report, "vuln_hunt_result", None)) or _coerce_dict(full_data_dict.get("vuln_hunt_result"))
    search_space = (
        _collect_dicts(full_data_dict)
        + _collect_dicts(structured_dict)
        + _collect_dicts(ai_raw_dict, 4)
        + _collect_dicts(verification_dict)
        + _collect_dicts(ai_vuln_dict)
        + _collect_dicts(pattern_scan_dict)
        + _collect_dicts(cve_lookup_dict)
        + _collect_dicts(vuln_hunt_dict)
    )

    functionality_report = _clean_text(getattr(report, "ai_functionality_report", None)) or _pick_text(search_space, [
        "ai_functionality_report",
        "ai_report_functionality",
        "functionality_report",
        "functionality",
        "report_functionality",
        "binary_functionality_report",
    ]) or _pick_text_by_hint(search_space, ["functionality", "purpose", "behavior", "what_does"])
    security_report = _clean_text(getattr(report, "ai_security_report", None)) or _pick_text(search_space, [
        "ai_security_report",
        "ai_report_security",
        "security_report",
        "security",
        "report_security",
        "binary_security_report",
    ]) or _pick_text_by_hint(search_space, ["security", "risk", "vulnerab", "threat"])
    architecture_diagram = _clean_text(getattr(report, "ai_architecture_diagram", None)) or _pick_mermaid(search_space, [
        "ai_architecture_diagram",
        "architecture_diagram",
        "ai_architecture",
        "architecture",
        "architecture_map",
        "system_architecture_diagram",
    ])
    attack_surface_map = _clean_text(getattr(report, "ai_attack_surface_map", None)) or _pick_mermaid(search_space, [
        "ai_attack_surface_map",
        "attack_surface_map",
        "attack_surface",
        "ai_attack_surface",
        "attack_tree",
        "attack_tree_diagram",
    ])

    if not functionality_report and ai_raw:
        functionality_report = ai_raw
    if not security_report and ai_raw:
        security_report = ai_raw

    docker_cve_scan_dict = _as_dict(getattr(report, "cve_scan_results", None))
    if not docker_cve_scan_dict:
        docker_cve_scan_dict = _as_dict(full_data_dict.get("cve_scan")) or _as_dict(full_data_dict.get("cve_scan_results"))
    docker_cve_findings = _coerce_list(docker_cve_scan_dict.get("vulnerabilities"))
    docker_trivy_scan_dict = _as_dict(getattr(report, "trivy_scan_results", None)) or _as_dict(getattr(report, "trivy_scan", None))
    if not docker_trivy_scan_dict:
        docker_trivy_scan_dict = _as_dict(full_data_dict.get("trivy_scan")) or _as_dict(full_data_dict.get("trivy_scan_results"))
    docker_trivy_vulns = _dict_items(docker_trivy_scan_dict.get("vulnerabilities"))
    docker_trivy_misconfigs = _dict_items(docker_trivy_scan_dict.get("misconfigurations"))
    docker_trivy_secrets = _dict_items(docker_trivy_scan_dict.get("secrets"))
    docker_trivy_licenses = _dict_items(docker_trivy_scan_dict.get("licenses"))
    docker_trivy_security_findings: List[Dict[str, Any]] = []
    for vuln in docker_trivy_vulns:
        docker_trivy_security_findings.append({
            "category": "trivy_vulnerability",
            "title": vuln.get("vulnerability_id") or vuln.get("title") or "Trivy Vulnerability",
            "description": vuln.get("description") or vuln.get("title") or "Native Trivy vulnerability finding.",
            "severity": vuln.get("severity") or "unknown",
            "location": vuln.get("target") or vuln.get("package_name") or "container image",
            "package_name": vuln.get("package_name"),
            "cvss_score": vuln.get("cvss_score") or vuln.get("cvss"),
            "attack_vector": "Known vulnerable package present in the container image.",
            "remediation": (
                f"Upgrade {vuln.get('package_name')} to {vuln.get('fixed_version')}"
                if vuln.get("package_name") and vuln.get("fixed_version")
                else "Upgrade or remove the affected package."
            ),
        })
    for misconfig in docker_trivy_misconfigs:
        docker_trivy_security_findings.append({
            "category": "trivy_misconfiguration",
            "title": misconfig.get("id") or misconfig.get("title") or "Trivy Misconfiguration",
            "description": misconfig.get("message") or misconfig.get("description") or misconfig.get("title") or "Native Trivy misconfiguration finding.",
            "severity": misconfig.get("severity") or "unknown",
            "location": misconfig.get("target") or "container image",
            "attack_vector": "Container misconfiguration weakens hardening or broadens attacker options.",
            "remediation": misconfig.get("resolution") or misconfig.get("recommendation") or "Apply the recommended hardening change and rebuild the image.",
        })
    for secret in docker_trivy_secrets:
        docker_trivy_security_findings.append({
            "category": "trivy_secret",
            "title": secret.get("rule_id") or secret.get("title") or "Trivy Secret Finding",
            "description": "Trivy detected secret material in the container image. See the Native Trivy Findings section for the raw match.",
            "severity": secret.get("severity") or "high",
            "location": secret.get("target") or "container image",
            "attack_vector": "Credential or secret material embedded in the container filesystem can enable unauthorized access.",
            "remediation": "Remove the secret from the image, rotate the credential, and rebuild from a clean source.",
        })

    # Gather comprehensive security findings across raw/verified/scanner sources.
    raw_security_findings: List[Dict[str, Any]] = []
    for source in [
        getattr(report, "security_issues", None),
        getattr(report, "decompiled_code_findings", None),
        full_data_dict.get("security_issues"),
        full_data_dict.get("decompiled_code_findings"),
        verification_dict.get("verified_vulnerabilities"),
        verification_dict.get("verified_findings"),
        verification_dict.get("findings"),
        ai_vuln_dict.get("combined_findings"),
        ai_vuln_dict.get("vulnerabilities"),
        ai_vuln_dict.get("findings"),
        pattern_scan_dict.get("findings"),
        vuln_hunt_dict.get("vulnerabilities"),
        full_data_dict.get("verified_findings"),
        full_data_dict.get("verified_vulnerabilities"),
        full_data_dict.get("combined_findings"),
        docker_cve_findings,
        docker_trivy_security_findings,
    ]:
        for item in _coerce_list(source):
            as_dict = _as_dict(item)
            if as_dict:
                raw_security_findings.append(as_dict)

    deduped_findings: List[Dict[str, Any]] = []
    seen_finding_keys: Set[str] = set()
    for item in raw_security_findings:
        severity = str(
            item.get("severity")
            or item.get("risk_level")
            or item.get("risk")
            or item.get("confidence_level")
            or "info"
        ).lower()
        title = str(
            item.get("title")
            or item.get("external_id")
            or item.get("id")
            or item.get("name")
            or item.get("type")
            or item.get("category")
            or "Security Finding"
        )
        description = str(item.get("description") or item.get("summary") or item.get("details") or item.get("reasoning") or item.get("finding") or "")
        location = str(
            item.get("file")
            or item.get("location")
            or item.get("affected_class")
            or item.get("package_name")
            or item.get("package")
            or item.get("source")
            or item.get("path")
            or ""
        )
        key = f"{title.strip().lower()}::{location.strip().lower()}::{description[:120].strip().lower()}"
        if key in seen_finding_keys:
            continue
        seen_finding_keys.add(key)
        normalized = dict(item)
        normalized["_normalized_severity"] = severity
        normalized["_normalized_title"] = title
        normalized["_normalized_description"] = description
        normalized["_normalized_location"] = location
        deduped_findings.append(normalized)

    all_security_findings = sorted(
        deduped_findings,
        key=lambda finding: (
            _severity_rank(finding.get("_normalized_severity", "info")),
            finding.get("_normalized_title", ""),
        ),
    )

    cve_findings = _coerce_list(getattr(report, "library_cves", None))
    if not cve_findings:
        cve_findings = _coerce_list(cve_lookup_dict.get("findings")) or _coerce_list(cve_lookup_dict.get("cves"))
    if not cve_findings:
        cve_findings = _coerce_list(full_data_dict.get("library_cves")) or _coerce_list(full_data_dict.get("cve_findings"))
    if not cve_findings and report.analysis_type == "docker":
        cve_findings = docker_cve_findings
    cve_findings = [_as_dict(item) for item in cve_findings if _as_dict(item)]
    
    # =====================================================
    # SECTION 1: Executive Summary - What Does This APK Do?
    # =====================================================
    md_content += """## Executive Summary

"""
    
    # Basic identification info (minimal)
    if report.analysis_type == 'apk':
        md_content += f"**Package:** `{report.package_name or 'N/A'}` | **Version:** {report.version_name or 'N/A'}\n\n"
    elif report.analysis_type == 'binary':
        md_content += f"**Type:** {report.file_type or 'N/A'} | **Architecture:** {report.architecture or 'N/A'}\n\n"
    elif report.analysis_type == 'docker':
        md_content += f"**Image:** `{report.image_name or 'N/A'}` | **Base:** {report.base_image or 'N/A'}\n\n"

    target_label = {
        "apk": "This APK",
        "binary": "This Binary",
        "docker": "This Container Image",
    }.get(report.analysis_type, "This Analysis Target")
    
    # Legitimate software detection
    if hasattr(report, 'is_legitimate_software') and report.is_legitimate_software:
        md_content += """### ✅ Legitimate Software Detected

This binary has been identified as **legitimate software** from a known publisher. Security findings have been contextually filtered to focus on configuration issues rather than false-positive malware indicators.

"""
        if hasattr(report, 'legitimacy_indicators') and report.legitimacy_indicators:
            md_content += "**Detection Indicators:**\n\n"
            for indicator in report.legitimacy_indicators[:10]:
                md_content += f"- {indicator}\n"
            md_content += "\n"
    
    # AI Functionality Report - What the app does
    if functionality_report:
        md_content += f"""### What Does {target_label} Do?

{functionality_report}

"""
    elif report.ai_analysis_raw:
        # Fallback to quick analysis if no functionality report
        md_content += f"""### Application Overview

{report.ai_analysis_raw}

"""
    
    # =====================================================
    # SECTION 2: Security Findings - Attacker's Perspective
    # =====================================================
    md_content += """---

## Security Assessment - Attack Surface Analysis

This section analyzes the application from an attacker's perspective, identifying exploitable weaknesses and potential attack vectors.

"""
    
    # AI Security Report (high-level security assessment)
    if security_report:
        md_content += f"""{security_report}

"""
    
    # Consolidated security issues summary with attacker focus
    critical_count = high_count = medium_count = low_count = 0
    if all_security_findings:
        for issue in all_security_findings:
            sev = str(issue.get('_normalized_severity', issue.get('severity', 'info'))).lower()
            if sev == 'critical': critical_count += 1
            elif sev == 'high': high_count += 1
            elif sev == 'medium': medium_count += 1
            elif sev == 'low': low_count += 1
        
        md_content += f"""### Vulnerability Summary

| Severity | Count | Attack Priority |
|----------|-------|-----------------|
| Critical | {critical_count} | Immediate exploitation possible |
| High | {high_count} | High-value targets |
| Medium | {medium_count} | Secondary targets |
| Low | {low_count} | Opportunistic |

"""
        # List critical and high issues with attacker-focused descriptions
        critical_high = [
            i for i in all_security_findings
            if str(i.get('_normalized_severity', i.get('severity', ''))).lower() in ['critical', 'high']
        ]
        if critical_high:
            md_content += "### Exploitable Vulnerabilities\n\n"
            md_content += "*These issues present immediate exploitation opportunities:*\n\n"
            for issue in critical_high:
                severity = issue.get('_normalized_severity', issue.get('severity', 'info'))
                title = issue.get('_normalized_title', issue.get('title', issue.get('category', 'Unknown Issue')))
                md_content += f"#### {title} ({severity.upper()})\n\n"
                
                description = issue.get('_normalized_description', issue.get('description'))
                if description:
                    md_content += f"**Finding:** {description}\n\n"
                
                # Add exploitation context
                category = str(issue.get('category', '')).lower()
                if 'sql' in category or 'injection' in category:
                    md_content += "**Attack Vector:** Database manipulation, data exfiltration, authentication bypass\n\n"
                elif 'crypto' in category or 'encrypt' in category:
                    md_content += "**Attack Vector:** Decrypt sensitive data, forge tokens, bypass integrity checks\n\n"
                elif 'auth' in category or 'permission' in category:
                    md_content += "**Attack Vector:** Privilege escalation, unauthorized access to protected resources\n\n"
                elif 'webview' in category or 'javascript' in category:
                    md_content += "**Attack Vector:** JavaScript injection, steal cookies/tokens, phishing\n\n"
                elif 'export' in category or 'component' in category:
                    md_content += "**Attack Vector:** Intent hijacking, data theft via exported components\n\n"
                elif 'network' in category or 'http' in category:
                    md_content += "**Attack Vector:** Man-in-the-middle, traffic interception, credential theft\n\n"
                elif 'secret' in category or 'key' in category or 'credential' in category:
                    md_content += "**Attack Vector:** Use hardcoded credentials for unauthorized API/service access\n\n"
                elif 'debug' in category:
                    md_content += "**Attack Vector:** Attach debugger, inspect memory, modify runtime behavior\n\n"
                
                # Include file/location for manual review
                file_loc = issue.get('_normalized_location') or issue.get('file', issue.get('location', issue.get('affected_class', '')))
                if file_loc:
                    md_content += f"**Target Location:** `{file_loc}`\n\n"

        # Complete finding list (all severities) for full export fidelity.
        md_content += f"""### Security Findings (Complete)

The following list includes **all {len(all_security_findings)} findings** captured in this scan.

"""
        for idx, finding in enumerate(all_security_findings, 1):
            severity = finding.get('_normalized_severity', finding.get('severity', 'info')).upper()
            title = finding.get('_normalized_title', finding.get('title', finding.get('category', 'Security Finding')))
            description = finding.get('_normalized_description', finding.get('description', ''))
            location = finding.get('_normalized_location', '')
            cwe = finding.get('cwe') or finding.get('cwe_id')
            cvss = finding.get('cvss') or finding.get('cvss_score')
            attack_vector = finding.get('attack_vector') or finding.get('exploitation_path')
            remediation = finding.get('remediation') or finding.get('recommendation') or finding.get('fix')

            md_content += f"#### {idx}. {title} ({severity})\n\n"
            if description:
                md_content += f"{description}\n\n"
            if location:
                md_content += f"- **Location:** `{location}`\n"
            if cwe:
                md_content += f"- **CWE:** {cwe}\n"
            if cvss:
                md_content += f"- **CVSS:** {cvss}\n"
            if attack_vector:
                md_content += f"- **Attack Vector:** {attack_vector}\n"
            if remediation:
                md_content += f"- **Remediation:** {remediation}\n"
            md_content += "\n"
    
    # Dangerous permissions with exploitation context
    permissions_list = _dict_items(getattr(report, "permissions", None))
    if permissions_list:
        dangerous = [p for p in permissions_list if p.get('is_dangerous')]
        if dangerous:
            md_content += "### Dangerous Permissions - Attack Enablers\n\n"
            md_content += "*These permissions grant capabilities that can be abused:*\n\n"
            for p in dangerous[:10]:
                perm_name = p.get('name', '')
                md_content += f"- **{perm_name}**\n"
                
                # Add exploitation context for common dangerous permissions
                if 'CAMERA' in perm_name:
                    md_content += "  - *Exploitation:* Covert photo/video capture, surveillance\n"
                elif 'MICROPHONE' in perm_name or 'RECORD_AUDIO' in perm_name:
                    md_content += "  - *Exploitation:* Audio surveillance, conversation recording\n"
                elif 'LOCATION' in perm_name:
                    md_content += "  - *Exploitation:* User tracking, location history theft\n"
                elif 'READ_CONTACTS' in perm_name or 'WRITE_CONTACTS' in perm_name:
                    md_content += "  - *Exploitation:* Contact list exfiltration, spam distribution\n"
                elif 'READ_SMS' in perm_name or 'SEND_SMS' in perm_name:
                    md_content += "  - *Exploitation:* 2FA bypass, premium SMS fraud\n"
                elif 'READ_EXTERNAL_STORAGE' in perm_name or 'WRITE_EXTERNAL_STORAGE' in perm_name:
                    md_content += "  - *Exploitation:* Data theft from shared storage, malware staging\n"
                elif 'READ_CALL_LOG' in perm_name:
                    md_content += "  - *Exploitation:* Call history surveillance, contact profiling\n"
                elif 'SYSTEM_ALERT_WINDOW' in perm_name:
                    md_content += "  - *Exploitation:* Overlay attacks, clickjacking, credential theft\n"
                else:
                    md_content += f"  - {p.get('description', 'Review for abuse potential')}\n"
            md_content += "\n"
    
    # =====================================================
    # SECTION 2B: Known CVEs in Third-Party Libraries
    # =====================================================
    if cve_findings:
        cves = cve_findings
        critical_cves = [c for c in cves if str(c.get('severity', '')).lower() == 'critical']
        high_cves = [c for c in cves if str(c.get('severity', '')).lower() == 'high']
        cve_heading = (
            "Known Vulnerabilities in Container Packages (CVE Scan)"
            if report.analysis_type == "docker"
            else "Known Vulnerabilities in Dependencies (CVE Lookup)"
        )
        cve_description = (
            "*These are known, published vulnerabilities in packages extracted from this container image.*"
            if report.analysis_type == "docker"
            else "*These are known, published vulnerabilities in the third-party libraries bundled with this application.*"
        )
        
        md_content += f"""### {cve_heading}

**Total CVEs Found:** {len(cves)} | **Critical:** {len(critical_cves)} | **High:** {len(high_cves)}

{cve_description}

"""
        
        # Critical CVEs first - immediate exploitation risk
        if critical_cves:
            md_content += "#### Critical CVEs - Immediate Exploitation Risk\n\n"
            for cve in critical_cves[:10]:
                cve_id = cve.get('cve_id') or cve.get('external_id') or cve.get('id') or 'Unknown'
                library = cve.get('library') or cve.get('package_name') or cve.get('package') or 'Unknown'
                summary = (cve.get('summary') or cve.get('title') or cve.get('description') or 'No description')[:200]
                exploitation = cve.get('exploitation_potential') or cve.get('priority_label') or 'Review required'
                attack_vector = cve.get('attack_vector') or ('Known exploited vulnerability' if cve.get('in_kev') else 'Unknown')
                
                md_content += f"**{cve_id}** in `{library}`\n\n"
                md_content += f"- **Summary:** {summary}\n"
                md_content += f"- **CVSS Score:** {cve.get('cvss_score', 'N/A')}\n"
                md_content += f"- **Exploitation:** {exploitation}\n"
                md_content += f"- **Attack Vector:** {attack_vector}\n"
                
                if cve.get('fixed_version'):
                    md_content += f"- **Fix:** Upgrade to {cve.get('fixed_version')}\n"
                elif cve.get('affected_versions'):
                    md_content += f"- **Fix:** {', '.join(cve.get('affected_versions', []))}\n"
                
                if cve.get('references'):
                    md_content += f"- **Reference:** {cve.get('references', [''])[0]}\n"
                
                md_content += "\n"
        
        # High severity CVEs
        if high_cves:
            md_content += "#### High Severity CVEs\n\n"
            for cve in high_cves[:10]:
                cve_id = cve.get('cve_id') or cve.get('external_id') or cve.get('id') or 'Unknown'
                library = cve.get('library') or cve.get('package_name') or cve.get('package') or 'Unknown'
                summary = (cve.get('summary') or cve.get('title') or cve.get('description') or 'No description')[:150]
                signal = cve.get('exploitation_potential') or cve.get('priority_label') or ('KEV' if cve.get('in_kev') else '')
                
                md_content += f"- **{cve_id}** in `{library}`: {summary}\n"
                md_content += f"  - CVSS: {cve.get('cvss_score', 'N/A')} | {signal}\n"
            
            md_content += "\n"
        
        # Summary table of all CVEs
        if len(cves) > 0:
            md_content += "#### All Detected CVEs\n\n"
            md_content += "| CVE ID | Package | Severity | CVSS |\n"
            md_content += "|--------|---------|----------|------|\n"
            for cve in cves[:25]:
                cve_id = cve.get('cve_id') or cve.get('external_id') or cve.get('id') or 'N/A'
                package_name = cve.get('library') or cve.get('package_name') or cve.get('package') or 'N/A'
                md_content += f"| {cve_id} | {str(package_name)[:30]} | {cve.get('severity', 'N/A').upper()} | {cve.get('cvss_score', 'N/A')} |\n"
            
            if len(cves) > 25:
                md_content += f"\n*...and {len(cves) - 25} more CVEs*\n"
            
            md_content += "\n"

    if report.analysis_type == "docker" and docker_trivy_scan_dict:
        enabled_scanners = ", ".join(str(scanner) for scanner in _coerce_list(docker_trivy_scan_dict.get("enabled_scanners")) if scanner) or "None"
        trivy_duration = _to_float(docker_trivy_scan_dict.get("scan_duration_seconds"), 0.0)
        md_content += f"""### Native Trivy Findings

**Enabled Scanners:** {enabled_scanners} | **Packages Observed:** {_to_int(docker_trivy_scan_dict.get('package_count', 0))} | **Vulnerabilities:** {len(docker_trivy_vulns)} | **Misconfigurations:** {len(docker_trivy_misconfigs)} | **Secrets:** {len(docker_trivy_secrets)} | **Licenses:** {len(docker_trivy_licenses)} | **Duration:** {trivy_duration:.1f}s

*This section preserves the native Trivy image-scan output alongside the Docker Inspector's enriched analysis.*

"""
        if docker_trivy_scan_dict.get("error"):
            md_content += f"**Scan Error:** {docker_trivy_scan_dict.get('error')}\n\n"

        if docker_trivy_vulns:
            md_content += "#### Native Trivy Vulnerabilities\n\n"
            md_content += "| ID | Severity | Package | Installed | Fixed |\n"
            md_content += "|----|----------|---------|-----------|-------|\n"
            for finding in docker_trivy_vulns[:25]:
                md_content += (
                    f"| {finding.get('vulnerability_id', 'N/A')} | "
                    f"{str(finding.get('severity', 'unknown')).upper()} | "
                    f"{finding.get('package_name', 'N/A')} | "
                    f"{finding.get('installed_version', 'N/A')} | "
                    f"{finding.get('fixed_version', 'N/A') or '-'} |\n"
                )
            md_content += "\n"

        if docker_trivy_misconfigs:
            md_content += "#### Trivy Misconfigurations\n\n"
            md_content += "| ID | Severity | Target | Message |\n"
            md_content += "|----|----------|--------|---------|\n"
            for finding in docker_trivy_misconfigs[:25]:
                md_content += (
                    f"| {finding.get('id', 'N/A')} | "
                    f"{str(finding.get('severity', 'unknown')).upper()} | "
                    f"{str(finding.get('target', 'N/A'))[:40]} | "
                    f"{str(finding.get('message') or finding.get('title') or 'N/A')[:100]} |\n"
                )
            md_content += "\n"

        if docker_trivy_secrets:
            md_content += "#### Trivy Secret Findings\n\n"
            md_content += "_Operator note: raw Trivy secret matches are intentionally included for remediation workflows._\n\n"
            for finding in docker_trivy_secrets[:20]:
                md_content += f"- **{finding.get('rule_id', 'Unknown')}** ({str(finding.get('severity', 'unknown')).upper()})\n"
                md_content += f"  - Target: `{_markdown_code(finding.get('target'))}`\n"
                md_content += f"  - Match: `{_markdown_code(finding.get('match'))}`\n"
            if len(docker_trivy_secrets) > 20:
                md_content += f"\n*...and {len(docker_trivy_secrets) - 20} more Trivy secret findings*\n"
            md_content += "\n"

        if docker_trivy_licenses:
            md_content += "#### Trivy License Findings\n\n"
            md_content += "| License | Severity | Package | Category |\n"
            md_content += "|---------|----------|---------|----------|\n"
            for finding in docker_trivy_licenses[:25]:
                md_content += (
                    f"| {finding.get('name', 'N/A')} | "
                    f"{str(finding.get('severity', 'unknown')).upper()} | "
                    f"{finding.get('package_name', 'N/A')} | "
                    f"{finding.get('category', 'N/A')} |\n"
                )
            md_content += "\n"

    # Detected libraries summary (even if no CVEs)
    libs = _dict_items(getattr(report, "detected_libraries", None))
    if libs and not cve_findings:
        high_risk = [l for l in libs if l.get('is_high_risk')]
        
        md_content += f"""### Detected Third-Party Libraries

**Total Libraries:** {len(libs)} | **High-Risk Libraries:** {len(high_risk)}

*No known CVEs found in OSV database for the detected library versions.*

"""
        if high_risk:
            md_content += "**High-Risk Libraries (historically vulnerable):**\n\n"
            for lib in high_risk:
                md_content += f"- `{lib.get('maven_coordinate', 'Unknown')}` - {lib.get('risk_reason', 'Review recommended')}\n"
            md_content += "\n"

    # =====================================================
    # SECTION 2C: Binary Static Analysis (Binary only)
    # =====================================================
    if report.analysis_type == 'binary':
        md_content += """---

## Binary Static Analysis

"""
        md_content += f"**File Type:** {report.file_type or 'N/A'} | **Architecture:** {report.architecture or 'N/A'}\n\n"
        md_content += f"- **File Size:** {format_size(_to_int(report.file_size, 0))}\n"
        md_content += f"- **Strings:** {report.strings_count or 0}\n"
        md_content += f"- **Imports:** {report.imports_count or 0}\n"
        md_content += f"- **Exports:** {report.exports_count or 0}\n"
        md_content += f"- **Secrets:** {report.secrets_count or 0}\n"
        md_content += f"- **Packed:** {report.is_packed or 'unknown'} {f'({report.packer_name})' if report.packer_name else ''}\n\n"

        suspicious_indicators = _dict_items(getattr(report, "suspicious_indicators", None))
        if suspicious_indicators:
            md_content += "### Suspicious Indicators\n\n"
            for indicator in suspicious_indicators[:25]:
                md_content += f"- **{indicator.get('severity', 'info').upper()}** {indicator.get('category', 'Unknown')}: {indicator.get('description', '')}\n"
            if len(suspicious_indicators) > 25:
                md_content += f"\n*...and {len(suspicious_indicators) - 25} more indicators*\n"
            md_content += "\n"

        ctx = report.full_analysis_data or {}
        imports = _dict_items(ctx.get("imports", []) if isinstance(ctx, dict) else [])
        suspicious_imports = [i for i in imports if i.get("is_suspicious")]
        if suspicious_imports:
            md_content += "### Suspicious Imports\n\n"
            for imp in suspicious_imports[:30]:
                md_content += f"- `{imp.get('name')}` ({imp.get('library')}): {imp.get('reason', 'Suspicious API')}\n"
            if len(suspicious_imports) > 30:
                md_content += f"\n*...and {len(suspicious_imports) - 30} more suspicious imports*\n"
            md_content += "\n"

        secrets = _dict_items(ctx.get("secrets", []) if isinstance(ctx, dict) else [])
        if secrets:
            md_content += "### Potential Secrets\n\n"
            include_raw_secret_values = str(getattr(report, "analysis_type", "")).lower() == "docker"
            if include_raw_secret_values:
                md_content += "_Operator note: raw secret values are intentionally included in this Docker report for remediation workflows._\n\n"
            for s in secrets[:20]:
                secret_label = s.get("type") or s.get("secret_type") or "Unknown"
                secret_value = s.get("value") if include_raw_secret_values else s.get("masked_value")
                display_value = secret_value or s.get("masked_value") or "N/A"
                md_content += f"- **{s.get('severity', 'medium').upper()}** {secret_label}: `{display_value}`\n"
                if include_raw_secret_values:
                    if s.get("context"):
                        md_content += f"  Context: {s.get('context')}\n"
                    if s.get("layer_command"):
                        md_content += f"  Layer Command: {s.get('layer_command')}\n"
            if len(secrets) > 20:
                md_content += f"\n*...and {len(secrets) - 20} more secrets*\n"
            md_content += "\n"

        ghidra = _coerce_dict(ctx.get("ghidra_analysis", {}) if isinstance(ctx, dict) else {})
        ghidra_ai = _dict_items(ctx.get("ghidra_ai_summaries", []) if isinstance(ctx, dict) else [])
        if ghidra:
            md_content += "### Ghidra Decompilation\n\n"
            if ghidra.get("error"):
                md_content += f"**Ghidra Error:** {ghidra.get('error')}\n\n"
            else:
                program = ghidra.get("program", {})
                md_content += "**Program Metadata**\n\n"
                md_content += f"- **Name:** `{program.get('name', 'N/A')}`\n"
                md_content += f"- **Processor:** `{program.get('processor', 'N/A')}`\n"
                md_content += f"- **Language ID:** `{program.get('language_id', 'N/A')}`\n"
                md_content += f"- **Compiler Spec:** `{program.get('compiler_spec', 'N/A')}`\n"
                md_content += f"- **Image Base:** `{program.get('image_base', 'N/A')}`\n\n"

                functions = _dict_items(ghidra.get("functions", []) or [])
                total_functions = ghidra.get("functions_total", len(functions))
                max_export = 20
                md_content += f"**Decompiled Functions (showing {min(len(functions), max_export)} of {total_functions})**\n\n"
                md_content += "*Raw decompiled source is omitted in exported reports for readability. Use the in-app decompiler view for full code.*\n\n"

                summary_map = {}
                for summary in ghidra_ai or []:
                    key = f"{summary.get('name')}:{summary.get('entry')}"
                    summary_map[key] = summary

                for fn in functions[:max_export]:
                    fn_name = fn.get("name", "unknown")
                    fn_entry = fn.get("entry", "0x0")
                    fn_size = fn.get("size", 0)
                    md_content += f"#### {fn_name} ({fn_entry})\n\n"
                    md_content += f"- **Size:** {fn_size} bytes\n"
                    if fn.get("called_functions"):
                        called = fn.get("called_functions", [])
                        md_content += f"- **Calls:** {', '.join(called[:12])}"
                        if len(called) > 12:
                            md_content += f" (+{len(called) - 12} more)"
                        md_content += "\n"

                    summary_key = f"{fn_name}:{fn_entry}"
                    summary = summary_map.get(summary_key)
                    if summary and summary.get("summary"):
                        md_content += "\n**Gemini Summary**\n\n"
                        md_content += f"{summary.get('summary')}\n\n"

                    # Intentionally omit raw decompiled code from exports (summary-only report).
                    if fn.get("decompiled") and not (summary and summary.get("summary")):
                        md_content += "- Decompiled source available in app view.\n\n"

                if total_functions > max_export:
                    md_content += f"*...and {total_functions - max_export} more functions*\n\n"
    
    # =====================================================
    # SECTION 3: Architecture Diagram
    # =====================================================
    if architecture_diagram:
        md_content += f"""---

## Application Architecture

The following diagram illustrates the high-level architecture and component relationships within the application.

```mermaid
{architecture_diagram}
```

"""
    
    # =====================================================
    # SECTION 4: AI Cross-Class Vulnerability Analysis
    # =====================================================
    vuln_data = ai_vuln_dict or verification_dict
    if vuln_data:
        
        # Check if we have enhanced security data embedded
        enhanced_security = _coerce_dict(vuln_data.get('enhanced_security'))
        
        if enhanced_security:
            risk_summary = _coerce_dict(enhanced_security.get('risk_summary', {}))
            analysis_metadata = _coerce_dict(enhanced_security.get('analysis_metadata', {}))
            # Use enhanced security data for the main analysis section
            md_content += f"""---

## Comprehensive Security Analysis

**Analysis Sources:** Pattern Detection + AI Analysis + CVE Lookup

**Overall Risk Level:** {enhanced_security.get('overall_risk', 'N/A').upper()}

### Executive Summary

{enhanced_security.get('executive_summary', 'No summary available.')}

### Risk Distribution

| Severity | Count |
|----------|-------|
| Critical | {risk_summary.get('critical', 0)} |
| High | {risk_summary.get('high', 0)} |
| Medium | {risk_summary.get('medium', 0)} |
| Low | {risk_summary.get('low', 0)} |
| Info | {risk_summary.get('info', 0)} |

### Analysis Metadata

- **Classes Scanned:** {analysis_metadata.get('classes_scanned', 0)}
- **Libraries Detected:** {analysis_metadata.get('libraries_detected', 0)}
- **CVEs Found:** {analysis_metadata.get('cves_found', 0)}

"""
            # Offensive Plan Summary (AI-generated assessment - this is the main export content)
            offensive_plan = _coerce_dict(enhanced_security.get('offensive_plan_summary'))
            if offensive_plan:
                md_content += """---

## Offensive Security Assessment

"""
                if offensive_plan.get('threat_assessment'):
                    md_content += f"""### Threat Assessment

{offensive_plan.get('threat_assessment')}

"""
                
                if offensive_plan.get('attack_surface_summary'):
                    md_content += f"""### Attack Surface

{offensive_plan.get('attack_surface_summary')}

"""
                
                # Primary Attack Vectors
                attack_vectors = _dict_items(offensive_plan.get('primary_attack_vectors', []))
                if attack_vectors:
                    md_content += "### Primary Attack Vectors\n\n"
                    for i, vector in enumerate(attack_vectors, 1):
                        md_content += f"#### {i}. {vector.get('vector', 'Unknown')}\n\n"
                        md_content += f"**Likelihood:** {vector.get('likelihood', 'Unknown')} | "
                        md_content += f"**Impact:** {vector.get('impact', 'Unknown')}\n\n"
                        md_content += f"{vector.get('description', '')}\n\n"
                        if vector.get('prerequisites'):
                            md_content += f"**Prerequisites:** {vector.get('prerequisites')}\n\n"
                
                # Recommended Test Scenarios
                test_scenarios = offensive_plan.get('recommended_test_scenarios', [])
                if test_scenarios:
                    md_content += "### Recommended Penetration Tests\n\n"
                    for i, scenario in enumerate(test_scenarios, 1):
                        md_content += f"{i}. {scenario}\n"
                    md_content += "\n"
                
                # Priority Targets
                priority_targets = offensive_plan.get('priority_targets', [])
                if priority_targets:
                    md_content += "### Priority Targets\n\n"
                    for target in priority_targets:
                        md_content += f"- {target}\n"
                    md_content += "\n"
                
                # Risk Rating and Confidence
                md_content += f"""### Assessment Summary

- **Risk Rating:** {offensive_plan.get('risk_rating', 'Unknown').upper()}
- **Confidence Level:** {offensive_plan.get('confidence_level', 'Unknown')}

"""
            
            # Recommendations from enhanced security
            recommendations = enhanced_security.get('recommendations', [])
            if recommendations:
                md_content += "### Remediation Recommendations\n\n"
                for rec in recommendations:
                    md_content += f"- {rec}\n"
                md_content += "\n"
            
            # Brief summary of findings (counts only, not individual details)
            combined_findings = _dict_items(enhanced_security.get('combined_findings', []))
            if combined_findings:
                md_content += f"""### Findings Summary

A total of **{len(combined_findings)}** security findings were identified:

"""
                # Group by severity for count summary
                for severity in ['critical', 'high', 'medium', 'low']:
                    severity_findings = [f for f in combined_findings if str(f.get('severity', '')).lower() == severity]
                    if severity_findings:
                        # Get unique titles for this severity
                        unique_titles = list(set(f.get('title', 'Unknown') for f in severity_findings))[:5]
                        md_content += f"- **{severity.upper()}** ({len(severity_findings)}): {', '.join(unique_titles)}"
                        if len(severity_findings) > 5:
                            md_content += f" and {len(severity_findings) - 5} more"
                        md_content += "\n"
                
                md_content += "\n*Note: Individual findings are available in the application for detailed review.*\n\n"
        
        else:
            # Fall back to original AI vuln scan format
            fallback_risk_summary = _coerce_dict(vuln_data.get('risk_summary', {}))
            md_content += f"""---

## AI Deep Vulnerability Analysis

**Classes Analyzed:** {vuln_data.get('classes_scanned', 0)} | **Overall Risk:** {vuln_data.get('overall_risk', 'N/A')}

### Risk Distribution
- Critical: {fallback_risk_summary.get('critical', 0)}
- High: {fallback_risk_summary.get('high', 0)}
- Medium: {fallback_risk_summary.get('medium', 0)}
- Low: {fallback_risk_summary.get('low', 0)}

"""
            if vuln_data.get('vulnerabilities'):
                md_content += "### Key Vulnerabilities\n\n"
                for vuln in _dict_items(vuln_data.get('vulnerabilities', []))[:15]:
                    severity = vuln.get('severity', 'N/A')
                    md_content += f"#### {vuln.get('title', 'Vulnerability')} ({severity.upper()})\n\n"
                    md_content += f"- **Category:** {vuln.get('category', 'N/A')}\n"
                    if vuln.get('cwe_id'):
                        md_content += f"- **CWE:** {vuln.get('cwe_id')}\n"
                    md_content += f"- **Description:** {vuln.get('description', 'N/A')}\n"
                    
                    # Include affected class/method for manual code review
                    affected_class = vuln.get('affected_class', '')
                    affected_method = vuln.get('affected_method', '')
                    if affected_class:
                        md_content += f"- **Affected Class:** `{affected_class}`\n"
                    if affected_method:
                        md_content += f"- **Affected Method:** `{affected_method}`\n"
                    
                    # Include code snippet if available
                    if vuln.get('code_snippet'):
                        snippet = vuln.get('code_snippet', '')[:500]
                        md_content += f"- **Code to Review:**\n```java\n{snippet}\n```\n"
                    
                    if vuln.get('impact'):
                        md_content += f"- **Impact:** {vuln.get('impact')}\n"
                    if vuln.get('remediation'):
                        md_content += f"- **Remediation:** {vuln.get('remediation')}\n"
                    md_content += "\n"
        
        # Attack chains if available (fallback for non-enhanced data)
        if not enhanced_security and vuln_data.get('attack_chains'):
            md_content += "### Attack Chains\n\n"
            for chain in _dict_items(vuln_data.get('attack_chains', []))[:5]:
                md_content += f"#### {chain.get('name', 'Attack Chain')}\n\n"
                md_content += f"**Likelihood:** {chain.get('likelihood', 'Unknown')} | **Impact:** {chain.get('impact', 'Unknown')}\n\n"
                if chain.get('steps'):
                    md_content += "**Steps:**\n"
                    for i, step in enumerate(chain.get('steps', []), 1):
                        md_content += f"{i}. {step}\n"
                md_content += "\n"
        
        # Recommendations for further research (fallback for non-enhanced data)
        if not enhanced_security and vuln_data.get('recommendations'):
            md_content += "### Recommendations for Further Analysis\n\n"
            for rec in vuln_data.get('recommendations', []):
                md_content += f"- {rec}\n"
            md_content += "\n"
    
    # =====================================================
    # SECTION 5: Attack Surface Map
    # =====================================================
    if attack_surface_map:
        md_content += f"""---

## Attack Surface Map

This attack tree visualizes entry points, vulnerabilities, and potential attack paths discovered through AI analysis of the decompiled source code.

```mermaid
{attack_surface_map}
```

"""
    
    # =====================================================
    # SECTION 6: Exposed Secrets
    # =====================================================
    secrets_found = []
    
    # Collect secrets from various sources
    jadx_data_dict = _coerce_dict(getattr(report, "jadx_data", None))
    if jadx_data_dict.get('secrets'):
        for raw_secret in _coerce_list(jadx_data_dict.get('secrets', [])):
            if isinstance(raw_secret, dict):
                secrets_found.append(raw_secret)
            else:
                secrets_found.append({
                    "type": "Secret",
                    "description": str(raw_secret)[:100],
                    "location": "Unknown",
                })
    
    for issue in all_security_findings:
        category_text = str(issue.get('category', '')).lower()
        if 'secret' in category_text or 'key' in category_text or 'credential' in category_text:
            secrets_found.append({
                'type': issue.get('category', 'Secret'),
                'description': issue.get('description', '')[:100],
                'location': issue.get('file', issue.get('location', 'Unknown'))
            })
    
    try:
        report_secrets_count = int(report.secrets_count) if report.secrets_count is not None else 0
    except (TypeError, ValueError):
        report_secrets_count = 0

    if secrets_found or report_secrets_count > 0:
        md_content += f"""---

## Exposed Secrets & Credentials

**Total Secrets Found:** {report_secrets_count or len(secrets_found)}

"""
        if secrets_found:
            md_content += "| Type | Location | Details |\n|------|----------|--------|\n"
            for secret in secrets_found[:20]:
                if isinstance(secret, dict):
                    stype = secret.get('type', secret.get('category', 'Secret'))
                    loc = secret.get('location', secret.get('file', 'Unknown'))[:50]
                    desc = secret.get('description', secret.get('value', ''))[:40]
                    md_content += f"| {stype} | `{loc}` | {desc}... |\n"
                else:
                    md_content += f"| Secret | - | {str(secret)[:50]}... |\n"
            md_content += "\n"
            if len(secrets_found) > 20:
                md_content += f"*...and {len(secrets_found) - 20} more secrets*\n\n"
    
    # =====================================================
    # SECTION 7: Privacy Analysis (if available)
    # =====================================================
    if report.ai_privacy_report:
        md_content += f"""---

## Privacy Analysis

{report.ai_privacy_report}

"""
    
    # =====================================================
    # SECTION 8: Threat Model (if available)
    # =====================================================
    tm = _coerce_dict(getattr(report, "ai_threat_model", None))
    if tm:
        md_content += """---

## Threat Model

"""
        if tm.get('threat_actors'):
            md_content += "### Potential Threat Actors\n\n"
            for actor in _dict_items(tm.get('threat_actors', []))[:5]:
                md_content += f"- **{actor.get('name', 'Unknown')}**: {actor.get('description', '')}\n"
            md_content += "\n"
        
        if tm.get('attack_scenarios'):
            md_content += "### Attack Scenarios\n\n"
            for scenario in _dict_items(tm.get('attack_scenarios', []))[:3]:
                md_content += f"**{scenario.get('name', 'Scenario')}:** {scenario.get('description', '')}\n\n"
    
    # =====================================================
    # SECTION 9: Manifest Visualization (Component Analysis)
    # =====================================================
    mv = _coerce_dict(getattr(report, "manifest_visualization", None))
    if mv:
        md_content += """---

## Manifest Visualization - Component Analysis

"""
        # Summary stats
        if mv.get('exported_count') is not None or mv.get('activities_count') is not None:
            md_content += "### Component Statistics\n\n"
            md_content += "| Component Type | Total | Exported |\n"
            md_content += "|----------------|-------|----------|\n"
            if mv.get('activities_count') is not None:
                md_content += f"| Activities | {mv.get('activities_count', 0)} | - |\n"
            if mv.get('services_count') is not None:
                md_content += f"| Services | {mv.get('services_count', 0)} | - |\n"
            if mv.get('receivers_count') is not None:
                md_content += f"| Receivers | {mv.get('receivers_count', 0)} | - |\n"
            if mv.get('providers_count') is not None:
                md_content += f"| Providers | {mv.get('providers_count', 0)} | - |\n"
            if mv.get('exported_count') is not None:
                md_content += f"| **Total Exported** | **{mv.get('exported_count', 0)}** | - |\n"
            md_content += "\n"
        
        # Deep Link Schemes
        deep_link_schemes = mv.get('deep_link_schemes', [])
        if deep_link_schemes:
            md_content += "### Deep Link Schemes (Attack Vectors)\n\n"
            md_content += "*These schemes can be used to launch components from external apps or web pages:*\n\n"
            for scheme in deep_link_schemes[:15]:
                md_content += f"- `{scheme}`\n"
            if len(deep_link_schemes) > 15:
                md_content += f"\n*...and {len(deep_link_schemes) - 15} more schemes*\n"
            md_content += "\n"
        
        # Security Risks
        security_risks = _dict_items(mv.get('security_risks', []))
        if security_risks:
            md_content += "### Component Security Risks\n\n"
            for risk in security_risks[:10]:
                md_content += f"- **{risk.get('severity', 'Medium').upper()}**: {risk.get('description', 'Unknown risk')}\n"
                if risk.get('component'):
                    md_content += f"  - Component: `{risk.get('component')}`\n"
            if len(security_risks) > 10:
                md_content += f"\n*...and {len(security_risks) - 10} more risks*\n"
            md_content += "\n"
        
        # AI Analysis Summary
        ai_summary = mv.get('ai_analysis_summary')
        if ai_summary:
            md_content += f"""### AI Component Analysis

{ai_summary}

"""
        
        # Mermaid Diagram
        mermaid_diagram = mv.get('mermaid_diagram')
        if mermaid_diagram:
            md_content += f"""### Component Interaction Diagram

```mermaid
{mermaid_diagram}
```

"""
    
    # =====================================================
    # SECTION 10: Obfuscation Analysis
    # =====================================================
    oa = _coerce_dict(getattr(report, "obfuscation_analysis", None))
    if oa:
        md_content += """---

## Obfuscation Analysis

"""
        # Overall Assessment
        level = oa.get('overall_obfuscation_level', 'Unknown')
        score = oa.get('obfuscation_score', 0)
        difficulty = oa.get('reverse_engineering_difficulty', 'Unknown')
        
        md_content += f"""### Assessment Summary

| Metric | Value |
|--------|-------|
| Obfuscation Level | **{level.upper()}** |
| Obfuscation Score | {score}/100 |
| Reverse Engineering Difficulty | {difficulty} |

"""
        
        # Class Naming Analysis
        class_naming = _coerce_dict(oa.get('class_naming', {}))
        if class_naming:
            md_content += "### Class Naming Analysis\n\n"
            total_classes = _to_int(class_naming.get('total_classes', 0))
            obfuscated_count = _to_int(class_naming.get('obfuscated_count', 0))
            readable_count = _to_int(class_naming.get('readable_count', 0))
            short_name_count = _to_int(class_naming.get('short_name_count', 0))
            obf_ratio = _to_float(class_naming.get('obfuscation_ratio', 0))
            md_content += f"- **Total Classes:** {total_classes:,}\n"
            md_content += f"- **Obfuscated Classes:** {obfuscated_count:,} ({obf_ratio:.1%})\n"
            md_content += f"- **Readable Classes:** {readable_count:,}\n"
            md_content += f"- **Short Names (a, b, c...):** {short_name_count:,}\n\n"
        
        # Detected Obfuscation Tools
        detected_tools = _dict_items(oa.get('detected_tools', []))
        if detected_tools:
            md_content += "### Detected Obfuscation Tools\n\n"
            for tool in detected_tools:
                md_content += f"- **{tool.get('name', 'Unknown')}**"
                if tool.get('confidence'):
                    md_content += f" (Confidence: {tool.get('confidence')})"
                md_content += "\n"
                if tool.get('indicators'):
                    for indicator in tool.get('indicators', [])[:3]:
                        md_content += f"  - {indicator}\n"
            md_content += "\n"
        
        # Obfuscation Indicators
        indicators = _dict_items(oa.get('indicators', []))
        if indicators:
            md_content += "### Obfuscation Indicators\n\n"
            high_confidence = [i for i in indicators if str(i.get('confidence', '')).lower() == 'high']
            for ind in high_confidence[:8]:
                md_content += f"- **{ind.get('indicator_type', 'Unknown')}**: {ind.get('description', '')}\n"
                if ind.get('evidence'):
                    md_content += f"  - Evidence: `{ind.get('evidence')[:80]}...`\n"
            if len(indicators) > 8:
                md_content += f"\n*...and {len(indicators) - 8} more indicators*\n"
            md_content += "\n"
        
        # Deobfuscation Strategies
        strategies = oa.get('deobfuscation_strategies', [])
        if strategies:
            md_content += "### Recommended Deobfuscation Strategies\n\n"
            for i, strategy in enumerate(strategies[:5], 1):
                md_content += f"{i}. {strategy}\n"
            md_content += "\n"
        
        # Recommended Tools
        recommended_tools = _dict_items(oa.get('recommended_tools', []))
        if recommended_tools:
            md_content += "### Recommended Deobfuscation Tools\n\n"
            for tool in recommended_tools[:5]:
                md_content += f"- **{tool.get('name', 'Unknown')}**: {tool.get('purpose', '')}\n"
            md_content += "\n"
        
        # AI Analysis Summary
        ai_obf_summary = oa.get('ai_analysis_summary')
        if ai_obf_summary:
            md_content += f"""### AI Obfuscation Analysis

{ai_obf_summary}

"""
        
        # Frida Hooks for Deobfuscation
        frida_hooks = _dict_items(oa.get('frida_hooks', []))
        if frida_hooks:
            md_content += "### Frida Hooks for Runtime Deobfuscation\n\n"
            md_content += "*Use these hooks to observe decrypted strings and deobfuscated values at runtime:*\n\n"
            for hook in frida_hooks[:3]:
                hook_name = hook.get('name', 'Hook')
                hook_code = hook.get('code', '')
                md_content += f"#### {hook_name}\n\n"
                md_content += f"```javascript\n{hook_code[:500]}\n```\n\n"
            if len(frida_hooks) > 3:
                md_content += f"*...and {len(frida_hooks) - 3} more hooks available in the application*\n\n"
    
    # =====================================================
    # SECTION 11: AI Finding Verification Results
    # =====================================================
    if verification_dict:
        vr = verification_dict
        stats = _coerce_dict(vr.get('verification_stats', {}))
        
        md_content += """---

## AI Finding Verification

*This section summarizes the AI verification pass that validates findings and filters false positives.*

"""
        # Verification Statistics
        md_content += f"""### Verification Summary

| Metric | Value |
|--------|-------|
| Total Findings Analyzed | {stats.get('total_input', 0)} |
| Verified Findings | {stats.get('verified', 0)} |
| False Positives Filtered | {stats.get('filtered', 0)} |
| Filter Rate | {_to_float(stats.get('filter_rate', 0)):.1f}% |
| Average Confidence | {_to_float(stats.get('avg_confidence', 0)):.0f}% |
| High Confidence (≥70%) | {stats.get('high_confidence_count', 0)} |

"""
        # Verdict Breakdown
        by_verdict = _coerce_dict(stats.get('by_verdict', {}))
        if by_verdict:
            md_content += """### Verdict Breakdown

| Verdict | Count | Description |
|---------|-------|-------------|
"""
            md_content += f"| CONFIRMED | {by_verdict.get('CONFIRMED', 0)} | Real vulnerability with clear exploitation path |\n"
            md_content += f"| LIKELY | {by_verdict.get('LIKELY', 0)} | Probably real, needs dynamic testing |\n"
            md_content += f"| SUSPICIOUS | {by_verdict.get('SUSPICIOUS', 0)} | Could be real, context unclear |\n"
            md_content += f"| UNVERIFIED | {by_verdict.get('UNVERIFIED', 0)} | AI verification unavailable |\n\n"
        
        # Attack Chains
        attack_chains = _dict_items(vr.get('attack_chains', []))
        if attack_chains:
            md_content += f"""### Attack Chains Detected ({len(attack_chains)})

*These are correlated finding patterns that indicate complete attack paths:*

"""
            for chain in attack_chains[:5]:
                chain_name = chain.get('chain_name', 'Unknown Chain')
                risk_level = chain.get('risk_level', 'unknown').upper()
                md_content += f"#### {chain_name} ({risk_level})\n\n"
                md_content += f"{chain.get('description', '')}\n\n"
                
                # Entry points
                entry_points = _dict_items(chain.get('entry_points', []))
                if entry_points:
                    md_content += "**Entry Points:**\n"
                    for ep in entry_points[:3]:
                        md_content += f"- {ep.get('title', 'Unknown')} (`{ep.get('class', '')}:{ep.get('line', '')}`)\n"
                    md_content += "\n"
                
                # Sinks
                sinks = _dict_items(chain.get('sinks', []))
                if sinks:
                    md_content += "**Vulnerable Sinks:**\n"
                    for sink in sinks[:3]:
                        md_content += f"- {sink.get('title', 'Unknown')} (`{sink.get('class', '')}:{sink.get('line', '')}`)\n"
                    md_content += "\n"
            
            if len(attack_chains) > 5:
                md_content += f"*...and {len(attack_chains) - 5} more attack chains detected*\n\n"
        
        # High confidence findings (top 10)
        verified_findings = _dict_items(vr.get('verified_findings', []))
        high_conf = [
            f for f in verified_findings
            if _coerce_dict(f.get('verification', {})).get('confidence', 0) >= 80
        ]
        if high_conf:
            md_content += f"""### High Confidence Findings (≥80%)

*These findings have been verified with high confidence by AI analysis:*

"""
            for f in high_conf[:10]:
                ver = _coerce_dict(f.get('verification', {}))
                title = f.get('title', f.get('type', 'Unknown'))
                severity = f.get('severity', 'medium').upper()
                confidence = ver.get('confidence', 0)
                verdict = ver.get('verdict', 'UNVERIFIED')
                
                md_content += f"- **{title}** ({severity}) - {confidence}% confidence [{verdict}]\n"
                if ver.get('reasoning'):
                    md_content += f"  - {ver.get('reasoning')[:100]}\n"
            
            if len(high_conf) > 10:
                md_content += f"\n*...and {len(high_conf) - 10} more high-confidence findings*\n"
            md_content += "\n"
    
    # =====================================================
    # SECTION 12: Sensitive Data Discovery
    # =====================================================
    sd = _coerce_dict(getattr(report, "sensitive_data_findings", None))
    if sd and sd.get('findings'):
        md_content += """---

## 12. Sensitive Data Discovery

*AI-verified scan for hardcoded credentials, API keys, PII, and other sensitive data.*

"""
        findings = _dict_items(sd.get('findings', []))
        summary = _coerce_dict(sd.get('summary', {}))
        scan_stats = _coerce_dict(sd.get('scan_stats', {}))
        
        # Summary table
        by_category = _coerce_dict(summary.get('by_category', {}))
        by_risk = _coerce_dict(summary.get('by_risk', {}))
        
        md_content += f"""### Discovery Summary

| Metric | Value |
|--------|-------|
| **Total Findings** | {summary.get('total', 0)} |
| **Files Scanned** | {_to_int(scan_stats.get('files_scanned', 0)):,} |
| **Raw Pattern Matches** | {_to_int(scan_stats.get('raw_matches', 0)):,} |
| **AI-Verified** | {scan_stats.get('verified', 0)} |
| **False Positives Filtered** | {scan_stats.get('filtered', 0)} |
| **High Confidence (≥80%)** | {summary.get('high_confidence_count', 0)} |

"""
        
        # By category breakdown
        if by_category:
            md_content += """### Findings by Category

| Category | Count |
|----------|-------|
"""
            for cat, count in sorted(by_category.items(), key=lambda x: -_to_int(x[1], 0)):
                cat_display = cat.replace('_', ' ').title()
                md_content += f"| {cat_display} | {_to_int(count, 0)} |\n"
            md_content += "\n"
        
        # By risk level
        if by_risk:
            md_content += """### Findings by Risk Level

| Risk | Count |
|------|-------|
"""
            for risk in ['critical', 'high', 'medium', 'low']:
                risk_count = _to_int(by_risk.get(risk, 0), 0)
                if risk_count > 0:
                    md_content += f"| {risk.upper()} | {risk_count} |\n"
            md_content += "\n"
        
        # Detailed findings by category
        md_content += """### Detailed Findings

"""
        # Group findings by category
        by_cat_findings = {}
        for f in findings:
            cat = f.get('category', 'unknown')
            if cat not in by_cat_findings:
                by_cat_findings[cat] = []
            by_cat_findings[cat].append(f)
        
        for cat, cat_findings in sorted(by_cat_findings.items()):
            cat_display = cat.replace('_', ' ').title()
            md_content += f"""#### {cat_display} ({len(cat_findings)} found)

"""
            for f in cat_findings[:10]:  # Limit per category
                ver = _coerce_dict(f.get('ai_verification', {}))
                confidence = ver.get('confidence', 0)
                risk = ver.get('risk_level', 'medium').upper()
                masked = f.get('masked_value', '****')
                file_path = f.get('file_path', 'Unknown')
                line = f.get('line', 0)
                
                md_content += f"""**{masked}** [{risk}] - {confidence}% confidence
- File: `{file_path}` (line {line})
"""
                if ver.get('reasoning'):
                    md_content += f"- AI Note: {ver.get('reasoning')[:100]}\n"
                md_content += "\n"
            
            if len(cat_findings) > 10:
                md_content += f"*...and {len(cat_findings) - 10} more {cat_display.lower()} findings*\n\n"
    
    elif sd and _to_int(_coerce_dict(sd.get('scan_stats', {})).get('files_scanned', 0), 0) > 0:
        # Scanned but found nothing
        md_content += """---

## 12. Sensitive Data Discovery

No sensitive data (passwords, API keys, emails, phone numbers) detected in the decompiled source code.
All potential matches were filtered as false positives or placeholders.

"""
    
    # =====================================================
    # SECTION 13: Appendix - Technical Details
    # =====================================================
    md_content += """---

## Appendix

"""
    
    # Analysis metadata
    if report.analysis_type == 'apk':
        md_content += f"""### Technical Details

| Property | Value |
|----------|-------|
| Package Name | `{report.package_name or 'N/A'}` |
| Version Name | {report.version_name or 'N/A'} |
| Min SDK | {report.min_sdk or 'N/A'} |
| Target SDK | {report.target_sdk or 'N/A'} |

"""
    
    # JADX stats if available
    if _to_int(report.jadx_total_classes, 0) > 0:
        jadx_total_classes = _to_int(report.jadx_total_classes, 0)
        jadx_total_files = _to_int(report.jadx_total_files, 0)
        md_content += f"""### Code Analysis Statistics

| Metric | Value |
|--------|-------|
| Total Classes | {jadx_total_classes:,} |
| Total Files | {jadx_total_files:,} |

"""
    
    # Notes
    if report.notes:
        md_content += f"""### Analyst Notes

{report.notes}

"""
    
    # Tags
    if report.tags:
        tags = report.tags if isinstance(report.tags, list) else [str(report.tags)]
        md_content += f"""### Tags

{', '.join(f'`{tag}`' for tag in tags)}

"""

    md_content += f"""---

*Report generated by VRAgent on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*
"""

    # Return based on format
    if format == "markdown":
        return md_content
    
    elif format == "pdf":
        # Generate properly formatted PDF
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem, HRFlowable
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            import re
            import unicodedata
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
            styles = getSampleStyleSheet()

            def _sanitize_pdf_text(value: str) -> str:
                """Strip glyphs Helvetica can't render (emoji/symbols) to avoid square boxes in PDF."""
                if not value:
                    return ""
                text = str(value)
                # Remove supplementary-plane symbols (most emoji) and misc dingbats.
                text = re.sub(r'[\U00010000-\U0010FFFF]', '', text)
                text = re.sub(r'[\u2600-\u27BF]', '', text)
                # Remove invisible/selector glyphs that often render as empty squares.
                text = re.sub(r'[\uFE00-\uFE0F\u200B-\u200D\u2060]', '', text)
                # Normalize punctuation for better PDF font compatibility.
                text = (
                    text.replace('’', "'")
                    .replace('“', '"')
                    .replace('”', '"')
                    .replace('–', '-')
                    .replace('—', '-')
                    .replace('…', '...')
                )
                # Force portable Latin text so exports render consistently across viewers.
                text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
                text = re.sub(r'[ \t]+', ' ', text)
                return text
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName='Helvetica-Bold',
                fontSize=21,
                leading=25,
                spaceAfter=12,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#0F172A"),
            )
            h1_style = ParagraphStyle(
                'H1',
                parent=styles['Heading1'],
                fontName='Helvetica-Bold',
                fontSize=16,
                leading=20,
                spaceBefore=18,
                spaceAfter=10,
                textColor=colors.HexColor("#111827"),
            )
            h2_style = ParagraphStyle(
                'H2',
                parent=styles['Heading2'],
                fontName='Helvetica-Bold',
                fontSize=14,
                leading=18,
                spaceBefore=14,
                spaceAfter=8,
                textColor=colors.HexColor("#1F2937"),
            )
            h3_style = ParagraphStyle(
                'H3',
                parent=styles['Heading3'],
                fontName='Helvetica-Bold',
                fontSize=12,
                leading=16,
                spaceBefore=12,
                spaceAfter=6,
                textColor=colors.HexColor("#374151"),
            )
            body_style = ParagraphStyle(
                'Body',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=10,
                spaceBefore=4,
                spaceAfter=4,
                leading=14,
                textColor=colors.HexColor("#111827"),
            )
            code_style = ParagraphStyle(
                'Code',
                parent=styles['Code'],
                fontName='Courier',
                fontSize=9,
                backColor=colors.HexColor("#F3F4F6"),
                leftIndent=10,
                rightIndent=10,
                spaceBefore=6,
                spaceAfter=6,
                textColor=colors.HexColor("#111827"),
            )
            bullet_style = ParagraphStyle('Bullet', parent=body_style, leftIndent=16, bulletIndent=8)
            
            story = []
            
            # Title
            story.append(Paragraph(_sanitize_pdf_text(report.title), title_style))
            story.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#D1D5DB"), spaceBefore=2, spaceAfter=10))
            story.append(Spacer(1, 12))
            
            # Parse markdown and convert to PDF elements
            lines = md_content.split('\n')
            i = 0
            in_code_block = False
            code_content = []
            code_language = None
            in_table = False
            table_rows = []
            
            while i < len(lines):
                line = lines[i]
                
                # Handle code blocks
                if line.strip().startswith('```'):
                    if in_code_block:
                        # End code block
                        if code_content:
                            if code_language == 'mermaid':
                                # Render mermaid as image
                                mermaid_code = '\n'.join(code_content)
                                img_bytes = _render_mermaid_to_image(mermaid_code, 'png')
                                if img_bytes:
                                    from reportlab.platypus import Image
                                    from reportlab.lib.utils import ImageReader
                                    img_buffer = BytesIO(img_bytes)
                                    try:
                                        img_reader = ImageReader(img_buffer)
                                        orig_w, orig_h = img_reader.getSize()
                                        if not orig_w or not orig_h:
                                            raise ValueError("Invalid mermaid image dimensions")

                                        # Fit diagram to page while preserving aspect ratio.
                                        max_w = float(doc.width)
                                        max_h = float(doc.height) * 0.70
                                        scale = min(max_w / float(orig_w), max_h / float(orig_h))
                                        # Allow moderate upscaling for small diagrams.
                                        scale = min(scale, 2.0)
                                        draw_w = max(120.0, float(orig_w) * scale)
                                        draw_h = max(90.0, float(orig_h) * scale)

                                        img_buffer.seek(0)
                                        img = Image(img_buffer, width=draw_w, height=draw_h)
                                        img.hAlign = 'CENTER'
                                        story.append(img)
                                        story.append(Spacer(1, 10))
                                    except Exception as img_err:
                                        logger.warning(f"Failed to add mermaid image to PDF: {img_err}")
                                        # Fallback to code
                                        code_text = '<br/>'.join([l.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') for l in code_content])
                                        story.append(Paragraph(code_text, code_style))
                                else:
                                    # Fallback: show as code block
                                    code_text = '<br/>'.join([l.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') for l in code_content])
                                    story.append(Paragraph(code_text, code_style))
                            else:
                                # Regular code block
                                code_text = '<br/>'.join([l.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') for l in code_content])
                                story.append(Paragraph(code_text, code_style))
                            story.append(Spacer(1, 6))
                        code_content = []
                        code_language = None
                        in_code_block = False
                    else:
                        in_code_block = True
                        # Check for language specification
                        lang_match = line.strip()[3:].strip()
                        code_language = lang_match if lang_match else None
                    i += 1
                    continue
                
                if in_code_block:
                    code_content.append(line)
                    i += 1
                    continue
                
                # Handle tables
                if '|' in line and line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                        table_rows = []
                    
                    # Skip separator lines (|---|---|)
                    if re.match(r'^\|[\s\-:]+\|', line):
                        i += 1
                        continue
                    
                    # Parse table row
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells:
                        table_rows.append(cells)
                    i += 1
                    continue
                elif in_table:
                    # End of table
                    if table_rows:
                        # Create table with proper styling
                        col_count = max(len(row) for row in table_rows)
                        # Normalize rows and clean cell content
                        normalized = []
                        for row in table_rows:
                            cleaned_row = []
                            for cell in row:
                                # Strip markdown and escape for PDF
                                clean = _sanitize_pdf_text(_strip_markdown_formatting(cell))
                                clean = clean.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                cleaned_row.append(clean)
                            # Pad row to match column count
                            cleaned_row += [''] * (col_count - len(cleaned_row))
                            normalized.append(cleaned_row)
                        t = Table(normalized, repeatRows=1)
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#E5E7EB")),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor("#111827")),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                            ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ]))
                        story.append(t)
                        story.append(Spacer(1, 10))
                    table_rows = []
                    in_table = False
                
                # Headers
                if line.startswith('# '):
                    text = _sanitize_pdf_text(line[2:])
                    text = _html_to_reportlab(text) if '<' in text else _format_inline_markdown(text)
                    story.append(Paragraph(text, h1_style))
                elif line.startswith('## '):
                    text = _sanitize_pdf_text(line[3:])
                    text = _html_to_reportlab(text) if '<' in text else _format_inline_markdown(text)
                    story.append(Paragraph(text, h1_style))
                elif line.startswith('### '):
                    text = _sanitize_pdf_text(line[4:])
                    text = _html_to_reportlab(text) if '<' in text else _format_inline_markdown(text)
                    story.append(Paragraph(text, h2_style))
                elif line.startswith('#### '):
                    text = _sanitize_pdf_text(line[5:])
                    text = _html_to_reportlab(text) if '<' in text else _format_inline_markdown(text)
                    story.append(Paragraph(text, h3_style))
                # Bullet points
                elif line.strip().startswith('- ') or line.strip().startswith('* ') or line.strip().startswith('• '):
                    text = _sanitize_pdf_text(line.strip()[2:])
                    # Check if content has HTML
                    if '<' in text and '>' in text:
                        text = _html_to_reportlab(text)
                    else:
                        text = _format_inline_markdown(text)
                    story.append(Paragraph(f"- {text}", bullet_style))
                # Numbered lists
                elif re.match(r'^\s*\d+\.\s+', line):
                    text = _sanitize_pdf_text(re.sub(r'^\s*\d+\.\s+', '', line))
                    if '<' in text and '>' in text:
                        text = _html_to_reportlab(text)
                    else:
                        text = _format_inline_markdown(text)
                    num = re.match(r'^\s*(\d+)\.', line).group(1)
                    story.append(Paragraph(f"{num}. {text}", bullet_style))
                # Horizontal rule
                elif line.strip() == '---':
                    story.append(Spacer(1, 10))
                # Regular paragraph
                elif line.strip():
                    text = _sanitize_pdf_text(line.strip())
                    # Check if content has HTML tags
                    if '<' in text and '>' in text:
                        text = _html_to_reportlab(text)
                    else:
                        text = _format_inline_markdown(text)
                    story.append(Paragraph(text, body_style))
                # Empty line
                else:
                    story.append(Spacer(1, 6))
                
                i += 1
            
            doc.build(story)
            return buffer.getvalue()
            
        except ImportError as e:
            logger.error(f"PDF generation failed - missing library: {e}")
            # Fallback: return markdown as plain text
            return md_content.encode('utf-8')
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return md_content.encode('utf-8')
    
    elif format == "docx":
        # Generate properly formatted Word document
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import re
            import unicodedata
            
            doc = Document()

            def _sanitize_word_text(value: str) -> str:
                if not value:
                    return ""
                text = str(value)
                text = re.sub(r'[\uFE00-\uFE0F\u200B-\u200D\u2060]', '', text)
                text = (
                    text.replace('’', "'")
                    .replace('“', '"')
                    .replace('”', '"')
                    .replace('–', '-')
                    .replace('—', '-')
                    .replace('…', '...')
                )
                text = unicodedata.normalize('NFKC', text)
                return re.sub(r'[ \t]+', ' ', text).strip()

            def _set_paragraph_shading(paragraph, fill: str = "F3F4F6"):
                p_pr = paragraph._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill)
                p_pr.append(shd)

            def _set_cell_shading(cell, fill: str = "E5E7EB"):
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill)
                tc_pr.append(shd)

            def _add_horizontal_rule():
                p = doc.add_paragraph()
                p_pr = p._p.get_or_add_pPr()
                p_bdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'D1D5DB')
                p_bdr.append(bottom)
                p_pr.append(p_bdr)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(10)

            def _add_code_block(content_lines: List[str]):
                code_text = '\n'.join(content_lines or [])
                code_text = re.sub(r'[\uFE00-\uFE0F\u200B-\u200D\u2060]', '', code_text)
                code_text = unicodedata.normalize('NFKC', code_text).strip('\n')
                if not code_text:
                    return
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(17, 24, 39)
                if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                _set_paragraph_shading(p, "F3F4F6")
             
            # Set document properties
            core_props = doc.core_properties
            core_props.title = _sanitize_word_text(report.title)
            core_props.author = "VRAgent Security Scanner"

            # Base typography
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(10.5)
            normal_style.font.color.rgb = RGBColor(17, 24, 39)
            if normal_style._element.rPr is not None and normal_style._element.rPr.rFonts is not None:
                normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.line_spacing = 1.25

            heading_specs = [
                ('Heading 1', 15, RGBColor(17, 24, 39), 18, 8),
                ('Heading 2', 13, RGBColor(31, 41, 55), 14, 6),
                ('Heading 3', 12, RGBColor(55, 65, 81), 12, 4),
                ('Heading 4', 11, RGBColor(75, 85, 99), 10, 4),
            ]
            for style_name, font_size, color, before, after in heading_specs:
                style = doc.styles[style_name]
                style.font.name = 'Calibri'
                style.font.bold = True
                style.font.size = Pt(font_size)
                style.font.color.rgb = color
                if style._element.rPr is not None and style._element.rPr.rFonts is not None:
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                style.paragraph_format.space_before = Pt(before)
                style.paragraph_format.space_after = Pt(after)

            title_style = doc.styles['Title']
            title_style.font.name = 'Calibri'
            title_style.font.bold = True
            title_style.font.size = Pt(24)
            title_style.font.color.rgb = RGBColor(15, 23, 42)
            if title_style._element.rPr is not None and title_style._element.rPr.rFonts is not None:
                title_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
             
            # Title
            title_para = doc.add_paragraph(_sanitize_word_text(report.title), style='Title')
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle = doc.add_paragraph(
                _sanitize_word_text(
                    f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} by VRAgent"
                )
            )
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if subtitle.runs:
                subtitle.runs[0].font.size = Pt(10)
                subtitle.runs[0].font.color.rgb = RGBColor(75, 85, 99)
                subtitle.runs[0].italic = True
            _add_horizontal_rule()
             
            # Parse markdown and convert to Word elements
            lines = md_content.split('\n')
            i = 0
            in_code_block = False
            code_content = []
            code_language = None
            in_table = False
            table_rows = []
            
            while i < len(lines):
                line = lines[i]
                
                # Handle code blocks
                if line.strip().startswith('```'):
                    if in_code_block:
                        # End code block
                        if code_content:
                            if code_language == 'mermaid':
                                # Render mermaid as image
                                mermaid_code = '\n'.join(code_content)
                                img_bytes = _render_mermaid_to_image(mermaid_code, 'png')
                                if img_bytes:
                                    try:
                                        img_buffer = BytesIO(img_bytes)
                                        doc.add_picture(img_buffer, width=Inches(6.5))
                                        # Center the image
                                        last_para = doc.paragraphs[-1]
                                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        last_para.paragraph_format.space_after = Pt(10)
                                    except Exception as img_err:
                                        logger.warning(f"Failed to add mermaid image to Word: {img_err}")
                                        # Fallback to code
                                        _add_code_block(code_content)
                                else:
                                    # Fallback: show as code block
                                    _add_code_block(code_content)
                            else:
                                # Regular code block
                                _add_code_block(code_content)
                        code_content = []
                        code_language = None
                        in_code_block = False
                    else:
                        in_code_block = True
                        # Check for language specification
                        lang_match = line.strip()[3:].strip()
                        code_language = lang_match if lang_match else None
                    i += 1
                    continue
                
                if in_code_block:
                    code_content.append(line)
                    i += 1
                    continue
                
                # Handle tables
                if '|' in line and line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                        table_rows = []
                    
                    # Skip separator lines
                    if re.match(r'^\|[\s\-:]+\|', line):
                        i += 1
                        continue
                    
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells:
                        table_rows.append(cells)
                    i += 1
                    continue
                elif in_table:
                    # End of table - create Word table
                    if table_rows:
                        col_count = max(len(row) for row in table_rows)
                        table = doc.add_table(rows=len(table_rows), cols=col_count)
                        table.style = 'Table Grid'
                        table.autofit = True
                        
                        for row_idx, row_data in enumerate(table_rows):
                            row = table.rows[row_idx]
                            for col_idx, cell_text in enumerate(row_data):
                                if col_idx < col_count:
                                    cell = row.cells[col_idx]
                                    # Clean markdown from cell text
                                    clean_text = _sanitize_word_text(_strip_markdown_formatting(cell_text))
                                    cell.text = clean_text
                                    for para in cell.paragraphs:
                                        para.paragraph_format.space_before = Pt(1)
                                        para.paragraph_format.space_after = Pt(1)
                                        for run in para.runs:
                                            run.font.name = 'Calibri'
                                            run.font.size = Pt(10)
                                            run.font.color.rgb = RGBColor(17, 24, 39)
                                            if row_idx == 0:
                                                run.bold = True
                                        if row_idx == 0:
                                            _set_cell_shading(cell, "E5E7EB")
                                        elif row_idx % 2 == 0:
                                            _set_cell_shading(cell, "F9FAFB")
                         
                        doc.add_paragraph()  # Space after table
                    table_rows = []
                    in_table = False
                 
                # Headers
                if line.startswith('# '):
                    text = _sanitize_word_text(_html_to_plain_text(line[2:]) if '<' in line else _strip_markdown_formatting(line[2:]))
                    doc.add_heading(text, 1)
                elif line.startswith('## '):
                    text = _sanitize_word_text(_html_to_plain_text(line[3:]) if '<' in line else _strip_markdown_formatting(line[3:]))
                    doc.add_heading(text, 2)
                elif line.startswith('### '):
                    text = _sanitize_word_text(_html_to_plain_text(line[4:]) if '<' in line else _strip_markdown_formatting(line[4:]))
                    doc.add_heading(text, 3)
                elif line.startswith('#### '):
                    text = _sanitize_word_text(_html_to_plain_text(line[5:]) if '<' in line else _strip_markdown_formatting(line[5:]))
                    doc.add_heading(text, 4)
                # Bullet points
                elif line.strip().startswith('- ') or line.strip().startswith('* ') or line.strip().startswith('• '):
                    text = _sanitize_word_text(line.strip()[2:])
                    p = doc.add_paragraph(style='List Bullet')
                    _add_formatted_text(p, text)
                    p.paragraph_format.space_after = Pt(3)
                # Numbered lists
                elif re.match(r'^\s*\d+\.\s+', line):
                    text = _sanitize_word_text(re.sub(r'^\s*\d+\.\s+', '', line))
                    p = doc.add_paragraph(style='List Number')
                    _add_formatted_text(p, text)
                    p.paragraph_format.space_after = Pt(3)
                # Horizontal rule
                elif line.strip() == '---':
                    _add_horizontal_rule()
                # Regular paragraph
                elif line.strip():
                    p = doc.add_paragraph()
                    _add_formatted_text(p, _sanitize_word_text(line))
                    p.paragraph_format.space_after = Pt(6)
                 
                i += 1
            
            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
            
        except ImportError as e:
            logger.error(f"DOCX generation failed - missing library: {e}")
            return md_content.encode('utf-8')
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return md_content.encode('utf-8')


def _format_inline_markdown(text: str) -> str:
    """Convert inline markdown to ReportLab XML tags for PDF."""
    import re
    
    # First, escape ampersands that aren't already part of entities
    text = re.sub(r'&(?!amp;|lt;|gt;|quot;|apos;)', '&amp;', text)
    
    # Escape < and > that aren't part of our formatting
    # We'll process markdown first, then escape remaining angle brackets
    
    # Process bold markdown: **text**
    def bold_replace(match):
        content = match.group(1)
        # Escape any angle brackets in the content
        content = content.replace('<', '&lt;').replace('>', '&gt;')
        return f'<b>{content}</b>'
    
    text = re.sub(r'\*\*([^*]+)\*\*', bold_replace, text)
    
    # Process italic markdown: *text* (but not **)
    def italic_replace(match):
        content = match.group(1)
        content = content.replace('<', '&lt;').replace('>', '&gt;')
        return f'<i>{content}</i>'
    
    text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', italic_replace, text)
    
    # Process inline code: `text`
    def code_replace(match):
        content = match.group(1)
        content = content.replace('<', '&lt;').replace('>', '&gt;')
        return f'<font face="Courier" size="9">{content}</font>'
    
    text = re.sub(r'`([^`]+)`', code_replace, text)
    
    # Now escape any remaining < and > that weren't part of our tags
    # Split by our known tags and escape content between them
    parts = re.split(r'(</?b>|</?i>|<font[^>]*>|</font>)', text)
    result = []
    for part in parts:
        if re.match(r'^</?b>$|^</?i>$|^<font[^>]*>$|^</font>$', part):
            # This is a tag, keep it
            result.append(part)
        else:
            # This is content, escape any remaining angle brackets
            part = part.replace('<', '&lt;').replace('>', '&gt;')
            result.append(part)
    
    return ''.join(result)


def _html_to_reportlab(text: str) -> str:
    """Convert HTML tags to ReportLab-compatible XML tags."""
    import re
    
    if not text:
        return ''
    
    # First escape ampersands
    text = re.sub(r'&(?!amp;|lt;|gt;|quot;|nbsp;)', '&amp;', text)
    
    # Convert HTML tags to ReportLab equivalents
    # Bold: <strong> -> <b>
    text = re.sub(r'<strong>([^<]*)</strong>', r'<b>\1</b>', text)
    text = re.sub(r'<b>([^<]*)</b>', r'<b>\1</b>', text)
    
    # Italic: <em> -> <i>
    text = re.sub(r'<em>([^<]*)</em>', r'<i>\1</i>', text)
    
    # Code: <code> -> monospace font
    text = re.sub(r'<code>([^<]*)</code>', r'<font face="Courier" size="9">\1</font>', text)
    
    # Remove unsupported tags but keep content
    text = re.sub(r'</?p>', '', text)
    text = re.sub(r'</?div>', '', text)
    text = re.sub(r'</?span[^>]*>', '', text)
    text = re.sub(r'<br\s*/?>', '<br/>', text)
    
    # Convert headers to bold
    text = re.sub(r'<h[1-6][^>]*>([^<]*)</h[1-6]>', r'<b>\1</b>', text)
    
    # Handle lists - convert to bullet points
    text = re.sub(r'<ul[^>]*>', '', text)
    text = re.sub(r'</ul>', '', text)
    text = re.sub(r'<ol[^>]*>', '', text)
    text = re.sub(r'</ol>', '', text)
    text = re.sub(r'<li[^>]*>([^<]*)</li>', r'• \1<br/>', text)
    text = re.sub(r'<li[^>]*>', '• ', text)
    text = re.sub(r'</li>', '<br/>', text)
    
    # Remove any remaining HTML tags we don't support
    text = re.sub(r'<(?!/?b>|/?i>|/?u>|font[^>]*>|/font>|br/>)[^>]+>', '', text)
    
    return text.strip()


def _html_to_plain_text(text: str) -> str:
    """Convert HTML to plain text for Word documents."""
    import re
    from html import unescape
    
    if not text:
        return ''
    
    # Decode HTML entities
    text = unescape(text)
    
    # Replace <br> with newlines
    text = re.sub(r'<br\s*/?>', '\n', text)
    
    # Replace </p> and </div> with newlines
    text = re.sub(r'</p>', '\n', text)
    text = re.sub(r'</div>', '\n', text)
    
    # Replace </li> with newlines
    text = re.sub(r'</li>', '\n', text)
    
    # Add bullet for list items
    text = re.sub(r'<li[^>]*>', '• ', text)
    
    # Remove all remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Clean up multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def _strip_markdown_formatting(text: str) -> str:
    """Remove markdown formatting for plain text."""
    import re
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Also strip HTML if present
    text = _html_to_plain_text(text)
    return text.strip()


def _render_mermaid_to_image(mermaid_code: str, output_format: str = 'png') -> Optional[bytes]:
    """
    Render mermaid diagram code to an image using Kroki API.
    Returns image bytes or None if rendering fails.
    """
    import base64
    import zlib
    import httpx
    
    try:
        # Clean up mermaid code
        mermaid_code = mermaid_code.strip()
        if not mermaid_code:
            return None
        
        # Encode the mermaid code for Kroki
        # Kroki accepts plain text POST or base64 encoded in URL
        encoded = base64.urlsafe_b64encode(
            zlib.compress(mermaid_code.encode('utf-8'), 9)
        ).decode('ascii')
        
        # Use GET with encoded diagram (more reliable).
        # Request 2x scale PNG for readability in exported documents.
        if output_format == "png":
            url = f"https://kroki.io/mermaid/{output_format}/{encoded}?scale=2"
        else:
            url = f"https://kroki.io/mermaid/{output_format}/{encoded}"
        
        with httpx.Client(timeout=30.0) as client:
            response = client.get(url)
            
            if response.status_code == 200:
                return response.content
            else:
                logger.warning(f"Mermaid rendering failed: {response.status_code} - {response.text[:200]}")
                return None
                
    except Exception as e:
        logger.warning(f"Failed to render mermaid diagram: {e}")
        return None


def _add_formatted_text(paragraph, text: str):
    """Add text to a Word paragraph with markdown/HTML formatting converted."""
    import re
    from docx.shared import Pt
    from html import unescape
    
    # First check if text has HTML - convert to plain text
    if '<' in text and '>' in text:
        text = _html_to_plain_text(text)
    
    # Decode any HTML entities
    text = unescape(text)
    
    # Pattern to find bold, italic, and code (markdown style)
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


class UpdateReportRequest(BaseModel):
    """Update notes, tags, or title for a saved report."""
    notes: Optional[str] = None
    tags: Optional[List[str]] = None
    title: Optional[str] = None


@router.delete("/reports/{report_id}")
def delete_report(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    Delete a saved reverse engineering report.
    """
    report = _ensure_reverse_engineering_report_access(
        db,
        _get_reverse_engineering_report_or_404(db, report_id),
        current_user.id,
        require_edit=True,
    )

    db.delete(report)
    db.commit()

    logger.info(f"Deleted RE report {report_id}")

    return {"message": "Report deleted successfully", "id": report_id}


@router.patch("/reports/{report_id}")
def update_report(
    report_id: int,
    update: Optional[UpdateReportRequest] = Body(None),
    notes: Optional[str] = None,
    tags: Optional[List[str]] = None,
    title: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    """
    Update notes, tags, or title of a saved report.
    """
    report = _ensure_reverse_engineering_report_access(
        db,
        _get_reverse_engineering_report_or_404(db, report_id),
        current_user.id,
        require_edit=True,
    )

    payload = update or UpdateReportRequest()

    notes_value = payload.notes if payload.notes is not None else notes
    tags_value = payload.tags if payload.tags is not None else tags
    title_value = payload.title if payload.title is not None else title

    if notes_value is not None:
        report.notes = notes_value
    if tags_value is not None:
        report.tags = tags_value
    if title_value is not None:
        report.title = title_value

    db.commit()
    db.refresh(report)

    return {"message": "Report updated successfully", "id": report_id}


# ============================================================================
# AI-Powered Analysis Endpoints
# ============================================================================

class ApkChatMessage(BaseModel):
    """A message in the APK analysis chat."""
    role: str  # 'user' or 'assistant'
    content: str
    timestamp: Optional[datetime] = None


class ApkChatRequest(BaseModel):
    """Request for APK chat interaction."""
    message: str
    conversation_history: List[ApkChatMessage] = []
    analysis_context: Dict[str, Any]  # The APK analysis result
    beginner_mode: bool = False


class ApkChatResponse(BaseModel):
    """Response from APK chat."""
    response: str
    suggested_questions: List[str] = []
    related_findings: List[str] = []
    learning_tip: Optional[str] = None


class ThreatModelRequest(BaseModel):
    """Request for threat modeling."""
    analysis_context: Dict[str, Any]
    focus_areas: List[str] = []  # e.g., ['data_exfiltration', 'authentication', 'injection']
    attacker_profile: str = "skilled"  # 'script_kiddie', 'skilled', 'nation_state'


class ThreatModelResponse(BaseModel):
    """Threat modeling response."""
    threat_actors: List[Dict[str, Any]]
    attack_scenarios: List[Dict[str, Any]]
    attack_tree: Dict[str, Any]
    mitre_attack_mappings: List[Dict[str, Any]]
    risk_matrix: Dict[str, Any]
    prioritized_threats: List[Dict[str, Any]]
    executive_summary: str


class ExploitSuggestionRequest(BaseModel):
    """Request for exploit suggestions."""
    analysis_context: Dict[str, Any]
    vulnerability_focus: Optional[str] = None  # Specific issue to focus on
    include_poc: bool = True
    skill_level: str = "intermediate"  # 'beginner', 'intermediate', 'advanced'


class ExploitSuggestionResponse(BaseModel):
    """Exploit suggestion response."""
    vulnerabilities: List[Dict[str, Any]]
    exploitation_paths: List[Dict[str, Any]]
    tools_required: List[Dict[str, Any]]
    poc_scripts: List[Dict[str, Any]]
    mitigation_bypasses: List[Dict[str, Any]]
    difficulty_assessment: Dict[str, Any]


class WalkthroughStep(BaseModel):
    """A step in the analysis walkthrough."""
    step_number: int
    phase: str
    title: str
    description: str
    technical_detail: str
    beginner_explanation: str
    why_it_matters: str
    findings_count: int = 0
    severity: Optional[str] = None
    progress_percent: int


class AnalysisWalkthroughResponse(BaseModel):
    """Complete walkthrough of analysis."""
    total_steps: int
    steps: List[WalkthroughStep]
    glossary: Dict[str, str]
    learning_resources: List[Dict[str, str]]
    next_steps: List[str]


@router.post("/apk/chat", response_model=ApkChatResponse)
async def chat_about_apk(
    request: ApkChatRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Interactive AI chat about APK analysis findings.
    
    Supports multi-turn conversations with context about the analyzed APK.
    Can answer questions about permissions, security issues, what they mean,
    and provide recommendations.
    """
    from backend.core.config import settings
    
    if not settings.gemini_api_key:
        raise HTTPException(status_code=503, detail="AI features require Gemini API key")
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # Build context from analysis
        ctx = request.analysis_context
        package_name = ctx.get('package_name', 'Unknown')
        
        # Build conversation history for Gemini
        contents = []
        
        # System context as first user message
        system_context = f"""You are an expert Android security analyst assistant helping users understand APK analysis results.

## APK BEING ANALYZED
- **Package:** {package_name}
- **Version:** {ctx.get('version_name', 'N/A')}
- **Target SDK:** {ctx.get('target_sdk', 'N/A')}
- **Min SDK:** {ctx.get('min_sdk', 'N/A')}
- **Debuggable:** {ctx.get('debuggable', False)}
- **Allow Backup:** {ctx.get('allow_backup', True)}

## PERMISSIONS ({len(ctx.get('permissions', []))})
{chr(10).join(f"- {p.get('name', 'Unknown')}{' [DANGEROUS]' if p.get('is_dangerous') else ''}" for p in ctx.get('permissions', [])[:20])}

## SECURITY ISSUES ({len(ctx.get('security_issues', []))})
{chr(10).join(f"- [{i.get('severity', 'INFO')}] {i.get('category', 'Unknown')}: {i.get('description', '')[:100]}" for i in ctx.get('security_issues', [])[:15])}

## SECRETS FOUND ({len(ctx.get('secrets', []))})
{chr(10).join(f"- {s.get('type', 'Unknown')}: {s.get('masked_value', '***')}" for s in ctx.get('secrets', [])[:10])}

## HARDENING SCORE
{f"Grade: {ctx.get('hardening_score', {}).get('grade', 'N/A')} ({ctx.get('hardening_score', {}).get('overall_score', 'N/A')}/100)" if ctx.get('hardening_score') else "Not calculated"}

## DYNAMIC ANALYSIS
- SSL Pinning Detected: {ctx.get('dynamic_analysis', {}).get('ssl_pinning_detected', False)}
- Root Detection: {ctx.get('dynamic_analysis', {}).get('root_detection_detected', False)}
- Emulator Detection: {ctx.get('dynamic_analysis', {}).get('emulator_detection_detected', False)}

{"## BEGINNER MODE ENABLED - Please explain concepts simply, use analogies, and define technical terms." if request.beginner_mode else ""}

Guidelines:
1. Be helpful and educational
2. Reference specific findings from the analysis
3. Suggest follow-up questions
4. {"Use simple language and analogies for beginners" if request.beginner_mode else "Be technically precise"}
5. Provide actionable recommendations
6. If asked about exploitation, focus on defensive understanding"""
        
        contents.append(types.Content(role="user", parts=[types.Part(text=system_context)]))
        contents.append(types.Content(role="model", parts=[types.Part(text="I understand. I'm ready to help you understand this APK analysis. What would you like to know?")]))
        
        # Add conversation history
        for msg in request.conversation_history[-10:]:  # Last 10 messages
            role = "user" if msg.role == "user" else "model"
            contents.append(types.Content(role=role, parts=[types.Part(text=msg.content)]))
        
        # Add current message
        contents.append(types.Content(role="user", parts=[types.Part(text=request.message)]))
        
        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=contents,
        )
        
        response_text = response.text
        
        # Generate suggested follow-up questions
        suggested_questions = []
        if "permission" in request.message.lower():
            suggested_questions = [
                "What data could this app access with these permissions?",
                "Are there any permission combinations that are concerning?",
                "How do these permissions compare to similar apps?",
            ]
        elif "security" in request.message.lower() or "issue" in request.message.lower():
            suggested_questions = [
                "How could an attacker exploit these issues?",
                "What's the priority order for fixing these?",
                "Are there any quick wins for improving security?",
            ]
        elif "secret" in request.message.lower() or "api key" in request.message.lower():
            suggested_questions = [
                "How can secrets be properly protected in Android apps?",
                "What's the risk if these secrets are exposed?",
                "How can I detect if these keys have been abused?",
            ]
        else:
            suggested_questions = [
                "What are the most critical security issues?",
                "Is this app safe to use?",
                "What would you recommend fixing first?",
            ]
        
        # Generate learning tip for beginner mode
        learning_tip = None
        if request.beginner_mode:
            tips = [
                "💡 Tip: Dangerous permissions don't mean the app is malicious - they just require extra scrutiny.",
                "💡 Tip: A low hardening score doesn't always mean the app is insecure - context matters!",
                "💡 Tip: SSL pinning is a defense mechanism that makes it harder to intercept app traffic.",
                "💡 Tip: Root detection helps apps protect themselves, but can be bypassed for security testing.",
                "💡 Tip: Exported components are entry points that other apps can interact with.",
            ]
            import random
            learning_tip = random.choice(tips)
        
        return ApkChatResponse(
            response=response_text,
            suggested_questions=suggested_questions,
            related_findings=[],
            learning_tip=learning_tip,
        )
        
    except Exception as e:
        logger.error(f"APK chat failed: {e}")
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")


@router.post("/apk/threat-model", response_model=ThreatModelResponse)
async def generate_threat_model(
    request: ThreatModelRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate AI-powered threat model for the APK.
    
    Creates attack scenarios, threat actors, MITRE ATT&CK mappings,
    and prioritized threat assessment.
    """
    from backend.core.config import settings
    
    if not settings.gemini_api_key:
        raise HTTPException(status_code=503, detail="AI features require Gemini API key")
    
    try:
        from google import genai
        from google.genai import types
        import json
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        ctx = request.analysis_context
        package_name = ctx.get('package_name', 'Unknown')
        
        attacker_profiles = {
            "script_kiddie": "Low skill, uses automated tools, opportunistic",
            "skilled": "Moderate skill, can develop custom exploits, targeted",
            "nation_state": "Advanced persistent threat, unlimited resources, sophisticated techniques"
        }
        
        prompt = f"""You are an expert mobile security threat modeler. Generate a comprehensive threat model for this Android application.

## APPLICATION CONTEXT
- **Package:** {package_name}
- **Version:** {ctx.get('version_name', 'N/A')}
- **Target SDK:** {ctx.get('target_sdk', 'N/A')}
- **Debuggable:** {ctx.get('debuggable', False)}

## ATTACK SURFACE
**Permissions ({len(ctx.get('permissions', []))}):**
{chr(10).join(f"- {p.get('name')}" for p in ctx.get('permissions', []) if p.get('is_dangerous'))[:10]}

**Exported Components:**
- Activities: {len([c for c in ctx.get('components', []) if c.get('component_type') == 'activity' and c.get('is_exported')])}
- Services: {len([c for c in ctx.get('components', []) if c.get('component_type') == 'service' and c.get('is_exported')])}
- Receivers: {len([c for c in ctx.get('components', []) if c.get('component_type') == 'receiver' and c.get('is_exported')])}
- Providers: {len([c for c in ctx.get('components', []) if c.get('component_type') == 'provider' and c.get('is_exported')])}

**Deep Links:** {len(ctx.get('intent_filter_analysis', {}).get('deep_links', []))}

**Security Issues ({len(ctx.get('security_issues', []))}):**
{chr(10).join(f"- [{i.get('severity')}] {i.get('description', '')[:80]}" for i in ctx.get('security_issues', [])[:10])}

**Native Libraries:** {len(ctx.get('native_libraries', []))}
**Secrets Found:** {len(ctx.get('secrets', []))}

## ATTACKER PROFILE
**Type:** {request.attacker_profile}
**Description:** {attacker_profiles.get(request.attacker_profile, 'Unknown')}

## FOCUS AREAS
{', '.join(request.focus_areas) if request.focus_areas else 'All attack vectors'}

Generate a JSON response with the following structure:
{{
    "threat_actors": [
        {{
            "name": "Actor name",
            "motivation": "Financial/Espionage/Disruption",
            "capability": "Low/Medium/High",
            "likelihood": "Low/Medium/High",
            "description": "Brief description"
        }}
    ],
    "attack_scenarios": [
        {{
            "id": "AS-001",
            "name": "Scenario name",
            "description": "Detailed attack scenario",
            "preconditions": ["Required conditions"],
            "attack_steps": ["Step 1", "Step 2"],
            "impact": "What damage could occur",
            "likelihood": "Low/Medium/High",
            "severity": "Low/Medium/High/Critical",
            "mitre_techniques": ["T1234"]
        }}
    ],
    "attack_tree": {{
        "goal": "Compromise application",
        "branches": [
            {{
                "method": "Attack vector",
                "sub_branches": ["Sub-attack 1", "Sub-attack 2"],
                "difficulty": "Easy/Medium/Hard"
            }}
        ]
    }},
    "mitre_attack_mappings": [
        {{
            "technique_id": "T1234",
            "technique_name": "Technique Name",
            "tactic": "Initial Access/Execution/etc",
            "relevance": "How it applies to this app",
            "finding_reference": "Which finding relates to this"
        }}
    ],
    "risk_matrix": {{
        "critical_risks": ["List of critical risks"],
        "high_risks": ["List of high risks"],
        "medium_risks": ["List of medium risks"],
        "low_risks": ["List of low risks"],
        "accepted_risks": ["Risks that may be acceptable"]
    }},
    "prioritized_threats": [
        {{
            "rank": 1,
            "threat": "Threat name",
            "risk_score": 85,
            "rationale": "Why this is prioritized",
            "recommendation": "What to do about it"
        }}
    ],
    "executive_summary": "2-3 paragraph executive summary of the threat landscape"
}}

Be thorough and realistic. Consider the specific findings from this APK."""

        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        # Parse JSON from response
        response_text = response.text
        
        # Extract JSON from markdown code blocks if present
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        elif "```" in response_text:
            json_start = response_text.find("```") + 3
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        
        try:
            threat_data = json.loads(response_text)
        except json.JSONDecodeError:
            # Fallback structure if JSON parsing fails
            threat_data = {
                "threat_actors": [{"name": "Unknown", "motivation": "Various", "capability": "Medium", "likelihood": "Medium", "description": response_text[:200]}],
                "attack_scenarios": [],
                "attack_tree": {"goal": "Compromise application", "branches": []},
                "mitre_attack_mappings": [],
                "risk_matrix": {"critical_risks": [], "high_risks": [], "medium_risks": [], "low_risks": [], "accepted_risks": []},
                "prioritized_threats": [],
                "executive_summary": response_text[:500]
            }
        
        return ThreatModelResponse(
            threat_actors=threat_data.get("threat_actors", []),
            attack_scenarios=threat_data.get("attack_scenarios", []),
            attack_tree=threat_data.get("attack_tree", {}),
            mitre_attack_mappings=threat_data.get("mitre_attack_mappings", []),
            risk_matrix=threat_data.get("risk_matrix", {}),
            prioritized_threats=threat_data.get("prioritized_threats", []),
            executive_summary=threat_data.get("executive_summary", ""),
        )
        
    except Exception as e:
        logger.error(f"Threat modeling failed: {e}")
        raise HTTPException(status_code=500, detail=f"Threat modeling failed: {str(e)}")


@router.post("/apk/exploit-suggestions", response_model=ExploitSuggestionResponse)
async def get_exploit_suggestions(
    request: ExploitSuggestionRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate AI-powered exploit suggestions for identified vulnerabilities.
    
    Provides exploitation paths, required tools, PoC scripts, and difficulty assessments.
    For educational/defensive purposes only.
    """
    from backend.core.config import settings
    
    if not settings.gemini_api_key:
        raise HTTPException(status_code=503, detail="AI features require Gemini API key")
    
    try:
        from google import genai
        from google.genai import types
        import json
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        ctx = request.analysis_context
        package_name = ctx.get('package_name', 'Unknown')
        
        skill_descriptions = {
            "beginner": "Basic Android knowledge, can follow step-by-step guides",
            "intermediate": "Familiar with Android internals, can modify existing tools",
            "advanced": "Expert level, can develop custom exploits and bypass protections"
        }
        
        prompt = f"""You are a mobile security penetration tester. Generate exploitation guidance for DEFENSIVE and EDUCATIONAL purposes to help developers understand and fix vulnerabilities.

## APPLICATION
- **Package:** {package_name}
- **Version:** {ctx.get('version_name', 'N/A')}
- **Debuggable:** {ctx.get('debuggable', False)}
- **Allow Backup:** {ctx.get('allow_backup', True)}

## IDENTIFIED VULNERABILITIES
{chr(10).join(f"- [{i.get('severity', 'INFO')}] {i.get('category', 'Unknown')}: {i.get('description', '')}" for i in ctx.get('security_issues', [])[:15])}

## SECRETS FOUND
{chr(10).join(f"- {s.get('type', 'Unknown')}: {s.get('masked_value', '***')}" for s in ctx.get('secrets', [])[:5])}

## ATTACK SURFACE
- Exported Activities: {len([c for c in ctx.get('components', []) if c.get('component_type') == 'activity' and c.get('is_exported')])}
- Deep Links: {ctx.get('intent_filter_analysis', {}).get('attack_surface_summary', {}).get('total_deep_links', 0)}
- Native Libraries: {len(ctx.get('native_libraries', []))}

## PROTECTIONS DETECTED
- SSL Pinning: {ctx.get('dynamic_analysis', {}).get('ssl_pinning_detected', False)}
- Root Detection: {ctx.get('dynamic_analysis', {}).get('root_detection_detected', False)}
- Emulator Detection: {ctx.get('dynamic_analysis', {}).get('emulator_detection_detected', False)}
- Anti-Tampering: {ctx.get('dynamic_analysis', {}).get('anti_tampering_detected', False)}

## TESTER SKILL LEVEL
**Level:** {request.skill_level}
**Description:** {skill_descriptions.get(request.skill_level, 'Unknown')}

{f"## FOCUS ON: {request.vulnerability_focus}" if request.vulnerability_focus else ""}

Generate a JSON response with DEFENSIVE exploitation guidance:
{{
    "vulnerabilities": [
        {{
            "id": "VULN-001",
            "name": "Vulnerability name",
            "category": "OWASP Mobile category",
            "severity": "Critical/High/Medium/Low",
            "description": "What the vulnerability is",
            "root_cause": "Why this vulnerability exists",
            "affected_component": "Which part of the app"
        }}
    ],
    "exploitation_paths": [
        {{
            "vulnerability_id": "VULN-001",
            "name": "Exploitation path name",
            "prerequisites": ["What's needed before exploitation"],
            "steps": [
                {{
                    "step": 1,
                    "action": "What to do",
                    "command": "adb shell command if applicable",
                    "expected_result": "What should happen"
                }}
            ],
            "success_indicators": ["How to know if it worked"],
            "impact": "What an attacker could achieve"
        }}
    ],
    "tools_required": [
        {{
            "name": "Tool name",
            "purpose": "What it's used for",
            "installation": "How to install it",
            "usage_example": "Example command"
        }}
    ],
    "poc_scripts": [
        {{
            "vulnerability_id": "VULN-001",
            "name": "PoC name",
            "language": "python/bash/frida",
            "description": "What the script does",
            "code": "The actual code",
            "usage": "How to run it"
        }}
    ],
    "mitigation_bypasses": [
        {{
            "protection": "SSL Pinning/Root Detection/etc",
            "bypass_method": "How to bypass",
            "tools": ["Required tools"],
            "difficulty": "Easy/Medium/Hard",
            "detection_risk": "How likely to be detected"
        }}
    ],
    "difficulty_assessment": {{
        "overall_difficulty": "Easy/Medium/Hard/Expert",
        "time_estimate": "Hours/days estimate",
        "skill_requirements": ["Required skills"],
        "resource_requirements": ["Required resources"],
        "success_probability": "High/Medium/Low"
    }}
}}

IMPORTANT: This is for DEFENSIVE purposes - helping developers understand how attackers think so they can build better defenses. Include remediation guidance."""

        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        # Parse JSON from response
        response_text = response.text
        
        # Extract JSON from markdown code blocks if present
        if "```json" in response_text:
            json_start = response_text.find("```json") + 7
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        elif "```" in response_text:
            json_start = response_text.find("```") + 3
            json_end = response_text.find("```", json_start)
            response_text = response_text[json_start:json_end].strip()
        
        try:
            exploit_data = json.loads(response_text)
        except json.JSONDecodeError:
            exploit_data = {
                "vulnerabilities": [],
                "exploitation_paths": [],
                "tools_required": [],
                "poc_scripts": [],
                "mitigation_bypasses": [],
                "difficulty_assessment": {
                    "overall_difficulty": "Unknown",
                    "time_estimate": "Unknown",
                    "skill_requirements": [],
                    "resource_requirements": [],
                    "success_probability": "Unknown"
                }
            }
        
        return ExploitSuggestionResponse(
            vulnerabilities=exploit_data.get("vulnerabilities", []),
            exploitation_paths=exploit_data.get("exploitation_paths", []),
            tools_required=exploit_data.get("tools_required", []),
            poc_scripts=exploit_data.get("poc_scripts", []) if request.include_poc else [],
            mitigation_bypasses=exploit_data.get("mitigation_bypasses", []),
            difficulty_assessment=exploit_data.get("difficulty_assessment", {}),
        )
        
    except Exception as e:
        logger.error(f"Exploit suggestions failed: {e}")
        raise HTTPException(status_code=500, detail=f"Exploit suggestions failed: {str(e)}")


@router.post("/apk/walkthrough", response_model=AnalysisWalkthroughResponse)
async def generate_walkthrough(
    analysis_context: Dict[str, Any],
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate a beginner-friendly walkthrough of the APK analysis.
    
    Provides step-by-step explanations of each analysis phase,
    what was found, and why it matters for security.
    """
    ctx = analysis_context
    
    # Security Glossary - expanded with new terms
    glossary = {
        "APK": "Android Package Kit - the file format used to distribute Android apps",
        "SDK": "Software Development Kit - tools for building Android apps. Min SDK is the oldest Android version supported, Target SDK is what the app is optimized for",
        "Permission": "A declaration that an app needs access to certain device features or data",
        "Dangerous Permission": "Permissions that could affect user privacy or device security, requiring explicit user approval",
        "Exported Component": "An app component (activity, service, etc.) that can be accessed by other apps",
        "Intent": "A messaging object used to request actions from other app components",
        "Deep Link": "A URL that opens a specific screen in an app",
        "SSL Pinning": "A security technique that ensures an app only trusts specific certificates",
        "Root Detection": "Code that checks if the device has been rooted (given superuser access)",
        "Obfuscation": "Making code harder to understand to prevent reverse engineering",
        "Hardcoded Secret": "Sensitive data (like API keys) embedded directly in the app code",
        "Native Library": "Code written in C/C++ compiled for specific processor architectures",
        "DEX": "Dalvik Executable - the compiled bytecode format for Android apps",
        "Smali": "Human-readable representation of DEX bytecode",
        "Frida": "A dynamic instrumentation toolkit for developers, reverse-engineers, and security researchers",
        "OWASP": "Open Web Application Security Project - organization that publishes security guidelines",
        "CVE": "Common Vulnerabilities and Exposures - standardized identifiers for security vulnerabilities",
        "CWE": "Common Weakness Enumeration - a catalog of software security weaknesses",
        "PII": "Personally Identifiable Information - data that can identify an individual (names, emails, phone numbers)",
        "Attack Surface": "All the points where an attacker could try to enter or extract data from an app",
        "Attack Chain": "A sequence of vulnerabilities that can be combined to achieve a larger security breach",
        "False Positive": "A security alert that turns out not to be a real vulnerability",
        "JADX": "A tool that decompiles Android apps back to readable Java source code",
        "SQL Injection": "A code injection technique that exploits database queries to access or modify data",
        "XSS": "Cross-Site Scripting - injecting malicious scripts into trusted websites or apps",
        "Path Traversal": "Accessing files outside the intended directory by manipulating file paths",
    }
    
    # Build walkthrough steps
    steps = []
    step_num = 0
    
    # Step 1: Basic Info
    step_num += 1
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Basic Information",
        title="📱 Extracting App Identity",
        description=f"Analyzed the AndroidManifest.xml to extract basic app information.",
        technical_detail=f"Package: {ctx.get('package_name', 'Unknown')}, Version: {ctx.get('version_name', 'N/A')}, Target SDK: {ctx.get('target_sdk', 'N/A')}, Min SDK: {ctx.get('min_sdk', 'N/A')}",
        beginner_explanation="Every Android app has an identity - its package name (like a unique address), version number, and the Android versions it supports. This is like checking someone's ID card before letting them in.",
        why_it_matters="The target SDK tells us if the app takes advantage of newer security features. Apps targeting older SDKs (below 28) may have weaker security. The min SDK shows what old devices are supported - older Android has more vulnerabilities.",
        findings_count=1,
        severity="info",
        progress_percent=5,
    ))
    
    # Step 2: Permissions
    step_num += 1
    permissions = ctx.get('permissions', [])
    dangerous = [p for p in permissions if p.get('is_dangerous')]
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Permission Analysis",
        title="🔐 Checking What the App Can Access",
        description=f"Found {len(permissions)} permissions, {len(dangerous)} are classified as dangerous.",
        technical_detail=f"Dangerous permissions: {', '.join(p.get('name', '').split('.')[-1] for p in dangerous[:5])}{'...' if len(dangerous) > 5 else ''}",
        beginner_explanation="Permissions are like keys to different parts of your phone. Camera permission lets the app use your camera, location permission lets it know where you are. 'Dangerous' permissions can access sensitive data like contacts, storage, or location.",
        why_it_matters=f"This app requests {len(dangerous)} dangerous permissions. Each one is a potential privacy concern if misused. We check if these make sense for what the app claims to do.",
        findings_count=len(dangerous),
        severity="high" if len(dangerous) > 5 else "medium" if len(dangerous) > 2 else "low",
        progress_percent=10,
    ))
    
    # Step 3: Secrets Detection (early stage)
    step_num += 1
    secrets = ctx.get('secrets', [])
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Secret Detection",
        title="🔑 Finding Hardcoded Secrets",
        description=f"Found {len(secrets)} potential hardcoded secrets in the app resources and code.",
        technical_detail=f"Types found: {', '.join(set(s.get('type', 'Unknown') for s in secrets[:10]))}",
        beginner_explanation="Developers sometimes accidentally leave passwords, API keys, or encryption keys directly in their code. This is like writing your house key on your front door - anyone who downloads the app can find them!",
        why_it_matters="Hardcoded secrets can be extracted by anyone who downloads the app. Attackers could use these to access backend services, steal data, or impersonate the app. This is one of the most common and dangerous mistakes.",
        findings_count=len(secrets),
        severity="critical" if len(secrets) > 3 else "high" if secrets else "low",
        progress_percent=15,
    ))
    
    # Step 4: JADX Decompilation
    step_num += 1
    total_classes = ctx.get('total_classes', 0)
    total_files = ctx.get('total_files', 0)
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="JADX Decompilation",
        title="☕ Decompiling to Java Source",
        description=f"Decompiled the app to {total_files:,} Java source files ({total_classes:,} classes).",
        technical_detail=f"JADX converted DEX bytecode back to readable Java. This allows deep code analysis that isn't possible with just APK inspection.",
        beginner_explanation="Android apps are compiled into a format machines can run (DEX). JADX is like a translator that converts this back to human-readable Java code, letting us examine exactly how the app works.",
        why_it_matters="Decompilation reveals the app's true behavior - not just what it claims to do. We can find hidden functionality, analyze security logic, and discover vulnerabilities that aren't visible from the outside.",
        findings_count=total_files,
        severity="info",
        progress_percent=25,
    ))
    
    # Step 5: Code Security Scan
    step_num += 1
    code_findings = ctx.get('decompiled_code_findings', [])
    code_summary = ctx.get('decompiled_code_summary', {})
    critical_code = len([f for f in code_findings if f.get('severity') == 'critical'])
    high_code = len([f for f in code_findings if f.get('severity') == 'high'])
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Code Security Scan",
        title="🔍 Pattern-Based Vulnerability Detection",
        description=f"Scanned decompiled code and found {len(code_findings)} potential vulnerabilities: {critical_code} critical, {high_code} high severity.",
        technical_detail=f"Categories: {', '.join(list(code_summary.get('by_category', {}).keys())[:6])}",
        beginner_explanation="We scan the Java code looking for dangerous patterns - like SQL queries built from user input (SQL injection), files accessed with user-controlled paths (path traversal), or data sent without encryption.",
        why_it_matters="Code-level vulnerabilities are the root cause of most security breaches. Finding them early prevents attackers from exploiting them to steal data or take control of user accounts.",
        findings_count=len(code_findings),
        severity="critical" if critical_code else "high" if high_code else "medium" if code_findings else "low",
        progress_percent=35,
    ))
    
    # Step 6: Sensitive Data Discovery (NEW!)
    step_num += 1
    sensitive_data = ctx.get('sensitive_data_findings', {})
    sensitive_findings = sensitive_data.get('findings', []) if sensitive_data else []
    sensitive_summary = sensitive_data.get('summary', {}) if sensitive_data else {}
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Sensitive Data Discovery",
        title="🔐 AI-Verified PII & Credentials Search",
        description=f"Found {len(sensitive_findings)} verified instances of sensitive data (passwords, API keys, emails, phone numbers).",
        technical_detail=f"Categories: {', '.join(sensitive_summary.get('by_category', {}).keys()) if sensitive_summary.get('by_category') else 'None found'}",
        beginner_explanation="We search for personal information like passwords, API keys, email addresses, and phone numbers in the code. AI verifies each finding to filter out false positives (things that look like passwords but aren't).",
        why_it_matters="Exposed PII and credentials can be used for identity theft, account takeover, or accessing backend systems. AI verification ensures we report real issues, not false alarms.",
        findings_count=len(sensitive_findings),
        severity="critical" if sensitive_summary.get('by_risk', {}).get('critical', 0) > 0 else "high" if sensitive_findings else "low",
        progress_percent=42,
    ))
    
    # Step 7: CVE Lookup
    step_num += 1
    cve_results = ctx.get('cve_scan_results', {})
    cve_findings = cve_results.get('findings', []) if cve_results else []
    cve_critical = len([c for c in cve_findings if c.get('severity', '').lower() == 'critical'])
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="CVE Database Lookup",
        title="🛡️ Checking Known Vulnerabilities",
        description=f"Checked dependencies against CVE databases. Found {len(cve_findings)} known vulnerabilities ({cve_critical} critical).",
        technical_detail=f"Libraries checked: {cve_results.get('libraries_checked', 0)}. CVEs found: {', '.join(c.get('cve_id', '') for c in cve_findings[:5])}",
        beginner_explanation="CVE is a database of publicly known security vulnerabilities. We check if the app uses any libraries with known security holes - like using a lock that burglars know how to pick.",
        why_it_matters="Known vulnerabilities have published exploits. Attackers actively scan for apps using vulnerable libraries. Updating these libraries is often a quick security win.",
        findings_count=len(cve_findings),
        severity="critical" if cve_critical else "high" if cve_findings else "low",
        progress_percent=50,
    ))
    
    # Step 8: AI Vulnerability Hunt (if enabled)
    step_num += 1
    vuln_hunt = ctx.get('vuln_hunt_results', {})
    ai_findings = vuln_hunt.get('findings', []) if vuln_hunt else []
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="AI Vulnerability Hunt",
        title="🐛 Multi-Pass AI Security Analysis",
        description=f"AI performed deep vulnerability hunting and found {len(ai_findings)} potential security issues.",
        technical_detail=f"AI analyzed code patterns, data flows, and security-sensitive operations across multiple passes to find vulnerabilities that pattern matching might miss.",
        beginner_explanation="Like a human security researcher, AI reads the code and looks for logical flaws - not just known patterns. It follows data through the app to see where user input could be dangerous.",
        why_it_matters="AI finds vulnerabilities that automated scanners miss - like business logic flaws, complex injection points, and authentication bypasses. These are often the most critical issues.",
        findings_count=len(ai_findings),
        severity="critical" if any(f.get('severity') == 'critical' for f in ai_findings) else "high" if ai_findings else "info",
        progress_percent=60,
    ))
    
    # Step 9: AI Verification (false positive filtering)
    step_num += 1
    verification = ctx.get('verification_results', {})
    verified_findings = verification.get('verified_findings', []) if verification else []
    verification_stats = verification.get('stats', {}) if verification else {}
    filter_rate = verification_stats.get('filter_rate', 0)
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="AI Finding Verification",
        title="✅ False Positive Elimination",
        description=f"AI verified findings: {len(verified_findings)} confirmed real ({filter_rate:.0%} filtered as false positives).",
        technical_detail=f"Confidence scores assigned. CONFIRMED: {verification_stats.get('by_verdict', {}).get('CONFIRMED', 0)}, LIKELY: {verification_stats.get('by_verdict', {}).get('LIKELY', 0)}",
        beginner_explanation="Not every security alert is a real problem. AI reviews each finding in context to determine if it's a genuine vulnerability or just code that looks suspicious but is actually safe.",
        why_it_matters="False positives waste time and create alert fatigue. By filtering them out, you can focus on issues that actually need fixing. Our AI typically filters 40-60% of false alerts.",
        findings_count=len(verified_findings),
        severity="info",
        progress_percent=70,
    ))
    
    # Step 10: Component Analysis
    step_num += 1
    components = ctx.get('components', [])
    exported = [c for c in components if c.get('is_exported')]
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Component Analysis",
        title="🚪 Mapping the Attack Surface",
        description=f"Found {len(components)} components, {len(exported)} are exported (accessible to other apps).",
        technical_detail=f"Exported: {len([c for c in exported if c.get('component_type') == 'activity'])} activities, {len([c for c in exported if c.get('component_type') == 'service'])} services, {len([c for c in exported if c.get('component_type') == 'receiver'])} receivers",
        beginner_explanation="Apps are made of building blocks called components. 'Exported' components can be triggered by other apps on your phone. It's like having multiple doors to your house - each one needs to be locked and secured.",
        why_it_matters="Exported components are entry points attackers can target. They might trigger functionality without authorization or inject malicious data through intents.",
        findings_count=len(exported),
        severity="medium" if len(exported) > 5 else "low",
        progress_percent=75,
    ))
    
    # Step 11: Protection Detection
    step_num += 1
    dynamic = ctx.get('dynamic_analysis', {})
    protections = sum([
        1 if dynamic.get('ssl_pinning_detected') else 0,
        1 if dynamic.get('root_detection_detected') else 0,
        1 if dynamic.get('emulator_detection_detected') else 0,
        1 if dynamic.get('anti_tampering_detected') else 0,
    ])
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Protection Detection",
        title="🛡️ Security Protection Analysis",
        description=f"Detected {protections}/4 security protection mechanisms.",
        technical_detail=f"SSL Pinning: {'✓' if dynamic.get('ssl_pinning_detected') else '✗'}, Root Detection: {'✓' if dynamic.get('root_detection_detected') else '✗'}, Emulator Detection: {'✓' if dynamic.get('emulator_detection_detected') else '✗'}, Anti-Tampering: {'✓' if dynamic.get('anti_tampering_detected') else '✗'}",
        beginner_explanation="Apps can include protections against tampering and analysis. SSL pinning ensures network connections can't be intercepted. Root/emulator detection stops the app on compromised devices. These are like security cameras and alarms.",
        why_it_matters=f"Missing protections ({4 - protections}/4) make it easier for attackers to reverse engineer, modify, or intercept communications with the app.",
        findings_count=4 - protections,
        severity="medium" if protections < 2 else "low",
        progress_percent=80,
    ))
    
    # Step 12: AI Reports (Architecture & Attack Surface)
    step_num += 1
    has_arch = bool(ctx.get('ai_architecture_diagram'))
    has_attack_surface = bool(ctx.get('ai_attack_surface_map'))
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="AI Report Generation",
        title="📊 Architecture & Attack Surface Maps",
        description=f"Generated AI-powered visual diagrams: Architecture ({'✓' if has_arch else '✗'}), Attack Surface ({'✓' if has_attack_surface else '✗'}).",
        technical_detail="AI analyzed the app structure to create visual maps showing how components connect and where attackers might target.",
        beginner_explanation="Visual diagrams help understand complex apps at a glance. The architecture diagram shows how different parts of the app connect. The attack surface map highlights the weakest points attackers would target.",
        why_it_matters="These diagrams help prioritize security efforts. The attack surface map shows where to focus testing and hardening. Architecture diagrams help understand data flows and trust boundaries.",
        findings_count=2 if has_arch and has_attack_surface else 1 if has_arch or has_attack_surface else 0,
        severity="info",
        progress_percent=90,
    ))
    
    # Step 13: Frida Scripts
    step_num += 1
    scripts_count = dynamic.get('total_scripts', 0) if dynamic else 0
    vuln_hooks = ctx.get('vulnerability_frida_hooks', {})
    vuln_hook_count = sum(len(hooks) for hooks in vuln_hooks.values()) if vuln_hooks else 0
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Dynamic Testing Scripts",
        title="⚡ Frida Scripts for Testing",
        description=f"Generated {scripts_count + vuln_hook_count} Frida scripts for dynamic security testing.",
        technical_detail=f"General scripts: {scripts_count}, Vulnerability-specific hooks: {vuln_hook_count}. Categories: SSL bypass, root bypass, crypto monitoring, auth hooks.",
        beginner_explanation="Frida is a tool that lets security researchers modify app behavior in real-time while it runs. We generate ready-to-use scripts that can bypass protections, monitor sensitive operations, and test vulnerabilities.",
        why_it_matters="These scripts let you verify if vulnerabilities are exploitable and test if protections actually work. Essential for hands-on security testing.",
        findings_count=scripts_count + vuln_hook_count,
        severity="info",
        progress_percent=95,
    ))
    
    # Step 14: Summary
    step_num += 1
    all_issues = ctx.get('security_issues', [])
    verified = verified_findings if verified_findings else []
    critical_total = len([f for f in verified if f.get('severity') == 'critical']) + cve_critical + critical_code
    high_total = len([f for f in verified if f.get('severity') == 'high']) + high_code
    steps.append(WalkthroughStep(
        step_number=step_num,
        phase="Analysis Complete",
        title="🎯 Summary & Action Items",
        description=f"Analysis complete. Found {critical_total} critical and {high_total} high severity issues requiring attention.",
        technical_detail=f"Total findings: {len(all_issues) + len(code_findings) + len(cve_findings)}, Verified: {len(verified)}, Secrets: {len(secrets)}, Sensitive data: {len(sensitive_findings)}",
        beginner_explanation="We've completed a comprehensive security analysis covering code, dependencies, protections, and data exposure. The findings show where the app could be improved from a security standpoint.",
        why_it_matters="Start with critical issues - these could be actively exploited. Then address high severity items. Use the generated reports and scripts for deeper testing and validation.",
        findings_count=critical_total + high_total,
        severity="critical" if critical_total else "high" if high_total else "medium",
        progress_percent=100,
    ))
    
    # Learning resources
    resources = [
        {"title": "OWASP Mobile Top 10", "url": "https://owasp.org/www-project-mobile-top-10/", "description": "Top 10 mobile security risks - essential reading"},
        {"title": "Mobile Security Testing Guide", "url": "https://mas.owasp.org/MASTG/", "description": "Comprehensive mobile pentesting guide from OWASP"},
        {"title": "Android Security Best Practices", "url": "https://developer.android.com/topic/security/best-practices", "description": "Official Android security guide from Google"},
        {"title": "Frida Documentation", "url": "https://frida.re/docs/", "description": "Learn dynamic instrumentation for mobile testing"},
        {"title": "HackTricks - Android Pentesting", "url": "https://book.hacktricks.xyz/mobile-pentesting/android-app-pentesting", "description": "Practical Android hacking techniques"},
    ]
    
    # Next steps based on findings
    next_steps = []
    if critical_total > 0:
        next_steps.append("🚨 Address critical vulnerabilities immediately - these are exploitable")
    if len(secrets) > 0:
        next_steps.append("🔑 Remove hardcoded secrets and migrate to secure storage")
    if len(sensitive_findings) > 0:
        next_steps.append("🔐 Review sensitive data findings - ensure PII is properly protected")
    if len(cve_findings) > 0:
        next_steps.append("📦 Update vulnerable libraries to patched versions")
    if len(exported) > 5:
        next_steps.append("🚪 Review exported components for proper access control")
    if not dynamic.get('ssl_pinning_detected'):
        next_steps.append("🔒 Implement SSL certificate pinning to prevent MITM attacks")
    if ctx.get('target_sdk', 99) < 28:
        next_steps.append("📱 Increase target SDK to 28+ for modern security features")
    if has_attack_surface:
        next_steps.append("🗺️ Use the Attack Surface Map to prioritize security testing")
    if not next_steps:
        next_steps.append("✅ App has good security posture - continue monitoring and keep dependencies updated")
    
    return AnalysisWalkthroughResponse(
        total_steps=len(steps),
        steps=steps,
        glossary=glossary,
        learning_resources=resources,
        next_steps=next_steps,
    )


# ============================================================================
# Unified APK Scan Endpoint (SSE Progress Streaming)
# ============================================================================

class UnifiedApkScanPhase(BaseModel):
    """A phase in the unified APK scan."""
    id: str
    label: str
    description: str
    status: str  # "pending", "in_progress", "completed", "error"
    progress: int = 0  # 0-100
    details: Optional[str] = None
    started_at: Optional[str] = None
    completed_at: Optional[str] = None


class UnifiedApkScanProgress(BaseModel):
    """Progress update for unified APK scan."""
    scan_id: str
    current_phase: str
    overall_progress: int  # 0-100
    phases: List[UnifiedApkScanPhase]
    message: str
    error: Optional[str] = None


class UnifiedApkScanResult(BaseModel):
    """Complete result from unified APK scan."""
    scan_id: str
    package_name: str
    version_name: Optional[str] = None
    version_code: Optional[int] = None
    min_sdk: Optional[int] = None
    target_sdk: Optional[int] = None
    
    # Quick Analysis Results
    permissions: List[Dict[str, Any]] = []
    dangerous_permissions_count: int = 0
    components: List[Dict[str, Any]] = []
    secrets: List[Dict[str, Any]] = []
    urls: List[str] = []
    native_libraries: List[str] = []
    security_issues: List[Dict[str, Any]] = []
    
    # JADX Decompilation Results
    jadx_session_id: Optional[str] = None
    total_classes: int = 0
    total_files: int = 0
    classes_summary: List[Dict[str, Any]] = []
    source_tree: Optional[Dict[str, Any]] = None
    jadx_security_issues: List[Dict[str, Any]] = []
    decompilation_time: float = 0
    
    # AI Analysis Results  
    ai_functionality_report: Optional[str] = None
    ai_security_report: Optional[str] = None
    ai_architecture_diagram: Optional[str] = None
    ai_attack_surface_map: Optional[str] = None  # Mermaid attack tree diagram
    
    # Decompiled Source Code Security Findings (pattern-based scanners)
    decompiled_code_findings: List[Dict[str, Any]] = []  # Individual vulnerability findings
    decompiled_code_summary: Dict[str, Any] = {}  # Summary with counts by severity/category
    
    # CVE Database Lookup Results
    cve_scan_results: Optional[Dict[str, Any]] = None  # Libraries and CVEs found
    
    # Dynamic Analysis - Frida Scripts (standard bypass scripts)
    dynamic_analysis: Optional[Dict[str, Any]] = None  # Standard Frida scripts (SSL bypass, root bypass, etc.)
    
    # Vulnerability-Specific Frida Hooks (auto-generated from findings)
    vulnerability_frida_hooks: Optional[Dict[str, Any]] = None  # Targeted Frida scripts based on discovered vulnerabilities
    
    # Multi-Pass AI Vulnerability Hunt Results
    vuln_hunt_result: Optional[Dict[str, Any]] = None  # Deep AI-guided vulnerability hunting results
    
    # Manifest Visualization (auto-generated during scan)
    manifest_visualization: Optional[Dict[str, Any]] = None  # Component graph and security analysis
    
    # Obfuscation Analysis (auto-generated during scan)
    obfuscation_analysis: Optional[Dict[str, Any]] = None  # Obfuscation detection and deobfuscation tips
    
    # Unified Verification Results (false positive filtering + attack chain detection)
    verification_results: Optional[Dict[str, Any]] = None  # Verified findings, filtered FPs, attack chains
    
    # Sensitive Data Discovery (AI-verified passwords, API keys, emails, phone numbers, etc.)
    sensitive_data_findings: Optional[Dict[str, Any]] = None  # Hardcoded credentials and PII found in source
    
    # Metadata
    scan_time: float = 0
    filename: str = ""
    file_size: int = 0


@router.post("/apk/unified-scan")
async def unified_apk_scan(
    file: UploadFile = File(..., description="APK file to analyze"),
    include_vuln_hunt: bool = Form(True, description="Enable multi-pass AI vulnerability hunting (default: enabled)"),
    vuln_hunt_max_passes: int = Form(5, description="Maximum passes for vulnerability hunt (2-8)", ge=2, le=8),
    vuln_hunt_max_targets: int = Form(50, description="Max targets per pass (10-100)", ge=10, le=100),
    current_user: User = Depends(get_current_active_user),
):
    """
    Perform a complete APK analysis with streaming progress updates.
    
    This unified scan combines:
    1. Manifest & permission analysis
    2. Secret & string extraction
    3. JADX decompilation to Java source
    4. Code security scan (pattern-based)
    5. CVE database lookup
    6. AI deep analysis
    7. Multi-pass AI vulnerability hunt (enabled by default)
    8. AI-powered reports (functionality, security, architecture, attack surface)
    
    Returns SSE stream with progress updates and final result.
    """
    scan_id = str(uuid.uuid4())
    filename = file.filename or "unknown.apk"
    suffix = Path(filename).suffix.lower()
    
    if suffix not in ALLOWED_APK_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {', '.join(ALLOWED_APK_EXTENSIONS)}"
        )
    
    # Check JADX availability
    if not check_jadx_available():
        raise HTTPException(
            status_code=503,
            detail="JADX is not available. Full analysis requires JADX."
        )
    
    # Save file to temp location with sanitized filename to prevent path traversal
    tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_unified_"))
    safe_filename = sanitize_filename(filename)
    tmp_path = tmp_dir / safe_filename
    
    file_size = 0
    try:
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    shutil.rmtree(tmp_dir, ignore_errors=True)
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
    except Exception as e:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
    
    logger.info(f"Starting unified APK scan: {filename} ({file_size:,} bytes)")
    
    # Initialize scan session with timestamp for cleanup tracking
    _unified_scan_sessions[scan_id] = {
        "cancelled": False,
        "tmp_dir": tmp_dir,
        "file_path": tmp_path,
    }
    _unified_scan_timestamps[scan_id] = datetime.now()
    
    # Define phases (vuln_hunt is optional, inserted before ai_reports)
    phases = [
        UnifiedApkScanPhase(
            id="manifest",
            label="Manifest Analysis",
            description="Extracting package info, permissions, and components",
            status="pending"
        ),
        UnifiedApkScanPhase(
            id="secrets",
            label="Secret Detection",
            description="Scanning for hardcoded secrets, URLs, and API keys",
            status="pending"
        ),
        UnifiedApkScanPhase(
            id="jadx",
            label="JADX Decompilation",
            description="Decompiling DEX to Java source code",
            status="pending"
        ),
        UnifiedApkScanPhase(
            id="code_scan",
            label="Code Security Scan",
            description="Pattern-based vulnerability scanning of decompiled code",
            status="pending"
        ),
        UnifiedApkScanPhase(
            id="sensitive_data",
            label="Sensitive Data Discovery",
            description="AI-verified scan for passwords, API keys, emails, phone numbers, and PII",
            status="pending"
        ),
        UnifiedApkScanPhase(
            id="cve_lookup",
            label="CVE Database Lookup",
            description="Checking dependencies against CVE/CWE vulnerability databases",
            status="pending"
        ),
    ]
    
    # Add vuln_hunt phase if enabled (runs BEFORE ai_analysis)
    if include_vuln_hunt:
        phases.append(UnifiedApkScanPhase(
            id="vuln_hunt",
            label="AI Vulnerability Hunt",
            description=f"Multi-pass AI-guided vulnerability hunting ({vuln_hunt_max_passes} passes)",
            status="pending"
        ))
    
    # Always add verification phase - validates findings and filters false positives
    # Runs AFTER vuln_hunt but BEFORE ai_analysis so reports use verified findings
    phases.append(UnifiedApkScanPhase(
        id="verification",
        label="AI Finding Verification",
        description="Validating findings with confidence scoring and false positive elimination",
        status="pending"
    ))
    
    # AI analysis phases use verified findings
    phases.append(UnifiedApkScanPhase(
        id="ai_analysis",
        label="AI Deep Analysis",
        description="Cross-reference analysis, parallel scan, and code sampling for AI",
        status="pending"
    ))
    phases.append(UnifiedApkScanPhase(
        id="advanced_analysis",
        label="Advanced Analysis",
        description="Manifest visualization and obfuscation detection",
        status="pending"
    ))
    
    # Final report generation using all verified data
    phases.append(UnifiedApkScanPhase(
        id="ai_reports",
        label="AI Report Generation",
        description="Generating functionality, security, architecture, and attack surface reports",
        status="pending"
    ))
    
    async def run_unified_scan():
        """Generator that yields SSE progress events."""
        result = UnifiedApkScanResult(
            scan_id=scan_id,
            package_name="",
            filename=filename,
            file_size=file_size,
        )
        start_time = datetime.now()
        current_phase_idx = 0
        apk_result = None
        jadx_result = None
        
        def make_progress(message: str, phase_progress: int = 0) -> str:
            nonlocal phases, current_phase_idx
            overall = (current_phase_idx * 100 // len(phases)) + (phase_progress // len(phases))
            progress = UnifiedApkScanProgress(
                scan_id=scan_id,
                current_phase=phases[current_phase_idx].id,
                overall_progress=min(overall, 100),
                phases=phases,
                message=message,
            )
            return f"data: {json.dumps({'type': 'progress', 'data': progress.model_dump()})}\n\n"
        
        def update_phase(phase_id: str, status: str, details: str = None, progress: int = 0):
            for p in phases:
                if p.id == phase_id:
                    p.status = status
                    p.progress = progress
                    if details:
                        p.details = details
                    if status == "in_progress" and not p.started_at:
                        p.started_at = datetime.now().isoformat()
                    if status == "completed":
                        p.completed_at = datetime.now().isoformat()
                        p.progress = 100
                    break
        
        try:
            # Check if cancelled
            if _unified_scan_sessions.get(scan_id, {}).get("cancelled"):
                yield f"data: {json.dumps({'type': 'cancelled'})}\n\n"
                return
            
            # =================================================================
            # Phase 1: Manifest Analysis
            # =================================================================
            current_phase_idx = 0
            update_phase("manifest", "in_progress")
            yield make_progress("Extracting AndroidManifest.xml...", 10)
            
            try:
                apk_result = re_service.analyze_apk(tmp_path)
                result.package_name = apk_result.package_name or ""
                result.version_name = apk_result.version_name
                result.version_code = apk_result.version_code
                result.min_sdk = apk_result.min_sdk
                result.target_sdk = apk_result.target_sdk
                result.permissions = [
                    {"name": p.name, "is_dangerous": p.is_dangerous, "description": p.description}
                    for p in apk_result.permissions
                ]
                result.dangerous_permissions_count = sum(1 for p in apk_result.permissions if p.is_dangerous)
                result.components = [
                    {"name": c.name, "component_type": c.component_type, "is_exported": c.is_exported, "intent_filters": c.intent_filters}
                    for c in apk_result.components
                ]
                result.native_libraries = apk_result.native_libraries
                update_phase("manifest", "completed", f"Found {len(result.permissions)} permissions, {len(result.components)} components")
            except Exception as e:
                update_phase("manifest", "error", str(e))
                logger.error(f"Manifest analysis failed: {e}")
            
            yield make_progress("Manifest analysis complete", 100)
            await asyncio.sleep(0.1)  # Allow UI to update
            
            # =================================================================
            # Phase 2: Secret Detection
            # =================================================================
            current_phase_idx = 1
            update_phase("secrets", "in_progress")
            yield make_progress("Scanning for secrets and URLs...", 10)
            
            if apk_result:
                result.secrets = [
                    {"type": s["type"], "value": s["value"], "masked_value": s["masked_value"], "severity": s["severity"]}
                    for s in apk_result.secrets
                ]
                result.urls = apk_result.urls[:100]
                result.security_issues = [
                    {"category": i["category"], "severity": i["severity"], "description": i["description"]}
                    for i in apk_result.security_issues
                ]
                update_phase("secrets", "completed", f"Found {len(result.secrets)} secrets, {len(result.urls)} URLs")
            else:
                update_phase("secrets", "completed", "Skipped - manifest failed")
            
            yield make_progress("Secret detection complete", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase 3: JADX Decompilation
            # =================================================================
            current_phase_idx = 2
            update_phase("jadx", "in_progress")
            yield make_progress("Starting JADX decompilation...", 5)
            
            try:
                # Run JADX in thread pool to avoid blocking the event loop
                # This allows progress messages to be sent during decompilation
                loop = asyncio.get_event_loop()
                
                yield make_progress("Running JADX (decompiling to Java source)...", 10)
                
                # Start a background task to send periodic progress updates
                jadx_done = False
                async def send_jadx_progress():
                    progress = 15
                    while not jadx_done and progress < 85:
                        await asyncio.sleep(3)  # Update every 3 seconds
                        if not jadx_done:
                            progress = min(progress + 5, 85)
                            # Can't yield from here, but this keeps connection alive
                    
                # Run JADX in executor (non-blocking)
                jadx_result = await loop.run_in_executor(
                    None, 
                    lambda: re_service.decompile_apk_with_jadx(tmp_path, max_classes_to_parse=2000)
                )
                jadx_done = True
                
                yield make_progress("JADX decompilation finished, processing results...", 90)
                
                # Store in cache for later queries with timestamp for TTL cleanup
                jadx_session_id = str(uuid.uuid4())
                _jadx_cache[jadx_session_id] = Path(jadx_result.output_directory)
                _jadx_cache_timestamps[jadx_session_id] = datetime.now()
                
                result.jadx_session_id = jadx_session_id
                result.total_classes = jadx_result.total_classes
                result.total_files = jadx_result.total_files
                result.decompilation_time = jadx_result.decompilation_time
                result.source_tree = jadx_result.source_tree
                
                # Log source tree info for debugging
                logger.info(f"JADX source tree has {len(jadx_result.source_tree)} top-level entries")
                logger.info(f"JADX decompiled {jadx_result.total_files} files, parsed {jadx_result.total_classes} classes in {jadx_result.decompilation_time:.1f}s")
                
                # Collect security issues
                all_jadx_issues = []
                for cls in jadx_result.classes:
                    all_jadx_issues.extend(cls.security_issues)
                result.jadx_security_issues = all_jadx_issues[:100]
                
                # Classes summary
                result.classes_summary = [
                    {
                        "class_name": c.class_name,
                        "package_name": c.package_name,
                        "file_path": c.file_path,
                        "line_count": c.line_count,
                        "is_activity": c.is_activity,
                        "is_service": c.is_service,
                        "security_issues_count": len(c.security_issues),
                    }
                    for c in jadx_result.classes[:500]
                ]
                
                update_phase("jadx", "completed", f"Decompiled {jadx_result.total_files} files ({jadx_result.total_classes} parsed) in {jadx_result.decompilation_time:.1f}s")
            except Exception as e:
                update_phase("jadx", "error", str(e))
                logger.error(f"JADX decompilation failed: {e}", exc_info=True)
            
            yield make_progress("JADX decompilation complete", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase 4: Code Security Scan (Pattern-based vulnerability detection)
            # =================================================================
            current_phase_idx = 3
            update_phase("code_scan", "in_progress")
            yield make_progress("Running decompiled code security scanners...", 10)
            
            try:
                if jadx_result and jadx_result.output_directory:
                    jadx_output_path = Path(jadx_result.output_directory)
                    if jadx_output_path.exists():
                        yield make_progress("Scanning WebView, crypto, SQL injection...", 20)
                        # Also run in executor to avoid blocking
                        code_scan_results = await loop.run_in_executor(
                            None,
                            lambda: re_service.scan_decompiled_source_comprehensive(jadx_output_path)
                        )
                        raw_findings = code_scan_results.get("findings", [])
                        
                        # FIX #3: Normalize field names for consistent processing
                        raw_findings = re_service.normalize_findings_batch(raw_findings, source="pattern")
                        
                        # Apply deduplication to remove duplicate findings
                        yield make_progress("Deduplicating findings...", 80)
                        if raw_findings:
                            try:
                                deduplicated = deduplication_service.deduplicate_findings_simple(raw_findings)
                                result.decompiled_code_findings = deduplicated
                                dedup_removed = len(raw_findings) - len(deduplicated)
                                if dedup_removed > 0:
                                    logger.info(f"Deduplication removed {dedup_removed} duplicate findings")
                            except Exception as dedup_err:
                                logger.warning(f"Deduplication failed, using raw findings: {dedup_err}")
                                result.decompiled_code_findings = raw_findings
                        else:
                            result.decompiled_code_findings = raw_findings
                        
                        result.decompiled_code_summary = code_scan_results.get("summary", {})
                        
                        finding_count = len(result.decompiled_code_findings)
                        critical_count = code_scan_results.get("summary", {}).get("by_severity", {}).get("critical", 0)
                        high_count = code_scan_results.get("summary", {}).get("by_severity", {}).get("high", 0)
                        
                        update_phase("code_scan", "completed", f"Found {finding_count} issues ({critical_count} critical, {high_count} high)")
                        logger.info(f"Decompiled code scan found {finding_count} issues ({critical_count} critical, {high_count} high)")
                else:
                    update_phase("code_scan", "completed", "Skipped - JADX not available")
            except Exception as e:
                update_phase("code_scan", "error", str(e))
                logger.error(f"Decompiled code security scan failed: {e}")
            
            yield make_progress("Code security scan complete", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase 5: Sensitive Data Discovery (AI-verified)
            # Scans for passwords, API keys, emails, phone numbers, PII
            # =================================================================
            current_phase_idx = 4
            update_phase("sensitive_data", "in_progress")
            yield make_progress("Scanning for sensitive data (passwords, API keys, emails, phone numbers)...", 10)
            
            try:
                if jadx_result and jadx_result.output_directory:
                    jadx_output_path = Path(jadx_result.output_directory)
                    sources_dir = jadx_output_path / "sources"
                    
                    if sources_dir.exists():
                        yield make_progress("Pattern matching for sensitive data...", 30)
                        
                        # Run sensitive data scan with AI verification
                        sensitive_data_result = await re_service.scan_sensitive_data_with_ai(
                            sources_dir=sources_dir,
                            package_name=result.package_name
                        )
                        
                        if sensitive_data_result:
                            result.sensitive_data_findings = sensitive_data_result
                            
                            findings_count = len(sensitive_data_result.get("findings", []))
                            filtered_count = len(sensitive_data_result.get("filtered_out", []))
                            summary = sensitive_data_result.get("summary", {})
                            
                            if findings_count > 0:
                                by_cat = summary.get("by_category", {})
                                cat_summary = ", ".join([f"{k}: {v}" for k, v in list(by_cat.items())[:3]])
                                update_phase("sensitive_data", "completed", 
                                    f"Found {findings_count} sensitive items ({cat_summary}), filtered {filtered_count} false positives")
                                logger.info(f"Sensitive data scan: {findings_count} findings, {filtered_count} filtered")
                            else:
                                update_phase("sensitive_data", "completed", f"No sensitive data found (filtered {filtered_count} false positives)")
                        else:
                            update_phase("sensitive_data", "completed", "Scan completed - no findings")
                    else:
                        update_phase("sensitive_data", "completed", "Skipped - no source files")
                else:
                    update_phase("sensitive_data", "completed", "Skipped - JADX not available")
            except Exception as e:
                update_phase("sensitive_data", "error", str(e))
                logger.error(f"Sensitive data scan failed: {e}", exc_info=True)
            
            yield make_progress("Sensitive data discovery complete", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase 6: CVE Database Lookup (Known vulnerabilities in dependencies)
            # =================================================================
            current_phase_idx = 5
            update_phase("cve_lookup", "in_progress")
            yield make_progress("Checking CVE/CWE databases for known vulnerabilities...", 10)
            
            cve_results = {"libraries": [], "cves": [], "stats": {}}
            try:
                if jadx_result and jadx_result.output_directory:
                    jadx_output_path = Path(jadx_result.output_directory)
                    sources_dir = jadx_output_path / "sources"
                    if sources_dir.exists():
                        yield make_progress("Extracting library dependencies...", 30)
                        
                        # Get all Java class names
                        class_names = []
                        for java_file in sources_dir.rglob("*.java"):
                            try:
                                rel_path = java_file.relative_to(sources_dir)
                                class_name = str(rel_path).replace("/", ".").replace("\\", ".").replace(".java", "")
                                class_names.append(class_name)
                            except (ValueError, OSError) as e:
                                logger.debug(f"Failed to process Java file {java_file}: {e}")
                        
                        # Extract dependencies and lookup CVEs
                        libraries = re_service.extract_apk_dependencies(class_names, None)
                        cve_results["libraries"] = [
                            {"name": lib.maven_coordinate, "version": lib.version, "is_high_risk": lib.is_high_risk}
                            for lib in libraries
                        ]
                        
                        if libraries:
                            yield make_progress(f"Looking up CVEs for {len(libraries)} libraries...", 60)
                            cves = await re_service.lookup_apk_cves(libraries)
                            
                            # FIX #1: Verify CVE reachability before adding to results
                            if cves:
                                yield make_progress(f"Verifying {len(cves)} CVE reachability...", 80)
                                cve_verification = await re_service.verify_cve_reachability(
                                    cve_findings=cves,
                                    sources_dir=sources_dir,
                                    package_name=result.package_name
                                )
                                
                                # Use only reachable CVEs in main results
                                verified_cves = cve_verification.get("verified_cves", cves)
                                unreachable_cves = cve_verification.get("unreachable_cves", [])
                                
                                cve_results["cves"] = verified_cves
                                cve_results["unreachable_cves"] = unreachable_cves
                                cve_results["reachability_stats"] = cve_verification.get("verification_stats", {})
                            else:
                                cve_results["cves"] = cves
                            
                            cve_results["stats"] = {
                                "total_libraries": len(libraries),
                                "high_risk_libraries": sum(1 for lib in libraries if lib.is_high_risk),
                                "total_cves": len(cve_results["cves"]),
                                "unreachable_cves": len(cve_results.get("unreachable_cves", [])),
                                "critical_cves": sum(1 for cve in cve_results["cves"] if cve.get("severity", "").lower() == "critical"),
                                "high_cves": sum(1 for cve in cve_results["cves"] if cve.get("severity", "").lower() == "high"),
                            }
                            
                            # Store CVE results in the scan result
                            result.cve_scan_results = cve_results
                            
                            unreachable_msg = f", {cve_results['stats']['unreachable_cves']} unreachable filtered" if cve_results['stats']['unreachable_cves'] > 0 else ""
                            update_phase("cve_lookup", "completed", 
                                f"Found {len(cve_results['cves'])} reachable CVEs ({cve_results['stats']['critical_cves']} critical, {cve_results['stats']['high_cves']} high){unreachable_msg}")
                            logger.info(f"CVE lookup found {len(cve_results['cves'])} reachable CVEs in {len(libraries)} libraries")
                        else:
                            update_phase("cve_lookup", "completed", "No third-party libraries detected")
                    else:
                        update_phase("cve_lookup", "completed", "Skipped - no source files")
                else:
                    update_phase("cve_lookup", "completed", "Skipped - JADX not available")
            except Exception as e:
                update_phase("cve_lookup", "error", str(e))
                logger.error(f"CVE lookup failed: {e}")
            
            yield make_progress("CVE database lookup complete", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase 6 (Optional): Multi-Pass AI Vulnerability Hunt
            # MOVED BEFORE AI Reports so findings feed into all reports
            # =================================================================
            vuln_hunt_findings = []
            if include_vuln_hunt:
                current_phase_idx += 1
                update_phase("vuln_hunt", "in_progress")
                yield make_progress("Starting multi-pass AI vulnerability hunt...", 5)
                
                try:
                    if jadx_result and jadx_result.output_directory:
                        jadx_output_path = Path(jadx_result.output_directory)
                        sources_dir = jadx_output_path / "sources"
                        
                        if sources_dir.exists():
                            yield make_progress("Pass 0: AI filename triage...", 10)
                            yield make_progress("Pass 1: Pattern scan with AI-prioritized files...", 25)
                            
                            # Run multi-pass vulnerability hunt on Java source code
                            vuln_hunt_result_data = await re_service.ai_vulnerability_hunt_java(
                                sources_dir=sources_dir,
                                package_name=result.package_name,
                                existing_findings=result.decompiled_code_findings,
                                max_passes=vuln_hunt_max_passes,
                                max_targets_per_pass=vuln_hunt_max_targets,
                                on_progress=lambda phase, prog, msg: None,
                            )
                            
                            if vuln_hunt_result_data:
                                result.vuln_hunt_result = vuln_hunt_result_data
                                vuln_hunt_findings = vuln_hunt_result_data.get("vulnerabilities", [])
                                
                                # FIX #3: Normalize vuln hunt findings
                                vuln_hunt_findings = re_service.normalize_findings_batch(vuln_hunt_findings, source="vuln_hunt")
                                
                                vuln_count = len(vuln_hunt_findings)
                                critical_count = sum(
                                    1 for v in vuln_hunt_findings
                                    if v.get("severity") == "critical"
                                )
                                update_phase("vuln_hunt", "completed", 
                                    f"Found {vuln_count} vulnerabilities ({critical_count} critical)")
                                logger.info(f"APK vuln hunt found {vuln_count} vulnerabilities")
                            else:
                                update_phase("vuln_hunt", "completed", "No additional vulnerabilities found")
                        else:
                            update_phase("vuln_hunt", "completed", "Skipped - no source files")
                    else:
                        update_phase("vuln_hunt", "completed", "Skipped - JADX not available")
                except Exception as e:
                    update_phase("vuln_hunt", "error", str(e))
                    logger.error(f"APK vulnerability hunt failed: {e}", exc_info=True)
                
                yield make_progress("Vulnerability hunt complete", 100)
                await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase: AI Finding Verification - Confidence scoring & FP elimination
            # Validates ALL findings (pattern + vuln hunt) before AI reports
            # =================================================================
            current_phase_idx += 1
            update_phase("verification", "in_progress")
            yield make_progress("Starting AI verification of findings...", 5)
            
            verified_findings = []
            verification_stats = {}
            
            try:
                if jadx_result and jadx_result.output_directory:
                    jadx_output_path = Path(jadx_result.output_directory)
                    sources_dir = jadx_output_path / "sources"
                    
                    if sources_dir.exists():
                        # Combine all findings for verification
                        all_findings_to_verify = (result.decompiled_code_findings or []).copy()
                        all_findings_to_verify.extend(vuln_hunt_findings)
                        
                        if all_findings_to_verify:
                            yield make_progress(f"Verifying {len(all_findings_to_verify)} findings with AI...", 15)
                            
                            # Run unified verification
                            verification_result = await re_service.verify_findings_unified(
                                findings=all_findings_to_verify,
                                sources_dir=sources_dir,
                                package_name=result.package_name
                            )
                            
                            if verification_result:
                                verified_findings = verification_result.get("verified_findings", [])
                                verification_stats = verification_result.get("verification_stats", {})
                                
                                # FIX #2: Merge sensitive data filtered_out with main verification results
                                if result.sensitive_data_findings:
                                    sensitive_filtered = result.sensitive_data_findings.get("filtered_out", [])
                                    if sensitive_filtered:
                                        # Normalize and add to main filtered_out
                                        normalized_sensitive_filtered = re_service.normalize_findings_batch(
                                            sensitive_filtered, source="sensitive_data"
                                        )
                                        existing_filtered = verification_result.get("filtered_out", [])
                                        verification_result["filtered_out"] = existing_filtered + normalized_sensitive_filtered
                                        
                                        # Update stats
                                        verification_stats["sensitive_data_filtered"] = len(sensitive_filtered)
                                        verification_stats["filtered"] = verification_stats.get("filtered", 0) + len(sensitive_filtered)
                                        logger.info(f"Merged {len(sensitive_filtered)} sensitive data false positives into verification results")
                                
                                # FIX #2: Also merge CVE unreachable findings
                                if result.cve_scan_results:
                                    unreachable_cves = result.cve_scan_results.get("unreachable_cves", [])
                                    if unreachable_cves:
                                        # Normalize and add to main filtered_out
                                        normalized_unreachable = re_service.normalize_findings_batch(
                                            unreachable_cves, source="cve_unreachable"
                                        )
                                        existing_filtered = verification_result.get("filtered_out", [])
                                        verification_result["filtered_out"] = existing_filtered + normalized_unreachable
                                        
                                        # Update stats
                                        verification_stats["cve_unreachable_filtered"] = len(unreachable_cves)
                                        verification_stats["filtered"] = verification_stats.get("filtered", 0) + len(unreachable_cves)
                                        logger.info(f"Merged {len(unreachable_cves)} unreachable CVEs into verification filtered results")
                                
                                result.verification_results = verification_result
                                
                                # Update the original findings with verified versions
                                result.decompiled_code_findings = [
                                    f for f in verified_findings 
                                    if f.get("source") != "vuln_hunt"
                                ]
                                vuln_hunt_findings = [
                                    f for f in verified_findings 
                                    if f.get("source") == "vuln_hunt"
                                ]
                                
                                # Update vuln_hunt_result with verified findings
                                if result.vuln_hunt_result:
                                    result.vuln_hunt_result["vulnerabilities"] = vuln_hunt_findings
                                
                                filtered_count = verification_stats.get("filtered", 0)
                                confirmed_count = verification_stats.get("by_verdict", {}).get("CONFIRMED", 0)
                                likely_count = verification_stats.get("by_verdict", {}).get("LIKELY", 0)
                                avg_conf = verification_stats.get("avg_confidence", 0)
                                
                                update_phase("verification", "completed", 
                                    f"Verified {len(verified_findings)} findings ({confirmed_count} confirmed, {likely_count} likely), filtered {filtered_count} false positives, avg confidence: {avg_conf:.0f}%")
                                logger.info(f"Finding verification: {len(verified_findings)} verified, {filtered_count} filtered as false positives")
                            else:
                                update_phase("verification", "completed", "No verification results")
                        else:
                            update_phase("verification", "completed", "No findings to verify")
                    else:
                        update_phase("verification", "completed", "Skipped - no source files")
                else:
                    update_phase("verification", "completed", "Skipped - JADX not available")
            except Exception as e:
                update_phase("verification", "error", str(e))
                logger.error(f"Finding verification failed: {e}", exc_info=True)
                # On error, keep original findings
                verified_findings = (result.decompiled_code_findings or []).copy()
                verified_findings.extend(vuln_hunt_findings)
            
            yield make_progress("Finding verification complete", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase 7: AI Deep Analysis (Cross-ref, parallel scan, code sampling)
            # Now includes vuln hunt findings for comprehensive reports
            # =================================================================
            current_phase_idx += 1
            update_phase("ai_analysis", "in_progress")
            yield make_progress("Running AI deep analysis pipeline...", 10)
            
            try:
                if apk_result:
                    jadx_output_path = Path(jadx_result.output_directory) if jadx_result else None
                    
                    yield make_progress("Parallel scanning source files...", 20)
                    yield make_progress("Analyzing cross-references and data flows...", 40)
                    yield make_progress("Extracting detailed code samples...", 60)
                    yield make_progress("Processing CVE data for AI context...", 80)
                    
                    # Use VERIFIED findings for AI analysis
                    all_findings = verified_findings if verified_findings else (result.decompiled_code_findings or []).copy()
                    if not verified_findings:
                        all_findings.extend(vuln_hunt_findings)
                    
                    # Pass ALL findings (pattern + vuln hunt) AND verification results for comprehensive AI reports
                    # The verification_results contains FP-filtered findings and confidence scores
                    ai_reports = await re_service.analyze_apk_with_ai(
                        apk_result, 
                        jadx_output_path,
                        decompiled_findings=all_findings,
                        verification_results=result.verification_results
                    )
                    
                    update_phase("ai_analysis", "completed", "Deep analysis pipeline complete")
                else:
                    update_phase("ai_analysis", "completed", "Skipped - no APK result")
            except Exception as e:
                update_phase("ai_analysis", "error", str(e))
                logger.error(f"AI deep analysis failed: {e}")
            
            yield make_progress("AI deep analysis complete", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase: Advanced Analysis (Manifest Visualization + Obfuscation)
            # Runs automatically - no user button press needed
            # =================================================================
            current_phase_idx += 1
            update_phase("advanced_analysis", "in_progress")
            yield make_progress("Running advanced analysis (manifest viz, obfuscation)...", 10)
            
            try:
                # Generate Manifest Visualization
                yield make_progress("Generating manifest visualization...", 20)
                manifest_viz = re_service.generate_manifest_visualization(tmp_path)
                if manifest_viz:
                    result.manifest_visualization = {
                        "package_name": manifest_viz.package_name,
                        "app_name": manifest_viz.app_name,
                        "version_name": manifest_viz.version_name,
                        "nodes": [
                            {"id": n.id, "name": n.name, "node_type": n.node_type, "label": n.label, "attributes": n.attributes}
                            for n in manifest_viz.nodes
                        ],
                        "edges": [
                            {"source": e.source, "target": e.target, "edge_type": e.edge_type, "label": e.label}
                            for e in manifest_viz.edges
                        ],
                        "component_counts": manifest_viz.component_counts,
                        "permission_summary": manifest_viz.permission_summary,
                        "exported_count": manifest_viz.exported_count,
                        "main_activity": manifest_viz.main_activity,
                        "deep_link_schemes": manifest_viz.deep_link_schemes,
                        "mermaid_diagram": manifest_viz.mermaid_diagram,
                        "ai_analysis": manifest_viz.ai_analysis,
                        "security_assessment": manifest_viz.security_assessment,
                    }
                    logger.info(f"Manifest visualization: {manifest_viz.exported_count} exported components, {len(manifest_viz.deep_link_schemes)} deep link schemes")
                
                # Generate Obfuscation Analysis
                yield make_progress("Analyzing obfuscation techniques...", 60)
                obfuscation_result = re_service.analyze_apk_obfuscation(tmp_path)
                if obfuscation_result:
                    # Use getattr with fallbacks for robustness against schema changes
                    class_naming = obfuscation_result.class_naming
                    result.obfuscation_analysis = {
                        "package_name": obfuscation_result.package_name,
                        "overall_obfuscation_level": obfuscation_result.overall_obfuscation_level,
                        "obfuscation_score": obfuscation_result.obfuscation_score,
                        "detected_tools": obfuscation_result.detected_tools,
                        "indicators": [
                            {"indicator_type": i.indicator_type, "confidence": i.confidence, "evidence": i.evidence, "description": i.description}
                            for i in obfuscation_result.indicators
                        ],
                        "class_naming": {
                            "total_classes": getattr(class_naming, 'total_classes', 0),
                            "short_name_count": getattr(class_naming, 'short_name_classes', getattr(class_naming, 'short_name_count', 0)),
                            "obfuscated_count": getattr(class_naming, 'single_letter_classes', getattr(class_naming, 'obfuscated_count', 0)),
                            "readable_count": getattr(class_naming, 'meaningful_name_classes', getattr(class_naming, 'readable_count', 0)),
                            "obfuscation_ratio": getattr(class_naming, 'obfuscation_ratio', 0.0),
                        },
                        "deobfuscation_strategies": obfuscation_result.deobfuscation_strategies,
                        "recommended_tools": obfuscation_result.recommended_tools,
                        "frida_hooks": obfuscation_result.frida_hooks,
                        "analysis_time": obfuscation_result.analysis_time,
                        "warnings": obfuscation_result.warnings,
                        "ai_analysis_summary": getattr(obfuscation_result, 'ai_analysis_summary', None),
                        "reverse_engineering_difficulty": getattr(obfuscation_result, 'reverse_engineering_difficulty', None),
                    }
                    logger.info(f"Obfuscation analysis: {obfuscation_result.overall_obfuscation_level} level, score {obfuscation_result.obfuscation_score}/100")
                
                update_phase("advanced_analysis", "completed", f"Manifest viz + obfuscation ({obfuscation_result.overall_obfuscation_level if obfuscation_result else 'unknown'} level)")
            except Exception as e:
                update_phase("advanced_analysis", "error", str(e))
                logger.error(f"Advanced analysis failed: {e}", exc_info=True)
            
            yield make_progress("Advanced analysis complete", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Phase 8: AI Report Generation (Functionality, Security, Architecture, Attack Surface)
            # =================================================================
            current_phase_idx += 1
            update_phase("ai_reports", "in_progress")
            yield make_progress("Generating AI reports...", 10)
            
            try:
                if apk_result:
                    # Get the reports that were generated in the previous phase
                    if apk_result.ai_report_functionality:
                        result.ai_functionality_report = apk_result.ai_report_functionality
                        yield make_progress("Functionality report ready", 25)
                    
                    if apk_result.ai_report_security:
                        result.ai_security_report = apk_result.ai_report_security
                        yield make_progress("Security report ready", 50)
                    
                    # Generate architecture diagram
                    jadx_output_path = Path(jadx_result.output_directory) if jadx_result else None
                    if jadx_output_path and jadx_output_path.exists():
                        jadx_summary = None
                        try:
                            jadx_summary = re_service.get_jadx_result_summary(jadx_output_path)
                        except Exception:
                            pass
                        
                        diagram = await re_service.generate_ai_architecture_diagram(apk_result, jadx_summary, output_dir=jadx_output_path)
                        result.ai_architecture_diagram = diagram
                        yield make_progress("Architecture diagram ready", 75)
                        
                        # Generate attack surface map using VERIFIED findings only
                        # (false positives already filtered out in verification phase)
                        attack_surface = re_service.generate_attack_surface_map(tmp_path)
                        
                        # Use verified findings - these have false positives removed
                        verified_for_attack = verified_findings if verified_findings else (result.decompiled_code_findings or []).copy()
                        if not verified_findings:
                            # Fallback: if no verification ran, combine original findings
                            verified_for_attack.extend(vuln_hunt_findings)
                        
                        attack_tree = await re_service.generate_ai_attack_tree_mermaid(
                            attack_surface, 
                            jadx_output_path,
                            decompiled_findings=verified_for_attack
                        )
                        result.ai_attack_surface_map = attack_tree
                        yield make_progress("Attack surface map ready (verified findings only)", 90)
                    
                    update_phase("ai_reports", "completed", "All 4 AI reports generated")
                else:
                    update_phase("ai_reports", "completed", "Skipped - no APK result")
            except Exception as e:
                update_phase("ai_reports", "error", str(e))
                logger.error(f"AI report generation failed: {e}")
            
            yield make_progress("All AI reports generated", 100)
            await asyncio.sleep(0.1)
            
            # =================================================================
            # Bonus: Generate Frida Scripts (Standard + Vulnerability-specific)
            # =================================================================
            yield make_progress("Generating Frida scripts for dynamic analysis...", 10)
            
            # Generate standard Frida scripts (SSL bypass, root bypass, crypto hooks, etc.)
            try:
                if apk_result:
                    yield make_progress("Generating SSL pinning bypass scripts...", 20)
                    
                    # Extract required data from apk_result for Frida generation
                    # Note: apk_result.strings may be dicts or objects - handle both
                    strings_list = []
                    for s in (apk_result.strings[:1000] if hasattr(apk_result, 'strings') and apk_result.strings else []):
                        try:
                            if isinstance(s, dict):
                                strings_list.append(re_service.ExtractedString(
                                    value=s.get("value", ""),
                                    offset=s.get("offset", 0),
                                    encoding=s.get("encoding", "utf-8"),
                                    category=s.get("category") or s.get("type", None),
                                    confidence=s.get("confidence", 1.0)
                                ))
                            elif hasattr(s, 'value'):
                                strings_list.append(s)  # Already an ExtractedString
                            else:
                                strings_list.append(re_service.ExtractedString(
                                    value=str(s),
                                    offset=0,
                                    encoding="utf-8"
                                ))
                        except Exception as str_err:
                            logger.debug(f"Skipping malformed string: {str_err}")
                            continue
                    
                    permissions_list = [
                        re_service.ApkPermission(
                            name=p.name, 
                            is_dangerous=p.is_dangerous, 
                            description=p.description
                        ) 
                        for p in apk_result.permissions
                    ]
                    
                    urls_list = apk_result.urls[:100] if apk_result.urls else []
                    
                    # Get dex_analysis if available
                    dex_analysis = None
                    if hasattr(apk_result, 'dex_analysis') and apk_result.dex_analysis:
                        dex_analysis = apk_result.dex_analysis
                    
                    yield make_progress("Generating root detection bypass scripts...", 40)
                    
                    # Generate standard Frida scripts
                    dynamic_analysis = re_service.generate_frida_scripts(
                        package_name=apk_result.package_name,
                        strings=strings_list,
                        dex_analysis=dex_analysis,
                        permissions=permissions_list,
                        urls=urls_list,
                        smali_analysis=None
                    )
                    
                    if dynamic_analysis:
                        result.dynamic_analysis = dynamic_analysis
                        scripts_count = dynamic_analysis.get("total_scripts", 0)
                        logger.info(f"Generated {scripts_count} standard Frida scripts")
                        yield make_progress(f"Generated {scripts_count} standard Frida scripts", 60)
                    
            except Exception as e:
                logger.error(f"Standard Frida script generation failed: {e}")
            
            # Generate vulnerability-specific Frida hooks (based on discovered vulnerabilities)
            yield make_progress("Generating vulnerability-specific Frida hooks...", 70)
            
            try:
                if result.decompiled_code_findings and apk_result:
                    vuln_frida_hooks = re_service.generate_vulnerability_specific_frida_hooks(
                        package_name=apk_result.package_name,
                        decompiled_findings=result.decompiled_code_findings,
                        manifest_analysis=apk_result.manifest_analysis if hasattr(apk_result, 'manifest_analysis') else None
                    )
                    
                    if vuln_frida_hooks and vuln_frida_hooks.get("vulnerability_scripts"):
                        result.vulnerability_frida_hooks = vuln_frida_hooks
                        hook_count = len(vuln_frida_hooks.get("vulnerability_scripts", []))
                        logger.info(f"Generated {hook_count} vulnerability-specific Frida scripts")
            except Exception as e:
                logger.error(f"Vulnerability Frida hook generation failed: {e}")
            
            yield make_progress("Frida hooks generated", 100)
            
            # =================================================================
            # Final Result
            # =================================================================
            result.scan_time = (datetime.now() - start_time).total_seconds()
            
            yield f"data: {json.dumps({'type': 'result', 'data': result.model_dump()})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
            
        except Exception as e:
            logger.error(f"Unified scan failed: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
        finally:
            # Cleanup session (but keep JADX cache for browsing)
            if scan_id in _unified_scan_sessions:
                del _unified_scan_sessions[scan_id]
            # Note: Don't delete tmp_dir yet - JADX cache needs it
    
    return StreamingResponse(
        run_unified_scan(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )


@router.post("/apk/unified-scan/{scan_id}/cancel")
async def cancel_unified_scan(
    scan_id: str,
    current_user: User = Depends(get_current_active_user),
):
    """Cancel an in-progress unified scan."""
    if scan_id in _unified_scan_sessions:
        _unified_scan_sessions[scan_id]["cancelled"] = True
        return {"status": "cancelled"}
    raise HTTPException(status_code=404, detail="Scan session not found")


# ============================================================================
# JADX Decompilation Endpoints
# ============================================================================

class JadxDecompilationResponse(BaseModel):
    """Response for JADX decompilation."""
    package_name: str
    total_classes: int
    total_files: int
    output_directory: str
    decompilation_time: float
    classes: List[Dict[str, Any]]
    source_tree: Dict[str, Any]
    security_issues: List[Dict[str, Any]]
    errors: List[str] = []
    warnings: List[str] = []


class JadxSourceResponse(BaseModel):
    """Response for getting decompiled source."""
    class_name: str
    package_name: str
    file_path: str
    source_code: str
    line_count: int
    is_activity: bool
    is_service: bool
    is_receiver: bool
    is_provider: bool
    extends: Optional[str] = None
    implements: List[str] = []
    methods: List[str] = []
    security_issues: List[Dict[str, Any]] = []


class JadxSearchResponse(BaseModel):
    """Response for searching decompiled sources."""
    query: str
    total_results: int
    results: List[Dict[str, Any]]


# Store JADX output directories for session with timestamps
_jadx_cache: Dict[str, Path] = {}
_jadx_cache_timestamps: Dict[str, datetime] = {}
JADX_CACHE_TTL_HOURS = 2  # Clean up after 2 hours


async def cleanup_old_jadx_cache():
    """Background task to clean up old JADX cache entries."""
    while True:
        try:
            await asyncio.sleep(1800)  # Run every 30 minutes
            now = datetime.now()
            expired_sessions = []
            
            for session_id, timestamp in list(_jadx_cache_timestamps.items()):
                age_hours = (now - timestamp).total_seconds() / 3600
                if age_hours > JADX_CACHE_TTL_HOURS:
                    expired_sessions.append(session_id)
            
            for session_id in expired_sessions:
                if session_id in _jadx_cache:
                    cache_path = _jadx_cache[session_id]
                    if cache_path.exists():
                        try:
                            shutil.rmtree(cache_path, ignore_errors=True)
                            logger.info(f"Cleaned up JADX cache: {session_id}")
                        except Exception as e:
                            logger.warning(f"Failed to cleanup JADX cache {session_id}: {e}")
                    del _jadx_cache[session_id]
                if session_id in _jadx_cache_timestamps:
                    del _jadx_cache_timestamps[session_id]
            
            if expired_sessions:
                logger.info(f"Cleaned up {len(expired_sessions)} old JADX cache entries")
        except Exception as e:
            logger.error(f"JADX cache cleanup error: {e}")


def _resolve_jadx_output_dir(output_directory: str) -> Path:
    """
    Resolve output_directory to an actual Path.
    
    The output_directory can be either:
    1. A JADX session ID (UUID) - looked up from _jadx_cache
    2. A direct filesystem path
    
    Returns:
        Path to the JADX output directory
        
    Raises:
        HTTPException if not found
    """
    # First, check if it's a session ID in the cache
    if output_directory in _jadx_cache:
        return _jadx_cache[output_directory]
    
    # Otherwise, try as a direct path
    output_path = Path(output_directory)
    if output_path.exists():
        return output_path
    
    # Check if it's a session ID that looks like a path (backwards compatibility)
    # Sometimes the frontend might send the session ID
    for session_id, cached_path in _jadx_cache.items():
        if str(cached_path) == output_directory or session_id in output_directory:
            return cached_path
    
    raise HTTPException(
        status_code=404, 
        detail=f"Decompiled sources not found. Session may have expired. Please run the scan again."
    )


@router.post("/apk/decompile", response_model=JadxDecompilationResponse)
async def decompile_apk(
    file: UploadFile = File(..., description="APK file to decompile"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Decompile an APK to Java source code using JADX.
    
    Returns:
    - Decompiled Java classes with metadata
    - Source code tree structure
    - Security issues found in code
    
    Note: This can take a while for large APKs.
    """
    if not check_jadx_available():
        raise HTTPException(
            status_code=503,
            detail="JADX is not available. Please ensure JADX is installed."
        )
    
    filename = file.filename or "unknown.apk"
    suffix = Path(filename).suffix.lower()
    
    if suffix not in ALLOWED_APK_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {', '.join(ALLOWED_APK_EXTENSIONS)}"
        )
    
    tmp_dir = None
    try:
        # Save file to temp location with sanitized filename to prevent path traversal
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_jadx_"))
        safe_filename = sanitize_filename(filename)
        tmp_path = tmp_dir / safe_filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Decompiling APK with JADX: {filename} ({file_size:,} bytes)")
        
        # Run JADX decompilation
        result = re_service.decompile_apk_with_jadx(tmp_path)
        
        # Store output directory for later queries with timestamp
        import uuid
        session_id = str(uuid.uuid4())
        _jadx_cache[session_id] = Path(result.output_directory)
        _jadx_cache_timestamps[session_id] = datetime.now()
        
        # Collect security issues from all classes
        all_security_issues = []
        for cls in result.classes:
            all_security_issues.extend(cls.security_issues)
        
        # Limit classes in response for performance
        classes_summary = [
            {
                "class_name": c.class_name,
                "package_name": c.package_name,
                "file_path": c.file_path,
                "line_count": c.line_count,
                "is_activity": c.is_activity,
                "is_service": c.is_service,
                "is_receiver": c.is_receiver,
                "is_provider": c.is_provider,
                "extends": c.extends,
                "security_issues_count": len(c.security_issues),
            }
            for c in result.classes[:500]  # Limit to 500 classes
        ]
        
        return JadxDecompilationResponse(
            package_name=result.package_name,
            total_classes=result.total_classes,
            total_files=result.total_files,
            output_directory=session_id,  # Return session ID, not actual path
            decompilation_time=result.decompilation_time,
            classes=classes_summary,
            source_tree=result.source_tree,
            security_issues=all_security_issues[:100],  # Limit issues
            errors=result.errors,
            warnings=result.warnings[:20],
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"JADX decompilation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Decompilation failed: {str(e)}")


@router.get("/apk/decompile/{session_id}/source/{class_path:path}", response_model=JadxSourceResponse)
async def get_decompiled_source(
    session_id: str,
    class_path: str,
    current_user: User = Depends(get_current_active_user),
):
    """
    Get the decompiled Java source code for a specific class.
    
    Args:
        session_id: The session ID from decompilation
        class_path: Path to the Java file (e.g., "com/example/MainActivity.java")
    """
    if session_id not in _jadx_cache:
        raise HTTPException(status_code=404, detail="Decompilation session not found. Please decompile the APK again.")
    
    output_dir = _jadx_cache[session_id]
    
    source_code = re_service.get_jadx_class_source(output_dir, class_path)
    
    if source_code is None:
        raise HTTPException(status_code=404, detail=f"Class not found: {class_path}")
    
    # Parse class info
    class_info = re_service._parse_java_class(source_code, class_path)
    
    return JadxSourceResponse(
        class_name=class_info.class_name,
        package_name=class_info.package_name,
        file_path=class_path,
        source_code=source_code,
        line_count=class_info.line_count,
        is_activity=class_info.is_activity,
        is_service=class_info.is_service,
        is_receiver=class_info.is_receiver,
        is_provider=class_info.is_provider,
        extends=class_info.extends,
        implements=class_info.implements,
        methods=class_info.methods,
        security_issues=class_info.security_issues,
    )


@router.get("/apk/decompile/{session_id}/search", response_model=JadxSearchResponse)
async def search_decompiled_sources(
    session_id: str,
    query: str = Query(..., min_length=2, max_length=100, description="Search string"),
    max_results: int = Query(50, ge=1, le=200),
    current_user: User = Depends(get_current_active_user),
):
    """
    Search for a string in decompiled Java sources.
    
    Useful for finding:
    - API endpoints and URLs
    - Method names
    - Hardcoded strings
    - Security-sensitive patterns
    """
    if session_id not in _jadx_cache:
        raise HTTPException(status_code=404, detail="Decompilation session not found.")
    
    output_dir = _jadx_cache[session_id]
    
    results = re_service.search_jadx_sources(output_dir, query, max_results)
    
    return JadxSearchResponse(
        query=query,
        total_results=len(results),
        results=results,
    )


@router.delete("/apk/decompile/{session_id}")
async def cleanup_decompilation(
    session_id: str,
    current_user: User = Depends(get_current_active_user),
):
    """Clean up decompiled sources to free disk space."""
    if session_id not in _jadx_cache:
        raise HTTPException(status_code=404, detail="Session not found.")
    
    output_dir = _jadx_cache[session_id]
    
    try:
        if output_dir.exists():
            shutil.rmtree(output_dir, ignore_errors=True)
        del _jadx_cache[session_id]
        return {"message": "Decompilation session cleaned up", "session_id": session_id}
    except Exception as e:
        logger.error(f"Failed to clean up JADX session: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class AiDiagramResponse(BaseModel):
    """Response for AI-generated Mermaid diagrams."""
    session_id: str
    architecture_diagram: Optional[str] = None
    data_flow_diagram: Optional[str] = None
    generation_time: float
    error: Optional[str] = None


@router.post("/apk/decompile/{session_id}/ai-diagrams", response_model=AiDiagramResponse)
async def generate_ai_diagrams_from_decompilation(
    session_id: str,
    include_architecture: bool = Query(True, description="Generate architecture diagram"),
    include_data_flow: bool = Query(True, description="Generate data flow diagram"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate AI-powered Mermaid diagrams from decompiled APK sources.
    
    Uses Gemini AI to analyze the decompiled Java source code and generate:
    - **Architecture Diagram**: Shows app components, activities, services, and their relationships
    - **Data Flow Diagram**: Shows how data moves through the app, including privacy-sensitive data
    
    The diagrams use Iconify icons for better visual representation:
    - fa6-brands:android, mdi:application, mdi:cog for components
    - fa6-solid:shield, fa6-solid:lock, fa6-solid:bug for security elements
    
    Note: Requires GEMINI_API_KEY to be configured.
    """
    import time
    start_time = time.time()
    
    if session_id not in _jadx_cache:
        raise HTTPException(status_code=404, detail="Decompilation session not found. Please decompile the APK first.")
    
    output_dir = _jadx_cache[session_id]
    
    if not output_dir.exists():
        raise HTTPException(status_code=404, detail="Decompilation output no longer exists.")
    
    try:
        # Get JADX result summary for AI context
        jadx_result = re_service.get_jadx_result_summary(output_dir)
        
        # Create a minimal ApkAnalysisResult for the diagram generators
        # We'll populate it from the JADX decompilation info
        from services.reverse_engineering_service import ApkAnalysisResult, ApkPermission, ApkComponent, ExtractedString
        
        result = ApkAnalysisResult(
            filename=f"{jadx_result.get('package_name', 'unknown')}.apk",
            package_name=jadx_result.get('package_name', 'unknown'),
            version_name=None,
            version_code=None,
            min_sdk=None,
            target_sdk=None,
            permissions=[],
            components=[
                ApkComponent(
                    name=cls.get('class_name', ''),
                    component_type='activity' if cls.get('is_activity') else 
                                   'service' if cls.get('is_service') else 
                                   'receiver' if cls.get('is_receiver') else 
                                   'provider' if cls.get('is_provider') else 'class',
                    is_exported=False,
                    intent_filters=[]
                )
                for cls in jadx_result.get('classes', [])[:100]  # Limit for context
            ],
            strings=[],
            secrets=[],
            urls=[],
            native_libraries=[],
            activities=[cls.get('class_name', '') for cls in jadx_result.get('classes', []) if cls.get('is_activity')],
            services=[cls.get('class_name', '') for cls in jadx_result.get('classes', []) if cls.get('is_service')],
            receivers=[cls.get('class_name', '') for cls in jadx_result.get('classes', []) if cls.get('is_receiver')],
            providers=[cls.get('class_name', '') for cls in jadx_result.get('classes', []) if cls.get('is_provider')],
        )
        
        architecture_diagram = None
        data_flow_diagram = None
        
        # Generate diagrams (these are async functions)
        if include_architecture:
            architecture_diagram = await re_service.generate_ai_architecture_diagram(result, jadx_result, output_dir=output_dir)
        
        if include_data_flow:
            data_flow_diagram = await re_service.generate_ai_data_flow_diagram(result, jadx_result, output_dir=output_dir)
        
        generation_time = time.time() - start_time
        
        return AiDiagramResponse(
            session_id=session_id,
            architecture_diagram=architecture_diagram,
            data_flow_diagram=data_flow_diagram,
            generation_time=generation_time,
        )
        
    except Exception as e:
        logger.error(f"AI diagram generation failed: {e}")
        return AiDiagramResponse(
            session_id=session_id,
            generation_time=time.time() - start_time,
            error=str(e),
        )


# ============================================================================
# AI Code Analysis Endpoints
# ============================================================================

class AICodeExplanationRequest(BaseModel):
    """Request for AI code explanation."""
    source_code: str
    class_name: str
    explanation_type: str = "general"  # general, security, method
    method_name: Optional[str] = None

class AICodeExplanationResponse(BaseModel):
    """Response with AI explanation."""
    class_name: str
    explanation_type: str
    explanation: str
    key_points: List[str]
    security_concerns: List[Dict[str, Any]]
    method_name: Optional[str] = None

class AIVulnerabilityAnalysisRequest(BaseModel):
    """Request for AI vulnerability analysis."""
    source_code: str
    class_name: str

class AIVulnerabilityAnalysisResponse(BaseModel):
    """Response with vulnerability analysis."""
    class_name: str
    risk_level: str  # critical, high, medium, low, info
    vulnerabilities: List[Dict[str, Any]]
    recommendations: List[str]
    exploitation_scenarios: List[str]
    summary: str


@router.post("/apk/decompile/ai/explain", response_model=AICodeExplanationResponse)
async def explain_code_with_ai(
    request: AICodeExplanationRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Use AI to explain decompiled Java/Kotlin code.
    
    Explanation types:
    - general: What does this class/code do?
    - security: Security-focused analysis
    - method: Explain a specific method
    """
    if not settings.gemini_api_key:
        raise HTTPException(status_code=503, detail="AI features require Gemini API key")
    
    try:
        result = await re_service.explain_code_with_ai(
            source_code=request.source_code,
            class_name=request.class_name,
            explanation_type=request.explanation_type,
            method_name=request.method_name
        )
        return AICodeExplanationResponse(**result)
    except Exception as e:
        logger.error(f"AI code explanation failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")


@router.post("/apk/decompile/ai/vulnerabilities", response_model=AIVulnerabilityAnalysisResponse)
async def analyze_vulnerabilities_with_ai(
    request: AIVulnerabilityAnalysisRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Use AI to perform deep vulnerability analysis on decompiled code.
    
    Returns:
    - Identified vulnerabilities with severity
    - Exploitation scenarios
    - Fix recommendations
    """
    if not settings.gemini_api_key:
        raise HTTPException(status_code=503, detail="AI features require Gemini API key")
    
    try:
        result = await re_service.analyze_code_vulnerabilities_with_ai(
            source_code=request.source_code,
            class_name=request.class_name
        )
        return AIVulnerabilityAnalysisResponse(**result)
    except Exception as e:
        logger.error(f"AI vulnerability analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")


# ============================================================================
# Data Flow Analysis Endpoints
# ============================================================================

class DataFlowSourceSink(BaseModel):
    """A data source or sink in the flow analysis."""
    type: str
    pattern: Optional[str] = None
    line: int
    code: str
    variable: Optional[str] = None


class DataFlow(BaseModel):
    """A data flow from source to sink."""
    source: Dict[str, Any]
    sink: Dict[str, Any]
    risk: str


class DataFlowSummary(BaseModel):
    """Summary of data flow analysis."""
    total_sources: int
    total_sinks: int
    potential_leaks: int
    risk_level: str


class DataFlowAnalysisRequest(BaseModel):
    """Request to analyze data flow in code."""
    source_code: str
    class_name: str


class DataFlowAnalysisResponse(BaseModel):
    """Response from data flow analysis."""
    class_name: str
    sources: List[Dict[str, Any]]
    sinks: List[Dict[str, Any]]
    flows: List[DataFlow]
    risk_flows: List[DataFlow]
    summary: DataFlowSummary


@router.post("/apk/decompile/dataflow", response_model=DataFlowAnalysisResponse)
async def analyze_data_flow(
    request: DataFlowAnalysisRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Analyze data flow in decompiled Java/Kotlin code.
    
    Performs lightweight taint analysis to track:
    - Data sources (user input, files, network, sensors)
    - Data sinks (logging, network, storage, IPC)
    - Potential data leakage paths
    """
    try:
        result = re_service.analyze_decompiled_code_data_flow(
            source_code=request.source_code,
            class_name=request.class_name
        )
        return DataFlowAnalysisResponse(**result)
    except Exception as e:
        logger.error(f"Data flow analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Data flow analysis failed: {str(e)}")


# ============================================================================
# Method Call Graph Endpoints
# ============================================================================

class MethodInfo(BaseModel):
    """Information about a method."""
    name: str
    return_type: str
    parameters: List[Dict[str, str]]
    line_start: int
    line_end: int
    is_entry_point: bool
    calls: List[Dict[str, Any]]
    called_by: List[str] = []
    modifiers: List[str] = []


class CallInfo(BaseModel):
    """Information about a method call."""
    caller: str
    caller_line: int
    callee: str
    callee_class: str
    line: int
    is_internal: bool


class GraphNode(BaseModel):
    """A node in the call graph."""
    id: str
    label: str
    type: str
    is_entry_point: bool
    line: Optional[int] = None


class GraphEdge(BaseModel):
    """An edge in the call graph."""
    source: str = Field(alias="from")
    target: str = Field(alias="to")
    label: str

    class Config:
        populate_by_name = True


class CallGraphStatistics(BaseModel):
    """Statistics about the call graph."""
    total_methods: int
    total_internal_calls: int
    total_external_calls: int
    max_depth: int
    cyclomatic_complexity: int


class CallGraphRequest(BaseModel):
    """Request to build method call graph."""
    source_code: str
    class_name: str


class CallGraphResponse(BaseModel):
    """Response from call graph analysis."""
    class_name: str
    methods: List[MethodInfo]
    calls: List[CallInfo]
    entry_points: List[Dict[str, Any]]
    external_calls: List[CallInfo]
    graph: Dict[str, Any]
    statistics: CallGraphStatistics


@router.post("/apk/decompile/callgraph", response_model=CallGraphResponse)
async def build_call_graph(
    request: CallGraphRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Build a method call graph from decompiled Java/Kotlin code.
    
    Returns:
    - Method definitions and signatures
    - Internal and external method calls
    - Entry points (lifecycle methods, callbacks)
    - Graph structure for visualization
    - Code complexity statistics
    """
    try:
        result = re_service.build_call_graph(
            source_code=request.source_code,
            class_name=request.class_name
        )
        return CallGraphResponse(**result)
    except Exception as e:
        logger.error(f"Call graph analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Call graph analysis failed: {str(e)}")


# ============================================================================
# Smart Search Endpoints
# ============================================================================

class SmartSearchMatch(BaseModel):
    """A match from smart search."""
    file: str
    line: int
    code: str
    match: str
    context: Optional[str] = None
    vuln_type: Optional[str] = None
    description: Optional[str] = None
    severity: Optional[str] = None


class VulnSummaryItem(BaseModel):
    """Vulnerability summary item."""
    count: int
    severity: str
    description: str


class SmartSearchRequest(BaseModel):
    """Request for smart search."""
    output_directory: str
    query: str
    search_type: str = "smart"  # smart, vuln, regex, exact
    max_results: int = 100


class SmartSearchResponse(BaseModel):
    """Response from smart search."""
    query: str
    search_type: str
    total_matches: int
    files_searched: int
    matches: List[SmartSearchMatch]
    vulnerability_summary: Dict[str, VulnSummaryItem] = {}
    expanded_terms: List[str] = []
    suggestions: List[str] = []
    error: Optional[str] = None


@router.post("/apk/decompile/smart-search", response_model=SmartSearchResponse)
async def smart_search(
    request: SmartSearchRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Perform smart/semantic search across decompiled sources.
    
    Search types:
    - smart: Expands query with related security terms
    - vuln: Searches for vulnerability patterns  
    - regex: Direct regex search
    - exact: Exact string match
    
    Returns matches with context and vulnerability classification.
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.smart_search(
            output_dir=output_dir,
            query=request.query,
            search_type=request.search_type,
            max_results=request.max_results
        )
        return SmartSearchResponse(**result)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Smart search failed: {e}")
        raise HTTPException(status_code=500, detail=f"Smart search failed: {str(e)}")


# ============================================================================
# AI Vulnerability Scan Endpoints
# ============================================================================

class AIVulnScanVulnerability(BaseModel):
    """A vulnerability from AI scan."""
    id: str = ""
    title: str
    severity: str
    category: str = ""
    affected_class: str = ""
    affected_method: str = ""
    description: str
    code_snippet: str = ""
    impact: str = ""
    remediation: str = ""
    cwe_id: str = ""


class AIVulnScanAttackChain(BaseModel):
    """An attack chain from AI scan."""
    name: str
    steps: List[str]
    impact: str
    likelihood: str = "medium"


class AIVulnScanRiskSummary(BaseModel):
    """Risk summary from AI scan."""
    critical: int = 0
    high: int = 0
    medium: int = 0
    low: int = 0


class AIVulnScanRequest(BaseModel):
    """Request for AI vulnerability scan."""
    output_directory: str
    scan_type: str = "quick"  # quick, deep, focused
    focus_areas: List[str] = []  # auth, crypto, network, storage


class AIVulnScanResponse(BaseModel):
    """Response from AI vulnerability scan."""
    scan_type: str
    focus_areas: List[str]
    classes_scanned: int
    vulnerabilities: List[AIVulnScanVulnerability]
    risk_summary: AIVulnScanRiskSummary
    attack_chains: List[AIVulnScanAttackChain] = []
    recommendations: List[str] = []
    summary: str
    overall_risk: str = "low"
    error: Optional[str] = None


@router.post("/apk/decompile/ai-vulnscan", response_model=AIVulnScanResponse)
async def ai_vulnerability_scan(
    request: AIVulnScanRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Perform AI-powered vulnerability scan across multiple classes.
    
    Scan types:
    - quick: Scan key classes (activities, services, network) - ~10 classes
    - deep: Scan all relevant classes - ~25 classes
    - focused: Scan specific areas (auth, crypto, network, storage)
    
    Returns comprehensive vulnerability analysis with attack chains.
    """
    if not settings.gemini_api_key:
        raise HTTPException(status_code=503, detail="AI features require Gemini API key")
    
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = await re_service.ai_vulnerability_scan(
            output_dir=output_dir,
            scan_type=request.scan_type,
            focus_areas=request.focus_areas if request.focus_areas else None
        )
        return AIVulnScanResponse(**result)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI vulnerability scan failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI vulnerability scan failed: {str(e)}")


# ============================================================================
# Library CVE Scan Endpoint
# ============================================================================

class LibraryCVEScanRequest(BaseModel):
    """Request for library CVE scan."""
    output_directory: str
    gradle_content: Optional[str] = None  # Optional gradle file content


class DetectedLibrary(BaseModel):
    """A detected library in the APK."""
    package_prefix: str
    maven_coordinate: str
    ecosystem: str
    version: Optional[str] = None
    class_count: int = 0
    is_high_risk: bool = False
    risk_reason: Optional[str] = None


class LibraryCVE(BaseModel):
    """A CVE found in a library."""
    library: str
    library_version: str
    cve_id: str
    aliases: List[str] = []
    summary: str
    details: str = ""
    severity: str
    cvss_score: Optional[float] = None
    affected_versions: List[str] = []
    references: List[str] = []
    published: str = ""
    modified: str = ""
    exploitation_potential: str = ""
    attack_vector: str = ""


class LibraryCVEScanResponse(BaseModel):
    """Response from library CVE scan."""
    total_libraries: int
    high_risk_libraries: int
    total_cves: int
    critical_cves: int
    high_cves: int
    libraries: List[DetectedLibrary]
    cves: List[LibraryCVE]
    error: Optional[str] = None


@router.post("/apk/decompile/library-cve-scan", response_model=LibraryCVEScanResponse)
async def library_cve_scan(
    request: LibraryCVEScanRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Scan decompiled APK for third-party libraries and look up known CVEs.
    
    This endpoint:
    1. Extracts class names from decompiled source
    2. Identifies third-party libraries by package prefixes
    3. Looks up CVEs for each library via OSV.dev API
    4. Returns exploitation-focused vulnerability assessment
    
    Use this to identify known vulnerabilities in bundled dependencies.
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        # Get all class names from decompiled source
        class_names = []
        sources_dir = output_dir / "sources"
        
        if sources_dir.exists():
            for java_file in sources_dir.rglob("*.java"):
                # Convert file path to class name
                try:
                    rel_path = java_file.relative_to(sources_dir)
                    class_name = str(rel_path).replace("/", ".").replace("\\", ".").replace(".java", "")
                    class_names.append(class_name)
                except Exception:
                    pass
        
        if not class_names:
            return LibraryCVEScanResponse(
                total_libraries=0,
                high_risk_libraries=0,
                total_cves=0,
                critical_cves=0,
                high_cves=0,
                libraries=[],
                cves=[],
                error="No decompiled classes found. Run JADX decompilation first."
            )
        
        # Extract dependencies from class names
        from backend.services.reverse_engineering_service import extract_apk_dependencies, lookup_apk_cves
        
        libraries = extract_apk_dependencies(class_names, request.gradle_content)
        
        # Convert to response format
        lib_responses = [
            DetectedLibrary(
                package_prefix=lib.package_prefix,
                maven_coordinate=lib.maven_coordinate,
                ecosystem=lib.ecosystem,
                version=lib.version,
                class_count=lib.class_count,
                is_high_risk=lib.is_high_risk,
                risk_reason=lib.risk_reason,
            )
            for lib in libraries
        ]
        
        # Look up CVEs for detected libraries
        cves_raw = await lookup_apk_cves(libraries)
        
        # Convert to response format
        cve_responses = [
            LibraryCVE(
                library=cve.get('library', ''),
                library_version=cve.get('library_version', 'unknown'),
                cve_id=cve.get('cve_id', ''),
                aliases=cve.get('aliases', []),
                summary=cve.get('summary', ''),
                details=cve.get('details', ''),
                severity=cve.get('severity', 'unknown'),
                cvss_score=cve.get('cvss_score'),
                affected_versions=cve.get('affected_versions', []),
                references=cve.get('references', []),
                published=cve.get('published', ''),
                modified=cve.get('modified', ''),
                exploitation_potential=cve.get('exploitation_potential', ''),
                attack_vector=cve.get('attack_vector', ''),
            )
            for cve in cves_raw
        ]
        
        # Calculate counts
        high_risk_libs = sum(1 for lib in libraries if lib.is_high_risk)
        critical_cves = sum(1 for cve in cves_raw if cve.get('severity', '').lower() == 'critical')
        high_cves = sum(1 for cve in cves_raw if cve.get('severity', '').lower() == 'high')
        
        logger.info(f"Library CVE scan: {len(libraries)} libraries, {len(cves_raw)} CVEs found")
        
        return LibraryCVEScanResponse(
            total_libraries=len(libraries),
            high_risk_libraries=high_risk_libs,
            total_cves=len(cves_raw),
            critical_cves=critical_cves,
            high_cves=high_cves,
            libraries=lib_responses,
            cves=cve_responses,
        )
        
    except Exception as e:
        logger.error(f"Library CVE scan failed: {e}")
        raise HTTPException(status_code=500, detail=f"Library CVE scan failed: {str(e)}")


# ============================================================================
# Enhanced Security Analysis Endpoint (Combined Pattern + AI + CVE)
# ============================================================================

class EnhancedSecurityFinding(BaseModel):
    """A unified security finding from any detection method."""
    source: str  # "pattern", "ai", or "cve"
    detection_method: str
    title: str
    severity: str
    category: str = ""
    affected_class: str = ""
    affected_method: str = ""
    description: str = ""
    code_snippet: str = ""
    line_number: int = 0
    impact: str = ""
    remediation: str = ""
    cve_id: str = ""
    cvss_score: Optional[float] = None
    cwe_id: str = ""
    affected_library: str = ""
    attack_vector: str = ""
    exploitation_potential: str = ""
    references: List[str] = []


class AttackChain(BaseModel):
    """A multi-step attack scenario."""
    name: str
    steps: List[str]
    impact: str
    likelihood: str = "medium"
    classes_involved: List[str] = []


class EnhancedSecurityRiskSummary(BaseModel):
    """Risk summary from enhanced security analysis."""
    critical: int = 0
    high: int = 0
    medium: int = 0
    low: int = 0
    info: int = 0


class EnhancedSecurityMetadata(BaseModel):
    """Metadata about the security analysis."""
    pattern_scan_enabled: bool = True
    ai_scan_enabled: bool = True
    cve_lookup_enabled: bool = True
    classes_scanned: int = 0
    libraries_detected: int = 0
    cves_found: int = 0
    ai_scan_error: Optional[str] = None
    cve_lookup_error: Optional[str] = None


class EnhancedSecurityRequest(BaseModel):
    """Request for enhanced security analysis."""
    output_directory: str
    include_ai_scan: bool = True
    include_cve_lookup: bool = True
    ai_scan_type: str = "quick"  # quick, deep, focused


class EnhancedSecurityResponse(BaseModel):
    """Combined security analysis response."""
    pattern_findings: List[Dict[str, Any]] = []
    ai_findings: List[EnhancedSecurityFinding] = []
    cve_findings: List[EnhancedSecurityFinding] = []
    combined_findings: List[EnhancedSecurityFinding] = []
    attack_chains: List[AttackChain] = []
    risk_summary: EnhancedSecurityRiskSummary
    overall_risk: str = "low"
    recommendations: List[str] = []
    executive_summary: str = ""
    analysis_metadata: EnhancedSecurityMetadata
    error: Optional[str] = None


@router.post("/apk/decompile/enhanced-security", response_model=EnhancedSecurityResponse)
async def enhanced_security_scan(
    request: EnhancedSecurityRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Perform comprehensive security analysis combining multiple detection methods.
    
    This unified scan combines:
    1. **Pattern-based Detection** (fast) - Regex scanning for known vulnerability patterns
    2. **AI-powered Analysis** (thorough) - Cross-class vulnerability detection using Gemini AI
    3. **CVE Lookup** (authoritative) - Known vulnerabilities in third-party libraries via OSV.dev
    
    Returns deduplicated, prioritized findings with attack chains and recommendations.
    
    Scan options:
    - include_ai_scan: Enable AI cross-class analysis (requires Gemini API key)
    - include_cve_lookup: Enable CVE lookup for dependencies
    - ai_scan_type: "quick" (fast), "deep" (thorough), or "focused" (specific areas)
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = await re_service.enhanced_security_analysis(
            output_dir=output_dir,
            include_ai_scan=request.include_ai_scan,
            include_cve_lookup=request.include_cve_lookup,
            ai_scan_type=request.ai_scan_type
        )
        
        if "error" in result and result.get("combined_findings") is None:
            return EnhancedSecurityResponse(
                risk_summary=EnhancedSecurityRiskSummary(),
                analysis_metadata=EnhancedSecurityMetadata(),
                error=result["error"]
            )
        
        return EnhancedSecurityResponse(
            pattern_findings=result.get("pattern_findings", []),
            ai_findings=[EnhancedSecurityFinding(**f) for f in result.get("ai_findings", []) if isinstance(f, dict)],
            cve_findings=[EnhancedSecurityFinding(**f) for f in result.get("cve_findings", []) if isinstance(f, dict)],
            combined_findings=[EnhancedSecurityFinding(**f) if isinstance(f, dict) else f for f in result.get("combined_findings", [])],
            attack_chains=[AttackChain(**c) if isinstance(c, dict) else c for c in result.get("attack_chains", [])],
            risk_summary=EnhancedSecurityRiskSummary(**result.get("risk_summary", {})),
            overall_risk=result.get("overall_risk", "low"),
            recommendations=result.get("recommendations", []),
            executive_summary=result.get("executive_summary", ""),
            analysis_metadata=EnhancedSecurityMetadata(**result.get("analysis_metadata", {})),
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Enhanced security scan failed: {e}")
        raise HTTPException(status_code=500, detail=f"Enhanced security scan failed: {str(e)}")


# ============================================================================
# Smali View Endpoints
# ============================================================================

class SmaliInstruction(BaseModel):
    """A Smali instruction."""
    method: str
    instruction: str
    category: str


class SmaliBytecodStats(BaseModel):
    """Smali bytecode statistics."""
    invocations: Dict[str, int] = {}
    field_ops: Dict[str, int] = {}
    control_flow: Dict[str, int] = {}
    suspicious_ops: Dict[str, int] = {}


class SmaliViewRequest(BaseModel):
    """Request for Smali view."""
    output_directory: str
    class_path: str


class SmaliViewResponse(BaseModel):
    """Response with Smali bytecode."""
    class_path: str
    smali_code: str
    bytecode_stats: SmaliBytecodStats = SmaliBytecodStats()
    registers_used: int = 0
    method_count: int = 0
    field_count: int = 0
    instructions: List[SmaliInstruction] = []
    is_pseudo: bool = False
    error: Optional[str] = None


@router.post("/apk/decompile/smali", response_model=SmaliViewResponse)
async def get_smali_view(
    request: SmaliViewRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Get Smali bytecode view for a class.
    
    Returns the Dalvik bytecode (Smali) representation of a class,
    which shows low-level operations and is useful for:
    - Analyzing obfuscated code
    - Understanding actual runtime behavior
    - Finding hidden functionality
    - Patching/modifying APKs
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.get_smali_for_class(output_dir, request.class_path)
        
        if result is None:
            return SmaliViewResponse(
                class_path=request.class_path,
                smali_code="# Smali not available\n# Try using baksmali on the original APK",
                is_pseudo=True,
                error="Smali bytecode not available for this class"
            )
        
        return SmaliViewResponse(
            class_path=result["class_path"],
            smali_code=result["smali_code"],
            bytecode_stats=SmaliBytecodStats(**result.get("bytecode_stats", {})) if isinstance(result.get("bytecode_stats"), dict) and "invocations" in result.get("bytecode_stats", {}) else SmaliBytecodStats(),
            registers_used=result.get("registers_used", 0),
            method_count=result.get("method_count", 0),
            field_count=result.get("field_count", 0),
            instructions=[SmaliInstruction(**i) for i in result.get("instructions", [])],
            is_pseudo=result.get("is_pseudo", False),
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Smali view failed: {e}")
        raise HTTPException(status_code=500, detail=f"Smali view failed: {str(e)}")


# ============================================================================
# String Extraction Endpoints
# ============================================================================

class ExtractedString(BaseModel):
    """An extracted string with metadata."""
    value: str
    file: str
    line: int
    categories: List[str]
    severity: str
    length: int
    is_resource: bool = False


class StringExtractionRequest(BaseModel):
    """Request for string extraction."""
    output_directory: str
    filters: Optional[List[str]] = None  # url, api_key, password, etc.


class StringExtractionResponse(BaseModel):
    """Response with extracted strings."""
    total_strings: int
    files_scanned: int
    strings: List[ExtractedString]
    stats: Dict[str, int]
    severity_counts: Dict[str, int]
    top_categories: List[List[Any]]
    error: Optional[str] = None


@router.post("/apk/decompile/strings", response_model=StringExtractionResponse)
async def extract_strings(
    request: StringExtractionRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Extract and categorize all strings from decompiled sources.
    
    Automatically classifies strings into categories:
    - url: HTTP/HTTPS URLs
    - api_key: API keys and secrets
    - password: Hardcoded passwords
    - firebase: Firebase URLs and keys
    - sql_query: SQL queries
    - file_path: File system paths
    - ip_address: IP addresses
    - email: Email addresses
    - jwt: JSON Web Tokens
    - And more...
    
    Use filters to narrow results (e.g., ["url", "api_key"]).
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.extract_all_strings(output_dir, request.filters)
        
        if "error" in result:
            return StringExtractionResponse(
                total_strings=0,
                files_scanned=0,
                strings=[],
                stats={},
                severity_counts={},
                top_categories=[],
                error=result["error"]
            )
        
        return StringExtractionResponse(
            total_strings=result["total_strings"],
            files_scanned=result["files_scanned"],
            strings=[ExtractedString(**s) for s in result["strings"]],
            stats=result["stats"],
            severity_counts=result["severity_counts"],
            top_categories=result["top_categories"],
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"String extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"String extraction failed: {str(e)}")


# ============================================================================
# Cross-Reference (XREF) Endpoints
# ============================================================================

class XrefCaller(BaseModel):
    """A caller reference."""
    class_name: str = Field(alias="class")
    file: str
    method: str
    line: int

    class Config:
        populate_by_name = True


class XrefCallee(BaseModel):
    """An outgoing call reference."""
    method: str
    object: str
    line: int


class XrefMethod(BaseModel):
    """Method with cross-references."""
    name: str
    return_type: str
    params: str
    signature: str
    line: int
    callers: List[Dict[str, Any]] = []
    callees: List[XrefCallee] = []
    caller_count: int = 0
    callee_count: int = 0


class XrefField(BaseModel):
    """Field with cross-references."""
    name: str
    type: str
    line: int
    readers: List[Dict[str, Any]] = []
    writers: List[Dict[str, Any]] = []
    read_count: int = 0
    write_count: int = 0


class XrefStatistics(BaseModel):
    """Cross-reference statistics."""
    method_count: int
    field_count: int
    total_incoming_refs: int
    total_outgoing_refs: int
    is_heavily_used: bool
    is_hub_class: bool


class CrossReferenceRequest(BaseModel):
    """Request for cross-references."""
    output_directory: str
    class_path: str


class CrossReferenceResponse(BaseModel):
    """Response with cross-references."""
    class_name: str
    package: str
    file_path: str
    methods: List[XrefMethod]
    fields: List[XrefField]
    statistics: XrefStatistics
    summary: str
    error: Optional[str] = None


@router.post("/apk/decompile/xref", response_model=CrossReferenceResponse)
async def get_cross_references(
    request: CrossReferenceRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Build cross-references for a class.
    
    Returns:
    - All methods defined in the class
    - Who calls each method (incoming references)
    - What each method calls (outgoing references)
    - Field read/write references
    - Statistics about class usage
    
    Useful for:
    - Understanding how a class is used
    - Finding entry points to functionality
    - Tracing data flow through the app
    - Identifying critical/hub classes
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.build_cross_references(output_dir, request.class_path)
        
        if "error" in result:
            raise HTTPException(status_code=404, detail=result["error"])
        
        return CrossReferenceResponse(
            class_name=result["class_name"],
            package=result["package"],
            file_path=result["file_path"],
            methods=[XrefMethod(
                name=m["name"],
                return_type=m["return_type"],
                params=m["params"],
                signature=m["signature"],
                line=m["line"],
                callers=m["callers"],
                callees=[XrefCallee(**c) for c in m["callees"]],
                caller_count=m["caller_count"],
                callee_count=m["callee_count"],
            ) for m in result["methods"]],
            fields=[XrefField(
                name=f["name"],
                type=f["type"],
                line=f["line"],
                readers=f["readers"],
                writers=f["writers"],
                read_count=f["read_count"],
                write_count=f["write_count"],
            ) for f in result["fields"]],
            statistics=XrefStatistics(**result["statistics"]),
            summary=result["summary"],
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Cross-reference failed: {e}")
        raise HTTPException(status_code=500, detail=f"Cross-reference failed: {str(e)}")


# ============================================================================
# Download Project ZIP Endpoint
# ============================================================================

class ProjectZipInfoResponse(BaseModel):
    """Information about project ZIP."""
    total_files: int
    total_size_bytes: int
    total_size_mb: float
    file_types: Dict[str, int]
    estimated_zip_size_mb: float
    error: Optional[str] = None


class DownloadProjectRequest(BaseModel):
    """Request to download project as ZIP."""
    output_directory: str


@router.post("/apk/decompile/zip-info", response_model=ProjectZipInfoResponse)
async def get_project_zip_info(
    request: DownloadProjectRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Get information about what would be in the project ZIP.
    
    Returns file counts, sizes, and estimated download size.
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.get_project_zip_info(output_dir)
        
        if "error" in result:
            raise HTTPException(status_code=404, detail=result["error"])
        
        return ProjectZipInfoResponse(**result)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get ZIP info failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/apk/decompile/download-zip")
async def download_project_zip(
    request: DownloadProjectRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Create and download the decompiled project as a ZIP file.
    
    Returns the ZIP file as a downloadable response.
    """
    from fastapi.responses import FileResponse
    
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        zip_path = re_service.create_project_zip(output_dir)
        
        return FileResponse(
            path=str(zip_path),
            filename=zip_path.name,
            media_type="application/zip",
            headers={
                "Content-Disposition": f"attachment; filename={zip_path.name}"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Create ZIP failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create ZIP: {str(e)}")


# ============================================================================
# Permission Analyzer Endpoint
# ============================================================================

class PermissionInfo(BaseModel):
    """Information about a single permission."""
    name: str
    short_name: str
    level: str  # dangerous, normal, signature, deprecated, unknown
    description: str
    category: str


class DangerousCombination(BaseModel):
    """A dangerous permission combination."""
    permissions: List[str]
    risk: str
    description: str


class PermissionAnalysisResponse(BaseModel):
    """Permission analysis results."""
    total_permissions: int
    permissions: List[PermissionInfo]
    by_level: Dict[str, List[PermissionInfo]]
    by_category: Dict[str, List[PermissionInfo]]
    dangerous_combinations: List[DangerousCombination]
    risk_score: int
    overall_risk: str
    summary: str
    error: Optional[str] = None


class PermissionAnalysisRequest(BaseModel):
    """Request for permission analysis."""
    output_directory: str


@router.post("/apk/decompile/permissions", response_model=PermissionAnalysisResponse)
async def analyze_permissions(
    request: PermissionAnalysisRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Analyze permissions from AndroidManifest.xml.
    
    Returns:
    - All requested permissions
    - Categorized by danger level (dangerous, normal, signature)
    - Categorized by type (location, camera, storage, etc.)
    - Dangerous permission combinations
    - Overall risk score and assessment
    
    Useful for:
    - Understanding what the app can access
    - Identifying privacy risks
    - Finding potential malware indicators
    - Security auditing
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.analyze_permissions(output_dir)
        
        if "error" in result:
            raise HTTPException(status_code=404, detail=result["error"])
        
        return PermissionAnalysisResponse(
            total_permissions=result["total_permissions"],
            permissions=[PermissionInfo(**p) for p in result["permissions"]],
            by_level={k: [PermissionInfo(**p) for p in v] for k, v in result["by_level"].items()},
            by_category={k: [PermissionInfo(**p) for p in v] for k, v in result["by_category"].items()},
            dangerous_combinations=[DangerousCombination(**c) for c in result["dangerous_combinations"]],
            risk_score=result["risk_score"],
            overall_risk=result["overall_risk"],
            summary=result["summary"],
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Permission analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Permission analysis failed: {str(e)}")


# ============================================================================
# Network Endpoint Extractor Endpoint
# ============================================================================

class NetworkEndpoint(BaseModel):
    """A single network endpoint."""
    value: str
    type: str
    category: str
    risk: str
    file: str
    line: int


class NetworkEndpointResponse(BaseModel):
    """Network endpoint extraction results."""
    total_endpoints: int
    endpoints: List[NetworkEndpoint]
    by_category: Dict[str, List[NetworkEndpoint]]
    by_risk: Dict[str, List[NetworkEndpoint]]
    unique_domains: List[str]
    domain_count: int
    summary: str
    error: Optional[str] = None


class NetworkEndpointRequest(BaseModel):
    """Request for network endpoint extraction."""
    output_directory: str


@router.post("/apk/decompile/network-endpoints", response_model=NetworkEndpointResponse)
async def extract_network_endpoints(
    request: NetworkEndpointRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Extract all network endpoints from decompiled sources.
    
    Scans for:
    - HTTP/HTTPS URLs
    - IP addresses (IPv4)
    - API endpoints and paths
    - WebSocket URLs
    - Cloud service URLs (Firebase, AWS, Azure, GCP)
    - Webhooks (Slack, Discord)
    - Payment APIs (Stripe, etc.)
    
    Returns:
    - All found endpoints with file locations
    - Categorized by type and risk level
    - List of unique domains
    - Risk assessment
    
    Useful for:
    - Finding API keys and secrets
    - Identifying C&C servers
    - Understanding app network behavior
    - Security auditing
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.extract_network_endpoints(output_dir)
        
        if "error" in result:
            raise HTTPException(status_code=404, detail=result["error"])
        
        return NetworkEndpointResponse(
            total_endpoints=result["total_endpoints"],
            endpoints=[NetworkEndpoint(**e) for e in result["endpoints"]],
            by_category={k: [NetworkEndpoint(**e) for e in v] for k, v in result["by_category"].items()},
            by_risk={k: [NetworkEndpoint(**e) for e in v] for k, v in result["by_risk"].items()},
            unique_domains=result["unique_domains"],
            domain_count=result["domain_count"],
            summary=result["summary"],
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Network endpoint extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"Network endpoint extraction failed: {str(e)}")


# ============================================================================
# Manifest Visualization Endpoints
# ============================================================================

class ManifestNodeResponse(BaseModel):
    """A node in the manifest visualization."""
    id: str
    name: str
    node_type: str
    label: str
    is_exported: bool = False
    is_main: bool = False
    is_dangerous: bool = False
    attributes: Dict[str, Any] = {}


class ManifestEdgeResponse(BaseModel):
    """An edge in the manifest visualization."""
    source: str
    target: str
    edge_type: str
    label: str = ""


class ManifestVisualizationResponse(BaseModel):
    """Complete manifest visualization response."""
    package_name: str
    app_name: Optional[str] = None
    version_name: Optional[str] = None
    nodes: List[ManifestNodeResponse]
    edges: List[ManifestEdgeResponse]
    component_counts: Dict[str, int]
    permission_summary: Dict[str, int]
    exported_count: int
    main_activity: Optional[str] = None
    deep_link_schemes: List[str] = []
    mermaid_diagram: str
    # AI-enhanced fields
    ai_analysis: Optional[str] = None
    component_purposes: Optional[Dict[str, str]] = None
    security_assessment: Optional[str] = None
    intent_filter_analysis: Optional[Dict[str, Any]] = None


@router.post("/apk/manifest-visualization", response_model=ManifestVisualizationResponse)
async def get_manifest_visualization(
    file: UploadFile = File(..., description="APK file to visualize"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate visualization data for an APK's AndroidManifest.
    
    Returns:
    - Graph nodes for all components and permissions
    - Graph edges showing relationships
    - Component counts by type
    - Mermaid diagram for rendering
    
    Use this data to render an interactive component graph.
    """
    filename = file.filename or "unknown.apk"
    suffix = Path(filename).suffix.lower()
    
    if suffix not in ALLOWED_APK_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {', '.join(ALLOWED_APK_EXTENSIONS)}"
        )
    
    tmp_dir = None
    try:
        # Save file to temp location with sanitized filename to prevent path traversal
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_manifest_"))
        safe_filename = sanitize_filename(filename)
        tmp_path = tmp_dir / safe_filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Generating manifest visualization: {filename}")
        
        # Generate visualization
        result = re_service.generate_manifest_visualization(tmp_path)
        
        return ManifestVisualizationResponse(
            package_name=result.package_name,
            app_name=result.app_name,
            version_name=result.version_name,
            nodes=[
                ManifestNodeResponse(
                    id=n.id,
                    name=n.name,
                    node_type=n.node_type,
                    label=n.label,
                    is_exported=n.is_exported,
                    is_main=n.is_main,
                    is_dangerous=n.is_dangerous,
                    attributes=n.attributes,
                )
                for n in result.nodes
            ],
            edges=[
                ManifestEdgeResponse(
                    source=e.source,
                    target=e.target,
                    edge_type=e.edge_type,
                    label=e.label,
                )
                for e in result.edges
            ],
            component_counts=result.component_counts,
            permission_summary=result.permission_summary,
            exported_count=result.exported_count,
            main_activity=result.main_activity,
            deep_link_schemes=result.deep_link_schemes,
            mermaid_diagram=result.mermaid_diagram,
            ai_analysis=getattr(result, 'ai_analysis', None),
            component_purposes=getattr(result, 'component_purposes', None),
            security_assessment=getattr(result, 'security_assessment', None),
            intent_filter_analysis=getattr(result, 'intent_filter_analysis', None),
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Manifest visualization failed: {e}")
        raise HTTPException(status_code=500, detail=f"Visualization failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


# ============================================================================
# Attack Surface Map Endpoints
# ============================================================================

class AttackVectorResponse(BaseModel):
    """An attack vector in the attack surface."""
    id: str
    name: str
    vector_type: str
    component: str
    severity: str
    description: str
    exploitation_steps: List[str]
    required_permissions: List[str] = []
    adb_command: Optional[str] = None
    intent_example: Optional[str] = None
    mitigation: Optional[str] = None


class DeepLinkResponse(BaseModel):
    """A deep link entry."""
    scheme: str
    host: str
    path: str
    full_url: str
    handling_activity: str
    parameters: List[str] = []
    is_verified: bool = False
    security_notes: List[str] = []


class ExposedDataPathResponse(BaseModel):
    """An exposed data path in content provider."""
    provider_name: str
    uri_pattern: str
    permissions_required: List[str]
    operations: List[str]
    is_exported: bool
    potential_data: str
    risk_level: str


class AttackSurfaceMapResponse(BaseModel):
    """Complete attack surface map response."""
    package_name: str
    total_attack_vectors: int
    attack_vectors: List[AttackVectorResponse]
    exposed_data_paths: List[ExposedDataPathResponse]
    deep_links: List[DeepLinkResponse]
    overall_exposure_score: int
    risk_level: str
    risk_breakdown: Dict[str, int]
    priority_targets: List[str]
    automated_tests: List[Dict[str, Any]]
    mermaid_attack_tree: str


@router.post("/apk/attack-surface", response_model=AttackSurfaceMapResponse)
async def get_attack_surface_map(
    file: UploadFile = File(..., description="APK file to analyze"),
    include_ai_analysis: bool = Query(False, description="Include AI analysis of decompiled source code (slower but more accurate)"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate a comprehensive attack surface map for an APK.
    
    Returns:
    - All attack vectors with exploitation steps
    - Deep links and their security implications
    - Exposed content provider paths
    - ADB commands for testing
    - Risk assessment and prioritization
    - Mermaid attack tree diagram
    
    Set include_ai_analysis=true to enable AI-powered analysis of decompiled source code
    for more accurate vulnerability detection (requires JADX decompilation).
    
    This provides a penetration tester's view of the app's attack surface.
    """
    filename = file.filename or "unknown.apk"
    suffix = Path(filename).suffix.lower()
    
    if suffix not in ALLOWED_APK_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {', '.join(ALLOWED_APK_EXTENSIONS)}"
        )
    
    tmp_dir = None
    try:
        # Save file to temp location with sanitized filename to prevent path traversal
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_attack_"))
        safe_filename = sanitize_filename(filename)
        tmp_path = tmp_dir / safe_filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Generating attack surface map: {filename}")
        
        # Generate basic attack surface map
        result = re_service.generate_attack_surface_map(tmp_path)
        
        # Enhance with AI analysis if requested
        if include_ai_analysis:
            logger.info("Running JADX decompilation for AI attack surface analysis...")
            try:
                jadx_result = re_service.decompile_apk_with_jadx(tmp_path)
                output_dir = Path(jadx_result.output_directory)
                
                logger.info("Enhancing attack surface with AI analysis of source code...")
                result = await re_service.enhance_attack_surface_with_ai(result, output_dir)
                
                # Generate AI-powered attack tree from source code
                logger.info("Generating AI-powered attack tree from source code analysis...")
                ai_attack_tree = await re_service.generate_ai_attack_tree_mermaid(result, output_dir)
                if ai_attack_tree:
                    result.mermaid_attack_tree = ai_attack_tree
                
                # Clean up JADX output
                shutil.rmtree(output_dir, ignore_errors=True)
            except Exception as e:
                logger.warning(f"AI attack surface enhancement failed, using basic analysis: {e}")
        
        return AttackSurfaceMapResponse(
            package_name=result.package_name,
            total_attack_vectors=result.total_attack_vectors,
            attack_vectors=[
                AttackVectorResponse(
                    id=v.id,
                    name=v.name,
                    vector_type=v.vector_type,
                    component=v.component,
                    severity=v.severity,
                    description=v.description,
                    exploitation_steps=v.exploitation_steps,
                    required_permissions=v.required_permissions,
                    adb_command=v.adb_command,
                    intent_example=v.intent_example,
                    mitigation=v.mitigation,
                )
                for v in result.attack_vectors
            ],
            exposed_data_paths=[
                ExposedDataPathResponse(
                    provider_name=p.provider_name,
                    uri_pattern=p.uri_pattern,
                    permissions_required=p.permissions_required,
                    operations=p.operations,
                    is_exported=p.is_exported,
                    potential_data=p.potential_data,
                    risk_level=p.risk_level,
                )
                for p in result.exposed_data_paths
            ],
            deep_links=[
                DeepLinkResponse(
                    scheme=d.scheme,
                    host=d.host,
                    path=d.path,
                    full_url=d.full_url,
                    handling_activity=d.handling_activity,
                    parameters=d.parameters,
                    is_verified=d.is_verified,
                    security_notes=d.security_notes,
                )
                for d in result.deep_links
            ],
            overall_exposure_score=result.overall_exposure_score,
            risk_level=result.risk_level,
            risk_breakdown=result.risk_breakdown,
            priority_targets=result.priority_targets,
            automated_tests=result.automated_tests,
            mermaid_attack_tree=result.mermaid_attack_tree,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Attack surface mapping failed: {e}")
        raise HTTPException(status_code=500, detail=f"Attack surface mapping failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


class AIAttackSurfaceRequest(BaseModel):
    """Request to enhance attack surface with AI analysis."""
    session_id: str = Field(..., description="JADX decompilation session ID")


@router.post("/apk/decompile/attack-surface-ai", response_model=AttackSurfaceMapResponse)
async def get_ai_attack_surface_map(
    request: AIAttackSurfaceRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate AI-enhanced attack surface map using an existing JADX session.
    
    This endpoint analyzes decompiled source code to find real vulnerabilities
    in exported components, providing:
    - Code-level vulnerability detection
    - Specific exploitation steps based on actual code
    - Attack chain analysis
    - AI-powered severity assessment
    
    Prerequisites: APK must be decompiled first using /apk/decompile endpoint.
    """
    session_id = request.session_id
    
    if session_id not in _jadx_cache:
        raise HTTPException(
            status_code=404, 
            detail="Decompilation session not found. Please decompile the APK first."
        )
    
    output_dir = _jadx_cache[session_id]
    
    if not output_dir.exists():
        raise HTTPException(status_code=404, detail="Decompilation output no longer exists.")
    
    try:
        # Find the APK path - look for it in the parent of the output dir
        apk_path = None
        for ext in ['.apk', '.APK']:
            for apk in output_dir.parent.glob(f"*{ext}"):
                apk_path = apk
                break
            if apk_path:
                break
        
        if not apk_path:
            # Create minimal attack surface from JADX info
            jadx_summary = re_service.get_jadx_result_summary(output_dir)
            
            # Build a basic attack surface from the JADX classes
            attack_vectors = []
            for cls in jadx_summary.get('classes', []):
                if cls.get('is_activity'):
                    attack_vectors.append(re_service.AttackVector(
                        id=f"activity_{cls.get('class_name', 'unknown').split('.')[-1]}",
                        name=f"Activity: {cls.get('class_name', 'unknown').split('.')[-1]}",
                        vector_type="exported_activity",
                        component=cls.get('class_name', 'unknown'),
                        severity="medium",
                        description=f"Activity found in decompiled code",
                        exploitation_steps=["Analyze source code for vulnerabilities"],
                        required_permissions=[],
                        adb_command="",
                        intent_example="",
                        mitigation=""
                    ))
            
            basic_surface = re_service.AttackSurfaceMap(
                package_name=jadx_summary.get('package_name', 'unknown'),
                total_attack_vectors=len(attack_vectors),
                attack_vectors=attack_vectors,
                exposed_data_paths=[],
                deep_links=[],
                ipc_endpoints=[],
                overall_exposure_score=50,
                risk_level="medium",
                risk_breakdown={'critical': 0, 'high': 0, 'medium': len(attack_vectors), 'low': 0},
                priority_targets=[],
                automated_tests=[],
                mermaid_attack_tree="flowchart TD\n  A[Analysis from decompiled code]"
            )
            result = await re_service.enhance_attack_surface_with_ai(basic_surface, output_dir)
            
            # Generate AI-powered attack tree from source code
            ai_attack_tree = await re_service.generate_ai_attack_tree_mermaid(result, output_dir)
            if ai_attack_tree:
                result.mermaid_attack_tree = ai_attack_tree
        else:
            # Generate full attack surface from APK
            basic_surface = re_service.generate_attack_surface_map(apk_path)
            result = await re_service.enhance_attack_surface_with_ai(basic_surface, output_dir)
            
            # Generate AI-powered attack tree from source code
            ai_attack_tree = await re_service.generate_ai_attack_tree_mermaid(result, output_dir)
            if ai_attack_tree:
                result.mermaid_attack_tree = ai_attack_tree
        
        return AttackSurfaceMapResponse(
            package_name=result.package_name,
            total_attack_vectors=result.total_attack_vectors,
            attack_vectors=[
                AttackVectorResponse(
                    id=v.id,
                    name=v.name,
                    vector_type=v.vector_type,
                    component=v.component,
                    severity=v.severity,
                    description=v.description,
                    exploitation_steps=v.exploitation_steps,
                    required_permissions=v.required_permissions,
                    adb_command=v.adb_command,
                    intent_example=v.intent_example,
                    mitigation=v.mitigation,
                )
                for v in result.attack_vectors
            ],
            exposed_data_paths=[
                ExposedDataPathResponse(
                    provider_name=p.provider_name,
                    uri_pattern=p.uri_pattern,
                    permissions_required=p.permissions_required,
                    operations=p.operations,
                    is_exported=p.is_exported,
                    potential_data=p.potential_data,
                    risk_level=p.risk_level,
                )
                for p in result.exposed_data_paths
            ],
            deep_links=[
                DeepLinkResponse(
                    scheme=d.scheme,
                    host=d.host,
                    path=d.path,
                    full_url=d.full_url,
                    handling_activity=d.handling_activity,
                    parameters=d.parameters,
                    is_verified=d.is_verified,
                    security_notes=d.security_notes,
                )
                for d in result.deep_links
            ],
            overall_exposure_score=result.overall_exposure_score,
            risk_level=result.risk_level,
            risk_breakdown=result.risk_breakdown,
            priority_targets=result.priority_targets,
            automated_tests=result.automated_tests,
            mermaid_attack_tree=result.mermaid_attack_tree,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI attack surface analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI attack surface analysis failed: {str(e)}")


# ============================================================================
# Obfuscation Analysis Response Models
# ============================================================================

class ObfuscationIndicatorResponse(BaseModel):
    """Response model for obfuscation indicator."""
    indicator_type: str
    confidence: str
    description: str
    evidence: List[str]
    location: Optional[str] = None
    deobfuscation_hint: Optional[str] = None


class StringEncryptionPatternResponse(BaseModel):
    """Response model for string encryption pattern."""
    pattern_name: str
    class_name: str
    method_name: str
    encrypted_strings_count: int
    decryption_method_signature: Optional[str] = None
    sample_encrypted_values: List[str] = []
    suggested_frida_hook: Optional[str] = None


class ClassNamingAnalysisResponse(BaseModel):
    """Response model for class naming analysis."""
    total_classes: int
    single_letter_classes: int
    short_name_classes: int
    meaningful_name_classes: int
    obfuscation_ratio: float
    sample_obfuscated_names: List[str]
    sample_original_names: List[str]


class ControlFlowObfuscationResponse(BaseModel):
    """Response model for control flow obfuscation."""
    pattern_type: str
    affected_methods: int
    sample_classes: List[str]
    complexity_score: float


class NativeProtectionResponse(BaseModel):
    """Response model for native protection analysis."""
    has_native_libs: bool
    native_lib_names: List[str]
    protection_indicators: List[str]
    jni_functions: List[str]


class ObfuscationAnalysisResponse(BaseModel):
    """Response model for complete obfuscation analysis."""
    package_name: str
    overall_obfuscation_level: str
    obfuscation_score: int
    detected_tools: List[str]
    
    indicators: List[ObfuscationIndicatorResponse]
    class_naming: ClassNamingAnalysisResponse
    string_encryption: List[StringEncryptionPatternResponse]
    control_flow: List[ControlFlowObfuscationResponse]
    native_protection: NativeProtectionResponse
    
    deobfuscation_strategies: List[str]
    recommended_tools: List[str]
    frida_hooks: List[str]
    
    analysis_time: float
    warnings: List[str]
    
    # AI-enhanced fields
    ai_analysis_summary: Optional[str] = None
    reverse_engineering_difficulty: Optional[str] = None
    ai_recommended_approach: Optional[str] = None


@router.post("/apk/obfuscation-analysis", response_model=ObfuscationAnalysisResponse)
async def analyze_apk_obfuscation(
    file: UploadFile = File(..., description="APK file to analyze for obfuscation"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Analyze an APK for obfuscation techniques.
    
    Detects:
    - ProGuard/R8 obfuscation patterns
    - DexGuard commercial protection
    - String encryption methods
    - Control flow obfuscation
    - Native library protection
    - Reflection-based API hiding
    
    Returns analysis with:
    - Detected obfuscation tools
    - Obfuscation score (0-100)
    - Deobfuscation strategies
    - Auto-generated Frida hooks
    """
    tmp_dir = None
    
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        filename = file.filename.lower()
        if not filename.endswith(('.apk', '.aab')):
            raise HTTPException(status_code=400, detail="Only APK/AAB files are supported")
        
        # Save file temporarily with sanitized filename to prevent path traversal
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_obfusc_"))
        safe_filename = sanitize_filename(filename)
        tmp_path = tmp_dir / safe_filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Analyzing obfuscation for: {filename}")
        
        # Perform obfuscation analysis
        result = re_service.analyze_apk_obfuscation(tmp_path)
        
        return ObfuscationAnalysisResponse(
            package_name=result.package_name,
            overall_obfuscation_level=result.overall_obfuscation_level,
            obfuscation_score=result.obfuscation_score,
            detected_tools=result.detected_tools,
            indicators=[
                ObfuscationIndicatorResponse(
                    indicator_type=i.indicator_type,
                    confidence=i.confidence,
                    description=i.description,
                    evidence=i.evidence,
                    location=i.location,
                    deobfuscation_hint=i.deobfuscation_hint,
                )
                for i in result.indicators
            ],
            class_naming=ClassNamingAnalysisResponse(
                total_classes=result.class_naming.total_classes,
                single_letter_classes=result.class_naming.single_letter_classes,
                short_name_classes=result.class_naming.short_name_classes,
                meaningful_name_classes=result.class_naming.meaningful_name_classes,
                obfuscation_ratio=result.class_naming.obfuscation_ratio,
                sample_obfuscated_names=result.class_naming.sample_obfuscated_names,
                sample_original_names=result.class_naming.sample_original_names,
            ),
            string_encryption=[
                StringEncryptionPatternResponse(
                    pattern_name=s.pattern_name,
                    class_name=s.class_name,
                    method_name=s.method_name,
                    encrypted_strings_count=s.encrypted_strings_count,
                    decryption_method_signature=s.decryption_method_signature,
                    sample_encrypted_values=s.sample_encrypted_values,
                    suggested_frida_hook=s.suggested_frida_hook,
                )
                for s in result.string_encryption
            ],
            control_flow=[
                ControlFlowObfuscationResponse(
                    pattern_type=c.pattern_type,
                    affected_methods=c.affected_methods,
                    sample_classes=c.sample_classes,
                    complexity_score=c.complexity_score,
                )
                for c in result.control_flow
            ],
            native_protection=NativeProtectionResponse(
                has_native_libs=result.native_protection.has_native_libs,
                native_lib_names=result.native_protection.native_lib_names,
                protection_indicators=result.native_protection.protection_indicators,
                jni_functions=result.native_protection.jni_functions,
            ),
            deobfuscation_strategies=result.deobfuscation_strategies,
            recommended_tools=result.recommended_tools,
            frida_hooks=result.frida_hooks,
            analysis_time=result.analysis_time,
            warnings=result.warnings,
            ai_analysis_summary=getattr(result, 'ai_analysis_summary', None),
            reverse_engineering_difficulty=getattr(result, 'reverse_engineering_difficulty', None),
            ai_recommended_approach=getattr(result, 'ai_recommended_approach', None),
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Obfuscation analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Obfuscation analysis failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


@router.post("/apk/obfuscation-analysis/ai-enhanced", response_model=ObfuscationAnalysisResponse)
async def analyze_apk_obfuscation_ai_enhanced(
    file: UploadFile = File(..., description="APK file to analyze for obfuscation"),
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-ENHANCED obfuscation analysis with deep code pattern recognition.
    
    This endpoint combines fast static analysis with AI-powered insights:
    - Identifies specific obfuscation tools (ProGuard, DexGuard, Allatori, etc.)
    - Analyzes code samples to detect obfuscation patterns
    - Provides tailored deobfuscation strategies
    - Generates custom Frida hooks for specific patterns
    - Assesses reverse engineering difficulty
    
    Returns enhanced analysis with AI insights and recommendations.
    """
    tmp_dir = None
    
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        filename = file.filename.lower()
        if not filename.endswith(('.apk', '.aab')):
            raise HTTPException(status_code=400, detail="Only APK/AAB files are supported")
        
        # Save file temporarily with sanitized filename to prevent path traversal
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_obfusc_ai_"))
        safe_filename = sanitize_filename(filename)
        tmp_path = tmp_dir / safe_filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"AI-enhanced obfuscation analysis for: {filename}")
        
        # Run JADX decompilation to get source code for AI analysis
        jadx_output_dir = None
        try:
            jadx_result = re_service.decompile_apk_with_jadx(tmp_path)
            jadx_output_dir = Path(jadx_result.output_directory)
            logger.info(f"JADX decompilation complete for obfuscation AI analysis")
        except Exception as e:
            logger.warning(f"JADX decompilation failed, continuing without source code: {e}")
        
        # Perform AI-enhanced obfuscation analysis with source code context
        result = await re_service.analyze_apk_obfuscation_ai_enhanced(tmp_path, jadx_output_dir)
        
        # Clean up JADX output
        if jadx_output_dir and jadx_output_dir.exists():
            shutil.rmtree(jadx_output_dir, ignore_errors=True)
        
        return ObfuscationAnalysisResponse(
            package_name=result.package_name,
            overall_obfuscation_level=result.overall_obfuscation_level,
            obfuscation_score=result.obfuscation_score,
            detected_tools=result.detected_tools,
            indicators=[
                ObfuscationIndicatorResponse(
                    indicator_type=i.indicator_type,
                    confidence=i.confidence,
                    description=i.description,
                    evidence=i.evidence,
                    location=getattr(i, 'location', None),
                    deobfuscation_hint=getattr(i, 'deobfuscation_hint', None),
                )
                for i in result.indicators
            ],
            class_naming=ClassNamingAnalysisResponse(
                total_classes=result.class_naming.total_classes,
                single_letter_classes=result.class_naming.single_letter_classes,
                short_name_classes=result.class_naming.short_name_classes,
                meaningful_name_classes=result.class_naming.meaningful_name_classes,
                obfuscation_ratio=result.class_naming.obfuscation_ratio,
                sample_obfuscated_names=result.class_naming.sample_obfuscated_names,
                sample_original_names=result.class_naming.sample_original_names,
            ),
            string_encryption=[
                StringEncryptionPatternResponse(
                    pattern_name=getattr(s, 'pattern_name', s.pattern_type),
                    class_name=getattr(s, 'class_name', ''),
                    method_name=getattr(s, 'method_name', ''),
                    encrypted_strings_count=getattr(s, 'encrypted_strings_count', s.occurrences),
                    decryption_method_signature=getattr(s, 'decryption_method_signature', ''),
                    sample_encrypted_values=getattr(s, 'sample_encrypted_values', s.sample_encrypted),
                    suggested_frida_hook=getattr(s, 'suggested_frida_hook', s.decryption_hint),
                )
                for s in result.string_encryption
            ],
            control_flow=[
                ControlFlowObfuscationResponse(
                    pattern_type=c.pattern_type,
                    affected_methods=c.affected_methods,
                    sample_classes=c.sample_classes,
                    complexity_score=c.complexity_score,
                )
                for c in result.control_flow
            ],
            native_protection=NativeProtectionResponse(
                has_native_libs=result.native_protection.has_native_libs,
                native_lib_names=result.native_protection.native_lib_names,
                protection_indicators=result.native_protection.protection_indicators,
                jni_functions=getattr(result.native_protection, 'jni_functions', []),
            ),
            deobfuscation_strategies=result.deobfuscation_strategies,
            recommended_tools=result.recommended_tools,
            frida_hooks=result.frida_hooks,
            analysis_time=result.analysis_time,
            warnings=result.warnings,
            ai_analysis_summary=getattr(result, 'ai_analysis_summary', None),
            reverse_engineering_difficulty=getattr(result, 'reverse_engineering_difficulty', None),
            ai_recommended_approach=getattr(result, 'ai_recommended_approach', None),
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI-enhanced obfuscation analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI-enhanced obfuscation analysis failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


@router.post("/apk/manifest-visualization/ai-enhanced", response_model=ManifestVisualizationResponse)
async def get_manifest_visualization_ai_enhanced(
    file: UploadFile = File(..., description="APK file to visualize"),
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-ENHANCED manifest visualization with deep component analysis.
    
    This endpoint provides:
    - Standard manifest visualization (graph, mermaid diagram)
    - AI analysis of component purposes based on names and code
    - Security assessment of exported components
    - Intent filter analysis for security implications
    - Component relationship mapping
    
    Returns visualization data enriched with AI insights.
    """
    filename = file.filename or "unknown.apk"
    suffix = Path(filename).suffix.lower()
    
    if suffix not in ALLOWED_APK_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {', '.join(ALLOWED_APK_EXTENSIONS)}"
        )
    
    tmp_dir = None
    try:
        # Save file to temp location with sanitized filename to prevent path traversal
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_manifest_ai_"))
        safe_filename = sanitize_filename(filename)
        tmp_path = tmp_dir / safe_filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Generating AI-enhanced manifest visualization: {filename}")
        
        # Run JADX decompilation to get source code for AI analysis
        jadx_output_dir = None
        try:
            jadx_result = re_service.decompile_apk_with_jadx(tmp_path)
            jadx_output_dir = Path(jadx_result.output_directory)
            logger.info(f"JADX decompilation complete for manifest AI analysis")
        except Exception as e:
            logger.warning(f"JADX decompilation failed, continuing without source code: {e}")
        
        # Generate AI-enhanced visualization with source code context
        result = await re_service.generate_ai_enhanced_manifest_visualization(tmp_path, jadx_output_dir)
        
        # Clean up JADX output
        if jadx_output_dir and jadx_output_dir.exists():
            shutil.rmtree(jadx_output_dir, ignore_errors=True)
        
        return ManifestVisualizationResponse(
            package_name=result.package_name,
            app_name=result.app_name,
            version_name=result.version_name,
            nodes=[
                ManifestNodeResponse(
                    id=n.id,
                    name=n.name,
                    node_type=n.node_type,
                    label=n.label,
                    is_exported=n.is_exported,
                    is_main=n.is_main,
                    is_dangerous=n.is_dangerous,
                    attributes=n.attributes,
                )
                for n in result.nodes
            ],
            edges=[
                ManifestEdgeResponse(
                    source=e.source,
                    target=e.target,
                    edge_type=e.edge_type,
                    label=e.label,
                )
                for e in result.edges
            ],
            component_counts=result.component_counts,
            permission_summary=result.permission_summary,
            exported_count=result.exported_count,
            main_activity=result.main_activity,
            deep_link_schemes=result.deep_link_schemes,
            mermaid_diagram=result.mermaid_diagram,
            ai_analysis=result.ai_analysis,
            component_purposes=result.component_purposes,
            security_assessment=result.security_assessment,
            intent_filter_analysis=result.intent_filter_analysis,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"AI-enhanced manifest visualization failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI-enhanced visualization failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


# ============================================================================
# Binary Entropy Analysis
# ============================================================================

class EntropyDataPointResponse(BaseModel):
    """Response model for entropy data point."""
    offset: int
    entropy: float
    size: int


class EntropyRegionResponse(BaseModel):
    """Response model for entropy region."""
    start_offset: int
    end_offset: int
    avg_entropy: float
    max_entropy: float
    min_entropy: float
    classification: str
    section_name: Optional[str] = None
    description: str = ""


class EntropyAnalysisResponse(BaseModel):
    """Response model for entropy analysis."""
    filename: str
    file_size: int
    overall_entropy: float
    entropy_data: List[EntropyDataPointResponse]
    regions: List[EntropyRegionResponse]
    is_likely_packed: bool
    packing_confidence: float
    detected_packers: List[str]
    section_entropy: List[Dict[str, Any]]
    analysis_notes: List[str]
    window_size: int
    step_size: int


@router.post("/binary/entropy", response_model=EntropyAnalysisResponse)
async def analyze_binary_entropy(
    file: UploadFile = File(..., description="Binary file to analyze"),
    window_size: int = Query(256, ge=64, le=4096, description="Entropy calculation window size"),
    step_size: int = Query(128, ge=32, le=2048, description="Step size between measurements"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Analyze entropy distribution across a binary file.
    
    Entropy analysis helps identify:
    - Packed or compressed code sections
    - Encrypted regions
    - Normal code vs data sections
    - Potential malware indicators
    
    Returns entropy data points for visualization and region classification.
    """
    tmp_dir = None
    
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Save file temporarily
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_entropy_"))
        tmp_path = tmp_dir / file.filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Analyzing entropy for: {file.filename}")
        
        # Perform entropy analysis
        result = re_service.analyze_binary_entropy(tmp_path, window_size, step_size)
        
        return EntropyAnalysisResponse(
            filename=result.filename,
            file_size=result.file_size,
            overall_entropy=result.overall_entropy,
            entropy_data=[
                EntropyDataPointResponse(
                    offset=p.offset,
                    entropy=p.entropy,
                    size=p.size
                )
                for p in result.entropy_data
            ],
            regions=[
                EntropyRegionResponse(
                    start_offset=r.start_offset,
                    end_offset=r.end_offset,
                    avg_entropy=r.avg_entropy,
                    max_entropy=r.max_entropy,
                    min_entropy=r.min_entropy,
                    classification=r.classification,
                    section_name=r.section_name,
                    description=r.description
                )
                for r in result.regions
            ],
            is_likely_packed=result.is_likely_packed,
            packing_confidence=result.packing_confidence,
            detected_packers=result.detected_packers,
            section_entropy=result.section_entropy,
            analysis_notes=result.analysis_notes,
            window_size=result.window_size,
            step_size=result.step_size
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Entropy analysis failed: {e}")
        raise HTTPException(status_code=500, detail=f"Entropy analysis failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


# ============================================================================
# APK Report Export Endpoints
# ============================================================================

@router.post("/apk/export")
async def export_apk_report(
    file: UploadFile = File(...),
    format: str = Query(..., description="Export format: markdown, pdf, docx"),
    report_type: str = Query("both", description="Report type: functionality, security, both"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Analyze an APK and export the report to Markdown, PDF, or Word format.
    
    - **format**: Export format (markdown, pdf, docx)
    - **report_type**: Which report to generate (functionality, security, both)
    """
    from fastapi.responses import Response
    
    if format not in ["markdown", "pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Invalid format. Use: markdown, pdf, docx")
    
    if report_type not in ["functionality", "security", "both"]:
        raise HTTPException(status_code=400, detail="Invalid report_type. Use: functionality, security, both")
    
    # Validate file
    filename = file.filename or "unknown.apk"
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_APK_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: {', '.join(ALLOWED_APK_EXTENSIONS)}"
        )
    
    tmp_dir = None
    try:
        # Save uploaded file
        tmp_dir = tempfile.mkdtemp()
        tmp_path = Path(tmp_dir) / filename
        
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail=f"File too large. Max: {MAX_FILE_SIZE // (1024*1024)}MB")
        
        with open(tmp_path, "wb") as f:
            f.write(content)
        
        # Analyze APK
        result = await re_service.analyze_apk(tmp_path)
        
        # Generate AI reports if Gemini is available
        await re_service.analyze_apk_with_ai(result)
        
        # Generate export based on format
        base_filename = Path(filename).stem
        
        if format == "markdown":
            markdown_content = re_service.generate_apk_markdown_report(result, report_type)
            return Response(
                content=markdown_content.encode('utf-8'),
                media_type="text/markdown",
                headers={"Content-Disposition": f'attachment; filename="{base_filename}_report.md"'}
            )
        
        elif format == "pdf":
            pdf_bytes = re_service.generate_apk_pdf_report(result, report_type)
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{base_filename}_report.pdf"'}
            )
        
        elif format == "docx":
            docx_bytes = re_service.generate_apk_docx_report(result, report_type)
            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{base_filename}_report.docx"'}
            )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"APK export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)


@router.post("/apk/export-from-result")
async def export_apk_report_from_result(
    result_data: Dict[str, Any],
    format: str = Query(..., description="Export format: markdown, pdf, docx"),
    report_type: str = Query("both", description="Report type: functionality, security, both"),
    report_title: Optional[str] = Query(None, description="Optional custom report title"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Export an existing APK analysis result to Markdown, PDF, or Word format.
    
    Use this when you already have analysis results and don't want to re-analyze.
    """
    from fastapi.responses import Response
    
    if format not in ["markdown", "pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Invalid format. Use: markdown, pdf, docx")
    
    if report_type not in ["functionality", "security", "both"]:
        raise HTTPException(status_code=400, detail="Invalid report_type. Use: functionality, security, both")
    
    try:
        def _normalize_dict_list(value: Any, nested_keys: Optional[List[str]] = None) -> List[Dict[str, Any]]:
            if isinstance(value, list):
                return [item for item in value if isinstance(item, dict)]
            if isinstance(value, dict):
                for key in nested_keys or []:
                    nested = value.get(key)
                    if isinstance(nested, list):
                        return [item for item in nested if isinstance(item, dict)]
            return []

        resolved_title = (report_title or result_data.get("title") or "").strip()
        if resolved_title:
            resolved_title = resolved_title[:180]

        # Create permission objects
        permissions = []
        for p in result_data.get('permissions', []):
            permissions.append(re_service.ApkPermission(
                name=p.get('name', ''),
                description=p.get('description'),
                is_dangerous=p.get('is_dangerous', False)
            ))

        components = []
        activities: List[str] = []
        services: List[str] = []
        receivers: List[str] = []
        providers: List[str] = []
        for c in result_data.get('components', []):
            component = re_service.ApkComponent(
                name=c.get('name', ''),
                component_type=c.get('component_type', ''),
                is_exported=c.get('is_exported', False),
                intent_filters=c.get('intent_filters') or [],
            )
            components.append(component)
            component_type = component.component_type.lower()
            if component_type == 'activity':
                activities.append(component.name)
            elif component_type == 'service':
                services.append(component.name)
            elif component_type == 'receiver':
                receivers.append(component.name)
            elif component_type == 'provider':
                providers.append(component.name)
        
        # Create certificate object if present
        certificate = None
        if result_data.get('certificate'):
            cert_data = result_data['certificate']
            certificate = re_service.ApkCertificate(
                subject=cert_data.get('subject', ''),
                issuer=cert_data.get('issuer', ''),
                fingerprint_sha256=cert_data.get('fingerprint_sha256', ''),
                fingerprint_sha1=cert_data.get('fingerprint_sha1'),
                fingerprint_md5=cert_data.get('fingerprint_md5'),
                serial_number=cert_data.get('serial_number'),
                valid_from=cert_data.get('valid_from'),
                valid_until=cert_data.get('valid_until'),
                is_debug_cert=cert_data.get('is_debug_cert', False),
                is_expired=cert_data.get('is_expired', False),
                is_self_signed=cert_data.get('is_self_signed', False),
                signature_version=cert_data.get('signature_version', 'v1'),
                public_key_algorithm=cert_data.get('public_key_algorithm'),
                public_key_bits=cert_data.get('public_key_bits')
            )
        
        # Create result object
        result = re_service.ApkAnalysisResult(
            filename=result_data.get('filename', 'unknown.apk'),
            package_name=result_data.get('package_name', ''),
            version_name=result_data.get('version_name'),
            version_code=result_data.get('version_code'),
            min_sdk=result_data.get('min_sdk'),
            target_sdk=result_data.get('target_sdk'),
            permissions=permissions,
            components=components,
            strings=[],  # Simplified for export
            secrets=result_data.get('secrets', []),
            urls=result_data.get('urls', []),
            native_libraries=result_data.get('native_libraries', []),
            certificate=certificate,
            activities=result_data.get('activities', []) or activities,
            services=result_data.get('services', []) or services,
            receivers=result_data.get('receivers', []) or receivers,
            providers=result_data.get('providers', []) or providers,
            uses_features=result_data.get('uses_features', []),
            app_name=result_data.get('app_name'),
            debuggable=result_data.get('debuggable', False),
            allow_backup=result_data.get('allow_backup', True),
            security_issues=result_data.get('security_issues', []),
            ai_analysis=result_data.get('ai_analysis'),
            ai_report_functionality=result_data.get('ai_report_functionality'),
            ai_report_security=result_data.get('ai_report_security'),
            hardening_score=result_data.get('hardening_score'),
        )
        
        # Get decompiled code findings if present
        decompiled_findings = _normalize_dict_list(
            result_data.get('decompiled_code_findings', []),
            nested_keys=['findings', 'issues', 'results'],
        )
        
        # Get CVE scan results if present
        cve_scan_results = _normalize_dict_list(
            result_data.get('cve_scan_results', []),
            nested_keys=['cves', 'findings', 'results', 'items'],
        )
        
        # Get AI diagrams if present
        ai_architecture_diagram = result_data.get('ai_architecture_diagram')
        ai_attack_surface_map = result_data.get('ai_attack_surface_map')
        
        # Get dynamic analysis (Frida scripts) if present
        dynamic_analysis = result_data.get('dynamic_analysis')
        
        # Generate export based on format
        if not resolved_title:
            default_name = result.package_name or result.filename or "APK"
            resolved_title = f"APK Analysis: {default_name}"
        base_name = "".join(ch if (ch.isalnum() or ch in {"-", "_"}) else "_" for ch in resolved_title)
        base_name = base_name.strip("_")[:120] or "apk_analysis_report"
        
        if format == "markdown":
            markdown_content = re_service.generate_apk_markdown_report(
                result, report_type, decompiled_findings,
                cve_scan_results=cve_scan_results,
                ai_architecture_diagram=ai_architecture_diagram,
                ai_attack_surface_map=ai_attack_surface_map,
                dynamic_analysis=dynamic_analysis,
                report_title=resolved_title,
            )
            return Response(
                content=markdown_content.encode('utf-8'),
                media_type="text/markdown",
                headers={"Content-Disposition": f'attachment; filename="{base_name}.md"'}
            )
        
        elif format == "pdf":
            pdf_bytes = re_service.generate_apk_pdf_report(
                result, report_type, decompiled_findings,
                cve_scan_results=cve_scan_results,
                ai_architecture_diagram=ai_architecture_diagram,
                ai_attack_surface_map=ai_attack_surface_map,
                dynamic_analysis=dynamic_analysis,
                report_title=resolved_title,
            )
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{base_name}.pdf"'}
            )
        
        elif format == "docx":
            docx_bytes = re_service.generate_apk_docx_report(
                result, report_type, decompiled_findings,
                cve_scan_results=cve_scan_results,
                ai_architecture_diagram=ai_architecture_diagram,
                ai_attack_surface_map=ai_attack_surface_map,
                dynamic_analysis=dynamic_analysis,
                report_title=resolved_title,
            )
            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{base_name}.docx"'}
            )
    
    except Exception as e:
        logger.exception("APK export from result failed")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


# ============================================================================
# Frida Script Export Endpoints
# ============================================================================

class FridaScriptExportRequest(BaseModel):
    """Request for exporting Frida scripts."""
    package_name: str
    dynamic_analysis: Optional[Dict[str, Any]] = None
    vulnerability_scripts: Optional[List[Dict[str, Any]]] = None
    selected_categories: Optional[List[str]] = None
    include_combined: bool = True
    export_format: str = "zip"  # "zip", "individual", "combined_only"


@router.post("/apk/frida-scripts/export")
async def export_frida_scripts(
    request: FridaScriptExportRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Export generated Frida scripts as downloadable files.
    
    Export formats:
    - zip: All scripts bundled in a ZIP file with README
    - individual: Returns JSON with individual script contents
    - combined_only: Returns just the combined all-in-one script
    """
    from fastapi.responses import Response
    import zipfile
    import io
    
    if not request.dynamic_analysis and not request.vulnerability_scripts:
        raise HTTPException(
            status_code=400, 
            detail="No Frida scripts to export. Run APK analysis first."
        )
    
    try:
        # Get standard scripts from dynamic_analysis
        scripts = []
        if request.dynamic_analysis:
            scripts = request.dynamic_analysis.get('frida_scripts', [])
        
        # Get vulnerability-specific scripts
        vuln_scripts = request.vulnerability_scripts or []
        
        # Export scripts
        export_result = re_service.export_frida_scripts_to_files(
            package_name=request.package_name,
            scripts=scripts,
            include_combined=request.include_combined,
            vuln_scripts=vuln_scripts
        )
        
        if request.export_format == "individual":
            # Return JSON with all script contents
            return {
                "success": True,
                "package_name": request.package_name,
                "total_files": export_result["total_files"],
                "files": [
                    {
                        "filename": f["filename"],
                        "category": f["category"],
                        "script_name": f["script_name"],
                        "content": f["content"],
                        "usage": f["usage"]
                    }
                    for f in export_result["files"]
                ]
            }
        
        elif request.export_format == "combined_only":
            # Return just the combined script
            combined_result = re_service.combine_frida_scripts(
                package_name=request.package_name,
                scripts=scripts,
                selected_categories=request.selected_categories,
                include_vulnerability_scripts=bool(vuln_scripts),
                vuln_scripts=vuln_scripts
            )
            
            if not combined_result.get("success"):
                raise HTTPException(status_code=400, detail=combined_result.get("error", "Failed to combine scripts"))
            
            return Response(
                content=combined_result["combined_script"].encode('utf-8'),
                media_type="application/javascript",
                headers={
                    "Content-Disposition": f'attachment; filename="frida_combined_{request.package_name.replace(".", "_")}.js"'
                }
            )
        
        else:  # zip format (default)
            # Create ZIP file in memory
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                # Create folder structure
                folder_name = f"frida_scripts_{request.package_name.replace('.', '_')}"
                
                for file_data in export_result["files"]:
                    # Put scripts in category subfolders
                    category = file_data["category"]
                    if category == "documentation":
                        file_path = f"{folder_name}/{file_data['filename']}"
                    elif category == "combined":
                        file_path = f"{folder_name}/{file_data['filename']}"
                    else:
                        file_path = f"{folder_name}/{category}/{file_data['filename']}"
                    
                    zf.writestr(file_path, file_data["content"])
            
            zip_buffer.seek(0)
            
            return Response(
                content=zip_buffer.getvalue(),
                media_type="application/zip",
                headers={
                    "Content-Disposition": f'attachment; filename="frida_scripts_{request.package_name.replace(".", "_")}.zip"'
                }
            )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Frida script export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


@router.post("/apk/frida-scripts/combine")
async def combine_frida_scripts_endpoint(
    package_name: str,
    scripts: List[Dict[str, Any]],
    selected_categories: Optional[List[str]] = None,
    vuln_scripts: Optional[List[Dict[str, Any]]] = None,
    current_user: User = Depends(get_current_active_user),
):
    """
    Combine multiple Frida scripts into a single all-in-one script.
    
    Useful for running all hooks at once without loading multiple files.
    """
    try:
        result = re_service.combine_frida_scripts(
            package_name=package_name,
            scripts=scripts,
            selected_categories=selected_categories,
            include_vulnerability_scripts=bool(vuln_scripts),
            vuln_scripts=vuln_scripts
        )
        
        if not result.get("success"):
            raise HTTPException(status_code=400, detail=result.get("error", "Failed to combine scripts"))
        
        return {
            "success": True,
            "package_name": package_name,
            "combined_script": result["combined_script"],
            "scripts_included": result["scripts_included"],
            "categories_included": result["categories_included"],
            "total_scripts": result["total_scripts"],
            "usage_command": result["usage_command"]
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Frida script combination failed: {e}")
        raise HTTPException(status_code=500, detail=f"Combination failed: {str(e)}")


# ============================================================================
# AI Chat Export Endpoints
# ============================================================================

class ChatExportRequest(BaseModel):
    """Request for exporting chat history."""
    messages: List[ApkChatMessage]
    analysis_context: Dict[str, Any]
    format: str = "markdown"  # "markdown", "json", "pdf"


class ChatExportResponse(BaseModel):
    """Response for chat export."""
    filename: str
    content_type: str


@router.post("/apk/chat/export")
async def export_apk_chat(
    request: ChatExportRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Export APK AI chat conversation to various formats.
    
    Supported formats:
    - markdown: Readable markdown format with conversation
    - json: Raw JSON export of messages and context
    - pdf: Formatted PDF document
    """
    from fastapi.responses import Response
    import json
    
    if not request.messages:
        raise HTTPException(status_code=400, detail="No messages to export")
    
    package_name = request.analysis_context.get('package_name', 'unknown')
    app_name = package_name.split('.')[-1] if package_name else 'apk'
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if request.format == "markdown":
        # Generate markdown chat export
        md_lines = [
            f"# APK Analysis Chat Export",
            f"",
            f"**Package:** {package_name}",
            f"**Exported:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"**Messages:** {len(request.messages)}",
            f"",
            f"---",
            f"",
            f"## Conversation",
            f"",
        ]
        
        for msg in request.messages:
            role_label = "**You:**" if msg.role == "user" else "**AI Assistant:**"
            timestamp_str = ""
            if msg.timestamp:
                ts = msg.timestamp if isinstance(msg.timestamp, datetime) else datetime.fromisoformat(str(msg.timestamp).replace('Z', '+00:00'))
                timestamp_str = f" *({ts.strftime('%H:%M:%S')})*"
            
            md_lines.append(f"### {role_label}{timestamp_str}")
            md_lines.append(f"")
            md_lines.append(msg.content)
            md_lines.append(f"")
            md_lines.append(f"---")
            md_lines.append(f"")
        
        # Add context summary
        md_lines.extend([
            f"## Analysis Context Summary",
            f"",
            f"- **Permissions:** {len(request.analysis_context.get('permissions', []))}",
            f"- **Security Issues:** {len(request.analysis_context.get('security_issues', []))}",
            f"- **Activities:** {len(request.analysis_context.get('activities', []))}",
            f"- **Services:** {len(request.analysis_context.get('services', []))}",
        ])
        
        content = "\n".join(md_lines)
        return Response(
            content=content.encode('utf-8'),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{app_name}_chat_{timestamp}.md"'}
        )
    
    elif request.format == "json":
        # JSON export
        export_data = {
            "exported_at": datetime.now().isoformat(),
            "package_name": package_name,
            "messages": [
                {
                    "role": msg.role,
                    "content": msg.content,
                    "timestamp": msg.timestamp.isoformat() if msg.timestamp else None
                }
                for msg in request.messages
            ],
            "analysis_summary": {
                "permissions_count": len(request.analysis_context.get('permissions', [])),
                "security_issues_count": len(request.analysis_context.get('security_issues', [])),
                "dangerous_permissions": [
                    p.get('name') for p in request.analysis_context.get('permissions', [])
                    if p.get('is_dangerous')
                ],
            }
        }
        
        content = json.dumps(export_data, indent=2)
        return Response(
            content=content.encode('utf-8'),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{app_name}_chat_{timestamp}.json"'}
        )
    
    elif request.format == "pdf":
        # PDF export using reportlab
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT
            import io
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
            
            styles = getSampleStyleSheet()
            title_style = styles['Heading1']
            heading_style = styles['Heading2']
            
            user_style = ParagraphStyle(
                'UserMessage',
                parent=styles['Normal'],
                backColor=colors.Color(0.9, 0.95, 1.0),
                borderPadding=8,
                leftIndent=20,
                rightIndent=20,
            )
            
            ai_style = ParagraphStyle(
                'AIMessage',
                parent=styles['Normal'],
                backColor=colors.Color(0.95, 0.95, 0.95),
                borderPadding=8,
                leftIndent=20,
                rightIndent=20,
            )
            
            elements = []
            
            # Title
            elements.append(Paragraph("APK Analysis Chat Export", title_style))
            elements.append(Spacer(1, 12))
            
            # Metadata
            elements.append(Paragraph(f"<b>Package:</b> {package_name}", styles['Normal']))
            elements.append(Paragraph(f"<b>Exported:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            elements.append(Paragraph(f"<b>Messages:</b> {len(request.messages)}", styles['Normal']))
            elements.append(Spacer(1, 20))
            
            # Messages
            elements.append(Paragraph("Conversation", heading_style))
            elements.append(Spacer(1, 12))
            
            for msg in request.messages:
                style = user_style if msg.role == "user" else ai_style
                role_label = "You" if msg.role == "user" else "AI Assistant"
                
                # Escape HTML characters
                safe_content = msg.content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                safe_content = safe_content.replace('\n', '<br/>')
                
                elements.append(Paragraph(f"<b>{role_label}:</b>", styles['Normal']))
                elements.append(Paragraph(safe_content, style))
                elements.append(Spacer(1, 12))
            
            doc.build(elements)
            
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{app_name}_chat_{timestamp}.pdf"'}
            )
            
        except ImportError:
            raise HTTPException(status_code=503, detail="PDF export requires reportlab library")
    
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {request.format}. Use markdown, json, or pdf.")


# ============================================================================
# AI-Assisted Code Explanation Endpoints
# ============================================================================

class CodeExplanationRequest(BaseModel):
    """Request for AI code explanation."""
    source_code: str
    class_name: str
    language: str = "java"  # java, smali, kotlin
    focus_area: Optional[str] = None  # "security", "functionality", "data_flow", None for general
    beginner_mode: bool = False


class CodeExplanationResponse(BaseModel):
    """Response with AI code explanation."""
    summary: str
    detailed_explanation: str
    security_concerns: List[Dict[str, Any]]
    interesting_findings: List[str]
    data_flow_analysis: Optional[str] = None
    suggested_focus_points: List[str]
    code_quality_notes: List[str]


@router.post("/apk/code/explain", response_model=CodeExplanationResponse)
async def explain_decompiled_code(
    request: CodeExplanationRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-powered explanation of decompiled code.
    
    Analyzes decompiled Java/Smali code and provides:
    - Natural language explanation of what the code does
    - Security vulnerability analysis
    - Data flow tracking
    - Interesting patterns and behaviors
    """
    from backend.core.config import settings
    
    if not settings.gemini_api_key:
        raise HTTPException(status_code=503, detail="AI features require Gemini API key")
    
    if not request.source_code.strip():
        raise HTTPException(status_code=400, detail="No source code provided")
    
    # Limit code size to prevent token overflow
    max_code_length = 15000
    code_to_analyze = request.source_code[:max_code_length]
    if len(request.source_code) > max_code_length:
        code_to_analyze += "\n\n// ... (code truncated for analysis)"
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        focus_prompt = ""
        if request.focus_area == "security":
            focus_prompt = """
Focus particularly on:
- Input validation and sanitization
- Authentication/authorization checks
- Cryptographic implementations
- Data leakage possibilities
- Injection vulnerabilities
- Insecure storage patterns
- Network security issues
"""
        elif request.focus_area == "functionality":
            focus_prompt = """
Focus particularly on:
- Main purpose and functionality
- User-facing features
- Data processing logic
- External integrations
- State management
"""
        elif request.focus_area == "data_flow":
            focus_prompt = """
Focus particularly on:
- Data sources and sinks
- Sensitive data handling
- Data transformations
- Storage locations
- Network transmissions
- Inter-component communication
"""
        
        beginner_note = """
Explain concepts simply using analogies. Define technical terms when you use them.
Assume the reader knows basic programming but not Android security.""" if request.beginner_mode else ""
        
        prompt = f"""You are an expert Android security researcher and code analyst.
Analyze this decompiled {request.language.upper()} code from class '{request.class_name}'.
{focus_prompt}
{beginner_note}

Provide your analysis in this JSON format:
{{
    "summary": "2-3 sentence summary of what this code does",
    "detailed_explanation": "Detailed explanation of the code's functionality, structure, and purpose",
    "security_concerns": [
        {{
            "severity": "critical|high|medium|low",
            "issue": "Description of the security issue",
            "location": "Method or line reference",
            "recommendation": "How to fix or exploit"
        }}
    ],
    "interesting_findings": ["List of interesting behaviors, patterns, or features"],
    "data_flow_analysis": "How data moves through this code (if applicable)",
    "suggested_focus_points": ["Areas worth investigating further"],
    "code_quality_notes": ["Observations about code quality, obfuscation, etc."]
}}

CODE TO ANALYZE:
```{request.language}
{code_to_analyze}
```

Return ONLY valid JSON, no other text."""

        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        # Parse JSON response
        import json
        response_text = response.text.strip()
        
        # Clean up potential markdown code blocks
        if response_text.startswith("```"):
            lines = response_text.split('\n')
            response_text = '\n'.join(lines[1:-1] if lines[-1].strip() == '```' else lines[1:])
        
        try:
            result = json.loads(response_text)
        except json.JSONDecodeError:
            # Try to extract JSON from response
            import re
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                result = json.loads(json_match.group())
            else:
                # Return a basic response if parsing fails
                result = {
                    "summary": "Code analysis completed but response parsing failed.",
                    "detailed_explanation": response_text,
                    "security_concerns": [],
                    "interesting_findings": [],
                    "data_flow_analysis": None,
                    "suggested_focus_points": ["Review the raw analysis above"],
                    "code_quality_notes": []
                }
        
        return CodeExplanationResponse(
            summary=result.get("summary", ""),
            detailed_explanation=result.get("detailed_explanation", ""),
            security_concerns=result.get("security_concerns", []),
            interesting_findings=result.get("interesting_findings", []),
            data_flow_analysis=result.get("data_flow_analysis"),
            suggested_focus_points=result.get("suggested_focus_points", []),
            code_quality_notes=result.get("code_quality_notes", [])
        )
        
    except Exception as e:
        logger.error(f"Code explanation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Code explanation failed: {str(e)}")


class CodeSearchAIRequest(BaseModel):
    """Request for AI-powered code search."""
    session_id: str
    query: str  # Natural language query like "find where API keys are stored"
    max_results: int = 20


class CodeSearchAIResponse(BaseModel):
    """Response for AI code search."""
    query: str
    interpreted_as: str  # How the AI interpreted the query
    search_patterns: List[str]  # Patterns used to search
    results: List[Dict[str, Any]]
    suggestions: List[str]  # Follow-up search suggestions


@router.post("/apk/code/search-ai", response_model=CodeSearchAIResponse)
async def ai_code_search(
    request: CodeSearchAIRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    AI-powered semantic code search in decompiled sources.
    
    Understands natural language queries like:
    - "Find where user passwords are handled"
    - "Show me network request code"
    - "Where is sensitive data stored"
    """
    from backend.core.config import settings
    
    if not settings.gemini_api_key:
        raise HTTPException(status_code=503, detail="AI features require Gemini API key")
    
    if request.session_id not in _jadx_cache:
        raise HTTPException(status_code=404, detail="Decompilation session not found. Please decompile the APK first.")
    
    output_dir = _jadx_cache[request.session_id]
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # First, use AI to understand the query and generate search patterns
        interpret_prompt = f"""You are an Android security researcher.
A user wants to search decompiled Java code for: "{request.query}"

Generate search patterns (regex-compatible strings) that would help find relevant code.
Think about:
- API method names that would be called
- Class names that might be involved
- Variable/field names commonly used
- String literals that might appear
- Android framework classes involved

Return JSON:
{{
    "interpretation": "What the user is looking for in plain English",
    "patterns": ["pattern1", "pattern2", ...],  // Up to 10 patterns
    "follow_up_suggestions": ["Other related things to search for"]
}}

Return ONLY valid JSON."""

        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=interpret_prompt)])],
        )
        
        import json
        response_text = response.text.strip()
        if response_text.startswith("```"):
            lines = response_text.split('\n')
            response_text = '\n'.join(lines[1:-1] if lines[-1].strip() == '```' else lines[1:])
        
        try:
            ai_interpretation = json.loads(response_text)
        except json.JSONDecodeError:
            ai_interpretation = {
                "interpretation": request.query,
                "patterns": [request.query],
                "follow_up_suggestions": []
            }
        
        # Search for each pattern
        results = []
        sources_dir = output_dir / "sources"
        
        if sources_dir.exists():
            for pattern in ai_interpretation.get("patterns", [])[:10]:
                for java_file in sources_dir.rglob("*.java"):
                    try:
                        content = java_file.read_text(encoding='utf-8', errors='ignore')
                        if pattern.lower() in content.lower():
                            # Find matching lines
                            lines = content.split('\n')
                            for line_num, line in enumerate(lines, 1):
                                if pattern.lower() in line.lower():
                                    results.append({
                                        "file_path": str(java_file.relative_to(sources_dir)),
                                        "line_number": line_num,
                                        "line_content": line.strip()[:200],
                                        "matched_pattern": pattern,
                                    })
                                    if len(results) >= request.max_results:
                                        break
                    except Exception:
                        continue
                    
                    if len(results) >= request.max_results:
                        break
                
                if len(results) >= request.max_results:
                    break
        
        return CodeSearchAIResponse(
            query=request.query,
            interpreted_as=ai_interpretation.get("interpretation", request.query),
            search_patterns=ai_interpretation.get("patterns", []),
            results=results[:request.max_results],
            suggestions=ai_interpretation.get("follow_up_suggestions", [])
        )
        
    except Exception as e:
        logger.error(f"AI code search failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI code search failed: {str(e)}")


# ============================================================================
# Crypto Audit Endpoints
# ============================================================================

class CryptoFinding(BaseModel):
    """A cryptographic vulnerability finding."""
    type: str
    category: str
    severity: str
    description: str
    recommendation: str
    file: str
    line: int
    match: str
    context: Optional[str] = None


class CryptoGoodPractice(BaseModel):
    """A good cryptographic practice found."""
    type: str
    file: str
    line: int
    match: str


class CryptoMethod(BaseModel):
    """A cryptographic method usage."""
    type: str
    algorithm: str
    file: str
    line: int


class CryptoAuditRequest(BaseModel):
    """Request for crypto audit."""
    output_directory: str


class CryptoAuditResponse(BaseModel):
    """Response with crypto audit results."""
    total_findings: int
    findings: List[CryptoFinding]
    by_severity: Dict[str, List[CryptoFinding]]
    by_category: Dict[str, List[CryptoFinding]]
    good_practices: List[CryptoGoodPractice]
    crypto_methods: List[CryptoMethod]
    files_scanned: int
    risk_score: int
    grade: str
    overall_risk: str
    top_recommendations: List[str]
    summary: str
    error: Optional[str] = None


@router.post("/apk/decompile/crypto-audit", response_model=CryptoAuditResponse)
async def crypto_audit(
    request: CryptoAuditRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Perform comprehensive cryptographic audit on decompiled APK sources.
    
    Detects:
    - Weak algorithms (MD5, SHA1, DES, 3DES, RC4)
    - ECB mode usage (insecure)
    - Hardcoded keys and IVs
    - Static/null IVs
    - Insecure random (java.util.Random)
    - RSA without OAEP padding
    - Weak PBKDF iterations
    - Certificate validation bypass
    
    Returns risk score, grade (A-F), and actionable recommendations.
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.crypto_audit(output_dir)
        
        if "error" in result:
            return CryptoAuditResponse(
                total_findings=0,
                findings=[],
                by_severity={},
                by_category={},
                good_practices=[],
                crypto_methods=[],
                files_scanned=0,
                risk_score=0,
                grade="?",
                overall_risk="unknown",
                top_recommendations=[],
                summary="",
                error=result["error"]
            )
        
        return CryptoAuditResponse(
            total_findings=result["total_findings"],
            findings=[CryptoFinding(**f) for f in result["findings"]],
            by_severity={k: [CryptoFinding(**f) for f in v] for k, v in result["by_severity"].items()},
            by_category={k: [CryptoFinding(**f) for f in v] for k, v in result["by_category"].items()},
            good_practices=[CryptoGoodPractice(**p) for p in result["good_practices"]],
            crypto_methods=[CryptoMethod(**m) for m in result["crypto_methods"]],
            files_scanned=result["files_scanned"],
            risk_score=result["risk_score"],
            grade=result["grade"],
            overall_risk=result["overall_risk"],
            top_recommendations=result["top_recommendations"],
            summary=result["summary"],
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Crypto audit failed: {e}")
        raise HTTPException(status_code=500, detail=f"Crypto audit failed: {str(e)}")


# ============================================================================
# Component Map Endpoints
# ============================================================================

class ComponentInfo(BaseModel):
    """Information about an Android component."""
    name: str
    full_name: str
    exported: bool
    risk: str


class ActivityInfo(ComponentInfo):
    """Activity component info."""
    launcher: bool = False
    actions: List[str] = []
    categories: List[str] = []
    data_schemes: List[str] = []
    theme: Optional[str] = None
    launch_mode: str = "standard"


class ServiceInfo(ComponentInfo):
    """Service component info."""
    actions: List[str] = []
    permission: Optional[str] = None
    foreground: bool = False


class ReceiverInfo(ComponentInfo):
    """Receiver component info."""
    actions: List[str] = []
    permission: Optional[str] = None
    system_broadcast: bool = False


class ProviderInfo(ComponentInfo):
    """Provider component info."""
    authorities: Optional[str] = None
    read_permission: Optional[str] = None
    write_permission: Optional[str] = None
    grant_uri_permissions: bool = False


class DeepLinkInfo(BaseModel):
    """Deep link info."""
    scheme: str
    host: Optional[str] = None
    path: Optional[str] = None
    component: str
    component_full: str
    type: str


class ConnectionInfo(BaseModel):
    """Component connection info."""
    source: str
    target: str
    type: str


class ComponentMapRequest(BaseModel):
    """Request for component map."""
    output_directory: str


class ComponentMapResponse(BaseModel):
    """Response with component map data."""
    package_name: str
    components: Dict[str, Any]  # activities, services, receivers, providers
    connections: List[ConnectionInfo]
    deep_links: List[DeepLinkInfo]
    stats: Dict[str, int]
    risk_counts: Dict[str, int]
    attack_surface_score: int
    summary: str
    error: Optional[str] = None


@router.post("/apk/decompile/component-map", response_model=ComponentMapResponse)
async def get_component_map(
    request: ComponentMapRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate visual component map showing activities, services, receivers,
    providers and their relationships.
    
    Returns:
    - All components with export status and risk levels
    - Deep links with schemes and hosts
    - Inter-component connections (intents)
    - Attack surface score
    - Statistics and risk breakdown
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.generate_component_map(output_dir)
        
        if "error" in result:
            return ComponentMapResponse(
                package_name="",
                components={},
                connections=[],
                deep_links=[],
                stats={},
                risk_counts={},
                attack_surface_score=0,
                summary="",
                error=result["error"]
            )
        
        return ComponentMapResponse(
            package_name=result["package_name"],
            components=result["components"],
            connections=[ConnectionInfo(**c) for c in result["connections"]],
            deep_links=[DeepLinkInfo(**d) for d in result["deep_links"]],
            stats=result["stats"],
            risk_counts=result["risk_counts"],
            attack_surface_score=result["attack_surface_score"],
            summary=result["summary"],
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Component map generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Component map generation failed: {str(e)}")


# ============================================================================
# Class Dependency Graph Endpoint
# ============================================================================

class GraphNode(BaseModel):
    """A node in the dependency graph."""
    id: str
    label: str
    full_name: str
    package: str
    type: str
    color: str
    size: int
    methods: int
    lines: int
    file_path: str


class GraphEdge(BaseModel):
    """An edge in the dependency graph."""
    from_: str = Field(..., alias="from")
    to: str
    type: str
    color: str
    dashes: Optional[Any] = None
    width: Optional[int] = None
    
    class Config:
        populate_by_name = True


class GraphStatistics(BaseModel):
    """Statistics about the dependency graph."""
    total_classes: int
    total_connections: int
    node_types: Dict[str, int]
    edge_types: Dict[str, int]
    packages: Dict[str, int]
    hub_classes: List[Dict[str, Any]]


class DependencyGraphRequest(BaseModel):
    """Request for class dependency graph."""
    output_directory: str
    max_classes: Optional[int] = 100


class DependencyGraphResponse(BaseModel):
    """Response with dependency graph data."""
    nodes: List[GraphNode]
    edges: List[Dict[str, Any]]  # Use Dict to avoid alias issues
    statistics: GraphStatistics
    legend: Dict[str, Dict[str, str]]
    error: Optional[str] = None


@router.post("/apk/decompile/dependency-graph")
async def get_dependency_graph(
    request: DependencyGraphRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Generate a class dependency graph showing how classes are interconnected.
    
    Analyzes:
    - Import statements (which classes depend on which)
    - Inheritance (extends relationships)
    - Interface implementation (implements)
    - Method calls between classes
    
    Returns graph data suitable for visualization with nodes and edges.
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.generate_class_dependency_graph(
            output_dir, 
            max_classes=request.max_classes or 100
        )
        
        if "error" in result:
            return {"error": result["error"]}
        
        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Dependency graph generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Dependency graph generation failed: {str(e)}")


# ============================================================================
# Symbol Lookup Endpoints (Jump to Definition)
# ============================================================================

class SymbolResult(BaseModel):
    """A symbol lookup result."""
    type: str  # class, method, field
    name: str
    file: str
    line: int
    # Optional fields depending on type
    package: Optional[str] = None
    full_name: Optional[str] = None
    class_name: Optional[str] = Field(None, alias="class")
    signature: Optional[str] = None
    return_type: Optional[str] = None
    params: Optional[str] = None
    field_type: Optional[str] = None
    
    class Config:
        populate_by_name = True


class SymbolLookupRequest(BaseModel):
    """Request for symbol lookup."""
    output_directory: str
    symbol: str
    symbol_type: Optional[str] = None  # class, method, field, or None for all


class SymbolLookupResponse(BaseModel):
    """Response with symbol lookup results."""
    symbol: str
    results: List[SymbolResult]
    total_found: int
    index_stats: Dict[str, int]
    error: Optional[str] = None


@router.post("/apk/decompile/symbol-lookup", response_model=SymbolLookupResponse)
async def lookup_symbol(
    request: SymbolLookupRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Look up a symbol (class, method, or field) and return its definition location.
    
    Enables jump-to-definition functionality in the source viewer.
    
    Args:
        symbol: Name to search for (supports partial matching)
        symbol_type: Filter by type (class/method/field) or None for all
    
    Returns file paths and line numbers for navigation.
    """
    try:
        output_dir = _resolve_jadx_output_dir(request.output_directory)
        
        result = re_service.lookup_symbol(output_dir, request.symbol, request.symbol_type)
        
        if "error" in result:
            return SymbolLookupResponse(
                symbol=request.symbol,
                results=[],
                total_found=0,
                index_stats={},
                error=result["error"]
            )
        
        return SymbolLookupResponse(
            symbol=result["symbol"],
            results=[SymbolResult(**r) for r in result["results"]],
            total_found=result["total_found"],
            index_stats=result["index_stats"],
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Symbol lookup failed: {e}")
        raise HTTPException(status_code=500, detail=f"Symbol lookup failed: {str(e)}")


# ============================================================================
# Enhanced Security Export
# ============================================================================

class ExportEnhancedSecurityRequest(BaseModel):
    """Request to export enhanced security results."""
    results: Dict[str, Any]
    format: str = "markdown"  # markdown, pdf, docx
    include_code_snippets: bool = True
    include_attack_chains: bool = True


@router.post("/apk/decompile/enhanced-security/export")
async def export_enhanced_security(
    request: ExportEnhancedSecurityRequest,
    current_user: User = Depends(get_current_active_user),
):
    """
    Export enhanced security analysis results to Markdown, PDF, or Word format.
    
    Args:
        results: The enhanced security results to export
        format: Export format (markdown, pdf, docx)
        include_code_snippets: Include code snippets in the export
        include_attack_chains: Include attack chain analysis
    
    Returns the exported document as a downloadable file.
    """
    from fastapi.responses import Response
    from io import BytesIO
    from datetime import datetime
    
    if request.format not in ["markdown", "pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Invalid format. Use: markdown, pdf, docx")
    
    try:
        results = request.results
        
        # Generate Markdown content
        md_lines = [
            "# Comprehensive Security Analysis Report",
            f"",
            f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"",
            "---",
            "",
            "## Executive Summary",
            "",
            f"**Overall Risk Level:** {results.get('overall_risk', 'unknown').upper()}",
            "",
            results.get('executive_summary', 'No executive summary available.'),
            "",
            "---",
            "",
            "## Risk Summary",
            "",
            "| Severity | Count |",
            "|----------|-------|",
        ]
        
        risk_summary = results.get('risk_summary', {})
        for severity in ['critical', 'high', 'medium', 'low', 'info']:
            count = risk_summary.get(severity, 0)
            emoji = {'critical': '🔴', 'high': '🟠', 'medium': '🟡', 'low': '🟢', 'info': '🔵'}.get(severity, '')
            md_lines.append(f"| {emoji} {severity.capitalize()} | {count} |")
        
        md_lines.extend(["", "---", "", "## Analysis Metadata", ""])
        
        metadata = results.get('analysis_metadata', {})
        md_lines.extend([
            f"- **Classes Scanned:** {metadata.get('classes_scanned', 0)}",
            f"- **Libraries Detected:** {metadata.get('libraries_detected', 0)}",
            f"- **CVEs Found:** {metadata.get('cves_found', 0)}",
            f"- **Pattern Scan:** {'✅ Enabled' if metadata.get('pattern_scan_enabled') else '❌ Disabled'}",
            f"- **AI Scan:** {'✅ Enabled' if metadata.get('ai_scan_enabled') else '❌ Disabled'}",
            f"- **CVE Lookup:** {'✅ Enabled' if metadata.get('cve_lookup_enabled') else '❌ Disabled'}",
            "",
            "---",
            "",
        ])
        
        # Combined Findings
        combined_findings = results.get('combined_findings', [])
        if combined_findings:
            md_lines.extend([
                "## Security Findings",
                "",
                f"Total: **{len(combined_findings)}** findings",
                "",
            ])
            
            # Group by severity
            for severity in ['critical', 'high', 'medium', 'low', 'info']:
                severity_findings = [f for f in combined_findings if f.get('severity', '').lower() == severity]
                if severity_findings:
                    md_lines.extend([
                        f"### {severity.upper()} Severity ({len(severity_findings)})",
                        "",
                    ])
                    
                    for i, finding in enumerate(severity_findings, 1):
                        source_badge = f"[{finding.get('source', 'unknown').upper()}]"
                        md_lines.extend([
                            f"#### {i}. {finding.get('title', 'Unknown Issue')} {source_badge}",
                            "",
                            f"**Description:** {finding.get('description', 'No description')}",
                            "",
                        ])
                        
                        if finding.get('affected_class'):
                            location = finding.get('affected_class')
                            if finding.get('affected_method'):
                                location += f" → {finding.get('affected_method')}"
                            if finding.get('line_number'):
                                location += f" (line {finding.get('line_number')})"
                            md_lines.append(f"**Location:** `{location}`")
                            md_lines.append("")
                        
                        if finding.get('affected_library'):
                            md_lines.append(f"**Library:** {finding.get('affected_library')}")
                            md_lines.append("")
                        
                        if finding.get('cve_id'):
                            md_lines.append(f"**CVE:** {finding.get('cve_id')}")
                            md_lines.append("")
                        
                        if finding.get('cwe_id'):
                            md_lines.append(f"**CWE:** {finding.get('cwe_id')}")
                            md_lines.append("")
                        
                        if finding.get('cvss_score'):
                            md_lines.append(f"**CVSS Score:** {finding.get('cvss_score')}")
                            md_lines.append("")
                        
                        if finding.get('impact'):
                            md_lines.append(f"**Impact:** {finding.get('impact')}")
                            md_lines.append("")
                        
                        if finding.get('exploitation_potential'):
                            md_lines.append(f"**Exploitation Potential:** {finding.get('exploitation_potential')}")
                            md_lines.append("")
                        
                        if finding.get('attack_vector'):
                            md_lines.append(f"**Attack Vector:** {finding.get('attack_vector')}")
                            md_lines.append("")
                        
                        if request.include_code_snippets and finding.get('code_snippet'):
                            md_lines.extend([
                                "**Code:**",
                                "```java",
                                finding.get('code_snippet'),
                                "```",
                                "",
                            ])
                        
                        if finding.get('remediation'):
                            md_lines.extend([
                                "**Remediation:**",
                                f"> {finding.get('remediation')}",
                                "",
                            ])
                        
                        md_lines.append("---")
                        md_lines.append("")
        
        # Attack Chains
        if request.include_attack_chains:
            attack_chains = results.get('attack_chains', [])
            if attack_chains:
                md_lines.extend([
                    "## Attack Chain Analysis",
                    "",
                    f"**{len(attack_chains)}** potential attack chains identified:",
                    "",
                ])
                
                for i, chain in enumerate(attack_chains, 1):
                    md_lines.extend([
                        f"### Chain {i}: {chain.get('name', 'Unknown Chain')}",
                        "",
                        f"**Risk Level:** {chain.get('risk_level', 'Unknown')}",
                        "",
                        f"**Description:** {chain.get('description', 'No description')}",
                        "",
                        "**Steps:**",
                    ])
                    
                    for step in chain.get('steps', []):
                        md_lines.append(f"1. {step}")
                    
                    if chain.get('impact'):
                        md_lines.extend(["", f"**Impact:** {chain.get('impact')}"])
                    
                    if chain.get('likelihood'):
                        md_lines.extend(["", f"**Likelihood:** {chain.get('likelihood')}"])
                    
                    md_lines.extend(["", "---", ""])
        
        # Recommendations
        recommendations = results.get('recommendations', [])
        if recommendations:
            md_lines.extend([
                "## Recommendations",
                "",
            ])
            for rec in recommendations:
                md_lines.append(f"- {rec}")
            md_lines.append("")
        
        # Footer
        md_lines.extend([
            "---",
            "",
            "*Generated by VRAgent Security Analyzer*",
        ])
        
        markdown_content = "\n".join(md_lines)
        
        # Return based on format
        if request.format == "markdown":
            return Response(
                content=markdown_content.encode('utf-8'),
                media_type="text/markdown",
                headers={"Content-Disposition": "attachment; filename=security_analysis_report.md"}
            )
        
        elif request.format == "pdf":
            try:
                from weasyprint import HTML, CSS
                from markdown import markdown
                
                html_content = markdown(markdown_content, extensions=['tables', 'fenced_code'])
                
                styled_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                        h1 {{ color: #1e3a5f; border-bottom: 3px solid #e74c3c; padding-bottom: 10px; }}
                        h2 {{ color: #2c5282; border-bottom: 2px solid #3182ce; padding-bottom: 5px; margin-top: 30px; }}
                        h3 {{ color: #2d3748; margin-top: 20px; }}
                        h4 {{ color: #4a5568; margin-top: 15px; }}
                        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
                        th, td {{ border: 1px solid #e2e8f0; padding: 10px; text-align: left; }}
                        th {{ background-color: #edf2f7; color: #2d3748; }}
                        code {{ background-color: #f7fafc; padding: 2px 6px; border-radius: 4px; font-family: 'Consolas', monospace; }}
                        pre {{ background-color: #2d3748; color: #e2e8f0; padding: 15px; border-radius: 8px; overflow-x: auto; }}
                        pre code {{ background: none; padding: 0; color: #e2e8f0; }}
                        blockquote {{ border-left: 4px solid #48bb78; padding-left: 15px; margin: 15px 0; color: #2f855a; background: #f0fff4; padding: 10px 15px; }}
                        hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 20px 0; }}
                        ul {{ padding-left: 25px; }}
                        li {{ margin: 5px 0; }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                
                pdf_buffer = BytesIO()
                HTML(string=styled_html).write_pdf(pdf_buffer)
                pdf_buffer.seek(0)
                
                return Response(
                    content=pdf_buffer.getvalue(),
                    media_type="application/pdf",
                    headers={"Content-Disposition": "attachment; filename=security_analysis_report.pdf"}
                )
            except ImportError:
                raise HTTPException(status_code=500, detail="PDF export requires weasyprint. Install with: pip install weasyprint")
        
        elif request.format == "docx":
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.style import WD_STYLE_TYPE
                
                doc = Document()
                
                # Title
                title = doc.add_heading('Comprehensive Security Analysis Report', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph()
                
                # Executive Summary
                doc.add_heading('Executive Summary', level=1)
                risk_para = doc.add_paragraph()
                risk_para.add_run(f"Overall Risk Level: ").bold = True
                risk_run = risk_para.add_run(results.get('overall_risk', 'unknown').upper())
                risk_run.bold = True
                risk_color = {'critical': RGBColor(220, 38, 38), 'high': RGBColor(234, 88, 12), 
                             'medium': RGBColor(202, 138, 4), 'low': RGBColor(22, 163, 74)}.get(
                    results.get('overall_risk', '').lower(), RGBColor(107, 114, 128))
                risk_run.font.color.rgb = risk_color
                
                doc.add_paragraph(results.get('executive_summary', 'No executive summary available.'))
                
                # Risk Summary Table
                doc.add_heading('Risk Summary', level=1)
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Severity'
                hdr_cells[1].text = 'Count'
                
                for severity in ['critical', 'high', 'medium', 'low', 'info']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = severity.capitalize()
                    row_cells[1].text = str(risk_summary.get(severity, 0))
                
                # Metadata
                doc.add_heading('Analysis Metadata', level=1)
                doc.add_paragraph(f"Classes Scanned: {metadata.get('classes_scanned', 0)}")
                doc.add_paragraph(f"Libraries Detected: {metadata.get('libraries_detected', 0)}")
                doc.add_paragraph(f"CVEs Found: {metadata.get('cves_found', 0)}")
                
                # Findings
                if combined_findings:
                    doc.add_heading('Security Findings', level=1)
                    doc.add_paragraph(f"Total: {len(combined_findings)} findings")
                    
                    for severity in ['critical', 'high', 'medium', 'low', 'info']:
                        severity_findings = [f for f in combined_findings if f.get('severity', '').lower() == severity]
                        if severity_findings:
                            doc.add_heading(f'{severity.upper()} Severity ({len(severity_findings)})', level=2)
                            
                            for finding in severity_findings:
                                p = doc.add_paragraph()
                                p.add_run(f"{finding.get('title', 'Unknown')} ").bold = True
                                p.add_run(f"[{finding.get('source', 'unknown').upper()}]")
                                
                                doc.add_paragraph(finding.get('description', 'No description'))
                                
                                if finding.get('affected_class'):
                                    doc.add_paragraph(f"Location: {finding.get('affected_class')}")
                                
                                if finding.get('remediation'):
                                    rem_para = doc.add_paragraph()
                                    rem_para.add_run("Remediation: ").bold = True
                                    rem_para.add_run(finding.get('remediation'))
                                
                                doc.add_paragraph()  # Spacing
                
                # Recommendations
                if recommendations:
                    doc.add_heading('Recommendations', level=1)
                    for rec in recommendations:
                        doc.add_paragraph(rec, style='List Bullet')
                
                # Save to buffer
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                return Response(
                    content=docx_buffer.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": "attachment; filename=security_analysis_report.docx"}
                )
            except ImportError:
                raise HTTPException(status_code=500, detail="DOCX export requires python-docx. Install with: pip install python-docx")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Enhanced security export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


# ============================================================================
# ADVANCED BINARY ANALYSIS ENDPOINTS
# Symbolic Execution, Binary Diffing, ROP Gadget Finding
# ============================================================================

class SymbolicInputResponse(BaseModel):
    name: str
    type: str
    size_bits: int
    constraints: List[str] = []
    concrete_examples: List[str] = []


class CrashInputResponse(BaseModel):
    input_type: str
    input_value: str
    crash_address: int
    crash_type: str
    vulnerability_type: str
    cwe_id: Optional[str] = None
    exploitability: str


class TargetReachResponse(BaseModel):
    target_address: int
    target_name: Optional[str] = None
    reached: bool
    input_to_reach: Optional[str] = None
    path_length: int
    constraints_solved: int


class VulnFoundResponse(BaseModel):
    type: str
    function: Optional[str] = None
    address: int
    cwe: Optional[str] = None
    input_sample: Optional[str] = None
    description: str


class SymbolicExecutionResponse(BaseModel):
    paths_explored: int
    paths_deadended: int
    paths_errored: int
    max_depth_reached: int
    symbolic_inputs: List[SymbolicInputResponse] = []
    crash_inputs: List[CrashInputResponse] = []
    target_reaches: List[TargetReachResponse] = []
    vulnerabilities_found: List[VulnFoundResponse] = []
    execution_time_seconds: float
    memory_used_mb: float
    timeout_reached: bool = False
    error: Optional[str] = None


@router.post("/binary/symbolic-execution", response_model=SymbolicExecutionResponse)
async def perform_symbolic_execution_endpoint(
    file: UploadFile = File(..., description="Binary file to analyze"),
    target_addresses: str = Query("", description="Comma-separated hex addresses to target (e.g., '0x401000,0x402000')"),
    max_time_seconds: int = Query(120, ge=10, le=600, description="Maximum execution time"),
    max_paths: int = Query(100, ge=10, le=1000, description="Maximum paths to explore"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Perform symbolic execution on a binary to discover vulnerabilities.
    
    Symbolic execution explores all possible program paths by treating inputs
    as symbolic values and using constraint solving to find inputs that:
    - Reach specific code addresses (like dangerous functions)
    - Trigger crashes or undefined behavior
    - Satisfy complex conditions
    
    This is a powerful technique for automatic vulnerability discovery.
    """
    tmp_dir = None
    
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Parse target addresses
        targets = []
        if target_addresses.strip():
            for addr_str in target_addresses.split(','):
                addr_str = addr_str.strip()
                if addr_str:
                    try:
                        targets.append(int(addr_str, 16) if addr_str.startswith('0x') else int(addr_str))
                    except ValueError:
                        raise HTTPException(status_code=400, detail=f"Invalid address: {addr_str}")
        
        # Save file temporarily
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_symbolic_"))
        tmp_path = tmp_dir / file.filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Performing symbolic execution on: {file.filename}")
        
        # Perform symbolic execution
        result = re_service.perform_symbolic_execution(
            tmp_path,
            target_addresses=targets if targets else None,
            max_time_seconds=max_time_seconds,
            max_paths=max_paths
        )
        
        return SymbolicExecutionResponse(
            paths_explored=result.paths_explored,
            paths_deadended=result.paths_deadended,
            paths_errored=result.paths_errored,
            max_depth_reached=result.max_depth_reached,
            symbolic_inputs=[
                SymbolicInputResponse(
                    name=si.name,
                    type=si.type,
                    size_bits=si.size_bits,
                    constraints=si.constraints,
                    concrete_examples=si.concrete_examples
                )
                for si in result.symbolic_inputs
            ],
            crash_inputs=[
                CrashInputResponse(
                    input_type=ci.input_type,
                    input_value=ci.input_value,
                    crash_address=ci.crash_address,
                    crash_type=ci.crash_type,
                    vulnerability_type=ci.vulnerability_type,
                    cwe_id=ci.cwe_id,
                    exploitability=ci.exploitability
                )
                for ci in result.crash_inputs
            ],
            target_reaches=[
                TargetReachResponse(
                    target_address=tr.target_address,
                    target_name=tr.target_name,
                    reached=tr.reached,
                    input_to_reach=tr.input_to_reach,
                    path_length=tr.path_length,
                    constraints_solved=tr.constraints_solved
                )
                for tr in result.target_reaches
            ],
            vulnerabilities_found=[
                VulnFoundResponse(
                    type=vf["type"],
                    function=vf.get("function"),
                    address=vf["address"],
                    cwe=vf.get("cwe"),
                    input_sample=vf.get("input_sample"),
                    description=vf["description"]
                )
                for vf in result.vulnerabilities_found
            ],
            execution_time_seconds=result.execution_time_seconds,
            memory_used_mb=result.memory_used_mb,
            timeout_reached=result.timeout_reached,
            error=result.error
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Symbolic execution failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Symbolic execution failed: {str(e)}")
    finally:
        if tmp_dir and tmp_dir.exists():
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass


class FunctionDiffResponse(BaseModel):
    address_a: int
    address_b: Optional[int] = None
    name: str
    match_type: str
    similarity_score: float
    size_a: int
    size_b: Optional[int] = None
    instructions_changed: int
    blocks_changed: int
    calls_added: List[str] = []
    calls_removed: List[str] = []
    is_security_relevant: bool
    diff_summary: Optional[str] = None


class StringDiffResponse(BaseModel):
    value: str
    status: str
    address_a: Optional[int] = None
    address_b: Optional[int] = None
    is_security_relevant: bool


class ImportDiffResponse(BaseModel):
    name: str
    library: str
    status: str
    is_security_relevant: bool


class SecurityChangeResponse(BaseModel):
    type: str
    function: Optional[str] = None
    value: Optional[str] = None
    description: str


class BinaryDiffResponse(BaseModel):
    file_a: str
    file_b: str
    architecture_a: str
    architecture_b: str
    functions_identical: int
    functions_modified: int
    functions_added: int
    functions_removed: int
    overall_similarity: float
    function_diffs: List[FunctionDiffResponse] = []
    string_diffs: List[StringDiffResponse] = []
    import_diffs: List[ImportDiffResponse] = []
    security_relevant_changes: List[SecurityChangeResponse] = []
    patch_analysis: Optional[str] = None
    is_same_binary: bool
    error: Optional[str] = None


@router.post("/binary/diff", response_model=BinaryDiffResponse)
async def diff_binaries_endpoint(
    file_a: UploadFile = File(..., description="First binary file (original/old version)"),
    file_b: UploadFile = File(..., description="Second binary file (patched/new version)"),
    detailed: bool = Query(True, description="Include detailed block-level diffs"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Compare two binary files to identify differences.
    
    Binary diffing is essential for:
    - Patch analysis (understanding what a security patch fixes)
    - Vulnerability research (1-day exploit development)
    - Malware analysis (comparing variants)
    - Software forensics (detecting unauthorized modifications)
    
    Identifies function-level, string-level, and import-level differences
    with automatic security-relevance detection.
    """
    tmp_dir = None
    
    try:
        if not file_a.filename or not file_b.filename:
            raise HTTPException(status_code=400, detail="Both files are required")
        
        # Save files temporarily
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_bindiff_"))
        tmp_path_a = tmp_dir / f"a_{file_a.filename}"
        tmp_path_b = tmp_dir / f"b_{file_b.filename}"
        
        # Save file A
        file_size = 0
        with tmp_path_a.open("wb") as f:
            while chunk := await file_a.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File A too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        # Save file B
        file_size = 0
        with tmp_path_b.open("wb") as f:
            while chunk := await file_b.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File B too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Diffing binaries: {file_a.filename} vs {file_b.filename}")
        
        # Perform binary diff
        result = re_service.diff_binaries(tmp_path_a, tmp_path_b, detailed=detailed)
        
        return BinaryDiffResponse(
            file_a=result.file_a,
            file_b=result.file_b,
            architecture_a=result.architecture_a,
            architecture_b=result.architecture_b,
            functions_identical=result.functions_identical,
            functions_modified=result.functions_modified,
            functions_added=result.functions_added,
            functions_removed=result.functions_removed,
            overall_similarity=result.overall_similarity,
            function_diffs=[
                FunctionDiffResponse(
                    address_a=fd.address_a,
                    address_b=fd.address_b,
                    name=fd.name,
                    match_type=fd.match_type,
                    similarity_score=fd.similarity_score,
                    size_a=fd.size_a,
                    size_b=fd.size_b,
                    instructions_changed=fd.instructions_changed,
                    blocks_changed=fd.blocks_changed,
                    calls_added=fd.calls_added,
                    calls_removed=fd.calls_removed,
                    is_security_relevant=fd.is_security_relevant,
                    diff_summary=fd.diff_summary
                )
                for fd in result.function_diffs
            ],
            string_diffs=[
                StringDiffResponse(
                    value=sd.value,
                    status=sd.status,
                    address_a=sd.address_a,
                    address_b=sd.address_b,
                    is_security_relevant=sd.is_security_relevant
                )
                for sd in result.string_diffs
            ],
            import_diffs=[
                ImportDiffResponse(
                    name=id.name,
                    library=id.library,
                    status=id.status,
                    is_security_relevant=id.is_security_relevant
                )
                for id in result.import_diffs
            ],
            security_relevant_changes=[
                SecurityChangeResponse(
                    type=sc["type"],
                    function=sc.get("function"),
                    value=sc.get("value"),
                    description=sc["description"]
                )
                for sc in result.security_relevant_changes
            ],
            patch_analysis=result.patch_analysis,
            is_same_binary=result.is_same_binary,
            error=result.error
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Binary diffing failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Binary diffing failed: {str(e)}")
    finally:
        if tmp_dir and tmp_dir.exists():
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass


class ROPGadgetResponse(BaseModel):
    address: int
    instructions: List[str]
    gadget_string: str
    gadget_type: str
    size_bytes: int
    registers_controlled: List[str] = []
    is_useful: bool
    quality_score: float


class ROPChainTemplateResponse(BaseModel):
    name: str
    description: str
    required_gadgets: List[str]
    available_gadgets: List[ROPGadgetResponse] = []
    is_buildable: bool
    chain_addresses: List[int] = []
    payload_template: Optional[str] = None


class ROPGadgetAnalysisResponse(BaseModel):
    total_gadgets: int
    unique_gadgets: int
    gadgets_by_type: Dict[str, int] = {}
    gadgets: List[ROPGadgetResponse] = []
    useful_gadgets: List[ROPGadgetResponse] = []
    pop_gadgets: List[ROPGadgetResponse] = []
    syscall_gadgets: List[ROPGadgetResponse] = []
    write_gadgets: List[ROPGadgetResponse] = []
    pivot_gadgets: List[ROPGadgetResponse] = []
    chain_templates: List[ROPChainTemplateResponse] = []
    nx_bypass_possible: bool
    execve_chain_buildable: bool
    mprotect_chain_buildable: bool
    rop_difficulty: str
    error: Optional[str] = None


@router.post("/binary/rop-gadgets", response_model=ROPGadgetAnalysisResponse)
async def find_rop_gadgets_endpoint(
    file: UploadFile = File(..., description="Binary file to analyze"),
    max_gadgets: int = Query(1000, ge=100, le=10000, description="Maximum gadgets to find"),
    max_gadget_length: int = Query(10, ge=2, le=20, description="Maximum instructions per gadget"),
    current_user: User = Depends(get_current_active_user),
):
    """
    Find ROP (Return-Oriented Programming) gadgets in a binary.
    
    ROP gadgets are small instruction sequences ending in 'ret' that can be
    chained together to bypass DEP/NX protections. This analysis:
    - Finds all usable gadgets
    - Classifies gadgets by type (pop, mov, syscall, etc.)
    - Evaluates which common exploit chains can be built
    - Assesses overall ROP-based exploitability
    
    Essential for exploit development and security assessment.
    """
    tmp_dir = None
    
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Save file temporarily
        tmp_dir = Path(tempfile.mkdtemp(prefix="vragent_rop_"))
        tmp_path = tmp_dir / file.filename
        
        file_size = 0
        with tmp_path.open("wb") as f:
            while chunk := await file.read(65536):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"File too large. Maximum: {MAX_FILE_SIZE // (1024*1024)}MB"
                    )
                f.write(chunk)
        
        logger.info(f"Finding ROP gadgets in: {file.filename}")
        
        # Find ROP gadgets
        result = re_service.find_rop_gadgets(
            tmp_path,
            max_gadgets=max_gadgets,
            max_gadget_length=max_gadget_length
        )
        
        def convert_gadget(g) -> ROPGadgetResponse:
            return ROPGadgetResponse(
                address=g.address,
                instructions=g.instructions,
                gadget_string=g.gadget_string,
                gadget_type=g.gadget_type,
                size_bytes=g.size_bytes,
                registers_controlled=g.registers_controlled,
                is_useful=g.is_useful,
                quality_score=g.quality_score
            )
        
        return ROPGadgetAnalysisResponse(
            total_gadgets=result.total_gadgets,
            unique_gadgets=result.unique_gadgets,
            gadgets_by_type=result.gadgets_by_type,
            gadgets=[convert_gadget(g) for g in result.gadgets[:200]],  # Limit response size
            useful_gadgets=[convert_gadget(g) for g in result.useful_gadgets],
            pop_gadgets=[convert_gadget(g) for g in result.pop_gadgets],
            syscall_gadgets=[convert_gadget(g) for g in result.syscall_gadgets],
            write_gadgets=[convert_gadget(g) for g in result.write_gadgets],
            pivot_gadgets=[convert_gadget(g) for g in result.pivot_gadgets],
            chain_templates=[
                ROPChainTemplateResponse(
                    name=ct.name,
                    description=ct.description,
                    required_gadgets=ct.required_gadgets,
                    available_gadgets=[convert_gadget(g) for g in ct.available_gadgets],
                    is_buildable=ct.is_buildable,
                    chain_addresses=ct.chain_addresses,
                    payload_template=ct.payload_template
                )
                for ct in result.chain_templates
            ],
            nx_bypass_possible=result.nx_bypass_possible,
            execve_chain_buildable=result.execve_chain_buildable,
            mprotect_chain_buildable=result.mprotect_chain_buildable,
            rop_difficulty=result.rop_difficulty,
            error=result.error
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ROP gadget finding failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"ROP gadget finding failed: {str(e)}")
    finally:
        if tmp_dir and tmp_dir.exists():
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass
