"""
Dynamic Security Scanner API Router

REST API endpoints for the unified Dynamic Security Scanner that
combines Nmap, ZAP, Nuclei, and Exploit DB into an AI-orchestrated
automated pentesting workflow.
"""

import asyncio
import json
import logging
from dataclasses import asdict
from datetime import datetime
from typing import Optional, List, Dict, Literal, Any

from fastapi import APIRouter, Depends, HTTPException, WebSocket, WebSocketDisconnect, Query
from pydantic import BaseModel, Field, field_validator
from sqlalchemy.orm import Session

from backend.core.auth import get_current_active_user, get_optional_user
from backend.core.database import get_db
from backend.models.models import User, DynamicScan, DynamicScanFinding
from backend.services.dynamic_scan_service import (
    get_dynamic_scan_service,
    ScanConcurrencyError,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/dynamic-scan", tags=["Dynamic Security Scanner"])


# ============== Helpers ==============

def _extract_next_action(response_text: str) -> Optional[str]:
    """Read the NEXT_ACTION token that signals what the agent wants to do next."""
    for line in response_text.splitlines():
        if line.strip().upper().startswith("NEXT_ACTION:"):
            return line.split(":", 1)[1].strip().lower()
    return None


def _extract_final_response(response_text: str) -> str:
    """Return the FINAL_RESPONSE text if present; otherwise return the assistant text."""
    marker = "FINAL_RESPONSE:"
    upper_text = response_text.upper()
    if marker in upper_text:
        start = upper_text.index(marker) + len(marker)
        return response_text[start:].strip()
    return response_text.strip()


def _summarize_scan_context(scan_context: str) -> str:
    """Generate a quick severity summary that can be fed back into the agent."""
    if not scan_context:
        return "No scan context is currently available."
    snippet = " ".join(scan_context.splitlines())[:400]
    severity_levels = ["Critical", "High", "Medium", "Low"]
    counts = {level: scan_context.lower().count(level.lower()) for level in severity_levels}
    return f"Severity counts: {counts}. Sample context: {snippet}..."


# ============== Request/Response Models ==============

class ZapAuthRequest(BaseModel):
    """Optional authentication configuration for ZAP authenticated scans."""
    method: str = Field(
        ...,
        description="Auth method: formBasedAuthentication, jsonBasedAuthentication, httpAuthentication, scriptBasedAuthentication"
    )
    login_url: Optional[str] = Field(default=None, description="Login endpoint URL")
    login_request_data: Optional[str] = Field(
        default=None,
        description="Login request data with {%username%} and {%password%} placeholders"
    )
    json_template: Optional[str] = Field(
        default=None,
        description="JSON template for login with placeholders"
    )
    hostname: Optional[str] = Field(default=None, description="Hostname for HTTP Basic auth")
    realm: Optional[str] = Field(default=None, description="Realm for HTTP Basic auth")
    port: Optional[int] = Field(default=None, description="Port for HTTP Basic auth")
    script_name: Optional[str] = Field(default=None, description="ZAP auth script name")
    script_params: Optional[Dict[str, str]] = Field(default=None, description="Auth script parameters")
    username: Optional[str] = Field(default=None, description="Username for authentication")
    password: Optional[str] = Field(default=None, description="Password for authentication")
    logged_in_indicator: Optional[str] = Field(default=None, description="Regex indicating logged-in state")
    logged_out_indicator: Optional[str] = Field(default=None, description="Regex indicating logged-out state")
    context_name: Optional[str] = Field(default=None, description="Custom ZAP context name")


class OpenVASCredentialRequest(BaseModel):
    """OpenVAS credential IDs for authenticated scanning."""
    ssh_credential_id: Optional[str] = Field(default=None, description="OpenVAS SSH credential ID")
    ssh_credential_port: Optional[int] = Field(default=22, description="SSH port for OpenVAS credential")
    smb_credential_id: Optional[str] = Field(default=None, description="OpenVAS SMB credential ID")
    snmp_credential_id: Optional[str] = Field(default=None, description="OpenVAS SNMP credential ID")
    esxi_credential_id: Optional[str] = Field(default=None, description="OpenVAS ESXi credential ID")


class BrowserCrawlRequest(BaseModel):
    """Optional headless browser crawl configuration."""
    enabled: bool = Field(default=False, description="Enable headless browser crawl and HAR capture")
    start_url: Optional[str] = Field(default=None, description="URL to start crawling from")
    login_url: Optional[str] = Field(default=None, description="Login page URL for browser crawl")
    username: Optional[str] = Field(default=None, description="Username for browser login")
    password: Optional[str] = Field(default=None, description="Password for browser login")
    username_selector: Optional[str] = Field(default=None, description="CSS selector for username field")
    password_selector: Optional[str] = Field(default=None, description="CSS selector for password field")
    submit_selector: Optional[str] = Field(default=None, description="CSS selector for login submit")
    wait_for_selector: Optional[str] = Field(default=None, description="Selector indicating login completed")
    max_pages: int = Field(default=15, ge=1, le=200, description="Maximum pages to crawl")
    max_duration_seconds: int = Field(default=120, ge=10, le=900, description="Maximum crawl duration")
    same_origin_only: bool = Field(default=True, description="Restrict crawl to the same origin")
    record_har: bool = Field(default=True, description="Record HAR during crawl")


class StartScanRequest(BaseModel):
    """Request to start a dynamic security scan.

    Supports three modes:
    - AI-led mode (recommended): Set ai_led=True, provide target and optional context.
      AI automatically determines scan strategy based on target analysis.
    - IP/Network mode: Provide IP address, CIDR range, or hostname for full reconnaissance
    - URL mode: Provide full URL (https://...) for direct web application scanning
    """
    # Scan identification
    scan_name: Optional[str] = Field(
        default=None,
        description="Optional name for this scan (e.g., 'Production Server Audit'). Makes it easier to identify in saved scans.",
        max_length=200
    )

    target: str = Field(
        ...,
        description="Target to scan. Can be: IP address (192.168.1.1), CIDR range (192.168.1.0/24), hostname (example.com), or full URL for web apps (https://app.example.com/path)"
    )

    @field_validator('target')
    @classmethod
    def validate_target(cls, v):
        """Validate scan target to prevent SSRF and invalid inputs."""
        import ipaddress
        from urllib.parse import urlparse
        from backend.core.config import get_settings
        
        _settings = get_settings()
        v = v.strip()
        if not v:
            raise ValueError("Target cannot be empty")

        # Allow localhost/private IPs if configured (for development/internal testing)
        if _settings.allow_localhost_scans:
            # Skip localhost validation - useful for testing against local targets like Juice Shop
            return v

        # Block localhost variants to prevent SSRF
        localhost_patterns = ['localhost', '127.', '0.0.0.0', '::1', '0:0:0:0:0:0:0:1']
        target_lower = v.lower()
        for pattern in localhost_patterns:
            if pattern in target_lower:
                raise ValueError(f"Localhost/loopback scanning not allowed for security reasons: {v}. Set ALLOW_LOCALHOST_SCANS=true to enable.")

        # If it's a URL, validate the hostname
        if v.startswith(('http://', 'https://')):
            try:
                parsed = urlparse(v)
                hostname = parsed.hostname
                if hostname:
                    # Check if hostname resolves to localhost
                    for pattern in localhost_patterns:
                        if pattern in hostname.lower():
                            raise ValueError(f"URL resolves to localhost: {v}")

                    # Try to parse hostname as IP to check for private ranges
                    try:
                        ip = ipaddress.ip_address(hostname)
                        if ip.is_private or ip.is_loopback or ip.is_link_local:
                            raise ValueError(f"URL points to private/internal IP address: {v}")
                    except ValueError:
                        pass  # Not an IP, it's a hostname - that's fine
            except Exception as e:
                if "localhost" in str(e) or "private" in str(e):
                    raise
                # Other parsing errors are ok
        else:
            # If it looks like an IP/CIDR, validate it
            try:
                # Extract IP from CIDR notation if present
                ip_part = v.split('/')[0] if '/' in v else v

                # Try parsing as IP network
                network = ipaddress.ip_network(ip_part, strict=False)

                # Block private/internal ranges (configurable via allow_internal_scan)
                if network.is_private:
                    raise ValueError(
                        f"Private IP range scanning blocked for security: {v}. "
                        f"If this is intentional, please contact administrator to allowlist this range."
                    )
                if network.is_link_local:
                    raise ValueError(f"Link-local addresses not allowed: {v}")
                if network.is_loopback:
                    raise ValueError(f"Loopback addresses not allowed: {v}")

            except ValueError as e:
                # If it's not an IP validation error, re-raise
                if "does not appear to be" not in str(e):
                    raise
                # Otherwise it's probably a hostname - allow it

        return v
    
    # AI-led mode options
    ai_led: bool = Field(
        default=False,
        description="Enable AI-led scanning. When true, AI decides scan strategy, scan types, and tools to use based on target analysis."
    )
    user_context: Optional[str] = Field(
        default=None,
        description="Optional context for AI to consider (e.g., 'This is a production e-commerce site' or 'Looking for SQL injection vulnerabilities'). Only used when ai_led=True."
    )
    
    # Manual configuration options (ignored when ai_led=True)
    scan_type: str = Field(
        default="service",
        description="Nmap scan type: ping, basic, service, comprehensive, stealth, udp. Ignored in URL mode and AI-led mode."
    )
    ports: Optional[str] = Field(
        default=None,
        description="Port specification (e.g., '22,80,443' or '1-1000'). Ignored in URL mode and AI-led mode."
    )
    include_web_scan: bool = Field(
        default=True,
        description="Run ZAP web vulnerability scanning on web services. AI decides when ai_led=True."
    )
    include_cve_scan: bool = Field(
        default=True,
        description="Run Nuclei CVE scanning on network services. Auto-disabled in URL mode. AI decides when ai_led=True."
    )
    include_exploit_mapping: bool = Field(
        default=True,
        description="Map findings to available exploits. AI decides when ai_led=True."
    )
    include_openvas: bool = Field(
        default=True,
        description="Run OpenVAS network vulnerability scanning. Auto-disabled in URL mode. AI decides when ai_led=True."
    )
    include_directory_enum: bool = Field(
        default=True,
        description="Run Gobuster/Dirbuster directory enumeration before ZAP."
    )
    directory_enum_engine: Literal["gobuster", "dirbuster"] = Field(
        default="gobuster",
        description="Directory enumeration engine preference."
    )
    directory_enum_wordlist: Optional[str] = Field(
        default=None,
        description="Wordlist (key or file) for directory enumeration."
    )
    directory_enum_extensions: Optional[List[str]] = Field(
        default=None,
        description="Optional extensions to probe (e.g., ['php','jsp'])."
    )
    directory_enum_threads: int = Field(
        default=25,
        description="Threads to use for Gobuster/Dirbuster."
    )
    include_sqlmap: bool = Field(
        default=True,
        description="Run SQLMap against discovered endpoints."
    )
    sqlmap_level: int = Field(
        default=2,
        ge=1,
        le=5,
        description="SQLMap level (1-5)."
    )
    sqlmap_risk: int = Field(
        default=2,
        ge=0,
        le=3,
        description="SQLMap risk (0-3)."
    )
    sqlmap_method: Literal["GET", "POST"] = Field(
        default="GET",
        description="HTTP method SQLMap should use."
    )
    sqlmap_data: Optional[str] = Field(
        default=None,
        description="POST payload template for SQLMap."
    )
    sqlmap_threads: int = Field(
        default=1,
        ge=1,
        le=10,
        description="SQLMap parallel threads."
    )
    include_wapiti: bool = Field(
        default=True,
        description="Run Wapiti web scanning alongside ZAP."
    )
    wapiti_level: int = Field(
        default=2,
        ge=1,
        le=5,
        description="Wapiti scan intensity level."
    )
    discover_js_endpoints: bool = Field(
        default=True,
        description="Extract endpoints from JavaScript content."
    )
    discover_parameters: bool = Field(
        default=True,
        description="Capture query/form parameters for additional testing."
    )
    aggressive_scan: bool = Field(
        default=True,
        description="Use aggressive/maximum scan intensity by default. Uncheck for thorough scanning."
    )
    zap_auth: Optional[ZapAuthRequest] = Field(
        default=None,
        description="Optional ZAP authentication configuration for authenticated scanning"
    )
    zap_forced_browse: bool = Field(
        default=False,
        description="Enable forced browsing with wordlists to discover hidden paths"
    )
    zap_wordlist: Optional[str] = Field(
        default=None,
        description="Wordlist key or filename for forced browsing (e.g., 'standard', 'api', 'directories_comprehensive.txt')"
    )
    openvas_credentials: Optional[OpenVASCredentialRequest] = Field(
        default=None,
        description="OpenVAS credential IDs for authenticated scanning"
    )
    openapi_spec_url: Optional[str] = Field(
        default=None,
        description="OpenAPI/Swagger spec URL to seed API scanning"
    )
    openapi_spec_content: Optional[str] = Field(
        default=None,
        description="Raw OpenAPI spec content (JSON/YAML) to seed API scanning"
    )
    openapi_base_url: Optional[str] = Field(
        default=None,
        description="Override base URL for OpenAPI spec if needed"
    )
    graphql_endpoint_url: Optional[str] = Field(
        default=None,
        description="GraphQL endpoint URL for schema import"
    )
    graphql_schema_url: Optional[str] = Field(
        default=None,
        description="GraphQL schema URL if separate from endpoint"
    )
    browser_crawl: Optional[BrowserCrawlRequest] = Field(
        default=None,
        description="Optional headless browser crawl configuration"
    )
    project_id: Optional[int] = Field(
        default=None,
        description="Project to associate scan with"
    )


class ScanProgressResponse(BaseModel):
    """Progress information for a scan."""
    phase: str
    phase_progress: int
    overall_progress: int
    message: str
    started_at: str
    hosts_discovered: int = 0
    web_targets: int = 0
    network_targets: int = 0
    findings_count: int = 0
    errors: List[str] = []


class ScanSummaryResponse(BaseModel):
    """Summary of a scan."""
    scan_id: str
    target: str
    status: str
    phase: str
    progress: int
    findings_count: int
    started_at: Optional[str]
    completed_at: Optional[str]
    duration_seconds: Optional[int]


class FindingResponse(BaseModel):
    """A vulnerability finding."""
    source: str
    severity: str
    title: str
    description: str
    host: str
    port: Optional[int]
    url: Optional[str]
    cve_id: Optional[str]
    cvss_score: Optional[float]
    exploit_available: bool
    exploit_info: Optional[dict]


class ScanResultResponse(BaseModel):
    """Full scan result."""
    scan_id: str
    target: str
    status: str
    progress: ScanProgressResponse
    hosts: List[dict]
    web_targets: List[dict]
    network_targets: List[dict]
    findings: List[FindingResponse]
    attack_narrative: str
    exploit_chains: List[dict]
    recommendations: List[str]
    exploit_commands: dict
    manual_guidance: List[str]
    started_at: Optional[str]
    completed_at: Optional[str]
    duration_seconds: Optional[int]


class ScanPlanResponse(BaseModel):
    """Agent plan details for a dynamic scan."""
    scan_id: str
    status: str
    agent_plan: List[str]
    agent_log: List[Dict[str, Any]]


class ManualGuidanceResponse(BaseModel):
    """Manual guidance for non-AI scans."""
    scan_profile: Dict[str, Any]
    manual_guidance: List[str]


class ExploitSearchRequest(BaseModel):
    """Request to search exploit database."""
    query: Optional[str] = None
    cve_id: Optional[str] = None
    product: Optional[str] = None
    version: Optional[str] = None
    platform: Optional[str] = None
    verified_only: bool = False


# ============== Endpoints ==============

@router.post("/start", response_model=ScanSummaryResponse)
async def start_dynamic_scan(
    request: StartScanRequest,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db),
):
    """
    Start a new Dynamic Security Scan.
    
    This initiates an AI-orchestrated automated pentesting workflow:
    
    **AI-Led Mode (recommended):** Set ai_led=True
    - AI analyzes the target and decides the optimal scan strategy
    - Determines if Nmap is needed (skips for URLs)
    - Chooses appropriate scan types and tools
    - May run multiple scans based on findings
    
    **Manual Mode:** Configure scan parameters yourself
    1. Nmap reconnaissance
    2. AI-driven service classification
    3. OpenVAS network vulnerability scanning
    4. ZAP web vulnerability scanning
    5. Nuclei CVE detection
    6. Exploit database mapping
    7. AI attack narrative generation
    """
    service = get_dynamic_scan_service()

    # Check scanner service availability before starting scan
    try:
        scanner_info = await service._get_scanner_info(max_age_seconds=30)
        if not scanner_info:
            raise HTTPException(
                status_code=503,
                detail="Scanner services are unavailable. Please ensure the scanner sidecar container is running. "
                       "Try: docker-compose up -d scanner"
            )

        # Check for required tools based on scan configuration
        if request.include_cve_scan:
            nuclei_available = service._scanner_tool_available(scanner_info, "nuclei")
            if nuclei_available is False:
                logger.warning("Nuclei not available, CVE scanning will be skipped")
                # Don't fail, just log warning

        if not request.target.startswith(('http://', 'https://')):
            # Network scan requires nmap
            nmap_available = service._scanner_tool_available(scanner_info, "nmap")
            if nmap_available is False:
                raise HTTPException(
                    status_code=503,
                    detail="Nmap is required for network scanning but is not available in the scanner sidecar. "
                           "Please ensure nmap is installed in the scanner container."
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to check scanner availability: {e}")
        raise HTTPException(
            status_code=503,
            detail=f"Unable to verify scanner availability: {str(e)}. "
                   f"The scanner sidecar may be down or unreachable."
        )

    try:
        result = await service.start_scan(
            target=request.target,
            scan_type=request.scan_type,
            ports=request.ports,
            include_web_scan=request.include_web_scan,
            include_cve_scan=request.include_cve_scan,
            include_exploit_mapping=request.include_exploit_mapping,
            include_openvas=request.include_openvas,
            include_directory_enum=request.include_directory_enum,
            directory_enum_engine=request.directory_enum_engine,
            directory_enum_wordlist=request.directory_enum_wordlist,
            directory_enum_extensions=request.directory_enum_extensions,
            directory_enum_threads=request.directory_enum_threads,
            include_sqlmap=request.include_sqlmap,
            sqlmap_level=request.sqlmap_level,
            sqlmap_risk=request.sqlmap_risk,
            sqlmap_method=request.sqlmap_method,
            sqlmap_data=request.sqlmap_data,
            sqlmap_threads=request.sqlmap_threads,
            include_wapiti=request.include_wapiti,
            wapiti_level=request.wapiti_level,
            aggressive_scan=request.aggressive_scan,
            zap_auth=request.zap_auth.dict(exclude_none=True) if request.zap_auth else None,
            zap_forced_browse=request.zap_forced_browse,
            zap_wordlist=request.zap_wordlist,
            openvas_credentials=request.openvas_credentials.dict(exclude_none=True) if request.openvas_credentials else None,
            openapi_spec_url=request.openapi_spec_url,
            openapi_spec_content=request.openapi_spec_content,
            openapi_base_url=request.openapi_base_url,
            graphql_endpoint_url=request.graphql_endpoint_url,
            graphql_schema_url=request.graphql_schema_url,
            browser_crawl=request.browser_crawl.dict(exclude_none=True) if request.browser_crawl else None,
            ai_led=request.ai_led,
            user_context=request.user_context,
            discover_js_endpoints=request.discover_js_endpoints,
            discover_parameters=request.discover_parameters,
            db=db,
            user_id=current_user.id,
            project_id=request.project_id,
            scan_name=request.scan_name,
        )
        
        return ScanSummaryResponse(
            scan_id=result.scan_id,
            target=result.target,
            status=result.status.value,
            phase=result.progress.phase.value,
            progress=result.progress.overall_progress,
            findings_count=len(result.findings),
            started_at=result.started_at,
            completed_at=result.completed_at,
            duration_seconds=result.duration_seconds,
        )
        
    except ScanConcurrencyError as exc:
        raise HTTPException(
            status_code=429,
            detail=f"Maximum concurrent scans reached. {str(exc)} Please wait for existing scans to complete or cancel them.",
        )
    except ValueError as e:
        # Input validation errors from target validator
        logger.warning(f"Invalid scan request: {e}")
        raise HTTPException(
            status_code=400,
            detail=f"Invalid scan parameters: {str(e)}"
        )
    except HTTPException:
        # Re-raise HTTP exceptions (like scanner unavailable)
        raise
    except Exception as e:
        logger.error(f"Failed to start scan for target '{request.target}': {e}", exc_info=True)
        error_detail = str(e)

        # Provide helpful error messages for common issues
        if "connection" in error_detail.lower():
            error_detail = f"Connection error: Unable to reach scanner services or target. {error_detail}"
        elif "timeout" in error_detail.lower():
            error_detail = f"Timeout error: Scanner services took too long to respond. {error_detail}"
        elif "permission" in error_detail.lower() or "denied" in error_detail.lower():
            error_detail = f"Permission error: {error_detail}"

        raise HTTPException(
            status_code=500,
            detail=f"Failed to start scan: {error_detail}"
        )


@router.get("/manual-guidance", response_model=ManualGuidanceResponse)
async def manual_scan_guidance(
    aggressive_scan: bool = Query(
        default=True,
        description="Whether the scan should use the aggressive/default scan profile",
    ),
    include_openvas: bool = Query(
        default=True,
        description="Include OpenVAS network scanning guidance",
    ),
    include_web_scan: bool = Query(
        default=True,
        description="Include OWASP ZAP guidance",
    ),
    include_cve_scan: bool = Query(
        default=True,
        description="Include Nuclei CVE scanning guidance",
    ),
    include_directory_enum: bool = Query(
        default=True,
        description="Include directory enumeration guidance",
    ),
    include_sqlmap: bool = Query(
        default=True,
        description="Include SQLMap guidance",
    ),
    include_wapiti: bool = Query(
        default=True,
        description="Include Wapiti guidance",
    ),
):
    """Return the manual scan profile and guidance for non-AI scans."""
    service = get_dynamic_scan_service()
    profile = service.build_scan_profile(aggressive_scan)
    guidance = service.build_manual_guidance(
        aggressive_scan=aggressive_scan,
        include_openvas=include_openvas,
        include_web_scan=include_web_scan,
        include_cve_scan=include_cve_scan,
        include_directory_enum=include_directory_enum,
        include_sqlmap=include_sqlmap,
        include_wapiti=include_wapiti,
        profile=profile,
    )
    return ManualGuidanceResponse(
        scan_profile=profile,
        manual_guidance=guidance,
    )


@router.get("/status/{scan_id}", response_model=ScanSummaryResponse)
async def get_scan_status(
    scan_id: str,
    current_user: User = Depends(get_optional_user),
    db: Session = Depends(get_db),
):
    """Get the current status of a scan."""
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    
    service = get_dynamic_scan_service()
    result = service.get_scan_status(scan_id)
    
    if not result:
        # Check if scan exists in database (may be running on different worker)
        db_scan = db.query(DynamicScan).filter(DynamicScan.scan_id == scan_id).first()
        if db_scan:
            findings_count = db.query(DynamicScanFinding).filter(
                DynamicScanFinding.scan_id == db_scan.id
            ).count()
            
            return ScanSummaryResponse(
                scan_id=db_scan.scan_id,
                target=db_scan.target,
                status=db_scan.status,
                phase=db_scan.status if db_scan.status in ["completed", "failed"] else "running",
                progress=100 if db_scan.status == "completed" else 0,
                findings_count=findings_count,
                started_at=db_scan.created_at.isoformat() if db_scan.created_at else None,
                completed_at=db_scan.completed_at.isoformat() if db_scan.completed_at else None,
                duration_seconds=db_scan.duration_seconds,
            )
        raise HTTPException(status_code=404, detail=f"Scan {scan_id} not found")
    
    return ScanSummaryResponse(
        scan_id=result.scan_id,
        target=result.target,
        status=result.status.value,
        phase=result.progress.phase.value,
        progress=result.progress.overall_progress,
        findings_count=len(result.findings),
        started_at=result.started_at,
        completed_at=result.completed_at,
        duration_seconds=result.duration_seconds,
    )


@router.get("/plan/{scan_id}", response_model=ScanPlanResponse)
async def get_scan_plan(
    scan_id: str,
    current_user: User = Depends(get_optional_user),
    db: Session = Depends(get_db),
):
    """Return the agentic plan/log for a given scan."""
    from backend.services.dynamic_scan_service import get_dynamic_scan_service

    service = get_dynamic_scan_service()
    scan = service.get_scan_status(scan_id)
    if scan:
        return ScanPlanResponse(
            scan_id=scan_id,
            status=scan.status.value,
            agent_plan=scan.agent_plan,
            agent_log=scan.agent_log,
        )

    db_scan = db.query(DynamicScan).filter(DynamicScan.scan_id == scan_id).first()
    if not db_scan:
        raise HTTPException(status_code=404, detail=f"Scan {scan_id} not found")

    results = {}
    if db_scan.results:
        try:
            results = json.loads(db_scan.results)
        except json.JSONDecodeError:
            results = {}

    agent_plan = results.get("agent_plan", [])
    agent_log = results.get("agent_log", [])
    if not agent_plan and not agent_log:
        raise HTTPException(status_code=404, detail="Agent plan not available for this scan")

    return ScanPlanResponse(
        scan_id=scan_id,
        status=db_scan.status,
        agent_plan=agent_plan,
        agent_log=agent_log,
    )


@router.get("/results/{scan_id}")
async def get_scan_results(
    scan_id: str,
    current_user: User = Depends(get_optional_user),
    db: Session = Depends(get_db),
):
    """Get full results of a completed scan."""
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    
    service = get_dynamic_scan_service()
    result = service.get_scan_status(scan_id)
    
    if not result:
        # Check if scan exists in database (may be running on different worker)
        db_scan = db.query(DynamicScan).filter(DynamicScan.scan_id == scan_id).first()
        if db_scan:
            # Try to load full results from the JSON blob first
            if db_scan.results:
                try:
                    full_results = json.loads(db_scan.results) if isinstance(db_scan.results, str) else db_scan.results
                    # Return the full stored results
                    return {
                        "scan_id": db_scan.scan_id,
                        "target": db_scan.target,
                        "status": db_scan.status,
                        "progress": full_results.get("progress", {
                            "phase": "completed" if db_scan.status == "completed" else "running",
                            "phase_progress": 100 if db_scan.status == "completed" else 0,
                            "overall_progress": 100 if db_scan.status == "completed" else 0,
                            "message": f"Scan {db_scan.status}",
                            "started_at": db_scan.created_at.isoformat() if db_scan.created_at else None,
                            "hosts_discovered": db_scan.hosts_discovered or 0,
                            "web_targets": db_scan.web_targets or 0,
                            "network_targets": db_scan.network_targets or 0,
                            "findings_count": db_scan.total_findings or 0,
                            "errors": [],
                        }),
                        "hosts": full_results.get("hosts", []),
                        "web_targets": full_results.get("web_targets", []),
                        "network_targets": full_results.get("network_targets", []),
                        "findings": full_results.get("findings", []),
                        "attack_narrative": full_results.get("attack_narrative", "") or db_scan.attack_narrative or "",
                        "executive_summary": full_results.get("executive_summary", ""),
                        "risk_summary": full_results.get("risk_summary", ""),
                        "exploit_chains": full_results.get("exploit_chains", []) or (json.loads(db_scan.exploit_chains) if db_scan.exploit_chains else []),
                        "recommendations": full_results.get("recommendations", []) or (json.loads(db_scan.recommendations) if db_scan.recommendations else []),
                        "exploit_commands": full_results.get("exploit_commands", {}),
                        "oob_summary": full_results.get("oob_summary", {}),
                        "validation_summary": full_results.get("validation_summary", {}),
                        "agent_plan": full_results.get("agent_plan", []),
                        "agent_log": full_results.get("agent_log", []),
                        "started_at": full_results.get("started_at") or (db_scan.created_at.isoformat() if db_scan.created_at else None),
                        "completed_at": full_results.get("completed_at") or (db_scan.completed_at.isoformat() if db_scan.completed_at else None),
                        "duration_seconds": full_results.get("duration_seconds") or db_scan.duration_seconds,
                        "manual_guidance": full_results.get("manual_guidance", []),
                    }
                except Exception as e:
                    logger.warning(f"Failed to parse results JSON for scan {scan_id}: {e}")
            
            # Fallback: Return minimal response with findings from separate table
            findings = db.query(DynamicScanFinding).filter(
                DynamicScanFinding.scan_id == db_scan.id
            ).all()
            
            return {
                "scan_id": db_scan.scan_id,
                "target": db_scan.target,
                "status": db_scan.status,
                "progress": {
                    "phase": db_scan.status if db_scan.status in ["completed", "failed"] else "running",
                    "phase_progress": 100 if db_scan.status == "completed" else 0,
                    "overall_progress": 100 if db_scan.status == "completed" else 0,
                    "message": "Scan in progress..." if db_scan.status == "pending" else f"Scan {db_scan.status}",
                    "started_at": db_scan.created_at.isoformat() if db_scan.created_at else None,
                    "hosts_discovered": 0,
                    "web_targets": 0,
                    "network_targets": 0,
                    "findings_count": len(findings),
                    "errors": [],
                },
                "hosts": [],
                "web_targets": [],
                "network_targets": [],
                "findings": [
                    {
                        "source": f.source or "unknown",
                        "severity": f.severity or "info",
                        "title": f.title,
                        "description": f.description or "",
                        "host": f.host or "",
                        "port": f.port,
                        "url": None,
                        "cve_id": f.cve_id,
                        "cvss_score": None,
                        "evidence": None,
                        "remediation": None,
                        "references": [],
                        "exploit_available": f.exploit_available or False,
                        "exploit_info": None,
                        "raw_data": json.loads(f.raw_data) if f.raw_data else None,
                    }
                    for f in findings
                ],
                "attack_narrative": db_scan.attack_narrative or "",
                "executive_summary": "",
                "risk_summary": "",
                "exploit_chains": json.loads(db_scan.exploit_chains) if db_scan.exploit_chains else [],
                "recommendations": json.loads(db_scan.recommendations) if db_scan.recommendations else [],
                "exploit_commands": {},
                "oob_summary": {},
                "validation_summary": {},
                "agent_plan": [],
                "agent_log": [],
                "started_at": db_scan.created_at.isoformat() if db_scan.created_at else None,
                "completed_at": db_scan.completed_at.isoformat() if db_scan.completed_at else None,
                "duration_seconds": db_scan.duration_seconds,
                "manual_guidance": [],
            }
        raise HTTPException(status_code=404, detail=f"Scan {scan_id} not found")
    
    # Convert dataclasses to dicts
    return {
        "scan_id": result.scan_id,
        "target": result.target,
        "status": result.status.value,
        "progress": asdict(result.progress),
        "hosts": [asdict(h) for h in result.hosts],
        "web_targets": [asdict(t) for t in result.web_targets],
        "network_targets": [asdict(t) for t in result.network_targets],
        "findings": [asdict(f) for f in result.findings],
        "attack_narrative": result.attack_narrative,
        "executive_summary": result.executive_summary,
        "risk_summary": result.risk_summary,
        "exploit_chains": result.exploit_chains,
        "recommendations": result.recommendations,
        "exploit_commands": result.exploit_commands,
        "oob_summary": result.oob_summary,
        "validation_summary": result.validation_summary,
        "agent_plan": result.agent_plan,
        "agent_log": result.agent_log,
        "started_at": result.started_at,
        "completed_at": result.completed_at,
        "duration_seconds": result.duration_seconds,
        "manual_guidance": result.manual_guidance,
    }


@router.post("/cancel/{scan_id}")
async def cancel_scan(
    scan_id: str,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db),
):
    """Cancel a running scan."""
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    
    service = get_dynamic_scan_service()
    
    if service.cancel_scan(scan_id, db):
        return {"message": f"Scan {scan_id} cancelled"}
    else:
        raise HTTPException(status_code=400, detail="Scan not found or already completed")


@router.post("/resume/{scan_id}")
async def resume_scan(
    scan_id: str,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db),
):
    """
    Resume a failed or cancelled scan by restarting it with the same configuration.
    
    This creates a new scan using the original scan's parameters.
    Returns the new scan ID.
    """
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    
    # Find the original scan in database
    original_scan = db.query(DynamicScan).filter(DynamicScan.scan_id == scan_id).first()
    
    if not original_scan:
        raise HTTPException(status_code=404, detail=f"Scan {scan_id} not found")
    
    # Check ownership
    if original_scan.user_id and original_scan.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized to resume this scan")
    
    # Only allow resume of failed/cancelled/completed scans
    if original_scan.status not in ["failed", "cancelled", "completed"]:
        raise HTTPException(
            status_code=400, 
            detail=f"Cannot resume scan with status '{original_scan.status}'. Only failed, cancelled, or completed scans can be resumed."
        )
    
    # Extract original configuration
    config = original_scan.scan_config or {}
    
    service = get_dynamic_scan_service()
    
    try:
        # Start a new scan with the same configuration
        new_scan_id = await service.start_scan(
            target=original_scan.target,
            scan_type=config.get("scan_type", "service"),
            ports=config.get("ports"),
            include_web_scan=config.get("include_web_scan", True),
            include_cve_scan=config.get("include_cve_scan", True),
            include_exploit_mapping=config.get("include_exploit_mapping", True),
            include_openvas=config.get("include_openvas", True),
            aggressive_scan=config.get("aggressive_scan", True),
            zap_auth=config.get("zap_auth"),
            db=db,
            user_id=current_user.id,
            scan_name=f"{original_scan.scan_name or 'Scan'} (resumed)",
            ai_led=config.get("ai_led", False),
            user_context=config.get("user_context"),
        )
        
        return {
            "message": f"Scan resumed successfully",
            "original_scan_id": scan_id,
            "new_scan_id": new_scan_id,
        }
    except ScanConcurrencyError as e:
        raise HTTPException(status_code=429, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to resume scan {scan_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to resume scan: {str(e)}")


@router.delete("/delete/{scan_id}")
async def delete_scan(
    scan_id: str,
    current_user: User = Depends(get_optional_user),
    db: Session = Depends(get_db),
):
    """Delete a saved scan from the database."""
    # Find the scan in database
    scan = db.query(DynamicScan).filter(DynamicScan.scan_id == scan_id).first()
    
    if not scan:
        raise HTTPException(status_code=404, detail=f"Scan {scan_id} not found")
    
    # Check ownership if user is authenticated
    if current_user and scan.user_id and scan.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized to delete this scan")
    
    # Delete associated findings first
    db.query(DynamicScanFinding).filter(DynamicScanFinding.scan_id == scan.id).delete()
    
    # Delete the scan
    db.delete(scan)
    db.commit()
    
    logger.info(f"Deleted scan {scan_id}")
    
    return {"status": "deleted", "scan_id": scan_id}


@router.get("/list")
async def list_scans(
    current_user: User = Depends(get_optional_user),
    status: Optional[str] = None,
    limit: int = Query(default=50, le=100),
    db: Session = Depends(get_db),
):
    """List all scans (active and from database)."""
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    
    service = get_dynamic_scan_service()
    
    # Get active scans
    active_scans = service.list_scans()
    
    # Get historical scans from database
    query = db.query(DynamicScan)
    if status:
        query = query.filter(DynamicScan.status == status)
    if current_user:
        query = query.filter(DynamicScan.user_id == current_user.id)
    
    db_scans = query.order_by(DynamicScan.created_at.desc()).limit(limit).all()
    
    historical = [
        {
            "scan_id": s.scan_id,
            "scan_name": s.scan_name,
            "target": s.target,
            "status": s.status,
            "phase": "completed" if s.status == "completed" else "unknown",
            "progress": 100 if s.status == "completed" else 0,
            "findings_count": db.query(DynamicScanFinding).filter(
                DynamicScanFinding.scan_id == s.id
            ).count(),
            "critical_count": s.critical_findings or 0,
            "high_count": s.high_findings or 0,
            "started_at": s.created_at.isoformat() if s.created_at else None,
            "completed_at": s.completed_at.isoformat() if s.completed_at else None,
            "duration_seconds": s.duration_seconds,
        }
        for s in db_scans
        if s.scan_id not in [a["scan_id"] for a in active_scans]
    ]
    
    return {
        "active_scans": active_scans,
        "historical_scans": historical,
        "total": len(active_scans) + len(historical),
    }


@router.get("/findings/{scan_id}")
async def get_scan_findings(
    scan_id: str,
    severity: Optional[str] = None,
    source: Optional[str] = None,
    exploitable_only: bool = False,
    current_user: User = Depends(get_optional_user),
):
    """Get findings from a scan with optional filtering."""
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    
    service = get_dynamic_scan_service()
    result = service.get_scan_status(scan_id)
    
    if not result:
        raise HTTPException(status_code=404, detail=f"Scan {scan_id} not found")
    
    findings = result.findings
    
    # Apply filters
    if severity:
        findings = [f for f in findings if f.severity.lower() == severity.lower()]
    
    if source:
        findings = [f for f in findings if f.source.lower() == source.lower()]
    
    if exploitable_only:
        findings = [f for f in findings if f.exploit_available]
    
    # Sort by severity
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    findings.sort(key=lambda f: severity_order.get(f.severity.lower(), 5))
    
    return {
        "scan_id": scan_id,
        "total": len(findings),
        "findings": [asdict(f) for f in findings],
    }


@router.get("/finding/{scan_id}/{finding_index}/explain")
async def explain_finding(
    scan_id: str,
    finding_index: int,
    current_user: User = Depends(get_optional_user),
):
    """Get AI explanation of a specific finding."""
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    from backend.services.dynamic_scan_agent import DynamicScanAgent
    
    service = get_dynamic_scan_service()
    result = service.get_scan_status(scan_id)
    
    if not result:
        raise HTTPException(status_code=404, detail=f"Scan {scan_id} not found")
    
    if finding_index >= len(result.findings):
        raise HTTPException(status_code=404, detail="Finding not found")
    
    finding = result.findings[finding_index]
    
    agent = DynamicScanAgent()
    explanation = await agent.explain_finding(finding)
    
    return {
        "finding": asdict(finding),
        "explanation": explanation,
    }


# ============== Exploit Database Endpoints ==============

@router.post("/exploits/search")
async def search_exploits(
    request: ExploitSearchRequest,
    current_user: User = Depends(get_optional_user),
):
    """Search the offline exploit database."""
    from backend.services.exploit_db_service import ExploitDBService
    
    service = ExploitDBService()
    
    results = []
    
    if request.cve_id:
        results = await service.search_by_cve(request.cve_id)
    elif request.product:
        results = await service.search_by_product(request.product, request.version)
    elif request.query:
        results = await service.search(
            query=request.query,
            platform=request.platform,
            verified_only=request.verified_only,
        )
    else:
        raise HTTPException(status_code=400, detail="Provide query, cve_id, or product")
    
    return {
        "total": len(results),
        "exploits": results,
    }


@router.get("/exploits/stats")
async def exploit_db_stats(
    current_user: User = Depends(get_optional_user),
):
    """Get exploit database statistics including connectivity status."""
    from backend.services.exploit_db_service import ExploitDBService
    
    service = ExploitDBService()
    return await service.get_stats()


@router.get("/cve/{cve_id}")
async def get_cve_details(
    cve_id: str,
    prefer_live: bool = True,
    current_user: User = Depends(get_optional_user),
):
    """
    Get CVE details with live API fallback.
    
    Tries NVD/CIRCL APIs first if online, falls back to local database.
    
    Args:
        cve_id: CVE identifier (e.g., CVE-2021-44228)
        prefer_live: If True, try live APIs first (default: True)
    """
    from backend.services.exploit_db_service import ExploitDBService
    
    service = ExploitDBService(prefer_live=prefer_live)
    result = await service.get_cve_details(cve_id)
    
    if not result:
        raise HTTPException(status_code=404, detail=f"CVE {cve_id} not found")
    
    # Also get exploits
    exploits = await service.search_by_cve(cve_id)
    result["exploits"] = exploits
    result["exploit_count"] = len(exploits)
    
    return result


@router.post("/exploits/import/exploitdb")
async def import_exploitdb(
    csv_path: str,
    current_user: User = Depends(get_current_active_user),
):
    """Import exploits from ExploitDB CSV file."""
    from backend.services.exploit_db_service import ExploitDBService
    
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Admin only")
    
    service = ExploitDBService()
    count = await service.import_exploitdb_csv(csv_path)
    
    return {"message": f"Imported {count} exploits"}


# ============== AI Tools Catalog Endpoint ==============

@router.get("/ai/tools")
async def get_available_ai_tools(
    current_user: User = Depends(get_optional_user),
):
    """
    Get catalog of all available scanning tools and options for AI-led scanning.
    
    This returns comprehensive information about:
    - Nmap scan types and NSE script categories
    - ZAP scan policies and spider options
    - OpenVAS scan configurations and port lists
    - Nuclei template categories
    - Directory discovery wordlists
    
    The AI uses this catalog to intelligently select the optimal tools
    and configurations based on the target being scanned.
    """
    from backend.services.dynamic_scan_agent import DynamicScanAgent
    
    return {
        "tools": DynamicScanAgent.get_available_tools(),
        "description": "Available scanning tools and configurations for AI-led security scanning",
    }


# ============== Scanner Status Endpoints ==============

@router.get("/scanner/status")
async def scanner_status(
    current_user: User = Depends(get_optional_user),
):
    """Check status of scanner sidecar."""
    import httpx
    from backend.core.config import settings
    
    scanner_url = getattr(settings, 'SCANNER_URL', 'http://localhost:9999')
    
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            response = await client.get(f"{scanner_url}/info")
            
            if response.status_code == 200:
                return {
                    "status": "connected",
                    "scanner": response.json(),
                }
    except Exception as e:
        logger.warning(f"Scanner not available: {e}")
    
    return {
        "status": "disconnected",
        "message": "Scanner sidecar not available. Using local scanning.",
    }


# ============== WebSocket for Real-time Updates ==============

@router.websocket("/ws/{scan_id}")
async def scan_websocket(
    websocket: WebSocket,
    scan_id: str,
):
    """WebSocket for real-time scan progress updates."""
    await websocket.accept()
    
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    
    service = get_dynamic_scan_service()
    
    try:
        # Check if scan exists
        result = service.get_scan_status(scan_id)
        if not result:
            await websocket.send_json({"error": "Scan not found"})
            await websocket.close()
            return
        
        # Send initial status
        await websocket.send_json({
            "type": "status",
            "data": {
                "scan_id": result.scan_id,
                "status": result.status.value,
                "progress": asdict(result.progress),
            }
        })
        
        # Poll for updates
        last_progress = -1
        last_findings = 0
        
        while True:
            await asyncio.sleep(2)
            
            result = service.get_scan_status(scan_id)
            if not result:
                break
            
            # Send update if progress changed
            current_progress = result.progress.overall_progress
            current_findings = len(result.findings)
            
            if current_progress != last_progress or current_findings != last_findings:
                await websocket.send_json({
                    "type": "progress",
                    "data": {
                        "phase": result.progress.phase.value,
                        "phase_progress": result.progress.phase_progress,
                        "overall_progress": current_progress,
                        "message": result.progress.message,
                        "findings_count": current_findings,
                        "hosts_discovered": result.progress.hosts_discovered,
                    }
                })
                last_progress = current_progress
                last_findings = current_findings
            
            # Check if complete
            if result.status.value in ["completed", "failed", "cancelled"]:
                await websocket.send_json({
                    "type": "complete",
                    "data": {
                        "status": result.status.value,
                        "findings_count": len(result.findings),
                        "duration_seconds": result.duration_seconds,
                    }
                })
                break
    
    except WebSocketDisconnect:
        logger.info(f"WebSocket disconnected for scan {scan_id}")
    except Exception as e:
        logger.error(f"WebSocket error: {e}")
        try:
            await websocket.send_json({"error": str(e)})
        except (ConnectionError, RuntimeError):
            pass  # WebSocket already closed
    finally:
        try:
            await websocket.close()
        except (ConnectionError, RuntimeError):
            pass  # WebSocket already closed


# ============== Export Endpoints ==============

@router.get("/export/{scan_id}")
async def export_scan_results(
    scan_id: str,
    format: str = Query(default="json", regex="^(json|markdown|html|pdf|docx)$"),
    current_user: User = Depends(get_optional_user),
):
    """
    Export scan results in various formats.
    
    Supported formats:
    - json: Raw JSON data
    - markdown: Formatted Markdown report
    - html: Styled HTML report
    - pdf: PDF document (requires weasyprint)
    - docx: Microsoft Word document (requires python-docx)
    """
    from backend.services.dynamic_scan_service import get_dynamic_scan_service
    
    service = get_dynamic_scan_service()
    result = service.get_scan_status(scan_id)
    
    if not result:
        raise HTTPException(status_code=404, detail=f"Scan {scan_id} not found")
    
    if format == "json":
        return {
            "scan_id": result.scan_id,
            "target": result.target,
            "status": result.status.value,
            "started_at": result.started_at,
            "completed_at": result.completed_at,
            "duration_seconds": result.duration_seconds,
            "hosts": [asdict(h) for h in result.hosts],
            "findings": [asdict(f) for f in result.findings],
            "attack_narrative": result.attack_narrative,
            "exploit_chains": result.exploit_chains,
            "recommendations": result.recommendations,
            "exploit_commands": result.exploit_commands,
        }
    
    elif format == "markdown":
        md = _generate_markdown_report(result)
        from fastapi.responses import PlainTextResponse
        return PlainTextResponse(md, media_type="text/markdown")
    
    elif format == "html":
        html = _generate_html_report(result)
        from fastapi.responses import HTMLResponse
        return HTMLResponse(html)
    
    elif format == "pdf":
        pdf_bytes = _generate_pdf_report(result)
        from fastapi.responses import Response
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=pentest-{scan_id}.pdf"}
        )
    
    elif format == "docx":
        docx_bytes = _generate_docx_report(result)
        from fastapi.responses import Response
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=pentest-{scan_id}.docx"}
        )


def _generate_markdown_report(result) -> str:
    """Generate markdown report from scan results."""
    md = f"""# Dynamic Security Scan Report

**Target:** {result.target}  
**Status:** {result.status.value}  
**Started:** {result.started_at}  
**Duration:** {result.duration_seconds or 'N/A'} seconds  

---

## Executive Summary

{result.attack_narrative or 'No attack narrative generated.'}

---

## Findings Summary

| Severity | Count |
|----------|-------|
"""
    
    # Count by severity
    severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in result.findings:
        sev = f.severity.lower()
        if sev in severity_counts:
            severity_counts[sev] += 1
    
    for sev, count in severity_counts.items():
        if count > 0:
            md += f"| {sev.capitalize()} | {count} |\n"
    
    md += f"\n**Total Findings:** {len(result.findings)}\n\n"
    
    # Hosts discovered
    md += "## Discovered Hosts\n\n"
    for host in result.hosts:
        open_ports = [p for p in host.ports if p.get("state") == "open"]
        md += f"### {host.ip}\n"
        md += f"- **Hostname:** {host.hostname or 'N/A'}\n"
        md += f"- **OS:** {host.os or 'Unknown'}\n"
        md += f"- **Open Ports:** {len(open_ports)}\n\n"
        
        if open_ports:
            md += "| Port | Service | Version |\n|------|---------|--------|\n"
            for p in open_ports[:20]:
                md += f"| {p.get('port')} | {p.get('service', 'unknown')} | {p.get('version', '')} |\n"
            md += "\n"
    
    # Critical/High findings
    critical_high = [f for f in result.findings if f.severity.lower() in ["critical", "high"]]
    if critical_high:
        md += "## Critical/High Findings\n\n"
        for f in critical_high:
            md += f"### {f.title}\n\n"
            md += f"- **Severity:** {f.severity}\n"
            md += f"- **Host:** {f.host}:{f.port or 'N/A'}\n"
            md += f"- **Source:** {f.source}\n"
            if f.cve_id:
                md += f"- **CVE:** {f.cve_id}\n"
            if f.exploit_available:
                md += f"- **Exploit Available:** ✅\n"
            md += f"\n{f.description}\n\n"
    
    # Exploit chains
    if result.exploit_chains:
        md += "## Attack Chains\n\n"
        for chain in result.exploit_chains:
            md += f"### {chain.get('name', 'Unnamed Chain')}\n\n"
            md += f"{chain.get('description', '')}\n\n"
            if chain.get('steps'):
                md += "**Steps:**\n"
                for i, step in enumerate(chain['steps'], 1):
                    md += f"{i}. {step}\n"
            md += "\n"
    
    # Exploitation commands
    if result.exploit_commands:
        md += "## Exploitation Commands\n\n"
        for tool, cmds in result.exploit_commands.items():
            md += f"### {tool.capitalize()}\n\n```bash\n"
            for cmd in cmds:
                md += f"{cmd}\n"
            md += "```\n\n"
    
    # Recommendations
    if result.recommendations:
        md += "## Recommendations\n\n"
        for i, rec in enumerate(result.recommendations, 1):
            md += f"{i}. {rec}\n"
    
    return md


def _generate_html_report(result) -> str:
    """Generate HTML report from scan results."""
    # Convert markdown to basic HTML
    md = _generate_markdown_report(result)
    
    # Basic conversion (in production, use a proper markdown library)
    style_lines = [
        "        body {",
        "            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;",
        "            max-width: 1200px;",
        "            margin: 0 auto;",
        "            padding: 20px;",
        "            background: #0d1117;",
        "            color: #c9d1d9;",
        "        }",
        "        h1, h2, h3 {",
        "            color: #58a6ff;",
        "        }",
        "        table {",
        "            border-collapse: collapse;",
        "            width: 100%;",
        "            margin: 10px 0;",
        "        }",
        "        th, td {",
        "            border: 1px solid #30363d;",
        "            padding: 8px;",
        "            text-align: left;",
        "        }",
        "        th {",
        "            background: #161b22;",
        "        }",
        "        code, pre {",
        "            background: #161b22;",
        "            padding: 2px 6px;",
        "            border-radius: 4px;",
        "        }",
        "        pre {",
        "            padding: 15px;",
        "            overflow-x: auto;",
        "        }",
        "        .critical {",
        "            color: #f85149;",
        "        }",
        "        .high {",
        "            color: #db6d28;",
        "        }",
        "        .medium {",
        "            color: #d29922;",
        "        }",
        "        .low {",
        "            color: #3fb950;",
        "        }",
    ]
    style = "\n".join(style_lines)
    html = f"""<!DOCTYPE html>
<html>
<head>
    <title>Dynamic Security Scan Report - {result.target}</title>
    <style>
{style}
    </style>
</head>
<body>
    <pre style="white-space: pre-wrap;">{md}</pre>
</body>
</html>"""
    
    return html


def _generate_pdf_report(result) -> bytes:
    """Generate PDF report from scan results using WeasyPrint."""
    try:
        from weasyprint import HTML, CSS
        
        # Generate HTML with print-friendly styling
        html_content = _generate_pdf_html(result)
        
        # Convert to PDF
        html = HTML(string=html_content)
        pdf_bytes = html.write_pdf()
        
        return pdf_bytes
        
    except ImportError:
        # Fallback: Return HTML as PDF-ready content with instructions
        logger.warning("WeasyPrint not installed. Install with: pip install weasyprint")
        html = _generate_html_report(result)
        return html.encode('utf-8')


def _generate_pdf_html(result) -> str:
    """Generate HTML optimized for PDF conversion."""
    
    # Count by severity
    severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in result.findings:
        sev = f.severity.lower()
        if sev in severity_counts:
            severity_counts[sev] += 1
    
    # Build findings HTML
    findings_html = ""
    for f in result.findings:
        severity_class = f.severity.lower()
        cve_badge = f'<span class="cve-badge">{f.cve_id}</span>' if f.cve_id else ''
        exploit_badge = '<span class="exploit-badge">⚠️ Exploit Available</span>' if f.exploit_available else ''
        
        findings_html += f'''
        <div class="finding {severity_class}">
            <div class="finding-header">
                <span class="severity-badge {severity_class}">{f.severity.upper()}</span>
                <span class="finding-title">{f.title}</span>
                {cve_badge}
                {exploit_badge}
            </div>
            <div class="finding-meta">
                <strong>Host:</strong> {f.host}:{f.port or 'N/A'} | 
                <strong>Source:</strong> {f.source}
            </div>
            <div class="finding-description">{f.description or 'No description available.'}</div>
        </div>
        '''
    
    # Build hosts HTML
    hosts_html = ""
    for host in result.hosts:
        open_ports = [p for p in host.ports if p.get("state") == "open"]
        ports_list = ", ".join([f"{p['port']}/{p.get('service', '?')}" for p in open_ports[:10]])
        hosts_html += f'''
        <tr>
            <td>{host.ip}</td>
            <td>{host.hostname or 'N/A'}</td>
            <td>{host.os or 'Unknown'}</td>
            <td>{ports_list}{'...' if len(open_ports) > 10 else ''}</td>
        </tr>
        '''
    
    # Build recommendations HTML
    recommendations_html = ""
    for i, rec in enumerate(result.recommendations or [], 1):
        recommendations_html += f"<li>{rec}</li>"
    
    # Build exploit chains HTML
    chains_html = ""
    for chain in result.exploit_chains or []:
        steps_html = "".join([f"<li>{step}</li>" for step in chain.get('steps', [])])
        chains_html += f'''
        <div class="attack-chain">
            <h4>{chain.get('name', 'Attack Chain')}</h4>
            <p>{chain.get('description', '')}</p>
            <ol>{steps_html}</ol>
            <p><strong>Impact:</strong> {chain.get('impact', 'Unknown')} | 
               <strong>Likelihood:</strong> {chain.get('likelihood', 'Unknown')}</p>
        </div>
        '''
    
    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Penetration Test Report - {result.target}</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        body {{
            font-family: 'Helvetica Neue', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            color: #1a365d;
            border-bottom: 3px solid #2563eb;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #1e40af;
            border-bottom: 1px solid #93c5fd;
            padding-bottom: 5px;
            margin-top: 30px;
        }}
        h3 {{ color: #3b82f6; }}
        
        .header-box {{
            background: linear-gradient(135deg, #1e3a8a, #3b82f6);
            color: white;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 20px;
        }}
        .header-box h1 {{
            color: white;
            border: none;
            margin: 0;
        }}
        
        .summary-grid {{
            display: grid;
            grid-template-columns: repeat(4, 1fr);
            gap: 10px;
            margin: 20px 0;
        }}
        .summary-card {{
            text-align: center;
            padding: 15px;
            border-radius: 8px;
        }}
        .summary-card.critical {{ background: #fef2f2; border: 2px solid #dc2626; }}
        .summary-card.high {{ background: #fff7ed; border: 2px solid #ea580c; }}
        .summary-card.medium {{ background: #fefce8; border: 2px solid #ca8a04; }}
        .summary-card.low {{ background: #f0fdf4; border: 2px solid #16a34a; }}
        .summary-card .count {{ font-size: 24pt; font-weight: bold; }}
        
        .finding {{
            border-left: 4px solid #ccc;
            padding: 10px 15px;
            margin: 10px 0;
            background: #f9fafb;
            page-break-inside: avoid;
        }}
        .finding.critical {{ border-color: #dc2626; background: #fef2f2; }}
        .finding.high {{ border-color: #ea580c; background: #fff7ed; }}
        .finding.medium {{ border-color: #ca8a04; background: #fefce8; }}
        .finding.low {{ border-color: #16a34a; background: #f0fdf4; }}
        
        .severity-badge {{
            display: inline-block;
            padding: 2px 8px;
            border-radius: 4px;
            font-size: 9pt;
            font-weight: bold;
            color: white;
            margin-right: 10px;
        }}
        .severity-badge.critical {{ background: #dc2626; }}
        .severity-badge.high {{ background: #ea580c; }}
        .severity-badge.medium {{ background: #ca8a04; }}
        .severity-badge.low {{ background: #16a34a; }}
        .severity-badge.info {{ background: #6b7280; }}
        
        .cve-badge {{
            background: #7c3aed;
            color: white;
            padding: 2px 6px;
            border-radius: 4px;
            font-size: 9pt;
            margin-left: 10px;
        }}
        .exploit-badge {{
            background: #dc2626;
            color: white;
            padding: 2px 6px;
            border-radius: 4px;
            font-size: 9pt;
            margin-left: 10px;
        }}
        
        .finding-title {{ font-weight: bold; }}
        .finding-meta {{ font-size: 10pt; color: #666; margin: 5px 0; }}
        .finding-description {{ margin-top: 10px; }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        th, td {{
            border: 1px solid #e5e7eb;
            padding: 8px 12px;
            text-align: left;
        }}
        th {{
            background: #f3f4f6;
            font-weight: bold;
        }}
        
        .attack-chain {{
            background: #fef3c7;
            border: 1px solid #f59e0b;
            border-radius: 8px;
            padding: 15px;
            margin: 10px 0;
        }}
        
        .narrative {{
            background: #eff6ff;
            border: 1px solid #3b82f6;
            border-radius: 8px;
            padding: 20px;
            margin: 20px 0;
        }}
        
        pre {{
            background: #1f2937;
            color: #10b981;
            padding: 15px;
            border-radius: 8px;
            overflow-x: auto;
            font-family: 'Courier New', monospace;
            font-size: 9pt;
        }}
        
        .footer {{
            text-align: center;
            font-size: 9pt;
            color: #666;
            margin-top: 30px;
            padding-top: 10px;
            border-top: 1px solid #e5e7eb;
        }}
    </style>
</head>
<body>
    <div class="header-box">
        <h1>🔒 Penetration Test Report</h1>
        <p><strong>Target:</strong> {result.target}<br>
        <strong>Scan Date:</strong> {result.started_at}<br>
        <strong>Duration:</strong> {result.duration_seconds or 'N/A'} seconds<br>
        <strong>Status:</strong> {result.status.value}</p>
    </div>
    
    <h2>📊 Executive Summary</h2>
    
    <div class="summary-grid">
        <div class="summary-card critical">
            <div class="count">{severity_counts['critical']}</div>
            <div>Critical</div>
        </div>
        <div class="summary-card high">
            <div class="count">{severity_counts['high']}</div>
            <div>High</div>
        </div>
        <div class="summary-card medium">
            <div class="count">{severity_counts['medium']}</div>
            <div>Medium</div>
        </div>
        <div class="summary-card low">
            <div class="count">{severity_counts['low']}</div>
            <div>Low</div>
        </div>
    </div>
    
    <div class="narrative">
        <h3>🎯 Attack Narrative</h3>
        <p>{result.attack_narrative or 'No attack narrative generated.'}</p>
    </div>
    
    <h2>🖥️ Discovered Hosts ({len(result.hosts)})</h2>
    <table>
        <tr><th>IP Address</th><th>Hostname</th><th>OS</th><th>Open Ports</th></tr>
        {hosts_html}
    </table>
    
    <h2>🚨 Vulnerability Findings ({len(result.findings)})</h2>
    {findings_html or '<p>No vulnerabilities detected.</p>'}
    
    {'<h2>⚔️ Attack Chains</h2>' + chains_html if chains_html else ''}
    
    <h2>📋 Recommendations</h2>
    <ol>{recommendations_html or '<li>No specific recommendations available.</li>'}</ol>
    
    <div class="footer">
        <p>Generated by VRAgent Dynamic Security Scanner | Confidential</p>
    </div>
</body>
</html>
'''


def _generate_docx_report(result) -> bytes:
    """Generate Word document from scan results using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import io
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Penetration Test Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Target info
        doc.add_paragraph(f'Target: {result.target}')
        doc.add_paragraph(f'Scan Date: {result.started_at}')
        doc.add_paragraph(f'Duration: {result.duration_seconds or "N/A"} seconds')
        doc.add_paragraph(f'Status: {result.status.value}')
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        
        # Severity counts
        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
        for f in result.findings:
            sev = f.severity.lower()
            if sev in severity_counts:
                severity_counts[sev] += 1
        
        summary_table = doc.add_table(rows=2, cols=4)
        summary_table.style = 'Table Grid'
        
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Critical'
        hdr_cells[1].text = 'High'
        hdr_cells[2].text = 'Medium'
        hdr_cells[3].text = 'Low'
        
        data_cells = summary_table.rows[1].cells
        data_cells[0].text = str(severity_counts['critical'])
        data_cells[1].text = str(severity_counts['high'])
        data_cells[2].text = str(severity_counts['medium'])
        data_cells[3].text = str(severity_counts['low'])
        
        doc.add_paragraph()
        
        # Attack Narrative
        doc.add_heading('Attack Narrative', level=1)
        doc.add_paragraph(result.attack_narrative or 'No attack narrative generated.')
        
        # Discovered Hosts
        doc.add_heading('Discovered Hosts', level=1)
        if result.hosts:
            hosts_table = doc.add_table(rows=1, cols=4)
            hosts_table.style = 'Table Grid'
            hdr = hosts_table.rows[0].cells
            hdr[0].text = 'IP Address'
            hdr[1].text = 'Hostname'
            hdr[2].text = 'OS'
            hdr[3].text = 'Open Ports'
            
            for host in result.hosts:
                open_ports = [p for p in host.ports if p.get("state") == "open"]
                ports_str = ", ".join([f"{p['port']}" for p in open_ports[:5]])
                if len(open_ports) > 5:
                    ports_str += "..."
                
                row = hosts_table.add_row().cells
                row[0].text = host.ip
                row[1].text = host.hostname or 'N/A'
                row[2].text = host.os or 'Unknown'
                row[3].text = ports_str
        else:
            doc.add_paragraph('No hosts discovered.')
        
        # Vulnerability Findings
        doc.add_heading('Vulnerability Findings', level=1)
        
        for f in result.findings:
            p = doc.add_paragraph()
            run = p.add_run(f'[{f.severity.upper()}] ')
            if f.severity.lower() == 'critical':
                run.font.color.rgb = RGBColor(220, 38, 38)
            elif f.severity.lower() == 'high':
                run.font.color.rgb = RGBColor(234, 88, 12)
            elif f.severity.lower() == 'medium':
                run.font.color.rgb = RGBColor(202, 138, 4)
            else:
                run.font.color.rgb = RGBColor(22, 163, 74)
            run.bold = True
            
            p.add_run(f.title).bold = True
            
            details = doc.add_paragraph()
            details.add_run(f'Host: ').bold = True
            details.add_run(f'{f.host}:{f.port or "N/A"} | ')
            details.add_run(f'Source: ').bold = True
            details.add_run(f'{f.source}')
            if f.cve_id:
                details.add_run(f' | ')
                details.add_run(f'CVE: ').bold = True
                details.add_run(f.cve_id)
            if f.exploit_available:
                details.add_run(' | ⚠️ Exploit Available')
            
            if f.description:
                doc.add_paragraph(f.description)
            
            doc.add_paragraph()
        
        if not result.findings:
            doc.add_paragraph('No vulnerabilities detected.')
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        if result.recommendations:
            for i, rec in enumerate(result.recommendations, 1):
                doc.add_paragraph(f'{i}. {rec}')
        else:
            doc.add_paragraph('No specific recommendations available.')
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except ImportError:
        logger.warning("python-docx not installed. Install with: pip install python-docx")
        # Return markdown as fallback
        md = _generate_markdown_report(result)
        return md.encode('utf-8')


# ============== AI Chat Endpoint ==============

class ChatMessage(BaseModel):
    role: str
    content: str


class DynamicScanChatRequest(BaseModel):
    """Request for AI chat about a dynamic scan."""
    message: str = Field(..., description="User's question about the scan")
    scan_context: str = Field(..., description="Full context of the scan results")
    scan_id: Optional[str] = Field(default=None, description="ID of the scan being discussed")
    conversation_history: Optional[List[ChatMessage]] = Field(default=[], description="Previous messages in conversation")


class DynamicScanChatResponse(BaseModel):
    """Response from AI chat."""
    response: str


@router.post("/chat", response_model=DynamicScanChatResponse)
async def dynamic_scan_chat(
    request: DynamicScanChatRequest,
    current_user: User = Depends(get_optional_user),
):
    # AI-powered chat for analyzing dynamic scan results.
    # The AI has full context of the scan and can answer questions about:
    # - Vulnerability findings and their severity
    # - Attack vectors and exploitation paths
    # - Remediation recommendations
    # - Technical details about discovered hosts and services
    try:
        import os
        from google import genai
        from backend.core.config import settings
        
        api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
        if not api_key:
            raise HTTPException(
                status_code=500,
                detail="AI API key not configured"
            )
        
        client = genai.Client(api_key=api_key)
        
        # Build the system prompt with prioritized chunks and agentic directives
        scan_context_chunk = request.scan_context[:15000] if request.scan_context else "No scan context has been recorded yet."
        system_prompt_parts = [
            "You are an expert penetration testing analyst AI assistant embedded in VRAgent, a professional security scanning tool.",
            (
                f"You have FULL ACCESS to the following dynamic security scan results:\n---\n{scan_context_chunk}\n---"
            ),
            (
                "Your expertise includes:\n"
                "1. Vulnerability analysis and risk assessment\n"
                "2. Exploitation techniques and attack chain construction\n"
                "3. Network reconnaissance and service enumeration\n"
                "4. Web application security testing\n"
                "5. CVE analysis and exploit research\n"
                "6. Remediation strategies and security hardening"
            ),
            (
                "Your role is to:\n"
                "- Answer questions about the scan findings in detail\n"
                "- Explain attack vectors and exploitation paths based on discovered vulnerabilities\n"
                "- Prioritize remediation efforts based on risk\n"
                "- Provide technical guidance for manual verification\n"
                "- Suggest additional testing approaches\n"
                "- Explain CVEs and their real-world impact\n"
                "- Help interpret the AI-generated attack narrative\n"
                "- Provide exploitation commands and techniques when asked (educational context)"
            ),
            (
                "Final directives:\n"
                "- Keep responses grounded in the provided scan context; clearly state when information is missing.\n"
                "- Maintain agentic multi-step reasoning when escalating findings or planning remediations.\n"
                "- Remain focused, robust, and traceable while offering actionable, severity-ranked guidance."
            ),
            (
                "Response format:\n"
                "NEXT_ACTION: <scan_summary|finalize>\n"
                "ANALYSIS: <thoughts for this step>\n"
                "FINAL_RESPONSE: <final answer once NEXT_ACTION=finalize>\n"
                "When requesting a tool, pause with NEXT_ACTION=scan_summary and the agent will feed you the tool result."
            ),
        ]
        system_prompt = "\n\n".join(part for part in system_prompt_parts if part)

        # Build conversation messages
        messages = []
        for msg in (request.conversation_history or [])[-10:]:
            messages.append({
                "role": "user" if msg.role == "user" else "model",
                "parts": [{"text": msg.content}]
            })
        messages.append({
            "role": "user",
            "parts": [{"text": request.message}]
        })

        last_assistant_text = ""
        max_iterations = 3
        from google.genai import types
        for iteration in range(max_iterations):
            response = client.models.generate_content(
                model=settings.gemini_model_id,
                contents=messages,
                config=types.GenerateContentConfig(
                    system_instruction=system_prompt,
                    thinking_config=types.ThinkingConfig(thinking_level="high"),
                    max_output_tokens=2048,
                )
            )
            assistant_text = (response.text or "").strip()
            if not assistant_text:
                break
            last_assistant_text = assistant_text
            messages.append({
                "role": "model",
                "parts": [{"text": assistant_text}]
            })

            action = _extract_next_action(assistant_text)
            if action == "scan_summary":
                tool_result = _summarize_scan_context(request.scan_context)
                messages.append({
                    "role": "user",
                    "parts": [{
                        "text": (
                            f"Tool result (scan_summary):\n{tool_result}\n"
                            "Please continue your analysis with the extra context."
                        )
                    }]
                })
                continue

            if action == "finalize":
                final_text = _extract_final_response(assistant_text)
                return DynamicScanChatResponse(response=final_text or assistant_text)

            break

        final_text = _extract_final_response(last_assistant_text) if last_assistant_text else "I apologize, but I couldn't generate a response. Please try again."
        return DynamicScanChatResponse(response=final_text)
            
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="AI module not available. Please check server configuration."
        )
    except Exception as e:
        logger.error(f"Dynamic scan chat error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate response: {str(e)}"
        )
