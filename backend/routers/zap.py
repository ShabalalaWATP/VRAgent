"""
OWASP ZAP Router

API endpoints for OWASP ZAP integration, providing:
- Full scan orchestration (Spider + Active + Passive)
- Quick scans for fast results
- Spider-only for endpoint discovery
- Alert retrieval and management
- Report generation
- Session management
- Integration with Agentic Fuzzer
"""

from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field
from typing import Any, Dict, List, Optional
from sqlalchemy.orm import Session as DBSession
from datetime import datetime as dt
import json
import logging

from backend.core.auth import get_current_active_user, get_optional_user
from backend.core.database import get_db
from backend.models.models import User, ZAPScan

from backend.services.zap_service import (
    get_zap_scanner,
    zap_health_check,
    zap_full_scan,
    zap_get_findings,
    merge_zap_findings_with_fuzzer,
    ZAPScanConfig,
    ZAPScanType,
    ZAPError,
    ZAPAuthConfig,
    ZAPAuthMethod,
    ZAPScanPolicy,
)

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/zap", tags=["OWASP ZAP"])


# =============================================================================
# REQUEST MODELS
# =============================================================================

class ZAPScanRequest(BaseModel):
    """Request to start a ZAP scan."""
    target_url: str = Field(..., description="Target URL to scan")
    scan_type: str = Field(default="full_scan", description="Scan type: spider, ajax_spider, active_scan, passive_scan, full_scan")
    project_id: Optional[int] = Field(default=None, description="Project to associate scan with")
    
    # Spider options
    max_depth: int = Field(default=5, ge=1, le=20, description="Maximum spider depth")
    max_children: int = Field(default=0, ge=0, description="Max children per node (0 = unlimited)")
    subtree_only: bool = Field(default=True, description="Only scan URLs under target path")
    
    # AJAX Spider options
    enable_ajax_spider: bool = Field(default=True, description="Enable AJAX spider for JS-heavy sites")
    ajax_spider_max_duration: int = Field(default=30, ge=5, le=120, description="AJAX spider max duration (minutes)")
    browser_id: str = Field(default="firefox-headless", description="Browser for AJAX spider")
    
    # Active scan options
    scan_policy: Optional[str] = Field(default=None, description="Scan policy name")
    recurse: bool = Field(default=True, description="Recurse into subdirectories")
    in_scope_only: bool = Field(default=True, description="Only scan in-scope URLs")
    delay_in_ms: int = Field(default=0, ge=0, description="Delay between requests (ms)")
    max_scan_duration_mins: int = Field(default=0, ge=0, description="Max scan duration (0 = unlimited)")
    
    # Scope
    include_regexes: List[str] = Field(default_factory=list, description="Regex patterns to include")
    exclude_regexes: List[str] = Field(default_factory=list, description="Regex patterns to exclude")


class QuickScanRequest(BaseModel):
    """Request for quick ZAP scan."""
    url: str = Field(..., description="Target URL")
    max_duration_mins: int = Field(default=5, ge=1, le=30, description="Max scan duration (minutes)")
    project_id: Optional[int] = Field(default=None, description="Project to associate scan with")


class SpiderScanRequest(BaseModel):
    """Request for spider-only scan."""
    url: str = Field(..., description="Target URL to spider")
    max_depth: int = Field(default=5, ge=1, le=20, description="Maximum crawl depth")
    include_ajax: bool = Field(default=True, description="Include AJAX spider")
    project_id: Optional[int] = Field(default=None, description="Project to associate scan with")


class SaveScanRequest(BaseModel):
    """Request to save a ZAP scan to database."""
    session_id: str = Field(..., description="ZAP session ID to save")
    title: Optional[str] = Field(default=None, description="Custom title")
    project_id: Optional[int] = Field(default=None, description="Associated project")


class AuthenticationRequest(BaseModel):
    """Request to configure authentication for scans."""
    target_url: str = Field(..., description="Target URL to authenticate against")
    method: str = Field(..., description="Auth method: formBasedAuthentication, httpAuthentication, jsonBasedAuthentication, scriptBasedAuthentication")
    
    # Form-based auth
    login_url: Optional[str] = Field(default=None, description="Login form URL")
    login_request_data: Optional[str] = Field(default=None, description="Login POST data with {%username%} and {%password%} placeholders")
    
    # HTTP Basic
    hostname: Optional[str] = Field(default=None, description="Hostname for HTTP Basic auth")
    realm: Optional[str] = Field(default=None, description="Realm for HTTP Basic auth")
    
    # JSON/API auth
    json_template: Optional[str] = Field(default=None, description="JSON template with placeholders")
    
    # Script-based
    script_name: Optional[str] = Field(default=None, description="Authentication script name")
    script_params: Optional[Dict[str, str]] = Field(default=None, description="Script parameters")
    
    # Credentials
    username: Optional[str] = Field(default=None, description="Username for authentication")
    password: Optional[str] = Field(default=None, description="Password for authentication")
    
    # Session indicators
    logged_in_indicator: Optional[str] = Field(default=None, description="Regex pattern indicating logged-in state")
    logged_out_indicator: Optional[str] = Field(default=None, description="Regex pattern indicating logged-out state")
    
    # Context
    context_name: Optional[str] = Field(default=None, description="Custom context name")


class OAuthRequest(BaseModel):
    """Request to configure OAuth2 authentication."""
    target_url: str = Field(..., description="Target URL to scan")
    token_endpoint: str = Field(..., description="OAuth2 token endpoint")
    client_id: str = Field(..., description="OAuth2 client ID")
    client_secret: str = Field(..., description="OAuth2 client secret")
    scope: str = Field(default="openid profile", description="OAuth2 scopes")
    grant_type: str = Field(default="client_credentials", description="OAuth2 grant type")
    context_name: Optional[str] = Field(default=None, description="Custom context name")


class ScanPolicyRequest(BaseModel):
    """Request to create a scan policy."""
    name: str = Field(..., description="Policy name")
    description: Optional[str] = Field(default=None, description="Policy description")
    attack_strength: str = Field(default="MEDIUM", description="Attack strength: LOW, MEDIUM, HIGH, INSANE")
    alert_threshold: str = Field(default="MEDIUM", description="Alert threshold: OFF, LOW, MEDIUM, HIGH")
    enabled_scanners: Optional[List[int]] = Field(default=None, description="Scanner IDs to enable")
    disabled_scanners: Optional[List[int]] = Field(default=None, description="Scanner IDs to disable")


class QuickPolicyRequest(BaseModel):
    """Request to create a quick scan policy by category."""
    name: str = Field(..., description="Policy name")
    strength: str = Field(default="MEDIUM", description="Attack strength")
    categories: List[str] = Field(..., description="Scanner categories: sql_injection, xss, path_traversal, command_injection, etc.")


# =============================================================================
# HEALTH & STATUS
# =============================================================================

@router.get("/health")
async def zap_health():
    """Check if ZAP is available and get status."""
    try:
        health = await zap_health_check()
        return health
    except Exception as e:
        logger.error(f"ZAP health check failed: {e}")
        return {
            "available": False,
            "error": str(e),
            "message": "ZAP service is not available. Make sure the ZAP container is running."
        }


@router.get("/version")
async def get_zap_version():
    """Get ZAP version information."""
    scanner = get_zap_scanner()
    available, version = await scanner.is_available()
    
    if not available:
        raise HTTPException(status_code=503, detail=f"ZAP not available: {version}")
    
    return {"version": version, "status": "running"}


# =============================================================================
# SCANNING ENDPOINTS
# =============================================================================

@router.post("/scan")
async def start_zap_scan(
    request: ZAPScanRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Start a full ZAP scan with streaming progress.
    
    Scan types:
    - spider: Crawl to discover endpoints
    - ajax_spider: Crawl JavaScript-heavy applications
    - active_scan: Actively probe for vulnerabilities
    - passive_scan: Analyze existing traffic (must have proxied traffic)
    - full_scan: Spider + Active + Passive (recommended)
    
    Returns Server-Sent Events for real-time progress.
    """
    # Validate scan type
    valid_types = [e.value for e in ZAPScanType]
    if request.scan_type not in valid_types:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid scan type. Must be one of: {valid_types}"
        )
    
    config = ZAPScanConfig(
        target_url=request.target_url,
        scan_type=ZAPScanType(request.scan_type),
        max_depth=request.max_depth,
        max_children=request.max_children,
        subtree_only=request.subtree_only,
        enable_ajax_spider=request.enable_ajax_spider,
        ajax_spider_max_duration=request.ajax_spider_max_duration,
        browser_id=request.browser_id,
        scan_policy=request.scan_policy,
        recurse=request.recurse,
        in_scope_only=request.in_scope_only,
        delay_in_ms=request.delay_in_ms,
        max_scan_duration_mins=request.max_scan_duration_mins,
        include_regexes=request.include_regexes,
        exclude_regexes=request.exclude_regexes,
    )
    
    scanner = get_zap_scanner()
    
    async def event_generator():
        try:
            async for event in scanner.full_scan(
                config,
                user_id=current_user.id,
                project_id=request.project_id,
                persist_to_db=True
            ):
                yield f"data: {json.dumps(event)}\n\n"
        except ZAPError as e:
            logger.error(f"ZAP scan error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
        except Exception as e:
            logger.exception(f"Unexpected error in ZAP scan: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@router.post("/quick-scan")
async def quick_zap_scan(
    request: QuickScanRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Run a quick ZAP scan with time limit.
    
    Good for fast security checks without full crawling.
    """
    scanner = get_zap_scanner()
    
    async def event_generator():
        try:
            async for event in scanner.quick_scan(
                url=request.url,
                max_duration_mins=request.max_duration_mins,
                user_id=current_user.id,
                project_id=request.project_id,
                persist_to_db=True
            ):
                yield f"data: {json.dumps(event)}\n\n"
        except Exception as e:
            logger.exception(f"Quick scan error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@router.post("/spider")
async def spider_scan(
    request: SpiderScanRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Run spider-only scan for endpoint discovery.
    
    Use this to discover all endpoints before targeted fuzzing.
    """
    scanner = get_zap_scanner()
    
    async def event_generator():
        try:
            async for event in scanner.spider_only(
                url=request.url,
                max_depth=request.max_depth,
                include_ajax=request.include_ajax,
                user_id=current_user.id,
                project_id=request.project_id,
                persist_to_db=True
            ):
                yield f"data: {json.dumps(event)}\n\n"
        except Exception as e:
            logger.exception(f"Spider scan error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@router.post("/scan/{session_id}/stop")
async def stop_zap_scan(
    session_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Stop a running ZAP scan."""
    scanner = get_zap_scanner()
    success = await scanner.stop_scan(session_id)
    
    if not success:
        raise HTTPException(status_code=404, detail="Scan session not found or already stopped")
    
    return {"message": "Scan stopped", "session_id": session_id}


# =============================================================================
# SESSIONS & RESULTS
# =============================================================================

@router.get("/sessions")
async def list_zap_sessions(current_user: User = Depends(get_current_active_user)):
    """List all ZAP scan sessions."""
    scanner = get_zap_scanner()
    sessions = await scanner.get_all_sessions()
    return {"sessions": sessions}


@router.get("/sessions/{session_id}")
async def get_zap_session(
    session_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get details of a specific ZAP session."""
    scanner = get_zap_scanner()
    session = await scanner.get_session(session_id)
    
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return session.to_dict()


@router.get("/alerts")
async def get_zap_alerts(
    url: Optional[str] = Query(None, description="Filter by base URL"),
    min_risk: int = Query(0, ge=0, le=3, description="Minimum risk level (0=info, 1=low, 2=medium, 3=high)"),
    current_user: User = Depends(get_current_active_user)
):
    """Get all ZAP alerts/findings."""
    scanner = get_zap_scanner()
    alerts = await scanner.get_alerts(url=url, min_risk=min_risk)
    
    return {
        "count": len(alerts),
        "alerts": [
            {
                "id": a.id,
                "name": a.alert_name,
                "risk": a.risk,
                "risk_code": a.risk_code,
                "confidence": a.confidence,
                "url": a.url,
                "method": a.method,
                "parameter": a.parameter,
                "attack": a.attack,
                "evidence": a.evidence,
                "description": a.description,
                "solution": a.solution,
                "cwe_id": a.cwe_id,
                "wasc_id": a.wasc_id,
            }
            for a in alerts
        ]
    }


@router.get("/findings")
async def get_zap_findings_formatted(
    url: Optional[str] = Query(None, description="Filter by base URL"),
    current_user: User = Depends(get_current_active_user)
):
    """
    Get ZAP findings in Agentic Fuzzer compatible format.
    
    This allows easy integration with existing fuzzer findings.
    """
    findings = await zap_get_findings(url=url)
    return {
        "count": len(findings),
        "source": "owasp_zap",
        "findings": findings
    }


# =============================================================================
# REPORTS
# =============================================================================

@router.get("/report")
async def generate_zap_report(
    format: str = Query("json", description="Report format: json, html, xml, markdown"),
    current_user: User = Depends(get_current_active_user)
):
    """Generate a ZAP scan report."""
    valid_formats = ["json", "html", "xml", "markdown"]
    if format not in valid_formats:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid format. Must be one of: {valid_formats}"
        )
    
    scanner = get_zap_scanner()
    
    try:
        report = await scanner.generate_report(format=format)
        
        if format == "json":
            return report
        elif format == "html":
            return Response(
                content=report,
                media_type="text/html",
                headers={"Content-Disposition": f"attachment; filename=zap_report.html"}
            )
        elif format == "xml":
            return Response(
                content=report,
                media_type="application/xml",
                headers={"Content-Disposition": f"attachment; filename=zap_report.xml"}
            )
        else:  # markdown
            return Response(
                content=report,
                media_type="text/markdown",
                headers={"Content-Disposition": f"attachment; filename=zap_report.md"}
            )
            
    except ZAPError as e:
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# SESSION MANAGEMENT
# =============================================================================

@router.post("/session/new")
async def new_zap_session(current_user: User = Depends(get_current_active_user)):
    """Create a new ZAP session (clears existing data)."""
    scanner = get_zap_scanner()
    success = await scanner.clear_session()
    
    if not success:
        raise HTTPException(status_code=500, detail="Failed to create new session")
    
    return {"message": "New ZAP session created", "status": "clean"}


@router.delete("/alerts")
async def clear_zap_alerts(current_user: User = Depends(get_current_active_user)):
    """Delete all ZAP alerts."""
    scanner = get_zap_scanner()
    success = await scanner.clear_session()
    
    if not success:
        raise HTTPException(status_code=500, detail="Failed to clear alerts")
    
    return {"message": "All alerts cleared"}


# =============================================================================
# DATABASE PERSISTENCE
# =============================================================================

@router.post("/scans/save")
async def save_zap_scan(
    request: SaveScanRequest,
    db: DBSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Save a ZAP scan session to the database."""
    scanner = get_zap_scanner()
    session = await scanner.get_session(request.session_id)
    
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Check if already saved
    existing = db.query(ZAPScan).filter(ZAPScan.session_id == request.session_id).first()
    if existing:
        # Update the title if a new one is provided
        if request.title and request.title != existing.title:
            existing.title = request.title
            db.commit()
            logger.info(f"Updated ZAP scan {existing.id} title to: {request.title}")
        return {
            "success": True,
            "message": "Scan already saved" if not request.title else "Scan title updated",
            "scan_id": existing.id,
            "session_id": existing.session_id,
            "alerts_count": (existing.alerts_high or 0) + (existing.alerts_medium or 0) + (existing.alerts_low or 0) + (existing.alerts_info or 0)
        }
    
    # Count alerts by severity
    alerts_by_risk = session._count_alerts_by_risk()
    
    # Create scan record
    scan = ZAPScan(
        session_id=session.id,
        user_id=current_user.id,
        project_id=request.project_id,
        title=request.title or f"ZAP Scan: {session.target_url}",
        target_url=session.target_url,
        scan_type=session.scan_type.value,
        status=session.status.value,
        started_at=dt.fromisoformat(session.started_at) if session.started_at else dt.utcnow(),
        completed_at=dt.fromisoformat(session.completed_at) if session.completed_at else None,
        urls_found=len(session.urls_found),
        alerts_high=alerts_by_risk.get("high", 0),
        alerts_medium=alerts_by_risk.get("medium", 0),
        alerts_low=alerts_by_risk.get("low", 0),
        alerts_info=alerts_by_risk.get("info", 0),
        alerts_data=[a.to_finding_dict() for a in session.alerts],
        urls_data=session.urls_found[:500],  # Limit stored URLs
        stats=session.stats,
    )
    
    db.add(scan)
    db.commit()
    db.refresh(scan)
    
    logger.info(f"Saved ZAP scan {scan.id} for session {request.session_id}")
    
    return {
        "success": True,
        "message": "Scan saved successfully",
        "scan_id": scan.id,
        "session_id": scan.session_id,
        "alerts_count": sum(alerts_by_risk.values())
    }


@router.get("/scans")
async def list_saved_zap_scans(
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
    db: DBSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """List saved ZAP scans for the current user."""
    query = db.query(ZAPScan).filter(
        ZAPScan.user_id == current_user.id
    ).order_by(ZAPScan.created_at.desc())
    
    total = query.count()
    scans = query.offset(skip).limit(limit).all()
    
    return {
        "total": total,
        "skip": skip,
        "limit": limit,
        "scans": [
            {
                "id": s.id,
                "session_id": s.session_id,
                "title": s.title,
                "target_url": s.target_url,
                "scan_type": s.scan_type,
                "status": s.status,
                "completed_at": s.completed_at.isoformat() if s.completed_at else None,
                "urls_found": s.urls_found,
                "alerts": {
                    "high": s.alerts_high,
                    "medium": s.alerts_medium,
                    "low": s.alerts_low,
                    "info": s.alerts_info,
                    "total": (s.alerts_high or 0) + (s.alerts_medium or 0) + (s.alerts_low or 0) + (s.alerts_info or 0)
                },
                "created_at": s.created_at.isoformat() if s.created_at else None,
            }
            for s in scans
        ]
    }


@router.get("/scans/{scan_id}")
async def get_saved_zap_scan(
    scan_id: int,
    db: DBSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Get a saved ZAP scan by ID."""
    scan = db.query(ZAPScan).filter(
        ZAPScan.id == scan_id,
        ZAPScan.user_id == current_user.id
    ).first()
    
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")
    
    return {
        "id": scan.id,
        "session_id": scan.session_id,
        "title": scan.title,
        "target_url": scan.target_url,
        "scan_type": scan.scan_type,
        "status": scan.status,
        "started_at": scan.started_at.isoformat() if scan.started_at else None,
        "completed_at": scan.completed_at.isoformat() if scan.completed_at else None,
        "urls_found": scan.urls_found,
        "alerts_summary": {
            "high": scan.alerts_high,
            "medium": scan.alerts_medium,
            "low": scan.alerts_low,
            "info": scan.alerts_info,
            "total": (scan.alerts_high or 0) + (scan.alerts_medium or 0) + (scan.alerts_low or 0) + (scan.alerts_info or 0)
        },
        "alerts": scan.alerts_data,
        "urls": scan.urls_data,
        "stats": scan.stats,
        "created_at": scan.created_at.isoformat() if scan.created_at else None,
    }


@router.delete("/scans/{scan_id}")
async def delete_saved_zap_scan(
    scan_id: int,
    db: DBSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """Delete a saved ZAP scan."""
    scan = db.query(ZAPScan).filter(
        ZAPScan.id == scan_id,
        ZAPScan.user_id == current_user.id
    ).first()
    
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")
    
    db.delete(scan)
    db.commit()
    
    return {"success": True, "message": "Scan deleted"}


@router.get("/scans/interrupted")
async def list_interrupted_scans(
    db: DBSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """
    List ZAP scans that were interrupted and may be recoverable.
    
    Returns scans with status 'scanning', 'spidering', or 'active_scanning'
    that may have been interrupted by a restart.
    """
    interrupted_statuses = ['scanning', 'spidering', 'active_scanning', 'ajax_spidering']
    
    scans = db.query(ZAPScan).filter(
        ZAPScan.user_id == current_user.id,
        ZAPScan.status.in_(interrupted_statuses)
    ).order_by(ZAPScan.created_at.desc()).all()
    
    return {
        "count": len(scans),
        "scans": [
            {
                "id": s.id,
                "session_id": s.session_id,
                "title": s.title,
                "target_url": s.target_url,
                "scan_type": s.scan_type,
                "status": s.status,
                "started_at": s.started_at.isoformat() if s.started_at else None,
                "urls_found": s.urls_found,
                "alerts_before_interrupt": {
                    "high": s.alerts_high or 0,
                    "medium": s.alerts_medium or 0,
                    "low": s.alerts_low or 0,
                    "info": s.alerts_info or 0,
                },
                "can_resume": True,  # These scans can potentially be resumed
            }
            for s in scans
        ]
    }


@router.post("/scans/{scan_id}/resume")
async def resume_zap_scan(
    scan_id: int,
    db: DBSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """
    Resume an interrupted ZAP scan.
    
    This restarts the scan from where it left off (approximately),
    using the same target URL and configuration.
    
    Returns Server-Sent Events for real-time progress.
    """
    scan = db.query(ZAPScan).filter(
        ZAPScan.id == scan_id,
        ZAPScan.user_id == current_user.id
    ).first()
    
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")
    
    # Create config from saved scan
    config = ZAPScanConfig(
        target_url=scan.target_url,
        scan_type=ZAPScanType(scan.scan_type),
    )
    
    scanner = get_zap_scanner()
    
    async def event_generator():
        try:
            # Emit a recovery started event
            yield f"data: {json.dumps({'type': 'recovery', 'message': f'Resuming scan {scan.session_id}', 'original_urls': scan.urls_found, 'original_alerts': scan.alerts_high + scan.alerts_medium + scan.alerts_low + scan.alerts_info})}\n\n"
            
            async for event in scanner.full_scan(
                config,
                user_id=current_user.id,
                project_id=scan.project_id,
                persist_to_db=True
            ):
                yield f"data: {json.dumps(event)}\n\n"
                
        except ZAPError as e:
            logger.error(f"ZAP scan resume error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
        except Exception as e:
            logger.exception(f"Unexpected error resuming ZAP scan: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
    
    # Mark old scan as superseded
    scan.status = 'superseded'
    db.commit()
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


# =============================================================================
# AUTHENTICATION
# =============================================================================

@router.post("/auth/setup")
async def setup_authentication(
    request: AuthenticationRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Configure authentication for ZAP scans.
    
    Supports:
    - **formBasedAuthentication**: Traditional login forms
    - **httpAuthentication**: HTTP Basic/Digest auth
    - **jsonBasedAuthentication**: API/JSON login endpoints
    - **scriptBasedAuthentication**: Custom scripts for OAuth/OIDC/etc.
    
    The configured authentication will be used for subsequent scans.
    """
    scanner = get_zap_scanner()
    
    try:
        # Convert method string to enum
        method_map = {
            "formBasedAuthentication": ZAPAuthMethod.FORM_BASED,
            "httpAuthentication": ZAPAuthMethod.HTTP_BASIC,
            "jsonBasedAuthentication": ZAPAuthMethod.JSON_BASED,
            "scriptBasedAuthentication": ZAPAuthMethod.SCRIPT_BASED,
            "manual": ZAPAuthMethod.MANUAL,
        }
        
        auth_method = method_map.get(request.method)
        if not auth_method:
            raise HTTPException(status_code=400, detail=f"Unknown auth method: {request.method}")
        
        auth_config = ZAPAuthConfig(
            method=auth_method,
            login_url=request.login_url,
            login_request_data=request.login_request_data,
            hostname=request.hostname,
            realm=request.realm,
            json_template=request.json_template,
            script_name=request.script_name,
            script_params=request.script_params,
            username=request.username,
            password=request.password,
            logged_in_indicator=request.logged_in_indicator,
            logged_out_indicator=request.logged_out_indicator,
        )
        
        result = await scanner.setup_authentication(
            auth_config,
            request.target_url,
            request.context_name
        )
        
        return {
            "success": True,
            "message": f"Authentication configured using {request.method}",
            **result
        }
        
    except ZAPError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        logger.exception(f"Error setting up authentication: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/auth/oauth")
async def setup_oauth_authentication(
    request: OAuthRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Configure OAuth2 client credentials authentication.
    
    This obtains an OAuth token and configures ZAP to include it
    in all requests to the target.
    
    For Authorization Code flow, use the script-based auth setup instead.
    """
    scanner = get_zap_scanner()
    
    try:
        result = await scanner.setup_oauth_authentication(
            target_url=request.target_url,
            token_endpoint=request.token_endpoint,
            client_id=request.client_id,
            client_secret=request.client_secret,
            scope=request.scope,
            grant_type=request.grant_type,
            context_name=request.context_name
        )
        
        return {
            "success": True,
            "message": "OAuth2 authentication configured",
            **result
        }
        
    except ZAPError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        logger.exception(f"Error setting up OAuth: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/auth/methods")
async def list_auth_methods(current_user: User = Depends(get_current_active_user)):
    """List supported authentication methods."""
    scanner = get_zap_scanner()
    
    try:
        methods = await scanner.list_auth_methods()
        return {
            "methods": methods,
            "recommended": {
                "web_forms": "formBasedAuthentication",
                "api_json": "jsonBasedAuthentication", 
                "oauth": "scriptBasedAuthentication",
                "basic": "httpAuthentication"
            }
        }
    except Exception as e:
        return {
            "methods": ["formBasedAuthentication", "httpAuthentication", "jsonBasedAuthentication", "scriptBasedAuthentication"],
            "error": str(e)
        }


@router.delete("/auth/context/{context_name}")
async def remove_auth_context(
    context_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Remove an authentication context."""
    scanner = get_zap_scanner()
    success = await scanner.remove_authentication_context(context_name)
    
    if not success:
        raise HTTPException(status_code=500, detail="Failed to remove context")
    
    return {"success": True, "message": f"Context '{context_name}' removed"}


# =============================================================================
# SCAN POLICIES
# =============================================================================

@router.get("/policies")
async def list_scan_policies(current_user: User = Depends(get_current_active_user)):
    """List all available scan policies."""
    scanner = get_zap_scanner()
    
    try:
        policies = await scanner.list_scan_policies()
        return {"policies": policies}
    except Exception as e:
        logger.error(f"Error listing policies: {e}")
        return {"policies": [], "error": str(e)}


@router.post("/policies")
async def create_scan_policy(
    request: ScanPolicyRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Create a custom scan policy.
    
    Policies control which scanners are enabled and their aggressiveness.
    
    Attack Strength: LOW, MEDIUM, HIGH, INSANE
    Alert Threshold: OFF (disabled), LOW, MEDIUM, HIGH
    """
    scanner = get_zap_scanner()
    
    try:
        policy = ZAPScanPolicy(
            name=request.name,
            description=request.description,
            default_attack_strength=request.attack_strength,
            default_alert_threshold=request.alert_threshold,
            enabled_scanners=request.enabled_scanners,
            disabled_scanners=request.disabled_scanners,
        )
        
        result = await scanner.create_scan_policy(policy)
        return {
            "success": True,
            "message": f"Policy '{request.name}' created",
            **result
        }
        
    except Exception as e:
        logger.exception(f"Error creating policy: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/policies/quick")
async def create_quick_policy(
    request: QuickPolicyRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Create a quick policy by scanner category.
    
    Available categories:
    - sql_injection
    - xss (cross-site scripting)
    - path_traversal
    - command_injection
    - remote_file_inclusion
    - ldap_injection
    - xml_injection
    - script_injection
    - server_side_include
    - information_disclosure
    - authentication
    """
    scanner = get_zap_scanner()
    
    try:
        result = await scanner.create_quick_policy(
            name=request.name,
            strength=request.strength,
            categories=request.categories
        )
        return {
            "success": True,
            "message": f"Quick policy '{request.name}' created",
            **result
        }
        
    except Exception as e:
        logger.exception(f"Error creating quick policy: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/policies/{policy_name}")
async def get_scan_policy_details(
    policy_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get detailed information about a scan policy."""
    scanner = get_zap_scanner()
    
    try:
        details = await scanner.get_scan_policy_details(policy_name)
        return details
    except Exception as e:
        logger.error(f"Error getting policy details: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/policies/{policy_name}")
async def delete_scan_policy(
    policy_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Delete a scan policy."""
    scanner = get_zap_scanner()
    success = await scanner.delete_scan_policy(policy_name)
    
    if not success:
        raise HTTPException(status_code=500, detail="Failed to delete policy")
    
    return {"success": True, "message": f"Policy '{policy_name}' deleted"}


@router.get("/scanners")
async def list_scanners(
    policy_name: Optional[str] = Query(None, description="Filter by policy name"),
    current_user: User = Depends(get_current_active_user)
):
    """List all available vulnerability scanners."""
    scanner = get_zap_scanner()
    
    try:
        scanners = await scanner.list_scanners(policy_name)
        return {
            "count": len(scanners),
            "scanners": scanners
        }
    except Exception as e:
        logger.error(f"Error listing scanners: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# SCAN RESUME (True Checkpoint-based)
# =============================================================================

@router.post("/scans/{session_id}/resume-from-checkpoint")
async def resume_scan_from_checkpoint(
    session_id: str,
    db: DBSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user)
):
    """
    Resume an interrupted scan from its last checkpoint.
    
    This performs true resume by:
    1. Loading checkpoint from database (discovered URLs, alerts, phase)
    2. Restoring ZAP state by seeding discovered URLs
    3. Continuing from the interrupted phase
    
    Unlike restart, this preserves previous progress.
    
    Returns Server-Sent Events for real-time progress.
    """
    scanner = get_zap_scanner()
    
    async def event_generator():
        try:
            async for event in scanner.resume_scan(
                session_id=session_id,
                user_id=current_user.id,
                project_id=None
            ):
                yield f"data: {json.dumps(event)}\n\n"
        except Exception as e:
            logger.exception(f"Error resuming scan: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


# =============================================================================
# INTEGRATION WITH AGENTIC FUZZER
# =============================================================================

@router.post("/integrate/merge-findings")
async def merge_findings_endpoint(
    fuzzer_session_id: Optional[str] = Query(None, description="Agentic Fuzzer session ID"),
    zap_session_id: Optional[str] = Query(None, description="ZAP session ID"),
    url: Optional[str] = Query(None, description="Filter by URL"),
    current_user: User = Depends(get_current_active_user)
):
    """
    Merge ZAP findings with Agentic Fuzzer findings.
    
    This creates a unified view of all security findings from both tools,
    with deduplication.
    """
    # Get ZAP findings
    zap_findings = await zap_get_findings(url=url)
    
    # Get fuzzer findings if session ID provided
    fuzzer_findings = []
    if fuzzer_session_id:
        try:
            from backend.services.agentic_fuzzer_service import get_session
            fuzzer_session = get_session(fuzzer_session_id)
            if fuzzer_session:
                fuzzer_findings = [f.to_dict() if hasattr(f, 'to_dict') else f 
                                   for f in fuzzer_session.findings]
        except Exception as e:
            logger.warning(f"Could not get fuzzer session: {e}")
    
    # Merge findings
    merged = await merge_zap_findings_with_fuzzer(zap_findings, fuzzer_findings)
    
    # Count by source
    zap_count = sum(1 for f in merged if f.get("source") == "owasp_zap")
    fuzzer_count = len(merged) - zap_count
    
    return {
        "total": len(merged),
        "by_source": {
            "owasp_zap": zap_count,
            "agentic_fuzzer": fuzzer_count
        },
        "by_severity": {
            "critical": sum(1 for f in merged if f.get("severity") == "critical"),
            "high": sum(1 for f in merged if f.get("severity") == "high"),
            "medium": sum(1 for f in merged if f.get("severity") == "medium"),
            "low": sum(1 for f in merged if f.get("severity") == "low"),
            "info": sum(1 for f in merged if f.get("severity") == "info"),
        },
        "findings": merged
    }


# =============================================================================
# AI ANALYSIS
# =============================================================================

class ScanStatistics(BaseModel):
    """Scan statistics to include in AI analysis."""
    total_messages: int = Field(default=0, description="Total HTTP messages captured")
    urls_discovered: int = Field(default=0, description="Number of URLs discovered")
    hosts_count: int = Field(default=0, description="Number of hosts scanned")
    passive_scan_queue: int = Field(default=0, description="Passive scan queue size")
    scan_duration_seconds: Optional[int] = Field(default=None, description="Scan duration in seconds")
    spider_progress: Optional[int] = Field(default=None, description="Spider progress percentage")
    active_scan_progress: Optional[int] = Field(default=None, description="Active scan progress percentage")

class ZAPAIAnalysisRequest(BaseModel):
    """Request for AI analysis of ZAP findings."""
    alerts: List[Dict[str, Any]] = Field(..., description="ZAP alerts to analyze")
    target_url: str = Field(..., description="Target URL that was scanned")
    include_exploit_chains: bool = Field(default=True, description="Analyze potential exploit chains")
    include_remediation: bool = Field(default=True, description="Generate remediation plan")
    include_business_impact: bool = Field(default=True, description="Assess business impact")
    additional_context: Optional[str] = Field(default=None, description="Additional context for AI analysis")
    scan_statistics: Optional[ScanStatistics] = Field(default=None, description="Scan statistics to include in report")


@router.post("/analyze")
async def ai_analyze_zap_findings(
    request: ZAPAIAnalysisRequest,
    current_user: User = Depends(get_current_active_user)
) -> Dict[str, Any]:
    """
    AI-powered analysis of ZAP security findings.
    
    Provides:
    - Executive summary
    - Risk assessment
    - Exploit chain analysis
    - Remediation prioritization
    - Business impact assessment
    """
    try:
        from backend.services.ai_analysis_service import AIAnalysisService
        
        ai_service = AIAnalysisService()
        alerts = request.alerts
        
        if not alerts:
            return {
                "status": "success",
                "summary": "No vulnerabilities found in the scan.",
                "risk_level": "info",
                "risk_score": 0,
                "findings_analyzed": 0,
                "exploit_chains": [],
                "remediation_plan": [],
                "business_impact": "No security issues detected."
            }
        
        # Convert ZAP alerts to standard finding format
        findings = []
        for alert in alerts:
            finding = {
                "id": alert.get("id", alert.get("alert_id", "")),
                "technique": alert.get("name", "Unknown"),
                "severity": _normalize_severity(alert.get("risk", "info")),
                "title": alert.get("name", "Unknown Alert"),
                "description": alert.get("description", ""),
                "endpoint": alert.get("url", request.target_url),
                "parameter": alert.get("parameter", ""),
                "evidence": alert.get("evidence", ""),
                "solution": alert.get("solution", ""),
                "cwe_id": alert.get("cwe_id", alert.get("cweid", "")),
                "wasc_id": alert.get("wasc_id", alert.get("wascid", "")),
                "source": "owasp_zap"
            }
            findings.append(finding)
        
        # Calculate risk score
        severity_scores = {"critical": 10, "high": 8, "medium": 5, "low": 2, "info": 1}
        total_score = sum(severity_scores.get(f["severity"], 1) for f in findings)
        max_score = len(findings) * 10
        risk_score = min(100, int((total_score / max_score) * 100)) if max_score > 0 else 0
        
        # Determine risk level
        if any(f["severity"] == "critical" for f in findings):
            risk_level = "critical"
        elif any(f["severity"] == "high" for f in findings):
            risk_level = "high"
        elif any(f["severity"] == "medium" for f in findings):
            risk_level = "medium"
        elif any(f["severity"] == "low" for f in findings):
            risk_level = "low"
        else:
            risk_level = "info"
        
        # Generate AI analysis
        result = {
            "status": "success",
            "findings_analyzed": len(findings),
            "risk_score": risk_score,
            "risk_level": risk_level,
            "summary": "",
            "exploit_chains": [],
            "remediation_plan": [],
            "business_impact": "",
            "scan_statistics": request.scan_statistics.dict() if request.scan_statistics else None,
            "findings_by_severity": {
                "critical": [f for f in findings if f["severity"] == "critical"],
                "high": [f for f in findings if f["severity"] == "high"],
                "medium": [f for f in findings if f["severity"] == "medium"],
                "low": [f for f in findings if f["severity"] == "low"],
                "info": [f for f in findings if f["severity"] == "info"],
            },
            "findings_by_type": {},
            "owasp_mapping": {},
        }
        
        # Group by technique/type
        for f in findings:
            tech = f["technique"]
            if tech not in result["findings_by_type"]:
                result["findings_by_type"][tech] = []
            result["findings_by_type"][tech].append(f)
        
        # Try to use AI service for enhanced analysis
        try:
            ai_analysis = await ai_service.analyze_zap_findings(
                findings=findings,
                target_url=request.target_url,
                include_exploit_chains=request.include_exploit_chains,
                include_remediation=request.include_remediation,
                include_business_impact=request.include_business_impact,
                additional_context=request.additional_context
            )
            
            if ai_analysis:
                result["summary"] = ai_analysis.get("summary", "")
                result["exploit_chains"] = ai_analysis.get("exploit_chains", [])
                result["remediation_plan"] = ai_analysis.get("remediation_plan", [])
                result["business_impact"] = ai_analysis.get("business_impact", "")
                result["owasp_mapping"] = ai_analysis.get("owasp_mapping", {})
                # Add attack narrative and offensive insights
                result["attack_narrative"] = ai_analysis.get("attack_narrative", "")
                result["offensive_insights"] = ai_analysis.get("offensive_insights", {})
                result["aggregate_statistics"] = ai_analysis.get("aggregate_statistics", {})
                
        except Exception as e:
            logger.warning(f"AI analysis enhancement failed, using basic analysis: {e}")
            # Generate basic summary if AI fails
            result["summary"] = _generate_basic_summary(findings, risk_level, request.target_url)
            result["remediation_plan"] = _generate_basic_remediation(findings)
            result["business_impact"] = _generate_basic_impact(risk_level, len(findings))
        
        return result
        
    except Exception as e:
        logger.error(f"ZAP AI analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _normalize_severity(risk: str) -> str:
    """Normalize ZAP risk levels to standard severity."""
    risk_lower = risk.lower()
    if risk_lower in ["critical", "high", "medium", "low", "info", "informational"]:
        return "info" if risk_lower == "informational" else risk_lower
    return "info"


def _generate_basic_summary(findings: List[Dict], risk_level: str, target_url: str) -> str:
    """Generate a basic summary when AI is unavailable."""
    severity_counts = {}
    for f in findings:
        sev = f["severity"]
        severity_counts[sev] = severity_counts.get(sev, 0) + 1
    
    parts = []
    if severity_counts.get("critical", 0) > 0:
        parts.append(f"{severity_counts['critical']} critical")
    if severity_counts.get("high", 0) > 0:
        parts.append(f"{severity_counts['high']} high")
    if severity_counts.get("medium", 0) > 0:
        parts.append(f"{severity_counts['medium']} medium")
    if severity_counts.get("low", 0) > 0:
        parts.append(f"{severity_counts['low']} low")
    
    severity_str = ", ".join(parts) if parts else "informational"
    
    return f"DAST scan of {target_url} identified {len(findings)} security findings ({severity_str}). Overall risk level: {risk_level.upper()}. Review high and critical findings immediately."


def _generate_basic_remediation(findings: List[Dict]) -> List[Dict]:
    """Generate basic remediation plan when AI is unavailable."""
    remediation = []
    seen_techniques = set()
    
    # Sort by severity
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    sorted_findings = sorted(findings, key=lambda f: severity_order.get(f["severity"], 4))
    
    for f in sorted_findings:
        tech = f["technique"]
        if tech in seen_techniques:
            continue
        seen_techniques.add(tech)
        
        remediation.append({
            "priority": len(remediation) + 1,
            "vulnerability": tech,
            "severity": f["severity"],
            "recommendation": f.get("solution", "Review and remediate this vulnerability."),
            "affected_urls": [f["endpoint"]],
            "effort": "medium",
        })
        
        if len(remediation) >= 10:
            break
    
    return remediation


def _generate_basic_impact(risk_level: str, finding_count: int) -> str:
    """Generate basic business impact when AI is unavailable."""
    impacts = {
        "critical": f"Critical security vulnerabilities detected ({finding_count} findings). Immediate action required to prevent potential data breach, system compromise, or regulatory violations.",
        "high": f"High-severity security issues identified ({finding_count} findings). These vulnerabilities could lead to unauthorized access or data exposure if exploited.",
        "medium": f"Medium-severity findings detected ({finding_count} findings). While not immediately critical, these should be addressed to maintain security posture.",
        "low": f"Low-severity issues found ({finding_count} findings). These represent minor security improvements that should be addressed during regular maintenance.",
        "info": f"Informational findings only ({finding_count} findings). No significant security risks detected, but review for potential hardening opportunities.",
    }
    return impacts.get(risk_level, impacts["info"])


# =============================================================================
# EXPORT AI ANALYSIS REPORT
# =============================================================================

class ExportAIReportRequest(BaseModel):
    """Request to export AI analysis report."""
    analysis: Dict[str, Any] = Field(..., description="AI analysis result to export")
    target_url: str = Field(..., description="Target URL that was scanned")
    scan_info: Optional[Dict[str, Any]] = Field(default=None, description="Scan metadata")
    format: str = Field(default="markdown", description="Export format: markdown, pdf, word")


@router.post("/export-report")
async def export_ai_analysis_report(
    request: ExportAIReportRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Export AI analysis report in various formats.
    
    Supports:
    - **markdown**: Clean markdown for documentation
    - **pdf**: Professional PDF report with styling
    - **word**: Microsoft Word document (.docx)
    
    Includes offensive security insights and attack narratives.
    """
    try:
        analysis = request.analysis
        target_url = request.target_url
        scan_info = request.scan_info or {}
        format_type = request.format.lower()
        
        if format_type not in ["markdown", "pdf", "word"]:
            raise HTTPException(status_code=400, detail="Invalid format. Use: markdown, pdf, or word")
        
        # Generate markdown content (base for all formats)
        md_content = _generate_zap_ai_report_markdown(analysis, target_url, scan_info)
        
        if format_type == "markdown":
            return Response(
                content=md_content,
                media_type="text/markdown",
                headers={
                    "Content-Disposition": f"attachment; filename=zap_security_report_{dt.utcnow().strftime('%Y%m%d_%H%M%S')}.md"
                }
            )
        
        elif format_type == "pdf":
            pdf_content = _generate_zap_pdf_report(md_content, analysis, target_url)
            return Response(
                content=pdf_content,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=zap_security_report_{dt.utcnow().strftime('%Y%m%d_%H%M%S')}.pdf"
                }
            )
        
        else:  # word
            docx_content = _generate_zap_word_report(md_content, analysis, target_url)
            return Response(
                content=docx_content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=zap_security_report_{dt.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
                }
            )
            
    except Exception as e:
        logger.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _generate_zap_ai_report_markdown(analysis: Dict, target_url: str, scan_info: Dict) -> str:
    """Generate comprehensive markdown report from AI analysis."""
    from datetime import datetime
    
    md = []
    md.append("# 🔒 DAST Security Assessment Report")
    md.append(f"\n**Target:** {target_url}")
    md.append(f"**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    md.append(f"**Scanner:** OWASP ZAP with AI Analysis")
    
    if scan_info:
        md.append(f"**Scan Type:** {scan_info.get('scan_type', 'Full Scan')}")
        if scan_info.get('completed_at'):
            md.append(f"**Completed:** {scan_info.get('completed_at')}")
    
    md.append("\n---\n")
    
    # Risk Score Banner
    risk_score = analysis.get('risk_score', 0)
    risk_level = analysis.get('risk_level', 'info').upper()
    findings_count = analysis.get('findings_analyzed', 0)
    
    md.append("## 📊 Executive Summary\n")
    md.append(f"| Metric | Value |")
    md.append(f"|--------|-------|")
    md.append(f"| **Risk Score** | {risk_score}/100 |")
    md.append(f"| **Risk Level** | {risk_level} |")
    md.append(f"| **Findings** | {findings_count} |")
    
    # Add scan statistics if available
    scan_stats = analysis.get('scan_statistics')
    if scan_stats:
        if scan_stats.get('total_messages'):
            md.append(f"| **HTTP Messages** | {scan_stats.get('total_messages'):,} |")
        if scan_stats.get('urls_discovered'):
            md.append(f"| **URLs Discovered** | {scan_stats.get('urls_discovered'):,} |")
        if scan_stats.get('hosts_count'):
            md.append(f"| **Hosts Scanned** | {scan_stats.get('hosts_count')} |")
        if scan_stats.get('passive_scan_queue') is not None:
            md.append(f"| **Passive Scan Queue** | {scan_stats.get('passive_scan_queue')} |")
    md.append("")
    
    # AI Summary
    if analysis.get('summary'):
        md.append("### Assessment Overview\n")
        md.append(analysis.get('summary'))
        md.append("")
    
    # Findings by Severity
    findings_by_sev = analysis.get('findings_by_severity', {})
    if findings_by_sev:
        md.append("\n### Findings by Severity\n")
        md.append("| Severity | Count |")
        md.append("|----------|-------|")
        for sev in ['critical', 'high', 'medium', 'low', 'info']:
            count = len(findings_by_sev.get(sev, []))
            emoji = {"critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🔵", "info": "⚪"}.get(sev, "⚪")
            md.append(f"| {emoji} {sev.title()} | {count} |")
        md.append("")
    
    # Aggregate Statistics from AI Analysis (if available)
    agg_stats = analysis.get('aggregate_statistics')
    if agg_stats:
        md.append("\n### 📈 Complete Vulnerability Statistics\n")
        md.append(f"**Total Findings Analyzed:** {agg_stats.get('total_findings', 0)}")
        md.append(f"**Unique Vulnerability Types:** {agg_stats.get('total_types', 0)}")
        md.append(f"**Detailed Samples Used:** {agg_stats.get('sampled_count', 0)}\n")
        
        # Severity breakdown
        sev_breakdown = agg_stats.get('severity_breakdown', {})
        if sev_breakdown:
            md.append("| Severity | Total Count |")
            md.append("|----------|-------------|")
            for sev in ['critical', 'high', 'medium', 'low', 'info']:
                count = sev_breakdown.get(sev, 0)
                emoji = {"critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🔵", "info": "⚪"}.get(sev, "⚪")
                md.append(f"| {emoji} {sev.title()} | {count} |")
            md.append("")
        
        # Top vulnerability types
        findings_by_type = agg_stats.get('findings_by_type', [])
        if findings_by_type:
            md.append("\n**Top Vulnerability Types:**\n")
            for vtype in findings_by_type[:10]:
                type_name = vtype.get('type', 'Unknown')
                total = vtype.get('total_count', 0)
                breakdown = vtype.get('severity_breakdown', {})
                breakdown_str = ', '.join([f"{s}: {c}" for s, c in breakdown.items() if c > 0])
                md.append(f"- **{type_name}**: {total} instances ({breakdown_str})")
            md.append("")
    
    # Attack Narrative (Offensive Perspective)
    if analysis.get('attack_narrative'):
        md.append("\n---\n")
        md.append("## ⚔️ Attack Narrative\n")
        md.append("*From an offensive security perspective:*\n")
        md.append(analysis.get('attack_narrative'))
        md.append("")
    
    # Offensive Insights
    offensive = analysis.get('offensive_insights', {})
    if offensive:
        md.append("\n---\n")
        md.append("## 🎯 Offensive Security Insights\n")
        
        if offensive.get('easiest_entry_point'):
            md.append(f"**🚪 Easiest Entry Point:** {offensive.get('easiest_entry_point')}")
        if offensive.get('most_valuable_target'):
            md.append(f"**💎 Most Valuable Target:** {offensive.get('most_valuable_target')}")
        if offensive.get('estimated_time_to_compromise'):
            md.append(f"**⏱️ Est. Time to Compromise:** {offensive.get('estimated_time_to_compromise')}")
        if offensive.get('required_skill_level'):
            md.append(f"**📚 Required Skill Level:** {offensive.get('required_skill_level')}")
        if offensive.get('detection_likelihood'):
            md.append(f"**👁️ Detection Likelihood:** {offensive.get('detection_likelihood')}")
        
        # New enhanced fields from smart sampling
        if offensive.get('attack_surface_assessment'):
            md.append(f"\n**🗺️ Attack Surface Assessment:**\n{offensive.get('attack_surface_assessment')}")
        
        if offensive.get('highest_risk_endpoints'):
            endpoints = offensive.get('highest_risk_endpoints')
            if isinstance(endpoints, list) and endpoints:
                md.append("\n**🎯 Highest Risk Endpoints:**")
                for ep in endpoints[:10]:
                    md.append(f"- `{ep}`")
        
        if offensive.get('recommended_attacker_toolkit'):
            toolkit = offensive.get('recommended_attacker_toolkit')
            if isinstance(toolkit, list) and toolkit:
                md.append(f"\n**🧰 Recommended Attacker Toolkit:** {', '.join(toolkit)}")
        
        md.append("")
    
    # Attack Chains
    exploit_chains = analysis.get('exploit_chains', [])
    if exploit_chains:
        md.append("\n---\n")
        md.append("## 🔗 Exploitation Pathways\n")
        
        for i, chain in enumerate(exploit_chains, 1):
            title = chain.get('title', f'Attack Chain {i}')
            md.append(f"### {i}. {title}\n")
            
            if chain.get('description'):
                md.append(f"{chain.get('description')}\n")
            
            # Chain metadata
            if chain.get('tools') or chain.get('difficulty') or chain.get('real_world_impact'):
                md.append(f"| Attribute | Value |")
                md.append(f"|-----------|-------|")
                if chain.get('tools'):
                    md.append(f"| **Tools** | {chain.get('tools')} |")
                if chain.get('difficulty'):
                    md.append(f"| **Difficulty** | {chain.get('difficulty')} |")
                if chain.get('real_world_impact'):
                    md.append(f"| **Real-World Impact** | {chain.get('real_world_impact')} |")
                md.append("")
            
            # Chain steps
            steps = chain.get('steps', [])
            if steps:
                md.append("**Attack Steps:**\n")
                for j, step in enumerate(steps, 1):
                    if isinstance(step, dict):
                        md.append(f"{j}. **{step.get('action', 'Step')}**")
                        if step.get('detail'):
                            md.append(f"   - {step.get('detail')}")
                    else:
                        md.append(f"{j}. {step}")
                md.append("")
    
    # Remediation Plan
    remediation = analysis.get('remediation_plan', [])
    if remediation:
        md.append("\n---\n")
        md.append("## 🛠️ Remediation Plan\n")
        
        for item in remediation:
            priority = item.get('priority', '-')
            vuln = item.get('vulnerability', 'Unknown')
            severity = item.get('severity', 'info')
            quick_win = "⚡ " if item.get('quick_win') else ""
            
            md.append(f"### {quick_win}Priority {priority}: {vuln}\n")
            md.append(f"**Severity:** {severity.title()}")
            
            if item.get('recommendation'):
                md.append(f"\n**Recommendation:**\n{item.get('recommendation')}")
            
            affected = item.get('affected_urls', [])
            if affected:
                md.append(f"\n**Affected URLs:** {len(affected)}")
                for url in affected[:5]:
                    md.append(f"- `{url}`")
                if len(affected) > 5:
                    md.append(f"- ... and {len(affected) - 5} more")
            
            md.append("")
    
    # Business Impact
    if analysis.get('business_impact'):
        md.append("\n---\n")
        md.append("## 💼 Business Impact Assessment\n")
        md.append(analysis.get('business_impact'))
        md.append("")
    
    # Detailed Findings
    findings_by_type = analysis.get('findings_by_type', {})
    if findings_by_type:
        md.append("\n---\n")
        md.append("## 📋 Detailed Findings\n")
        
        for vuln_type, findings in findings_by_type.items():
            md.append(f"### {vuln_type}\n")
            md.append(f"*{len(findings)} instance(s) found*\n")
            
            for f in findings[:5]:  # Limit to 5 per type
                md.append(f"- **{f.get('severity', 'info').title()}**: `{f.get('endpoint', 'N/A')}`")
                if f.get('parameter'):
                    md.append(f"  - Parameter: `{f.get('parameter')}`")
                if f.get('evidence'):
                    evidence = f.get('evidence')[:100]
                    md.append(f"  - Evidence: `{evidence}...`" if len(f.get('evidence', '')) > 100 else f"  - Evidence: `{evidence}`")
            
            if len(findings) > 5:
                md.append(f"\n*...and {len(findings) - 5} more instances*")
            md.append("")
    
    # Footer
    md.append("\n---\n")
    md.append("*Report generated by VRAgent DAST Scanner with AI-powered analysis*")
    
    return "\n".join(md)


def _generate_zap_pdf_report(md_content: str, analysis: Dict, target_url: str) -> bytes:
    """Generate PDF report from markdown content using WeasyPrint."""
    import io
    
    try:
        from weasyprint import HTML, CSS
        import markdown
        
        # Convert markdown to HTML
        html_body = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code', 'codehilite']
        )
        
        # Professional CSS styling
        css = CSS(string='''
            @page {
                size: A4;
                margin: 2cm;
                @bottom-center {
                    content: "Page " counter(page) " of " counter(pages);
                    font-size: 10px;
                    color: #666;
                }
            }
            body {
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #333;
            }
            h1 {
                color: #1a1a2e;
                border-bottom: 3px solid #e94560;
                padding-bottom: 10px;
                font-size: 24pt;
            }
            h2 {
                color: #16213e;
                border-bottom: 2px solid #0f3460;
                padding-bottom: 8px;
                margin-top: 24px;
                font-size: 18pt;
            }
            h3 {
                color: #0f3460;
                margin-top: 16px;
                font-size: 14pt;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 16px 0;
                font-size: 10pt;
            }
            th, td {
                border: 1px solid #ddd;
                padding: 8px 12px;
                text-align: left;
            }
            th {
                background-color: #16213e;
                color: white;
                font-weight: bold;
            }
            tr:nth-child(even) {
                background-color: #f9f9f9;
            }
            code {
                background-color: #f4f4f4;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: "Courier New", monospace;
                font-size: 9pt;
            }
            pre {
                background-color: #1a1a2e;
                color: #e94560;
                padding: 12px;
                border-radius: 6px;
                overflow-x: auto;
            }
            hr {
                border: none;
                border-top: 1px solid #ddd;
                margin: 20px 0;
            }
            blockquote {
                border-left: 4px solid #e94560;
                padding-left: 16px;
                color: #666;
                font-style: italic;
            }
            ul, ol {
                margin-left: 20px;
            }
            li {
                margin-bottom: 4px;
            }
        ''')
        
        # Wrap in HTML document
        full_html = f'''
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>ZAP Security Report</title>
        </head>
        <body>
            {html_body}
        </body>
        </html>
        '''
        
        # Generate PDF
        pdf_buffer = io.BytesIO()
        HTML(string=full_html).write_pdf(pdf_buffer, stylesheets=[css])
        
        return pdf_buffer.getvalue()
        
    except ImportError as e:
        logger.error(f"WeasyPrint not available: {e}")
        raise HTTPException(status_code=500, detail="PDF generation requires WeasyPrint")
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")


def _generate_zap_word_report(md_content: str, analysis: Dict, target_url: str) -> bytes:
    """Generate Word document from analysis data."""
    import io
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # Title
        title = doc.add_heading('DAST Security Assessment Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        p = doc.add_paragraph()
        p.add_run(f'Target: ').bold = True
        p.add_run(target_url)
        
        p = doc.add_paragraph()
        p.add_run(f'Generated: ').bold = True
        p.add_run(dt.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC'))
        
        p = doc.add_paragraph()
        p.add_run(f'Scanner: ').bold = True
        p.add_run('OWASP ZAP with AI Analysis')
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        
        risk_score = analysis.get('risk_score', 0)
        risk_level = analysis.get('risk_level', 'info').upper()
        findings_count = analysis.get('findings_analyzed', 0)
        
        # Summary table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        rows = [
            ('Risk Score', f'{risk_score}/100'),
            ('Risk Level', risk_level),
            ('Findings', str(findings_count)),
            ('Status', 'Analyzed'),
        ]
        for i, (label, value) in enumerate(rows):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Summary text
        if analysis.get('summary'):
            doc.add_heading('Assessment Overview', level=2)
            doc.add_paragraph(analysis.get('summary'))
        
        # Findings by Severity
        findings_by_sev = analysis.get('findings_by_severity', {})
        if findings_by_sev:
            doc.add_heading('Findings by Severity', level=2)
            
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            
            header = table.rows[0]
            header.cells[0].text = 'Severity'
            header.cells[0].paragraphs[0].runs[0].bold = True
            header.cells[1].text = 'Count'
            header.cells[1].paragraphs[0].runs[0].bold = True
            
            for i, sev in enumerate(['critical', 'high', 'medium', 'low', 'info'], 1):
                count = len(findings_by_sev.get(sev, []))
                table.rows[i].cells[0].text = sev.title()
                table.rows[i].cells[1].text = str(count)
            
            doc.add_paragraph()
        
        # Attack Narrative
        if analysis.get('attack_narrative'):
            doc.add_heading('Attack Narrative', level=1)
            doc.add_paragraph('From an offensive security perspective:', style='Intense Quote')
            doc.add_paragraph(analysis.get('attack_narrative'))
        
        # Offensive Insights
        offensive = analysis.get('offensive_insights', {})
        if offensive:
            doc.add_heading('Offensive Security Insights', level=1)
            
            if offensive.get('easiest_entry_point'):
                p = doc.add_paragraph()
                p.add_run('Easiest Entry Point: ').bold = True
                p.add_run(offensive.get('easiest_entry_point'))
            
            if offensive.get('most_valuable_target'):
                p = doc.add_paragraph()
                p.add_run('Most Valuable Target: ').bold = True
                p.add_run(offensive.get('most_valuable_target'))
            
            if offensive.get('estimated_time_to_compromise'):
                p = doc.add_paragraph()
                p.add_run('Est. Time to Compromise: ').bold = True
                p.add_run(offensive.get('estimated_time_to_compromise'))
            
            if offensive.get('required_skill_level'):
                p = doc.add_paragraph()
                p.add_run('Required Skill Level: ').bold = True
                p.add_run(offensive.get('required_skill_level'))
            
            if offensive.get('detection_likelihood'):
                p = doc.add_paragraph()
                p.add_run('Detection Likelihood: ').bold = True
                p.add_run(offensive.get('detection_likelihood'))
        
        # Attack Chains
        exploit_chains = analysis.get('exploit_chains', [])
        if exploit_chains:
            doc.add_heading('Exploitation Pathways', level=1)
            
            for i, chain in enumerate(exploit_chains, 1):
                title = chain.get('title', f'Attack Chain {i}')
                doc.add_heading(f'{i}. {title}', level=2)
                
                if chain.get('description'):
                    doc.add_paragraph(chain.get('description'))
                
                # Chain metadata
                if chain.get('tools'):
                    p = doc.add_paragraph()
                    p.add_run('Tools: ').bold = True
                    p.add_run(chain.get('tools'))
                
                if chain.get('difficulty'):
                    p = doc.add_paragraph()
                    p.add_run('Difficulty: ').bold = True
                    p.add_run(chain.get('difficulty'))
                
                if chain.get('real_world_impact'):
                    p = doc.add_paragraph()
                    p.add_run('Real-World Impact: ').bold = True
                    p.add_run(chain.get('real_world_impact'))
                
                # Steps
                steps = chain.get('steps', [])
                if steps:
                    doc.add_paragraph('Attack Steps:', style='Intense Quote')
                    for j, step in enumerate(steps, 1):
                        if isinstance(step, dict):
                            step_text = f"{j}. {step.get('action', 'Step')}"
                            if step.get('detail'):
                                step_text += f" - {step.get('detail')}"
                        else:
                            step_text = f"{j}. {step}"
                        doc.add_paragraph(step_text, style='List Number')
        
        # Remediation Plan
        remediation = analysis.get('remediation_plan', [])
        if remediation:
            doc.add_heading('Remediation Plan', level=1)
            
            for item in remediation:
                priority = item.get('priority', '-')
                vuln = item.get('vulnerability', 'Unknown')
                quick_win = ' ⚡ Quick Win' if item.get('quick_win') else ''
                
                doc.add_heading(f'Priority {priority}: {vuln}{quick_win}', level=2)
                
                p = doc.add_paragraph()
                p.add_run('Severity: ').bold = True
                p.add_run(item.get('severity', 'info').title())
                
                if item.get('recommendation'):
                    doc.add_paragraph('Recommendation:', style='Intense Quote')
                    doc.add_paragraph(item.get('recommendation'))
                
                affected = item.get('affected_urls', [])
                if affected:
                    p = doc.add_paragraph()
                    p.add_run(f'Affected URLs ({len(affected)}):').bold = True
                    for url in affected[:5]:
                        doc.add_paragraph(url, style='List Bullet')
                    if len(affected) > 5:
                        doc.add_paragraph(f'...and {len(affected) - 5} more')
        
        # Business Impact
        if analysis.get('business_impact'):
            doc.add_heading('Business Impact Assessment', level=1)
            doc.add_paragraph(analysis.get('business_impact'))
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph('Report generated by VRAgent DAST Scanner with AI-powered analysis', style='Intense Quote')
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except ImportError as e:
        logger.error(f"python-docx not available: {e}")
        raise HTTPException(status_code=500, detail="Word generation requires python-docx")
    except Exception as e:
        logger.error(f"Word generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Word generation failed: {e}")


# =============================================================================
# AI CHAT ENDPOINT
# =============================================================================

class ZAPChatRequest(BaseModel):
    """Request for ZAP AI chat."""
    message: str = Field(..., description="User message")
    context: str = Field(default="", description="Scan context summary")
    history: List[Dict[str, str]] = Field(default=[], description="Chat history")
    alerts: List[Dict[str, Any]] = Field(default=[], description="Sample alerts for context")


class ZAPChatResponse(BaseModel):
    """Response from ZAP AI chat."""
    response: str


@router.post("/chat", response_model=ZAPChatResponse)
async def zap_ai_chat(
    request: ZAPChatRequest,
    current_user: Optional[User] = Depends(get_optional_user)
):
    """
    AI-powered chat assistant for ZAP scan results.
    
    Ask questions about vulnerabilities, attack vectors, remediation strategies,
    and get contextual help understanding scan results.
    """
    try:
        from google import genai
        import os

        api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="AI service not configured")

        client = genai.Client(api_key=api_key)
        model_id = "gemini-3-flash-preview"
        
        # Build system prompt
        system_prompt = """You are an expert penetration tester and security analyst assistant. 
You help users understand and remediate vulnerabilities found during security scans.

Your expertise includes:
- OWASP Top 10 vulnerabilities
- Web application security testing
- Attack vectors and exploit chains
- Remediation strategies and secure coding practices
- Risk assessment and prioritization

Be concise but thorough. Use markdown formatting for code snippets and lists.
When discussing specific vulnerabilities, explain:
1. What the vulnerability is
2. How it can be exploited
3. The potential impact
4. How to fix it with code examples when relevant

If the user asks about vulnerabilities in their scan, reference the specific findings provided."""

        # Build context from alerts
        alerts_context = ""
        if request.alerts:
            alerts_context = "\n\nScan Findings Summary:\n"
            for alert in request.alerts[:15]:
                alerts_context += f"- [{alert.get('risk', 'Unknown')}] {alert.get('name', 'Unknown')}: {alert.get('url', 'N/A')}\n"
                if alert.get('description'):
                    alerts_context += f"  Description: {alert.get('description')[:150]}...\n"
        
        # Build conversation history
        history_text = ""
        if request.history:
            history_text = "\n\nPrevious conversation:\n"
            for msg in request.history[-6:]:  # Last 6 messages
                role = "User" if msg.get('role') == 'user' else "Assistant"
                history_text += f"{role}: {msg.get('content', '')[:500]}\n"
        
        # Build full prompt
        full_prompt = f"""{system_prompt}

Current Scan Context:
{request.context}
{alerts_context}
{history_text}

User Question: {request.message}

Provide a helpful, security-focused response:"""

        # Generate response
        response = client.models.generate_content(
            model=model_id,
            contents=full_prompt
        )

        return ZAPChatResponse(response=response.text)
        
    except ImportError:
        # Fallback to OpenAI if available
        try:
            import openai
            import os
            
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                raise HTTPException(status_code=500, detail="AI service not configured")
            
            client = openai.OpenAI(api_key=api_key)
            
            messages = [
                {"role": "system", "content": """You are an expert penetration tester and security analyst assistant.
You help users understand and remediate vulnerabilities found during security scans.
Be concise but thorough. Use markdown formatting."""},
                {"role": "user", "content": f"""Context: {request.context}

Alerts: {json.dumps(request.alerts[:10]) if request.alerts else 'None'}

Question: {request.message}"""}
            ]
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=messages,
                max_tokens=1000
            )
            
            return ZAPChatResponse(response=response.choices[0].message.content)
            
        except Exception as e:
            logger.error(f"OpenAI fallback failed: {e}")
            raise HTTPException(status_code=500, detail="AI service unavailable")
            
    except Exception as e:
        logger.error(f"AI chat failed: {e}")
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")


# =============================================================================
# WEBSOCKET TESTING ENDPOINTS
# =============================================================================

class WebSocketChannelResponse(BaseModel):
    """Response containing WebSocket channels."""
    channels: List[Dict[str, Any]]


class WebSocketMessageResponse(BaseModel):
    """Response containing WebSocket messages."""
    messages: List[Dict[str, Any]]


class SendWebSocketMessageRequest(BaseModel):
    """Request to send a WebSocket message."""
    channel_id: int = Field(..., description="Channel ID to send message to")
    message: str = Field(..., description="Message content to send")
    outgoing: bool = Field(default=True, description="True for outgoing, False for incoming")


class SetWebSocketBreakRequest(BaseModel):
    """Request to set WebSocket break pattern."""
    message: str = Field(..., description="Message pattern to break on")
    outgoing: bool = Field(default=True, description="True for outgoing, False for incoming")


@router.get("/websocket/channels", response_model=WebSocketChannelResponse)
async def get_websocket_channels(
    current_user: User = Depends(get_current_active_user)
):
    """Get all WebSocket channels discovered by ZAP."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            channels = await client.websocket_channels()
            return WebSocketChannelResponse(channels=channels)
    except Exception as e:
        logger.error(f"Failed to get WebSocket channels: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/websocket/messages", response_model=WebSocketMessageResponse)
async def get_websocket_messages(
    channel_id: Optional[int] = Query(default=None, description="Filter by channel ID"),
    start: int = Query(default=0, ge=0, description="Start index"),
    count: int = Query(default=100, ge=1, le=1000, description="Number of messages"),
    current_user: User = Depends(get_current_active_user)
):
    """Get WebSocket messages, optionally filtered by channel."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            messages = await client.websocket_messages(
                channel_id=channel_id,
                start=start,
                count=count
            )
            return WebSocketMessageResponse(messages=messages)
    except Exception as e:
        logger.error(f"Failed to get WebSocket messages: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/websocket/messages/{channel_id}/{message_id}")
async def get_websocket_message(
    channel_id: int,
    message_id: int,
    current_user: User = Depends(get_current_active_user)
):
    """Get a specific WebSocket message."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            message = await client.websocket_message(message_id, channel_id)
            return {"message": message}
    except Exception as e:
        logger.error(f"Failed to get WebSocket message: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/websocket/send")
async def send_websocket_message(
    request: SendWebSocketMessageRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Send a WebSocket message through ZAP."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.websocket_send_text_message(
                channel_id=request.channel_id,
                outgoing=request.outgoing,
                message=request.message
            )
            return {"status": "sent", "result": result}
    except Exception as e:
        logger.error(f"Failed to send WebSocket message: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/websocket/break")
async def set_websocket_break(
    request: SetWebSocketBreakRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Set a break pattern for WebSocket messages."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.websocket_set_break_text_message(
                message=request.message,
                outgoing=request.outgoing
            )
            return {"status": "break_set", "result": result}
    except Exception as e:
        logger.error(f"Failed to set WebSocket break: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# GRAPHQL TESTING ENDPOINTS
# =============================================================================

class GraphQLOptionsResponse(BaseModel):
    """Response containing GraphQL configuration options."""
    args_type: str = "INLINE"
    max_args_depth: int = 5
    max_query_depth: int = 5
    optional_args_enabled: bool = False
    query_split_type: str = "LEAF"
    request_method: str = "POST_JSON"
    lenient_max_query_depth: bool = False


class GraphQLImportUrlRequest(BaseModel):
    """Request to import GraphQL schema via introspection."""
    url: str = Field(..., description="GraphQL endpoint URL")
    endpoint_url: Optional[str] = Field(default=None, description="Optional different endpoint URL for testing")


class GraphQLImportFileRequest(BaseModel):
    """Request to import GraphQL schema from file content."""
    schema_content: str = Field(..., description="GraphQL schema SDL content")
    endpoint_url: str = Field(..., description="Endpoint URL for testing")


class GraphQLOptionsRequest(BaseModel):
    """Request to update GraphQL testing options."""
    args_type: Optional[str] = Field(default=None, description="INLINE, JSON, or BOTH")
    max_args_depth: Optional[int] = Field(default=None, ge=1, le=20, description="Max argument depth")
    max_query_depth: Optional[int] = Field(default=None, ge=1, le=20, description="Max query depth")
    optional_args_enabled: Optional[bool] = Field(default=None, description="Include optional args")
    query_split_type: Optional[str] = Field(default=None, description="LEAF, ROOT_FIELD, or OPERATION")
    request_method: Optional[str] = Field(default=None, description="POST_JSON, POST_GRAPHQL, or GET")
    lenient_max_query_depth: Optional[bool] = Field(default=None, description="Enable lenient depth")


@router.get("/graphql/options", response_model=GraphQLOptionsResponse)
async def get_graphql_options(
    current_user: User = Depends(get_current_active_user)
):
    """Get current GraphQL testing configuration."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            options = GraphQLOptionsResponse(
                max_args_depth=await client.graphql_option_max_args_depth(),
                max_query_depth=await client.graphql_option_max_query_depth(),
                optional_args_enabled=await client.graphql_option_optional_args_enabled(),
                query_split_type=await client.graphql_option_query_split_type(),
                request_method=await client.graphql_option_request_method(),
            )
            return options
    except Exception as e:
        logger.error(f"Failed to get GraphQL options: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/graphql/options")
async def update_graphql_options(
    request: GraphQLOptionsRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Update GraphQL testing configuration."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            results = []
            
            if request.args_type:
                result = await client.graphql_set_option_args_type(request.args_type)
                results.append({"args_type": result})
            
            if request.max_args_depth is not None:
                result = await client.graphql_set_option_max_args_depth(request.max_args_depth)
                results.append({"max_args_depth": result})
            
            if request.max_query_depth is not None:
                result = await client.graphql_set_option_max_query_depth(request.max_query_depth)
                results.append({"max_query_depth": result})
            
            if request.optional_args_enabled is not None:
                result = await client.graphql_set_option_optional_args(request.optional_args_enabled)
                results.append({"optional_args_enabled": result})
            
            if request.query_split_type:
                result = await client.graphql_set_option_query_split_type(request.query_split_type)
                results.append({"query_split_type": result})
            
            if request.request_method:
                result = await client.graphql_set_option_request_method(request.request_method)
                results.append({"request_method": result})
            
            if request.lenient_max_query_depth is not None:
                result = await client.graphql_set_option_lenient_max_query_depth(request.lenient_max_query_depth)
                results.append({"lenient_max_query_depth": result})
            
            return {"status": "updated", "results": results}
    except Exception as e:
        logger.error(f"Failed to update GraphQL options: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/graphql/import/url")
async def import_graphql_url(
    request: GraphQLImportUrlRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Import a GraphQL schema via introspection query.
    
    This will send an introspection query to the target URL and use the
    resulting schema for security testing.
    """
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.graphql_import_url(
                url=request.url,
                endpoint_url=request.endpoint_url
            )
            return {
                "status": "imported",
                "url": request.url,
                "result": result
            }
    except Exception as e:
        logger.error(f"Failed to import GraphQL schema: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/graphql/import/schema")
async def import_graphql_schema(
    request: GraphQLImportFileRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Import a GraphQL schema from SDL content.
    
    This allows importing schemas when introspection is disabled.
    The schema content should be valid GraphQL SDL.
    """
    try:
        import tempfile
        import os
        
        # Write schema to temp file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.graphql', delete=False) as f:
            f.write(request.schema_content)
            temp_path = f.name
        
        try:
            scanner = get_zap_scanner()
            async with scanner.client() as client:
                result = await client.graphql_import_file(
                    file_path=temp_path,
                    endpoint_url=request.endpoint_url
                )
                return {
                    "status": "imported",
                    "endpoint_url": request.endpoint_url,
                    "result": result
                }
        finally:
            # Clean up temp file
            os.unlink(temp_path)
    except Exception as e:
        logger.error(f"Failed to import GraphQL schema: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# OPENAPI/SWAGGER IMPORT ENDPOINTS
# =============================================================================

class OpenAPIImportUrlRequest(BaseModel):
    """Request to import OpenAPI/Swagger from URL."""
    url: str = Field(..., description="URL to OpenAPI/Swagger definition")
    host_override: Optional[str] = Field(default=None, description="Override host from definition")
    context_id: Optional[int] = Field(default=None, description="ZAP context ID")


class OpenAPIImportFileRequest(BaseModel):
    """Request to import OpenAPI/Swagger from content."""
    content: str = Field(..., description="OpenAPI/Swagger definition content (JSON or YAML)")
    target_url: str = Field(..., description="Target URL for API testing")
    context_id: Optional[int] = Field(default=None, description="ZAP context ID")


@router.post("/openapi/import/url")
async def import_openapi_url(
    request: OpenAPIImportUrlRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Import an OpenAPI/Swagger definition from a URL.
    
    This automatically discovers all API endpoints defined in the specification
    and adds them to ZAP for security testing.
    """
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.openapi_import_url(
                url=request.url,
                host_override=request.host_override,
                context_id=request.context_id
            )
            return {
                "status": "imported",
                "url": request.url,
                "result": result
            }
    except Exception as e:
        logger.error(f"Failed to import OpenAPI definition: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/openapi/import/file")
async def import_openapi_file(
    request: OpenAPIImportFileRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Import an OpenAPI/Swagger definition from content.
    
    Supports both JSON and YAML formats.
    """
    try:
        import tempfile
        import os
        
        # Determine file extension based on content
        content = request.content.strip()
        suffix = '.yaml' if content.startswith(('openapi:', 'swagger:', '---')) else '.json'
        
        # Write to temp file
        with tempfile.NamedTemporaryFile(mode='w', suffix=suffix, delete=False) as f:
            f.write(request.content)
            temp_path = f.name
        
        try:
            scanner = get_zap_scanner()
            async with scanner.client() as client:
                result = await client.openapi_import_file(
                    file_path=temp_path,
                    target_url=request.target_url,
                    context_id=request.context_id
                )
                return {
                    "status": "imported",
                    "target_url": request.target_url,
                    "result": result
                }
        finally:
            os.unlink(temp_path)
    except Exception as e:
        logger.error(f"Failed to import OpenAPI definition: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# MANUAL REQUEST EDITOR ENDPOINTS
# =============================================================================

class SendRequestRequest(BaseModel):
    """Request to send a manual HTTP request."""
    request: str = Field(..., description="Full HTTP request (method, URL, headers, body)")
    follow_redirects: bool = Field(default=True, description="Follow redirects")


class MessageResponse(BaseModel):
    """Response containing a message (request/response pair)."""
    id: int
    request_header: str
    request_body: str
    response_header: str
    response_body: str
    timestamp: Optional[str] = None


@router.post("/request/send")
async def send_manual_request(
    request: SendRequestRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Send a manual HTTP request through ZAP proxy.
    
    The request should be in raw HTTP format:
    ```
    GET /api/users HTTP/1.1
    Host: example.com
    Authorization: Bearer token123
    Content-Type: application/json
    
    {"key": "value"}
    ```
    """
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.send_request(
                request=request.request,
                follow_redirects=request.follow_redirects
            )
            return {
                "status": "sent",
                "result": result
            }
    except Exception as e:
        logger.error(f"Failed to send request: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/request/messages")
async def get_messages(
    base_url: Optional[str] = Query(default=None, description="Filter by base URL"),
    start: int = Query(default=0, ge=0, description="Start index"),
    count: int = Query(default=50, ge=1, le=500, description="Number of messages"),
    current_user: User = Depends(get_current_active_user)
):
    """Get messages from ZAP history."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            messages = await client.get_messages(
                base_url=base_url,
                start=start,
                count=count
            )
            return {"messages": messages}
    except Exception as e:
        logger.error(f"Failed to get messages: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/request/messages/{message_id}")
async def get_message_detail(
    message_id: int,
    current_user: User = Depends(get_current_active_user)
):
    """Get detailed information about a specific message."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            message = await client.get_message(message_id)
            request_header = await client.get_request_header(message_id)
            request_body = await client.get_request_body(message_id)
            response_header = await client.get_response_header(message_id)
            response_body = await client.get_response_body(message_id)
            
            return {
                "id": message_id,
                "message": message,
                "request_header": request_header,
                "request_body": request_body,
                "response_header": response_header,
                "response_body": response_body,
            }
    except Exception as e:
        logger.error(f"Failed to get message: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# CUSTOM SCAN POLICIES ENDPOINTS
# =============================================================================

class ScanPolicyResponse(BaseModel):
    """Response containing scan policies."""
    policies: List[Dict[str, Any]]


class CreateScanPolicyRequest(BaseModel):
    """Request to create a new scan policy."""
    name: str = Field(..., description="Policy name")
    alert_threshold: Optional[str] = Field(default="MEDIUM", description="Alert threshold: OFF, LOW, MEDIUM, HIGH")
    attack_strength: Optional[str] = Field(default="MEDIUM", description="Attack strength: LOW, MEDIUM, HIGH, INSANE")


class UpdateScannerRequest(BaseModel):
    """Request to update scanner settings in a policy."""
    scanner_id: int = Field(..., description="Scanner ID")
    alert_threshold: Optional[str] = Field(default=None, description="Alert threshold")
    attack_strength: Optional[str] = Field(default=None, description="Attack strength")
    enabled: Optional[bool] = Field(default=None, description="Enable/disable scanner")


class BulkUpdateScannersRequest(BaseModel):
    """Request to enable/disable multiple scanners."""
    scanner_ids: List[int] = Field(..., description="Scanner IDs")
    enabled: bool = Field(..., description="Enable or disable")
    policy_name: Optional[str] = Field(default=None, description="Policy name")


@router.get("/policies", response_model=ScanPolicyResponse)
async def list_scan_policies(
    current_user: User = Depends(get_current_active_user)
):
    """List all available scan policies."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            policies = await client.ascan_list_policies()
            return ScanPolicyResponse(policies=policies)
    except Exception as e:
        logger.error(f"Failed to list policies: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/policies")
async def create_scan_policy(
    request: CreateScanPolicyRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Create a new scan policy."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            policy_id = await client.ascan_add_policy(
                name=request.name,
                alert_threshold=request.alert_threshold,
                attack_strength=request.attack_strength
            )
            return {
                "status": "created",
                "name": request.name,
                "policy_id": policy_id
            }
    except Exception as e:
        logger.error(f"Failed to create policy: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/policies/{policy_name}")
async def delete_scan_policy(
    policy_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Delete a scan policy."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.ascan_remove_policy(policy_name)
            return {"status": "deleted", "name": policy_name, "result": result}
    except Exception as e:
        logger.error(f"Failed to delete policy: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/policies/{policy_name}/scanners")
async def list_policy_scanners(
    policy_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """List all scanners in a policy with their settings."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            scanners = await client.ascan_list_scanners(policy_name)
            return {"policy_name": policy_name, "scanners": scanners}
    except Exception as e:
        logger.error(f"Failed to list scanners: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/policies/{policy_name}/scanners")
async def update_policy_scanner(
    policy_name: str,
    request: UpdateScannerRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Update a scanner's settings in a policy."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            results = []
            
            if request.alert_threshold:
                result = await client.ascan_set_scanner_alert_threshold(
                    scanner_id=request.scanner_id,
                    alert_threshold=request.alert_threshold,
                    policy_name=policy_name
                )
                results.append({"alert_threshold": result})
            
            if request.attack_strength:
                result = await client.ascan_set_scanner_attack_strength(
                    scanner_id=request.scanner_id,
                    attack_strength=request.attack_strength,
                    policy_name=policy_name
                )
                results.append({"attack_strength": result})
            
            if request.enabled is not None:
                if request.enabled:
                    result = await client.ascan_enable_scanners([request.scanner_id], policy_name)
                else:
                    result = await client.ascan_disable_scanners([request.scanner_id], policy_name)
                results.append({"enabled": result})
            
            return {"status": "updated", "results": results}
    except Exception as e:
        logger.error(f"Failed to update scanner: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/policies/{policy_name}/scanners/bulk")
async def bulk_update_scanners(
    policy_name: str,
    request: BulkUpdateScannersRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Enable or disable multiple scanners at once."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            if request.enabled:
                result = await client.ascan_enable_scanners(request.scanner_ids, policy_name)
            else:
                result = await client.ascan_disable_scanners(request.scanner_ids, policy_name)
            
            return {
                "status": "updated",
                "scanner_count": len(request.scanner_ids),
                "enabled": request.enabled,
                "result": result
            }
    except Exception as e:
        logger.error(f"Failed to bulk update scanners: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/policies/{policy_name}/enable-all")
async def enable_all_scanners(
    policy_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Enable all scanners in a policy."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.ascan_enable_all_scanners(policy_name)
            return {"status": "enabled_all", "policy_name": policy_name, "result": result}
    except Exception as e:
        logger.error(f"Failed to enable all scanners: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/policies/{policy_name}/disable-all")
async def disable_all_scanners(
    policy_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Disable all scanners in a policy."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.ascan_disable_all_scanners(policy_name)
            return {"status": "disabled_all", "policy_name": policy_name, "result": result}
    except Exception as e:
        logger.error(f"Failed to disable all scanners: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# CONTEXT MANAGEMENT ENDPOINTS
# =============================================================================

class CreateContextRequest(BaseModel):
    """Request to create a new context."""
    name: str = Field(..., description="Context name")


class ContextRegexRequest(BaseModel):
    """Request to add include/exclude regex."""
    regex: str = Field(..., description="Regex pattern to match URLs")


class ContextTechnologyRequest(BaseModel):
    """Request to include/exclude technologies."""
    technologies: List[str] = Field(..., description="List of technology names")


class ContextScopeRequest(BaseModel):
    """Request to set context scope."""
    in_scope: bool = Field(..., description="Whether context is in scope")


class ContextExportRequest(BaseModel):
    """Request to export context."""
    file_path: str = Field(..., description="File path to export to")


class ContextImportRequest(BaseModel):
    """Request to import context."""
    file_path: str = Field(..., description="File path to import from")


@router.get("/contexts")
async def list_contexts(
    current_user: User = Depends(get_current_active_user)
):
    """List all ZAP contexts."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            contexts = await client.context_list()
            return {"contexts": contexts}
    except Exception as e:
        logger.error(f"Failed to list contexts: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts")
async def create_context(
    request: CreateContextRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Create a new context."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            context_id = await client.context_new(request.name)
            return {"status": "created", "name": request.name, "context_id": context_id}
    except Exception as e:
        logger.error(f"Failed to create context: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/contexts/{context_name}")
async def get_context(
    context_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get context details."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            context = await client.context_get(context_name)
            include_regexes = await client.context_get_include_regexes(context_name)
            exclude_regexes = await client.context_get_exclude_regexes(context_name)
            included_tech = await client.context_get_included_technology(context_name)
            excluded_tech = await client.context_get_excluded_technology(context_name)
            
            return {
                "name": context_name,
                "context": context,
                "include_regexes": include_regexes,
                "exclude_regexes": exclude_regexes,
                "included_technologies": included_tech,
                "excluded_technologies": excluded_tech,
            }
    except Exception as e:
        logger.error(f"Failed to get context: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/contexts/{context_name}")
async def delete_context(
    context_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Delete a context."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_remove(context_name)
            return {"status": "deleted", "name": context_name, "result": result}
    except Exception as e:
        logger.error(f"Failed to delete context: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/{context_name}/include")
async def add_include_regex(
    context_name: str,
    request: ContextRegexRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Add a URL pattern to include in context."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_include_in_context(context_name, request.regex)
            return {"status": "added", "regex": request.regex, "result": result}
    except Exception as e:
        logger.error(f"Failed to add include regex: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/{context_name}/exclude")
async def add_exclude_regex(
    context_name: str,
    request: ContextRegexRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Add a URL pattern to exclude from context."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_exclude_from_context(context_name, request.regex)
            return {"status": "added", "regex": request.regex, "result": result}
    except Exception as e:
        logger.error(f"Failed to add exclude regex: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/{context_name}/scope")
async def set_context_scope(
    context_name: str,
    request: ContextScopeRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Set whether context is in scope."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_set_in_scope(context_name, request.in_scope)
            return {"status": "updated", "in_scope": request.in_scope, "result": result}
    except Exception as e:
        logger.error(f"Failed to set context scope: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/contexts/technologies/available")
async def list_available_technologies(
    current_user: User = Depends(get_current_active_user)
):
    """List all available technologies for context filtering."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            technologies = await client.context_get_technology_list()
            return {"technologies": technologies}
    except Exception as e:
        logger.error(f"Failed to list technologies: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/{context_name}/technologies/include")
async def include_technologies(
    context_name: str,
    request: ContextTechnologyRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Include technologies in context scan."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_include_technology(context_name, request.technologies)
            return {"status": "included", "technologies": request.technologies, "result": result}
    except Exception as e:
        logger.error(f"Failed to include technologies: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/{context_name}/technologies/exclude")
async def exclude_technologies(
    context_name: str,
    request: ContextTechnologyRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Exclude technologies from context scan."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_exclude_technology(context_name, request.technologies)
            return {"status": "excluded", "technologies": request.technologies, "result": result}
    except Exception as e:
        logger.error(f"Failed to exclude technologies: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/{context_name}/technologies/include-all")
async def include_all_technologies(
    context_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Include all technologies in context."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_include_all_technologies(context_name)
            return {"status": "included_all", "result": result}
    except Exception as e:
        logger.error(f"Failed to include all technologies: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/{context_name}/technologies/exclude-all")
async def exclude_all_technologies(
    context_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Exclude all technologies from context."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_exclude_all_technologies(context_name)
            return {"status": "excluded_all", "result": result}
    except Exception as e:
        logger.error(f"Failed to exclude all technologies: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/{context_name}/export")
async def export_context(
    context_name: str,
    request: ContextExportRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Export context to file."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_export(context_name, request.file_path)
            return {"status": "exported", "file_path": request.file_path, "result": result}
    except Exception as e:
        logger.error(f"Failed to export context: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/contexts/import")
async def import_context(
    request: ContextImportRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Import context from file."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.context_import(request.file_path)
            return {"status": "imported", "file_path": request.file_path, "result": result}
    except Exception as e:
        logger.error(f"Failed to import context: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# FORCED BROWSE / DIRECTORY DISCOVERY ENDPOINTS
# =============================================================================

class ForcedBrowseRequest(BaseModel):
    """Request to start forced browse scan."""
    url: str = Field(..., description="Target URL to scan")
    recurse: bool = Field(default=True, description="Scan recursively")


class ForcedBrowseOptionsRequest(BaseModel):
    """Request to set forced browse options."""
    threads: Optional[int] = Field(default=None, ge=1, le=50, description="Number of threads")
    recursive: Optional[bool] = Field(default=None, description="Enable recursive scanning")
    fail_case_string: Optional[str] = Field(default=None, description="String indicating 404/failed response")


class ForcedBrowseWordlistRequest(BaseModel):
    """Request to add custom wordlist."""
    file_path: str = Field(..., description="Path to wordlist file")


class LocalForcedBrowseRequest(BaseModel):
    """Request to start local forced browse scan."""
    url: str = Field(..., description="Target URL to scan")
    wordlist: str = Field(default="directories_comprehensive.txt", description="Wordlist to use")
    recursive: bool = Field(default=False, description="Recursively scan discovered directories")
    threads: int = Field(default=10, ge=1, le=50, description="Concurrent threads")
    extensions: Optional[List[str]] = Field(default=None, description="File extensions to try")


@router.get("/forcedBrowse/status")
async def get_forced_browse_status(
    session_id: str = Query(..., description="Session ID from start scan"),
    current_user: User = Depends(get_current_active_user)
):
    """Get forced browse scan status and progress."""
    try:
        from backend.services.forced_browse_service import get_forced_browse_service
        service = get_forced_browse_service()
        status = service.get_status(session_id)
        if "error" in status:
            raise HTTPException(status_code=404, detail=status["error"])
        return status
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get forced browse status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/forcedBrowse/results")
async def get_forced_browse_results(
    session_id: str = Query(..., description="Session ID from start scan"),
    current_user: User = Depends(get_current_active_user)
):
    """Get discovered paths from forced browse scan."""
    try:
        from backend.services.forced_browse_service import get_forced_browse_service
        service = get_forced_browse_service()
        results = service.get_results(session_id)
        return {"results": results, "count": len(results)}
    except Exception as e:
        logger.error(f"Failed to get forced browse results: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/forcedBrowse/scan")
async def start_forced_browse(
    request: LocalForcedBrowseRequest,
    current_user: User = Depends(get_current_active_user)
):
    """
    Start a local forced browse scan to discover hidden files and directories.
    
    Uses local wordlists to brute-force common file/directory names.
    Works entirely offline without requiring external APIs.
    Useful for finding admin panels, backup files, configuration files, etc.
    """
    try:
        from backend.services.forced_browse_service import get_forced_browse_service
        service = get_forced_browse_service()
        
        session_id = await service.start_scan(
            target_url=request.url,
            wordlist=request.wordlist,
            recursive=request.recursive,
            threads=request.threads,
            extensions=request.extensions
        )
        
        return {
            "status": "started",
            "session_id": session_id,
            "url": request.url,
            "wordlist": request.wordlist
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to start forced browse: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/forcedBrowse/stop")
async def stop_forced_browse(
    session_id: str = Query(..., description="Session ID to stop"),
    current_user: User = Depends(get_current_active_user)
):
    """Stop the current forced browse scan."""
    try:
        from backend.services.forced_browse_service import get_forced_browse_service
        service = get_forced_browse_service()
        success = service.stop_scan(session_id)
        return {"status": "stopped" if success else "not_running", "session_id": session_id}
    except Exception as e:
        logger.error(f"Failed to stop forced browse: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/forcedBrowse/pause")
async def pause_forced_browse(
    session_id: str = Query(..., description="Session ID to pause"),
    current_user: User = Depends(get_current_active_user)
):
    """Pause the current forced browse scan."""
    try:
        from backend.services.forced_browse_service import get_forced_browse_service
        service = get_forced_browse_service()
        success = service.pause_scan(session_id)
        return {"status": "paused" if success else "not_running", "session_id": session_id}
    except Exception as e:
        logger.error(f"Failed to pause forced browse: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/forcedBrowse/resume")
async def resume_forced_browse(
    session_id: str = Query(..., description="Session ID to resume"),
    current_user: User = Depends(get_current_active_user)
):
    """Resume a paused forced browse scan."""
    try:
        from backend.services.forced_browse_service import get_forced_browse_service
        service = get_forced_browse_service()
        success = service.resume_scan(session_id)
        return {"status": "resumed" if success else "not_paused", "session_id": session_id}
    except Exception as e:
        logger.error(f"Failed to resume forced browse: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/forcedBrowse/wordlists")
async def list_wordlists(
    current_user: User = Depends(get_current_active_user)
):
    """List available wordlists for forced browsing from local storage."""
    try:
        from backend.services.forced_browse_service import get_forced_browse_service
        service = get_forced_browse_service()
        wordlists = service.get_available_wordlists()
        
        # Return both detailed list and simple list for compatibility
        simple_list = [w["name"] for w in wordlists]
        default_file = "directories_comprehensive.txt" if "directories_comprehensive.txt" in simple_list else (simple_list[0] if simple_list else "")
        
        return {
            "wordlists": simple_list,
            "default": default_file,
            "details": wordlists
        }
    except Exception as e:
        logger.error(f"Failed to list wordlists: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/forcedBrowse/wordlists")
async def add_custom_wordlist(
    request: ForcedBrowseWordlistRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Add a custom wordlist file for forced browsing."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.forcedBrowse_add_custom_file(request.file_path)
            return {"status": "added", "file_path": request.file_path, "result": result}
    except Exception as e:
        logger.error(f"Failed to add custom wordlist: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/forcedBrowse/options")
async def set_forced_browse_options(
    request: ForcedBrowseOptionsRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Configure forced browse options."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            results = []
            
            if request.threads is not None:
                result = await client.forcedBrowse_set_option_threads(request.threads)
                results.append({"threads": result})
            
            if request.recursive is not None:
                result = await client.forcedBrowse_set_option_recursive(request.recursive)
                results.append({"recursive": result})
            
            if request.fail_case_string is not None:
                result = await client.forcedBrowse_set_option_fail_case_string(request.fail_case_string)
                results.append({"fail_case_string": result})
            
            return {"status": "updated", "results": results}
    except Exception as e:
        logger.error(f"Failed to set forced browse options: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# SCRIPT CONSOLE ENDPOINTS
# =============================================================================

class LoadScriptRequest(BaseModel):
    """Request to load a script."""
    name: str = Field(..., description="Script name")
    type: str = Field(..., description="Script type (standalone, proxy, active, passive, etc.)")
    engine: str = Field(..., description="Script engine (ECMAScript, Graal.js, etc.)")
    file_path: str = Field(..., description="Path to script file")
    description: Optional[str] = Field(default="", description="Script description")


class RunScriptRequest(BaseModel):
    """Request to run a standalone script."""
    name: str = Field(..., description="Script name to run")


class ScriptVarRequest(BaseModel):
    """Request to set a script variable."""
    key: str = Field(..., description="Variable key")
    value: str = Field(..., description="Variable value")


@router.get("/scripts")
async def list_scripts(
    current_user: User = Depends(get_current_active_user)
):
    """List all available ZAP scripts (local templates + loaded scripts)."""
    import os
    try:
        # Serve local ZAP script templates
        scripts_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "zap_scripts")
        local_scripts = []
        
        if os.path.exists(scripts_dir):
            for f in os.listdir(scripts_dir):
                if f.endswith(('.js', '.py', '.zst')):
                    script_path = os.path.join(scripts_dir, f)
                    local_scripts.append({
                        "name": f,
                        "type": "template",
                        "engine": "ECMAScript" if f.endswith('.js') else "python" if f.endswith('.py') else "zest",
                        "enabled": False,
                        "description": f"Local template: {f}",
                        "source": "local"
                    })
        
        # Also try to get scripts loaded in ZAP
        zap_scripts = []
        try:
            scanner = get_zap_scanner()
            async with scanner.client() as client:
                zap_scripts = await client.script_list()
                # Mark ZAP scripts with source
                for s in zap_scripts:
                    s["source"] = "zap"
        except Exception:
            pass  # ZAP might not be available
        
        return {"scripts": local_scripts + zap_scripts}
    except Exception as e:
        logger.error(f"Failed to list scripts: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/scripts/engines")
async def list_script_engines(
    current_user: User = Depends(get_current_active_user)
):
    """List available script engines (JavaScript, Python, etc.)."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            engines = await client.script_list_engines()
            return {"engines": engines}
    except Exception as e:
        logger.error(f"Failed to list script engines: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/scripts/types")
async def list_script_types(
    current_user: User = Depends(get_current_active_user)
):
    """List available script types (standalone, proxy, active, passive, etc.)."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            types = await client.script_list_types()
            return {"types": types}
    except Exception as e:
        logger.error(f"Failed to list script types: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/scripts/load")
async def load_script(
    request: LoadScriptRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Load a script from file."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.script_load(
                script_name=request.name,
                script_type=request.type,
                script_engine=request.engine,
                file_name=request.file_path,
                description=request.description or ""
            )
            return {"status": "loaded", "name": request.name, "result": result}
    except Exception as e:
        logger.error(f"Failed to load script: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/scripts/{script_name}/enable")
async def enable_script(
    script_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Enable a script."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.script_enable(script_name)
            return {"status": "enabled", "name": script_name, "result": result}
    except Exception as e:
        logger.error(f"Failed to enable script: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/scripts/{script_name}/disable")
async def disable_script(
    script_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Disable a script."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.script_disable(script_name)
            return {"status": "disabled", "name": script_name, "result": result}
    except Exception as e:
        logger.error(f"Failed to disable script: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/scripts/{script_name}/run")
async def run_script(
    script_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Run a standalone script."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.script_run_standalone(script_name)
            return {"status": "executed", "name": script_name, "result": result}
    except Exception as e:
        logger.error(f"Failed to run script: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/scripts/{script_name}")
async def remove_script(
    script_name: str,
    current_user: User = Depends(get_current_active_user)
):
    """Remove a script."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.script_remove(script_name)
            return {"status": "removed", "name": script_name, "result": result}
    except Exception as e:
        logger.error(f"Failed to remove script: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/scripts/vars")
async def list_script_variables(
    current_user: User = Depends(get_current_active_user)
):
    """List all global script variables."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            global_vars = await client.script_global_vars()
            custom_vars = await client.script_global_custom_vars()
            return {"global_vars": global_vars, "custom_vars": custom_vars}
    except Exception as e:
        logger.error(f"Failed to list script variables: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/scripts/vars")
async def set_script_variable(
    request: ScriptVarRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Set a global script variable."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.script_set_global_var(request.key, request.value)
            return {"status": "set", "key": request.key, "result": result}
    except Exception as e:
        logger.error(f"Failed to set script variable: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/scripts/vars/{var_key}")
async def clear_script_variable(
    var_key: str,
    current_user: User = Depends(get_current_active_user)
):
    """Clear a global script variable."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.script_clear_global_var(var_key)
            return {"status": "cleared", "key": var_key, "result": result}
    except Exception as e:
        logger.error(f"Failed to clear script variable: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/scripts/vars")
async def clear_all_script_variables(
    current_user: User = Depends(get_current_active_user)
):
    """Clear all global script variables."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.script_clear_global_vars()
            return {"status": "cleared_all", "result": result}
    except Exception as e:
        logger.error(f"Failed to clear all script variables: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# PASSIVE SCAN RULE CONFIGURATION ENDPOINTS
# =============================================================================

class SetPscanThresholdRequest(BaseModel):
    """Request to set passive scanner alert threshold."""
    scanner_id: int = Field(..., description="Scanner ID")
    threshold: str = Field(..., description="Threshold: OFF, DEFAULT, LOW, MEDIUM, HIGH")


class SetPscanOptionsRequest(BaseModel):
    """Request to set passive scan options."""
    max_alerts_per_rule: Optional[int] = Field(default=None, description="Max alerts per rule")
    scan_only_in_scope: Optional[bool] = Field(default=None, description="Scan only in scope")


@router.get("/pscan/scanners")
async def list_passive_scanners(
    current_user: User = Depends(get_current_active_user)
):
    """List all passive scan rules/scanners."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            scanners = await client.pscan_scanners()
            return {"scanners": scanners}
    except Exception as e:
        logger.error(f"Failed to list passive scanners: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pscan/status")
async def get_passive_scan_status(
    current_user: User = Depends(get_current_active_user)
):
    """Get passive scan status including queue length and current rule."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            records = await client.pscan_records_to_scan()
            current_rule = await client.pscan_current_rule()
            max_alerts = await client.pscan_max_alerts_per_rule()
            only_in_scope = await client.pscan_scan_only_in_scope()
            return {
                "records_to_scan": records,
                "current_rule": current_rule,
                "max_alerts_per_rule": max_alerts,
                "scan_only_in_scope": only_in_scope
            }
    except Exception as e:
        logger.error(f"Failed to get passive scan status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pscan/scanners/{scanner_id}/enable")
async def enable_passive_scanner(
    scanner_id: int,
    current_user: User = Depends(get_current_active_user)
):
    """Enable a specific passive scanner."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.pscan_enable_scanners(str(scanner_id))
            return {"status": "enabled", "scanner_id": scanner_id, "result": result}
    except Exception as e:
        logger.error(f"Failed to enable passive scanner: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pscan/scanners/{scanner_id}/disable")
async def disable_passive_scanner(
    scanner_id: int,
    current_user: User = Depends(get_current_active_user)
):
    """Disable a specific passive scanner."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.pscan_disable_scanners(str(scanner_id))
            return {"status": "disabled", "scanner_id": scanner_id, "result": result}
    except Exception as e:
        logger.error(f"Failed to disable passive scanner: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pscan/scanners/enable-all")
async def enable_all_passive_scanners(
    current_user: User = Depends(get_current_active_user)
):
    """Enable all passive scanners."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.pscan_enable_all_scanners()
            return {"status": "enabled_all", "result": result}
    except Exception as e:
        logger.error(f"Failed to enable all passive scanners: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pscan/scanners/disable-all")
async def disable_all_passive_scanners(
    current_user: User = Depends(get_current_active_user)
):
    """Disable all passive scanners."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.pscan_disable_all_scanners()
            return {"status": "disabled_all", "result": result}
    except Exception as e:
        logger.error(f"Failed to disable all passive scanners: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pscan/scanners/{scanner_id}/threshold")
async def set_passive_scanner_threshold(
    scanner_id: int,
    threshold: str = Query(..., description="Threshold: OFF, DEFAULT, LOW, MEDIUM, HIGH"),
    current_user: User = Depends(get_current_active_user)
):
    """Set alert threshold for a passive scanner."""
    try:
        if threshold.upper() not in ["OFF", "DEFAULT", "LOW", "MEDIUM", "HIGH"]:
            raise HTTPException(status_code=400, detail="Invalid threshold value")
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.pscan_set_scanner_alert_threshold(scanner_id, threshold.upper())
            return {"status": "threshold_set", "scanner_id": scanner_id, "threshold": threshold.upper(), "result": result}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to set passive scanner threshold: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/pscan/options")
async def set_passive_scan_options(
    request: SetPscanOptionsRequest,
    current_user: User = Depends(get_current_active_user)
):
    """Set passive scan options."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            results = {}
            if request.max_alerts_per_rule is not None:
                results["max_alerts"] = await client.pscan_set_max_alerts_per_rule(request.max_alerts_per_rule)
            if request.scan_only_in_scope is not None:
                results["only_in_scope"] = await client.pscan_set_scan_only_in_scope(request.scan_only_in_scope)
            return {"status": "options_set", "results": results}
    except Exception as e:
        logger.error(f"Failed to set passive scan options: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pscan/clear-queue")
async def clear_passive_scan_queue(
    current_user: User = Depends(get_current_active_user)
):
    """Clear the passive scan queue."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.pscan_clear_queue()
            return {"status": "queue_cleared", "result": result}
    except Exception as e:
        logger.error(f"Failed to clear passive scan queue: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# STATISTICS & PROGRESS DASHBOARD ENDPOINTS
# =============================================================================

@router.get("/stats/overview")
async def get_stats_overview(
    current_user: User = Depends(get_current_active_user)
):
    """Get comprehensive scan statistics overview."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            # Gather all statistics
            hosts = await client.core_hosts()
            sites = await client.core_sites()
            num_messages = await client.core_number_of_messages()
            mode = await client.core_mode()
            version = await client.core_version()
            
            # Alert counts by risk
            total_alerts = await client.core_number_of_alerts()
            high_alerts = await client.core_number_of_alerts(risk_id=3)
            medium_alerts = await client.core_number_of_alerts(risk_id=2)
            low_alerts = await client.core_number_of_alerts(risk_id=1)
            info_alerts = await client.core_number_of_alerts(risk_id=0)
            
            # Passive scan queue
            pscan_queue = await client.pscan_records_to_scan()
            
            return {
                "zap_version": version,
                "mode": mode,
                "hosts_count": len(hosts),
                "sites_count": len(sites),
                "total_messages": num_messages,
                "alerts": {
                    "total": total_alerts,
                    "high": high_alerts,
                    "medium": medium_alerts,
                    "low": low_alerts,
                    "info": info_alerts
                },
                "passive_scan_queue": pscan_queue,
                "hosts": hosts[:20],  # Limit for performance
                "sites": sites[:20]   # Limit for performance
            }
    except Exception as e:
        logger.error(f"Failed to get stats overview: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/stats/sites")
async def get_site_stats(
    site: Optional[str] = Query(default=None, description="Specific site URL"),
    in_scope: bool = Query(default=False, description="Only in-scope sites"),
    current_user: User = Depends(get_current_active_user)
):
    """Get statistics for sites."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            if site:
                stats = await client.stats_site_stats(site=site, in_scope=in_scope)
                return {"site": site, "stats": stats}
            else:
                all_stats = await client.stats_all_sites_stats()
                return {"sites_stats": all_stats}
    except Exception as e:
        logger.error(f"Failed to get site stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/stats/scans")
async def get_scan_stats(
    current_user: User = Depends(get_current_active_user)
):
    """Get active and spider scan statistics."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            active_scans = await client.ascan_scans()
            spider_status = await client.spider_status()
            
            return {
                "active_scans": active_scans,
                "spider_progress": spider_status
            }
    except Exception as e:
        logger.error(f"Failed to get scan stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/stats/urls")
async def get_discovered_urls(
    base_url: Optional[str] = Query(default=None, description="Filter by base URL"),
    limit: int = Query(default=100, ge=1, le=1000, description="Max URLs to return"),
    current_user: User = Depends(get_current_active_user)
):
    """Get discovered URLs."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            urls = await client.core_urls(base_url=base_url)
            return {
                "total_urls": len(urls),
                "urls": urls[:limit]
            }
    except Exception as e:
        logger.error(f"Failed to get discovered URLs: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/stats/clear")
async def clear_stats(
    site: Optional[str] = Query(default=None, description="Clear stats for specific site only"),
    current_user: User = Depends(get_current_active_user)
):
    """Clear statistics."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            result = await client.stats_clear(site=site)
            return {"status": "cleared", "site": site or "all", "result": result}
    except Exception as e:
        logger.error(f"Failed to clear stats: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/stats/progress/{scan_id}")
async def get_scan_progress(
    scan_id: str,
    current_user: User = Depends(get_current_active_user)
):
    """Get detailed progress for a specific scan."""
    try:
        scanner = get_zap_scanner()
        async with scanner.client() as client:
            spider_progress = await client.spider_status(scan_id)
            spider_urls = await client.spider_results(scan_id)
            active_progress = await client.ascan_status(scan_id)
            active_alerts = await client.ascan_alerts_ids(scan_id)
            
            return {
                "scan_id": scan_id,
                "spider": {
                    "progress": spider_progress,
                    "urls_found": len(spider_urls)
                },
                "active_scan": {
                    "progress": active_progress,
                    "alerts_found": len(active_alerts)
                }
            }
    except Exception as e:
        logger.error(f"Failed to get scan progress: {e}")
        raise HTTPException(status_code=500, detail=str(e))

