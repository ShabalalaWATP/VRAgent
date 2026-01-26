"""
API Endpoint Tester Service

Comprehensive API security testing including:
- Authentication testing (missing auth, weak auth, auth bypass)
- Rate limiting detection
- Input validation testing (SQLi, XSS, command injection payloads)
- HTTP method testing (verb tampering)
- CORS configuration analysis
- Security header analysis
- Response analysis (sensitive data exposure, error leakage)
- GraphQL introspection testing
"""

import asyncio
import httpx
import json
import re
import time
import logging
from dataclasses import dataclass, field, asdict
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlparse, urljoin, parse_qs
from enum import Enum

logger = logging.getLogger(__name__)


class Severity(str, Enum):
    CRITICAL = "critical"
    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"
    INFO = "info"


class TestCategory(str, Enum):
    AUTHENTICATION = "authentication"
    AUTHORIZATION = "authorization"
    INPUT_VALIDATION = "input_validation"
    RATE_LIMITING = "rate_limiting"
    CORS = "cors"
    HEADERS = "headers"
    INFORMATION_DISCLOSURE = "information_disclosure"
    HTTP_METHODS = "http_methods"
    GRAPHQL = "graphql"
    GENERAL = "general"


@dataclass
class Finding:
    """A security finding from API testing."""
    title: str
    description: str
    severity: str
    category: str
    evidence: str = ""
    remediation: str = ""
    cwe: Optional[str] = None
    endpoint: str = ""
    owasp_api: Optional[str] = None  # OWASP API Security Top 10 mapping
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        return result


@dataclass
class EndpointResult:
    """Result of testing a single endpoint."""
    url: str
    method: str
    status_code: Optional[int] = None
    response_time_ms: float = 0
    content_type: Optional[str] = None
    response_size: int = 0
    headers: Dict[str, str] = field(default_factory=dict)
    findings: List[Finding] = field(default_factory=list)
    error: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "url": self.url,
            "method": self.method,
            "status_code": self.status_code,
            "response_time_ms": self.response_time_ms,
            "content_type": self.content_type,
            "response_size": self.response_size,
            "headers": self.headers,
            "findings": [f.to_dict() for f in self.findings],
            "error": self.error,
        }


@dataclass 
class APITestResult:
    """Complete result of API testing."""
    base_url: str
    endpoints_tested: int = 0
    total_findings: int = 0
    critical_count: int = 0
    high_count: int = 0
    medium_count: int = 0
    low_count: int = 0
    info_count: int = 0
    test_duration_seconds: float = 0
    endpoint_results: List[EndpointResult] = field(default_factory=list)
    all_findings: List[Finding] = field(default_factory=list)
    security_score: int = 100
    summary: str = ""
    error: Optional[str] = None
    owasp_api_breakdown: Dict[str, int] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        # Auto-calculate owasp breakdown if not set
        owasp_breakdown = self.owasp_api_breakdown or self.get_owasp_breakdown()
        return {
            "base_url": self.base_url,
            "endpoints_tested": self.endpoints_tested,
            "total_findings": self.total_findings,
            "critical_count": self.critical_count,
            "high_count": self.high_count,
            "medium_count": self.medium_count,
            "low_count": self.low_count,
            "info_count": self.info_count,
            "test_duration_seconds": self.test_duration_seconds,
            "endpoint_results": [e.to_dict() for e in self.endpoint_results],
            "all_findings": [f.to_dict() for f in self.all_findings],
            "security_score": self.security_score,
            "summary": self.summary,
            "error": self.error,
            "owasp_api_breakdown": owasp_breakdown,
        }
    
    def get_owasp_breakdown(self) -> Dict[str, int]:
        """Get count of findings by OWASP API Top 10 category."""
        breakdown = {}
        for finding in self.all_findings:
            if hasattr(finding, 'owasp_api') and finding.owasp_api:
                owasp = finding.owasp_api
                breakdown[owasp] = breakdown.get(owasp, 0) + 1
        return breakdown


# OWASP API Security Top 10 (2023) mapping
OWASP_API_TOP_10 = {
    "API1:2023": {
        "name": "Broken Object Level Authorization",
        "description": "APIs expose endpoints that handle object identifiers, creating a wide attack surface of Object Level Access Control issues.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa1-broken-object-level-authorization/",
    },
    "API2:2023": {
        "name": "Broken Authentication",
        "description": "Authentication mechanisms are often implemented incorrectly, allowing attackers to compromise authentication tokens.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa2-broken-authentication/",
    },
    "API3:2023": {
        "name": "Broken Object Property Level Authorization",
        "description": "APIs expose endpoints that return all object properties without considering sensitivity level.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa3-broken-object-property-level-authorization/",
    },
    "API4:2023": {
        "name": "Unrestricted Resource Consumption",
        "description": "APIs do not limit the number/size of resources that can be requested, leading to DoS or cost spikes.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa4-unrestricted-resource-consumption/",
    },
    "API5:2023": {
        "name": "Broken Function Level Authorization",
        "description": "Complex access control policies with hierarchies and roles are not properly implemented.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa5-broken-function-level-authorization/",
    },
    "API6:2023": {
        "name": "Unrestricted Access to Sensitive Business Flows",
        "description": "APIs vulnerable to abuse of business logic if not protected against automated threats.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa6-unrestricted-access-to-sensitive-business-flows/",
    },
    "API7:2023": {
        "name": "Server Side Request Forgery",
        "description": "SSRF flaws occur when an API fetches a remote resource without validating the user-supplied URI.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa7-server-side-request-forgery/",
    },
    "API8:2023": {
        "name": "Security Misconfiguration",
        "description": "Misconfigured security settings are common and often stem from insecure defaults or incomplete configurations.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa8-security-misconfiguration/",
    },
    "API9:2023": {
        "name": "Improper Inventory Management",
        "description": "APIs tend to expose more endpoints than traditional web apps, making proper documentation important.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xa9-improper-inventory-management/",
    },
    "API10:2023": {
        "name": "Unsafe Consumption of APIs",
        "description": "Developers trust data received from third-party APIs more than user input without proper validation.",
        "url": "https://owasp.org/API-Security/editions/2023/en/0xaa-unsafe-consumption-of-apis/",
    },
}

# Map CWE to OWASP API Top 10
CWE_TO_OWASP_API = {
    "CWE-284": "API1:2023",  # Improper Access Control -> BOLA
    "CWE-285": "API1:2023",  # Improper Authorization -> BOLA
    "CWE-639": "API1:2023",  # Authorization Bypass -> BOLA
    "CWE-287": "API2:2023",  # Improper Authentication -> Broken Auth
    "CWE-306": "API2:2023",  # Missing Authentication -> Broken Auth
    "CWE-307": "API2:2023",  # Improper Restriction of Auth Attempts -> Broken Auth
    "CWE-200": "API3:2023",  # Exposure of Sensitive Info -> Broken Property Auth
    "CWE-359": "API3:2023",  # Exposure of Private Info -> Broken Property Auth
    "CWE-770": "API4:2023",  # Allocation of Resources Without Limits -> Unrestricted Resource
    "CWE-400": "API4:2023",  # Uncontrolled Resource Consumption -> Unrestricted Resource
    "CWE-269": "API5:2023",  # Improper Privilege Management -> Broken Function Auth
    "CWE-650": "API5:2023",  # Trusting HTTP Methods -> Broken Function Auth
    "CWE-799": "API6:2023",  # Improper Control of Interaction Frequency
    "CWE-918": "API7:2023",  # Server-Side Request Forgery
    "CWE-16": "API8:2023",   # Configuration -> Security Misconfiguration
    "CWE-319": "API8:2023",  # Cleartext Transmission -> Security Misconfiguration
    "CWE-693": "API8:2023",  # Protection Mechanism Failure -> Security Misconfiguration
    "CWE-942": "API8:2023",  # CORS Misconfiguration -> Security Misconfiguration
    "CWE-1021": "API8:2023", # Improper Restriction of Rendered UI -> Security Misconfiguration
    "CWE-525": "API8:2023",  # Browser Cache -> Security Misconfiguration
    "CWE-749": "API8:2023",  # Exposed Dangerous Method -> Security Misconfiguration
    "CWE-79": "API8:2023",   # XSS -> Security Misconfiguration
    "CWE-209": "API9:2023",  # Error Message Info Exposure -> Improper Inventory
    "CWE-89": "API8:2023",   # SQL Injection -> Security Misconfiguration
}


def get_owasp_api_mapping(cwe: Optional[str]) -> Optional[str]:
    """Map a CWE to OWASP API Top 10 category."""
    if not cwe:
        return None
    return CWE_TO_OWASP_API.get(cwe)


def create_finding(
    title: str,
    description: str,
    severity: str,
    category: str,
    evidence: str = "",
    remediation: str = "",
    cwe: Optional[str] = None,
    endpoint: str = "",
) -> Finding:
    """Create a Finding with automatic OWASP API mapping."""
    return Finding(
        title=title,
        description=description,
        severity=severity,
        category=category,
        evidence=evidence,
        remediation=remediation,
        cwe=cwe,
        endpoint=endpoint,
        owasp_api=get_owasp_api_mapping(cwe),
    )


# Security headers to check
SECURITY_HEADERS = {
    "Strict-Transport-Security": {
        "description": "HSTS - Forces HTTPS connections",
        "severity": Severity.MEDIUM,
        "remediation": "Add 'Strict-Transport-Security: max-age=31536000; includeSubDomains' header",
        "cwe": "CWE-319",
    },
    "X-Content-Type-Options": {
        "description": "Prevents MIME type sniffing",
        "severity": Severity.LOW,
        "remediation": "Add 'X-Content-Type-Options: nosniff' header",
        "cwe": "CWE-693",
    },
    "X-Frame-Options": {
        "description": "Prevents clickjacking attacks",
        "severity": Severity.MEDIUM,
        "remediation": "Add 'X-Frame-Options: DENY' or 'SAMEORIGIN' header",
        "cwe": "CWE-1021",
    },
    "Content-Security-Policy": {
        "description": "Controls resources the browser can load",
        "severity": Severity.MEDIUM,
        "remediation": "Implement a Content-Security-Policy header appropriate for your application",
        "cwe": "CWE-693",
    },
    "X-XSS-Protection": {
        "description": "Legacy XSS filter (use CSP instead)",
        "severity": Severity.INFO,
        "remediation": "Consider 'X-XSS-Protection: 0' if CSP is implemented, or '1; mode=block' otherwise",
        "cwe": "CWE-79",
    },
    "Referrer-Policy": {
        "description": "Controls referrer information sent with requests",
        "severity": Severity.LOW,
        "remediation": "Add 'Referrer-Policy: strict-origin-when-cross-origin' or more restrictive",
        "cwe": "CWE-200",
    },
    "Permissions-Policy": {
        "description": "Controls browser features and APIs",
        "severity": Severity.LOW,
        "remediation": "Add Permissions-Policy header to restrict unnecessary browser features",
        "cwe": "CWE-693",
    },
    "Cache-Control": {
        "description": "Controls caching of sensitive responses",
        "severity": Severity.LOW,
        "remediation": "For sensitive APIs, use 'Cache-Control: no-store, private'",
        "cwe": "CWE-525",
    },
}

# Sensitive data patterns to detect in responses
SENSITIVE_PATTERNS = [
    (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', "Email address", Severity.LOW),
    (r'\b\d{3}-\d{2}-\d{4}\b', "SSN pattern", Severity.CRITICAL),
    (r'\b\d{16}\b', "Credit card number pattern", Severity.CRITICAL),
    (r'\b(?:api[_-]?key|apikey|api[_-]?secret)\s*[=:]\s*["\']?[\w-]{20,}', "API key exposure", Severity.HIGH),
    (r'\b(?:password|passwd|pwd)\s*[=:]\s*["\']?[^\s"\']{4,}', "Password in response", Severity.CRITICAL),
    (r'\b(?:secret|token|bearer)\s*[=:]\s*["\']?[\w-]{10,}', "Secret/token exposure", Severity.HIGH),
    (r'-----BEGIN (?:RSA |DSA |EC )?PRIVATE KEY-----', "Private key exposure", Severity.CRITICAL),
    (r'\beyJ[a-zA-Z0-9_-]*\.eyJ[a-zA-Z0-9_-]*\.[a-zA-Z0-9_-]*\b', "JWT token in response", Severity.MEDIUM),
    (r'(?:mysql|postgresql|mongodb|redis)://[^\s<>"]+', "Database connection string", Severity.CRITICAL),
    (r'\b(?:aws_access_key_id|aws_secret_access_key)\s*[=:]\s*[\w/+=]+', "AWS credentials", Severity.CRITICAL),
]

# SQL injection test payloads
SQLI_PAYLOADS = [
    "' OR '1'='1",
    "' OR '1'='1' --",
    "1' OR '1'='1",
    "' UNION SELECT NULL--",
    "'; DROP TABLE users--",
    "1; SELECT * FROM users",
    "' AND 1=1--",
    "' AND 1=2--",
]

# XSS test payloads
XSS_PAYLOADS = [
    "<script>alert('XSS')</script>",
    "<img src=x onerror=alert('XSS')>",
    "javascript:alert('XSS')",
    "<svg/onload=alert('XSS')>",
    "'\"><script>alert('XSS')</script>",
    "<body onload=alert('XSS')>",
]

# Command injection payloads
COMMAND_INJECTION_PAYLOADS = [
    "; ls -la",
    "| cat /etc/passwd",
    "& whoami",
    "`id`",
    "$(whoami)",
    "; ping -c 1 127.0.0.1",
]

# Path traversal payloads
PATH_TRAVERSAL_PAYLOADS = [
    "../../../etc/passwd",
    "..\\..\\..\\windows\\system32\\config\\sam",
    "....//....//....//etc/passwd",
    "%2e%2e%2f%2e%2e%2f%2e%2e%2fetc%2fpasswd",
    "..%252f..%252f..%252fetc/passwd",
]

# Error patterns indicating vulnerabilities
ERROR_PATTERNS = [
    (r'SQL syntax.*?MySQL', "MySQL error", "sql_injection"),
    (r'Warning.*?\Wmysqli?_', "MySQL error", "sql_injection"),
    (r'PostgreSQL.*?ERROR', "PostgreSQL error", "sql_injection"),
    (r'ORA-\d{5}', "Oracle error", "sql_injection"),
    (r'Microsoft SQL Server', "MSSQL error", "sql_injection"),
    (r'sqlite3\.OperationalError', "SQLite error", "sql_injection"),
    (r'Traceback \(most recent call last\)', "Python traceback", "information_disclosure"),
    (r'at \w+\.\w+\([\w\.]+:\d+\)', "Java stack trace", "information_disclosure"),
    (r'System\.(\w+)?Exception', ".NET exception", "information_disclosure"),
    (r'PHP (?:Parse|Fatal) error', "PHP error", "information_disclosure"),
    (r'<b>Warning</b>:.*?on line <b>\d+</b>', "PHP warning", "information_disclosure"),
]

# HTTP methods to test
HTTP_METHODS = ["GET", "POST", "PUT", "DELETE", "PATCH", "OPTIONS", "HEAD", "TRACE", "CONNECT"]

# Dangerous methods
DANGEROUS_METHODS = ["TRACE", "CONNECT", "DELETE", "PUT", "PATCH"]


async def test_endpoint(
    client: httpx.AsyncClient,
    url: str,
    method: str = "GET",
    headers: Optional[Dict[str, str]] = None,
    body: Optional[Any] = None,
    auth_header: Optional[str] = None,
) -> EndpointResult:
    """Test a single endpoint and return results."""
    result = EndpointResult(url=url, method=method)
    request_headers = headers or {}
    
    if auth_header:
        request_headers["Authorization"] = auth_header
    
    try:
        start_time = time.time()
        
        kwargs = {"headers": request_headers, "timeout": 30.0}
        if body and method in ["POST", "PUT", "PATCH"]:
            if isinstance(body, dict):
                kwargs["json"] = body
            else:
                kwargs["content"] = body
        
        response = await client.request(method, url, **kwargs)
        
        result.response_time_ms = (time.time() - start_time) * 1000
        result.status_code = response.status_code
        result.content_type = response.headers.get("content-type", "")
        result.response_size = len(response.content)
        result.headers = dict(response.headers)
        
    except httpx.TimeoutException:
        result.error = "Request timed out"
    except httpx.ConnectError as e:
        result.error = f"Connection failed: {str(e)}"
    except Exception as e:
        result.error = f"Request failed: {str(e)}"
    
    return result


async def check_security_headers(result: EndpointResult) -> List[Finding]:
    """Check for missing or misconfigured security headers."""
    findings = []
    headers_lower = {k.lower(): v for k, v in result.headers.items()}
    
    for header, config in SECURITY_HEADERS.items():
        header_lower = header.lower()
        if header_lower not in headers_lower:
            findings.append(create_finding(
                title=f"Missing Security Header: {header}",
                description=config["description"],
                severity=config["severity"].value,
                category=TestCategory.HEADERS.value,
                evidence=f"Header '{header}' not present in response",
                remediation=config["remediation"],
                cwe=config["cwe"],
                endpoint=result.url,
            ))
    
    # Check for dangerous headers that leak info
    dangerous_headers = {
        "server": "Server version disclosure",
        "x-powered-by": "Technology stack disclosure",
        "x-aspnet-version": "ASP.NET version disclosure",
        "x-aspnetmvc-version": "ASP.NET MVC version disclosure",
    }
    
    for header, description in dangerous_headers.items():
        if header in headers_lower:
            findings.append(create_finding(
                title=f"Information Disclosure: {header}",
                description=description,
                severity=Severity.LOW.value,
                category=TestCategory.INFORMATION_DISCLOSURE.value,
                evidence=f"{header}: {headers_lower[header]}",
                remediation=f"Remove or obfuscate the '{header}' header",
                cwe="CWE-200",
                endpoint=result.url,
            ))
    
    return findings


async def check_cors(client: httpx.AsyncClient, url: str) -> List[Finding]:
    """Test CORS configuration."""
    findings = []
    
    # Test with malicious origin
    test_origins = [
        "https://evil.com",
        "https://attacker.com",
        "null",
    ]
    
    for origin in test_origins:
        try:
            response = await client.options(
                url,
                headers={
                    "Origin": origin,
                    "Access-Control-Request-Method": "GET",
                },
                timeout=10.0,
            )
            
            acao = response.headers.get("access-control-allow-origin", "")
            acac = response.headers.get("access-control-allow-credentials", "")
            
            # Check for wildcard with credentials (critical vulnerability)
            if acao == "*" and acac.lower() == "true":
                findings.append(create_finding(
                    title="Critical CORS Misconfiguration",
                    description="CORS allows any origin with credentials - allows credential theft",
                    severity=Severity.CRITICAL.value,
                    category=TestCategory.CORS.value,
                    evidence=f"Access-Control-Allow-Origin: *, Access-Control-Allow-Credentials: true",
                    remediation="Never use wildcard origin with credentials. Whitelist specific trusted origins.",
                    cwe="CWE-942",
                    endpoint=url,
                ))
                break
            
            # Check if arbitrary origin is reflected
            if acao == origin and origin != "null":
                severity = Severity.HIGH if acac.lower() == "true" else Severity.MEDIUM
                findings.append(create_finding(
                    title="CORS Origin Reflection",
                    description=f"CORS reflects arbitrary origin '{origin}'",
                    severity=severity.value,
                    category=TestCategory.CORS.value,
                    evidence=f"Origin '{origin}' was reflected in Access-Control-Allow-Origin",
                    remediation="Implement a whitelist of allowed origins instead of reflecting the Origin header",
                    cwe="CWE-942",
                    endpoint=url,
                ))
                break
            
            # Check for null origin allowed
            if acao == "null":
                findings.append(create_finding(
                    title="CORS Allows Null Origin",
                    description="CORS accepts 'null' origin which can be exploited via sandboxed iframes",
                    severity=Severity.MEDIUM.value,
                    category=TestCategory.CORS.value,
                    evidence="Access-Control-Allow-Origin: null",
                    remediation="Do not allow 'null' as a valid origin",
                    cwe="CWE-942",
                    endpoint=url,
                ))
                break
                
        except Exception as e:
            logger.debug(f"CORS check failed for {origin}: {e}")
    
    return findings


async def check_authentication(
    client: httpx.AsyncClient,
    url: str,
    method: str = "GET",
    auth_header: Optional[str] = None,
) -> List[Finding]:
    """Test authentication requirements."""
    findings = []
    
    # Test without authentication
    try:
        response_no_auth = await client.request(method, url, timeout=10.0)
        
        # If the endpoint works without auth but we expect it to need auth
        if response_no_auth.status_code == 200:
            # Check if response contains data (not just an empty success)
            if len(response_no_auth.content) > 10:
                findings.append(create_finding(
                    title="Endpoint Accessible Without Authentication",
                    description="API endpoint returns data without requiring authentication",
                    severity=Severity.HIGH.value,
                    category=TestCategory.AUTHENTICATION.value,
                    evidence=f"HTTP {response_no_auth.status_code} with {len(response_no_auth.content)} bytes response",
                    remediation="Implement authentication for this endpoint",
                    cwe="CWE-306",
                    endpoint=url,
                ))
        
        # Check for auth bypass via method tampering
        if method == "GET" and auth_header:
            for alt_method in ["POST", "PUT", "HEAD"]:
                try:
                    alt_response = await client.request(alt_method, url, timeout=10.0)
                    if alt_response.status_code == 200 and response_no_auth.status_code in [401, 403]:
                        findings.append(create_finding(
                            title="Authentication Bypass via HTTP Method",
                            description=f"Authentication can be bypassed using {alt_method} instead of {method}",
                            severity=Severity.CRITICAL.value,
                            category=TestCategory.AUTHENTICATION.value,
                            evidence=f"{method} returns {response_no_auth.status_code}, {alt_method} returns {alt_response.status_code}",
                            remediation="Enforce authentication consistently across all HTTP methods",
                            cwe="CWE-287",
                            endpoint=url,
                        ))
                except:
                    pass
                    
    except Exception as e:
        logger.debug(f"Auth check failed: {e}")
    
    return findings


async def check_rate_limiting(
    client: httpx.AsyncClient,
    url: str,
    method: str = "GET",
    num_requests: int = 20,
) -> List[Finding]:
    """Test for rate limiting implementation."""
    findings = []
    
    try:
        responses = []
        for _ in range(num_requests):
            response = await client.request(method, url, timeout=5.0)
            responses.append(response.status_code)
            # Small delay to not overwhelm
            await asyncio.sleep(0.05)
        
        # Check if any rate limiting kicked in
        rate_limited = any(code == 429 for code in responses)
        
        if not rate_limited:
            # Check response headers for rate limit info
            last_response = await client.request(method, url, timeout=5.0)
            has_rate_headers = any(
                h.lower() in ["x-ratelimit-limit", "x-rate-limit-limit", "ratelimit-limit", "retry-after"]
                for h in last_response.headers
            )
            
            if not has_rate_headers:
                findings.append(create_finding(
                    title="No Rate Limiting Detected",
                    description=f"Endpoint accepted {num_requests} requests without rate limiting",
                    severity=Severity.MEDIUM.value,
                    category=TestCategory.RATE_LIMITING.value,
                    evidence=f"All {num_requests} requests returned success status",
                    remediation="Implement rate limiting to prevent abuse and DoS attacks",
                    cwe="CWE-770",
                    endpoint=url,
                ))
        else:
            # Rate limiting exists - info finding
            rate_limit_index = next(i for i, code in enumerate(responses) if code == 429)
            findings.append(create_finding(
                title="Rate Limiting Detected",
                description=f"Rate limiting activated after {rate_limit_index + 1} requests",
                severity=Severity.INFO.value,
                category=TestCategory.RATE_LIMITING.value,
                evidence=f"HTTP 429 returned after {rate_limit_index + 1} requests",
                remediation="Rate limiting is properly configured",
                endpoint=url,
            ))
            
    except Exception as e:
        logger.debug(f"Rate limit check failed: {e}")
    
    return findings


async def check_input_validation(
    client: httpx.AsyncClient,
    url: str,
    method: str = "GET",
    params: Optional[Dict[str, str]] = None,
) -> List[Finding]:
    """Test input validation with various payloads."""
    findings = []
    
    # Parse URL for query parameters
    parsed = urlparse(url)
    existing_params = parse_qs(parsed.query)
    
    # Get parameter names to test
    test_params = list(existing_params.keys()) if existing_params else ["id", "q", "search", "query", "name", "user"]
    
    for param in test_params[:3]:  # Limit to first 3 params
        # SQL Injection tests
        for payload in SQLI_PAYLOADS[:3]:  # Limit payloads
            try:
                test_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}?{param}={payload}"
                response = await client.request(method, test_url, timeout=10.0)
                response_text = response.text.lower()
                
                # Check for SQL error patterns
                for pattern, error_type, vuln_type in ERROR_PATTERNS:
                    if re.search(pattern, response.text, re.IGNORECASE):
                        findings.append(create_finding(
                            title=f"Potential SQL Injection ({error_type})",
                            description=f"SQL error detected when testing parameter '{param}'",
                            severity=Severity.CRITICAL.value,
                            category=TestCategory.INPUT_VALIDATION.value,
                            evidence=f"Payload: {payload}, Error pattern: {error_type}",
                            remediation="Use parameterized queries/prepared statements",
                            cwe="CWE-89",
                            endpoint=url,
                        ))
                        break
            except:
                pass
        
        # XSS tests (check if payload is reflected)
        for payload in XSS_PAYLOADS[:2]:
            try:
                test_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}?{param}={payload}"
                response = await client.request(method, test_url, timeout=10.0)
                
                if payload in response.text:
                    findings.append(create_finding(
                        title="Reflected XSS Potential",
                        description=f"Input is reflected in response without encoding for parameter '{param}'",
                        severity=Severity.HIGH.value,
                        category=TestCategory.INPUT_VALIDATION.value,
                        evidence=f"Payload '{payload[:30]}...' reflected in response",
                        remediation="Encode all user input before including in responses",
                        cwe="CWE-79",
                        endpoint=url,
                    ))
                    break
            except:
                pass
    
    return findings


async def check_http_methods(
    client: httpx.AsyncClient,
    url: str,
) -> List[Finding]:
    """Test which HTTP methods are allowed."""
    findings = []
    allowed_methods = []
    
    # First try OPTIONS to get allowed methods
    try:
        response = await client.options(url, timeout=10.0)
        allow_header = response.headers.get("allow", "")
        if allow_header:
            allowed_methods = [m.strip().upper() for m in allow_header.split(",")]
    except:
        pass
    
    # If OPTIONS didn't work, test each method
    if not allowed_methods:
        for method in HTTP_METHODS:
            try:
                response = await client.request(method, url, timeout=5.0)
                if response.status_code not in [405, 501]:
                    allowed_methods.append(method)
            except:
                pass
    
    # Check for dangerous methods
    for method in DANGEROUS_METHODS:
        if method in allowed_methods:
            severity = Severity.HIGH if method == "TRACE" else Severity.MEDIUM
            findings.append(create_finding(
                title=f"Dangerous HTTP Method Allowed: {method}",
                description=f"HTTP {method} method is enabled on this endpoint",
                severity=severity.value,
                category=TestCategory.HTTP_METHODS.value,
                evidence=f"Allowed methods: {', '.join(allowed_methods)}",
                remediation=f"Disable {method} method if not required",
                cwe="CWE-749" if method == "TRACE" else "CWE-650",
                endpoint=url,
            ))
    
    # TRACE can enable XST attacks
    if "TRACE" in allowed_methods:
        findings.append(create_finding(
            title="Cross-Site Tracing (XST) Possible",
            description="TRACE method enabled which can be used to steal cookies via XST attacks",
            severity=Severity.HIGH.value,
            category=TestCategory.HTTP_METHODS.value,
            evidence="TRACE method returns request headers in response body",
            remediation="Disable TRACE method on the web server",
            cwe="CWE-693",
            endpoint=url,
        ))
    
    return findings


async def check_sensitive_data(response_text: str, url: str) -> List[Finding]:
    """Check response for sensitive data exposure."""
    findings = []
    
    for pattern, description, severity in SENSITIVE_PATTERNS:
        matches = re.findall(pattern, response_text, re.IGNORECASE)
        if matches:
            # Redact the actual sensitive data in evidence
            redacted_matches = [m[:4] + "..." + m[-4:] if len(m) > 10 else "***" for m in matches[:3]]
            findings.append(create_finding(
                title=f"Sensitive Data Exposure: {description}",
                description=f"Response contains potentially sensitive data matching {description} pattern",
                severity=severity.value,
                category=TestCategory.INFORMATION_DISCLOSURE.value,
                evidence=f"Found {len(matches)} match(es): {redacted_matches}",
                remediation="Review and remove sensitive data from API responses",
                cwe="CWE-200",
                endpoint=url,
            ))
    
    return findings


async def check_error_handling(
    client: httpx.AsyncClient,
    url: str,
) -> List[Finding]:
    """Test error handling and information disclosure."""
    findings = []
    
    # Test with malformed requests
    test_cases = [
        (f"{url}/../../../etc/passwd", "Path traversal"),
        (f"{url}?<script>alert(1)</script>", "XSS in query"),
        (f"{url}?id=-1", "Negative ID"),
        (f"{url}?id=99999999999", "Large ID"),
        (f"{url}?id=abc", "String instead of number"),
    ]
    
    for test_url, test_type in test_cases:
        try:
            response = await client.get(test_url, timeout=10.0)
            
            # Check for detailed error messages
            for pattern, error_type, vuln_type in ERROR_PATTERNS:
                if re.search(pattern, response.text, re.IGNORECASE):
                    findings.append(create_finding(
                        title=f"Verbose Error Message: {error_type}",
                        description=f"Detailed error information exposed via {test_type} test",
                        severity=Severity.MEDIUM.value,
                        category=TestCategory.INFORMATION_DISCLOSURE.value,
                        evidence=f"Test: {test_type}, Error type: {error_type}",
                        remediation="Implement generic error messages for production environments",
                        cwe="CWE-209",
                        endpoint=url,
                    ))
                    break
        except:
            pass
    
    return findings


async def check_graphql(
    client: httpx.AsyncClient,
    url: str,
) -> List[Finding]:
    """Test GraphQL-specific vulnerabilities."""
    findings = []
    
    # Test for introspection
    introspection_query = {
        "query": """
        query IntrospectionQuery {
            __schema {
                queryType { name }
                types { name kind }
            }
        }
        """
    }
    
    try:
        response = await client.post(
            url,
            json=introspection_query,
            headers={"Content-Type": "application/json"},
            timeout=10.0,
        )
        
        if response.status_code == 200 and "__schema" in response.text:
            findings.append(create_finding(
                title="GraphQL Introspection Enabled",
                description="GraphQL introspection is enabled, exposing the entire schema",
                severity=Severity.MEDIUM.value,
                category=TestCategory.GRAPHQL.value,
                evidence="Introspection query returned schema information",
                remediation="Disable introspection in production environments",
                cwe="CWE-200",
                endpoint=url,
            ))
            
            # Check if mutations are exposed
            if "Mutation" in response.text:
                findings.append(create_finding(
                    title="GraphQL Mutations Exposed",
                    description="GraphQL mutations are discoverable via introspection",
                    severity=Severity.MEDIUM.value,
                    category=TestCategory.GRAPHQL.value,
                    evidence="Mutation type found in schema",
                    remediation="Review and secure all mutation operations",
                    cwe="CWE-200",
                    endpoint=url,
                ))
    except:
        pass
    
    # Test for batching attacks
    batch_query = [
        {"query": "{ __typename }"},
        {"query": "{ __typename }"},
        {"query": "{ __typename }"},
    ]
    
    try:
        response = await client.post(
            url,
            json=batch_query,
            headers={"Content-Type": "application/json"},
            timeout=10.0,
        )
        
        if response.status_code == 200:
            try:
                data = response.json()
                if isinstance(data, list) and len(data) == 3:
                    findings.append(create_finding(
                        title="GraphQL Batch Queries Allowed",
                        description="GraphQL endpoint accepts batched queries which can be abused for DoS",
                        severity=Severity.LOW.value,
                        category=TestCategory.GRAPHQL.value,
                        evidence="Batch query with 3 operations was accepted",
                        remediation="Implement query complexity limits and batch size restrictions",
                        cwe="CWE-770",
                        endpoint=url,
                    ))
            except:
                pass
    except:
        pass
    
    return findings


def calculate_security_score(findings: List[Finding]) -> int:
    """Calculate a security score based on findings."""
    score = 100
    
    severity_penalties = {
        Severity.CRITICAL.value: 25,
        Severity.HIGH.value: 15,
        Severity.MEDIUM.value: 8,
        Severity.LOW.value: 3,
        Severity.INFO.value: 0,
    }
    
    for finding in findings:
        penalty = severity_penalties.get(finding.severity, 0)
        score -= penalty
    
    return max(0, score)


async def test_api_endpoints(
    base_url: str,
    endpoints: List[Dict[str, Any]],
    auth_header: Optional[str] = None,
    test_auth: bool = True,
    test_cors: bool = True,
    test_rate_limit: bool = True,
    test_input_validation: bool = True,
    test_methods: bool = True,
    test_graphql: bool = False,
    proxy_url: Optional[str] = None,
    timeout: float = 30.0,
    verify_ssl: bool = True,
) -> APITestResult:
    """
    Main function to test API endpoints.

    Args:
        base_url: Base URL of the API
        endpoints: List of endpoint configs with url, method, params, body
        auth_header: Optional authorization header value
        test_auth: Whether to test authentication
        test_cors: Whether to test CORS configuration
        test_rate_limit: Whether to test rate limiting
        test_input_validation: Whether to test input validation
        test_methods: Whether to test HTTP methods
        test_graphql: Whether to run GraphQL-specific tests
        proxy_url: Optional HTTP/HTTPS proxy URL (e.g., http://proxy:8080)
        timeout: Request timeout in seconds (default 30)
        verify_ssl: Whether to verify SSL certificates (default True, disable only for self-signed certs)

    Returns:
        APITestResult with all findings
    """
    result = APITestResult(base_url=base_url)
    start_time = time.time()
    
    # Validate base URL
    try:
        parsed = urlparse(base_url)
        if not parsed.scheme or not parsed.netloc:
            result.error = "Invalid base URL format"
            return result
    except:
        result.error = "Failed to parse base URL"
        return result
    
    # Configure proxy if provided
    proxy_config = None
    if proxy_url:
        proxy_config = {
            "http://": proxy_url,
            "https://": proxy_url,
        }

    # Warn if SSL verification is disabled
    if not verify_ssl:
        logger.warning("SSL verification disabled for API testing - only use for self-signed certificates in controlled environments")

    async with httpx.AsyncClient(
        verify=verify_ssl,
        follow_redirects=True,
        proxy=proxy_config,
        timeout=timeout,
    ) as client:
        all_findings: List[Finding] = []
        
        for endpoint_config in endpoints:
            endpoint_url = endpoint_config.get("url", "")
            if not endpoint_url.startswith("http"):
                endpoint_url = urljoin(base_url, endpoint_url)
            
            method = endpoint_config.get("method", "GET").upper()
            params = endpoint_config.get("params", {})
            body = endpoint_config.get("body")
            
            logger.info(f"Testing endpoint: {method} {endpoint_url}")
            
            # Basic endpoint test
            endpoint_result = await test_endpoint(
                client, endpoint_url, method, 
                auth_header=auth_header,
                body=body,
            )
            
            if endpoint_result.error:
                result.endpoint_results.append(endpoint_result)
                continue
            
            # Security header check
            header_findings = await check_security_headers(endpoint_result)
            endpoint_result.findings.extend(header_findings)
            all_findings.extend(header_findings)
            
            # CORS check
            if test_cors:
                cors_findings = await check_cors(client, endpoint_url)
                endpoint_result.findings.extend(cors_findings)
                all_findings.extend(cors_findings)
            
            # Authentication check
            if test_auth:
                auth_findings = await check_authentication(
                    client, endpoint_url, method, auth_header
                )
                endpoint_result.findings.extend(auth_findings)
                all_findings.extend(auth_findings)
            
            # Rate limiting check (only on one endpoint to avoid excessive requests)
            if test_rate_limit and endpoints.index(endpoint_config) == 0:
                rate_findings = await check_rate_limiting(client, endpoint_url, method)
                endpoint_result.findings.extend(rate_findings)
                all_findings.extend(rate_findings)
            
            # Input validation check
            if test_input_validation:
                input_findings = await check_input_validation(
                    client, endpoint_url, method, params
                )
                endpoint_result.findings.extend(input_findings)
                all_findings.extend(input_findings)
            
            # HTTP methods check
            if test_methods:
                method_findings = await check_http_methods(client, endpoint_url)
                endpoint_result.findings.extend(method_findings)
                all_findings.extend(method_findings)
            
            # Sensitive data check
            try:
                response = await client.request(method, endpoint_url, timeout=10.0)
                sensitive_findings = await check_sensitive_data(response.text, endpoint_url)
                endpoint_result.findings.extend(sensitive_findings)
                all_findings.extend(sensitive_findings)
            except:
                pass
            
            # Error handling check
            error_findings = await check_error_handling(client, endpoint_url)
            endpoint_result.findings.extend(error_findings)
            all_findings.extend(error_findings)
            
            # GraphQL check
            if test_graphql:
                graphql_findings = await check_graphql(client, endpoint_url)
                endpoint_result.findings.extend(graphql_findings)
                all_findings.extend(graphql_findings)
            
            result.endpoint_results.append(endpoint_result)
        
        # Deduplicate findings
        seen_findings = set()
        unique_findings = []
        for f in all_findings:
            key = (f.title, f.endpoint, f.severity)
            if key not in seen_findings:
                seen_findings.add(key)
                unique_findings.append(f)
        
        result.all_findings = unique_findings
        result.total_findings = len(unique_findings)
        result.endpoints_tested = len(endpoints)
        
        # Count by severity
        for f in unique_findings:
            if f.severity == Severity.CRITICAL.value:
                result.critical_count += 1
            elif f.severity == Severity.HIGH.value:
                result.high_count += 1
            elif f.severity == Severity.MEDIUM.value:
                result.medium_count += 1
            elif f.severity == Severity.LOW.value:
                result.low_count += 1
            else:
                result.info_count += 1
        
        result.security_score = calculate_security_score(unique_findings)
        result.test_duration_seconds = time.time() - start_time
        
        # Generate summary
        result.summary = generate_summary(result)
    
    return result


def generate_summary(result: APITestResult) -> str:
    """Generate a text summary of the test results."""
    lines = [
        f"API Security Test Summary for {result.base_url}",
        f"=" * 50,
        f"Endpoints Tested: {result.endpoints_tested}",
        f"Total Findings: {result.total_findings}",
        f"Security Score: {result.security_score}/100",
        "",
        "Findings by Severity:",
        f"  Critical: {result.critical_count}",
        f"  High: {result.high_count}",
        f"  Medium: {result.medium_count}",
        f"  Low: {result.low_count}",
        f"  Info: {result.info_count}",
        "",
        f"Test Duration: {result.test_duration_seconds:.2f} seconds",
    ]
    
    if result.critical_count > 0:
        lines.append("")
        lines.append("⚠️ CRITICAL ISSUES FOUND - Immediate attention required!")
    
    return "\n".join(lines)


async def quick_scan(url: str, proxy_url: Optional[str] = None) -> APITestResult:
    """
    Quick scan of a single URL with all tests enabled.
    """
    return await test_api_endpoints(
        base_url=url,
        endpoints=[{"url": url, "method": "GET"}],
        test_auth=True,
        test_cors=True,
        test_rate_limit=False,  # Skip rate limit for quick scan
        test_input_validation=True,
        test_methods=True,
        test_graphql=False,
        proxy_url=proxy_url,
    )


# ============================================================================
# WebSocket Security Testing
# ============================================================================

@dataclass
class WebSocketFinding:
    """A security finding from WebSocket testing."""
    title: str
    description: str
    severity: str
    category: str
    evidence: str = ""
    remediation: str = ""
    cwe: Optional[str] = None
    owasp_api: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class WebSocketTestResult:
    """Result of WebSocket security testing."""
    url: str
    connected: bool = False
    connection_time_ms: float = 0
    protocol: Optional[str] = None
    subprotocol: Optional[str] = None
    findings: List[WebSocketFinding] = field(default_factory=list)
    messages_sent: int = 0
    messages_received: int = 0
    error: Optional[str] = None
    test_duration_seconds: float = 0
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "url": self.url,
            "connected": self.connected,
            "connection_time_ms": self.connection_time_ms,
            "protocol": self.protocol,
            "subprotocol": self.subprotocol,
            "findings": [f.to_dict() for f in self.findings],
            "messages_sent": self.messages_sent,
            "messages_received": self.messages_received,
            "error": self.error,
            "test_duration_seconds": self.test_duration_seconds,
            "security_score": self.calculate_score(),
            "owasp_api_breakdown": self.get_owasp_breakdown(),
        }
    
    def calculate_score(self) -> int:
        """Calculate security score based on findings."""
        score = 100
        penalties = {"critical": 25, "high": 15, "medium": 8, "low": 3, "info": 0}
        for f in self.findings:
            score -= penalties.get(f.severity, 0)
        return max(0, score)
    
    def get_owasp_breakdown(self) -> Dict[str, int]:
        """Get count by OWASP API category."""
        breakdown = {}
        for f in self.findings:
            if f.owasp_api:
                breakdown[f.owasp_api] = breakdown.get(f.owasp_api, 0) + 1
        return breakdown


# WebSocket test payloads
WS_INJECTION_PAYLOADS = [
    '{"type": "subscribe", "channel": "../../../etc/passwd"}',
    '{"type": "message", "content": "<script>alert(1)</script>"}',
    '{"type": "auth", "token": "' + "A" * 10000 + '"}',  # Buffer overflow test
    '{"__proto__": {"admin": true}}',  # Prototype pollution
    '{"constructor": {"prototype": {"admin": true}}}',
    '{"type": "subscribe", "channel": "*"}',  # Wildcard subscription
]


async def test_websocket(
    url: str,
    auth_token: Optional[str] = None,
    test_messages: Optional[List[str]] = None,
    timeout: float = 10.0,
    proxy_url: Optional[str] = None,
) -> WebSocketTestResult:
    """
    Test WebSocket endpoint for security vulnerabilities.
    
    Tests:
    - Connection without authentication
    - Origin header validation
    - Message injection (XSS, SQLi patterns)
    - Prototype pollution
    - Buffer overflow with large messages
    - Rate limiting on messages
    - Cross-Site WebSocket Hijacking (CSWSH)
    
    Args:
        url: WebSocket URL (ws:// or wss://)
        auth_token: Optional authentication token
        test_messages: Custom messages to send for testing
        timeout: Connection timeout in seconds
        proxy_url: Optional proxy URL
    
    Returns:
        WebSocketTestResult with findings
    """
    import websockets
    from websockets.exceptions import (
        ConnectionClosed, 
        InvalidStatusCode,
        InvalidHandshake,
    )
    
    result = WebSocketTestResult(url=url)
    start_time = time.time()
    findings: List[WebSocketFinding] = []
    
    # Validate URL
    if not url.startswith(("ws://", "wss://")):
        result.error = "Invalid WebSocket URL. Must start with ws:// or wss://"
        return result
    
    # Check for unencrypted WebSocket
    if url.startswith("ws://") and not ("localhost" in url or "127.0.0.1" in url):
        findings.append(WebSocketFinding(
            title="Unencrypted WebSocket Connection",
            description="WebSocket connection uses ws:// instead of wss://, allowing traffic interception",
            severity=Severity.HIGH.value,
            category="transport",
            evidence=f"URL: {url}",
            remediation="Use wss:// (WebSocket Secure) for all WebSocket connections",
            cwe="CWE-319",
            owasp_api="API8:2023",
        ))
    
    # Test 1: Connection without authentication
    try:
        connect_start = time.time()
        async with websockets.connect(
            url,
            close_timeout=timeout,
            open_timeout=timeout,
        ) as ws:
            result.connected = True
            result.connection_time_ms = (time.time() - connect_start) * 1000
            # websockets 13.x: protocol info is different from older versions
            result.protocol = "WebSocket"  # Generic protocol name
            result.subprotocol = ws.subprotocol if hasattr(ws, 'subprotocol') and ws.subprotocol else None
            
            # If connected without auth token, that might be an issue
            if not auth_token:
                findings.append(WebSocketFinding(
                    title="WebSocket Accepts Unauthenticated Connections",
                    description="WebSocket endpoint allows connections without authentication",
                    severity=Severity.MEDIUM.value,
                    category="authentication",
                    evidence="Successfully connected without authentication token",
                    remediation="Implement authentication for WebSocket connections",
                    cwe="CWE-306",
                    owasp_api="API2:2023",
                ))
            
            # Test 2: Send test messages and check responses
            messages_to_test = test_messages or WS_INJECTION_PAYLOADS
            
            for msg in messages_to_test[:5]:  # Limit to 5 test messages
                try:
                    await ws.send(msg)
                    result.messages_sent += 1
                    
                    # Try to receive response
                    try:
                        response = await asyncio.wait_for(ws.recv(), timeout=2.0)
                        result.messages_received += 1
                        
                        # Check for XSS reflection
                        if "<script>" in msg and "<script>" in str(response):
                            findings.append(WebSocketFinding(
                                title="XSS Payload Reflected in WebSocket Response",
                                description="Server reflects XSS payloads without sanitization",
                                severity=Severity.HIGH.value,
                                category="input_validation",
                                evidence=f"Sent: {msg[:100]}, Received: {str(response)[:100]}",
                                remediation="Sanitize all WebSocket message content before processing",
                                cwe="CWE-79",
                                owasp_api="API8:2023",
                            ))
                        
                        # Check for error disclosure
                        response_lower = str(response).lower()
                        if any(x in response_lower for x in ["traceback", "exception", "error", "stack"]):
                            findings.append(WebSocketFinding(
                                title="Verbose Error Messages in WebSocket",
                                description="Server returns detailed error information",
                                severity=Severity.LOW.value,
                                category="information_disclosure",
                                evidence=f"Response contains error details: {str(response)[:200]}",
                                remediation="Return generic error messages in production",
                                cwe="CWE-209",
                                owasp_api="API9:2023",
                            ))
                            
                    except asyncio.TimeoutError:
                        pass  # No response is okay
                        
                except Exception as e:
                    logger.debug(f"Message send failed: {e}")
            
            # Test 3: Rate limiting (send rapid messages)
            rate_test_count = 20
            rate_blocked = False
            for i in range(rate_test_count):
                try:
                    await ws.send('{"type": "ping"}')
                    result.messages_sent += 1
                except Exception:
                    rate_blocked = True
                    break
            
            if not rate_blocked:
                findings.append(WebSocketFinding(
                    title="No Rate Limiting on WebSocket Messages",
                    description=f"Server accepted {rate_test_count} rapid messages without rate limiting",
                    severity=Severity.MEDIUM.value,
                    category="rate_limiting",
                    evidence=f"Sent {rate_test_count} messages in rapid succession",
                    remediation="Implement rate limiting for WebSocket messages",
                    cwe="CWE-770",
                    owasp_api="API4:2023",
                ))
                
    except InvalidStatusCode as e:
        if e.status_code == 401:
            findings.append(WebSocketFinding(
                title="WebSocket Requires Authentication",
                description="WebSocket properly rejects unauthenticated connections",
                severity=Severity.INFO.value,
                category="authentication",
                evidence=f"HTTP {e.status_code} returned",
                remediation="Authentication is properly implemented",
                cwe=None,
                owasp_api=None,
            ))
        elif e.status_code == 403:
            findings.append(WebSocketFinding(
                title="WebSocket Access Forbidden",
                description="WebSocket endpoint returned 403 Forbidden",
                severity=Severity.INFO.value,
                category="authorization",
                evidence=f"HTTP {e.status_code} returned",
                remediation="Access control is implemented",
                cwe=None,
                owasp_api=None,
            ))
        else:
            result.error = f"Connection failed with status {e.status_code}"
            
    except InvalidHandshake as e:
        result.error = f"WebSocket handshake failed: {str(e)}"
        
    except ConnectionClosed as e:
        findings.append(WebSocketFinding(
            title="WebSocket Connection Closed Unexpectedly",
            description="Server closed connection during testing",
            severity=Severity.INFO.value,
            category="general",
            evidence=f"Close code: {e.code}, reason: {e.reason}",
            remediation="Review connection handling logic",
            cwe=None,
            owasp_api=None,
        ))
        
    except asyncio.TimeoutError:
        result.error = f"Connection timeout after {timeout} seconds"
        
    except Exception as e:
        result.error = f"WebSocket test failed: {str(e)}"
    
    # Test 4: CSWSH - Cross-Site WebSocket Hijacking (origin check)
    try:
        evil_origins = ["https://evil.com", "https://attacker.com", "null"]
        for origin in evil_origins:
            try:
                async with websockets.connect(
                    url,
                    additional_headers={"Origin": origin},
                    close_timeout=5.0,
                    open_timeout=5.0,
                ) as ws:
                    findings.append(WebSocketFinding(
                        title="Cross-Site WebSocket Hijacking (CSWSH) Possible",
                        description=f"WebSocket accepts connections from untrusted origin: {origin}",
                        severity=Severity.HIGH.value,
                        category="cors",
                        evidence=f"Successfully connected with Origin: {origin}",
                        remediation="Validate Origin header and reject untrusted origins",
                        cwe="CWE-942",
                        owasp_api="API8:2023",
                    ))
                    break
            except:
                pass
    except:
        pass
    
    result.findings = findings
    result.test_duration_seconds = time.time() - start_time
    
    return result


def get_owasp_api_reference() -> Dict[str, Dict[str, str]]:
    """Return OWASP API Security Top 10 reference information."""
    return OWASP_API_TOP_10


# =============================================================================
# Network Discovery for Air-Gapped Environments
# =============================================================================

@dataclass
class DiscoveredService:
    """A discovered HTTP/API service on the network."""
    ip: str
    port: int
    url: str
    status_code: Optional[int] = None
    server: Optional[str] = None
    title: Optional[str] = None
    is_api: bool = False
    api_indicators: List[str] = field(default_factory=list)
    response_time_ms: float = 0
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class NetworkDiscoveryResult:
    """Result of network discovery scan."""
    subnet: str
    total_hosts_scanned: int
    services_found: List[DiscoveredService] = field(default_factory=list)
    scan_duration_seconds: float = 0
    errors: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "subnet": self.subnet,
            "total_hosts_scanned": self.total_hosts_scanned,
            "services_found": [s.to_dict() for s in self.services_found],
            "scan_duration_seconds": self.scan_duration_seconds,
            "errors": self.errors,
        }


async def discover_http_services(
    subnet: str,
    ports: Optional[List[int]] = None,
    timeout: float = 1.5,
    max_concurrent: int = 100,
    max_hosts: int = 256,
    overall_timeout: float = 120.0,
) -> NetworkDiscoveryResult:
    """
    Discover HTTP/API services on a subnet.
    
    Args:
        subnet: CIDR notation (e.g., "192.168.1.0/24") or IP range
        ports: Ports to scan (default: common HTTP ports)
        timeout: Connection timeout per host (default: 1.5s)
        max_concurrent: Max concurrent connections (default: 100)
        max_hosts: Maximum number of hosts to scan (default: 256)
        overall_timeout: Maximum total scan time in seconds (default: 120s)
    
    Returns:
        NetworkDiscoveryResult with discovered services
    """
    import ipaddress
    
    if ports is None:
        ports = [80, 443, 8080, 8443, 3000, 5000, 8000]
    
    result = NetworkDiscoveryResult(subnet=subnet, total_hosts_scanned=0)
    start_time = time.time()
    
    # Parse subnet
    try:
        if "/" in subnet:
            network = ipaddress.ip_network(subnet, strict=False)
            hosts = list(network.hosts())
        elif "-" in subnet:
            # Handle IP range like "192.168.1.1-192.168.1.254"
            start_ip, end_ip = subnet.split("-")
            start = ipaddress.ip_address(start_ip.strip())
            end = ipaddress.ip_address(end_ip.strip())
            hosts = []
            current = start
            while current <= end:
                hosts.append(current)
                current = ipaddress.ip_address(int(current) + 1)
        else:
            # Single IP
            hosts = [ipaddress.ip_address(subnet)]
    except Exception as e:
        result.errors.append(f"Invalid subnet format: {e}")
        return result
    
    # Limit hosts to prevent excessive scanning
    if len(hosts) > max_hosts:
        result.errors.append(f"Subnet limited to {max_hosts} hosts (requested {len(hosts)})")
        hosts = hosts[:max_hosts]
    
    result.total_hosts_scanned = len(hosts) * len(ports)
    logger.info(f"Network discovery: scanning {len(hosts)} hosts x {len(ports)} ports = {result.total_hosts_scanned} probes")
    
    # API indicators to look for
    api_indicators = [
        "application/json",
        "api",
        "swagger",
        "openapi",
        "graphql",
        "rest",
        "v1",
        "v2",
        "endpoint",
    ]
    
    semaphore = asyncio.Semaphore(max_concurrent)
    
    async def check_service(ip: str, port: int) -> Optional[DiscoveredService]:
        async with semaphore:
            scheme = "https" if port in [443, 8443] else "http"
            url = f"{scheme}://{ip}:{port}"

            try:
                start = time.time()
                # Note: SSL verification disabled for service discovery - scanning unknown services
                async with httpx.AsyncClient(
                    timeout=timeout,
                    verify=True,  # Enable SSL verification by default
                    follow_redirects=True,
                ) as client:
                    response = await client.get(url)
                    response_time = (time.time() - start) * 1000
                    
                    # Check for API indicators
                    content_type = response.headers.get("content-type", "").lower()
                    server = response.headers.get("server", "")
                    body = response.text[:2000].lower()
                    
                    found_indicators = []
                    for indicator in api_indicators:
                        if indicator in content_type or indicator in body:
                            found_indicators.append(indicator)
                    
                    # Try to extract page title
                    title = None
                    title_match = re.search(r"<title[^>]*>([^<]+)</title>", body, re.IGNORECASE)
                    if title_match:
                        title = title_match.group(1).strip()[:100]
                    
                    return DiscoveredService(
                        ip=ip,
                        port=port,
                        url=url,
                        status_code=response.status_code,
                        server=server or None,
                        title=title,
                        is_api=len(found_indicators) > 0 or "json" in content_type,
                        api_indicators=found_indicators,
                        response_time_ms=response_time,
                    )
            except Exception:
                return None
    
    # Run discovery with overall timeout
    tasks = []
    for host in hosts:
        for port in ports:
            tasks.append(check_service(str(host), port))
    
    try:
        # Apply overall timeout to prevent hanging
        remaining_time = overall_timeout - (time.time() - start_time)
        if remaining_time <= 0:
            result.errors.append("Scan timed out before starting probes")
            return result
            
        results = await asyncio.wait_for(
            asyncio.gather(*tasks, return_exceptions=True),
            timeout=remaining_time
        )
        
        for r in results:
            if isinstance(r, DiscoveredService):
                result.services_found.append(r)
            elif isinstance(r, Exception):
                # Don't log every failed connection
                pass
    except asyncio.TimeoutError:
        result.errors.append(f"Scan timed out after {overall_timeout}s")
        logger.warning(f"Network discovery timed out for {subnet}")
    
    # Sort by IP and port
    result.services_found.sort(key=lambda s: (s.ip, s.port))
    result.scan_duration_seconds = time.time() - start_time
    
    return result


# =============================================================================
# Target Presets for Quick Access
# =============================================================================

@dataclass
class TargetPreset:
    """A saved target preset for quick testing."""
    id: str
    name: str
    description: str
    base_url: str
    endpoints: List[Dict[str, str]] = field(default_factory=list)
    auth_type: Optional[str] = None
    auth_value: Optional[str] = None
    headers: Dict[str, str] = field(default_factory=dict)
    tags: List[str] = field(default_factory=list)
    created_at: str = ""
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


# In-memory storage for presets (could be moved to DB later)
_target_presets: Dict[str, TargetPreset] = {}

# Default presets for common testing scenarios
DEFAULT_PRESETS = [
    TargetPreset(
        id="httpbin",
        name="HTTPBin (Public)",
        description="Public HTTP testing service - great for testing your setup",
        base_url="https://httpbin.org",
        endpoints=[
            {"url": "/get", "method": "GET"},
            {"url": "/post", "method": "POST"},
            {"url": "/headers", "method": "GET"},
            {"url": "/status/200", "method": "GET"},
        ],
        tags=["public", "testing"],
    ),
    TargetPreset(
        id="local-api",
        name="Local API (localhost:8080)",
        description="Common local development API endpoint",
        base_url="http://localhost:8080",
        endpoints=[
            {"url": "/api/health", "method": "GET"},
            {"url": "/api/v1/users", "method": "GET"},
        ],
        tags=["local", "development"],
    ),
    TargetPreset(
        id="vm-template",
        name="VM Template",
        description="Template for testing VM APIs - customize the IP",
        base_url="http://192.168.1.100",
        endpoints=[
            {"url": "/", "method": "GET"},
            {"url": "/api", "method": "GET"},
        ],
        tags=["vm", "template"],
    ),
]

# Initialize default presets
for preset in DEFAULT_PRESETS:
    _target_presets[preset.id] = preset


def get_all_presets() -> List[TargetPreset]:
    """Get all saved target presets."""
    return list(_target_presets.values())


def get_preset(preset_id: str) -> Optional[TargetPreset]:
    """Get a specific preset by ID."""
    return _target_presets.get(preset_id)


def save_preset(preset: TargetPreset) -> TargetPreset:
    """Save or update a target preset."""
    import uuid
    from datetime import datetime
    
    if not preset.id:
        preset.id = str(uuid.uuid4())[:8]
    if not preset.created_at:
        preset.created_at = datetime.now().isoformat()
    
    _target_presets[preset.id] = preset
    return preset


def delete_preset(preset_id: str) -> bool:
    """Delete a target preset."""
    if preset_id in _target_presets and preset_id not in ["httpbin", "local-api", "vm-template"]:
        del _target_presets[preset_id]
        return True
    return False


# =============================================================================
# Batch Testing for Multiple Targets
# =============================================================================

@dataclass
class BatchTestTarget:
    """A target in a batch test."""
    url: str
    name: Optional[str] = None
    auth_type: Optional[str] = None
    auth_value: Optional[str] = None


@dataclass
class BatchTestResult:
    """Result of batch testing multiple targets."""
    total_targets: int
    successful: int
    failed: int
    results: List[Dict[str, Any]] = field(default_factory=list)
    total_findings: int = 0
    critical_findings: int = 0
    high_findings: int = 0
    scan_duration_seconds: float = 0
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


async def batch_test_targets(
    targets: List[BatchTestTarget],
    test_options: Optional[Dict[str, bool]] = None,
    proxy_url: Optional[str] = None,
    max_concurrent: int = 5,
) -> BatchTestResult:
    """
    Test multiple API targets in batch.
    
    Args:
        targets: List of targets to test
        test_options: Test configuration options
        proxy_url: Optional proxy for all requests
        max_concurrent: Max concurrent target tests
    
    Returns:
        BatchTestResult with aggregated results
    """
    if test_options is None:
        test_options = {
            "test_auth": True,
            "test_cors": True,
            "test_rate_limit": False,  # Disabled by default for batch
            "test_input_validation": True,
            "test_methods": True,
        }
    
    result = BatchTestResult(
        total_targets=len(targets),
        successful=0,
        failed=0,
    )
    start_time = time.time()
    
    semaphore = asyncio.Semaphore(max_concurrent)
    
    async def test_target(target: BatchTestTarget) -> Dict[str, Any]:
        async with semaphore:
            try:
                # Run quick scan on the target
                scan_result = await quick_scan(
                    url=target.url,
                    proxy_url=proxy_url,
                )
                
                return {
                    "target": target.url,
                    "name": target.name or target.url,
                    "success": True,
                    "security_score": scan_result.security_score,
                    "total_findings": scan_result.total_findings,
                    "critical_count": scan_result.critical_count,
                    "high_count": scan_result.high_count,
                    "medium_count": scan_result.medium_count,
                    "low_count": scan_result.low_count,
                    "findings": [f.to_dict() for f in scan_result.all_findings[:10]],  # Top 10
                    "error": None,
                }
            except Exception as e:
                return {
                    "target": target.url,
                    "name": target.name or target.url,
                    "success": False,
                    "security_score": 0,
                    "total_findings": 0,
                    "critical_count": 0,
                    "high_count": 0,
                    "medium_count": 0,
                    "low_count": 0,
                    "findings": [],
                    "error": str(e),
                }
    
    # Run all tests (with return_exceptions for robustness)
    tasks = [test_target(t) for t in targets]
    results = await asyncio.gather(*tasks, return_exceptions=True)
    
    # Aggregate results
    for r in results:
        # Skip exceptions from gather
        if isinstance(r, Exception):
            result.failed += 1
            result.results.append({
                "target": "Unknown",
                "name": "Unknown",
                "success": False,
                "security_score": 0,
                "total_findings": 0,
                "critical_count": 0,
                "high_count": 0,
                "medium_count": 0,
                "low_count": 0,
                "findings": [],
                "error": str(r),
            })
            continue
        
        result.results.append(r)
        if r["success"]:
            result.successful += 1
            result.total_findings += r["total_findings"]
            result.critical_findings += r["critical_count"]
            result.high_findings += r["high_count"]
        else:
            result.failed += 1
    
    result.scan_duration_seconds = time.time() - start_time
    
    return result


# =============================================================================
# OpenAPI/Swagger Import
# =============================================================================

@dataclass
class OpenAPIEndpoint:
    """An endpoint discovered from OpenAPI spec."""
    path: str
    method: str
    summary: str = ""
    description: str = ""
    parameters: List[Dict[str, Any]] = field(default_factory=list)
    request_body: Optional[Dict[str, Any]] = None
    security: List[str] = field(default_factory=list)
    tags: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class OpenAPIParseResult:
    """Result of parsing an OpenAPI spec."""
    title: str
    version: str
    base_url: str
    endpoints: List[OpenAPIEndpoint]
    security_schemes: Dict[str, Any]
    total_endpoints: int
    methods_breakdown: Dict[str, int]
    tags: List[str]
    errors: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "version": self.version,
            "base_url": self.base_url,
            "endpoints": [e.to_dict() for e in self.endpoints],
            "security_schemes": self.security_schemes,
            "total_endpoints": self.total_endpoints,
            "methods_breakdown": self.methods_breakdown,
            "tags": self.tags,
            "errors": self.errors,
        }


def parse_openapi_spec(spec_content: str, spec_url: Optional[str] = None) -> OpenAPIParseResult:
    """
    Parse an OpenAPI/Swagger specification and extract endpoints.
    
    Supports OpenAPI 3.0.x and Swagger 2.0.
    
    Args:
        spec_content: JSON or YAML content of the spec
        spec_url: Optional URL where spec was loaded from (for base URL resolution)
    
    Returns:
        OpenAPIParseResult with discovered endpoints
    """
    import yaml
    
    errors = []
    
    # Parse the spec
    try:
        if spec_content.strip().startswith("{"):
            spec = json.loads(spec_content)
        else:
            spec = yaml.safe_load(spec_content)
    except Exception as e:
        return OpenAPIParseResult(
            title="Parse Error",
            version="",
            base_url="",
            endpoints=[],
            security_schemes={},
            total_endpoints=0,
            methods_breakdown={},
            tags=[],
            errors=[f"Failed to parse spec: {str(e)}"],
        )
    
    # Detect spec version
    is_openapi3 = "openapi" in spec and spec.get("openapi", "").startswith("3.")
    is_swagger2 = "swagger" in spec and spec.get("swagger", "").startswith("2.")
    
    if not is_openapi3 and not is_swagger2:
        errors.append("Unknown spec format. Expected OpenAPI 3.x or Swagger 2.0")
    
    # Extract info
    info = spec.get("info", {})
    title = info.get("title", "Unknown API")
    version = info.get("version", "")
    
    # Extract base URL
    base_url = ""
    if is_openapi3:
        servers = spec.get("servers", [])
        if servers:
            base_url = servers[0].get("url", "")
    elif is_swagger2:
        host = spec.get("host", "")
        base_path = spec.get("basePath", "")
        schemes = spec.get("schemes", ["https"])
        if host:
            base_url = f"{schemes[0]}://{host}{base_path}"
    
    if spec_url and not base_url:
        parsed = urlparse(spec_url)
        base_url = f"{parsed.scheme}://{parsed.netloc}"
    
    # Extract security schemes
    security_schemes = {}
    if is_openapi3:
        components = spec.get("components", {})
        security_schemes = components.get("securitySchemes", {})
    elif is_swagger2:
        security_schemes = spec.get("securityDefinitions", {})
    
    # Extract endpoints
    endpoints = []
    methods_breakdown = {"GET": 0, "POST": 0, "PUT": 0, "DELETE": 0, "PATCH": 0, "OPTIONS": 0, "HEAD": 0}
    all_tags = set()
    
    paths = spec.get("paths", {})
    for path, path_item in paths.items():
        if not isinstance(path_item, dict):
            continue
            
        for method in ["get", "post", "put", "delete", "patch", "options", "head"]:
            if method not in path_item:
                continue
            
            operation = path_item[method]
            if not isinstance(operation, dict):
                continue
            
            # Extract parameters
            parameters = []
            params = operation.get("parameters", []) + path_item.get("parameters", [])
            for param in params:
                if isinstance(param, dict):
                    parameters.append({
                        "name": param.get("name", ""),
                        "in": param.get("in", ""),
                        "required": param.get("required", False),
                        "type": param.get("type", param.get("schema", {}).get("type", "")),
                    })
            
            # Extract request body (OpenAPI 3)
            request_body = None
            if "requestBody" in operation:
                rb = operation["requestBody"]
                content = rb.get("content", {})
                if content:
                    first_content_type = list(content.keys())[0]
                    schema = content[first_content_type].get("schema", {})
                    request_body = {
                        "content_type": first_content_type,
                        "required": rb.get("required", False),
                        "schema": schema,
                    }
            
            # Extract security requirements
            security = []
            op_security = operation.get("security", spec.get("security", []))
            for sec_req in op_security:
                if isinstance(sec_req, dict):
                    security.extend(sec_req.keys())
            
            # Extract tags
            tags = operation.get("tags", [])
            all_tags.update(tags)
            
            endpoint = OpenAPIEndpoint(
                path=path,
                method=method.upper(),
                summary=operation.get("summary", ""),
                description=operation.get("description", ""),
                parameters=parameters,
                request_body=request_body,
                security=security,
                tags=tags,
            )
            endpoints.append(endpoint)
            
            method_upper = method.upper()
            if method_upper in methods_breakdown:
                methods_breakdown[method_upper] += 1
    
    return OpenAPIParseResult(
        title=title,
        version=version,
        base_url=base_url,
        endpoints=endpoints,
        security_schemes=security_schemes,
        total_endpoints=len(endpoints),
        methods_breakdown=methods_breakdown,
        tags=sorted(all_tags),
        errors=errors,
    )


async def fetch_openapi_spec(url: str, verify_ssl: bool = True) -> str:
    """Fetch OpenAPI spec from a URL.

    Args:
        url: URL to fetch the OpenAPI spec from
        verify_ssl: Whether to verify SSL certificates (default True)
    """
    if not verify_ssl:
        logger.warning("Fetching OpenAPI spec with SSL verification disabled")
    async with httpx.AsyncClient(timeout=30.0, verify=verify_ssl) as client:
        response = await client.get(url)
        response.raise_for_status()
        return response.text


# =============================================================================
# JWT Token Analyzer
# =============================================================================

@dataclass
class JWTAnalysisResult:
    """Result of analyzing a JWT token."""
    valid_structure: bool
    header: Dict[str, Any]
    payload: Dict[str, Any]
    signature: str
    algorithm: str
    findings: List[Dict[str, Any]]
    is_expired: bool
    expiry_time: Optional[str]
    issued_at: Optional[str]
    issuer: Optional[str]
    audience: Optional[str]
    subject: Optional[str]
    raw_parts: List[str]
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


def analyze_jwt(token: str, test_weak_secrets: bool = True) -> JWTAnalysisResult:
    """
    Analyze a JWT token for security issues.
    
    Checks for:
    - Algorithm vulnerabilities (none, HS256 with weak secrets)
    - Expiration status
    - Missing recommended claims
    - Sensitive data in payload
    
    Args:
        token: The JWT token to analyze
        test_weak_secrets: Whether to test common weak secrets
    
    Returns:
        JWTAnalysisResult with detailed analysis
    """
    import base64
    from datetime import datetime
    
    findings = []
    header = {}
    payload = {}
    signature = ""
    algorithm = ""
    is_expired = False
    expiry_time = None
    issued_at = None
    issuer = None
    audience = None
    subject = None
    raw_parts = []
    
    # Split token
    parts = token.strip().split(".")
    raw_parts = parts
    
    if len(parts) != 3:
        return JWTAnalysisResult(
            valid_structure=False,
            header={},
            payload={},
            signature="",
            algorithm="",
            findings=[{
                "title": "Invalid JWT Structure",
                "description": f"JWT should have 3 parts separated by dots, found {len(parts)}",
                "severity": "high",
            }],
            is_expired=False,
            expiry_time=None,
            issued_at=None,
            issuer=None,
            audience=None,
            subject=None,
            raw_parts=raw_parts,
        )
    
    # Decode header
    try:
        header_b64 = parts[0] + "=" * (4 - len(parts[0]) % 4)  # Add padding
        header = json.loads(base64.urlsafe_b64decode(header_b64))
        algorithm = header.get("alg", "")
    except Exception as e:
        findings.append({
            "title": "Invalid Header",
            "description": f"Failed to decode JWT header: {str(e)}",
            "severity": "high",
        })
    
    # Decode payload
    try:
        payload_b64 = parts[1] + "=" * (4 - len(parts[1]) % 4)
        payload = json.loads(base64.urlsafe_b64decode(payload_b64))
    except Exception as e:
        findings.append({
            "title": "Invalid Payload",
            "description": f"Failed to decode JWT payload: {str(e)}",
            "severity": "high",
        })
    
    signature = parts[2]
    
    # Check algorithm vulnerabilities
    if algorithm.lower() == "none":
        findings.append({
            "title": "Algorithm 'none' Used",
            "description": "JWT uses 'none' algorithm which provides no signature verification. This is a critical vulnerability.",
            "severity": "critical",
            "cwe": "CWE-327",
            "remediation": "Use a secure algorithm like RS256 or ES256. Never accept 'none' algorithm.",
        })
    elif algorithm.upper() in ["HS256", "HS384", "HS512"]:
        findings.append({
            "title": "Symmetric Algorithm Used",
            "description": f"JWT uses symmetric algorithm {algorithm}. Ensure the secret key is strong and securely stored.",
            "severity": "info",
            "remediation": "Consider using asymmetric algorithms (RS256, ES256) for better security in distributed systems.",
        })
        
        # Test weak secrets
        if test_weak_secrets:
            weak_secrets = [
                "secret", "password", "123456", "key", "private",
                "jwt_secret", "your-256-bit-secret", "your-secret-key",
                "changeme", "supersecret", "admin", "test", "development",
            ]
            
            import hmac
            import hashlib
            
            for weak_secret in weak_secrets:
                try:
                    # Create signature with weak secret
                    message = f"{parts[0]}.{parts[1]}".encode()
                    
                    if algorithm.upper() == "HS256":
                        expected_sig = base64.urlsafe_b64encode(
                            hmac.new(weak_secret.encode(), message, hashlib.sha256).digest()
                        ).rstrip(b"=").decode()
                    elif algorithm.upper() == "HS384":
                        expected_sig = base64.urlsafe_b64encode(
                            hmac.new(weak_secret.encode(), message, hashlib.sha384).digest()
                        ).rstrip(b"=").decode()
                    elif algorithm.upper() == "HS512":
                        expected_sig = base64.urlsafe_b64encode(
                            hmac.new(weak_secret.encode(), message, hashlib.sha512).digest()
                        ).rstrip(b"=").decode()
                    else:
                        continue
                    
                    if expected_sig == signature:
                        findings.append({
                            "title": "Weak Secret Key Detected",
                            "description": f"JWT is signed with a weak/common secret: '{weak_secret}'",
                            "severity": "critical",
                            "cwe": "CWE-521",
                            "evidence": f"Secret: {weak_secret}",
                            "remediation": "Use a cryptographically strong random secret of at least 256 bits.",
                        })
                        break
                except Exception:
                    pass
    
    # Check expiration
    if "exp" in payload:
        try:
            exp_timestamp = payload["exp"]
            exp_dt = datetime.fromtimestamp(exp_timestamp)
            expiry_time = exp_dt.isoformat()
            
            if datetime.now() > exp_dt:
                is_expired = True
                findings.append({
                    "title": "Token Expired",
                    "description": f"JWT token expired at {expiry_time}",
                    "severity": "medium",
                })
        except Exception:
            pass
    else:
        findings.append({
            "title": "Missing Expiration Claim",
            "description": "JWT does not have an 'exp' (expiration) claim. Tokens should have a limited lifetime.",
            "severity": "medium",
            "cwe": "CWE-613",
            "remediation": "Add an expiration claim (exp) to limit token lifetime.",
        })
    
    # Extract other claims
    if "iat" in payload:
        try:
            issued_at = datetime.fromtimestamp(payload["iat"]).isoformat()
        except Exception:
            pass
    
    issuer = payload.get("iss")
    audience = payload.get("aud")
    subject = payload.get("sub")
    
    # Check for sensitive data in payload
    sensitive_keys = ["password", "passwd", "pwd", "secret", "ssn", "credit_card", "cc", "cvv", "pin"]
    for key in payload.keys():
        if any(sensitive in key.lower() for sensitive in sensitive_keys):
            findings.append({
                "title": "Potential Sensitive Data in Payload",
                "description": f"JWT payload contains potentially sensitive field: '{key}'",
                "severity": "high",
                "cwe": "CWE-200",
                "remediation": "Do not store sensitive data in JWT payloads as they are only base64 encoded, not encrypted.",
            })
    
    # Check for missing recommended claims
    if "iat" not in payload:
        findings.append({
            "title": "Missing Issued At Claim",
            "description": "JWT does not have an 'iat' (issued at) claim.",
            "severity": "low",
            "remediation": "Add an issued at claim (iat) for token freshness validation.",
        })
    
    if "jti" not in payload:
        findings.append({
            "title": "Missing JWT ID Claim",
            "description": "JWT does not have a 'jti' (JWT ID) claim for uniqueness.",
            "severity": "info",
            "remediation": "Consider adding a unique identifier (jti) to prevent token replay attacks.",
        })
    
    return JWTAnalysisResult(
        valid_structure=True,
        header=header,
        payload=payload,
        signature=signature,
        algorithm=algorithm,
        findings=findings,
        is_expired=is_expired,
        expiry_time=expiry_time,
        issued_at=issued_at,
        issuer=issuer,
        audience=audience,
        subject=subject,
        raw_parts=raw_parts,
    )


# =============================================================================
# Export Reports
# =============================================================================

def export_test_result_json(result: APITestResult) -> str:
    """Export test result as JSON."""
    return json.dumps(result.to_dict(), indent=2)


def export_test_result_markdown(result: APITestResult, title: str = "API Security Test Report") -> str:
    """Export test result as Markdown report."""
    lines = []
    
    # Header
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"**Generated:** {time.strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"**Base URL:** {result.base_url}")
    lines.append("")
    
    # Summary
    lines.append("## Summary")
    lines.append("")
    lines.append(f"| Metric | Value |")
    lines.append("|--------|-------|")
    lines.append(f"| Security Score | **{result.security_score}/100** |")
    lines.append(f"| Endpoints Tested | {result.endpoints_tested} |")
    lines.append(f"| Total Findings | {result.total_findings} |")
    lines.append(f"| Critical | {result.critical_count} |")
    lines.append(f"| High | {result.high_count} |")
    lines.append(f"| Medium | {result.medium_count} |")
    lines.append(f"| Low | {result.low_count} |")
    lines.append(f"| Info | {result.info_count} |")
    lines.append(f"| Test Duration | {result.test_duration_seconds:.2f}s |")
    lines.append("")
    
    # OWASP API Top 10 Breakdown
    if result.owasp_api_breakdown:
        lines.append("## OWASP API Security Top 10")
        lines.append("")
        lines.append("| Category | Count |")
        lines.append("|----------|-------|")
        for category, count in sorted(result.owasp_api_breakdown.items(), key=lambda x: -x[1]):
            lines.append(f"| {category} | {count} |")
        lines.append("")
    
    # Findings by Severity
    if result.all_findings:
        lines.append("## Security Findings")
        lines.append("")
        
        # Group by severity
        by_severity = {"critical": [], "high": [], "medium": [], "low": [], "info": []}
        for finding in result.all_findings:
            sev = finding.severity.lower()
            if sev in by_severity:
                by_severity[sev].append(finding)
        
        for severity in ["critical", "high", "medium", "low", "info"]:
            findings = by_severity[severity]
            if findings:
                severity_emoji = {
                    "critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🔵", "info": "⚪"
                }
                lines.append(f"### {severity_emoji.get(severity, '')} {severity.upper()} ({len(findings)})")
                lines.append("")
                
                for i, finding in enumerate(findings, 1):
                    lines.append(f"#### {i}. {finding.title}")
                    lines.append("")
                    lines.append(f"**Category:** {finding.category}")
                    if finding.cwe:
                        lines.append(f"**CWE:** {finding.cwe}")
                    if finding.owasp_api:
                        lines.append(f"**OWASP API:** {finding.owasp_api}")
                    if finding.endpoint:
                        lines.append(f"**Endpoint:** `{finding.endpoint}`")
                    lines.append("")
                    lines.append(f"**Description:** {finding.description}")
                    lines.append("")
                    if finding.evidence:
                        lines.append("**Evidence:**")
                        lines.append("```")
                        lines.append(finding.evidence[:500])
                        lines.append("```")
                        lines.append("")
                    lines.append(f"**Remediation:** {finding.remediation}")
                    lines.append("")
                    lines.append("---")
                    lines.append("")
    else:
        lines.append("## Security Findings")
        lines.append("")
        lines.append("✅ No security issues found!")
        lines.append("")
    
    # Endpoint Details
    if result.endpoint_results:
        lines.append("## Endpoint Details")
        lines.append("")
        lines.append("| Method | URL | Status | Response Time | Findings |")
        lines.append("|--------|-----|--------|---------------|----------|")
        for ep in result.endpoint_results:
            lines.append(f"| {ep.method} | `{ep.url}` | {ep.status_code or 'Error'} | {ep.response_time_ms:.0f}ms | {len(ep.findings)} |")
        lines.append("")
    
    return "\n".join(lines)


def export_batch_result_markdown(result: BatchTestResult, title: str = "Batch API Test Report") -> str:
    """Export batch test result as Markdown report."""
    lines = []
    
    # Header
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"**Generated:** {time.strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    
    # Summary
    lines.append("## Summary")
    lines.append("")
    lines.append(f"| Metric | Value |")
    lines.append("|--------|-------|")
    lines.append(f"| Total Targets | {result.total_targets} |")
    lines.append(f"| Successful | {result.successful} |")
    lines.append(f"| Failed | {result.failed} |")
    lines.append(f"| Total Findings | {result.total_findings} |")
    lines.append(f"| Critical Findings | {result.critical_findings} |")
    lines.append(f"| High Findings | {result.high_findings} |")
    lines.append(f"| Scan Duration | {result.scan_duration_seconds:.2f}s |")
    lines.append("")
    
    # Results Table
    lines.append("## Results by Target")
    lines.append("")
    lines.append("| Target | Status | Score | Findings | Critical | High |")
    lines.append("|--------|--------|-------|----------|----------|------|")
    
    for r in sorted(result.results, key=lambda x: x.get("security_score", 0)):
        status = "✅" if r.get("success") else "❌"
        lines.append(
            f"| {r.get('name', r.get('target', ''))} | {status} | "
            f"{r.get('security_score', 0)}/100 | {r.get('total_findings', 0)} | "
            f"{r.get('critical_count', 0)} | {r.get('high_count', 0)} |"
        )
    lines.append("")
    
    # Individual Target Details
    lines.append("## Target Details")
    lines.append("")
    
    for r in result.results:
        lines.append(f"### {r.get('name', r.get('target', 'Unknown'))}")
        lines.append("")
        lines.append(f"**URL:** `{r.get('target', '')}`")
        lines.append(f"**Status:** {'Success' if r.get('success') else 'Failed'}")
        lines.append(f"**Security Score:** {r.get('security_score', 0)}/100")
        
        if r.get("error"):
            lines.append(f"**Error:** {r.get('error')}")
        
        if r.get("findings"):
            lines.append("")
            lines.append("**Top Findings:**")
            for finding in r.get("findings", [])[:5]:
                lines.append(f"- [{finding.get('severity', 'info').upper()}] {finding.get('title', 'Unknown')}")
        
        lines.append("")
        lines.append("---")
        lines.append("")
    
    return "\n".join(lines)


def export_jwt_analysis_markdown(result: JWTAnalysisResult) -> str:
    """Export JWT analysis as Markdown report."""
    lines = []
    
    lines.append("# JWT Token Analysis Report")
    lines.append("")
    lines.append(f"**Generated:** {time.strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    
    # Structure
    lines.append("## Token Structure")
    lines.append("")
    lines.append(f"**Valid Structure:** {'✅ Yes' if result.valid_structure else '❌ No'}")
    lines.append(f"**Algorithm:** `{result.algorithm}`")
    lines.append(f"**Expired:** {'❌ Yes' if result.is_expired else '✅ No'}")
    lines.append("")
    
    # Header
    lines.append("## Header")
    lines.append("")
    lines.append("```json")
    lines.append(json.dumps(result.header, indent=2))
    lines.append("```")
    lines.append("")
    
    # Payload
    lines.append("## Payload")
    lines.append("")
    lines.append("```json")
    lines.append(json.dumps(result.payload, indent=2))
    lines.append("```")
    lines.append("")
    
    # Claims
    lines.append("## Claims")
    lines.append("")
    lines.append(f"| Claim | Value |")
    lines.append("|-------|-------|")
    if result.issuer:
        lines.append(f"| Issuer (iss) | {result.issuer} |")
    if result.subject:
        lines.append(f"| Subject (sub) | {result.subject} |")
    if result.audience:
        lines.append(f"| Audience (aud) | {result.audience} |")
    if result.issued_at:
        lines.append(f"| Issued At (iat) | {result.issued_at} |")
    if result.expiry_time:
        lines.append(f"| Expires (exp) | {result.expiry_time} |")
    lines.append("")
    
    # Findings
    lines.append("## Security Findings")
    lines.append("")
    
    if result.findings:
        for finding in result.findings:
            severity_emoji = {
                "critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🔵", "info": "⚪"
            }
            sev = finding.get("severity", "info").lower()
            lines.append(f"### {severity_emoji.get(sev, '⚪')} {finding.get('title', 'Unknown')}")
            lines.append("")
            lines.append(f"**Severity:** {sev.upper()}")
            lines.append("")
            lines.append(finding.get("description", ""))
            if finding.get("evidence"):
                lines.append("")
                lines.append(f"**Evidence:** `{finding.get('evidence')}`")
            if finding.get("remediation"):
                lines.append("")
                lines.append(f"**Remediation:** {finding.get('remediation')}")
            lines.append("")
            lines.append("---")
            lines.append("")
    else:
        lines.append("✅ No security issues found!")
        lines.append("")
    
    return "\n".join(lines)


# =============================================================================
# PDF and Word Export Functions for API Tester
# =============================================================================

def export_test_result_pdf(result: APITestResult, title: str = "API Security Test Report") -> bytes:
    """Export API test result as a professional PDF report."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from io import BytesIO
        import time
    except ImportError:
        logger.error("reportlab not installed")
        return b"%PDF-1.4 placeholder - install reportlab for PDF generation"
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=16,
        textColor=colors.HexColor('#1a365d'),
        alignment=1
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.HexColor('#2c5282')
    )
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor('#4a5568')
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceBefore=4,
        spaceAfter=4,
        leading=14
    )
    
    # Severity colors
    severity_colors = {
        "critical": colors.HexColor('#dc2626'),
        "high": colors.HexColor('#ea580c'),
        "medium": colors.HexColor('#ca8a04'),
        "low": colors.HexColor('#2563eb'),
        "info": colors.HexColor('#6b7280')
    }
    
    story = []
    
    # Title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}", body_style))
    story.append(Paragraph(f"Target: {result.base_url}", body_style))
    story.append(Spacer(1, 24))
    
    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))
    
    # Score indicator
    score = result.security_score
    score_color = colors.green if score >= 80 else colors.orange if score >= 50 else colors.red
    story.append(Paragraph(f"<b>Security Score:</b> <font color='{score_color.hexval()}'>{score}/100</font>", body_style))
    story.append(Spacer(1, 12))
    
    # Summary table
    summary_data = [
        ['Metric', 'Value'],
        ['Endpoints Tested', str(result.endpoints_tested)],
        ['Total Findings', str(result.total_findings)],
        ['Critical', str(result.critical_count)],
        ['High', str(result.high_count)],
        ['Medium', str(result.medium_count)],
        ['Low', str(result.low_count)],
        ['Info', str(result.info_count)],
        ['Test Duration', f'{result.test_duration_seconds:.2f}s'],
    ]
    
    summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a365d')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 24))
    
    # OWASP API Top 10 Breakdown
    if result.owasp_api_breakdown:
        story.append(Paragraph("OWASP API Security Top 10", heading_style))
        owasp_data = [['Category', 'Findings']]
        for category, count in sorted(result.owasp_api_breakdown.items(), key=lambda x: -x[1]):
            owasp_data.append([category, str(count)])
        
        if len(owasp_data) > 1:
            owasp_table = Table(owasp_data, colWidths=[5*inch, 1.5*inch])
            owasp_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
            ]))
            story.append(owasp_table)
            story.append(Spacer(1, 24))
    
    # Security Findings
    story.append(Paragraph("Security Findings", heading_style))
    
    if result.all_findings:
        by_severity = {"critical": [], "high": [], "medium": [], "low": [], "info": []}
        for finding in result.all_findings:
            sev = finding.severity.lower()
            if sev in by_severity:
                by_severity[sev].append(finding)
        
        for severity in ["critical", "high", "medium", "low", "info"]:
            findings_list = by_severity[severity]
            if findings_list:
                sev_color = severity_colors.get(severity, colors.grey)
                story.append(Paragraph(
                    f"<font color='{sev_color.hexval()}'><b>{severity.upper()} ({len(findings_list)})</b></font>",
                    subheading_style
                ))
                
                for finding in findings_list:
                    story.append(Paragraph(f"<b>{finding.title}</b>", body_style))
                    story.append(Paragraph(f"<i>Category:</i> {finding.category}", body_style))
                    if finding.endpoint:
                        story.append(Paragraph(f"<i>Endpoint:</i> {finding.endpoint}", body_style))
                    story.append(Paragraph(finding.description[:500], body_style))
                    if finding.remediation:
                        story.append(Paragraph(f"<i>Remediation:</i> {finding.remediation[:300]}", body_style))
                    story.append(Spacer(1, 8))
    else:
        story.append(Paragraph("✓ No security issues found!", body_style))
    
    doc.build(story)
    return buffer.getvalue()


def export_test_result_docx(result: APITestResult, title: str = "API Security Test Report") -> bytes:
    """Export API test result as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from io import BytesIO
        import time
    except ImportError:
        logger.error("python-docx not installed")
        return b"Install python-docx for Word document generation"
    
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Target: {result.base_url}")
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    score = result.security_score
    score_text = doc.add_paragraph()
    score_text.add_run('Security Score: ').bold = True
    score_run = score_text.add_run(f'{score}/100')
    if score >= 80:
        score_run.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)  # Green
    elif score >= 50:
        score_run.font.color.rgb = RGBColor(0xf5, 0x97, 0x00)  # Orange
    else:
        score_run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)  # Red
    score_run.bold = True
    
    # Summary table
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Metric', 'Value']
    data = [
        ['Endpoints Tested', str(result.endpoints_tested)],
        ['Total Findings', str(result.total_findings)],
        ['Critical', str(result.critical_count)],
        ['High', str(result.high_count)],
        ['Medium', str(result.medium_count)],
        ['Low', str(result.low_count)],
        ['Info', str(result.info_count)],
        ['Test Duration', f'{result.test_duration_seconds:.2f}s'],
    ]
    
    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = headers[0]
    hdr[1].text = headers[1]
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    
    for i, row_data in enumerate(data):
        row = table.rows[i + 1].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()  # Spacer
    
    # OWASP API Top 10
    if result.owasp_api_breakdown:
        doc.add_heading('OWASP API Security Top 10', level=1)
        owasp_table = doc.add_table(rows=len(result.owasp_api_breakdown) + 1, cols=2)
        owasp_table.style = 'Table Grid'
        
        owasp_hdr = owasp_table.rows[0].cells
        owasp_hdr[0].text = 'Category'
        owasp_hdr[1].text = 'Findings'
        for cell in owasp_hdr:
            cell.paragraphs[0].runs[0].bold = True
        
        for i, (cat, count) in enumerate(sorted(result.owasp_api_breakdown.items(), key=lambda x: -x[1])):
            row = owasp_table.rows[i + 1].cells
            row[0].text = cat
            row[1].text = str(count)
    
    # Security Findings
    doc.add_heading('Security Findings', level=1)
    
    if result.all_findings:
        severity_colors = {
            "critical": RGBColor(0xdc, 0x26, 0x26),
            "high": RGBColor(0xea, 0x58, 0x0c),
            "medium": RGBColor(0xca, 0x8a, 0x04),
            "low": RGBColor(0x25, 0x63, 0xeb),
            "info": RGBColor(0x6b, 0x72, 0x80)
        }
        
        by_severity = {"critical": [], "high": [], "medium": [], "low": [], "info": []}
        for finding in result.all_findings:
            sev = finding.severity.lower()
            if sev in by_severity:
                by_severity[sev].append(finding)
        
        for severity in ["critical", "high", "medium", "low", "info"]:
            findings_list = by_severity[severity]
            if findings_list:
                heading = doc.add_heading(f'{severity.upper()} ({len(findings_list)})', level=2)
                for run in heading.runs:
                    run.font.color.rgb = severity_colors.get(severity, RGBColor(0, 0, 0))
                
                for finding in findings_list:
                    # Title
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(finding.title)
                    title_run.bold = True
                    
                    # Details
                    doc.add_paragraph(f"Category: {finding.category}")
                    if finding.endpoint:
                        doc.add_paragraph(f"Endpoint: {finding.endpoint}")
                    doc.add_paragraph(finding.description[:500])
                    if finding.remediation:
                        rem_para = doc.add_paragraph()
                        rem_para.add_run("Remediation: ").italic = True
                        rem_para.add_run(finding.remediation[:300])
                    doc.add_paragraph()  # Spacer
    else:
        doc.add_paragraph("✓ No security issues found!")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_auto_test_pdf(result: "AIAutoTestResult", title: str = "AI Auto-Test Security Report") -> bytes:
    """Export AI Auto-Test result as a professional PDF report."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from io import BytesIO
        import time
    except ImportError:
        logger.error("reportlab not installed")
        return b"%PDF-1.4 placeholder - install reportlab for PDF generation"
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'],
        fontSize=24, spaceAfter=16,
        textColor=colors.HexColor('#1a365d'), alignment=1
    )
    heading_style = ParagraphStyle(
        'CustomHeading', parent=styles['Heading2'],
        fontSize=16, spaceBefore=20, spaceAfter=10,
        textColor=colors.HexColor('#2c5282')
    )
    subheading_style = ParagraphStyle(
        'CustomSubheading', parent=styles['Heading3'],
        fontSize=12, spaceBefore=12, spaceAfter=6,
        textColor=colors.HexColor('#4a5568')
    )
    body_style = ParagraphStyle(
        'CustomBody', parent=styles['Normal'],
        fontSize=10, spaceBefore=4, spaceAfter=4, leading=14
    )
    
    severity_colors = {
        "critical": colors.HexColor('#dc2626'),
        "high": colors.HexColor('#ea580c'),
        "medium": colors.HexColor('#ca8a04'),
        "low": colors.HexColor('#2563eb'),
        "info": colors.HexColor('#6b7280')
    }
    
    story = []
    
    # Title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}", body_style))
    story.append(Spacer(1, 24))
    
    # Target Information
    story.append(Paragraph("Target Information", heading_style))
    target_data = [
        ['Property', 'Value'],
        ['Target', result.target],
        ['Type', result.target_type.upper()],
        ['Scan Duration', f'{result.scan_duration_seconds:.1f}s'],
        ['Endpoints Discovered', str(len(result.discovered_endpoints))],
        ['Services Found', str(len(result.discovered_services))],
    ]
    
    target_table = Table(target_data, colWidths=[2.5*inch, 4*inch])
    target_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a365d')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(target_table)
    story.append(Spacer(1, 24))
    
    # Discovered Services
    if result.discovered_services:
        story.append(Paragraph("Discovered Services", heading_style))
        svc_data = [['Port', 'Protocol', 'Status', 'Server']]
        for svc in result.discovered_services:
            svc_data.append([
                str(svc.get('port', '')),
                svc.get('scheme', '').upper(),
                str(svc.get('status_code', '')),
                svc.get('server', 'Unknown')[:30]
            ])
        
        svc_table = Table(svc_data, colWidths=[1*inch, 1.5*inch, 1*inch, 3*inch])
        svc_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
        ]))
        story.append(svc_table)
        story.append(Spacer(1, 24))
    
    # Discovered Endpoints
    if result.discovered_endpoints:
        story.append(Paragraph("Discovered Endpoints", heading_style))
        ep_data = [['Path', 'Status', 'Type']]
        for ep in result.discovered_endpoints[:15]:
            ep_type = 'JSON' if ep.get('is_json') else 'Other'
            if ep.get('requires_auth'):
                ep_type += ' (Auth)'
            ep_data.append([ep.get('path', ''), str(ep.get('status_code', '')), ep_type])
        
        if len(result.discovered_endpoints) > 15:
            ep_data.append([f'... and {len(result.discovered_endpoints) - 15} more', '', ''])
        
        ep_table = Table(ep_data, colWidths=[4*inch, 1*inch, 1.5*inch])
        ep_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
        ]))
        story.append(ep_table)
        story.append(Spacer(1, 24))
    
    # Security Assessment
    story.append(Paragraph("Security Assessment", heading_style))
    
    score = result.security_score
    score_color = colors.green if score >= 80 else colors.orange if score >= 50 else colors.red
    story.append(Paragraph(f"<b>Security Score:</b> <font color='{score_color.hexval()}'>{score}/100</font>", body_style))
    story.append(Spacer(1, 12))
    
    # Findings summary
    findings_data = [
        ['Severity', 'Count', 'Impact'],
        ['Critical', str(result.critical_count), 'Immediate action required'],
        ['High', str(result.high_count), 'Address urgently'],
        ['Medium', str(result.medium_count), 'Plan to fix soon'],
        ['Low', str(result.low_count), 'Consider fixing'],
        ['Info', str(result.info_count), 'Informational'],
    ]
    
    findings_table = Table(findings_data, colWidths=[1.5*inch, 1*inch, 3.5*inch])
    findings_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a365d')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(findings_table)
    story.append(Spacer(1, 24))
    
    # Top Security Issues
    if result.all_findings:
        story.append(Paragraph("Top Security Issues", heading_style))
        
        top_findings = sorted(
            result.all_findings,
            key=lambda x: {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}.get(x.get("severity", "info").lower(), 5)
        )[:10]
        
        for finding in top_findings:
            sev = finding.get('severity', 'info').lower()
            sev_color = severity_colors.get(sev, colors.grey)
            
            story.append(Paragraph(
                f"<font color='{sev_color.hexval()}'><b>[{sev.upper()}]</b></font> {finding.get('title', 'Unknown')}",
                subheading_style
            ))
            if finding.get('endpoint'):
                story.append(Paragraph(f"<i>Endpoint:</i> {finding.get('endpoint')}", body_style))
            if finding.get('description'):
                desc = finding['description'][:300] + "..." if len(finding.get('description', '')) > 300 else finding.get('description', '')
                story.append(Paragraph(desc, body_style))
            story.append(Spacer(1, 8))
    
    doc.build(story)
    return buffer.getvalue()


def export_auto_test_docx(result: "AIAutoTestResult", title: str = "AI Auto-Test Security Report") -> bytes:
    """Export AI Auto-Test result as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from io import BytesIO
        import time
    except ImportError:
        logger.error("python-docx not installed")
        return b"Install python-docx for Word document generation"
    
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Target Information
    doc.add_heading('Target Information', level=1)
    
    target_table = doc.add_table(rows=6, cols=2)
    target_table.style = 'Table Grid'
    
    target_data = [
        ['Property', 'Value'],
        ['Target', result.target],
        ['Type', result.target_type.upper()],
        ['Scan Duration', f'{result.scan_duration_seconds:.1f}s'],
        ['Endpoints Discovered', str(len(result.discovered_endpoints))],
        ['Services Found', str(len(result.discovered_services))],
    ]
    
    for i, row_data in enumerate(target_data):
        row = target_table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        if i == 0:
            for cell in row:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Discovered Services
    if result.discovered_services:
        doc.add_heading('Discovered Services', level=1)
        svc_table = doc.add_table(rows=len(result.discovered_services) + 1, cols=4)
        svc_table.style = 'Table Grid'
        
        svc_hdr = svc_table.rows[0].cells
        for j, h in enumerate(['Port', 'Protocol', 'Status', 'Server']):
            svc_hdr[j].text = h
            svc_hdr[j].paragraphs[0].runs[0].bold = True
        
        for i, svc in enumerate(result.discovered_services):
            row = svc_table.rows[i + 1].cells
            row[0].text = str(svc.get('port', ''))
            row[1].text = svc.get('scheme', '').upper()
            row[2].text = str(svc.get('status_code', ''))
            row[3].text = svc.get('server', 'Unknown')[:30]
    
    # Discovered Endpoints
    if result.discovered_endpoints:
        doc.add_heading('Discovered Endpoints', level=1)
        
        endpoints_to_show = result.discovered_endpoints[:20]
        ep_table = doc.add_table(rows=len(endpoints_to_show) + 1, cols=3)
        ep_table.style = 'Table Grid'
        
        ep_hdr = ep_table.rows[0].cells
        for j, h in enumerate(['Path', 'Status', 'Type']):
            ep_hdr[j].text = h
            ep_hdr[j].paragraphs[0].runs[0].bold = True
        
        for i, ep in enumerate(endpoints_to_show):
            row = ep_table.rows[i + 1].cells
            row[0].text = ep.get('path', '')
            row[1].text = str(ep.get('status_code', ''))
            ep_type = 'JSON' if ep.get('is_json') else 'Other'
            if ep.get('requires_auth'):
                ep_type += ' (Auth)'
            row[2].text = ep_type
        
        if len(result.discovered_endpoints) > 20:
            doc.add_paragraph(f"... and {len(result.discovered_endpoints) - 20} more endpoints")
    
    # Security Assessment
    doc.add_heading('Security Assessment', level=1)
    
    score = result.security_score
    score_para = doc.add_paragraph()
    score_para.add_run('Security Score: ').bold = True
    score_run = score_para.add_run(f'{score}/100')
    if score >= 80:
        score_run.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)
    elif score >= 50:
        score_run.font.color.rgb = RGBColor(0xf5, 0x97, 0x00)
    else:
        score_run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
    score_run.bold = True
    
    # Findings Summary
    findings_table = doc.add_table(rows=6, cols=3)
    findings_table.style = 'Table Grid'
    
    findings_data = [
        ['Severity', 'Count', 'Impact'],
        ['Critical', str(result.critical_count), 'Immediate action required'],
        ['High', str(result.high_count), 'Address urgently'],
        ['Medium', str(result.medium_count), 'Plan to fix soon'],
        ['Low', str(result.low_count), 'Consider fixing'],
        ['Info', str(result.info_count), 'Informational'],
    ]
    
    for i, row_data in enumerate(findings_data):
        row = findings_table.rows[i].cells
        for j, val in enumerate(row_data):
            row[j].text = val
            if i == 0:
                row[j].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Top Security Issues
    if result.all_findings:
        doc.add_heading('Top Security Issues', level=1)
        
        severity_colors = {
            "critical": RGBColor(0xdc, 0x26, 0x26),
            "high": RGBColor(0xea, 0x58, 0x0c),
            "medium": RGBColor(0xca, 0x8a, 0x04),
            "low": RGBColor(0x25, 0x63, 0xeb),
            "info": RGBColor(0x6b, 0x72, 0x80)
        }
        
        top_findings = sorted(
            result.all_findings,
            key=lambda x: {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}.get(x.get("severity", "info").lower(), 5)
        )[:10]
        
        for finding in top_findings:
            sev = finding.get('severity', 'info').lower()
            
            title_para = doc.add_paragraph()
            sev_run = title_para.add_run(f'[{sev.upper()}] ')
            sev_run.bold = True
            sev_run.font.color.rgb = severity_colors.get(sev, RGBColor(0, 0, 0))
            title_para.add_run(finding.get('title', 'Unknown')).bold = True
            
            if finding.get('endpoint'):
                doc.add_paragraph(f"Endpoint: {finding.get('endpoint')}")
            if finding.get('description'):
                desc = finding['description'][:300]
                doc.add_paragraph(desc)
            doc.add_paragraph()
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# =============================================================================
# WebSocket Test Export Functions
# =============================================================================

def export_websocket_markdown(result: "WebSocketTestResult", title: str = "WebSocket Security Test Report") -> str:
    """Export WebSocket test result as a Markdown report."""
    import time
    
    lines = []
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"*Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}*")
    lines.append("")
    lines.append("---")
    lines.append("")
    
    # Target Information
    lines.append("## Target Information")
    lines.append("")
    lines.append("| Property | Value |")
    lines.append("|----------|-------|")
    lines.append(f"| **URL** | `{result.url}` |")
    lines.append(f"| **Connected** | {'✅ Yes' if result.connected else '❌ No'} |")
    lines.append(f"| **Connection Time** | {result.connection_time_ms:.0f} ms |")
    lines.append(f"| **Protocol** | {result.protocol or 'N/A'} |")
    lines.append(f"| **Test Duration** | {result.test_duration_seconds:.1f}s |")
    lines.append("")
    
    # Security Score
    lines.append("## Security Assessment")
    lines.append("")
    score = result.security_score
    if score >= 80:
        score_label = "🟢 GOOD"
    elif score >= 60:
        score_label = "🟡 MODERATE"
    elif score >= 40:
        score_label = "🟠 CONCERNING"
    else:
        score_label = "🔴 CRITICAL"
    
    lines.append(f"**Security Score:** {score}/100 ({score_label})")
    lines.append("")
    
    # Findings Summary
    if result.findings:
        lines.append("## Security Findings")
        lines.append("")
        lines.append(f"**Total Findings:** {len(result.findings)}")
        lines.append("")
        
        severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
        sorted_findings = sorted(result.findings, key=lambda f: severity_order.get(f.get("severity", "info").lower(), 5))
        
        for i, finding in enumerate(sorted_findings, 1):
            sev = finding.get("severity", "info").upper()
            sev_emoji = {"CRITICAL": "🔴", "HIGH": "🟠", "MEDIUM": "🟡", "LOW": "🔵", "INFO": "⚪"}.get(sev, "⚪")
            
            lines.append(f"### {i}. {sev_emoji} [{sev}] {finding.get('title', 'Unknown')}")
            lines.append("")
            if finding.get("description"):
                lines.append(finding["description"])
                lines.append("")
            if finding.get("evidence"):
                lines.append("**Evidence:**")
                lines.append(f"```")
                lines.append(finding["evidence"])
                lines.append(f"```")
                lines.append("")
            if finding.get("remediation"):
                lines.append(f"**Remediation:** {finding['remediation']}")
                lines.append("")
            if finding.get("cwe"):
                lines.append(f"**CWE:** {finding['cwe']}")
            if finding.get("owasp_api"):
                lines.append(f"**OWASP API:** {finding['owasp_api']}")
            lines.append("")
    else:
        lines.append("## Security Findings")
        lines.append("")
        lines.append("✅ No security issues found! The WebSocket endpoint passed all tests.")
        lines.append("")
    
    # OWASP Breakdown
    if result.owasp_api_breakdown:
        lines.append("## OWASP API Security Breakdown")
        lines.append("")
        lines.append("| Category | Count |")
        lines.append("|----------|-------|")
        for cat, count in sorted(result.owasp_api_breakdown.items(), key=lambda x: -x[1]):
            if count > 0:
                lines.append(f"| {cat} | {count} |")
        lines.append("")
    
    lines.append("---")
    lines.append("")
    lines.append("*Report generated by VRAgent WebSocket Security Tester*")
    
    return "\n".join(lines)


def export_websocket_pdf(result: "WebSocketTestResult", title: str = "WebSocket Security Test Report") -> bytes:
    """Export WebSocket test result as a PDF report."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from io import BytesIO
        import time
    except ImportError:
        return b"%PDF-1.4 placeholder - install reportlab"
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, spaceAfter=16, textColor=colors.HexColor('#1a365d'), alignment=1)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=16, spaceBefore=20, spaceAfter=10, textColor=colors.HexColor('#2c5282'))
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, spaceBefore=4, spaceAfter=4)
    
    severity_colors = {
        "critical": colors.HexColor('#dc2626'),
        "high": colors.HexColor('#ea580c'),
        "medium": colors.HexColor('#ca8a04'),
        "low": colors.HexColor('#2563eb'),
        "info": colors.HexColor('#6b7280')
    }
    
    story = []
    
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}", body_style))
    story.append(Spacer(1, 24))
    
    # Target Information
    story.append(Paragraph("Target Information", heading_style))
    target_data = [
        ['Property', 'Value'],
        ['URL', result.url],
        ['Connected', 'Yes' if result.connected else 'No'],
        ['Connection Time', f'{result.connection_time_ms:.0f} ms'],
        ['Protocol', result.protocol or 'N/A'],
        ['Test Duration', f'{result.test_duration_seconds:.1f}s'],
    ]
    
    target_table = Table(target_data, colWidths=[2*inch, 4.5*inch])
    target_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a365d')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(target_table)
    story.append(Spacer(1, 24))
    
    # Security Score
    story.append(Paragraph("Security Assessment", heading_style))
    score = result.security_score
    score_color = colors.green if score >= 80 else colors.orange if score >= 50 else colors.red
    story.append(Paragraph(f"<b>Security Score:</b> <font color='{score_color.hexval()}'>{score}/100</font>", body_style))
    story.append(Spacer(1, 12))
    
    # Findings
    if result.findings:
        story.append(Paragraph(f"Security Findings ({len(result.findings)})", heading_style))
        
        for finding in result.findings:
            sev = finding.get('severity', 'info').lower()
            sev_color = severity_colors.get(sev, colors.grey)
            
            story.append(Paragraph(
                f"<font color='{sev_color.hexval()}'><b>[{sev.upper()}]</b></font> {finding.get('title', 'Unknown')}",
                body_style
            ))
            if finding.get('description'):
                story.append(Paragraph(finding['description'], body_style))
            if finding.get('remediation'):
                story.append(Paragraph(f"<i>Fix:</i> {finding['remediation']}", body_style))
            story.append(Spacer(1, 8))
    else:
        story.append(Paragraph("No security issues found!", body_style))
    
    doc.build(story)
    return buffer.getvalue()


def export_websocket_docx(result: "WebSocketTestResult", title: str = "WebSocket Security Test Report") -> bytes:
    """Export WebSocket test result as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        import time
    except ImportError:
        return b"Install python-docx for Word generation"
    
    doc = Document()
    
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Target Information
    doc.add_heading('Target Information', level=1)
    
    target_table = doc.add_table(rows=6, cols=2)
    target_table.style = 'Table Grid'
    
    target_data = [
        ['Property', 'Value'],
        ['URL', result.url],
        ['Connected', 'Yes' if result.connected else 'No'],
        ['Connection Time', f'{result.connection_time_ms:.0f} ms'],
        ['Protocol', result.protocol or 'N/A'],
        ['Test Duration', f'{result.test_duration_seconds:.1f}s'],
    ]
    
    for i, row_data in enumerate(target_data):
        row = target_table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        if i == 0:
            for cell in row:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Security Score
    doc.add_heading('Security Assessment', level=1)
    score_para = doc.add_paragraph()
    score_para.add_run('Security Score: ').bold = True
    score_run = score_para.add_run(f'{result.security_score}/100')
    score_run.bold = True
    if result.security_score >= 80:
        score_run.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e)
    elif result.security_score >= 50:
        score_run.font.color.rgb = RGBColor(0xf5, 0x97, 0x00)
    else:
        score_run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
    
    # Findings
    doc.add_heading('Security Findings', level=1)
    
    if result.findings:
        severity_colors = {
            "critical": RGBColor(0xdc, 0x26, 0x26),
            "high": RGBColor(0xea, 0x58, 0x0c),
            "medium": RGBColor(0xca, 0x8a, 0x04),
            "low": RGBColor(0x25, 0x63, 0xeb),
            "info": RGBColor(0x6b, 0x72, 0x80)
        }
        
        for finding in result.findings:
            sev = finding.get('severity', 'info').lower()
            
            title_para = doc.add_paragraph()
            sev_run = title_para.add_run(f'[{sev.upper()}] ')
            sev_run.bold = True
            sev_run.font.color.rgb = severity_colors.get(sev, RGBColor(0, 0, 0))
            title_para.add_run(finding.get('title', 'Unknown')).bold = True
            
            if finding.get('description'):
                doc.add_paragraph(finding['description'])
            if finding.get('remediation'):
                doc.add_paragraph(f"Remediation: {finding['remediation']}")
            doc.add_paragraph()
    else:
        doc.add_paragraph("✓ No security issues found!")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_batch_result_pdf(result: BatchTestResult, title: str = "Batch API Test Report") -> bytes:
    """Export batch test result as a PDF report."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from io import BytesIO
        import time
    except ImportError:
        return b"%PDF-1.4 placeholder"
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, spaceAfter=16, textColor=colors.HexColor('#1a365d'), alignment=1)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=16, spaceBefore=20, spaceAfter=10, textColor=colors.HexColor('#2c5282'))
    body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=10, spaceBefore=4, spaceAfter=4)
    
    story = []
    
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}", body_style))
    story.append(Spacer(1, 24))
    
    # Summary
    story.append(Paragraph("Summary", heading_style))
    summary_data = [
        ['Metric', 'Value'],
        ['Total Targets', str(result.total_targets)],
        ['Successful', str(result.successful)],
        ['Failed', str(result.failed)],
        ['Total Findings', str(result.total_findings)],
        ['Critical', str(result.critical_findings)],
        ['High', str(result.high_findings)],
        ['Duration', f'{result.scan_duration_seconds:.2f}s'],
    ]
    
    summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a365d')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 24))
    
    # Results by Target
    story.append(Paragraph("Results by Target", heading_style))
    results_data = [['Target', 'Status', 'Score', 'Findings', 'Critical']]
    for r in sorted(result.results, key=lambda x: x.get("security_score", 0)):
        status = "✓" if r.get("success") else "✗"
        results_data.append([
            r.get('name', r.get('target', ''))[:40],
            status,
            f"{r.get('security_score', 0)}/100",
            str(r.get('total_findings', 0)),
            str(r.get('critical_count', 0))
        ])
    
    results_table = Table(results_data, colWidths=[2.5*inch, 0.8*inch, 1*inch, 1*inch, 1*inch])
    results_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f7fafc')]),
    ]))
    story.append(results_table)
    
    doc.build(story)
    return buffer.getvalue()


def export_batch_result_docx(result: BatchTestResult, title: str = "Batch API Test Report") -> bytes:
    """Export batch test result as a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        import time
    except ImportError:
        return b"Install python-docx"
    
    doc = Document()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Summary
    doc.add_heading('Summary', level=1)
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ['Metric', 'Value'],
        ['Total Targets', str(result.total_targets)],
        ['Successful', str(result.successful)],
        ['Failed', str(result.failed)],
        ['Total Findings', str(result.total_findings)],
        ['Critical', str(result.critical_findings)],
        ['High', str(result.high_findings)],
        ['Duration', f'{result.scan_duration_seconds:.2f}s'],
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i].cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        if i == 0:
            for cell in row:
                cell.paragraphs[0].runs[0].bold = True
    
    # Results
    doc.add_heading('Results by Target', level=1)
    results_table = doc.add_table(rows=len(result.results) + 1, cols=5)
    results_table.style = 'Table Grid'
    
    hdr = results_table.rows[0].cells
    for j, h in enumerate(['Target', 'Status', 'Score', 'Findings', 'Critical']):
        hdr[j].text = h
        hdr[j].paragraphs[0].runs[0].bold = True
    
    for i, r in enumerate(sorted(result.results, key=lambda x: x.get("security_score", 0))):
        row = results_table.rows[i + 1].cells
        row[0].text = r.get('name', r.get('target', ''))[:40]
        row[1].text = "✓" if r.get("success") else "✗"
        row[2].text = f"{r.get('security_score', 0)}/100"
        row[3].text = str(r.get('total_findings', 0))
        row[4].text = str(r.get('critical_count', 0))
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_jwt_analysis_pdf(result: JWTAnalysisResult) -> bytes:
    """Export JWT analysis as a PDF report."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted
        from io import BytesIO
        import time
    except ImportError:
        return b"%PDF-1.4 placeholder"
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, spaceAfter=16, textColor=colors.HexColor('#1a365d'), alignment=1)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=16, spaceBefore=20, spaceAfter=10, textColor=colors.HexColor('#2c5282'))
    body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontSize=10, spaceBefore=4, spaceAfter=4)
    code_style = ParagraphStyle('CodeStyle', parent=styles['Code'], fontSize=8, backColor=colors.HexColor('#f7fafc'))
    
    story = []
    
    story.append(Paragraph("JWT Token Analysis Report", title_style))
    story.append(Paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}", body_style))
    story.append(Spacer(1, 24))
    
    # Token Structure
    story.append(Paragraph("Token Structure", heading_style))
    valid_color = colors.green if result.valid_structure else colors.red
    story.append(Paragraph(f"<b>Valid Structure:</b> <font color='{valid_color.hexval()}'>{'Yes' if result.valid_structure else 'No'}</font>", body_style))
    story.append(Paragraph(f"<b>Algorithm:</b> {result.algorithm}", body_style))
    expired_color = colors.red if result.is_expired else colors.green
    story.append(Paragraph(f"<b>Expired:</b> <font color='{expired_color.hexval()}'>{'Yes' if result.is_expired else 'No'}</font>", body_style))
    story.append(Spacer(1, 12))
    
    # Header
    story.append(Paragraph("Header", heading_style))
    story.append(Preformatted(json.dumps(result.header, indent=2), code_style))
    story.append(Spacer(1, 12))
    
    # Payload
    story.append(Paragraph("Payload", heading_style))
    story.append(Preformatted(json.dumps(result.payload, indent=2), code_style))
    story.append(Spacer(1, 12))
    
    # Claims
    story.append(Paragraph("Claims", heading_style))
    claims_data = [['Claim', 'Value']]
    if result.issuer:
        claims_data.append(['Issuer (iss)', result.issuer])
    if result.subject:
        claims_data.append(['Subject (sub)', result.subject])
    if result.audience:
        claims_data.append(['Audience (aud)', str(result.audience)])
    if result.issued_at:
        claims_data.append(['Issued At (iat)', result.issued_at])
    if result.expiry_time:
        claims_data.append(['Expires (exp)', result.expiry_time])
    
    if len(claims_data) > 1:
        claims_table = Table(claims_data, colWidths=[2*inch, 4.5*inch])
        claims_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(claims_table)
    story.append(Spacer(1, 24))
    
    # Findings
    story.append(Paragraph("Security Findings", heading_style))
    if result.findings:
        severity_colors = {
            "critical": colors.HexColor('#dc2626'),
            "high": colors.HexColor('#ea580c'),
            "medium": colors.HexColor('#ca8a04'),
            "low": colors.HexColor('#2563eb'),
            "info": colors.HexColor('#6b7280')
        }
        for finding in result.findings:
            sev = finding.get('severity', 'info').lower()
            sev_color = severity_colors.get(sev, colors.grey)
            story.append(Paragraph(f"<font color='{sev_color.hexval()}'><b>[{sev.upper()}]</b></font> {finding.get('title', 'Unknown')}", body_style))
            story.append(Paragraph(finding.get('description', ''), body_style))
            story.append(Spacer(1, 8))
    else:
        story.append(Paragraph("✓ No security issues found!", body_style))
    
    doc.build(story)
    return buffer.getvalue()


def export_jwt_analysis_docx(result: JWTAnalysisResult) -> bytes:
    """Export JWT analysis as a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        import time
    except ImportError:
        return b"Install python-docx"
    
    doc = Document()
    doc.add_heading("JWT Token Analysis Report", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Token Structure
    doc.add_heading('Token Structure', level=1)
    
    struct_para = doc.add_paragraph()
    struct_para.add_run('Valid Structure: ').bold = True
    valid_run = struct_para.add_run('Yes' if result.valid_structure else 'No')
    valid_run.font.color.rgb = RGBColor(0x22, 0xc5, 0x5e) if result.valid_structure else RGBColor(0xdc, 0x26, 0x26)
    
    doc.add_paragraph(f"Algorithm: {result.algorithm}")
    
    exp_para = doc.add_paragraph()
    exp_para.add_run('Expired: ').bold = True
    exp_run = exp_para.add_run('Yes' if result.is_expired else 'No')
    exp_run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26) if result.is_expired else RGBColor(0x22, 0xc5, 0x5e)
    
    # Header
    doc.add_heading('Header', level=1)
    doc.add_paragraph(json.dumps(result.header, indent=2))
    
    # Payload
    doc.add_heading('Payload', level=1)
    doc.add_paragraph(json.dumps(result.payload, indent=2))
    
    # Claims
    doc.add_heading('Claims', level=1)
    claims = []
    if result.issuer:
        claims.append(['Issuer (iss)', result.issuer])
    if result.subject:
        claims.append(['Subject (sub)', result.subject])
    if result.audience:
        claims.append(['Audience (aud)', str(result.audience)])
    if result.issued_at:
        claims.append(['Issued At (iat)', result.issued_at])
    if result.expiry_time:
        claims.append(['Expires (exp)', result.expiry_time])
    
    if claims:
        table = doc.add_table(rows=len(claims) + 1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Claim'
        hdr[1].text = 'Value'
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        for i, (claim, val) in enumerate(claims):
            row = table.rows[i + 1].cells
            row[0].text = claim
            row[1].text = val
    
    # Findings
    doc.add_heading('Security Findings', level=1)
    if result.findings:
        severity_colors = {
            "critical": RGBColor(0xdc, 0x26, 0x26),
            "high": RGBColor(0xea, 0x58, 0x0c),
            "medium": RGBColor(0xca, 0x8a, 0x04),
            "low": RGBColor(0x25, 0x63, 0xeb),
            "info": RGBColor(0x6b, 0x72, 0x80)
        }
        for finding in result.findings:
            sev = finding.get('severity', 'info').lower()
            title_para = doc.add_paragraph()
            sev_run = title_para.add_run(f'[{sev.upper()}] ')
            sev_run.bold = True
            sev_run.font.color.rgb = severity_colors.get(sev, RGBColor(0, 0, 0))
            title_para.add_run(finding.get('title', 'Unknown')).bold = True
            doc.add_paragraph(finding.get('description', ''))
    else:
        doc.add_paragraph("✓ No security issues found!")
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ============================================================================

@dataclass
class AIAutoTestResult:
    """Result of AI-driven automated security testing."""
    target: str
    target_type: str  # "ip", "url", "domain"
    discovered_services: List[Dict[str, Any]] = field(default_factory=list)
    discovered_endpoints: List[Dict[str, Any]] = field(default_factory=list)
    test_results: List[Dict[str, Any]] = field(default_factory=list)
    all_findings: List[Dict[str, Any]] = field(default_factory=list)
    total_findings: int = 0
    critical_count: int = 0
    high_count: int = 0
    medium_count: int = 0
    low_count: int = 0
    info_count: int = 0
    security_score: int = 100
    ai_summary: str = ""
    scan_duration_seconds: float = 0
    error: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "target": self.target,
            "target_type": self.target_type,
            "discovered_services": self.discovered_services,
            "discovered_endpoints": self.discovered_endpoints,
            "test_results": self.test_results,
            "all_findings": self.all_findings,
            "total_findings": self.total_findings,
            "critical_count": self.critical_count,
            "high_count": self.high_count,
            "medium_count": self.medium_count,
            "low_count": self.low_count,
            "info_count": self.info_count,
            "security_score": self.security_score,
            "ai_summary": self.ai_summary,
            "scan_duration_seconds": self.scan_duration_seconds,
            "error": self.error,
        }


# Common API paths to probe
COMMON_API_PATHS = [
    # Root and API versions
    "/", "/api", "/api/v1", "/api/v2", "/api/v3", "/v1", "/v2", "/v3",
    # Health and status
    "/health", "/healthz", "/health/live", "/health/ready", "/healthcheck",
    "/status", "/info", "/version", "/ping", "/_health", "/__health",
    # API Documentation
    "/swagger", "/swagger.json", "/swagger.yaml", "/swagger/v1/swagger.json",
    "/openapi", "/openapi.json", "/openapi.yaml", "/api-docs", "/api-docs.json",
    "/docs", "/redoc", "/rapidoc", "/scalar",
    # GraphQL
    "/graphql", "/graphiql", "/playground", "/altair", "/gql",
    "/graphql/console", "/api/graphql", "/v1/graphql",
    # Authentication
    "/admin", "/admin/login", "/login", "/signin", "/signup",
    "/auth", "/oauth", "/oauth/token", "/token", "/jwt",
    "/auth/login", "/api/auth", "/api/login", "/.auth",
    "/forgot-password", "/reset-password", "/register",
    # Common resources
    "/users", "/user", "/accounts", "/account", "/profile", "/me",
    "/products", "/items", "/orders", "/data", "/resources",
    "/customers", "/clients", "/employees", "/members",
    # Configuration and debug (sensitive)
    "/config", "/settings", "/env", "/debug", "/trace",
    "/metrics", "/prometheus", "/stats", "/statistics",
    "/actuator", "/actuator/health", "/actuator/info", "/actuator/env",
    "/management", "/manage", "/console", "/dashboard",
    # Well-known and standard files
    "/.well-known/openid-configuration", "/.well-known/security.txt",
    "/.env", "/robots.txt", "/sitemap.xml", "/favicon.ico",
    "/wp-json", "/wp-admin", "/xmlrpc.php",  # WordPress
    # Internal/debug endpoints
    "/internal", "/private", "/test", "/demo", "/dev",
    "/phpinfo.php", "/server-status", "/server-info",
    "/.git/config", "/.svn/entries", "/backup", "/backups",
]

# Common ports for web services
COMMON_WEB_PORTS = [80, 443, 8080, 8443, 3000, 5000, 8000, 9000, 4000, 5001, 8888]


def detect_target_type(target: str) -> Tuple[str, str]:
    """
    Detect if target is IP, URL, domain, or CIDR network.
    Returns (target_type, normalized_target)
    """
    import ipaddress
    
    # Remove protocol if present
    cleaned = target.strip()
    
    # Check if it's already a URL
    if cleaned.startswith(("http://", "https://")):
        return ("url", cleaned)
    
    # Check if it's a CIDR network notation (e.g., 192.168.1.0/24)
    if "/" in cleaned and not cleaned.startswith("/"):
        try:
            network = ipaddress.ip_network(cleaned, strict=False)
            return ("cidr", cleaned)
        except ValueError:
            pass
    
    # Check if it's an IP range (e.g., 192.168.1.1-192.168.1.254)
    if "-" in cleaned and "." in cleaned:
        parts = cleaned.split("-")
        if len(parts) == 2:
            try:
                ipaddress.ip_address(parts[0].strip())
                ipaddress.ip_address(parts[1].strip())
                return ("ip_range", cleaned)
            except ValueError:
                pass
    
    # Check if it's an IP address (with optional port)
    ip_part = cleaned.split(":")[0]
    try:
        ipaddress.ip_address(ip_part)
        return ("ip", cleaned)
    except ValueError:
        pass
    
    # Check if it looks like a domain
    if "." in cleaned and not cleaned.startswith("/"):
        return ("domain", cleaned)
    
    # Default to treating as partial URL
    return ("url", cleaned)


async def probe_port(host: str, port: int, timeout: float = 2.0, verify_ssl: bool = True) -> Optional[Dict[str, Any]]:
    """
    Probe a single port to check if a web service is running.
    Uses fast connection-based checking for efficiency.

    Args:
        host: Target host to probe
        port: Port number to probe
        timeout: Connection timeout in seconds
        verify_ssl: Whether to verify SSL certificates (default True)
    """
    for scheme in ["https", "http"]:
        url = f"{scheme}://{host}:{port}/"
        try:
            async with httpx.AsyncClient(verify=verify_ssl, timeout=timeout) as client:
                resp = await client.get(url, follow_redirects=True)
                return {
                    "port": port,
                    "scheme": scheme,
                    "url": f"{scheme}://{host}:{port}",
                    "status_code": resp.status_code,
                    "content_type": resp.headers.get("content-type", ""),
                    "server": resp.headers.get("server", ""),
                    "content_length": len(resp.content),
                    "host": host,
                }
        except Exception:
            continue
    return None


async def fast_scan_network(
    network_str: str,
    ports: Optional[List[int]] = None,
    timeout: float = 1.0,
    max_concurrent: int = 200,
    max_hosts: int = 256,
) -> List[Dict[str, Any]]:
    """
    Fast network scanner optimized for speed on local networks.
    Uses aggressive concurrency and short timeouts.
    
    Args:
        network_str: CIDR notation (e.g., "192.168.1.0/24") or IP range
        ports: Ports to scan (default: common web ports)
        timeout: Connection timeout per host (keep low for speed)
        max_concurrent: Maximum concurrent connections (high = faster)
        max_hosts: Maximum number of hosts to scan
        
    Returns:
        List of discovered services
    """
    import ipaddress
    
    if ports is None:
        ports = [80, 443, 8080, 8443, 3000, 5000, 8000]
    
    # Parse network
    hosts = []
    try:
        if "/" in network_str:
            network = ipaddress.ip_network(network_str, strict=False)
            hosts = list(network.hosts())[:max_hosts]
        elif "-" in network_str:
            start_ip, end_ip = network_str.split("-")
            start = ipaddress.ip_address(start_ip.strip())
            end = ipaddress.ip_address(end_ip.strip())
            current = start
            while current <= end and len(hosts) < max_hosts:
                hosts.append(current)
                current = ipaddress.ip_address(int(current) + 1)
        else:
            hosts = [ipaddress.ip_address(network_str)]
    except Exception as e:
        logger.error(f"Failed to parse network {network_str}: {e}")
        return []
    
    logger.info(f"Fast network scan: {len(hosts)} hosts, {len(ports)} ports, {len(hosts) * len(ports)} total probes")
    
    discovered = []
    semaphore = asyncio.Semaphore(max_concurrent)
    
    async def check_host_port(ip: str, port: int) -> Optional[Dict[str, Any]]:
        async with semaphore:
            # Try HTTP first (faster), then HTTPS
            for scheme in ["http", "https"]:
                url = f"{scheme}://{ip}:{port}/"
                try:
                    async with httpx.AsyncClient(
                        verify=True,  # SSL verification enabled
                        timeout=httpx.Timeout(timeout, connect=timeout)
                    ) as client:
                        resp = await client.get(url, follow_redirects=False)
                        return {
                            "host": ip,
                            "port": port,
                            "scheme": scheme,
                            "url": f"{scheme}://{ip}:{port}",
                            "status_code": resp.status_code,
                            "content_type": resp.headers.get("content-type", ""),
                            "server": resp.headers.get("server", ""),
                        }
                except Exception:
                    continue
            return None
    
    # Create all tasks
    tasks = []
    for host in hosts:
        for port in ports:
            tasks.append(check_host_port(str(host), port))
    
    # Run with progress tracking
    start_time = time.time()
    results = await asyncio.gather(*tasks, return_exceptions=True)
    elapsed = time.time() - start_time
    
    for r in results:
        if isinstance(r, dict):
            discovered.append(r)
    
    logger.info(f"Fast network scan completed in {elapsed:.1f}s: found {len(discovered)} services")
    return discovered


async def discover_endpoints(base_url: str, timeout: float = 2.0, max_paths: int = 30, verify_ssl: bool = True) -> List[Dict[str, Any]]:
    """
    Discover API endpoints by probing common paths.
    Uses short timeouts and limits paths for speed.

    Args:
        base_url: Base URL to probe
        timeout: Request timeout in seconds
        max_paths: Maximum number of paths to check
        verify_ssl: Whether to verify SSL certificates (default True)
    """
    discovered = []
    paths_to_check = COMMON_API_PATHS[:max_paths]  # Limit paths for speed

    async with httpx.AsyncClient(
        verify=verify_ssl,
        timeout=httpx.Timeout(timeout, connect=timeout)
    ) as client:
        for path in paths_to_check:
            url = urljoin(base_url.rstrip("/") + "/", path.lstrip("/"))
            try:
                resp = await client.get(url, follow_redirects=False)
                if resp.status_code not in [404, 502, 503]:
                    discovered.append({
                        "path": path,
                        "url": url,
                        "method": "GET",
                        "status_code": resp.status_code,
                        "content_type": resp.headers.get("content-type", ""),
                        "content_length": len(resp.content),
                        "is_json": "json" in resp.headers.get("content-type", "").lower(),
                        "is_html": "html" in resp.headers.get("content-type", "").lower(),
                        "requires_auth": resp.status_code in [401, 403],
                    })
            except Exception:
                continue
    
    return discovered


async def ai_auto_test(
    target: str,
    ports: Optional[List[int]] = None,
    probe_common_paths: bool = True,
    run_security_tests: bool = True,
    max_endpoints: int = 20,
    timeout: float = 10.0,
    proxy_url: Optional[str] = None,
    network_timeout: float = 1.0,
    max_concurrent: int = 200,
    overall_timeout: float = 120.0,  # 2 minute max for entire operation
) -> AIAutoTestResult:
    """
    AI-driven automated security testing.
    
    1. Detects target type (IP, URL, domain, CIDR network, IP range)
    2. Discovers running services/ports (fast parallel scan for networks)
    3. Probes for common API endpoints
    4. Runs security tests on discovered endpoints
    5. Aggregates findings and generates AI summary
    
    For network scans (/24 = 256 hosts), uses aggressive parallel scanning
    with short timeouts for speed.
    
    Args:
        overall_timeout: Maximum time in seconds for the entire operation (default: 120s)
    """
    start_time = time.time()
    result = AIAutoTestResult(target=target, target_type="unknown")
    
    def check_timeout():
        """Check if we've exceeded overall timeout."""
        elapsed = time.time() - start_time
        if elapsed > overall_timeout:
            raise asyncio.TimeoutError(f"AI Auto-Test exceeded {overall_timeout}s timeout")
        return elapsed
    
    try:
        # Step 1: Detect target type
        target_type, normalized = detect_target_type(target)
        result.target_type = target_type
        logger.info(f"AI Auto-Test: Target type={target_type}, normalized={normalized}")
        
        # Step 2: Build base URLs to test
        base_urls = []
        
        if target_type == "url":
            base_urls.append(normalized)
            
        elif target_type in ("cidr", "ip_range"):
            # FAST NETWORK SCAN for CIDR ranges like 192.168.1.0/24
            logger.info(f"AI Auto-Test: Starting fast network scan for {normalized}")
            port_list = ports or [80, 443, 8080, 8443, 3000, 5000, 8000]
            
            discovered_services = await fast_scan_network(
                network_str=normalized,
                ports=port_list,
                timeout=network_timeout,  # Short timeout for speed
                max_concurrent=max_concurrent,  # High concurrency
                max_hosts=256,  # Limit to /24 equivalent
            )
            
            # Add discovered services to result
            for svc in discovered_services:
                result.discovered_services.append(svc)
                if svc["url"] not in base_urls:
                    base_urls.append(svc["url"])
            
            logger.info(f"AI Auto-Test: Network scan found {len(discovered_services)} services")
            
        elif target_type == "ip":
            # For single IPs, probe common ports with short timeouts
            host = normalized.split(":")[0]
            port_list = ports or COMMON_WEB_PORTS
            
            # Probe ports concurrently with 1s timeout for speed
            tasks = [probe_port(host, port, timeout=1.0) for port in port_list[:10]]
            try:
                # Wrap port probing in timeout
                port_results = await asyncio.wait_for(
                    asyncio.gather(*tasks, return_exceptions=True),
                    timeout=15.0  # Max 15s for all port probes
                )
            except asyncio.TimeoutError:
                logger.warning(f"AI Auto-Test: Port probing timed out for {host}")
                port_results = []
            
            for pr in port_results:
                if isinstance(pr, dict):
                    result.discovered_services.append(pr)
                    base_urls.append(pr["url"])
            
            # If no ports found, try default HTTP/HTTPS
            if not base_urls:
                base_urls.append(f"http://{normalized}")
                base_urls.append(f"https://{normalized}")
        else:
            # Domain - try common schemes
            base_urls.append(f"https://{normalized}")
            base_urls.append(f"http://{normalized}")
        
        logger.info(f"AI Auto-Test: Found {len(base_urls)} base URLs to probe")
        
        # Step 3: Discover endpoints on each base URL
        check_timeout()  # Check before endpoint discovery
        all_endpoints = []
        if probe_common_paths:
            for base_url in base_urls[:3]:  # Limit to first 3 base URLs
                check_timeout()  # Check before each base URL
                try:
                    endpoints = await discover_endpoints(base_url, timeout=min(timeout, 5.0))
                    for ep in endpoints:
                        ep["base_url"] = base_url
                    all_endpoints.extend(endpoints)
                except Exception as e:
                    logger.warning(f"Failed to probe {base_url}: {e}")
        
        # De-duplicate by URL
        seen_urls = set()
        unique_endpoints = []
        for ep in all_endpoints:
            if ep["url"] not in seen_urls:
                seen_urls.add(ep["url"])
                unique_endpoints.append(ep)
        
        result.discovered_endpoints = unique_endpoints[:max_endpoints]
        logger.info(f"AI Auto-Test: Discovered {len(result.discovered_endpoints)} unique endpoints")
        
        # Step 4: Run security tests on discovered endpoints
        check_timeout()  # Check before security tests
        if run_security_tests and result.discovered_endpoints:
            # Group endpoints by base URL for efficient testing
            by_base = {}
            for ep in result.discovered_endpoints:
                base = ep.get("base_url", ep["url"])
                if base not in by_base:
                    by_base[base] = []
                by_base[base].append({"url": ep["url"], "method": "GET"})
            
            # Run tests for each base URL
            for base_url, endpoints in list(by_base.items())[:5]:  # Limit to 5 base URLs
                check_timeout()  # Check before each base URL test
                try:
                    # Check if GraphQL endpoint was discovered
                    has_graphql = any(
                        "graphql" in ep.get("path", "").lower() or "gql" in ep.get("path", "").lower()
                        for ep in result.discovered_endpoints
                    )
                    
                    test_result = await test_api_endpoints(
                        base_url=base_url,
                        endpoints=endpoints[:10],  # Limit endpoints per base
                        test_auth=True,
                        test_cors=True,
                        test_rate_limit=True,  # Enable rate limit testing
                        test_input_validation=True,
                        test_methods=True,
                        test_graphql=has_graphql,  # Auto-enable if GraphQL found
                        proxy_url=proxy_url,
                        timeout=timeout,
                    )
                    
                    result.test_results.append({
                        "base_url": base_url,
                        "endpoints_tested": test_result.endpoints_tested,
                        "security_score": test_result.security_score,
                        "findings_count": test_result.total_findings,
                    })
                    
                    # Aggregate findings (convert to dict if needed)
                    for finding in test_result.all_findings:
                        if hasattr(finding, 'to_dict'):
                            result.all_findings.append(finding.to_dict())
                        elif isinstance(finding, dict):
                            result.all_findings.append(finding)
                        else:
                            result.all_findings.append(asdict(finding))
                        
                except Exception as e:
                    logger.warning(f"Failed to test {base_url}: {e}")
        
        # Step 4b: Check for WebSocket endpoints and test them
        ws_endpoints = [
            ep for ep in result.discovered_endpoints
            if any(ws in ep.get("path", "").lower() for ws in ["ws", "socket", "websocket", "realtime", "stream"])
        ]
        
        if ws_endpoints:
            logger.info(f"AI Auto-Test: Found {len(ws_endpoints)} potential WebSocket endpoints")
            for ws_ep in ws_endpoints[:3]:  # Limit to first 3 WS endpoints
                try:
                    base = ws_ep.get("base_url", "").replace("http://", "ws://").replace("https://", "wss://")
                    ws_url = urljoin(base, ws_ep["path"])
                    ws_result = await test_websocket(ws_url, timeout=5.0)
                    
                    # Add WebSocket findings
                    for finding in ws_result.findings:
                        if hasattr(finding, 'to_dict'):
                            result.all_findings.append(finding.to_dict())
                        elif isinstance(finding, dict):
                            result.all_findings.append(finding)
                        else:
                            result.all_findings.append(asdict(finding))
                except Exception as e:
                    logger.debug(f"WebSocket test failed for {ws_ep['path']}: {e}")
        
        # Step 5: Calculate totals
        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
        for f in result.all_findings:
            sev = f.get("severity", "info").lower()
            severity_counts[sev] = severity_counts.get(sev, 0) + 1
        
        result.total_findings = len(result.all_findings)
        result.critical_count = severity_counts["critical"]
        result.high_count = severity_counts["high"]
        result.medium_count = severity_counts["medium"]
        result.low_count = severity_counts["low"]
        result.info_count = severity_counts["info"]
        
        # Calculate score
        score = 100
        score -= result.critical_count * 25
        score -= result.high_count * 15
        score -= result.medium_count * 8
        score -= result.low_count * 3
        result.security_score = max(0, min(100, score))
        
    except asyncio.TimeoutError as e:
        logger.warning(f"AI Auto-Test timeout: {e}")
        result.error = f"Scan timed out after {overall_timeout}s (partial results may be available)"
    except Exception as e:
        logger.error(f"AI Auto-Test error: {e}")
        result.error = str(e)
    
    # Always calculate duration and generate summary (even on error/timeout)
    result.scan_duration_seconds = time.time() - start_time
    
    # Generate summary - always run this even if there were errors
    try:
        result.ai_summary = _generate_auto_test_summary(result)
    except Exception as e:
        logger.error(f"Failed to generate AI summary: {e}")
        result.ai_summary = f"# ⚠️ AI Auto-Test Report\n\n**Target:** `{result.target}`\n\n**Error:** {result.error or 'Unknown error during scan'}\n\n**Duration:** {result.scan_duration_seconds:.1f}s"
    
    return result


def _generate_auto_test_summary(result: AIAutoTestResult) -> str:
    """Generate a rich, well-formatted summary of auto-test results."""
    lines = []
    
    # Header with visual styling
    lines.append("# 🔍 AI Auto-Test Security Report")
    lines.append("")
    lines.append("---")
    lines.append("")
    
    # Show error banner if there was an error
    if result.error:
        lines.append("## ⚠️ Scan Status: Partial Results")
        lines.append("")
        lines.append(f"> **Note:** {result.error}")
        lines.append(">")
        lines.append("> The scan encountered an issue but partial results are shown below.")
        lines.append("")
    
    # Target Information Box
    lines.append("## 📋 Target Information")
    lines.append("")
    lines.append(f"| Property | Value |")
    lines.append(f"|----------|-------|")
    lines.append(f"| **Target** | `{result.target}` |")
    lines.append(f"| **Type** | {result.target_type.upper()} |")
    lines.append(f"| **Scan Duration** | {result.scan_duration_seconds:.1f} seconds |")
    lines.append(f"| **Endpoints Discovered** | {len(result.discovered_endpoints)} |")
    lines.append(f"| **Services Found** | {len(result.discovered_services)} |")
    lines.append("")
    
    # Discovered Services
    if result.discovered_services:
        lines.append("## 🌐 Discovered Services")
        lines.append("")
        lines.append("| Port | Protocol | Status | Server |")
        lines.append("|------|----------|--------|--------|")
        for svc in result.discovered_services:
            server = svc.get('server', 'Unknown')[:30]
            lines.append(f"| **{svc['port']}** | {svc['scheme'].upper()} | {svc['status_code']} | {server} |")
        lines.append("")
    
    # Discovered Endpoints
    if result.discovered_endpoints:
        lines.append("## 🔗 Discovered Endpoints")
        lines.append("")
        
        auth_required = [e for e in result.discovered_endpoints if e.get("requires_auth")]
        json_endpoints = [e for e in result.discovered_endpoints if e.get("is_json")]
        graphql_found = any("graphql" in e.get("path", "").lower() or "gql" in e.get("path", "").lower() for e in result.discovered_endpoints)
        
        lines.append("### Endpoint Statistics")
        lines.append("")
        lines.append(f"- **Total Discovered:** {len(result.discovered_endpoints)}")
        lines.append(f"- **JSON API Endpoints:** {len(json_endpoints)}")
        lines.append(f"- **Auth Required:** {len(auth_required)}")
        lines.append(f"- **GraphQL Detected:** {'✅ Yes' if graphql_found else '❌ No'}")
        lines.append("")
        
        # Categorize endpoints
        interesting_keywords = {
            "🔐 Authentication": ["auth", "login", "signin", "signup", "oauth", "token", "jwt", "session"],
            "⚙️ Configuration": ["config", "settings", "env", "debug", "admin"],
            "📚 Documentation": ["swagger", "openapi", "docs", "redoc", "api-docs"],
            "🔮 GraphQL": ["graphql", "gql", "graphiql", "playground"],
            "📊 Monitoring": ["health", "metrics", "status", "actuator", "prometheus"],
        }
        
        categorized = {cat: [] for cat in interesting_keywords}
        other_endpoints = []
        
        for ep in result.discovered_endpoints:
            path_lower = ep.get("path", "").lower()
            matched = False
            for category, keywords in interesting_keywords.items():
                if any(kw in path_lower for kw in keywords):
                    categorized[category].append(ep)
                    matched = True
                    break
            if not matched:
                other_endpoints.append(ep)
        
        # Show categorized endpoints
        for category, endpoints in categorized.items():
            if endpoints:
                lines.append(f"### {category}")
                lines.append("")
                for ep in endpoints[:5]:
                    status_emoji = "✅" if ep["status_code"] < 400 else "⚠️" if ep["status_code"] < 500 else "❌"
                    lines.append(f"- {status_emoji} `{ep['path']}` → **{ep['status_code']}**")
                if len(endpoints) > 5:
                    lines.append(f"- *...and {len(endpoints) - 5} more*")
                lines.append("")
        
        # Other endpoints summary
        if other_endpoints:
            lines.append("### 📁 Other Endpoints")
            lines.append("")
            for ep in other_endpoints[:8]:
                status_emoji = "✅" if ep["status_code"] < 400 else "⚠️" if ep["status_code"] < 500 else "❌"
                lines.append(f"- {status_emoji} `{ep['path']}` → **{ep['status_code']}**")
            if len(other_endpoints) > 8:
                lines.append(f"- *...and {len(other_endpoints) - 8} more*")
            lines.append("")
    
    # Security Assessment Section
    lines.append("---")
    lines.append("")
    lines.append("## 🛡️ Security Assessment")
    lines.append("")
    
    # Score visualization
    score = result.security_score
    if score >= 80:
        score_emoji = "🟢"
        score_label = "GOOD"
    elif score >= 60:
        score_emoji = "🟡"
        score_label = "MODERATE"
    elif score >= 40:
        score_emoji = "🟠"
        score_label = "CONCERNING"
    else:
        score_emoji = "🔴"
        score_label = "CRITICAL"
    
    lines.append(f"### Security Score: {score_emoji} **{score}/100** ({score_label})")
    lines.append("")
    
    # Progress bar visualization
    filled = int(score / 10)
    empty = 10 - filled
    bar = "█" * filled + "░" * empty
    lines.append(f"```")
    lines.append(f"[{bar}] {score}%")
    lines.append(f"```")
    lines.append("")
    
    # Findings breakdown
    if result.total_findings > 0:
        lines.append("### 📊 Findings Summary")
        lines.append("")
        lines.append(f"**Total Findings:** {result.total_findings}")
        lines.append("")
        lines.append("| Severity | Count | Impact |")
        lines.append("|----------|-------|--------|")
        if result.critical_count:
            lines.append(f"| 🔴 **Critical** | **{result.critical_count}** | Immediate action required |")
        if result.high_count:
            lines.append(f"| 🟠 **High** | **{result.high_count}** | Address urgently |")
        if result.medium_count:
            lines.append(f"| 🟡 **Medium** | **{result.medium_count}** | Plan to fix soon |")
        if result.low_count:
            lines.append(f"| 🔵 **Low** | **{result.low_count}** | Consider fixing |")
        if result.info_count:
            lines.append(f"| ⚪ **Info** | **{result.info_count}** | Informational |")
        lines.append("")
        
        # Top findings with details
        top_findings = sorted(
            result.all_findings,
            key=lambda x: {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}.get(x.get("severity", "info").lower(), 5)
        )[:7]
        
        if top_findings:
            lines.append("### 🚨 Top Security Issues")
            lines.append("")
            for i, f in enumerate(top_findings, 1):
                sev = f.get('severity', 'info').upper()
                sev_emoji = {"CRITICAL": "🔴", "HIGH": "🟠", "MEDIUM": "🟡", "LOW": "🔵", "INFO": "⚪"}.get(sev, "⚪")
                title = f.get('title', 'Unknown Issue')
                endpoint = f.get('endpoint', f.get('url', 'N/A'))
                
                lines.append(f"**{i}. {sev_emoji} [{sev}] {title}**")
                if endpoint and endpoint != 'N/A':
                    lines.append(f"   - *Endpoint:* `{endpoint}`")
                if f.get('description'):
                    desc = f['description'][:150] + "..." if len(f.get('description', '')) > 150 else f.get('description', '')
                    lines.append(f"   - *Details:* {desc}")
                lines.append("")
    else:
        lines.append("### ✅ No Security Issues Detected")
        lines.append("")
        lines.append("Great news! The automated security scan did not find any significant vulnerabilities.")
        lines.append("")
        lines.append("> **Note:** This doesn't guarantee the target is completely secure. Manual testing and")
        lines.append("> deeper analysis are recommended for comprehensive security assurance.")
        lines.append("")
    
    # Recommendations Section
    lines.append("---")
    lines.append("")
    lines.append("## 💡 Recommendations")
    lines.append("")
    
    recommendations = []
    
    if result.critical_count > 0:
        recommendations.append({
            "priority": "🔴 CRITICAL",
            "text": f"**{result.critical_count} critical vulnerabilities** require immediate attention. These can lead to system compromise."
        })
    
    if result.high_count > 0:
        recommendations.append({
            "priority": "🟠 HIGH",
            "text": f"**{result.high_count} high-severity issues** should be addressed urgently as part of your security roadmap."
        })
    
    # Endpoint-specific recommendations
    swagger_found = any("swagger" in e.get("path", "").lower() or "openapi" in e.get("path", "").lower() for e in result.discovered_endpoints)
    debug_found = any(x in e.get("path", "").lower() for e in result.discovered_endpoints for x in ["debug", "env", "config", "trace"])
    admin_found = any("admin" in e.get("path", "").lower() for e in result.discovered_endpoints)
    graphql_found = any("graphql" in e.get("path", "").lower() or "gql" in e.get("path", "").lower() for e in result.discovered_endpoints)
    actuator_found = any("actuator" in e.get("path", "").lower() for e in result.discovered_endpoints)
    
    if swagger_found:
        recommendations.append({
            "priority": "🟡 MEDIUM",
            "text": "**API Documentation Exposed** - Swagger/OpenAPI docs are publicly accessible. Consider restricting access in production."
        })
    
    if debug_found:
        recommendations.append({
            "priority": "🟠 HIGH",
            "text": "**Debug Endpoints Detected** - Configuration or debug endpoints are exposed. Disable these in production environments."
        })
    
    if admin_found:
        recommendations.append({
            "priority": "🟡 MEDIUM",
            "text": "**Admin Panel Found** - Ensure admin endpoints have proper authentication, rate limiting, and access logging."
        })
    
    if graphql_found:
        recommendations.append({
            "priority": "🟡 MEDIUM",
            "text": "**GraphQL Endpoint Detected** - Review introspection settings, implement query depth limiting, and add proper authentication."
        })
    
    if actuator_found:
        recommendations.append({
            "priority": "🟠 HIGH",
            "text": "**Spring Actuator Exposed** - Management endpoints may leak sensitive information. Secure or disable in production."
        })
    
    # General recommendations
    recommendations.append({
        "priority": "ℹ️ INFO",
        "text": "**Regular Testing** - Schedule periodic security scans to catch new vulnerabilities early."
    })
    
    if recommendations:
        for rec in recommendations:
            lines.append(f"### {rec['priority']}")
            lines.append(f"{rec['text']}")
            lines.append("")
    
    # Footer
    lines.append("---")
    lines.append("")
    lines.append("*Report generated by VRAgent AI Auto-Test • For comprehensive security assessment, combine with manual penetration testing.*")
    
    return "\n".join(lines)
