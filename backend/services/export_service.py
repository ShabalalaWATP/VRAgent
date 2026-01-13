"""
Report export service - generates PDF, DOCX, and Markdown reports.
"""
from io import BytesIO
from typing import List, Dict, Any, Optional
from collections import defaultdict

from sqlalchemy.orm import Session

from backend import models
from backend.core.logging import get_logger
from backend.core.config import settings

logger = get_logger(__name__)


def get_codebase_statistics(db: Session, report: models.Report) -> Dict[str, Any]:
    """Get codebase statistics for a report."""
    project_id = report.project_id
    
    # Get all code chunks
    chunks = db.query(models.CodeChunk).filter(
        models.CodeChunk.project_id == project_id
    ).all()
    
    if not chunks:
        return {
            "total_files": 0,
            "total_lines": 0,
            "languages": {}
        }
    
    # Calculate statistics
    file_lines = {}
    languages = defaultdict(int)
    
    for chunk in chunks:
        path = chunk.file_path
        if path not in file_lines:
            file_lines[path] = 0
        if chunk.end_line:
            file_lines[path] = max(file_lines[path], chunk.end_line)
        if chunk.language:
            languages[chunk.language] += 1
    
    return {
        "total_files": len(file_lines),
        "total_lines": sum(file_lines.values()),
        "languages": dict(sorted(languages.items(), key=lambda x: -x[1]))
    }


def get_exploit_scenarios(db: Session, report_id: int) -> List[models.ExploitScenario]:
    """Get exploit scenarios for a report."""
    return db.query(models.ExploitScenario).filter(
        models.ExploitScenario.report_id == report_id
    ).all()


def get_exploit_context(scenario: models.ExploitScenario) -> Dict[str, Any]:
    """
    Get beginner-friendly exploit development context for a scenario.
    Returns attack type, difficulty, skills needed, tools, phases, glossary, and tips.
    """
    title = (scenario.title or "").lower()
    narrative = (scenario.narrative or "").lower()
    
    # SQL Injection
    if "sql" in title or "sql injection" in narrative:
        return {
            "attack_type": "SQL Injection",
            "difficulty": "Beginner",
            "time_estimate": "30 min - 2 hours",
            "skills_needed": ["Basic SQL knowledge", "Understanding HTTP requests", "Using browser dev tools"],
            "tools_needed": ["sqlmap", "Burp Suite (optional)", "Browser with dev tools"],
            "phases": [
                {"name": "Reconnaissance", "description": "Identify input fields that interact with the database", "emoji": "🔍"},
                {"name": "Testing", "description": "Send test payloads like ' OR '1'='1 to find injection points", "emoji": "🧪"},
                {"name": "Exploitation", "description": "Extract data using UNION SELECT or error-based techniques", "emoji": "💉"},
                {"name": "Post-Exploitation", "description": "Dump database contents, escalate privileges if possible", "emoji": "🎯"},
            ],
            "glossary": [
                {"term": "Payload", "definition": "The malicious SQL code you inject into the vulnerable input"},
                {"term": "UNION attack", "definition": "A technique to combine results from multiple SELECT queries"},
                {"term": "Blind SQLi", "definition": "When you can't see query results but can infer them from app behavior"},
            ],
            "tips": [
                "Start with simple payloads: ' OR '1'='1'--",
                "Check for different SQL dialects (MySQL, PostgreSQL, MSSQL)",
                "Use time-based payloads if results aren't visible",
            ],
        }
    
    # XSS
    if "xss" in title or "cross-site" in title or "script injection" in narrative:
        return {
            "attack_type": "Cross-Site Scripting (XSS)",
            "difficulty": "Beginner",
            "time_estimate": "15 min - 1 hour",
            "skills_needed": ["Basic JavaScript", "HTML knowledge", "Understanding of DOM"],
            "tools_needed": ["Browser dev tools", "XSS Hunter (optional)", "Burp Suite (optional)"],
            "phases": [
                {"name": "Identify Input", "description": "Find fields where user input is reflected in the page", "emoji": "🔍"},
                {"name": "Test Reflection", "description": "Input test strings to see how they appear in the response", "emoji": "🧪"},
                {"name": "Bypass Filters", "description": "Try encoding or alternative payloads if blocked", "emoji": "🔓"},
                {"name": "Weaponize", "description": "Craft payload to steal cookies, redirect users, or deface page", "emoji": "⚔️"},
            ],
            "glossary": [
                {"term": "Reflected XSS", "definition": "Payload is in the URL and executes when victim clicks the link"},
                {"term": "Stored XSS", "definition": "Payload is saved (e.g., in comments) and runs for all viewers"},
                {"term": "DOM XSS", "definition": "Payload manipulates the page's JavaScript without server involvement"},
            ],
            "tips": [
                "Start with <script>alert(1)</script>",
                "Try event handlers: <img src=x onerror=alert(1)>",
                "Use encoding to bypass filters",
            ],
        }
    
    # Command Injection
    if "command" in title or "exec" in title or "shell" in narrative or "os command" in narrative:
        return {
            "attack_type": "Command Injection",
            "difficulty": "Intermediate",
            "time_estimate": "1 - 4 hours",
            "skills_needed": ["Linux/Windows CLI basics", "Understanding shell operators", "Networking basics"],
            "tools_needed": ["Netcat", "Burp Suite", "Reverse shell generators"],
            "phases": [
                {"name": "Identify Entry Point", "description": "Find features that might execute system commands", "emoji": "🔍"},
                {"name": "Test Injection", "description": "Try command separators: ; | & || &&", "emoji": "🧪"},
                {"name": "Confirm Execution", "description": "Use sleep/ping to confirm blind command execution", "emoji": "✅"},
                {"name": "Get Shell", "description": "Establish reverse shell for interactive access", "emoji": "🐚"},
            ],
            "glossary": [
                {"term": "Reverse Shell", "definition": "Target connects back to your machine, giving you a command prompt"},
                {"term": "Blind Injection", "definition": "When you can't see output but can infer success from timing/behavior"},
                {"term": "Out-of-band", "definition": "Using DNS/HTTP requests to exfiltrate data when direct output is blocked"},
            ],
            "tips": [
                "Try different separators: ; | && || `command`",
                "Use $(whoami) or `whoami` for command substitution",
                "Encode spaces as ${IFS} if blocked",
            ],
        }
    
    # Path Traversal
    if "path" in title or "traversal" in title or "directory" in title or "../" in narrative:
        return {
            "attack_type": "Path Traversal",
            "difficulty": "Beginner",
            "time_estimate": "15 min - 1 hour",
            "skills_needed": ["File system basics", "Understanding URL encoding", "HTTP fundamentals"],
            "tools_needed": ["Burp Suite", "curl", "Browser"],
            "phases": [
                {"name": "Find File Parameters", "description": "Look for URLs with file=, path=, doc= parameters", "emoji": "🔍"},
                {"name": "Test Traversal", "description": "Replace value with ../../../etc/passwd", "emoji": "🧪"},
                {"name": "Bypass Filters", "description": "Try encoding: %2e%2e%2f, double encoding, or ..\\ ", "emoji": "🔓"},
                {"name": "Extract Files", "description": "Read sensitive configs, source code, or credentials", "emoji": "📂"},
            ],
            "glossary": [
                {"term": "../ (dot-dot-slash)", "definition": "Moves up one directory level in the file system"},
                {"term": "Null byte", "definition": "Using %00 to truncate file paths (works on older systems)"},
                {"term": "Absolute path", "definition": "Full path from root: /etc/passwd vs relative path"},
            ],
            "tips": [
                "Try ../../../../etc/passwd on Linux",
                "Try ....//....//etc/passwd to bypass simple filters",
                "URL encode: %2e%2e%2f = ../",
            ],
        }
    
    # Authentication/Session
    if "auth" in title or "session" in title or "token" in title or "bypass" in narrative:
        return {
            "attack_type": "Authentication Bypass",
            "difficulty": "Intermediate",
            "time_estimate": "1 - 3 hours",
            "skills_needed": ["Session management concepts", "Cookie mechanics", "HTTP headers"],
            "tools_needed": ["Burp Suite", "Browser dev tools", "JWT.io (for JWTs)"],
            "phases": [
                {"name": "Analyze Flow", "description": "Map out the authentication process and session handling", "emoji": "🗺️"},
                {"name": "Test Weaknesses", "description": "Try default creds, SQL injection in login, or token manipulation", "emoji": "🧪"},
                {"name": "Session Attack", "description": "Attempt session fixation, hijacking, or prediction", "emoji": "🔐"},
                {"name": "Escalate", "description": "Access admin functions or other users' accounts", "emoji": "⬆️"},
            ],
            "glossary": [
                {"term": "Session Token", "definition": "A unique identifier that proves you're logged in (usually a cookie)"},
                {"term": "JWT", "definition": "JSON Web Token - a structured token format often used for stateless auth"},
                {"term": "IDOR", "definition": "Insecure Direct Object Reference - accessing others' data by changing IDs"},
            ],
            "tips": [
                "Check if session tokens are predictable",
                "Try changing user IDs in requests (IDOR)",
                "Look for role parameters that can be manipulated",
            ],
        }
    
    # Deserialization
    if "deserial" in title or "pickle" in title or "yaml" in title:
        return {
            "attack_type": "Insecure Deserialization",
            "difficulty": "Advanced",
            "time_estimate": "2 - 6 hours",
            "skills_needed": ["Language-specific serialization", "Object-oriented programming", "Gadget chains"],
            "tools_needed": ["ysoserial (Java)", "pickle-payload (Python)", "Burp Suite"],
            "phases": [
                {"name": "Identify Format", "description": "Determine serialization format (Java, PHP, Python, .NET)", "emoji": "🔍"},
                {"name": "Find Gadgets", "description": "Research gadget chains available in the target libraries", "emoji": "🔧"},
                {"name": "Craft Payload", "description": "Build serialized object that triggers code execution", "emoji": "📦"},
                {"name": "Execute", "description": "Send payload and achieve RCE or other impact", "emoji": "💥"},
            ],
            "glossary": [
                {"term": "Serialization", "definition": "Converting objects to bytes/strings for storage or transmission"},
                {"term": "Gadget Chain", "definition": "A sequence of existing code that performs malicious actions when deserialized"},
                {"term": "Magic Methods", "definition": "Special methods called during deserialization (e.g., __reduce__ in Python)"},
            ],
            "tips": [
                "Python pickle: Never unpickle untrusted data",
                "Java: Look for Base64-encoded data starting with rO0",
                "Use ysoserial to generate Java payloads",
            ],
        }
    
    # SSRF
    if "ssrf" in title or "server-side request" in narrative:
        return {
            "attack_type": "Server-Side Request Forgery",
            "difficulty": "Intermediate",
            "time_estimate": "1 - 3 hours",
            "skills_needed": ["URL manipulation", "Cloud metadata knowledge", "Internal network concepts"],
            "tools_needed": ["Burp Suite", "ngrok (for callbacks)", "Collaborator/webhook.site"],
            "phases": [
                {"name": "Find URL Input", "description": "Locate features that fetch URLs (webhooks, import, preview)", "emoji": "🔍"},
                {"name": "Test Localhost", "description": "Try http://127.0.0.1, http://localhost, or 0.0.0.0", "emoji": "🧪"},
                {"name": "Cloud Metadata", "description": "Access http://169.254.169.254 for AWS/GCP/Azure secrets", "emoji": "☁️"},
                {"name": "Port Scan", "description": "Use the server to scan internal network services", "emoji": "🔌"},
            ],
            "glossary": [
                {"term": "Metadata Service", "definition": "Cloud provider endpoint exposing credentials and instance info"},
                {"term": "Blind SSRF", "definition": "When you can't see the response but can confirm the request was made"},
                {"term": "Protocol Smuggling", "definition": "Using file://, gopher://, or dict:// for advanced attacks"},
            ],
            "tips": [
                "Try AWS metadata: http://169.254.169.254/latest/meta-data/",
                "Use DNS rebinding to bypass allow-lists",
                "Try different URL formats: http://127.1, http://[::1]",
            ],
        }
    
    # Buffer Overflow
    if "buffer" in title or "overflow" in title or "memory" in title:
        return {
            "attack_type": "Buffer Overflow",
            "difficulty": "Expert",
            "time_estimate": "4 - 12 hours",
            "skills_needed": ["Assembly language", "Memory layout", "Debugging (GDB)", "Binary exploitation"],
            "tools_needed": ["GDB/PEDA", "pwntools", "ROPgadget", "Binary analysis tools"],
            "phases": [
                {"name": "Find Overflow", "description": "Identify buffer that can be overwritten with controlled data", "emoji": "🔍"},
                {"name": "Control EIP/RIP", "description": "Determine offset to overwrite the return address", "emoji": "🎯"},
                {"name": "Bypass Protections", "description": "Handle ASLR, NX, stack canaries, PIE", "emoji": "🛡️"},
                {"name": "Shellcode/ROP", "description": "Execute payload via shellcode or return-oriented programming", "emoji": "💉"},
            ],
            "glossary": [
                {"term": "Return Address", "definition": "Memory location that tells the CPU where to go after a function returns"},
                {"term": "Shellcode", "definition": "Small piece of machine code that spawns a shell"},
                {"term": "ROP Chain", "definition": "Chaining existing code snippets to bypass NX protection"},
            ],
            "tips": [
                "Use pattern_create to find exact offset",
                "Check protections with checksec",
                "Modern systems need ROP or ret2libc techniques",
            ],
        }
    
    # Crypto/Secret
    if "crypto" in title or "secret" in title or "password" in title or "credential" in title:
        return {
            "attack_type": "Credential/Secret Exposure",
            "difficulty": "Beginner",
            "time_estimate": "5 - 30 minutes",
            "skills_needed": ["Reading code", "Understanding API keys/tokens", "Basic enumeration"],
            "tools_needed": ["git (for history)", "truffleHog", "gitleaks"],
            "phases": [
                {"name": "Extract Secret", "description": "Copy the exposed credential from the source code", "emoji": "🔑"},
                {"name": "Identify Service", "description": "Determine what system the credential accesses", "emoji": "🔍"},
                {"name": "Test Access", "description": "Try using the credential to authenticate", "emoji": "🧪"},
                {"name": "Enumerate", "description": "Discover what data or functions are accessible", "emoji": "📋"},
            ],
            "glossary": [
                {"term": "API Key", "definition": "A secret string used to authenticate API requests"},
                {"term": "Token", "definition": "A credential that grants access to a service (often time-limited)"},
                {"term": "Secret Rotation", "definition": "Regularly changing credentials to limit exposure window"},
            ],
            "tips": [
                "Check git history for previously committed secrets",
                "Search for environment variable references that might leak",
                "Determine scope: what does this key access?",
            ],
        }
    
    # Default for unknown types
    return {
        "attack_type": "Security Vulnerability Exploitation",
        "difficulty": "Intermediate",
        "time_estimate": "1 - 4 hours",
        "skills_needed": ["Security fundamentals", "HTTP/networking basics", "Reading technical documentation"],
        "tools_needed": ["Burp Suite", "Browser dev tools", "curl/wget"],
        "phases": [
            {"name": "Understand", "description": "Read the vulnerability description and research the attack type", "emoji": "📚"},
            {"name": "Reproduce", "description": "Confirm the vulnerability exists with a simple proof", "emoji": "🔄"},
            {"name": "Exploit", "description": "Develop a working proof-of-concept attack", "emoji": "⚔️"},
            {"name": "Document", "description": "Record steps, impact, and evidence for the report", "emoji": "📝"},
        ],
        "glossary": [
            {"term": "PoC", "definition": "Proof of Concept - minimal code/steps to demonstrate an exploit"},
            {"term": "Attack Vector", "definition": "The method used to exploit a vulnerability"},
            {"term": "Impact", "definition": "The potential damage from successful exploitation"},
        ],
        "tips": [
            "Start by understanding exactly what the vulnerability allows",
            "Search for existing PoCs or write-ups for similar issues",
            "Always document your steps for reporting",
        ],
    }


def generate_ai_summaries(db: Session, report: models.Report, findings: List[models.Finding]) -> Dict[str, Optional[str]]:
    """Generate COMPREHENSIVE AI summaries for app and security analysis.
    
    Matches the quality and depth of the APK Analyzer reports with:
    - Detailed HTML formatting with styled headers and badges
    - 15+ sections covering all aspects of the application
    - Attacker-focused security analysis with PoC ideas
    - Attack chains, playbooks, and exploitation scenarios
    
    First checks if summaries are cached in report.data, otherwise generates new ones.
    """
    # Check for cached summaries first
    if report.data and report.data.get("ai_summaries"):
        cached = report.data["ai_summaries"]
        logger.info(f"Using cached AI summaries for report {report.id}")
        return {
            "app_summary": cached.get("app_summary"),
            "security_summary": cached.get("security_summary")
        }
    
    app_summary = None
    security_summary = None
    
    if not settings.gemini_api_key:
        return {"app_summary": None, "security_summary": None}
    
    try:
        from google import genai
        
        # Get MORE codebase info for comprehensive context
        chunks = db.query(models.CodeChunk).filter(
            models.CodeChunk.project_id == report.project_id
        ).limit(100).all()  # Increased from 50 to 100
        
        if not chunks:
            return {"app_summary": None, "security_summary": None}
        
        # Build comprehensive code context with more samples
        code_samples = []
        for chunk in chunks[:50]:  # Increased from 30 to 50
            code_samples.append(f"File: {chunk.file_path}\n```\n{chunk.code[:800]}...\n```")  # Increased from 500 to 800
        
        project = db.get(models.Project, report.project_id)
        
        # Gather additional context for richer analysis
        language_breakdown = {}
        file_types = set()
        for chunk in chunks:
            lang = chunk.language or "Unknown"
            language_breakdown[lang] = language_breakdown.get(lang, 0) + 1
            if chunk.file_path:
                ext = chunk.file_path.split('.')[-1] if '.' in chunk.file_path else 'unknown'
                file_types.add(ext)
        
        # ==================== COMPREHENSIVE APP SUMMARY PROMPT ====================
        app_prompt = f"""You are an expert software architect and code analyst. Your task is to provide a COMPREHENSIVE analysis of what this application does based on DEEP code inspection.

## PROJECT INFORMATION
Project Name: {project.name if project else 'Unknown'}
Total Code Chunks Analyzed: {len(chunks)}
Languages Detected: {', '.join(f'{k}: {v} files' for k, v in sorted(language_breakdown.items(), key=lambda x: -x[1]))}
File Types: {', '.join(sorted(file_types))}

## COMPLETE CODE SAMPLES
{chr(10).join(code_samples[:40])}

## YOUR TASK
Perform a THOROUGH analysis of the codebase by:
1. Reading ALL provided source code carefully
2. Tracing data flows and feature implementations
3. Identifying ALL major features and capabilities
4. Understanding how different components work together
5. Noting any hidden or non-obvious functionality

Generate a DETAILED, COMPREHENSIVE report. Use HTML formatting.

FORMAT YOUR RESPONSE AS CLEAN HTML (no markdown, no code blocks):
- Use <h3> for section headers
- Use <h4> for sub-sections
- Use <ul> and <li> for bullet points
- Use <strong> for emphasis
- Use <p> for paragraphs
- Use <code> for class/method names

REQUIRED SECTIONS (BE THOROUGH):

<h3>📱 Application Overview</h3>
<p>[2-3 sentences describing what this application is, its purpose, and target audience - base this on actual code analysis]</p>

<h3>🎯 Core Features & Functionality</h3>
<p>Based on deep code analysis, this application provides the following features:</p>

<h4>Main Features</h4>
<ul>
<li><strong>[Feature Name]:</strong> [Detailed description of what it does, HOW it works based on the code you see, which classes/modules implement it]</li>
[List ALL major features you can identify from the code]
</ul>

<h4>User Interface & Navigation</h4>
<ul>
<li>[Describe the app's UI patterns, components, pages/views, and navigation flow based on code analysis]</li>
</ul>

<h3>🔐 Authentication & User Management</h3>
<ul>
<li><strong>Authentication:</strong> [How authentication works - JWT? OAuth? Session-based? API keys?]</li>
<li><strong>Authorization:</strong> [Role-based? Permission system? Access control patterns?]</li>
<li><strong>Session Management:</strong> [How sessions/tokens are handled]</li>
<li><strong>User Data:</strong> [What user data is collected/stored]</li>
</ul>

<h3>📡 API & Network Communication</h3>
<h4>API Endpoints</h4>
<ul>
<li><strong>Endpoint Patterns:</strong> [REST? GraphQL? WebSocket?]</li>
<li><strong>Key Routes:</strong> [List important API endpoints found in code]</li>
<li><strong>Data Formats:</strong> [JSON? XML? How is data serialized?]</li>
</ul>

<h4>External Services</h4>
<ul>
<li>[Third-party APIs, cloud services, external dependencies]</li>
</ul>

<h3>💾 Data Storage & Persistence</h3>
<ul>
<li><strong>Database:</strong> [SQL? NoSQL? What ORM/driver? Schema patterns?]</li>
<li><strong>Caching:</strong> [Redis? Memcached? In-memory?]</li>
<li><strong>File Storage:</strong> [Local files? Cloud storage? How files are handled?]</li>
<li><strong>Data Models:</strong> [Key entities and relationships]</li>
</ul>

<h3>⚙️ Technology Stack</h3>
<ul>
<li><strong>Languages:</strong> [Programming languages with specific versions if apparent]</li>
<li><strong>Frameworks:</strong> [Web frameworks, libraries used]</li>
<li><strong>Build Tools:</strong> [Package managers, build systems]</li>
<li><strong>Runtime:</strong> [Node.js? Python? Docker?]</li>
</ul>

<h3>🏗️ Architecture & Design Patterns</h3>
<ul>
<li><strong>Architecture Style:</strong> [Monolith? Microservices? MVC? Clean Architecture?]</li>
<li><strong>Design Patterns:</strong> [Repository? Factory? Observer? etc.]</li>
<li><strong>Code Organization:</strong> [How is the codebase structured?]</li>
<li><strong>Module Boundaries:</strong> [How do components interact?]</li>
</ul>

<h3>🔌 Background Processing & Jobs</h3>
<ul>
<li>[Scheduled tasks? Message queues? Worker processes? Cron jobs?]</li>
</ul>

<h3>🔗 Integrations & External Dependencies</h3>
<ul>
<li>[Email services? Payment processors? Analytics? Logging?]</li>
</ul>

<h3>🔍 Notable Implementation Details</h3>
<ul>
<li>[Interesting technical patterns you noticed in the code]</li>
<li>[Unique approaches or clever solutions]</li>
<li>[Technical debt or areas that could be improved]</li>
</ul>

<h3>📊 Application Complexity Assessment</h3>
<p>[Simple utility / Medium complexity / Complex enterprise application - justify based on code analysis]</p>
<ul>
<li><strong>Lines of Code Estimate:</strong> [Based on chunks analyzed]</li>
<li><strong>Component Count:</strong> [Approximate number of major modules]</li>
<li><strong>Technical Sophistication:</strong> [Basic / Intermediate / Advanced]</li>
</ul>

BE THOROUGH AND SPECIFIC. Reference actual classes, functions, and files you see in the code. Don't make assumptions - only report what you can verify from the code."""

        from google import genai
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # Generate app summary
        response = client.models.generate_content(
            model=settings.gemini_model_id,
            contents=app_prompt
        )
        if response and response.text:
            app_summary = response.text
            # Clean up any markdown wrappers
            app_summary = app_summary.strip()
            if app_summary.startswith("```html"):
                app_summary = app_summary[7:]
            if app_summary.startswith("```"):
                app_summary = app_summary[3:]
            if app_summary.endswith("```"):
                app_summary = app_summary[:-3]
            app_summary = app_summary.strip()
        
        # Generate COMPREHENSIVE security summary if there are findings
        if findings:
            severity_counts = defaultdict(int)
            for f in findings:
                severity_counts[f.severity.lower() if f.severity else "info"] += 1
            
            # Build detailed findings context for comprehensive analysis
            findings_by_type = {}
            for f in findings:
                ftype = f.type or "Unknown"
                if ftype not in findings_by_type:
                    findings_by_type[ftype] = []
                findings_by_type[ftype].append({
                    "severity": f.severity,
                    "summary": f.summary[:200] if f.summary else "",
                    "description": f.description[:300] if f.description else "",
                    "file_path": f.file_path,
                    "line_number": f.line_number,
                    "code": f.vulnerable_code[:200] if f.vulnerable_code else ""
                })
            
            findings_context = "\n".join(
                f"[{ft}] ({len(items)} findings): " + 
                "; ".join(f"{i['severity']}: {i['summary'][:80]}" for i in items[:5])
                for ft, items in sorted(findings_by_type.items(), key=lambda x: -len(x[1]))[:15]
            )
            
            detailed_findings = []
            for f in findings[:30]:
                detailed_findings.append(f"""
- [{f.severity.upper() if f.severity else 'INFO'}] {f.type}: {f.summary[:150] if f.summary else 'N/A'}
  File: {f.file_path or 'Unknown'}:{f.line_number or '?'}
  Code: {f.vulnerable_code[:100] if f.vulnerable_code else 'N/A'}
""")

            # ==================== COMPREHENSIVE SECURITY PROMPT ====================
            security_prompt = f"""You are an elite RED TEAM OPERATOR performing offensive security assessment of this codebase. Your goal is to find EXPLOITABLE vulnerabilities and demonstrate HOW to attack this application.

Think like an ATTACKER, not a compliance auditor. Focus on:
- What can I ACTUALLY exploit?
- How do I chain vulnerabilities together?
- What's the realistic attack path to compromise the system?
- What would I do FIRST if I wanted to hack this app?

## CODEBASE CONTEXT
Project: {project.name if project else 'Unknown'}
Languages: {', '.join(language_breakdown.keys())}
Total Files Analyzed: {len(chunks)}

## SECURITY FINDINGS IDENTIFIED

### Severity Distribution
- Critical: {severity_counts['critical']}
- High: {severity_counts['high']}
- Medium: {severity_counts['medium']}
- Low: {severity_counts['low']}
- Info: {severity_counts['info']}
Total: {len(findings)} security weaknesses identified

### Findings by Category
{findings_context}

### Detailed Top Findings
{chr(10).join(detailed_findings)}

## RELEVANT CODE SAMPLES
{chr(10).join(code_samples[:20])}

## YOUR MISSION
Perform an OFFENSIVE security assessment:
1. Identify the ATTACK SURFACE - where can an attacker get in?
2. Find EXPLOITABLE vulnerabilities with working attack scenarios
3. Build ATTACK CHAINS - how do multiple weaknesses combine?
4. Prioritize by REAL-WORLD EXPLOITABILITY, not theoretical risk
5. Provide PROOF-OF-CONCEPT ideas for each finding
6. Consider both REMOTE attacks (network) and LOCAL attacks (insider, supply chain)

Generate an ATTACKER-FOCUSED security report. Use HTML formatting.

FORMAT YOUR RESPONSE AS CLEAN HTML (no markdown, no code blocks):
- Use <h3> for section headers
- Use <h4> for sub-sections
- Use <ul> and <li> for bullet points
- Use <strong> for emphasis
- Use <code> for code/commands
- Severity badges: <span style="color: #dc2626; font-weight: bold;">CRITICAL</span>, <span style="color: #ea580c; font-weight: bold;">HIGH</span>, <span style="color: #ca8a04; font-weight: bold;">MEDIUM</span>, <span style="color: #16a34a; font-weight: bold;">LOW</span>

REQUIRED SECTIONS:

<h3>⚔️ Attack Summary</h3>
<p><strong>Hackability Score:</strong> [X/10] - How easy is this application to compromise?</p>
<p><strong>Most Dangerous Finding:</strong> [One-liner of the worst issue]</p>
<p><strong>Recommended Attack Path:</strong> [The attack chain I would use]</p>
<p><strong>Attacker Value:</strong> [What's worth stealing? User data? Credentials? Money? Business logic?]</p>

<h3>🎯 Attack Surface Analysis</h3>
<h4>Entry Points (How I Get In)</h4>
<ul>
<li><strong>API Endpoints:</strong> [Which endpoints are vulnerable? Authentication bypasses?]</li>
<li><strong>User Input:</strong> [Form fields, file uploads, query parameters at risk]</li>
<li><strong>Authentication:</strong> [Login weaknesses, session handling issues]</li>
<li><strong>File Operations:</strong> [Path traversal, file upload, file inclusion risks]</li>
<li><strong>Database Queries:</strong> [SQL injection, NoSQL injection points]</li>
<li><strong>External Dependencies:</strong> [Vulnerable libraries, supply chain risks]</li>
</ul>

<h4>Sensitive Assets (What I Want)</h4>
<ul>
<li>[User credentials, tokens, personal data, financial info, business data]</li>
<li>[Where is each asset stored? How is it protected?]</li>
</ul>

<h3>💀 Critical Exploits</h3>
<p>Vulnerabilities I can exploit RIGHT NOW:</p>

<h4>Exploit 1: [Catchy Attack Name]</h4>
<ul>
<li><span style="color: #dc2626; font-weight: bold;">CRITICAL</span> or appropriate severity</li>
<li><strong>What:</strong> [Technical description]</li>
<li><strong>Where:</strong> <code>[File:Line or Component]</code></li>
<li><strong>Vulnerable Code:</strong> <pre><code>[The actual vulnerable code snippet]</code></pre></li>
<li><strong>Attack Scenario:</strong> [Step-by-step how I would exploit this]</li>
<li><strong>PoC Idea:</strong> <code>[curl command / payload / script snippet]</code></li>
<li><strong>Impact:</strong> [What I gain - RCE? Data theft? Account takeover? Privilege escalation?]</li>
<li><strong>Difficulty:</strong> [Easy/Medium/Hard] - [Why]</li>
</ul>

[Repeat for top 3-5 most critical exploits]

<h3>🔗 Attack Chains</h3>
<p>How I combine multiple weaknesses for maximum impact:</p>

<h4>Chain 1: [Attack Chain Name]</h4>
<ol>
<li><strong>Step 1:</strong> [First exploit/technique - initial access]</li>
<li><strong>Step 2:</strong> [Second exploit/technique - escalation]</li>
<li><strong>Step 3:</strong> [Third technique - lateral movement or data access]</li>
<li><strong>Result:</strong> [What attacker achieves - full compromise, data exfil, etc.]</li>
</ol>

<h3>🔓 Authentication & Session Attacks</h3>
<ul>
<li><strong>Token Weakness:</strong> [Can I forge/steal/reuse tokens?]</li>
<li><strong>Session Hijacking:</strong> [How would I steal a session?]</li>
<li><strong>Credential Extraction:</strong> [Where are creds stored? Can I get them?]</li>
<li><strong>Privilege Escalation:</strong> [Can regular user become admin?]</li>
<li><strong>Password Reset:</strong> [Flaws in password reset flow?]</li>
</ul>

<h3>💉 Injection Attack Vectors</h3>
<ul>
<li><strong>SQL Injection:</strong> [Where? PoC query?]</li>
<li><strong>Command Injection:</strong> [Any exec/system calls?]</li>
<li><strong>XSS (Cross-Site Scripting):</strong> [User input reflected unsanitized?]</li>
<li><strong>Template Injection:</strong> [SSTI opportunities?]</li>
<li><strong>Path Traversal:</strong> [File operations I can abuse?]</li>
<li><strong>LDAP/XML/Header Injection:</strong> [Other injection points?]</li>
</ul>

<h3>🔑 Secrets & Sensitive Data Exposure</h3>
<ul>
<li><strong>[Secret Type]:</strong> <code>[Location]</code>
  <ul>
  <li>Risk: [What access does this give me?]</li>
  <li>Exploitation: [How I'd extract and use this]</li>
  </ul>
</li>
</ul>

<h3>📦 Vulnerable Dependencies</h3>
<ul>
<li><strong>[Library Name] v[X.X]:</strong> [Known CVEs or issues]
  <ul>
  <li>Exploitability: [Can I actually exploit this in context?]</li>
  <li>Attack: [How would I leverage this vulnerability?]</li>
  </ul>
</li>
</ul>

<h3>🌐 Network & API Attack Vectors</h3>
<ul>
<li><strong>MITM Possibility:</strong> [Insecure communications?]</li>
<li><strong>API Abuse:</strong> [Rate limiting? IDOR? Broken access control?]</li>
<li><strong>SSRF:</strong> [Can I make the server fetch internal resources?]</li>
<li><strong>Request Smuggling:</strong> [HTTP parsing issues?]</li>
</ul>

<h3>⚙️ Business Logic Flaws</h3>
<ul>
<li>[Logic bypasses, workflow manipulation, race conditions]</li>
<li>[Price manipulation, discount abuse, feature abuse]</li>
</ul>

<h3>📋 Attack Playbook</h3>
<p>If I had 1 hour to hack this application, I would:</p>
<ol>
<li><strong>[First 15 min]:</strong> [Initial reconnaissance and access technique]</li>
<li><strong>[Next 15 min]:</strong> [Privilege escalation / deeper access]</li>
<li><strong>[Next 15 min]:</strong> [Data exfiltration / persistence]</li>
<li><strong>[Final 15 min]:</strong> [Cover tracks / maximize impact]</li>
</ol>

<h3>🛡️ Security Controls I'd Need to Bypass</h3>
<ul>
<li><strong>Authentication:</strong> [Strength assessment, bypass techniques]</li>
<li><strong>Authorization:</strong> [Access control gaps]</li>
<li><strong>Input Validation:</strong> [What's not being validated?]</li>
<li><strong>Rate Limiting:</strong> [Present? Bypassable?]</li>
<li><strong>Logging/Monitoring:</strong> [Would my attacks be detected?]</li>
</ul>

<h3>🎯 Bug Bounty Priority List</h3>
<p>If this app had a bug bounty, I'd focus on:</p>
<ol>
<li>[Highest-impact vulnerability class] - [Why it's exploitable here]</li>
<li>[Second target] - [Specific weakness and approach]</li>
<li>[Third target] - [Attack vector and expected payout]</li>
</ol>

<h3>⚠️ Quick Wins for Defenders</h3>
<p>Immediate fixes that would significantly reduce risk:</p>
<ol>
<li>[Critical fix #1 with specific code location]</li>
<li>[Critical fix #2 with implementation guidance]</li>
<li>[Critical fix #3 with detection/prevention tips]</li>
</ol>

## CRITICAL: REAL EXPLOITS ONLY
- Only report vulnerabilities you could ACTUALLY EXPLOIT based on the findings
- Provide SPECIFIC attack scenarios, not theoretical risks
- Include PROOF-OF-CONCEPT ideas (payloads, curl commands, scripts)
- Rate by EXPLOITABILITY (Easy/Medium/Hard), not just severity
- Skip compliance issues that aren't really attackable
- Focus on REAL, EXPLOITABLE issues that would affect users

BE EXHAUSTIVE but ACCURATE. Analyze all provided findings. Reference specific code locations and files."""

            response = client.models.generate_content(
                model=settings.gemini_model_id,
                contents=security_prompt
            )
            if response and response.text:
                security_summary = response.text
                # Clean up any markdown wrappers
                security_summary = security_summary.strip()
                if security_summary.startswith("```html"):
                    security_summary = security_summary[7:]
                if security_summary.startswith("```"):
                    security_summary = security_summary[3:]
                if security_summary.endswith("```"):
                    security_summary = security_summary[:-3]
                security_summary = security_summary.strip()
        
        # Cache the generated summaries in report.data
        if app_summary or security_summary:
            report_data = dict(report.data) if report.data else {}
            report_data["ai_summaries"] = {
                "app_summary": app_summary,
                "security_summary": security_summary
            }
            report.data = report_data
            from sqlalchemy.orm.attributes import flag_modified
            flag_modified(report, "data")
            db.commit()
            logger.info(f"Cached AI summaries for report {report.id}")
                
    except Exception as e:
        logger.error(f"Failed to generate AI summaries: {e}")
    
    return {"app_summary": app_summary, "security_summary": security_summary}


def generate_markdown(
    report: models.Report, 
    findings: List[models.Finding],
    db: Session = None,
    include_ai_summaries: bool = True,
    filtered_findings: List[models.Finding] = None
) -> str:
    """
    Generate a comprehensive Markdown report with proper structure.
    
    Args:
        report: Report model
        findings: List of ACTIVE findings for the report (excludes filtered/duplicates)
        db: Database session (optional, for fetching additional data)
        include_ai_summaries: Whether to generate AI summaries
        filtered_findings: Optional list of filtered findings for transparency section
        
    Returns:
        Markdown string
    """
    lines = []
    
    # Get additional data if db session is provided
    codebase_stats = None
    exploit_scenarios = []
    ai_summaries = {"app_summary": None, "security_summary": None}
    
    if db:
        codebase_stats = get_codebase_statistics(db, report)
        exploit_scenarios = get_exploit_scenarios(db, report.id)
        if include_ai_summaries:
            ai_summaries = generate_ai_summaries(db, report, findings)
    
    # Calculate severity counts for ACTIVE findings only
    severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        sev = (f.severity or "info").lower()
        if sev in severity_counts:
            severity_counts[sev] += 1
    
    severity_emoji = {
        "critical": "🔴",
        "high": "🟠", 
        "medium": "🟡",
        "low": "🟢",
        "info": "🔵"
    }
    
    # ==================== HEADER ====================
    lines.extend([
        "# 🎯 Offensive Security Assessment",
        "",
        f"## {report.title}",
        "",
        "> **For Defensive Application Security** - This report uses an attacker's perspective to help identify and remediate vulnerabilities before they can be exploited.",
        "",
        "---",
        "",
        "| Property | Value |",
        "|----------|-------|",
        f"| **Report ID** | {report.id} |",
        f"| **Generated** | {report.created_at.strftime('%Y-%m-%d %H:%M:%S') if report.created_at else 'N/A'} |",
        f"| **Attack Surface Score** | {report.overall_risk_score or 'N/A'}/100 |",
        f"| **Exploitable Weaknesses** | {len(findings)} |",
        "",
        "---",
        "",
    ])
    
    # ==================== ATTACK SURFACE SUMMARY ====================
    lines.extend([
        "## 🎯 Attack Surface Summary",
        "",
    ])
    
    if report.summary:
        lines.extend([report.summary, ""])
    
    # Quick severity breakdown in summary
    if any(severity_counts.values()):
        lines.extend([
            "### Exploitability Matrix",
            "",
            "| Severity | Count | Exploitation Priority |",
            "|----------|-------|----------------------|",
        ])
        for severity in ["critical", "high", "medium", "low", "info"]:
            count = severity_counts.get(severity, 0)
            if count > 0:
                emoji = severity_emoji.get(severity, "⚪")
                lines.append(f"| {emoji} **{severity.upper()}** | {count} | {'🔥 EXPLOIT IMMEDIATELY' if severity == 'critical' else '⚡ HIGH VALUE TARGET' if severity == 'high' else '🎯 OPPORTUNISTIC' if severity == 'medium' else '📝 LOW PRIORITY'} |")
        lines.append("")
    
    # ==================== AI TARGET ANALYSIS ====================
    if ai_summaries.get("app_summary"):
        lines.extend([
            "---",
            "",
            "## 🔍 Target Reconnaissance (AI Analysis)",
            "",
            "> *AI-powered target analysis for exploit development*",
            "",
            ai_summaries["app_summary"],
            "",
        ])
    elif include_ai_summaries:
        lines.extend([
            "---",
            "",
            "## 🔍 Target Reconnaissance",
            "",
            "> ⚠️ *AI analysis not available. Set GEMINI_API_KEY environment variable to enable AI-powered target analysis.*",
            "",
        ])
    
    # ==================== CODEBASE STATISTICS ====================
    if codebase_stats and codebase_stats.get("total_files", 0) > 0:
        lines.extend([
            "---",
            "",
            "## 📊 Codebase Statistics",
            "",
            "| Metric | Value |",
            "|--------|-------|",
            f"| **Total Files Analyzed** | {codebase_stats['total_files']} |",
            f"| **Total Lines of Code** | {codebase_stats['total_lines']:,} |",
            f"| **Languages Detected** | {len(codebase_stats['languages'])} |",
            "",
        ])
        
        if codebase_stats['languages']:
            lines.extend([
                "### Language Breakdown",
                "",
                "| Language | Files |",
                "|----------|-------|",
            ])
            for lang, count in list(codebase_stats['languages'].items())[:10]:
                lines.append(f"| {lang} | {count} |")
            lines.append("")
    
    # ==================== AI ATTACK VECTOR ANALYSIS ====================
    if ai_summaries.get("security_summary"):
        lines.extend([
            "---",
            "",
            "## ⚔️ Attack Vector Analysis (AI Analysis)",
            "",
            "> *AI-generated offensive analysis for penetration testing*",
            "",
            ai_summaries["security_summary"],
            "",
        ])
    elif include_ai_summaries and findings:
        lines.extend([
            "---",
            "",
            "## ⚔️ Attack Vector Analysis",
            "",
            "> ⚠️ *AI analysis not available. Set GEMINI_API_KEY environment variable to enable AI-powered attack analysis.*",
            "",
        ])
    
    # ==================== AI INSIGHTS (FALSE POSITIVES & ATTACK CHAINS) ====================
    # Extract AI analysis from findings
    false_positives = []
    severity_adjustments = []
    for f in findings:
        if f.details and f.details.get("ai_analysis"):
            ai = f.details["ai_analysis"]
            if ai.get("is_false_positive"):
                false_positives.append({
                    "summary": f.summary,
                    "file": f.file_path,
                    "line": f.start_line,
                    "reason": ai.get("false_positive_reason", "No reason provided"),
                    "confidence": ai.get("confidence", 0)
                })
            if ai.get("severity_adjustment") and ai.get("severity_adjustment") != f.severity:
                severity_adjustments.append({
                    "summary": f.summary,
                    "file": f.file_path,
                    "line": f.start_line,
                    "original": f.severity,
                    "adjusted": ai.get("severity_adjustment"),
                    "reason": ai.get("adjustment_reason", "No reason provided")
                })
    
    # Extract attack chains from report data
    attack_chains = []
    if report.data and report.data.get("ai_insights"):
        attack_chains = report.data["ai_insights"].get("attack_chains", [])
    
    # Only show section if we have AI insights
    if false_positives or severity_adjustments or attack_chains:
        lines.extend([
            "---",
            "",
            "## 🤖 AI-Powered Insights",
            "",
            "> *The following insights were generated by AI analysis to help prioritize remediation efforts and reduce noise.*",
            "",
        ])
        
        # False Positives Section
        if false_positives:
            lines.extend([
                "### 🎯 Likely False Positives",
                "",
                f"The AI analysis identified **{len(false_positives)}** finding(s) that are likely false positives. ",
                "These findings may not represent actual security risks and should be reviewed manually before being dismissed.",
                "",
                "| Finding | Location | Confidence | Reason |",
                "|---------|----------|------------|--------|",
            ])
            for fp in false_positives:
                location = f"`{fp['file']}:{fp['line']}`" if fp['file'] else "N/A"
                confidence = f"{fp['confidence']*100:.0f}%" if fp.get('confidence') else "N/A"
                reason = fp['reason'][:80] + "..." if len(fp['reason']) > 80 else fp['reason']
                lines.append(f"| {fp['summary'][:50]} | {location} | {confidence} | {reason} |")
            lines.extend(["", ""])
        
        # Severity Adjustments Section  
        if severity_adjustments:
            lines.extend([
                "### ⚖️ Severity Adjustments",
                "",
                f"The AI analysis suggests **{len(severity_adjustments)}** severity adjustment(s) based on context analysis:",
                "",
                "| Finding | Original | Suggested | Reason |",
                "|---------|----------|-----------|--------|",
            ])
            for adj in severity_adjustments:
                orig_emoji = severity_emoji.get((adj['original'] or '').lower(), "⚪")
                adj_emoji = severity_emoji.get((adj['adjusted'] or '').lower(), "⚪")
                reason = adj['reason'][:80] + "..." if len(adj['reason']) > 80 else adj['reason']
                lines.append(f"| {adj['summary'][:50]} | {orig_emoji} {adj['original'].upper()} | {adj_emoji} {adj['adjusted'].upper()} | {reason} |")
            lines.extend(["", ""])
        
        # Attack Chains Section
        if attack_chains:
            lines.extend([
                "### ⛓️ Attack Chains Discovered",
                "",
                f"The AI analysis discovered **{len(attack_chains)}** potential attack chain(s) where multiple vulnerabilities could be chained together:",
                "",
            ])
            for i, chain in enumerate(attack_chains, 1):
                chain_sev = chain.get('severity', 'medium').upper()
                chain_emoji = severity_emoji.get(chain.get('severity', 'medium').lower(), "⚪")
                lines.extend([
                    f"#### {i}. {chain_emoji} {chain.get('title', 'Unnamed Chain')}",
                    "",
                    f"**Severity:** {chain_sev} | **Likelihood:** {chain.get('likelihood', 'Unknown')}",
                    "",
                    f"**Description:** {chain.get('chain_description', 'No description')}",
                    "",
                    f"**Potential Impact:** {chain.get('impact', 'Unknown impact')}",
                    "",
                    f"**Findings Involved:** {', '.join([f'#{fid}' for fid in chain.get('finding_ids', [])])}",
                    "",
                ])
            lines.append("")
    
    # ==================== EXPLOITABLE WEAKNESSES ====================
    lines.extend([
        "---",
        "",
        "## 💣 Exploitable Weaknesses",
        "",
    ])
    
    if not findings:
        lines.extend([
            "🛡️ **Target Hardened** - No exploitable weaknesses identified in this assessment.",
            "",
            "The target application appears to have a limited attack surface. Consider:",
            "- Deeper manual testing for logic flaws",
            "- Social engineering vectors",
            "- Infrastructure-level attacks",
            "",
        ])
    else:
        # Group findings by severity for better organization
        findings_by_severity = {"critical": [], "high": [], "medium": [], "low": [], "info": []}
        for f in findings:
            sev = (f.severity or "info").lower()
            if sev in findings_by_severity:
                findings_by_severity[sev].append(f)
        
        finding_number = 1
        for severity in ["critical", "high", "medium", "low", "info"]:
            sev_findings = findings_by_severity[severity]
            if not sev_findings:
                continue
                
            emoji = severity_emoji.get(severity, "⚪")
            lines.extend([
                f"### {emoji} {severity.upper()} Severity ({len(sev_findings)} finding{'s' if len(sev_findings) != 1 else ''})",
                "",
            ])
            
            for finding in sev_findings:
                lines.extend([
                    f"#### {finding_number}. {finding.summary}",
                    "",
                ])
                
                # Create a details table for each finding
                lines.extend([
                    "| Property | Value |",
                    "|----------|-------|",
                    f"| **Type** | `{finding.type}` |",
                    f"| **Severity** | {severity.upper()} |",
                ])
                
                if finding.file_path:
                    lines.append(f"| **File** | `{finding.file_path}` |")
                if finding.start_line:
                    line_info = str(finding.start_line)
                    if finding.end_line and finding.end_line != finding.start_line:
                        line_info += f"-{finding.end_line}"
                    lines.append(f"| **Line(s)** | {line_info} |")
                
                # Add vulnerability details if available
                if finding.details:
                    if finding.details.get("external_id"):
                        ext_id = finding.details["external_id"]
                        if ext_id.startswith("CVE-"):
                            lines.append(f"| **CVE ID** | [{ext_id}](https://nvd.nist.gov/vuln/detail/{ext_id}) |")
                        elif ext_id.startswith("GHSA-"):
                            lines.append(f"| **Advisory** | [{ext_id}](https://github.com/advisories/{ext_id}) |")
                        else:
                            lines.append(f"| **ID** | {ext_id} |")
                    
                    if finding.details.get("cvss_score"):
                        cvss = finding.details["cvss_score"]
                        lines.append(f"| **CVSS Score** | {cvss} ({_cvss_rating(cvss)}) |")
                    
                    if finding.details.get("epss_score") is not None:
                        epss = finding.details["epss_score"]
                        percentile = finding.details.get("epss_percentile", 0)
                        lines.append(f"| **EPSS Score** | {epss:.2%} ({percentile:.0%} percentile) |")
                    
                    if finding.details.get("cwe"):
                        cwe = finding.details["cwe"]
                        if isinstance(cwe, list):
                            cwe = cwe[0] if cwe else None
                        if cwe:
                            cwe = str(cwe)  # Ensure it's a string
                            cwe_num = cwe.replace("CWE-", "")
                            cwe_display = f"CWE-{cwe_num}" if not cwe.startswith("CWE-") else cwe
                            lines.append(f"| **CWE** | [{cwe_display}](https://cwe.mitre.org/data/definitions/{cwe_num}.html) |")
                    
                    if finding.details.get("dependency"):
                        lines.append(f"| **Affected Package** | `{finding.details['dependency']}` |")
                    
                    if finding.details.get("rule_id"):
                        lines.append(f"| **Rule ID** | `{finding.details['rule_id']}` |")
                
                lines.append("")
                
                # Add description if available
                if finding.details and finding.details.get("description"):
                    lines.extend([
                        "**Description:**",
                        "",
                        finding.details["description"],
                        "",
                    ])
                
                # Add NVD description for CVE findings
                if finding.details and finding.details.get("nvd_description"):
                    lines.extend([
                        "**NVD Description:**",
                        "",
                        finding.details["nvd_description"],
                        "",
                    ])
                
                # Add exploitation notes if available (was "fix")
                if finding.details and finding.details.get("fix"):
                    lines.extend([
                        "**Exploitation Notes:**",
                        "",
                        f"> 🎯 {finding.details['fix']}",
                        "",
                    ])
                
                # Add references if available
                if finding.details and finding.details.get("references"):
                    refs = finding.details["references"]
                    if refs and len(refs) > 0:
                        lines.append("**References:**")
                        lines.append("")
                        for ref in refs[:5]:  # Limit to 5 references
                            if isinstance(ref, dict):
                                url = ref.get("url", "")
                                source = ref.get("source", "")
                                if url:
                                    lines.append(f"- [{source or url}]({url})")
                            elif isinstance(ref, str):
                                lines.append(f"- {ref}")
                        lines.append("")
                
                lines.append("---")
                lines.append("")
                finding_number += 1
    
    # ==================== EXPLOIT DEVELOPMENT GUIDE ====================
    if exploit_scenarios:
        lines.extend([
            "## 🔥 Exploit Development Guide (AI-Generated)",
            "",
            "> **For Defensive App Security:** Understanding how attackers exploit vulnerabilities helps ",
            "> security teams prioritize remediation and build stronger defenses.",
            "",
        ])
        
        for i, scenario in enumerate(exploit_scenarios, 1):
            sev_emoji = severity_emoji.get((scenario.severity or "").lower(), "⚪")
            context = get_exploit_context(scenario)
            
            lines.extend([
                f"### {i}. {sev_emoji} {scenario.title}",
                "",
            ])
            
            # Severity and quick info badges
            if scenario.severity:
                lines.extend([
                    f"**Severity:** {scenario.severity.upper()} | **Attack Type:** {context['attack_type']} | **Difficulty:** {context['difficulty']} | **Time Estimate:** {context['time_estimate']}",
                    "",
                ])
            
            # Beginner's Quick Start Guide
            lines.extend([
                "#### 📚 Beginner's Quick Start Guide",
                "",
                "**Skills You'll Need:**",
                "",
            ])
            for skill in context['skills_needed']:
                lines.append(f"- ✓ {skill}")
            lines.append("")
            
            lines.append("**Tools You'll Need:**")
            lines.append("")
            for tool in context['tools_needed']:
                lines.append(f"- 🔧 {tool}")
            lines.append("")
            
            lines.append("**Pro Tips:**")
            lines.append("")
            for tip in context['tips']:
                lines.append(f"- 💡 {tip}")
            lines.append("")
            
            # Attack Development Phases
            lines.extend([
                "#### 🎯 Attack Development Phases",
                "",
                "| Phase | Name | Description |",
                "|-------|------|-------------|",
            ])
            for idx, phase in enumerate(context['phases'], 1):
                lines.append(f"| {idx} | {phase['emoji']} {phase['name']} | {phase['description']} |")
            lines.append("")
            
            # Key Terms Glossary
            lines.extend([
                "#### 📖 Key Terms Glossary",
                "",
            ])
            for term in context['glossary']:
                lines.append(f"- **{term['term']}:** {term['definition']}")
            lines.append("")
            
            # Original AI-generated content
            if scenario.narrative:
                lines.extend([
                    "#### 🎯 Attack Narrative (AI-Generated)",
                    "",
                    scenario.narrative,
                    "",
                ])
            
            if scenario.preconditions:
                lines.extend([
                    "#### 📋 Preconditions",
                    "",
                    scenario.preconditions,
                    "",
                ])
            
            if scenario.impact:
                lines.extend([
                    "#### 💥 Potential Impact",
                    "",
                    scenario.impact,
                    "",
                ])
            
            if scenario.poc_outline:
                lines.extend([
                    "#### 🔧 Proof of Concept Outline",
                    "",
                    "```",
                    scenario.poc_outline,
                    "```",
                    "",
                ])
            
            # Removed mitigation notes - offensive focus
            
            lines.extend(["---", ""])
    
    # ==================== VULNERABLE DEPENDENCIES ====================
    if report.data and report.data.get("affected_packages"):
        packages = [p for p in report.data["affected_packages"] if p]
        if packages:
            lines.extend([
                "## 📦 Vulnerable Dependencies (Attack Surface)",
                "",
                "The following packages have known CVEs that may be exploitable:",
                "",
            ])
            for pkg in sorted(set(packages)):
                lines.append(f"- `{pkg}` - Research exploit-db/GitHub for public exploits")
            lines.extend(["", ""])
    
    # ==================== ATTACK STRATEGY ====================
    lines.extend([
        "## ⚔️ Attack Strategy Notes",
        "",
        "> **For Defensive App Security:** Use this attacker perspective to understand and prioritize your remediation efforts.",
        "",
    ])
    
    critical_count = severity_counts.get("critical", 0)
    high_count = severity_counts.get("high", 0)
    
    if critical_count > 0 or high_count > 0:
        lines.extend([
            "### 🔥 High-Value Targets",
            "",
            f"**{critical_count + high_count} critical/high severity weaknesses** identified for immediate exploitation:",
            "",
            "1. **Prioritize by Impact** - Focus on vulnerabilities that provide:",
            "   - Remote Code Execution (RCE)",
            "   - Authentication Bypass",
            "   - Data Exfiltration paths",
            "",
            "2. **Chain for Maximum Impact** - Look for combinations that:",
            "   - Escalate privileges",
            "   - Enable lateral movement",
            "   - Bypass security controls",
            "",
        ])
    
    lines.extend([
        "### 🎯 Exploitation Methodology",
        "",
        "1. **Reconnaissance Phase**",
        "   - Map the complete attack surface",
        "   - Identify all entry points and data flows",
        "   - Research public exploits for identified CVEs",
        "",
        "2. **Weaponization**",
        "   - Develop or adapt PoC exploits",
        "   - Test in isolated environment first",
        "   - Prepare payload delivery mechanisms",
        "",
        "3. **Exploitation**",
        "   - Execute against highest-value targets first",
        "   - Document all successful attack paths",
        "   - Capture evidence for reporting",
        "",
        "4. **Post-Exploitation**",
        "   - Establish persistence if authorized",
        "   - Enumerate internal resources",
        "   - Identify paths to crown jewels",
        "",
    ])
    
    # ==================== FILTERED FINDINGS (TRANSPARENCY) ====================
    if filtered_findings:
        lines.extend([
            "---",
            "",
            "## 🔍 Filtered Findings (Transparency)",
            "",
            "> **Note:** These findings were filtered out by AI analysis as likely false positives.",
            "> They are included here for transparency and manual review if needed.",
            "",
            f"| Severity | Type | File | FP Score | Reason |",
            "|----------|------|------|----------|--------|",
        ])
        
        for f in filtered_findings[:20]:  # Limit to 20 for readability
            ai_analysis = f.details.get("ai_analysis", {}) if f.details else {}
            fp_score = ai_analysis.get("false_positive_score", 0)
            fp_reason = ai_analysis.get("false_positive_reason", "Not corroborated by AI")[:50]
            severity = f.severity or "info"
            finding_type = f.type or "unknown"
            file_path = f.file_path or "unknown"
            
            # Truncate long values
            if len(finding_type) > 25:
                finding_type = finding_type[:22] + "..."
            if len(file_path) > 30:
                file_path = "..." + file_path[-27:]
            
            lines.append(f"| {severity.upper()} | {finding_type} | {file_path} | {fp_score:.0%} | {fp_reason} |")
        
        if len(filtered_findings) > 20:
            lines.append(f"| ... | *{len(filtered_findings) - 20} more filtered findings* | ... | ... | ... |")
        
        lines.extend([
            "",
            f"**Total Filtered:** {len(filtered_findings)} findings",
            "",
        ])
    
    # ==================== FOOTER ====================
    lines.extend([
        "---",
        "",
        "*Offensive Security Assessment generated by VRAgent*",
        "",
        f"*Assessment completed: {report.created_at.strftime('%Y-%m-%d %H:%M:%S UTC') if report.created_at else 'N/A'}*",
        "",
        "> **Purpose:** This offensive security report is intended for defensive application security - helping teams understand and fix vulnerabilities before attackers can exploit them.",
        "",
    ])
    
    return "\n".join(lines)


def _cvss_rating(score: float) -> str:
    """Convert CVSS score to rating string."""
    if score >= 9.0:
        return "Critical"
    elif score >= 7.0:
        return "High"
    elif score >= 4.0:
        return "Medium"
    elif score >= 0.1:
        return "Low"
    return "None"


def _strip_markdown(text: str) -> str:
    """Strip markdown formatting for PDF/DOCX output.
    
    Converts markdown to plain text suitable for non-markdown renderers.
    """
    if not text:
        return text
    
    import re
    
    # Remove bold/italic markers
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # ***bold italic***
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # **bold**
    text = re.sub(r'\*(.+?)\*', r'\1', text)  # *italic*
    text = re.sub(r'___(.+?)___', r'\1', text)  # ___bold italic___
    text = re.sub(r'__(.+?)__', r'\1', text)  # __bold__
    text = re.sub(r'_(.+?)_', r'\1', text)  # _italic_
    
    # Remove headers but keep the text
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)  # # Header
    
    # Remove code blocks but keep content
    text = re.sub(r'```[\w]*\n?', '', text)  # ```code```
    text = re.sub(r'`([^`]+)`', r'\1', text)  # `inline code`
    
    # Convert bullet points to plain text with dashes
    text = re.sub(r'^\s*[\*\-\+]\s+', '- ', text, flags=re.MULTILINE)
    
    # Remove link formatting, keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # [text](url)
    
    # Remove blockquotes marker
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    
    # Clean up multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def _escape_html(text: str) -> str:
    """Escape HTML/XML special characters for safe use in reportlab Paragraphs.
    
    Reportlab's Paragraph parses text as XML, so we must escape any HTML-like content.
    """
    if not text:
        return text
    
    from html import escape
    return escape(text, quote=True)


def _add_formatted_pdf_text(story, text: str, styles, body_style, subheading_style, ai_style):
    """Convert markdown-like text to formatted PDF paragraphs with proper structure.
    
    Handles headers, bullet points, paragraphs, and maintains visual structure.
    """
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    import re
    
    if not text:
        return
    
    # First strip markdown formatting but keep structure
    stripped = _strip_markdown(text)
    
    # Bullet style for lists
    bullet_style = ParagraphStyle(
        'bullet',
        parent=body_style,
        leftIndent=20,
        bulletIndent=8,
        spaceBefore=2,
        spaceAfter=2
    )
    
    # Split into lines and process
    lines = stripped.split('\n')
    current_paragraph = []
    
    def flush_paragraph():
        """Output accumulated paragraph text."""
        if current_paragraph:
            para_text = ' '.join(current_paragraph)
            para_text = _escape_html(para_text)
            if para_text.strip():
                story.append(Paragraph(para_text, ai_style))
                story.append(Spacer(1, 4))
            current_paragraph.clear()
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - flush current paragraph
            flush_paragraph()
            continue
        
        # Check for section headers (lines that look like titles - all caps or ending with colon)
        if (line.isupper() and len(line) > 5 and len(line) < 100) or \
           (line.endswith(':') and len(line) < 80 and not line.startswith('-')):
            flush_paragraph()
            # Clean up colons at end for display
            header_text = line.rstrip(':')
            story.append(Paragraph(f"<b>{_escape_html(header_text)}</b>", 
                ParagraphStyle('section_header', parent=body_style, 
                               fontSize=10, spaceBefore=8, spaceAfter=4,
                               textColor=colors.HexColor('#2d3748'))))
            continue
        
        # Check for bullet points (- or * at start)
        if re.match(r'^[\-\*\•]\s+', line):
            flush_paragraph()
            bullet_text = re.sub(r'^[\-\*\•]\s+', '', line)
            story.append(Paragraph(f"• {_escape_html(bullet_text)}", bullet_style))
            continue
        
        # Check for numbered points (1. 2. etc)
        num_match = re.match(r'^(\d+)[\.\)]\s+(.+)$', line)
        if num_match:
            flush_paragraph()
            num = num_match.group(1)
            item_text = num_match.group(2)
            story.append(Paragraph(f"{num}. {_escape_html(item_text)}", bullet_style))
            continue
        
        # Regular text - accumulate into paragraph
        current_paragraph.append(line)
    
    # Flush any remaining paragraph
    flush_paragraph()


def _add_markdown_to_docx(doc, text: str):
    """Convert markdown text to formatted DOCX paragraphs.
    
    Handles bold, italic, headers, bullet points, and preserves structure.
    """
    if not text:
        return
    
    import re
    from docx.shared import Pt
    
    lines = text.split('\n')
    current_para = None
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            # Empty line - add spacing
            if current_para:
                current_para = None
            continue
        
        # Check for headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if header_match:
            level = min(len(header_match.group(1)), 4)  # DOCX max heading level
            header_text = header_match.group(2)
            # Strip any remaining markdown from header
            header_text = re.sub(r'\*\*(.+?)\*\*', r'\1', header_text)
            header_text = re.sub(r'\*(.+?)\*', r'\1', header_text)
            doc.add_heading(header_text, level=level)
            current_para = None
            continue
        
        # Check for bullet points
        bullet_match = re.match(r'^[\*\-\+]\s+(.+)$', stripped)
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, bullet_match.group(1))
            current_para = None
            continue
        
        # Check for numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            _add_formatted_runs(p, num_match.group(2))
            current_para = None
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_runs(p, stripped)
        current_para = p


def _add_formatted_runs(paragraph, text: str):
    """Add formatted runs to a paragraph, parsing markdown bold/italic."""
    import re
    from docx.shared import Pt
    
    if not text:
        return
    
    # Pattern to match **bold**, *italic*, ***bold italic***, `code`
    # Process in order of specificity
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`|[^*`]+)'
    
    pos = 0
    while pos < len(text):
        # Try bold italic first (***text***)
        bold_italic_match = re.match(r'\*\*\*(.+?)\*\*\*', text[pos:])
        if bold_italic_match:
            run = paragraph.add_run(bold_italic_match.group(1))
            run.bold = True
            run.italic = True
            pos += bold_italic_match.end()
            continue
        
        # Try bold (**text**)
        bold_match = re.match(r'\*\*(.+?)\*\*', text[pos:])
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
            pos += bold_match.end()
            continue
        
        # Try italic (*text*)
        italic_match = re.match(r'\*(.+?)\*', text[pos:])
        if italic_match:
            run = paragraph.add_run(italic_match.group(1))
            run.italic = True
            pos += italic_match.end()
            continue
        
        # Try code (`text`)
        code_match = re.match(r'`([^`]+)`', text[pos:])
        if code_match:
            run = paragraph.add_run(code_match.group(1))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            pos += code_match.end()
            continue
        
        # Find next special character or end
        next_special = len(text)
        for char in ['*', '`']:
            idx = text.find(char, pos)
            if idx != -1 and idx < next_special:
                next_special = idx
        
        # Add plain text up to next special or end
        if next_special > pos:
            paragraph.add_run(text[pos:next_special])
            pos = next_special
        else:
            # Single special char that didn't match - add it
            paragraph.add_run(text[pos])
            pos += 1


def generate_pdf(
    report: models.Report, 
    findings: List[models.Finding],
    db: Session = None,
    include_ai_summaries: bool = True,
    filtered_findings: List[models.Finding] = None
) -> bytes:
    """
    Generate a comprehensive PDF report with proper structure and AI findings.
    
    Args:
        report: Report model
        findings: List of ACTIVE findings for the report (excludes filtered/duplicates)
        db: Database session (optional, for fetching additional data)
        include_ai_summaries: Whether to generate AI summaries
        filtered_findings: Optional list of filtered findings for transparency
        
    Returns:
        PDF bytes
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, ListFlowable, ListItem
    except ImportError:
        logger.error("reportlab not installed, returning placeholder PDF")
        return b"%PDF-1.4 placeholder - install reportlab for real PDF generation"
    
    # Get additional data if db session is provided
    codebase_stats = None
    exploit_scenarios = []
    ai_summaries = {"app_summary": None, "security_summary": None}
    
    if db:
        codebase_stats = get_codebase_statistics(db, report)
        exploit_scenarios = get_exploit_scenarios(db, report.id)
        if include_ai_summaries:
            ai_summaries = generate_ai_summaries(db, report, findings)
    
    # Calculate severity counts
    severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        sev = (f.severity or "info").lower()
        if sev in severity_counts:
            severity_counts[sev] += 1
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter, 
        topMargin=0.75*inch, 
        bottomMargin=0.75*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=16,
        textColor=colors.HexColor('#1a365d'),
        alignment=1  # Center
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.HexColor('#2c5282')
    )
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor('#4a5568')
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceBefore=4,
        spaceAfter=4,
        leading=14
    )
    ai_style = ParagraphStyle(
        'AIStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceBefore=8,
        spaceAfter=8,
        leading=14,
        leftIndent=10,
        rightIndent=10,
        backColor=colors.HexColor('#f7fafc'),
        borderPadding=8
    )
    
    # Helper to safely create paragraphs with HTML escaping
    def safe_para(text: str, style) -> Paragraph:
        """Create a Paragraph with escaped HTML content."""
        return Paragraph(_escape_html(text or ""), style)
    
    story = []
    
    # ==================== TITLE PAGE ====================
    story.append(Spacer(1, 1*inch))
    story.append(Paragraph("🛡️ Security Assessment Report", title_style))
    story.append(Spacer(1, 20))
    story.append(safe_para(report.title, styles['Heading2']))
    story.append(Spacer(1, 30))
    
    # Report metadata table
    meta_data = [
        ["Report ID", str(report.id)],
        ["Generated", report.created_at.strftime('%Y-%m-%d %H:%M:%S') if report.created_at else 'N/A'],
        ["Risk Score", f"{report.overall_risk_score or 'N/A'}/100"],
        ["Total Findings", str(len(findings))]
    ]
    meta_table = Table(meta_data, colWidths=[2*inch, 3*inch])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#edf2f7')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e0')),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(meta_table)
    story.append(PageBreak())
    
    # ==================== EXECUTIVE SUMMARY ====================
    story.append(Paragraph("📋 Executive Summary", heading_style))
    story.append(safe_para(report.summary or "No summary available.", body_style))
    story.append(Spacer(1, 12))
    
    # Severity breakdown
    if any(severity_counts.values()):
        story.append(Paragraph("Vulnerability Distribution", subheading_style))
        severity_colors = {
            'critical': colors.HexColor('#c53030'),
            'high': colors.HexColor('#dd6b20'),
            'medium': colors.HexColor('#d69e2e'),
            'low': colors.HexColor('#38a169'),
            'info': colors.HexColor('#3182ce')
        }
        
        sev_data = [["Severity", "Count", "Priority"]]
        for sev in ["critical", "high", "medium", "low", "info"]:
            count = severity_counts.get(sev, 0)
            if count > 0:
                priority = "IMMEDIATE" if sev == "critical" else "HIGH" if sev == "high" else "MODERATE" if sev == "medium" else "LOW"
                sev_data.append([sev.upper(), str(count), priority])
        
        if len(sev_data) > 1:
            sev_table = Table(sev_data, colWidths=[1.5*inch, 1*inch, 1.5*inch])
            sev_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2d3748')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e0')),
                ('PADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(sev_table)
            story.append(Spacer(1, 12))
    
    # ==================== AI APPLICATION ANALYSIS ====================
    if ai_summaries.get("app_summary"):
        story.append(Paragraph("🎯 Application Overview (AI Analysis)", heading_style))
        story.append(Paragraph("<i>Analysis generated by Gemini AI based on code review</i>", 
                               ParagraphStyle('italic', parent=styles['Normal'], fontSize=9, textColor=colors.grey)))
        story.append(Spacer(1, 8))
        # Render AI summary with proper paragraph formatting
        _add_formatted_pdf_text(story, ai_summaries["app_summary"], styles, body_style, subheading_style, ai_style)
        story.append(Spacer(1, 12))
    elif include_ai_summaries:
        story.append(Paragraph("🎯 Application Overview", heading_style))
        story.append(Paragraph("<i>⚠️ AI analysis not available. Set GEMINI_API_KEY environment variable to enable AI-powered insights.</i>", 
                               ParagraphStyle('italic', parent=styles['Normal'], fontSize=9, textColor=colors.grey)))
        story.append(Spacer(1, 12))
    
    # ==================== CODEBASE STATISTICS ====================
    if codebase_stats and codebase_stats.get("total_files", 0) > 0:
        story.append(Paragraph("📊 Codebase Statistics", heading_style))
        stats_data = [
            ["Metric", "Value"],
            ["Total Files Analyzed", str(codebase_stats['total_files'])],
            ["Total Lines of Code", f"{codebase_stats['total_lines']:,}"],
            ["Languages Detected", str(len(codebase_stats['languages']))]
        ]
        stats_table = Table(stats_data, colWidths=[2.5*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4299e1')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e0')),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ebf8ff')),
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 12))
    
    # ==================== AI SECURITY ANALYSIS ====================
    if ai_summaries.get("security_summary"):
        story.append(Paragraph("🔒 Security Risk Analysis (AI Analysis)", heading_style))
        story.append(Paragraph("<i>Security assessment generated by Gemini AI</i>", 
                               ParagraphStyle('italic', parent=styles['Normal'], fontSize=9, textColor=colors.grey)))
        story.append(Spacer(1, 8))
        # Render AI summary with proper paragraph formatting
        _add_formatted_pdf_text(story, ai_summaries["security_summary"], styles, body_style, subheading_style, ai_style)
        story.append(Spacer(1, 12))
    elif include_ai_summaries and findings:
        story.append(Paragraph("🔒 Security Risk Analysis", heading_style))
        story.append(Paragraph("<i>⚠️ AI security analysis not available. Set GEMINI_API_KEY to enable.</i>", 
                               ParagraphStyle('italic', parent=styles['Normal'], fontSize=9, textColor=colors.grey)))
        story.append(Spacer(1, 12))
    
    # ==================== DETAILED FINDINGS ====================
    story.append(PageBreak())
    story.append(Paragraph("🔍 Detailed Findings", heading_style))
    
    if not findings:
        story.append(Paragraph("✅ No vulnerabilities found! The scan completed without identifying security issues.", body_style))
    else:
        # Group by severity
        findings_by_severity = {"critical": [], "high": [], "medium": [], "low": [], "info": []}
        for f in findings:
            sev = (f.severity or "info").lower()
            if sev in findings_by_severity:
                findings_by_severity[sev].append(f)
        
        finding_num = 1
        for severity in ["critical", "high", "medium", "low", "info"]:
            sev_findings = findings_by_severity[severity]
            if not sev_findings:
                continue
            
            story.append(Paragraph(f"{severity.upper()} Severity ({len(sev_findings)} findings)", subheading_style))
            
            for finding in sev_findings[:25]:  # Limit per severity
                # Escape HTML in finding summary
                escaped_summary = _escape_html(finding.summary or "")
                story.append(Paragraph(f"<b>{finding_num}. {escaped_summary}</b>", body_style))
                
                details = [f"Type: {_escape_html(finding.type or '')} | Severity: {finding.severity}"]
                if finding.file_path:
                    details.append(f"File: {_escape_html(finding.file_path)}")
                if finding.start_line:
                    details.append(f"Line: {finding.start_line}")
                
                if finding.details:
                    if finding.details.get("external_id"):
                        details.append(f"ID: {_escape_html(str(finding.details['external_id']))}")
                    if finding.details.get("cvss_score"):
                        details.append(f"CVSS: {finding.details['cvss_score']} ({_cvss_rating(finding.details['cvss_score'])})")
                    if finding.details.get("epss_score") is not None:
                        details.append(f"EPSS: {finding.details['epss_score']:.2%}")
                    if finding.details.get("cwe"):
                        cwe = finding.details["cwe"]
                        if isinstance(cwe, list):
                            cwe = cwe[0] if cwe else ""
                        if cwe:
                            cwe = str(cwe)  # Ensure it's a string
                            cwe_display = f"CWE-{cwe}" if not cwe.startswith("CWE-") else cwe
                            details.append(f"CWE: {cwe_display}")
                
                story.append(Paragraph(" | ".join(details), 
                    ParagraphStyle('details', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#4a5568'))))
                story.append(Spacer(1, 8))
                finding_num += 1
    
    # ==================== EXPLOIT DEVELOPMENT GUIDE ====================
    if exploit_scenarios:
        story.append(PageBreak())
        story.append(Paragraph("🔥 Exploit Development Guide (AI-Generated)", heading_style))
        story.append(Paragraph(
            "<i>For Defensive App Security: Understanding exploitation helps prioritize remediation</i>", 
            ParagraphStyle('disclaimer', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#2b6cb0'))
        ))
        story.append(Spacer(1, 12))
        
        for i, scenario in enumerate(exploit_scenarios[:15], 1):  # Limit to 15
            context = get_exploit_context(scenario)
            
            # Escape HTML in all scenario fields to prevent XML parsing errors
            escaped_title = _escape_html(scenario.title or "")
            story.append(Paragraph(f"<b>{i}. {escaped_title}</b>", subheading_style))
            
            # Severity and quick info
            if scenario.severity:
                info_line = f"Severity: {scenario.severity.upper()} | Attack Type: {context['attack_type']} | Difficulty: {context['difficulty']} | Time: {context['time_estimate']}"
                story.append(Paragraph(info_line, 
                    ParagraphStyle('info', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#2b6cb0'))))
            
            # Beginner's Quick Start - compact for PDF
            story.append(Paragraph("<b>📚 Quick Start Guide:</b>", body_style))
            skills_text = "Skills: " + ", ".join(context['skills_needed'][:3])
            tools_text = "Tools: " + ", ".join(context['tools_needed'][:3])
            story.append(Paragraph(_escape_html(skills_text), 
                ParagraphStyle('skills', parent=styles['Normal'], fontSize=9, leftIndent=10)))
            story.append(Paragraph(_escape_html(tools_text), 
                ParagraphStyle('tools', parent=styles['Normal'], fontSize=9, leftIndent=10)))
            
            # Attack Phases - compact table
            story.append(Paragraph("<b>🎯 Attack Phases:</b>", body_style))
            phases_data = [[f"{idx+1}. {p['emoji']} {p['name']}", p['description'][:50] + "..." if len(p['description']) > 50 else p['description']] 
                          for idx, p in enumerate(context['phases'])]
            if phases_data:
                phases_table = Table(phases_data, colWidths=[1.8*inch, 4*inch])
                phases_table.setStyle(TableStyle([
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                    ('TOPPADDING', (0, 0), (-1, -1), 2),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ]))
                story.append(phases_table)
            
            # Glossary - compact
            story.append(Paragraph("<b>📖 Key Terms:</b>", body_style))
            for term in context['glossary'][:2]:  # Limit to 2 for PDF space
                story.append(Paragraph(f"• <b>{_escape_html(term['term'])}:</b> {_escape_html(term['definition'][:60])}...", 
                    ParagraphStyle('glossary', parent=styles['Normal'], fontSize=8, leftIndent=10)))
            
            # Original AI content
            if scenario.narrative:
                story.append(Paragraph("<b>🎯 Attack Narrative:</b>", body_style))
                narrative_text = _escape_html(scenario.narrative[:500] + ("..." if len(scenario.narrative) > 500 else ""))
                story.append(Paragraph(narrative_text, body_style))
            
            if scenario.impact:
                story.append(Paragraph("<b>💥 Impact:</b>", body_style))
                impact_text = _escape_html(scenario.impact[:300] + ("..." if len(scenario.impact) > 300 else ""))
                story.append(Paragraph(impact_text, body_style))
            
            if scenario.poc_outline:
                story.append(Paragraph("<b>🔧 PoC Outline:</b>", body_style))
                poc_text = _escape_html(scenario.poc_outline[:400] + ("..." if len(scenario.poc_outline) > 400 else ""))
                story.append(Paragraph(f"<font face='Courier' size='8'>{poc_text}</font>", body_style))
            
            # Removed mitigation - offensive focus
            
            story.append(Spacer(1, 16))
    
    # ==================== ATTACK STRATEGY ====================
    story.append(PageBreak())
    story.append(Paragraph("⚔️ Attack Strategy Notes", heading_style))
    story.append(Paragraph(
        "<i>For Defensive App Security: Use this attacker perspective to prioritize your remediation</i>",
        ParagraphStyle('note', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#2b6cb0'))
    ))
    story.append(Spacer(1, 8))
    
    critical_high = severity_counts.get("critical", 0) + severity_counts.get("high", 0)
    if critical_high > 0:
        story.append(Paragraph(f"<b>🔥 High-Value Targets:</b> {critical_high} critical/high severity weaknesses identified", body_style))
        story.append(Spacer(1, 8))
    
    attack_methodology = [
        "Reconnaissance: Map complete attack surface and identify entry points",
        "Weaponization: Develop PoC exploits and prepare payload delivery",
        "Exploitation: Execute against highest-value targets first",
        "Post-Exploitation: Establish persistence and enumerate internal resources",
        "Document all successful attack paths for reporting"
    ]
    
    for item in attack_methodology:
        story.append(Paragraph(f"• {item}", body_style))
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph(
        f"<i>Offensive Security Assessment - VRAgent - For Defensive App Security Purposes</i>",
        ParagraphStyle('footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey, alignment=1)
    ))
    story.append(Paragraph(
        f"<i>{report.created_at.strftime('%Y-%m-%d %H:%M:%S') if report.created_at else 'N/A'}</i>",
        ParagraphStyle('footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey, alignment=1)
    ))
    
    doc.build(story)
    logger.info(f"Generated PDF report for report {report.id}")
    return buffer.getvalue()


def generate_docx(
    report: models.Report, 
    findings: List[models.Finding],
    db: Session = None,
    include_ai_summaries: bool = True,
    filtered_findings: List[models.Finding] = None
) -> bytes:
    """
    Generate a comprehensive DOCX report with proper structure and AI findings.
    
    Args:
        report: Report model
        findings: List of ACTIVE findings for the report (excludes filtered/duplicates)
        db: Database session (optional, for fetching additional data)
        include_ai_summaries: Whether to generate AI summaries
        filtered_findings: Optional list of filtered findings for transparency
        
    Returns:
        DOCX bytes
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        logger.error("python-docx not installed, returning placeholder DOCX")
        return b"PK placeholder - install python-docx for real DOCX generation"
    
    # Get additional data if db session is provided
    codebase_stats = None
    exploit_scenarios = []
    ai_summaries = {"app_summary": None, "security_summary": None}
    
    if db:
        codebase_stats = get_codebase_statistics(db, report)
        exploit_scenarios = get_exploit_scenarios(db, report.id)
        if include_ai_summaries:
            ai_summaries = generate_ai_summaries(db, report, findings)
    
    # Calculate severity counts
    severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        sev = (f.severity or "info").lower()
        if sev in severity_counts:
            severity_counts[sev] += 1
    
    doc = Document()
    
    # ==================== TITLE PAGE ====================
    title = doc.add_heading("🛡️ Security Assessment Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(report.title, level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata table
    doc.add_paragraph()
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ("Report ID", str(report.id)),
        ("Generated", report.created_at.strftime('%Y-%m-%d %H:%M:%S') if report.created_at else 'N/A'),
        ("Risk Score", f"{report.overall_risk_score or 'N/A'}/100"),
        ("Total Findings", str(len(findings)))
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        # Bold the label
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # ==================== EXECUTIVE SUMMARY ====================
    doc.add_heading("📋 Executive Summary", level=1)
    doc.add_paragraph(report.summary or "No summary available.")
    
    # Severity breakdown
    if any(severity_counts.values()):
        doc.add_heading("Vulnerability Distribution", level=2)
        sev_table = doc.add_table(rows=1, cols=3)
        sev_table.style = 'Table Grid'
        hdr_cells = sev_table.rows[0].cells
        hdr_cells[0].text = "Severity"
        hdr_cells[1].text = "Count"
        hdr_cells[2].text = "Priority"
        
        for sev in ["critical", "high", "medium", "low", "info"]:
            count = severity_counts.get(sev, 0)
            if count > 0:
                row_cells = sev_table.add_row().cells
                row_cells[0].text = sev.upper()
                row_cells[1].text = str(count)
                priority = "IMMEDIATE" if sev == "critical" else "HIGH" if sev == "high" else "MODERATE" if sev == "medium" else "LOW"
                row_cells[2].text = priority
    
    # ==================== AI APPLICATION ANALYSIS ====================
    if ai_summaries.get("app_summary"):
        doc.add_heading("🎯 Application Overview (AI Analysis)", level=1)
        p = doc.add_paragraph()
        p.add_run("Analysis generated by Gemini AI based on code review").italic = True
        doc.add_paragraph()
        # Convert markdown to formatted DOCX
        _add_markdown_to_docx(doc, ai_summaries["app_summary"])
    elif include_ai_summaries:
        doc.add_heading("🎯 Application Overview", level=1)
        p = doc.add_paragraph()
        p.add_run("⚠️ AI analysis not available. Set GEMINI_API_KEY environment variable to enable AI-powered insights.").italic = True
    
    # ==================== CODEBASE STATISTICS ====================
    if codebase_stats and codebase_stats.get("total_files", 0) > 0:
        doc.add_heading("📊 Codebase Statistics", level=1)
        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ("Total Files Analyzed", str(codebase_stats['total_files'])),
            ("Total Lines of Code", f"{codebase_stats['total_lines']:,}"),
            ("Languages Detected", str(len(codebase_stats['languages']))),
            ("Top Languages", ", ".join(list(codebase_stats['languages'].keys())[:5]))
        ]
        
        for i, (metric, value) in enumerate(stats_data):
            stats_table.rows[i].cells[0].text = metric
            stats_table.rows[i].cells[1].text = value
            stats_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    # ==================== AI SECURITY ANALYSIS ====================
    if ai_summaries.get("security_summary"):
        doc.add_heading("🔒 Security Risk Analysis (AI Analysis)", level=1)
        p = doc.add_paragraph()
        p.add_run("Security assessment generated by Gemini AI").italic = True
        doc.add_paragraph()
        # Convert markdown to formatted DOCX
        _add_markdown_to_docx(doc, ai_summaries["security_summary"])
    elif include_ai_summaries and findings:
        doc.add_heading("🔒 Security Risk Analysis", level=1)
        p = doc.add_paragraph()
        p.add_run("⚠️ AI security analysis not available. Set GEMINI_API_KEY to enable.").italic = True
    
    # ==================== DETAILED FINDINGS ====================
    doc.add_heading("🔍 Detailed Findings", level=1)
    
    if not findings:
        p = doc.add_paragraph()
        p.add_run("✅ No vulnerabilities found!").bold = True
        doc.add_paragraph("The scan completed without identifying security issues.")
    else:
        # Group by severity
        findings_by_severity = {"critical": [], "high": [], "medium": [], "low": [], "info": []}
        for f in findings:
            sev = (f.severity or "info").lower()
            if sev in findings_by_severity:
                findings_by_severity[sev].append(f)
        
        finding_num = 1
        for severity in ["critical", "high", "medium", "low", "info"]:
            sev_findings = findings_by_severity[severity]
            if not sev_findings:
                continue
            
            doc.add_heading(f"{severity.upper()} Severity ({len(sev_findings)} findings)", level=2)
            
            for finding in sev_findings:
                # Finding title
                doc.add_heading(f"{finding_num}. {finding.summary}", level=3)
                
                # Create details table
                details_table = doc.add_table(rows=0, cols=2)
                details_table.style = 'Table Grid'
                
                # Add rows for each detail
                details_rows = [
                    ("Type", finding.type),
                    ("Severity", (finding.severity or "info").upper()),
                ]
                
                if finding.file_path:
                    details_rows.append(("File", finding.file_path))
                if finding.start_line:
                    line_info = str(finding.start_line)
                    if finding.end_line and finding.end_line != finding.start_line:
                        line_info += f"-{finding.end_line}"
                    details_rows.append(("Line(s)", line_info))
                
                if finding.details:
                    if finding.details.get("external_id"):
                        details_rows.append(("ID", finding.details["external_id"]))
                    if finding.details.get("cvss_score"):
                        score = finding.details["cvss_score"]
                        details_rows.append(("CVSS Score", f"{score} ({_cvss_rating(score)})"))
                    if finding.details.get("epss_score") is not None:
                        epss = finding.details["epss_score"]
                        details_rows.append(("EPSS Score", f"{epss:.2%}"))
                    if finding.details.get("cwe"):
                        cwe = finding.details["cwe"]
                        if isinstance(cwe, list):
                            cwe = cwe[0] if cwe else ""
                        if cwe:
                            cwe = str(cwe)  # Ensure it's a string
                            cwe_display = f"CWE-{cwe}" if not cwe.startswith("CWE-") else cwe
                            details_rows.append(("CWE", cwe_display))
                    if finding.details.get("dependency"):
                        details_rows.append(("Affected Package", finding.details["dependency"]))
                
                for label, value in details_rows:
                    row = details_table.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = str(value)
                    row.cells[0].paragraphs[0].runs[0].bold = True
                
                # Add description if available
                if finding.details and finding.details.get("description"):
                    p = doc.add_paragraph()
                    p.add_run("Description: ").bold = True
                    p.add_run(finding.details["description"])
                
                # Add fix if available
                if finding.details and finding.details.get("fix"):
                    p = doc.add_paragraph()
                    p.add_run("Recommended Fix: ").bold = True
                    p.add_run(finding.details["fix"])
                
                doc.add_paragraph()  # Spacing
                finding_num += 1
    
    # ==================== RED TEAM ANALYSIS ====================
    if exploit_scenarios:
        doc.add_heading("🔥 Red Team Analysis (AI-Generated)", level=1)
        
        # Disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer.add_run("Defensive Purpose: ").bold = True
        disclaimer.add_run("Understanding how attackers might exploit these vulnerabilities helps "
                          "security teams prioritize remediation and implement effective detection controls.")
        doc.add_paragraph()
        
        for i, scenario in enumerate(exploit_scenarios, 1):
            context = get_exploit_context(scenario)
            
            doc.add_heading(f"{i}. {scenario.title}", level=2)
            
            # Severity and quick info
            if scenario.severity:
                p = doc.add_paragraph()
                p.add_run("Severity: ").bold = True
                p.add_run(f"{scenario.severity.upper()} | ")
                p.add_run("Attack Type: ").bold = True
                p.add_run(f"{context['attack_type']} | ")
                p.add_run("Difficulty: ").bold = True
                p.add_run(f"{context['difficulty']} | ")
                p.add_run("Time Estimate: ").bold = True
                p.add_run(context['time_estimate'])
            
            # Beginner's Quick Start Guide
            doc.add_heading("📚 Beginner's Quick Start Guide", level=3)
            
            # Skills needed
            p = doc.add_paragraph()
            p.add_run("Skills You'll Need:").bold = True
            for skill in context['skills_needed']:
                doc.add_paragraph(f"✓ {skill}", style='List Bullet')
            
            # Tools needed
            p = doc.add_paragraph()
            p.add_run("Tools You'll Need:").bold = True
            for tool in context['tools_needed']:
                doc.add_paragraph(f"🔧 {tool}", style='List Bullet')
            
            # Pro tips
            p = doc.add_paragraph()
            p.add_run("Pro Tips:").bold = True
            for tip in context['tips']:
                doc.add_paragraph(f"💡 {tip}", style='List Bullet')
            
            # Attack Development Phases
            doc.add_heading("🎯 Attack Development Phases", level=3)
            phases_table = doc.add_table(rows=len(context['phases']) + 1, cols=3)
            phases_table.style = 'Table Grid'
            
            # Header row
            header_cells = phases_table.rows[0].cells
            header_cells[0].text = "Phase"
            header_cells[1].text = "Name"
            header_cells[2].text = "Description"
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            for idx, phase in enumerate(context['phases']):
                row_cells = phases_table.rows[idx + 1].cells
                row_cells[0].text = str(idx + 1)
                row_cells[1].text = f"{phase['emoji']} {phase['name']}"
                row_cells[2].text = phase['description']
            
            doc.add_paragraph()
            
            # Key Terms Glossary
            doc.add_heading("📖 Key Terms Glossary", level=3)
            for term in context['glossary']:
                p = doc.add_paragraph()
                p.add_run(f"{term['term']}: ").bold = True
                p.add_run(term['definition'])
            
            doc.add_paragraph()
            
            # Original AI-generated content
            if scenario.narrative:
                doc.add_heading("🎯 Attack Narrative (AI-Generated)", level=3)
                doc.add_paragraph(scenario.narrative)
            
            if scenario.preconditions:
                doc.add_heading("📋 Preconditions", level=3)
                doc.add_paragraph(scenario.preconditions)
            
            if scenario.impact:
                doc.add_heading("💥 Potential Impact", level=3)
                doc.add_paragraph(scenario.impact)
            
            if scenario.poc_outline:
                doc.add_heading("🔧 Proof of Concept Outline", level=3)
                # Add as code-like paragraph
                p = doc.add_paragraph()
                run = p.add_run(scenario.poc_outline)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            # Removed mitigation - offensive focus
            
            doc.add_paragraph()  # Spacing between scenarios
    
    # ==================== AFFECTED DEPENDENCIES ====================
    if report.data and report.data.get("affected_packages"):
        packages = [p for p in report.data["affected_packages"] if p]
        if packages:
            doc.add_heading("📦 Vulnerable Dependencies", level=1)
            for pkg in sorted(set(packages)):
                doc.add_paragraph(f"{pkg} - Check for available patches or updates", style='List Bullet')
    
    # ==================== ATTACK STRATEGY ====================
    doc.add_heading("⚔️ Attack Strategy Notes", level=1)
    
    # Defensive purpose notice
    notice = doc.add_paragraph()
    notice.add_run("For Defensive App Security: ").bold = True
    notice.add_run("Use this attacker perspective to understand and prioritize your remediation efforts.")
    doc.add_paragraph()
    
    critical_high = severity_counts.get("critical", 0) + severity_counts.get("high", 0)
    if critical_high > 0:
        p = doc.add_paragraph()
        p.add_run(f"🔥 High-Value Targets: ").bold = True
        p.add_run(f"{critical_high} critical/high severity weaknesses identified for immediate exploitation")
        doc.add_paragraph()
    
    doc.add_heading("Exploitation Methodology", level=2)
    attack_methodology = [
        "Reconnaissance: Map complete attack surface and identify entry points",
        "Weaponization: Develop PoC exploits and prepare payload delivery",
        "Exploitation: Execute against highest-value targets first",
        "Post-Exploitation: Establish persistence and enumerate internal resources",
        "Document all successful attack paths for reporting"
    ]
    for item in attack_methodology:
        doc.add_paragraph(item, style='List Bullet')
    
    # ==================== FOOTER ====================
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Offensive Security Assessment - VRAgent").italic = True
    footer.add_run(f"\n{report.created_at.strftime('%Y-%m-%d %H:%M:%S') if report.created_at else 'N/A'}").italic = True
    footer.add_run("\nFor Defensive Application Security Purposes").italic = True
    
    buffer = BytesIO()
    doc.save(buffer)
    logger.info(f"Generated DOCX report for report {report.id}")
    return buffer.getvalue()


def get_report_with_findings(db: Session, report_id: int, include_filtered: bool = False) -> tuple[models.Report, List[models.Finding], List[models.Finding]]:
    """
    Fetch a report and its associated findings.
    
    Args:
        db: Database session
        report_id: Report ID to fetch
        include_filtered: If True, returns filtered findings as third element
        
    Returns:
        Tuple of (report, active_findings, filtered_findings)
        - active_findings: Findings not filtered out by AI analysis and not duplicates
        - filtered_findings: Findings marked as likely false positives (for transparency)
        
    Raises:
        ValueError: If report not found
    """
    report = db.get(models.Report, report_id)
    if not report:
        raise ValueError(f"Report {report_id} not found")
    
    all_findings = (
        db.query(models.Finding)
        .filter(models.Finding.scan_run_id == report.scan_run_id)
        .order_by(
            # Order by severity (critical first)
            models.Finding.severity.desc(),
            models.Finding.id
        )
        .all()
    )
    
    # Separate active findings from filtered/duplicate findings
    active_findings = []
    filtered_findings = []
    
    for f in all_findings:
        # Check if finding is a duplicate (merged into another finding)
        if f.is_duplicate:
            continue  # Skip duplicates entirely - they're represented by the primary finding
        
        # Check if finding was filtered out by AI analysis as likely false positive
        ai_analysis = f.details.get("ai_analysis", {}) if f.details else {}
        is_filtered = ai_analysis.get("filtered_out", False)
        
        if is_filtered:
            filtered_findings.append(f)
        else:
            active_findings.append(f)
    
    return report, active_findings, filtered_findings
