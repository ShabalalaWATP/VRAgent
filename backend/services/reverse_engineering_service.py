"""
Reverse Engineering Service for VRAgent.

Provides analysis capabilities for:
- Binary files (EXE, ELF, DLL, SO)
- Android APK files
- Docker image layers
"""

import os
import re
import json
import struct
import tempfile
import shutil
import zipfile
import subprocess
import math
import asyncio
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field, asdict
from xml.etree import ElementTree

from backend.core.logging import get_logger
from backend.core.config import settings

logger = get_logger(__name__)

# Suppress Androguard's extremely verbose DEBUG logging (uses loguru)
# This dramatically improves performance by avoiding thousands of log entries
try:
    from loguru import logger as loguru_logger
    loguru_logger.disable("androguard")
    loguru_logger.disable("androguard.core")
    loguru_logger.disable("androguard.core.axml")
    loguru_logger.disable("androguard.core.apk")
    loguru_logger.disable("androguard.core.analysis")
    loguru_logger.disable("androguard.core.dex")
except ImportError:
    pass

# Also suppress standard logging just in case
logging.getLogger('androguard').setLevel(logging.ERROR)
logging.getLogger('androguard.core').setLevel(logging.ERROR)

# Try to import pefile for better PE analysis
try:
    import pefile
    PEFILE_AVAILABLE = True
except ImportError:
    PEFILE_AVAILABLE = False
    logger.warning("pefile not installed - using basic PE parsing")

# Try to import pyelftools for comprehensive ELF analysis
try:
    from elftools.elf.elffile import ELFFile
    from elftools.elf.sections import SymbolTableSection
    from elftools.elf.dynamic import DynamicSection
    from elftools.elf.relocation import RelocationSection
    from elftools.elf.gnuversions import GNUVerDefSection, GNUVerNeedSection
    from elftools.dwarf.dwarfinfo import DWARFInfo
    PYELFTOOLS_AVAILABLE = True
except ImportError:
    PYELFTOOLS_AVAILABLE = False
    logger.warning("pyelftools not installed - using basic ELF parsing")

# Try to import Capstone for disassembly
try:
    import capstone
    from capstone import Cs, CS_ARCH_X86, CS_ARCH_ARM, CS_ARCH_ARM64, CS_ARCH_MIPS
    from capstone import CS_MODE_32, CS_MODE_64, CS_MODE_ARM, CS_MODE_THUMB
    CAPSTONE_AVAILABLE = True
except ImportError:
    CAPSTONE_AVAILABLE = False
    logger.warning("capstone not installed - disassembly not available")

# Try to import androguard for APK analysis
try:
    from androguard.core.apk import APK
    from androguard.core.axml import AXMLPrinter
    ANDROGUARD_AVAILABLE = True
except ImportError:
    ANDROGUARD_AVAILABLE = False
    logger.warning("androguard not installed - using basic APK parsing")


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class ExtractedString:
    """A string extracted from a binary."""
    value: str
    offset: int
    encoding: str  # "ascii" or "utf16"
    category: Optional[str] = None  # "url", "email", "path", "api_key", etc.
    confidence: float = 1.0


@dataclass
class ImportedFunction:
    """An imported function from a binary."""
    name: str
    library: str
    ordinal: Optional[int] = None
    is_suspicious: bool = False
    reason: Optional[str] = None


@dataclass
class RichHeaderEntry:
    """An entry in the PE Rich header."""
    product_id: int
    build_id: int
    count: int
    product_name: Optional[str] = None
    vs_version: Optional[str] = None


@dataclass
class RichHeader:
    """PE Rich header information."""
    entries: List[RichHeaderEntry]
    rich_hash: str  # MD5 hash of Rich header for malware identification
    checksum: int
    raw_data: str  # Hex representation
    clear_data: str  # Decrypted hex representation


@dataclass
class BinaryMetadata:
    """Metadata extracted from a binary file."""
    file_type: str
    architecture: str
    file_size: int
    entry_point: Optional[int] = None
    is_packed: bool = False
    packer_name: Optional[str] = None
    compile_time: Optional[str] = None
    sections: List[Dict[str, Any]] = field(default_factory=list)
    headers: Dict[str, Any] = field(default_factory=dict)
    # PE-specific fields
    rich_header: Optional[RichHeader] = None
    imphash: Optional[str] = None
    # ELF-specific fields
    interpreter: Optional[str] = None
    linked_libraries: List[str] = field(default_factory=list)
    relro: Optional[str] = None  # "Full", "Partial", "None"
    stack_canary: bool = False
    nx_enabled: bool = False
    pie_enabled: bool = False


@dataclass
class ELFSymbol:
    """A symbol from an ELF binary."""
    name: str
    address: int
    size: int
    symbol_type: str  # "FUNC", "OBJECT", "NOTYPE", etc.
    binding: str  # "LOCAL", "GLOBAL", "WEAK"
    section: str
    is_imported: bool = False
    is_exported: bool = False
    is_suspicious: bool = False
    reason: Optional[str] = None


@dataclass
class DisassemblyInstruction:
    """A disassembled instruction."""
    address: int
    mnemonic: str
    op_str: str
    bytes_hex: str
    size: int
    is_call: bool = False
    is_jump: bool = False
    is_suspicious: bool = False
    comment: Optional[str] = None


@dataclass
class DisassemblyFunction:
    """Disassembly of a function."""
    name: str
    address: int
    size: int
    instructions: List[DisassemblyInstruction]
    calls: List[str]  # Functions called
    suspicious_patterns: List[str]


@dataclass
class DisassemblyResult:
    """Complete disassembly result."""
    entry_point_disasm: List[DisassemblyInstruction]
    functions: List[DisassemblyFunction]
    suspicious_instructions: List[Dict[str, Any]]
    architecture: str
    mode: str


@dataclass
class BinaryAnalysisResult:
    """Complete analysis result for a binary file."""
    filename: str
    metadata: BinaryMetadata
    strings: List[ExtractedString]
    imports: List[ImportedFunction]
    exports: List[str]
    secrets: List[Dict[str, Any]]
    suspicious_indicators: List[Dict[str, Any]]
    # Enhanced ELF fields
    symbols: List[ELFSymbol] = field(default_factory=list)
    disassembly: Optional[DisassemblyResult] = None
    dwarf_info: Optional[Dict[str, Any]] = None
    ai_analysis: Optional[str] = None
    error: Optional[str] = None


@dataclass
class ApkCertificate:
    """APK signing certificate information."""
    subject: str
    issuer: str
    serial_number: str
    fingerprint_sha256: str
    fingerprint_sha1: str
    fingerprint_md5: str
    valid_from: str
    valid_until: str
    is_debug_cert: bool = False
    is_expired: bool = False
    is_self_signed: bool = False
    signature_version: str = "v1"  # v1, v2, v3
    public_key_algorithm: Optional[str] = None
    public_key_bits: Optional[int] = None


@dataclass
class ApkPermission:
    """An Android permission."""
    name: str
    is_dangerous: bool
    description: Optional[str] = None


@dataclass
class ApkComponent:
    """An Android app component."""
    name: str
    component_type: str  # "activity", "service", "receiver", "provider"
    is_exported: bool
    intent_filters: List[str] = field(default_factory=list)


@dataclass
class ApkAnalysisResult:
    """Complete analysis result for an APK file."""
    filename: str
    package_name: str
    version_name: Optional[str]
    version_code: Optional[int]
    min_sdk: Optional[int]
    target_sdk: Optional[int]
    permissions: List[ApkPermission]
    components: List[ApkComponent]
    strings: List[ExtractedString]
    secrets: List[Dict[str, Any]]
    urls: List[str]
    native_libraries: List[str]
    certificate: Optional[ApkCertificate] = None
    activities: List[str] = field(default_factory=list)
    services: List[str] = field(default_factory=list)
    receivers: List[str] = field(default_factory=list)
    providers: List[str] = field(default_factory=list)
    uses_features: List[str] = field(default_factory=list)
    app_name: Optional[str] = None
    debuggable: bool = False
    allow_backup: bool = True
    network_security_config: Optional[str] = None
    # New analysis fields
    dex_analysis: Optional[Dict[str, Any]] = None
    resource_analysis: Optional[Dict[str, Any]] = None
    intent_filter_analysis: Optional[Dict[str, Any]] = None
    network_config_analysis: Optional[Dict[str, Any]] = None
    smali_analysis: Optional[Dict[str, Any]] = None  # Smali/bytecode decompilation
    dynamic_analysis: Optional[Dict[str, Any]] = None  # Frida scripts for dynamic testing
    native_analysis: Optional[Dict[str, Any]] = None  # Native library (.so) analysis
    hardening_score: Optional[Dict[str, Any]] = None  # Security hardening score
    data_flow_analysis: Optional[Dict[str, Any]] = None  # Data flow/taint analysis
    security_issues: List[Dict[str, Any]] = field(default_factory=list)
    ai_analysis: Optional[str] = None
    # New structured AI reports
    ai_report_functionality: Optional[str] = None  # "What does this APK do" report
    ai_report_security: Optional[str] = None  # "Security Findings" report
    # AI-Generated Mermaid Diagrams (with icons)
    ai_architecture_diagram: Optional[str] = None  # App architecture flowchart
    ai_data_flow_diagram: Optional[str] = None  # Data flow and privacy diagram
    error: Optional[str] = None


@dataclass
class SmaliMethodCode:
    """Decompiled Smali bytecode for a method."""
    class_name: str
    method_name: str
    method_signature: str
    access_flags: str
    return_type: str
    parameters: List[str]
    registers_count: int
    instructions: List[str]  # Smali bytecode instructions
    instruction_count: int
    has_try_catch: bool = False
    is_native: bool = False
    is_abstract: bool = False


@dataclass
class DexClassInfo:
    """Information about a class in DEX."""
    name: str
    access_flags: str
    superclass: Optional[str]
    interfaces: List[str]
    methods_count: int
    fields_count: int
    is_suspicious: bool = False
    suspicious_reasons: List[str] = field(default_factory=list)


@dataclass
class DexMethodInfo:
    """Information about a method in DEX."""
    class_name: str
    method_name: str
    access_flags: str
    return_type: str
    parameters: List[str]
    is_suspicious: bool = False
    suspicious_reason: Optional[str] = None


@dataclass
class ApkResourceInfo:
    """Information about APK resources."""
    string_resources: Dict[str, str]  # name -> value
    asset_files: List[str]
    raw_resources: List[str]
    drawable_count: int
    layout_count: int
    potential_secrets_in_resources: List[Dict[str, Any]]


@dataclass
class IntentFilterInfo:
    """Deep link and intent filter information."""
    component_name: str
    component_type: str
    actions: List[str]
    categories: List[str]
    data_schemes: List[str]
    data_hosts: List[str]
    data_paths: List[str]
    is_browsable: bool
    is_exported: bool
    deep_links: List[str]


@dataclass
class NetworkSecurityConfig:
    """Parsed network security configuration."""
    cleartext_permitted: bool
    cleartext_domains: List[str]
    trust_anchors: List[Dict[str, Any]]
    certificate_pins: List[Dict[str, Any]]
    domain_configs: List[Dict[str, Any]]
    security_issues: List[str]


@dataclass
class DockerLayerSecret:
    """A potential secret found in a Docker layer."""
    layer_id: str
    layer_command: str
    secret_type: str
    value: str
    masked_value: str
    context: str
    severity: str


@dataclass
class DockerLayerAnalysisResult:
    """Analysis result for Docker image layers."""
    image_name: str
    image_id: str
    total_layers: int
    total_size: int
    base_image: Optional[str]
    layers: List[Dict[str, Any]]
    secrets: List[DockerLayerSecret]
    deleted_files: List[Dict[str, Any]]
    security_issues: List[Dict[str, Any]]
    ai_analysis: Optional[str] = None
    error: Optional[str] = None


@dataclass
class FridaScript:
    """A generated Frida script for dynamic analysis."""
    name: str
    category: str  # ssl_bypass, root_bypass, crypto_hook, method_trace, etc.
    description: str
    script_code: str
    target_classes: List[str]
    target_methods: List[str]
    is_dangerous: bool = False  # Scripts that modify app behavior
    usage_instructions: str = ""


@dataclass
class DynamicAnalysisResult:
    """Dynamic analysis data including generated Frida scripts."""
    package_name: str
    frida_scripts: List[FridaScript]
    ssl_pinning_detected: bool
    root_detection_detected: bool
    crypto_methods: List[Dict[str, Any]]
    interesting_hooks: List[Dict[str, Any]]
    suggested_test_cases: List[str]
    frida_spawn_command: str
    frida_attach_command: str


@dataclass
class NativeFunction:
    """A function found in a native library."""
    name: str
    address: str
    size: int
    is_jni: bool = False
    is_exported: bool = False
    is_suspicious: bool = False
    suspicious_reason: Optional[str] = None


@dataclass
class NativeLibraryInfo:
    """Analysis of a single native library (.so file)."""
    name: str
    architecture: str
    size: int
    is_stripped: bool
    has_debug_info: bool
    exported_functions: List[NativeFunction]
    jni_functions: List[str]
    imported_libraries: List[str]
    strings: List[str]  # Interesting strings found
    hardcoded_secrets: List[Dict[str, Any]]
    anti_debug_detected: bool
    anti_debug_techniques: List[str]
    crypto_functions: List[str]
    suspicious_patterns: List[Dict[str, Any]]


@dataclass  
class NativeAnalysisResult:
    """Complete native library analysis result."""
    total_libraries: int
    libraries: List[NativeLibraryInfo]
    total_jni_functions: int
    total_exported_functions: int
    architectures: List[str]
    security_findings: List[Dict[str, Any]]
    overall_native_risk: str  # low, medium, high, critical


@dataclass
class HardeningCategory:
    """A category in the hardening score."""
    name: str
    score: int  # 0-100
    max_score: int
    weight: float
    findings: List[Dict[str, Any]]
    recommendations: List[str]


@dataclass
class HardeningScore:
    """Overall APK hardening/security score."""
    overall_score: int  # 0-100
    grade: str  # A, B, C, D, F
    risk_level: str  # Low, Medium, High, Critical
    categories: List[HardeningCategory]
    attack_surface_score: int
    protection_score: int
    data_security_score: int
    summary: str
    top_risks: List[str]
    top_recommendations: List[str]


# ============================================================================
# JADX Decompilation Types
# ============================================================================

@dataclass
class JadxDecompiledClass:
    """A decompiled Java class from JADX."""
    class_name: str
    package_name: str
    file_path: str  # Relative path in decompiled output
    source_code: str
    line_count: int
    is_activity: bool = False
    is_service: bool = False
    is_receiver: bool = False
    is_provider: bool = False
    is_application: bool = False
    extends: Optional[str] = None
    implements: List[str] = field(default_factory=list)
    methods: List[str] = field(default_factory=list)
    fields: List[str] = field(default_factory=list)
    security_issues: List[Dict[str, Any]] = field(default_factory=list)


@dataclass
class JadxDecompilationResult:
    """Complete JADX decompilation result."""
    package_name: str
    total_classes: int
    total_files: int
    output_directory: str
    classes: List[JadxDecompiledClass]
    resources_dir: str
    manifest_path: str
    source_tree: Dict[str, Any]  # Directory structure
    decompilation_time: float
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


# ============================================================================
# Manifest Visualization Types
# ============================================================================

@dataclass
class ManifestNode:
    """A node in the manifest visualization graph."""
    id: str
    name: str
    node_type: str  # activity, service, receiver, provider, permission, feature
    label: str
    is_exported: bool = False
    is_main: bool = False
    is_dangerous: bool = False
    attributes: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ManifestEdge:
    """An edge in the manifest visualization graph."""
    source: str
    target: str
    edge_type: str  # uses_permission, intent_filter, data_scheme, category
    label: str


@dataclass
class ManifestVisualization:
    """Complete manifest visualization data."""
    package_name: str
    app_name: Optional[str]
    version_name: Optional[str]
    nodes: List[ManifestNode]
    edges: List[ManifestEdge]
    component_counts: Dict[str, int]
    permission_summary: Dict[str, int]  # dangerous, normal, signature, etc.
    exported_count: int
    main_activity: Optional[str]
    deep_link_schemes: List[str]
    mermaid_diagram: str  # Pre-rendered mermaid flowchart


# ============================================================================
# Attack Surface Map Types
# ============================================================================

@dataclass
class AttackVector:
    """A potential attack vector/entry point."""
    id: str
    name: str
    vector_type: str  # exported_activity, deep_link, content_provider, broadcast, etc.
    component: str
    severity: str  # low, medium, high, critical
    description: str
    exploitation_steps: List[str]
    required_permissions: List[str]
    adb_command: Optional[str] = None
    intent_example: Optional[str] = None
    mitigation: str = ""


@dataclass
class ExposedDataPath:
    """An exposed data path through content providers."""
    provider_name: str
    uri_pattern: str
    permissions_required: List[str]
    operations: List[str]  # read, write, delete
    is_exported: bool
    potential_data: str
    risk_level: str


@dataclass
class DeepLinkEntry:
    """A deep link entry point."""
    scheme: str
    host: str
    path: str
    full_url: str
    handling_activity: str
    parameters: List[str]
    is_verified: bool  # App Links verification
    security_notes: List[str]


@dataclass
class AttackSurfaceMap:
    """Complete attack surface analysis."""
    package_name: str
    total_attack_vectors: int
    attack_vectors: List[AttackVector]
    exposed_data_paths: List[ExposedDataPath]
    deep_links: List[DeepLinkEntry]
    ipc_endpoints: List[Dict[str, Any]]  # Inter-Process Communication endpoints
    overall_exposure_score: int  # 0-100
    risk_level: str  # low, medium, high, critical
    risk_breakdown: Dict[str, int]  # vectors by severity
    priority_targets: List[str]  # Top items to investigate
    automated_tests: List[Dict[str, Any]]  # adb commands to test each vector
    mermaid_attack_tree: str  # Visual attack tree diagram


# ============================================================================
# Secret Patterns
# ============================================================================

SECRET_PATTERNS = {
    "api_key": re.compile(r'(?:api[_-]?key|apikey)["\']?\s*[:=]\s*["\']?([a-zA-Z0-9_\-]{20,})["\']?', re.IGNORECASE),
    "aws_key": re.compile(r'(?:AKIA|ABIA|ACCA|ASIA)[A-Z0-9]{16}'),
    "aws_secret": re.compile(r'(?:aws[_-]?secret)["\']?\s*[:=]\s*["\']?([a-zA-Z0-9/+=]{40})["\']?', re.IGNORECASE),
    "password": re.compile(r'(?:password|passwd|pwd)["\']?\s*[:=]\s*["\']?([^\s"\']{6,})["\']?', re.IGNORECASE),
    "private_key": re.compile(r'-----BEGIN (?:RSA |EC |DSA |OPENSSH )?PRIVATE KEY-----'),
    "github_token": re.compile(r'gh[pousr]_[A-Za-z0-9_]{36,}'),
    "jwt": re.compile(r'eyJ[a-zA-Z0-9_-]*\.eyJ[a-zA-Z0-9_-]*\.[a-zA-Z0-9_-]*'),
    "connection_string": re.compile(r'(?:mongodb|mysql|postgres|redis|mssql)://[^\s"\']+', re.IGNORECASE),
    "bearer_token": re.compile(r'[Bb]earer\s+[a-zA-Z0-9_\-\.]+'),
    "base64_secret": re.compile(r'(?:secret|key|token|password)["\']?\s*[:=]\s*["\']?([A-Za-z0-9+/]{40,}={0,2})["\']?', re.IGNORECASE),
}

URL_PATTERN = re.compile(r'https?://[^\s<>"\']+')
EMAIL_PATTERN = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
IP_PATTERN = re.compile(r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b')
PATH_PATTERN = re.compile(r'(?:/[a-zA-Z0-9_\-\.]+){2,}|(?:[A-Z]:\\[a-zA-Z0-9_\-\.\\ ]+)')

# Dangerous Android permissions
DANGEROUS_PERMISSIONS = {
    "android.permission.READ_CONTACTS": "Access contacts",
    "android.permission.WRITE_CONTACTS": "Modify contacts",
    "android.permission.READ_CALL_LOG": "Access call logs",
    "android.permission.WRITE_CALL_LOG": "Modify call logs",
    "android.permission.READ_CALENDAR": "Access calendar",
    "android.permission.WRITE_CALENDAR": "Modify calendar",
    "android.permission.CAMERA": "Access camera",
    "android.permission.RECORD_AUDIO": "Record audio",
    "android.permission.READ_PHONE_STATE": "Access phone state",
    "android.permission.READ_PHONE_NUMBERS": "Access phone numbers",
    "android.permission.CALL_PHONE": "Make phone calls",
    "android.permission.READ_SMS": "Read SMS messages",
    "android.permission.SEND_SMS": "Send SMS messages",
    "android.permission.RECEIVE_SMS": "Receive SMS messages",
    "android.permission.READ_EXTERNAL_STORAGE": "Read external storage",
    "android.permission.WRITE_EXTERNAL_STORAGE": "Write external storage",
    "android.permission.ACCESS_FINE_LOCATION": "Access precise location",
    "android.permission.ACCESS_COARSE_LOCATION": "Access approximate location",
    "android.permission.ACCESS_BACKGROUND_LOCATION": "Access location in background",
    "android.permission.BODY_SENSORS": "Access body sensors",
    "android.permission.ACTIVITY_RECOGNITION": "Activity recognition",
    "android.permission.INTERNET": "Full network access",
    "android.permission.SYSTEM_ALERT_WINDOW": "Draw over other apps",
    "android.permission.REQUEST_INSTALL_PACKAGES": "Install packages",
    "android.permission.BIND_ACCESSIBILITY_SERVICE": "Accessibility service",
    "android.permission.BIND_DEVICE_ADMIN": "Device admin",
}

# Suspicious Windows API imports
SUSPICIOUS_IMPORTS = {
    "CreateRemoteThread": "Can inject code into other processes",
    "VirtualAllocEx": "Can allocate memory in other processes",
    "WriteProcessMemory": "Can write to other processes' memory",
    "ReadProcessMemory": "Can read other processes' memory",
    "NtUnmapViewOfSection": "Process hollowing technique",
    "SetWindowsHookEx": "Keylogger capability",
    "GetAsyncKeyState": "Keylogger capability",
    "InternetOpen": "Network communication",
    "URLDownloadToFile": "Can download files from internet",
    "WinExec": "Can execute commands",
    "ShellExecute": "Can execute programs",
    "CreateProcess": "Can spawn processes",
    "RegSetValue": "Can modify registry",
    "CryptEncrypt": "Encryption capability (ransomware indicator)",
    "CryptDecrypt": "Decryption capability",
    "IsDebuggerPresent": "Anti-debugging technique",
    "CheckRemoteDebuggerPresent": "Anti-debugging technique",
    "OutputDebugString": "Anti-debugging technique",
    "GetTickCount": "Anti-sandbox technique",
    "Sleep": "Anti-sandbox technique (long sleep)",
}

# Suspicious Linux/ELF function imports
SUSPICIOUS_ELF_FUNCTIONS = {
    # Process manipulation
    "ptrace": "Can debug/trace processes (anti-debugging or injection)",
    "fork": "Creates child processes",
    "execve": "Executes programs",
    "execl": "Executes programs",
    "execlp": "Executes programs",
    "execv": "Executes programs",
    "execvp": "Executes programs",
    "system": "Executes shell commands",
    "popen": "Opens pipe to shell command",
    "dlopen": "Dynamic library loading",
    "dlsym": "Dynamic symbol resolution",
    # Network
    "socket": "Network communication",
    "connect": "Network connection",
    "bind": "Network binding (server)",
    "listen": "Network listening (server)",
    "accept": "Accepts network connections",
    "send": "Sends network data",
    "recv": "Receives network data",
    "sendto": "Sends UDP data",
    "recvfrom": "Receives UDP data",
    "gethostbyname": "DNS resolution",
    "getaddrinfo": "Address resolution",
    # File operations
    "unlink": "Deletes files",
    "rmdir": "Removes directories",
    "chmod": "Changes file permissions",
    "chown": "Changes file ownership",
    "mmap": "Memory mapping (code injection)",
    "mprotect": "Memory protection change (code injection)",
    # Privilege escalation
    "setuid": "Changes user ID",
    "setgid": "Changes group ID",
    "seteuid": "Changes effective user ID",
    "setegid": "Changes effective group ID",
    # Crypto (potential ransomware)
    "EVP_EncryptInit": "OpenSSL encryption",
    "EVP_DecryptInit": "OpenSSL decryption",
    "AES_encrypt": "AES encryption",
    "AES_decrypt": "AES decryption",
    "RSA_public_encrypt": "RSA encryption",
    # Anti-debugging/evasion
    "prctl": "Process control (can hide from ps)",
    "getenv": "Environment variable access",
    "uname": "System information gathering",
    "geteuid": "Check if running as root",
    "getpid": "Get process ID",
    "kill": "Send signals to processes",
}

# Suspicious x86/x64 instruction patterns
SUSPICIOUS_INSTRUCTIONS = {
    "int 0x80": "Linux syscall (x86)",
    "syscall": "Linux syscall (x64)",
    "sysenter": "Fast syscall entry",
    "int 0x2e": "Windows syscall (legacy)",
    "int3": "Debugger breakpoint",
    "cpuid": "CPU identification (VM detection)",
    "rdtsc": "Timestamp counter (timing attacks/anti-debug)",
    "in al": "I/O port access (rootkit behavior)",
    "out": "I/O port write (rootkit behavior)",
}


# ============================================================================
# Binary Analysis Functions
# ============================================================================

def extract_strings(data: bytes, min_length: int = 4, max_strings: int = 5000) -> List[ExtractedString]:
    """Extract ASCII and UTF-16 strings from binary data."""
    strings = []
    
    # Extract ASCII strings
    ascii_pattern = re.compile(rb'[\x20-\x7e]{' + str(min_length).encode() + rb',}')
    for match in ascii_pattern.finditer(data):
        if len(strings) >= max_strings:
            break
        try:
            value = match.group().decode('ascii')
            strings.append(ExtractedString(
                value=value,
                offset=match.start(),
                encoding="ascii",
                category=categorize_string(value),
            ))
        except:
            pass
    
    # Extract UTF-16 LE strings
    utf16_pattern = re.compile(rb'(?:[\x20-\x7e]\x00){' + str(min_length).encode() + rb',}')
    for match in utf16_pattern.finditer(data):
        if len(strings) >= max_strings:
            break
        try:
            value = match.group().decode('utf-16-le')
            strings.append(ExtractedString(
                value=value,
                offset=match.start(),
                encoding="utf16",
                category=categorize_string(value),
            ))
        except:
            pass
    
    # Deduplicate by value
    seen = set()
    unique_strings = []
    for s in strings:
        if s.value not in seen:
            seen.add(s.value)
            unique_strings.append(s)
    
    return unique_strings


def categorize_string(value: str) -> Optional[str]:
    """Categorize a string based on its content."""
    if URL_PATTERN.search(value):
        return "url"
    if EMAIL_PATTERN.search(value):
        return "email"
    if IP_PATTERN.search(value):
        return "ip_address"
    if PATH_PATTERN.search(value):
        return "path"
    for secret_type, pattern in SECRET_PATTERNS.items():
        if pattern.search(value):
            return secret_type
    return None


def detect_secrets_in_strings(strings: List[ExtractedString]) -> List[Dict[str, Any]]:
    """Find potential secrets in extracted strings."""
    secrets = []
    seen = set()
    
    for s in strings:
        for secret_type, pattern in SECRET_PATTERNS.items():
            matches = pattern.finditer(s.value)
            for match in matches:
                value = match.group(1) if match.lastindex else match.group(0)
                
                # Skip common false positives
                if value.lower() in {'password', 'secret', 'token', 'api_key', 'example', 'test'}:
                    continue
                if len(value) < 8:
                    continue
                
                dedup_key = f"{secret_type}:{value}"
                if dedup_key in seen:
                    continue
                seen.add(dedup_key)
                
                # Mask the value
                if len(value) > 8:
                    masked = value[:4] + '*' * (len(value) - 8) + value[-4:]
                else:
                    masked = value[:2] + '*' * (len(value) - 2)
                
                secrets.append({
                    "type": secret_type,
                    "value": value,
                    "masked_value": masked,
                    "offset": s.offset,
                    "context": s.value[:200],
                    "severity": get_secret_severity(secret_type),
                })
    
    return secrets


def get_secret_severity(secret_type: str) -> str:
    """Get severity level for a secret type."""
    critical = {"private_key", "aws_secret", "password", "connection_string"}
    high = {"api_key", "aws_key", "github_token", "jwt", "bearer_token", "base64_secret"}
    return "critical" if secret_type in critical else "high" if secret_type in high else "medium"


def parse_pe_header(data: bytes) -> Optional[BinaryMetadata]:
    """Parse PE (Windows executable) header."""
    try:
        # Check MZ signature
        if data[:2] != b'MZ':
            return None
        
        # Get PE header offset
        pe_offset = struct.unpack('<I', data[0x3C:0x40])[0]
        
        # Check PE signature
        if data[pe_offset:pe_offset+4] != b'PE\x00\x00':
            return None
        
        # Parse COFF header
        coff_offset = pe_offset + 4
        machine = struct.unpack('<H', data[coff_offset:coff_offset+2])[0]
        num_sections = struct.unpack('<H', data[coff_offset+2:coff_offset+4])[0]
        timestamp = struct.unpack('<I', data[coff_offset+4:coff_offset+8])[0]
        
        # Determine architecture
        arch_map = {0x14c: "x86", 0x8664: "x64", 0x1c0: "ARM", 0xaa64: "ARM64"}
        architecture = arch_map.get(machine, f"unknown (0x{machine:x})")
        
        # Parse optional header
        optional_offset = coff_offset + 20
        magic = struct.unpack('<H', data[optional_offset:optional_offset+2])[0]
        
        if magic == 0x10b:  # PE32
            entry_point = struct.unpack('<I', data[optional_offset+16:optional_offset+20])[0]
        elif magic == 0x20b:  # PE32+
            entry_point = struct.unpack('<I', data[optional_offset+16:optional_offset+20])[0]
        else:
            entry_point = None
        
        # Parse sections
        section_offset = optional_offset + (112 if magic == 0x10b else 128) + 16 * 16
        sections = []
        for i in range(min(num_sections, 20)):  # Limit to 20 sections
            sec_data = data[section_offset + i*40:section_offset + (i+1)*40]
            if len(sec_data) < 40:
                break
            name = sec_data[:8].rstrip(b'\x00').decode('ascii', errors='ignore')
            virtual_size = struct.unpack('<I', sec_data[8:12])[0]
            raw_size = struct.unpack('<I', sec_data[16:20])[0]
            characteristics = struct.unpack('<I', sec_data[36:40])[0]
            
            sections.append({
                "name": name,
                "virtual_size": virtual_size,
                "raw_size": raw_size,
                "characteristics": f"0x{characteristics:08x}",
            })
        
        # Check for packing indicators
        is_packed = False
        packer_name = None
        
        section_names = [s["name"].lower() for s in sections]
        if "upx0" in section_names or "upx1" in section_names:
            is_packed = True
            packer_name = "UPX"
        elif ".aspack" in section_names:
            is_packed = True
            packer_name = "ASPack"
        elif ".themida" in section_names:
            is_packed = True
            packer_name = "Themida"
        elif any(s["name"] == "" or s["virtual_size"] > s["raw_size"] * 10 for s in sections if s["raw_size"] > 0):
            is_packed = True
            packer_name = "Unknown (high entropy or unusual sections)"
        
        return BinaryMetadata(
            file_type="PE (Windows Executable)",
            architecture=architecture,
            file_size=len(data),
            entry_point=entry_point,
            is_packed=is_packed,
            packer_name=packer_name,
            compile_time=str(timestamp),
            sections=sections,
            headers={"pe_offset": pe_offset, "machine": f"0x{machine:x}", "magic": f"0x{magic:x}"},
        )
    except Exception as e:
        logger.error(f"PE parsing error: {e}")
        return None


# Rich header product ID to name mapping
RICH_PRODUCT_IDS = {
    0: "Unknown",
    1: "Import0 (VS 6.0)",
    2: "Linker510",
    3: "Cvtomf510",
    4: "Linker600",
    5: "Cvtomf600",
    6: "Cvtres500",
    7: "Utc11_Basic",
    8: "Utc11_C",
    9: "Utc12_Basic",
    10: "Utc12_C",
    11: "Utc12_CPP",
    12: "AliasObj60",
    13: "VisualBasic60",
    14: "Masm613",
    15: "Masm710",
    16: "Linker511",
    17: "Cvtomf511",
    18: "Masm614",
    19: "Linker512",
    20: "Cvtomf512",
    21: "Utc12_C_Std",
    22: "Utc12_CPP_Std",
    23: "Utc12_C_Book",
    24: "Utc12_CPP_Book",
    25: "Implib700",
    26: "Cvtomf700",
    27: "Utc13_Basic",
    28: "Utc13_C",
    29: "Utc13_CPP",
    30: "Linker610",
    31: "Cvtomf610",
    32: "Linker601",
    33: "Cvtomf601",
    34: "Utc12_1_Basic",
    35: "Utc12_1_C",
    36: "Utc12_1_CPP",
    37: "Linker620",
    38: "Cvtomf620",
    39: "AliasObj70",
    40: "Linker621",
    41: "Cvtomf621",
    42: "Masm615",
    43: "Utc13_LTCG_C",
    44: "Utc13_LTCG_CPP",
    45: "Masm620",
    46: "ILAsm100",
    47: "Utc12_2_Basic",
    48: "Utc12_2_C",
    49: "Utc12_2_CPP",
    50: "Utc12_2_C_Std",
    51: "Utc12_2_CPP_Std",
    52: "Utc12_2_C_Book",
    53: "Utc12_2_CPP_Book",
    54: "Implib622",
    55: "Cvtomf622",
    56: "Cvtres501",
    57: "Utc13_C_Std",
    58: "Utc13_CPP_Std",
    59: "Cvtpgd1300",
    60: "Linker622",
    61: "Linker700",
    62: "Export622",
    63: "Export700",
    64: "Masm700",
    65: "Utc13_POGO_I_C",
    66: "Utc13_POGO_I_CPP",
    67: "Utc13_POGO_O_C",
    68: "Utc13_POGO_O_CPP",
    69: "Cvtres700",
    70: "Cvtres710p",
    71: "Linker710p",
    72: "Cvtomf710p",
    73: "Export710p",
    74: "Implib710p",
    75: "Masm710p",
    76: "Utc13_POGO_I_C",
    77: "Utc13_POGO_I_CPP",
    78: "Linker624",
    79: "Cvtomf624",
    80: "Export624",
    81: "Implib624",
    82: "Linker710",
    83: "Cvtomf710",
    84: "Export710",
    85: "Implib710",
    86: "Cvtres710",
    87: "Utc14_C",
    88: "Utc14_CPP",
    89: "Utc14_C_Std",
    90: "Utc14_CPP_Std",
    91: "Utc14_LTCG_C",
    92: "Utc14_LTCG_CPP",
    93: "Utc14_POGO_I_C",
    94: "Utc14_POGO_I_CPP",
    95: "Utc14_POGO_O_C",
    96: "Utc14_POGO_O_CPP",
    # VS 2005+
    104: "Linker800",
    105: "Cvtomf800",
    106: "Export800",
    107: "Implib800",
    108: "Cvtres800",
    109: "Masm800",
    # VS 2008
    128: "Utc15_C (VS2008)",
    129: "Utc15_CPP (VS2008)",
    # VS 2010
    147: "Linker900 (VS2010)",
    148: "Cvtres900 (VS2010)",
    157: "Utc16_C (VS2010)",
    158: "Utc16_CPP (VS2010)",
    # VS 2012
    170: "Linker1000 (VS2012)",
    175: "Utc17_C (VS2012)",
    176: "Utc17_CPP (VS2012)",
    # VS 2013
    183: "Linker1100 (VS2013)",
    190: "Utc18_C (VS2013)",
    191: "Utc18_CPP (VS2013)",
    # VS 2015
    199: "Linker1200 (VS2015)",
    210: "Utc19_C (VS2015)",
    211: "Utc19_CPP (VS2015)",
    # VS 2017+
    220: "Linker1400 (VS2017)",
    257: "Utc1900_C (VS2017)",
    258: "Utc1900_CPP (VS2017)",
    259: "Utc1911_C (VS2017 15.3)",
    260: "Utc1911_CPP (VS2017 15.3)",
    261: "Utc1912_C (VS2017 15.5)",
    262: "Utc1912_CPP (VS2017 15.5)",
}


def parse_rich_header(data: bytes) -> Optional[RichHeader]:
    """Parse the PE Rich header to extract compiler/linker information."""
    import hashlib
    
    try:
        # Find "Rich" marker
        rich_offset = data.find(b'Rich')
        if rich_offset == -1:
            return None
        
        # The XOR key follows "Rich"
        xor_key = struct.unpack('<I', data[rich_offset + 4:rich_offset + 8])[0]
        
        # Find "DanS" marker (start of Rich header, after XOR decryption)
        # The Rich header starts at the DOS stub end and is XOR encrypted
        # Look for the encrypted "DanS" signature
        dans_encrypted = struct.unpack('<I', b'DanS')[0] ^ xor_key
        
        # Search backwards from "Rich" to find the start
        start_offset = None
        for i in range(rich_offset - 4, 0x40, -4):  # Don't go before DOS header
            if struct.unpack('<I', data[i:i+4])[0] == dans_encrypted:
                start_offset = i
                break
        
        if start_offset is None:
            return None
        
        # Extract and decrypt the Rich header
        rich_data_encrypted = data[start_offset:rich_offset + 8]
        rich_data_decrypted = bytearray()
        
        for i in range(0, len(rich_data_encrypted) - 8, 4):  # Exclude "Rich" + key
            dword = struct.unpack('<I', rich_data_encrypted[i:i+4])[0]
            decrypted = dword ^ xor_key
            rich_data_decrypted.extend(struct.pack('<I', decrypted))
        
        # Parse entries (skip "DanS" + 3 padding DWORDs)
        entries = []
        for i in range(16, len(rich_data_decrypted), 8):  # Start after header (4 DWORDs = 16 bytes)
            if i + 8 > len(rich_data_decrypted):
                break
            
            compid = struct.unpack('<I', rich_data_decrypted[i:i+4])[0]
            count = struct.unpack('<I', rich_data_decrypted[i+4:i+8])[0]
            
            if compid == 0 and count == 0:
                continue
            
            # Extract product ID (high 16 bits) and build ID (low 16 bits)
            product_id = (compid >> 16) & 0xFFFF
            build_id = compid & 0xFFFF
            
            product_name = RICH_PRODUCT_IDS.get(product_id, f"Unknown ({product_id})")
            
            entries.append(RichHeaderEntry(
                product_id=product_id,
                build_id=build_id,
                count=count,
                product_name=product_name,
                vs_version=None,
            ))
        
        if not entries:
            return None
        
        # Calculate Rich hash (MD5 of clear text for malware identification)
        rich_hash = hashlib.md5(bytes(rich_data_decrypted)).hexdigest()
        
        return RichHeader(
            entries=entries,
            rich_hash=rich_hash,
            checksum=xor_key,
            raw_data=rich_data_encrypted.hex()[:200],  # Limit size
            clear_data=bytes(rich_data_decrypted).hex()[:200],  # Limit size
        )
        
    except Exception as e:
        logger.warning(f"Rich header parsing error: {e}")
        return None


def calculate_imphash(pe) -> Optional[str]:
    """Calculate the import hash (imphash) for PE malware identification."""
    try:
        return pe.get_imphash()
    except Exception as e:
        logger.warning(f"Imphash calculation error: {e}")
        return None


def parse_pe_with_pefile(file_path: Path) -> tuple[Optional[BinaryMetadata], List[ImportedFunction], List[str]]:
    """Parse PE file using pefile library for comprehensive analysis."""
    if not PEFILE_AVAILABLE:
        return None, [], []
    
    try:
        pe = pefile.PE(str(file_path))
        
        # Architecture
        arch_map = {
            pefile.MACHINE_TYPE['IMAGE_FILE_MACHINE_I386']: "x86",
            pefile.MACHINE_TYPE['IMAGE_FILE_MACHINE_AMD64']: "x64",
            pefile.MACHINE_TYPE['IMAGE_FILE_MACHINE_ARM']: "ARM",
            pefile.MACHINE_TYPE['IMAGE_FILE_MACHINE_ARM64']: "ARM64",
        }
        architecture = arch_map.get(pe.FILE_HEADER.Machine, f"unknown (0x{pe.FILE_HEADER.Machine:x})")
        
        # Entry point
        entry_point = pe.OPTIONAL_HEADER.AddressOfEntryPoint
        
        # Compile time
        import datetime
        compile_time = datetime.datetime.utcfromtimestamp(pe.FILE_HEADER.TimeDateStamp).isoformat()
        
        # Sections with entropy calculation
        sections = []
        high_entropy_sections = 0
        for section in pe.sections:
            name = section.Name.decode('utf-8', errors='ignore').rstrip('\x00')
            entropy = section.get_entropy()
            sections.append({
                "name": name,
                "virtual_address": section.VirtualAddress,
                "virtual_size": section.Misc_VirtualSize,
                "raw_size": section.SizeOfRawData,
                "entropy": round(entropy, 2),
                "characteristics": f"0x{section.Characteristics:08x}",
            })
            if entropy > 7.0:
                high_entropy_sections += 1
        
        # Check for packing
        is_packed = False
        packer_name = None
        section_names_lower = [s["name"].lower() for s in sections]
        
        if "upx0" in section_names_lower or "upx1" in section_names_lower:
            is_packed = True
            packer_name = "UPX"
        elif ".aspack" in section_names_lower:
            is_packed = True
            packer_name = "ASPack"
        elif ".themida" in section_names_lower:
            is_packed = True
            packer_name = "Themida"
        elif ".vmp" in section_names_lower:
            is_packed = True
            packer_name = "VMProtect"
        elif ".nsp" in section_names_lower:
            is_packed = True
            packer_name = "NSPack"
        elif high_entropy_sections > len(sections) / 2:
            is_packed = True
            packer_name = f"Unknown (high entropy in {high_entropy_sections}/{len(sections)} sections)"
        
        # Parse imports
        imports = []
        if hasattr(pe, 'DIRECTORY_ENTRY_IMPORT'):
            for entry in pe.DIRECTORY_ENTRY_IMPORT:
                dll_name = entry.dll.decode('utf-8', errors='ignore')
                for imp in entry.imports:
                    if imp.name:
                        func_name = imp.name.decode('utf-8', errors='ignore')
                        is_suspicious = func_name in SUSPICIOUS_IMPORTS
                        imports.append(ImportedFunction(
                            name=func_name,
                            library=dll_name,
                            ordinal=imp.ordinal,
                            is_suspicious=is_suspicious,
                            reason=SUSPICIOUS_IMPORTS.get(func_name),
                        ))
        
        # Parse exports
        exports = []
        if hasattr(pe, 'DIRECTORY_ENTRY_EXPORT'):
            for exp in pe.DIRECTORY_ENTRY_EXPORT.symbols:
                if exp.name:
                    exports.append(exp.name.decode('utf-8', errors='ignore'))
        
        # Parse Rich header for compiler/linker information
        with open(file_path, 'rb') as f:
            raw_data = f.read(0x1000)  # Rich header is in first 4KB
        rich_header = parse_rich_header(raw_data)
        
        # Calculate imphash for malware identification
        imphash = calculate_imphash(pe)
        
        metadata = BinaryMetadata(
            file_type="PE (Windows Executable)",
            architecture=architecture,
            file_size=pe.DOS_HEADER.e_lfanew + pe.FILE_HEADER.sizeof() + pe.OPTIONAL_HEADER.sizeof(),
            entry_point=entry_point,
            is_packed=is_packed,
            packer_name=packer_name,
            compile_time=compile_time,
            sections=sections,
            headers={
                "machine": f"0x{pe.FILE_HEADER.Machine:x}",
                "characteristics": f"0x{pe.FILE_HEADER.Characteristics:x}",
                "subsystem": pe.OPTIONAL_HEADER.Subsystem,
                "dll_characteristics": f"0x{pe.OPTIONAL_HEADER.DllCharacteristics:x}",
            },
            rich_header=rich_header,
            imphash=imphash,
        )
        
        pe.close()
        return metadata, imports, exports
        
    except Exception as e:
        logger.error(f"pefile parsing error: {e}")
        return None, [], []


def parse_elf_header(data: bytes) -> Optional[BinaryMetadata]:
    """Parse ELF (Linux executable) header - basic fallback parser."""
    try:
        # Check ELF magic
        if data[:4] != b'\x7fELF':
            return None
        
        # Get ELF class (32 or 64 bit)
        ei_class = data[4]
        is_64bit = ei_class == 2
        
        # Get architecture
        if is_64bit:
            e_machine = struct.unpack('<H', data[18:20])[0]
            e_entry = struct.unpack('<Q', data[24:32])[0]
        else:
            e_machine = struct.unpack('<H', data[18:20])[0]
            e_entry = struct.unpack('<I', data[24:28])[0]
        
        arch_map = {
            0x03: "x86",
            0x3E: "x86_64",
            0x28: "ARM",
            0xB7: "ARM64",
            0x08: "MIPS",
            0x14: "PowerPC",
        }
        architecture = arch_map.get(e_machine, f"unknown (0x{e_machine:x})")
        
        # Get ELF type
        e_type = struct.unpack('<H', data[16:18])[0]
        type_map = {1: "Relocatable", 2: "Executable", 3: "Shared Object", 4: "Core"}
        file_type = f"ELF {type_map.get(e_type, 'Unknown')} ({64 if is_64bit else 32}-bit)"
        
        return BinaryMetadata(
            file_type=file_type,
            architecture=architecture,
            file_size=len(data),
            entry_point=e_entry,
            is_packed=False,
            sections=[],
            headers={"e_machine": f"0x{e_machine:x}", "e_entry": f"0x{e_entry:x}"},
        )
    except Exception as e:
        logger.error(f"ELF parsing error: {e}")
        return None


def parse_elf_with_pyelftools(file_path: Path) -> tuple[Optional[BinaryMetadata], List[ELFSymbol], List[ImportedFunction], List[str], Optional[Dict[str, Any]]]:
    """Parse ELF file using pyelftools for comprehensive analysis."""
    if not PYELFTOOLS_AVAILABLE:
        return None, [], [], [], None
    
    try:
        with open(file_path, 'rb') as f:
            elf = ELFFile(f)
            
            # Basic metadata
            is_64bit = elf.elfclass == 64
            arch_map = {
                'EM_386': 'x86',
                'EM_X86_64': 'x86_64',
                'EM_ARM': 'ARM',
                'EM_AARCH64': 'ARM64',
                'EM_MIPS': 'MIPS',
                'EM_PPC': 'PowerPC',
                'EM_PPC64': 'PowerPC64',
                'EM_RISCV': 'RISC-V',
            }
            architecture = arch_map.get(elf['e_machine'], elf['e_machine'])
            
            # ELF type
            type_map = {
                'ET_REL': 'Relocatable',
                'ET_EXEC': 'Executable',
                'ET_DYN': 'Shared Object/PIE',
                'ET_CORE': 'Core Dump',
            }
            elf_type = type_map.get(elf['e_type'], elf['e_type'])
            file_type = f"ELF {elf_type} ({64 if is_64bit else 32}-bit)"
            
            # Entry point
            entry_point = elf['e_entry']
            
            # Parse sections with entropy
            sections = []
            text_section_data = None
            text_section_addr = None
            
            for section in elf.iter_sections():
                section_name = section.name
                section_data = section.data() if hasattr(section, 'data') else b''
                entropy = calculate_entropy(section_data) if section_data else 0.0
                
                sections.append({
                    "name": section_name,
                    "type": section['sh_type'],
                    "address": section['sh_addr'],
                    "size": section['sh_size'],
                    "entropy": round(entropy, 2),
                    "flags": f"0x{section['sh_flags']:x}",
                })
                
                # Save .text section for disassembly
                if section_name == '.text':
                    text_section_data = section_data
                    text_section_addr = section['sh_addr']
            
            # Check for security features
            relro = "None"
            stack_canary = False
            nx_enabled = False
            pie_enabled = elf['e_type'] == 'ET_DYN'
            interpreter = None
            linked_libraries = []
            
            for segment in elf.iter_segments():
                seg_type = segment['p_type']
                
                # Check for interpreter (dynamic linker)
                if seg_type == 'PT_INTERP':
                    interpreter = segment.get_interp_name()
                
                # Check for GNU_RELRO
                if seg_type == 'PT_GNU_RELRO':
                    relro = "Partial"
                
                # Check for GNU_STACK (NX)
                if seg_type == 'PT_GNU_STACK':
                    # If execute flag is not set, NX is enabled
                    if not (segment['p_flags'] & 0x1):  # PF_X
                        nx_enabled = True
            
            # Parse dynamic section for BIND_NOW (Full RELRO) and libraries
            for section in elf.iter_sections():
                if isinstance(section, DynamicSection):
                    for tag in section.iter_tags():
                        if tag.entry.d_tag == 'DT_BIND_NOW':
                            relro = "Full"
                        elif tag.entry.d_tag == 'DT_FLAGS' and tag.entry.d_val & 0x8:  # DF_BIND_NOW
                            relro = "Full"
                        elif tag.entry.d_tag == 'DT_NEEDED':
                            linked_libraries.append(tag.needed)
            
            # Parse symbol tables
            symbols = []
            imports = []
            exports = []
            
            for section in elf.iter_sections():
                if isinstance(section, SymbolTableSection):
                    for symbol in section.iter_symbols():
                        sym_name = symbol.name
                        if not sym_name:
                            continue
                        
                        sym_type = symbol['st_info']['type']
                        sym_bind = symbol['st_info']['bind']
                        sym_value = symbol['st_value']
                        sym_size = symbol['st_size']
                        sym_shndx = symbol['st_shndx']
                        
                        # Determine if imported or exported
                        is_imported = sym_shndx == 'SHN_UNDEF' and sym_bind == 'STB_GLOBAL'
                        is_exported = sym_shndx != 'SHN_UNDEF' and sym_bind in ('STB_GLOBAL', 'STB_WEAK')
                        
                        # Check for suspicious functions
                        is_suspicious = sym_name in SUSPICIOUS_ELF_FUNCTIONS
                        reason = SUSPICIOUS_ELF_FUNCTIONS.get(sym_name)
                        
                        # Check stack canary
                        if sym_name == '__stack_chk_fail':
                            stack_canary = True
                        
                        elf_symbol = ELFSymbol(
                            name=sym_name,
                            address=sym_value,
                            size=sym_size,
                            symbol_type=sym_type,
                            binding=sym_bind,
                            section=str(sym_shndx),
                            is_imported=is_imported,
                            is_exported=is_exported,
                            is_suspicious=is_suspicious,
                            reason=reason,
                        )
                        symbols.append(elf_symbol)
                        
                        # Build imports/exports lists
                        if is_imported:
                            imports.append(ImportedFunction(
                                name=sym_name,
                                library="(dynamic)",
                                is_suspicious=is_suspicious,
                                reason=reason,
                            ))
                        if is_exported and sym_type == 'STT_FUNC':
                            exports.append(sym_name)
            
            # Try to parse DWARF info
            dwarf_info = None
            if elf.has_dwarf_info():
                try:
                    dwarf = elf.get_dwarf_info()
                    dwarf_info = {
                        "has_debug_info": True,
                        "compilation_units": [],
                        "source_files": [],
                    }
                    
                    for cu in dwarf.iter_CUs():
                        die = cu.get_top_DIE()
                        cu_info = {
                            "name": die.attributes.get('DW_AT_name', {}).value if 'DW_AT_name' in die.attributes else "unknown",
                            "producer": die.attributes.get('DW_AT_producer', {}).value if 'DW_AT_producer' in die.attributes else None,
                            "language": die.attributes.get('DW_AT_language', {}).value if 'DW_AT_language' in die.attributes else None,
                        }
                        # Decode bytes to string if needed
                        if isinstance(cu_info["name"], bytes):
                            cu_info["name"] = cu_info["name"].decode('utf-8', errors='ignore')
                        if isinstance(cu_info["producer"], bytes):
                            cu_info["producer"] = cu_info["producer"].decode('utf-8', errors='ignore')
                        dwarf_info["compilation_units"].append(cu_info)
                    
                    # Limit to first 10 CUs
                    dwarf_info["compilation_units"] = dwarf_info["compilation_units"][:10]
                except Exception as e:
                    logger.warning(f"Failed to parse DWARF info: {e}")
                    dwarf_info = {"has_debug_info": True, "error": str(e)}
            
            # Check for packing (high entropy in code sections)
            is_packed = False
            packer_name = None
            code_sections = [s for s in sections if s["name"] in ('.text', '.code', '.init', '.fini')]
            high_entropy_code = [s for s in code_sections if s["entropy"] > 7.0]
            if high_entropy_code:
                is_packed = True
                packer_name = "Unknown (high entropy in code sections)"
            
            # Check for UPX
            section_names = [s["name"] for s in sections]
            if 'UPX0' in section_names or 'UPX1' in section_names:
                is_packed = True
                packer_name = "UPX"
            
            metadata = BinaryMetadata(
                file_type=file_type,
                architecture=architecture,
                file_size=file_path.stat().st_size,
                entry_point=entry_point,
                is_packed=is_packed,
                packer_name=packer_name,
                sections=sections,
                headers={
                    "e_machine": elf['e_machine'],
                    "e_type": elf['e_type'],
                    "e_entry": f"0x{entry_point:x}",
                    "e_phnum": elf['e_phnum'],
                    "e_shnum": elf['e_shnum'],
                },
                interpreter=interpreter,
                linked_libraries=linked_libraries,
                relro=relro,
                stack_canary=stack_canary,
                nx_enabled=nx_enabled,
                pie_enabled=pie_enabled,
            )
            
            return metadata, symbols, imports, exports, dwarf_info
            
    except Exception as e:
        logger.error(f"pyelftools parsing error: {e}")
        return None, [], [], [], None


def calculate_entropy(data: bytes) -> float:
    """Calculate Shannon entropy of data."""
    if not data:
        return 0.0
    
    byte_counts = {}
    for byte in data:
        byte_counts[byte] = byte_counts.get(byte, 0) + 1
    
    entropy = 0.0
    length = len(data)
    for count in byte_counts.values():
        if count > 0:
            p = count / length
            entropy -= p * math.log2(p)
    
    return entropy


# ============================================================================
# Capstone Disassembly Functions
# ============================================================================

def get_capstone_instance(architecture: str) -> Optional[Any]:
    """Get a Capstone disassembler instance for the given architecture."""
    if not CAPSTONE_AVAILABLE:
        return None
    
    arch_mode_map = {
        'x86': (CS_ARCH_X86, CS_MODE_32),
        'x86_64': (CS_ARCH_X86, CS_MODE_64),
        'ARM': (CS_ARCH_ARM, CS_MODE_ARM),
        'ARM64': (CS_ARCH_ARM64, CS_MODE_ARM),
        'MIPS': (CS_ARCH_MIPS, CS_MODE_32),
    }
    
    if architecture not in arch_mode_map:
        return None
    
    arch, mode = arch_mode_map[architecture]
    try:
        md = Cs(arch, mode)
        md.detail = True
        return md
    except Exception as e:
        logger.error(f"Failed to create Capstone instance: {e}")
        return None


def disassemble_at_address(data: bytes, base_address: int, start_offset: int, 
                           architecture: str, max_instructions: int = 100) -> List[DisassemblyInstruction]:
    """Disassemble code at a given offset."""
    if not CAPSTONE_AVAILABLE:
        return []
    
    md = get_capstone_instance(architecture)
    if not md:
        return []
    
    instructions = []
    call_mnemonics = {'call', 'bl', 'blx', 'jal', 'jalr'}
    jump_mnemonics = {'jmp', 'je', 'jne', 'jz', 'jnz', 'ja', 'jb', 'jg', 'jl', 'jge', 'jle',
                      'b', 'beq', 'bne', 'bgt', 'blt', 'bge', 'ble'}
    
    try:
        code = data[start_offset:start_offset + 4096]  # Disassemble up to 4KB
        
        for insn in md.disasm(code, base_address + start_offset):
            if len(instructions) >= max_instructions:
                break
            
            mnemonic_lower = insn.mnemonic.lower()
            is_call = mnemonic_lower in call_mnemonics
            is_jump = mnemonic_lower in jump_mnemonics
            
            # Check for suspicious instructions
            is_suspicious = False
            comment = None
            
            full_insn = f"{insn.mnemonic} {insn.op_str}"
            for pattern, desc in SUSPICIOUS_INSTRUCTIONS.items():
                if pattern.lower() in full_insn.lower():
                    is_suspicious = True
                    comment = desc
                    break
            
            instructions.append(DisassemblyInstruction(
                address=insn.address,
                mnemonic=insn.mnemonic,
                op_str=insn.op_str,
                bytes_hex=insn.bytes.hex(),
                size=insn.size,
                is_call=is_call,
                is_jump=is_jump,
                is_suspicious=is_suspicious,
                comment=comment,
            ))
            
    except Exception as e:
        logger.error(f"Disassembly failed: {e}")
    
    return instructions


def disassemble_function(data: bytes, base_address: int, func_offset: int, func_size: int,
                        func_name: str, architecture: str, symbols: List[ELFSymbol]) -> DisassemblyFunction:
    """Disassemble a complete function."""
    instructions = disassemble_at_address(
        data, base_address, func_offset, architecture, 
        max_instructions=min(func_size // 2, 500)  # Rough estimate
    )
    
    # Build symbol address map for call resolution
    symbol_map = {s.address: s.name for s in symbols if s.symbol_type == 'STT_FUNC'}
    
    calls = []
    suspicious_patterns = []
    
    for insn in instructions:
        if insn.is_call:
            # Try to resolve call target
            try:
                # Extract address from operand
                if insn.op_str.startswith('0x'):
                    target_addr = int(insn.op_str, 16)
                    if target_addr in symbol_map:
                        calls.append(symbol_map[target_addr])
                    else:
                        calls.append(f"sub_{target_addr:x}")
            except:
                calls.append(insn.op_str)
        
        if insn.is_suspicious:
            suspicious_patterns.append(f"{insn.mnemonic} {insn.op_str}: {insn.comment}")
    
    return DisassemblyFunction(
        name=func_name,
        address=base_address + func_offset,
        size=func_size,
        instructions=instructions[:200],  # Limit for response size
        calls=list(set(calls)),
        suspicious_patterns=suspicious_patterns,
    )


def disassemble_binary(file_path: Path, metadata: BinaryMetadata, 
                       symbols: List[ELFSymbol]) -> Optional[DisassemblyResult]:
    """Perform disassembly analysis of a binary."""
    if not CAPSTONE_AVAILABLE:
        return None
    
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
        
        architecture = metadata.architecture
        if architecture not in ('x86', 'x86_64', 'ARM', 'ARM64'):
            logger.info(f"Disassembly not supported for architecture: {architecture}")
            return None
        
        # Find .text section
        text_section = None
        for section in metadata.sections:
            if section.get("name") == ".text":
                text_section = section
                break
        
        if not text_section:
            # Try to disassemble from entry point
            entry_point = metadata.entry_point or 0
            entry_disasm = disassemble_at_address(data, 0, entry_point, architecture, 50)
            return DisassemblyResult(
                entry_point_disasm=entry_disasm,
                functions=[],
                suspicious_instructions=[{"note": "Could not find .text section"}],
                architecture=architecture,
                mode="64-bit" if "64" in architecture else "32-bit",
            )
        
        text_addr = text_section.get("address", 0)
        text_size = text_section.get("size", 0)
        
        # Calculate file offset for .text section
        # For ELF, we need to find the section in the file
        text_offset = 0
        if PYELFTOOLS_AVAILABLE:
            with open(file_path, 'rb') as f:
                elf = ELFFile(f)
                for section in elf.iter_sections():
                    if section.name == '.text':
                        text_offset = section['sh_offset']
                        break
        
        # Disassemble entry point (first 50 instructions)
        entry_point = metadata.entry_point or text_addr
        entry_offset = entry_point - text_addr + text_offset if entry_point >= text_addr else text_offset
        entry_disasm = disassemble_at_address(data, text_addr, entry_offset, architecture, 50)
        
        # Disassemble interesting functions
        functions = []
        suspicious_instructions = []
        
        # Find functions with symbols
        func_symbols = [s for s in symbols if s.symbol_type == 'STT_FUNC' and s.size > 0]
        
        # Prioritize suspicious functions and main/init functions
        priority_funcs = []
        other_funcs = []
        
        for sym in func_symbols:
            if sym.is_suspicious or sym.name in ('main', '_start', '__libc_start_main', 'init', '_init'):
                priority_funcs.append(sym)
            else:
                other_funcs.append(sym)
        
        # Disassemble up to 10 priority functions and 5 other functions
        for sym in priority_funcs[:10] + other_funcs[:5]:
            try:
                func_offset = sym.address - text_addr + text_offset
                if 0 <= func_offset < len(data):
                    func_disasm = disassemble_function(
                        data, text_addr, func_offset, sym.size,
                        sym.name, architecture, symbols
                    )
                    functions.append(func_disasm)
                    
                    # Collect suspicious patterns
                    for pattern in func_disasm.suspicious_patterns:
                        suspicious_instructions.append({
                            "function": sym.name,
                            "pattern": pattern,
                        })
            except Exception as e:
                logger.warning(f"Failed to disassemble function {sym.name}: {e}")
        
        return DisassemblyResult(
            entry_point_disasm=entry_disasm,
            functions=functions,
            suspicious_instructions=suspicious_instructions,
            architecture=architecture,
            mode="64-bit" if "64" in architecture else "32-bit",
        )
        
    except Exception as e:
        logger.error(f"Binary disassembly failed: {e}")
        return None


def analyze_imports(data: bytes, metadata: BinaryMetadata) -> List[ImportedFunction]:
    """Extract and analyze imported functions."""
    imports = []
    
    # For PE files, we'd need to parse the import directory
    # For now, we'll extract function names from strings
    strings = extract_strings(data, min_length=4, max_strings=10000)
    
    for s in strings:
        value = s.value
        # Check against suspicious imports
        for import_name, reason in SUSPICIOUS_IMPORTS.items():
            if import_name.lower() in value.lower():
                imports.append(ImportedFunction(
                    name=import_name,
                    library="(detected in strings)",
                    is_suspicious=True,
                    reason=reason,
                ))
    
    # Deduplicate
    seen = set()
    unique_imports = []
    for imp in imports:
        if imp.name not in seen:
            seen.add(imp.name)
            unique_imports.append(imp)
    
    return unique_imports


def analyze_binary(file_path: Path) -> BinaryAnalysisResult:
    """Perform complete analysis of a binary file."""
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
        
        filename = file_path.name
        imports = []
        exports = []
        symbols = []
        dwarf_info = None
        disassembly = None
        
        # Try pefile first for PE files (better analysis)
        if PEFILE_AVAILABLE and data[:2] == b'MZ':
            metadata, imports, exports = parse_pe_with_pefile(file_path)
            if metadata:
                metadata.file_size = len(data)  # Update with actual size
        else:
            metadata = None
        
        # Try pyelftools for ELF files
        if not metadata and data[:4] == b'\x7fELF':
            if PYELFTOOLS_AVAILABLE:
                metadata, symbols, imports, exports, dwarf_info = parse_elf_with_pyelftools(file_path)
                
                # Perform disassembly if we have metadata
                if metadata and CAPSTONE_AVAILABLE:
                    disassembly = disassemble_binary(file_path, metadata, symbols)
            
            # Fall back to basic ELF parsing
            if not metadata:
                metadata = parse_elf_header(data)
        
        # Fall back to basic PE parsing
        if not metadata and data[:2] == b'MZ':
            metadata = parse_pe_header(data)
        
        if not metadata:
            # Unknown binary format
            metadata = BinaryMetadata(
                file_type="Unknown binary",
                architecture="unknown",
                file_size=len(data),
            )
        
        # Extract strings
        strings = extract_strings(data, min_length=6, max_strings=3000)
        
        # Detect secrets
        secrets = detect_secrets_in_strings(strings)
        
        # Analyze imports if not already done via pefile/pyelftools
        if not imports:
            imports = analyze_imports(data, metadata)
        
        # Identify suspicious indicators
        suspicious = []
        
        # Check for suspicious imports
        suspicious_imports = [imp for imp in imports if imp.is_suspicious]
        if suspicious_imports:
            suspicious.append({
                "category": "Suspicious API Calls",
                "severity": "high",
                "description": f"Found {len(suspicious_imports)} suspicious API calls",
                "details": [{"name": imp.name, "reason": imp.reason} for imp in suspicious_imports[:10]],
            })
        
        # Check if packed
        if metadata.is_packed:
            suspicious.append({
                "category": "Packed/Obfuscated",
                "severity": "medium",
                "description": f"Binary appears to be packed",
                "details": {"packer": metadata.packer_name},
            })
        
        # Check for anti-debugging
        anti_debug = any(s.value for s in strings if any(x in s.value for x in ["IsDebuggerPresent", "CheckRemoteDebuggerPresent", "OutputDebugString"]))
        if anti_debug:
            suspicious.append({
                "category": "Anti-Debugging",
                "severity": "medium",
                "description": "Binary contains anti-debugging techniques",
            })
        
        # Extract URLs
        urls = list(set(s.value for s in strings if s.category == "url"))
        if urls:
            suspicious.append({
                "category": "Network Indicators",
                "severity": "info",
                "description": f"Found {len(urls)} URLs",
                "details": urls[:20],
            })
        
        # Add ELF-specific security indicators
        if metadata.file_type.startswith("ELF"):
            security_features = []
            if metadata.nx_enabled:
                security_features.append("NX (No-Execute) enabled")
            else:
                suspicious.append({
                    "category": "Missing Security Feature",
                    "severity": "medium",
                    "description": "NX (No-Execute) not enabled - stack may be executable",
                })
            
            if metadata.pie_enabled:
                security_features.append("PIE (Position Independent Executable) enabled")
            else:
                suspicious.append({
                    "category": "Missing Security Feature",
                    "severity": "low",
                    "description": "PIE not enabled - ASLR less effective",
                })
            
            if metadata.stack_canary:
                security_features.append("Stack canary enabled")
            else:
                suspicious.append({
                    "category": "Missing Security Feature",
                    "severity": "medium",
                    "description": "Stack canary not detected - vulnerable to buffer overflows",
                })
            
            if metadata.relro == "Full":
                security_features.append("Full RELRO enabled")
            elif metadata.relro == "Partial":
                security_features.append("Partial RELRO enabled")
            else:
                suspicious.append({
                    "category": "Missing Security Feature",
                    "severity": "medium",
                    "description": "RELRO not enabled - GOT may be writable",
                })
            
            if security_features:
                suspicious.append({
                    "category": "Security Features",
                    "severity": "info",
                    "description": "Enabled security mitigations",
                    "details": security_features,
                })
        
        # Add disassembly findings
        if disassembly and disassembly.suspicious_instructions:
            suspicious.append({
                "category": "Suspicious Instructions",
                "severity": "high",
                "description": f"Found {len(disassembly.suspicious_instructions)} suspicious instruction patterns",
                "details": disassembly.suspicious_instructions[:10],
            })
        
        return BinaryAnalysisResult(
            filename=filename,
            metadata=metadata,
            strings=strings[:500],  # Limit for response size
            imports=imports,
            exports=exports,
            secrets=secrets,
            suspicious_indicators=suspicious,
            symbols=symbols[:500],  # Limit for response size
            disassembly=disassembly,
            dwarf_info=dwarf_info,
        )
        
    except Exception as e:
        logger.error(f"Binary analysis failed: {e}")
        return BinaryAnalysisResult(
            filename=file_path.name,
            metadata=BinaryMetadata(file_type="error", architecture="unknown", file_size=0),
            strings=[],
            imports=[],
            exports=[],
            secrets=[],
            suspicious_indicators=[],
            symbols=[],
            disassembly=None,
            dwarf_info=None,
            error=str(e),
        )


# ============================================================================
# APK Analysis Functions
# ============================================================================

# Suspicious method patterns for DEX analysis
SUSPICIOUS_DEX_PATTERNS = {
    "reflection": [
        "java.lang.reflect", "getDeclaredMethod", "getDeclaredField",
        "setAccessible", "invoke", "getMethod", "forName",
    ],
    "crypto": [
        "javax.crypto", "Cipher", "SecretKey", "AES", "DES", "RSA",
        "MessageDigest", "Mac", "KeyGenerator", "KeyPairGenerator",
    ],
    "native": [
        "System.loadLibrary", "System.load", "Runtime.exec",
        "ProcessBuilder", "nativeLibraryDir",
    ],
    "dynamic_loading": [
        "DexClassLoader", "PathClassLoader", "dalvik.system.DexFile",
        "InMemoryDexClassLoader", "loadClass",
    ],
    "obfuscation": [
        "ProGuard", "DexGuard", "Allatori", "StringFog",
    ],
    "anti_analysis": [
        "isDebuggerConnected", "Debug.isDebuggerConnected",
        "android.os.Debug", "Xposed", "frida", "substrate",
    ],
    "data_exfiltration": [
        "getDeviceId", "getSubscriberId", "getSimSerialNumber",
        "getLine1Number", "getAccounts", "READ_CONTACTS",
    ],
    "root_detection": [
        "su", "/system/xbin/su", "/system/bin/su", "Superuser",
        "RootBeer", "isRooted", "checkRoot",
    ],
}

# Known tracker/SDK signatures
KNOWN_TRACKERS = {
    "com.google.firebase": "Firebase Analytics",
    "com.google.android.gms.analytics": "Google Analytics",
    "com.facebook.": "Facebook SDK",
    "com.appsflyer": "AppsFlyer",
    "com.adjust.sdk": "Adjust",
    "io.branch": "Branch.io",
    "com.amplitude": "Amplitude",
    "com.mixpanel": "Mixpanel",
    "com.segment": "Segment",
    "com.crashlytics": "Crashlytics",
    "io.sentry": "Sentry",
    "com.newrelic": "New Relic",
    "com.appdynamics": "AppDynamics",
    "com.google.ads": "Google Ads",
    "com.mopub": "MoPub",
    "com.unity3d.ads": "Unity Ads",
    "com.chartboost": "Chartboost",
    "com.vungle": "Vungle",
    "com.ironsource": "ironSource",
    "com.applovin": "AppLovin",
}


# ============================================================================
# Dynamic Analysis - Frida Script Generation
# ============================================================================

def generate_frida_scripts(
    package_name: str,
    strings: List[ExtractedString],
    dex_analysis: Optional[Dict[str, Any]],
    permissions: List[ApkPermission],
    urls: List[str],
    smali_analysis: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Generate Frida scripts for dynamic analysis based on static APK analysis.
    
    Creates targeted scripts for:
    - SSL pinning bypass
    - Root detection bypass
    - Crypto method hooking
    - Authentication hooks
    - Method tracing
    - Custom hooks for suspicious methods
    """
    scripts = []
    crypto_methods = []
    interesting_hooks = []
    ssl_pinning_detected = False
    root_detection_detected = False
    emulator_detection_detected = False
    anti_tampering_detected = False
    debugger_detection_detected = False
    
    # Detect SSL Pinning patterns
    ssl_patterns = [
        'certificatePinner', 'CertificatePinner', 'X509TrustManager',
        'TrustManagerFactory', 'ssl_pinning', 'checkServerTrusted',
        'SSLContext', 'HostnameVerifier', 'checkClientTrusted',
        'OkHostnameVerifier', 'pinCertificate'
    ]
    ssl_classes_found = set()
    for s in strings:
        for pattern in ssl_patterns:
            if pattern.lower() in s.value.lower():
                ssl_pinning_detected = True
                ssl_classes_found.add(pattern)
    
    # Detect Root Detection patterns
    root_patterns = [
        'isRooted', 'checkRoot', 'RootBeer', 'detectRoot', 'rootCheck',
        '/system/app/Superuser', '/system/xbin/su', 'test-keys',
        'com.noshufou.android.su', 'com.thirdparty.superuser',
        'eu.chainfire.supersu', 'com.koushikdutta.superuser',
        'com.topjohnwu.magisk', 'isDeviceRooted', 'checkForRoot'
    ]
    root_classes_found = set()
    for s in strings:
        for pattern in root_patterns:
            if pattern.lower() in s.value.lower():
                root_detection_detected = True
                root_classes_found.add(pattern)
    
    # Detect Emulator Detection patterns
    emulator_patterns = [
        'isEmulator', 'checkEmulator', 'detectEmulator', 'EMULATOR',
        'goldfish', 'sdk_gphone', 'generic_x86', 'vbox86',
        'google_sdk', 'nox', 'bluestacks', 'genymotion',
        'Andy', 'Droid4X', 'ro.kernel.qemu', 'ro.hardware.virtual'
    ]
    emulator_classes_found = set()
    for s in strings:
        for pattern in emulator_patterns:
            if pattern.lower() in s.value.lower():
                emulator_detection_detected = True
                emulator_classes_found.add(pattern)
    
    # Detect Debugger Detection patterns
    debugger_patterns = [
        'isDebuggerConnected', 'Debug.isDebuggerConnected', 'waitForDebugger',
        'android.os.Debug', 'JDWP', 'TracerPid', 'ptrace',
        'checkDebugger', 'detectDebugger', 'isDebuggable'
    ]
    debugger_classes_found = set()
    for s in strings:
        for pattern in debugger_patterns:
            if pattern.lower() in s.value.lower():
                debugger_detection_detected = True
                debugger_classes_found.add(pattern)
    
    # Detect Anti-Tampering patterns
    tampering_patterns = [
        'signature', 'checkSignature', 'verifySignature', 'PackageInfo',
        'signatures', 'GET_SIGNATURES', 'hashCode', 'checkIntegrity',
        'SafetyNet', 'Attestation', 'integrity', 'tamper'
    ]
    tampering_classes_found = set()
    for s in strings:
        for pattern in tampering_patterns:
            if pattern.lower() in s.value.lower():
                anti_tampering_detected = True
                tampering_classes_found.add(pattern)
    
    # Detect crypto patterns
    crypto_patterns = [
        ('Cipher', 'javax.crypto.Cipher'),
        ('SecretKeySpec', 'javax.crypto.spec.SecretKeySpec'),
        ('AES', 'AES encryption'),
        ('RSA', 'RSA encryption'),
        ('MessageDigest', 'java.security.MessageDigest'),
        ('Mac', 'javax.crypto.Mac'),
        ('PBKDF2', 'Key derivation'),
        ('KeyGenerator', 'javax.crypto.KeyGenerator'),
        ('EncryptedSharedPreferences', 'AndroidX encrypted storage'),
    ]
    for pattern, desc in crypto_patterns:
        for s in strings:
            if pattern in s.value:
                crypto_methods.append({
                    "pattern": pattern,
                    "description": desc,
                    "context": s.value[:100]
                })
                break
    
    # Detect authentication patterns
    auth_patterns = [
        'login', 'authenticate', 'verifyPassword', 'checkPassword',
        'validateToken', 'refreshToken', 'biometric', 'fingerprint',
        'FingerprintManager', 'BiometricPrompt', 'KeyguardManager'
    ]
    auth_methods_found = set()
    for s in strings:
        for pattern in auth_patterns:
            if pattern.lower() in s.value.lower():
                auth_methods_found.add(pattern)
    
    # =========================================================================
    # Generate SSL Pinning Bypass Script
    # =========================================================================
    ssl_bypass_script = '''// Universal SSL Pinning Bypass for ''' + package_name + '''
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f ''' + package_name + ''' -l ssl_bypass.js --no-pause

Java.perform(function() {
    console.log("[*] SSL Pinning Bypass Script Loaded");
    
    // ===== OkHttp3 CertificatePinner Bypass =====
    try {
        var CertificatePinner = Java.use('okhttp3.CertificatePinner');
        CertificatePinner.check.overload('java.lang.String', 'java.util.List').implementation = function(hostname, peerCertificates) {
            console.log('[+] OkHttp3 CertificatePinner.check() bypassed for: ' + hostname);
            return;
        };
        CertificatePinner.check$okhttp.overload('java.lang.String', 'kotlin.jvm.functions.Function0').implementation = function(hostname, peerCertificates) {
            console.log('[+] OkHttp3 CertificatePinner.check$okhttp() bypassed for: ' + hostname);
            return;
        };
        console.log('[*] OkHttp3 CertificatePinner hooks installed');
    } catch (e) {
        console.log('[-] OkHttp3 CertificatePinner not found: ' + e);
    }
    
    // ===== TrustManagerImpl (Android 7+) =====
    try {
        var TrustManagerImpl = Java.use('com.android.org.conscrypt.TrustManagerImpl');
        TrustManagerImpl.verifyChain.implementation = function(untrustedChain, trustAnchorChain, host, clientAuth, ocspData, tlsSctData) {
            console.log('[+] TrustManagerImpl.verifyChain() bypassed for: ' + host);
            return untrustedChain;
        };
        console.log('[*] TrustManagerImpl hooks installed');
    } catch (e) {
        console.log('[-] TrustManagerImpl not found: ' + e);
    }
    
    // ===== X509TrustManager =====
    try {
        var X509TrustManager = Java.use('javax.net.ssl.X509TrustManager');
        var TrustManager = Java.registerClass({
            name: 'com.vragent.TrustManager',
            implements: [X509TrustManager],
            methods: {
                checkClientTrusted: function(chain, authType) { },
                checkServerTrusted: function(chain, authType) { },
                getAcceptedIssuers: function() { return []; }
            }
        });
        console.log('[*] Custom X509TrustManager registered');
    } catch (e) {
        console.log('[-] X509TrustManager registration failed: ' + e);
    }
    
    // ===== SSLContext =====
    try {
        var SSLContext = Java.use('javax.net.ssl.SSLContext');
        SSLContext.init.overload('[Ljavax.net.ssl.KeyManager;', '[Ljavax.net.ssl.TrustManager;', 'java.security.SecureRandom').implementation = function(km, tm, sr) {
            console.log('[+] SSLContext.init() - Installing permissive TrustManager');
            var TrustManagerImpl = Java.use('com.vragent.TrustManager');
            this.init(km, [TrustManagerImpl.$new()], sr);
        };
        console.log('[*] SSLContext hooks installed');
    } catch (e) {
        console.log('[-] SSLContext hook failed: ' + e);
    }
    
    // ===== HostnameVerifier =====
    try {
        var HostnameVerifier = Java.use('javax.net.ssl.HostnameVerifier');
        var AllowAllHostnameVerifier = Java.registerClass({
            name: 'com.vragent.AllowAllHostnameVerifier',
            implements: [HostnameVerifier],
            methods: {
                verify: function(hostname, session) {
                    console.log('[+] HostnameVerifier.verify() bypassed for: ' + hostname);
                    return true;
                }
            }
        });
        console.log('[*] HostnameVerifier bypass registered');
    } catch (e) {
        console.log('[-] HostnameVerifier registration failed: ' + e);
    }
    
    // ===== Network Security Config (Android 7+) =====
    try {
        var NetworkSecurityConfig = Java.use('android.security.net.config.NetworkSecurityConfig');
        NetworkSecurityConfig.isCleartextTrafficPermitted.overload().implementation = function() {
            console.log('[+] Cleartext traffic permitted');
            return true;
        };
        console.log('[*] NetworkSecurityConfig hooks installed');
    } catch (e) {
        console.log('[-] NetworkSecurityConfig not available');
    }
    
    console.log("[*] SSL Pinning Bypass Complete - Ready to intercept traffic");
});
'''
    scripts.append(FridaScript(
        name="SSL Pinning Bypass",
        category="ssl_bypass",
        description="Universal SSL pinning bypass for OkHttp3, TrustManager, and Network Security Config",
        script_code=ssl_bypass_script,
        target_classes=["okhttp3.CertificatePinner", "javax.net.ssl.X509TrustManager", "javax.net.ssl.SSLContext"],
        target_methods=["check", "checkServerTrusted", "verify", "init"],
        is_dangerous=True,
        usage_instructions=f"frida -U -f {package_name} -l ssl_bypass.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Root Detection Bypass Script
    # =========================================================================
    root_bypass_script = '''// Root Detection Bypass for ''' + package_name + '''
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f ''' + package_name + ''' -l root_bypass.js --no-pause

Java.perform(function() {
    console.log("[*] Root Detection Bypass Script Loaded");
    
    // ===== RootBeer Library Bypass =====
    try {
        var RootBeer = Java.use('com.scottyab.rootbeer.RootBeer');
        RootBeer.isRooted.implementation = function() {
            console.log('[+] RootBeer.isRooted() bypassed - returning false');
            return false;
        };
        RootBeer.isRootedWithoutBusyBoxCheck.implementation = function() {
            console.log('[+] RootBeer.isRootedWithoutBusyBoxCheck() bypassed');
            return false;
        };
        RootBeer.detectRootManagementApps.implementation = function() {
            console.log('[+] RootBeer.detectRootManagementApps() bypassed');
            return false;
        };
        RootBeer.detectPotentiallyDangerousApps.implementation = function() {
            console.log('[+] RootBeer.detectPotentiallyDangerousApps() bypassed');
            return false;
        };
        RootBeer.detectTestKeys.implementation = function() {
            console.log('[+] RootBeer.detectTestKeys() bypassed');
            return false;
        };
        RootBeer.checkForBusyBoxBinary.implementation = function() {
            console.log('[+] RootBeer.checkForBusyBoxBinary() bypassed');
            return false;
        };
        RootBeer.checkForSuBinary.implementation = function() {
            console.log('[+] RootBeer.checkForSuBinary() bypassed');
            return false;
        };
        RootBeer.checkSuExists.implementation = function() {
            console.log('[+] RootBeer.checkSuExists() bypassed');
            return false;
        };
        RootBeer.checkForRWPaths.implementation = function() {
            console.log('[+] RootBeer.checkForRWPaths() bypassed');
            return false;
        };
        RootBeer.checkForDangerousProps.implementation = function() {
            console.log('[+] RootBeer.checkForDangerousProps() bypassed');
            return false;
        };
        RootBeer.checkForRootNative.implementation = function() {
            console.log('[+] RootBeer.checkForRootNative() bypassed');
            return false;
        };
        RootBeer.detectRootCloakingApps.implementation = function() {
            console.log('[+] RootBeer.detectRootCloakingApps() bypassed');
            return false;
        };
        console.log('[*] RootBeer library hooks installed');
    } catch (e) {
        console.log('[-] RootBeer library not found');
    }
    
    // ===== File.exists() for common root paths =====
    try {
        var File = Java.use('java.io.File');
        var originalExists = File.exists;
        File.exists.implementation = function() {
            var path = this.getAbsolutePath();
            var rootPaths = [
                '/system/app/Superuser.apk',
                '/system/xbin/su',
                '/system/bin/su',
                '/sbin/su',
                '/data/local/xbin/su',
                '/data/local/bin/su',
                '/data/local/su',
                '/system/sd/xbin/su',
                '/system/bin/failsafe/su',
                '/su/bin/su',
                '/magisk',
                '/sbin/.magisk',
                '/data/adb/magisk'
            ];
            for (var i = 0; i < rootPaths.length; i++) {
                if (path.indexOf(rootPaths[i]) !== -1) {
                    console.log('[+] File.exists() bypassed for root path: ' + path);
                    return false;
                }
            }
            return originalExists.call(this);
        };
        console.log('[*] File.exists() hooks installed');
    } catch (e) {
        console.log('[-] File.exists() hook failed: ' + e);
    }
    
    // ===== Runtime.exec() for su commands =====
    try {
        var Runtime = Java.use('java.lang.Runtime');
        var originalExec = Runtime.exec.overload('java.lang.String');
        Runtime.exec.overload('java.lang.String').implementation = function(cmd) {
            if (cmd.indexOf('su') !== -1 || cmd.indexOf('which su') !== -1) {
                console.log('[+] Runtime.exec() blocked su command: ' + cmd);
                throw new Error('su not found');
            }
            return originalExec.call(this, cmd);
        };
        console.log('[*] Runtime.exec() hooks installed');
    } catch (e) {
        console.log('[-] Runtime.exec() hook failed: ' + e);
    }
    
    // ===== Build.TAGS =====
    try {
        var Build = Java.use('android.os.Build');
        var originalTags = Build.TAGS.value;
        Build.TAGS.value = 'release-keys';
        console.log('[*] Build.TAGS changed from "' + originalTags + '" to "release-keys"');
    } catch (e) {
        console.log('[-] Build.TAGS modification failed: ' + e);
    }
    
    // ===== System.getProperty =====
    try {
        var System = Java.use('java.lang.System');
        var originalGetProperty = System.getProperty.overload('java.lang.String');
        System.getProperty.overload('java.lang.String').implementation = function(key) {
            if (key === 'ro.build.tags') {
                console.log('[+] System.getProperty("ro.build.tags") bypassed');
                return 'release-keys';
            }
            if (key === 'ro.debuggable') {
                console.log('[+] System.getProperty("ro.debuggable") bypassed');
                return '0';
            }
            return originalGetProperty.call(this, key);
        };
        console.log('[*] System.getProperty() hooks installed');
    } catch (e) {
        console.log('[-] System.getProperty() hook failed: ' + e);
    }
    
    console.log("[*] Root Detection Bypass Complete");
});
'''
    scripts.append(FridaScript(
        name="Root Detection Bypass",
        category="root_bypass",
        description="Comprehensive root detection bypass for RootBeer, file checks, and build properties",
        script_code=root_bypass_script,
        target_classes=["com.scottyab.rootbeer.RootBeer", "java.io.File", "java.lang.Runtime"],
        target_methods=["isRooted", "exists", "exec"],
        is_dangerous=True,
        usage_instructions=f"frida -U -f {package_name} -l root_bypass.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Crypto Hooking Script
    # =========================================================================
    crypto_hook_script = '''// Cryptographic Operations Hook for ''' + package_name + '''
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f ''' + package_name + ''' -l crypto_hook.js --no-pause

Java.perform(function() {
    console.log("[*] Crypto Hook Script Loaded");
    
    // ===== javax.crypto.Cipher =====
    try {
        var Cipher = Java.use('javax.crypto.Cipher');
        
        Cipher.getInstance.overload('java.lang.String').implementation = function(transformation) {
            console.log('[CIPHER] getInstance("' + transformation + '")');
            return this.getInstance(transformation);
        };
        
        Cipher.init.overload('int', 'java.security.Key').implementation = function(mode, key) {
            var modeStr = mode === 1 ? 'ENCRYPT' : (mode === 2 ? 'DECRYPT' : mode);
            console.log('[CIPHER] init(' + modeStr + ', key)');
            console.log('  Algorithm: ' + key.getAlgorithm());
            console.log('  Key (hex): ' + bytesToHex(key.getEncoded()));
            return this.init(mode, key);
        };
        
        Cipher.doFinal.overload('[B').implementation = function(input) {
            console.log('[CIPHER] doFinal() - Input (' + input.length + ' bytes):');
            console.log('  Hex: ' + bytesToHex(input).substring(0, 100) + '...');
            try {
                console.log('  UTF8: ' + Java.use('java.lang.String').$new(input, 'UTF-8'));
            } catch(e) {}
            var result = this.doFinal(input);
            console.log('[CIPHER] doFinal() - Output (' + result.length + ' bytes):');
            console.log('  Hex: ' + bytesToHex(result).substring(0, 100) + '...');
            return result;
        };
        console.log('[*] Cipher hooks installed');
    } catch (e) {
        console.log('[-] Cipher hook failed: ' + e);
    }
    
    // ===== javax.crypto.spec.SecretKeySpec =====
    try {
        var SecretKeySpec = Java.use('javax.crypto.spec.SecretKeySpec');
        SecretKeySpec.$init.overload('[B', 'java.lang.String').implementation = function(key, algorithm) {
            console.log('[KEY] SecretKeySpec created');
            console.log('  Algorithm: ' + algorithm);
            console.log('  Key (hex): ' + bytesToHex(key));
            console.log('  Key length: ' + key.length + ' bytes');
            return this.$init(key, algorithm);
        };
        console.log('[*] SecretKeySpec hooks installed');
    } catch (e) {
        console.log('[-] SecretKeySpec hook failed: ' + e);
    }
    
    // ===== javax.crypto.spec.IvParameterSpec =====
    try {
        var IvParameterSpec = Java.use('javax.crypto.spec.IvParameterSpec');
        IvParameterSpec.$init.overload('[B').implementation = function(iv) {
            console.log('[IV] IvParameterSpec created');
            console.log('  IV (hex): ' + bytesToHex(iv));
            console.log('  IV length: ' + iv.length + ' bytes');
            return this.$init(iv);
        };
        console.log('[*] IvParameterSpec hooks installed');
    } catch (e) {
        console.log('[-] IvParameterSpec hook failed: ' + e);
    }
    
    // ===== java.security.MessageDigest =====
    try {
        var MessageDigest = Java.use('java.security.MessageDigest');
        
        MessageDigest.getInstance.overload('java.lang.String').implementation = function(algorithm) {
            console.log('[HASH] MessageDigest.getInstance("' + algorithm + '")');
            return this.getInstance(algorithm);
        };
        
        MessageDigest.digest.overload('[B').implementation = function(input) {
            console.log('[HASH] digest() - Input (' + input.length + ' bytes):');
            try {
                console.log('  UTF8: ' + Java.use('java.lang.String').$new(input, 'UTF-8'));
            } catch(e) {
                console.log('  Hex: ' + bytesToHex(input).substring(0, 64) + '...');
            }
            var result = this.digest(input);
            console.log('[HASH] digest() - Output: ' + bytesToHex(result));
            return result;
        };
        console.log('[*] MessageDigest hooks installed');
    } catch (e) {
        console.log('[-] MessageDigest hook failed: ' + e);
    }
    
    // ===== javax.crypto.Mac =====
    try {
        var Mac = Java.use('javax.crypto.Mac');
        
        Mac.getInstance.overload('java.lang.String').implementation = function(algorithm) {
            console.log('[MAC] Mac.getInstance("' + algorithm + '")');
            return this.getInstance(algorithm);
        };
        
        Mac.init.overload('java.security.Key').implementation = function(key) {
            console.log('[MAC] init() with key');
            console.log('  Algorithm: ' + key.getAlgorithm());
            console.log('  Key (hex): ' + bytesToHex(key.getEncoded()));
            return this.init(key);
        };
        
        Mac.doFinal.overload('[B').implementation = function(input) {
            console.log('[MAC] doFinal() - Input: ');
            try {
                console.log('  UTF8: ' + Java.use('java.lang.String').$new(input, 'UTF-8'));
            } catch(e) {
                console.log('  Hex: ' + bytesToHex(input).substring(0, 64) + '...');
            }
            var result = this.doFinal(input);
            console.log('[MAC] doFinal() - Output: ' + bytesToHex(result));
            return result;
        };
        console.log('[*] Mac hooks installed');
    } catch (e) {
        console.log('[-] Mac hook failed: ' + e);
    }
    
    // ===== PBKDF2 Key Derivation =====
    try {
        var SecretKeyFactory = Java.use('javax.crypto.SecretKeyFactory');
        SecretKeyFactory.generateSecret.implementation = function(keySpec) {
            console.log('[PBKDF] generateSecret() called');
            var result = this.generateSecret(keySpec);
            if (keySpec.$className.indexOf('PBEKeySpec') !== -1) {
                var PBEKeySpec = Java.use('javax.crypto.spec.PBEKeySpec');
                var spec = Java.cast(keySpec, PBEKeySpec);
                console.log('  Password length: ' + spec.getPassword().length);
                console.log('  Salt (hex): ' + bytesToHex(spec.getSalt()));
                console.log('  Iterations: ' + spec.getIterationCount());
                console.log('  Key length: ' + spec.getKeyLength());
            }
            console.log('  Derived key (hex): ' + bytesToHex(result.getEncoded()));
            return result;
        };
        console.log('[*] SecretKeyFactory hooks installed');
    } catch (e) {
        console.log('[-] SecretKeyFactory hook failed: ' + e);
    }
    
    // Helper function
    function bytesToHex(bytes) {
        if (!bytes) return 'null';
        var hex = '';
        for (var i = 0; i < bytes.length; i++) {
            var b = (bytes[i] & 0xFF).toString(16);
            hex += (b.length === 1 ? '0' : '') + b;
        }
        return hex;
    }
    
    console.log("[*] Crypto Hooks Ready - All cryptographic operations will be logged");
});
'''
    scripts.append(FridaScript(
        name="Crypto Operations Hook",
        category="crypto_hook",
        description="Log all cryptographic operations including keys, IVs, and encrypted data",
        script_code=crypto_hook_script,
        target_classes=["javax.crypto.Cipher", "javax.crypto.spec.SecretKeySpec", "java.security.MessageDigest", "javax.crypto.Mac"],
        target_methods=["getInstance", "init", "doFinal", "digest", "generateSecret"],
        is_dangerous=False,
        usage_instructions=f"frida -U -f {package_name} -l crypto_hook.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Authentication Hook Script
    # =========================================================================
    auth_hook_script = '''// Authentication & Login Hook for ''' + package_name + '''
// Generated by VRAgent APK Analyzer  
// Usage: frida -U -f ''' + package_name + ''' -l auth_hook.js --no-pause

Java.perform(function() {
    console.log("[*] Authentication Hook Script Loaded");
    
    // ===== SharedPreferences (Token Storage) =====
    try {
        var SharedPreferences = Java.use('android.content.SharedPreferences');
        var Editor = Java.use('android.content.SharedPreferences$Editor');
        
        Editor.putString.implementation = function(key, value) {
            if (key.toLowerCase().indexOf('token') !== -1 || 
                key.toLowerCase().indexOf('session') !== -1 ||
                key.toLowerCase().indexOf('auth') !== -1 ||
                key.toLowerCase().indexOf('jwt') !== -1 ||
                key.toLowerCase().indexOf('cookie') !== -1) {
                console.log('[AUTH] SharedPreferences.putString()');
                console.log('  Key: ' + key);
                console.log('  Value: ' + value);
            }
            return this.putString(key, value);
        };
        
        console.log('[*] SharedPreferences hooks installed');
    } catch (e) {
        console.log('[-] SharedPreferences hook failed: ' + e);
    }
    
    // ===== HTTP Headers (Auth tokens) =====
    try {
        var OkHttpClient = Java.use('okhttp3.OkHttpClient');
        var Request = Java.use('okhttp3.Request');
        var RequestBuilder = Java.use('okhttp3.Request$Builder');
        
        RequestBuilder.header.implementation = function(name, value) {
            var nameLower = name.toLowerCase();
            if (nameLower === 'authorization' || 
                nameLower === 'x-auth-token' ||
                nameLower === 'x-api-key' ||
                nameLower === 'cookie' ||
                nameLower.indexOf('bearer') !== -1) {
                console.log('[AUTH] HTTP Header Set');
                console.log('  Header: ' + name);
                console.log('  Value: ' + value);
            }
            return this.header(name, value);
        };
        
        RequestBuilder.addHeader.implementation = function(name, value) {
            var nameLower = name.toLowerCase();
            if (nameLower === 'authorization' || 
                nameLower === 'x-auth-token' ||
                nameLower === 'x-api-key' ||
                nameLower === 'cookie') {
                console.log('[AUTH] HTTP Header Added');
                console.log('  Header: ' + name);
                console.log('  Value: ' + value);
            }
            return this.addHeader(name, value);
        };
        console.log('[*] OkHttp header hooks installed');
    } catch (e) {
        console.log('[-] OkHttp hooks failed: ' + e);
    }
    
    // ===== Biometric Authentication =====
    try {
        var BiometricPrompt = Java.use('androidx.biometric.BiometricPrompt');
        BiometricPrompt.authenticate.overload('androidx.biometric.BiometricPrompt$PromptInfo').implementation = function(promptInfo) {
            console.log('[BIOMETRIC] BiometricPrompt.authenticate() called');
            console.log('  Title: ' + promptInfo.getTitle());
            console.log('  Description: ' + promptInfo.getDescription());
            return this.authenticate(promptInfo);
        };
        console.log('[*] BiometricPrompt hooks installed');
    } catch (e) {
        console.log('[-] BiometricPrompt not found');
    }
    
    // ===== KeyStore (Certificate/Key Access) =====
    try {
        var KeyStore = Java.use('java.security.KeyStore');
        
        KeyStore.getKey.implementation = function(alias, password) {
            console.log('[KEYSTORE] getKey() called');
            console.log('  Alias: ' + alias);
            if (password) {
                console.log('  Password length: ' + password.length);
            }
            return this.getKey(alias, password);
        };
        
        KeyStore.getCertificate.implementation = function(alias) {
            console.log('[KEYSTORE] getCertificate("' + alias + '")');
            return this.getCertificate(alias);
        };
        console.log('[*] KeyStore hooks installed');
    } catch (e) {
        console.log('[-] KeyStore hooks failed: ' + e);
    }
    
    // ===== JWT Token Parsing =====
    try {
        // Hook Base64 decode for JWT detection
        var Base64 = Java.use('android.util.Base64');
        var originalDecode = Base64.decode.overload('[B', 'int');
        Base64.decode.overload('[B', 'int').implementation = function(input, flags) {
            var result = originalDecode.call(this, input, flags);
            try {
                var inputStr = Java.use('java.lang.String').$new(input, 'UTF-8');
                if (inputStr.indexOf('eyJ') === 0) {
                    console.log('[JWT] Potential JWT detected');
                    console.log('  Encoded: ' + inputStr.substring(0, 50) + '...');
                    var decoded = Java.use('java.lang.String').$new(result, 'UTF-8');
                    console.log('  Decoded: ' + decoded);
                }
            } catch(e) {}
            return result;
        };
        console.log('[*] Base64/JWT hooks installed');
    } catch (e) {
        console.log('[-] Base64 hooks failed: ' + e);
    }
    
    console.log("[*] Authentication Hooks Ready - Monitoring auth operations");
});
'''
    scripts.append(FridaScript(
        name="Authentication Hook",
        category="auth_hook",
        description="Monitor authentication operations, tokens, biometrics, and HTTP auth headers",
        script_code=auth_hook_script,
        target_classes=["android.content.SharedPreferences", "okhttp3.Request$Builder", "java.security.KeyStore"],
        target_methods=["putString", "header", "addHeader", "authenticate", "getKey"],
        is_dangerous=False,
        usage_instructions=f"frida -U -f {package_name} -l auth_hook.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Method Tracer Script
    # =========================================================================
    method_trace_script = '''// Method Tracing Script for ''' + package_name + '''
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f ''' + package_name + ''' -l method_trace.js --no-pause
// Modify CLASS_PATTERNS and METHOD_PATTERNS to trace specific classes/methods

var CLASS_PATTERNS = [
    // Add your target class patterns here
    "''' + package_name.replace('.', '/') + '''",
];

var METHOD_PATTERNS = [
    // Methods to always trace regardless of class
    "login", "auth", "verify", "validate", "check",
    "encrypt", "decrypt", "sign", "hash",
    "getToken", "setToken", "refreshToken",
    "sendRequest", "doPost", "doGet"
];

var EXCLUDE_PATTERNS = [
    "toString", "hashCode", "equals", "valueOf",
    "<init>", "<clinit>", "access$"
];

Java.perform(function() {
    console.log("[*] Method Tracer Script Loaded");
    console.log("[*] Tracing classes matching: " + CLASS_PATTERNS.join(", "));
    
    var tracedMethods = 0;
    var maxMethods = 500; // Prevent overload
    
    Java.enumerateLoadedClasses({
        onMatch: function(className) {
            if (tracedMethods >= maxMethods) return;
            
            var shouldTrace = false;
            for (var i = 0; i < CLASS_PATTERNS.length; i++) {
                if (className.indexOf(CLASS_PATTERNS[i]) !== -1) {
                    shouldTrace = true;
                    break;
                }
            }
            
            if (!shouldTrace) return;
            
            try {
                var clazz = Java.use(className);
                var methods = clazz.class.getDeclaredMethods();
                
                methods.forEach(function(method) {
                    var methodName = method.getName();
                    
                    // Skip excluded methods
                    for (var i = 0; i < EXCLUDE_PATTERNS.length; i++) {
                        if (methodName.indexOf(EXCLUDE_PATTERNS[i]) !== -1) return;
                    }
                    
                    if (tracedMethods >= maxMethods) return;
                    
                    try {
                        var overloads = clazz[methodName].overloads;
                        overloads.forEach(function(overload) {
                            overload.implementation = function() {
                                var args = [];
                                for (var i = 0; i < arguments.length; i++) {
                                    try {
                                        args.push(String(arguments[i]).substring(0, 100));
                                    } catch(e) {
                                        args.push('[unprintable]');
                                    }
                                }
                                console.log('[TRACE] ' + className + '.' + methodName + '(' + args.join(', ') + ')');
                                
                                var retval = this[methodName].apply(this, arguments);
                                
                                if (retval !== undefined) {
                                    try {
                                        console.log('[TRACE] └─ Return: ' + String(retval).substring(0, 200));
                                    } catch(e) {
                                        console.log('[TRACE] └─ Return: [unprintable]');
                                    }
                                }
                                return retval;
                            };
                            tracedMethods++;
                        });
                    } catch (e) {
                        // Method hook failed, skip
                    }
                });
            } catch (e) {
                // Class use failed, skip
            }
        },
        onComplete: function() {
            console.log("[*] Method Tracer Setup Complete - Traced " + tracedMethods + " methods");
        }
    });
});
'''
    scripts.append(FridaScript(
        name="Method Tracer",
        category="method_trace",
        description="Trace method calls in app classes with arguments and return values",
        script_code=method_trace_script,
        target_classes=[package_name],
        target_methods=["*"],
        is_dangerous=False,
        usage_instructions=f"frida -U -f {package_name} -l method_trace.js --no-pause\n# Edit CLASS_PATTERNS in script to target specific classes"
    ))
    
    # =========================================================================
    # Generate Network Traffic Logger Script
    # =========================================================================
    network_hook_script = '''// Network Traffic Logger for ''' + package_name + '''
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f ''' + package_name + ''' -l network_hook.js --no-pause

Java.perform(function() {
    console.log("[*] Network Traffic Logger Loaded");
    
    // ===== OkHttp Interceptor =====
    try {
        var OkHttpClient = Java.use('okhttp3.OkHttpClient');
        var Request = Java.use('okhttp3.Request');
        var Response = Java.use('okhttp3.Response');
        var RequestBody = Java.use('okhttp3.RequestBody');
        var Buffer = Java.use('okio.Buffer');
        
        var Interceptor = Java.use('okhttp3.Interceptor');
        var Chain = Java.use('okhttp3.Interceptor$Chain');
        
        // Hook newCall to log requests
        var RealCall = Java.use('okhttp3.RealCall');
        RealCall.execute.implementation = function() {
            var request = this.request();
            console.log('\\n[HTTP] ═══════════════════════════════════════');
            console.log('[HTTP] ' + request.method() + ' ' + request.url().toString());
            
            // Log headers
            var headers = request.headers();
            for (var i = 0; i < headers.size(); i++) {
                console.log('[HTTP] ' + headers.name(i) + ': ' + headers.value(i));
            }
            
            // Log body if present
            var body = request.body();
            if (body !== null) {
                try {
                    var buffer = Buffer.$new();
                    body.writeTo(buffer);
                    console.log('[HTTP] Body: ' + buffer.readUtf8());
                } catch(e) {
                    console.log('[HTTP] Body: [binary or empty]');
                }
            }
            
            var response = this.execute();
            console.log('[HTTP] Response: ' + response.code() + ' ' + response.message());
            return response;
        };
        console.log('[*] OkHttp hooks installed');
    } catch (e) {
        console.log('[-] OkHttp hooks failed: ' + e);
    }
    
    // ===== HttpURLConnection =====
    try {
        var HttpURLConnection = Java.use('java.net.HttpURLConnection');
        var URL = Java.use('java.net.URL');
        
        HttpURLConnection.connect.implementation = function() {
            console.log('[HTTP] HttpURLConnection.connect()');
            console.log('  URL: ' + this.getURL().toString());
            console.log('  Method: ' + this.getRequestMethod());
            return this.connect();
        };
        
        HttpURLConnection.getInputStream.implementation = function() {
            console.log('[HTTP] HttpURLConnection.getInputStream()');
            console.log('  URL: ' + this.getURL().toString());
            console.log('  Response Code: ' + this.getResponseCode());
            return this.getInputStream();
        };
        console.log('[*] HttpURLConnection hooks installed');
    } catch (e) {
        console.log('[-] HttpURLConnection hooks failed: ' + e);
    }
    
    // ===== Retrofit =====
    try {
        var Retrofit = Java.use('retrofit2.Retrofit');
        Retrofit.create.implementation = function(service) {
            console.log('[RETROFIT] Creating service: ' + service);
            console.log('  Base URL: ' + this.baseUrl().toString());
            return this.create(service);
        };
        console.log('[*] Retrofit hooks installed');
    } catch (e) {
        console.log('[-] Retrofit not found');
    }
    
    // ===== WebView URL Loading =====
    try {
        var WebView = Java.use('android.webkit.WebView');
        WebView.loadUrl.overload('java.lang.String').implementation = function(url) {
            console.log('[WEBVIEW] loadUrl: ' + url);
            return this.loadUrl(url);
        };
        WebView.loadUrl.overload('java.lang.String', 'java.util.Map').implementation = function(url, headers) {
            console.log('[WEBVIEW] loadUrl: ' + url);
            console.log('  Headers: ' + headers);
            return this.loadUrl(url, headers);
        };
        console.log('[*] WebView hooks installed');
    } catch (e) {
        console.log('[-] WebView hooks failed: ' + e);
    }
    
    console.log("[*] Network Logger Ready - All HTTP traffic will be logged");
});
'''
    scripts.append(FridaScript(
        name="Network Traffic Logger",
        category="network_hook",
        description="Log all HTTP/HTTPS requests including URLs, headers, and bodies",
        script_code=network_hook_script,
        target_classes=["okhttp3.RealCall", "java.net.HttpURLConnection", "android.webkit.WebView"],
        target_methods=["execute", "connect", "loadUrl"],
        is_dangerous=False,
        usage_instructions=f"frida -U -f {package_name} -l network_hook.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Combined "All-in-One" Script
    # =========================================================================
    combined_script = f'''// Combined Security Testing Script for {package_name}
// Generated by VRAgent APK Analyzer
// This script combines SSL bypass, root bypass, and essential hooks
// Usage: frida -U -f {package_name} -l combined.js --no-pause

Java.perform(function() {{
    console.log("╔════════════════════════════════════════════════════════════╗");
    console.log("║       VRAgent Combined Security Testing Script             ║");
    console.log("║       Package: {package_name:<36} ║");
    console.log("╚════════════════════════════════════════════════════════════╝");
    
    // ---------- SSL PINNING BYPASS ----------
    console.log("\\n[*] Installing SSL Pinning Bypass...");
    try {{
        var CertificatePinner = Java.use('okhttp3.CertificatePinner');
        CertificatePinner.check.overload('java.lang.String', 'java.util.List').implementation = function(h, c) {{
            console.log('[SSL] Bypassed for: ' + h);
        }};
    }} catch(e) {{}}
    
    try {{
        var TrustManagerImpl = Java.use('com.android.org.conscrypt.TrustManagerImpl');
        TrustManagerImpl.verifyChain.implementation = function(u, t, h, c, o, s) {{
            console.log('[SSL] TrustManager bypassed');
            return u;
        }};
    }} catch(e) {{}}
    console.log("[+] SSL bypass installed");
    
    // ---------- ROOT DETECTION BYPASS ----------
    console.log("\\n[*] Installing Root Detection Bypass...");
    try {{
        var RootBeer = Java.use('com.scottyab.rootbeer.RootBeer');
        RootBeer.isRooted.implementation = function() {{ return false; }};
        RootBeer.isRootedWithoutBusyBoxCheck.implementation = function() {{ return false; }};
    }} catch(e) {{}}
    
    try {{
        var File = Java.use('java.io.File');
        var orig = File.exists;
        File.exists.implementation = function() {{
            var p = this.getAbsolutePath();
            if (p.indexOf('su') !== -1 || p.indexOf('magisk') !== -1) return false;
            return orig.call(this);
        }};
    }} catch(e) {{}}
    console.log("[+] Root bypass installed");
    
    // ---------- CRYPTO HOOKS ----------
    console.log("\\n[*] Installing Crypto Hooks...");
    try {{
        var Cipher = Java.use('javax.crypto.Cipher');
        Cipher.doFinal.overload('[B').implementation = function(i) {{
            console.log('[CRYPTO] Cipher.doFinal - ' + i.length + ' bytes');
            return this.doFinal(i);
        }};
        
        var SecretKeySpec = Java.use('javax.crypto.spec.SecretKeySpec');
        SecretKeySpec.$init.overload('[B', 'java.lang.String').implementation = function(k, a) {{
            function hex(b) {{ var h=''; for(var i=0;i<b.length;i++) h+=('0'+(b[i]&0xFF).toString(16)).slice(-2); return h; }}
            console.log('[CRYPTO] Key: ' + hex(k) + ' (' + a + ')');
            return this.$init(k, a);
        }};
    }} catch(e) {{}}
    console.log("[+] Crypto hooks installed");
    
    // ---------- AUTH HOOKS ----------
    console.log("\\n[*] Installing Auth Hooks...");
    try {{
        var Editor = Java.use('android.content.SharedPreferences$Editor');
        Editor.putString.implementation = function(k, v) {{
            if (k.toLowerCase().match(/token|session|auth|jwt/)) {{
                console.log('[AUTH] ' + k + ' = ' + v);
            }}
            return this.putString(k, v);
        }};
    }} catch(e) {{}}
    console.log("[+] Auth hooks installed");
    
    console.log("\\n[*] ════════════════════════════════════════════════════════");
    console.log("[*] All hooks installed successfully!");
    console.log("[*] SSL: Bypassed | Root: Bypassed | Crypto: Logged | Auth: Logged");
    console.log("[*] ════════════════════════════════════════════════════════\\n");
}});
'''
    scripts.append(FridaScript(
        name="Combined Security Script",
        category="combined",
        description="All-in-one script: SSL bypass + Root bypass + Crypto hooks + Auth monitoring",
        script_code=combined_script,
        target_classes=["*"],
        target_methods=["*"],
        is_dangerous=True,
        usage_instructions=f"frida -U -f {package_name} -l combined.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Emulator Detection Bypass Script
    # =========================================================================
    emulator_bypass_script = f'''// Emulator Detection Bypass for {package_name}
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f {package_name} -l emulator_bypass.js --no-pause

Java.perform(function() {{
    console.log("[*] Emulator Detection Bypass Script Loaded");
    
    // ===== Build Properties Spoofing =====
    try {{
        var Build = Java.use('android.os.Build');
        
        // Spoof device properties
        Build.FINGERPRINT.value = 'google/sunfish/sunfish:11/RQ3A.210805.001.A1/7474174:user/release-keys';
        Build.MODEL.value = 'Pixel 4a';
        Build.MANUFACTURER.value = 'Google';
        Build.BRAND.value = 'google';
        Build.DEVICE.value = 'sunfish';
        Build.PRODUCT.value = 'sunfish';
        Build.HARDWARE.value = 'sunfish';
        Build.BOARD.value = 'sunfish';
        Build.HOST.value = 'abfarm';
        Build.TAGS.value = 'release-keys';
        
        console.log('[+] Build properties spoofed to Pixel 4a');
    }} catch (e) {{
        console.log('[-] Build spoofing failed: ' + e);
    }}
    
    // ===== System Properties =====
    try {{
        var SystemProperties = Java.use('android.os.SystemProperties');
        var originalGet = SystemProperties.get.overload('java.lang.String');
        
        SystemProperties.get.overload('java.lang.String').implementation = function(key) {{
            var emulatorProps = {{
                'ro.kernel.qemu': '0',
                'ro.hardware.virtual_device': '',
                'ro.product.device': 'sunfish',
                'ro.product.model': 'Pixel 4a',
                'ro.product.brand': 'google',
                'ro.boot.qemu': '0',
                'init.svc.qemu-props': '',
                'qemu.hw.mainkeys': '',
                'ro.kernel.android.qemud': '',
                'ro.kernel.qemu.gles': '',
            }};
            
            if (key in emulatorProps) {{
                console.log('[+] SystemProperties.get("' + key + '") spoofed');
                return emulatorProps[key];
            }}
            return originalGet.call(this, key);
        }};
        console.log('[*] SystemProperties hooks installed');
    }} catch (e) {{
        console.log('[-] SystemProperties hook failed: ' + e);
    }}
    
    // ===== TelephonyManager =====
    try {{
        var TelephonyManager = Java.use('android.telephony.TelephonyManager');
        
        TelephonyManager.getDeviceId.overload().implementation = function() {{
            console.log('[+] TelephonyManager.getDeviceId() spoofed');
            return '358240051111110';
        }};
        
        TelephonyManager.getSubscriberId.implementation = function() {{
            console.log('[+] TelephonyManager.getSubscriberId() spoofed');
            return '310260000000000';
        }};
        
        TelephonyManager.getLine1Number.implementation = function() {{
            console.log('[+] TelephonyManager.getLine1Number() spoofed');
            return '+15555555555';
        }};
        
        TelephonyManager.getNetworkOperator.implementation = function() {{
            console.log('[+] TelephonyManager.getNetworkOperator() spoofed');
            return '310260';
        }};
        
        TelephonyManager.getNetworkOperatorName.implementation = function() {{
            console.log('[+] TelephonyManager.getNetworkOperatorName() spoofed');
            return 'T-Mobile';
        }};
        
        TelephonyManager.getSimOperator.implementation = function() {{
            console.log('[+] TelephonyManager.getSimOperator() spoofed');
            return '310260';
        }};
        
        console.log('[*] TelephonyManager hooks installed');
    }} catch (e) {{
        console.log('[-] TelephonyManager hooks failed: ' + e);
    }}
    
    // ===== Sensors (emulators often have fake sensors) =====
    try {{
        var SensorManager = Java.use('android.hardware.SensorManager');
        var originalGetSensorList = SensorManager.getSensorList;
        // Most emulator detection checks if sensor list is empty or has specific values
        console.log('[*] SensorManager context established');
    }} catch (e) {{
        console.log('[-] SensorManager hook failed: ' + e);
    }}
    
    // ===== File checks for emulator artifacts =====
    try {{
        var File = Java.use('java.io.File');
        var originalExists = File.exists;
        File.exists.implementation = function() {{
            var path = this.getAbsolutePath();
            var emulatorPaths = [
                '/dev/socket/qemud',
                '/dev/qemu_pipe',
                '/system/lib/libc_malloc_debug_qemu.so',
                '/sys/qemu_trace',
                '/system/bin/qemu-props',
                '/dev/socket/genyd',
                '/dev/socket/baseband_genyd',
                'ueventd.android_x86.rc',
                'x86.prop',
                'ueventd.ttVM_x86.rc',
                'init.ttVM_x86.rc',
                'fstab.ttVM_x86',
                'fstab.vbox86',
                'init.vbox86.rc',
                'ueventd.vbox86.rc',
            ];
            for (var i = 0; i < emulatorPaths.length; i++) {{
                if (path.indexOf(emulatorPaths[i]) !== -1) {{
                    console.log('[+] File.exists() bypassed for emulator path: ' + path);
                    return false;
                }}
            }}
            return originalExists.call(this);
        }};
        console.log('[*] File.exists() emulator bypass installed');
    }} catch (e) {{
        console.log('[-] File.exists() hook failed: ' + e);
    }}
    
    console.log("[*] Emulator Detection Bypass Complete");
}});
'''
    scripts.append(FridaScript(
        name="Emulator Detection Bypass",
        category="emulator_bypass",
        description="Bypass emulator detection by spoofing Build properties, system props, and telephony",
        script_code=emulator_bypass_script,
        target_classes=["android.os.Build", "android.os.SystemProperties", "android.telephony.TelephonyManager"],
        target_methods=["get", "getDeviceId", "getSubscriberId"],
        is_dangerous=True,
        usage_instructions=f"frida -U -f {package_name} -l emulator_bypass.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Debugger Detection Bypass Script  
    # =========================================================================
    debugger_bypass_script = f'''// Debugger Detection Bypass for {package_name}
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f {package_name} -l debugger_bypass.js --no-pause

Java.perform(function() {{
    console.log("[*] Debugger Detection Bypass Script Loaded");
    
    // ===== android.os.Debug =====
    try {{
        var Debug = Java.use('android.os.Debug');
        
        Debug.isDebuggerConnected.implementation = function() {{
            console.log('[+] Debug.isDebuggerConnected() bypassed - returning false');
            return false;
        }};
        
        Debug.waitingForDebugger.implementation = function() {{
            console.log('[+] Debug.waitingForDebugger() bypassed - returning false');
            return false;
        }};
        
        console.log('[*] android.os.Debug hooks installed');
    }} catch (e) {{
        console.log('[-] Debug hooks failed: ' + e);
    }}
    
    // ===== ApplicationInfo flags =====
    try {{
        var ApplicationInfo = Java.use('android.content.pm.ApplicationInfo');
        ApplicationInfo.flags.value = ApplicationInfo.flags.value & ~2; // Remove FLAG_DEBUGGABLE
        console.log('[*] ApplicationInfo.FLAG_DEBUGGABLE cleared');
    }} catch (e) {{
        console.log('[-] ApplicationInfo patch failed: ' + e);
    }}
    
    // ===== /proc/self/status TracerPid check =====
    try {{
        var BufferedReader = Java.use('java.io.BufferedReader');
        var originalReadLine = BufferedReader.readLine;
        BufferedReader.readLine.implementation = function() {{
            var line = originalReadLine.call(this);
            if (line && line.indexOf('TracerPid') !== -1) {{
                console.log('[+] TracerPid line intercepted, returning 0');
                return 'TracerPid:\\t0';
            }}
            return line;
        }};
        console.log('[*] TracerPid bypass installed');
    }} catch (e) {{
        console.log('[-] TracerPid bypass failed: ' + e);
    }}
    
    // ===== Timer-based anti-debugging =====
    try {{
        var System = Java.use('java.lang.System');
        var lastTime = 0;
        var originalNanoTime = System.nanoTime;
        System.nanoTime.implementation = function() {{
            // Prevent large time gaps that indicate debugging
            var currentTime = originalNanoTime.call(this);
            if (lastTime !== 0 && (currentTime - lastTime) > 1000000000) {{ // > 1 second
                console.log('[+] Large time gap detected, normalizing');
                currentTime = lastTime + 1000000; // Add 1ms instead
            }}
            lastTime = currentTime;
            return currentTime;
        }};
        console.log('[*] Timer anti-debug bypass installed');
    }} catch (e) {{
        console.log('[-] Timer bypass failed: ' + e);
    }}
    
    // ===== Runtime.exec for ps/pidof commands =====
    try {{
        var Runtime = Java.use('java.lang.Runtime');
        var originalExec = Runtime.exec.overload('java.lang.String');
        Runtime.exec.overload('java.lang.String').implementation = function(cmd) {{
            if (cmd.indexOf('ps') !== -1 || cmd.indexOf('pidof') !== -1 || cmd.indexOf('gdb') !== -1) {{
                console.log('[+] Blocked debug-detection command: ' + cmd);
                // Return a fake process that does nothing
                return originalExec.call(this, 'echo');
            }}
            return originalExec.call(this, cmd);
        }};
        console.log('[*] Runtime.exec debug bypass installed');
    }} catch (e) {{
        console.log('[-] Runtime.exec bypass failed: ' + e);
    }}
    
    console.log("[*] Debugger Detection Bypass Complete");
}});
'''
    scripts.append(FridaScript(
        name="Debugger Detection Bypass",
        category="debugger_bypass",
        description="Bypass debugger detection including Debug.isDebuggerConnected, TracerPid, timing checks",
        script_code=debugger_bypass_script,
        target_classes=["android.os.Debug", "java.io.BufferedReader", "java.lang.System"],
        target_methods=["isDebuggerConnected", "waitingForDebugger", "readLine", "nanoTime"],
        is_dangerous=True,
        usage_instructions=f"frida -U -f {package_name} -l debugger_bypass.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Anti-Tampering/Integrity Bypass Script
    # =========================================================================
    tampering_bypass_script = f'''// Anti-Tampering & Integrity Bypass for {package_name}
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f {package_name} -l tampering_bypass.js --no-pause

Java.perform(function() {{
    console.log("[*] Anti-Tampering Bypass Script Loaded");
    
    // ===== PackageManager Signature Spoofing =====
    try {{
        var PackageManager = Java.use('android.content.pm.PackageManager');
        var PackageInfo = Java.use('android.content.pm.PackageInfo');
        var Signature = Java.use('android.content.pm.Signature');
        
        // Store original signature for spoofing
        var originalSignature = null;
        
        var ApplicationPackageManager = Java.use('android.app.ApplicationPackageManager');
        ApplicationPackageManager.getPackageInfo.overload('java.lang.String', 'int').implementation = function(packageName, flags) {{
            var result = this.getPackageInfo(packageName, flags);
            
            // Check if signatures are requested
            if ((flags & 0x40) !== 0) {{ // GET_SIGNATURES = 0x40
                console.log('[+] Signature request intercepted for: ' + packageName);
                // You can spoof the signature here if needed
                // result.signatures = [Signature.$new("original_hex_signature")];
            }}
            return result;
        }};
        console.log('[*] PackageManager signature hooks installed');
    }} catch (e) {{
        console.log('[-] PackageManager hooks failed: ' + e);
    }}
    
    // ===== File Integrity Checks =====
    try {{
        var MessageDigest = Java.use('java.security.MessageDigest');
        var digestResults = {{}};
        
        MessageDigest.digest.overload('[B').implementation = function(input) {{
            var result = this.digest(input);
            var algo = this.getAlgorithm();
            
            // Log checksum calculations (potential integrity checks)
            if (input.length > 1000) {{ // Likely a file being checksummed
                console.log('[INTEGRITY] ' + algo + ' digest calculated on ' + input.length + ' bytes');
            }}
            return result;
        }};
        console.log('[*] Integrity check logging installed');
    }} catch (e) {{
        console.log('[-] Integrity check hooks failed: ' + e);
    }}
    
    // ===== SafetyNet/Play Integrity Bypass =====
    try {{
        // SafetyNet Attestation
        var SafetyNetClient = Java.use('com.google.android.gms.safetynet.SafetyNetClient');
        SafetyNetClient.attest.implementation = function(nonce, apiKey) {{
            console.log('[SAFETYNET] attest() called - this would need a valid response');
            return this.attest(nonce, apiKey);
        }};
        console.log('[*] SafetyNet hooks installed');
    }} catch (e) {{
        console.log('[-] SafetyNet not found or hooks failed');
    }}
    
    // ===== Installer Package Check =====
    try {{
        var ApplicationPackageManager = Java.use('android.app.ApplicationPackageManager');
        ApplicationPackageManager.getInstallerPackageName.implementation = function(packageName) {{
            console.log('[+] getInstallerPackageName() spoofed to com.android.vending');
            return 'com.android.vending'; // Spoof as installed from Play Store
        }};
        console.log('[*] Installer package spoof installed');
    }} catch (e) {{
        console.log('[-] Installer spoof failed: ' + e);
    }}
    
    // ===== APK Path Verification =====
    try {{
        var ApplicationInfo = Java.use('android.content.pm.ApplicationInfo');
        // Some apps check if APK is in expected location
        console.log('[*] ApplicationInfo context established');
    }} catch (e) {{
        console.log('[-] ApplicationInfo hooks failed');
    }}
    
    // ===== Xposed/Frida Detection Bypass =====
    try {{
        // Bypass common Frida detection methods
        var Module = Java.use('java.lang.reflect.Module');
        console.log('[*] Reflection module context established');
    }} catch (e) {{}}
    
    try {{
        // Hide frida-server
        var File = Java.use('java.io.File');
        var originalExists = File.exists;
        File.exists.implementation = function() {{
            var path = this.getAbsolutePath();
            var fridaPaths = [
                'frida', 'frida-server', 'frida-agent',
                'libfrida', 'gmain', 'gum-js-loop',
                'linjector', '/data/local/tmp/re.frida'
            ];
            for (var i = 0; i < fridaPaths.length; i++) {{
                if (path.toLowerCase().indexOf(fridaPaths[i]) !== -1) {{
                    console.log('[+] Frida file check bypassed: ' + path);
                    return false;
                }}
            }}
            return originalExists.call(this);
        }};
        console.log('[*] Frida detection bypass installed');
    }} catch (e) {{
        console.log('[-] Frida bypass failed: ' + e);
    }}
    
    console.log("[*] Anti-Tampering Bypass Complete");
}});
'''
    scripts.append(FridaScript(
        name="Anti-Tampering Bypass",
        category="tampering_bypass",
        description="Bypass signature checks, SafetyNet, integrity verification, and Frida detection",
        script_code=tampering_bypass_script,
        target_classes=["android.app.ApplicationPackageManager", "java.security.MessageDigest"],
        target_methods=["getPackageInfo", "getInstallerPackageName", "digest"],
        is_dangerous=True,
        usage_instructions=f"frida -U -f {package_name} -l tampering_bypass.js --no-pause"
    ))
    
    # =========================================================================
    # Generate Screenshot/Screen Capture Bypass Script
    # =========================================================================
    screenshot_bypass_script = f'''// Screenshot/Screen Capture Bypass for {package_name}
// Generated by VRAgent APK Analyzer
// Usage: frida -U -f {package_name} -l screenshot_bypass.js --no-pause

Java.perform(function() {{
    console.log("[*] Screenshot Bypass Script Loaded");
    
    // ===== Window FLAG_SECURE Bypass =====
    try {{
        var Window = Java.use('android.view.Window');
        var originalSetFlags = Window.setFlags;
        Window.setFlags.implementation = function(flags, mask) {{
            // FLAG_SECURE = 0x2000
            if ((flags & 0x2000) !== 0) {{
                console.log('[+] FLAG_SECURE detected, removing it');
                flags = flags & ~0x2000;
            }}
            return originalSetFlags.call(this, flags, mask);
        }};
        console.log('[*] Window.setFlags() hook installed');
    }} catch (e) {{
        console.log('[-] Window.setFlags() hook failed: ' + e);
    }}
    
    // ===== SurfaceView Secure Layer Bypass =====
    try {{
        var SurfaceView = Java.use('android.view.SurfaceView');
        SurfaceView.setSecure.implementation = function(isSecure) {{
            console.log('[+] SurfaceView.setSecure(' + isSecure + ') bypassed');
            return this.setSecure(false);
        }};
        console.log('[*] SurfaceView.setSecure() hook installed');
    }} catch (e) {{
        console.log('[-] SurfaceView hook failed: ' + e);
    }}
    
    // ===== TextureView Secure Bypass =====
    try {{
        var TextureView = Java.use('android.view.TextureView');
        // TextureView doesn't have setSecure but check for secure surface
        console.log('[*] TextureView context established');
    }} catch (e) {{}}
    
    // ===== Activity Window Flag Bypass =====
    try {{
        var Activity = Java.use('android.app.Activity');
        Activity.onCreate.overload('android.os.Bundle').implementation = function(savedInstanceState) {{
            this.onCreate(savedInstanceState);
            
            // Remove FLAG_SECURE from window after creation
            try {{
                var window = this.getWindow();
                window.clearFlags(0x2000); // FLAG_SECURE
                console.log('[+] Cleared FLAG_SECURE from Activity: ' + this.getClass().getName());
            }} catch (e) {{}}
        }};
        console.log('[*] Activity onCreate hook for FLAG_SECURE installed');
    }} catch (e) {{
        console.log('[-] Activity hook failed: ' + e);
    }}
    
    // ===== Fragment Window Flag Bypass =====
    try {{
        var Fragment = Java.use('androidx.fragment.app.Fragment');
        Fragment.onViewCreated.overload('android.view.View', 'android.os.Bundle').implementation = function(view, savedInstanceState) {{
            this.onViewCreated(view, savedInstanceState);
            try {{
                var activity = this.getActivity();
                if (activity !== null) {{
                    var window = activity.getWindow();
                    window.clearFlags(0x2000);
                    console.log('[+] Cleared FLAG_SECURE from Fragment');
                }}
            }} catch (e) {{}}
        }};
        console.log('[*] Fragment hook for FLAG_SECURE installed');
    }} catch (e) {{
        console.log('[-] Fragment hook failed');
    }}
    
    console.log("[*] Screenshot Bypass Complete - Screen capture should now work");
}});
'''
    scripts.append(FridaScript(
        name="Screenshot Bypass",
        category="screenshot_bypass",
        description="Bypass FLAG_SECURE and other screen capture protections for testing",
        script_code=screenshot_bypass_script,
        target_classes=["android.view.Window", "android.view.SurfaceView", "android.app.Activity"],
        target_methods=["setFlags", "setSecure", "clearFlags"],
        is_dangerous=True,
        usage_instructions=f"frida -U -f {package_name} -l screenshot_bypass.js --no-pause"
    ))
    
    # Build interesting hooks list from analysis
    if dex_analysis and "classes" in dex_analysis:
        for cls in dex_analysis.get("classes", [])[:20]:
            if any(pattern in cls.get("name", "").lower() for pattern in ["crypto", "auth", "login", "key", "token", "password", "secret"]):
                interesting_hooks.append({
                    "class": cls.get("name"),
                    "reason": "Contains sensitive-sounding name",
                    "methods": cls.get("methods", [])[:5]
                })
    
    # Generate suggested test cases
    suggested_tests = [
        "Test SSL pinning bypass by proxying traffic through Burp Suite",
        "Verify root detection bypass on rooted device",
        "Monitor crypto operations during login flow",
        "Trace authentication token generation and storage",
        "Intercept and modify API requests/responses"
    ]
    
    if ssl_pinning_detected:
        suggested_tests.append("SSL pinning detected - test certificate pinning bypass script")
    if root_detection_detected:
        suggested_tests.append("Root detection found - verify root bypass script effectiveness")
    if crypto_methods:
        suggested_tests.append("Cryptographic operations detected - monitor key material with crypto hook")
    if 'biometric' in str(auth_methods_found).lower():
        suggested_tests.append("Biometric auth found - test biometric bypass scenarios")
    
    # Convert scripts to dict format
    scripts_dict = [asdict(s) for s in scripts]
    
    return {
        "package_name": package_name,
        "frida_scripts": scripts_dict,
        "ssl_pinning_detected": ssl_pinning_detected,
        "ssl_patterns_found": list(ssl_classes_found),
        "root_detection_detected": root_detection_detected,
        "root_patterns_found": list(root_classes_found),
        "crypto_methods": crypto_methods[:10],
        "auth_patterns_found": list(auth_methods_found),
        "interesting_hooks": interesting_hooks,
        "suggested_test_cases": suggested_tests,
        "frida_spawn_command": f"frida -U -f {package_name} -l <script.js> --no-pause",
        "frida_attach_command": f"frida -U {package_name} -l <script.js>",
        "total_scripts": len(scripts),
        "emulator_detection_detected": emulator_detection_detected,
        "anti_tampering_detected": anti_tampering_detected,
        "debugger_detection_detected": debugger_detection_detected,
    }


# ============================================================================
# Native Library Analysis
# ============================================================================

def analyze_native_libraries(apk, file_path: Path) -> Dict[str, Any]:
    """
    Analyze native libraries (.so files) in the APK.
    
    Extracts:
    - JNI function signatures
    - Exported symbols
    - Hardcoded strings and secrets
    - Anti-debugging techniques
    - Cryptographic function usage
    """
    import zipfile
    import struct
    
    result = {
        "total_libraries": 0,
        "libraries": [],
        "total_jni_functions": 0,
        "total_exported_functions": 0,
        "architectures": [],
        "security_findings": [],
        "overall_native_risk": "low"
    }
    
    # Architecture mapping
    ARCH_MAP = {
        'armeabi-v7a': 'ARM 32-bit',
        'arm64-v8a': 'ARM 64-bit',
        'x86': 'x86 32-bit',
        'x86_64': 'x86 64-bit',
        'armeabi': 'ARM (legacy)',
        'mips': 'MIPS',
        'mips64': 'MIPS 64-bit',
    }
    
    # Anti-debugging patterns to look for in native code
    ANTI_DEBUG_PATTERNS = [
        (b'ptrace', 'ptrace() anti-debugging'),
        (b'/proc/self/status', 'TracerPid check'),
        (b'/proc/self/maps', 'Memory maps inspection'),
        (b'SIGTRAP', 'SIGTRAP signal handler'),
        (b'SIGSTOP', 'SIGSTOP signal handler'),
        (b'android_dlopen_ext', 'Dynamic library loading'),
        (b'dlsym', 'Dynamic symbol resolution'),
        (b'inotify', 'File monitoring (anti-tampering)'),
        (b'/data/local/tmp', 'Temp directory check (frida detection)'),
        (b'frida', 'Frida detection'),
        (b'xposed', 'Xposed detection'),
        (b'substrate', 'Cydia Substrate detection'),
        (b'libhoudini', 'ARM translation detection'),
        (b'ro.debuggable', 'Debug build check'),
        (b'ro.secure', 'Secure boot check'),
        (b'/system/bin/su', 'Root check (native)'),
        (b'magisk', 'Magisk detection (native)'),
    ]
    
    # Crypto patterns in native code
    NATIVE_CRYPTO_PATTERNS = [
        (b'AES', 'AES encryption'),
        (b'DES', 'DES encryption (weak)'),
        (b'RSA', 'RSA encryption'),
        (b'SHA256', 'SHA-256 hashing'),
        (b'SHA1', 'SHA-1 hashing (weak)'),
        (b'MD5', 'MD5 hashing (weak)'),
        (b'HMAC', 'HMAC authentication'),
        (b'EVP_', 'OpenSSL EVP functions'),
        (b'OPENSSL', 'OpenSSL library'),
        (b'mbedtls', 'mbed TLS library'),
        (b'boringssl', 'BoringSSL library'),
        (b'sodium', 'libsodium library'),
        (b'curve25519', 'Curve25519 (modern crypto)'),
        (b'chacha', 'ChaCha20 encryption'),
        (b'poly1305', 'Poly1305 MAC'),
    ]
    
    # JNI function prefixes
    JNI_PREFIXES = [
        'Java_', 
        'JNI_OnLoad',
        'JNI_OnUnload',
        'nativeInit',
        'native_',
    ]
    
    try:
        with zipfile.ZipFile(str(file_path), 'r') as zf:
            native_files = [f for f in zf.namelist() if f.endswith('.so') and '/lib/' in f]
            
            architectures_found = set()
            total_jni = 0
            total_exported = 0
            risk_score = 0
            
            for so_path in native_files:
                # Extract architecture from path
                parts = so_path.split('/')
                arch = None
                for part in parts:
                    if part in ARCH_MAP:
                        arch = part
                        architectures_found.add(ARCH_MAP[part])
                        break
                
                lib_name = so_path.split('/')[-1]
                
                try:
                    so_data = zf.read(so_path)
                    lib_info = {
                        "name": lib_name,
                        "path": so_path,
                        "architecture": ARCH_MAP.get(arch, arch or 'unknown'),
                        "size": len(so_data),
                        "is_stripped": True,  # Will update if we find debug info
                        "has_debug_info": False,
                        "exported_functions": [],
                        "jni_functions": [],
                        "imported_libraries": [],
                        "strings": [],
                        "hardcoded_secrets": [],
                        "anti_debug_detected": False,
                        "anti_debug_techniques": [],
                        "crypto_functions": [],
                        "suspicious_patterns": [],
                    }
                    
                    # Try to parse ELF structure for more details
                    try:
                        from elftools.elf.elffile import ELFFile
                        from elftools.elf.sections import SymbolTableSection
                        from io import BytesIO
                        
                        elf = ELFFile(BytesIO(so_data))
                        
                        # Check for debug sections
                        for section in elf.iter_sections():
                            if section.name.startswith('.debug'):
                                lib_info["has_debug_info"] = True
                                lib_info["is_stripped"] = False
                                result["security_findings"].append({
                                    "library": lib_name,
                                    "finding": "Debug symbols present",
                                    "severity": "medium",
                                    "description": "Library contains debug information which aids reverse engineering"
                                })
                                break
                        
                        # Extract symbols
                        for section in elf.iter_sections():
                            if isinstance(section, SymbolTableSection):
                                for symbol in section.iter_symbols():
                                    sym_name = symbol.name
                                    if not sym_name:
                                        continue
                                    
                                    # Check for JNI functions
                                    is_jni = any(sym_name.startswith(prefix) for prefix in JNI_PREFIXES)
                                    if is_jni:
                                        lib_info["jni_functions"].append(sym_name)
                                        total_jni += 1
                                    
                                    # Exported functions (global symbols)
                                    if symbol['st_info']['bind'] == 'STB_GLOBAL' and symbol['st_shndx'] != 'SHN_UNDEF':
                                        lib_info["exported_functions"].append({
                                            "name": sym_name,
                                            "address": hex(symbol['st_value']),
                                            "size": symbol['st_size'],
                                            "is_jni": is_jni,
                                        })
                                        total_exported += 1
                        
                        # Get imported libraries
                        for section in elf.iter_sections():
                            if section.name == '.dynamic':
                                for tag in section.iter_tags():
                                    if tag.entry.d_tag == 'DT_NEEDED':
                                        lib_info["imported_libraries"].append(tag.needed)
                        
                    except Exception as elf_error:
                        logger.debug(f"ELF parsing failed for {lib_name}: {elf_error}")
                    
                    # Extract printable strings (minimum 6 chars)
                    strings_found = []
                    current_string = b''
                    for byte in so_data:
                        if 32 <= byte < 127:  # Printable ASCII
                            current_string += bytes([byte])
                        else:
                            if len(current_string) >= 6:
                                try:
                                    decoded = current_string.decode('ascii')
                                    strings_found.append(decoded)
                                except:
                                    pass
                            current_string = b''
                    
                    # Filter interesting strings
                    interesting_strings = []
                    secret_patterns_found = []
                    
                    for s in strings_found:
                        s_lower = s.lower()
                        
                        # Check for URLs
                        if s.startswith('http://') or s.startswith('https://'):
                            interesting_strings.append(s)
                            if s.startswith('http://'):
                                result["security_findings"].append({
                                    "library": lib_name,
                                    "finding": "HTTP URL in native code",
                                    "severity": "medium",
                                    "description": f"Insecure HTTP URL found: {s[:50]}..."
                                })
                        
                        # Check for potential secrets
                        if any(kw in s_lower for kw in ['api_key', 'apikey', 'secret', 'password', 'token', 'private']):
                            if len(s) > 10 and '=' in s or ':' in s:
                                secret_patterns_found.append({
                                    "type": "potential_secret",
                                    "value": s[:50] + "..." if len(s) > 50 else s,
                                    "context": "native_string"
                                })
                        
                        # Firebase/Cloud configs
                        if 'firebase' in s_lower or 'google-services' in s_lower:
                            interesting_strings.append(s)
                        
                        # Encryption keys (base64-like long strings)
                        if len(s) >= 32 and s.replace('+', '').replace('/', '').replace('=', '').isalnum():
                            if not any(c in s for c in [' ', '\n', '\t']):
                                secret_patterns_found.append({
                                    "type": "potential_key",
                                    "value": s[:20] + "..." if len(s) > 20 else s,
                                    "length": len(s)
                                })
                    
                    lib_info["strings"] = interesting_strings[:50]
                    lib_info["hardcoded_secrets"] = secret_patterns_found[:10]
                    
                    if secret_patterns_found:
                        risk_score += len(secret_patterns_found) * 5
                        result["security_findings"].append({
                            "library": lib_name,
                            "finding": f"{len(secret_patterns_found)} potential secrets in native code",
                            "severity": "high",
                            "description": "Hardcoded secrets found in native library"
                        })
                    
                    # Check for anti-debugging patterns
                    for pattern, description in ANTI_DEBUG_PATTERNS:
                        if pattern in so_data:
                            lib_info["anti_debug_detected"] = True
                            lib_info["anti_debug_techniques"].append(description)
                    
                    if lib_info["anti_debug_detected"]:
                        risk_score += 10
                        result["security_findings"].append({
                            "library": lib_name,
                            "finding": "Anti-debugging techniques detected",
                            "severity": "info",
                            "description": f"Techniques: {', '.join(lib_info['anti_debug_techniques'][:5])}"
                        })
                    
                    # Check for crypto patterns
                    for pattern, description in NATIVE_CRYPTO_PATTERNS:
                        if pattern in so_data:
                            lib_info["crypto_functions"].append(description)
                    
                    # Check for suspicious patterns
                    suspicious_checks = [
                        (b'system(', 'system() call - command execution'),
                        (b'exec', 'exec() family - process execution'),
                        (b'popen', 'popen() - command execution'),
                        (b'dlopen', 'dlopen() - dynamic loading'),
                        (b'mprotect', 'mprotect() - memory protection changes'),
                        (b'mmap', 'mmap() - memory mapping'),
                        (b'fork', 'fork() - process creation'),
                        (b'socket', 'socket() - network operations'),
                        (b'connect', 'connect() - network connections'),
                        (b'send', 'send() - network transmission'),
                        (b'recv', 'recv() - network reception'),
                    ]
                    
                    for pattern, description in suspicious_checks:
                        if pattern in so_data:
                            lib_info["suspicious_patterns"].append({
                                "pattern": pattern.decode('ascii', errors='ignore'),
                                "description": description
                            })
                    
                    result["libraries"].append(lib_info)
                    
                except Exception as lib_error:
                    logger.debug(f"Error analyzing {so_path}: {lib_error}")
                    result["libraries"].append({
                        "name": lib_name,
                        "path": so_path,
                        "error": str(lib_error)
                    })
            
            result["total_libraries"] = len(native_files)
            result["total_jni_functions"] = total_jni
            result["total_exported_functions"] = total_exported
            result["architectures"] = list(architectures_found)
            
            # Determine overall risk
            if risk_score >= 50:
                result["overall_native_risk"] = "critical"
            elif risk_score >= 30:
                result["overall_native_risk"] = "high"
            elif risk_score >= 15:
                result["overall_native_risk"] = "medium"
            else:
                result["overall_native_risk"] = "low"
            
    except Exception as e:
        logger.error(f"Native library analysis failed: {e}")
        result["error"] = str(e)
    
    return result


# ============================================================================
# Hardening Score Calculator
# ============================================================================

def calculate_hardening_score(
    package_name: str,
    permissions: List[ApkPermission],
    components: List[ApkComponent],
    debuggable: bool,
    allow_backup: bool,
    min_sdk: Optional[int],
    target_sdk: Optional[int],
    certificate: Optional[ApkCertificate],
    strings: List[ExtractedString],
    urls: List[str],
    native_analysis: Optional[Dict[str, Any]],
    dex_analysis: Optional[Dict[str, Any]],
    network_config_analysis: Optional[Dict[str, Any]],
    security_issues: List[Dict[str, Any]],
    dynamic_analysis: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Calculate a comprehensive security hardening score for the APK.
    
    Categories:
    1. Code Protection (25%) - Obfuscation, anti-tampering, native protections
    2. Network Security (20%) - SSL pinning, cleartext, HTTPS usage
    3. Data Storage (20%) - Backup, encryption, secure storage
    4. Authentication & Crypto (15%) - Crypto strength, auth mechanisms
    5. Platform Security (20%) - Permissions, SDK versions, exported components
    """
    categories = []
    
    # ========================================================================
    # 1. CODE PROTECTION (25% weight)
    # ========================================================================
    code_score = 100
    code_findings = []
    code_recommendations = []
    
    # Check debuggable
    if debuggable:
        code_score -= 40
        code_findings.append({
            "issue": "App is debuggable",
            "severity": "critical",
            "impact": -40
        })
        code_recommendations.append("Set android:debuggable=\"false\" in release builds")
    
    # Check for obfuscation (look for short class names in dex analysis)
    obfuscation_detected = False
    if dex_analysis:
        class_names = [c.get("name", "") for c in dex_analysis.get("classes", [])]
        short_names = [n for n in class_names if len(n.split('.')[-1]) <= 2]
        if len(short_names) > len(class_names) * 0.3:
            obfuscation_detected = True
            code_score += 15
            code_findings.append({
                "issue": "Code obfuscation detected (ProGuard/R8)",
                "severity": "positive",
                "impact": 15
            })
    
    if not obfuscation_detected:
        code_score -= 15
        code_findings.append({
            "issue": "No code obfuscation detected",
            "severity": "medium",
            "impact": -15
        })
        code_recommendations.append("Enable ProGuard/R8 obfuscation for release builds")
    
    # Check for native protections
    if native_analysis:
        if native_analysis.get("total_libraries", 0) > 0:
            has_anti_debug = any(
                lib.get("anti_debug_detected", False) 
                for lib in native_analysis.get("libraries", [])
            )
            if has_anti_debug:
                code_score += 10
                code_findings.append({
                    "issue": "Native anti-debugging detected",
                    "severity": "positive",
                    "impact": 10
                })
            
            # Check for stripped binaries
            stripped_count = sum(
                1 for lib in native_analysis.get("libraries", [])
                if lib.get("is_stripped", False)
            )
            if stripped_count == native_analysis["total_libraries"]:
                code_score += 5
                code_findings.append({
                    "issue": "All native libraries are stripped",
                    "severity": "positive",
                    "impact": 5
                })
    
    # Check for anti-tampering indicators
    if dex_analysis:
        anti_tamper_patterns = dex_analysis.get("anti_analysis_detected", [])
        if anti_tamper_patterns:
            code_score += 10
            code_findings.append({
                "issue": f"Anti-analysis techniques detected ({len(anti_tamper_patterns)})",
                "severity": "positive",
                "impact": 10
            })
    
    # Check for root detection
    if dynamic_analysis and dynamic_analysis.get("root_detection_detected"):
        code_score += 10
        code_findings.append({
            "issue": "Root detection implemented",
            "severity": "positive",
            "impact": 10
        })
    else:
        code_score -= 10
        code_recommendations.append("Implement root detection for sensitive apps")
    
    code_score = max(0, min(100, code_score))
    categories.append({
        "name": "Code Protection",
        "score": code_score,
        "max_score": 100,
        "weight": 0.25,
        "findings": code_findings,
        "recommendations": code_recommendations,
        "icon": "🛡️"
    })
    
    # ========================================================================
    # 2. NETWORK SECURITY (20% weight)
    # ========================================================================
    network_score = 100
    network_findings = []
    network_recommendations = []
    
    # Check for HTTP URLs
    http_urls = [u for u in urls if u.startswith('http://') and 'localhost' not in u]
    if http_urls:
        penalty = min(30, len(http_urls) * 5)
        network_score -= penalty
        network_findings.append({
            "issue": f"{len(http_urls)} insecure HTTP URLs found",
            "severity": "high",
            "impact": -penalty
        })
        network_recommendations.append("Replace all HTTP URLs with HTTPS")
    
    # Check for SSL pinning
    if dynamic_analysis and dynamic_analysis.get("ssl_pinning_detected"):
        network_score += 20
        network_findings.append({
            "issue": "SSL certificate pinning detected",
            "severity": "positive",
            "impact": 20
        })
    else:
        network_score -= 20
        network_findings.append({
            "issue": "No SSL pinning detected",
            "severity": "medium",
            "impact": -20
        })
        network_recommendations.append("Implement certificate pinning for API connections")
    
    # Check network security config
    if network_config_analysis:
        if network_config_analysis.get("cleartext_permitted"):
            network_score -= 20
            network_findings.append({
                "issue": "Cleartext traffic permitted",
                "severity": "high",
                "impact": -20
            })
            network_recommendations.append("Disable cleartext traffic in network_security_config.xml")
        
        if network_config_analysis.get("certificate_pins"):
            network_score += 10
            network_findings.append({
                "issue": "Certificate pins configured in network security config",
                "severity": "positive",
                "impact": 10
            })
    
    network_score = max(0, min(100, network_score))
    categories.append({
        "name": "Network Security",
        "score": network_score,
        "max_score": 100,
        "weight": 0.20,
        "findings": network_findings,
        "recommendations": network_recommendations,
        "icon": "🌐"
    })
    
    # ========================================================================
    # 3. DATA STORAGE (20% weight)
    # ========================================================================
    storage_score = 100
    storage_findings = []
    storage_recommendations = []
    
    # Check backup setting
    if allow_backup:
        storage_score -= 20
        storage_findings.append({
            "issue": "App data backup enabled",
            "severity": "medium",
            "impact": -20
        })
        storage_recommendations.append("Set android:allowBackup=\"false\" or implement BackupAgent")
    else:
        storage_score += 10
        storage_findings.append({
            "issue": "App data backup disabled",
            "severity": "positive",
            "impact": 10
        })
    
    # Check for hardcoded secrets
    secrets = [s for s in strings if s.category in ['api_key', 'password', 'private_key', 'jwt']]
    if secrets:
        penalty = min(40, len(secrets) * 10)
        storage_score -= penalty
        storage_findings.append({
            "issue": f"{len(secrets)} hardcoded secrets found",
            "severity": "critical",
            "impact": -penalty
        })
        storage_recommendations.append("Remove hardcoded secrets, use Android Keystore or secure storage")
    
    # Check for encrypted storage indicators
    encrypted_storage_patterns = ['EncryptedSharedPreferences', 'AndroidKeyStore', 'KeyStore']
    has_encrypted_storage = any(
        any(p in s.value for p in encrypted_storage_patterns)
        for s in strings
    )
    if has_encrypted_storage:
        storage_score += 15
        storage_findings.append({
            "issue": "Encrypted storage mechanisms detected",
            "severity": "positive",
            "impact": 15
        })
    
    # Check native secrets
    if native_analysis:
        native_secrets = sum(
            len(lib.get("hardcoded_secrets", []))
            for lib in native_analysis.get("libraries", [])
        )
        if native_secrets:
            penalty = min(25, native_secrets * 5)
            storage_score -= penalty
            storage_findings.append({
                "issue": f"{native_secrets} potential secrets in native code",
                "severity": "high",
                "impact": -penalty
            })
            storage_recommendations.append("Review and remove secrets from native libraries")
    
    storage_score = max(0, min(100, storage_score))
    categories.append({
        "name": "Data Storage",
        "score": storage_score,
        "max_score": 100,
        "weight": 0.20,
        "findings": storage_findings,
        "recommendations": storage_recommendations,
        "icon": "💾"
    })
    
    # ========================================================================
    # 4. AUTHENTICATION & CRYPTO (15% weight)
    # ========================================================================
    crypto_score = 100
    crypto_findings = []
    crypto_recommendations = []
    
    # Check for weak crypto
    weak_crypto_patterns = ['DES', 'RC4', 'MD5', 'SHA1', 'ECB']
    weak_crypto_found = []
    for s in strings:
        for pattern in weak_crypto_patterns:
            if pattern in s.value and pattern not in weak_crypto_found:
                weak_crypto_found.append(pattern)
    
    if weak_crypto_found:
        penalty = len(weak_crypto_found) * 10
        crypto_score -= penalty
        crypto_findings.append({
            "issue": f"Weak cryptographic algorithms: {', '.join(weak_crypto_found)}",
            "severity": "high",
            "impact": -penalty
        })
        crypto_recommendations.append("Use AES-256-GCM for encryption, SHA-256+ for hashing")
    
    # Check for strong crypto indicators
    strong_crypto = ['AES-256', 'ChaCha20', 'curve25519', 'PBKDF2', 'Argon2']
    strong_found = [p for p in strong_crypto if any(p.lower() in s.value.lower() for s in strings)]
    if strong_found:
        crypto_score += 15
        crypto_findings.append({
            "issue": f"Strong cryptography detected: {', '.join(strong_found)}",
            "severity": "positive",
            "impact": 15
        })
    
    # Check for biometric auth
    biometric_patterns = ['BiometricPrompt', 'FingerprintManager', 'biometric']
    has_biometric = any(any(p in s.value for p in biometric_patterns) for s in strings)
    if has_biometric:
        crypto_score += 10
        crypto_findings.append({
            "issue": "Biometric authentication supported",
            "severity": "positive",
            "impact": 10
        })
    
    crypto_score = max(0, min(100, crypto_score))
    categories.append({
        "name": "Authentication & Crypto",
        "score": crypto_score,
        "max_score": 100,
        "weight": 0.15,
        "findings": crypto_findings,
        "recommendations": crypto_recommendations,
        "icon": "🔐"
    })
    
    # ========================================================================
    # 5. PLATFORM SECURITY (20% weight)
    # ========================================================================
    platform_score = 100
    platform_findings = []
    platform_recommendations = []
    
    # Check target SDK
    if target_sdk:
        if target_sdk < 28:  # Android 9
            platform_score -= 30
            platform_findings.append({
                "issue": f"Outdated target SDK: {target_sdk} (< Android 9)",
                "severity": "high",
                "impact": -30
            })
            platform_recommendations.append(f"Update targetSdkVersion to at least 33 (Android 13)")
        elif target_sdk < 30:  # Android 11
            platform_score -= 15
            platform_findings.append({
                "issue": f"Old target SDK: {target_sdk} (< Android 11)",
                "severity": "medium",
                "impact": -15
            })
            platform_recommendations.append("Consider updating targetSdkVersion to 33+")
        else:
            platform_score += 10
            platform_findings.append({
                "issue": f"Modern target SDK: {target_sdk}",
                "severity": "positive",
                "impact": 10
            })
    
    # Check dangerous permissions
    dangerous_perms = [p for p in permissions if p.is_dangerous]
    if len(dangerous_perms) > 5:
        penalty = min(25, (len(dangerous_perms) - 5) * 3)
        platform_score -= penalty
        platform_findings.append({
            "issue": f"Excessive dangerous permissions: {len(dangerous_perms)}",
            "severity": "medium",
            "impact": -penalty
        })
        platform_recommendations.append("Review and minimize dangerous permission requests")
    
    # Check exported components
    exported_unprotected = [
        c for c in components 
        if c.is_exported and c.component_type != 'activity'
    ]
    if exported_unprotected:
        penalty = min(20, len(exported_unprotected) * 4)
        platform_score -= penalty
        platform_findings.append({
            "issue": f"{len(exported_unprotected)} exported components without protection",
            "severity": "medium",
            "impact": -penalty
        })
        platform_recommendations.append("Add permission requirements to exported components")
    
    # Check certificate validity
    if certificate:
        from datetime import datetime
        try:
            not_after = datetime.strptime(certificate.valid_until, "%Y-%m-%d %H:%M:%S")
            if not_after < datetime.now():
                platform_score -= 20
                platform_findings.append({
                    "issue": "Certificate has expired",
                    "severity": "high",
                    "impact": -20
                })
            elif (not_after - datetime.now()).days < 365:
                platform_score -= 5
                platform_findings.append({
                    "issue": "Certificate expires within 1 year",
                    "severity": "low",
                    "impact": -5
                })
        except:
            pass
    
    platform_score = max(0, min(100, platform_score))
    categories.append({
        "name": "Platform Security",
        "score": platform_score,
        "max_score": 100,
        "weight": 0.20,
        "findings": platform_findings,
        "recommendations": platform_recommendations,
        "icon": "📱"
    })
    
    # ========================================================================
    # Calculate Overall Score
    # ========================================================================
    overall_score = sum(cat["score"] * cat["weight"] for cat in categories)
    overall_score = int(overall_score)
    
    # Determine grade and risk level
    if overall_score >= 90:
        grade = "A"
        risk_level = "Low"
    elif overall_score >= 80:
        grade = "B"
        risk_level = "Low"
    elif overall_score >= 70:
        grade = "C"
        risk_level = "Medium"
    elif overall_score >= 60:
        grade = "D"
        risk_level = "Medium"
    elif overall_score >= 50:
        grade = "D"
        risk_level = "High"
    else:
        grade = "F"
        risk_level = "Critical"
    
    # Compile top risks and recommendations
    all_findings = []
    for cat in categories:
        for f in cat["findings"]:
            if f.get("severity") in ["critical", "high"]:
                all_findings.append({
                    "category": cat["name"],
                    **f
                })
    
    top_risks = [f["issue"] for f in sorted(all_findings, key=lambda x: x.get("impact", 0))[:5]]
    
    all_recommendations = []
    for cat in categories:
        all_recommendations.extend(cat["recommendations"])
    top_recommendations = all_recommendations[:7]
    
    # Generate summary
    if grade in ["A", "B"]:
        summary = f"This APK demonstrates strong security practices with a score of {overall_score}/100. "
        summary += "It implements multiple protection mechanisms and follows security best practices."
    elif grade == "C":
        summary = f"This APK has moderate security with a score of {overall_score}/100. "
        summary += "While some protections are in place, there are areas that need improvement."
    elif grade == "D":
        summary = f"This APK has weak security with a score of {overall_score}/100. "
        summary += "Significant security improvements are recommended before production deployment."
    else:
        summary = f"This APK has critical security issues with a score of {overall_score}/100. "
        summary += "Immediate security remediation is required. The app is vulnerable to common attacks."
    
    # Calculate sub-scores
    attack_surface_score = int((platform_score * 0.5 + network_score * 0.3 + storage_score * 0.2))
    protection_score = int((code_score * 0.6 + crypto_score * 0.4))
    data_security_score = int((storage_score * 0.5 + crypto_score * 0.3 + network_score * 0.2))
    
    return {
        "overall_score": overall_score,
        "grade": grade,
        "risk_level": risk_level,
        "categories": categories,
        "attack_surface_score": attack_surface_score,
        "protection_score": protection_score,
        "data_security_score": data_security_score,
        "summary": summary,
        "top_risks": top_risks,
        "top_recommendations": top_recommendations,
    }


def analyze_smali_bytecode(apk, file_path: Path, 
                           target_classes: Optional[List[str]] = None,
                           max_methods: int = 100,
                           max_instructions_per_method: int = 500) -> Dict[str, Any]:
    """
    Decompile DEX bytecode to Smali format for reverse engineering analysis.
    
    Args:
        apk: androguard APK object
        file_path: Path to APK file
        target_classes: Optional list of specific class names to decompile (e.g., ['com.example.MainActivity'])
        max_methods: Maximum number of methods to decompile
        max_instructions_per_method: Maximum instructions per method to prevent huge outputs
    
    Returns:
        Dict containing:
            - decompiled_methods: List of SmaliMethodCode objects as dicts
            - class_smali: Dict mapping class names to full smali class definitions
            - statistics: Counts and metadata
            - interesting_methods: Methods containing suspicious patterns
    """
    result = {
        "decompiled_methods": [],
        "class_smali": {},
        "statistics": {
            "total_methods_analyzed": 0,
            "total_instructions": 0,
            "native_methods": 0,
            "abstract_methods": 0,
            "classes_analyzed": 0,
        },
        "interesting_methods": [],
        "search_index": [],  # For frontend search functionality
    }
    
    if not ANDROGUARD_AVAILABLE:
        result["error"] = "androguard not available"
        return result
    
    try:
        from androguard.core.dex import DEX
        
        methods_count = 0
        
        with zipfile.ZipFile(file_path, 'r') as zf:
            dex_files = [n for n in zf.namelist() if n.endswith('.dex')]
            
            for dex_name in dex_files:
                if methods_count >= max_methods:
                    break
                    
                try:
                    dex_data = zf.read(dex_name)
                    dex = DEX(dex_data)
                    
                    for cls in dex.get_classes():
                        if methods_count >= max_methods:
                            break
                        
                        class_name = cls.get_name()
                        readable_class = class_name.replace('/', '.').strip('L;')
                        
                        # If target classes specified, only process those
                        if target_classes:
                            if not any(t in readable_class for t in target_classes):
                                continue
                        
                        result["statistics"]["classes_analyzed"] += 1
                        
                        # Build class-level smali
                        class_smali_lines = []
                        class_smali_lines.append(f".class {_get_access_flags_string(cls)} {class_name}")
                        
                        superclass = cls.get_superclassname()
                        if superclass:
                            class_smali_lines.append(f".super {superclass}")
                        
                        # Add interfaces
                        if hasattr(cls, 'get_interfaces'):
                            for iface in cls.get_interfaces():
                                class_smali_lines.append(f".implements {iface}")
                        
                        class_smali_lines.append("")
                        
                        # Process fields
                        if hasattr(cls, 'get_fields'):
                            for field in cls.get_fields():
                                field_line = _format_smali_field(field)
                                if field_line:
                                    class_smali_lines.append(field_line)
                        
                        class_smali_lines.append("")
                        
                        # Process methods
                        for method in cls.get_methods():
                            if methods_count >= max_methods:
                                break
                            
                            method_name = method.get_name()
                            method_desc = method.get_descriptor()
                            full_method = f"{readable_class}.{method_name}"
                            
                            # Get method details
                            access_flags = _get_method_access_flags(method)
                            return_type, param_types = _parse_method_descriptor(method_desc)
                            
                            # Check for special methods
                            is_native = 'native' in access_flags.lower()
                            is_abstract = 'abstract' in access_flags.lower()
                            
                            if is_native:
                                result["statistics"]["native_methods"] += 1
                            if is_abstract:
                                result["statistics"]["abstract_methods"] += 1
                            
                            # Get bytecode instructions
                            instructions = []
                            registers_count = 0
                            has_try_catch = False
                            
                            code = method.get_code()
                            if code and not is_native and not is_abstract:
                                registers_count = code.get_registers_size() if hasattr(code, 'get_registers_size') else 0
                                
                                # Get bytecode
                                bc = code.get_bc() if hasattr(code, 'get_bc') else None
                                if bc:
                                    instruction_count = 0
                                    for ins in bc.get_instructions():
                                        if instruction_count >= max_instructions_per_method:
                                            instructions.append(f"    # ... truncated ({instruction_count}+ instructions)")
                                            break
                                        
                                        # Format instruction
                                        ins_name = ins.get_name()
                                        ins_output = ins.get_output() if hasattr(ins, 'get_output') else ""
                                        
                                        if ins_output:
                                            instructions.append(f"    {ins_name} {ins_output}")
                                        else:
                                            instructions.append(f"    {ins_name}")
                                        
                                        instruction_count += 1
                                        result["statistics"]["total_instructions"] += 1
                                
                                # Check for try-catch blocks
                                if hasattr(code, 'get_tries_size') and code.get_tries_size() > 0:
                                    has_try_catch = True
                            
                            # Build method smali
                            method_signature = f"{method_name}{method_desc}"
                            method_smali = _build_smali_method(
                                access_flags, method_name, method_desc,
                                registers_count, instructions, has_try_catch
                            )
                            class_smali_lines.extend(method_smali)
                            class_smali_lines.append("")
                            
                            # Create method record
                            method_record = {
                                "class_name": readable_class,
                                "method_name": method_name,
                                "method_signature": method_signature,
                                "access_flags": access_flags,
                                "return_type": return_type,
                                "parameters": param_types,
                                "registers_count": registers_count,
                                "instructions": instructions,
                                "instruction_count": len(instructions),
                                "has_try_catch": has_try_catch,
                                "is_native": is_native,
                                "is_abstract": is_abstract,
                            }
                            result["decompiled_methods"].append(method_record)
                            
                            # Build search index entry
                            result["search_index"].append({
                                "class": readable_class,
                                "method": method_name,
                                "signature": method_signature,
                                "preview": instructions[0] if instructions else "(empty)",
                            })
                            
                            # Check for interesting patterns
                            instructions_text = ' '.join(instructions).lower()
                            interesting_patterns = [
                                ("invoke-virtual.*crypto", "Cryptographic operation"),
                                ("invoke-static.*cipher", "Cipher usage"),
                                ("invoke-virtual.*reflect", "Reflection"),
                                ("invoke-static.*dexclassloader", "Dynamic DEX loading"),
                                ("invoke-virtual.*runtime.*exec", "Runtime execution"),
                                ("invoke-static.*base64", "Base64 encoding"),
                                ("invoke-virtual.*ssl", "SSL/TLS operation"),
                                ("invoke-static.*processbuilder", "Process execution"),
                                ("const-string.*http", "HTTP URL"),
                                ("const-string.*password", "Password reference"),
                                ("const-string.*api.?key", "API key reference"),
                            ]
                            
                            for pattern, description in interesting_patterns:
                                if re.search(pattern, instructions_text):
                                    result["interesting_methods"].append({
                                        "class": readable_class,
                                        "method": method_name,
                                        "pattern": description,
                                        "preview": instructions[:5] if instructions else [],
                                    })
                                    break
                            
                            methods_count += 1
                            result["statistics"]["total_methods_analyzed"] += 1
                        
                        # Store class smali
                        class_smali_lines.append(".end class")
                        result["class_smali"][readable_class] = '\n'.join(class_smali_lines)
                        
                except Exception as e:
                    logger.warning(f"Failed to decompile DEX {dex_name}: {e}")
                    
    except Exception as e:
        logger.error(f"Smali decompilation failed: {e}")
        result["error"] = str(e)
    
    # Limit interesting methods
    result["interesting_methods"] = result["interesting_methods"][:50]
    result["search_index"] = result["search_index"][:500]  # Limit search index
    
    return result


def _get_access_flags_string(cls_or_method) -> str:
    """Convert access flags to smali format string."""
    flags = []
    
    try:
        access = 0
        if hasattr(cls_or_method, 'get_access_flags'):
            access = cls_or_method.get_access_flags()
        elif hasattr(cls_or_method, 'get_access_flags_string'):
            return cls_or_method.get_access_flags_string()
        
        if access & 0x0001: flags.append("public")
        if access & 0x0002: flags.append("private")
        if access & 0x0004: flags.append("protected")
        if access & 0x0008: flags.append("static")
        if access & 0x0010: flags.append("final")
        if access & 0x0020: flags.append("synchronized")
        if access & 0x0040: flags.append("volatile")  # or bridge for methods
        if access & 0x0080: flags.append("transient")  # or varargs for methods
        if access & 0x0100: flags.append("native")
        if access & 0x0200: flags.append("interface")
        if access & 0x0400: flags.append("abstract")
        if access & 0x0800: flags.append("strict")
        if access & 0x1000: flags.append("synthetic")
        if access & 0x2000: flags.append("annotation")
        if access & 0x4000: flags.append("enum")
        
    except Exception:
        pass
    
    return ' '.join(flags) if flags else "public"


def _get_method_access_flags(method) -> str:
    """Get method access flags as a string."""
    try:
        if hasattr(method, 'get_access_flags_string'):
            return method.get_access_flags_string()
        return _get_access_flags_string(method)
    except:
        return "public"


def _parse_method_descriptor(descriptor: str) -> tuple:
    """Parse method descriptor to get return type and parameter types."""
    try:
        # Format: (params)return
        # e.g., (Ljava/lang/String;I)V -> params: [String, int], return: void
        params_match = re.match(r'\(([^)]*)\)(.+)', descriptor)
        if not params_match:
            return "void", []
        
        params_str, return_str = params_match.groups()
        
        # Parse parameters
        params = []
        i = 0
        while i < len(params_str):
            param_type, consumed = _parse_type(params_str[i:])
            if param_type:
                params.append(param_type)
            i += consumed if consumed > 0 else 1
        
        # Parse return type
        return_type, _ = _parse_type(return_str)
        
        return return_type or "void", params
        
    except Exception:
        return "void", []


def _parse_type(type_str: str) -> tuple:
    """Parse a single type from a descriptor."""
    if not type_str:
        return None, 0
    
    type_map = {
        'V': ('void', 1),
        'Z': ('boolean', 1),
        'B': ('byte', 1),
        'S': ('short', 1),
        'C': ('char', 1),
        'I': ('int', 1),
        'J': ('long', 1),
        'F': ('float', 1),
        'D': ('double', 1),
    }
    
    if type_str[0] in type_map:
        return type_map[type_str[0]]
    
    if type_str[0] == '[':
        inner_type, consumed = _parse_type(type_str[1:])
        return f"{inner_type}[]", consumed + 1
    
    if type_str[0] == 'L':
        end = type_str.find(';')
        if end != -1:
            class_name = type_str[1:end].replace('/', '.')
            return class_name, end + 1
    
    return type_str, 1


def _format_smali_field(field) -> Optional[str]:
    """Format a field definition in smali syntax."""
    try:
        name = field.get_name()
        field_type = field.get_descriptor() if hasattr(field, 'get_descriptor') else "?"
        access = _get_access_flags_string(field)
        
        return f".field {access} {name}:{field_type}"
    except:
        return None


def _build_smali_method(access_flags: str, method_name: str, descriptor: str,
                        registers: int, instructions: List[str], has_try: bool) -> List[str]:
    """Build complete smali method definition."""
    lines = []
    lines.append(f".method {access_flags} {method_name}{descriptor}")
    
    if registers > 0:
        lines.append(f"    .registers {registers}")
    
    lines.append("")
    
    if instructions:
        lines.extend(instructions)
    else:
        lines.append("    # (empty or native method)")
    
    lines.append("")
    lines.append(".end method")
    
    return lines


def analyze_dex_classes_and_methods(apk, file_path: Path) -> Dict[str, Any]:
    """Analyze DEX files for classes, methods, and suspicious patterns."""
    result = {
        "total_classes": 0,
        "total_methods": 0,
        "suspicious_classes": [],
        "suspicious_methods": [],
        "detected_trackers": [],
        "class_hierarchy": [],
        "reflection_usage": [],
        "crypto_usage": [],
        "native_calls": [],
        "dynamic_loading": [],
        "anti_analysis_detected": [],
    }
    
    if not ANDROGUARD_AVAILABLE:
        return result
    
    try:
        from androguard.core.dex import DEX
        
        # Get all DEX files from APK
        with zipfile.ZipFile(file_path, 'r') as zf:
            dex_files = [n for n in zf.namelist() if n.endswith('.dex')]
            
            for dex_name in dex_files:
                try:
                    dex_data = zf.read(dex_name)
                    dex = DEX(dex_data)
                    
                    for cls in dex.get_classes():
                        class_name = cls.get_name()
                        result["total_classes"] += 1
                        
                        # Convert class name format
                        readable_name = class_name.replace('/', '.').strip('L;')
                        
                        # Check for known trackers
                        for tracker_pkg, tracker_name in KNOWN_TRACKERS.items():
                            if tracker_pkg in readable_name:
                                if tracker_name not in result["detected_trackers"]:
                                    result["detected_trackers"].append({
                                        "name": tracker_name,
                                        "package": tracker_pkg,
                                        "class": readable_name,
                                    })
                        
                        # Analyze methods
                        methods = list(cls.get_methods())
                        result["total_methods"] += len(methods)
                        
                        for method in methods[:100]:  # Limit per class
                            method_name = method.get_name()
                            full_method = f"{readable_name}.{method_name}"
                            
                            # Check for suspicious patterns
                            for category, patterns in SUSPICIOUS_DEX_PATTERNS.items():
                                for pattern in patterns:
                                    if pattern.lower() in full_method.lower():
                                        suspicious_entry = {
                                            "class": readable_name,
                                            "method": method_name,
                                            "category": category,
                                            "pattern": pattern,
                                        }
                                        
                                        if category == "reflection":
                                            result["reflection_usage"].append(suspicious_entry)
                                        elif category == "crypto":
                                            result["crypto_usage"].append(suspicious_entry)
                                        elif category == "native":
                                            result["native_calls"].append(suspicious_entry)
                                        elif category == "dynamic_loading":
                                            result["dynamic_loading"].append(suspicious_entry)
                                        elif category == "anti_analysis":
                                            result["anti_analysis_detected"].append(suspicious_entry)
                                        
                                        if suspicious_entry not in result["suspicious_methods"]:
                                            result["suspicious_methods"].append(suspicious_entry)
                        
                        # Build class hierarchy (limited)
                        if len(result["class_hierarchy"]) < 100:
                            superclass = cls.get_superclassname()
                            interfaces = list(cls.get_interfaces()) if hasattr(cls, 'get_interfaces') else []
                            
                            result["class_hierarchy"].append({
                                "name": readable_name,
                                "superclass": superclass.replace('/', '.').strip('L;') if superclass else None,
                                "interfaces": [i.replace('/', '.').strip('L;') for i in interfaces[:5]],
                                "methods_count": len(methods),
                            })
                            
                except Exception as e:
                    logger.warning(f"Failed to analyze DEX {dex_name}: {e}")
                    
    except Exception as e:
        logger.warning(f"DEX analysis failed: {e}")
    
    # Limit results to prevent huge responses
    result["suspicious_methods"] = result["suspicious_methods"][:50]
    result["reflection_usage"] = result["reflection_usage"][:20]
    result["crypto_usage"] = result["crypto_usage"][:20]
    result["native_calls"] = result["native_calls"][:20]
    result["dynamic_loading"] = result["dynamic_loading"][:10]
    result["anti_analysis_detected"] = result["anti_analysis_detected"][:10]
    result["class_hierarchy"] = result["class_hierarchy"][:100]
    
    return result


def analyze_apk_resources(apk, file_path: Path) -> Dict[str, Any]:
    """Analyze APK resources for strings, assets, and potential secrets."""
    result = {
        "string_resources": {},
        "string_count": 0,
        "asset_files": [],
        "raw_resources": [],
        "drawable_count": 0,
        "layout_count": 0,
        "potential_secrets": [],
        "interesting_assets": [],
        "database_files": [],
        "config_files": [],
    }
    
    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            for name in zf.namelist():
                # Count drawables and layouts
                if name.startswith('res/drawable'):
                    result["drawable_count"] += 1
                elif name.startswith('res/layout'):
                    result["layout_count"] += 1
                
                # Track raw resources
                elif name.startswith('res/raw/'):
                    result["raw_resources"].append(name)
                
                # Track assets
                elif name.startswith('assets/'):
                    asset_name = name[7:]  # Remove 'assets/' prefix
                    result["asset_files"].append(asset_name)
                    
                    # Identify interesting assets
                    lower_name = asset_name.lower()
                    if any(ext in lower_name for ext in ['.db', '.sqlite', '.sqlite3']):
                        result["database_files"].append(asset_name)
                    elif any(ext in lower_name for ext in ['.json', '.xml', '.yml', '.yaml', '.properties', '.conf', '.config']):
                        result["config_files"].append(asset_name)
                        
                        # Try to read config files for secrets
                        try:
                            if zf.getinfo(name).file_size < 100000:  # 100KB limit
                                content = zf.read(name).decode('utf-8', errors='ignore')
                                secrets_found = _scan_content_for_secrets(content, asset_name)
                                result["potential_secrets"].extend(secrets_found)
                        except:
                            pass
                    
                    # Check for embedded APKs or DEX
                    if lower_name.endswith('.apk') or lower_name.endswith('.dex'):
                        result["interesting_assets"].append({
                            "name": asset_name,
                            "type": "embedded_code",
                            "risk": "high",
                        })
        
        # Try to get string resources from androguard
        if ANDROGUARD_AVAILABLE:
            try:
                arsc = apk.get_android_resources()
                if arsc:
                    # Get string resources
                    packages = arsc.get_packages_names()
                    for pkg in packages:
                        try:
                            strings = arsc.get_strings_resources()
                            for locale, string_dict in strings.items():
                                if locale == 'DEFAULT' or locale == '':
                                    for str_name, str_value in list(string_dict.items())[:500]:
                                        result["string_resources"][str_name] = str_value
                                        result["string_count"] += 1
                                        
                                        # Check for secrets in string resources
                                        secrets = _scan_content_for_secrets(str_value, f"strings/{str_name}")
                                        result["potential_secrets"].extend(secrets)
                                    break
                        except Exception as e:
                            logger.debug(f"Error getting strings for package {pkg}: {e}")
            except Exception as e:
                logger.debug(f"Could not get ARSC resources: {e}")
                
    except Exception as e:
        logger.warning(f"Resource analysis failed: {e}")
    
    # Limit results
    result["potential_secrets"] = result["potential_secrets"][:30]
    result["asset_files"] = result["asset_files"][:100]
    result["raw_resources"] = result["raw_resources"][:50]
    
    return result


def _scan_content_for_secrets(content: str, source: str) -> List[Dict[str, Any]]:
    """Scan content for potential secrets."""
    secrets = []
    
    # Secret patterns to look for
    patterns = {
        "api_key": r'(?:api[_-]?key|apikey)["\']?\s*[:=]\s*["\']?([a-zA-Z0-9_\-]{16,})["\']?',
        "aws_key": r'(AKIA[A-Z0-9]{16})',
        "google_api": r'AIza[0-9A-Za-z_-]{35}',
        "firebase_url": r'https://[a-z0-9-]+\.firebaseio\.com',
        "private_key": r'-----BEGIN (?:RSA |EC |DSA )?PRIVATE KEY-----',
        "password_field": r'(?:password|passwd|secret)["\']?\s*[:=]\s*["\']?([^\s"\']{6,})["\']?',
        "bearer_token": r'[Bb]earer\s+[a-zA-Z0-9_\-\.]{20,}',
        "base64_key": r'(?:key|secret|token)["\']?\s*[:=]\s*["\']?([A-Za-z0-9+/]{32,}={0,2})["\']?',
    }
    
    for secret_type, pattern in patterns.items():
        matches = re.findall(pattern, content, re.IGNORECASE)
        for match in matches[:5]:  # Limit matches per type
            value = match if isinstance(match, str) else match[0]
            if len(value) > 6:  # Minimum length
                secrets.append({
                    "type": secret_type,
                    "source": source,
                    "value_preview": value[:20] + "..." if len(value) > 20 else value,
                    "severity": "high" if secret_type in ["private_key", "aws_key"] else "medium",
                })
    
    return secrets


def analyze_intent_filters(apk) -> Dict[str, Any]:
    """Analyze intent filters for deep links and attack surfaces."""
    result = {
        "deep_links": [],
        "browsable_activities": [],
        "exported_components": [],
        "uri_schemes": [],
        "data_handlers": [],
        "implicit_intents": [],
        "attack_surface_summary": {},
    }
    
    if not ANDROGUARD_AVAILABLE:
        return result
    
    try:
        manifest_xml = apk.get_android_manifest_xml()
        if manifest_xml is None:
            return result
        
        ns = '{http://schemas.android.com/apk/res/android}'
        
        # Process each component type
        for component_type in ['activity', 'service', 'receiver', 'provider']:
            components = manifest_xml.findall(f'.//{component_type}')
            
            for comp in components:
                comp_name = comp.get(f'{ns}name', 'unknown')
                is_exported = comp.get(f'{ns}exported', 'false') == 'true'
                
                # Get intent filters
                intent_filters = comp.findall('.//intent-filter')
                
                for intent_filter in intent_filters:
                    filter_info = {
                        "component": comp_name,
                        "component_type": component_type,
                        "actions": [],
                        "categories": [],
                        "data_schemes": [],
                        "data_hosts": [],
                        "data_paths": [],
                        "is_browsable": False,
                        "is_exported": is_exported,
                    }
                    
                    # Get actions
                    for action in intent_filter.findall('.//action'):
                        action_name = action.get(f'{ns}name')
                        if action_name:
                            filter_info["actions"].append(action_name)
                    
                    # Get categories
                    for category in intent_filter.findall('.//category'):
                        cat_name = category.get(f'{ns}name')
                        if cat_name:
                            filter_info["categories"].append(cat_name)
                            if 'BROWSABLE' in cat_name:
                                filter_info["is_browsable"] = True
                    
                    # Get data elements
                    for data in intent_filter.findall('.//data'):
                        scheme = data.get(f'{ns}scheme')
                        host = data.get(f'{ns}host')
                        path = data.get(f'{ns}path')
                        pathPrefix = data.get(f'{ns}pathPrefix')
                        pathPattern = data.get(f'{ns}pathPattern')
                        
                        if scheme:
                            filter_info["data_schemes"].append(scheme)
                            if scheme not in result["uri_schemes"]:
                                result["uri_schemes"].append(scheme)
                        if host:
                            filter_info["data_hosts"].append(host)
                        if path:
                            filter_info["data_paths"].append(path)
                        if pathPrefix:
                            filter_info["data_paths"].append(f"{pathPrefix}*")
                        if pathPattern:
                            filter_info["data_paths"].append(pathPattern)
                    
                    # Build deep links
                    if filter_info["data_schemes"] and filter_info["is_browsable"]:
                        for scheme in filter_info["data_schemes"]:
                            for host in filter_info["data_hosts"] or ['*']:
                                deep_link = f"{scheme}://{host}"
                                if filter_info["data_paths"]:
                                    for path in filter_info["data_paths"]:
                                        result["deep_links"].append({
                                            "url": f"{deep_link}{path}",
                                            "component": comp_name,
                                            "type": component_type,
                                        })
                                else:
                                    result["deep_links"].append({
                                        "url": deep_link,
                                        "component": comp_name,
                                        "type": component_type,
                                    })
                    
                    # Track browsable activities
                    if filter_info["is_browsable"] and component_type == "activity":
                        result["browsable_activities"].append({
                            "name": comp_name,
                            "schemes": filter_info["data_schemes"],
                            "hosts": filter_info["data_hosts"],
                        })
                    
                    # Track implicit intent handlers
                    if filter_info["actions"]:
                        result["implicit_intents"].append(filter_info)
                    
                    # Track exported components with intent filters
                    if is_exported or intent_filters:
                        result["exported_components"].append({
                            "name": comp_name,
                            "type": component_type,
                            "exported": is_exported,
                            "has_intent_filter": bool(intent_filters),
                            "actions": filter_info["actions"],
                        })
        
        # Build attack surface summary
        result["attack_surface_summary"] = {
            "total_deep_links": len(result["deep_links"]),
            "browsable_activities_count": len(result["browsable_activities"]),
            "custom_uri_schemes": [s for s in result["uri_schemes"] if s not in ['http', 'https']],
            "exported_activities": len([c for c in result["exported_components"] if c["type"] == "activity"]),
            "exported_services": len([c for c in result["exported_components"] if c["type"] == "service"]),
            "exported_receivers": len([c for c in result["exported_components"] if c["type"] == "receiver"]),
            "exported_providers": len([c for c in result["exported_components"] if c["type"] == "provider"]),
        }
        
    except Exception as e:
        logger.warning(f"Intent filter analysis failed: {e}")
    
    # Limit results
    result["deep_links"] = result["deep_links"][:50]
    result["implicit_intents"] = result["implicit_intents"][:50]
    result["exported_components"] = result["exported_components"][:50]
    
    return result


def analyze_network_security_config(apk, file_path: Path) -> Dict[str, Any]:
    """Parse and analyze network security configuration."""
    result = {
        "has_config": False,
        "cleartext_permitted": True,  # Default for older SDKs
        "cleartext_domains": [],
        "trust_anchors": [],
        "certificate_pins": [],
        "domain_configs": [],
        "security_issues": [],
        "config_xml": None,
    }
    
    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            # Look for network_security_config.xml in res/xml/
            config_paths = [
                'res/xml/network_security_config.xml',
                'res/xml/network-security-config.xml',
            ]
            
            config_content = None
            for config_path in config_paths:
                if config_path in zf.namelist():
                    try:
                        config_content = zf.read(config_path)
                        result["has_config"] = True
                        break
                    except:
                        pass
            
            if not config_content and ANDROGUARD_AVAILABLE:
                # Try to get from androguard's parsed resources
                try:
                    files = apk.get_files()
                    for f in files:
                        if 'network_security_config' in f.lower() or 'network-security-config' in f.lower():
                            config_content = apk.get_file(f)
                            result["has_config"] = True
                            break
                except:
                    pass
            
            if config_content:
                # Parse the XML
                # Note: In APK, this might be binary XML - try to decode
                try:
                    if ANDROGUARD_AVAILABLE:
                        from androguard.core.axml import AXMLPrinter
                        # Check if it's binary XML
                        if config_content[:4] == b'\x03\x00\x08\x00':
                            axml = AXMLPrinter(config_content)
                            xml_string = axml.get_xml()
                            result["config_xml"] = xml_string
                            root = ET.fromstring(xml_string)
                        else:
                            result["config_xml"] = config_content.decode('utf-8')
                            root = ET.fromstring(config_content)
                    else:
                        result["config_xml"] = config_content.decode('utf-8')
                        root = ET.fromstring(config_content)
                    
                    # Parse base-config
                    base_config = root.find('.//base-config')
                    if base_config is not None:
                        cleartext = base_config.get('cleartextTrafficPermitted', 'true')
                        result["cleartext_permitted"] = cleartext.lower() == 'true'
                        
                        # Get trust anchors
                        for anchor in base_config.findall('.//certificates'):
                            src = anchor.get('src', 'unknown')
                            result["trust_anchors"].append({
                                "source": src,
                                "scope": "base",
                            })
                    
                    # Parse domain-config
                    for domain_config in root.findall('.//domain-config'):
                        cleartext = domain_config.get('cleartextTrafficPermitted', 'false')
                        
                        domains = []
                        for domain in domain_config.findall('.//domain'):
                            include_subdomains = domain.get('includeSubdomains', 'false')
                            domain_name = domain.text or ''
                            domains.append({
                                "name": domain_name,
                                "include_subdomains": include_subdomains.lower() == 'true',
                            })
                            
                            if cleartext.lower() == 'true':
                                result["cleartext_domains"].append(domain_name)
                        
                        # Get pin-set
                        pin_set = domain_config.find('.//pin-set')
                        if pin_set is not None:
                            expiration = pin_set.get('expiration')
                            pins = []
                            for pin in pin_set.findall('.//pin'):
                                digest = pin.get('digest', 'SHA-256')
                                pin_value = pin.text or ''
                                pins.append({
                                    "digest": digest,
                                    "value": pin_value[:20] + "..." if len(pin_value) > 20 else pin_value,
                                })
                            
                            result["certificate_pins"].append({
                                "domains": [d["name"] for d in domains],
                                "expiration": expiration,
                                "pins": pins,
                            })
                        
                        result["domain_configs"].append({
                            "domains": domains,
                            "cleartext_permitted": cleartext.lower() == 'true',
                            "has_pins": pin_set is not None,
                        })
                    
                except Exception as e:
                    logger.warning(f"Failed to parse network security config XML: {e}")
        
        # Analyze for security issues
        if not result["has_config"]:
            result["security_issues"].append(
                "No network_security_config.xml found - cleartext traffic allowed by default on older SDKs"
            )
        
        if result["cleartext_permitted"]:
            result["security_issues"].append(
                "Base config allows cleartext (HTTP) traffic - vulnerable to MITM attacks"
            )
        
        if result["cleartext_domains"]:
            result["security_issues"].append(
                f"Cleartext traffic explicitly allowed for {len(result['cleartext_domains'])} domain(s)"
            )
        
        if not result["certificate_pins"]:
            result["security_issues"].append(
                "No certificate pinning configured - app trusts all valid certificates"
            )
        
        # Check for user-installed certificates trust
        for anchor in result["trust_anchors"]:
            if anchor["source"] == "user":
                result["security_issues"].append(
                    "App trusts user-installed certificates - can be bypassed with custom CA"
                )
                break
                
    except Exception as e:
        logger.warning(f"Network security config analysis failed: {e}")
    
    return result


def analyze_apk(file_path: Path) -> ApkAnalysisResult:
    """Analyze an Android APK file using androguard for comprehensive analysis."""
    
    # Use androguard if available for proper analysis
    if ANDROGUARD_AVAILABLE:
        return analyze_apk_with_androguard(file_path)
    else:
        return analyze_apk_basic(file_path)


def analyze_apk_with_androguard(file_path: Path) -> ApkAnalysisResult:
    """Analyze APK using androguard library for comprehensive analysis."""
    try:
        apk = APK(str(file_path))
        
        # Basic info
        package_name = apk.get_package() or "unknown"
        version_name = apk.get_androidversion_name()
        version_code = None
        try:
            vc = apk.get_androidversion_code()
            version_code = int(vc) if vc else None
        except:
            pass
        
        min_sdk = None
        target_sdk = None
        try:
            min_sdk = int(apk.get_min_sdk_version()) if apk.get_min_sdk_version() else None
            target_sdk = int(apk.get_target_sdk_version()) if apk.get_target_sdk_version() else None
        except:
            pass
        
        app_name = apk.get_app_name()
        
        # Permissions
        permissions = []
        for perm in apk.get_permissions():
            is_dangerous = perm in DANGEROUS_PERMISSIONS
            permissions.append(ApkPermission(
                name=perm,
                is_dangerous=is_dangerous,
                description=DANGEROUS_PERMISSIONS.get(perm),
            ))
        
        # Components
        components = []
        activities = apk.get_activities() or []
        services = apk.get_services() or []
        receivers = apk.get_receivers() or []
        providers = apk.get_providers() or []
        
        # Check for exported components
        for activity in activities:
            is_exported = apk.get_intent_filters("activity", activity) is not None
            intent_filters = []
            filters = apk.get_intent_filters("activity", activity)
            if filters:
                for f in filters:
                    if hasattr(f, 'get_action'):
                        intent_filters.extend(f.get_action() or [])
            components.append(ApkComponent(
                name=activity,
                component_type="activity",
                is_exported=is_exported or _is_main_activity(intent_filters),
                intent_filters=intent_filters[:5],
            ))
        
        for service in services:
            is_exported = apk.get_intent_filters("service", service) is not None
            components.append(ApkComponent(
                name=service,
                component_type="service",
                is_exported=is_exported,
                intent_filters=[],
            ))
        
        for receiver in receivers:
            is_exported = apk.get_intent_filters("receiver", receiver) is not None
            components.append(ApkComponent(
                name=receiver,
                component_type="receiver",
                is_exported=is_exported,
                intent_filters=[],
            ))
        
        for provider in providers:
            components.append(ApkComponent(
                name=provider,
                component_type="provider",
                is_exported=False,  # Need deeper analysis
                intent_filters=[],
            ))
        
        # Security flags from manifest
        debuggable = apk.get_effective_target_sdk_version() is not None and _check_debuggable(apk)
        allow_backup = _check_allow_backup(apk)
        
        # Extract certificate info
        certificate = extract_apk_certificate(apk)
        
        # Extract strings from DEX files
        strings = []
        urls = []
        
        with zipfile.ZipFile(file_path, 'r') as zf:
            for name in zf.namelist():
                if name.endswith('.dex'):
                    try:
                        dex_data = zf.read(name)
                        dex_strings = extract_strings(dex_data, min_length=6, max_strings=2000)
                        strings.extend(dex_strings)
                        for s in dex_strings:
                            if s.category == "url":
                                urls.append(s.value)
                    except Exception as e:
                        logger.warning(f"Failed to extract strings from {name}: {e}")
        
        # Detect secrets
        secrets = detect_secrets_in_strings(strings)
        
        # Find native libraries
        native_libs = apk.get_libraries() or []
        
        # Uses features
        uses_features = apk.get_features() or []
        
        # NEW: Run additional analyses
        # 1. DEX class/method analysis
        dex_analysis = analyze_dex_classes_and_methods(apk, file_path)
        
        # 2. Resource analysis
        resource_analysis = analyze_apk_resources(apk, file_path)
        
        # 3. Intent filter / deep link analysis
        intent_filter_analysis = analyze_intent_filters(apk)
        
        # 4. Network security config analysis
        network_config_analysis = analyze_network_security_config(apk, file_path)
        
        # 5. Smali/bytecode decompilation (limited to interesting classes)
        # Focus on main app code, not libraries
        target_classes = None
        if package_name and package_name != "unknown":
            # Target the app's main package and common interesting classes
            target_classes = [package_name.replace('.', '/')]
        smali_analysis = analyze_smali_bytecode(apk, file_path, target_classes=target_classes, max_methods=150)
        
        # Comprehensive security analysis
        security_issues = analyze_apk_security_comprehensive(
            apk=apk,
            package_name=package_name,
            permissions=permissions,
            components=components,
            certificate=certificate,
            strings=strings,
            urls=urls,
            min_sdk=min_sdk,
            target_sdk=target_sdk,
            debuggable=debuggable,
            allow_backup=allow_backup,
        )
        
        # Add security issues from new analyses
        if dex_analysis.get("anti_analysis_detected"):
            security_issues.append({
                "category": "Anti-Analysis",
                "severity": "medium",
                "description": f"App uses {len(dex_analysis['anti_analysis_detected'])} anti-analysis techniques",
                "details": dex_analysis["anti_analysis_detected"][:5],
                "recommendation": "App may be attempting to detect debugging/hooking tools.",
            })
        
        if dex_analysis.get("dynamic_loading"):
            security_issues.append({
                "category": "Dynamic Code Loading",
                "severity": "high",
                "description": f"App uses dynamic class loading ({len(dex_analysis['dynamic_loading'])} instances)",
                "details": dex_analysis["dynamic_loading"][:5],
                "recommendation": "Dynamic code loading can be used to hide malicious behavior.",
            })
        
        if resource_analysis.get("potential_secrets"):
            for secret in resource_analysis["potential_secrets"][:3]:
                security_issues.append({
                    "category": "M2: Secret in Resources",
                    "severity": secret.get("severity", "medium"),
                    "description": f"Potential {secret['type']} found in {secret['source']}",
                    "recommendation": "Remove secrets from APK resources. Use secure storage.",
                })
        
        if intent_filter_analysis.get("deep_links"):
            custom_schemes = intent_filter_analysis.get("attack_surface_summary", {}).get("custom_uri_schemes", [])
            if custom_schemes:
                security_issues.append({
                    "category": "Attack Surface: Deep Links",
                    "severity": "medium",
                    "description": f"App registers {len(custom_schemes)} custom URI scheme(s): {', '.join(custom_schemes[:5])}",
                    "details": intent_filter_analysis["deep_links"][:5],
                    "recommendation": "Validate all data received via deep links to prevent injection attacks.",
                })
        
        for issue in network_config_analysis.get("security_issues", []):
            security_issues.append({
                "category": "M3: Network Security",
                "severity": "medium",
                "description": issue,
                "recommendation": "Configure network_security_config.xml properly for secure communications.",
            })
        
        # 6. Generate Frida scripts for dynamic analysis
        dynamic_analysis = generate_frida_scripts(
            package_name=package_name,
            strings=strings,
            dex_analysis=dex_analysis,
            permissions=permissions,
            urls=urls,
            smali_analysis=smali_analysis
        )
        
        # 7. Analyze native libraries (.so files)
        native_analysis_result = analyze_native_libraries(file_path, native_libs)
        
        # 8. Calculate security hardening score
        hardening_score_result = calculate_hardening_score(
            package_name=package_name,
            permissions=permissions,
            components=components,
            certificate=certificate,
            strings=strings,
            urls=urls,
            min_sdk=min_sdk,
            target_sdk=target_sdk,
            debuggable=debuggable,
            allow_backup=allow_backup,
            dex_analysis=dex_analysis,
            network_config_analysis=network_config_analysis,
            native_analysis=native_analysis_result,
            security_issues=security_issues
        )
        
        # Add security issues from native analysis
        if native_analysis_result and native_analysis_result.get("risk_level") in ["high", "critical"]:
            native_issues = []
            for lib in native_analysis_result.get("libraries", []):
                if lib.get("has_anti_debug"):
                    native_issues.append(f"{lib['name']}: Anti-debugging detected")
                if lib.get("secrets_found"):
                    native_issues.append(f"{lib['name']}: {len(lib['secrets_found'])} potential secrets")
            
            if native_issues:
                security_issues.append({
                    "category": "Native Code Security",
                    "severity": native_analysis_result.get("risk_level", "medium"),
                    "description": f"Native library analysis found {len(native_issues)} security concerns",
                    "details": native_issues[:10],
                    "recommendation": "Review native code for hardcoded secrets and implement proper security measures.",
                })
        
        # 9. Data flow / taint analysis
        data_flow_result = None
        data_flow_dict = None
        try:
            data_flow_result = analyze_data_flow(file_path, package_name)
            data_flow_dict = dataflow_result_to_dict(data_flow_result)
            
            # Add data flow security issues
            if data_flow_result.critical_flows > 0:
                security_issues.append({
                    "category": "M2: Data Flow - Critical",
                    "severity": "critical",
                    "description": f"Found {data_flow_result.critical_flows} critical data flow paths (sensitive data to insecure sinks)",
                    "details": [
                        f"{p.source.source_type} → {p.sink.sink_type}: {p.description}"
                        for p in data_flow_result.data_flow_paths[:5]
                        if p.severity == "critical"
                    ],
                    "recommendation": "Review critical data flows. Ensure sensitive data is encrypted before transmission/storage.",
                })
            
            if data_flow_result.high_risk_flows > 0:
                security_issues.append({
                    "category": "M2: Data Flow - High Risk",
                    "severity": "high",
                    "description": f"Found {data_flow_result.high_risk_flows} high-risk data flow paths",
                    "details": [
                        f"{p.source.source_type} → {p.sink.sink_type}: {p.description}"
                        for p in data_flow_result.data_flow_paths[:5]
                        if p.severity == "high"
                    ],
                    "recommendation": "Validate and sanitize sensitive data before passing to external sinks.",
                })
            
            if data_flow_result.privacy_violations:
                security_issues.append({
                    "category": "Privacy Violation",
                    "severity": "high",
                    "description": f"Found {len(data_flow_result.privacy_violations)} potential privacy/GDPR violations",
                    "details": [v.get("description", "") for v in data_flow_result.privacy_violations[:5]],
                    "recommendation": "Review data handling practices. Ensure user consent for PII collection and transmission.",
                })
                
            logger.info(f"Data flow analysis complete: {data_flow_result.total_flows} flows found ({data_flow_result.critical_flows} critical)")
        except Exception as e:
            logger.warning(f"Data flow analysis failed: {e}")
            data_flow_dict = None
        
        return ApkAnalysisResult(
            filename=file_path.name,
            package_name=package_name,
            version_name=version_name,
            version_code=version_code,
            min_sdk=min_sdk,
            target_sdk=target_sdk,
            permissions=permissions,
            components=components,
            strings=strings[:300],
            secrets=secrets,
            urls=list(set(urls))[:100],
            native_libraries=native_libs,
            certificate=certificate,
            activities=list(activities),
            services=list(services),
            receivers=list(receivers),
            providers=list(providers),
            uses_features=list(uses_features),
            app_name=app_name,
            debuggable=debuggable,
            allow_backup=allow_backup,
            dex_analysis=dex_analysis,
            resource_analysis=resource_analysis,
            intent_filter_analysis=intent_filter_analysis,
            network_config_analysis=network_config_analysis,
            smali_analysis=smali_analysis,
            dynamic_analysis=dynamic_analysis,
            native_analysis=native_analysis_result,
            hardening_score=hardening_score_result,
            security_issues=security_issues,
            data_flow_analysis=data_flow_dict,
        )
        
    except Exception as e:
        logger.error(f"APK analysis with androguard failed: {e}")
        # Fallback to basic analysis
        return analyze_apk_basic(file_path)


def _is_main_activity(intent_filters: List[str]) -> bool:
    """Check if activity is the main launcher activity."""
    return "android.intent.action.MAIN" in intent_filters


def _check_debuggable(apk) -> bool:
    """Check if APK is debuggable."""
    try:
        # Try to get from manifest
        manifest_xml = apk.get_android_manifest_xml()
        if manifest_xml is not None:
            app_elem = manifest_xml.find('.//application')
            if app_elem is not None:
                debuggable = app_elem.get('{http://schemas.android.com/apk/res/android}debuggable')
                return debuggable == 'true'
    except:
        pass
    return False


def _check_allow_backup(apk) -> bool:
    """Check if APK allows backup (default is true)."""
    try:
        manifest_xml = apk.get_android_manifest_xml()
        if manifest_xml is not None:
            app_elem = manifest_xml.find('.//application')
            if app_elem is not None:
                allow_backup = app_elem.get('{http://schemas.android.com/apk/res/android}allowBackup')
                return allow_backup != 'false'  # Default is true
    except:
        pass
    return True  # Default


def extract_apk_certificate(apk) -> Optional[ApkCertificate]:
    """Extract certificate information from APK."""
    try:
        certs = apk.get_certificates()
        if not certs:
            return None
        
        cert = certs[0]  # Primary certificate
        
        # Get certificate details
        import hashlib
        from datetime import datetime
        
        # Get the raw certificate bytes
        cert_der = cert.public_bytes(encoding=None) if hasattr(cert, 'public_bytes') else None
        
        # Calculate fingerprints
        cert_bytes = cert.dump() if hasattr(cert, 'dump') else bytes(cert)
        sha256_fp = hashlib.sha256(cert_bytes).hexdigest().upper()
        sha1_fp = hashlib.sha1(cert_bytes).hexdigest().upper()
        md5_fp = hashlib.md5(cert_bytes).hexdigest().upper()
        
        # Format fingerprints with colons
        sha256_fp = ':'.join(sha256_fp[i:i+2] for i in range(0, len(sha256_fp), 2))
        sha1_fp = ':'.join(sha1_fp[i:i+2] for i in range(0, len(sha1_fp), 2))
        md5_fp = ':'.join(md5_fp[i:i+2] for i in range(0, len(md5_fp), 2))
        
        # Extract subject and issuer
        subject = str(cert.subject) if hasattr(cert, 'subject') else "Unknown"
        issuer = str(cert.issuer) if hasattr(cert, 'issuer') else "Unknown"
        serial = str(cert.serial_number) if hasattr(cert, 'serial_number') else "Unknown"
        
        # Validity dates
        valid_from = ""
        valid_until = ""
        is_expired = False
        
        if hasattr(cert, 'not_valid_before'):
            valid_from = str(cert.not_valid_before)
        if hasattr(cert, 'not_valid_after'):
            valid_until = str(cert.not_valid_after)
            try:
                expiry = cert.not_valid_after
                is_expired = datetime.now() > expiry if hasattr(expiry, '__gt__') else False
            except:
                pass
        
        # Check for debug certificate indicators
        is_debug = _is_debug_certificate(subject, issuer)
        
        # Check if self-signed
        is_self_signed = subject == issuer
        
        # Signature version
        sig_version = "v1"
        if apk.is_signed_v2():
            sig_version = "v2"
        if apk.is_signed_v3():
            sig_version = "v3"
        
        # Public key info
        pub_key_algo = None
        pub_key_bits = None
        if hasattr(cert, 'public_key'):
            pk = cert.public_key()
            if hasattr(pk, 'key_size'):
                pub_key_bits = pk.key_size
            pub_key_algo = type(pk).__name__.replace('_', ' ')
        
        return ApkCertificate(
            subject=subject,
            issuer=issuer,
            serial_number=serial,
            fingerprint_sha256=sha256_fp,
            fingerprint_sha1=sha1_fp,
            fingerprint_md5=md5_fp,
            valid_from=valid_from,
            valid_until=valid_until,
            is_debug_cert=is_debug,
            is_expired=is_expired,
            is_self_signed=is_self_signed,
            signature_version=sig_version,
            public_key_algorithm=pub_key_algo,
            public_key_bits=pub_key_bits,
        )
        
    except Exception as e:
        logger.warning(f"Failed to extract certificate: {e}")
        return None


def _is_debug_certificate(subject: str, issuer: str) -> bool:
    """Check if certificate appears to be a debug/development certificate."""
    debug_indicators = [
        "CN=Android Debug",
        "CN=Debug",
        "O=Android",
        "CN=unknown",
        "OU=Android",
    ]
    combined = (subject + issuer).lower()
    return any(ind.lower() in combined for ind in debug_indicators)


def analyze_apk_security_comprehensive(
    apk,
    package_name: str,
    permissions: List[ApkPermission],
    components: List[ApkComponent],
    certificate: Optional[ApkCertificate],
    strings: List[ExtractedString],
    urls: List[str],
    min_sdk: Optional[int],
    target_sdk: Optional[int],
    debuggable: bool,
    allow_backup: bool,
) -> List[Dict[str, Any]]:
    """Comprehensive security analysis based on OWASP Mobile Top 10."""
    issues = []
    
    # ========== M1: Improper Platform Usage ==========
    
    # Check for dangerous permissions
    dangerous_perms = [p for p in permissions if p.is_dangerous]
    if dangerous_perms:
        severity = "critical" if len(dangerous_perms) > 5 else "high"
        issues.append({
            "category": "M1: Dangerous Permissions",
            "severity": severity,
            "description": f"App requests {len(dangerous_perms)} dangerous permissions that could compromise user privacy",
            "details": [{"name": p.name, "description": p.description} for p in dangerous_perms],
            "recommendation": "Review each permission and remove unnecessary ones. Implement runtime permission requests.",
        })
    
    # Check for exported components without protection
    exported_unprotected = [c for c in components if c.is_exported and c.component_type != "activity"]
    if exported_unprotected:
        issues.append({
            "category": "M1: Exported Components",
            "severity": "high",
            "description": f"App has {len(exported_unprotected)} exported components (services/receivers/providers) accessible by other apps",
            "details": [{"name": c.name, "type": c.component_type} for c in exported_unprotected[:10]],
            "recommendation": "Add android:permission attribute or set android:exported=false for internal components.",
        })
    
    # ========== M2: Insecure Data Storage ==========
    
    # Check allowBackup
    if allow_backup:
        issues.append({
            "category": "M2: Insecure Backup",
            "severity": "medium",
            "description": "App allows data backup (android:allowBackup=true), which could expose sensitive data",
            "recommendation": "Set android:allowBackup=\"false\" or implement BackupAgent to control what gets backed up.",
        })
    
    # Check for hardcoded sensitive patterns
    sensitive_patterns = [s for s in strings if s.category in ["api_key", "password", "private_key"]]
    if sensitive_patterns:
        issues.append({
            "category": "M2: Hardcoded Secrets",
            "severity": "critical",
            "description": f"Found {len(sensitive_patterns)} potential hardcoded secrets in the APK",
            "details": [{"type": s.category, "preview": s.value[:30] + "..."} for s in sensitive_patterns[:5]],
            "recommendation": "Never hardcode secrets. Use Android Keystore or secure key management solutions.",
        })
    
    # Check for SharedPreferences patterns (potential insecure storage)
    shared_prefs = [s for s in strings if 'SharedPreferences' in s.value or 'getSharedPreferences' in s.value]
    if shared_prefs:
        issues.append({
            "category": "M2: SharedPreferences Usage",
            "severity": "low",
            "description": "App uses SharedPreferences - ensure sensitive data is encrypted",
            "recommendation": "Use EncryptedSharedPreferences from AndroidX Security library for sensitive data.",
        })
    
    # ========== M3: Insecure Communication ==========
    
    # Check for HTTP URLs
    http_urls = [u for u in urls if u.startswith('http://') and 'localhost' not in u and '127.0.0.1' not in u]
    if http_urls:
        issues.append({
            "category": "M3: Insecure Communication",
            "severity": "high",
            "description": f"App contains {len(http_urls)} insecure HTTP URLs instead of HTTPS",
            "details": http_urls[:10],
            "recommendation": "Use HTTPS for all network communication. Implement certificate pinning for sensitive connections.",
        })
    
    # Check for certificate pinning indicators
    pinning_patterns = ['certificatePinner', 'CertificatePinner', 'ssl_pinning', 'TrustManager']
    has_pinning = any(any(p in s.value for p in pinning_patterns) for s in strings)
    if not has_pinning and http_urls:
        issues.append({
            "category": "M3: No Certificate Pinning",
            "severity": "medium",
            "description": "No certificate pinning detected - app may be vulnerable to MITM attacks",
            "recommendation": "Implement certificate pinning using OkHttp CertificatePinner or Network Security Config.",
        })
    
    # ========== M4: Insecure Authentication ==========
    
    # Check for weak auth patterns
    weak_auth = [s for s in strings if any(w in s.value.lower() for w in ['password=', 'pwd=', 'basic auth', 'md5('])]
    if weak_auth:
        issues.append({
            "category": "M4: Weak Authentication",
            "severity": "high",
            "description": "Potential weak authentication patterns detected",
            "recommendation": "Use strong authentication mechanisms. Avoid storing passwords in plain text.",
        })
    
    # ========== M5: Insufficient Cryptography ==========
    
    # Check for weak crypto
    weak_crypto_patterns = ['DES', 'RC4', 'MD5', 'SHA1', 'ECB']
    weak_crypto = [s for s in strings if any(wc in s.value for wc in weak_crypto_patterns)]
    if weak_crypto:
        issues.append({
            "category": "M5: Weak Cryptography",
            "severity": "high",
            "description": "Potentially weak cryptographic algorithms detected (DES, RC4, MD5, SHA1, ECB mode)",
            "details": list(set([s.value[:50] for s in weak_crypto[:5]])),
            "recommendation": "Use AES-256-GCM for encryption, SHA-256+ for hashing, and avoid ECB mode.",
        })
    
    # ========== M6: Insecure Authorization ==========
    
    # Check for root detection bypass patterns
    root_patterns = ['isRooted', 'checkRoot', 'RootBeer', 'detectRoot', '/system/app/Superuser']
    has_root_detection = any(any(p in s.value for p in root_patterns) for s in strings)
    if not has_root_detection:
        issues.append({
            "category": "M6: No Root Detection",
            "severity": "low",
            "description": "No root detection mechanism found - app may run on compromised devices",
            "recommendation": "Implement root detection to protect sensitive functionality on rooted devices.",
        })
    
    # ========== M7: Client Code Quality ==========
    
    # Check for debugging enabled
    if debuggable:
        issues.append({
            "category": "M7: Debuggable App",
            "severity": "critical",
            "description": "App is debuggable - allows attackers to attach debuggers and inspect/modify runtime",
            "recommendation": "Set android:debuggable=\"false\" in release builds.",
        })
    
    # Check target SDK version
    if target_sdk and target_sdk < 30:  # Android 11
        issues.append({
            "category": "M7: Outdated Target SDK",
            "severity": "medium",
            "description": f"App targets SDK {target_sdk} - missing security improvements from newer Android versions",
            "recommendation": f"Update targetSdkVersion to at least 33 (Android 13) to benefit from security enhancements.",
        })
    
    if min_sdk and min_sdk < 23:  # Android 6.0
        issues.append({
            "category": "M7: Low Minimum SDK",
            "severity": "medium",
            "description": f"App supports devices with SDK {min_sdk} - these lack important security features",
            "recommendation": "Consider raising minSdkVersion to 23+ for runtime permissions and better security.",
        })
    
    # ========== M8: Code Tampering ==========
    
    # Check for anti-tampering
    tamper_patterns = ['SafetyNet', 'Play Integrity', 'signature verification', 'checkSignature']
    has_tampering_protection = any(any(p in s.value for p in tamper_patterns) for s in strings)
    if not has_tampering_protection:
        issues.append({
            "category": "M8: No Tampering Protection",
            "severity": "medium",
            "description": "No code tampering protection detected",
            "recommendation": "Implement Play Integrity API or SafetyNet Attestation to detect tampered apps.",
        })
    
    # ========== M9: Reverse Engineering ==========
    
    # Check for obfuscation (ProGuard/R8/DexGuard)
    obfuscation_indicators = _detect_obfuscation(strings)
    if not obfuscation_indicators['is_obfuscated']:
        issues.append({
            "category": "M9: No Obfuscation",
            "severity": "medium",
            "description": "App does not appear to be obfuscated - code can be easily reverse engineered",
            "recommendation": "Enable R8/ProGuard minification and obfuscation for release builds.",
        })
    
    # ========== M10: Extraneous Functionality ==========
    
    # Check for logging
    log_patterns = ['Log.d(', 'Log.v(', 'Log.i(', 'System.out.print', 'console.log']
    has_logging = any(any(p in s.value for p in log_patterns) for s in strings)
    if has_logging:
        issues.append({
            "category": "M10: Debug Logging",
            "severity": "low",
            "description": "Debug logging statements detected - may leak sensitive information",
            "recommendation": "Remove or disable debug logging in release builds using BuildConfig.DEBUG checks.",
        })
    
    # Check certificate issues
    if certificate:
        if certificate.is_debug_cert:
            issues.append({
                "category": "Certificate: Debug Signing",
                "severity": "critical",
                "description": "App is signed with a debug certificate - not suitable for production",
                "recommendation": "Sign the app with a proper release keystore before distribution.",
            })
        
        if certificate.is_expired:
            issues.append({
                "category": "Certificate: Expired",
                "severity": "high",
                "description": "App signing certificate has expired",
                "recommendation": "Re-sign the app with a valid certificate.",
            })
        
        if certificate.signature_version == "v1":
            issues.append({
                "category": "Certificate: Legacy Signature",
                "severity": "medium",
                "description": "App only uses v1 signature scheme - vulnerable to Janus vulnerability (CVE-2017-13156)",
                "recommendation": "Enable v2 and v3 signature schemes for better security.",
            })
        
        if certificate.public_key_bits and certificate.public_key_bits < 2048:
            issues.append({
                "category": "Certificate: Weak Key",
                "severity": "high",
                "description": f"Certificate uses weak {certificate.public_key_bits}-bit key",
                "recommendation": "Use at least 2048-bit RSA or 256-bit ECDSA keys.",
            })
    
    # Sort by severity
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    issues.sort(key=lambda x: severity_order.get(x["severity"], 4))
    
    return issues


def _detect_obfuscation(strings: List[ExtractedString]) -> Dict[str, Any]:
    """Detect if APK is obfuscated."""
    # Count single-letter class/method names
    single_letter_count = 0
    for s in strings:
        if len(s.value) == 1 and s.value.isalpha():
            single_letter_count += 1
    
    # Check for ProGuard/R8 patterns
    proguard_indicators = ['proguard', 'r8', '-keep', '-dontwarn']
    has_proguard = any(any(p in s.value.lower() for p in proguard_indicators) for s in strings)
    
    # Many single letter names = likely obfuscated
    is_obfuscated = single_letter_count > 50 or has_proguard
    
    return {
        "is_obfuscated": is_obfuscated,
        "single_letter_count": single_letter_count,
        "has_proguard_config": has_proguard,
    }


def analyze_apk_basic(file_path: Path) -> ApkAnalysisResult:
    """Basic APK analysis fallback when androguard is not available."""
    try:
        with zipfile.ZipFile(file_path, 'r') as apk:
            # Parse AndroidManifest.xml (basic parsing)
            manifest_data = extract_android_manifest(apk)
            
            # Extract strings from DEX files
            strings = []
            secrets = []
            urls = []
            
            for name in apk.namelist():
                if name.endswith('.dex'):
                    try:
                        dex_data = apk.read(name)
                        dex_strings = extract_strings(dex_data, min_length=6, max_strings=2000)
                        strings.extend(dex_strings)
                        
                        # Extract URLs
                        for s in dex_strings:
                            if s.category == "url":
                                urls.append(s.value)
                    except Exception as e:
                        logger.warning(f"Failed to extract strings from {name}: {e}")
            
            # Detect secrets
            secrets = detect_secrets_in_strings(strings)
            
            # Find native libraries
            native_libs = [n for n in apk.namelist() if n.startswith('lib/') and n.endswith('.so')]
            
            # Basic security analysis
            security_issues = analyze_apk_security(manifest_data, strings, urls)
            
            return ApkAnalysisResult(
                filename=file_path.name,
                package_name=manifest_data.get("package", "unknown"),
                version_name=manifest_data.get("version_name"),
                version_code=manifest_data.get("version_code"),
                min_sdk=manifest_data.get("min_sdk"),
                target_sdk=manifest_data.get("target_sdk"),
                permissions=manifest_data.get("permissions", []),
                components=manifest_data.get("components", []),
                strings=strings[:300],
                secrets=secrets,
                urls=list(set(urls))[:100],
                native_libraries=native_libs,
                certificate=None,  # Not available without androguard
                activities=[],
                services=[],
                receivers=[],
                providers=[],
                uses_features=[],
                app_name=None,
                debuggable=False,
                allow_backup=True,
                security_issues=security_issues,
            )
            
    except Exception as e:
        logger.error(f"APK analysis failed: {e}")
        return ApkAnalysisResult(
            filename=file_path.name,
            package_name="error",
            version_name=None,
            version_code=None,
            min_sdk=None,
            target_sdk=None,
            permissions=[],
            components=[],
            strings=[],
            secrets=[],
            urls=[],
            native_libraries=[],
            certificate=None,
            activities=[],
            services=[],
            receivers=[],
            providers=[],
            uses_features=[],
            app_name=None,
            debuggable=False,
            allow_backup=True,
            security_issues=[],
            error=str(e),
        )


def extract_android_manifest(apk: zipfile.ZipFile) -> Dict[str, Any]:
    """Extract and parse AndroidManifest.xml from APK."""
    result = {
        "package": "unknown",
        "permissions": [],
        "components": [],
    }
    
    try:
        # AndroidManifest.xml in APKs is in binary XML format
        # We need to decode it or use aapt/apktool
        # For now, extract what we can from strings
        
        manifest_binary = apk.read('AndroidManifest.xml')
        
        # Extract strings that look like permissions
        strings = extract_strings(manifest_binary, min_length=10)
        
        permissions = []
        for s in strings:
            if 'android.permission.' in s.value:
                perm_name = s.value
                if perm_name.startswith('android.permission.'):
                    is_dangerous = perm_name in DANGEROUS_PERMISSIONS
                    permissions.append(ApkPermission(
                        name=perm_name,
                        is_dangerous=is_dangerous,
                        description=DANGEROUS_PERMISSIONS.get(perm_name),
                    ))
        
        # Deduplicate permissions
        seen = set()
        unique_perms = []
        for p in permissions:
            if p.name not in seen:
                seen.add(p.name)
                unique_perms.append(p)
        
        result["permissions"] = unique_perms
        
        # Try to extract package name
        for s in strings:
            if '.' in s.value and not s.value.startswith('android.') and s.value.count('.') >= 2:
                # Looks like a package name
                if all(c.isalnum() or c in '._' for c in s.value):
                    result["package"] = s.value
                    break
        
    except Exception as e:
        logger.warning(f"Failed to parse AndroidManifest.xml: {e}")
    
    return result


def analyze_apk_security(manifest_data: Dict, strings: List[ExtractedString], urls: List[str]) -> List[Dict[str, Any]]:
    """Analyze APK for security issues."""
    issues = []
    
    # Check for dangerous permissions
    dangerous_perms = [p for p in manifest_data.get("permissions", []) if p.is_dangerous]
    if dangerous_perms:
        issues.append({
            "category": "Dangerous Permissions",
            "severity": "high",
            "description": f"App requests {len(dangerous_perms)} dangerous permissions",
            "details": [{"name": p.name, "description": p.description} for p in dangerous_perms],
        })
    
    # Check for exported components
    exported = [c for c in manifest_data.get("components", []) if c.is_exported]
    if exported:
        issues.append({
            "category": "Exported Components",
            "severity": "medium",
            "description": f"App has {len(exported)} exported components",
            "details": [{"name": c.name, "type": c.component_type} for c in exported[:10]],
        })
    
    # Check for HTTP URLs (not HTTPS)
    http_urls = [u for u in urls if u.startswith('http://') and not u.startswith('http://localhost')]
    if http_urls:
        issues.append({
            "category": "Insecure Network",
            "severity": "medium",
            "description": f"App contains {len(http_urls)} insecure HTTP URLs",
            "details": http_urls[:10],
        })
    
    # Check for hardcoded IPs
    ip_strings = [s.value for s in strings if s.category == "ip_address"]
    if ip_strings:
        issues.append({
            "category": "Hardcoded IPs",
            "severity": "low",
            "description": f"App contains {len(set(ip_strings))} hardcoded IP addresses",
            "details": list(set(ip_strings))[:10],
        })
    
    # Check for debugging flags
    debug_strings = [s for s in strings if 'debuggable' in s.value.lower() or 'debug' in s.value.lower()]
    if debug_strings:
        issues.append({
            "category": "Debug Mode",
            "severity": "medium",
            "description": "App may have debugging enabled",
        })
    
    return issues


# ============================================================================
# Docker Layer Analysis Functions
# ============================================================================

def analyze_docker_image(image_name: str) -> DockerLayerAnalysisResult:
    """Analyze Docker image layers for secrets and security issues."""
    try:
        # Get image history
        result = subprocess.run(
            ["docker", "history", "--no-trunc", "--format", "{{.ID}}|||{{.CreatedBy}}|||{{.Size}}", image_name],
            capture_output=True,
            text=True,
            timeout=60,
        )
        
        if result.returncode != 0:
            return DockerLayerAnalysisResult(
                image_name=image_name,
                image_id="unknown",
                total_layers=0,
                total_size=0,
                base_image=None,
                layers=[],
                secrets=[],
                deleted_files=[],
                security_issues=[],
                error=f"Docker command failed: {result.stderr}",
            )
        
        # Parse history
        layers = []
        secrets = []
        security_issues = []
        total_size = 0
        
        for line in result.stdout.strip().split('\n'):
            if not line.strip():
                continue
            
            parts = line.split('|||')
            if len(parts) < 3:
                continue
            
            layer_id = parts[0].strip()
            command = parts[1].strip()
            size_str = parts[2].strip()
            
            # Parse size
            try:
                if 'GB' in size_str:
                    size = int(float(size_str.replace('GB', '')) * 1024 * 1024 * 1024)
                elif 'MB' in size_str:
                    size = int(float(size_str.replace('MB', '')) * 1024 * 1024)
                elif 'KB' in size_str or 'kB' in size_str:
                    size = int(float(size_str.replace('KB', '').replace('kB', '')) * 1024)
                elif 'B' in size_str:
                    size = int(float(size_str.replace('B', '')))
                else:
                    size = 0
            except:
                size = 0
            
            total_size += size
            
            layer = {
                "id": layer_id[:12] if layer_id != "<missing>" else "inherited",
                "command": command[:500],
                "size": size,
            }
            layers.append(layer)
            
            # Check for secrets in layer commands
            layer_secrets = detect_secrets_in_layer(layer_id[:12], command)
            secrets.extend(layer_secrets)
            
            # Check for security issues
            layer_issues = check_layer_security(command)
            security_issues.extend(layer_issues)
        
        # Get image ID
        inspect_result = subprocess.run(
            ["docker", "inspect", "--format", "{{.Id}}", image_name],
            capture_output=True,
            text=True,
            timeout=30,
        )
        image_id = inspect_result.stdout.strip()[:12] if inspect_result.returncode == 0 else "unknown"
        
        # Try to identify base image
        base_image = None
        for layer in reversed(layers):
            if 'FROM' in layer["command"]:
                parts = layer["command"].split()
                for i, p in enumerate(parts):
                    if p == 'FROM' and i + 1 < len(parts):
                        base_image = parts[i + 1]
                        break
        
        return DockerLayerAnalysisResult(
            image_name=image_name,
            image_id=image_id,
            total_layers=len(layers),
            total_size=total_size,
            base_image=base_image,
            layers=layers,
            secrets=secrets,
            deleted_files=[],
            security_issues=security_issues,
        )
        
    except subprocess.TimeoutExpired:
        return DockerLayerAnalysisResult(
            image_name=image_name,
            image_id="unknown",
            total_layers=0,
            total_size=0,
            base_image=None,
            layers=[],
            secrets=[],
            deleted_files=[],
            security_issues=[],
            error="Docker command timed out",
        )
    except Exception as e:
        logger.error(f"Docker analysis failed: {e}")
        return DockerLayerAnalysisResult(
            image_name=image_name,
            image_id="unknown",
            total_layers=0,
            total_size=0,
            base_image=None,
            layers=[],
            secrets=[],
            deleted_files=[],
            security_issues=[],
            error=str(e),
        )


def detect_secrets_in_layer(layer_id: str, command: str) -> List[DockerLayerSecret]:
    """Detect secrets in Docker layer commands."""
    secrets = []
    
    # Check for secrets in ENV commands
    env_patterns = [
        (r'ENV\s+(\w*(?:PASSWORD|SECRET|KEY|TOKEN|API_KEY|CREDENTIAL)\w*)\s*=\s*["\']?([^\s"\']+)', "env_secret"),
        (r'ARG\s+(\w*(?:PASSWORD|SECRET|KEY|TOKEN|API_KEY|CREDENTIAL)\w*)\s*=\s*["\']?([^\s"\']+)', "arg_secret"),
        (r'--password[=\s]+["\']?([^\s"\']+)', "password_arg"),
        (r'--secret[=\s]+["\']?([^\s"\']+)', "secret_arg"),
        (r'--api-key[=\s]+["\']?([^\s"\']+)', "api_key_arg"),
    ]
    
    for pattern, secret_type in env_patterns:
        matches = re.finditer(pattern, command, re.IGNORECASE)
        for match in matches:
            if match.lastindex and match.lastindex >= 1:
                value = match.group(match.lastindex)
                
                # Skip placeholders
                if value.startswith('$') or value in {'password', 'secret', 'changeme', 'your-'}:
                    continue
                
                # Mask value
                if len(value) > 8:
                    masked = value[:3] + '*' * (len(value) - 6) + value[-3:]
                else:
                    masked = value[:2] + '*' * max(len(value) - 2, 1)
                
                secrets.append(DockerLayerSecret(
                    layer_id=layer_id,
                    layer_command=command[:200],
                    secret_type=secret_type,
                    value=value,
                    masked_value=masked,
                    context=match.group(0)[:100],
                    severity="critical" if "password" in secret_type.lower() else "high",
                ))
    
    # Check for secret patterns in commands
    for secret_type, pattern in SECRET_PATTERNS.items():
        matches = pattern.finditer(command)
        for match in matches:
            value = match.group(1) if match.lastindex else match.group(0)
            
            if value.lower() in {'password', 'secret', 'token', 'example'}:
                continue
            
            if len(value) > 8:
                masked = value[:3] + '*' * (len(value) - 6) + value[-3:]
            else:
                masked = value[:2] + '*' * max(len(value) - 2, 1)
            
            secrets.append(DockerLayerSecret(
                layer_id=layer_id,
                layer_command=command[:200],
                secret_type=secret_type,
                value=value,
                masked_value=masked,
                context=match.group(0)[:100],
                severity=get_secret_severity(secret_type),
            ))
    
    return secrets


def check_layer_security(command: str) -> List[Dict[str, Any]]:
    """Check Docker layer command for security issues."""
    issues = []
    
    # Running as root
    if 'USER root' in command:
        issues.append({
            "category": "Running as Root",
            "severity": "medium",
            "description": "Container runs as root user",
            "command": command[:200],
        })
    
    # Sensitive file operations
    sensitive_paths = ['/etc/passwd', '/etc/shadow', '.ssh/', 'id_rsa', '.aws/', '.kube/']
    for path in sensitive_paths:
        if path in command.lower():
            issues.append({
                "category": "Sensitive File Access",
                "severity": "high",
                "description": f"Layer accesses sensitive path: {path}",
                "command": command[:200],
            })
            break
    
    # curl/wget piped to shell
    if re.search(r'(curl|wget)\s+.*\|\s*(sh|bash)', command, re.IGNORECASE):
        issues.append({
            "category": "Remote Code Execution",
            "severity": "high",
            "description": "Script downloaded and executed directly",
            "command": command[:200],
        })
    
    # chmod 777
    if 'chmod 777' in command or 'chmod -R 777' in command:
        issues.append({
            "category": "Insecure Permissions",
            "severity": "medium",
            "description": "Files given world-writable permissions",
            "command": command[:200],
        })
    
    return issues


# ============================================================================
# AI Analysis Functions
# ============================================================================

# Structured AI analysis response model
@dataclass
class AIAnalysisStructured:
    """Structured AI analysis result."""
    risk_level: str  # Critical, High, Medium, Low, Clean
    risk_score: int  # 0-100
    summary: str
    key_findings: List[Dict[str, Any]]
    malware_indicators: List[str]
    recommendations: List[str]
    iocs: Dict[str, List[str]]  # Indicators of Compromise (urls, ips, domains, hashes)
    attack_techniques: List[str]  # MITRE ATT&CK techniques if applicable
    confidence: float  # 0.0-1.0


async def analyze_binary_with_ai(result: BinaryAnalysisResult) -> Optional[str]:
    """Use Gemini to provide comprehensive security analysis of binary."""
    if not settings.gemini_api_key:
        return None
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # Build detailed context
        suspicious_summary = "\n".join([
            f"- [{s['severity'].upper()}] {s['category']}: {s['description']}" 
            for s in result.suspicious_indicators
        ])
        
        secrets_summary = "\n".join([
            f"- [{s['severity'].upper()}] {s['type']}: {s['masked_value']}"
            for s in result.secrets[:15]
        ])
        
        # Get section entropy info
        section_info = ""
        if result.metadata.sections:
            section_info = "\n".join([
                f"- {s.get('name', 'N/A')}: size={s.get('raw_size', 0):,}B, entropy={s.get('entropy', 'N/A')}"
                for s in result.metadata.sections[:8]
            ])
        
        # Get notable imports
        suspicious_imports = [imp for imp in result.imports if imp.is_suspicious]
        imports_summary = "\n".join([
            f"- {imp.name} ({imp.library}): {imp.reason}"
            for imp in suspicious_imports[:20]
        ])
        
        # Get interesting strings
        interesting_strings = [s for s in result.strings if s.category in ('url', 'ip', 'email', 'path') or any(kw in s.value.lower() for kw in ('http', 'password', 'key', 'token', 'api', 'secret', 'cmd', 'exec', 'shell', 'powershell'))]
        strings_summary = "\n".join([
            f"- [{s.category or 'other'}] {s.value[:100]}"
            for s in interesting_strings[:30]
        ])
        
        prompt = f"""You are a malware analyst. Analyze this binary file and provide a comprehensive security assessment.

## FILE INFORMATION
- **Filename:** {result.filename}
- **Type:** {result.metadata.file_type}
- **Architecture:** {result.metadata.architecture}
- **Size:** {result.metadata.file_size:,} bytes
- **Entry Point:** {hex(result.metadata.entry_point) if result.metadata.entry_point else 'N/A'}
- **Compile Time:** {result.metadata.compile_time or 'Unknown'}
- **Packed:** {result.metadata.is_packed} {f'(Packer: {result.metadata.packer_name})' if result.metadata.packer_name else ''}

## SECTION ANALYSIS
{section_info or "No section data available"}

## SUSPICIOUS INDICATORS ({len(result.suspicious_indicators)})
{suspicious_summary or "None detected"}

## SUSPICIOUS IMPORTS ({len(suspicious_imports)})
{imports_summary or "None detected"}

## POTENTIAL SECRETS ({len(result.secrets)})
{secrets_summary or "None detected"}

## INTERESTING STRINGS ({len(interesting_strings)} of {len(result.strings)} total)
{strings_summary or "None of interest"}

## ANALYSIS INSTRUCTIONS
Provide your analysis in the following structured format:

**RISK ASSESSMENT**
- Risk Level: [Critical/High/Medium/Low/Clean]
- Risk Score: [0-100]
- Confidence: [High/Medium/Low]

**EXECUTIVE SUMMARY**
[2-3 sentence summary of findings]

**KEY FINDINGS**
1. [Finding with severity and details]
2. [Continue for each significant finding]

**MALWARE INDICATORS**
- [List any indicators suggesting malicious behavior]

**MITRE ATT&CK TECHNIQUES**
- [List applicable techniques like T1055 Process Injection if detected]

**INDICATORS OF COMPROMISE (IOCs)**
- URLs: [list any malicious/suspicious URLs]
- IPs: [list any suspicious IPs]
- Domains: [list any suspicious domains]
- File Hashes: [if relevant]

**RECOMMENDATIONS**
1. [Specific actionable recommendation]
2. [Continue as needed]

Be thorough but concise. Focus on actionable intelligence."""

        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        return response.text
        
    except Exception as e:
        logger.error(f"AI analysis failed: {e}")
        return f"AI analysis unavailable: {str(e)}"


@dataclass
class ApkAIReports:
    """AI-generated reports for APK analysis."""
    functionality_report: Optional[str] = None
    security_report: Optional[str] = None
    legacy_report: Optional[str] = None  # Combined for backwards compatibility


async def analyze_apk_with_ai(result: ApkAnalysisResult) -> Optional[str]:
    """Use Gemini to provide comprehensive security analysis of APK."""
    reports = await generate_apk_ai_reports(result)
    if reports:
        # Return combined report for backwards compatibility
        result.ai_report_functionality = reports.functionality_report
        result.ai_report_security = reports.security_report
        return reports.legacy_report
    return None


async def generate_apk_ai_reports(result: ApkAnalysisResult) -> Optional[ApkAIReports]:
    """Generate two separate AI reports: functionality and security."""
    if not settings.gemini_api_key:
        return None
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # Categorize permissions
        dangerous_perms = [p for p in result.permissions if p.is_dangerous]
        privacy_perms = [p for p in result.permissions if any(kw in p.name.lower() for kw in ('camera', 'microphone', 'location', 'contacts', 'sms', 'call', 'calendar', 'storage'))]
        network_perms = [p for p in result.permissions if any(kw in p.name.lower() for kw in ('internet', 'network', 'wifi', 'bluetooth'))]
        
        # Get notable URLs
        suspicious_urls = [u for u in result.urls if not any(safe in u.lower() for safe in ('google.com', 'android.com', 'googleapis.com', 'gstatic.com'))]
        
        # Get data flow info if available
        data_flow_context = ""
        if result.data_flow_analysis:
            dfa = result.data_flow_analysis
            data_flow_context = f"""
Data Flow Analysis:
- Sensitive Data Sources Found: {dfa.get('total_sources', 0)}
- Data Sinks (exit points): {dfa.get('total_sinks', 0)}
- Data Flow Paths: {dfa.get('total_flows', 0)}
- Critical Flows: {dfa.get('critical_flows', 0)}
- High-Risk Flows: {dfa.get('high_risk_flows', 0)}
- Privacy Violations: {len(dfa.get('privacy_violations', []))}

Top Data Flow Concerns:
{chr(10).join(f'- {p["source"]["source_type"]} → {p["sink"]["sink_type"]}: {p["description"]}' for p in dfa.get('data_flow_paths', [])[:5]) or 'None'}
"""
        
        # Build context for AI
        app_context = f"""
Package: {result.package_name}
App Name: {result.app_name or 'Unknown'}
Version: {result.version_name} (code: {result.version_code})
Min SDK: {result.min_sdk}
Target SDK: {result.target_sdk}

Activities ({len(result.activities)}): {', '.join(result.activities[:10])}
Services ({len(result.services)}): {', '.join(result.services[:10])}
Receivers ({len(result.receivers)}): {', '.join(result.receivers[:10])}
Providers ({len(result.providers)}): {', '.join(result.providers[:10])}

Permissions ({len(result.permissions)}):
{chr(10).join(f'- {p.name} (dangerous: {p.is_dangerous})' for p in result.permissions[:20])}

Native Libraries: {', '.join(result.native_libraries[:10]) or 'None'}
URLs Found: {chr(10).join(suspicious_urls[:15]) or 'None'}
{data_flow_context}
"""

        # ==================== REPORT 1: What Does This APK Do ====================
        functionality_prompt = f"""You are an Android app analyst explaining what an app does to a non-technical user.

## APP DATA
{app_context}

## YOUR TASK
Generate a clear, user-friendly report explaining what this app does. Use HTML formatting for proper display.

FORMAT YOUR RESPONSE AS CLEAN HTML (no markdown, no code blocks):
- Use <h3> for section headers
- Use <ul> and <li> for bullet points  
- Use <strong> for emphasis
- Use <p> for paragraphs
- Keep it readable and well-organized

REQUIRED SECTIONS:

<h3>📱 App Overview</h3>
<p>[1-2 sentences describing what this app appears to be and its main purpose]</p>

<h3>🎯 Primary Functions</h3>
<ul>
<li><strong>[Function 1]:</strong> [Brief description]</li>
<li><strong>[Function 2]:</strong> [Brief description]</li>
[Continue for main features]
</ul>

<h3>🔌 System Integration</h3>
<p>[How it interacts with other apps/system - based on components]</p>
<ul>
<li>[Integration point 1]</li>
<li>[Integration point 2]</li>
</ul>

<h3>📡 Network & Communication</h3>
<ul>
<li>[Describe network usage based on URLs and permissions]</li>
</ul>

<h3>💾 Data Access</h3>
<ul>
<li>[What data the app can access based on permissions]</li>
</ul>

<h3>⚠️ User Should Know</h3>
<ul>
<li>[Important things users should be aware of]</li>
</ul>

Write in plain English. Be factual and objective."""

        # ==================== REPORT 2: Security Findings ====================
        security_context = f"""
{app_context}

Security Issues ({len(result.security_issues)}):
{chr(10).join(f'- [{i.get("severity", "INFO").upper()}] {i["category"]}: {i["description"]}' for i in result.security_issues[:15]) or "None"}

Secrets/Keys Found ({len(result.secrets)}):
{chr(10).join(f'- {s["type"]}: {s["masked_value"]}' for s in result.secrets[:10]) or "None"}

Debuggable: {result.debuggable}
Allows Backup: {result.allow_backup}
Certificate: {"Debug certificate" if result.certificate and result.certificate.is_debug_cert else "Production certificate" if result.certificate else "Unknown"}
"""

        security_prompt = f"""You are a mobile security expert providing a security assessment.

## APP DATA
{security_context}

## YOUR TASK
Generate a professional security assessment report. Use HTML formatting for proper display.

FORMAT YOUR RESPONSE AS CLEAN HTML (no markdown, no code blocks):
- Use <h3> for section headers
- Use <ul> and <li> for bullet points
- Use <strong> for emphasis and labels
- Use colored badges for severity: <span style="color: #dc2626; font-weight: bold;">CRITICAL</span>, <span style="color: #ea580c; font-weight: bold;">HIGH</span>, <span style="color: #ca8a04; font-weight: bold;">MEDIUM</span>, <span style="color: #16a34a; font-weight: bold;">LOW</span>

REQUIRED SECTIONS:

<h3>🎯 Risk Assessment</h3>
<p><strong>Overall Risk:</strong> [<span style="color: #XX;">LEVEL</span>] - [Brief rationale]</p>
<p><strong>Risk Score:</strong> [X/100]</p>

<h3>🔴 Critical & High Findings</h3>
<ul>
<li><span style="color: #dc2626; font-weight: bold;">CRITICAL</span> - <strong>[Finding]:</strong> [Description and impact]</li>
<li><span style="color: #ea580c; font-weight: bold;">HIGH</span> - <strong>[Finding]:</strong> [Description]</li>
[Include all critical/high issues, including data flow vulnerabilities]
</ul>

<h3>🔀 Data Flow & Privacy Issues</h3>
<p>Analyze the data flow paths showing how sensitive data moves through the app.</p>
<ul>
<li><strong>[Source → Sink]:</strong> [Description of sensitive data flow]</li>
</ul>

<h3>🟡 Medium & Low Findings</h3>
<ul>
<li><span style="color: #ca8a04; font-weight: bold;">MEDIUM</span> - <strong>[Finding]:</strong> [Description]</li>
<li><span style="color: #16a34a; font-weight: bold;">LOW</span> - <strong>[Finding]:</strong> [Description]</li>
</ul>

<h3>🔐 Privacy Concerns</h3>
<ul>
<li><strong>[Concern 1]:</strong> [Description based on permissions and data flows]</li>
</ul>

<h3>🔑 Exposed Secrets</h3>
<ul>
<li><strong>[Secret type]:</strong> [Risk and recommendation]</li>
</ul>

<h3>✅ Security Positives</h3>
<ul>
<li>[Good security practices observed]</li>
</ul>

<h3>📋 Recommendations</h3>
<ol>
<li><strong>[Priority]:</strong> [Specific recommendation]</li>
<li><strong>[Next priority]:</strong> [Recommendation]</li>
</ol>

Be thorough and actionable. Focus on real risks."""

        # Generate both reports in parallel
        functionality_response, security_response = await asyncio.gather(
            client.aio.models.generate_content(
                model=settings.gemini_model_id,
                contents=[types.Content(role="user", parts=[types.Part(text=functionality_prompt)])],
            ),
            client.aio.models.generate_content(
                model=settings.gemini_model_id,
                contents=[types.Content(role="user", parts=[types.Part(text=security_prompt)])],
            )
        )
        
        functionality_report = functionality_response.text if functionality_response else None
        security_report = security_response.text if security_response else None
        
        # Clean up any markdown code block wrappers
        if functionality_report:
            functionality_report = functionality_report.strip()
            if functionality_report.startswith("```html"):
                functionality_report = functionality_report[7:]
            if functionality_report.startswith("```"):
                functionality_report = functionality_report[3:]
            if functionality_report.endswith("```"):
                functionality_report = functionality_report[:-3]
            functionality_report = functionality_report.strip()
            
        if security_report:
            security_report = security_report.strip()
            if security_report.startswith("```html"):
                security_report = security_report[7:]
            if security_report.startswith("```"):
                security_report = security_report[3:]
            if security_report.endswith("```"):
                security_report = security_report[:-3]
            security_report = security_report.strip()
        
        # Combine for legacy support
        legacy_report = f"""== WHAT DOES THIS APK DO ==

{functionality_report or 'Analysis unavailable'}

== SECURITY FINDINGS ==

{security_report or 'Analysis unavailable'}
"""
        
        return ApkAIReports(
            functionality_report=functionality_report,
            security_report=security_report,
            legacy_report=legacy_report
        )
        
    except Exception as e:
        logger.error(f"AI analysis failed: {e}")
        return None


async def analyze_docker_with_ai(result: DockerLayerAnalysisResult) -> Optional[str]:
    """Use Gemini to provide comprehensive security analysis of Docker image."""
    if not settings.gemini_api_key:
        return None
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # Group secrets by type
        secrets_by_type = {}
        for s in result.secrets:
            st = s.secret_type
            if st not in secrets_by_type:
                secrets_by_type[st] = []
            secrets_by_type[st].append(s)
        
        secrets_summary = "\n".join([
            f"- **{stype}** ({len(secrets)}): {', '.join(s.masked_value for s in secrets[:3])}"
            for stype, secrets in secrets_by_type.items()
        ])
        
        # Layer commands analysis
        layer_commands = "\n".join([
            f"- Layer {i+1}: {layer.get('created_by', 'unknown')[:100]}"
            for i, layer in enumerate(result.layers[:10]) if layer.get('created_by')
        ])
        
        prompt = f"""You are a container security analyst. Analyze this Docker image and provide a comprehensive security assessment.

## IMAGE INFORMATION
- **Name:** {result.image_name}
- **ID:** {result.image_id}
- **Layers:** {result.total_layers}
- **Total Size:** {result.total_size / (1024*1024):.1f} MB
- **Base Image:** {result.base_image or "Unknown"}

## LAYER COMMANDS (build history)
{layer_commands or "No layer history available"}

## SECRETS FOUND ({len(result.secrets)})
{secrets_summary or "None detected"}

**Detailed Secret Locations:**
{chr(10).join(f'- {s.secret_type} in layer {s.layer_id}: {s.file_path or "unknown file"}' for s in result.secrets[:10]) or "None"}

## SECURITY ISSUES ({len(result.security_issues)})
{chr(10).join(f'- [{i.get("severity", "INFO").upper()}] {i["category"]}: {i["description"]}' for i in result.security_issues[:15]) or "None"}

## ANALYSIS INSTRUCTIONS
Provide your analysis in the following structured format:

**RISK ASSESSMENT**
- Risk Level: [Critical/High/Medium/Low/Clean]
- Risk Score: [0-100]
- Confidence: [High/Medium/Low]

**EXECUTIVE SUMMARY**
[2-3 sentence summary of findings]

**SECRET EXPOSURE ANALYSIS**
[Analysis of exposed secrets, their impact, and which layers contain them]

**SUPPLY CHAIN CONCERNS**
[Analysis of base image, dependencies, and supply chain risks]

**KEY SECURITY FINDINGS**
1. [Finding with severity and details]
2. [Continue for significant findings]

**CONTAINER HARDENING RECOMMENDATIONS**
1. [Specific Dockerfile improvements]
2. [Runtime security recommendations]
3. [Secret management recommendations]

**REMEDIATION STEPS**
1. [Step to remove/rotate exposed secrets]
2. [Steps to rebuild image securely]

Be thorough but actionable. Focus on practical remediation."""

        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        return response.text
        
    except Exception as e:
        logger.error(f"AI analysis failed: {e}")
        return f"AI analysis unavailable: {str(e)}"


# ============================================================================
# AI-Generated Mermaid Architecture Diagrams
# ============================================================================

async def generate_ai_architecture_diagram(
    result: ApkAnalysisResult,
    jadx_result: Optional[Dict[str, Any]] = None
) -> Optional[str]:
    """
    Generate an AI-powered Mermaid architecture diagram after APK analysis.
    
    Uses Gemini to create a comprehensive architecture visualization with icons.
    
    Args:
        result: The APK analysis result
        jadx_result: Optional JADX decompilation result dict (from get_jadx_result_summary)
    
    Returns:
        Mermaid diagram code string, or None if generation fails
    """
    if not settings.gemini_api_key:
        return None
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # Build context about the app
        components_context = f"""
Package: {result.package_name}
Activities ({len(result.activities)}): {', '.join(result.activities[:10])}
Services ({len(result.services)}): {', '.join(result.services[:5])}
Receivers ({len(result.receivers)}): {', '.join(result.receivers[:5])}
Providers ({len(result.providers)}): {', '.join(result.providers[:5])}
"""
        
        # Add decompilation insights if available (dict format from get_jadx_result_summary)
        decompile_context = ""
        if jadx_result and isinstance(jadx_result, dict):
            classes = jadx_result.get('classes', [])
            key_classes = [c for c in classes[:30] if c.get('is_activity') or c.get('is_service') or c.get('is_receiver')]
            if key_classes:
                decompile_context = f"""
Key Classes Found:
{chr(10).join(f"- {c.get('class_name', 'Unknown')} ({'Activity' if c.get('is_activity') else 'Service' if c.get('is_service') else 'Receiver' if c.get('is_receiver') else 'Provider'})" for c in key_classes[:15])}
"""
            # Add sample code context for better AI understanding
            sample_code = jadx_result.get('sample_code', [])
            if sample_code:
                decompile_context += f"""
Sample Source Code Context:
"""
                for sample in sample_code[:3]:  # Limit to 3 samples
                    decompile_context += f"""
--- {sample.get('class_name', 'Unknown')} ({sample.get('type', 'class')}) ---
{sample.get('code_snippet', '')[:1500]}
"""
        
        # Security context
        security_context = f"""
Security Issues: {len(result.security_issues)}
Dangerous Permissions: {len([p for p in result.permissions if p.is_dangerous])}
Exported Components: {len([a for a in result.activities if 'exported' in str(result.manifest_raw).lower()])}
Secrets Found: {len(result.secrets)}
"""
        
        # URLs and network
        network_context = ""
        if result.urls:
            network_context = f"""
URLs Found ({len(result.urls)}): {', '.join(result.urls[:5])}
"""

        prompt = f"""You are an Android app architecture analyst. Generate a Mermaid flowchart showing the app's architecture and data flow.

## APP DATA
{components_context}
{decompile_context}
{security_context}
{network_context}

## AVAILABLE ICONS (use these in your diagram)
Use icons in node labels like: A[prefix:icon-name Label Text]

- fa6-brands:android - Android app/APK
- mdi:application - Activity
- mdi:rocket-launch - Main/Launcher Activity
- mdi:cog - Service
- mdi:broadcast - Broadcast Receiver
- mdi:database - Content Provider / Database
- fa6-solid:shield - Security component
- fa6-solid:lock - Authentication
- fa6-solid:key - Crypto/Keys
- fa6-solid:bug - Vulnerability
- fa6-solid:network-wired - Network
- mdi:api - API calls
- fa6-solid:server - Server/Backend
- fa6-solid:triangle-exclamation - Warning/Risk
- mdi:file-code - Code/Files
- fa6-solid:user - User
- mdi:cellphone - Device

## INSTRUCTIONS
Generate a Mermaid flowchart that shows:
1. Main entry points (launcher activities)
2. Key components and their relationships
3. Data flow (user input → processing → storage/network)
4. Security-relevant areas (auth, crypto, network)
5. Any concerning patterns (exported components, sensitive data flows)

Return ONLY the Mermaid diagram code, starting with "flowchart" - no explanation or markdown code blocks.

Example format:
flowchart TD
    subgraph User[\"fa6-solid:user User Interaction\"]
        A[mdi:rocket-launch MainActivity]
    end
    A --> B[mdi:application LoginActivity]
    B --> C{{fa6-solid:lock Auth Check}}
    C -->|Success| D[mdi:database User Data]
    C -->|Fail| E[fa6-solid:triangle-exclamation Error]
"""

        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        diagram = response.text if response else None
        
        if diagram:
            # Clean up response
            diagram = diagram.strip()
            if diagram.startswith("```mermaid"):
                diagram = diagram[10:]
            if diagram.startswith("```"):
                diagram = diagram[3:]
            if diagram.endswith("```"):
                diagram = diagram[:-3]
            diagram = diagram.strip()
            
            # Validate it starts with flowchart
            if not diagram.startswith(("flowchart", "graph")):
                logger.warning("AI generated invalid Mermaid diagram")
                return None
        
        return diagram
        
    except Exception as e:
        logger.error(f"AI architecture diagram generation failed: {e}")
        return None


async def generate_ai_data_flow_diagram(
    result: ApkAnalysisResult,
    jadx_result: Optional[Dict[str, Any]] = None
) -> Optional[str]:
    """
    Generate an AI-powered Mermaid diagram showing data flow and privacy concerns.
    
    Args:
        result: The APK analysis result
        jadx_result: Optional JADX decompilation result dict (from get_jadx_result_summary)
    
    Returns:
        Mermaid diagram code string, or None if generation fails
    """
    if not settings.gemini_api_key:
        return None
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        # Build data flow context
        permissions_context = ""
        dangerous_perms = [p for p in result.permissions if p.is_dangerous]
        if dangerous_perms:
            permissions_context = f"""
Dangerous Permissions:
{chr(10).join(f'- {p.name}' for p in dangerous_perms[:10])}
"""
        
        # Data flow context from JADX if available
        flow_context = ""
        if jadx_result and isinstance(jadx_result, dict):
            sample_code = jadx_result.get('sample_code', [])
            if sample_code:
                flow_context = f"""
Source Code Analysis:
Total Classes: {len(jadx_result.get('classes', []))}
"""
                # Add code snippets for data flow analysis
                for sample in sample_code[:2]:
                    code = sample.get('code_snippet', '')[:1000]
                    flow_context += f"""
--- {sample.get('class_name', 'Unknown')} ---
{code}
"""
        
        # Secrets context
        secrets_context = ""
        if result.secrets:
            secrets_context = f"""
Exposed Secrets ({len(result.secrets)}):
{chr(10).join(f'- {s["type"]}: {s["masked_value"]}' for s in result.secrets[:5])}
"""

        prompt = f"""You are a privacy and data flow analyst. Generate a Mermaid flowchart showing how sensitive data flows through the app.

## APP: {result.package_name}

{permissions_context}
{flow_context}
{secrets_context}

URLs: {', '.join(result.urls[:5]) if result.urls else 'None found'}

## AVAILABLE ICONS
- fa6-solid:user - User input
- mdi:cellphone - Device data (IMEI, etc)
- fa6-solid:location-dot - Location
- mdi:camera - Camera
- mdi:microphone - Microphone
- mdi:contacts - Contacts
- mdi:database - Local storage
- fa6-solid:server - Remote server
- fa6-solid:cloud - Cloud service
- fa6-solid:lock - Encrypted
- fa6-solid:unlock - Unencrypted/Exposed
- fa6-solid:triangle-exclamation - Privacy risk
- fa6-solid:shield-halved - Protected
- mdi:api - API endpoint
- fa6-solid:file-export - Data exfiltration risk

## INSTRUCTIONS
Generate a Mermaid flowchart showing:
1. Data sources (user input, device sensors, storage)
2. Data processing within the app
3. Data destinations (storage, network, third parties)
4. Privacy risk indicators
5. Encryption status where apparent

Return ONLY the Mermaid diagram code - no explanation or markdown.

Example:
flowchart LR
    subgraph Sources[\"fa6-solid:database Data Sources\"]
        U[fa6-solid:user User Input]
        L[fa6-solid:location-dot Location]
    end
    U --> P[mdi:application Processing]
    L --> P
    P --> S[fa6-solid:lock Encrypted Storage]
    P --> N[fa6-solid:triangle-exclamation Unencrypted Network]
"""

        response = await client.aio.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        diagram = response.text if response else None
        
        if diagram:
            diagram = diagram.strip()
            if diagram.startswith("```mermaid"):
                diagram = diagram[10:]
            if diagram.startswith("```"):
                diagram = diagram[3:]
            if diagram.endswith("```"):
                diagram = diagram[:-3]
            diagram = diagram.strip()
            
            if not diagram.startswith(("flowchart", "graph")):
                return None
        
        return diagram
        
    except Exception as e:
        logger.error(f"AI data flow diagram generation failed: {e}")
        return None


# ============================================================================
# JADX Decompilation Functions
# ============================================================================

def decompile_apk_with_jadx(apk_path: Path, output_dir: Optional[Path] = None) -> JadxDecompilationResult:
    """
    Decompile APK to Java source code using JADX.
    
    Args:
        apk_path: Path to the APK file
        output_dir: Optional output directory (temp dir if not specified)
    
    Returns:
        JadxDecompilationResult with decompiled classes and metadata
    """
    import subprocess
    import time
    
    start_time = time.time()
    
    # Create output directory if not specified
    if output_dir is None:
        output_dir = Path(tempfile.mkdtemp(prefix="jadx_"))
    
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    errors = []
    warnings = []
    
    try:
        # Run JADX decompilation with performance optimizations
        import os
        cpu_count = os.cpu_count() or 4
        thread_count = max(4, cpu_count - 1)  # Use all but 1 core
        
        result = subprocess.run(
            [
                "jadx",
                "-d", str(output_dir),
                "--show-bad-code",  # Show decompiled code even with errors
                "--deobf",  # Deobfuscate names
                "--deobf-min", "2",
                "--deobf-max", "64",
                "--threads-count", str(thread_count),  # Dynamic thread count
                "--no-debug-info",  # Skip debug info for speed
                str(apk_path)
            ],
            capture_output=True,
            text=True,
            timeout=1800  # 30 minute timeout for very large APKs (games, enterprise apps)
        )
        
        if result.returncode != 0:
            warnings.append(f"JADX warnings: {result.stderr[:500]}")
        
    except subprocess.TimeoutExpired:
        errors.append("JADX decompilation timed out after 30 minutes. The APK may be extremely large or complex.")
        return JadxDecompilationResult(
            package_name="unknown",
            total_classes=0,
            total_files=0,
            output_directory=str(output_dir),
            classes=[],
            resources_dir=str(output_dir / "resources"),
            manifest_path="",
            source_tree={},
            decompilation_time=time.time() - start_time,
            errors=errors
        )
    except FileNotFoundError:
        errors.append("JADX not found. Please install JADX.")
        return JadxDecompilationResult(
            package_name="unknown",
            total_classes=0,
            total_files=0,
            output_directory=str(output_dir),
            classes=[],
            resources_dir=str(output_dir / "resources"),
            manifest_path="",
            source_tree={},
            decompilation_time=time.time() - start_time,
            errors=errors
        )
    
    # Parse decompiled output
    sources_dir = output_dir / "sources"
    resources_dir = output_dir / "resources"
    manifest_path = resources_dir / "AndroidManifest.xml"
    
    # Build source tree and collect classes
    classes = []
    source_tree = {}
    package_name = "unknown"
    
    if sources_dir.exists():
        for java_file in sources_dir.rglob("*.java"):
            try:
                source_code = java_file.read_text(encoding='utf-8', errors='ignore')
                rel_path = str(java_file.relative_to(sources_dir))
                
                # Parse class info from source
                class_info = _parse_java_class(source_code, rel_path)
                classes.append(class_info)
                
                # Extract package name from first class
                if package_name == "unknown" and class_info.package_name:
                    # Try to get base package (first 2-3 segments)
                    parts = class_info.package_name.split('.')
                    if len(parts) >= 2:
                        package_name = '.'.join(parts[:min(3, len(parts))])
                
                # Build tree structure
                _add_to_source_tree(source_tree, rel_path)
                
            except Exception as e:
                warnings.append(f"Failed to parse {java_file.name}: {str(e)}")
    
    decompilation_time = time.time() - start_time
    
    return JadxDecompilationResult(
        package_name=package_name,
        total_classes=len(classes),
        total_files=len(list(sources_dir.rglob("*.java"))) if sources_dir.exists() else 0,
        output_directory=str(output_dir),
        classes=classes,
        resources_dir=str(resources_dir),
        manifest_path=str(manifest_path) if manifest_path.exists() else "",
        source_tree=source_tree,
        decompilation_time=decompilation_time,
        errors=errors,
        warnings=warnings
    )


def _parse_java_class(source_code: str, file_path: str) -> JadxDecompiledClass:
    """Parse a Java source file to extract class information."""
    import re
    
    lines = source_code.split('\n')
    line_count = len(lines)
    
    # Extract package
    package_match = re.search(r'package\s+([\w.]+)\s*;', source_code)
    package_name = package_match.group(1) if package_match else ""
    
    # Extract class name
    class_match = re.search(r'(?:public\s+)?(?:abstract\s+)?(?:final\s+)?class\s+(\w+)', source_code)
    class_name = class_match.group(1) if class_match else Path(file_path).stem
    
    # Extract extends
    extends_match = re.search(r'class\s+\w+\s+extends\s+([\w.]+)', source_code)
    extends = extends_match.group(1) if extends_match else None
    
    # Extract implements
    implements_match = re.search(r'implements\s+([\w.,\s]+)(?:\s*\{)', source_code)
    implements = []
    if implements_match:
        implements = [i.strip() for i in implements_match.group(1).split(',')]
    
    # Check component types - ensure boolean, not None
    is_activity = bool(extends and ('Activity' in extends or 'AppCompatActivity' in extends))
    is_service = bool(extends and 'Service' in extends)
    is_receiver = bool(extends and 'BroadcastReceiver' in extends)
    is_provider = bool(extends and 'ContentProvider' in extends)
    is_application = bool(extends and 'Application' in extends)
    
    # Extract methods
    method_pattern = r'(?:public|private|protected)?\s*(?:static\s+)?(?:final\s+)?(?:synchronized\s+)?(?:[\w<>\[\],\s]+)\s+(\w+)\s*\([^)]*\)'
    methods = re.findall(method_pattern, source_code)
    
    # Extract fields
    field_pattern = r'(?:public|private|protected)\s+(?:static\s+)?(?:final\s+)?[\w<>\[\],\s]+\s+(\w+)\s*[;=]'
    fields = re.findall(field_pattern, source_code)
    
    # Check for security issues
    security_issues = _scan_java_security_issues(source_code, class_name)
    
    return JadxDecompiledClass(
        class_name=class_name,
        package_name=package_name,
        file_path=file_path,
        source_code=source_code,
        line_count=line_count,
        is_activity=is_activity,
        is_service=is_service,
        is_receiver=is_receiver,
        is_provider=is_provider,
        is_application=is_application,
        extends=extends,
        implements=implements,
        methods=methods[:50],  # Limit to avoid huge responses
        fields=fields[:50],
        security_issues=security_issues
    )


def _scan_java_security_issues(source_code: str, class_name: str) -> List[Dict[str, Any]]:
    """Scan Java source for common security issues."""
    issues = []
    
    # Comprehensive security patterns - 60+ patterns covering OWASP Mobile Top 10
    patterns = [
        # ========== M1: Improper Platform Usage ==========
        (r'\.exec\s*\(', 'Command Execution', 'high', 'Potential command injection vulnerability'),
        (r'Runtime\.getRuntime\(\)\.exec', 'Runtime Exec', 'high', 'Direct runtime command execution - shell injection risk'),
        (r'ProcessBuilder', 'Process Builder', 'high', 'Process execution - potential command injection'),
        (r'android:exported\s*=\s*"true"', 'Exported Component', 'medium', 'Exported component may be accessible by other apps'),
        (r'PendingIntent\.get(Activity|Service|Broadcast)\s*\([^,]+,\s*0', 'Insecure PendingIntent', 'high', 'PendingIntent without FLAG_IMMUTABLE is mutable'),
        (r'registerReceiver\s*\([^,]+,\s*new\s+IntentFilter', 'Dynamic Receiver', 'medium', 'Dynamically registered receiver may be exploitable'),
        (r'sendBroadcast\s*\([^)]+\)', 'Broadcast Without Permission', 'medium', 'Broadcasting without permission restriction'),
        (r'bindService\s*\(.*BIND_AUTO_CREATE', 'Service Binding', 'low', 'Service binding - verify permission requirements'),
        
        # ========== M2: Insecure Data Storage ==========
        (r'MODE_WORLD_READABLE|MODE_WORLD_WRITEABLE', 'World Accessible', 'critical', 'World-accessible file permissions - data exposure'),
        (r'getSharedPreferences.*\.edit\(\)', 'SharedPreferences Write', 'info', 'Writing to SharedPreferences - verify no sensitive data'),
        (r'putString\s*\([^,]*password|putString\s*\([^,]*token|putString\s*\([^,]*secret|putString\s*\([^,]*key', 'Sensitive Data Storage', 'high', 'Sensitive data stored in SharedPreferences'),
        (r'openFileOutput\s*\([^,]+,\s*Context\.MODE_PRIVATE', 'File Storage', 'low', 'File storage - verify encryption for sensitive data'),
        (r'SQLiteDatabase\.openOrCreateDatabase', 'SQLite Database', 'medium', 'SQLite database - verify encryption and parameterized queries'),
        (r'\.getWritableDatabase|\.getReadableDatabase', 'Database Access', 'info', 'Database access - review for SQL injection'),
        (r'rawQuery\s*\([^,]*\+', 'SQL Concatenation', 'high', 'String concatenation in SQL query - SQL injection risk'),
        (r'execSQL\s*\([^,]*\+', 'SQL Exec Concatenation', 'high', 'String concatenation in execSQL - SQL injection risk'),
        (r'getExternalFilesDir|getExternalStorageDirectory', 'External Storage', 'medium', 'External storage access - data accessible by other apps'),
        (r'Environment\.getExternalStoragePublicDirectory', 'Public Storage', 'high', 'Public external storage - data exposed to all apps'),
        (r'openFileInput|openFileOutput', 'Internal File Access', 'low', 'Internal file access - verify sensitive data handling'),
        (r'\.write\s*\([^)]*password|\.write\s*\([^)]*secret', 'Writing Secrets', 'high', 'Potentially writing secrets to file'),
        
        # ========== M3: Insecure Communication ==========
        (r'TrustManager.*X509', 'Custom TrustManager', 'critical', 'Custom SSL/TLS trust manager - may bypass validation'),
        (r'SSLSocketFactory.*ALLOW_ALL', 'SSL Bypass', 'critical', 'SSL certificate validation completely disabled'),
        (r'setHostnameVerifier.*ALLOW_ALL', 'Hostname Bypass', 'critical', 'Hostname verification disabled'),
        (r'TrustAllCertificates|TrustAllManager', 'Trust All Certs', 'critical', 'Trusting all certificates - MITM vulnerability'),
        (r'checkServerTrusted.*\{\s*\}|checkServerTrusted.*return;', 'Empty TrustManager', 'critical', 'Empty certificate validation - critical vulnerability'),
        (r'ALLOW_ALL_HOSTNAME_VERIFIER', 'Hostname Verifier Bypass', 'critical', 'Accepting all hostnames - MITM vulnerability'),
        (r'http://', 'Cleartext HTTP', 'medium', 'Cleartext HTTP connection - data may be intercepted'),
        (r'\.setRequestProperty\s*\([^,]*[Aa]uthorization', 'Auth Header', 'info', 'Authorization header - verify HTTPS usage'),
        (r'BasicAuth|Basic\s+Auth', 'Basic Authentication', 'medium', 'Basic authentication - credentials may be exposed'),
        (r'WebSocket\s*\(\s*"ws://', 'Cleartext WebSocket', 'medium', 'Unencrypted WebSocket connection'),
        
        # ========== M4: Insecure Authentication ==========
        (r'BiometricPrompt.*setNegativeButtonText', 'Biometric Auth', 'info', 'Biometric authentication implementation'),
        (r'checkSelfPermission.*FINGERPRINT', 'Fingerprint Auth', 'info', 'Fingerprint authentication check'),
        (r'\.equals\s*\([^)]*password', 'Password Comparison', 'medium', 'String equals for password - timing attack risk'),
        (r'hardcoded.*password|password\s*=\s*"[^"]+"', 'Hardcoded Password', 'critical', 'Hardcoded password detected'),
        (r'api[_-]?key\s*=\s*"[^"]+"', 'Hardcoded API Key', 'high', 'Hardcoded API key detected'),
        (r'secret\s*=\s*"[^"]+"', 'Hardcoded Secret', 'high', 'Hardcoded secret value detected'),
        (r'token\s*=\s*"[A-Za-z0-9+/=]{20,}"', 'Hardcoded Token', 'high', 'Hardcoded authentication token'),
        
        # ========== M5: Insufficient Cryptography ==========
        (r'SecretKeySpec\s*\([^,]+,\s*"DES"', 'Weak Crypto DES', 'high', 'Using weak DES encryption - use AES instead'),
        (r'Cipher\.getInstance\s*\(\s*"DES', 'DES Cipher', 'high', 'DES cipher is deprecated and weak'),
        (r'Cipher\.getInstance\s*\(\s*"AES/ECB', 'ECB Mode', 'high', 'ECB mode leaks patterns - use GCM or CBC'),
        (r'Cipher\.getInstance\s*\(\s*"AES"\s*\)', 'AES Default Mode', 'medium', 'AES without mode defaults to ECB'),
        (r'Cipher\.getInstance\s*\(\s*"RC4', 'RC4 Cipher', 'high', 'RC4 is broken - do not use'),
        (r'Cipher\.getInstance\s*\(\s*"Blowfish', 'Blowfish Cipher', 'medium', 'Blowfish is outdated - use AES'),
        (r'MessageDigest\.getInstance\s*\(\s*"MD5"', 'MD5 Hash', 'high', 'MD5 is broken for security purposes'),
        (r'MessageDigest\.getInstance\s*\(\s*"SHA-1"', 'SHA1 Hash', 'medium', 'SHA-1 is deprecated - use SHA-256+'),
        (r'new\s+Random\s*\(\)', 'Weak Random', 'high', 'Using non-cryptographic Random - use SecureRandom'),
        (r'Math\.random\s*\(\)', 'Math Random', 'high', 'Math.random is not cryptographically secure'),
        (r'IvParameterSpec\s*\([^)]*new byte\[\]\s*\{0', 'Zero IV', 'critical', 'Using zero/static IV - breaks encryption security'),
        (r'SecretKeySpec\s*\([^)]*"[^"]+"\s*\.getBytes', 'Static Key', 'critical', 'Using static/hardcoded encryption key'),
        (r'KeyGenerator\.getInstance.*\.init\s*\(\s*(56|64|128)\s*\)', 'Weak Key Size', 'medium', 'Verify encryption key size is adequate'),
        
        # ========== M6: Insecure Authorization ==========
        (r'ContentProvider.*exported.*true', 'Exported Provider', 'high', 'Exported ContentProvider - verify permissions'),
        (r'grantUriPermission', 'URI Permission Grant', 'medium', 'Granting URI permissions - verify scope'),
        (r'checkCallingPermission.*PERMISSION_DENIED', 'Permission Check', 'info', 'Permission check implementation'),
        (r'enforceCallingPermission', 'Enforce Permission', 'info', 'Permission enforcement present'),
        
        # ========== M7: Client Code Quality ==========
        (r'WebView.*setJavaScriptEnabled\s*\(\s*true', 'JavaScript Enabled', 'medium', 'WebView with JavaScript - XSS risk if loading untrusted content'),
        (r'addJavascriptInterface', 'JavaScript Interface', 'high', 'WebView JavaScript bridge - code execution risk on API < 17'),
        (r'loadUrl\s*\([^)]*getIntent\(\)', 'Intent URL Loading', 'critical', 'Loading URLs from intent - injection vulnerability'),
        (r'loadUrl\s*\([^)]*\+', 'Dynamic URL Loading', 'high', 'Dynamic URL construction - verify input validation'),
        (r'evaluateJavascript\s*\([^)]*\+', 'JS Evaluation', 'high', 'Dynamic JavaScript evaluation - XSS risk'),
        (r'setAllowFileAccess\s*\(\s*true', 'WebView File Access', 'high', 'WebView file access enabled - local file theft risk'),
        (r'setAllowUniversalAccessFromFileURLs\s*\(\s*true', 'Universal File Access', 'critical', 'Universal file access in WebView - critical vulnerability'),
        (r'setAllowFileAccessFromFileURLs\s*\(\s*true', 'File URL Access', 'high', 'File URL access enabled - security risk'),
        
        # ========== M8: Code Tampering ==========
        (r'getPackageInfo\s*\([^)]*GET_SIGNATURES', 'Signature Check', 'info', 'App signature verification - anti-tampering'),
        (r'checkSignatures', 'Signature Comparison', 'info', 'Signature comparison for integrity'),
        (r'PackageManager\.GET_SIGNING_CERTIFICATES', 'Certificate Check', 'info', 'Certificate verification present'),
        
        # ========== M9: Reverse Engineering ==========
        (r'BuildConfig\.DEBUG', 'Debug Check', 'low', 'Debug build check - may expose debug functionality'),
        (r'android\.os\.Debug\.isDebuggerConnected', 'Debugger Detection', 'info', 'Debugger detection implemented'),
        (r'Debug\.waitingForDebugger', 'Debug Wait', 'info', 'Waiting for debugger attachment'),
        
        # ========== M10: Extraneous Functionality ==========
        (r'Log\.(d|v|i|w|e)\s*\(', 'Logging', 'low', 'Logging present - verify no sensitive data'),
        (r'System\.out\.print', 'System Print', 'low', 'System.out print - remove in production'),
        (r'printStackTrace\(\)', 'Stack Trace', 'low', 'Printing stack traces - information disclosure'),
        (r'Throwable.*getMessage\(\)', 'Exception Message', 'low', 'Exception messages may leak information'),
        
        # ========== Additional Security Patterns ==========
        (r'intent\.getStringExtra', 'Intent String Extra', 'info', 'Reading intent extras - validate input'),
        (r'getIntent\(\)\.getData\(\)', 'Intent Data', 'medium', 'Reading intent data URI - validate source'),
        (r'Class\.forName\s*\([^)]*\+', 'Dynamic Class Loading', 'high', 'Dynamic class loading - code injection risk'),
        (r'DexClassLoader|PathClassLoader', 'Custom ClassLoader', 'high', 'Custom class loader - verify loaded code source'),
        (r'ContentResolver.*query.*ContactsContract', 'Contacts Access', 'info', 'Accessing contacts data'),
        (r'ContentResolver.*query.*CallLog', 'Call Log Access', 'medium', 'Accessing call log data'),
        (r'SmsManager\.sendTextMessage', 'SMS Sending', 'medium', 'Sending SMS - verify user consent'),
        (r'TelephonyManager.*getDeviceId', 'Device ID Access', 'medium', 'Accessing device ID - privacy concern'),
        (r'LocationManager.*requestLocationUpdates', 'Location Tracking', 'medium', 'Location tracking - verify user consent'),
        (r'Camera\.open|CameraManager', 'Camera Access', 'medium', 'Camera access - verify user consent'),
        (r'MediaRecorder.*setAudioSource', 'Audio Recording', 'medium', 'Audio recording - verify user consent'),
        (r'ClipboardManager.*getPrimaryClip', 'Clipboard Read', 'medium', 'Reading clipboard - may contain sensitive data'),
        (r'createSocket.*\d+\.\d+\.\d+\.\d+', 'Hardcoded IP', 'medium', 'Hardcoded IP address detected'),
        (r'InetAddress\.getByName\s*\([^)]*"[^"]*\d+\.\d+\.\d+\.\d+', 'Hardcoded IP Address', 'medium', 'Hardcoded IP in network code'),
        (r'new\s+URL\s*\([^)]*"http:', 'Hardcoded HTTP URL', 'medium', 'Hardcoded HTTP URL - use HTTPS'),
        (r'@SuppressLint\s*\(\s*"[^"]*Security', 'Suppressed Security Lint', 'high', 'Security lint warning suppressed'),
        (r'catch\s*\(\s*Exception\s+\w+\s*\)\s*\{\s*\}', 'Empty Catch', 'medium', 'Empty exception handler - may hide errors'),
        (r'TODO|FIXME|HACK|XXX', 'Code Comment', 'low', 'Development comment found - review before release'),
    ]
    
    for pattern, issue_type, severity, description in patterns:
        if re.search(pattern, source_code, re.IGNORECASE):
            # Find line number
            for i, line in enumerate(source_code.split('\n'), 1):
                if re.search(pattern, line, re.IGNORECASE):
                    issues.append({
                        'type': issue_type,
                        'severity': severity,
                        'description': description,
                        'class': class_name,
                        'line': i,
                        'code_snippet': line.strip()[:100]
                    })
                    break
    
    return issues


# ============================================================================
# AI Code Analysis Functions
# ============================================================================

async def explain_code_with_ai(
    source_code: str,
    class_name: str,
    explanation_type: str = "general",
    method_name: Optional[str] = None
) -> Dict[str, Any]:
    """
    Use Gemini AI to explain decompiled Java/Kotlin code.
    
    Args:
        source_code: The decompiled source code
        class_name: Name of the class being analyzed
        explanation_type: Type of explanation (general, security, method)
        method_name: Optional method name for method-specific explanation
    
    Returns:
        Dictionary with explanation, key points, and security concerns
    """
    from google import genai
    from google.genai import types
    import json
    
    client = genai.Client(api_key=settings.gemini_api_key)
    
    # Truncate very long source code
    max_code_length = 15000
    if len(source_code) > max_code_length:
        source_code = source_code[:max_code_length] + "\n\n// ... [truncated for analysis]"
    
    if explanation_type == "method" and method_name:
        prompt = f"""You are an expert Android reverse engineer. Analyze this specific method from decompiled Java/Android code.

Class: {class_name}
Method: {method_name}

```java
{source_code}
```

Provide a detailed analysis in the following JSON format:
{{
    "explanation": "Clear explanation of what this method does, its purpose, and how it works (2-4 paragraphs)",
    "key_points": [
        "Key point 1 about the method's functionality",
        "Key point 2 about inputs/outputs",
        "Key point 3 about side effects or behavior"
    ],
    "security_concerns": [
        {{
            "concern": "Security concern description",
            "severity": "high/medium/low",
            "line_hint": "relevant code pattern",
            "recommendation": "How to mitigate"
        }}
    ]
}}

Focus on:
1. What the method accomplishes
2. Input validation and data handling
3. Any security implications
4. Interaction with Android APIs"""

    elif explanation_type == "security":
        prompt = f"""You are an expert Android security researcher. Perform a security-focused analysis of this decompiled Java/Android code.

Class: {class_name}

```java
{source_code}
```

Provide a security analysis in the following JSON format:
{{
    "explanation": "Security-focused overview of what this class does and potential risks (2-3 paragraphs)",
    "key_points": [
        "Key security-relevant functionality 1",
        "Key security-relevant functionality 2",
        "Data handling patterns",
        "Permission usage"
    ],
    "security_concerns": [
        {{
            "concern": "Specific security vulnerability or concern",
            "severity": "critical/high/medium/low",
            "line_hint": "code pattern that indicates this",
            "recommendation": "How to exploit or mitigate"
        }}
    ]
}}

Focus on:
1. Sensitive data handling
2. Authentication/authorization
3. Cryptographic operations
4. Network communications
5. Intent handling and IPC
6. File operations
7. SQL queries
8. WebView usage"""

    else:  # general
        prompt = f"""You are an expert Android reverse engineer helping to understand decompiled Java/Android code.

Class: {class_name}

```java
{source_code}
```

Provide a clear explanation in the following JSON format:
{{
    "explanation": "Clear, detailed explanation of what this class does and its purpose (2-4 paragraphs). Explain as if to someone who wants to understand the app's behavior.",
    "key_points": [
        "Main purpose/responsibility of this class",
        "Key methods and what they do",
        "Data this class handles",
        "How it interacts with other components"
    ],
    "security_concerns": [
        {{
            "concern": "Any notable security concern",
            "severity": "high/medium/low",
            "line_hint": "relevant code pattern",
            "recommendation": "Mitigation suggestion"
        }}
    ]
}}

Consider:
1. Is this an Activity, Service, BroadcastReceiver, or ContentProvider?
2. What user-facing functionality does it provide?
3. What data does it process or store?
4. What Android APIs does it use?"""

    try:
        response = client.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        # Parse JSON from response
        response_text = response.text
        
        # Try to extract JSON
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            result = json.loads(json_match.group())
        else:
            # Fallback - create structured response from text
            result = {
                "explanation": response_text,
                "key_points": [],
                "security_concerns": []
            }
        
        return {
            "class_name": class_name,
            "explanation_type": explanation_type,
            "explanation": result.get("explanation", "Analysis complete."),
            "key_points": result.get("key_points", []),
            "security_concerns": result.get("security_concerns", []),
            "method_name": method_name
        }
        
    except Exception as e:
        logger.error(f"AI code explanation failed: {e}")
        return {
            "class_name": class_name,
            "explanation_type": explanation_type,
            "explanation": f"AI analysis failed: {str(e)}",
            "key_points": [],
            "security_concerns": [],
            "method_name": method_name
        }


async def analyze_code_vulnerabilities_with_ai(
    source_code: str,
    class_name: str
) -> Dict[str, Any]:
    """
    Use Gemini AI to perform deep vulnerability analysis on decompiled code.
    
    Args:
        source_code: The decompiled source code
        class_name: Name of the class being analyzed
    
    Returns:
        Dictionary with vulnerabilities, recommendations, and exploitation scenarios
    """
    from google import genai
    from google.genai import types
    import json
    
    client = genai.Client(api_key=settings.gemini_api_key)
    
    # Truncate very long source code
    max_code_length = 15000
    if len(source_code) > max_code_length:
        source_code = source_code[:max_code_length] + "\n\n// ... [truncated for analysis]"
    
    prompt = f"""You are an expert Android security researcher and penetration tester. Perform a comprehensive vulnerability analysis of this decompiled Android code.

Class: {class_name}

```java
{source_code}
```

Analyze for ALL potential security vulnerabilities and provide your findings in this JSON format:
{{
    "risk_level": "critical/high/medium/low/info",
    "summary": "Executive summary of security posture (2-3 sentences)",
    "vulnerabilities": [
        {{
            "id": "VULN-001",
            "title": "Vulnerability title",
            "severity": "critical/high/medium/low",
            "category": "OWASP category or type",
            "description": "Detailed description of the vulnerability",
            "affected_code": "The vulnerable code pattern or line",
            "impact": "What an attacker could achieve",
            "cvss_estimate": "Estimated CVSS score if applicable"
        }}
    ],
    "recommendations": [
        "Specific recommendation 1 with code fix suggestion",
        "Specific recommendation 2",
        "General security improvement"
    ],
    "exploitation_scenarios": [
        "Step-by-step exploitation scenario 1",
        "How an attacker could chain vulnerabilities"
    ]
}}

Check for vulnerabilities including:
1. **Injection Flaws**: SQL injection, command injection, LDAP injection
2. **Insecure Data Storage**: SharedPreferences, databases, files, logs
3. **Insecure Communication**: HTTP, certificate validation, hostname verification
4. **Insecure Authentication**: Weak passwords, hardcoded credentials, biometric bypass
5. **Insufficient Cryptography**: Weak algorithms, static keys, predictable IVs
6. **Insecure Authorization**: Exported components, permission issues
7. **Client Code Quality**: WebView vulnerabilities, JavaScript interfaces
8. **Code Tampering**: Missing integrity checks
9. **Reverse Engineering**: Debug code, logging sensitive data
10. **Extraneous Functionality**: Hidden backdoors, test code

Be thorough and identify ALL potential security issues, even low-severity ones."""

    try:
        response = client.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        response_text = response.text
        
        # Try to extract JSON
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            result = json.loads(json_match.group())
        else:
            result = {
                "risk_level": "unknown",
                "summary": response_text[:500],
                "vulnerabilities": [],
                "recommendations": [],
                "exploitation_scenarios": []
            }
        
        return {
            "class_name": class_name,
            "risk_level": result.get("risk_level", "unknown"),
            "vulnerabilities": result.get("vulnerabilities", []),
            "recommendations": result.get("recommendations", []),
            "exploitation_scenarios": result.get("exploitation_scenarios", []),
            "summary": result.get("summary", "Analysis complete.")
        }
        
    except Exception as e:
        logger.error(f"AI vulnerability analysis failed: {e}")
        return {
            "class_name": class_name,
            "risk_level": "error",
            "vulnerabilities": [],
            "recommendations": [],
            "exploitation_scenarios": [],
            "summary": f"AI analysis failed: {str(e)}"
        }


# ============================================================================
# Data Flow Analysis
# ============================================================================

# Data sources - where sensitive data originates
DATA_SOURCES = {
    # User Input Sources
    "user_input": [
        r"getIntent\(\)",
        r"getExtras\(\)",
        r"getStringExtra\(",
        r"getBundleExtra\(",
        r"getParcelableExtra\(",
        r"getText\(\)",
        r"getEditableText\(\)",
        r"EditText.*\.getText\(",
        r"getClipboardManager\(\)",
        r"ClipData",
        r"onActivityResult\(",
        r"ContentResolver.*query\(",
    ],
    # File Sources
    "file_input": [
        r"FileInputStream",
        r"BufferedReader",
        r"InputStreamReader",
        r"openFileInput\(",
        r"getAssets\(\)\.open\(",
        r"getResources\(\)\.openRawResource\(",
        r"readFile\(",
        r"Files\.readAllBytes\(",
        r"Scanner.*new File\(",
    ],
    # Network Sources
    "network_input": [
        r"HttpURLConnection.*getInputStream\(",
        r"URLConnection.*getInputStream\(",
        r"Socket.*getInputStream\(",
        r"OkHttpClient",
        r"Retrofit",
        r"Volley",
        r"WebSocket.*onMessage\(",
        r"Response\.body\(",
    ],
    # Database Sources
    "database_input": [
        r"SQLiteDatabase.*query\(",
        r"rawQuery\(",
        r"Cursor.*getString\(",
        r"Cursor.*getInt\(",
        r"Room.*Dao",
        r"ContentProvider.*query\(",
    ],
    # Sensitive Data Sources
    "sensitive_data": [
        r"getDeviceId\(",
        r"getSubscriberId\(",
        r"getLine1Number\(",
        r"getSimSerialNumber\(",
        r"getMacAddress\(",
        r"getLastKnownLocation\(",
        r"getLatitude\(",
        r"getLongitude\(",
        r"AccountManager.*getAccounts\(",
        r"getInstalledPackages\(",
        r"SmsManager.*getAllMessagesFromIcc\(",
        r"ContactsContract",
        r"CalendarContract",
        r"MediaStore",
    ],
    # Crypto Sources
    "crypto_input": [
        r"SecretKey",
        r"PrivateKey",
        r"getEncoded\(",
        r"KeyStore.*getKey\(",
        r"Cipher.*doFinal\(",
    ],
}

# Data sinks - where data flows to (potential leakage points)
DATA_SINKS = {
    # Logging Sinks
    "logging": [
        r"Log\.[vdiwea]\(",
        r"System\.out\.print",
        r"System\.err\.print",
        r"printStackTrace\(",
        r"logger\.",
        r"Timber\.[dviwe]\(",
    ],
    # Network Sinks
    "network_output": [
        r"HttpURLConnection.*getOutputStream\(",
        r"URLConnection.*getOutputStream\(",
        r"Socket.*getOutputStream\(",
        r"DataOutputStream.*write",
        r"OutputStream.*write\(",
        r"OkHttpClient.*newCall\(",
        r"RequestBody\.create\(",
        r"WebSocket.*send\(",
    ],
    # File Sinks
    "file_output": [
        r"FileOutputStream",
        r"BufferedWriter",
        r"OutputStreamWriter",
        r"openFileOutput\(",
        r"FileWriter",
        r"Files\.write\(",
        r"PrintWriter",
    ],
    # Database Sinks
    "database_output": [
        r"SQLiteDatabase.*insert\(",
        r"SQLiteDatabase.*update\(",
        r"execSQL\(",
        r"ContentValues.*put\(",
        r"Room.*insert\(",
    ],
    # SharedPreferences Sinks
    "preferences": [
        r"SharedPreferences.*edit\(",
        r"putString\(",
        r"putInt\(",
        r"putBoolean\(",
        r"commit\(\)",
        r"apply\(\)",
    ],
    # IPC Sinks
    "ipc_output": [
        r"startActivity\(",
        r"startService\(",
        r"sendBroadcast\(",
        r"sendOrderedBroadcast\(",
        r"bindService\(",
        r"setResult\(",
        r"ContentResolver.*insert\(",
        r"ContentResolver.*update\(",
    ],
    # Clipboard Sinks
    "clipboard": [
        r"ClipboardManager.*setPrimaryClip\(",
        r"ClipData\.newPlainText\(",
    ],
    # WebView Sinks
    "webview": [
        r"loadUrl\(",
        r"loadData\(",
        r"evaluateJavascript\(",
        r"addJavascriptInterface\(",
    ],
}

# Taint propagation patterns - how data flows between variables
TAINT_PROPAGATORS = [
    r"(\w+)\s*=\s*(\w+)\.toString\(\)",
    r"(\w+)\s*=\s*String\.valueOf\((\w+)\)",
    r"(\w+)\s*=\s*(\w+)\s*\+",
    r"(\w+)\s*=\s*new\s+String\((\w+)\)",
    r"(\w+)\s*=\s*(\w+)\.getBytes\(",
    r"(\w+)\s*=\s*Base64\.encode\((\w+)",
    r"(\w+)\s*=\s*URLEncoder\.encode\((\w+)",
    r"(\w+)\s*=\s*(\w+)\.substring\(",
    r"(\w+)\s*=\s*(\w+)\.split\(",
    r"(\w+)\s*=\s*(\w+)\.trim\(",
    r"(\w+)\s*=\s*(\w+)\.toLowerCase\(",
    r"(\w+)\s*=\s*(\w+)\.toUpperCase\(",
    r"StringBuilder.*append\((\w+)\)",
    r"(\w+)\.put\([^,]+,\s*(\w+)\)",
]


def analyze_data_flow(source_code: str, class_name: str) -> Dict[str, Any]:
    """
    Analyze data flow in decompiled Java source code.
    
    Performs lightweight taint analysis to track:
    - Data sources (where sensitive data originates)
    - Data sinks (where data flows to)
    - Potential data leakage paths
    """
    results = {
        "class_name": class_name,
        "sources": [],
        "sinks": [],
        "flows": [],
        "risk_flows": [],
        "summary": {
            "total_sources": 0,
            "total_sinks": 0,
            "potential_leaks": 0,
            "risk_level": "low"
        }
    }
    
    lines = source_code.split('\n')
    
    # Track tainted variables (variable_name -> source_type, line)
    tainted_vars: Dict[str, tuple] = {}
    
    # Find all data sources
    for category, patterns in DATA_SOURCES.items():
        for pattern in patterns:
            for line_num, line in enumerate(lines, 1):
                matches = re.finditer(pattern, line)
                for match in matches:
                    # Try to find the variable being assigned
                    assign_match = re.search(r'(\w+)\s*=', line[:match.start()])
                    var_name = assign_match.group(1) if assign_match else None
                    
                    source_entry = {
                        "type": category,
                        "pattern": pattern,
                        "line": line_num,
                        "code": line.strip(),
                        "variable": var_name
                    }
                    results["sources"].append(source_entry)
                    
                    # Mark variable as tainted
                    if var_name:
                        tainted_vars[var_name] = (category, line_num)
    
    # Propagate taint through assignments
    for line_num, line in enumerate(lines, 1):
        for prop_pattern in TAINT_PROPAGATORS:
            prop_matches = re.finditer(prop_pattern, line)
            for match in prop_matches:
                groups = match.groups()
                if len(groups) >= 2:
                    target_var = groups[0]
                    source_var = groups[1]
                    if source_var in tainted_vars:
                        tainted_vars[target_var] = tainted_vars[source_var]
    
    # Find all data sinks and check if tainted data flows to them
    for category, patterns in DATA_SINKS.items():
        for pattern in patterns:
            for line_num, line in enumerate(lines, 1):
                matches = re.finditer(pattern, line)
                for match in matches:
                    sink_entry = {
                        "type": category,
                        "pattern": pattern,
                        "line": line_num,
                        "code": line.strip()
                    }
                    results["sinks"].append(sink_entry)
                    
                    # Check if any tainted variable is used in this sink
                    for var_name, (source_type, source_line) in tainted_vars.items():
                        if re.search(rf'\b{re.escape(var_name)}\b', line):
                            flow = {
                                "source": {
                                    "type": source_type,
                                    "variable": var_name,
                                    "line": source_line
                                },
                                "sink": {
                                    "type": category,
                                    "line": line_num,
                                    "code": line.strip()
                                },
                                "risk": _calculate_flow_risk(source_type, category)
                            }
                            results["flows"].append(flow)
                            
                            # Track high-risk flows separately
                            if flow["risk"] in ["high", "critical"]:
                                results["risk_flows"].append(flow)
    
    # Calculate summary
    results["summary"]["total_sources"] = len(results["sources"])
    results["summary"]["total_sinks"] = len(results["sinks"])
    results["summary"]["potential_leaks"] = len(results["risk_flows"])
    
    # Determine overall risk level
    critical_flows = sum(1 for f in results["flows"] if f["risk"] == "critical")
    high_flows = sum(1 for f in results["flows"] if f["risk"] == "high")
    
    if critical_flows > 0:
        results["summary"]["risk_level"] = "critical"
    elif high_flows > 0:
        results["summary"]["risk_level"] = "high"
    elif len(results["flows"]) > 5:
        results["summary"]["risk_level"] = "medium"
    else:
        results["summary"]["risk_level"] = "low"
    
    return results


def _calculate_flow_risk(source_type: str, sink_type: str) -> str:
    """Calculate the risk level of a data flow from source to sink."""
    # High-risk combinations
    critical_combinations = [
        ("sensitive_data", "network_output"),
        ("sensitive_data", "logging"),
        ("crypto_input", "logging"),
        ("crypto_input", "network_output"),
        ("user_input", "webview"),  # XSS risk
    ]
    
    high_risk_combinations = [
        ("sensitive_data", "file_output"),
        ("sensitive_data", "ipc_output"),
        ("database_input", "logging"),
        ("user_input", "database_output"),  # SQL injection risk
        ("network_input", "webview"),
        ("file_input", "webview"),
    ]
    
    medium_risk_combinations = [
        ("user_input", "logging"),
        ("database_input", "network_output"),
        ("file_input", "network_output"),
    ]
    
    combo = (source_type, sink_type)
    
    if combo in critical_combinations:
        return "critical"
    elif combo in high_risk_combinations:
        return "high"
    elif combo in medium_risk_combinations:
        return "medium"
    else:
        return "low"


# ============================================================================
# Method Call Graph Analysis
# ============================================================================

def build_call_graph(source_code: str, class_name: str) -> Dict[str, Any]:
    """
    Build a method call graph from decompiled Java source code.
    
    Extracts:
    - Method definitions and their signatures
    - Method calls within each method
    - Call relationships (caller -> callee)
    - Entry points (lifecycle methods, exported components)
    """
    results = {
        "class_name": class_name,
        "methods": [],
        "calls": [],
        "entry_points": [],
        "external_calls": [],
        "graph": {
            "nodes": [],
            "edges": []
        },
        "statistics": {
            "total_methods": 0,
            "total_internal_calls": 0,
            "total_external_calls": 0,
            "max_depth": 0,
            "cyclomatic_complexity": 0
        }
    }
    
    lines = source_code.split('\n')
    
    # Method definition patterns
    method_pattern = re.compile(
        r'(?:public|private|protected|static|final|native|synchronized|abstract|transient|\s)*'
        r'(?:<[\w\s,<>?]+>\s+)?'  # Generic types
        r'(\w+(?:<[\w\s,<>?]+>)?)\s+'  # Return type
        r'(\w+)\s*'  # Method name
        r'\(([^)]*)\)\s*'  # Parameters
        r'(?:throws\s+[\w\s,]+)?'  # Throws clause
        r'\s*\{'  # Opening brace
    )
    
    # Entry point patterns (Android lifecycle methods)
    entry_point_methods = {
        "onCreate", "onStart", "onResume", "onPause", "onStop", "onDestroy",
        "onCreateView", "onViewCreated", "onAttach", "onDetach",
        "onReceive", "onBind", "onStartCommand", "onHandleIntent",
        "onClick", "onLongClick", "onTouch", "onItemClick",
        "onOptionsItemSelected", "onCreateOptionsMenu",
        "handleMessage", "run", "call",
        "doInBackground", "onPreExecute", "onPostExecute",
        "query", "insert", "update", "delete", "getType",  # ContentProvider
        "onNewIntent", "onActivityResult",
    }
    
    # Parse methods
    current_method = None
    brace_count = 0
    method_start_line = 0
    method_body_lines = []
    
    for line_num, line in enumerate(lines, 1):
        # Check for method definition
        method_match = method_pattern.search(line)
        if method_match and brace_count == 0:
            return_type = method_match.group(1)
            method_name = method_match.group(2)
            params = method_match.group(3)
            
            # Skip constructors that look like class name
            if method_name == class_name.split('.')[-1]:
                method_name = "<init>"
            
            current_method = {
                "name": method_name,
                "return_type": return_type,
                "parameters": _parse_parameters(params),
                "line_start": line_num,
                "line_end": 0,
                "is_entry_point": method_name in entry_point_methods,
                "calls": [],
                "called_by": [],
                "modifiers": _extract_modifiers(line)
            }
            method_start_line = line_num
            method_body_lines = []
            brace_count = line.count('{') - line.count('}')
            
        elif current_method:
            method_body_lines.append(line)
            brace_count += line.count('{') - line.count('}')
            
            if brace_count == 0:
                current_method["line_end"] = line_num
                current_method["calls"] = _extract_method_calls(
                    '\n'.join(method_body_lines), 
                    class_name
                )
                results["methods"].append(current_method)
                
                if current_method["is_entry_point"]:
                    results["entry_points"].append({
                        "name": current_method["name"],
                        "line": current_method["line_start"],
                        "type": _get_entry_point_type(current_method["name"])
                    })
                
                current_method = None
    
    # Build call relationships
    method_names = {m["name"] for m in results["methods"]}
    
    for method in results["methods"]:
        for call in method["calls"]:
            call_entry = {
                "caller": method["name"],
                "caller_line": method["line_start"],
                "callee": call["method"],
                "callee_class": call.get("class", class_name),
                "line": call["line"],
                "is_internal": call["method"] in method_names
            }
            
            results["calls"].append(call_entry)
            
            if not call_entry["is_internal"]:
                results["external_calls"].append(call_entry)
            
            # Add edge to graph
            results["graph"]["edges"].append({
                "from": method["name"],
                "to": call["method"],
                "label": f"line {call['line']}"
            })
    
    # Build graph nodes
    for method in results["methods"]:
        node = {
            "id": method["name"],
            "label": method["name"],
            "type": "internal",
            "is_entry_point": method["is_entry_point"],
            "line": method["line_start"]
        }
        results["graph"]["nodes"].append(node)
    
    # Add external method nodes
    external_methods = set()
    for call in results["external_calls"]:
        external_id = f"{call['callee_class']}.{call['callee']}"
        if external_id not in external_methods:
            external_methods.add(external_id)
            results["graph"]["nodes"].append({
                "id": external_id,
                "label": call["callee"],
                "type": "external",
                "is_entry_point": False,
                "class": call["callee_class"]
            })
    
    # Calculate statistics
    results["statistics"]["total_methods"] = len(results["methods"])
    results["statistics"]["total_internal_calls"] = sum(
        1 for c in results["calls"] if c["is_internal"]
    )
    results["statistics"]["total_external_calls"] = len(results["external_calls"])
    results["statistics"]["max_depth"] = _calculate_max_call_depth(results["methods"])
    results["statistics"]["cyclomatic_complexity"] = _estimate_cyclomatic_complexity(source_code)
    
    return results


def _parse_parameters(params_str: str) -> List[Dict[str, str]]:
    """Parse method parameters from string."""
    if not params_str.strip():
        return []
    
    params = []
    # Split by comma, but handle generics
    depth = 0
    current = ""
    
    for char in params_str:
        if char == '<':
            depth += 1
        elif char == '>':
            depth -= 1
        elif char == ',' and depth == 0:
            if current.strip():
                params.append(_parse_single_param(current.strip()))
            current = ""
            continue
        current += char
    
    if current.strip():
        params.append(_parse_single_param(current.strip()))
    
    return params


def _parse_single_param(param_str: str) -> Dict[str, str]:
    """Parse a single parameter."""
    parts = param_str.split()
    if len(parts) >= 2:
        return {"type": ' '.join(parts[:-1]), "name": parts[-1]}
    elif len(parts) == 1:
        return {"type": parts[0], "name": ""}
    return {"type": "unknown", "name": ""}


def _extract_modifiers(line: str) -> List[str]:
    """Extract method modifiers from line."""
    modifiers = []
    modifier_keywords = ["public", "private", "protected", "static", "final", 
                        "native", "synchronized", "abstract"]
    for mod in modifier_keywords:
        if re.search(rf'\b{mod}\b', line):
            modifiers.append(mod)
    return modifiers


def _extract_method_calls(code: str, current_class: str) -> List[Dict[str, Any]]:
    """Extract method calls from code block."""
    calls = []
    lines = code.split('\n')
    
    # Method call patterns
    # this.method() or method()
    internal_call = re.compile(r'(?:this\.)?(\w+)\s*\(')
    # object.method()
    object_call = re.compile(r'(\w+)\.(\w+)\s*\(')
    # Class.staticMethod()
    static_call = re.compile(r'([A-Z]\w+)\.(\w+)\s*\(')
    # new Constructor()
    constructor_call = re.compile(r'new\s+(\w+)\s*\(')
    # super.method()
    super_call = re.compile(r'super\.(\w+)\s*\(')
    
    for line_num, line in enumerate(lines, 1):
        # Skip comments
        if line.strip().startswith('//') or line.strip().startswith('*'):
            continue
        
        # Constructor calls
        for match in constructor_call.finditer(line):
            calls.append({
                "method": f"<init>",
                "class": match.group(1),
                "line": line_num,
                "type": "constructor"
            })
        
        # Static method calls
        for match in static_call.finditer(line):
            class_name = match.group(1)
            method_name = match.group(2)
            # Skip common non-method patterns
            if method_name not in ['class', 'this']:
                calls.append({
                    "method": method_name,
                    "class": class_name,
                    "line": line_num,
                    "type": "static"
                })
        
        # Object method calls (excluding static ones already captured)
        for match in object_call.finditer(line):
            obj_name = match.group(1)
            method_name = match.group(2)
            # Skip if it's a static call we already captured
            if not obj_name[0].isupper():
                calls.append({
                    "method": method_name,
                    "class": obj_name,
                    "line": line_num,
                    "type": "instance"
                })
        
        # Super calls
        for match in super_call.finditer(line):
            calls.append({
                "method": match.group(1),
                "class": "super",
                "line": line_num,
                "type": "super"
            })
    
    return calls


def _get_entry_point_type(method_name: str) -> str:
    """Get the type of entry point based on method name."""
    activity_methods = {"onCreate", "onStart", "onResume", "onPause", "onStop", 
                       "onDestroy", "onNewIntent", "onActivityResult"}
    fragment_methods = {"onCreateView", "onViewCreated", "onAttach", "onDetach"}
    service_methods = {"onBind", "onStartCommand", "onHandleIntent"}
    receiver_methods = {"onReceive"}
    provider_methods = {"query", "insert", "update", "delete", "getType"}
    ui_methods = {"onClick", "onLongClick", "onTouch", "onItemClick"}
    thread_methods = {"run", "call", "doInBackground"}
    
    if method_name in activity_methods:
        return "activity_lifecycle"
    elif method_name in fragment_methods:
        return "fragment_lifecycle"
    elif method_name in service_methods:
        return "service_lifecycle"
    elif method_name in receiver_methods:
        return "broadcast_receiver"
    elif method_name in provider_methods:
        return "content_provider"
    elif method_name in ui_methods:
        return "ui_callback"
    elif method_name in thread_methods:
        return "async_task"
    else:
        return "other"


def _calculate_max_call_depth(methods: List[Dict]) -> int:
    """Calculate the maximum call depth in the call graph."""
    # Build adjacency list
    adjacency = {}
    for method in methods:
        adjacency[method["name"]] = [c["method"] for c in method["calls"]]
    
    def dfs_depth(node: str, visited: set) -> int:
        if node in visited or node not in adjacency:
            return 0
        visited.add(node)
        max_child = 0
        for child in adjacency.get(node, []):
            max_child = max(max_child, dfs_depth(child, visited.copy()))
        return 1 + max_child
    
    max_depth = 0
    for method in methods:
        if method["is_entry_point"]:
            depth = dfs_depth(method["name"], set())
            max_depth = max(max_depth, depth)
    
    return max_depth


def _estimate_cyclomatic_complexity(source_code: str) -> int:
    """Estimate cyclomatic complexity of the code."""
    complexity = 1  # Base complexity
    
    # Decision points
    decision_patterns = [
        r'\bif\s*\(',
        r'\belse\s+if\s*\(',
        r'\bfor\s*\(',
        r'\bwhile\s*\(',
        r'\bcase\s+',
        r'\bcatch\s*\(',
        r'\?\s*.*\s*:',  # Ternary operator
        r'\&\&',
        r'\|\|',
    ]
    
    for pattern in decision_patterns:
        complexity += len(re.findall(pattern, source_code))
    
    return complexity


# ============================================================================
# Smart Search - Semantic Code Search
# ============================================================================

# Security-related keywords and their synonyms/related terms
SECURITY_KEYWORD_EXPANSIONS = {
    "password": ["password", "passwd", "pwd", "secret", "credential", "auth", "login"],
    "crypto": ["encrypt", "decrypt", "cipher", "aes", "rsa", "des", "hash", "md5", "sha", "crypto"],
    "network": ["http", "https", "url", "socket", "connection", "request", "response", "api", "endpoint"],
    "storage": ["file", "database", "sqlite", "sharedpreferences", "cache", "store", "save", "write", "read"],
    "auth": ["authentication", "authorization", "login", "logout", "session", "token", "jwt", "oauth", "sso"],
    "injection": ["sql", "query", "execute", "command", "shell", "runtime", "exec", "eval"],
    "sensitive": ["private", "secret", "key", "api_key", "apikey", "token", "credential", "ssn", "credit"],
    "webview": ["webview", "javascript", "loadurl", "addjavascriptinterface", "evaluatejavascript"],
    "intent": ["intent", "broadcast", "startactivity", "startservice", "bindservice", "pendingintent"],
    "permission": ["permission", "checkpermission", "requestpermission", "granted", "denied"],
}

# Vulnerability patterns with descriptions
SMART_SEARCH_VULN_PATTERNS = {
    "hardcoded_secret": {
        "patterns": [
            r'(?:password|passwd|pwd|secret|key|token|api[_-]?key)\s*=\s*["\'][^"\']{4,}["\']',
            r'private\s+(?:static\s+)?(?:final\s+)?String\s+\w*(?:KEY|SECRET|PASSWORD|TOKEN)\w*\s*=',
        ],
        "description": "Hardcoded secrets or credentials",
        "severity": "high"
    },
    "sql_injection": {
        "patterns": [
            r'rawQuery\s*\([^)]*\+',
            r'execSQL\s*\([^)]*\+',
            r'query\s*\([^)]*\+\s*\w+',
        ],
        "description": "Potential SQL injection vulnerability",
        "severity": "critical"
    },
    "insecure_random": {
        "patterns": [
            r'new\s+Random\s*\(',
            r'Math\.random\s*\(',
        ],
        "description": "Insecure random number generation",
        "severity": "medium"
    },
    "weak_crypto": {
        "patterns": [
            r'DES[/"\'.]',
            r'MD5',
            r'SHA-?1[^0-9]',
            r'ECB',
            r'RC4',
        ],
        "description": "Weak cryptographic algorithm",
        "severity": "high"
    },
    "insecure_network": {
        "patterns": [
            r'http://[^"\']*(?!localhost)',
            r'setHostnameVerifier\s*\(\s*.*ALLOW_ALL',
            r'trustAllCerts',
            r'X509TrustManager.*checkServerTrusted.*\{\s*\}',
        ],
        "description": "Insecure network communication",
        "severity": "high"
    },
    "exported_component": {
        "patterns": [
            r'android:exported\s*=\s*["\']true["\']',
        ],
        "description": "Exported Android component",
        "severity": "medium"
    },
    "webview_risk": {
        "patterns": [
            r'setJavaScriptEnabled\s*\(\s*true\s*\)',
            r'addJavascriptInterface\s*\(',
            r'setAllowFileAccess\s*\(\s*true\s*\)',
        ],
        "description": "Potentially risky WebView configuration",
        "severity": "medium"
    },
    "logging_sensitive": {
        "patterns": [
            r'Log\.[dviwe]\s*\([^)]*(?:password|secret|key|token|credential)',
        ],
        "description": "Sensitive data in logs",
        "severity": "high"
    },
}


def smart_search(output_dir: Path, query: str, search_type: str = "smart", max_results: int = 100) -> Dict[str, Any]:
    """
    Perform smart/semantic search across decompiled sources.
    
    Search types:
    - "smart": Expands query with related security terms
    - "vuln": Searches for vulnerability patterns
    - "regex": Direct regex search
    - "exact": Exact string match
    """
    results = {
        "query": query,
        "search_type": search_type,
        "total_matches": 0,
        "files_searched": 0,
        "matches": [],
        "vulnerability_summary": {},
        "suggestions": []
    }
    
    sources_dir = output_dir / "sources"
    if not sources_dir.exists():
        return results
    
    # Build search patterns based on type
    search_patterns = []
    
    if search_type == "smart":
        # Expand query with related terms
        query_lower = query.lower()
        expanded_terms = [query]
        
        for key, synonyms in SECURITY_KEYWORD_EXPANSIONS.items():
            if key in query_lower or any(s in query_lower for s in synonyms):
                expanded_terms.extend(synonyms)
        
        # Remove duplicates and create patterns
        expanded_terms = list(set(expanded_terms))
        search_patterns = [re.compile(rf'\b{re.escape(term)}\b', re.IGNORECASE) for term in expanded_terms]
        results["expanded_terms"] = expanded_terms
        
    elif search_type == "vuln":
        # Search for vulnerability patterns
        vuln_results = {}
        for vuln_name, vuln_info in SMART_SEARCH_VULN_PATTERNS.items():
            if query.lower() in vuln_name or query.lower() == "all":
                for pattern in vuln_info["patterns"]:
                    search_patterns.append((
                        re.compile(pattern, re.IGNORECASE),
                        vuln_name,
                        vuln_info["description"],
                        vuln_info["severity"]
                    ))
        
    elif search_type == "regex":
        try:
            search_patterns = [re.compile(query, re.IGNORECASE)]
        except re.error:
            results["error"] = "Invalid regex pattern"
            return results
            
    else:  # exact
        search_patterns = [re.compile(re.escape(query), re.IGNORECASE)]
    
    # Search through all Java files
    for java_file in sources_dir.rglob("*.java"):
        if len(results["matches"]) >= max_results:
            break
            
        results["files_searched"] += 1
        
        try:
            content = java_file.read_text(encoding='utf-8', errors='ignore')
            lines = content.split('\n')
            rel_path = str(java_file.relative_to(sources_dir))
            
            for line_num, line in enumerate(lines, 1):
                if len(results["matches"]) >= max_results:
                    break
                    
                for pattern_item in search_patterns:
                    # Handle vuln search with extra info
                    if search_type == "vuln" and isinstance(pattern_item, tuple):
                        pattern, vuln_name, description, severity = pattern_item
                        match = pattern.search(line)
                        if match:
                            results["matches"].append({
                                "file": rel_path,
                                "line": line_num,
                                "code": line.strip(),
                                "match": match.group(),
                                "vuln_type": vuln_name,
                                "description": description,
                                "severity": severity
                            })
                            results["total_matches"] += 1
                            
                            # Track vulnerability summary
                            if vuln_name not in results["vulnerability_summary"]:
                                results["vulnerability_summary"][vuln_name] = {
                                    "count": 0,
                                    "severity": severity,
                                    "description": description
                                }
                            results["vulnerability_summary"][vuln_name]["count"] += 1
                    else:
                        # Regular search
                        pattern = pattern_item
                        match = pattern.search(line)
                        if match:
                            # Get context lines
                            context_start = max(0, line_num - 3)
                            context_end = min(len(lines), line_num + 2)
                            context = '\n'.join(lines[context_start:context_end])
                            
                            results["matches"].append({
                                "file": rel_path,
                                "line": line_num,
                                "code": line.strip(),
                                "match": match.group(),
                                "context": context
                            })
                            results["total_matches"] += 1
                            break  # Don't duplicate matches on same line
                            
        except Exception as e:
            logger.warning(f"Error searching file {java_file}: {e}")
    
    # Add search suggestions
    if results["total_matches"] == 0:
        results["suggestions"] = [
            "Try using 'smart' search type to expand your query",
            "Use 'vuln' search type to find security vulnerabilities",
            "Check if the search term exists in the codebase"
        ]
    
    return results


# ============================================================================
# AI Vulnerability Scan - Cross-Class Analysis
# ============================================================================

async def ai_vulnerability_scan(
    output_dir: Path,
    scan_type: str = "quick",
    focus_areas: List[str] = None
) -> Dict[str, Any]:
    """
    Perform AI-powered vulnerability scan across multiple classes.
    
    Scan types:
    - "quick": Scan key classes (activities, services, network)
    - "deep": Scan all classes for vulnerabilities
    - "focused": Scan specific areas (auth, crypto, network, storage)
    """
    from google import genai
    from google.genai import types
    
    if not settings.gemini_api_key:
        return {
            "error": "AI features require Gemini API key",
            "scan_type": scan_type,
            "vulnerabilities": [],
            "summary": "AI scan unavailable"
        }
    
    client = genai.Client(api_key=settings.gemini_api_key)
    
    results = {
        "scan_type": scan_type,
        "focus_areas": focus_areas or [],
        "classes_scanned": 0,
        "vulnerabilities": [],
        "risk_summary": {
            "critical": 0,
            "high": 0,
            "medium": 0,
            "low": 0
        },
        "attack_chains": [],
        "recommendations": [],
        "summary": ""
    }
    
    sources_dir = output_dir / "sources"
    if not sources_dir.exists():
        results["error"] = "No decompiled sources found"
        return results
    
    # Collect relevant classes based on scan type
    classes_to_scan = []
    
    # Priority patterns for quick scan
    priority_patterns = {
        "activities": [r"Activity\.java$", r"Fragment\.java$"],
        "auth": [r"(?i)login|auth|signin|signup|credential|password"],
        "network": [r"(?i)api|http|network|request|retrofit|okhttp"],
        "crypto": [r"(?i)crypto|cipher|encrypt|decrypt|key|hash"],
        "storage": [r"(?i)database|sqlite|preference|storage|cache"],
        "receivers": [r"Receiver\.java$", r"(?i)broadcast"],
        "services": [r"Service\.java$"],
    }
    
    # Determine which patterns to use
    if scan_type == "quick":
        patterns_to_use = ["activities", "auth", "network"]
    elif scan_type == "focused" and focus_areas:
        patterns_to_use = [a for a in focus_areas if a in priority_patterns]
    else:  # deep
        patterns_to_use = list(priority_patterns.keys())
    
    # Find matching files
    all_java_files = list(sources_dir.rglob("*.java"))
    
    for java_file in all_java_files:
        rel_path = str(java_file.relative_to(sources_dir))
        
        for pattern_key in patterns_to_use:
            for pattern in priority_patterns[pattern_key]:
                if re.search(pattern, rel_path):
                    try:
                        content = java_file.read_text(encoding='utf-8', errors='ignore')
                        # Limit content size for API
                        if len(content) < 50000:
                            classes_to_scan.append({
                                "path": rel_path,
                                "content": content,
                                "category": pattern_key
                            })
                    except Exception as e:
                        logger.warning(f"Error reading {java_file}: {e}")
                    break
        
        # Limit number of classes for quick scan
        if scan_type == "quick" and len(classes_to_scan) >= 10:
            break
        elif scan_type != "deep" and len(classes_to_scan) >= 25:
            break
    
    results["classes_scanned"] = len(classes_to_scan)
    
    if not classes_to_scan:
        results["summary"] = "No relevant classes found to scan"
        return results
    
    # Build combined code context for AI analysis
    code_context = "\n\n".join([
        f"// === {c['path']} ({c['category']}) ===\n{c['content'][:8000]}"
        for c in classes_to_scan[:15]  # Limit context size
    ])
    
    # AI prompt for vulnerability analysis
    prompt = f"""You are an expert Android security auditor. Analyze the following decompiled Java code from an Android app for security vulnerabilities.

SCAN TYPE: {scan_type}
FOCUS AREAS: {', '.join(focus_areas) if focus_areas else 'General security'}

Analyze the code and return a JSON response with the following structure:
{{
    "vulnerabilities": [
        {{
            "id": "VULN-001",
            "title": "Clear vulnerability title",
            "severity": "critical|high|medium|low",
            "category": "Category (e.g., Authentication, Cryptography, Data Storage)",
            "affected_class": "ClassName.java",
            "affected_method": "methodName()",
            "description": "Detailed description of the vulnerability",
            "code_snippet": "The vulnerable code",
            "impact": "What an attacker could achieve",
            "remediation": "How to fix the vulnerability",
            "cwe_id": "CWE-XXX if applicable"
        }}
    ],
    "attack_chains": [
        {{
            "name": "Attack chain name",
            "steps": ["Step 1", "Step 2", "Step 3"],
            "impact": "Final impact of the attack chain",
            "likelihood": "high|medium|low"
        }}
    ],
    "recommendations": [
        "Priority recommendation 1",
        "Priority recommendation 2"
    ],
    "summary": "Executive summary of security posture"
}}

Focus on:
1. **Authentication Issues**: Weak auth, hardcoded credentials, insecure session management
2. **Cryptographic Failures**: Weak algorithms, hardcoded keys, improper IV usage
3. **Data Exposure**: Sensitive data in logs, insecure storage, data leakage
4. **Injection Vulnerabilities**: SQL injection, command injection, path traversal
5. **Insecure Communication**: HTTP usage, certificate validation bypass
6. **Access Control**: Exported components, missing permission checks
7. **WebView Vulnerabilities**: JavaScript interfaces, file access
8. **Code Quality**: Debug code, error handling, null checks

Be thorough and identify ALL security issues, including subtle ones that could be chained together.

CODE TO ANALYZE:
{code_context}"""

    try:
        response = client.models.generate_content(
            model=settings.gemini_model_id,
            contents=[types.Content(role="user", parts=[types.Part(text=prompt)])],
        )
        
        response_text = response.text
        
        # Extract JSON from response
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            ai_result = json.loads(json_match.group())
            
            # Process vulnerabilities
            for vuln in ai_result.get("vulnerabilities", []):
                severity = vuln.get("severity", "medium").lower()
                results["vulnerabilities"].append(vuln)
                results["risk_summary"][severity] = results["risk_summary"].get(severity, 0) + 1
            
            # Add attack chains
            results["attack_chains"] = ai_result.get("attack_chains", [])
            
            # Add recommendations
            results["recommendations"] = ai_result.get("recommendations", [])
            
            # Add summary
            results["summary"] = ai_result.get("summary", "Analysis complete")
        else:
            results["summary"] = response_text[:500]
            
    except Exception as e:
        logger.error(f"AI vulnerability scan failed: {e}")
        results["error"] = str(e)
        results["summary"] = f"AI analysis failed: {str(e)}"
    
    # Calculate overall risk
    if results["risk_summary"]["critical"] > 0:
        results["overall_risk"] = "critical"
    elif results["risk_summary"]["high"] > 0:
        results["overall_risk"] = "high"
    elif results["risk_summary"]["medium"] > 0:
        results["overall_risk"] = "medium"
    else:
        results["overall_risk"] = "low"
    
    return results


def _add_to_source_tree(tree: Dict, path: str) -> None:
    """Add a file path to the source tree structure."""
    parts = path.split('/')
    current = tree
    for i, part in enumerate(parts):
        if i == len(parts) - 1:
            # File
            if 'files' not in current:
                current['files'] = []
            current['files'].append(part)
        else:
            # Directory
            if part not in current:
                current[part] = {}
            current = current[part]


def get_jadx_class_source(output_dir: Path, class_path: str) -> Optional[str]:
    """Get the source code for a specific class from JADX output."""
    source_file = output_dir / "sources" / class_path
    if source_file.exists():
        return source_file.read_text(encoding='utf-8', errors='ignore')
    return None


def get_jadx_result_summary(output_dir: Path) -> Dict[str, Any]:
    """
    Get a summary of JADX decompilation results for AI diagram generation.
    
    Args:
        output_dir: Path to JADX output directory
        
    Returns:
        Dict with package_name, classes, source_tree, and sample code
    """
    sources_dir = output_dir / "sources"
    
    summary = {
        'package_name': 'unknown',
        'classes': [],
        'source_tree': {},
        'sample_code': [],  # Sample source code for AI context
    }
    
    if not sources_dir.exists():
        return summary
    
    # Find all Java files and categorize them
    java_files = list(sources_dir.rglob("*.java"))
    
    # Infer package name from directory structure
    if java_files:
        first_file = java_files[0]
        relative_path = first_file.relative_to(sources_dir)
        if len(relative_path.parts) >= 2:
            # Take first two package parts
            summary['package_name'] = '.'.join(relative_path.parts[:2])
    
    # Process classes (limit to first 200 for performance)
    for java_file in java_files[:200]:
        try:
            content = java_file.read_text(encoding='utf-8', errors='ignore')[:5000]  # First 5KB
            class_name = java_file.stem
            relative_path = str(java_file.relative_to(sources_dir))
            
            # Detect component type
            is_activity = 'extends Activity' in content or 'extends AppCompatActivity' in content or 'extends FragmentActivity' in content
            is_service = 'extends Service' in content or 'extends IntentService' in content
            is_receiver = 'extends BroadcastReceiver' in content
            is_provider = 'extends ContentProvider' in content
            
            # Extract package and extends info
            package_match = re.search(r'package\s+([\w.]+);', content)
            extends_match = re.search(r'class\s+\w+\s+extends\s+(\w+)', content)
            
            class_info = {
                'class_name': class_name,
                'package_name': package_match.group(1) if package_match else 'unknown',
                'file_path': relative_path,
                'is_activity': is_activity,
                'is_service': is_service,
                'is_receiver': is_receiver,
                'is_provider': is_provider,
                'extends': extends_match.group(1) if extends_match else None,
                'line_count': content.count('\n')
            }
            
            summary['classes'].append(class_info)
            
            # Collect sample code for main components (for AI context)
            if (is_activity or is_service or is_receiver or is_provider) and len(summary['sample_code']) < 10:
                summary['sample_code'].append({
                    'class_name': class_name,
                    'type': 'activity' if is_activity else 'service' if is_service else 'receiver' if is_receiver else 'provider',
                    'code_snippet': content[:3000]  # First 3KB
                })
                
        except Exception:
            continue
    
    return summary


def search_jadx_sources(output_dir: Path, query: str, max_results: int = 50) -> List[Dict[str, Any]]:
    """Search for a string in JADX decompiled sources."""
    import re
    
    results = []
    sources_dir = output_dir / "sources"
    
    if not sources_dir.exists():
        return results
    
    pattern = re.compile(re.escape(query), re.IGNORECASE)
    
    for java_file in sources_dir.rglob("*.java"):
        if len(results) >= max_results:
            break
            
        try:
            content = java_file.read_text(encoding='utf-8', errors='ignore')
            for i, line in enumerate(content.split('\n'), 1):
                if pattern.search(line):
                    results.append({
                        'file': str(java_file.relative_to(sources_dir)),
                        'line': i,
                        'content': line.strip()[:200],
                        'class_name': java_file.stem
                    })
                    if len(results) >= max_results:
                        break
        except Exception:
            continue
    
    return results


# ============================================================================
# Manifest Visualization Functions
# ============================================================================

def generate_manifest_visualization(apk_path: Path) -> ManifestVisualization:
    """
    Generate visualization data for an APK's AndroidManifest.
    
    Args:
        apk_path: Path to the APK file
    
    Returns:
        ManifestVisualization with nodes, edges, and mermaid diagram
    """
    try:
        from androguard.core.apk import APK
        apk = APK(str(apk_path))
    except Exception as e:
        logger.error(f"Failed to parse APK for manifest visualization: {e}")
        return ManifestVisualization(
            package_name="unknown",
            app_name=None,
            version_name=None,
            nodes=[],
            edges=[],
            component_counts={},
            permission_summary={},
            exported_count=0,
            main_activity=None,
            deep_link_schemes=[],
            mermaid_diagram="flowchart LR\n  Error[Failed to parse APK]"
        )
    
    nodes = []
    edges = []
    package_name = apk.get_package() or "unknown"
    app_name = apk.get_app_name()
    version_name = apk.get_androidversion_name()
    
    # Component counts
    component_counts = {
        'activities': 0,
        'services': 0,
        'receivers': 0,
        'providers': 0,
        'permissions': 0
    }
    
    # Permission analysis
    permission_summary = {
        'dangerous': 0,
        'normal': 0,
        'signature': 0,
        'total': 0
    }
    
    deep_link_schemes = set()
    main_activity = None
    exported_count = 0
    
    # Add app node
    nodes.append(ManifestNode(
        id="app",
        name=package_name,
        node_type="application",
        label=app_name or package_name,
        attributes={'version': version_name}
    ))
    
    # Process activities
    for activity in apk.get_activities():
        activity_short = activity.split('.')[-1]
        is_exported = _is_component_exported(apk, activity, 'activity')
        is_main = _is_main_activity(apk, activity)
        
        if is_main:
            main_activity = activity
        if is_exported:
            exported_count += 1
        
        node = ManifestNode(
            id=f"act_{activity_short}",
            name=activity,
            node_type="activity",
            label=activity_short,
            is_exported=is_exported,
            is_main=is_main,
            attributes=_get_component_attributes(apk, activity, 'activity')
        )
        nodes.append(node)
        component_counts['activities'] += 1
        
        # Edge from app to activity
        edges.append(ManifestEdge(
            source="app",
            target=f"act_{activity_short}",
            edge_type="contains",
            label="activity"
        ))
        
        # Process intent filters for deep links
        try:
            intent_filters = apk.get_intent_filters('activity', activity)
            if intent_filters and isinstance(intent_filters, dict):
                # androguard returns dict with 'action', 'category', 'data' keys
                data_list = intent_filters.get('data', [])
                if isinstance(data_list, list):
                    for data in data_list:
                        if isinstance(data, dict):
                            scheme = data.get('scheme', '')
                            if scheme and scheme not in ('http', 'https'):
                                deep_link_schemes.add(scheme)
                        elif isinstance(data, str):
                            # Handle if data is a simple string
                            if data and data not in ('http', 'https'):
                                deep_link_schemes.add(data)
        except Exception as e:
            logger.debug(f"Could not process intent filters for {activity}: {e}")
    
    # Process services
    for service in apk.get_services():
        service_short = service.split('.')[-1]
        is_exported = _is_component_exported(apk, service, 'service')
        
        if is_exported:
            exported_count += 1
        
        node = ManifestNode(
            id=f"svc_{service_short}",
            name=service,
            node_type="service",
            label=service_short,
            is_exported=is_exported,
            attributes=_get_component_attributes(apk, service, 'service')
        )
        nodes.append(node)
        component_counts['services'] += 1
        
        edges.append(ManifestEdge(
            source="app",
            target=f"svc_{service_short}",
            edge_type="contains",
            label="service"
        ))
    
    # Process receivers
    for receiver in apk.get_receivers():
        receiver_short = receiver.split('.')[-1]
        is_exported = _is_component_exported(apk, receiver, 'receiver')
        
        if is_exported:
            exported_count += 1
        
        node = ManifestNode(
            id=f"rcv_{receiver_short}",
            name=receiver,
            node_type="receiver",
            label=receiver_short,
            is_exported=is_exported,
            attributes=_get_component_attributes(apk, receiver, 'receiver')
        )
        nodes.append(node)
        component_counts['receivers'] += 1
        
        edges.append(ManifestEdge(
            source="app",
            target=f"rcv_{receiver_short}",
            edge_type="contains",
            label="receiver"
        ))
    
    # Process providers
    for provider in apk.get_providers():
        provider_short = provider.split('.')[-1]
        is_exported = _is_component_exported(apk, provider, 'provider')
        
        if is_exported:
            exported_count += 1
        
        node = ManifestNode(
            id=f"prv_{provider_short}",
            name=provider,
            node_type="provider",
            label=provider_short,
            is_exported=is_exported,
            is_dangerous=is_exported,  # Exported providers are dangerous
            attributes=_get_component_attributes(apk, provider, 'provider')
        )
        nodes.append(node)
        component_counts['providers'] += 1
        
        edges.append(ManifestEdge(
            source="app",
            target=f"prv_{provider_short}",
            edge_type="contains",
            label="provider"
        ))
    
    # Process permissions
    for permission in apk.get_permissions():
        perm_short = permission.split('.')[-1]
        is_dangerous = permission in DANGEROUS_PERMISSIONS
        
        if is_dangerous:
            permission_summary['dangerous'] += 1
        else:
            permission_summary['normal'] += 1
        permission_summary['total'] += 1
        
        node = ManifestNode(
            id=f"perm_{perm_short}",
            name=permission,
            node_type="permission",
            label=perm_short,
            is_dangerous=is_dangerous
        )
        nodes.append(node)
        component_counts['permissions'] += 1
        
        edges.append(ManifestEdge(
            source="app",
            target=f"perm_{perm_short}",
            edge_type="uses_permission",
            label=""
        ))
    
    # Generate mermaid diagram
    mermaid_diagram = _generate_manifest_mermaid(
        package_name, nodes, edges, main_activity, exported_count
    )
    
    return ManifestVisualization(
        package_name=package_name,
        app_name=app_name,
        version_name=version_name,
        nodes=nodes,
        edges=edges,
        component_counts=component_counts,
        permission_summary=permission_summary,
        exported_count=exported_count,
        main_activity=main_activity,
        deep_link_schemes=list(deep_link_schemes),
        mermaid_diagram=mermaid_diagram
    )


def _is_component_exported(apk, component_name: str, component_type: str) -> bool:
    """Check if a component is exported."""
    try:
        # Get android:exported attribute
        manifest = apk.get_android_manifest_xml()
        if manifest is None:
            return False
        
        # Simple check - if it has intent filters, it's typically exported
        intent_filters = apk.get_intent_filters(component_type, component_name)
        if intent_filters:
            return True
        
        return False
    except Exception:
        return False


def _is_main_activity(apk, activity_name: str) -> bool:
    """Check if an activity is the main launcher activity."""
    try:
        main = apk.get_main_activity()
        return main == activity_name
    except Exception:
        return False


def _get_component_attributes(apk, component_name: str, component_type: str) -> Dict[str, Any]:
    """Get additional attributes for a component."""
    attrs = {}
    try:
        intent_filters = apk.get_intent_filters(component_type, component_name)
        if intent_filters:
            # androguard returns a dict with 'action', 'category', 'data' keys
            if isinstance(intent_filters, dict):
                attrs['intent_filters'] = 1
                actions = intent_filters.get('action', [])
                if isinstance(actions, list):
                    attrs['actions'] = actions
                else:
                    attrs['actions'] = [str(actions)] if actions else []
            elif isinstance(intent_filters, list):
                attrs['intent_filters'] = len(intent_filters)
                attrs['actions'] = []
                for f in intent_filters:
                    if isinstance(f, dict):
                        attrs['actions'].extend(f.get('action', []))
                    elif isinstance(f, str):
                        attrs['actions'].append(f)
            else:
                attrs['intent_filters'] = 1
                attrs['actions'] = []
    except Exception:
        pass
    return attrs


def _generate_manifest_mermaid(
    package_name: str,
    nodes: List[ManifestNode],
    edges: List[ManifestEdge],
    main_activity: Optional[str],
    exported_count: int
) -> str:
    """Generate a Mermaid flowchart for the manifest with iconify icons."""
    lines = ["flowchart TB"]
    lines.append(f"    subgraph APP[\"fa6-brands:android {package_name}\"]")
    lines.append("    direction TB")
    
    # Group nodes by type
    activities = [n for n in nodes if n.node_type == "activity"]
    services = [n for n in nodes if n.node_type == "service"]
    receivers = [n for n in nodes if n.node_type == "receiver"]
    providers = [n for n in nodes if n.node_type == "provider"]
    permissions = [n for n in nodes if n.node_type == "permission"]
    
    # Activities subgraph
    if activities:
        lines.append("    subgraph Activities[\"mdi:application Activities\"]")
        for node in activities[:10]:  # Limit to avoid huge diagrams
            style = ":::exported" if node.is_exported else ""
            icon = "mdi:rocket-launch" if node.is_main else "mdi:application"
            lines.append(f"        {node.id}[{icon} {node.label}]{style}")
        if len(activities) > 10:
            lines.append(f"        act_more[mdi:dots-horizontal +{len(activities) - 10} more]")
        lines.append("    end")
    
    # Services subgraph
    if services:
        lines.append("    subgraph Services[\"mdi:cog Services\"]")
        for node in services[:5]:
            style = ":::exported" if node.is_exported else ""
            lines.append(f"        {node.id}[mdi:cog {node.label}]{style}")
        if len(services) > 5:
            lines.append(f"        svc_more[mdi:dots-horizontal +{len(services) - 5} more]")
        lines.append("    end")
    
    # Receivers subgraph
    if receivers:
        lines.append("    subgraph Receivers[\"mdi:broadcast Receivers\"]")
        for node in receivers[:5]:
            style = ":::exported" if node.is_exported else ""
            lines.append(f"        {node.id}[mdi:broadcast {node.label}]{style}")
        if len(receivers) > 5:
            lines.append(f"        rcv_more[mdi:dots-horizontal +{len(receivers) - 5} more]")
        lines.append("    end")
    
    # Providers subgraph
    if providers:
        lines.append("    subgraph Providers[\"mdi:database Providers\"]")
        for node in providers[:5]:
            style = ":::danger" if node.is_exported else ""
            lines.append(f"        {node.id}[mdi:database {node.label}]{style}")
        lines.append("    end")
    
    lines.append("    end")  # Close APP subgraph
    
    # Dangerous permissions outside
    dangerous_perms = [n for n in permissions if n.is_dangerous]
    if dangerous_perms:
        lines.append("    subgraph DangerousPermissions[\"fa6-solid:triangle-exclamation Dangerous Permissions\"]")
        for node in dangerous_perms[:8]:
            lines.append(f"        {node.id}[fa6-solid:unlock {node.label}]:::danger")
        if len(dangerous_perms) > 8:
            lines.append(f"        perm_more[mdi:dots-horizontal +{len(dangerous_perms) - 8} more]")
        lines.append("    end")
        lines.append("    APP --> DangerousPermissions")
    
    # Styles
    lines.append("    classDef exported fill:#ff9800,stroke:#e65100,color:#000")
    lines.append("    classDef danger fill:#f44336,stroke:#b71c1c,color:#fff")
    
    return "\n".join(lines)


# ============================================================================
# Attack Surface Map Functions
# ============================================================================

def generate_attack_surface_map(apk_path: Path) -> AttackSurfaceMap:
    """
    Generate a comprehensive attack surface map for an APK.
    
    Args:
        apk_path: Path to the APK file
    
    Returns:
        AttackSurfaceMap with attack vectors, deep links, and automated tests
    """
    try:
        from androguard.core.apk import APK
        apk = APK(str(apk_path))
    except Exception as e:
        logger.error(f"Failed to parse APK for attack surface: {e}")
        return AttackSurfaceMap(
            package_name="unknown",
            total_attack_vectors=0,
            attack_vectors=[],
            exposed_data_paths=[],
            deep_links=[],
            ipc_endpoints=[],
            overall_exposure_score=0,
            risk_level="unknown",
            risk_breakdown={},
            priority_targets=[],
            automated_tests=[],
            mermaid_attack_tree="flowchart TD\n  Error[Failed to parse APK]"
        )
    
    package_name = apk.get_package() or "unknown"
    attack_vectors = []
    exposed_data_paths = []
    deep_links = []
    ipc_endpoints = []
    automated_tests = []
    
    risk_breakdown = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0}
    
    # Analyze exported activities
    for activity in apk.get_activities():
        is_exported = _is_component_exported(apk, activity, 'activity')
        if is_exported:
            vector = _create_activity_attack_vector(apk, activity, package_name)
            attack_vectors.append(vector)
            risk_breakdown[vector.severity] = risk_breakdown.get(vector.severity, 0) + 1
            
            # Add automated test
            automated_tests.append({
                'name': f'Launch {activity.split(".")[-1]}',
                'command': f'adb shell am start -n {package_name}/{activity}',
                'description': 'Attempt to launch exported activity'
            })
    
    # Analyze exported services
    for service in apk.get_services():
        is_exported = _is_component_exported(apk, service, 'service')
        if is_exported:
            vector = _create_service_attack_vector(apk, service, package_name)
            attack_vectors.append(vector)
            risk_breakdown[vector.severity] = risk_breakdown.get(vector.severity, 0) + 1
            
            automated_tests.append({
                'name': f'Start {service.split(".")[-1]}',
                'command': f'adb shell am startservice -n {package_name}/{service}',
                'description': 'Attempt to start exported service'
            })
    
    # Analyze exported receivers
    for receiver in apk.get_receivers():
        is_exported = _is_component_exported(apk, receiver, 'receiver')
        if is_exported:
            vector = _create_receiver_attack_vector(apk, receiver, package_name)
            attack_vectors.append(vector)
            risk_breakdown[vector.severity] = risk_breakdown.get(vector.severity, 0) + 1
            
            # Get intent actions for receiver
            actions = []
            try:
                intent_filters = apk.get_intent_filters('receiver', receiver)
                if isinstance(intent_filters, dict):
                    action_list = intent_filters.get('action', [])
                    if isinstance(action_list, list):
                        actions.extend(action_list)
                elif isinstance(intent_filters, list):
                    for intent_filter in intent_filters:
                        if isinstance(intent_filter, dict):
                            actions.extend(intent_filter.get('action', []))
            except Exception:
                pass
            
            if actions:
                automated_tests.append({
                    'name': f'Broadcast to {receiver.split(".")[-1]}',
                    'command': f'adb shell am broadcast -a {actions[0]} -n {package_name}/{receiver}',
                    'description': 'Send broadcast to exported receiver'
                })
    
    # Analyze content providers
    for provider in apk.get_providers():
        is_exported = _is_component_exported(apk, provider, 'provider')
        if is_exported:
            vector = _create_provider_attack_vector(apk, provider, package_name)
            attack_vectors.append(vector)
            risk_breakdown['high'] += 1
            
            # Create exposed data path
            data_path = _create_exposed_data_path(apk, provider)
            if data_path:
                exposed_data_paths.append(data_path)
            
            automated_tests.append({
                'name': f'Query {provider.split(".")[-1]}',
                'command': f'adb shell content query --uri content://{package_name}.provider',
                'description': 'Attempt to query exported content provider'
            })
    
    # Analyze deep links
    deep_links = _extract_deep_links(apk, package_name)
    for deep_link in deep_links:
        vector = AttackVector(
            id=f"deeplink_{len(attack_vectors)}",
            name=f"Deep Link: {deep_link.scheme}://{deep_link.host}",
            vector_type="deep_link",
            component=deep_link.handling_activity,
            severity="medium",
            description=f"Custom URL scheme {deep_link.scheme}:// can be invoked from other apps or web pages",
            exploitation_steps=[
                f"Create HTML page with link: <a href=\"{deep_link.full_url}\">Click</a>",
                "Or use intent: Intent.parseUri() with the deep link URL",
                "Test parameter injection in path/query parameters"
            ],
            required_permissions=[],
            adb_command=f'adb shell am start -W -a android.intent.action.VIEW -d "{deep_link.full_url}"',
            intent_example=f'new Intent(Intent.ACTION_VIEW, Uri.parse("{deep_link.full_url}"))',
            mitigation="Validate all deep link parameters, use App Links with verification"
        )
        attack_vectors.append(vector)
        risk_breakdown['medium'] += 1
        
        automated_tests.append({
            'name': f'Open deep link {deep_link.scheme}://',
            'command': f'adb shell am start -W -a android.intent.action.VIEW -d "{deep_link.full_url}"',
            'description': f'Open deep link to {deep_link.handling_activity}'
        })
    
    # Calculate overall exposure score
    total_vectors = len(attack_vectors)
    exposure_score = min(100, (
        risk_breakdown['critical'] * 25 +
        risk_breakdown['high'] * 15 +
        risk_breakdown['medium'] * 8 +
        risk_breakdown['low'] * 3
    ))
    
    # Determine risk level
    if exposure_score >= 70 or risk_breakdown['critical'] > 0:
        risk_level = "critical"
    elif exposure_score >= 50 or risk_breakdown['high'] >= 3:
        risk_level = "high"
    elif exposure_score >= 25 or risk_breakdown['medium'] >= 3:
        risk_level = "medium"
    else:
        risk_level = "low"
    
    # Priority targets
    priority_targets = []
    for vector in sorted(attack_vectors, key=lambda v: {'critical': 0, 'high': 1, 'medium': 2, 'low': 3}.get(v.severity, 4)):
        if len(priority_targets) >= 5:
            break
        priority_targets.append(f"[{vector.severity.upper()}] {vector.name}")
    
    # Generate mermaid attack tree
    mermaid_attack_tree = _generate_attack_tree_mermaid(
        package_name, attack_vectors, deep_links, exposed_data_paths
    )
    
    return AttackSurfaceMap(
        package_name=package_name,
        total_attack_vectors=total_vectors,
        attack_vectors=attack_vectors,
        exposed_data_paths=exposed_data_paths,
        deep_links=deep_links,
        ipc_endpoints=ipc_endpoints,
        overall_exposure_score=exposure_score,
        risk_level=risk_level,
        risk_breakdown=risk_breakdown,
        priority_targets=priority_targets,
        automated_tests=automated_tests,
        mermaid_attack_tree=mermaid_attack_tree
    )


def _create_activity_attack_vector(apk, activity: str, package: str) -> AttackVector:
    """Create an attack vector for an exported activity."""
    activity_short = activity.split('.')[-1]
    
    # Check for sensitive activity names
    severity = "medium"
    if any(kw in activity.lower() for kw in ('login', 'auth', 'password', 'settings', 'admin', 'debug', 'internal')):
        severity = "high"
    if any(kw in activity.lower() for kw in ('webview', 'browser', 'url')):
        severity = "high"  # Potential URL injection
    
    return AttackVector(
        id=f"activity_{activity_short}",
        name=f"Exported Activity: {activity_short}",
        vector_type="exported_activity",
        component=activity,
        severity=severity,
        description=f"Activity {activity_short} is exported and can be launched by other apps",
        exploitation_steps=[
            "Launch activity directly using adb or Intent",
            "Check for sensitive functionality accessible without authentication",
            "Test intent extra parameters for injection",
            "Look for data returned in onActivityResult"
        ],
        required_permissions=[],
        adb_command=f'adb shell am start -n {package}/{activity}',
        intent_example=f'Intent intent = new Intent(); intent.setComponent(new ComponentName("{package}", "{activity}"));',
        mitigation="Remove android:exported=true if not needed, or add permission requirements"
    )


def _create_service_attack_vector(apk, service: str, package: str) -> AttackVector:
    """Create an attack vector for an exported service."""
    service_short = service.split('.')[-1]
    
    severity = "high"  # Services are generally higher risk
    if any(kw in service.lower() for kw in ('sync', 'background', 'job')):
        severity = "medium"
    
    return AttackVector(
        id=f"service_{service_short}",
        name=f"Exported Service: {service_short}",
        vector_type="exported_service",
        component=service,
        severity=severity,
        description=f"Service {service_short} is exported and can be started/bound by other apps",
        exploitation_steps=[
            "Start service with malicious intent data",
            "Bind to service and call exposed methods",
            "Analyze AIDL interface if available",
            "Check for sensitive operations performed by service"
        ],
        required_permissions=[],
        adb_command=f'adb shell am startservice -n {package}/{service}',
        intent_example=f'Intent intent = new Intent(); intent.setComponent(new ComponentName("{package}", "{service}"));',
        mitigation="Set android:exported=false or add signature-level permission requirement"
    )


def _create_receiver_attack_vector(apk, receiver: str, package: str) -> AttackVector:
    """Create an attack vector for an exported broadcast receiver."""
    receiver_short = receiver.split('.')[-1]
    
    # Get actions
    actions = []
    try:
        intent_filters = apk.get_intent_filters('receiver', receiver)
        if isinstance(intent_filters, dict):
            action_list = intent_filters.get('action', [])
            if isinstance(action_list, list):
                actions.extend(action_list)
        elif isinstance(intent_filters, list):
            for intent_filter in intent_filters:
                if isinstance(intent_filter, dict):
                    actions.extend(intent_filter.get('action', []))
    except Exception:
        pass
    
    severity = "medium"
    if any(kw in receiver.lower() for kw in ('sms', 'push', 'notification', 'message')):
        severity = "high"
    
    return AttackVector(
        id=f"receiver_{receiver_short}",
        name=f"Exported Receiver: {receiver_short}",
        vector_type="exported_receiver",
        component=receiver,
        severity=severity,
        description=f"Broadcast receiver {receiver_short} is exported and can receive broadcasts from other apps",
        exploitation_steps=[
            "Send broadcast with malicious intent data",
            "Test registered actions: " + ", ".join(actions[:3]) if actions else "Check for registered actions",
            "Look for command injection in extras",
            "Check for sensitive operations triggered by broadcast"
        ],
        required_permissions=[],
        adb_command=f'adb shell am broadcast -n {package}/{receiver}' + (f' -a {actions[0]}' if actions else ''),
        intent_example=f'sendBroadcast(new Intent("{actions[0] if actions else "ACTION"}").setComponent(new ComponentName("{package}", "{receiver}")));',
        mitigation="Use LocalBroadcastManager for internal broadcasts, add permission requirements"
    )


def _create_provider_attack_vector(apk, provider: str, package: str) -> AttackVector:
    """Create an attack vector for an exported content provider."""
    provider_short = provider.split('.')[-1]
    
    return AttackVector(
        id=f"provider_{provider_short}",
        name=f"Exported Provider: {provider_short}",
        vector_type="exported_provider",
        component=provider,
        severity="high",  # Providers are high risk due to data access
        description=f"Content provider {provider_short} is exported and may expose sensitive data",
        exploitation_steps=[
            "Query provider for data: content://authority/table",
            "Test SQL injection in selection parameters",
            "Check for path traversal in URI paths",
            "Look for sensitive data (user info, tokens, etc.)"
        ],
        required_permissions=[],
        adb_command=f'adb shell content query --uri content://{package}',
        intent_example=f'getContentResolver().query(Uri.parse("content://{package}/data"), null, null, null, null);',
        mitigation="Set android:exported=false, implement proper permission checks, sanitize queries"
    )


def _create_exposed_data_path(apk, provider: str) -> Optional[ExposedDataPath]:
    """Create an exposed data path entry for a content provider."""
    return ExposedDataPath(
        provider_name=provider,
        uri_pattern=f"content://{provider.split('.')[-1].lower()}/*",
        permissions_required=[],
        operations=["query", "insert", "update", "delete"],
        is_exported=True,
        potential_data="Database rows, files, or app-specific data",
        risk_level="high"
    )


def _extract_deep_links(apk, package: str) -> List[DeepLinkEntry]:
    """Extract all deep links from the APK."""
    deep_links = []
    
    for activity in apk.get_activities():
        try:
            intent_filters = apk.get_intent_filters('activity', activity)
            
            # Handle dict format (single filter) or list format (multiple filters)
            if isinstance(intent_filters, dict):
                # Process single intent filter dict
                actions = intent_filters.get('action', [])
                categories = intent_filters.get('category', [])
                
                if not isinstance(actions, list):
                    actions = [actions] if actions else []
                if not isinstance(categories, list):
                    categories = [categories] if categories else []
                
                # Check for VIEW action with BROWSABLE category
                is_browsable = 'android.intent.action.VIEW' in actions and \
                               'android.intent.category.BROWSABLE' in categories
                
                data_list = intent_filters.get('data', [])
                if not isinstance(data_list, list):
                    data_list = [data_list] if data_list else []
                
                for data in data_list:
                    if isinstance(data, dict):
                        scheme = data.get('scheme', '')
                        host = data.get('host', '')
                        path = data.get('path', data.get('pathPrefix', data.get('pathPattern', '')))
                        
                        if scheme:
                            # Build full URL
                            full_url = f"{scheme}://"
                            if host:
                                full_url += host
                            if path:
                                full_url += path
                            
                            deep_links.append(DeepLinkEntry(
                                scheme=scheme,
                                host=host or "*",
                                path=path or "/",
                                full_url=full_url,
                                handling_activity=activity,
                                parameters=[],
                                is_verified=scheme in ('http', 'https'),  # App Links
                                security_notes=_get_deep_link_security_notes(scheme, host, is_browsable)
                            ))
                    elif isinstance(data, str) and data:
                        # Simple scheme string
                        deep_links.append(DeepLinkEntry(
                            scheme=data,
                            host="*",
                            path="/",
                            full_url=f"{data}://",
                            handling_activity=activity,
                            parameters=[],
                            is_verified=False,
                            security_notes=[]
                        ))
            elif isinstance(intent_filters, list):
                # Process list of intent filters
                for intent_filter in intent_filters:
                    if not isinstance(intent_filter, dict):
                        continue
                    
                    actions = intent_filter.get('action', [])
                    categories = intent_filter.get('category', [])
                    
                    if not isinstance(actions, list):
                        actions = [actions] if actions else []
                    if not isinstance(categories, list):
                        categories = [categories] if categories else []
                    
                    is_browsable = 'android.intent.action.VIEW' in actions and \
                                   'android.intent.category.BROWSABLE' in categories
                    
                    data_list = intent_filter.get('data', [])
                    if not isinstance(data_list, list):
                        data_list = [data_list] if data_list else []
                    
                    for data in data_list:
                        if isinstance(data, dict):
                            scheme = data.get('scheme', '')
                            host = data.get('host', '')
                            path = data.get('path', data.get('pathPrefix', data.get('pathPattern', '')))
                            
                            if scheme:
                                full_url = f"{scheme}://"
                                if host:
                                    full_url += host
                                if path:
                                    full_url += path
                                
                                deep_links.append(DeepLinkEntry(
                                    scheme=scheme,
                                    host=host or "*",
                                    path=path or "/",
                                    full_url=full_url,
                                    handling_activity=activity,
                                    parameters=[],
                                    is_verified=scheme in ('http', 'https'),
                                    security_notes=_get_deep_link_security_notes(scheme, host, is_browsable)
                                ))
        except Exception as e:
            logger.debug(f"Could not process intent filters for activity {activity}: {e}")
    
    return deep_links


def _get_deep_link_security_notes(scheme: str, host: str, is_browsable: bool) -> List[str]:
    """Get security notes for a deep link."""
    notes = []
    
    if scheme in ('http', 'https'):
        if not host:
            notes.append("⚠️ No host specified - handles all URLs with this scheme")
        notes.append("Consider implementing App Links verification")
    else:
        notes.append(f"Custom scheme '{scheme}' - any app can trigger this")
        if is_browsable:
            notes.append("BROWSABLE category - can be triggered from web pages")
    
    if not host or host == "*":
        notes.append("⚠️ Wildcard host - very broad attack surface")
    
    return notes


def _generate_attack_tree_mermaid(
    package: str,
    vectors: List[AttackVector],
    deep_links: List[DeepLinkEntry],
    data_paths: List[ExposedDataPath]
) -> str:
    """Generate a mermaid attack tree diagram with iconify icons."""
    lines = ["flowchart TD"]
    lines.append(f"    ROOT[fa6-solid:crosshairs {package}]")
    
    # Group by vector type
    activities = [v for v in vectors if v.vector_type == "exported_activity"]
    services = [v for v in vectors if v.vector_type == "exported_service"]
    receivers = [v for v in vectors if v.vector_type == "exported_receiver"]
    providers = [v for v in vectors if v.vector_type == "exported_provider"]
    dl_vectors = [v for v in vectors if v.vector_type == "deep_link"]
    
    if activities:
        lines.append("    ROOT --> ACT[mdi:application-export Exported Activities]")
        for i, v in enumerate(activities[:5]):
            severity_icon = "fa6-solid:skull-crossbones" if v.severity == "critical" else "fa6-solid:circle-exclamation" if v.severity == "high" else "fa6-solid:triangle-exclamation"
            lines.append(f"    ACT --> ACT{i}[{severity_icon} {v.component.split('.')[-1]}]")
    
    if services:
        lines.append("    ROOT --> SVC[mdi:cog-transfer Exported Services]")
        for i, v in enumerate(services[:5]):
            severity_icon = "fa6-solid:skull-crossbones" if v.severity == "critical" else "fa6-solid:circle-exclamation" if v.severity == "high" else "fa6-solid:triangle-exclamation"
            lines.append(f"    SVC --> SVC{i}[{severity_icon} {v.component.split('.')[-1]}]")
    
    if receivers:
        lines.append("    ROOT --> RCV[mdi:broadcast Broadcast Receivers]")
        for i, v in enumerate(receivers[:5]):
            severity_icon = "fa6-solid:skull-crossbones" if v.severity == "critical" else "fa6-solid:circle-exclamation" if v.severity == "high" else "fa6-solid:triangle-exclamation"
            lines.append(f"    RCV --> RCV{i}[{severity_icon} {v.component.split('.')[-1]}]")
    
    if providers:
        lines.append("    ROOT --> PRV[mdi:database-export Content Providers]")
        for i, v in enumerate(providers[:3]):
            lines.append(f"    PRV --> PRV{i}[fa6-solid:skull-crossbones {v.component.split('.')[-1]}]")
        lines.append("    PRV --> SQLI[mdi:database-alert SQL Injection]")
        lines.append("    PRV --> PATH[mdi:folder-alert Path Traversal]")
    
    if deep_links:
        lines.append("    ROOT --> DL[mdi:link-variant Deep Links]")
        schemes = set(dl.scheme for dl in deep_links if dl.scheme not in ('http', 'https'))
        for scheme in list(schemes)[:3]:
            lines.append(f"    DL --> DL_{scheme}[lucide:link {scheme}://]")
    
    # Add attack outcomes
    lines.append("    ")
    lines.append("    subgraph IMPACTS[\"fa6-solid:explosion Potential Impacts\"]")
    lines.append("    IMP1[fa6-solid:door-open Auth Bypass]")
    lines.append("    IMP2[fa6-solid:file-export Data Exfil]")
    lines.append("    IMP3[mdi:needle Injection]")
    lines.append("    IMP4[fa6-solid:stairs Priv Esc]")
    lines.append("    end")
    
    return "\n".join(lines)


# ============================================================================
# Obfuscation Detection and Analysis
# ============================================================================

@dataclass
class ObfuscationIndicator:
    """A single obfuscation indicator found in the APK."""
    indicator_type: str  # proguard, dexguard, string_encryption, control_flow, etc.
    confidence: str  # high, medium, low
    description: str
    evidence: List[str]  # Specific examples found
    location: Optional[str] = None  # File or class where found
    deobfuscation_hint: Optional[str] = None


@dataclass
class StringEncryptionPattern:
    """Detected string encryption pattern."""
    pattern_name: str
    class_name: str
    method_name: str
    encrypted_strings_count: int
    decryption_method_signature: Optional[str] = None
    sample_encrypted_values: List[str] = field(default_factory=list)
    suggested_frida_hook: Optional[str] = None


@dataclass
class ClassNamingAnalysis:
    """Analysis of class naming patterns for obfuscation detection."""
    total_classes: int
    single_letter_classes: int
    short_name_classes: int  # 2-3 chars
    meaningful_name_classes: int
    obfuscation_ratio: float  # 0.0 to 1.0
    sample_obfuscated_names: List[str]
    sample_original_names: List[str]


@dataclass
class ControlFlowObfuscation:
    """Detected control flow obfuscation patterns."""
    pattern_type: str  # switch_dispatch, opaque_predicates, dead_code, etc.
    affected_methods: int
    sample_classes: List[str]
    complexity_score: float  # Higher = more obfuscated


@dataclass
class NativeProtection:
    """Detected native code protection."""
    has_native_libs: bool
    native_lib_names: List[str]
    protection_indicators: List[str]  # packing, anti-debug, integrity checks
    jni_functions: List[str]


@dataclass
class ObfuscationAnalysisResult:
    """Complete obfuscation analysis result."""
    package_name: str
    overall_obfuscation_level: str  # none, light, moderate, heavy, extreme
    obfuscation_score: int  # 0-100
    detected_tools: List[str]  # ProGuard, DexGuard, Allatori, etc.
    
    # Detailed analysis
    indicators: List[ObfuscationIndicator]
    class_naming: ClassNamingAnalysis
    string_encryption: List[StringEncryptionPattern]
    control_flow: List[ControlFlowObfuscation]
    native_protection: NativeProtection
    
    # Recommendations
    deobfuscation_strategies: List[str]
    recommended_tools: List[str]
    frida_hooks: List[str]  # Auto-generated Frida hooks
    
    # Metadata
    analysis_time: float
    warnings: List[str]


def analyze_apk_obfuscation(apk_path: Path) -> ObfuscationAnalysisResult:
    """
    Analyze an APK for obfuscation techniques using FAST lightweight analysis.
    
    Uses direct DEX parsing instead of full AnalyzeAPK() to complete in seconds
    instead of minutes for large APKs.
    
    Detects:
    - ProGuard/R8 class renaming
    - DexGuard commercial protection
    - String encryption patterns
    - Native code protection
    """
    import time
    import zipfile
    import struct
    from androguard.core.apk import APK
    
    start_time = time.time()
    warnings = []
    
    try:
        # Use lightweight APK parsing (fast - only manifest/resources)
        apk = APK(str(apk_path))
        package_name = apk.get_package() or "unknown"
    except Exception as e:
        return ObfuscationAnalysisResult(
            package_name="unknown",
            overall_obfuscation_level="unknown",
            obfuscation_score=0,
            detected_tools=[],
            indicators=[],
            class_naming=ClassNamingAnalysis(0, 0, 0, 0, 0.0, [], []),
            string_encryption=[],
            control_flow=[],
            native_protection=NativeProtection(False, [], [], []),
            deobfuscation_strategies=[],
            recommended_tools=[],
            frida_hooks=[],
            analysis_time=time.time() - start_time,
            warnings=[f"Failed to analyze APK: {str(e)}"]
        )
    
    indicators = []
    detected_tools = []
    
    # Fast class name extraction from DEX without full analysis
    class_names = _extract_class_names_fast(apk_path)
    class_naming = _analyze_class_naming_fast(class_names, package_name)
    
    # Fast obfuscation detection based on class names and APK structure
    indicators.extend(_detect_proguard_indicators_fast(apk, class_naming))
    indicators.extend(_detect_dexguard_indicators_fast(apk))
    
    # Fast string pattern analysis (sample DEX strings)
    string_encryption = _analyze_string_encryption_fast(apk_path)
    
    # Control flow analysis skipped (requires full DEX analysis - too slow)
    control_flow = []
    
    # Analyze native protection (fast - just checks for .so files)
    native_protection = _analyze_native_protection(apk)
    
    # Determine detected tools
    detected_tools = _identify_obfuscation_tools_fast(indicators, class_naming, string_encryption)
    
    # Calculate overall score
    obfuscation_score = _calculate_obfuscation_score_fast(
        class_naming, indicators, string_encryption, native_protection
    )
    
    # Determine level
    if obfuscation_score < 15:
        level = "none"
    elif obfuscation_score < 35:
        level = "light"
    elif obfuscation_score < 60:
        level = "moderate"
    elif obfuscation_score < 85:
        level = "heavy"
    else:
        level = "extreme"
    
    # Generate recommendations
    strategies = _generate_deobfuscation_strategies_fast(detected_tools, level)
    recommended_tools = _recommend_deobfuscation_tools(detected_tools, level)
    frida_hooks = _generate_frida_hooks_fast(package_name, level)
    
    return ObfuscationAnalysisResult(
        package_name=package_name,
        overall_obfuscation_level=level,
        obfuscation_score=obfuscation_score,
        detected_tools=detected_tools,
        indicators=indicators,
        class_naming=class_naming,
        string_encryption=string_encryption,
        control_flow=control_flow,
        native_protection=native_protection,
        deobfuscation_strategies=strategies,
        recommended_tools=recommended_tools,
        frida_hooks=frida_hooks,
        analysis_time=time.time() - start_time,
        warnings=warnings
    )


def _extract_class_names_fast(apk_path: Path) -> List[str]:
    """Extract class names from DEX files without full analysis - very fast."""
    import zipfile
    class_names = []
    
    try:
        with zipfile.ZipFile(str(apk_path), 'r') as zf:
            for name in zf.namelist():
                if name.endswith('.dex'):
                    try:
                        dex_data = zf.read(name)
                        # Parse DEX header to find string table
                        if len(dex_data) < 112:
                            continue
                        
                        # DEX file magic check
                        if dex_data[:4] != b'dex\n':
                            continue
                        
                        # Read class_defs section info from header
                        class_defs_size = int.from_bytes(dex_data[96:100], 'little')
                        class_defs_off = int.from_bytes(dex_data[100:104], 'little')
                        
                        # Read string IDs info
                        string_ids_size = int.from_bytes(dex_data[56:60], 'little')
                        string_ids_off = int.from_bytes(dex_data[60:64], 'little')
                        
                        # Read type IDs info  
                        type_ids_size = int.from_bytes(dex_data[64:68], 'little')
                        type_ids_off = int.from_bytes(dex_data[68:72], 'little')
                        
                        # Extract class names from class_defs (sample first 1000)
                        max_classes = min(class_defs_size, 1000)
                        for i in range(max_classes):
                            class_def_off = class_defs_off + (i * 32)
                            if class_def_off + 4 > len(dex_data):
                                break
                            
                            class_idx = int.from_bytes(dex_data[class_def_off:class_def_off+4], 'little')
                            
                            # Get type descriptor
                            if class_idx < type_ids_size:
                                type_id_off = type_ids_off + (class_idx * 4)
                                if type_id_off + 4 <= len(dex_data):
                                    descriptor_idx = int.from_bytes(dex_data[type_id_off:type_id_off+4], 'little')
                                    
                                    # Get string from string table
                                    if descriptor_idx < string_ids_size:
                                        string_id_off = string_ids_off + (descriptor_idx * 4)
                                        if string_id_off + 4 <= len(dex_data):
                                            string_data_off = int.from_bytes(dex_data[string_id_off:string_id_off+4], 'little')
                                            
                                            # Read ULEB128 length and string
                                            if string_data_off < len(dex_data):
                                                # Skip ULEB128 length
                                                pos = string_data_off
                                                while pos < len(dex_data) and dex_data[pos] & 0x80:
                                                    pos += 1
                                                pos += 1
                                                
                                                # Read null-terminated string
                                                end = pos
                                                while end < len(dex_data) and dex_data[end] != 0:
                                                    end += 1
                                                
                                                try:
                                                    class_name = dex_data[pos:end].decode('utf-8', errors='ignore')
                                                    if class_name.startswith('L') and class_name.endswith(';'):
                                                        class_names.append(class_name[1:-1])  # Remove L and ;
                                                except:
                                                    pass
                    except Exception as e:
                        logger.debug(f"Error parsing DEX {name}: {e}")
    except Exception as e:
        logger.warning(f"Error extracting class names: {e}")
    
    return class_names


def _analyze_class_naming_fast(class_names: List[str], package_name: str) -> ClassNamingAnalysis:
    """Fast class naming analysis based on extracted class names."""
    single_letter = []
    short_name = []
    meaningful = []
    
    # Filter to app classes only (exclude framework)
    app_classes = [c for c in class_names if not any(
        c.startswith(p) for p in ['android/', 'androidx/', 'java/', 'kotlin/', 'com/google/android/']
    )]
    
    for class_name in app_classes[:500]:  # Sample up to 500 classes
        simple_name = class_name.split('/')[-1]
        # Handle inner classes
        if '$' in simple_name:
            simple_name = simple_name.split('$')[-1]
        
        if len(simple_name) == 1 and simple_name.isalpha():
            single_letter.append(class_name)
        elif len(simple_name) <= 3 and simple_name.replace('_', '').isalnum():
            short_name.append(class_name)
        elif len(simple_name) > 6:
            meaningful.append(class_name)
    
    total = len(app_classes)
    obfuscated = len(single_letter) + len(short_name)
    obfuscation_ratio = obfuscated / max(total, 1)
    
    return ClassNamingAnalysis(
        total_classes=total,
        single_letter_classes=len(single_letter),
        short_name_classes=len(short_name),
        meaningful_name_classes=len(meaningful),
        obfuscation_ratio=obfuscation_ratio,
        sample_obfuscated_names=single_letter[:10] + short_name[:10],
        sample_original_names=meaningful[:10]
    )


def _detect_proguard_indicators_fast(apk, class_naming: ClassNamingAnalysis) -> List[ObfuscationIndicator]:
    """Fast ProGuard detection without full DEX analysis."""
    indicators = []
    
    # Check class naming patterns
    if class_naming.obfuscation_ratio > 0.3:
        indicators.append(ObfuscationIndicator(
            indicator_type="class_renaming",
            description="High ratio of short/single-letter class names indicates ProGuard/R8",
            confidence="high" if class_naming.obfuscation_ratio > 0.5 else "medium",
            evidence=[f"Obfuscation ratio: {class_naming.obfuscation_ratio:.1%}"],
            affected_classes=class_naming.sample_obfuscated[:5]
        ))
    
    # Check for mapping file references
    try:
        files = apk.get_files()
        for f in files:
            if 'proguard' in f.lower() or 'mapping' in f.lower():
                indicators.append(ObfuscationIndicator(
                    indicator_type="proguard_config",
                    description=f"ProGuard-related file found: {f}",
                    confidence="high",
                    evidence=[f],
                    affected_classes=[]
                ))
                break
    except:
        pass
    
    return indicators


def _detect_dexguard_indicators_fast(apk) -> List[ObfuscationIndicator]:
    """Fast DexGuard detection."""
    indicators = []
    
    try:
        # Check for DexGuard-specific files or classes
        files = apk.get_files()
        dexguard_signs = ['dexguard', 'guardsquare']
        
        for f in files:
            f_lower = f.lower()
            if any(sign in f_lower for sign in dexguard_signs):
                indicators.append(ObfuscationIndicator(
                    indicator_type="dexguard",
                    description="DexGuard commercial protection detected",
                    confidence="high",
                    evidence=[f],
                    affected_classes=[]
                ))
                break
    except:
        pass
    
    return indicators


def _analyze_string_encryption_fast(apk_path: Path) -> List[StringEncryptionPattern]:
    """Fast string encryption pattern detection by sampling DEX strings."""
    import zipfile
    patterns = []
    
    try:
        with zipfile.ZipFile(str(apk_path), 'r') as zf:
            for name in zf.namelist():
                if name.endswith('.dex'):
                    dex_data = zf.read(name)
                    
                    # Look for common encryption patterns in raw bytes
                    # Base64-like encoded strings
                    base64_count = dex_data.count(b'==')
                    if base64_count > 50:
                        patterns.append(StringEncryptionPattern(
                            pattern_type="base64_encoding",
                            description="High number of Base64-encoded strings detected",
                            sample_encrypted=["[Base64 patterns detected]"],
                            decryption_hint="Strings may be Base64 encoded at runtime",
                            occurrences=base64_count
                        ))
                    
                    # Check for AES/DES encryption class references
                    if b'Cipher' in dex_data or b'AES' in dex_data or b'DES' in dex_data:
                        patterns.append(StringEncryptionPattern(
                            pattern_type="crypto_usage",
                            description="Cryptographic APIs found - may be used for string encryption",
                            sample_encrypted=["[Cipher/AES/DES references]"],
                            decryption_hint="Hook javax.crypto.Cipher to intercept decryption",
                            occurrences=1
                        ))
                    
                    break  # Only check first DEX
    except:
        pass
    
    return patterns


def _identify_obfuscation_tools_fast(indicators: List[ObfuscationIndicator], 
                                     class_naming: ClassNamingAnalysis,
                                     string_encryption: List[StringEncryptionPattern]) -> List[str]:
    """Identify likely obfuscation tools used."""
    tools = []
    
    for ind in indicators:
        if 'dexguard' in ind.indicator_type.lower():
            if 'DexGuard' not in tools:
                tools.append('DexGuard')
        elif 'proguard' in ind.indicator_type.lower() or ind.indicator_type == 'class_renaming':
            if 'ProGuard/R8' not in tools:
                tools.append('ProGuard/R8')
    
    if class_naming.obfuscation_ratio > 0.3 and not tools:
        tools.append('ProGuard/R8 (likely)')
    
    return tools


def _calculate_obfuscation_score_fast(class_naming: ClassNamingAnalysis,
                                      indicators: List[ObfuscationIndicator],
                                      string_encryption: List[StringEncryptionPattern],
                                      native_protection: NativeProtection) -> int:
    """Calculate obfuscation score (0-100)."""
    score = 0
    
    # Class naming (up to 40 points)
    score += min(40, int(class_naming.obfuscation_ratio * 60))
    
    # Indicators (up to 30 points)
    for ind in indicators:
        if ind.confidence == "high":
            score += 10
        elif ind.confidence == "medium":
            score += 5
    score = min(score, 70)  # Cap at 70 so far
    
    # String encryption (up to 15 points)
    score += min(15, len(string_encryption) * 5)
    
    # Native protection (up to 15 points)
    if native_protection.has_native_libs:
        score += 5
        if native_protection.protection_indicators:
            score += 10
    
    return min(100, score)


def _generate_deobfuscation_strategies_fast(detected_tools: List[str], level: str) -> List[str]:
    """Generate deobfuscation strategies based on detected tools."""
    strategies = []
    
    if 'DexGuard' in detected_tools:
        strategies.extend([
            "DexGuard uses commercial protection - consider using dex2jar + CFR decompiler",
            "Use Frida to hook decryption methods at runtime",
            "Try jadx with --deobf flag for automatic renaming"
        ])
    
    if any('ProGuard' in t for t in detected_tools):
        strategies.extend([
            "Standard ProGuard/R8 - use jadx --deobf for automatic renaming",
            "Look for mapping.txt file if available for original names",
            "Use regex patterns to identify renamed classes by behavior"
        ])
    
    if level in ['heavy', 'extreme']:
        strategies.extend([
            "Consider dynamic analysis with Frida for runtime inspection",
            "Use apktool for smali-level analysis",
            "Run in emulator with API monitoring"
        ])
    
    if not strategies:
        strategies.append("Minimal obfuscation detected - standard decompilers should work well")
    
    return strategies


def _generate_frida_hooks_fast(package_name: str, level: str) -> List[str]:
    """Generate basic Frida hooks for obfuscation bypass."""
    hooks = []
    
    # Basic string decryption hook
    hooks.append(f'''// Hook String class to catch decrypted strings
Java.perform(function() {{
    var String = Java.use('java.lang.String');
    String.$init.overload('[B').implementation = function(bytes) {{
        var result = this.$init(bytes);
        console.log('[String] ' + result);
        return result;
    }};
}});''')
    
    if level in ['moderate', 'heavy', 'extreme']:
        hooks.append(f'''// Hook Cipher for encryption/decryption
Java.perform(function() {{
    var Cipher = Java.use('javax.crypto.Cipher');
    Cipher.doFinal.overload('[B').implementation = function(input) {{
        var result = this.doFinal(input);
        console.log('[Cipher] Input: ' + input.length + ' bytes, Output: ' + result.length + ' bytes');
        return result;
    }};
}});''')
    
    return hooks


# Keep original slow function for deep analysis if needed
def analyze_apk_obfuscation_deep(apk_path: Path) -> ObfuscationAnalysisResult:
    """Analyze class naming patterns to detect obfuscation."""
    single_letter = []
    short_name = []
    meaningful = []
    all_classes = []
    
    for cls in dx.get_classes():
        class_name = cls.name
        if class_name.startswith("L") and class_name.endswith(";"):
            class_name = class_name[1:-1]
        
        # Skip Android/Java framework classes
        if any(class_name.startswith(prefix) for prefix in 
               ['android/', 'androidx/', 'java/', 'kotlin/', 'com/google/android/']):
            continue
        
        all_classes.append(class_name)
        
        # Get simple class name
        simple_name = class_name.split('/')[-1]
        
        if len(simple_name) == 1:
            single_letter.append(class_name)
        elif len(simple_name) <= 3 and simple_name.isalnum():
            short_name.append(class_name)
        else:
            meaningful.append(class_name)
    
    total = len(all_classes)
    obfuscated_count = len(single_letter) + len(short_name)
    obfuscation_ratio = obfuscated_count / total if total > 0 else 0.0
    
    return ClassNamingAnalysis(
        total_classes=total,
        single_letter_classes=len(single_letter),
        short_name_classes=len(short_name),
        meaningful_name_classes=len(meaningful),
        obfuscation_ratio=round(obfuscation_ratio, 3),
        sample_obfuscated_names=(single_letter + short_name)[:10],
        sample_original_names=meaningful[:10]
    )


def _detect_proguard_indicators(apk, dx, class_naming: ClassNamingAnalysis) -> List[ObfuscationIndicator]:
    """Detect ProGuard/R8 obfuscation indicators."""
    indicators = []
    
    # Check for mapping file residue
    if class_naming.obfuscation_ratio > 0.3:
        indicators.append(ObfuscationIndicator(
            indicator_type="proguard_naming",
            confidence="high" if class_naming.obfuscation_ratio > 0.6 else "medium",
            description=f"Class naming suggests ProGuard/R8: {class_naming.obfuscation_ratio*100:.1f}% obfuscated",
            evidence=class_naming.sample_obfuscated_names[:5],
            deobfuscation_hint="Look for mapping.txt file or use jadx's deobfuscation features"
        ))
    
    # Check for sequential class names (a, b, c pattern)
    sequential_pattern = _check_sequential_naming(class_naming.sample_obfuscated_names)
    if sequential_pattern:
        indicators.append(ObfuscationIndicator(
            indicator_type="proguard_sequential",
            confidence="high",
            description="Sequential alphabetical class naming detected (ProGuard default)",
            evidence=sequential_pattern[:5],
            deobfuscation_hint="Classes are likely renamed alphabetically; mapping file would reveal originals"
        ))
    
    # Check for removed line numbers
    if _check_removed_line_numbers(dx):
        indicators.append(ObfuscationIndicator(
            indicator_type="proguard_no_linenumbers",
            confidence="medium",
            description="Line number information removed (ProGuard optimization)",
            evidence=["Debug info stripped from bytecode"],
            deobfuscation_hint="Debugging will be harder; use dynamic analysis with Frida"
        ))
    
    return indicators


def _detect_dexguard_indicators(apk, dx) -> List[ObfuscationIndicator]:
    """Detect DexGuard-specific obfuscation."""
    indicators = []
    
    # DexGuard-specific patterns
    dexguard_classes = [
        'com/guardsquare/',
        'o/a/', 'o/b/', 'o/c/',  # Common DexGuard output patterns
    ]
    
    for cls in dx.get_classes():
        class_name = cls.name
        for pattern in dexguard_classes:
            if pattern in class_name:
                indicators.append(ObfuscationIndicator(
                    indicator_type="dexguard",
                    confidence="high",
                    description="DexGuard commercial obfuscation detected",
                    evidence=[class_name],
                    location=class_name,
                    deobfuscation_hint="DexGuard uses advanced protection; dynamic analysis recommended"
                ))
                break
    
    # Check for encrypted assets (DexGuard feature)
    encrypted_assets = _check_encrypted_assets(apk)
    if encrypted_assets:
        indicators.append(ObfuscationIndicator(
            indicator_type="encrypted_assets",
            confidence="medium",
            description="Encrypted or obfuscated assets detected",
            evidence=encrypted_assets[:5],
            deobfuscation_hint="Assets may be decrypted at runtime; hook decryption methods"
        ))
    
    # Check for integrity checks
    integrity_patterns = ['checksum', 'signature', 'integrity', 'tamper']
    for cls in dx.get_classes():
        for method in cls.get_methods():
            method_name = method.name.lower()
            if any(p in method_name for p in integrity_patterns):
                indicators.append(ObfuscationIndicator(
                    indicator_type="integrity_check",
                    confidence="medium",
                    description="Integrity/tamper detection found",
                    evidence=[f"{cls.name}->{method.name}"],
                    location=cls.name,
                    deobfuscation_hint="May need to bypass integrity checks for modification"
                ))
                break
    
    return indicators


def _detect_reflection_usage(dx) -> List[ObfuscationIndicator]:
    """Detect heavy reflection usage (API hiding)."""
    indicators = []
    reflection_methods = []
    
    reflection_apis = [
        'Ljava/lang/reflect/Method;->invoke',
        'Ljava/lang/Class;->forName',
        'Ljava/lang/Class;->getDeclaredMethod',
        'Ljava/lang/Class;->getDeclaredField',
        'Ljava/lang/Class;->getMethod',
    ]
    
    reflection_count = 0
    for cls in dx.get_classes():
        for method in cls.get_methods():
            try:
                if method.get_method() is None:
                    continue
                code = method.get_method().get_code()
                if code is None:
                    continue
                
                # Check for reflection API calls
                for instruction in code.get_bc().get_instructions():
                    if hasattr(instruction, 'get_output'):
                        output = instruction.get_output()
                        if any(api in output for api in reflection_apis):
                            reflection_count += 1
                            if len(reflection_methods) < 10:
                                reflection_methods.append(f"{cls.name}->{method.name}")
            except:
                continue
    
    if reflection_count > 20:
        indicators.append(ObfuscationIndicator(
            indicator_type="reflection_hiding",
            confidence="high" if reflection_count > 50 else "medium",
            description=f"Heavy reflection usage detected ({reflection_count} calls) - API hiding",
            evidence=reflection_methods,
            deobfuscation_hint="Hook Class.forName and Method.invoke to trace API calls"
        ))
    
    return indicators


def _analyze_string_encryption(dx) -> List[StringEncryptionPattern]:
    """Detect and analyze string encryption patterns."""
    patterns = []
    
    # Look for common string decryption patterns
    decryptor_signatures = [
        ('xor_decrypt', ['xor', 'decrypt', 'decode']),
        ('base64_decode', ['base64', 'b64', 'decode']),
        ('aes_decrypt', ['aes', 'cipher', 'decrypt']),
        ('custom_decrypt', ['deobfuscate', 'unprotect', 'reveal']),
    ]
    
    for cls in dx.get_classes():
        class_name = cls.name
        
        # Skip framework classes
        if any(class_name.startswith(f"L{prefix}") for prefix in 
               ['android/', 'androidx/', 'java/', 'kotlin/']):
            continue
        
        for method in cls.get_methods():
            method_name = method.name.lower()
            
            for pattern_name, keywords in decryptor_signatures:
                if any(kw in method_name for kw in keywords):
                    # Check if method takes String and returns String
                    try:
                        proto = method.get_method().get_descriptor() if method.get_method() else ""
                        if 'Ljava/lang/String;' in proto:
                            frida_hook = _generate_string_decrypt_hook(class_name, method.name)
                            patterns.append(StringEncryptionPattern(
                                pattern_name=pattern_name,
                                class_name=class_name,
                                method_name=method.name,
                                encrypted_strings_count=0,  # Would need deeper analysis
                                decryption_method_signature=proto,
                                sample_encrypted_values=[],
                                suggested_frida_hook=frida_hook
                            ))
                    except:
                        pass
    
    # Look for classes with many short static string fields (encrypted constants)
    for cls in dx.get_classes():
        static_strings = []
        for field in cls.get_fields():
            try:
                if 'Ljava/lang/String;' in str(field.get_field().get_descriptor()):
                    static_strings.append(field.name)
            except:
                pass
        
        if len(static_strings) > 20:
            # Many string constants might indicate encrypted strings
            class_name = cls.name
            simple_name = class_name.split('/')[-1].rstrip(';')
            if len(simple_name) <= 3:  # Likely obfuscated class
                patterns.append(StringEncryptionPattern(
                    pattern_name="encrypted_constants",
                    class_name=class_name,
                    method_name="<clinit>",
                    encrypted_strings_count=len(static_strings),
                    sample_encrypted_values=[],
                    suggested_frida_hook=f"// Hook static initializer of {class_name}"
                ))
    
    return patterns


def _analyze_control_flow(dx) -> List[ControlFlowObfuscation]:
    """Analyze control flow obfuscation patterns."""
    patterns = []
    
    switch_dispatch_methods = []
    high_complexity_methods = []
    
    for cls in dx.get_classes():
        class_name = cls.name
        
        # Skip framework
        if any(class_name.startswith(f"L{prefix}") for prefix in 
               ['android/', 'androidx/', 'java/', 'kotlin/']):
            continue
        
        for method in cls.get_methods():
            try:
                m = method.get_method()
                if m is None:
                    continue
                code = m.get_code()
                if code is None:
                    continue
                
                # Count control flow complexity
                switch_count = 0
                goto_count = 0
                
                for instruction in code.get_bc().get_instructions():
                    op_name = instruction.get_name()
                    if 'switch' in op_name:
                        switch_count += 1
                    if 'goto' in op_name:
                        goto_count += 1
                
                # Large switch statements often indicate dispatcher pattern
                if switch_count > 3:
                    switch_dispatch_methods.append(f"{class_name}->{method.name}")
                
                # Excessive gotos indicate control flow obfuscation
                if goto_count > 20:
                    high_complexity_methods.append(f"{class_name}->{method.name}")
                    
            except:
                continue
    
    if switch_dispatch_methods:
        patterns.append(ControlFlowObfuscation(
            pattern_type="switch_dispatch",
            affected_methods=len(switch_dispatch_methods),
            sample_classes=switch_dispatch_methods[:5],
            complexity_score=min(len(switch_dispatch_methods) / 10, 1.0)
        ))
    
    if high_complexity_methods:
        patterns.append(ControlFlowObfuscation(
            pattern_type="goto_obfuscation",
            affected_methods=len(high_complexity_methods),
            sample_classes=high_complexity_methods[:5],
            complexity_score=min(len(high_complexity_methods) / 5, 1.0)
        ))
    
    return patterns


def _analyze_native_protection(apk) -> NativeProtection:
    """Analyze native library protection."""
    native_libs = []
    protection_indicators = []
    jni_functions = []
    
    # Find native libraries
    for f in apk.get_files():
        if f.endswith('.so'):
            native_libs.append(f)
    
    # Check for known protection patterns in lib names
    protection_patterns = {
        'libjiagu': 'Baidu Jiagu packer',
        'libsecexe': 'Alibaba/Bangcle protection',
        'libDexHelper': 'Dex protection helper',
        'libprotect': 'Generic protection library',
        'libtprt': 'Tencent protection',
        'libexec': 'Code execution protection',
        'libsecmain': 'Security main library',
    }
    
    for lib in native_libs:
        lib_name = lib.split('/')[-1]
        for pattern, description in protection_patterns.items():
            if pattern in lib_name:
                protection_indicators.append(f"{description}: {lib_name}")
    
    # Look for JNI registration in class names
    jni_related = ['JNI', 'Native', 'nativeLib']
    # This would need deeper analysis of the actual native libs
    
    return NativeProtection(
        has_native_libs=len(native_libs) > 0,
        native_lib_names=native_libs,
        protection_indicators=protection_indicators,
        jni_functions=jni_functions
    )


def _check_sequential_naming(names: List[str]) -> List[str]:
    """Check for sequential alphabetical naming pattern."""
    sequential = []
    simple_names = [n.split('/')[-1].rstrip(';') for n in names if n]
    
    # Look for a, b, c or aa, ab, ac patterns
    sorted_names = sorted(simple_names)
    for i, name in enumerate(sorted_names[:-1]):
        if len(name) == 1 and len(sorted_names[i+1]) == 1:
            if ord(sorted_names[i+1]) == ord(name) + 1:
                if name not in sequential:
                    sequential.append(name)
                sequential.append(sorted_names[i+1])
    
    return sequential


def _check_removed_line_numbers(dx) -> bool:
    """Check if line number information has been removed."""
    checked = 0
    no_lines = 0
    
    for cls in dx.get_classes():
        for method in cls.get_methods():
            try:
                m = method.get_method()
                if m and m.get_code():
                    checked += 1
                    debug_info = m.get_code().get_debug()
                    if debug_info is None:
                        no_lines += 1
            except:
                pass
            
            if checked > 100:
                break
        if checked > 100:
            break
    
    return no_lines > checked * 0.8 if checked > 0 else False


def _check_encrypted_assets(apk) -> List[str]:
    """Check for potentially encrypted assets."""
    encrypted = []
    
    for f in apk.get_files():
        if f.startswith('assets/'):
            # Check for unusual extensions or patterns
            if any(ext in f for ext in ['.enc', '.dat', '.bin', '.encrypted']):
                encrypted.append(f)
            elif f.endswith('.so') or f.endswith('.dex'):
                encrypted.append(f"Suspicious asset: {f}")
    
    return encrypted


def _identify_obfuscation_tools(
    indicators: List[ObfuscationIndicator],
    class_naming: ClassNamingAnalysis,
    string_encryption: List[StringEncryptionPattern]
) -> List[str]:
    """Identify which obfuscation tools were likely used."""
    tools = []
    
    indicator_types = [i.indicator_type for i in indicators]
    
    # ProGuard/R8 detection
    if any(t.startswith('proguard') for t in indicator_types):
        if class_naming.obfuscation_ratio > 0.5:
            tools.append("ProGuard/R8 (aggressive)")
        else:
            tools.append("ProGuard/R8 (standard)")
    
    # DexGuard detection
    if 'dexguard' in indicator_types:
        tools.append("DexGuard (commercial)")
    
    # String encryption
    if string_encryption:
        encryption_types = set(p.pattern_name for p in string_encryption)
        if 'aes_decrypt' in encryption_types:
            tools.append("AES String Encryption")
        if 'xor_decrypt' in encryption_types:
            tools.append("XOR String Encryption")
    
    # Integrity checks
    if 'integrity_check' in indicator_types:
        tools.append("Integrity/Tamper Detection")
    
    # Reflection hiding
    if 'reflection_hiding' in indicator_types:
        tools.append("Reflection-based API Hiding")
    
    if not tools and class_naming.obfuscation_ratio > 0.2:
        tools.append("Basic Minification")
    
    return tools


def _calculate_obfuscation_score(
    class_naming: ClassNamingAnalysis,
    indicators: List[ObfuscationIndicator],
    string_encryption: List[StringEncryptionPattern],
    control_flow: List[ControlFlowObfuscation],
    native_protection: NativeProtection
) -> int:
    """Calculate an overall obfuscation score (0-100)."""
    score = 0
    
    # Class naming (up to 30 points)
    score += int(class_naming.obfuscation_ratio * 30)
    
    # Indicators (up to 25 points)
    high_conf = sum(1 for i in indicators if i.confidence == "high")
    med_conf = sum(1 for i in indicators if i.confidence == "medium")
    score += min(high_conf * 5 + med_conf * 2, 25)
    
    # String encryption (up to 20 points)
    if string_encryption:
        score += min(len(string_encryption) * 4, 20)
    
    # Control flow (up to 15 points)
    for cf in control_flow:
        score += int(cf.complexity_score * 7)
    score = min(score, 85)  # Cap at 85 before native
    
    # Native protection (up to 15 points)
    if native_protection.protection_indicators:
        score += min(len(native_protection.protection_indicators) * 5, 15)
    
    return min(score, 100)


def _generate_deobfuscation_strategies(
    tools: List[str],
    indicators: List[ObfuscationIndicator],
    string_encryption: List[StringEncryptionPattern]
) -> List[str]:
    """Generate recommended deobfuscation strategies."""
    strategies = []
    
    if any('ProGuard' in t for t in tools):
        strategies.append("Use JADX with deobfuscation enabled (--deobf flag)")
        strategies.append("Look for mapping.txt in the APK or build artifacts")
        strategies.append("Use JEB or similar tool for advanced rename suggestions")
    
    if any('DexGuard' in t for t in tools):
        strategies.append("Dynamic analysis with Frida recommended for DexGuard")
        strategies.append("Hook String decryption methods at runtime")
        strategies.append("Consider using DexGuard's commercial tools for analysis")
    
    if string_encryption:
        strategies.append("Hook string decryption methods with Frida to dump decrypted strings")
        strategies.append("Trace method calls to identify decryption routines")
        strategies.append("Memory dump analysis after app initialization")
    
    if any(i.indicator_type == 'reflection_hiding' for i in indicators):
        strategies.append("Hook Class.forName() and Method.invoke() to trace hidden API calls")
        strategies.append("Use Frida's Java.choose() to enumerate instantiated classes")
    
    if any(i.indicator_type == 'integrity_check' for i in indicators):
        strategies.append("Identify and bypass integrity checks before modification")
        strategies.append("Hook signature verification methods")
    
    if not strategies:
        strategies.append("APK appears to have minimal obfuscation")
        strategies.append("Standard static analysis with JADX should be effective")
    
    return strategies


def _recommend_deobfuscation_tools(tools: List[str], level: str) -> List[str]:
    """Recommend tools based on obfuscation level."""
    recommended = []
    
    # Always recommend JADX
    recommended.append("JADX - Primary decompiler with deobfuscation support")
    
    if level in ['moderate', 'heavy', 'extreme']:
        recommended.append("Frida - Dynamic instrumentation for runtime analysis")
        recommended.append("Objection - Frida-based runtime exploration")
    
    if level in ['heavy', 'extreme']:
        recommended.append("JEB Decompiler - Commercial tool with advanced deobfuscation")
        recommended.append("Ghidra with DEX support - For native library analysis")
    
    if any('DexGuard' in t for t in tools):
        recommended.append("Consider commercial tools (JEB, IDA Pro) for DexGuard")
    
    if level == 'extreme':
        recommended.append("Custom Frida scripts for targeted analysis")
        recommended.append("Memory forensics tools (Volatility, etc.)")
    
    return recommended


def _generate_string_decrypt_hook(class_name: str, method_name: str) -> str:
    """Generate a Frida hook for a string decryption method."""
    java_class = class_name.replace('/', '.').lstrip('L').rstrip(';')
    
    return f'''Java.perform(function() {{
    var DecryptClass = Java.use("{java_class}");
    DecryptClass.{method_name}.overload("java.lang.String").implementation = function(encrypted) {{
        var result = this.{method_name}(encrypted);
        console.log("[DECRYPT] " + encrypted + " -> " + result);
        return result;
    }};
}});'''


def _generate_frida_hooks(
    string_encryption: List[StringEncryptionPattern],
    dx,
    package_name: str
) -> List[str]:
    """Generate useful Frida hooks based on analysis."""
    hooks = []
    
    # String decryption hooks
    for pattern in string_encryption[:3]:  # Limit to 3
        if pattern.suggested_frida_hook:
            hooks.append(pattern.suggested_frida_hook)
    
    # Generic hooks
    hooks.append(f'''// Hook app launch
Java.perform(function() {{
    var Activity = Java.use("android.app.Activity");
    Activity.onCreate.overload("android.os.Bundle").implementation = function(bundle) {{
        console.log("[ACTIVITY] " + this.getClass().getName() + " created");
        this.onCreate(bundle);
    }};
}});''')
    
    hooks.append('''// Dump all loaded classes
Java.perform(function() {
    Java.enumerateLoadedClasses({
        onMatch: function(className) {
            if (className.includes("''' + package_name.split('.')[0] + '''")) {
                console.log("[CLASS] " + className);
            }
        },
        onComplete: function() {}
    });
});''')
    
    hooks.append('''// Hook reflection (Class.forName)
Java.perform(function() {
    var Class = Java.use("java.lang.Class");
    Class.forName.overload("java.lang.String").implementation = function(name) {
        console.log("[REFLECTION] Class.forName: " + name);
        return this.forName(name);
    };
});''')
    
    return hooks


# ============================================================================
# Binary Entropy Analysis
# ============================================================================

@dataclass
class EntropyDataPoint:
    """A single entropy measurement at an offset."""
    offset: int
    entropy: float  # 0.0 to 8.0 (bits per byte)
    size: int  # Window size used


@dataclass
class EntropyRegion:
    """A region with notable entropy characteristics."""
    start_offset: int
    end_offset: int
    avg_entropy: float
    max_entropy: float
    min_entropy: float
    classification: str  # "packed", "encrypted", "code", "data", "sparse"
    section_name: Optional[str] = None
    description: str = ""


@dataclass
class EntropyAnalysisResult:
    """Complete entropy analysis result for a binary."""
    filename: str
    file_size: int
    overall_entropy: float
    entropy_data: List[EntropyDataPoint]  # For visualization
    regions: List[EntropyRegion]
    is_likely_packed: bool
    packing_confidence: float  # 0.0 to 1.0
    detected_packers: List[str]
    section_entropy: List[Dict[str, Any]]  # Per-section entropy for PE/ELF
    analysis_notes: List[str]
    window_size: int
    step_size: int


def calculate_entropy(data: bytes) -> float:
    """Calculate Shannon entropy of data (bits per byte, 0-8)."""
    if not data:
        return 0.0
    
    # Count byte frequencies
    freq = [0] * 256
    for byte in data:
        freq[byte] += 1
    
    # Calculate entropy
    length = len(data)
    entropy = 0.0
    for count in freq:
        if count > 0:
            p = count / length
            entropy -= p * math.log2(p)
    
    return entropy


def analyze_binary_entropy(
    file_path: Path,
    window_size: int = 256,
    step_size: int = 128
) -> EntropyAnalysisResult:
    """
    Analyze entropy distribution across a binary file.
    
    Args:
        file_path: Path to the binary file
        window_size: Size of sliding window for entropy calculation
        step_size: Step size between measurements
    
    Returns:
        EntropyAnalysisResult with entropy data and classification
    """
    import time
    start_time = time.time()
    
    filename = file_path.name
    file_size = file_path.stat().st_size
    analysis_notes = []
    
    # Read the entire file
    with open(file_path, 'rb') as f:
        data = f.read()
    
    # Calculate overall entropy
    overall_entropy = calculate_entropy(data)
    
    # Calculate entropy at regular intervals
    entropy_data = []
    offset = 0
    while offset + window_size <= len(data):
        window = data[offset:offset + window_size]
        entropy = calculate_entropy(window)
        entropy_data.append(EntropyDataPoint(
            offset=offset,
            entropy=round(entropy, 4),
            size=window_size
        ))
        offset += step_size
    
    # Analyze per-section entropy for PE/ELF files
    section_entropy = []
    regions = []
    
    # Try PE analysis
    if PEFILE_AVAILABLE and data[:2] == b'MZ':
        try:
            pe = pefile.PE(data=data)
            for section in pe.sections:
                section_name = section.Name.decode('utf-8', errors='ignore').rstrip('\x00')
                section_data = section.get_data()
                sect_entropy = calculate_entropy(section_data)
                
                section_entropy.append({
                    'name': section_name,
                    'virtual_address': section.VirtualAddress,
                    'raw_size': section.SizeOfRawData,
                    'virtual_size': section.Misc_VirtualSize,
                    'entropy': round(sect_entropy, 4),
                    'characteristics': hex(section.Characteristics)
                })
                
                # Classify section
                classification = _classify_entropy_region(sect_entropy, section_name)
                regions.append(EntropyRegion(
                    start_offset=section.PointerToRawData,
                    end_offset=section.PointerToRawData + section.SizeOfRawData,
                    avg_entropy=sect_entropy,
                    max_entropy=sect_entropy,
                    min_entropy=sect_entropy,
                    classification=classification,
                    section_name=section_name,
                    description=_get_entropy_description(sect_entropy, section_name)
                ))
            pe.close()
        except Exception as e:
            analysis_notes.append(f"PE parsing error: {str(e)}")
    
    # Try ELF analysis
    elif PYELFTOOLS_AVAILABLE and data[:4] == b'\x7fELF':
        try:
            from io import BytesIO
            elf = ELFFile(BytesIO(data))
            for section in elf.iter_sections():
                section_name = section.name
                if section.data_size > 0:
                    section_data = section.data()
                    sect_entropy = calculate_entropy(section_data)
                    
                    section_entropy.append({
                        'name': section_name,
                        'address': section['sh_addr'],
                        'size': section.data_size,
                        'entropy': round(sect_entropy, 4),
                        'type': section['sh_type']
                    })
                    
                    classification = _classify_entropy_region(sect_entropy, section_name)
                    regions.append(EntropyRegion(
                        start_offset=section['sh_offset'],
                        end_offset=section['sh_offset'] + section.data_size,
                        avg_entropy=sect_entropy,
                        max_entropy=sect_entropy,
                        min_entropy=sect_entropy,
                        classification=classification,
                        section_name=section_name,
                        description=_get_entropy_description(sect_entropy, section_name)
                    ))
        except Exception as e:
            analysis_notes.append(f"ELF parsing error: {str(e)}")
    
    # Identify high-entropy regions if no section info
    if not regions:
        regions = _find_entropy_regions(entropy_data)
    
    # Detect packing
    is_packed, pack_confidence, packers = _detect_packing(
        overall_entropy, section_entropy, data, regions
    )
    
    if is_packed:
        analysis_notes.append(f"Binary appears to be packed (confidence: {pack_confidence:.0%})")
        if packers:
            analysis_notes.append(f"Possible packers: {', '.join(packers)}")
    
    # Add entropy classification notes
    if overall_entropy > 7.5:
        analysis_notes.append("Very high overall entropy - likely encrypted or compressed")
    elif overall_entropy > 7.0:
        analysis_notes.append("High overall entropy - possibly packed or obfuscated")
    elif overall_entropy < 4.0:
        analysis_notes.append("Low overall entropy - mostly data or sparse content")
    
    return EntropyAnalysisResult(
        filename=filename,
        file_size=file_size,
        overall_entropy=round(overall_entropy, 4),
        entropy_data=entropy_data,
        regions=regions,
        is_likely_packed=is_packed,
        packing_confidence=round(pack_confidence, 3),
        detected_packers=packers,
        section_entropy=section_entropy,
        analysis_notes=analysis_notes,
        window_size=window_size,
        step_size=step_size
    )


def _classify_entropy_region(entropy: float, section_name: str = "") -> str:
    """Classify a region based on its entropy."""
    section_lower = section_name.lower()
    
    # Known section patterns
    if section_lower in ['.text', '__text', 'code']:
        if entropy > 6.5:
            return "packed_code"
        return "code"
    
    if section_lower in ['.data', '__data', '.rdata', '__const']:
        if entropy > 7.0:
            return "encrypted"
        return "data"
    
    if section_lower in ['.rsrc', '__DATA']:
        return "resources"
    
    # Entropy-based classification
    if entropy > 7.8:
        return "encrypted"
    elif entropy > 7.2:
        return "packed"
    elif entropy > 6.0:
        return "code"
    elif entropy > 4.0:
        return "data"
    elif entropy > 1.0:
        return "sparse"
    else:
        return "empty"


def _get_entropy_description(entropy: float, section_name: str = "") -> str:
    """Get a human-readable description of entropy level."""
    if entropy > 7.8:
        return "Very high entropy - encrypted or heavily compressed"
    elif entropy > 7.2:
        return "High entropy - likely packed or compressed"
    elif entropy > 6.5:
        return "Elevated entropy - compiled code or light compression"
    elif entropy > 5.5:
        return "Moderate entropy - typical executable code"
    elif entropy > 4.0:
        return "Low-moderate entropy - mixed code and data"
    elif entropy > 2.0:
        return "Low entropy - mostly data or strings"
    else:
        return "Very low entropy - sparse or repetitive data"


def _find_entropy_regions(entropy_data: List[EntropyDataPoint]) -> List[EntropyRegion]:
    """Find distinct entropy regions in the data."""
    if not entropy_data:
        return []
    
    regions = []
    current_class = _classify_entropy_region(entropy_data[0].entropy)
    start_offset = entropy_data[0].offset
    entropies = [entropy_data[0].entropy]
    
    for point in entropy_data[1:]:
        point_class = _classify_entropy_region(point.entropy)
        
        if point_class != current_class:
            # End current region
            regions.append(EntropyRegion(
                start_offset=start_offset,
                end_offset=point.offset,
                avg_entropy=sum(entropies) / len(entropies),
                max_entropy=max(entropies),
                min_entropy=min(entropies),
                classification=current_class,
                description=_get_entropy_description(sum(entropies) / len(entropies))
            ))
            # Start new region
            current_class = point_class
            start_offset = point.offset
            entropies = [point.entropy]
        else:
            entropies.append(point.entropy)
    
    # Add final region
    if entropies:
        regions.append(EntropyRegion(
            start_offset=start_offset,
            end_offset=entropy_data[-1].offset + entropy_data[-1].size,
            avg_entropy=sum(entropies) / len(entropies),
            max_entropy=max(entropies),
            min_entropy=min(entropies),
            classification=current_class,
            description=_get_entropy_description(sum(entropies) / len(entropies))
        ))
    
    return regions


def _detect_packing(
    overall_entropy: float,
    section_entropy: List[Dict],
    data: bytes,
    regions: List[EntropyRegion]
) -> tuple:
    """Detect if binary is likely packed and identify possible packers."""
    is_packed = False
    confidence = 0.0
    packers = []
    
    # Check overall entropy
    if overall_entropy > 7.2:
        is_packed = True
        confidence += 0.3
    elif overall_entropy > 6.8:
        confidence += 0.15
    
    # Check for high entropy in code sections
    for sect in section_entropy:
        if sect.get('name', '').lower() in ['.text', '__text', 'code', '.code']:
            if sect['entropy'] > 7.0:
                is_packed = True
                confidence += 0.3
            elif sect['entropy'] > 6.5:
                confidence += 0.15
    
    # Check for UPX signature
    if b'UPX!' in data or b'UPX0' in data or b'UPX1' in data:
        is_packed = True
        packers.append('UPX')
        confidence += 0.4
    
    # Check for other packer signatures
    packer_signatures = {
        b'PEC2': 'PECompact',
        b'ASPack': 'ASPack',
        b'.aspack': 'ASPack',
        b'.adata': 'ASProtect',
        b'FSG!': 'FSG',
        b'MPRESS': 'MPRESS',
        b'.nsp0': 'NsPack',
        b'.nsp1': 'NsPack',
        b'MEW': 'MEW',
        b'.petite': 'Petite',
        b'PEtite': 'Petite',
        b'Themida': 'Themida',
        b'.themida': 'Themida',
        b'VMProtect': 'VMProtect',
        b'.vmp0': 'VMProtect',
        b'.vmp1': 'VMProtect',
        b'Obsidium': 'Obsidium',
        b'.enigma': 'Enigma Protector',
    }
    
    for sig, name in packer_signatures.items():
        if sig in data:
            is_packed = True
            if name not in packers:
                packers.append(name)
            confidence += 0.3
    
    # Check for small code section with high entropy data
    if section_entropy:
        text_size = 0
        data_entropy = 0
        for sect in section_entropy:
            name = sect.get('name', '').lower()
            if name in ['.text', '__text', 'code']:
                text_size = sect.get('raw_size', sect.get('size', 0))
            elif name in ['.data', '.rdata'] and sect['entropy'] > 7.0:
                data_entropy = sect['entropy']
        
        if text_size < 10000 and data_entropy > 7.0:
            is_packed = True
            confidence += 0.2
    
    # Cap confidence at 1.0
    confidence = min(confidence, 1.0)
    
    # Set packed flag based on confidence
    if confidence > 0.5:
        is_packed = True
    
    return is_packed, confidence, packers


# ============================================================================
# APK Report Export Functions
# ============================================================================

def generate_apk_markdown_report(
    result: ApkAnalysisResult,
    report_type: str = "both"
) -> str:
    """
    Generate a formatted Markdown report for APK analysis.
    
    Args:
        result: APK analysis result
        report_type: "functionality", "security", or "both"
    
    Returns:
        Markdown formatted string
    """
    md = []
    
    # Header
    md.append(f"# APK Analysis Report")
    md.append(f"**Package:** {result.package_name}")
    md.append(f"**App Name:** {result.app_name or 'Unknown'}")
    md.append(f"**Version:** {result.version_name} (code: {result.version_code})")
    md.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    md.append("")
    md.append("---")
    md.append("")
    
    # ==================== FUNCTIONALITY REPORT ====================
    if report_type in ["functionality", "both"]:
        md.append("## 📱 What Does This APK Do")
        md.append("")
        
        # App Overview
        md.append("### App Overview")
        md.append("")
        md.append(f"- **Package:** `{result.package_name}`")
        md.append(f"- **Min SDK:** {result.min_sdk} (Android {5 + (result.min_sdk - 21) if result.min_sdk and result.min_sdk >= 21 else 'Legacy'})")
        md.append(f"- **Target SDK:** {result.target_sdk}")
        md.append(f"- **Debuggable:** {'Yes ⚠️' if result.debuggable else 'No ✅'}")
        md.append(f"- **Allows Backup:** {'Yes ⚠️' if result.allow_backup else 'No ✅'}")
        md.append("")
        
        # Components
        md.append("### App Components")
        md.append("")
        md.append(f"- **Activities:** {len(result.activities)}")
        md.append(f"- **Services:** {len(result.services)}")
        md.append(f"- **Broadcast Receivers:** {len(result.receivers)}")
        md.append(f"- **Content Providers:** {len(result.providers)}")
        md.append("")
        
        if result.activities:
            md.append("#### Activities")
            for act in result.activities[:10]:
                md.append(f"- `{act}`")
            if len(result.activities) > 10:
                md.append(f"- ... and {len(result.activities) - 10} more")
            md.append("")
        
        if result.services:
            md.append("#### Services")
            for svc in result.services[:10]:
                md.append(f"- `{svc}`")
            if len(result.services) > 10:
                md.append(f"- ... and {len(result.services) - 10} more")
            md.append("")
        
        # Permissions Summary
        md.append("### Permissions Summary")
        md.append("")
        dangerous_perms = [p for p in result.permissions if p.is_dangerous]
        md.append(f"- **Total Permissions:** {len(result.permissions)}")
        md.append(f"- **Dangerous Permissions:** {len(dangerous_perms)}")
        md.append("")
        
        if dangerous_perms:
            md.append("#### Dangerous Permissions")
            for perm in dangerous_perms:
                desc = perm.description or "No description"
                md.append(f"- **{perm.name.replace('android.permission.', '')}**: {desc}")
            md.append("")
        
        # Network Indicators
        if result.urls:
            md.append("### Network Communication")
            md.append("")
            md.append("URLs found in the APK:")
            for url in result.urls[:15]:
                md.append(f"- `{url[:100]}`")
            if len(result.urls) > 15:
                md.append(f"- ... and {len(result.urls) - 15} more")
            md.append("")
        
        # Native Libraries
        if result.native_libraries:
            md.append("### Native Libraries")
            md.append("")
            for lib in result.native_libraries:
                md.append(f"- `{lib}`")
            md.append("")
        
        # AI Functionality Report
        if result.ai_report_functionality:
            md.append("### AI Analysis: App Functionality")
            md.append("")
            # Convert HTML to basic markdown
            html_content = result.ai_report_functionality
            # Strip HTML tags but keep content
            import re
            html_content = re.sub(r'<h3[^>]*>', '\n#### ', html_content)
            html_content = re.sub(r'</h3>', '\n', html_content)
            html_content = re.sub(r'<strong>', '**', html_content)
            html_content = re.sub(r'</strong>', '**', html_content)
            html_content = re.sub(r'<li>', '- ', html_content)
            html_content = re.sub(r'</li>', '\n', html_content)
            html_content = re.sub(r'<[^>]+>', '', html_content)
            md.append(html_content)
            md.append("")
        
        md.append("---")
        md.append("")
    
    # ==================== SECURITY REPORT ====================
    if report_type in ["security", "both"]:
        md.append("## 🔒 Security Findings")
        md.append("")
        
        # Security Overview
        md.append("### Security Overview")
        md.append("")
        md.append(f"- **Security Issues Found:** {len(result.security_issues)}")
        md.append(f"- **Secrets Detected:** {len(result.secrets)}")
        md.append(f"- **Dangerous Permissions:** {len([p for p in result.permissions if p.is_dangerous])}")
        md.append("")
        
        # Certificate Analysis
        if result.certificate:
            md.append("### Certificate Analysis")
            md.append("")
            cert = result.certificate
            status = "⚠️ DEBUG CERTIFICATE" if cert.is_debug_cert else "✅ Production Certificate"
            md.append(f"- **Status:** {status}")
            md.append(f"- **Subject:** `{cert.subject}`")
            md.append(f"- **Issuer:** `{cert.issuer}`")
            md.append(f"- **Valid From:** {cert.valid_from}")
            md.append(f"- **Valid Until:** {cert.valid_until}")
            md.append(f"- **Signature Version:** {cert.signature_version}")
            if cert.is_expired:
                md.append("- **⚠️ EXPIRED**")
            md.append("")
        
        # Security Issues
        if result.security_issues:
            md.append("### Security Issues")
            md.append("")
            
            # Group by severity
            by_severity = {"critical": [], "high": [], "medium": [], "low": [], "info": []}
            for issue in result.security_issues:
                sev = issue.get("severity", "info").lower()
                if sev in by_severity:
                    by_severity[sev].append(issue)
                else:
                    by_severity["info"].append(issue)
            
            severity_emoji = {
                "critical": "🔴",
                "high": "🟠",
                "medium": "🟡",
                "low": "🟢",
                "info": "🔵"
            }
            
            for sev in ["critical", "high", "medium", "low", "info"]:
                issues = by_severity[sev]
                if issues:
                    md.append(f"#### {severity_emoji[sev]} {sev.upper()} ({len(issues)})")
                    md.append("")
                    for issue in issues:
                        md.append(f"- **{issue.get('category', 'Unknown')}:** {issue.get('description', 'No description')}")
                        if issue.get('recommendation'):
                            md.append(f"  - 💡 *Recommendation: {issue['recommendation']}*")
                    md.append("")
        
        # Secrets Found
        if result.secrets:
            md.append("### Secrets & Hardcoded Credentials")
            md.append("")
            md.append("| Type | Value (Masked) | Severity |")
            md.append("|------|----------------|----------|")
            for secret in result.secrets[:20]:
                md.append(f"| {secret.get('type', 'Unknown')} | `{secret.get('masked_value', '***')}` | {secret.get('severity', 'Unknown')} |")
            if len(result.secrets) > 20:
                md.append(f"| ... | {len(result.secrets) - 20} more | ... |")
            md.append("")
        
        # Hardening Score
        if result.hardening_score:
            hs = result.hardening_score
            md.append("### Security Hardening Score")
            md.append("")
            md.append(f"- **Overall Score:** {hs.get('score', 0)}/100")
            md.append(f"- **Grade:** {hs.get('grade', 'N/A')}")
            md.append("")
            
            if hs.get('passed_checks'):
                md.append("#### ✅ Passed Checks")
                for check in hs['passed_checks'][:10]:
                    md.append(f"- {check}")
                md.append("")
            
            if hs.get('failed_checks'):
                md.append("#### ❌ Failed Checks")
                for check in hs['failed_checks'][:10]:
                    md.append(f"- {check}")
                md.append("")
        
        # AI Security Report
        if result.ai_report_security:
            md.append("### AI Security Analysis")
            md.append("")
            # Convert HTML to basic markdown
            import re
            html_content = result.ai_report_security
            html_content = re.sub(r'<h3[^>]*>', '\n#### ', html_content)
            html_content = re.sub(r'</h3>', '\n', html_content)
            html_content = re.sub(r'<strong>', '**', html_content)
            html_content = re.sub(r'</strong>', '**', html_content)
            html_content = re.sub(r'<li>', '- ', html_content)
            html_content = re.sub(r'</li>', '\n', html_content)
            html_content = re.sub(r'<span[^>]*>([^<]+)</span>', r'\1', html_content)
            html_content = re.sub(r'<[^>]+>', '', html_content)
            md.append(html_content)
            md.append("")
    
    return "\n".join(md)


def generate_apk_pdf_report(result: ApkAnalysisResult, report_type: str = "both") -> bytes:
    """
    Generate a PDF report for APK analysis.
    
    Args:
        result: APK analysis result
        report_type: "functionality", "security", or "both"
    
    Returns:
        PDF bytes
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        import io
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        styles.add(ParagraphStyle(
            name='Title',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=20,
            textColor=colors.HexColor("#1e40af")
        ))
        styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor("#1e40af")
        ))
        styles.add(ParagraphStyle(
            name='SubHeader',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=15,
            spaceAfter=8,
            textColor=colors.HexColor("#374151")
        ))
        styles.add(ParagraphStyle(
            name='Body',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            leading=14
        ))
        styles.add(ParagraphStyle(
            name='Code',
            parent=styles['Normal'],
            fontSize=9,
            fontName='Courier',
            backColor=colors.HexColor("#f3f4f6"),
            leftIndent=10
        ))
        styles.add(ParagraphStyle(
            name='BulletItem',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceAfter=4
        ))
        
        story = []
        
        # Title
        story.append(Paragraph("APK Analysis Report", styles['Title']))
        story.append(Spacer(1, 10))
        
        # App Info Table
        app_info = [
            ["Package", result.package_name],
            ["App Name", result.app_name or "Unknown"],
            ["Version", f"{result.version_name} ({result.version_code})"],
            ["Min SDK", str(result.min_sdk)],
            ["Target SDK", str(result.target_sdk)],
        ]
        t = Table(app_info, colWidths=[1.5*inch, 4*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#f3f4f6")),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('PADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
        ]))
        story.append(t)
        story.append(Spacer(1, 20))
        
        # ==================== FUNCTIONALITY SECTION ====================
        if report_type in ["functionality", "both"]:
            story.append(Paragraph("What Does This APK Do", styles['SectionHeader']))
            story.append(Spacer(1, 10))
            
            # Components summary
            story.append(Paragraph("App Components", styles['SubHeader']))
            components = [
                ["Component Type", "Count"],
                ["Activities", str(len(result.activities))],
                ["Services", str(len(result.services))],
                ["Receivers", str(len(result.receivers))],
                ["Providers", str(len(result.providers))],
            ]
            t = Table(components, colWidths=[2*inch, 1*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1e40af")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ]))
            story.append(t)
            story.append(Spacer(1, 15))
            
            # Permissions
            story.append(Paragraph("Permissions", styles['SubHeader']))
            dangerous_perms = [p for p in result.permissions if p.is_dangerous]
            story.append(Paragraph(f"Total: {len(result.permissions)} | Dangerous: {len(dangerous_perms)}", styles['Body']))
            
            if dangerous_perms:
                story.append(Spacer(1, 5))
                story.append(Paragraph("Dangerous Permissions:", styles['Body']))
                for perm in dangerous_perms[:10]:
                    name = perm.name.replace('android.permission.', '')
                    story.append(Paragraph(f"• {name}", styles['BulletItem']))
            
            story.append(Spacer(1, 15))
            
            # Native libraries
            if result.native_libraries:
                story.append(Paragraph("Native Libraries", styles['SubHeader']))
                for lib in result.native_libraries[:8]:
                    story.append(Paragraph(f"• {lib}", styles['BulletItem']))
            
            if report_type == "both":
                story.append(PageBreak())
        
        # ==================== SECURITY SECTION ====================
        if report_type in ["security", "both"]:
            story.append(Paragraph("Security Findings", styles['SectionHeader']))
            story.append(Spacer(1, 10))
            
            # Security Summary
            summary_data = [
                ["Metric", "Value", "Status"],
                ["Security Issues", str(len(result.security_issues)), "⚠️" if result.security_issues else "✅"],
                ["Secrets Found", str(len(result.secrets)), "⚠️" if result.secrets else "✅"],
                ["Debuggable", "Yes" if result.debuggable else "No", "⚠️" if result.debuggable else "✅"],
                ["Allows Backup", "Yes" if result.allow_backup else "No", "⚠️" if result.allow_backup else "✅"],
            ]
            t = Table(summary_data, colWidths=[2*inch, 1.5*inch, 0.5*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#dc2626")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('PADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
            ]))
            story.append(t)
            story.append(Spacer(1, 20))
            
            # Security Issues by Severity
            if result.security_issues:
                story.append(Paragraph("Security Issues", styles['SubHeader']))
                
                severity_colors = {
                    'critical': colors.HexColor("#dc2626"),
                    'high': colors.HexColor("#ea580c"),
                    'medium': colors.HexColor("#ca8a04"),
                    'low': colors.HexColor("#16a34a"),
                    'info': colors.HexColor("#6b7280"),
                }
                
                for issue in result.security_issues[:15]:
                    sev = issue.get('severity', 'info').lower()
                    color = severity_colors.get(sev, colors.black)
                    story.append(Paragraph(
                        f"<font color='{color}'>[{sev.upper()}]</font> <b>{issue.get('category', 'Unknown')}:</b> {issue.get('description', '')}",
                        styles['BulletItem']
                    ))
                story.append(Spacer(1, 15))
            
            # Secrets
            if result.secrets:
                story.append(Paragraph("Hardcoded Secrets", styles['SubHeader']))
                secrets_data = [["Type", "Value (Masked)", "Severity"]]
                for secret in result.secrets[:10]:
                    secrets_data.append([
                        secret.get('type', 'Unknown'),
                        secret.get('masked_value', '***')[:40],
                        secret.get('severity', 'Unknown')
                    ])
                t = Table(secrets_data, colWidths=[1.5*inch, 3*inch, 1*inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#7f1d1d")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('PADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                ]))
                story.append(t)
                story.append(Spacer(1, 15))
            
            # Certificate
            if result.certificate:
                story.append(Paragraph("Certificate Analysis", styles['SubHeader']))
                cert = result.certificate
                cert_status = "DEBUG CERTIFICATE - INSECURE" if cert.is_debug_cert else "Production Certificate"
                cert_data = [
                    ["Status", cert_status],
                    ["Subject", cert.subject[:60] if cert.subject else "Unknown"],
                    ["Valid From", cert.valid_from or "Unknown"],
                    ["Valid Until", cert.valid_until or "Unknown"],
                    ["Signature", cert.signature_version or "Unknown"],
                ]
                t = Table(cert_data, colWidths=[1.5*inch, 4*inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#fef2f2") if cert.is_debug_cert else colors.HexColor("#f0fdf4")),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('PADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                ]))
                story.append(t)
        
        doc.build(story)
        return buffer.getvalue()
        
    except ImportError as e:
        logger.error(f"PDF generation failed - missing dependency: {e}")
        raise RuntimeError("PDF export requires reportlab. Install with: pip install reportlab")


def generate_apk_docx_report(result: ApkAnalysisResult, report_type: str = "both") -> bytes:
    """
    Generate a Word document report for APK analysis.
    
    Args:
        result: APK analysis result
        report_type: "functionality", "security", or "both"
    
    Returns:
        DOCX bytes
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import io
        
        doc = Document()
        
        # Title
        title = doc.add_heading('APK Analysis Report', 0)
        title.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        
        # App Info
        doc.add_paragraph()
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        info_data = [
            ("Package", result.package_name),
            ("App Name", result.app_name or "Unknown"),
            ("Version", f"{result.version_name} ({result.version_code})"),
            ("Min SDK", str(result.min_sdk)),
            ("Target SDK", str(result.target_sdk)),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # ==================== FUNCTIONALITY SECTION ====================
        if report_type in ["functionality", "both"]:
            doc.add_heading('What Does This APK Do', 1)
            
            # Components
            doc.add_heading('App Components', 2)
            comp_table = doc.add_table(rows=5, cols=2)
            comp_table.style = 'Table Grid'
            comp_data = [
                ("Component Type", "Count"),
                ("Activities", str(len(result.activities))),
                ("Services", str(len(result.services))),
                ("Receivers", str(len(result.receivers))),
                ("Providers", str(len(result.providers))),
            ]
            for i, (label, value) in enumerate(comp_data):
                comp_table.rows[i].cells[0].text = label
                comp_table.rows[i].cells[1].text = value
                if i == 0:
                    comp_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                    comp_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
            
            doc.add_paragraph()
            
            # Permissions
            doc.add_heading('Permissions', 2)
            dangerous_perms = [p for p in result.permissions if p.is_dangerous]
            doc.add_paragraph(f"Total: {len(result.permissions)} | Dangerous: {len(dangerous_perms)}")
            
            if dangerous_perms:
                doc.add_paragraph("Dangerous Permissions:", style='Intense Quote')
                for perm in dangerous_perms[:10]:
                    name = perm.name.replace('android.permission.', '')
                    doc.add_paragraph(name, style='List Bullet')
            
            # Native libraries
            if result.native_libraries:
                doc.add_heading('Native Libraries', 2)
                for lib in result.native_libraries[:8]:
                    doc.add_paragraph(lib, style='List Bullet')
            
            if report_type == "both":
                doc.add_page_break()
        
        # ==================== SECURITY SECTION ====================
        if report_type in ["security", "both"]:
            sec_heading = doc.add_heading('Security Findings', 1)
            sec_heading.runs[0].font.color.rgb = RGBColor(220, 38, 38)
            
            # Summary
            doc.add_heading('Security Summary', 2)
            summary_table = doc.add_table(rows=5, cols=3)
            summary_table.style = 'Table Grid'
            summary_data = [
                ("Metric", "Value", "Status"),
                ("Security Issues", str(len(result.security_issues)), "⚠️" if result.security_issues else "✅"),
                ("Secrets Found", str(len(result.secrets)), "⚠️" if result.secrets else "✅"),
                ("Debuggable", "Yes" if result.debuggable else "No", "⚠️" if result.debuggable else "✅"),
                ("Allows Backup", "Yes" if result.allow_backup else "No", "⚠️" if result.allow_backup else "✅"),
            ]
            for i, (metric, value, status) in enumerate(summary_data):
                summary_table.rows[i].cells[0].text = metric
                summary_table.rows[i].cells[1].text = value
                summary_table.rows[i].cells[2].text = status
                if i == 0:
                    for cell in summary_table.rows[i].cells:
                        cell.paragraphs[0].runs[0].bold = True
            
            doc.add_paragraph()
            
            # Security Issues
            if result.security_issues:
                doc.add_heading('Security Issues', 2)
                for issue in result.security_issues[:15]:
                    sev = issue.get('severity', 'info').upper()
                    cat = issue.get('category', 'Unknown')
                    desc = issue.get('description', '')
                    p = doc.add_paragraph()
                    run = p.add_run(f"[{sev}] ")
                    run.bold = True
                    if sev == 'CRITICAL':
                        run.font.color.rgb = RGBColor(220, 38, 38)
                    elif sev == 'HIGH':
                        run.font.color.rgb = RGBColor(234, 88, 12)
                    elif sev == 'MEDIUM':
                        run.font.color.rgb = RGBColor(202, 138, 4)
                    else:
                        run.font.color.rgb = RGBColor(22, 163, 74)
                    p.add_run(f"{cat}: ").bold = True
                    p.add_run(desc)
            
            # Secrets
            if result.secrets:
                doc.add_heading('Hardcoded Secrets', 2)
                secrets_table = doc.add_table(rows=len(result.secrets[:10]) + 1, cols=3)
                secrets_table.style = 'Table Grid'
                secrets_table.rows[0].cells[0].text = "Type"
                secrets_table.rows[0].cells[1].text = "Value (Masked)"
                secrets_table.rows[0].cells[2].text = "Severity"
                for cell in secrets_table.rows[0].cells:
                    cell.paragraphs[0].runs[0].bold = True
                
                for i, secret in enumerate(result.secrets[:10]):
                    secrets_table.rows[i+1].cells[0].text = secret.get('type', 'Unknown')
                    secrets_table.rows[i+1].cells[1].text = secret.get('masked_value', '***')[:40]
                    secrets_table.rows[i+1].cells[2].text = secret.get('severity', 'Unknown')
            
            # Certificate
            if result.certificate:
                doc.add_heading('Certificate Analysis', 2)
                cert = result.certificate
                cert_status = "DEBUG CERTIFICATE - INSECURE" if cert.is_debug_cert else "Production Certificate"
                
                p = doc.add_paragraph()
                p.add_run("Status: ").bold = True
                status_run = p.add_run(cert_status)
                if cert.is_debug_cert:
                    status_run.font.color.rgb = RGBColor(220, 38, 38)
                
                doc.add_paragraph(f"Subject: {cert.subject}")
                doc.add_paragraph(f"Valid: {cert.valid_from} to {cert.valid_until}")
                doc.add_paragraph(f"Signature: {cert.signature_version}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
        
    except ImportError as e:
        logger.error(f"DOCX generation failed - missing dependency: {e}")
        raise RuntimeError("Word export requires python-docx. Install with: pip install python-docx")


# Import datetime at the top level if not already
from datetime import datetime


# ============================================================================
# DATA FLOW ANALYSIS - Taint Tracking for Android Apps
# ============================================================================

@dataclass
class TaintSource:
    """A source of sensitive/tainted data."""
    source_type: str  # "user_input", "device_info", "sensor", "location", "contacts", etc.
    class_name: str
    method_name: str
    description: str
    sensitivity: str  # "low", "medium", "high", "critical"
    owasp_category: str  # M1-M10 mapping


@dataclass
class TaintSink:
    """A sink where tainted data flows to."""
    sink_type: str  # "network", "storage", "log", "ipc", "sms", etc.
    class_name: str
    method_name: str
    description: str
    risk_level: str  # "low", "medium", "high", "critical"
    owasp_category: str


@dataclass
class DataFlowPath:
    """A path from source to sink."""
    source: TaintSource
    sink: TaintSink
    intermediate_methods: List[str]  # Methods in the call chain
    affected_class: str
    affected_method: str
    tainted_variable: Optional[str]
    code_snippet: str
    severity: str
    description: str
    recommendation: str


@dataclass
class DataFlowAnalysisResult:
    """Complete data flow analysis result."""
    total_sources: int
    total_sinks: int
    total_flows: int
    critical_flows: int
    high_risk_flows: int
    sources_found: List[Dict[str, Any]]
    sinks_found: List[Dict[str, Any]]
    data_flow_paths: List[DataFlowPath]
    privacy_violations: List[Dict[str, Any]]
    data_leak_risks: List[Dict[str, Any]]
    summary: str
    recommendations: List[str]


# Android API Sources - Where sensitive data enters the app
ANDROID_SOURCES = {
    # Location Sources
    "location": [
        ("Landroid/location/LocationManager;", "getLastKnownLocation", "Last known device location", "critical"),
        ("Landroid/location/LocationManager;", "requestLocationUpdates", "Real-time location tracking", "critical"),
        ("Lcom/google/android/gms/location/FusedLocationProviderClient;", "getLastLocation", "Google Play location", "critical"),
        ("Lcom/google/android/gms/location/FusedLocationProviderClient;", "requestLocationUpdates", "Google Play location updates", "critical"),
    ],
    # Device/Hardware ID Sources
    "device_info": [
        ("Landroid/telephony/TelephonyManager;", "getDeviceId", "Device IMEI/MEID (deprecated)", "critical"),
        ("Landroid/telephony/TelephonyManager;", "getImei", "Device IMEI", "critical"),
        ("Landroid/telephony/TelephonyManager;", "getMeid", "Device MEID", "critical"),
        ("Landroid/telephony/TelephonyManager;", "getLine1Number", "Phone number", "critical"),
        ("Landroid/telephony/TelephonyManager;", "getSimSerialNumber", "SIM serial number", "critical"),
        ("Landroid/telephony/TelephonyManager;", "getSubscriberId", "IMSI (subscriber ID)", "critical"),
        ("Landroid/provider/Settings$Secure;", "getString", "Android ID/Secure settings", "high"),
        ("Landroid/os/Build;", "SERIAL", "Hardware serial (deprecated)", "high"),
        ("Landroid/net/wifi/WifiInfo;", "getMacAddress", "WiFi MAC address", "high"),
        ("Landroid/bluetooth/BluetoothAdapter;", "getAddress", "Bluetooth MAC address", "high"),
    ],
    # User Input Sources
    "user_input": [
        ("Landroid/widget/EditText;", "getText", "Text input from user", "medium"),
        ("Landroid/widget/TextView;", "getText", "Text content", "low"),
        ("Landroid/content/Intent;", "getStringExtra", "Intent string data", "medium"),
        ("Landroid/content/Intent;", "getExtras", "Intent bundle data", "medium"),
        ("Landroid/content/Intent;", "getData", "Intent URI data", "medium"),
        ("Landroid/webkit/WebView;", "getUrl", "WebView URL", "medium"),
    ],
    # Contacts Sources
    "contacts": [
        ("Landroid/content/ContentResolver;", "query", "Content provider query (contacts, calendar, etc.)", "high"),
        ("Landroid/provider/ContactsContract;", "CONTENT_URI", "Contacts database access", "critical"),
        ("Landroid/provider/CalendarContract;", "CONTENT_URI", "Calendar database access", "high"),
    ],
    # Camera/Media Sources
    "media": [
        ("Landroid/hardware/Camera;", "takePicture", "Camera capture", "high"),
        ("Landroid/hardware/camera2/CameraCaptureSession;", "capture", "Camera2 capture", "high"),
        ("Landroid/media/MediaRecorder;", "start", "Audio/Video recording", "critical"),
        ("Landroid/media/AudioRecord;", "read", "Microphone audio data", "critical"),
    ],
    # SMS Sources
    "sms": [
        ("Landroid/telephony/SmsManager;", "getAllMessagesFromIcc", "ICC SMS messages", "critical"),
        ("Landroid/provider/Telephony$Sms;", "CONTENT_URI", "SMS database", "critical"),
    ],
    # Account Sources
    "accounts": [
        ("Landroid/accounts/AccountManager;", "getAccounts", "User accounts", "high"),
        ("Landroid/accounts/AccountManager;", "getAccountsByType", "Accounts by type", "high"),
        ("Landroid/accounts/AccountManager;", "peekAuthToken", "Auth tokens", "critical"),
    ],
    # Clipboard
    "clipboard": [
        ("Landroid/content/ClipboardManager;", "getPrimaryClip", "Clipboard content", "medium"),
        ("Landroid/content/ClipboardManager;", "getText", "Clipboard text", "medium"),
    ],
    # File System
    "file_system": [
        ("Ljava/io/FileInputStream;", "<init>", "File read", "medium"),
        ("Ljava/io/BufferedReader;", "readLine", "File line read", "medium"),
        ("Landroid/content/SharedPreferences;", "getString", "SharedPrefs read", "medium"),
        ("Landroid/content/SharedPreferences;", "getAll", "All SharedPrefs", "high"),
    ],
    # Network Responses
    "network_response": [
        ("Ljava/net/HttpURLConnection;", "getInputStream", "HTTP response", "medium"),
        ("Lokhttp3/Response;", "body", "OkHttp response body", "medium"),
        ("Lretrofit2/Response;", "body", "Retrofit response body", "medium"),
    ],
    # Crypto Keys
    "crypto": [
        ("Ljavax/crypto/SecretKey;", "getEncoded", "Secret key bytes", "critical"),
        ("Ljava/security/PrivateKey;", "getEncoded", "Private key bytes", "critical"),
        ("Landroid/security/keystore/KeyGenParameterSpec;", "Builder", "Keystore key generation", "high"),
    ],
}

# Android API Sinks - Where data exits or is stored
ANDROID_SINKS = {
    # Network Sinks
    "network": [
        ("Ljava/net/HttpURLConnection;", "getOutputStream", "HTTP request output", "high"),
        ("Ljava/net/HttpURLConnection;", "connect", "HTTP connection", "medium"),
        ("Lokhttp3/OkHttpClient;", "newCall", "OkHttp request", "high"),
        ("Lokhttp3/RequestBody;", "create", "OkHttp request body", "high"),
        ("Lretrofit2/Call;", "execute", "Retrofit sync request", "high"),
        ("Lretrofit2/Call;", "enqueue", "Retrofit async request", "high"),
        ("Ljava/net/Socket;", "getOutputStream", "Socket output", "high"),
        ("Landroid/webkit/WebView;", "loadUrl", "WebView URL load", "medium"),
        ("Landroid/webkit/WebView;", "postUrl", "WebView POST", "high"),
        ("Landroid/webkit/WebView;", "evaluateJavascript", "WebView JS execution", "high"),
    ],
    # Storage Sinks
    "storage": [
        ("Ljava/io/FileOutputStream;", "write", "File write", "medium"),
        ("Ljava/io/BufferedWriter;", "write", "Buffered file write", "medium"),
        ("Landroid/content/SharedPreferences$Editor;", "putString", "SharedPrefs write", "medium"),
        ("Landroid/content/SharedPreferences$Editor;", "commit", "SharedPrefs commit", "medium"),
        ("Landroid/database/sqlite/SQLiteDatabase;", "insert", "SQLite insert", "medium"),
        ("Landroid/database/sqlite/SQLiteDatabase;", "execSQL", "SQLite raw query", "high"),
        ("Landroid/database/sqlite/SQLiteDatabase;", "rawQuery", "SQLite raw query", "medium"),
    ],
    # Logging Sinks (data leak risk)
    "logging": [
        ("Landroid/util/Log;", "d", "Debug log", "medium"),
        ("Landroid/util/Log;", "v", "Verbose log", "low"),
        ("Landroid/util/Log;", "i", "Info log", "low"),
        ("Landroid/util/Log;", "w", "Warning log", "medium"),
        ("Landroid/util/Log;", "e", "Error log", "medium"),
        ("Ljava/io/PrintStream;", "println", "System.out.println", "medium"),
    ],
    # IPC Sinks (sending to other apps)
    "ipc": [
        ("Landroid/content/Context;", "sendBroadcast", "Send broadcast", "high"),
        ("Landroid/content/Context;", "sendOrderedBroadcast", "Send ordered broadcast", "high"),
        ("Landroid/content/Context;", "startActivity", "Start activity", "medium"),
        ("Landroid/content/Context;", "startService", "Start service", "medium"),
        ("Landroid/content/Intent;", "putExtra", "Intent data attachment", "medium"),
        ("Landroid/content/ContentResolver;", "insert", "Content provider insert", "medium"),
        ("Landroid/content/ContentResolver;", "update", "Content provider update", "medium"),
    ],
    # SMS Sinks
    "sms": [
        ("Landroid/telephony/SmsManager;", "sendTextMessage", "Send SMS", "critical"),
        ("Landroid/telephony/SmsManager;", "sendMultipartTextMessage", "Send multipart SMS", "critical"),
        ("Landroid/telephony/SmsManager;", "sendDataMessage", "Send data SMS", "critical"),
    ],
    # Clipboard Sinks
    "clipboard": [
        ("Landroid/content/ClipboardManager;", "setPrimaryClip", "Set clipboard", "medium"),
    ],
    # External Storage
    "external_storage": [
        ("Landroid/os/Environment;", "getExternalStorageDirectory", "External storage write", "high"),
        ("Landroid/os/Environment;", "getExternalStoragePublicDirectory", "Public external storage", "high"),
    ],
    # Process Execution (command injection risk)
    "process_execution": [
        ("Ljava/lang/Runtime;", "exec", "Runtime command execution", "critical"),
        ("Ljava/lang/ProcessBuilder;", "start", "Process execution", "critical"),
    ],
    # Reflection (potential security bypass)
    "reflection": [
        ("Ljava/lang/reflect/Method;", "invoke", "Reflective method call", "high"),
        ("Ljava/lang/Class;", "forName", "Dynamic class loading", "medium"),
        ("Ldalvik/system/DexClassLoader;", "loadClass", "DEX class loading", "critical"),
    ],
    # Native Code
    "native": [
        ("Ljava/lang/System;", "loadLibrary", "Load native library", "medium"),
        ("Ljava/lang/System;", "load", "Load native library path", "medium"),
    ],
}

# Privacy-sensitive data categories for GDPR/privacy compliance
PRIVACY_DATA_CATEGORIES = {
    "PII": ["device_info", "contacts", "accounts", "location"],
    "Financial": ["accounts", "clipboard"],
    "Health": ["sensor"],
    "Biometric": ["media"],
    "Communication": ["sms", "contacts"],
    "Behavioral": ["location", "user_input"],
}


def analyze_data_flow(file_path: Path, package_name: str = None) -> DataFlowAnalysisResult:
    """
    Perform comprehensive data flow analysis on an APK.
    
    Identifies:
    1. Sources of sensitive data (location, device ID, contacts, etc.)
    2. Sinks where data exits (network, storage, logs, etc.)
    3. Paths from sources to sinks (potential data leaks)
    4. Privacy violations (PII being sent to network)
    
    Args:
        file_path: Path to the APK file
        package_name: Optional package name to focus analysis on app code
        
    Returns:
        DataFlowAnalysisResult with complete analysis
    """
    sources_found = []
    sinks_found = []
    data_flow_paths = []
    privacy_violations = []
    data_leak_risks = []
    
    if not ANDROGUARD_AVAILABLE:
        return DataFlowAnalysisResult(
            total_sources=0,
            total_sinks=0,
            total_flows=0,
            critical_flows=0,
            high_risk_flows=0,
            sources_found=[],
            sinks_found=[],
            data_flow_paths=[],
            privacy_violations=[],
            data_leak_risks=[],
            summary="Androguard not available for data flow analysis",
            recommendations=["Install androguard for data flow analysis"],
        )
    
    try:
        from androguard.core.dex import DEX
        
        # Track method calls and their contexts
        method_calls = {}  # method -> list of (callee_class, callee_method)
        source_locations = []  # List of (class, method, source_type, details)
        sink_locations = []  # List of (class, method, sink_type, details)
        
        with zipfile.ZipFile(file_path, 'r') as zf:
            dex_files = [n for n in zf.namelist() if n.endswith('.dex')]
            
            for dex_name in dex_files:
                try:
                    dex_data = zf.read(dex_name)
                    dex = DEX(dex_data)
                    
                    for cls in dex.get_classes():
                        class_name = cls.get_name()
                        readable_class = class_name.replace('/', '.').strip('L;')
                        
                        # Focus on app code if package name provided
                        if package_name:
                            # Include app package and common library patterns
                            if not any(pattern in readable_class for pattern in [
                                package_name.replace('.', '/'),
                                package_name,
                                # Include common libraries that might be interesting
                                'retrofit', 'okhttp', 'volley', 'apache.http',
                            ]):
                                # Skip Android framework and other libraries
                                if readable_class.startswith(('android.', 'androidx.', 'java.', 'kotlin.', 'com.google.')):
                                    continue
                        
                        for method in cls.get_methods():
                            method_name = method.get_name()
                            method_desc = method.get_descriptor()
                            full_method = f"{readable_class}.{method_name}"
                            
                            code = method.get_code()
                            if not code:
                                continue
                            
                            # Get bytecode instructions
                            bc = code.get_bc() if hasattr(code, 'get_bc') else None
                            if not bc:
                                continue
                            
                            method_sources = []
                            method_sinks = []
                            instruction_context = []
                            
                            for ins in bc.get_instructions():
                                ins_name = ins.get_name()
                                ins_output = ins.get_output() if hasattr(ins, 'get_output') else ""
                                instruction_context.append(f"{ins_name} {ins_output}")
                                
                                # Check for invoke instructions (method calls)
                                if 'invoke' in ins_name.lower():
                                    # Parse the called method
                                    called_info = _parse_invoke_instruction(ins_output)
                                    if called_info:
                                        callee_class, callee_method = called_info
                                        
                                        # Check if it's a source
                                        source_match = _match_source(callee_class, callee_method)
                                        if source_match:
                                            source_type, description, sensitivity = source_match
                                            method_sources.append({
                                                "type": source_type,
                                                "class": callee_class,
                                                "method": callee_method,
                                                "description": description,
                                                "sensitivity": sensitivity,
                                                "instruction": f"{ins_name} {ins_output}",
                                            })
                                            source_locations.append((readable_class, method_name, source_type, source_match))
                                        
                                        # Check if it's a sink
                                        sink_match = _match_sink(callee_class, callee_method)
                                        if sink_match:
                                            sink_type, description, risk_level = sink_match
                                            method_sinks.append({
                                                "type": sink_type,
                                                "class": callee_class,
                                                "method": callee_method,
                                                "description": description,
                                                "risk_level": risk_level,
                                                "instruction": f"{ins_name} {ins_output}",
                                            })
                                            sink_locations.append((readable_class, method_name, sink_type, sink_match))
                            
                            # Record sources found
                            for src in method_sources:
                                sources_found.append({
                                    "source_type": src["type"],
                                    "class_name": readable_class,
                                    "method_name": method_name,
                                    "api_class": src["class"],
                                    "api_method": src["method"],
                                    "description": src["description"],
                                    "sensitivity": src["sensitivity"],
                                })
                            
                            # Record sinks found
                            for snk in method_sinks:
                                sinks_found.append({
                                    "sink_type": snk["type"],
                                    "class_name": readable_class,
                                    "method_name": method_name,
                                    "api_class": snk["class"],
                                    "api_method": snk["method"],
                                    "description": snk["description"],
                                    "risk_level": snk["risk_level"],
                                })
                            
                            # Detect potential data flows (source and sink in same method)
                            # This is a simplified intra-procedural analysis
                            if method_sources and method_sinks:
                                for src in method_sources:
                                    for snk in method_sinks:
                                        # Calculate severity based on source sensitivity and sink risk
                                        severity = _calculate_flow_severity(src["sensitivity"], snk["risk_level"])
                                        
                                        flow_path = DataFlowPath(
                                            source=TaintSource(
                                                source_type=src["type"],
                                                class_name=src["class"],
                                                method_name=src["method"],
                                                description=src["description"],
                                                sensitivity=src["sensitivity"],
                                                owasp_category=_get_owasp_category(src["type"], "source"),
                                            ),
                                            sink=TaintSink(
                                                sink_type=snk["type"],
                                                class_name=snk["class"],
                                                method_name=snk["method"],
                                                description=snk["description"],
                                                risk_level=snk["risk_level"],
                                                owasp_category=_get_owasp_category(snk["type"], "sink"),
                                            ),
                                            intermediate_methods=[],
                                            affected_class=readable_class,
                                            affected_method=method_name,
                                            tainted_variable=None,
                                            code_snippet='\n'.join(instruction_context[-10:]),
                                            severity=severity,
                                            description=f"{src['description']} may flow to {snk['description']}",
                                            recommendation=_get_flow_recommendation(src["type"], snk["type"]),
                                        )
                                        data_flow_paths.append(flow_path)
                                        
                                        # Check for privacy violations
                                        if src["type"] in ["device_info", "location", "contacts", "sms"]:
                                            if snk["type"] in ["network", "logging", "ipc"]:
                                                privacy_violations.append({
                                                    "violation_type": "PII_LEAKAGE",
                                                    "data_category": src["type"],
                                                    "destination": snk["type"],
                                                    "class": readable_class,
                                                    "method": method_name,
                                                    "severity": severity,
                                                    "description": f"Sensitive {src['type']} data may be sent to {snk['type']}",
                                                    "gdpr_relevant": True,
                                                    "recommendation": f"Ensure proper consent before sending {src['type']} data",
                                                })
                                        
                                        # Check for data leak risks
                                        if snk["type"] in ["logging", "external_storage"]:
                                            if src["sensitivity"] in ["high", "critical"]:
                                                data_leak_risks.append({
                                                    "risk_type": "DATA_EXPOSURE",
                                                    "data_sensitivity": src["sensitivity"],
                                                    "exposure_vector": snk["type"],
                                                    "class": readable_class,
                                                    "method": method_name,
                                                    "severity": severity,
                                                    "description": f"Sensitive data ({src['type']}) may be exposed via {snk['type']}",
                                                    "recommendation": f"Remove {snk['type']} of sensitive data in production",
                                                })
                                
                except Exception as e:
                    logger.warning(f"Data flow analysis failed for {dex_name}: {e}")
        
        # Calculate statistics
        critical_flows = sum(1 for p in data_flow_paths if p.severity == "critical")
        high_risk_flows = sum(1 for p in data_flow_paths if p.severity == "high")
        
        # Generate summary
        summary = _generate_flow_summary(sources_found, sinks_found, data_flow_paths, privacy_violations)
        
        # Generate recommendations
        recommendations = _generate_flow_recommendations(data_flow_paths, privacy_violations, data_leak_risks)
        
        return DataFlowAnalysisResult(
            total_sources=len(sources_found),
            total_sinks=len(sinks_found),
            total_flows=len(data_flow_paths),
            critical_flows=critical_flows,
            high_risk_flows=high_risk_flows,
            sources_found=sources_found,
            sinks_found=sinks_found,
            data_flow_paths=data_flow_paths,
            privacy_violations=privacy_violations,
            data_leak_risks=data_leak_risks,
            summary=summary,
            recommendations=recommendations,
        )
        
    except Exception as e:
        logger.error(f"Data flow analysis failed: {e}")
        return DataFlowAnalysisResult(
            total_sources=0,
            total_sinks=0,
            total_flows=0,
            critical_flows=0,
            high_risk_flows=0,
            sources_found=[],
            sinks_found=[],
            data_flow_paths=[],
            privacy_violations=[],
            data_leak_risks=[],
            summary=f"Analysis failed: {str(e)}",
            recommendations=[],
        )


def _parse_invoke_instruction(output: str) -> Optional[tuple]:
    """Parse invoke instruction to extract called class and method."""
    try:
        # Format: Lcom/example/Class;->methodName(params)ReturnType
        match = re.search(r'(L[^;]+;)->([^(]+)', output)
        if match:
            return match.group(1), match.group(2)
    except:
        pass
    return None


def _match_source(callee_class: str, callee_method: str) -> Optional[tuple]:
    """Match a method call against known sources."""
    for source_type, sources in ANDROID_SOURCES.items():
        for src_class, src_method, description, sensitivity in sources:
            if src_class in callee_class and src_method == callee_method:
                return source_type, description, sensitivity
    return None


def _match_sink(callee_class: str, callee_method: str) -> Optional[tuple]:
    """Match a method call against known sinks."""
    for sink_type, sinks in ANDROID_SINKS.items():
        for snk_class, snk_method, description, risk_level in sinks:
            if snk_class in callee_class and snk_method == callee_method:
                return sink_type, description, risk_level
    return None


def _calculate_flow_severity(source_sensitivity: str, sink_risk: str) -> str:
    """Calculate overall severity based on source sensitivity and sink risk."""
    severity_matrix = {
        ("critical", "critical"): "critical",
        ("critical", "high"): "critical",
        ("critical", "medium"): "high",
        ("critical", "low"): "high",
        ("high", "critical"): "critical",
        ("high", "high"): "high",
        ("high", "medium"): "high",
        ("high", "low"): "medium",
        ("medium", "critical"): "high",
        ("medium", "high"): "medium",
        ("medium", "medium"): "medium",
        ("medium", "low"): "low",
        ("low", "critical"): "medium",
        ("low", "high"): "medium",
        ("low", "medium"): "low",
        ("low", "low"): "low",
    }
    return severity_matrix.get((source_sensitivity, sink_risk), "medium")


def _get_owasp_category(data_type: str, flow_type: str) -> str:
    """Map data type to OWASP Mobile Top 10 category."""
    owasp_map = {
        # Sources
        "location": "M1: Improper Platform Usage",
        "device_info": "M1: Improper Platform Usage",
        "contacts": "M2: Insecure Data Storage",
        "sms": "M1: Improper Platform Usage",
        "user_input": "M3: Insecure Communication",
        "crypto": "M5: Insufficient Cryptography",
        "accounts": "M4: Insecure Authentication",
        # Sinks
        "network": "M3: Insecure Communication",
        "storage": "M2: Insecure Data Storage",
        "logging": "M2: Insecure Data Storage",
        "ipc": "M1: Improper Platform Usage",
        "external_storage": "M2: Insecure Data Storage",
        "process_execution": "M7: Client Code Quality",
        "reflection": "M8: Code Tampering",
    }
    return owasp_map.get(data_type, "M10: Extraneous Functionality")


def _get_flow_recommendation(source_type: str, sink_type: str) -> str:
    """Get recommendation for a specific source-to-sink flow."""
    recommendations = {
        ("location", "network"): "Ensure location data is transmitted over HTTPS with certificate pinning. Implement user consent before sharing location.",
        ("location", "logging"): "Remove location logging in production builds. Use debug-only logging guards.",
        ("device_info", "network"): "Do not send device identifiers without user consent. Consider using resettable advertising IDs instead.",
        ("device_info", "storage"): "Encrypt device identifiers before storing. Consider using Android Keystore for secure storage.",
        ("contacts", "network"): "Implement explicit user consent before syncing contacts. Use secure transmission (TLS 1.3).",
        ("user_input", "network"): "Validate and sanitize user input before transmission. Use HTTPS with certificate pinning.",
        ("user_input", "storage"): "Encrypt sensitive user data before storage. Use Android Keystore for keys.",
        ("crypto", "logging"): "Never log cryptographic keys or sensitive data. Remove all crypto logging.",
        ("accounts", "network"): "Use secure authentication protocols. Never transmit passwords in plain text.",
        ("sms", "network"): "Do not transmit SMS content without explicit user consent. Ensure compliance with privacy laws.",
    }
    return recommendations.get((source_type, sink_type), 
        f"Review data flow from {source_type} to {sink_type}. Ensure proper validation, encryption, and user consent.")


def _generate_flow_summary(sources, sinks, flows, violations) -> str:
    """Generate a human-readable summary of data flow analysis."""
    source_types = list(set(s["source_type"] for s in sources))
    sink_types = list(set(s["sink_type"] for s in sinks))
    
    summary_parts = []
    
    if flows:
        critical = sum(1 for f in flows if f.severity == "critical")
        high = sum(1 for f in flows if f.severity == "high")
        
        summary_parts.append(f"Found {len(flows)} potential data flows")
        if critical:
            summary_parts.append(f"including {critical} CRITICAL risk flows")
        if high:
            summary_parts.append(f"and {high} HIGH risk flows")
    else:
        summary_parts.append("No direct data flows detected between sources and sinks")
    
    if violations:
        summary_parts.append(f". Identified {len(violations)} potential privacy violations")
    
    if source_types:
        summary_parts.append(f". Sensitive data sources: {', '.join(source_types[:5])}")
    
    if sink_types:
        summary_parts.append(f". Data destinations: {', '.join(sink_types[:5])}")
    
    return ''.join(summary_parts) + "."


def _generate_flow_recommendations(flows, violations, leaks) -> List[str]:
    """Generate prioritized recommendations based on analysis."""
    recommendations = []
    seen = set()
    
    # Critical flow recommendations
    critical_flows = [f for f in flows if f.severity == "critical"]
    for flow in critical_flows[:5]:
        rec = flow.recommendation
        if rec not in seen:
            recommendations.append(f"🔴 CRITICAL: {rec}")
            seen.add(rec)
    
    # Privacy violation recommendations
    for violation in violations[:5]:
        rec = violation.get("recommendation", "")
        if rec and rec not in seen:
            recommendations.append(f"🟠 PRIVACY: {rec}")
            seen.add(rec)
    
    # Data leak recommendations
    for leak in leaks[:5]:
        rec = leak.get("recommendation", "")
        if rec and rec not in seen:
            recommendations.append(f"🟡 DATA LEAK: {rec}")
            seen.add(rec)
    
    # General recommendations
    if any(f.sink.sink_type == "logging" for f in flows):
        if "Remove sensitive data from logs" not in seen:
            recommendations.append("📋 Remove all sensitive data from logs before production release")
    
    if any(f.source.source_type == "device_info" for f in flows):
        if "device identifiers" not in ' '.join(seen).lower():
            recommendations.append("📱 Review device identifier usage for privacy compliance (GDPR, CCPA)")
    
    if any(f.sink.sink_type == "network" for f in flows):
        recommendations.append("🔒 Implement certificate pinning for all network communications")
        recommendations.append("🔐 Use TLS 1.3 for all data transmissions")
    
    return recommendations[:15]  # Limit to top 15


def dataflow_result_to_dict(result: DataFlowAnalysisResult) -> dict:
    """Convert DataFlowAnalysisResult to dictionary for JSON serialization."""
    return {
        "total_sources": result.total_sources,
        "total_sinks": result.total_sinks,
        "total_flows": result.total_flows,
        "critical_flows": result.critical_flows,
        "high_risk_flows": result.high_risk_flows,
        "sources_found": result.sources_found,
        "sinks_found": result.sinks_found,
        "data_flow_paths": [
            {
                "source": {
                    "type": p.source.source_type,
                    "class": p.source.class_name,
                    "method": p.source.method_name,
                    "description": p.source.description,
                    "sensitivity": p.source.sensitivity,
                },
                "sink": {
                    "type": p.sink.sink_type,
                    "class": p.sink.class_name,
                    "method": p.sink.method_name,
                    "description": p.sink.description,
                    "risk_level": p.sink.risk_level,
                },
                "affected_class": p.affected_class,
                "affected_method": p.affected_method,
                "severity": p.severity,
                "description": p.description,
                "recommendation": p.recommendation,
                "code_snippet": p.code_snippet,
            }
            for p in result.data_flow_paths
        ],
        "privacy_violations": result.privacy_violations,
        "data_leak_risks": result.data_leak_risks,
        "summary": result.summary,
        "recommendations": result.recommendations,
    }


# ============================================================================
# Smali View Functions
# ============================================================================

def get_smali_for_class(output_dir: Path, class_path: str) -> Optional[Dict[str, Any]]:
    """
    Get Smali bytecode for a specific class.
    
    Args:
        output_dir: JADX output directory
        class_path: Path to Java class file (e.g., "com/example/MainActivity.java")
    
    Returns:
        Dictionary with smali_code, class_info, and bytecode_analysis
    """
    # Convert Java path to Smali path
    # Java: sources/com/example/MainActivity.java
    # Smali would be in: sources/com/example/MainActivity.smali (if we decompiled with --show-bad-code)
    
    # First try to find smali from a separate baksmali run
    smali_dir = output_dir / "smali"
    
    # Convert class path to package path
    class_path_normalized = class_path.replace('.java', '').replace('/', '.')
    parts = class_path_normalized.split('.')
    class_name = parts[-1] if parts else ""
    package_path = '/'.join(parts[:-1]) if len(parts) > 1 else ""
    
    smali_file = None
    
    # Try multiple locations
    possible_paths = [
        smali_dir / f"{package_path}/{class_name}.smali",
        smali_dir / "classes" / f"{package_path}/{class_name}.smali",
        output_dir / "sources" / class_path.replace('.java', '.smali'),
    ]
    
    for path in possible_paths:
        if path.exists():
            smali_file = path
            break
    
    # If no pre-existing smali, try to generate from DEX
    if not smali_file or not smali_file.exists():
        # We need to run baksmali on the original APK/DEX
        # Check if we have the original APK path stored
        apk_path = output_dir / ".apk_path"
        
        if apk_path.exists():
            original_apk = Path(apk_path.read_text().strip())
            if original_apk.exists():
                smali_code = _extract_smali_from_apk(original_apk, class_path_normalized, output_dir)
                if smali_code:
                    return {
                        "class_path": class_path,
                        "smali_code": smali_code,
                        "bytecode_stats": _analyze_smali_bytecode(smali_code),
                        "registers_used": _count_registers(smali_code),
                        "method_count": smali_code.count(".method"),
                        "field_count": smali_code.count(".field"),
                        "instructions": _extract_smali_instructions(smali_code),
                    }
        
        # If no APK, try to generate pseudo-smali from Java
        java_source = get_jadx_class_source(output_dir, class_path)
        if java_source:
            return {
                "class_path": class_path,
                "smali_code": _generate_pseudo_smali(java_source, class_name),
                "bytecode_stats": {"note": "Pseudo-Smali generated from Java source"},
                "registers_used": 0,
                "method_count": java_source.count(" void ") + java_source.count(" int ") + java_source.count(" String "),
                "field_count": 0,
                "instructions": [],
                "is_pseudo": True,
            }
    
    if smali_file and smali_file.exists():
        smali_code = smali_file.read_text(encoding='utf-8', errors='ignore')
        return {
            "class_path": class_path,
            "smali_code": smali_code,
            "bytecode_stats": _analyze_smali_bytecode(smali_code),
            "registers_used": _count_registers(smali_code),
            "method_count": smali_code.count(".method"),
            "field_count": smali_code.count(".field"),
            "instructions": _extract_smali_instructions(smali_code),
        }
    
    return None


def _extract_smali_from_apk(apk_path: Path, class_name: str, output_dir: Path) -> Optional[str]:
    """Extract Smali for a specific class using baksmali."""
    import subprocess
    import zipfile
    
    try:
        # Create temp directory for smali output
        smali_out = output_dir / "smali"
        smali_out.mkdir(exist_ok=True)
        
        # Try to find dex file in APK
        with zipfile.ZipFile(apk_path, 'r') as apk:
            dex_files = [f for f in apk.namelist() if f.endswith('.dex')]
            
            if not dex_files:
                return None
            
            # Extract classes.dex
            for dex_name in dex_files:
                dex_path = output_dir / dex_name
                apk.extract(dex_name, output_dir)
                
                # Run baksmali
                try:
                    result = subprocess.run(
                        ["baksmali", "d", str(dex_path), "-o", str(smali_out)],
                        capture_output=True,
                        text=True,
                        timeout=60
                    )
                except FileNotFoundError:
                    # baksmali not installed, try dex2jar approach
                    logger.warning("baksmali not found, using fallback")
                    break
        
        # Find the smali file for our class
        class_path = class_name.replace('.', '/') + '.smali'
        smali_file = smali_out / class_path
        
        if smali_file.exists():
            return smali_file.read_text(encoding='utf-8', errors='ignore')
        
        # Try with inner classes
        for f in smali_out.rglob(f"*{class_name.split('.')[-1]}*.smali"):
            return f.read_text(encoding='utf-8', errors='ignore')
            
    except Exception as e:
        logger.warning(f"Failed to extract smali: {e}")
    
    return None


def _generate_pseudo_smali(java_source: str, class_name: str) -> str:
    """Generate pseudo-Smali representation from Java source for visualization."""
    lines = []
    lines.append(f".class public L{class_name};")
    lines.append(".super Ljava/lang/Object;")
    lines.append("")
    
    # Extract fields
    field_pattern = r'(?:public|private|protected)\s+(?:static\s+)?(?:final\s+)?([\w<>\[\]]+)\s+(\w+)\s*[;=]'
    for match in re.finditer(field_pattern, java_source):
        java_type, field_name = match.groups()
        smali_type = _java_type_to_smali(java_type)
        lines.append(f".field private {field_name}:{smali_type}")
    
    lines.append("")
    
    # Extract methods
    method_pattern = r'(?:public|private|protected)\s+(?:static\s+)?(?:final\s+)?(?:synchronized\s+)?([\w<>\[\]]+)\s+(\w+)\s*\(([^)]*)\)'
    for match in re.finditer(method_pattern, java_source):
        return_type, method_name, params = match.groups()
        smali_return = _java_type_to_smali(return_type)
        smali_params = _parse_java_params_to_smali(params)
        
        lines.append(f".method public {method_name}({smali_params}){smali_return}")
        lines.append("    .registers 5")
        lines.append("")
        lines.append("    # Pseudo-bytecode (actual bytecode requires APK)")
        lines.append("    return-void")
        lines.append(".end method")
        lines.append("")
    
    return '\n'.join(lines)


def _java_type_to_smali(java_type: str) -> str:
    """Convert Java type to Smali type descriptor."""
    type_map = {
        'void': 'V',
        'boolean': 'Z',
        'byte': 'B',
        'char': 'C',
        'short': 'S',
        'int': 'I',
        'long': 'J',
        'float': 'F',
        'double': 'D',
        'String': 'Ljava/lang/String;',
        'Object': 'Ljava/lang/Object;',
    }
    
    if java_type in type_map:
        return type_map[java_type]
    
    if java_type.endswith('[]'):
        return '[' + _java_type_to_smali(java_type[:-2])
    
    # Assume it's a class
    return f'L{java_type.replace(".", "/")};'


def _parse_java_params_to_smali(params: str) -> str:
    """Convert Java method parameters to Smali format."""
    if not params.strip():
        return ""
    
    result = []
    for param in params.split(','):
        param = param.strip()
        if not param:
            continue
        parts = param.split()
        if parts:
            java_type = parts[0]
            result.append(_java_type_to_smali(java_type))
    
    return ''.join(result)


def _analyze_smali_bytecode(smali_code: str) -> Dict[str, Any]:
    """Analyze Smali bytecode for interesting patterns."""
    stats = {
        "invocations": {
            "virtual": smali_code.count("invoke-virtual"),
            "static": smali_code.count("invoke-static"),
            "direct": smali_code.count("invoke-direct"),
            "interface": smali_code.count("invoke-interface"),
            "super": smali_code.count("invoke-super"),
        },
        "field_ops": {
            "iget": smali_code.count("iget"),
            "iput": smali_code.count("iput"),
            "sget": smali_code.count("sget"),
            "sput": smali_code.count("sput"),
        },
        "control_flow": {
            "if_statements": len(re.findall(r'if-\w+', smali_code)),
            "goto": smali_code.count("goto"),
            "switch": smali_code.count("packed-switch") + smali_code.count("sparse-switch"),
            "try_catch": smali_code.count(".catch"),
        },
        "suspicious_ops": {
            "reflection": smali_code.count("Ljava/lang/reflect/"),
            "runtime_exec": smali_code.count("Ljava/lang/Runtime;->exec"),
            "class_loader": smali_code.count("ClassLoader"),
            "dex_load": smali_code.count("DexClassLoader") + smali_code.count("PathClassLoader"),
            "native_calls": smali_code.count(".native"),
            "crypto": smali_code.count("Ljavax/crypto/"),
        },
    }
    return stats


def _count_registers(smali_code: str) -> int:
    """Count total registers used in Smali code."""
    matches = re.findall(r'\.registers\s+(\d+)', smali_code)
    return sum(int(m) for m in matches)


def _extract_smali_instructions(smali_code: str) -> List[Dict[str, Any]]:
    """Extract notable Smali instructions for display."""
    instructions = []
    
    # Find method boundaries
    method_pattern = r'\.method\s+([^\n]+)\n(.*?)\.end method'
    for match in re.finditer(method_pattern, smali_code, re.DOTALL):
        method_sig = match.group(1)
        method_body = match.group(2)
        
        # Extract interesting instructions
        for line_num, line in enumerate(method_body.split('\n')):
            line = line.strip()
            if not line or line.startswith('#') or line.startswith('.'):
                continue
            
            # Categorize instruction
            category = None
            if line.startswith('invoke-'):
                category = "invocation"
            elif line.startswith('const-string'):
                category = "string"
            elif line.startswith(('iget', 'iput', 'sget', 'sput')):
                category = "field"
            elif line.startswith('if-'):
                category = "branch"
            elif line.startswith('new-'):
                category = "allocation"
            
            if category:
                instructions.append({
                    "method": method_sig.split('(')[0].strip(),
                    "instruction": line[:100],
                    "category": category,
                })
        
        if len(instructions) > 100:  # Limit
            break
    
    return instructions[:100]


# ============================================================================
# String Extraction Functions
# ============================================================================

# String categories and patterns for classification
STRING_PATTERNS = {
    "url": (re.compile(r'https?://[^\s"\'<>]+', re.IGNORECASE), "high"),
    "ip_address": (re.compile(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b'), "medium"),
    "email": (re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+'), "medium"),
    "api_key": (re.compile(r'(?:api[_-]?key|apikey|api_secret|api_token)["\']?\s*[:=]\s*["\']?[\w-]{20,}', re.IGNORECASE), "critical"),
    "aws_key": (re.compile(r'AKIA[0-9A-Z]{16}'), "critical"),
    "private_key": (re.compile(r'-----BEGIN\s+(?:RSA\s+)?PRIVATE\s+KEY-----'), "critical"),
    "jwt": (re.compile(r'eyJ[A-Za-z0-9_-]+\.eyJ[A-Za-z0-9_-]+\.[A-Za-z0-9_-]+'), "high"),
    "base64_secret": (re.compile(r'[A-Za-z0-9+/]{40,}={0,2}'), "low"),
    "password_field": (re.compile(r'(?:password|passwd|pwd|secret)["\']?\s*[:=]\s*["\'][^"\']+["\']', re.IGNORECASE), "critical"),
    "firebase": (re.compile(r'[a-z0-9-]+\.firebaseio\.com', re.IGNORECASE), "high"),
    "firebase_key": (re.compile(r'AIza[0-9A-Za-z_-]{35}'), "critical"),
    "sql_query": (re.compile(r'(?:SELECT|INSERT|UPDATE|DELETE)\s+.{10,}(?:FROM|INTO|SET)\s+', re.IGNORECASE), "medium"),
    "file_path": (re.compile(r'/(?:data|sdcard|storage|system)/[\w/.-]+'), "low"),
    "package_name": (re.compile(r'com\.[a-z][a-z0-9_]*(?:\.[a-z][a-z0-9_]*)+', re.IGNORECASE), "low"),
    "content_uri": (re.compile(r'content://[\w./]+'), "medium"),
    "intent_action": (re.compile(r'android\.intent\.action\.\w+'), "low"),
    "permission": (re.compile(r'android\.permission\.\w+'), "medium"),
    "phone_number": (re.compile(r'(?:\+\d{1,3})?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'), "low"),
    "crypto_algo": (re.compile(r'(?:AES|DES|RSA|MD5|SHA\d*|HMAC)', re.IGNORECASE), "medium"),
    "hardcoded_iv": (re.compile(r'(?:iv|IV|nonce)["\']?\s*[:=]\s*["\'][A-Za-z0-9+/=]{8,}["\']'), "high"),
}


def extract_all_strings(output_dir: Path, filters: Optional[List[str]] = None) -> Dict[str, Any]:
    """
    Extract all strings from decompiled sources with classification.
    
    Args:
        output_dir: JADX output directory
        filters: Optional list of categories to filter (url, api_key, etc.)
    
    Returns:
        Dictionary with categorized strings and statistics
    """
    sources_dir = output_dir / "sources"
    if not sources_dir.exists():
        return {"error": "Decompiled sources not found", "strings": [], "stats": {}}
    
    all_strings = []
    stats = {cat: 0 for cat in STRING_PATTERNS.keys()}
    stats["uncategorized"] = 0
    files_scanned = 0
    
    # Also check resources
    resources_dir = output_dir / "resources"
    
    for java_file in sources_dir.rglob("*.java"):
        try:
            content = java_file.read_text(encoding='utf-8', errors='ignore')
            rel_path = str(java_file.relative_to(sources_dir))
            files_scanned += 1
            
            # Extract string literals
            string_literals = re.findall(r'"([^"\\]*(\\.[^"\\]*)*)"', content)
            
            for match in string_literals:
                string_val = match[0] if isinstance(match, tuple) else match
                if len(string_val) < 3:  # Skip very short strings
                    continue
                
                # Classify string
                categories = []
                severity = "low"
                
                for cat_name, (pattern, cat_severity) in STRING_PATTERNS.items():
                    if pattern.search(string_val):
                        categories.append(cat_name)
                        stats[cat_name] += 1
                        if cat_severity == "critical":
                            severity = "critical"
                        elif cat_severity == "high" and severity not in ["critical"]:
                            severity = "high"
                        elif cat_severity == "medium" and severity not in ["critical", "high"]:
                            severity = "medium"
                
                if not categories:
                    categories = ["uncategorized"]
                    stats["uncategorized"] += 1
                
                # Apply filters
                if filters:
                    if not any(cat in filters for cat in categories):
                        continue
                
                # Find line number
                line_num = 1
                for i, line in enumerate(content.split('\n')):
                    if string_val in line:
                        line_num = i + 1
                        break
                
                all_strings.append({
                    "value": string_val[:500],  # Truncate long strings
                    "file": rel_path,
                    "line": line_num,
                    "categories": categories,
                    "severity": severity,
                    "length": len(string_val),
                })
        except Exception as e:
            logger.warning(f"Failed to extract strings from {java_file}: {e}")
    
    # Also extract from resources (strings.xml, etc.)
    if resources_dir.exists():
        for xml_file in resources_dir.rglob("*.xml"):
            try:
                content = xml_file.read_text(encoding='utf-8', errors='ignore')
                rel_path = str(xml_file.relative_to(resources_dir))
                
                # Find string values
                for match in re.finditer(r'>([^<]+)</', content):
                    string_val = match.group(1).strip()
                    if len(string_val) < 3:
                        continue
                    
                    categories = []
                    severity = "low"
                    
                    for cat_name, (pattern, cat_severity) in STRING_PATTERNS.items():
                        if pattern.search(string_val):
                            categories.append(cat_name)
                            stats[cat_name] += 1
                            if cat_severity == "critical":
                                severity = "critical"
                            elif cat_severity == "high" and severity != "critical":
                                severity = "high"
                    
                    if not categories:
                        continue  # Skip uncategorized XML strings
                    
                    if filters and not any(cat in filters for cat in categories):
                        continue
                    
                    all_strings.append({
                        "value": string_val[:500],
                        "file": f"resources/{rel_path}",
                        "line": 0,
                        "categories": categories,
                        "severity": severity,
                        "length": len(string_val),
                        "is_resource": True,
                    })
            except Exception as e:
                pass
    
    # Sort by severity
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    all_strings.sort(key=lambda x: (severity_order.get(x["severity"], 4), -x["length"]))
    
    return {
        "total_strings": len(all_strings),
        "files_scanned": files_scanned,
        "strings": all_strings[:1000],  # Limit to 1000
        "stats": stats,
        "severity_counts": {
            "critical": sum(1 for s in all_strings if s["severity"] == "critical"),
            "high": sum(1 for s in all_strings if s["severity"] == "high"),
            "medium": sum(1 for s in all_strings if s["severity"] == "medium"),
            "low": sum(1 for s in all_strings if s["severity"] == "low"),
        },
        "top_categories": sorted(
            [(k, v) for k, v in stats.items() if v > 0],
            key=lambda x: -x[1]
        )[:10],
    }


# ============================================================================
# Cross-Reference (XREF) Functions
# ============================================================================

def build_cross_references(output_dir: Path, class_path: str) -> Dict[str, Any]:
    """
    Build cross-references for a specific class.
    
    Returns:
    - Methods defined in this class
    - What calls each method (callers/incoming refs)
    - What each method calls (callees/outgoing refs)
    - Field references
    """
    sources_dir = output_dir / "sources"
    if not sources_dir.exists():
        return {"error": "Decompiled sources not found"}
    
    # Get the target class source
    target_source = get_jadx_class_source(output_dir, class_path)
    if not target_source:
        return {"error": f"Class not found: {class_path}"}
    
    # Parse target class info
    target_class_info = _parse_java_class(target_source, class_path)
    target_class_name = target_class_info.class_name
    target_package = target_class_info.package_name
    target_fqn = f"{target_package}.{target_class_name}" if target_package else target_class_name
    
    # Extract methods from target class with signatures
    method_pattern = r'(?:public|private|protected)?\s*(?:static\s+)?(?:final\s+)?(?:synchronized\s+)?([\w<>\[\],\s]+)\s+(\w+)\s*\(([^)]*)\)'
    target_methods = []
    for match in re.finditer(method_pattern, target_source):
        return_type, method_name, params = match.groups()
        target_methods.append({
            "name": method_name,
            "return_type": return_type.strip(),
            "params": params.strip(),
            "signature": f"{method_name}({params})",
            "callers": [],
            "callees": [],
            "line": _find_line_number(target_source, match.group(0)),
        })
    
    # Extract fields from target class
    field_pattern = r'(?:public|private|protected)\s+(?:static\s+)?(?:final\s+)?([\w<>\[\],\s]+)\s+(\w+)\s*[;=]'
    target_fields = []
    for match in re.finditer(field_pattern, target_source):
        field_type, field_name = match.groups()
        target_fields.append({
            "name": field_name,
            "type": field_type.strip(),
            "readers": [],
            "writers": [],
            "line": _find_line_number(target_source, match.group(0)),
        })
    
    # Find what target methods call (outgoing references)
    for method in target_methods:
        # Extract method body
        method_body = _extract_method_body(target_source, method["name"])
        if method_body:
            # Find method calls in body
            call_pattern = r'(?:(\w+)\.)?(\w+)\s*\('
            for call_match in re.finditer(call_pattern, method_body):
                obj_name, called_method = call_match.groups()
                if called_method not in ['if', 'for', 'while', 'switch', 'catch', 'synchronized']:
                    method["callees"].append({
                        "method": called_method,
                        "object": obj_name or "this",
                        "line": _find_line_number(method_body, call_match.group(0)),
                    })
    
    # Scan all other classes for references to this class
    callers_map = {m["name"]: [] for m in target_methods}
    field_readers_map = {f["name"]: [] for f in target_fields}
    field_writers_map = {f["name"]: [] for f in target_fields}
    
    for java_file in sources_dir.rglob("*.java"):
        if str(java_file).endswith(class_path):
            continue  # Skip self
        
        try:
            source = java_file.read_text(encoding='utf-8', errors='ignore')
            rel_path = str(java_file.relative_to(sources_dir))
            other_class = Path(rel_path).stem
            
            # Check for imports of target class
            imports_target = target_fqn in source or f"import {target_fqn}" in source
            uses_simple_name = target_class_name in source
            
            if not (imports_target or uses_simple_name):
                continue
            
            # Find method calls to target class
            for method in target_methods:
                # Look for TargetClass.method() or instance.method() patterns
                patterns = [
                    f'{target_class_name}.{method["name"]}\\s*\\(',
                    f'\\.{method["name"]}\\s*\\(',  # Any call to this method name
                ]
                
                for pattern in patterns:
                    for match in re.finditer(pattern, source):
                        line_num = _find_line_number(source, match.group(0))
                        caller_method = _find_enclosing_method(source, match.start())
                        
                        callers_map[method["name"]].append({
                            "class": other_class,
                            "file": rel_path,
                            "method": caller_method,
                            "line": line_num,
                        })
            
            # Find field accesses
            for field in target_fields:
                # Reader pattern: target.field or TargetClass.field
                reader_pattern = f'(?:{target_class_name}|\\w+)\\.{field["name"]}(?!\\s*=)'
                for match in re.finditer(reader_pattern, source):
                    line_num = _find_line_number(source, match.group(0))
                    field_readers_map[field["name"]].append({
                        "class": other_class,
                        "file": rel_path,
                        "line": line_num,
                    })
                
                # Writer pattern: target.field = 
                writer_pattern = f'(?:{target_class_name}|\\w+)\\.{field["name"]}\\s*='
                for match in re.finditer(writer_pattern, source):
                    line_num = _find_line_number(source, match.group(0))
                    field_writers_map[field["name"]].append({
                        "class": other_class,
                        "file": rel_path,
                        "line": line_num,
                    })
                    
        except Exception as e:
            logger.debug(f"Error scanning {java_file}: {e}")
    
    # Update methods with callers
    for method in target_methods:
        method["callers"] = callers_map.get(method["name"], [])[:50]  # Limit
        method["caller_count"] = len(callers_map.get(method["name"], []))
        method["callee_count"] = len(method["callees"])
    
    # Update fields with references
    for field in target_fields:
        field["readers"] = field_readers_map.get(field["name"], [])[:30]
        field["writers"] = field_writers_map.get(field["name"], [])[:30]
        field["read_count"] = len(field_readers_map.get(field["name"], []))
        field["write_count"] = len(field_writers_map.get(field["name"], []))
    
    # Calculate statistics
    total_incoming = sum(m["caller_count"] for m in target_methods)
    total_outgoing = sum(m["callee_count"] for m in target_methods)
    
    return {
        "class_name": target_class_name,
        "package": target_package,
        "file_path": class_path,
        "methods": target_methods,
        "fields": target_fields,
        "statistics": {
            "method_count": len(target_methods),
            "field_count": len(target_fields),
            "total_incoming_refs": total_incoming,
            "total_outgoing_refs": total_outgoing,
            "is_heavily_used": total_incoming > 10,
            "is_hub_class": total_outgoing > 20,
        },
        "summary": _generate_xref_summary(target_class_name, target_methods, target_fields, total_incoming, total_outgoing),
    }


def _find_line_number(source: str, target: str) -> int:
    """Find the line number of a target string in source."""
    pos = source.find(target)
    if pos == -1:
        return 0
    return source[:pos].count('\n') + 1


def _extract_method_body(source: str, method_name: str) -> Optional[str]:
    """Extract the body of a method from source code."""
    # Find method start
    pattern = rf'(?:public|private|protected)?\s*(?:static\s+)?[\w<>\[\],\s]+\s+{method_name}\s*\([^)]*\)\s*(?:throws\s+[\w,\s]+)?\s*\{{'
    match = re.search(pattern, source)
    if not match:
        return None
    
    start = match.end()
    brace_count = 1
    end = start
    
    while end < len(source) and brace_count > 0:
        if source[end] == '{':
            brace_count += 1
        elif source[end] == '}':
            brace_count -= 1
        end += 1
    
    return source[start:end-1]


def _find_enclosing_method(source: str, position: int) -> str:
    """Find the method name that contains a given position."""
    # Look backwards for method declaration
    search_area = source[:position]
    method_pattern = r'(?:public|private|protected)?\s*(?:static\s+)?[\w<>\[\],\s]+\s+(\w+)\s*\([^)]*\)\s*(?:throws\s+[\w,\s]+)?\s*\{'
    
    matches = list(re.finditer(method_pattern, search_area))
    if matches:
        return matches[-1].group(1)
    return "unknown"


def _generate_xref_summary(class_name: str, methods: List, fields: List, incoming: int, outgoing: int) -> str:
    """Generate a summary of cross-references."""
    parts = [f"Class {class_name}:"]
    
    # Most called methods
    most_called = sorted(methods, key=lambda m: m.get("caller_count", 0), reverse=True)[:3]
    if most_called and most_called[0].get("caller_count", 0) > 0:
        parts.append(f"Most called methods: {', '.join(m['name'] for m in most_called)}")
    
    # Hub methods (lots of outgoing calls)
    hub_methods = [m for m in methods if m.get("callee_count", 0) > 5]
    if hub_methods:
        parts.append(f"Hub methods (many calls): {', '.join(m['name'] for m in hub_methods[:3])}")
    
    # Heavily accessed fields
    hot_fields = [f for f in fields if f.get("read_count", 0) + f.get("write_count", 0) > 5]
    if hot_fields:
        parts.append(f"Frequently accessed fields: {', '.join(f['name'] for f in hot_fields[:3])}")
    
    if incoming > 10:
        parts.append(f"⚡ This is a heavily-used class ({incoming} incoming references)")
    
    if outgoing > 20:
        parts.append(f"🔗 This is a hub class ({outgoing} outgoing calls)")
    
    return " | ".join(parts)


# ============================================================================
# Feature: Download Project as ZIP
# ============================================================================

def create_project_zip(output_dir: Path) -> Path:
    """
    Create a ZIP file of the entire decompiled project.
    
    Args:
        output_dir: JADX output directory
    
    Returns:
        Path to the created ZIP file
    """
    import zipfile
    import shutil
    
    output_dir = Path(output_dir)
    if not output_dir.exists():
        raise ValueError(f"Output directory not found: {output_dir}")
    
    # Create zip filename based on directory name
    zip_name = f"{output_dir.name}_decompiled.zip"
    zip_path = output_dir.parent / zip_name
    
    # Remove existing zip if present
    if zip_path.exists():
        zip_path.unlink()
    
    # Create the ZIP file
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(output_dir):
            # Skip .git and __pycache__ directories
            dirs[:] = [d for d in dirs if d not in ['.git', '__pycache__', '.gradle']]
            
            for file in files:
                file_path = Path(root) / file
                arcname = file_path.relative_to(output_dir)
                
                # Skip very large files (>10MB)
                if file_path.stat().st_size > 10 * 1024 * 1024:
                    continue
                
                zipf.write(file_path, arcname)
    
    return zip_path


def get_project_zip_info(output_dir: Path) -> Dict[str, Any]:
    """
    Get information about what would be in the project ZIP.
    
    Args:
        output_dir: JADX output directory
    
    Returns:
        Dictionary with file counts and sizes
    """
    output_dir = Path(output_dir)
    if not output_dir.exists():
        return {"error": "Output directory not found"}
    
    total_files = 0
    total_size = 0
    file_types = {}
    
    for root, dirs, files in os.walk(output_dir):
        dirs[:] = [d for d in dirs if d not in ['.git', '__pycache__', '.gradle']]
        
        for file in files:
            file_path = Path(root) / file
            file_size = file_path.stat().st_size
            
            if file_size > 10 * 1024 * 1024:  # Skip >10MB
                continue
            
            total_files += 1
            total_size += file_size
            
            ext = file_path.suffix.lower() or '.noext'
            file_types[ext] = file_types.get(ext, 0) + 1
    
    return {
        "total_files": total_files,
        "total_size_bytes": total_size,
        "total_size_mb": round(total_size / (1024 * 1024), 2),
        "file_types": file_types,
        "estimated_zip_size_mb": round(total_size * 0.3 / (1024 * 1024), 2)  # Rough compression estimate
    }


# ============================================================================
# Feature: Permission Analyzer
# ============================================================================

# Android permission danger levels and descriptions
ANDROID_PERMISSIONS = {
    # Dangerous permissions - require explicit user approval
    "android.permission.READ_CALENDAR": {"level": "dangerous", "description": "Read calendar events", "category": "calendar"},
    "android.permission.WRITE_CALENDAR": {"level": "dangerous", "description": "Modify calendar events", "category": "calendar"},
    "android.permission.CAMERA": {"level": "dangerous", "description": "Access camera", "category": "camera"},
    "android.permission.READ_CONTACTS": {"level": "dangerous", "description": "Read contacts", "category": "contacts"},
    "android.permission.WRITE_CONTACTS": {"level": "dangerous", "description": "Modify contacts", "category": "contacts"},
    "android.permission.GET_ACCOUNTS": {"level": "dangerous", "description": "Get device accounts", "category": "contacts"},
    "android.permission.ACCESS_FINE_LOCATION": {"level": "dangerous", "description": "Precise GPS location", "category": "location"},
    "android.permission.ACCESS_COARSE_LOCATION": {"level": "dangerous", "description": "Approximate location", "category": "location"},
    "android.permission.ACCESS_BACKGROUND_LOCATION": {"level": "dangerous", "description": "Background location access", "category": "location"},
    "android.permission.RECORD_AUDIO": {"level": "dangerous", "description": "Record audio/microphone", "category": "microphone"},
    "android.permission.READ_PHONE_STATE": {"level": "dangerous", "description": "Read phone state/identity", "category": "phone"},
    "android.permission.READ_PHONE_NUMBERS": {"level": "dangerous", "description": "Read phone numbers", "category": "phone"},
    "android.permission.CALL_PHONE": {"level": "dangerous", "description": "Make phone calls", "category": "phone"},
    "android.permission.ANSWER_PHONE_CALLS": {"level": "dangerous", "description": "Answer phone calls", "category": "phone"},
    "android.permission.READ_CALL_LOG": {"level": "dangerous", "description": "Read call history", "category": "phone"},
    "android.permission.WRITE_CALL_LOG": {"level": "dangerous", "description": "Modify call history", "category": "phone"},
    "android.permission.ADD_VOICEMAIL": {"level": "dangerous", "description": "Add voicemail", "category": "phone"},
    "android.permission.USE_SIP": {"level": "dangerous", "description": "Use SIP calls", "category": "phone"},
    "android.permission.PROCESS_OUTGOING_CALLS": {"level": "dangerous", "description": "Process outgoing calls", "category": "phone"},
    "android.permission.BODY_SENSORS": {"level": "dangerous", "description": "Access body sensors", "category": "sensors"},
    "android.permission.ACTIVITY_RECOGNITION": {"level": "dangerous", "description": "Recognize physical activity", "category": "sensors"},
    "android.permission.SEND_SMS": {"level": "dangerous", "description": "Send SMS messages", "category": "sms"},
    "android.permission.RECEIVE_SMS": {"level": "dangerous", "description": "Receive SMS messages", "category": "sms"},
    "android.permission.READ_SMS": {"level": "dangerous", "description": "Read SMS messages", "category": "sms"},
    "android.permission.RECEIVE_WAP_PUSH": {"level": "dangerous", "description": "Receive WAP messages", "category": "sms"},
    "android.permission.RECEIVE_MMS": {"level": "dangerous", "description": "Receive MMS messages", "category": "sms"},
    "android.permission.READ_EXTERNAL_STORAGE": {"level": "dangerous", "description": "Read external storage", "category": "storage"},
    "android.permission.WRITE_EXTERNAL_STORAGE": {"level": "dangerous", "description": "Write to external storage", "category": "storage"},
    "android.permission.READ_MEDIA_IMAGES": {"level": "dangerous", "description": "Read images", "category": "storage"},
    "android.permission.READ_MEDIA_VIDEO": {"level": "dangerous", "description": "Read videos", "category": "storage"},
    "android.permission.READ_MEDIA_AUDIO": {"level": "dangerous", "description": "Read audio files", "category": "storage"},
    
    # Signature/System permissions - very high risk
    "android.permission.INSTALL_PACKAGES": {"level": "signature", "description": "Install apps silently", "category": "system"},
    "android.permission.DELETE_PACKAGES": {"level": "signature", "description": "Delete apps", "category": "system"},
    "android.permission.MOUNT_UNMOUNT_FILESYSTEMS": {"level": "signature", "description": "Mount filesystems", "category": "system"},
    "android.permission.CHANGE_COMPONENT_ENABLED_STATE": {"level": "signature", "description": "Enable/disable components", "category": "system"},
    "android.permission.REQUEST_INSTALL_PACKAGES": {"level": "dangerous", "description": "Request app installation", "category": "system"},
    
    # Network permissions - moderate risk
    "android.permission.INTERNET": {"level": "normal", "description": "Internet access", "category": "network"},
    "android.permission.ACCESS_NETWORK_STATE": {"level": "normal", "description": "View network state", "category": "network"},
    "android.permission.ACCESS_WIFI_STATE": {"level": "normal", "description": "View Wi-Fi state", "category": "network"},
    "android.permission.CHANGE_WIFI_STATE": {"level": "normal", "description": "Change Wi-Fi state", "category": "network"},
    "android.permission.CHANGE_NETWORK_STATE": {"level": "normal", "description": "Change network state", "category": "network"},
    "android.permission.BLUETOOTH": {"level": "normal", "description": "Bluetooth access", "category": "network"},
    "android.permission.BLUETOOTH_ADMIN": {"level": "normal", "description": "Bluetooth admin", "category": "network"},
    "android.permission.BLUETOOTH_CONNECT": {"level": "dangerous", "description": "Connect to Bluetooth devices", "category": "network"},
    "android.permission.BLUETOOTH_SCAN": {"level": "dangerous", "description": "Scan for Bluetooth devices", "category": "network"},
    "android.permission.NFC": {"level": "normal", "description": "NFC access", "category": "network"},
    
    # Other notable permissions
    "android.permission.VIBRATE": {"level": "normal", "description": "Vibrate device", "category": "hardware"},
    "android.permission.WAKE_LOCK": {"level": "normal", "description": "Keep device awake", "category": "hardware"},
    "android.permission.FLASHLIGHT": {"level": "normal", "description": "Use flashlight", "category": "hardware"},
    "android.permission.RECEIVE_BOOT_COMPLETED": {"level": "normal", "description": "Auto-start on boot", "category": "system"},
    "android.permission.FOREGROUND_SERVICE": {"level": "normal", "description": "Run foreground service", "category": "system"},
    "android.permission.SYSTEM_ALERT_WINDOW": {"level": "signature", "description": "Draw over other apps", "category": "system"},
    "android.permission.BIND_ACCESSIBILITY_SERVICE": {"level": "signature", "description": "Accessibility service (can monitor screen)", "category": "system"},
    "android.permission.BIND_DEVICE_ADMIN": {"level": "signature", "description": "Device administrator", "category": "system"},
    "android.permission.BIND_NOTIFICATION_LISTENER_SERVICE": {"level": "signature", "description": "Read all notifications", "category": "system"},
    "android.permission.PACKAGE_USAGE_STATS": {"level": "signature", "description": "Track app usage", "category": "system"},
    "android.permission.QUERY_ALL_PACKAGES": {"level": "normal", "description": "See all installed apps", "category": "system"},
    "android.permission.REQUEST_DELETE_PACKAGES": {"level": "normal", "description": "Request app deletion", "category": "system"},
    "android.permission.GET_TASKS": {"level": "deprecated", "description": "Get running tasks (deprecated)", "category": "system"},
    "android.permission.REORDER_TASKS": {"level": "normal", "description": "Reorder tasks", "category": "system"},
    "android.permission.DISABLE_KEYGUARD": {"level": "normal", "description": "Disable lock screen", "category": "system"},
    "android.permission.USE_BIOMETRIC": {"level": "normal", "description": "Use biometric authentication", "category": "hardware"},
    "android.permission.USE_FINGERPRINT": {"level": "normal", "description": "Use fingerprint sensor", "category": "hardware"},
}

# Security concerns for specific permission combinations
PERMISSION_COMBINATIONS = [
    {
        "permissions": ["android.permission.INTERNET", "android.permission.READ_CONTACTS"],
        "risk": "high",
        "description": "Can exfiltrate contacts over network"
    },
    {
        "permissions": ["android.permission.INTERNET", "android.permission.READ_SMS"],
        "risk": "critical",
        "description": "Can exfiltrate SMS messages (2FA codes!)"
    },
    {
        "permissions": ["android.permission.INTERNET", "android.permission.ACCESS_FINE_LOCATION"],
        "risk": "high",
        "description": "Can track and transmit precise location"
    },
    {
        "permissions": ["android.permission.INTERNET", "android.permission.CAMERA"],
        "risk": "high",
        "description": "Can capture and transmit photos"
    },
    {
        "permissions": ["android.permission.INTERNET", "android.permission.RECORD_AUDIO"],
        "risk": "high",
        "description": "Can record and transmit audio"
    },
    {
        "permissions": ["android.permission.READ_EXTERNAL_STORAGE", "android.permission.INTERNET"],
        "risk": "high",
        "description": "Can exfiltrate files over network"
    },
    {
        "permissions": ["android.permission.RECEIVE_BOOT_COMPLETED", "android.permission.INTERNET"],
        "risk": "medium",
        "description": "Auto-starts and connects to network (potential C&C)"
    },
    {
        "permissions": ["android.permission.SYSTEM_ALERT_WINDOW", "android.permission.RECORD_AUDIO"],
        "risk": "critical",
        "description": "Can overlay screens and record (phishing/spyware)"
    },
    {
        "permissions": ["android.permission.BIND_ACCESSIBILITY_SERVICE"],
        "risk": "critical",
        "description": "Full screen monitoring - can capture passwords"
    },
    {
        "permissions": ["android.permission.BIND_NOTIFICATION_LISTENER_SERVICE", "android.permission.INTERNET"],
        "risk": "critical",
        "description": "Can read and exfiltrate all notifications (including 2FA)"
    },
    {
        "permissions": ["android.permission.REQUEST_INSTALL_PACKAGES", "android.permission.INTERNET"],
        "risk": "high",
        "description": "Can download and install additional malware"
    },
]


def analyze_permissions(output_dir: Path) -> Dict[str, Any]:
    """
    Analyze permissions from AndroidManifest.xml.
    
    Args:
        output_dir: JADX output directory
    
    Returns:
        Dictionary with permission analysis results
    """
    import xml.etree.ElementTree as ET
    
    output_dir = Path(output_dir)
    manifest_path = output_dir / "resources" / "AndroidManifest.xml"
    
    # Also check root for manifest
    if not manifest_path.exists():
        manifest_path = output_dir / "AndroidManifest.xml"
    
    if not manifest_path.exists():
        return {"error": "AndroidManifest.xml not found"}
    
    try:
        tree = ET.parse(manifest_path)
        root = tree.getroot()
    except ET.ParseError as e:
        return {"error": f"Failed to parse manifest: {str(e)}"}
    
    # Extract namespace
    ns = {'android': 'http://schemas.android.com/apk/res/android'}
    
    # Find all uses-permission elements
    permissions = []
    permission_names = set()
    
    for elem in root.iter():
        if elem.tag == 'uses-permission' or elem.tag.endswith('}uses-permission'):
            perm_name = elem.get('{http://schemas.android.com/apk/res/android}name') or elem.get('android:name') or elem.get('name')
            if perm_name:
                permission_names.add(perm_name)
                
                # Look up permission info
                perm_info = ANDROID_PERMISSIONS.get(perm_name, {
                    "level": "unknown",
                    "description": "Unknown permission",
                    "category": "unknown"
                })
                
                permissions.append({
                    "name": perm_name,
                    "short_name": perm_name.split('.')[-1] if '.' in perm_name else perm_name,
                    "level": perm_info["level"],
                    "description": perm_info["description"],
                    "category": perm_info["category"]
                })
    
    # Sort by danger level
    level_order = {"signature": 0, "dangerous": 1, "normal": 2, "deprecated": 3, "unknown": 4}
    permissions.sort(key=lambda p: level_order.get(p["level"], 5))
    
    # Categorize permissions
    by_level = {"signature": [], "dangerous": [], "normal": [], "deprecated": [], "unknown": []}
    by_category = {}
    
    for perm in permissions:
        level = perm["level"]
        if level in by_level:
            by_level[level].append(perm)
        
        cat = perm["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(perm)
    
    # Check for dangerous combinations
    dangerous_combinations = []
    for combo in PERMISSION_COMBINATIONS:
        if all(p in permission_names for p in combo["permissions"]):
            dangerous_combinations.append({
                "permissions": combo["permissions"],
                "risk": combo["risk"],
                "description": combo["description"]
            })
    
    # Calculate risk score
    risk_score = 0
    risk_score += len(by_level["signature"]) * 25  # Signature perms are highest risk
    risk_score += len(by_level["dangerous"]) * 10
    risk_score += len(dangerous_combinations) * 15
    risk_score = min(risk_score, 100)  # Cap at 100
    
    # Determine overall risk level
    if risk_score >= 70 or len(by_level["signature"]) >= 2:
        overall_risk = "critical"
    elif risk_score >= 50 or len(by_level["dangerous"]) >= 5:
        overall_risk = "high"
    elif risk_score >= 25 or len(by_level["dangerous"]) >= 2:
        overall_risk = "medium"
    else:
        overall_risk = "low"
    
    # Generate summary
    summary_parts = []
    if by_level["signature"]:
        summary_parts.append(f"⚠️ {len(by_level['signature'])} system/signature permissions")
    if by_level["dangerous"]:
        summary_parts.append(f"🔴 {len(by_level['dangerous'])} dangerous permissions")
    if dangerous_combinations:
        summary_parts.append(f"💀 {len(dangerous_combinations)} risky permission combinations")
    
    return {
        "total_permissions": len(permissions),
        "permissions": permissions,
        "by_level": by_level,
        "by_category": by_category,
        "dangerous_combinations": dangerous_combinations,
        "risk_score": risk_score,
        "overall_risk": overall_risk,
        "summary": " | ".join(summary_parts) if summary_parts else "✅ Low-risk permission profile"
    }


# ============================================================================
# Feature: Network Endpoint Extractor
# ============================================================================

# Patterns for network endpoint extraction
NETWORK_PATTERNS = {
    # URLs
    "url_https": r'https://[^\s"\'<>)\]]+',
    "url_http": r'http://[^\s"\'<>)\]]+',
    
    # IP addresses
    "ipv4": r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b',
    "ipv4_port": r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?):\d{1,5}\b',
    
    # Domains (more specific patterns)
    "api_domain": r'["\']([a-zA-Z0-9][-a-zA-Z0-9]*\.)+[a-zA-Z]{2,}["\']',
    
    # API patterns
    "api_path": r'/api/v?\d*[/\w\-\.]+',
    "rest_endpoint": r'/(?:api|rest|v\d+|graphql|query|mutation)/[\w\-/\.]+',
    
    # WebSocket
    "websocket": r'wss?://[^\s"\'<>)\]]+',
    
    # Firebase
    "firebase_url": r'https://[\w\-]+\.firebaseio\.com[^\s"\'<>)\]]*',
    "firebase_storage": r'gs://[\w\-\.]+',
    "firebase_config": r'[\w\-]+\.firebaseapp\.com',
    
    # AWS
    "aws_s3": r's3://[\w\-\.]+',
    "aws_s3_url": r'https://[\w\-]+\.s3[\.\-][\w\-]+\.amazonaws\.com[^\s"\'<>)\]]*',
    "aws_api_gateway": r'https://[\w]+\.execute-api\.[\w\-]+\.amazonaws\.com[^\s"\'<>)\]]*',
    "aws_lambda": r'https://[\w]+\.lambda-url\.[\w\-]+\.on\.aws[^\s"\'<>)\]]*',
    
    # Azure
    "azure_blob": r'https://[\w]+\.blob\.core\.windows\.net[^\s"\'<>)\]]*',
    "azure_api": r'https://[\w\-]+\.azure-api\.net[^\s"\'<>)\]]*',
    
    # Google Cloud
    "gcp_storage": r'https://storage\.googleapis\.com/[\w\-]+[^\s"\'<>)\]]*',
    "gcp_functions": r'https://[\w\-]+\.cloudfunctions\.net[^\s"\'<>)\]]*',
    
    # Common API services
    "stripe_api": r'https://api\.stripe\.com[^\s"\'<>)\]]*',
    "twilio_api": r'https://api\.twilio\.com[^\s"\'<>)\]]*',
    "sendgrid_api": r'https://api\.sendgrid\.com[^\s"\'<>)\]]*',
    "slack_webhook": r'https://hooks\.slack\.com/[\w/]+',
    "discord_webhook": r'https://discord(?:app)?\.com/api/webhooks/[\w/]+',
    
    # GraphQL
    "graphql_endpoint": r'/graphql(?:/[\w\-]+)?',
}

# Categories for endpoints
ENDPOINT_CATEGORIES = {
    "url_https": "https_url",
    "url_http": "http_url",
    "ipv4": "ip_address",
    "ipv4_port": "ip_address",
    "api_domain": "domain",
    "api_path": "api_endpoint",
    "rest_endpoint": "api_endpoint",
    "websocket": "websocket",
    "firebase_url": "firebase",
    "firebase_storage": "firebase",
    "firebase_config": "firebase",
    "aws_s3": "aws",
    "aws_s3_url": "aws",
    "aws_api_gateway": "aws",
    "aws_lambda": "aws",
    "azure_blob": "azure",
    "azure_api": "azure",
    "gcp_storage": "gcp",
    "gcp_functions": "gcp",
    "stripe_api": "payment_service",
    "twilio_api": "communication_service",
    "sendgrid_api": "email_service",
    "slack_webhook": "webhook",
    "discord_webhook": "webhook",
    "graphql_endpoint": "graphql",
}

# Risk levels for endpoint types
ENDPOINT_RISKS = {
    "http_url": "high",  # Unencrypted
    "ip_address": "medium",  # Hardcoded IPs
    "webhook": "high",  # Sensitive webhooks
    "firebase": "medium",
    "aws": "medium",
    "azure": "medium",
    "gcp": "medium",
    "payment_service": "high",
    "https_url": "low",
    "api_endpoint": "low",
    "websocket": "low",
    "domain": "low",
    "graphql": "low",
    "communication_service": "medium",
    "email_service": "medium",
}

# Domains to exclude (common SDK/library domains)
EXCLUDED_DOMAINS = {
    "schemas.android.com",
    "www.w3.org",
    "ns.adobe.com",
    "xmlpull.org",
    "json.org",
    "apache.org",
    "google.com",  # Generic google
    "googleapis.com",  # Generic APIs
    "gstatic.com",
    "googleusercontent.com",
    "android.com",
    "example.com",
    "localhost",
    "127.0.0.1",
    "0.0.0.0",
}


def extract_network_endpoints(output_dir: Path) -> Dict[str, Any]:
    """
    Extract all network endpoints from decompiled sources.
    
    Args:
        output_dir: JADX output directory
    
    Returns:
        Dictionary with extracted endpoints and analysis
    """
    output_dir = Path(output_dir)
    sources_dir = output_dir / "sources"
    
    if not sources_dir.exists():
        return {"error": "Decompiled sources not found"}
    
    endpoints = []
    seen = set()  # Deduplicate
    
    # Process all source files
    for java_file in sources_dir.rglob("*.java"):
        try:
            content = java_file.read_text(encoding='utf-8', errors='ignore')
            rel_path = str(java_file.relative_to(sources_dir))
            
            # Search for each pattern
            for pattern_name, pattern in NETWORK_PATTERNS.items():
                for match in re.finditer(pattern, content, re.IGNORECASE):
                    value = match.group(0).strip('"\'')
                    
                    # Skip if already seen
                    if value in seen:
                        continue
                    
                    # Skip excluded domains
                    skip = False
                    for excluded in EXCLUDED_DOMAINS:
                        if excluded in value.lower():
                            skip = True
                            break
                    if skip:
                        continue
                    
                    # Skip very short matches
                    if len(value) < 8:
                        continue
                    
                    seen.add(value)
                    
                    category = ENDPOINT_CATEGORIES.get(pattern_name, "other")
                    risk = ENDPOINT_RISKS.get(category, "low")
                    
                    # Find line number
                    line_num = content[:match.start()].count('\n') + 1
                    
                    endpoints.append({
                        "value": value,
                        "type": pattern_name,
                        "category": category,
                        "risk": risk,
                        "file": rel_path,
                        "line": line_num
                    })
        except Exception as e:
            continue
    
    # Also search in resources (strings.xml, etc.)
    resources_dir = output_dir / "resources"
    if resources_dir.exists():
        for xml_file in resources_dir.rglob("*.xml"):
            try:
                content = xml_file.read_text(encoding='utf-8', errors='ignore')
                rel_path = f"resources/{xml_file.relative_to(resources_dir)}"
                
                for pattern_name, pattern in NETWORK_PATTERNS.items():
                    for match in re.finditer(pattern, content, re.IGNORECASE):
                        value = match.group(0).strip('"\'')
                        
                        if value in seen:
                            continue
                        
                        skip = False
                        for excluded in EXCLUDED_DOMAINS:
                            if excluded in value.lower():
                                skip = True
                                break
                        if skip:
                            continue
                        
                        if len(value) < 8:
                            continue
                        
                        seen.add(value)
                        
                        category = ENDPOINT_CATEGORIES.get(pattern_name, "other")
                        risk = ENDPOINT_RISKS.get(category, "low")
                        line_num = content[:match.start()].count('\n') + 1
                        
                        endpoints.append({
                            "value": value,
                            "type": pattern_name,
                            "category": category,
                            "risk": risk,
                            "file": rel_path,
                            "line": line_num
                        })
            except:
                continue
    
    # Sort by risk then value
    risk_order = {"high": 0, "medium": 1, "low": 2}
    endpoints.sort(key=lambda e: (risk_order.get(e["risk"], 3), e["value"]))
    
    # Categorize results
    by_category = {}
    by_risk = {"high": [], "medium": [], "low": []}
    
    for ep in endpoints:
        cat = ep["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(ep)
        
        risk = ep["risk"]
        if risk in by_risk:
            by_risk[risk].append(ep)
    
    # Find unique domains
    domains = set()
    for ep in endpoints:
        if "://" in ep["value"]:
            try:
                from urllib.parse import urlparse
                parsed = urlparse(ep["value"])
                if parsed.netloc:
                    domains.add(parsed.netloc)
            except:
                pass
    
    # Generate summary
    summary_parts = []
    if by_risk["high"]:
        summary_parts.append(f"🔴 {len(by_risk['high'])} high-risk endpoints")
    if by_risk["medium"]:
        summary_parts.append(f"🟡 {len(by_risk['medium'])} medium-risk")
    if any("http_url" in ep["type"] for ep in endpoints):
        summary_parts.append("⚠️ HTTP (unencrypted) URLs found")
    if any(ep["category"] == "webhook" for ep in endpoints):
        summary_parts.append("🔗 Webhooks detected")
    
    return {
        "total_endpoints": len(endpoints),
        "endpoints": endpoints,
        "by_category": by_category,
        "by_risk": by_risk,
        "unique_domains": sorted(list(domains)),
        "domain_count": len(domains),
        "summary": " | ".join(summary_parts) if summary_parts else "✅ No high-risk network endpoints found"
    }


# ============================================================================
# Feature: Crypto Audit
# ============================================================================

# Weak crypto algorithms that should be flagged
WEAK_CRYPTO_PATTERNS = {
    # Weak hashing algorithms
    "md5": {
        "patterns": [
            r'MessageDigest\.getInstance\s*\(\s*["\']MD5["\']\s*\)',
            r'\.md5\s*\(',
            r'MD5\.digest',
            r'DigestUtils\.md5',
            r'Hashing\.md5\s*\(',
        ],
        "severity": "high",
        "category": "weak_hash",
        "description": "MD5 is cryptographically broken and should not be used for security purposes",
        "recommendation": "Use SHA-256 or SHA-3 instead of MD5"
    },
    "sha1": {
        "patterns": [
            r'MessageDigest\.getInstance\s*\(\s*["\']SHA-?1["\']\s*\)',
            r'\.sha1\s*\(',
            r'SHA1\.digest',
            r'DigestUtils\.sha1',
            r'Hashing\.sha1\s*\(',
        ],
        "severity": "medium",
        "category": "weak_hash",
        "description": "SHA-1 is deprecated and vulnerable to collision attacks",
        "recommendation": "Use SHA-256 or SHA-3 instead of SHA-1"
    },
    
    # Weak encryption algorithms
    "des": {
        "patterns": [
            r'Cipher\.getInstance\s*\(\s*["\']DES[^E]',
            r'DESKeySpec',
            r'DES/ECB',
            r'DES/CBC',
            r'SecretKeyFactory\.getInstance\s*\(\s*["\']DES["\']\s*\)',
        ],
        "severity": "critical",
        "category": "weak_encryption",
        "description": "DES is obsolete with only 56-bit key strength",
        "recommendation": "Use AES-256 instead of DES"
    },
    "3des": {
        "patterns": [
            r'Cipher\.getInstance\s*\(\s*["\']DESede',
            r'DESedeKeySpec',
            r'TripleDES',
            r'3DES',
        ],
        "severity": "medium",
        "category": "weak_encryption",
        "description": "3DES is deprecated with known vulnerabilities",
        "recommendation": "Use AES-256 instead of 3DES"
    },
    "rc4": {
        "patterns": [
            r'Cipher\.getInstance\s*\(\s*["\']RC4',
            r'ARCFOUR',
            r'RC4/None',
        ],
        "severity": "critical",
        "category": "weak_encryption",
        "description": "RC4 has critical vulnerabilities and is prohibited in TLS",
        "recommendation": "Use AES-GCM instead of RC4"
    },
    "blowfish": {
        "patterns": [
            r'Cipher\.getInstance\s*\(\s*["\']Blowfish',
        ],
        "severity": "low",
        "category": "weak_encryption",
        "description": "Blowfish has a 64-bit block size which may be vulnerable",
        "recommendation": "Consider using AES-256 for new implementations"
    },
    
    # ECB mode (insecure)
    "ecb_mode": {
        "patterns": [
            r'Cipher\.getInstance\s*\(\s*["\'][A-Z0-9]+/ECB/',
            r'/ECB/PKCS[57]Padding',
            r'/ECB/NoPadding',
        ],
        "severity": "high",
        "category": "insecure_mode",
        "description": "ECB mode does not provide semantic security - identical plaintext blocks produce identical ciphertext",
        "recommendation": "Use CBC, CTR, or GCM mode instead of ECB"
    },
    
    # Hardcoded keys and IVs
    "hardcoded_key": {
        "patterns": [
            r'SecretKeySpec\s*\(\s*["\'][^"\']+["\']\.getBytes\(',
            r'SecretKeySpec\s*\(\s*new\s+byte\s*\[\s*\]\s*\{[^}]+\}',
            r'\.getBytes\s*\(\s*\)\s*,\s*["\']AES["\']\s*\)',
            r'IvParameterSpec\s*\(\s*["\'][^"\']+["\']\.getBytes\(',
            r'IvParameterSpec\s*\(\s*new\s+byte\s*\[\s*\]\s*\{[^}]+\}',
        ],
        "severity": "critical",
        "category": "hardcoded_secret",
        "description": "Cryptographic keys or IVs appear to be hardcoded in source code",
        "recommendation": "Store keys securely using Android Keystore or secure key derivation"
    },
    
    # Static IVs
    "static_iv": {
        "patterns": [
            r'new\s+byte\s*\[\s*\]\s*\{\s*0\s*,\s*0\s*,\s*0',
            r'IvParameterSpec\s*\(\s*["\'][0]+["\']\s*\)',
            r'IV\s*=\s*["\'][^"\']{16,}["\']',
        ],
        "severity": "high",
        "category": "static_iv",
        "description": "Static or null IVs compromise encryption security",
        "recommendation": "Generate a random IV for each encryption operation"
    },
    
    # Insecure random
    "insecure_random": {
        "patterns": [
            r'new\s+Random\s*\(',
            r'java\.util\.Random',
            r'Math\.random\s*\(',
        ],
        "severity": "medium",
        "category": "weak_random",
        "description": "java.util.Random is not cryptographically secure",
        "recommendation": "Use SecureRandom for cryptographic purposes"
    },
    
    # No padding
    "no_padding": {
        "patterns": [
            r'Cipher\.getInstance\s*\(\s*["\'][A-Z0-9]+/[A-Z]+/NoPadding',
        ],
        "severity": "medium",
        "category": "insecure_padding",
        "description": "NoPadding may be vulnerable to padding oracle attacks",
        "recommendation": "Use proper padding like PKCS7 or use authenticated encryption (GCM)"
    },
    
    # RSA without OAEP
    "rsa_pkcs1": {
        "patterns": [
            r'Cipher\.getInstance\s*\(\s*["\']RSA/[^/]*/PKCS1Padding',
            r'Cipher\.getInstance\s*\(\s*["\']RSA["\']\s*\)',
        ],
        "severity": "medium",
        "category": "weak_rsa",
        "description": "RSA with PKCS#1 v1.5 padding is vulnerable to attacks",
        "recommendation": "Use RSA with OAEP padding"
    },
    
    # Short key lengths
    "short_rsa_key": {
        "patterns": [
            r'KeyPairGenerator\.getInstance\s*\([^)]+\)\s*;\s*[^;]*\.initialize\s*\(\s*(?:512|768|1024)\s*[,)]',
        ],
        "severity": "high",
        "category": "weak_key_length",
        "description": "RSA keys shorter than 2048 bits are considered weak",
        "recommendation": "Use at least 2048-bit RSA keys, preferably 4096"
    },
    
    # Password-based encryption with weak iteration
    "weak_pbkdf": {
        "patterns": [
            r'PBEKeySpec\s*\([^,]+,\s*[^,]+,\s*(?:1|[0-9]{1,3})\s*[,)]',
            r'\.setIterationCount\s*\(\s*(?:1|[0-9]{1,3})\s*\)',
        ],
        "severity": "high",
        "category": "weak_kdf",
        "description": "Password-based key derivation with low iteration count",
        "recommendation": "Use at least 100,000 iterations for PBKDF2"
    },
    
    # Certificate pinning bypass potential
    "cert_pinning_bypass": {
        "patterns": [
            r'TrustManager\s*\[\s*\]\s*\{\s*new\s+X509TrustManager',
            r'checkServerTrusted.*\{\s*\}',
            r'getAcceptedIssuers.*return\s+null',
            r'trustAllCerts',
            r'ALLOW_ALL_HOSTNAME_VERIFIER',
            r'\.setHostnameVerifier\s*\(\s*SSLSocketFactory\.ALLOW_ALL',
        ],
        "severity": "critical",
        "category": "certificate_validation",
        "description": "Certificate validation appears to be disabled or bypassed",
        "recommendation": "Implement proper certificate pinning and validation"
    },
}

# Good crypto practices to highlight
GOOD_CRYPTO_PATTERNS = {
    "secure_random": [r'SecureRandom', r'new\s+SecureRandom\s*\('],
    "aes_gcm": [r'/GCM/', r'AES/GCM'],
    "sha256": [r'SHA-?256', r'MessageDigest\.getInstance\s*\(\s*["\']SHA-256'],
    "sha512": [r'SHA-?512', r'MessageDigest\.getInstance\s*\(\s*["\']SHA-512'],
    "keystore": [r'KeyStore\.getInstance', r'AndroidKeyStore'],
    "bcrypt": [r'BCrypt', r'bcrypt'],
    "argon2": [r'Argon2', r'argon2'],
}


def crypto_audit(output_dir: Path) -> Dict[str, Any]:
    """
    Perform a comprehensive cryptographic audit on decompiled APK sources.
    
    Args:
        output_dir: JADX output directory
    
    Returns:
        Dictionary with audit results, findings, and recommendations
    """
    output_dir = Path(output_dir)
    sources_dir = output_dir / "sources"
    
    if not sources_dir.exists():
        return {"error": "Decompiled sources not found"}
    
    findings = []
    good_practices = []
    crypto_methods = []  # Track all crypto usage for overview
    files_scanned = 0
    
    # Scan all Java source files
    for java_file in sources_dir.rglob("*.java"):
        try:
            content = java_file.read_text(encoding='utf-8', errors='ignore')
            rel_path = str(java_file.relative_to(sources_dir))
            files_scanned += 1
            
            # Check for weak patterns
            for vuln_name, vuln_info in WEAK_CRYPTO_PATTERNS.items():
                for pattern in vuln_info["patterns"]:
                    for match in re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE):
                        line_num = content[:match.start()].count('\n') + 1
                        
                        # Get context (surrounding lines)
                        lines = content.split('\n')
                        start_line = max(0, line_num - 3)
                        end_line = min(len(lines), line_num + 2)
                        context = '\n'.join(lines[start_line:end_line])
                        
                        findings.append({
                            "type": vuln_name,
                            "category": vuln_info["category"],
                            "severity": vuln_info["severity"],
                            "description": vuln_info["description"],
                            "recommendation": vuln_info["recommendation"],
                            "file": rel_path,
                            "line": line_num,
                            "match": match.group(0)[:100],  # Truncate long matches
                            "context": context[:500],  # Truncate context
                        })
            
            # Check for good practices
            for practice_name, patterns in GOOD_CRYPTO_PATTERNS.items():
                for pattern in patterns:
                    for match in re.finditer(pattern, content, re.IGNORECASE):
                        line_num = content[:match.start()].count('\n') + 1
                        good_practices.append({
                            "type": practice_name,
                            "file": rel_path,
                            "line": line_num,
                            "match": match.group(0)[:50]
                        })
            
            # Track crypto API usage for overview
            crypto_api_patterns = [
                (r'Cipher\.getInstance', "encryption"),
                (r'MessageDigest\.getInstance', "hashing"),
                (r'KeyGenerator\.getInstance', "key_generation"),
                (r'KeyPairGenerator\.getInstance', "asymmetric_key"),
                (r'Mac\.getInstance', "mac"),
                (r'Signature\.getInstance', "digital_signature"),
                (r'KeyStore\.getInstance', "keystore"),
                (r'SecretKeyFactory\.getInstance', "key_derivation"),
            ]
            
            for pattern, usage_type in crypto_api_patterns:
                for match in re.finditer(pattern, content):
                    line_num = content[:match.start()].count('\n') + 1
                    # Try to extract algorithm name
                    algo_match = re.search(r'\(\s*["\']([^"\']+)["\']\s*\)', content[match.start():match.start()+100])
                    algorithm = algo_match.group(1) if algo_match else "unknown"
                    
                    crypto_methods.append({
                        "type": usage_type,
                        "algorithm": algorithm,
                        "file": rel_path,
                        "line": line_num
                    })
        
        except Exception as e:
            continue
    
    # Deduplicate findings (same type + file + line)
    seen_findings = set()
    unique_findings = []
    for f in findings:
        key = (f["type"], f["file"], f["line"])
        if key not in seen_findings:
            seen_findings.add(key)
            unique_findings.append(f)
    findings = unique_findings
    
    # Sort findings by severity
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    findings.sort(key=lambda f: (severity_order.get(f["severity"], 4), f["file"]))
    
    # Categorize findings
    by_severity = {"critical": [], "high": [], "medium": [], "low": []}
    by_category = {}
    
    for finding in findings:
        sev = finding["severity"]
        if sev in by_severity:
            by_severity[sev].append(finding)
        
        cat = finding["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(finding)
    
    # Calculate risk score
    risk_score = 0
    risk_score += len(by_severity["critical"]) * 30
    risk_score += len(by_severity["high"]) * 15
    risk_score += len(by_severity["medium"]) * 5
    risk_score += len(by_severity["low"]) * 1
    risk_score = min(risk_score, 100)  # Cap at 100
    
    # Determine overall grade
    if risk_score >= 70:
        grade = "F"
        overall_risk = "critical"
    elif risk_score >= 50:
        grade = "D"
        overall_risk = "high"
    elif risk_score >= 30:
        grade = "C"
        overall_risk = "medium"
    elif risk_score >= 10:
        grade = "B"
        overall_risk = "low"
    else:
        grade = "A"
        overall_risk = "minimal"
    
    # Generate summary
    summary_parts = []
    if by_severity["critical"]:
        summary_parts.append(f"💀 {len(by_severity['critical'])} critical issues")
    if by_severity["high"]:
        summary_parts.append(f"🔴 {len(by_severity['high'])} high-risk issues")
    if by_severity["medium"]:
        summary_parts.append(f"🟡 {len(by_severity['medium'])} medium issues")
    if good_practices:
        summary_parts.append(f"✅ {len(set(p['type'] for p in good_practices))} good practices found")
    
    # Generate recommendations
    top_recommendations = []
    if by_category.get("hardcoded_secret"):
        top_recommendations.append("Move cryptographic keys to Android Keystore or use secure key derivation")
    if by_category.get("weak_hash"):
        top_recommendations.append("Replace MD5/SHA-1 with SHA-256 or SHA-3")
    if by_category.get("weak_encryption"):
        top_recommendations.append("Upgrade from DES/3DES/RC4 to AES-256-GCM")
    if by_category.get("insecure_mode"):
        top_recommendations.append("Use authenticated encryption modes (GCM) instead of ECB")
    if by_category.get("certificate_validation"):
        top_recommendations.append("Implement proper certificate validation and pinning")
    if by_category.get("weak_random"):
        top_recommendations.append("Use SecureRandom instead of java.util.Random for cryptographic purposes")
    
    return {
        "total_findings": len(findings),
        "findings": findings,
        "by_severity": by_severity,
        "by_category": by_category,
        "good_practices": good_practices[:50],  # Limit
        "crypto_methods": crypto_methods[:100],  # Limit
        "files_scanned": files_scanned,
        "risk_score": risk_score,
        "grade": grade,
        "overall_risk": overall_risk,
        "top_recommendations": top_recommendations[:5],
        "summary": " | ".join(summary_parts) if summary_parts else "✅ No cryptographic issues detected"
    }


# ============================================================================
# Feature: Activity/Service Map (Component Map)
# ============================================================================

def generate_component_map(output_dir: Path) -> Dict[str, Any]:
    """
    Generate a visual component map showing activities, services, receivers,
    providers and their relationships (intents, exports, etc.)
    
    Args:
        output_dir: JADX output directory
    
    Returns:
        Dictionary with component map data for visualization
    """
    output_dir = Path(output_dir)
    resources_dir = output_dir / "resources"
    sources_dir = output_dir / "sources"
    
    # Find AndroidManifest.xml
    manifest_path = None
    for candidate in [
        resources_dir / "AndroidManifest.xml",
        output_dir / "AndroidManifest.xml",
    ]:
        if candidate.exists():
            manifest_path = candidate
            break
    
    if not manifest_path:
        return {"error": "AndroidManifest.xml not found"}
    
    try:
        import xml.etree.ElementTree as ET
        tree = ET.parse(manifest_path)
        root = tree.getroot()
        
        # Handle Android namespace
        ns = {"android": "http://schemas.android.com/apk/res/android"}
        
        def get_attr(elem, attr):
            """Get attribute with Android namespace."""
            return elem.get(f"{{{ns['android']}}}{attr}", elem.get(attr, ""))
        
        # Extract package name
        package_name = root.get("package", "unknown")
        
        # Component containers
        components = {
            "activities": [],
            "services": [],
            "receivers": [],
            "providers": [],
        }
        
        # Intent filter relationships
        intent_connections = []
        deep_links = []
        
        # Find application element
        app_elem = root.find("application")
        if app_elem is None:
            return {"error": "No application element in manifest"}
        
        # Process activities
        for activity in app_elem.findall("activity"):
            name = get_attr(activity, "name")
            if name.startswith("."):
                name = package_name + name
            
            exported = get_attr(activity, "exported")
            is_exported = exported.lower() == "true"
            
            # Check if it has intent filters (implicitly exported if no explicit export)
            intent_filters = activity.findall("intent-filter")
            if intent_filters and exported == "":
                is_exported = True
            
            # Determine if it's a launcher activity
            is_launcher = False
            actions = []
            categories = []
            data_schemes = []
            
            for intent_filter in intent_filters:
                for action in intent_filter.findall("action"):
                    action_name = get_attr(action, "name")
                    actions.append(action_name)
                    if action_name == "android.intent.action.MAIN":
                        for cat in intent_filter.findall("category"):
                            if get_attr(cat, "name") == "android.intent.category.LAUNCHER":
                                is_launcher = True
                
                for category in intent_filter.findall("category"):
                    categories.append(get_attr(category, "name"))
                
                for data in intent_filter.findall("data"):
                    scheme = get_attr(data, "scheme")
                    host = get_attr(data, "host")
                    path = get_attr(data, "path") or get_attr(data, "pathPrefix") or get_attr(data, "pathPattern")
                    if scheme:
                        deep_link = f"{scheme}://{host or '*'}{path or ''}"
                        data_schemes.append(scheme)
                        deep_links.append({
                            "scheme": scheme,
                            "host": host,
                            "path": path,
                            "component": name.split(".")[-1],
                            "component_full": name,
                            "type": "activity"
                        })
            
            # Risk assessment
            risk = "low"
            if is_exported and not is_launcher:
                risk = "medium"
                if actions and any("BROWSABLE" in c for c in categories):
                    risk = "high"  # Deep link exposed
            
            components["activities"].append({
                "name": name.split(".")[-1],
                "full_name": name,
                "exported": is_exported,
                "launcher": is_launcher,
                "actions": actions,
                "categories": categories,
                "data_schemes": data_schemes,
                "risk": risk,
                "theme": get_attr(activity, "theme"),
                "launch_mode": get_attr(activity, "launchMode") or "standard",
            })
        
        # Process services
        for service in app_elem.findall("service"):
            name = get_attr(service, "name")
            if name.startswith("."):
                name = package_name + name
            
            exported = get_attr(service, "exported")
            intent_filters = service.findall("intent-filter")
            is_exported = exported.lower() == "true" or (intent_filters and exported == "")
            
            actions = []
            for intent_filter in intent_filters:
                for action in intent_filter.findall("action"):
                    actions.append(get_attr(action, "name"))
            
            # Check for bound service patterns
            permission = get_attr(service, "permission")
            
            risk = "low"
            if is_exported:
                risk = "medium" if permission else "high"
            
            components["services"].append({
                "name": name.split(".")[-1],
                "full_name": name,
                "exported": is_exported,
                "actions": actions,
                "permission": permission,
                "foreground": get_attr(service, "foregroundServiceType") != "",
                "risk": risk,
            })
        
        # Process receivers
        for receiver in app_elem.findall("receiver"):
            name = get_attr(receiver, "name")
            if name.startswith("."):
                name = package_name + name
            
            exported = get_attr(receiver, "exported")
            intent_filters = receiver.findall("intent-filter")
            is_exported = exported.lower() == "true" or (intent_filters and exported == "")
            
            actions = []
            for intent_filter in intent_filters:
                for action in intent_filter.findall("action"):
                    actions.append(get_attr(action, "name"))
            
            permission = get_attr(receiver, "permission")
            
            # System broadcast check
            is_system = any("android.intent.action" in a for a in actions)
            
            risk = "low"
            if is_exported and not is_system:
                risk = "medium" if permission else "high"
            
            components["receivers"].append({
                "name": name.split(".")[-1],
                "full_name": name,
                "exported": is_exported,
                "actions": actions,
                "permission": permission,
                "system_broadcast": is_system,
                "risk": risk,
            })
        
        # Process providers
        for provider in app_elem.findall("provider"):
            name = get_attr(provider, "name")
            if name.startswith("."):
                name = package_name + name
            
            exported = get_attr(provider, "exported")
            # Providers default to exported=true if targetSdk < 17
            is_exported = exported.lower() != "false"
            
            authorities = get_attr(provider, "authorities")
            read_perm = get_attr(provider, "readPermission")
            write_perm = get_attr(provider, "writePermission")
            grant_uri = get_attr(provider, "grantUriPermissions")
            
            risk = "low"
            if is_exported:
                if not read_perm and not write_perm:
                    risk = "critical"
                elif not read_perm or not write_perm:
                    risk = "high"
                else:
                    risk = "medium"
            
            components["providers"].append({
                "name": name.split(".")[-1],
                "full_name": name,
                "exported": is_exported,
                "authorities": authorities,
                "read_permission": read_perm,
                "write_permission": write_perm,
                "grant_uri_permissions": grant_uri.lower() == "true",
                "risk": risk,
            })
        
        # Analyze inter-component connections from source code
        connections = []
        if sources_dir.exists():
            # Patterns for finding component invocations
            intent_patterns = [
                r'new\s+Intent\s*\(\s*[^,]+,\s*(\w+)\.class\s*\)',
                r'Intent\s*\(\s*[^,]+,\s*(\w+)\.class\s*\)',
                r'startActivity\s*\([^)]*(\w+)\.class',
                r'startService\s*\([^)]*(\w+)\.class',
                r'bindService\s*\([^)]*(\w+)\.class',
                r'startActivityForResult\s*\([^)]*(\w+)\.class',
            ]
            
            for java_file in list(sources_dir.rglob("*.java"))[:200]:  # Limit for performance
                try:
                    content = java_file.read_text(encoding='utf-8', errors='ignore')
                    source_name = java_file.stem
                    
                    for pattern in intent_patterns:
                        for match in re.finditer(pattern, content):
                            target = match.group(1)
                            if target != source_name:  # Skip self-references
                                connections.append({
                                    "source": source_name,
                                    "target": target,
                                    "type": "intent"
                                })
                except:
                    continue
        
        # Deduplicate connections
        seen = set()
        unique_connections = []
        for conn in connections:
            key = (conn["source"], conn["target"])
            if key not in seen:
                seen.add(key)
                unique_connections.append(conn)
        
        # Statistics
        stats = {
            "total_activities": len(components["activities"]),
            "total_services": len(components["services"]),
            "total_receivers": len(components["receivers"]),
            "total_providers": len(components["providers"]),
            "exported_activities": len([a for a in components["activities"] if a["exported"]]),
            "exported_services": len([s for s in components["services"] if s["exported"]]),
            "exported_receivers": len([r for r in components["receivers"] if r["exported"]]),
            "exported_providers": len([p for p in components["providers"] if p["exported"]]),
            "deep_links": len(deep_links),
            "connections": len(unique_connections),
        }
        
        # Risk summary
        all_components = (
            components["activities"] + 
            components["services"] + 
            components["receivers"] + 
            components["providers"]
        )
        
        risk_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
        for comp in all_components:
            risk_counts[comp.get("risk", "low")] += 1
        
        # Calculate attack surface score
        attack_surface_score = 0
        attack_surface_score += stats["exported_activities"] * 5
        attack_surface_score += stats["exported_services"] * 10
        attack_surface_score += stats["exported_receivers"] * 3
        attack_surface_score += stats["exported_providers"] * 15
        attack_surface_score += len(deep_links) * 8
        attack_surface_score = min(attack_surface_score, 100)
        
        return {
            "package_name": package_name,
            "components": components,
            "connections": unique_connections,
            "deep_links": deep_links,
            "stats": stats,
            "risk_counts": risk_counts,
            "attack_surface_score": attack_surface_score,
            "summary": f"📊 {stats['total_activities']} activities, {stats['total_services']} services, {stats['total_receivers']} receivers, {stats['total_providers']} providers | 🔓 {sum([stats['exported_activities'], stats['exported_services'], stats['exported_receivers'], stats['exported_providers']])} exported | 🔗 {len(deep_links)} deep links"
        }
        
    except Exception as e:
        return {"error": f"Failed to parse manifest: {str(e)}"}


# ============================================================================
# Feature: Jump to Definition (Symbol Lookup)
# ============================================================================

def build_symbol_index(output_dir: Path) -> Dict[str, Any]:
    """
    Build an index of all class and method definitions for jump-to-definition.
    
    Args:
        output_dir: JADX output directory
    
    Returns:
        Dictionary with symbol index
    """
    output_dir = Path(output_dir)
    sources_dir = output_dir / "sources"
    
    if not sources_dir.exists():
        return {"error": "Decompiled sources not found"}
    
    symbols = {
        "classes": {},  # class_name -> {file, line, package}
        "methods": {},  # method_name -> [{class, file, line, signature}]
        "fields": {},   # field_name -> [{class, file, line, type}]
    }
    
    # Patterns for extraction
    class_pattern = re.compile(r'^(?:public|private|protected)?\s*(?:abstract|final|static)?\s*(?:class|interface|enum)\s+(\w+)', re.MULTILINE)
    method_pattern = re.compile(r'^\s*(?:public|private|protected)?\s*(?:abstract|static|final|synchronized|native)?\s*(?:<[\w\s,]+>\s*)?(\w+(?:<[^>]+>)?|\w+)\s+(\w+)\s*\(([^)]*)\)\s*(?:throws\s+[\w\s,]+)?\s*\{?', re.MULTILINE)
    field_pattern = re.compile(r'^\s*(?:public|private|protected)?\s*(?:static|final|volatile|transient)?\s*(\w+(?:<[^>]+>)?)\s+(\w+)\s*[;=]', re.MULTILINE)
    
    files_indexed = 0
    
    for java_file in sources_dir.rglob("*.java"):
        try:
            content = java_file.read_text(encoding='utf-8', errors='ignore')
            rel_path = str(java_file.relative_to(sources_dir))
            files_indexed += 1
            
            # Extract package
            package_match = re.search(r'^package\s+([\w\.]+)\s*;', content, re.MULTILINE)
            package_name = package_match.group(1) if package_match else ""
            
            # Find class definitions
            for match in class_pattern.finditer(content):
                class_name = match.group(1)
                line_num = content[:match.start()].count('\n') + 1
                full_name = f"{package_name}.{class_name}" if package_name else class_name
                
                symbols["classes"][class_name] = {
                    "file": rel_path,
                    "line": line_num,
                    "package": package_name,
                    "full_name": full_name
                }
                
                # Also index by full name
                symbols["classes"][full_name] = symbols["classes"][class_name]
            
            # Find method definitions (only in class context)
            for match in method_pattern.finditer(content):
                return_type = match.group(1)
                method_name = match.group(2)
                params = match.group(3).strip()
                line_num = content[:match.start()].count('\n') + 1
                
                # Skip constructors (they match too)
                if method_name in symbols["classes"]:
                    continue
                
                # Find containing class
                class_context = None
                for class_name, class_info in symbols["classes"].items():
                    if class_info["file"] == rel_path:
                        class_context = class_name
                        break
                
                if method_name not in symbols["methods"]:
                    symbols["methods"][method_name] = []
                
                symbols["methods"][method_name].append({
                    "class": class_context,
                    "file": rel_path,
                    "line": line_num,
                    "return_type": return_type,
                    "params": params,
                    "signature": f"{return_type} {method_name}({params})"
                })
            
            # Find field definitions
            for match in field_pattern.finditer(content):
                field_type = match.group(1)
                field_name = match.group(2)
                line_num = content[:match.start()].count('\n') + 1
                
                # Skip common names that are likely not fields
                if field_name in ["if", "for", "while", "return", "new", "this"]:
                    continue
                
                class_context = None
                for class_name, class_info in symbols["classes"].items():
                    if class_info["file"] == rel_path:
                        class_context = class_name
                        break
                
                if field_name not in symbols["fields"]:
                    symbols["fields"][field_name] = []
                
                symbols["fields"][field_name].append({
                    "class": class_context,
                    "file": rel_path,
                    "line": line_num,
                    "type": field_type
                })
        
        except Exception as e:
            continue
    
    return {
        "symbols": symbols,
        "stats": {
            "classes": len(symbols["classes"]),
            "methods": len(symbols["methods"]),
            "fields": len(symbols["fields"]),
            "files_indexed": files_indexed
        }
    }


def lookup_symbol(output_dir: Path, symbol: str, symbol_type: Optional[str] = None) -> Dict[str, Any]:
    """
    Look up a symbol (class, method, or field) and return its definition location.
    
    Args:
        output_dir: JADX output directory
        symbol: The symbol name to look up
        symbol_type: Optional type filter ('class', 'method', 'field')
    
    Returns:
        Dictionary with lookup results
    """
    # Build index (could be cached in production)
    index_result = build_symbol_index(output_dir)
    
    if "error" in index_result:
        return index_result
    
    symbols = index_result["symbols"]
    results = []
    
    # Search classes
    if symbol_type is None or symbol_type == "class":
        if symbol in symbols["classes"]:
            info = symbols["classes"][symbol]
            results.append({
                "type": "class",
                "name": symbol,
                "file": info["file"],
                "line": info["line"],
                "package": info.get("package", ""),
                "full_name": info.get("full_name", symbol)
            })
        else:
            # Partial match
            for class_name, info in symbols["classes"].items():
                if symbol.lower() in class_name.lower():
                    results.append({
                        "type": "class",
                        "name": class_name,
                        "file": info["file"],
                        "line": info["line"],
                        "package": info.get("package", ""),
                        "full_name": info.get("full_name", class_name)
                    })
    
    # Search methods
    if symbol_type is None or symbol_type == "method":
        if symbol in symbols["methods"]:
            for method_info in symbols["methods"][symbol]:
                results.append({
                    "type": "method",
                    "name": symbol,
                    "class": method_info.get("class"),
                    "file": method_info["file"],
                    "line": method_info["line"],
                    "signature": method_info.get("signature", ""),
                    "return_type": method_info.get("return_type", ""),
                    "params": method_info.get("params", "")
                })
        else:
            # Partial match for methods
            for method_name, method_list in symbols["methods"].items():
                if symbol.lower() in method_name.lower():
                    for method_info in method_list[:3]:  # Limit matches
                        results.append({
                            "type": "method",
                            "name": method_name,
                            "class": method_info.get("class"),
                            "file": method_info["file"],
                            "line": method_info["line"],
                            "signature": method_info.get("signature", "")
                        })
    
    # Search fields
    if symbol_type is None or symbol_type == "field":
        if symbol in symbols["fields"]:
            for field_info in symbols["fields"][symbol]:
                results.append({
                    "type": "field",
                    "name": symbol,
                    "class": field_info.get("class"),
                    "file": field_info["file"],
                    "line": field_info["line"],
                    "field_type": field_info.get("type", "")
                })
    
    # Limit and sort results
    results = results[:50]  # Limit total results
    results.sort(key=lambda r: (
        0 if r["type"] == "class" else 1 if r["type"] == "method" else 2,
        r["name"].lower() != symbol.lower(),  # Exact matches first
        len(r["name"])  # Shorter names first
    ))
    
    return {
        "symbol": symbol,
        "results": results,
        "total_found": len(results),
        "index_stats": index_result["stats"]
    }


def generate_class_dependency_graph(output_dir: Path, max_classes: int = 100) -> Dict[str, Any]:
    """
    Generate a class dependency graph showing how classes are interconnected.
    
    This analyzes:
    - Import statements (which classes depend on which)
    - Inheritance (extends)
    - Interface implementation (implements)
    - Method calls to other classes
    
    Args:
        output_dir: JADX output directory
        max_classes: Maximum number of classes to include (for performance)
    
    Returns:
        Dictionary with nodes (classes) and edges (dependencies) for visualization
    """
    import re
    from collections import defaultdict
    
    output_dir = Path(output_dir)
    sources_dir = output_dir / "sources"
    
    if not sources_dir.exists():
        return {"error": "Decompiled sources not found"}
    
    # Data structures
    nodes = []
    edges = []
    classes_info = {}
    package_stats = defaultdict(int)
    
    # First pass: collect all class names
    all_classes = set()
    for java_file in sources_dir.rglob("*.java"):
        rel_path = str(java_file.relative_to(sources_dir))
        class_name = java_file.stem
        all_classes.add(class_name)
    
    # Second pass: analyze dependencies
    processed = 0
    for java_file in sources_dir.rglob("*.java"):
        if processed >= max_classes:
            break
            
        try:
            source_code = java_file.read_text(encoding='utf-8', errors='ignore')
            rel_path = str(java_file.relative_to(sources_dir))
            class_name = java_file.stem
            
            # Extract package
            package_match = re.search(r'package\s+([\w.]+)\s*;', source_code)
            package_name = package_match.group(1) if package_match else "default"
            full_class_name = f"{package_name}.{class_name}"
            
            # Track package statistics
            base_pkg = '.'.join(package_name.split('.')[:3])
            package_stats[base_pkg] += 1
            
            # Determine class type
            is_activity = bool(re.search(r'extends\s+\w*Activity', source_code))
            is_service = bool(re.search(r'extends\s+\w*Service', source_code))
            is_receiver = bool(re.search(r'extends\s+\w*Receiver', source_code))
            is_provider = bool(re.search(r'extends\s+\w*Provider', source_code))
            is_fragment = bool(re.search(r'extends\s+\w*Fragment', source_code))
            is_adapter = bool(re.search(r'extends\s+\w*Adapter', source_code))
            is_interface = bool(re.search(r'^\s*(?:public\s+)?interface\s+', source_code, re.MULTILINE))
            is_abstract = bool(re.search(r'^\s*(?:public\s+)?abstract\s+class', source_code, re.MULTILINE))
            
            # Determine node type and color
            if is_activity:
                node_type = "activity"
                color = "#4CAF50"  # Green
            elif is_service:
                node_type = "service"
                color = "#2196F3"  # Blue
            elif is_receiver:
                node_type = "receiver"
                color = "#FF9800"  # Orange
            elif is_provider:
                node_type = "provider"
                color = "#9C27B0"  # Purple
            elif is_fragment:
                node_type = "fragment"
                color = "#00BCD4"  # Cyan
            elif is_adapter:
                node_type = "adapter"
                color = "#795548"  # Brown
            elif is_interface:
                node_type = "interface"
                color = "#607D8B"  # Blue Grey
            elif is_abstract:
                node_type = "abstract"
                color = "#9E9E9E"  # Grey
            else:
                node_type = "class"
                color = "#78909C"  # Blue Grey Light
            
            # Count methods and lines
            methods = re.findall(r'(?:public|private|protected)\s+[\w<>\[\]]+\s+(\w+)\s*\(', source_code)
            line_count = source_code.count('\n')
            
            # Store class info
            classes_info[class_name] = {
                "full_name": full_class_name,
                "package": package_name,
                "type": node_type,
                "file_path": rel_path,
            }
            
            # Create node
            nodes.append({
                "id": class_name,
                "label": class_name,
                "full_name": full_class_name,
                "package": package_name,
                "type": node_type,
                "color": color,
                "size": min(30, 10 + len(methods)),  # Size based on methods
                "methods": len(methods),
                "lines": line_count,
                "file_path": rel_path,
            })
            
            # Extract imports
            imports = re.findall(r'import\s+([\w.]+);', source_code)
            for imp in imports:
                imported_class = imp.split('.')[-1]
                if imported_class in all_classes and imported_class != class_name:
                    edges.append({
                        "from": class_name,
                        "to": imported_class,
                        "type": "imports",
                        "color": "#90A4AE",
                        "dashes": True,
                    })
            
            # Extract extends
            extends_match = re.search(r'extends\s+([\w.]+)', source_code)
            if extends_match:
                parent = extends_match.group(1).split('.')[-1]
                if parent in all_classes:
                    edges.append({
                        "from": class_name,
                        "to": parent,
                        "type": "extends",
                        "color": "#4CAF50",
                        "width": 3,
                    })
            
            # Extract implements
            implements_match = re.search(r'implements\s+([\w.,\s]+)(?:\s*\{)', source_code)
            if implements_match:
                interfaces = [i.strip().split('.')[-1] for i in implements_match.group(1).split(',')]
                for iface in interfaces:
                    if iface in all_classes:
                        edges.append({
                            "from": class_name,
                            "to": iface,
                            "type": "implements",
                            "color": "#2196F3",
                            "dashes": [5, 5],
                        })
            
            # Extract method calls to other classes (simplified)
            # Look for patterns like ClassName.method() or new ClassName()
            method_calls = re.findall(r'(?:new\s+|(\w+)\.)\s*(\w+)\s*\(', source_code)
            called_classes = set()
            for caller, method in method_calls:
                if caller and caller[0].isupper() and caller in all_classes and caller != class_name:
                    called_classes.add(caller)
            
            for called in called_classes:
                edges.append({
                    "from": class_name,
                    "to": called,
                    "type": "calls",
                    "color": "#FF5722",
                    "dashes": [2, 2],
                    "width": 1,
                })
            
            processed += 1
            
        except Exception as e:
            continue
    
    # Calculate statistics
    edge_type_counts = defaultdict(int)
    for edge in edges:
        edge_type_counts[edge["type"]] += 1
    
    node_type_counts = defaultdict(int)
    for node in nodes:
        node_type_counts[node["type"]] += 1
    
    # Find hub classes (most connections)
    connection_counts = defaultdict(int)
    for edge in edges:
        connection_counts[edge["from"]] += 1
        connection_counts[edge["to"]] += 1
    
    hub_classes = sorted(connection_counts.items(), key=lambda x: x[1], reverse=True)[:10]
    
    return {
        "nodes": nodes,
        "edges": edges,
        "statistics": {
            "total_classes": len(nodes),
            "total_connections": len(edges),
            "node_types": dict(node_type_counts),
            "edge_types": dict(edge_type_counts),
            "packages": dict(sorted(package_stats.items(), key=lambda x: x[1], reverse=True)[:10]),
            "hub_classes": [{"name": name, "connections": count} for name, count in hub_classes],
        },
        "legend": {
            "node_colors": {
                "activity": "#4CAF50",
                "service": "#2196F3",
                "receiver": "#FF9800",
                "provider": "#9C27B0",
                "fragment": "#00BCD4",
                "adapter": "#795548",
                "interface": "#607D8B",
                "abstract": "#9E9E9E",
                "class": "#78909C",
            },
            "edge_types": {
                "extends": "Inheritance (solid green)",
                "implements": "Interface (dashed blue)",
                "imports": "Import (dashed grey)",
                "calls": "Method call (dotted orange)",
            }
        }
    }