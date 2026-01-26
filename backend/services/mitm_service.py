"""
Man-in-the-Middle Workbench Service
Intercept, inspect, and modify network traffic between application components.

Features:
- TCP/HTTP/HTTPS proxy with traffic interception
- WebSocket deep inspection with frame-level analysis
- Certificate generation and management for HTTPS MITM
- Rule-based traffic modification
- AI-powered traffic analysis
"""

import asyncio
import socket
import ssl
import threading
import time
import uuid
import json
import re
import gzip
import zlib
import struct
import hashlib
import base64
import ast
import operator
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path
from urllib.parse import urlsplit, parse_qs, urlencode
from typing import Dict, List, Optional, Any, Callable, Tuple


def _safe_evaluate_condition(condition: str, variables: Dict[str, Any]) -> bool:
    """
    Safely evaluate a simple condition expression without using eval().

    Supports:
    - Variable references (looked up in variables dict)
    - String/number literals
    - Comparison operators: ==, !=, <, >, <=, >=
    - Logical operators: and, or, not
    - Membership: in, not in
    - Truthiness checks for single variables

    Returns False if the condition cannot be safely parsed.
    """
    if not condition or not condition.strip():
        return True

    condition = condition.strip()

    # Define safe operators
    safe_operators = {
        ast.Eq: operator.eq,
        ast.NotEq: operator.ne,
        ast.Lt: operator.lt,
        ast.LtE: operator.le,
        ast.Gt: operator.gt,
        ast.GtE: operator.ge,
        ast.In: lambda a, b: a in b,
        ast.NotIn: lambda a, b: a not in b,
        ast.And: lambda a, b: a and b,
        ast.Or: lambda a, b: a or b,
        ast.Not: operator.not_,
    }

    def _resolve_value(node: ast.AST) -> Any:
        """Resolve an AST node to a Python value."""
        if isinstance(node, ast.Constant):
            return node.value
        elif isinstance(node, ast.Str):  # Python 3.7 compatibility
            return node.s
        elif isinstance(node, ast.Num):  # Python 3.7 compatibility
            return node.n
        elif isinstance(node, ast.Name):
            # Look up variable name
            var_name = node.id
            if var_name in ('True', 'true'):
                return True
            elif var_name in ('False', 'false'):
                return False
            elif var_name in ('None', 'null'):
                return None
            return variables.get(var_name, '')
        elif isinstance(node, ast.List):
            return [_resolve_value(elt) for elt in node.elts]
        elif isinstance(node, ast.Tuple):
            return tuple(_resolve_value(elt) for elt in node.elts)
        else:
            raise ValueError(f"Unsupported node type: {type(node).__name__}")

    def _evaluate_node(node: ast.AST) -> Any:
        """Recursively evaluate an AST node."""
        if isinstance(node, ast.Expression):
            return _evaluate_node(node.body)
        elif isinstance(node, ast.BoolOp):
            # Handle 'and' / 'or'
            op_func = safe_operators.get(type(node.op))
            if not op_func:
                raise ValueError(f"Unsupported boolean operator: {type(node.op).__name__}")
            result = _evaluate_node(node.values[0])
            for value in node.values[1:]:
                result = op_func(result, _evaluate_node(value))
            return result
        elif isinstance(node, ast.UnaryOp):
            # Handle 'not'
            if isinstance(node.op, ast.Not):
                return not _evaluate_node(node.operand)
            raise ValueError(f"Unsupported unary operator: {type(node.op).__name__}")
        elif isinstance(node, ast.Compare):
            # Handle comparisons like a == b, x in y, etc.
            left = _resolve_value(node.left)
            for op, comparator in zip(node.ops, node.comparators):
                op_func = safe_operators.get(type(op))
                if not op_func:
                    raise ValueError(f"Unsupported comparison operator: {type(op).__name__}")
                right = _resolve_value(comparator)
                if not op_func(left, right):
                    return False
                left = right
            return True
        elif isinstance(node, (ast.Constant, ast.Str, ast.Num, ast.Name, ast.List, ast.Tuple)):
            # Single value - check truthiness
            return bool(_resolve_value(node))
        else:
            raise ValueError(f"Unsupported expression type: {type(node).__name__}")

    try:
        tree = ast.parse(condition, mode='eval')
        return bool(_evaluate_node(tree))
    except (SyntaxError, ValueError, TypeError, KeyError) as e:
        # Log the error and return False for safety
        import logging
        logging.getLogger(__name__).warning(
            f"Failed to evaluate condition '{condition}': {e}"
        )
        return False
from dataclasses import dataclass, field, asdict
from enum import Enum
from collections import defaultdict
import logging
from backend.core.mitm_ws_manager import mitm_stream_manager

logger = logging.getLogger(__name__)


# ============================================================================
# WebSocket Frame Constants and Types
# ============================================================================

class WebSocketOpcode(int, Enum):
    """WebSocket frame opcodes"""
    CONTINUATION = 0x0
    TEXT = 0x1
    BINARY = 0x2
    CLOSE = 0x8
    PING = 0x9
    PONG = 0xA


@dataclass
class WebSocketFrame:
    """Parsed WebSocket frame"""
    id: str
    timestamp: datetime
    direction: str  # "client_to_server" or "server_to_client"
    opcode: int
    opcode_name: str
    fin: bool
    masked: bool
    payload_length: int
    payload: bytes
    payload_text: Optional[str] = None
    payload_json: Optional[Any] = None
    is_control: bool = False
    connection_id: str = ""
    
    # For modification
    modified: bool = False
    original_payload: Optional[bytes] = None


@dataclass
class WebSocketConnection:
    """Tracks a WebSocket connection for deep inspection"""
    id: str
    proxy_id: str
    created_at: datetime
    client_ip: str
    client_port: int
    target_host: str
    target_port: int
    upgrade_request_id: str
    status: str = "active"  # active, closed
    frames: List[WebSocketFrame] = field(default_factory=list)
    total_frames: int = 0
    bytes_sent: int = 0
    bytes_received: int = 0
    closed_at: Optional[datetime] = None
    close_code: Optional[int] = None
    close_reason: Optional[str] = None


@dataclass  
class WebSocketRule:
    """Rule for WebSocket frame modification"""
    id: str
    name: str
    enabled: bool = True
    priority: int = 100
    
    # Match conditions
    match_direction: str = "both"  # "client_to_server", "server_to_client", "both"
    match_opcode: Optional[int] = None  # TEXT, BINARY, etc.
    match_payload_pattern: Optional[str] = None  # Regex pattern
    match_json_path: Optional[str] = None  # JSON path for matching
    
    # Actions
    action: str = "modify"  # "modify", "drop", "delay"
    
    # Modifications
    payload_find_replace: Optional[Dict[str, str]] = None
    json_path_edits: Optional[List[Dict[str, Any]]] = None
    delay_ms: int = 0
    
    # Stats
    hit_count: int = 0


# ============================================================================
# Certificate Management
# ============================================================================

@dataclass
class CACertificate:
    """Root CA certificate for MITM HTTPS interception"""
    serial: int
    common_name: str
    organization: str
    country: str
    valid_from: datetime
    valid_until: datetime
    private_key_pem: str
    certificate_pem: str
    fingerprint_sha256: str
    created_at: datetime


@dataclass
class HostCertificate:
    """Generated certificate for a specific host"""
    hostname: str
    serial: int
    valid_from: datetime
    valid_until: datetime
    private_key_pem: str
    certificate_pem: str
    fingerprint_sha256: str
    created_at: datetime
    ca_fingerprint: str  # Links to the CA that signed this


class CertificateManager:
    """Manages CA and host certificates for HTTPS MITM interception"""
    
    def __init__(self, storage_dir: Path):
        self.storage_dir = storage_dir
        self.ca_dir = storage_dir / "ca"
        self.hosts_dir = storage_dir / "hosts"
        self.ca_dir.mkdir(parents=True, exist_ok=True)
        self.hosts_dir.mkdir(parents=True, exist_ok=True)
        
        self._ca_cert: Optional[CACertificate] = None
        self._host_certs: Dict[str, HostCertificate] = {}
        self._lock = threading.Lock()
        
        # Load existing CA if present
        self._load_ca()
    
    def _load_ca(self) -> None:
        """Load existing CA certificate from disk"""
        ca_meta_path = self.ca_dir / "ca_meta.json"
        ca_key_path = self.ca_dir / "ca_key.pem"
        ca_cert_path = self.ca_dir / "ca_cert.pem"
        
        if ca_meta_path.exists() and ca_key_path.exists() and ca_cert_path.exists():
            try:
                meta = json.loads(ca_meta_path.read_text())
                private_key_pem = ca_key_path.read_text()
                certificate_pem = ca_cert_path.read_text()
                
                self._ca_cert = CACertificate(
                    serial=meta["serial"],
                    common_name=meta["common_name"],
                    organization=meta["organization"],
                    country=meta["country"],
                    valid_from=datetime.fromisoformat(meta["valid_from"]),
                    valid_until=datetime.fromisoformat(meta["valid_until"]),
                    private_key_pem=private_key_pem,
                    certificate_pem=certificate_pem,
                    fingerprint_sha256=meta["fingerprint_sha256"],
                    created_at=datetime.fromisoformat(meta["created_at"])
                )
                logger.info(f"Loaded existing CA certificate: {self._ca_cert.common_name}")
            except Exception as e:
                logger.warning(f"Failed to load CA certificate: {e}")
    
    def get_ca_certificate(self) -> Optional[CACertificate]:
        """Get the current CA certificate"""
        return self._ca_cert
    
    def generate_ca_certificate(
        self,
        common_name: str = "VRAgent MITM CA",
        organization: str = "VRAgent Security",
        country: str = "US",
        validity_days: int = 3650  # 10 years
    ) -> CACertificate:
        """Generate a new root CA certificate for MITM interception"""
        try:
            from cryptography import x509
            from cryptography.x509.oid import NameOID
            from cryptography.hazmat.primitives import hashes, serialization
            from cryptography.hazmat.primitives.asymmetric import rsa
            from cryptography.hazmat.backends import default_backend
        except ImportError:
            raise RuntimeError("cryptography library required. Install with: pip install cryptography")
        
        with self._lock:
            # Generate RSA key pair
            private_key = rsa.generate_private_key(
                public_exponent=65537,
                key_size=4096,
                backend=default_backend()
            )
            
            # Certificate details
            serial = x509.random_serial_number()
            valid_from = datetime.utcnow()
            valid_until = valid_from + timedelta(days=validity_days)
            
            # Build certificate
            subject = issuer = x509.Name([
                x509.NameAttribute(NameOID.COUNTRY_NAME, country),
                x509.NameAttribute(NameOID.ORGANIZATION_NAME, organization),
                x509.NameAttribute(NameOID.COMMON_NAME, common_name),
            ])
            
            cert_builder = x509.CertificateBuilder()
            cert_builder = cert_builder.subject_name(subject)
            cert_builder = cert_builder.issuer_name(issuer)
            cert_builder = cert_builder.public_key(private_key.public_key())
            cert_builder = cert_builder.serial_number(serial)
            cert_builder = cert_builder.not_valid_before(valid_from)
            cert_builder = cert_builder.not_valid_after(valid_until)
            
            # CA extensions
            cert_builder = cert_builder.add_extension(
                x509.BasicConstraints(ca=True, path_length=0),
                critical=True
            )
            cert_builder = cert_builder.add_extension(
                x509.KeyUsage(
                    digital_signature=True,
                    key_encipherment=False,
                    content_commitment=False,
                    data_encipherment=False,
                    key_agreement=False,
                    key_cert_sign=True,
                    crl_sign=True,
                    encipher_only=False,
                    decipher_only=False
                ),
                critical=True
            )
            cert_builder = cert_builder.add_extension(
                x509.SubjectKeyIdentifier.from_public_key(private_key.public_key()),
                critical=False
            )
            
            # Sign the certificate
            certificate = cert_builder.sign(private_key, hashes.SHA256(), default_backend())
            
            # Serialize to PEM
            private_key_pem = private_key.private_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PrivateFormat.TraditionalOpenSSL,
                encryption_algorithm=serialization.NoEncryption()
            ).decode('utf-8')
            
            certificate_pem = certificate.public_bytes(
                serialization.Encoding.PEM
            ).decode('utf-8')
            
            # Calculate fingerprint
            fingerprint = hashlib.sha256(
                certificate.public_bytes(serialization.Encoding.DER)
            ).hexdigest().upper()
            fingerprint_formatted = ':'.join(fingerprint[i:i+2] for i in range(0, len(fingerprint), 2))
            
            # Create CA certificate object
            ca_cert = CACertificate(
                serial=serial,
                common_name=common_name,
                organization=organization,
                country=country,
                valid_from=valid_from,
                valid_until=valid_until,
                private_key_pem=private_key_pem,
                certificate_pem=certificate_pem,
                fingerprint_sha256=fingerprint_formatted,
                created_at=datetime.utcnow()
            )
            
            # Save to disk
            self._save_ca(ca_cert)
            self._ca_cert = ca_cert
            
            # Clear cached host certificates (they need to be re-signed)
            self._host_certs.clear()
            for host_file in self.hosts_dir.glob("*.json"):
                host_file.unlink()
            
            logger.info(f"Generated new CA certificate: {common_name}")
            return ca_cert
    
    def _save_ca(self, ca_cert: CACertificate) -> None:
        """Save CA certificate to disk"""
        meta = {
            "serial": ca_cert.serial,
            "common_name": ca_cert.common_name,
            "organization": ca_cert.organization,
            "country": ca_cert.country,
            "valid_from": ca_cert.valid_from.isoformat(),
            "valid_until": ca_cert.valid_until.isoformat(),
            "fingerprint_sha256": ca_cert.fingerprint_sha256,
            "created_at": ca_cert.created_at.isoformat()
        }
        
        (self.ca_dir / "ca_meta.json").write_text(json.dumps(meta, indent=2))
        (self.ca_dir / "ca_key.pem").write_text(ca_cert.private_key_pem)
        (self.ca_dir / "ca_cert.pem").write_text(ca_cert.certificate_pem)
    
    def generate_host_certificate(self, hostname: str, validity_days: int = 365) -> HostCertificate:
        """Generate a certificate for a specific hostname, signed by our CA"""
        if not self._ca_cert:
            raise ValueError("No CA certificate. Generate a CA certificate first.")
        
        # Check cache
        if hostname in self._host_certs:
            cached = self._host_certs[hostname]
            if cached.valid_until > datetime.utcnow():
                return cached
        
        try:
            from cryptography import x509
            from cryptography.x509.oid import NameOID
            from cryptography.hazmat.primitives import hashes, serialization
            from cryptography.hazmat.primitives.asymmetric import rsa
            from cryptography.hazmat.backends import default_backend
        except ImportError:
            raise RuntimeError("cryptography library required")
        
        with self._lock:
            # Load CA private key
            ca_key = serialization.load_pem_private_key(
                self._ca_cert.private_key_pem.encode(),
                password=None,
                backend=default_backend()
            )
            
            ca_cert = x509.load_pem_x509_certificate(
                self._ca_cert.certificate_pem.encode(),
                default_backend()
            )
            
            # Generate host key pair
            host_key = rsa.generate_private_key(
                public_exponent=65537,
                key_size=2048,
                backend=default_backend()
            )
            
            serial = x509.random_serial_number()
            valid_from = datetime.utcnow()
            valid_until = valid_from + timedelta(days=validity_days)
            
            # Build certificate
            subject = x509.Name([
                x509.NameAttribute(NameOID.COMMON_NAME, hostname),
            ])
            
            cert_builder = x509.CertificateBuilder()
            cert_builder = cert_builder.subject_name(subject)
            cert_builder = cert_builder.issuer_name(ca_cert.subject)
            cert_builder = cert_builder.public_key(host_key.public_key())
            cert_builder = cert_builder.serial_number(serial)
            cert_builder = cert_builder.not_valid_before(valid_from)
            cert_builder = cert_builder.not_valid_after(valid_until)
            
            # Add SAN extension for the hostname
            san_list = [x509.DNSName(hostname)]
            
            # Add wildcard if it's a simple hostname
            if '.' in hostname and not hostname.startswith('*.'):
                parts = hostname.split('.')
                if len(parts) >= 2:
                    wildcard = '*.' + '.'.join(parts[1:])
                    san_list.append(x509.DNSName(wildcard))
            
            # Also add IP if it looks like an IP address
            try:
                import ipaddress
                ip = ipaddress.ip_address(hostname)
                san_list.append(x509.IPAddress(ip))
            except ValueError:
                pass
            
            cert_builder = cert_builder.add_extension(
                x509.SubjectAlternativeName(san_list),
                critical=False
            )
            
            cert_builder = cert_builder.add_extension(
                x509.BasicConstraints(ca=False, path_length=None),
                critical=True
            )
            
            cert_builder = cert_builder.add_extension(
                x509.KeyUsage(
                    digital_signature=True,
                    key_encipherment=True,
                    content_commitment=False,
                    data_encipherment=False,
                    key_agreement=False,
                    key_cert_sign=False,
                    crl_sign=False,
                    encipher_only=False,
                    decipher_only=False
                ),
                critical=True
            )
            
            cert_builder = cert_builder.add_extension(
                x509.ExtendedKeyUsage([
                    x509.oid.ExtendedKeyUsageOID.SERVER_AUTH,
                    x509.oid.ExtendedKeyUsageOID.CLIENT_AUTH,
                ]),
                critical=False
            )
            
            # Sign with CA key
            certificate = cert_builder.sign(ca_key, hashes.SHA256(), default_backend())
            
            # Serialize
            private_key_pem = host_key.private_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PrivateFormat.TraditionalOpenSSL,
                encryption_algorithm=serialization.NoEncryption()
            ).decode('utf-8')
            
            certificate_pem = certificate.public_bytes(
                serialization.Encoding.PEM
            ).decode('utf-8')
            
            fingerprint = hashlib.sha256(
                certificate.public_bytes(serialization.Encoding.DER)
            ).hexdigest().upper()
            fingerprint_formatted = ':'.join(fingerprint[i:i+2] for i in range(0, len(fingerprint), 2))
            
            host_cert = HostCertificate(
                hostname=hostname,
                serial=serial,
                valid_from=valid_from,
                valid_until=valid_until,
                private_key_pem=private_key_pem,
                certificate_pem=certificate_pem,
                fingerprint_sha256=fingerprint_formatted,
                created_at=datetime.utcnow(),
                ca_fingerprint=self._ca_cert.fingerprint_sha256
            )
            
            # Cache and save
            self._host_certs[hostname] = host_cert
            self._save_host_cert(host_cert)
            
            logger.info(f"Generated certificate for host: {hostname}")
            return host_cert
    
    def _save_host_cert(self, cert: HostCertificate) -> None:
        """Save host certificate to disk"""
        safe_name = re.sub(r'[^a-zA-Z0-9.-]', '_', cert.hostname)
        meta_path = self.hosts_dir / f"{safe_name}.json"
        key_path = self.hosts_dir / f"{safe_name}.key.pem"
        cert_path = self.hosts_dir / f"{safe_name}.cert.pem"
        
        meta = {
            "hostname": cert.hostname,
            "serial": cert.serial,
            "valid_from": cert.valid_from.isoformat(),
            "valid_until": cert.valid_until.isoformat(),
            "fingerprint_sha256": cert.fingerprint_sha256,
            "created_at": cert.created_at.isoformat(),
            "ca_fingerprint": cert.ca_fingerprint
        }
        
        meta_path.write_text(json.dumps(meta, indent=2))
        key_path.write_text(cert.private_key_pem)
        cert_path.write_text(cert.certificate_pem)
    
    def get_host_certificate(self, hostname: str) -> Optional[HostCertificate]:
        """Get an existing host certificate, or generate if not exists"""
        if hostname in self._host_certs:
            return self._host_certs[hostname]
        
        # Try to load from disk
        safe_name = re.sub(r'[^a-zA-Z0-9.-]', '_', hostname)
        meta_path = self.hosts_dir / f"{safe_name}.json"
        
        if meta_path.exists():
            try:
                meta = json.loads(meta_path.read_text())
                key_pem = (self.hosts_dir / f"{safe_name}.key.pem").read_text()
                cert_pem = (self.hosts_dir / f"{safe_name}.cert.pem").read_text()
                
                cert = HostCertificate(
                    hostname=meta["hostname"],
                    serial=meta["serial"],
                    valid_from=datetime.fromisoformat(meta["valid_from"]),
                    valid_until=datetime.fromisoformat(meta["valid_until"]),
                    private_key_pem=key_pem,
                    certificate_pem=cert_pem,
                    fingerprint_sha256=meta["fingerprint_sha256"],
                    created_at=datetime.fromisoformat(meta["created_at"]),
                    ca_fingerprint=meta["ca_fingerprint"]
                )
                
                # Check if still valid and signed by current CA
                if cert.valid_until > datetime.utcnow():
                    if self._ca_cert and cert.ca_fingerprint == self._ca_cert.fingerprint_sha256:
                        self._host_certs[hostname] = cert
                        return cert
            except Exception as e:
                logger.warning(f"Failed to load host cert for {hostname}: {e}")
        
        # Generate new certificate
        if self._ca_cert:
            return self.generate_host_certificate(hostname)
        
        return None
    
    def list_host_certificates(self) -> List[Dict[str, Any]]:
        """List all generated host certificates"""
        certs = []
        for meta_path in self.hosts_dir.glob("*.json"):
            try:
                meta = json.loads(meta_path.read_text())
                certs.append({
                    "hostname": meta["hostname"],
                    "valid_from": meta["valid_from"],
                    "valid_until": meta["valid_until"],
                    "fingerprint": meta["fingerprint_sha256"],
                    "created_at": meta["created_at"]
                })
            except Exception:
                continue
        return sorted(certs, key=lambda x: x["hostname"])
    
    def delete_host_certificate(self, hostname: str) -> bool:
        """Delete a host certificate"""
        safe_name = re.sub(r'[^a-zA-Z0-9.-]', '_', hostname)
        deleted = False
        
        for suffix in [".json", ".key.pem", ".cert.pem"]:
            path = self.hosts_dir / f"{safe_name}{suffix}"
            if path.exists():
                path.unlink()
                deleted = True
        
        self._host_certs.pop(hostname, None)
        return deleted
    
    def get_ssl_context_for_host(self, hostname: str) -> Optional[ssl.SSLContext]:
        """Get an SSL context configured for the given hostname"""
        host_cert = self.get_host_certificate(hostname)
        if not host_cert:
            return None
        
        # Create SSL context
        context = ssl.SSLContext(ssl.PROTOCOL_TLS_SERVER)
        
        # Write temp files for SSL context (it requires files)
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w', suffix='.pem', delete=False) as key_file:
            key_file.write(host_cert.private_key_pem)
            key_path = key_file.name
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.pem', delete=False) as cert_file:
            cert_file.write(host_cert.certificate_pem)
            cert_path = cert_file.name
        
        try:
            context.load_cert_chain(cert_path, key_path)
        finally:
            Path(key_path).unlink(missing_ok=True)
            Path(cert_path).unlink(missing_ok=True)
        
        return context
    
    def get_installation_instructions(self) -> Dict[str, Any]:
        """Get instructions for installing the CA certificate on various platforms"""
        if not self._ca_cert:
            return {"error": "No CA certificate generated yet"}
        
        return {
            "ca_certificate": {
                "common_name": self._ca_cert.common_name,
                "fingerprint": self._ca_cert.fingerprint_sha256,
                "valid_until": self._ca_cert.valid_until.isoformat(),
                "pem": self._ca_cert.certificate_pem
            },
            "instructions": {
                "windows": {
                    "title": "Windows Installation",
                    "steps": [
                        "1. Download the CA certificate (ca_cert.pem)",
                        "2. Rename to ca_cert.crt",
                        "3. Double-click the file",
                        "4. Click 'Install Certificate'",
                        "5. Select 'Local Machine' and click Next",
                        "6. Select 'Place all certificates in the following store'",
                        "7. Click Browse and select 'Trusted Root Certification Authorities'",
                        "8. Click Next, then Finish"
                    ],
                    "command": "certutil -addstore Root ca_cert.crt"
                },
                "macos": {
                    "title": "macOS Installation",
                    "steps": [
                        "1. Download the CA certificate",
                        "2. Double-click to open in Keychain Access",
                        "3. Add to 'System' keychain",
                        "4. Double-click the certificate",
                        "5. Expand 'Trust' section",
                        "6. Set 'When using this certificate' to 'Always Trust'",
                        "7. Close and authenticate"
                    ],
                    "command": "sudo security add-trusted-cert -d -r trustRoot -k /Library/Keychains/System.keychain ca_cert.pem"
                },
                "linux": {
                    "title": "Linux Installation (Ubuntu/Debian)",
                    "steps": [
                        "1. Download the CA certificate",
                        "2. Copy to /usr/local/share/ca-certificates/",
                        "3. Run update-ca-certificates"
                    ],
                    "command": "sudo cp ca_cert.pem /usr/local/share/ca-certificates/vragent-mitm.crt && sudo update-ca-certificates"
                },
                "firefox": {
                    "title": "Firefox (uses its own cert store)",
                    "steps": [
                        "1. Open Firefox Settings",
                        "2. Search for 'Certificates'",
                        "3. Click 'View Certificates'",
                        "4. Go to 'Authorities' tab",
                        "5. Click 'Import'",
                        "6. Select the CA certificate file",
                        "7. Check 'Trust this CA to identify websites'",
                        "8. Click OK"
                    ]
                },
                "android": {
                    "title": "Android Installation",
                    "steps": [
                        "1. Transfer certificate to device",
                        "2. Go to Settings > Security > Encryption & Credentials",
                        "3. Tap 'Install a certificate'",
                        "4. Select 'CA certificate'",
                        "5. Select the certificate file",
                        "6. Confirm installation"
                    ],
                    "note": "Android 7+ may require additional steps for apps targeting SDK 24+"
                },
                "ios": {
                    "title": "iOS Installation",
                    "steps": [
                        "1. Email or AirDrop the certificate to your device",
                        "2. Open the certificate file",
                        "3. Go to Settings > General > Profile",
                        "4. Tap on the profile to install",
                        "5. Go to Settings > General > About > Certificate Trust Settings",
                        "6. Enable full trust for the certificate"
                    ]
                }
            }
        }


class InterceptionMode(str, Enum):
    PASSTHROUGH = "passthrough"  # Just observe, don't modify
    INTERCEPT = "intercept"      # Hold for manual inspection/modification
    AUTO_MODIFY = "auto_modify"  # Apply rules automatically


class Protocol(str, Enum):
    HTTP = "http"
    HTTPS = "https"
    TCP = "tcp"
    WEBSOCKET = "websocket"


@dataclass
class InterceptionRule:
    """Rule for automatic traffic modification"""
    id: str
    name: str
    enabled: bool = True
    priority: int = 100
    group: Optional[str] = None
    
    # Match conditions
    match_host: Optional[str] = None  # Regex pattern
    match_path: Optional[str] = None  # Regex pattern
    match_method: Optional[str] = None
    match_content_type: Optional[str] = None
    match_body: Optional[str] = None  # Regex pattern
    match_header: Optional[Dict[str, str]] = None
    match_status_code: Optional[int] = None
    match_direction: str = "both"  # "request", "response", "both"
    match_query: Optional[Dict[str, str]] = None  # Query param patterns
    
    # Actions
    action: str = "modify"  # "modify", "drop", "delay", "replace"
    
    # Modification actions
    modify_headers: Optional[Dict[str, str]] = None  # Headers to add/replace
    remove_headers: Optional[List[str]] = None
    modify_body: Optional[str] = None  # New body content
    body_find_replace: Optional[Dict[str, str]] = None  # Find/replace in body
    body_find_replace_regex: bool = False
    json_path_edits: Optional[List[Dict[str, Any]]] = None
    modify_status_code: Optional[int] = None
    modify_path: Optional[str] = None
    
    # Delay action
    delay_ms: int = 0
    
    # Stats
    hit_count: int = 0


@dataclass
class InterceptedRequest:
    """Captured HTTP request"""
    id: str
    timestamp: datetime
    client_ip: str
    client_port: int
    
    # Request details
    method: str
    url: str
    path: str
    host: str
    port: int
    protocol: Protocol
    http_version: str
    headers: Dict[str, str]
    body: Optional[bytes] = None
    body_text: Optional[str] = None
    
    # State
    intercepted: bool = False
    modified: bool = False
    forwarded: bool = False
    
    # Timing
    received_at: float = 0
    forwarded_at: float = 0


@dataclass
class InterceptedResponse:
    """Captured HTTP response"""
    id: str
    request_id: str
    timestamp: datetime
    
    # Response details
    status_code: int
    status_message: str
    http_version: str
    headers: Dict[str, str]
    body: Optional[bytes] = None
    body_text: Optional[str] = None
    content_length: int = 0
    
    # State
    intercepted: bool = False
    modified: bool = False
    forwarded: bool = False
    
    # Timing
    received_at: float = 0
    forwarded_at: float = 0
    response_time_ms: float = 0


@dataclass
class TrafficEntry:
    """Complete request/response pair"""
    id: str
    request: InterceptedRequest
    response: Optional[InterceptedResponse] = None
    error: Optional[str] = None
    tags: List[str] = field(default_factory=list)
    notes: str = ""
    rules_applied: List[str] = field(default_factory=list)


class TrafficStore:
    """Disk-backed storage for MITM traffic with retention and sessions."""

    def __init__(self, root_dir: Path, max_entries: int = 5000):
        self.root_dir = root_dir
        self.max_entries = max_entries
        self._lock = threading.Lock()
        self._counts: Dict[str, int] = {}
        self.root_dir.mkdir(parents=True, exist_ok=True)

    def _proxy_dir(self, proxy_id: str) -> Path:
        safe_id = re.sub(r"[^a-zA-Z0-9_.-]", "_", proxy_id)
        return self.root_dir / safe_id

    def _traffic_file(self, proxy_id: str) -> Path:
        return self._proxy_dir(proxy_id) / "traffic.jsonl"

    def _sessions_dir(self, proxy_id: str) -> Path:
        return self._proxy_dir(proxy_id) / "sessions"

    def _load_count(self, proxy_id: str) -> int:
        path = self._traffic_file(proxy_id)
        if not path.exists():
            return 0
        with path.open("r", encoding="utf-8") as handle:
            return sum(1 for _ in handle)

    def append_entry(self, proxy_id: str, entry: Dict[str, Any]) -> None:
        path = self._traffic_file(proxy_id)
        path.parent.mkdir(parents=True, exist_ok=True)
        line = json.dumps(entry, ensure_ascii=True)
        with self._lock:
            with path.open("a", encoding="utf-8") as handle:
                handle.write(line + "\n")
            current = self._counts.get(proxy_id)
            if current is None:
                current = self._load_count(proxy_id)
            current += 1
            self._counts[proxy_id] = current
            if current > self.max_entries:
                self._trim_file(proxy_id)

    def list_entries(self, proxy_id: str, limit: int = 100, offset: int = 0) -> List[Dict[str, Any]]:
        path = self._traffic_file(proxy_id)
        if not path.exists():
            return []
        with path.open("r", encoding="utf-8") as handle:
            lines = handle.readlines()
        total = len(lines)
        if offset > 0:
            start = max(total - offset - limit, 0)
            end = total - offset
        else:
            start = max(total - limit, 0)
            end = total
        slice_lines = lines[start:end]
        entries = []
        for line in slice_lines:
            try:
                entries.append(json.loads(line))
            except json.JSONDecodeError:
                continue
        return entries

    def list_entries_range(self, proxy_id: str, start: int, end: Optional[int] = None) -> List[Dict[str, Any]]:
        path = self._traffic_file(proxy_id)
        if not path.exists():
            return []
        with path.open("r", encoding="utf-8") as handle:
            lines = handle.readlines()
        slice_lines = lines[start:end]
        entries = []
        for line in slice_lines:
            try:
                entries.append(json.loads(line))
            except json.JSONDecodeError:
                continue
        return entries

    def count_entries(self, proxy_id: str) -> int:
        current = self._counts.get(proxy_id)
        if current is None:
            current = self._load_count(proxy_id)
            self._counts[proxy_id] = current
        return current

    def get_entry(self, proxy_id: str, entry_id: str) -> Optional[Dict[str, Any]]:
        path = self._traffic_file(proxy_id)
        if not path.exists():
            return None
        with path.open("r", encoding="utf-8") as handle:
            for line in handle:
                try:
                    data = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if data.get("id") == entry_id:
                    return data
        return None

    def update_entry(self, proxy_id: str, entry_id: str, updated: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        path = self._traffic_file(proxy_id)
        if not path.exists():
            return None
        updated_entry = None
        with self._lock:
            with path.open("r", encoding="utf-8") as handle:
                lines = handle.readlines()
            new_lines = []
            for line in lines:
                try:
                    data = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if data.get("id") == entry_id:
                    data.update(updated)
                    updated_entry = data
                new_lines.append(json.dumps(data, ensure_ascii=True))
            with path.open("w", encoding="utf-8") as handle:
                handle.write("\n".join(new_lines) + ("\n" if new_lines else ""))
        return updated_entry

    def clear_entries(self, proxy_id: str) -> None:
        path = self._traffic_file(proxy_id)
        with self._lock:
            if path.exists():
                path.unlink()
            self._counts[proxy_id] = 0

    def _trim_file(self, proxy_id: str) -> None:
        path = self._traffic_file(proxy_id)
        if not path.exists():
            return
        with path.open("r", encoding="utf-8") as handle:
            lines = handle.readlines()
        if len(lines) <= self.max_entries:
            self._counts[proxy_id] = len(lines)
            return
        trimmed = lines[-self.max_entries:]
        with path.open("w", encoding="utf-8") as handle:
            handle.writelines(trimmed)
        self._counts[proxy_id] = len(trimmed)

    def save_session(self, proxy_id: str, name: Optional[str] = None) -> Dict[str, Any]:
        sessions_dir = self._sessions_dir(proxy_id)
        sessions_dir.mkdir(parents=True, exist_ok=True)
        session_id = f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}"
        safe_name = re.sub(r"[^a-zA-Z0-9_. -]", "", name or "session").strip() or "session"
        session_file = sessions_dir / f"{session_id}.jsonl"
        meta_file = sessions_dir / f"{session_id}.meta.json"

        traffic_file = self._traffic_file(proxy_id)
        entries = []
        if traffic_file.exists():
            with traffic_file.open("r", encoding="utf-8") as handle:
                entries = handle.readlines()
            with session_file.open("w", encoding="utf-8") as handle:
                handle.writelines(entries)

        meta = {
            "id": session_id,
            "name": safe_name,
            "created_at": datetime.utcnow().isoformat() + "Z",
            "entries": len(entries),
        }
        meta_file.write_text(json.dumps(meta, ensure_ascii=True), encoding="utf-8")
        return meta

    def list_sessions(self, proxy_id: str) -> List[Dict[str, Any]]:
        sessions_dir = self._sessions_dir(proxy_id)
        if not sessions_dir.exists():
            return []
        sessions = []
        for meta_path in sorted(sessions_dir.glob("*.meta.json")):
            try:
                sessions.append(json.loads(meta_path.read_text(encoding="utf-8")))
            except json.JSONDecodeError:
                continue
        return sessions

    def load_session_entries(self, proxy_id: str, session_id: str, limit: int = 100, offset: int = 0) -> List[Dict[str, Any]]:
        sessions_dir = self._sessions_dir(proxy_id)
        path = sessions_dir / f"{session_id}.jsonl"
        if not path.exists():
            return []
        with path.open("r", encoding="utf-8") as handle:
            lines = handle.readlines()
        total = len(lines)
        if offset > 0:
            start = max(total - offset - limit, 0)
            end = total - offset
        else:
            start = max(total - limit, 0)
            end = total
        slice_lines = lines[start:end]
        entries = []
        for line in slice_lines:
            try:
                entries.append(json.loads(line))
            except json.JSONDecodeError:
                continue
        return entries

    def get_session_meta(self, proxy_id: str, session_id: str) -> Optional[Dict[str, Any]]:
        sessions_dir = self._sessions_dir(proxy_id)
        meta_path = sessions_dir / f"{session_id}.meta.json"
        if not meta_path.exists():
            return None
        try:
            return json.loads(meta_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return None

    def update_session_meta(self, proxy_id: str, session_id: str, meta: Dict[str, Any]) -> None:
        """Update session metadata with additional info like analysis."""
        sessions_dir = self._sessions_dir(proxy_id)
        meta_path = sessions_dir / f"{session_id}.meta.json"
        if meta_path.exists():
            meta_path.write_text(json.dumps(meta, ensure_ascii=True, default=str), encoding="utf-8")

    def delete_session(self, proxy_id: str, session_id: str) -> None:
        """Delete a saved session and its metadata."""
        sessions_dir = self._sessions_dir(proxy_id)
        session_file = sessions_dir / f"{session_id}.jsonl"
        meta_file = sessions_dir / f"{session_id}.meta.json"
        if session_file.exists():
            session_file.unlink()
        if meta_file.exists():
            meta_file.unlink()


class MITMProxy:
    """TCP/HTTP Proxy for traffic interception"""
    
    def __init__(
        self,
        listen_host: str = "127.0.0.1",
        listen_port: int = 8080,
        target_host: str = "localhost",
        target_port: int = 80,
        mode: InterceptionMode = InterceptionMode.AUTO_MODIFY,  # Default to auto-modify for attack tools
        tls_enabled: bool = False
    ):
        self.listen_host = listen_host
        self.listen_port = listen_port
        self.target_host = target_host
        self.target_port = target_port
        self.mode = mode
        self.tls_enabled = tls_enabled
        
        self.running = False
        self.server_socket: Optional[socket.socket] = None
        self.traffic_log: List[TrafficEntry] = []
        self.rules: List[InterceptionRule] = []
        self.pending_requests: Dict[str, InterceptedRequest] = {}
        self.pending_responses: Dict[str, InterceptedResponse] = {}
        self._traffic_index: Dict[str, TrafficEntry] = {}
        self._connection_requests: Dict[str, List[str]] = {}
        
        # Callbacks for real-time streaming
        self.on_request: Optional[Callable] = None
        self.on_response: Optional[Callable] = None
        self.on_error: Optional[Callable] = None
        self.on_entry: Optional[Callable[[TrafficEntry, str], None]] = None
        
        # Stats
        self.stats = {
            "requests_total": 0,
            "responses_total": 0,
            "bytes_sent": 0,
            "bytes_received": 0,
            "errors": 0,
            "rules_applied": 0,
            "start_time": None
        }
        
        self._lock = threading.Lock()
        self._buffers: Dict[str, Dict[str, bytearray]] = {}
        self._connection_protocols: Dict[str, str] = {}
    
    def start(self):
        """Start the proxy server"""
        if self.running:
            return
        
        self.running = True
        self.stats["start_time"] = time.time()
        
        self.server_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        self.server_socket.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        self.server_socket.bind((self.listen_host, self.listen_port))
        self.server_socket.listen(100)
        self.server_socket.settimeout(1.0)
        
        # Start accept thread
        self._accept_thread = threading.Thread(target=self._accept_connections, daemon=True)
        self._accept_thread.start()
        
        logger.info(f"MITM Proxy started on {self.listen_host}:{self.listen_port} -> {self.target_host}:{self.target_port}")
    
    def stop(self):
        """Stop the proxy server"""
        self.running = False
        if self.server_socket:
            try:
                self.server_socket.close()
            except OSError:
                pass  # Socket already closed
        logger.info("MITM Proxy stopped")
    
    def _accept_connections(self):
        """Accept incoming connections"""
        while self.running:
            try:
                client_socket, client_addr = self.server_socket.accept()
                # Handle each connection in a new thread
                handler = threading.Thread(
                    target=self._handle_connection,
                    args=(client_socket, client_addr),
                    daemon=True
                )
                handler.start()
            except socket.timeout:
                continue
            except Exception as e:
                if self.running:
                    logger.error(f"Accept error: {e}")
    
    def _handle_connection(self, client_socket: socket.socket, client_addr: tuple):
        """Handle a single client connection"""
        target_socket = None
        connection_id = f"{client_addr[0]}:{client_addr[1]}-{uuid.uuid4()}"
        try:
            # Connect to target
            target_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            target_socket.connect((self.target_host, self.target_port))
            
            if self.tls_enabled:
                context = ssl.create_default_context()
                context.check_hostname = False
                context.verify_mode = ssl.CERT_NONE
                target_socket = context.wrap_socket(target_socket, server_hostname=self.target_host)
            
            # Bidirectional proxy
            client_to_target = threading.Thread(
                target=self._proxy_data,
                args=(client_socket, target_socket, client_addr, "request", connection_id),
                daemon=True
            )
            target_to_client = threading.Thread(
                target=self._proxy_data,
                args=(target_socket, client_socket, client_addr, "response", connection_id),
                daemon=True
            )
            
            client_to_target.start()
            target_to_client.start()
            
            client_to_target.join()
            target_to_client.join()
            
        except Exception as e:
            self.stats["errors"] += 1
            if self.on_error:
                self.on_error(str(e), client_addr)
        finally:
            try:
                client_socket.close()
            except:
                pass
            try:
                if target_socket:
                    target_socket.close()
            except:
                pass
            with self._lock:
                self._connection_requests.pop(connection_id, None)
                self._buffers.pop(connection_id, None)
                self._connection_protocols.pop(connection_id, None)
    
    def _proxy_data(self, src: socket.socket, dst: socket.socket, client_addr: tuple, direction: str, connection_id: str):
        """Proxy data between sockets with interception"""
        try:
            buffers = self._buffers.setdefault(connection_id, {"request": bytearray(), "response": bytearray()})
            while self.running:
                src.settimeout(30.0)
                data = src.recv(65536)
                if not data:
                    break

                protocol = self._connection_protocols.get(connection_id)
                if protocol in ("websocket", "http2", "stream"):
                    dst.sendall(data)
                    continue

                if protocol is None and data.startswith(b"PRI * HTTP/2.0"):
                    self._connection_protocols[connection_id] = "http2"
                    dst.sendall(data)
                    continue

                buffers[direction].extend(data)

                while True:
                    message, remaining = self._extract_http_message(buffers[direction])
                    if message is None:
                        buffers[direction] = remaining
                        break
                    buffers[direction] = remaining

                    if direction == "request":
                        modified, entry = self._process_request(message, client_addr, connection_id)
                        self.stats["requests_total"] += 1
                        self.stats["bytes_sent"] += len(modified)
                    else:
                        modified, entry = self._process_response(message, client_addr, connection_id)
                        self.stats["responses_total"] += 1
                        self.stats["bytes_received"] += len(modified)

                    dst.sendall(modified)

        except socket.timeout:
            logger.debug(f"Socket timeout in {direction} for {connection_id}")
        except Exception as e:
            if self.running:
                logger.warning(f"Proxy data error in {direction} for {connection_id}: {type(e).__name__}: {e}")

    def _extract_http_message(self, buffer: bytearray) -> Tuple[Optional[bytes], bytearray]:
        """Extract a complete HTTP message from the buffer if available."""
        header_end = buffer.find(b"\r\n\r\n")
        if header_end == -1:
            return None, buffer

        header_bytes = bytes(buffer[:header_end])
        headers = self._parse_headers(header_bytes)
        body_start = header_end + 4

        transfer_encoding = headers.get("transfer-encoding", "").lower()
        if "chunked" in transfer_encoding:
            body = bytes(buffer[body_start:])
            consumed = self._chunked_length(body)
            if consumed is None:
                return None, buffer
            total_len = body_start + consumed
        else:
            content_length = 0
            if "content-length" in headers:
                try:
                    content_length = int(headers.get("content-length", "0"))
                except ValueError:
                    content_length = 0
            total_len = body_start + content_length
            if len(buffer) < total_len:
                return None, buffer

        message = bytes(buffer[:total_len])
        remaining = bytearray(buffer[total_len:])
        return message, remaining

    def _chunked_length(self, body: bytes) -> Optional[int]:
        """Return total length of a chunked body if complete, otherwise None."""
        idx = 0
        total_len = 0
        while True:
            line_end = body.find(b"\r\n", idx)
            if line_end == -1:
                return None
            size_line = body[idx:line_end].split(b";", 1)[0]
            try:
                size = int(size_line.strip(), 16)
            except ValueError:
                return None
            idx = line_end + 2
            if len(body) < idx + size + 2:
                return None
            idx += size
            if body[idx:idx + 2] != b"\r\n":
                return None
            idx += 2
            total_len = idx
            if size == 0:
                return total_len

    def _decode_chunked_body(self, body: bytes) -> Optional[bytes]:
        """Decode chunked HTTP body; returns None if incomplete."""
        idx = 0
        decoded = bytearray()
        while True:
            line_end = body.find(b"\r\n", idx)
            if line_end == -1:
                return None
            size_line = body[idx:line_end].split(b";", 1)[0]
            try:
                size = int(size_line.strip(), 16)
            except ValueError:
                return None
            idx = line_end + 2
            if len(body) < idx + size + 2:
                return None
            decoded.extend(body[idx:idx + size])
            idx += size
            if body[idx:idx + 2] != b"\r\n":
                return None
            idx += 2
            if size == 0:
                return bytes(decoded)

    def _encode_chunked_body(self, body: bytes, chunk_size: int = 4096) -> bytes:
        """Encode body using chunked transfer encoding."""
        chunks: List[bytes] = []
        for i in range(0, len(body), chunk_size):
            chunk = body[i:i + chunk_size]
            chunks.append(f"{len(chunk):X}\r\n".encode("ascii") + chunk + b"\r\n")
        chunks.append(b"0\r\n\r\n")
        return b"".join(chunks)

    def _parse_headers(self, header_bytes: bytes) -> Dict[str, str]:
        """Parse HTTP headers into a case-insensitive dict."""
        header_text = header_bytes.decode("utf-8", errors="replace")
        lines = header_text.split("\r\n")
        headers: Dict[str, str] = {}
        for line in lines[1:]:
            if ":" not in line:
                continue
            key, value = line.split(":", 1)
            headers[key.strip().lower()] = value.strip()
        return headers
    
    def _process_request(self, data: bytes, client_addr: tuple, connection_id: str) -> tuple:
        """Process and potentially modify a request"""
        entry_id = str(uuid.uuid4())
        
        try:
            # Parse HTTP request
            request = self._parse_http_request(data, client_addr, entry_id)
            header_lower = {k.lower(): v for k, v in request.headers.items()}
            if header_lower.get("upgrade", "").lower() == "websocket":
                self._connection_protocols[connection_id] = "websocket_pending"
            
            # Apply rules
            modified_data = data
            applied_rules: List[str] = []
            if self.mode == InterceptionMode.AUTO_MODIFY:
                modified_data, applied_rules = self._apply_rules_to_request(data, request)
                if applied_rules:
                    request.modified = True
            
            # Create traffic entry
            entry = TrafficEntry(id=entry_id, request=request, rules_applied=applied_rules)
            self._append_entry(entry, connection_id)
            
            # Callback
            if self.on_request:
                self.on_request(request)
            if self.on_entry:
                self.on_entry(entry, "request")
            
            return modified_data, entry
            
        except Exception as e:
            logger.error(f"Request processing error: {e}")
            return data, None
    
    def _process_response(self, data: bytes, client_addr: tuple, connection_id: str) -> tuple:
        """Process and potentially modify a response"""
        try:
            # Parse HTTP response
            response = self._parse_http_response(data, str(uuid.uuid4()))
            
            # Apply rules
            modified_data = data
            applied_rules: List[str] = []
            if self.mode == InterceptionMode.AUTO_MODIFY:
                modified_data, applied_rules = self._apply_rules_to_response(data, response)
                if applied_rules:
                    response.modified = True
            
            # Attach response to pending request
            entry = None
            with self._lock:
                queue = self._connection_requests.get(connection_id, [])
                entry_id = queue.pop(0) if queue else None
                if entry_id:
                    entry = self._traffic_index.get(entry_id)
            if entry:
                response.request_id = entry.request.id
                response.response_time_ms = (response.received_at - entry.request.received_at) * 1000.0
                entry.response = response
                if applied_rules:
                    entry.rules_applied.extend(applied_rules)
                if response.status_code == 101 and self._connection_protocols.get(connection_id) == "websocket_pending":
                    self._connection_protocols[connection_id] = "websocket"
                elif self._connection_protocols.get(connection_id) == "websocket_pending":
                    self._connection_protocols.pop(connection_id, None)

            # Callback
            if self.on_response:
                self.on_response(response)
            if entry and self.on_entry:
                self.on_entry(entry, "response")
            
            return modified_data, entry
            
        except Exception as e:
            logger.error(f"Response processing error: {e}")
            return data, None
    
    def _parse_http_request(self, data: bytes, client_addr: tuple, entry_id: str) -> InterceptedRequest:
        """Parse raw HTTP request data"""
        header_end = data.find(b"\r\n\r\n")
        if header_end == -1:
            header_end = len(data)
        header_bytes = data[:header_end]
        body = data[header_end + 4:] if header_end + 4 <= len(data) else b""

        header_text = header_bytes.decode("utf-8", errors="replace")
        lines = header_text.split("\r\n")
        request_line = lines[0] if lines else ""
        parts = request_line.split(" ")

        method = parts[0] if len(parts) > 0 else "UNKNOWN"
        path = parts[1] if len(parts) > 1 else "/"
        http_version = parts[2] if len(parts) > 2 else "HTTP/1.1"

        headers: Dict[str, str] = {}
        for line in lines[1:]:
            if ":" not in line:
                continue
            key, value = line.split(":", 1)
            headers[key.strip()] = value.strip()

        header_lower = {k.lower(): v for k, v in headers.items()}
        if "chunked" in header_lower.get("transfer-encoding", "").lower():
            decoded = self._decode_chunked_body(body)
            if decoded is not None:
                body = decoded

        content_encoding = header_lower.get("content-encoding", "").lower()
        if content_encoding in ("gzip", "deflate"):
            try:
                if content_encoding == "gzip":
                    body = gzip.decompress(body)
                else:
                    body = zlib.decompress(body)
            except Exception:
                pass

        body_text = None
        if body:
            try:
                body_text = body.decode("utf-8", errors="replace")
            except Exception:
                body_text = None

        host = headers.get("Host", self.target_host)

        return InterceptedRequest(
            id=entry_id,
            timestamp=datetime.now(),
            client_ip=client_addr[0],
            client_port=client_addr[1],
            method=method,
            url=f"http://{host}{path}",
            path=path,
            host=host,
            port=self.target_port,
            protocol=Protocol.HTTPS if self.tls_enabled else Protocol.HTTP,
            http_version=http_version,
            headers=headers,
            body=body or None,
            body_text=body_text,
            received_at=time.time()
        )
    
    def _parse_http_response(self, data: bytes, entry_id: str) -> InterceptedResponse:
        """Parse raw HTTP response data"""
        header_end = data.find(b"\r\n\r\n")
        if header_end == -1:
            header_end = len(data)
        header_bytes = data[:header_end]
        body = data[header_end + 4:] if header_end + 4 <= len(data) else b""

        header_text = header_bytes.decode("utf-8", errors="replace")
        lines = header_text.split("\r\n")
        status_line = lines[0] if lines else ""
        parts = status_line.split(" ", 2)

        http_version = parts[0] if len(parts) > 0 else "HTTP/1.1"
        status_code = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0
        status_message = parts[2] if len(parts) > 2 else ""

        headers: Dict[str, str] = {}
        for line in lines[1:]:
            if ":" not in line:
                continue
            key, value = line.split(":", 1)
            headers[key.strip()] = value.strip()

        header_lower = {k.lower(): v for k, v in headers.items()}
        if "chunked" in header_lower.get("transfer-encoding", "").lower():
            decoded = self._decode_chunked_body(body)
            if decoded is not None:
                body = decoded

        content_encoding = header_lower.get("content-encoding", "").lower()
        if content_encoding in ("gzip", "deflate"):
            try:
                if content_encoding == "gzip":
                    body = gzip.decompress(body)
                else:
                    body = zlib.decompress(body)
            except Exception:
                pass

        body_text = None
        content_length = len(body)
        if body:
            try:
                body_text = body.decode("utf-8", errors="replace")
            except Exception:
                body_text = None

        return InterceptedResponse(
            id=str(uuid.uuid4()),
            request_id=entry_id,
            timestamp=datetime.now(),
            status_code=status_code,
            status_message=status_message,
            http_version=http_version,
            headers=headers,
            body=body or None,
            body_text=body_text,
            content_length=content_length,
            received_at=time.time()
        )
    
    def _apply_rules_to_request(self, data: bytes, request: InterceptedRequest) -> tuple:
        """Apply interception rules to a request"""
        applied_rules: List[str] = []
        result = data
        
        for rule in sorted(self.rules, key=lambda r: r.priority):
            if not rule.enabled:
                continue
            if rule.match_direction not in ["request", "both"]:
                continue
            
            matches, context = self._rule_matches_request(rule, request)
            if matches:
                result = self._apply_rule_modifications(result, rule, "request", context)
                rule.hit_count += 1
                self.stats["rules_applied"] += 1
                applied_rules.append(rule.name)
        
        return result, applied_rules
    
    def _apply_rules_to_response(self, data: bytes, response: InterceptedResponse) -> tuple:
        """Apply interception rules to a response"""
        applied_rules: List[str] = []
        result = data
        
        for rule in sorted(self.rules, key=lambda r: r.priority):
            if not rule.enabled:
                continue
            if rule.match_direction not in ["response", "both"]:
                continue
            
            matches, context = self._rule_matches_response(rule, response)
            if matches:
                result = self._apply_rule_modifications(result, rule, "response", context)
                rule.hit_count += 1
                self.stats["rules_applied"] += 1
                applied_rules.append(rule.name)
        
        return result, applied_rules
    
    def _rule_matches_request(self, rule: InterceptionRule, request: InterceptedRequest) -> Tuple[bool, Dict[str, Any]]:
        """Check if a rule matches a request."""
        context: Dict[str, Any] = {}
        if rule.match_host:
            match = re.search(rule.match_host, request.host, re.IGNORECASE)
            if not match:
                return False, {}
            context["host_match"] = match
        if rule.match_path:
            match = re.search(rule.match_path, request.path, re.IGNORECASE)
            if not match:
                return False, {}
            context["path_match"] = match
        if rule.match_method and request.method.upper() != rule.match_method.upper():
            return False, {}
        if rule.match_body and request.body_text:
            match = re.search(rule.match_body, request.body_text, re.IGNORECASE)
            if not match:
                return False, {}
            context["body_match"] = match
        if rule.match_header:
            for key, pattern in rule.match_header.items():
                if key not in request.headers:
                    return False, {}
                if not re.search(pattern, request.headers[key], re.IGNORECASE):
                    return False, {}
        if rule.match_query:
            split = urlsplit(request.path)
            query_params = parse_qs(split.query)
            for key, pattern in rule.match_query.items():
                values = query_params.get(key)
                if not values:
                    return False, {}
                if pattern:
                    if not any(re.search(pattern, value, re.IGNORECASE) for value in values):
                        return False, {}
        return True, context
    
    def _rule_matches_response(self, rule: InterceptionRule, response: InterceptedResponse) -> Tuple[bool, Dict[str, Any]]:
        """Check if a rule matches a response."""
        context: Dict[str, Any] = {}
        if rule.match_status_code and response.status_code != rule.match_status_code:
            return False, {}
        if rule.match_content_type:
            ct = response.headers.get("Content-Type", "")
            match = re.search(rule.match_content_type, ct, re.IGNORECASE)
            if not match:
                return False, {}
            context["content_type_match"] = match
        if rule.match_body and response.body_text:
            match = re.search(rule.match_body, response.body_text, re.IGNORECASE)
            if not match:
                return False, {}
            context["body_match"] = match
        return True, context
    
    def _apply_rule_modifications(self, data: bytes, rule: InterceptionRule, direction: str, context: Dict[str, Any]) -> bytes:
        """Apply rule modifications to data."""
        result = data

        if rule.delay_ms > 0:
            time.sleep(rule.delay_ms / 1000.0)

        if rule.action == "drop":
            return b""

        header_end = result.find(b"\r\n\r\n")
        if header_end == -1:
            return result

        header_bytes = result[:header_end]
        body = result[header_end + 4:]

        header_text = header_bytes.decode("utf-8", errors="replace")
        header_lines = header_text.split("\r\n")
        start_line = header_lines[0] if header_lines else ""

        headers: List[List[str]] = []
        header_map: Dict[str, str] = {}
        for line in header_lines[1:]:
            if ":" not in line:
                continue
            key, value = line.split(":", 1)
            key = key.strip()
            value = value.strip()
            headers.append([key, value])
            header_map[key.lower()] = value

        transfer_encoding = header_map.get("transfer-encoding", "").lower()
        content_encoding = header_map.get("content-encoding", "").lower()

        decoded_body = body
        was_chunked = False
        was_compressed = False
        decompression_succeeded = True

        if "chunked" in transfer_encoding:
            decoded = self._decode_chunked_body(body)
            if decoded is not None:
                decoded_body = decoded
                was_chunked = True

        if content_encoding in ("gzip", "deflate"):
            was_compressed = True
            try:
                if content_encoding == "gzip":
                    decoded_body = gzip.decompress(decoded_body)
                else:
                    decoded_body = zlib.decompress(decoded_body)
            except Exception:
                # Decompression failed - skip body modifications and pass through unchanged
                decompression_succeeded = False
                logger.debug(f"Decompression failed for {content_encoding}, passing through unchanged")

        body_text = None
        body_modified = False

        # Only attempt body modifications if we successfully decompressed (or content wasn't compressed)
        if decompression_succeeded and decoded_body:
            try:
                body_text = decoded_body.decode("utf-8", errors="replace")
            except Exception:
                body_text = None

        if direction == "request" and rule.modify_path and start_line:
            parts = start_line.split(" ")
            if len(parts) >= 2:
                parts[1] = rule.modify_path
                start_line = " ".join(parts)

        if direction == "response" and rule.modify_status_code and start_line:
            parts = start_line.split(" ", 2)
            if len(parts) >= 2:
                parts[1] = str(rule.modify_status_code)
                start_line = " ".join(parts)

        if decompression_succeeded and rule.json_path_edits and decoded_body:
            updated_json = self._apply_json_path_edits(decoded_body, rule.json_path_edits)
            if updated_json is not None:
                decoded_body = updated_json
                body_text = decoded_body.decode("utf-8", errors="replace")
                body_modified = True

        if decompression_succeeded and rule.modify_body is not None:
            decoded_body = rule.modify_body.encode("utf-8")
            body_text = rule.modify_body
            body_modified = True

        if decompression_succeeded and rule.body_find_replace and body_text is not None:
            try:
                if rule.body_find_replace_regex:
                    for find, replace in rule.body_find_replace.items():
                        body_text = re.sub(find, replace, body_text, flags=re.IGNORECASE)
                else:
                    for find, replace in rule.body_find_replace.items():
                        body_text = body_text.replace(find, replace)
                decoded_body = body_text.encode("utf-8")
                body_modified = True
            except Exception:
                pass

        # Recompress if we successfully decompressed and modified the body
        if was_compressed and decompression_succeeded and body_modified:
            try:
                if content_encoding == "gzip":
                    decoded_body = gzip.compress(decoded_body)
                else:
                    decoded_body = zlib.compress(decoded_body)
            except Exception:
                # Recompression failed - remove Content-Encoding header and send uncompressed
                logger.debug(f"Recompression failed for {content_encoding}, removing encoding header")
                self._remove_header(headers, "Content-Encoding")
        elif was_compressed and not decompression_succeeded:
            # Decompression failed - use original body unchanged
            decoded_body = body if not was_chunked else self._decode_chunked_body(body) or body

        if "chunked" in transfer_encoding:
            body = self._encode_chunked_body(decoded_body)
            self._update_header(headers, "Content-Length", None)
        else:
            body = decoded_body
            self._update_header(headers, "Content-Length", str(len(body)))

        if rule.modify_headers:
            for header, value in rule.modify_headers.items():
                self._update_header(headers, header, value)

        if rule.remove_headers:
            for header in rule.remove_headers:
                self._remove_header(headers, header)

        rebuilt_headers = "\r\n".join([start_line] + [f"{k}: {v}" for k, v in headers]) + "\r\n\r\n"
        return rebuilt_headers.encode("utf-8") + body

    def _update_header(self, headers: List[List[str]], name: str, value: Optional[str]) -> None:
        """Update or insert a header in the header list."""
        lowered = name.lower()
        for pair in headers:
            if pair[0].lower() == lowered:
                if value is None:
                    headers.remove(pair)
                else:
                    pair[1] = value
                return
        if value is not None:
            headers.append([name, value])

    def _remove_header(self, headers: List[List[str]], name: str) -> None:
        """Remove a header from the header list."""
        lowered = name.lower()
        for pair in list(headers):
            if pair[0].lower() == lowered:
                headers.remove(pair)

    def _apply_json_path_edits(self, body: bytes, edits: List[Dict[str, Any]]) -> Optional[bytes]:
        """Apply simple JSON path edits to a body."""
        try:
            data = json.loads(body.decode("utf-8", errors="replace"))
        except Exception:
            return None

        for edit in edits:
            path = str(edit.get("path", "")).strip()
            if not path:
                continue
            op = edit.get("op", "set")
            value = edit.get("value")
            tokens: List[Any] = []
            for part in path.split("."):
                while part:
                    if "[" in part:
                        name, rest = part.split("[", 1)
                        if name:
                            tokens.append(name)
                        idx_str, remainder = rest.split("]", 1)
                        if idx_str.isdigit():
                            tokens.append(int(idx_str))
                        part = remainder
                    else:
                        tokens.append(part)
                        part = ""

            if not tokens:
                continue

            current = data
            for token in tokens[:-1]:
                if isinstance(token, int):
                    if not isinstance(current, list):
                        break
                    if token >= len(current):
                        current.extend([None] * (token - len(current) + 1))
                    if current[token] is None:
                        current[token] = {}
                    current = current[token]
                else:
                    if not isinstance(current, dict):
                        break
                    current = current.setdefault(token, {})
            else:
                last = tokens[-1]
                if op == "remove":
                    if isinstance(last, int) and isinstance(current, list) and last < len(current):
                        current.pop(last)
                    elif isinstance(last, str) and isinstance(current, dict):
                        current.pop(last, None)
                else:
                    if isinstance(last, int) and isinstance(current, list):
                        if last >= len(current):
                            current.extend([None] * (last - len(current) + 1))
                        current[last] = value
                    elif isinstance(last, str) and isinstance(current, dict):
                        current[last] = value

        return json.dumps(data, ensure_ascii=True).encode("utf-8")
    
    def add_rule(self, rule: InterceptionRule):
        """Add an interception rule"""
        self.rules.append(rule)
    
    def remove_rule(self, rule_id: str):
        """Remove an interception rule"""
        self.rules = [r for r in self.rules if r.id != rule_id]
    
    def get_traffic_log(self, limit: int = 100, offset: int = 0) -> List[Dict]:
        """Get traffic log entries"""
        with self._lock:
            entries = self.traffic_log[-(offset + limit):]
            if offset > 0:
                entries = entries[:-offset]
            entries = entries[-limit:]
        
        return [self._traffic_entry_to_dict(e) for e in entries]
    
    def _traffic_entry_to_dict(self, entry: TrafficEntry) -> Dict:
        """Convert traffic entry to dictionary"""
        duration_ms = entry.response.response_time_ms if entry.response else 0
        modified = entry.request.modified or (entry.response.modified if entry.response else False)
        return {
            "id": entry.id,
            "timestamp": entry.request.timestamp.isoformat(),
            "request": {
                "id": entry.request.id,
                "timestamp": entry.request.timestamp.isoformat(),
                "client_ip": entry.request.client_ip,
                "client_port": entry.request.client_port,
                "method": entry.request.method,
                "url": entry.request.url,
                "path": entry.request.path,
                "host": entry.request.host,
                "port": entry.request.port,
                "protocol": entry.request.protocol.value,
                "http_version": entry.request.http_version,
                "headers": entry.request.headers,
                "body": entry.request.body_text[:10000] if entry.request.body_text else None,
                "body_text": entry.request.body_text[:10000] if entry.request.body_text else None,
                "modified": entry.request.modified
            },
            "response": {
                "id": entry.response.id,
                "status_code": entry.response.status_code,
                "status_text": entry.response.status_message,
                "status_message": entry.response.status_message,
                "headers": entry.response.headers,
                "body": entry.response.body_text[:10000] if entry.response.body_text else None,
                "body_text": entry.response.body_text[:10000] if entry.response.body_text else None,
                "content_length": entry.response.content_length,
                "response_time_ms": entry.response.response_time_ms,
                "modified": entry.response.modified
            } if entry.response else None,
            "duration_ms": duration_ms,
            "modified": modified,
            "rules_applied": entry.rules_applied,
            "error": entry.error,
            "tags": entry.tags,
            "notes": entry.notes
        }

    def _append_entry(self, entry: TrafficEntry, connection_id: Optional[str] = None) -> None:
        with self._lock:
            self.traffic_log.append(entry)
            self._traffic_index[entry.id] = entry
            if connection_id:
                self._connection_requests.setdefault(connection_id, []).append(entry.id)
            if len(self.traffic_log) > 1000:
                removed_entries = self.traffic_log[:-500]
                self.traffic_log = self.traffic_log[-500:]
                for old_entry in removed_entries:
                    self._traffic_index.pop(old_entry.id, None)
    
    def clear_traffic_log(self):
        """Clear the traffic log"""
        with self._lock:
            self.traffic_log.clear()
            self._traffic_index.clear()
            self._connection_requests.clear()

    def update_traffic_entry(self, entry_id: str, notes: Optional[str] = None, tags: Optional[List[str]] = None) -> TrafficEntry:
        """Update notes or tags for a traffic entry"""
        with self._lock:
            entry = self._traffic_index.get(entry_id)
            if not entry:
                raise ValueError(f"Traffic entry {entry_id} not found")
            if notes is not None:
                entry.notes = notes
            if tags is not None:
                entry.tags = tags
            return entry
    
    def get_stats(self) -> Dict:
        """Get proxy statistics"""
        uptime = 0
        if self.stats["start_time"]:
            uptime = time.time() - self.stats["start_time"]

        return {
            **self.stats,
            "requests": self.stats.get("requests_total", 0),
            "responses": self.stats.get("responses_total", 0),
            "uptime_seconds": uptime,
            "running": self.running,
            "mode": self.mode.value,
            "rules_count": len(self.rules),
            "traffic_log_size": len(self.traffic_log),
            "websocket_connections": len(getattr(self, '_websocket_connections', {})),
            "websocket_frames_total": self.stats.get("websocket_frames_total", 0)
        }


# ============================================================================
# Request/Response Diff Viewer
# ============================================================================

@dataclass
class DiffResult:
    """Result of comparing original vs modified data"""
    has_changes: bool
    change_type: str  # "headers", "body", "both", "none"
    header_changes: List[Dict[str, Any]] = field(default_factory=list)
    body_changes: List[Dict[str, Any]] = field(default_factory=list)
    summary: str = ""
    original_size: int = 0
    modified_size: int = 0


class TrafficDiffViewer:
    """
    Visual diff viewer for comparing original vs modified HTTP traffic.
    Provides line-by-line and character-level diff highlighting.
    """
    
    @staticmethod
    def diff_headers(
        original: Dict[str, str],
        modified: Dict[str, str]
    ) -> List[Dict[str, Any]]:
        """Compare two sets of HTTP headers"""
        changes = []
        all_keys = set(original.keys()) | set(modified.keys())
        
        for key in sorted(all_keys):
            orig_val = original.get(key)
            mod_val = modified.get(key)
            
            if orig_val is None:
                changes.append({
                    "type": "added",
                    "header": key,
                    "value": mod_val,
                    "original": None
                })
            elif mod_val is None:
                changes.append({
                    "type": "removed",
                    "header": key,
                    "value": None,
                    "original": orig_val
                })
            elif orig_val != mod_val:
                changes.append({
                    "type": "modified",
                    "header": key,
                    "value": mod_val,
                    "original": orig_val
                })
        
        return changes
    
    @staticmethod
    def diff_body(
        original: str,
        modified: str,
        context_lines: int = 3
    ) -> List[Dict[str, Any]]:
        """Generate line-by-line diff of body content"""
        import difflib
        
        if original == modified:
            return []
        
        orig_lines = original.splitlines(keepends=True)
        mod_lines = modified.splitlines(keepends=True)
        
        differ = difflib.unified_diff(
            orig_lines,
            mod_lines,
            fromfile='original',
            tofile='modified',
            lineterm='',
            n=context_lines
        )
        
        changes = []
        for line in differ:
            if line.startswith('+++') or line.startswith('---'):
                continue
            elif line.startswith('@@'):
                changes.append({"type": "range", "content": line.strip()})
            elif line.startswith('+'):
                changes.append({"type": "added", "content": line[1:]})
            elif line.startswith('-'):
                changes.append({"type": "removed", "content": line[1:]})
            else:
                changes.append({"type": "context", "content": line[1:] if line.startswith(' ') else line})
        
        return changes
    
    @staticmethod
    def diff_json(
        original: Any,
        modified: Any,
        path: str = ""
    ) -> List[Dict[str, Any]]:
        """Deep diff for JSON objects"""
        changes = []
        
        if type(original) != type(modified):
            changes.append({
                "type": "type_change",
                "path": path or "$",
                "original": type(original).__name__,
                "modified": type(modified).__name__,
                "original_value": original,
                "modified_value": modified
            })
            return changes
        
        if isinstance(original, dict):
            all_keys = set(original.keys()) | set(modified.keys())
            for key in sorted(all_keys):
                key_path = f"{path}.{key}" if path else f"$.{key}"
                if key not in modified:
                    changes.append({"type": "removed", "path": key_path, "value": original[key]})
                elif key not in original:
                    changes.append({"type": "added", "path": key_path, "value": modified[key]})
                else:
                    changes.extend(TrafficDiffViewer.diff_json(original[key], modified[key], key_path))
        
        elif isinstance(original, list):
            for i in range(max(len(original), len(modified))):
                idx_path = f"{path}[{i}]"
                if i >= len(modified):
                    changes.append({"type": "removed", "path": idx_path, "value": original[i]})
                elif i >= len(original):
                    changes.append({"type": "added", "path": idx_path, "value": modified[i]})
                else:
                    changes.extend(TrafficDiffViewer.diff_json(original[i], modified[i], idx_path))
        
        elif original != modified:
            changes.append({
                "type": "modified",
                "path": path or "$",
                "original": original,
                "modified": modified
            })
        
        return changes
    
    @classmethod
    def compare_traffic(
        cls,
        original_request: Dict[str, Any],
        modified_request: Optional[Dict[str, Any]],
        original_response: Optional[Dict[str, Any]] = None,
        modified_response: Optional[Dict[str, Any]] = None
    ) -> Dict[str, DiffResult]:
        """Compare full traffic entry (request and response)"""
        results = {}
        
        # Compare request
        if modified_request:
            req_header_changes = cls.diff_headers(
                original_request.get("headers", {}),
                modified_request.get("headers", {})
            )
            
            orig_body = original_request.get("body", "") or ""
            mod_body = modified_request.get("body", "") or ""
            
            # Try JSON diff for JSON content
            req_body_changes = []
            content_type = original_request.get("headers", {}).get("content-type", "")
            if "json" in content_type.lower():
                try:
                    orig_json = json.loads(orig_body) if orig_body else {}
                    mod_json = json.loads(mod_body) if mod_body else {}
                    req_body_changes = cls.diff_json(orig_json, mod_json)
                except json.JSONDecodeError:
                    req_body_changes = cls.diff_body(orig_body, mod_body)
            else:
                req_body_changes = cls.diff_body(orig_body, mod_body)
            
            has_changes = bool(req_header_changes) or bool(req_body_changes)
            change_type = "none"
            if req_header_changes and req_body_changes:
                change_type = "both"
            elif req_header_changes:
                change_type = "headers"
            elif req_body_changes:
                change_type = "body"
            
            results["request"] = DiffResult(
                has_changes=has_changes,
                change_type=change_type,
                header_changes=req_header_changes,
                body_changes=req_body_changes,
                summary=f"{len(req_header_changes)} header changes, {len(req_body_changes)} body changes",
                original_size=len(orig_body),
                modified_size=len(mod_body)
            )
        
        # Compare response
        if original_response and modified_response:
            resp_header_changes = cls.diff_headers(
                original_response.get("headers", {}),
                modified_response.get("headers", {})
            )
            
            orig_body = original_response.get("body", "") or ""
            mod_body = modified_response.get("body", "") or ""
            
            resp_body_changes = []
            content_type = original_response.get("headers", {}).get("content-type", "")
            if "json" in content_type.lower():
                try:
                    orig_json = json.loads(orig_body) if orig_body else {}
                    mod_json = json.loads(mod_body) if mod_body else {}
                    resp_body_changes = cls.diff_json(orig_json, mod_json)
                except json.JSONDecodeError:
                    resp_body_changes = cls.diff_body(orig_body, mod_body)
            else:
                resp_body_changes = cls.diff_body(orig_body, mod_body)
            
            has_changes = bool(resp_header_changes) or bool(resp_body_changes)
            change_type = "none"
            if resp_header_changes and resp_body_changes:
                change_type = "both"
            elif resp_header_changes:
                change_type = "headers"
            elif resp_body_changes:
                change_type = "body"
            
            results["response"] = DiffResult(
                has_changes=has_changes,
                change_type=change_type,
                header_changes=resp_header_changes,
                body_changes=resp_body_changes,
                summary=f"{len(resp_header_changes)} header changes, {len(resp_body_changes)} body changes",
                original_size=len(orig_body),
                modified_size=len(mod_body)
            )
        
        return results


# ============================================================================
# HTTP/2 & gRPC Support
# ============================================================================

class HTTP2FrameType(int, Enum):
    """HTTP/2 frame types"""
    DATA = 0x0
    HEADERS = 0x1
    PRIORITY = 0x2
    RST_STREAM = 0x3
    SETTINGS = 0x4
    PUSH_PROMISE = 0x5
    PING = 0x6
    GOAWAY = 0x7
    WINDOW_UPDATE = 0x8
    CONTINUATION = 0x9


@dataclass
class HTTP2Frame:
    """Parsed HTTP/2 frame"""
    id: str
    timestamp: datetime
    stream_id: int
    frame_type: int
    frame_type_name: str
    flags: int
    length: int
    payload: bytes
    headers: Optional[Dict[str, str]] = None  # For HEADERS frames
    data: Optional[bytes] = None  # For DATA frames
    
    # gRPC specific
    is_grpc: bool = False
    grpc_message: Optional[bytes] = None
    grpc_compressed: bool = False


@dataclass  
class GRPCMessage:
    """Decoded gRPC message"""
    id: str
    timestamp: datetime
    stream_id: int
    direction: str
    compressed: bool
    length: int
    payload: bytes
    payload_decoded: Optional[Dict[str, Any]] = None  # Decoded protobuf
    method: Optional[str] = None
    service: Optional[str] = None


class HTTP2Parser:
    """
    Parser for HTTP/2 frames and gRPC messages.
    Supports inspecting modern HTTP/2 and gRPC traffic.
    """
    
    FRAME_HEADER_SIZE = 9
    
    @classmethod
    def parse_frame_header(cls, data: bytes) -> Optional[Tuple[int, int, int, int]]:
        """Parse HTTP/2 frame header (9 bytes)"""
        if len(data) < cls.FRAME_HEADER_SIZE:
            return None
        
        length = (data[0] << 16) | (data[1] << 8) | data[2]
        frame_type = data[3]
        flags = data[4]
        stream_id = struct.unpack(">I", data[5:9])[0] & 0x7FFFFFFF
        
        return length, frame_type, flags, stream_id
    
    @classmethod
    def parse_frame(cls, data: bytes, offset: int = 0) -> Optional[Tuple[HTTP2Frame, int]]:
        """Parse a complete HTTP/2 frame"""
        if len(data) - offset < cls.FRAME_HEADER_SIZE:
            return None
        
        header_result = cls.parse_frame_header(data[offset:offset + cls.FRAME_HEADER_SIZE])
        if not header_result:
            return None
        
        length, frame_type, flags, stream_id = header_result
        
        if len(data) - offset - cls.FRAME_HEADER_SIZE < length:
            return None
        
        payload = data[offset + cls.FRAME_HEADER_SIZE:offset + cls.FRAME_HEADER_SIZE + length]
        
        frame_type_names = {
            0: "DATA", 1: "HEADERS", 2: "PRIORITY", 3: "RST_STREAM",
            4: "SETTINGS", 5: "PUSH_PROMISE", 6: "PING", 7: "GOAWAY",
            8: "WINDOW_UPDATE", 9: "CONTINUATION"
        }
        
        frame = HTTP2Frame(
            id=str(uuid.uuid4()),
            timestamp=datetime.now(),
            stream_id=stream_id,
            frame_type=frame_type,
            frame_type_name=frame_type_names.get(frame_type, f"UNKNOWN({frame_type})"),
            flags=flags,
            length=length,
            payload=payload
        )
        
        # Extract DATA frame payload
        if frame_type == HTTP2FrameType.DATA:
            frame.data = payload
            # Check for gRPC
            if len(payload) >= 5:
                frame.is_grpc = True
                frame.grpc_compressed = bool(payload[0])
                frame.grpc_message = payload[5:]
        
        return frame, offset + cls.FRAME_HEADER_SIZE + length
    
    @classmethod
    def parse_grpc_message(cls, data: bytes, stream_id: int, direction: str) -> Optional[GRPCMessage]:
        """Parse a gRPC message from DATA frame payload"""
        if len(data) < 5:
            return None
        
        compressed = bool(data[0])
        message_length = struct.unpack(">I", data[1:5])[0]
        
        if len(data) < 5 + message_length:
            return None
        
        payload = data[5:5 + message_length]
        
        return GRPCMessage(
            id=str(uuid.uuid4()),
            timestamp=datetime.now(),
            stream_id=stream_id,
            direction=direction,
            compressed=compressed,
            length=message_length,
            payload=payload
        )
    
    @classmethod
    def detect_http2(cls, data: bytes) -> bool:
        """Check if data looks like HTTP/2 connection preface or frames"""
        # HTTP/2 connection preface
        PREFACE = b"PRI * HTTP/2.0\r\n\r\nSM\r\n\r\n"
        if data.startswith(PREFACE):
            return True
        
        # Check for valid frame header
        if len(data) >= cls.FRAME_HEADER_SIZE:
            header = cls.parse_frame_header(data)
            if header:
                length, frame_type, _, _ = header
                # Valid frame types are 0-9
                if 0 <= frame_type <= 9 and length < 16777216:  # Max frame size
                    return True
        
        return False


# ============================================================================
# Match & Replace Templates Library
# ============================================================================

@dataclass
class MatchReplaceTemplate:
    """Pre-built template for common MITM modifications"""
    id: str
    name: str
    category: str
    description: str
    match_type: str  # "header", "body", "path", "query"
    match_pattern: str
    replace_pattern: str
    is_regex: bool = False
    case_sensitive: bool = True
    enabled: bool = True
    # Optional conditions
    match_host: Optional[str] = None
    match_content_type: Optional[str] = None
    direction: str = "both"  # "request", "response", "both"
    tags: List[str] = field(default_factory=list)
    
    # Usage tracking
    hit_count: int = 0
    created_at: datetime = field(default_factory=datetime.now)


class MatchReplaceLibrary:
    """
    Library of pre-built match/replace templates for common MITM modifications.
    Includes security testing, debugging, and development templates.
    """
    
    BUILT_IN_TEMPLATES: List[Dict[str, Any]] = [
        # Security Testing Templates
        {
            "id": "sec-cors-wildcard",
            "name": "CORS Wildcard Origin",
            "category": "Security Testing",
            "description": "Replace Access-Control-Allow-Origin with wildcard",
            "match_type": "header",
            "match_pattern": "access-control-allow-origin",
            "replace_pattern": "*",
            "direction": "response",
            "tags": ["cors", "security"]
        },
        {
            "id": "sec-remove-csp",
            "name": "Remove Content-Security-Policy",
            "category": "Security Testing", 
            "description": "Remove CSP header to allow script injection",
            "match_type": "header",
            "match_pattern": "content-security-policy",
            "replace_pattern": "",
            "direction": "response",
            "tags": ["csp", "xss", "security"]
        },
        {
            "id": "sec-remove-x-frame-options",
            "name": "Remove X-Frame-Options",
            "category": "Security Testing",
            "description": "Remove X-Frame-Options to allow framing",
            "match_type": "header",
            "match_pattern": "x-frame-options",
            "replace_pattern": "",
            "direction": "response",
            "tags": ["clickjacking", "security"]
        },
        {
            "id": "sec-jwt-alg-none",
            "name": "JWT Algorithm None Attack",
            "category": "Security Testing",
            "description": "Change JWT algorithm to 'none'",
            "match_type": "body",
            "match_pattern": '"alg"\\s*:\\s*"[^"]+"',
            "replace_pattern": '"alg": "none"',
            "is_regex": True,
            "direction": "request",
            "tags": ["jwt", "auth", "security"]
        },
        {
            "id": "sec-sql-inject-test",
            "name": "SQL Injection Test Payload",
            "category": "Security Testing",
            "description": "Add SQL injection test payload to ID parameters",
            "match_type": "query",
            "match_pattern": "id=(\\d+)",
            "replace_pattern": "id=\\1'--",
            "is_regex": True,
            "direction": "request",
            "tags": ["sqli", "security"]
        },
        {
            "id": "sec-xss-test",
            "name": "XSS Test Payload",
            "category": "Security Testing",
            "description": "Replace input with XSS test payload",
            "match_type": "body",
            "match_pattern": '"value"\\s*:\\s*"([^"]*)"',
            "replace_pattern": '"value": "<script>alert(1)</script>"',
            "is_regex": True,
            "direction": "request",
            "tags": ["xss", "security"]
        },
        
        # Debugging Templates
        {
            "id": "debug-add-timing",
            "name": "Add Request Timing Header",
            "category": "Debugging",
            "description": "Add X-Request-Time header for debugging",
            "match_type": "header",
            "match_pattern": "__add_header__",
            "replace_pattern": "X-Request-Time: {{timestamp}}",
            "direction": "request",
            "tags": ["debugging", "timing"]
        },
        {
            "id": "debug-force-no-cache",
            "name": "Force No Cache",
            "category": "Debugging",
            "description": "Add no-cache headers to bypass caching",
            "match_type": "header",
            "match_pattern": "cache-control",
            "replace_pattern": "no-cache, no-store, must-revalidate",
            "direction": "request",
            "tags": ["cache", "debugging"]
        },
        {
            "id": "debug-remove-gzip",
            "name": "Disable Compression",
            "category": "Debugging",
            "description": "Remove Accept-Encoding to get uncompressed responses",
            "match_type": "header",
            "match_pattern": "accept-encoding",
            "replace_pattern": "identity",
            "direction": "request",
            "tags": ["compression", "debugging"]
        },
        
        # Development Templates
        {
            "id": "dev-mock-auth",
            "name": "Mock Admin Authorization",
            "category": "Development",
            "description": "Add admin role to authorization header",
            "match_type": "header",
            "match_pattern": "authorization",
            "replace_pattern": "Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJyb2xlIjoiYWRtaW4ifQ.mock",
            "direction": "request",
            "tags": ["auth", "development"]
        },
        {
            "id": "dev-add-cors-headers",
            "name": "Add Full CORS Headers",
            "category": "Development",
            "description": "Add permissive CORS headers for local development",
            "match_type": "header",
            "match_pattern": "__add_headers__",
            "replace_pattern": "Access-Control-Allow-Origin: *\nAccess-Control-Allow-Methods: GET, POST, PUT, DELETE, OPTIONS\nAccess-Control-Allow-Headers: *",
            "direction": "response",
            "tags": ["cors", "development"]
        },
        {
            "id": "dev-delay-response",
            "name": "Simulate Slow Network",
            "category": "Development",
            "description": "Add artificial delay to responses",
            "match_type": "body",
            "match_pattern": "__delay__",
            "replace_pattern": "3000",
            "direction": "response",
            "tags": ["latency", "development"]
        },
        
        # API Testing Templates
        {
            "id": "api-change-method",
            "name": "HTTP Method Override",
            "category": "API Testing",
            "description": "Add X-HTTP-Method-Override header",
            "match_type": "header",
            "match_pattern": "__add_header__",
            "replace_pattern": "X-HTTP-Method-Override: DELETE",
            "direction": "request",
            "tags": ["api", "method"]
        },
        {
            "id": "api-json-to-xml",
            "name": "Force XML Response",
            "category": "API Testing",
            "description": "Change Accept header to request XML",
            "match_type": "header",
            "match_pattern": "accept",
            "replace_pattern": "application/xml",
            "direction": "request",
            "tags": ["api", "content-type"]
        },
        {
            "id": "api-version-downgrade",
            "name": "API Version Downgrade",
            "category": "API Testing",
            "description": "Change API version in path",
            "match_type": "path",
            "match_pattern": "/api/v\\d+/",
            "replace_pattern": "/api/v1/",
            "is_regex": True,
            "direction": "request",
            "tags": ["api", "version"]
        },
        
        # Mobile App Testing
        {
            "id": "mobile-ssl-pinning-bypass",
            "name": "SSL Pinning Bypass Headers",
            "category": "Mobile Testing",
            "description": "Add headers that some apps use to disable pinning in debug mode",
            "match_type": "header",
            "match_pattern": "__add_header__",
            "replace_pattern": "X-Debug-Mode: true\nX-SSL-Pinning: disabled",
            "direction": "request",
            "tags": ["mobile", "ssl", "security"]
        },
        {
            "id": "mobile-device-spoof",
            "name": "Spoof Mobile Device",
            "category": "Mobile Testing",
            "description": "Change User-Agent to mobile device",
            "match_type": "header",
            "match_pattern": "user-agent",
            "replace_pattern": "Mozilla/5.0 (iPhone; CPU iPhone OS 16_0 like Mac OS X) AppleWebKit/605.1.15",
            "direction": "request",
            "tags": ["mobile", "user-agent"]
        }
    ]
    
    def __init__(self):
        self.templates: Dict[str, MatchReplaceTemplate] = {}
        self.custom_templates: Dict[str, MatchReplaceTemplate] = {}
        self._load_built_in_templates()
    
    def _load_built_in_templates(self) -> None:
        """Load built-in templates"""
        for tmpl_data in self.BUILT_IN_TEMPLATES:
            tmpl = MatchReplaceTemplate(
                id=tmpl_data["id"],
                name=tmpl_data["name"],
                category=tmpl_data["category"],
                description=tmpl_data["description"],
                match_type=tmpl_data["match_type"],
                match_pattern=tmpl_data["match_pattern"],
                replace_pattern=tmpl_data["replace_pattern"],
                is_regex=tmpl_data.get("is_regex", False),
                direction=tmpl_data.get("direction", "both"),
                tags=tmpl_data.get("tags", [])
            )
            self.templates[tmpl.id] = tmpl
    
    def get_all_templates(self) -> List[MatchReplaceTemplate]:
        """Get all available templates (built-in + custom)"""
        return list(self.templates.values()) + list(self.custom_templates.values())
    
    def get_templates_by_category(self, category: str) -> List[MatchReplaceTemplate]:
        """Get templates filtered by category"""
        return [t for t in self.get_all_templates() if t.category == category]
    
    def get_templates_by_tag(self, tag: str) -> List[MatchReplaceTemplate]:
        """Get templates filtered by tag"""
        return [t for t in self.get_all_templates() if tag in t.tags]
    
    def get_categories(self) -> List[str]:
        """Get list of available categories"""
        return list(set(t.category for t in self.get_all_templates()))
    
    def add_custom_template(self, template: MatchReplaceTemplate) -> None:
        """Add a custom template"""
        self.custom_templates[template.id] = template
    
    def remove_custom_template(self, template_id: str) -> bool:
        """Remove a custom template"""
        if template_id in self.custom_templates:
            del self.custom_templates[template_id]
            return True
        return False
    
    def apply_template(
        self,
        template_id: str,
        data: Dict[str, Any],
        direction: str
    ) -> Tuple[Dict[str, Any], bool]:
        """Apply a template to request/response data"""
        template = self.templates.get(template_id) or self.custom_templates.get(template_id)
        if not template or not template.enabled:
            return data, False
        
        if template.direction not in ["both", direction]:
            return data, False
        
        modified = False
        result = data.copy()
        
        if template.match_type == "header":
            headers = result.get("headers", {})
            new_headers = {}
            for key, value in headers.items():
                if key.lower() == template.match_pattern.lower():
                    if template.replace_pattern:
                        new_headers[key] = template.replace_pattern
                        modified = True
                    # Empty replace_pattern means remove the header
                    else:
                        modified = True
                        continue
                else:
                    new_headers[key] = value
            result["headers"] = new_headers
        
        elif template.match_type == "body":
            body = result.get("body", "")
            if body and template.is_regex:
                new_body = re.sub(
                    template.match_pattern,
                    template.replace_pattern,
                    body,
                    flags=0 if template.case_sensitive else re.IGNORECASE
                )
                if new_body != body:
                    result["body"] = new_body
                    modified = True
            elif body:
                if template.case_sensitive:
                    new_body = body.replace(template.match_pattern, template.replace_pattern)
                else:
                    pattern = re.compile(re.escape(template.match_pattern), re.IGNORECASE)
                    new_body = pattern.sub(template.replace_pattern, body)
                if new_body != body:
                    result["body"] = new_body
                    modified = True
        
        elif template.match_type == "path":
            path = result.get("path", "")
            if path:
                if template.is_regex:
                    new_path = re.sub(template.match_pattern, template.replace_pattern, path)
                else:
                    new_path = path.replace(template.match_pattern, template.replace_pattern)
                if new_path != path:
                    result["path"] = new_path
                    modified = True
        
        elif template.match_type == "query":
            path = result.get("path", "")
            if "?" in path:
                base, query = path.split("?", 1)
                if template.is_regex:
                    new_query = re.sub(template.match_pattern, template.replace_pattern, query)
                else:
                    new_query = query.replace(template.match_pattern, template.replace_pattern)
                if new_query != query:
                    result["path"] = f"{base}?{new_query}"
                    modified = True
        
        if modified:
            template.hit_count += 1
        
        return result, modified


# ============================================================================
# Network Throttling Simulation
# ============================================================================

@dataclass
class ThrottleProfile:
    """Network throttling profile for simulating various connection types"""
    id: str
    name: str
    description: str
    bandwidth_kbps: int  # Kilobits per second (0 = unlimited)
    latency_ms: int  # Added latency in milliseconds
    packet_loss_percent: float = 0.0  # Percentage of dropped packets
    jitter_ms: int = 0  # Random latency variation
    enabled: bool = True
    is_builtin: bool = False
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "name": self.name,
            "description": self.description,
            "bandwidth_kbps": self.bandwidth_kbps,
            "latency_ms": self.latency_ms,
            "packet_loss_percent": self.packet_loss_percent,
            "jitter_ms": self.jitter_ms,
            "enabled": self.enabled,
            "is_builtin": self.is_builtin
        }


class NetworkThrottler:
    """
    Network throttling simulation for testing application behavior under various conditions.
    
    Features:
    - Bandwidth limiting
    - Latency injection
    - Packet loss simulation
    - Pre-built profiles for common scenarios
    """
    
    def __init__(self):
        self.profiles: Dict[str, ThrottleProfile] = {}
        self.custom_profiles: Dict[str, ThrottleProfile] = {}
        self.active_profile: Optional[str] = None
        self._init_builtin_profiles()
    
    def _init_builtin_profiles(self):
        """Initialize built-in throttling profiles"""
        builtin = [
            ThrottleProfile(
                id="no-throttle",
                name="No Throttling",
                description="Full speed, no restrictions",
                bandwidth_kbps=0,
                latency_ms=0,
                is_builtin=True
            ),
            ThrottleProfile(
                id="3g-slow",
                name="Slow 3G",
                description="Typical slow 3G mobile connection",
                bandwidth_kbps=400,
                latency_ms=300,
                packet_loss_percent=1.0,
                jitter_ms=50,
                is_builtin=True
            ),
            ThrottleProfile(
                id="3g-fast",
                name="Fast 3G",
                description="Fast 3G mobile connection",
                bandwidth_kbps=1500,
                latency_ms=150,
                packet_loss_percent=0.5,
                jitter_ms=30,
                is_builtin=True
            ),
            ThrottleProfile(
                id="4g-lte",
                name="4G LTE",
                description="4G LTE mobile connection",
                bandwidth_kbps=12000,
                latency_ms=50,
                packet_loss_percent=0.1,
                jitter_ms=10,
                is_builtin=True
            ),
            ThrottleProfile(
                id="wifi-slow",
                name="Slow WiFi",
                description="Congested or distant WiFi",
                bandwidth_kbps=2000,
                latency_ms=100,
                packet_loss_percent=2.0,
                jitter_ms=30,
                is_builtin=True
            ),
            ThrottleProfile(
                id="wifi-fast",
                name="Fast WiFi",
                description="Strong WiFi connection",
                bandwidth_kbps=30000,
                latency_ms=10,
                packet_loss_percent=0.1,
                jitter_ms=5,
                is_builtin=True
            ),
            ThrottleProfile(
                id="offline",
                name="Offline",
                description="Simulate complete network failure",
                bandwidth_kbps=0,
                latency_ms=10000,
                packet_loss_percent=100.0,
                is_builtin=True
            ),
            ThrottleProfile(
                id="high-latency",
                name="High Latency",
                description="Satellite or intercontinental connection",
                bandwidth_kbps=10000,
                latency_ms=500,
                jitter_ms=100,
                is_builtin=True
            ),
            ThrottleProfile(
                id="lossy",
                name="Lossy Connection",
                description="Unreliable network with high packet loss",
                bandwidth_kbps=5000,
                latency_ms=100,
                packet_loss_percent=10.0,
                jitter_ms=50,
                is_builtin=True
            ),
            ThrottleProfile(
                id="edge",
                name="EDGE (2G)",
                description="Very slow 2G EDGE connection",
                bandwidth_kbps=100,
                latency_ms=500,
                packet_loss_percent=2.0,
                jitter_ms=100,
                is_builtin=True
            )
        ]
        for profile in builtin:
            self.profiles[profile.id] = profile
    
    def get_profile(self, profile_id: str) -> Optional[ThrottleProfile]:
        """Get a throttle profile by ID"""
        return self.profiles.get(profile_id) or self.custom_profiles.get(profile_id)
    
    def get_all_profiles(self) -> List[ThrottleProfile]:
        """Get all available profiles"""
        return list(self.profiles.values()) + list(self.custom_profiles.values())
    
    def add_custom_profile(self, profile: ThrottleProfile) -> None:
        """Add a custom throttle profile"""
        self.custom_profiles[profile.id] = profile
    
    def remove_custom_profile(self, profile_id: str) -> bool:
        """Remove a custom profile"""
        if profile_id in self.custom_profiles:
            del self.custom_profiles[profile_id]
            if self.active_profile == profile_id:
                self.active_profile = None
            return True
        return False
    
    def set_active_profile(self, profile_id: Optional[str]) -> bool:
        """Set the active throttle profile"""
        if profile_id is None:
            self.active_profile = None
            return True
        if self.get_profile(profile_id):
            self.active_profile = profile_id
            return True
        return False
    
    def get_active_profile(self) -> Optional[ThrottleProfile]:
        """Get the currently active profile"""
        if self.active_profile:
            return self.get_profile(self.active_profile)
        return None
    
    async def apply_throttle(self, data: bytes) -> Tuple[bytes, Dict[str, Any]]:
        """
        Apply throttling to data transfer.
        Returns the data (possibly empty if dropped) and stats about the throttling applied.
        """
        profile = self.get_active_profile()
        if not profile or not profile.enabled:
            return data, {"throttled": False}
        
        stats = {
            "throttled": True,
            "profile": profile.name,
            "latency_applied_ms": 0,
            "bandwidth_delay_ms": 0,
            "dropped": False
        }
        
        import random
        
        # Check for packet loss
        if profile.packet_loss_percent > 0:
            if random.random() * 100 < profile.packet_loss_percent:
                stats["dropped"] = True
                return b"", stats
        
        # Apply latency
        total_latency = profile.latency_ms
        if profile.jitter_ms > 0:
            total_latency += random.randint(-profile.jitter_ms, profile.jitter_ms)
        total_latency = max(0, total_latency)
        
        if total_latency > 0:
            await asyncio.sleep(total_latency / 1000.0)
            stats["latency_applied_ms"] = total_latency
        
        # Apply bandwidth limiting
        if profile.bandwidth_kbps > 0:
            data_size_bits = len(data) * 8
            transfer_time_sec = data_size_bits / (profile.bandwidth_kbps * 1000)
            bandwidth_delay_ms = int(transfer_time_sec * 1000)
            if bandwidth_delay_ms > 0:
                await asyncio.sleep(transfer_time_sec)
                stats["bandwidth_delay_ms"] = bandwidth_delay_ms
        
        return data, stats


# Global throttler instance
network_throttler = NetworkThrottler()


# ============================================================================
# Macro Recorder - Request Sequence Recording and Replay
# ============================================================================

@dataclass
class MacroStep:
    """A single step in a macro sequence"""
    id: str
    order: int
    name: str
    request: Dict[str, Any]  # Request template
    delay_before_ms: int = 0  # Delay before executing this step
    extract_from_response: Optional[Dict[str, str]] = None  # Variables to extract
    condition: Optional[str] = None  # Condition to check before executing
    repeat_count: int = 1
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "order": self.order,
            "name": self.name,
            "request": self.request,
            "delay_before_ms": self.delay_before_ms,
            "extract_from_response": self.extract_from_response,
            "condition": self.condition,
            "repeat_count": self.repeat_count
        }


@dataclass
class Macro:
    """A recorded sequence of requests for replay"""
    id: str
    name: str
    description: str
    steps: List[MacroStep] = field(default_factory=list)
    variables: Dict[str, str] = field(default_factory=dict)  # User-defined variables
    created_at: datetime = field(default_factory=datetime.utcnow)
    last_run: Optional[datetime] = None
    run_count: int = 0
    tags: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "name": self.name,
            "description": self.description,
            "steps": [s.to_dict() for s in self.steps],
            "variables": self.variables,
            "created_at": self.created_at.isoformat(),
            "last_run": self.last_run.isoformat() if self.last_run else None,
            "run_count": self.run_count,
            "tags": self.tags,
            "step_count": len(self.steps)
        }


@dataclass
class MacroRunResult:
    """Result of running a macro"""
    macro_id: str
    started_at: datetime
    completed_at: Optional[datetime]
    status: str  # "running", "completed", "failed", "cancelled"
    steps_completed: int
    steps_total: int
    results: List[Dict[str, Any]]  # Results for each step
    variables: Dict[str, str]  # Final variable values
    error: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "macro_id": self.macro_id,
            "started_at": self.started_at.isoformat(),
            "completed_at": self.completed_at.isoformat() if self.completed_at else None,
            "status": self.status,
            "steps_completed": self.steps_completed,
            "steps_total": self.steps_total,
            "results": self.results,
            "variables": self.variables,
            "error": self.error,
            "duration_ms": int((self.completed_at - self.started_at).total_seconds() * 1000) if self.completed_at else None
        }


class MacroRecorder:
    """
    Record and replay sequences of HTTP requests.
    
    Features:
    - Record traffic as macro steps
    - Variable extraction from responses
    - Variable substitution in requests
    - Conditional execution
    - Repeat steps
    """
    
    def __init__(self):
        self.macros: Dict[str, Macro] = {}
        self.recording: bool = False
        self.recording_macro_id: Optional[str] = None
        self.running_results: Dict[str, MacroRunResult] = {}
        self._lock = threading.Lock()
    
    def start_recording(self, name: str, description: str = "") -> Macro:
        """Start recording a new macro"""
        macro_id = f"macro_{uuid.uuid4().hex[:8]}"
        macro = Macro(
            id=macro_id,
            name=name,
            description=description
        )
        self.macros[macro_id] = macro
        self.recording = True
        self.recording_macro_id = macro_id
        return macro
    
    def stop_recording(self) -> Optional[Macro]:
        """Stop recording and return the macro"""
        if not self.recording or not self.recording_macro_id:
            return None
        
        macro = self.macros.get(self.recording_macro_id)
        self.recording = False
        self.recording_macro_id = None
        return macro
    
    def record_step(
        self,
        request: Dict[str, Any],
        name: Optional[str] = None,
        extract_vars: Optional[Dict[str, str]] = None
    ) -> Optional[MacroStep]:
        """Record a request as a macro step"""
        if not self.recording or not self.recording_macro_id:
            return None
        
        macro = self.macros.get(self.recording_macro_id)
        if not macro:
            return None
        
        step_id = f"step_{uuid.uuid4().hex[:6]}"
        step_order = len(macro.steps) + 1
        
        step = MacroStep(
            id=step_id,
            order=step_order,
            name=name or f"Step {step_order}: {request.get('method', 'GET')} {request.get('path', '/')}",
            request=request,
            extract_from_response=extract_vars
        )
        
        macro.steps.append(step)
        return step
    
    def create_macro_from_traffic(
        self,
        traffic_entries: List[Dict[str, Any]],
        name: str,
        description: str = ""
    ) -> Macro:
        """Create a macro from captured traffic entries"""
        macro_id = f"macro_{uuid.uuid4().hex[:8]}"
        macro = Macro(
            id=macro_id,
            name=name,
            description=description
        )
        
        for i, entry in enumerate(traffic_entries):
            request = entry.get("request", {})
            step = MacroStep(
                id=f"step_{uuid.uuid4().hex[:6]}",
                order=i + 1,
                name=f"Step {i + 1}: {request.get('method', 'GET')} {request.get('path', '/')}",
                request={
                    "method": request.get("method", "GET"),
                    "path": request.get("path", "/"),
                    "headers": request.get("headers", {}),
                    "body": request.get("body")
                }
            )
            macro.steps.append(step)
        
        self.macros[macro_id] = macro
        return macro
    
    def get_macro(self, macro_id: str) -> Optional[Macro]:
        """Get a macro by ID"""
        return self.macros.get(macro_id)
    
    def list_macros(self) -> List[Macro]:
        """List all macros"""
        return list(self.macros.values())
    
    def delete_macro(self, macro_id: str) -> bool:
        """Delete a macro"""
        if macro_id in self.macros:
            del self.macros[macro_id]
            return True
        return False
    
    def update_macro(
        self,
        macro_id: str,
        name: Optional[str] = None,
        description: Optional[str] = None,
        variables: Optional[Dict[str, str]] = None,
        tags: Optional[List[str]] = None
    ) -> Optional[Macro]:
        """Update macro metadata"""
        macro = self.macros.get(macro_id)
        if not macro:
            return None
        
        if name:
            macro.name = name
        if description is not None:
            macro.description = description
        if variables is not None:
            macro.variables = variables
        if tags is not None:
            macro.tags = tags
        
        return macro
    
    def _substitute_variables(self, text: str, variables: Dict[str, str]) -> str:
        """Substitute {{variable}} placeholders with values"""
        result = text
        for key, value in variables.items():
            result = result.replace(f"{{{{{key}}}}}", value)
        return result
    
    def _extract_variables(
        self,
        response: Dict[str, Any],
        extract_config: Dict[str, str]
    ) -> Dict[str, str]:
        """Extract variables from response using JSONPath-like expressions"""
        extracted = {}
        
        for var_name, path in extract_config.items():
            try:
                if path.startswith("header:"):
                    header_name = path[7:]
                    headers = response.get("headers", {})
                    for key, value in headers.items():
                        if key.lower() == header_name.lower():
                            extracted[var_name] = value
                            break
                
                elif path.startswith("body:"):
                    json_path = path[5:]
                    body = response.get("body", "")
                    if isinstance(body, str):
                        try:
                            body = json.loads(body)
                        except json.JSONDecodeError:
                            continue
                    
                    # Simple dot-notation path
                    value = body
                    for part in json_path.split("."):
                        if isinstance(value, dict):
                            value = value.get(part)
                        elif isinstance(value, list) and part.isdigit():
                            value = value[int(part)]
                        else:
                            value = None
                            break
                    
                    if value is not None:
                        extracted[var_name] = str(value)
                
                elif path.startswith("status:"):
                    extracted[var_name] = str(response.get("status_code", ""))
            
            except Exception as e:
                logger.warning(f"Failed to extract variable {var_name}: {e}")
        
        return extracted
    
    async def run_macro(
        self,
        macro_id: str,
        base_url: str,
        initial_variables: Optional[Dict[str, str]] = None,
        timeout_per_step: float = 30.0
    ) -> MacroRunResult:
        """Run a macro and return results"""
        macro = self.macros.get(macro_id)
        if not macro:
            return MacroRunResult(
                macro_id=macro_id,
                started_at=datetime.utcnow(),
                completed_at=datetime.utcnow(),
                status="failed",
                steps_completed=0,
                steps_total=0,
                results=[],
                variables={},
                error="Macro not found"
            )
        
        result = MacroRunResult(
            macro_id=macro_id,
            started_at=datetime.utcnow(),
            completed_at=None,
            status="running",
            steps_completed=0,
            steps_total=len(macro.steps),
            results=[],
            variables={**(macro.variables or {}), **(initial_variables or {})}
        )
        
        self.running_results[macro_id] = result
        
        try:
            import aiohttp
            
            async with aiohttp.ClientSession() as session:
                for step in sorted(macro.steps, key=lambda s: s.order):
                    # Apply delay
                    if step.delay_before_ms > 0:
                        await asyncio.sleep(step.delay_before_ms / 1000.0)
                    
                    # Check condition
                    if step.condition:
                        condition = self._substitute_variables(step.condition, result.variables)
                        # Safe condition check (no eval)
                        if not _safe_evaluate_condition(condition, result.variables):
                            result.results.append({
                                "step_id": step.id,
                                "skipped": True,
                                "reason": "Condition not met"
                            })
                            continue
                    
                    for repeat in range(step.repeat_count):
                        # Prepare request with variable substitution
                        request = step.request.copy()
                        request["path"] = self._substitute_variables(
                            request.get("path", "/"),
                            result.variables
                        )
                        if request.get("body"):
                            request["body"] = self._substitute_variables(
                                request["body"],
                                result.variables
                            )
                        
                        headers = {}
                        for key, value in request.get("headers", {}).items():
                            headers[key] = self._substitute_variables(value, result.variables)
                        
                        url = f"{base_url.rstrip('/')}{request['path']}"
                        method = request.get("method", "GET")
                        
                        step_result = {
                            "step_id": step.id,
                            "step_name": step.name,
                            "repeat": repeat + 1 if step.repeat_count > 1 else None,
                            "request": {
                                "method": method,
                                "url": url,
                                "headers": headers
                            }
                        }
                        
                        try:
                            async with session.request(
                                method,
                                url,
                                headers=headers,
                                data=request.get("body"),
                                timeout=aiohttp.ClientTimeout(total=timeout_per_step),
                                ssl=False
                            ) as resp:
                                response_body = await resp.text()
                                step_result["response"] = {
                                    "status_code": resp.status,
                                    "headers": dict(resp.headers),
                                    "body": response_body[:10000]  # Limit body size
                                }
                                step_result["success"] = True
                                
                                # Extract variables
                                if step.extract_from_response:
                                    extracted = self._extract_variables(
                                        step_result["response"],
                                        step.extract_from_response
                                    )
                                    result.variables.update(extracted)
                                    step_result["extracted_variables"] = extracted
                        
                        except Exception as e:
                            step_result["success"] = False
                            step_result["error"] = str(e)
                        
                        result.results.append(step_result)
                    
                    result.steps_completed += 1
            
            result.status = "completed"
        
        except Exception as e:
            result.status = "failed"
            result.error = str(e)
        
        finally:
            result.completed_at = datetime.utcnow()
            macro.last_run = result.completed_at
            macro.run_count += 1
        
        return result
    
    def get_run_result(self, macro_id: str) -> Optional[MacroRunResult]:
        """Get the latest run result for a macro"""
        return self.running_results.get(macro_id)


# Global macro recorder instance
macro_recorder = MacroRecorder()


# ============================================================================
# HAR (HTTP Archive) Export
# ============================================================================

class HARExporter:
    """
    Export MITM traffic to HAR (HTTP Archive) format.
    HAR files can be imported into browser developer tools.
    """
    
    @staticmethod
    def traffic_to_har(
        traffic_entries: List[Dict[str, Any]],
        creator_name: str = "VRAgent MITM Workbench",
        creator_version: str = "1.0"
    ) -> Dict[str, Any]:
        """Convert traffic entries to HAR format"""
        entries = []
        
        for entry in traffic_entries:
            request = entry.get("request", {})
            response = entry.get("response")
            
            # Build request
            har_request = {
                "method": request.get("method", "GET"),
                "url": request.get("url") or f"http://{request.get('host', 'unknown')}{request.get('path', '/')}",
                "httpVersion": request.get("http_version", "HTTP/1.1"),
                "cookies": HARExporter._parse_cookies(request.get("headers", {}).get("Cookie", "")),
                "headers": [
                    {"name": k, "value": v}
                    for k, v in request.get("headers", {}).items()
                ],
                "queryString": HARExporter._parse_query_string(request.get("path", "")),
                "headersSize": -1,
                "bodySize": len(request.get("body", "")) if request.get("body") else 0
            }
            
            if request.get("body"):
                har_request["postData"] = {
                    "mimeType": request.get("headers", {}).get("Content-Type", "application/octet-stream"),
                    "text": request.get("body") if isinstance(request.get("body"), str) else ""
                }
            
            # Build response
            har_response = {
                "status": 0,
                "statusText": "",
                "httpVersion": "HTTP/1.1",
                "cookies": [],
                "headers": [],
                "content": {
                    "size": 0,
                    "mimeType": "application/octet-stream"
                },
                "redirectURL": "",
                "headersSize": -1,
                "bodySize": 0
            }
            
            if response:
                har_response["status"] = response.get("status_code", 0)
                har_response["statusText"] = response.get("status_text", "")
                har_response["headers"] = [
                    {"name": k, "value": v}
                    for k, v in response.get("headers", {}).items()
                ]
                har_response["cookies"] = HARExporter._parse_set_cookies(response.get("headers", {}))
                
                body = response.get("body", "")
                har_response["content"] = {
                    "size": len(body) if body else 0,
                    "mimeType": response.get("headers", {}).get("Content-Type", "application/octet-stream"),
                    "text": body if isinstance(body, str) else ""
                }
                har_response["bodySize"] = len(body) if body else 0
            
            # Build timing
            duration_ms = entry.get("duration_ms", 0)
            har_timing = {
                "send": 0,
                "wait": duration_ms,
                "receive": 0
            }
            
            # Build entry
            har_entry = {
                "startedDateTime": entry.get("timestamp", datetime.utcnow().isoformat()),
                "time": duration_ms,
                "request": har_request,
                "response": har_response,
                "cache": {},
                "timings": har_timing,
                "serverIPAddress": request.get("host", ""),
                "connection": entry.get("id", "")
            }
            
            # Add custom fields
            if entry.get("modified"):
                har_entry["_modified"] = True
            if entry.get("rules_applied"):
                har_entry["_rulesApplied"] = entry["rules_applied"]
            if entry.get("tags"):
                har_entry["_tags"] = entry["tags"]
            if entry.get("notes"):
                har_entry["_notes"] = entry["notes"]
            
            entries.append(har_entry)
        
        return {
            "log": {
                "version": "1.2",
                "creator": {
                    "name": creator_name,
                    "version": creator_version
                },
                "entries": entries,
                "comment": f"Exported from VRAgent MITM Workbench - {len(entries)} entries"
            }
        }
    
    @staticmethod
    def _parse_cookies(cookie_header: str) -> List[Dict[str, str]]:
        """Parse Cookie header into HAR cookies array"""
        cookies = []
        if cookie_header:
            for part in cookie_header.split(";"):
                if "=" in part:
                    name, value = part.strip().split("=", 1)
                    cookies.append({"name": name, "value": value})
        return cookies
    
    @staticmethod
    def _parse_set_cookies(headers: Dict[str, str]) -> List[Dict[str, Any]]:
        """Parse Set-Cookie headers into HAR cookies array"""
        cookies = []
        for key, value in headers.items():
            if key.lower() == "set-cookie":
                cookie = {"name": "", "value": ""}
                parts = value.split(";")
                if parts:
                    name_value = parts[0]
                    if "=" in name_value:
                        cookie["name"], cookie["value"] = name_value.split("=", 1)
                    
                    for part in parts[1:]:
                        part = part.strip()
                        if part.lower().startswith("expires="):
                            cookie["expires"] = part[8:]
                        elif part.lower().startswith("path="):
                            cookie["path"] = part[5:]
                        elif part.lower().startswith("domain="):
                            cookie["domain"] = part[7:]
                        elif part.lower() == "httponly":
                            cookie["httpOnly"] = True
                        elif part.lower() == "secure":
                            cookie["secure"] = True
                
                cookies.append(cookie)
        return cookies
    
    @staticmethod
    def _parse_query_string(path: str) -> List[Dict[str, str]]:
        """Parse query string from path"""
        params = []
        if "?" in path:
            query = path.split("?", 1)[1]
            for part in query.split("&"):
                if "=" in part:
                    name, value = part.split("=", 1)
                    params.append({"name": name, "value": value})
                else:
                    params.append({"name": part, "value": ""})
        return params


# Global HAR exporter instance
har_exporter = HARExporter()


# ============================================================================
# Custom Protocol Decoder Plugin System
# ============================================================================

@dataclass
class ProtocolDecoder:
    """Base class for custom protocol decoders"""
    id: str
    name: str
    description: str
    content_types: List[str] = field(default_factory=list)  # Matching Content-Types
    magic_bytes: Optional[bytes] = None  # Magic bytes to detect the protocol
    enabled: bool = True
    is_builtin: bool = False
    
    def can_decode(self, content_type: str, data: bytes) -> bool:
        """Check if this decoder can handle the data"""
        if self.content_types:
            for ct in self.content_types:
                if ct.lower() in content_type.lower():
                    return True
        if self.magic_bytes and data.startswith(self.magic_bytes):
            return True
        return False
    
    def decode(self, data: bytes) -> Dict[str, Any]:
        """Decode the data - override in subclasses"""
        return {"raw": base64.b64encode(data).decode()}
    
    def encode(self, data: Dict[str, Any]) -> bytes:
        """Encode data back - override in subclasses"""
        if "raw" in data:
            return base64.b64decode(data["raw"])
        return b""
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "name": self.name,
            "description": self.description,
            "content_types": self.content_types,
            "enabled": self.enabled,
            "is_builtin": self.is_builtin
        }


class MessagePackDecoder(ProtocolDecoder):
    """Decoder for MessagePack binary format"""
    
    def __init__(self):
        super().__init__(
            id="msgpack",
            name="MessagePack",
            description="Decode MessagePack binary serialization format",
            content_types=["application/msgpack", "application/x-msgpack"],
            is_builtin=True
        )
    
    def decode(self, data: bytes) -> Dict[str, Any]:
        try:
            import msgpack
            decoded = msgpack.unpackb(data, raw=False)
            return {"decoded": decoded, "format": "msgpack"}
        except ImportError:
            return {"error": "msgpack library not installed", "raw": base64.b64encode(data).decode()}
        except Exception as e:
            return {"error": str(e), "raw": base64.b64encode(data).decode()}
    
    def encode(self, data: Dict[str, Any]) -> bytes:
        try:
            import msgpack
            return msgpack.packb(data.get("decoded", data))
        except (TypeError, ValueError, ImportError):
            return super().encode(data)


class ProtobufDecoder(ProtocolDecoder):
    """Decoder for Protocol Buffers (without schema - raw field parsing)"""
    
    def __init__(self):
        super().__init__(
            id="protobuf",
            name="Protocol Buffers",
            description="Decode Protocol Buffers messages (schema-less)",
            content_types=["application/protobuf", "application/x-protobuf", "application/grpc"],
            is_builtin=True
        )
    
    def decode(self, data: bytes) -> Dict[str, Any]:
        """Decode protobuf without schema (best effort field extraction)"""
        try:
            fields = []
            pos = 0
            
            while pos < len(data):
                if pos >= len(data):
                    break
                
                # Read tag (varint)
                tag_byte = data[pos]
                field_number = tag_byte >> 3
                wire_type = tag_byte & 0x07
                pos += 1
                
                field = {"field_number": field_number, "wire_type": wire_type}
                
                if wire_type == 0:  # Varint
                    value = 0
                    shift = 0
                    while pos < len(data):
                        b = data[pos]
                        pos += 1
                        value |= (b & 0x7F) << shift
                        if not (b & 0x80):
                            break
                        shift += 7
                    field["value"] = value
                    field["type"] = "varint"
                
                elif wire_type == 1:  # 64-bit
                    if pos + 8 <= len(data):
                        field["value"] = struct.unpack("<Q", data[pos:pos+8])[0]
                        field["type"] = "fixed64"
                        pos += 8
                
                elif wire_type == 2:  # Length-delimited
                    length = 0
                    shift = 0
                    while pos < len(data):
                        b = data[pos]
                        pos += 1
                        length |= (b & 0x7F) << shift
                        if not (b & 0x80):
                            break
                        shift += 7
                    
                    if pos + length <= len(data):
                        raw_value = data[pos:pos+length]
                        try:
                            field["value"] = raw_value.decode('utf-8')
                            field["type"] = "string"
                        except:
                            field["value"] = base64.b64encode(raw_value).decode()
                            field["type"] = "bytes"
                        pos += length
                
                elif wire_type == 5:  # 32-bit
                    if pos + 4 <= len(data):
                        field["value"] = struct.unpack("<I", data[pos:pos+4])[0]
                        field["type"] = "fixed32"
                        pos += 4
                
                else:
                    field["error"] = f"Unknown wire type: {wire_type}"
                    break
                
                fields.append(field)
            
            return {"format": "protobuf", "fields": fields}
        
        except Exception as e:
            return {"error": str(e), "raw": base64.b64encode(data).decode()}


class BSONDecoder(ProtocolDecoder):
    """Decoder for BSON (Binary JSON) format"""
    
    def __init__(self):
        super().__init__(
            id="bson",
            name="BSON",
            description="Decode BSON (Binary JSON) format used by MongoDB",
            content_types=["application/bson"],
            is_builtin=True
        )
    
    def decode(self, data: bytes) -> Dict[str, Any]:
        try:
            import bson
            decoded = bson.decode(data)
            return {"decoded": decoded, "format": "bson"}
        except ImportError:
            return {"error": "bson library not installed", "raw": base64.b64encode(data).decode()}
        except Exception as e:
            return {"error": str(e), "raw": base64.b64encode(data).decode()}


class ProtocolDecoderManager:
    """Manages protocol decoders for custom protocol support"""
    
    def __init__(self):
        self.decoders: Dict[str, ProtocolDecoder] = {}
        self.custom_decoders: Dict[str, ProtocolDecoder] = {}
        self._init_builtin_decoders()
    
    def _init_builtin_decoders(self):
        """Initialize built-in decoders"""
        decoders = [
            MessagePackDecoder(),
            ProtobufDecoder(),
            BSONDecoder()
        ]
        for decoder in decoders:
            self.decoders[decoder.id] = decoder
    
    def get_decoder(self, decoder_id: str) -> Optional[ProtocolDecoder]:
        """Get a decoder by ID"""
        return self.decoders.get(decoder_id) or self.custom_decoders.get(decoder_id)
    
    def get_all_decoders(self) -> List[ProtocolDecoder]:
        """Get all available decoders"""
        return list(self.decoders.values()) + list(self.custom_decoders.values())
    
    def find_decoder(self, content_type: str, data: bytes) -> Optional[ProtocolDecoder]:
        """Find a decoder that can handle the given data"""
        for decoder in self.get_all_decoders():
            if decoder.enabled and decoder.can_decode(content_type, data):
                return decoder
        return None
    
    def decode(self, content_type: str, data: bytes) -> Dict[str, Any]:
        """Attempt to decode data using available decoders"""
        decoder = self.find_decoder(content_type, data)
        if decoder:
            result = decoder.decode(data)
            result["decoder_id"] = decoder.id
            result["decoder_name"] = decoder.name
            return result
        return {"raw": base64.b64encode(data).decode(), "decoder_id": None}
    
    def register_decoder(self, decoder: ProtocolDecoder) -> None:
        """Register a custom decoder"""
        self.custom_decoders[decoder.id] = decoder
    
    def unregister_decoder(self, decoder_id: str) -> bool:
        """Unregister a custom decoder"""
        if decoder_id in self.custom_decoders:
            del self.custom_decoders[decoder_id]
            return True
        return False


# Global protocol decoder manager
protocol_decoder_manager = ProtocolDecoderManager()


# ============================================================================
# Collaborative Session Sharing
# ============================================================================

@dataclass
class SharedSession:
    """A shared proxy session for collaborative testing"""
    id: str
    proxy_id: str
    name: str
    description: str
    owner_id: str
    owner_name: str
    shared_with: List[str] = field(default_factory=list)  # User IDs
    access_level: str = "view"  # "view", "interact", "full"
    created_at: datetime = field(default_factory=datetime.utcnow)
    expires_at: Optional[datetime] = None
    active_viewers: List[str] = field(default_factory=list)
    share_token: Optional[str] = None  # For link sharing
    settings: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "proxy_id": self.proxy_id,
            "name": self.name,
            "description": self.description,
            "owner_id": self.owner_id,
            "owner_name": self.owner_name,
            "shared_with": self.shared_with,
            "access_level": self.access_level,
            "created_at": self.created_at.isoformat(),
            "expires_at": self.expires_at.isoformat() if self.expires_at else None,
            "active_viewers": self.active_viewers,
            "share_token": self.share_token,
            "viewer_count": len(self.active_viewers),
            "settings": self.settings
        }


class SessionSharingManager:
    """Manages collaborative session sharing"""
    
    def __init__(self):
        self.shared_sessions: Dict[str, SharedSession] = {}
        self.token_to_session: Dict[str, str] = {}  # share_token -> session_id
        self._lock = threading.Lock()
    
    def create_shared_session(
        self,
        proxy_id: str,
        name: str,
        owner_id: str,
        owner_name: str,
        description: str = "",
        access_level: str = "view",
        expires_hours: Optional[int] = None,
        enable_link_sharing: bool = False
    ) -> SharedSession:
        """Create a new shared session"""
        session_id = f"share_{uuid.uuid4().hex[:12]}"
        
        expires_at = None
        if expires_hours:
            expires_at = datetime.utcnow() + timedelta(hours=expires_hours)
        
        share_token = None
        if enable_link_sharing:
            share_token = uuid.uuid4().hex
            self.token_to_session[share_token] = session_id
        
        session = SharedSession(
            id=session_id,
            proxy_id=proxy_id,
            name=name,
            description=description,
            owner_id=owner_id,
            owner_name=owner_name,
            access_level=access_level,
            expires_at=expires_at,
            share_token=share_token
        )
        
        with self._lock:
            self.shared_sessions[session_id] = session
        
        return session
    
    def get_session(self, session_id: str) -> Optional[SharedSession]:
        """Get a shared session by ID"""
        session = self.shared_sessions.get(session_id)
        if session and session.expires_at and datetime.utcnow() > session.expires_at:
            self.delete_session(session_id)
            return None
        return session
    
    def get_session_by_token(self, token: str) -> Optional[SharedSession]:
        """Get a shared session by share token"""
        session_id = self.token_to_session.get(token)
        if session_id:
            return self.get_session(session_id)
        return None
    
    def list_sessions_for_user(self, user_id: str) -> List[SharedSession]:
        """List sessions owned by or shared with a user"""
        sessions = []
        for session in self.shared_sessions.values():
            if session.owner_id == user_id or user_id in session.shared_with:
                if not session.expires_at or datetime.utcnow() <= session.expires_at:
                    sessions.append(session)
        return sessions
    
    def list_sessions_for_proxy(self, proxy_id: str) -> List[SharedSession]:
        """List all shared sessions for a proxy"""
        return [s for s in self.shared_sessions.values() if s.proxy_id == proxy_id]
    
    def share_with_user(self, session_id: str, user_id: str, by_owner: str) -> bool:
        """Share a session with another user"""
        session = self.get_session(session_id)
        if not session or session.owner_id != by_owner:
            return False
        
        if user_id not in session.shared_with:
            session.shared_with.append(user_id)
        return True
    
    def revoke_user_access(self, session_id: str, user_id: str, by_owner: str) -> bool:
        """Revoke a user's access to a session"""
        session = self.get_session(session_id)
        if not session or session.owner_id != by_owner:
            return False
        
        if user_id in session.shared_with:
            session.shared_with.remove(user_id)
        if user_id in session.active_viewers:
            session.active_viewers.remove(user_id)
        return True
    
    def join_session(self, session_id: str, user_id: str) -> bool:
        """Join a shared session as a viewer"""
        session = self.get_session(session_id)
        if not session:
            return False
        
        if session.owner_id != user_id and user_id not in session.shared_with:
            return False
        
        if user_id not in session.active_viewers:
            session.active_viewers.append(user_id)
        return True
    
    def leave_session(self, session_id: str, user_id: str) -> bool:
        """Leave a shared session"""
        session = self.get_session(session_id)
        if session and user_id in session.active_viewers:
            session.active_viewers.remove(user_id)
            return True
        return False
    
    def update_session(
        self,
        session_id: str,
        owner_id: str,
        name: Optional[str] = None,
        description: Optional[str] = None,
        access_level: Optional[str] = None,
        settings: Optional[Dict[str, Any]] = None
    ) -> Optional[SharedSession]:
        """Update session settings"""
        session = self.get_session(session_id)
        if not session or session.owner_id != owner_id:
            return None
        
        if name:
            session.name = name
        if description is not None:
            session.description = description
        if access_level:
            session.access_level = access_level
        if settings is not None:
            session.settings.update(settings)
        
        return session
    
    def delete_session(self, session_id: str, owner_id: Optional[str] = None) -> bool:
        """Delete a shared session"""
        session = self.shared_sessions.get(session_id)
        if not session:
            return False
        
        if owner_id and session.owner_id != owner_id:
            return False
        
        with self._lock:
            if session.share_token and session.share_token in self.token_to_session:
                del self.token_to_session[session.share_token]
            del self.shared_sessions[session_id]
        
        return True
    
    def check_access(self, session_id: str, user_id: str, required_level: str = "view") -> bool:
        """Check if a user has required access level"""
        session = self.get_session(session_id)
        if not session:
            return False
        
        if session.owner_id == user_id:
            return True
        
        if user_id not in session.shared_with:
            return False
        
        level_order = ["view", "interact", "full"]
        session_level_idx = level_order.index(session.access_level) if session.access_level in level_order else 0
        required_level_idx = level_order.index(required_level) if required_level in level_order else 0
        
        return session_level_idx >= required_level_idx


# Global session sharing manager
session_sharing_manager = SessionSharingManager()


# ============================================================================
# WebSocket Deep Inspection Implementation
# ============================================================================

class WebSocketInspector:
    """
    WebSocket frame parser and inspector for deep inspection of WebSocket traffic.
    
    Features:
    - Parse WebSocket frames (text, binary, control)
    - Track active WebSocket connections
    - Apply rules to WebSocket messages
    - JSON parsing and modification for text frames
    """
    
    def __init__(self):
        self.connections: Dict[str, WebSocketConnection] = {}
        self.rules: List[WebSocketRule] = []
        self._lock = threading.Lock()
        
        # Callbacks
        self.on_frame: Optional[Callable[[WebSocketFrame, str], None]] = None
        self.on_connection_open: Optional[Callable[[WebSocketConnection], None]] = None
        self.on_connection_close: Optional[Callable[[WebSocketConnection], None]] = None
        
        # Stats
        self.stats = {
            "frames_client_to_server": 0,
            "frames_server_to_client": 0,
            "bytes_client_to_server": 0,
            "bytes_server_to_client": 0,
            "text_frames": 0,
            "binary_frames": 0,
            "control_frames": 0,
            "rules_applied": 0
        }
    
    def register_connection(
        self,
        connection_id: str,
        proxy_id: str,
        client_ip: str,
        client_port: int,
        target_host: str,
        target_port: int,
        upgrade_request_id: str
    ) -> WebSocketConnection:
        """Register a new WebSocket connection for tracking"""
        conn = WebSocketConnection(
            id=connection_id,
            proxy_id=proxy_id,
            created_at=datetime.now(),
            client_ip=client_ip,
            client_port=client_port,
            target_host=target_host,
            target_port=target_port,
            upgrade_request_id=upgrade_request_id
        )
        
        with self._lock:
            self.connections[connection_id] = conn
        
        if self.on_connection_open:
            self.on_connection_open(conn)
        
        logger.info(f"WebSocket connection registered: {connection_id}")
        return conn
    
    def close_connection(
        self,
        connection_id: str,
        close_code: Optional[int] = None,
        close_reason: Optional[str] = None
    ) -> None:
        """Mark a WebSocket connection as closed"""
        with self._lock:
            conn = self.connections.get(connection_id)
            if conn:
                conn.status = "closed"
                conn.closed_at = datetime.now()
                conn.close_code = close_code
                conn.close_reason = close_reason
        
        if conn and self.on_connection_close:
            self.on_connection_close(conn)
        
        logger.info(f"WebSocket connection closed: {connection_id}, code={close_code}")
    
    def parse_frame(self, data: bytes, direction: str, connection_id: str) -> Tuple[Optional[WebSocketFrame], bytes]:
        """
        Parse a WebSocket frame from raw bytes.
        
        Returns:
            Tuple of (parsed_frame, remaining_bytes)
            If frame is incomplete, returns (None, original_data)
        """
        if len(data) < 2:
            return None, data
        
        # First byte: FIN + RSV + Opcode
        byte1 = data[0]
        fin = bool(byte1 & 0x80)
        opcode = byte1 & 0x0F
        
        # Second byte: MASK + Payload length
        byte2 = data[1]
        masked = bool(byte2 & 0x80)
        payload_len = byte2 & 0x7F
        
        header_len = 2
        
        # Extended payload length
        if payload_len == 126:
            if len(data) < 4:
                return None, data
            payload_len = struct.unpack(">H", data[2:4])[0]
            header_len = 4
        elif payload_len == 127:
            if len(data) < 10:
                return None, data
            payload_len = struct.unpack(">Q", data[2:10])[0]
            header_len = 10
        
        # Masking key (if masked)
        mask_key = None
        if masked:
            if len(data) < header_len + 4:
                return None, data
            mask_key = data[header_len:header_len + 4]
            header_len += 4
        
        # Check if we have the full payload
        total_len = header_len + payload_len
        if len(data) < total_len:
            return None, data
        
        # Extract payload
        payload = bytearray(data[header_len:total_len])
        
        # Unmask if necessary
        if masked and mask_key:
            for i in range(len(payload)):
                payload[i] ^= mask_key[i % 4]
        
        payload = bytes(payload)
        remaining = data[total_len:]
        
        # Determine opcode name
        opcode_names = {
            0x0: "CONTINUATION",
            0x1: "TEXT",
            0x2: "BINARY",
            0x8: "CLOSE",
            0x9: "PING",
            0xA: "PONG"
        }
        opcode_name = opcode_names.get(opcode, f"UNKNOWN({opcode})")
        
        # Parse payload for text frames
        payload_text = None
        payload_json = None
        
        if opcode == 0x1:  # TEXT
            try:
                payload_text = payload.decode('utf-8')
                try:
                    payload_json = json.loads(payload_text)
                except json.JSONDecodeError:
                    pass
            except UnicodeDecodeError:
                payload_text = payload.decode('utf-8', errors='replace')
        
        # Parse close frame
        if opcode == 0x8 and len(payload) >= 2:
            close_code = struct.unpack(">H", payload[:2])[0]
            close_reason = payload[2:].decode('utf-8', errors='replace') if len(payload) > 2 else None
            payload_text = f"Close Code: {close_code}"
            if close_reason:
                payload_text += f", Reason: {close_reason}"
        
        frame = WebSocketFrame(
            id=str(uuid.uuid4()),
            timestamp=datetime.now(),
            direction=direction,
            opcode=opcode,
            opcode_name=opcode_name,
            fin=fin,
            masked=masked,
            payload_length=payload_len,
            payload=payload,
            payload_text=payload_text,
            payload_json=payload_json,
            is_control=opcode >= 0x8,
            connection_id=connection_id
        )
        
        return frame, remaining
    
    def encode_frame(
        self,
        payload: bytes,
        opcode: int = 0x1,
        fin: bool = True,
        masked: bool = False
    ) -> bytes:
        """Encode a WebSocket frame"""
        frame = bytearray()
        
        # First byte
        byte1 = (0x80 if fin else 0x00) | opcode
        frame.append(byte1)
        
        # Payload length
        payload_len = len(payload)
        if payload_len <= 125:
            byte2 = (0x80 if masked else 0x00) | payload_len
            frame.append(byte2)
        elif payload_len <= 65535:
            byte2 = (0x80 if masked else 0x00) | 126
            frame.append(byte2)
            frame.extend(struct.pack(">H", payload_len))
        else:
            byte2 = (0x80 if masked else 0x00) | 127
            frame.append(byte2)
            frame.extend(struct.pack(">Q", payload_len))
        
        # Mask key and payload
        if masked:
            import os
            mask_key = os.urandom(4)
            frame.extend(mask_key)
            masked_payload = bytearray(payload)
            for i in range(len(masked_payload)):
                masked_payload[i] ^= mask_key[i % 4]
            frame.extend(masked_payload)
        else:
            frame.extend(payload)
        
        return bytes(frame)
    
    def process_frame(
        self,
        frame: WebSocketFrame,
        apply_rules: bool = True
    ) -> Tuple[WebSocketFrame, bool]:
        """
        Process a WebSocket frame, optionally applying rules.
        
        Returns:
            Tuple of (processed_frame, should_drop)
        """
        should_drop = False
        
        # Track stats
        if frame.direction == "client_to_server":
            self.stats["frames_client_to_server"] += 1
            self.stats["bytes_client_to_server"] += frame.payload_length
        else:
            self.stats["frames_server_to_client"] += 1
            self.stats["bytes_server_to_client"] += frame.payload_length
        
        if frame.is_control:
            self.stats["control_frames"] += 1
        elif frame.opcode == 0x1:
            self.stats["text_frames"] += 1
        elif frame.opcode == 0x2:
            self.stats["binary_frames"] += 1
        
        # Apply rules
        if apply_rules:
            for rule in sorted(self.rules, key=lambda r: r.priority):
                if not rule.enabled:
                    continue
                
                matches, _ = self._rule_matches_frame(rule, frame)
                if not matches:
                    continue
                
                rule.hit_count += 1
                self.stats["rules_applied"] += 1
                
                # Handle delay
                if rule.delay_ms > 0:
                    time.sleep(rule.delay_ms / 1000.0)
                
                # Handle drop
                if rule.action == "drop":
                    should_drop = True
                    break
                
                # Handle modify
                if rule.action == "modify":
                    frame = self._apply_rule_to_frame(frame, rule)
        
        # Record frame in connection
        with self._lock:
            conn = self.connections.get(frame.connection_id)
            if conn:
                conn.frames.append(frame)
                conn.total_frames += 1
                if frame.direction == "client_to_server":
                    conn.bytes_sent += frame.payload_length
                else:
                    conn.bytes_received += frame.payload_length
                
                # Handle close frame
                if frame.opcode == 0x8:
                    if len(frame.payload) >= 2:
                        conn.close_code = struct.unpack(">H", frame.payload[:2])[0]
                        if len(frame.payload) > 2:
                            conn.close_reason = frame.payload[2:].decode('utf-8', errors='replace')
        
        # Callback
        if self.on_frame:
            self.on_frame(frame, frame.direction)
        
        return frame, should_drop
    
    def _rule_matches_frame(
        self,
        rule: WebSocketRule,
        frame: WebSocketFrame
    ) -> Tuple[bool, Dict[str, Any]]:
        """Check if a rule matches a frame"""
        context: Dict[str, Any] = {}
        
        # Direction match
        if rule.match_direction != "both":
            if rule.match_direction != frame.direction:
                return False, {}
        
        # Opcode match
        if rule.match_opcode is not None:
            if frame.opcode != rule.match_opcode:
                return False, {}
        
        # Payload pattern match
        if rule.match_payload_pattern and frame.payload_text:
            match = re.search(rule.match_payload_pattern, frame.payload_text, re.IGNORECASE)
            if not match:
                return False, {}
            context["payload_match"] = match
        
        # JSON path match
        if rule.match_json_path and frame.payload_json:
            value = self._get_json_path_value(frame.payload_json, rule.match_json_path)
            if value is None:
                return False, {}
            context["json_value"] = value
        
        return True, context
    
    def _apply_rule_to_frame(
        self,
        frame: WebSocketFrame,
        rule: WebSocketRule
    ) -> WebSocketFrame:
        """Apply rule modifications to a frame"""
        # Only modify text/binary frames
        if frame.opcode not in (0x1, 0x2):
            return frame
        
        frame.modified = True
        frame.original_payload = frame.payload
        
        payload_text = frame.payload_text
        payload_json = frame.payload_json
        
        # Apply find/replace
        if rule.payload_find_replace and payload_text:
            for find, replace in rule.payload_find_replace.items():
                payload_text = payload_text.replace(find, replace)
            frame.payload_text = payload_text
            try:
                frame.payload_json = json.loads(payload_text)
                payload_json = frame.payload_json
            except json.JSONDecodeError:
                pass
        
        # Apply JSON path edits
        if rule.json_path_edits and payload_json:
            for edit in rule.json_path_edits:
                path = edit.get("path", "")
                op = edit.get("op", "set")
                value = edit.get("value")
                payload_json = self._apply_json_path_edit(payload_json, path, op, value)
            frame.payload_json = payload_json
            payload_text = json.dumps(payload_json)
            frame.payload_text = payload_text
        
        # Update payload bytes
        if payload_text:
            frame.payload = payload_text.encode('utf-8')
            frame.payload_length = len(frame.payload)
        
        return frame
    
    def _get_json_path_value(self, data: Any, path: str) -> Optional[Any]:
        """Get a value from JSON data using a simple path notation"""
        if not path:
            return data
        
        parts = path.split('.')
        current = data
        
        for part in parts:
            # Handle array index
            if '[' in part:
                name, rest = part.split('[', 1)
                idx = int(rest.rstrip(']'))
                if name and isinstance(current, dict):
                    current = current.get(name)
                if isinstance(current, list) and idx < len(current):
                    current = current[idx]
                else:
                    return None
            elif isinstance(current, dict):
                current = current.get(part)
            else:
                return None
            
            if current is None:
                return None
        
        return current
    
    def _apply_json_path_edit(
        self,
        data: Any,
        path: str,
        op: str,
        value: Any
    ) -> Any:
        """Apply an edit to JSON data at the given path"""
        if not path:
            return value if op == "set" else data
        
        parts = path.split('.')
        
        def navigate_and_edit(current: Any, parts_remaining: List[str]) -> Any:
            if not parts_remaining:
                if op == "set":
                    return value
                elif op == "remove":
                    return None
                return current
            
            part = parts_remaining[0]
            
            # Handle array index
            if '[' in part:
                name, rest = part.split('[', 1)
                idx = int(rest.rstrip(']'))
                
                if name:
                    if not isinstance(current, dict):
                        current = {}
                    if name not in current:
                        current[name] = []
                    target = current[name]
                else:
                    target = current
                
                if isinstance(target, list):
                    while len(target) <= idx:
                        target.append(None)
                    target[idx] = navigate_and_edit(target[idx], parts_remaining[1:])
                    if name:
                        current[name] = target
                    else:
                        current = target
            elif isinstance(current, dict):
                current[part] = navigate_and_edit(current.get(part, {}), parts_remaining[1:])
            
            return current
        
        return navigate_and_edit(data, parts)
    
    def get_connection(self, connection_id: str) -> Optional[WebSocketConnection]:
        """Get a WebSocket connection by ID"""
        return self.connections.get(connection_id)
    
    def get_connections(self, proxy_id: Optional[str] = None) -> List[WebSocketConnection]:
        """Get all WebSocket connections, optionally filtered by proxy"""
        with self._lock:
            if proxy_id:
                return [c for c in self.connections.values() if c.proxy_id == proxy_id]
            return list(self.connections.values())
    
    def get_frames(
        self,
        connection_id: str,
        limit: int = 100,
        offset: int = 0
    ) -> List[Dict[str, Any]]:
        """Get frames for a connection"""
        conn = self.connections.get(connection_id)
        if not conn:
            return []
        
        frames = conn.frames[-(offset + limit):]
        if offset > 0:
            frames = frames[:-offset]
        frames = frames[-limit:]
        
        return [self._frame_to_dict(f) for f in frames]
    
    def _frame_to_dict(self, frame: WebSocketFrame) -> Dict[str, Any]:
        """Convert a frame to a dictionary for JSON serialization"""
        return {
            "id": frame.id,
            "timestamp": frame.timestamp.isoformat(),
            "direction": frame.direction,
            "opcode": frame.opcode,
            "opcode_name": frame.opcode_name,
            "fin": frame.fin,
            "masked": frame.masked,
            "payload_length": frame.payload_length,
            "payload_text": frame.payload_text[:10000] if frame.payload_text else None,
            "payload_json": frame.payload_json,
            "payload_hex": frame.payload[:1000].hex() if not frame.payload_text and frame.payload else None,
            "is_control": frame.is_control,
            "connection_id": frame.connection_id,
            "modified": frame.modified
        }
    
    def add_rule(self, rule: WebSocketRule) -> None:
        """Add a WebSocket interception rule"""
        self.rules.append(rule)
    
    def remove_rule(self, rule_id: str) -> None:
        """Remove a WebSocket rule"""
        self.rules = [r for r in self.rules if r.id != rule_id]
    
    def get_rules(self) -> List[Dict[str, Any]]:
        """Get all WebSocket rules"""
        return [
            {
                "id": r.id,
                "name": r.name,
                "enabled": r.enabled,
                "priority": r.priority,
                "match_direction": r.match_direction,
                "match_opcode": r.match_opcode,
                "match_payload_pattern": r.match_payload_pattern,
                "match_json_path": r.match_json_path,
                "action": r.action,
                "payload_find_replace": r.payload_find_replace,
                "json_path_edits": r.json_path_edits,
                "delay_ms": r.delay_ms,
                "hit_count": r.hit_count
            }
            for r in self.rules
        ]
    
    def get_stats(self) -> Dict[str, Any]:
        """Get WebSocket inspection statistics"""
        return {
            **self.stats,
            "active_connections": sum(1 for c in self.connections.values() if c.status == "active"),
            "total_connections": len(self.connections)
        }


# Global WebSocket inspector instance
websocket_inspector = WebSocketInspector()


class MITMService:
    """Service for managing MITM proxy instances"""
    
    def __init__(self):
        self.proxies: Dict[str, MITMProxy] = {}
        self._lock = threading.Lock()
        from backend.core.config import settings
        self.traffic_store = TrafficStore(Path(settings.mitm_storage_dir), settings.mitm_max_entries)
    
    def create_proxy(
        self,
        proxy_id: str,
        listen_host: str = "127.0.0.1",
        listen_port: int = 8080,
        target_host: str = "localhost",
        target_port: int = 80,
        mode: str = "auto_modify",
        tls_enabled: bool = False
    ) -> Dict:
        """Create a new MITM proxy"""
        with self._lock:
            if proxy_id in self.proxies:
                raise ValueError(f"Proxy {proxy_id} already exists")
            
            proxy = MITMProxy(
                listen_host=listen_host,
                listen_port=listen_port,
                target_host=target_host,
                target_port=target_port,
                mode=InterceptionMode(mode),
                tls_enabled=tls_enabled
            )

            def handle_entry(entry: TrafficEntry, event_type: str) -> None:
                entry_dict = proxy._traffic_entry_to_dict(entry)
                if event_type == "request":
                    self.traffic_store.append_entry(proxy_id, entry_dict)
                else:
                    self.traffic_store.update_entry(proxy_id, entry.id, entry_dict)
                mitm_stream_manager.emit(proxy_id, {"type": "traffic", "entry": entry_dict})
                mitm_stream_manager.emit(proxy_id, {"type": "stats", "stats": proxy.get_stats()})

            proxy.on_entry = handle_entry
            
            self.proxies[proxy_id] = proxy
            
            return {
                "id": proxy_id,
                "listen_host": listen_host,
                "listen_port": listen_port,
                "target_host": target_host,
                "target_port": target_port,
                "mode": mode,
                "tls_enabled": tls_enabled,
                "status": "created"
            }
    
    def start_proxy(self, proxy_id: str) -> Dict:
        """Start a proxy"""
        proxy = self._get_proxy(proxy_id)
        proxy.start()
        mitm_stream_manager.emit(proxy_id, {"type": "status", "running": True})
        return {"status": "started", "id": proxy_id}
    
    def stop_proxy(self, proxy_id: str) -> Dict:
        """Stop a proxy"""
        proxy = self._get_proxy(proxy_id)
        proxy.stop()
        mitm_stream_manager.emit(proxy_id, {"type": "status", "running": False})
        return {"status": "stopped", "id": proxy_id}
    
    def delete_proxy(self, proxy_id: str) -> Dict:
        """Delete a proxy"""
        with self._lock:
            if proxy_id in self.proxies:
                self.proxies[proxy_id].stop()
                del self.proxies[proxy_id]
        mitm_stream_manager.emit(proxy_id, {"type": "status", "deleted": True})
        return {"status": "deleted", "id": proxy_id}
    
    def get_proxy_status(self, proxy_id: str) -> Dict:
        """Get proxy status and stats"""
        proxy = self._get_proxy(proxy_id)
        stats = proxy.get_stats()
        return {
            "id": proxy_id,
            "listen_host": proxy.listen_host,
            "listen_port": proxy.listen_port,
            "target_host": proxy.target_host,
            "target_port": proxy.target_port,
            "mode": proxy.mode.value,
            "tls_enabled": proxy.tls_enabled,
            "running": proxy.running,
            "stats": {
                "requests": stats.get("requests", 0),
                "responses": stats.get("responses", 0),
                "bytes_sent": stats.get("bytes_sent", 0),
                "bytes_received": stats.get("bytes_received", 0),
                "errors": stats.get("errors", 0),
                "rules_applied": stats.get("rules_applied", 0),
            },
            **stats
        }
    
    def list_proxies(self) -> List[Dict]:
        """List all proxies"""
        with self._lock:
            return [
                {
                    "id": pid,
                    "listen_host": p.listen_host,
                    "listen_port": p.listen_port,
                    "target_host": p.target_host,
                    "target_port": p.target_port,
                    "running": p.running,
                    "mode": p.mode.value,
                    "tls_enabled": p.tls_enabled,
                    "stats": {
                        "requests": p.stats.get("requests_total", 0),
                        "responses": p.stats.get("responses_total", 0),
                        "bytes_sent": p.stats.get("bytes_sent", 0),
                        "bytes_received": p.stats.get("bytes_received", 0),
                        "errors": p.stats.get("errors", 0),
                        "rules_applied": p.stats.get("rules_applied", 0),
                    }
                }
                for pid, p in self.proxies.items()
            ]
    
    def get_traffic(self, proxy_id: str, limit: int = 100, offset: int = 0) -> Dict:
        """Get traffic log for a proxy"""
        self._get_proxy(proxy_id)
        entries = self.traffic_store.list_entries(proxy_id, limit, offset)
        total = self.traffic_store.count_entries(proxy_id)
        return {
            "entries": entries,
            "total": total
        }

    def get_traffic_range(self, proxy_id: str, start: int, end: Optional[int] = None) -> Dict:
        """Get traffic log for a proxy by absolute index range."""
        self._get_proxy(proxy_id)
        entries = self.traffic_store.list_entries_range(proxy_id, start, end)
        total = self.traffic_store.count_entries(proxy_id)
        return {
            "entries": entries,
            "total": total
        }
    
    def clear_traffic(self, proxy_id: str) -> Dict:
        """Clear traffic log for a proxy"""
        proxy = self._get_proxy(proxy_id)
        proxy.clear_traffic_log()
        self.traffic_store.clear_entries(proxy_id)
        return {"status": "cleared", "id": proxy_id}

    def update_traffic_entry(self, proxy_id: str, entry_id: str, notes: Optional[str] = None, tags: Optional[List[str]] = None) -> Dict:
        """Update notes or tags for a traffic entry"""
        proxy = self._get_proxy(proxy_id)
        entry = proxy.update_traffic_entry(entry_id, notes=notes, tags=tags)
        updated_dict = proxy._traffic_entry_to_dict(entry)
        stored = self.traffic_store.update_entry(proxy_id, entry_id, updated_dict)
        return stored or updated_dict

    def save_session(self, proxy_id: str, name: Optional[str] = None) -> Dict:
        """Save current traffic log as a session snapshot."""
        self._get_proxy(proxy_id)
        return self.traffic_store.save_session(proxy_id, name)

    def save_session_with_analysis(self, proxy_id: str, name: Optional[str] = None, analysis: Optional[Dict[str, Any]] = None) -> Dict:
        """Save current traffic log as a session with AI analysis data."""
        proxy = self._get_proxy(proxy_id)
        session_meta = self.traffic_store.save_session(proxy_id, name)
        
        # Add proxy info and analysis to session metadata
        session_meta["proxy_id"] = proxy_id
        session_meta["target_host"] = proxy.target_host
        session_meta["target_port"] = proxy.target_port
        
        if analysis:
            # Extract key analysis info
            analysis_summary = {
                "summary": analysis.get("summary", ""),
                "risk_score": analysis.get("risk_score", 0),
                "findings_count": len(analysis.get("findings", [])),
                "critical_count": sum(1 for f in analysis.get("findings", []) if f.get("severity") == "Critical"),
                "high_count": sum(1 for f in analysis.get("findings", []) if f.get("severity") == "High"),
                "medium_count": sum(1 for f in analysis.get("findings", []) if f.get("severity") == "Medium"),
                "low_count": sum(1 for f in analysis.get("findings", []) if f.get("severity") == "Low"),
                "findings": [
                    {"severity": f.get("severity"), "title": f.get("title"), "description": f.get("description", "")[:200]}
                    for f in analysis.get("findings", [])[:20]  # Store top 20 findings summaries
                ],
                "ai_writeup": analysis.get("ai_writeup", "")[:2000] if analysis.get("ai_writeup") else None,
                "attack_paths": analysis.get("attack_paths", [])[:10] if analysis.get("attack_paths") else None,
            }
            session_meta["analysis"] = analysis_summary
            session_meta["has_analysis"] = True
        else:
            session_meta["has_analysis"] = False
        
        # Update the session metadata file
        self.traffic_store.update_session_meta(proxy_id, session_meta["id"], session_meta)
        return session_meta

    def delete_session(self, proxy_id: str, session_id: str) -> None:
        """Delete a saved session."""
        self._get_proxy(proxy_id)
        self.traffic_store.delete_session(proxy_id, session_id)

    def list_all_sessions(self) -> List[Dict[str, Any]]:
        """List all saved sessions across all proxies."""
        all_sessions = []
        for proxy_id in self.proxies.keys():
            try:
                sessions = self.traffic_store.list_sessions(proxy_id)
                for session in sessions:
                    session["proxy_id"] = proxy_id
                    proxy = self.proxies.get(proxy_id)
                    if proxy:
                        session["target_host"] = proxy.target_host
                        session["target_port"] = proxy.target_port
                all_sessions.extend(sessions)
            except Exception:
                continue
        # Sort by created_at descending
        all_sessions.sort(key=lambda s: s.get("created_at", ""), reverse=True)
        return all_sessions

    def list_sessions(self, proxy_id: str) -> List[Dict[str, Any]]:
        """List saved sessions for a proxy."""
        self._get_proxy(proxy_id)
        return self.traffic_store.list_sessions(proxy_id)

    def load_session(self, proxy_id: str, session_id: str, limit: int = 100, offset: int = 0) -> Dict:
        """Load a saved session's traffic."""
        self._get_proxy(proxy_id)
        entries = self.traffic_store.load_session_entries(proxy_id, session_id, limit, offset)
        meta = self.traffic_store.get_session_meta(proxy_id, session_id) or {}
        return {"entries": entries, "total": meta.get("entries", len(entries)), "meta": meta}

    async def replay_entry(
        self,
        proxy_id: str,
        entry_id: str,
        overrides: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Replay a captured request with optional overrides."""
        proxy = self._get_proxy(proxy_id)
        overrides = overrides or {}
        entry = proxy._traffic_index.get(entry_id)
        entry_dict = proxy._traffic_entry_to_dict(entry) if entry else self.traffic_store.get_entry(proxy_id, entry_id)
        if not entry_dict:
            raise ValueError(f"Traffic entry {entry_id} not found")

        request = entry_dict.get("request", {})
        method = overrides.get("method") or request.get("method", "GET")
        path = overrides.get("path") or request.get("path", "/")
        headers = dict(request.get("headers", {}) or {})
        body = overrides.get("body") if overrides.get("body") is not None else request.get("body") or ""
        verify_tls = overrides.get("verify_tls", False)

        add_headers = overrides.get("add_headers", {}) or {}
        remove_headers = overrides.get("remove_headers", []) or []
        for key, value in add_headers.items():
            headers[key] = value
        for key in remove_headers:
            headers.pop(key, None)

        base_url = overrides.get("base_url")
        if base_url:
            url = base_url.rstrip("/") + path
        else:
            scheme = "https" if proxy.tls_enabled else "http"
            host = request.get("host") or proxy.target_host
            url = f"{scheme}://{host}{path}"

        timeout = overrides.get("timeout", 20)

        import httpx

        async with httpx.AsyncClient(verify=verify_tls, timeout=timeout) as client:
            response = await client.request(
                method=method,
                url=url,
                headers=headers,
                content=body.encode("utf-8") if isinstance(body, str) else body
            )

        replay_request = InterceptedRequest(
            id=str(uuid.uuid4()),
            timestamp=datetime.now(),
            client_ip="replay",
            client_port=0,
            method=method,
            url=url,
            path=path,
            host=urlsplit(url).hostname or "",
            port=urlsplit(url).port or (443 if proxy.tls_enabled else 80),
            protocol=Protocol.HTTPS if proxy.tls_enabled else Protocol.HTTP,
            http_version="HTTP/1.1",
            headers=headers,
            body=body.encode("utf-8") if isinstance(body, str) else body,
            body_text=body if isinstance(body, str) else None,
            received_at=time.time()
        )

        replay_response = InterceptedResponse(
            id=str(uuid.uuid4()),
            request_id=replay_request.id,
            timestamp=datetime.now(),
            status_code=response.status_code,
            status_message=response.reason_phrase or "",
            http_version="HTTP/1.1",
            headers=dict(response.headers),
            body=response.content,
            body_text=response.text if response.text else None,
            content_length=len(response.content),
            received_at=time.time(),
            response_time_ms=response.elapsed.total_seconds() * 1000.0
        )

        replay_entry = TrafficEntry(
            id=str(uuid.uuid4()),
            request=replay_request,
            response=replay_response,
            tags=["replay"],
            notes="Replayed request"
        )
        proxy._append_entry(replay_entry)
        if proxy.on_entry:
            proxy.on_entry(replay_entry, "request")
        if proxy.on_entry:
            proxy.on_entry(replay_entry, "response")

        return {
            "status": "replayed",
            "entry": proxy._traffic_entry_to_dict(replay_entry),
            "response": {
                "status_code": response.status_code,
                "headers": dict(response.headers),
                "body": response.text[:10000] if response.text else ""
            }
        }
    
    def set_mode(self, proxy_id: str, mode: str) -> Dict:
        """Set proxy interception mode"""
        proxy = self._get_proxy(proxy_id)
        proxy.mode = InterceptionMode(mode)
        mitm_stream_manager.emit(proxy_id, {"type": "mode", "mode": mode})
        return {"status": "updated", "mode": mode}
    
    def add_rule(self, proxy_id: str, rule_data: Dict) -> Dict:
        """Add an interception rule. Prevents duplicates by checking rule name."""
        proxy = self._get_proxy(proxy_id)

        # Check for existing rule with same name to prevent duplicates
        rule_name = rule_data.get("name", "Unnamed Rule")
        for existing_rule in proxy.rules:
            if existing_rule.name == rule_name:
                logger.debug(f"Rule '{rule_name}' already exists, skipping duplicate")
                return {
                    "status": "exists",
                    "rule_id": existing_rule.id,
                    "rule_name": existing_rule.name,
                    "message": "Rule with this name already exists"
                }

        rule = InterceptionRule(
            id=str(uuid.uuid4()),
            name=rule_data.get("name", "Unnamed Rule"),
            enabled=rule_data.get("enabled", True),
            priority=rule_data.get("priority", 100),
            group=rule_data.get("group"),
            match_host=rule_data.get("match_host"),
            match_path=rule_data.get("match_path"),
            match_method=rule_data.get("match_method"),
            match_content_type=rule_data.get("match_content_type"),
            match_body=rule_data.get("match_body"),
            match_header=rule_data.get("match_header"),
            match_status_code=rule_data.get("match_status_code"),
            match_direction=rule_data.get("match_direction", "both"),
            match_query=rule_data.get("match_query"),
            action=rule_data.get("action", "modify"),
            modify_headers=rule_data.get("modify_headers"),
            remove_headers=rule_data.get("remove_headers"),
            modify_body=rule_data.get("modify_body"),
            body_find_replace=rule_data.get("body_find_replace"),
            body_find_replace_regex=rule_data.get("body_find_replace_regex", False),
            json_path_edits=rule_data.get("json_path_edits"),
            modify_status_code=rule_data.get("modify_status_code"),
            modify_path=rule_data.get("modify_path"),
            delay_ms=rule_data.get("delay_ms", 0)
        )
        
        proxy.add_rule(rule)
        mitm_stream_manager.emit(proxy_id, {"type": "rules", "action": "added", "rule_id": rule.id})
        
        return {
            "status": "added",
            "rule_id": rule.id,
            "rule_name": rule.name
        }
    
    def remove_rule(self, proxy_id: str, rule_id: str) -> Dict:
        """Remove an interception rule"""
        proxy = self._get_proxy(proxy_id)
        proxy.remove_rule(rule_id)
        mitm_stream_manager.emit(proxy_id, {"type": "rules", "action": "removed", "rule_id": rule_id})
        return {"status": "removed", "rule_id": rule_id}
    
    def get_rules(self, proxy_id: str) -> List[Dict]:
        """Get all rules for a proxy"""
        proxy = self._get_proxy(proxy_id)
        return [
            {
                "id": r.id,
                "name": r.name,
                "enabled": r.enabled,
                "priority": r.priority,
                "group": r.group,
                "match_host": r.match_host,
                "match_path": r.match_path,
                "match_method": r.match_method,
                "match_content_type": r.match_content_type,
                "match_status_code": r.match_status_code,
                "match_direction": r.match_direction,
                "match_query": r.match_query,
                "action": r.action,
                "delay_ms": r.delay_ms,
                "modify_headers": r.modify_headers,
                "remove_headers": r.remove_headers,
                "body_find_replace": r.body_find_replace,
                "body_find_replace_regex": r.body_find_replace_regex,
                "json_path_edits": r.json_path_edits,
                "modify_path": r.modify_path,
                "modify_status_code": r.modify_status_code,
                "hit_count": r.hit_count
            }
            for r in proxy.rules
        ]
    
    def toggle_rule(self, proxy_id: str, rule_id: str, enabled: bool) -> Dict:
        """Enable/disable a rule"""
        proxy = self._get_proxy(proxy_id)
        for rule in proxy.rules:
            if rule.id == rule_id:
                rule.enabled = enabled
                mitm_stream_manager.emit(proxy_id, {"type": "rules", "action": "toggled", "rule_id": rule_id, "enabled": enabled})
                return {"status": "updated", "rule_id": rule_id, "enabled": enabled}
        raise ValueError(f"Rule {rule_id} not found")

    def toggle_rule_group(self, proxy_id: str, group: str, enabled: bool) -> Dict:
        """Enable/disable all rules in a group."""
        proxy = self._get_proxy(proxy_id)
        updated = 0
        for rule in proxy.rules:
            if rule.group == group:
                rule.enabled = enabled
                updated += 1
        if updated == 0:
            raise ValueError(f"Rule group {group} not found")
        mitm_stream_manager.emit(proxy_id, {"type": "rules", "action": "group_toggled", "group": group, "enabled": enabled})
        return {"status": "updated", "group": group, "enabled": enabled, "count": updated}
    
    # ========================================================================
    # WebSocket Deep Inspection Methods
    # ========================================================================
    
    def get_websocket_connections(self, proxy_id: str) -> List[Dict[str, Any]]:
        """Get all WebSocket connections for a proxy"""
        self._get_proxy(proxy_id)  # Validate proxy exists
        connections = websocket_inspector.get_connections(proxy_id)
        return [
            {
                "id": c.id,
                "proxy_id": c.proxy_id,
                "created_at": c.created_at.isoformat(),
                "client_ip": c.client_ip,
                "client_port": c.client_port,
                "target_host": c.target_host,
                "target_port": c.target_port,
                "upgrade_request_id": c.upgrade_request_id,
                "status": c.status,
                "total_frames": c.total_frames,
                "bytes_sent": c.bytes_sent,
                "bytes_received": c.bytes_received,
                "closed_at": c.closed_at.isoformat() if c.closed_at else None,
                "close_code": c.close_code,
                "close_reason": c.close_reason
            }
            for c in connections
        ]
    
    def get_websocket_frames(
        self,
        proxy_id: str,
        connection_id: str,
        limit: int = 100,
        offset: int = 0
    ) -> Dict[str, Any]:
        """Get WebSocket frames for a connection"""
        self._get_proxy(proxy_id)
        frames = websocket_inspector.get_frames(connection_id, limit, offset)
        conn = websocket_inspector.get_connection(connection_id)
        total = conn.total_frames if conn else len(frames)
        return {
            "frames": frames,
            "total": total,
            "connection_id": connection_id
        }
    
    def add_websocket_rule(self, proxy_id: str, rule_data: Dict[str, Any]) -> Dict[str, Any]:
        """Add a WebSocket interception rule"""
        self._get_proxy(proxy_id)
        
        rule = WebSocketRule(
            id=str(uuid.uuid4()),
            name=rule_data.get("name", "Unnamed Rule"),
            enabled=rule_data.get("enabled", True),
            priority=rule_data.get("priority", 100),
            match_direction=rule_data.get("match_direction", "both"),
            match_opcode=rule_data.get("match_opcode"),
            match_payload_pattern=rule_data.get("match_payload_pattern"),
            match_json_path=rule_data.get("match_json_path"),
            action=rule_data.get("action", "modify"),
            payload_find_replace=rule_data.get("payload_find_replace"),
            json_path_edits=rule_data.get("json_path_edits"),
            delay_ms=rule_data.get("delay_ms", 0)
        )
        
        websocket_inspector.add_rule(rule)
        mitm_stream_manager.emit(proxy_id, {"type": "ws_rules", "action": "added", "rule_id": rule.id})
        
        return {
            "status": "added",
            "rule_id": rule.id,
            "rule_name": rule.name
        }
    
    def remove_websocket_rule(self, proxy_id: str, rule_id: str) -> Dict[str, Any]:
        """Remove a WebSocket rule"""
        self._get_proxy(proxy_id)
        websocket_inspector.remove_rule(rule_id)
        mitm_stream_manager.emit(proxy_id, {"type": "ws_rules", "action": "removed", "rule_id": rule_id})
        return {"status": "removed", "rule_id": rule_id}
    
    def get_websocket_rules(self, proxy_id: str) -> List[Dict[str, Any]]:
        """Get all WebSocket rules"""
        self._get_proxy(proxy_id)
        return websocket_inspector.get_rules()
    
    def get_websocket_stats(self, proxy_id: str) -> Dict[str, Any]:
        """Get WebSocket inspection statistics"""
        self._get_proxy(proxy_id)
        return websocket_inspector.get_stats()
    
    # ========================================================================
    # Certificate Management Methods
    # ========================================================================
    
    def get_certificate_manager(self) -> CertificateManager:
        """Get or create the certificate manager"""
        if not hasattr(self, '_cert_manager'):
            from backend.core.config import settings
            cert_dir = Path(settings.mitm_storage_dir) / "certificates"
            self._cert_manager = CertificateManager(cert_dir)
        return self._cert_manager
    
    def get_ca_certificate(self) -> Optional[Dict[str, Any]]:
        """Get the current CA certificate info"""
        cert_mgr = self.get_certificate_manager()
        ca = cert_mgr.get_ca_certificate()
        if not ca:
            return None
        
        return {
            "common_name": ca.common_name,
            "organization": ca.organization,
            "country": ca.country,
            "serial": ca.serial,
            "valid_from": ca.valid_from.isoformat(),
            "valid_until": ca.valid_until.isoformat(),
            "fingerprint_sha256": ca.fingerprint_sha256,
            "created_at": ca.created_at.isoformat(),
            "certificate_pem": ca.certificate_pem
        }
    
    def generate_ca_certificate(
        self,
        common_name: str = "VRAgent MITM CA",
        organization: str = "VRAgent Security",
        country: str = "US",
        validity_days: int = 3650
    ) -> Dict[str, Any]:
        """Generate a new CA certificate"""
        cert_mgr = self.get_certificate_manager()
        ca = cert_mgr.generate_ca_certificate(
            common_name=common_name,
            organization=organization,
            country=country,
            validity_days=validity_days
        )
        
        return {
            "status": "generated",
            "common_name": ca.common_name,
            "fingerprint_sha256": ca.fingerprint_sha256,
            "valid_from": ca.valid_from.isoformat(),
            "valid_until": ca.valid_until.isoformat()
        }
    
    def get_host_certificate(self, hostname: str) -> Optional[Dict[str, Any]]:
        """Get or generate a certificate for a specific host"""
        cert_mgr = self.get_certificate_manager()
        cert = cert_mgr.get_host_certificate(hostname)
        if not cert:
            return None
        
        return {
            "hostname": cert.hostname,
            "serial": cert.serial,
            "valid_from": cert.valid_from.isoformat(),
            "valid_until": cert.valid_until.isoformat(),
            "fingerprint_sha256": cert.fingerprint_sha256,
            "created_at": cert.created_at.isoformat(),
            "certificate_pem": cert.certificate_pem
        }
    
    def list_host_certificates(self) -> List[Dict[str, Any]]:
        """List all generated host certificates"""
        cert_mgr = self.get_certificate_manager()
        return cert_mgr.list_host_certificates()
    
    def delete_host_certificate(self, hostname: str) -> Dict[str, Any]:
        """Delete a host certificate"""
        cert_mgr = self.get_certificate_manager()
        deleted = cert_mgr.delete_host_certificate(hostname)
        return {
            "status": "deleted" if deleted else "not_found",
            "hostname": hostname
        }
    
    def get_certificate_installation_instructions(self) -> Dict[str, Any]:
        """Get instructions for installing the CA certificate"""
        cert_mgr = self.get_certificate_manager()
        return cert_mgr.get_installation_instructions()
    
    def download_ca_certificate(self, format: str = "pem") -> Tuple[bytes, str, str]:
        """
        Download the CA certificate in various formats.
        
        Returns:
            Tuple of (content_bytes, media_type, filename)
        """
        cert_mgr = self.get_certificate_manager()
        ca = cert_mgr.get_ca_certificate()
        if not ca:
            raise ValueError("No CA certificate generated yet")
        
        if format == "pem":
            return (
                ca.certificate_pem.encode('utf-8'),
                "application/x-pem-file",
                "vragent-mitm-ca.pem"
            )
        elif format == "crt":
            # Same as PEM but with .crt extension (Windows-friendly)
            return (
                ca.certificate_pem.encode('utf-8'),
                "application/x-x509-ca-cert",
                "vragent-mitm-ca.crt"
            )
        elif format == "der":
            # Binary DER format
            try:
                from cryptography import x509
                from cryptography.hazmat.primitives import serialization
                from cryptography.hazmat.backends import default_backend
                
                cert = x509.load_pem_x509_certificate(
                    ca.certificate_pem.encode(),
                    default_backend()
                )
                der_bytes = cert.public_bytes(serialization.Encoding.DER)
                return (
                    der_bytes,
                    "application/x-x509-ca-cert",
                    "vragent-mitm-ca.der"
                )
            except ImportError:
                raise ValueError("cryptography library required for DER format")
        else:
            raise ValueError(f"Unknown format: {format}")
    
    # ========================================================================
    # Traffic Diff Viewer Methods
    # ========================================================================
    
    def get_traffic_diff(self, proxy_id: str, entry_id: str) -> Dict[str, Any]:
        """Get diff between original and modified traffic entry"""
        proxy = self._get_proxy(proxy_id)
        
        # Find the traffic entry
        entry = None
        for e in proxy.traffic_log:
            if e.get("id") == entry_id:
                entry = e
                break
        
        if not entry:
            raise KeyError(f"Traffic entry {entry_id} not found")
        
        # Check if entry was modified
        if not entry.get("modified", False):
            return {
                "has_changes": False,
                "entry_id": entry_id,
                "message": "Entry was not modified"
            }
        
        # Get original data (stored when rule is applied)
        original_request = entry.get("original_request", entry.get("request", {}))
        original_response = entry.get("original_response", entry.get("response"))
        
        modified_request = entry.get("request", {})
        modified_response = entry.get("response")
        
        results = TrafficDiffViewer.compare_traffic(
            original_request,
            modified_request,
            original_response,
            modified_response
        )
        
        return {
            "entry_id": entry_id,
            "proxy_id": proxy_id,
            "modified": True,
            "rules_applied": entry.get("rules_applied", []),
            "diff": {
                key: {
                    "has_changes": result.has_changes,
                    "change_type": result.change_type,
                    "header_changes": result.header_changes,
                    "body_changes": result.body_changes,
                    "summary": result.summary,
                    "original_size": result.original_size,
                    "modified_size": result.modified_size
                }
                for key, result in results.items()
            }
        }
    
    # ========================================================================
    # HTTP/2 & gRPC Methods
    # ========================================================================
    
    def get_http2_frames(
        self,
        proxy_id: str,
        stream_id: Optional[int] = None,
        frame_type: Optional[str] = None,
        limit: int = 100,
        offset: int = 0
    ) -> Dict[str, Any]:
        """Get HTTP/2 frames captured for a proxy"""
        proxy = self._get_proxy(proxy_id)
        
        # HTTP/2 frames are stored in proxy.http2_frames if supported
        frames = getattr(proxy, 'http2_frames', [])
        
        # Apply filters
        filtered = frames
        if stream_id is not None:
            filtered = [f for f in filtered if f.get("stream_id") == stream_id]
        if frame_type:
            filtered = [f for f in filtered if f.get("frame_type_name", "").upper() == frame_type.upper()]
        
        total = len(filtered)
        paginated = filtered[offset:offset + limit]
        
        return {
            "proxy_id": proxy_id,
            "frames": paginated,
            "total": total,
            "limit": limit,
            "offset": offset,
            "http2_supported": getattr(proxy, 'http2_enabled', False)
        }
    
    def get_http2_streams(self, proxy_id: str) -> Dict[str, Any]:
        """Get active HTTP/2 streams for a proxy"""
        proxy = self._get_proxy(proxy_id)
        
        streams = getattr(proxy, 'http2_streams', {})
        
        return {
            "proxy_id": proxy_id,
            "streams": [
                {
                    "stream_id": sid,
                    "state": info.get("state", "unknown"),
                    "method": info.get("method"),
                    "path": info.get("path"),
                    "frame_count": info.get("frame_count", 0),
                    "bytes_sent": info.get("bytes_sent", 0),
                    "bytes_received": info.get("bytes_received", 0)
                }
                for sid, info in streams.items()
            ],
            "total_streams": len(streams),
            "http2_supported": getattr(proxy, 'http2_enabled', False)
        }
    
    def get_grpc_messages(
        self,
        proxy_id: str,
        service: Optional[str] = None,
        method: Optional[str] = None,
        limit: int = 100,
        offset: int = 0
    ) -> Dict[str, Any]:
        """Get gRPC messages captured for a proxy"""
        proxy = self._get_proxy(proxy_id)
        
        messages = getattr(proxy, 'grpc_messages', [])
        
        # Apply filters
        filtered = messages
        if service:
            filtered = [m for m in filtered if m.get("service") == service]
        if method:
            filtered = [m for m in filtered if m.get("method") == method]
        
        total = len(filtered)
        paginated = filtered[offset:offset + limit]
        
        return {
            "proxy_id": proxy_id,
            "messages": paginated,
            "total": total,
            "limit": limit,
            "offset": offset,
            "grpc_supported": getattr(proxy, 'grpc_enabled', False)
        }
    
    # ========================================================================
    # Match & Replace Templates Methods
    # ========================================================================
    
    def get_match_replace_library(self) -> MatchReplaceLibrary:
        """Get the match/replace template library"""
        if not hasattr(self, '_template_library'):
            self._template_library = MatchReplaceLibrary()
        return self._template_library
    
    def get_match_replace_templates(
        self,
        category: Optional[str] = None,
        tag: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Get available match/replace templates"""
        library = self.get_match_replace_library()
        
        if category:
            templates = library.get_templates_by_category(category)
        elif tag:
            templates = library.get_templates_by_tag(tag)
        else:
            templates = library.get_all_templates()
        
        return [
            {
                "id": t.id,
                "name": t.name,
                "category": t.category,
                "description": t.description,
                "match_type": t.match_type,
                "match_pattern": t.match_pattern,
                "replace_pattern": t.replace_pattern,
                "is_regex": t.is_regex,
                "direction": t.direction,
                "tags": t.tags,
                "enabled": t.enabled,
                "hit_count": t.hit_count,
                "is_built_in": t.id in library.templates
            }
            for t in templates
        ]
    
    def get_template_categories(self) -> List[str]:
        """Get available template categories"""
        library = self.get_match_replace_library()
        return library.get_categories()
    
    def get_template(self, template_id: str) -> Optional[Dict[str, Any]]:
        """Get a specific template by ID"""
        library = self.get_match_replace_library()
        template = library.templates.get(template_id) or library.custom_templates.get(template_id)
        
        if not template:
            return None
        
        return {
            "id": template.id,
            "name": template.name,
            "category": template.category,
            "description": template.description,
            "match_type": template.match_type,
            "match_pattern": template.match_pattern,
            "replace_pattern": template.replace_pattern,
            "is_regex": template.is_regex,
            "case_sensitive": template.case_sensitive,
            "direction": template.direction,
            "match_host": template.match_host,
            "match_content_type": template.match_content_type,
            "tags": template.tags,
            "enabled": template.enabled,
            "hit_count": template.hit_count,
            "created_at": template.created_at.isoformat(),
            "is_built_in": template.id in library.templates
        }
    
    def create_custom_template(self, config: Dict[str, Any]) -> Dict[str, Any]:
        """Create a custom match/replace template"""
        library = self.get_match_replace_library()
        
        template_id = f"custom-{uuid.uuid4().hex[:8]}"
        
        template = MatchReplaceTemplate(
            id=template_id,
            name=config["name"],
            category=config.get("category", "Custom"),
            description=config["description"],
            match_type=config["match_type"],
            match_pattern=config["match_pattern"],
            replace_pattern=config["replace_pattern"],
            is_regex=config.get("is_regex", False),
            case_sensitive=config.get("case_sensitive", True),
            direction=config.get("direction", "both"),
            match_host=config.get("match_host"),
            match_content_type=config.get("match_content_type"),
            tags=config.get("tags", []),
            created_at=datetime.now()
        )
        
        library.add_custom_template(template)
        
        return {
            "status": "created",
            "template_id": template_id,
            "template": self.get_template(template_id)
        }
    
    def delete_custom_template(self, template_id: str) -> bool:
        """Delete a custom template"""
        library = self.get_match_replace_library()
        return library.remove_custom_template(template_id)
    
    def apply_template_to_proxy(self, proxy_id: str, template_id: str) -> Dict[str, Any]:
        """Apply a template to a proxy as an interception rule"""
        proxy = self._get_proxy(proxy_id)
        library = self.get_match_replace_library()
        
        template = library.templates.get(template_id) or library.custom_templates.get(template_id)
        if not template:
            raise KeyError(f"Template {template_id} not found")
        
        # Convert template to a rule
        rule_id = f"tmpl-{template_id}-{uuid.uuid4().hex[:6]}"
        
        rule_config = {
            "id": rule_id,
            "name": f"[Template] {template.name}",
            "enabled": True,
            "priority": 50,  # Templates get medium priority
            "match_direction": template.direction,
            "source": f"template:{template_id}"
        }
        
        # Map template match_type to rule config
        if template.match_type == "header":
            if template.replace_pattern:
                rule_config["modify_headers"] = {template.match_pattern: template.replace_pattern}
            else:
                rule_config["remove_headers"] = [template.match_pattern]
        elif template.match_type == "body":
            rule_config["body_find_replace"] = {template.match_pattern: template.replace_pattern}
            rule_config["body_find_replace_regex"] = template.is_regex
        elif template.match_type == "path":
            rule_config["modify_path"] = template.replace_pattern
            rule_config["match_path"] = template.match_pattern
        elif template.match_type == "query":
            rule_config["match_path"] = f"*?*{template.match_pattern}*"
        
        if template.match_host:
            rule_config["match_host"] = template.match_host
        if template.match_content_type:
            rule_config["match_content_type"] = template.match_content_type
        
        # Add the rule to the proxy
        rule = InterceptionRule(**rule_config)
        proxy.rules.append(rule)
        
        return {
            "status": "applied",
            "proxy_id": proxy_id,
            "template_id": template_id,
            "rule_id": rule_id,
            "rule_name": rule_config["name"]
        }
    
    def test_template(
        self,
        template_id: str,
        request_data: Dict[str, Any],
        response_data: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Test a template against sample data"""
        library = self.get_match_replace_library()
        
        # Test on request
        modified_request, request_changed = library.apply_template(
            template_id, request_data, "request"
        )
        
        # Test on response if provided
        modified_response = None
        response_changed = False
        if response_data:
            modified_response, response_changed = library.apply_template(
                template_id, response_data, "response"
            )
        
        # Generate diff
        results = {}
        if request_changed:
            results["request"] = TrafficDiffViewer.compare_traffic(
                request_data, modified_request, None, None
            ).get("request")
        if response_changed and response_data:
            results["response"] = TrafficDiffViewer.compare_traffic(
                {}, {}, response_data, modified_response
            ).get("response")
        
        return {
            "template_id": template_id,
            "request_modified": request_changed,
            "response_modified": response_changed,
            "modified_request": modified_request if request_changed else None,
            "modified_response": modified_response if response_changed else None,
            "diff": {
                key: {
                    "has_changes": r.has_changes,
                    "change_type": r.change_type,
                    "header_changes": r.header_changes,
                    "body_changes": r.body_changes,
                    "summary": r.summary
                } if r else None
                for key, r in results.items()
            } if results else None
        }
    
    def _get_proxy(self, proxy_id: str) -> MITMProxy:
        """Get a proxy by ID"""
        with self._lock:
            if proxy_id not in self.proxies:
                raise ValueError(f"Proxy {proxy_id} not found")
            return self.proxies[proxy_id]


# Global service instance
mitm_service = MITMService()


# ============================================================================
# AI Analysis for MITM Traffic - Enhanced Multi-Pass System
# ============================================================================

async def analyze_mitm_traffic(
    traffic_log: List[Dict],
    rules: List[Dict],
    proxy_config: Dict,
    agent_activity: Optional[Dict] = None
) -> Dict:
    """
    Enhanced AI-powered analysis of MITM intercepted traffic.
    
    Uses a 3-pass analysis system:
    - Pass 1: Pattern-based vulnerability detection
    - Pass 2: AI-powered contextual analysis and additional findings
    - Pass 3: Verification, deduplication, and false positive removal
    
    Includes:
    - Security vulnerabilities with exploitation details
    - CVE database lookup
    - Exploit database references
    - Attack path generation
    - Detailed exploitation steps
    - Professional penetration test writeup
    """
    from backend.core.config import settings
    from backend.services.vuln_intelligence_service import (
        enrich_finding_with_intelligence,
        generate_attack_paths,
        generate_ai_exploitation_writeup,
        search_cve_database,
        get_offline_exploit_references
    )
    
    if not traffic_log:
        return {
            "summary": "No traffic to analyze",
            "risk_score": 0,
            "findings": [],
            "recommendations": [],
            "attack_paths": [],
            "exploit_references": [],
            "cve_references": [],
            "analysis_passes": 0
        }
    
    logger.info(f"Starting 3-pass MITM analysis on {len(traffic_log)} traffic entries")
    
    # =========================================================================
    # PASS 1: Pattern-based vulnerability detection
    # =========================================================================
    logger.info("Pass 1: Pattern-based vulnerability detection")
    
    pass1_findings = []
    detected_technologies = set()
    endpoint_findings = defaultdict(list)  # Track findings per endpoint
    
    for entry in traffic_log[:250]:  # Analyze up to 250 traffic entries
        request = entry.get("request", {})
        response = entry.get("response", {})
        
        # Check for sensitive data in requests
        request_body = request.get("body_text", "") or ""
        request_headers = request.get("headers", {})
        request_path = request.get("path", "")
        
        # Detect technologies from headers/paths
        if response:
            server = response.get("headers", {}).get("Server", "")
            x_powered = response.get("headers", {}).get("X-Powered-By", "")
            if server:
                detected_technologies.add(f"Server: {server}")
            if x_powered:
                detected_technologies.add(f"Framework: {x_powered}")
        
        # Enhanced sensitive data patterns
        sensitive_patterns = [
            ("password", "high", "Password transmitted in request"),
            ("passwd", "high", "Password field detected"),
            ("api_key", "critical", "API key exposed in request"),
            ("apikey", "critical", "API key exposed in request"),
            ("secret", "high", "Secret value in request"),
            ("token", "medium", "Token in request (verify if sensitive)"),
            ("bearer", "high", "Bearer token exposed"),
            ("authorization", "high", "Authorization credentials exposed"),
            ("credit_card", "critical", "Credit card data detected"),
            ("card_number", "critical", "Card number detected"),
            ("cvv", "critical", "CVV code detected"),
            ("ssn", "critical", "SSN pattern detected"),
            ("social_security", "critical", "Social security number detected"),
            ("private_key", "critical", "Private key exposed"),
        ]
        
        for pattern, severity, message in sensitive_patterns:
            if pattern in request_body.lower() or pattern in str(request_headers).lower():
                pass1_findings.append({
                    "severity": severity,
                    "category": "sensitive_data",
                    "title": message,
                    "description": f"Sensitive data pattern '{pattern}' found in request",
                    "evidence": f"Request: {request.get('method', 'GET')} {request_path}",
                    "recommendation": "Encrypt sensitive data, use HTTPS, avoid logging sensitive values",
                    "affected_endpoint": request_path
                })
        
        # Check for missing security headers in responses
        if response:
            response_headers = response.get("headers", {})
            header_lower = {k.lower(): v for k, v in response_headers.items()}
            
            # Enhanced security header checks with categories
            security_headers = [
                ("content-security-policy", "missing_csp", "Missing Content-Security-Policy header", "critical"),
                ("x-content-type-options", "missing_x_content_type", "Missing X-Content-Type-Options header", "medium"),
                ("x-frame-options", "missing_x_frame", "Missing X-Frame-Options header", "medium"),
                ("strict-transport-security", "missing_hsts", "Missing HSTS header", "high"),
                ("x-xss-protection", "missing_x_xss", "Missing X-XSS-Protection header", "low"),
                ("referrer-policy", "missing_referrer", "Missing Referrer-Policy header", "low"),
                ("permissions-policy", "missing_permissions", "Missing Permissions-Policy header", "low"),
            ]
            
            for header, category, message, severity in security_headers:
                if header not in header_lower:
                    pass1_findings.append({
                        "severity": severity,
                        "category": category,
                        "title": message,
                        "description": f"Response missing security header: {header}",
                        "evidence": f"Response from {request_path}",
                        "recommendation": f"Add {header} header to all responses",
                        "affected_endpoint": request_path
                    })
        
        # Check for insecure cookies with detailed analysis
        if response:
            for header_name, header_value in response.get("headers", {}).items():
                if header_name.lower() == "set-cookie":
                    cookie_lower = header_value.lower()
                    cookie_name = header_value.split("=")[0] if "=" in header_value else "unknown"
                    
                    if "httponly" not in cookie_lower:
                        pass1_findings.append({
                            "severity": "high",
                            "category": "cookie_no_httponly",
                            "title": f"Cookie '{cookie_name}' missing HttpOnly flag",
                            "description": "Cookie accessible via JavaScript, enabling theft via XSS attacks",
                            "evidence": f"Set-Cookie: {header_value[:80]}...",
                            "recommendation": "Add HttpOnly flag: Set-Cookie: {name}=value; HttpOnly",
                            "affected_endpoint": request_path
                        })
                    
                    if "secure" not in cookie_lower:
                        pass1_findings.append({
                            "severity": "medium",
                            "category": "cookie_no_secure",
                            "title": f"Cookie '{cookie_name}' missing Secure flag",
                            "description": "Cookie may be transmitted over unencrypted HTTP connections",
                            "evidence": f"Set-Cookie: {header_value[:80]}...",
                            "recommendation": "Add Secure flag: Set-Cookie: {name}=value; Secure",
                            "affected_endpoint": request_path
                        })
                    
                    if "samesite" not in cookie_lower:
                        pass1_findings.append({
                            "severity": "medium",
                            "category": "cookie_no_samesite",
                            "title": f"Cookie '{cookie_name}' missing SameSite attribute",
                            "description": "Cookie may be sent in cross-site requests, enabling CSRF attacks",
                            "evidence": f"Set-Cookie: {header_value[:80]}...",
                            "recommendation": "Add SameSite=Strict or SameSite=Lax",
                            "affected_endpoint": request_path
                        })
        
        # Check for CORS misconfigurations
        if response:
            acao = response.get("headers", {}).get("Access-Control-Allow-Origin", "")
            acac = response.get("headers", {}).get("Access-Control-Allow-Credentials", "")
            
            if acao == "*":
                pass1_findings.append({
                    "severity": "high",
                    "category": "cors",
                    "title": "Overly permissive CORS: Access-Control-Allow-Origin: *",
                    "description": "Any website can make cross-origin requests and read responses",
                    "evidence": "Access-Control-Allow-Origin: *",
                    "recommendation": "Restrict to specific trusted origins",
                    "affected_endpoint": request_path
                })
            elif acao and acac.lower() == "true":
                pass1_findings.append({
                    "severity": "critical",
                    "category": "cors",
                    "title": "CORS allows credentials with dynamic origin",
                    "description": "If origin is reflected from request, attackers can steal authenticated data",
                    "evidence": f"ACAO: {acao}, ACAC: true",
                    "recommendation": "Never combine Allow-Credentials with dynamic/reflected origins",
                    "affected_endpoint": request_path
                })
        
        # Check for error information disclosure
        if response and response.get("status_code", 0) >= 400:
            body = response.get("body_text", "") or ""
            body_lower = body.lower()
            
            disclosure_patterns = [
                ("stacktrace", "Stack trace exposed"),
                ("exception", "Exception details exposed"),
                ("traceback", "Python traceback exposed"),
                ("error at line", "Line number disclosed"),
                ("sqlstate", "SQL error disclosed"),
                ("mysql", "MySQL error exposed"),
                ("postgresql", "PostgreSQL error exposed"),
                ("mongodb", "MongoDB error exposed"),
                ("syntax error", "Syntax error exposed"),
                ("undefined variable", "Debug info exposed"),
                ("/var/www/", "Server path disclosed"),
                ("/home/", "Server path disclosed"),
                ("c:\\", "Windows path disclosed"),
            ]
            
            for pattern, message in disclosure_patterns:
                if pattern in body_lower:
                    pass1_findings.append({
                        "severity": "medium",
                        "category": "information_disclosure",
                        "title": message,
                        "description": f"Error response reveals internal information: {pattern}",
                        "evidence": f"{response.get('status_code')} error on {request_path}",
                        "recommendation": "Use generic error messages in production",
                        "affected_endpoint": request_path
                    })
                    break  # One disclosure per response
        
        # Check for potential injection points
        if "?" in request_path or request_body:
            params = request_path.split("?")[1] if "?" in request_path else ""
            if any(x in params.lower() or x in request_body.lower() for x in ["id=", "user=", "name=", "search=", "query=", "q="]):
                pass1_findings.append({
                    "severity": "info",
                    "category": "injection_point",
                    "title": "Potential injection point identified",
                    "description": "User-controllable parameters detected - test for SQLi, XSS, etc.",
                    "evidence": f"Parameters in {request_path}",
                    "recommendation": "Test these parameters with injection payloads",
                    "affected_endpoint": request_path
                })
    
    logger.info(f"Pass 1 complete: {len(pass1_findings)} raw findings")
    
    # =========================================================================
    # PASS 2: AI-Powered Contextual Analysis
    # =========================================================================
    logger.info("Pass 2: AI-powered contextual analysis")
    
    pass2_findings = []
    
    if settings.gemini_api_key and len(traffic_log) > 5:
        try:
            from google import genai
            from google.genai import types
            
            client = genai.Client(api_key=settings.gemini_api_key)
            
            # Build traffic context for AI
            traffic_context = []
            for entry in traffic_log[:100]:  # Analyze up to 100 entries for AI context
                req = entry.get("request", {})
                resp = entry.get("response", {})
                traffic_context.append({
                    "method": req.get("method"),
                    "path": req.get("path"),
                    "status": resp.get("status_code") if resp else None,
                    "response_headers": list(resp.get("headers", {}).keys()) if resp else [],
                    "has_body": bool(req.get("body_text"))
                })
            
            pass2_prompt = f"""You are a senior penetration tester analyzing HTTP traffic for security vulnerabilities.

TARGET: {proxy_config.get('target_host', 'Unknown')}:{proxy_config.get('target_port', 'Unknown')}

TRAFFIC SAMPLE (first 50 requests):
{json.dumps(traffic_context[:30], indent=2)}

INITIAL FINDINGS COUNT: {len(pass1_findings)}

Analyze this traffic for ADDITIONAL security issues that pattern matching might miss:
1. Authentication/authorization weaknesses (missing auth on sensitive endpoints)
2. Business logic flaws (predictable IDs, race conditions potential)
3. API security issues (excessive data exposure, lack of rate limiting indicators)
4. Session management weaknesses
5. Information leakage in URL patterns or headers

Return ONLY a JSON array of new findings. Each finding must have:
{{"severity": "critical|high|medium|low", "category": "string", "title": "string", "description": "string", "evidence": "string", "affected_endpoint": "string"}}

Return [] if no additional issues found. NO explanation, ONLY valid JSON array."""

            response = client.models.generate_content(
                model=settings.gemini_model_id,
                contents=pass2_prompt,
                config=types.GenerateContentConfig(
                    thinking_config={"thinking_level": "medium"},
                    max_output_tokens=1500,
                )
            )
            
            if response and response.text:
                try:
                    # Extract JSON from response
                    response_text = response.text.strip()
                    if response_text.startswith("```json"):
                        response_text = response_text[7:]
                    if response_text.startswith("```"):
                        response_text = response_text[3:]
                    if response_text.endswith("```"):
                        response_text = response_text[:-3]
                    
                    ai_findings = json.loads(response_text.strip())
                    if isinstance(ai_findings, list):
                        for af in ai_findings:
                            if isinstance(af, dict) and "title" in af and "severity" in af:
                                af["source"] = "ai_analysis"
                                pass2_findings.append(af)
                        logger.info(f"Pass 2 found {len(pass2_findings)} additional findings via AI")
                except json.JSONDecodeError:
                    logger.warning("Pass 2 AI response was not valid JSON")
                    
        except Exception as e:
            logger.warning(f"Pass 2 AI analysis failed: {e}")
    
    # Combine Pass 1 and Pass 2 findings
    all_findings = pass1_findings + pass2_findings
    
    # =========================================================================
    # PASS 3: Verification, Deduplication & False Positive Removal
    # =========================================================================
    logger.info("Pass 3: Verification and deduplication")
    
    # Step 3a: Deduplicate by category+title
    seen = set()
    deduped_findings = []
    for f in all_findings:
        key = (f.get("title", ""), f.get("category", ""))
        if key not in seen:
            seen.add(key)
            deduped_findings.append(f)
    
    logger.info(f"After basic dedup: {len(deduped_findings)} findings (from {len(all_findings)})")
    
    # Step 3b: Deduplicate similar findings per endpoint (preserve endpoint-specific findings)
    # Use category + endpoint as key to keep findings on different endpoints separate
    endpoint_dedup = {}
    severity_order = {"critical": 4, "high": 3, "medium": 2, "low": 1, "info": 0}

    for f in deduped_findings:
        category = f.get("category", "")
        endpoint = f.get("affected_endpoint", f.get("path", "unknown"))
        # Use category+endpoint as key to preserve endpoint-specific findings
        dedup_key = (category, endpoint)

        if dedup_key not in endpoint_dedup:
            endpoint_dedup[dedup_key] = f
        else:
            # Same category AND endpoint - keep most severe
            existing_severity = endpoint_dedup[dedup_key].get("severity", "low")
            new_severity = f.get("severity", "low")
            if severity_order.get(new_severity, 0) > severity_order.get(existing_severity, 0):
                endpoint_dedup[dedup_key] = f

    # Group security header findings for summary while keeping individual instances
    header_findings_by_category = {}
    other_findings = []
    all_individual_findings = []  # Keep ALL individual findings

    for f in endpoint_dedup.values():
        category = f.get("category", "")
        all_individual_findings.append(f)  # Preserve individual finding

        if category.startswith("missing_"):
            # Track for summary but don't replace individual findings
            if category not in header_findings_by_category:
                header_findings_by_category[category] = {
                    "endpoints": [f.get("affected_endpoint", "unknown")],
                    "count": 1,
                    "severity": f.get("severity", "low")
                }
            else:
                header_findings_by_category[category]["endpoints"].append(f.get("affected_endpoint", "unknown"))
                header_findings_by_category[category]["count"] += 1

    # Add occurrence count to individual findings for context
    for f in all_individual_findings:
        category = f.get("category", "")
        if category in header_findings_by_category:
            total_count = header_findings_by_category[category]["count"]
            if total_count > 1:
                f["total_occurrences"] = total_count

    consolidated_findings = all_individual_findings
    
    logger.info(f"After consolidation: {len(consolidated_findings)} findings")
    
    # Step 3c: False Positive Filtering with AI (if available)
    verified_findings = consolidated_findings
    false_positives_removed = 0
    
    if settings.gemini_api_key and len(consolidated_findings) > 3:
        try:
            from google import genai
            from google.genai import types
            
            client = genai.Client(api_key=settings.gemini_api_key)
            
            findings_for_review = [
                {"index": i, "severity": f.get("severity"), "title": f.get("title"), "category": f.get("category")}
                for i, f in enumerate(consolidated_findings)
            ]
            
            verification_prompt = f"""As a security expert, review these findings for false positives.

TARGET: {proxy_config.get('target_host', 'Unknown')}

FINDINGS TO VERIFY:
{json.dumps(findings_for_review, indent=2)}

Return a JSON array of indices (integers) that are likely FALSE POSITIVES and should be removed.
Consider:
- "info" severity items that add no security value
- Duplicate concepts with different wording
- Items that are not actual security vulnerabilities
- Low-risk findings that would overwhelm the report

Only mark clear false positives. When in doubt, keep the finding.
Return [] if all findings are valid.
ONLY return a JSON array of integers, nothing else."""

            response = client.models.generate_content(
                model=settings.gemini_model_id,
                contents=verification_prompt,
                config=types.GenerateContentConfig(
                    thinking_config={"thinking_level": "low"},
                    max_output_tokens=200,
                )
            )
            
            if response and response.text:
                try:
                    response_text = response.text.strip()
                    if response_text.startswith("```"):
                        response_text = response_text.split("```")[1]
                        if response_text.startswith("json"):
                            response_text = response_text[4:]
                    
                    fp_indices = json.loads(response_text.strip())
                    if isinstance(fp_indices, list):
                        fp_set = set(fp_indices)
                        verified_findings = [
                            f for i, f in enumerate(consolidated_findings) 
                            if i not in fp_set
                        ]
                        false_positives_removed = len(fp_set)
                        logger.info(f"Pass 3 verification removed {false_positives_removed} false positives")
                except json.JSONDecodeError:
                    logger.warning("Pass 3 verification response was not valid JSON")
                    
        except Exception as e:
            logger.warning(f"Pass 3 AI verification failed: {e}")
    
    logger.info(f"Final finding count: {len(verified_findings)} (removed {false_positives_removed} FPs)")
    
    # =========================================================================
    # Enrich and Generate Final Report
    # =========================================================================
    
    # Enrich findings with vulnerability intelligence
    enriched_findings = []
    for finding in verified_findings:
        enriched = enrich_finding_with_intelligence(finding)
        enriched_findings.append(enriched)
    
    # Calculate risk score (weighted by severity)
    severity_weights = {"critical": 40, "high": 25, "medium": 10, "low": 5, "info": 1}
    risk_score = min(100, sum(severity_weights.get(f["severity"], 5) for f in enriched_findings))
    
    # Generate attack paths
    attack_paths = generate_attack_paths(enriched_findings)
    
    # Get exploit references based on findings
    exploit_keywords = []
    for f in enriched_findings:
        if "csp" in f.get("title", "").lower():
            exploit_keywords.extend(["xss", "csp bypass"])
        if "cors" in f.get("title", "").lower():
            exploit_keywords.extend(["cors", "cross-origin"])
        if "hsts" in f.get("title", "").lower():
            exploit_keywords.extend(["ssl stripping", "hsts"])
        if "cookie" in f.get("title", "").lower():
            exploit_keywords.extend(["session hijacking", "cookie theft"])
    
    exploit_references = get_offline_exploit_references(list(set(exploit_keywords)))
    
    # Search CVE database for related vulnerabilities
    cve_references = []
    if detected_technologies:
        try:
            cve_keywords = list(detected_technologies)[:3]
            cve_references = await search_cve_database(cve_keywords, max_results=5)
        except Exception as e:
            logger.warning(f"CVE lookup failed: {e}")
    
    # Generate enhanced recommendations
    recommendations = []
    categories_found = set(f["category"] for f in enriched_findings)
    
    if "sensitive_data" in categories_found:
        recommendations.append({
            "priority": "critical",
            "title": "Protect Sensitive Data in Transit",
            "description": "Implement TLS 1.3, avoid logging sensitive data, use secure storage"
        })
    if any("csp" in cat for cat in categories_found):
        recommendations.append({
            "priority": "high",
            "title": "Implement Content Security Policy",
            "description": "Add strict CSP to prevent XSS: default-src 'self'; script-src 'self'"
        })
    if any("hsts" in cat for cat in categories_found):
        recommendations.append({
            "priority": "high",
            "title": "Enable HTTP Strict Transport Security",
            "description": "Add HSTS header: Strict-Transport-Security: max-age=31536000; includeSubDomains"
        })
    if "cors" in categories_found:
        recommendations.append({
            "priority": "high",
            "title": "Restrict CORS Policy",
            "description": "Whitelist specific origins, never use * with credentials"
        })
    if any("cookie" in cat for cat in categories_found):
        recommendations.append({
            "priority": "medium",
            "title": "Secure Cookie Configuration",
            "description": "Add HttpOnly, Secure, and SameSite=Strict to all session cookies"
        })
    if "information_disclosure" in categories_found:
        recommendations.append({
            "priority": "medium",
            "title": "Implement Proper Error Handling",
            "description": "Return generic error messages, log details server-side only"
        })
    
    # AI-powered comprehensive writeup
    ai_writeup = None
    ai_analysis = None
    
    if settings.gemini_api_key and len(traffic_log) > 0:
        # Generate comprehensive AI writeup
        traffic_summary = []
        for entry in traffic_log[:30]:
            req = entry.get("request", {})
            resp = entry.get("response", {})
            traffic_summary.append({
                "method": req.get("method"),
                "path": req.get("path"),
                "status": resp.get("status_code") if resp else None
            })
        
        ai_writeup = await generate_ai_exploitation_writeup(
            enriched_findings,
            traffic_summary,
            proxy_config,
            agent_activity
        )

        # Append deterministic agentic activity summary for report clarity
        if agent_activity and ai_writeup:
            exec_log = agent_activity.get("execution_log", [])
            verifications = agent_activity.get("verification_results", [])
            verification_lookup = {v.get("tool_id"): v for v in verifications}
            captured = agent_activity.get("captured_data_summary", {})
            decision_log = agent_activity.get("decision_log", [])

            summary_lines = [
                "",
                "### Agentic Tool Activity Summary",
                "",
                f"- **Monitoring Active:** {agent_activity.get('monitoring_active', False)}",
                f"- **Captured Data:** Credentials={captured.get('credentials', 0)}, Tokens={captured.get('tokens', 0)}, Cookies={captured.get('cookies', 0)}",
                "",
                "**Executed Tools:**",
            ]

            if exec_log:
                for log in exec_log:
                    verification = verification_lookup.get(log.get("tool_id"), {})
                    summary_lines.append(
                        f"- {log.get('tool_name', log.get('tool_id'))}: "
                        f"success={log.get('success')} | findings={log.get('findings_count', 0)} | "
                        f"verified={verification.get('success', False)}"
                    )
            else:
                summary_lines.append("- None")

            summary_lines.extend(["", "**Decision Log (latest session):**"]) 
            if decision_log:
                for entry in decision_log[:10]:
                    summary_lines.append(
                        f"- {entry.get('step', 'step')}: {entry.get('decision', 'decision')}"
                        f"{(' | tool=' + entry.get('tool')) if entry.get('tool') else ''}"
                        f"{(' | reason=' + entry.get('reason')) if entry.get('reason') else ''}"
                    )
            else:
                summary_lines.append("- None")

            ai_writeup = ai_writeup + "\n" + "\n".join(summary_lines)
        
        # Also generate quick summary
        try:
            from google import genai
            from google.genai import types
            
            client = genai.Client(api_key=settings.gemini_api_key)
            
            findings_summary = [{"title": f["title"], "severity": f["severity"]} for f in enriched_findings[:10]]
            
            prompt = f"""As a penetration tester, provide a 2-sentence security posture summary for this application.

Target: {proxy_config.get('target_host', 'Unknown')}:{proxy_config.get('target_port', 'Unknown')}
Findings: {json.dumps(findings_summary)}
Attack Paths Identified: {len(attack_paths)}

Be direct and actionable. Example: "The application exhibits weak browser-side security with missing CSP and HttpOnly flags, creating a high-risk XSS attack surface. Immediate remediation of security headers is recommended before production deployment."
"""
            
            response = client.models.generate_content(
                model=settings.gemini_model_id,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config={"thinking_level": "medium"},
                    max_output_tokens=200,
                )
            )
            
            if response and response.text:
                ai_analysis = response.text
                
        except Exception as e:
            logger.warning(f"AI quick analysis failed: {e}")
    
    # Determine risk level
    if risk_score >= 70:
        risk_level = "critical"
    elif risk_score >= 50:
        risk_level = "high"
    elif risk_score >= 25:
        risk_level = "medium"
    else:
        risk_level = "low"
    
    # Calculate analysis stats
    analysis_stats = {
        "pass1_findings": len(pass1_findings),
        "pass2_ai_findings": len(pass2_findings),
        "after_dedup": len(consolidated_findings),
        "false_positives_removed": false_positives_removed,
        "final_count": len(enriched_findings)
    }
    
    logger.info(f"3-pass analysis complete: {analysis_stats}")
    
    return {
        "summary": f"Analyzed {len(traffic_log)} traffic entries with 3-pass analysis, found {len(enriched_findings)} verified security issues",
        "risk_score": risk_score,
        "risk_level": risk_level,
        "findings": enriched_findings,
        "recommendations": recommendations,
        "ai_analysis": ai_analysis,
        "ai_writeup": ai_writeup,
        "agent_activity": agent_activity or {},
        "attack_paths": attack_paths,
        "exploit_references": exploit_references[:5],
        "cve_references": cve_references,
        "detected_technologies": list(detected_technologies),
        "traffic_analyzed": len(traffic_log),
        "rules_active": len(rules),
        "analysis_passes": 3,
        "analysis_stats": analysis_stats
    }


# ============================================================================
# Export MITM Analysis Reports
# ============================================================================

def _format_evidence_as_code(evidence: str) -> str:
    """
    Format evidence as a proper markdown code block with language detection.

    Detects common code patterns and wraps in appropriate code block.
    """
    if not evidence or evidence == "N/A":
        return "`N/A`"

    # Check if it's multi-line or contains code patterns
    is_multiline = "\n" in evidence

    # Detect language patterns
    lang = ""
    evidence_lower = evidence.lower()

    if any(p in evidence_lower for p in ["cookie:", "set-cookie:", "authorization:", "content-type:"]):
        lang = "http"
    elif evidence.strip().startswith("{") or evidence.strip().startswith("["):
        lang = "json"
    elif any(p in evidence_lower for p in ["<script", "<html", "<div", "</", "/>"]):
        lang = "html"
    elif any(p in evidence_lower for p in ["select ", "insert ", "update ", "delete ", " from ", " where "]):
        lang = "sql"
    elif any(p in evidence_lower for p in ["def ", "import ", "class ", "print(", "self."]):
        lang = "python"
    elif any(p in evidence_lower for p in ["function ", "const ", "let ", "var ", "=>"]):
        lang = "javascript"
    elif any(p in evidence_lower for p in ["curl ", "wget ", "bash", "#!/"]):
        lang = "bash"
    elif any(p in evidence_lower for p in ["get ", "post ", "put ", "patch ", "delete "]) and "http" in evidence_lower:
        lang = "http"

    # If multi-line or detected as code, use code block
    if is_multiline or lang:
        # Escape any existing code block markers
        evidence_safe = evidence.replace("```", "\\`\\`\\`")
        return f"```{lang}\n{evidence_safe}\n```"
    else:
        # Single line, use inline code
        return f"`{evidence}`"


def _format_execution_log_entry(entry: Dict) -> str:
    """Format an execution log entry for markdown output."""
    tool_id = entry.get("tool_id", "unknown")
    success = "Success" if entry.get("success") else "Failed"
    timestamp = entry.get("timestamp", "")

    lines = [f"- **{tool_id}** ({success})"]

    if timestamp:
        lines[0] += f" - {timestamp}"

    if entry.get("findings"):
        lines.append(f"  - Findings: {len(entry['findings'])}")

    if entry.get("captured_data"):
        captured = entry["captured_data"]
        creds = len(captured.get("credentials", []))
        tokens = len(captured.get("tokens", []))
        if creds or tokens:
            lines.append(f"  - Captured: {creds} credentials, {tokens} tokens")

    if entry.get("error"):
        lines.append(f"  - Error: {entry['error']}")

    return "\n".join(lines)


def generate_mitm_markdown_report(
    proxy_config: Dict,
    traffic_log: List[Dict],
    rules: List[Dict],
    analysis: Dict,
    agentic_data: Optional[Dict] = None,
) -> str:
    """Generate comprehensive Markdown report for MITM analysis.

    Args:
        proxy_config: Proxy configuration details
        traffic_log: Captured traffic entries
        rules: Active interception rules
        analysis: Security analysis results
        agentic_data: Optional dict with phases_executed, mitre_techniques,
                      attack_chains_executed, reasoning_traces, execution_log
    """
    lines = []
    
    # Header
    lines.extend([
        "# 🔐 Man-in-the-Middle Traffic Analysis Report",
        "",
        "---",
        "",
        "## 📋 Executive Summary",
        "",
        f"**Risk Level:** {analysis.get('risk_level', 'unknown').upper()}",
        f"**Risk Score:** {analysis.get('risk_score', 0)}/100",
        f"**Traffic Analyzed:** {analysis.get('traffic_analyzed', 0)} requests",
        f"**Security Issues Found:** {len(analysis.get('findings', []))}",
        "",
    ])
    
    if analysis.get("summary"):
        lines.extend([analysis["summary"], ""])
    
    # Proxy Configuration
    lines.extend([
        "---",
        "",
        "## ⚙️ Proxy Configuration",
        "",
        "| Property | Value |",
        "|----------|-------|",
        f"| Listen Address | `{proxy_config.get('listen_host', 'N/A')}:{proxy_config.get('listen_port', 'N/A')}` |",
        f"| Target Address | `{proxy_config.get('target_host', 'N/A')}:{proxy_config.get('target_port', 'N/A')}` |",
        f"| Mode | {proxy_config.get('mode', 'N/A')} |",
        f"| TLS Enabled | {'Yes' if proxy_config.get('tls_enabled') else 'No'} |",
        "",
    ])
    
    # AI Analysis (if available)
    if analysis.get("ai_analysis"):
        lines.extend([
            "---",
            "",
            "## 🤖 AI Security Analysis",
            "",
            analysis["ai_analysis"],
            "",
        ])
    
    # Comprehensive AI Pentest Writeup (if available)
    if analysis.get("ai_writeup"):
        lines.extend([
            "---",
            "",
            "## 📝 Comprehensive Penetration Test Report",
            "",
            analysis["ai_writeup"],
            "",
        ])
    
    # Attack Paths (if available)
    attack_paths = analysis.get("attack_paths", [])
    if attack_paths:
        lines.extend([
            "---",
            "",
            "## ⚔️ Attack Paths",
            "",
        ])
        
        for i, path in enumerate(attack_paths, 1):
            lines.extend([
                f"### Attack Path {i}: {path.get('name', 'Unknown')}",
                "",
                f"**Risk:** {path.get('risk', 'N/A')}",
                "",
                f"**Description:** {path.get('description', 'N/A')}",
                "",
            ])
            steps = path.get('steps', [])
            if steps:
                lines.append("**Steps:**")
                lines.append("")
                for j, step in enumerate(steps, 1):
                    lines.append(f"{j}. {step}")
                lines.append("")
        lines.append("")
    
    # CVE References (if available)
    cve_refs = analysis.get("cve_references", [])
    if cve_refs:
        lines.extend([
            "---",
            "",
            "## 🔍 Related CVE References",
            "",
            "| CVE ID | Description | Severity |",
            "|--------|-------------|----------|",
        ])
        
        for cve in cve_refs[:10]:
            cve_id = cve.get('id', 'N/A')
            desc = cve.get('description', 'N/A')[:80] + "..." if len(cve.get('description', '')) > 80 else cve.get('description', 'N/A')
            severity = cve.get('severity', 'N/A')
            lines.append(f"| {cve_id} | {desc} | {severity} |")
        
        lines.append("")
    
    # Exploit References (if available)
    exploit_refs = analysis.get("exploit_references", [])
    if exploit_refs:
        lines.extend([
            "---",
            "",
            "## 💣 Exploit Database References",
            "",
        ])
        
        for exploit in exploit_refs[:5]:
            lines.extend([
                f"- **{exploit.get('title', 'Unknown')}** ({exploit.get('platform', 'N/A')})",
                f"  - Type: {exploit.get('type', 'N/A')}",
                f"  - Source: {exploit.get('source', 'N/A')}",
                "",
            ])
    
    # Security Findings
    findings = analysis.get("findings", [])
    if findings:
        lines.extend([
            "---",
            "",
            "## 🚨 Security Findings",
            "",
        ])
        
        severity_order = ["critical", "high", "medium", "low", "info"]
        findings_by_severity = {s: [] for s in severity_order}
        for f in findings:
            sev = f.get("severity", "info").lower()
            if sev in findings_by_severity:
                findings_by_severity[sev].append(f)
        
        severity_emoji = {
            "critical": "🔴",
            "high": "🟠",
            "medium": "🟡",
            "low": "🟢",
            "info": "🔵"
        }
        
        for severity in severity_order:
            sev_findings = findings_by_severity[severity]
            if sev_findings:
                emoji = severity_emoji.get(severity, "⚪")
                lines.extend([
                    f"### {emoji} {severity.upper()} ({len(sev_findings)})",
                    "",
                ])
                
                for i, f in enumerate(sev_findings, 1):
                    evidence = f.get('evidence', 'N/A')
                    # Format evidence as code block if it's multi-line or contains code patterns
                    evidence_formatted = _format_evidence_as_code(evidence)

                    lines.extend([
                        f"#### {i}. {f.get('title', 'Unknown')}",
                        "",
                        f"**Category:** {f.get('category', 'N/A')}",
                        "",
                        f"**Affected Endpoint:** `{f.get('affected_endpoint', 'N/A')}`",
                        "",
                        f"**Description:** {f.get('description', 'N/A')}",
                        "",
                        "**Evidence:**",
                        "",
                        evidence_formatted,
                        "",
                        f"**Recommendation:** {f.get('recommendation', 'N/A')}",
                        "",
                        "---",
                        "",
                    ])
    
    # Active Rules
    if rules:
        lines.extend([
            "## 📝 Active Interception Rules",
            "",
            "| Rule Name | Direction | Action | Enabled |",
            "|-----------|-----------|--------|---------|",
        ])
        
        for rule in rules:
            enabled = "✅" if rule.get("enabled") else "❌"
            lines.append(f"| {rule.get('name', 'N/A')} | {rule.get('match_direction', 'N/A')} | {rule.get('action', 'N/A')} | {enabled} |")
        
        lines.append("")
    
    # Traffic Sample
    if traffic_log:
        lines.extend([
            "---",
            "",
            "## 📊 Traffic Sample",
            "",
            "| Time | Method | Path | Status | Modified |",
            "|------|--------|------|--------|----------|",
        ])
        
        for entry in traffic_log[:20]:
            req = entry.get("request", {})
            resp = entry.get("response", {})
            timestamp = entry.get("timestamp", "N/A")
            if isinstance(timestamp, str) and "T" in timestamp:
                timestamp = timestamp.split("T")[1].split(".")[0]
            
            method = req.get("method", "?")
            path = req.get("path", "/")[:40]
            status = resp.get("status_code", "-") if resp else "-"
            modified = "✏️" if entry.get("modified") else ""
            
            lines.append(f"| {timestamp} | {method} | `{path}` | {status} | {modified} |")
        
        if len(traffic_log) > 20:
            lines.append(f"| ... | *{len(traffic_log) - 20} more entries* | ... | ... | |")
        
        lines.append("")
    
    # Recommendations
    recommendations = analysis.get("recommendations", [])
    if recommendations:
        lines.extend([
            "---",
            "",
            "## Recommendations",
            "",
        ])

        for i, rec in enumerate(recommendations, 1):
            lines.append(f"{i}. {rec}")

        lines.append("")

    # =========================================================================
    # Agentic Session Data (if available)
    # =========================================================================
    if agentic_data:
        # MITRE ATT&CK Techniques
        mitre = agentic_data.get("mitre_techniques")
        if mitre and mitre.get("techniques_used"):
            lines.extend([
                "---",
                "",
                "## MITRE ATT&CK Mapping",
                "",
            ])

            techniques = mitre.get("technique_details", [])
            if techniques:
                lines.append("| Technique ID | Name | Tactic | Description |")
                lines.append("|--------------|------|--------|-------------|")
                for tech in techniques:
                    tid = tech.get("technique_id", "N/A")
                    name = tech.get("name", "N/A")
                    tactic = tech.get("tactic", "N/A")
                    desc = tech.get("description", "")[:60] + "..." if len(tech.get("description", "")) > 60 else tech.get("description", "N/A")
                    lines.append(f"| {tid} | {name} | {tactic} | {desc} |")
                lines.append("")
            else:
                for tid in mitre.get("techniques_used", []):
                    lines.append(f"- {tid}")
                lines.append("")

        # Attack Phases
        phases = agentic_data.get("phases_executed")
        if phases:
            lines.extend([
                "---",
                "",
                "## Attack Phase Progression",
                "",
                f"**Current Phase:** {phases.get('current_phase', 'N/A')}",
                "",
            ])

            phase_history = phases.get("phase_history", [])
            if phase_history:
                lines.append("### Phase History")
                lines.append("")
                for ph in phase_history:
                    phase_name = ph.get("phase", "N/A")
                    timestamp = ph.get("timestamp", "")
                    goals = ph.get("goals_achieved", [])
                    lines.append(f"- **{phase_name}** ({timestamp})")
                    for goal in goals:
                        lines.append(f"  - Goal achieved: {goal}")
                lines.append("")

        # Attack Chains Executed
        chains = agentic_data.get("attack_chains_executed")
        if chains and chains.get("execution_history"):
            lines.extend([
                "---",
                "",
                "## Attack Chains Executed",
                "",
            ])

            for exec_entry in chains.get("execution_history", [])[:10]:
                chain_id = exec_entry.get("chain_id", "N/A")
                status = exec_entry.get("status", "N/A")
                started = exec_entry.get("started_at", "")
                tools_run = len(exec_entry.get("tools_executed", []))
                lines.append(f"- **{chain_id}** - {status} ({tools_run} tools) - {started}")
            lines.append("")

            stats = chains.get("stats")
            if stats:
                lines.append(f"**Total Chains Executed:** {stats.get('total_executions', 0)}")
                lines.append(f"**Successful:** {stats.get('successful', 0)}")
                lines.append(f"**Failed:** {stats.get('failed', 0)}")
                lines.append("")

        # Execution Log
        exec_log = agentic_data.get("execution_log")
        if exec_log:
            lines.extend([
                "---",
                "",
                "## Tool Execution Log",
                "",
            ])

            for entry in exec_log[:20]:
                lines.append(_format_execution_log_entry(entry))
            if len(exec_log) > 20:
                lines.append(f"- *... and {len(exec_log) - 20} more executions*")
            lines.append("")

        # Reasoning Traces Summary
        reasoning = agentic_data.get("reasoning_traces")
        if reasoning:
            lines.extend([
                "---",
                "",
                "## Agent Reasoning Traces",
                "",
                f"**Total Reasoning Chains:** {len(reasoning)}",
                "",
            ])

            for trace in reasoning[:5]:
                tool_id = trace.get("tool_id", "N/A")
                decision = trace.get("decision", "N/A")
                confidence = trace.get("confidence", 0)
                lines.append(f"### {tool_id} (Confidence: {confidence:.1%})")
                lines.append("")
                lines.append(f"**Decision:** {decision}")
                lines.append("")

                steps = trace.get("steps", [])
                if steps:
                    lines.append("**Reasoning Steps:**")
                    lines.append("")
                    for i, step in enumerate(steps, 1):
                        step_type = step.get("type", "")
                        content = step.get("content", "")
                        lines.append(f"{i}. **{step_type}:** {content[:200]}...")
                    lines.append("")

            if len(reasoning) > 5:
                lines.append(f"*... and {len(reasoning) - 5} more reasoning chains*")
                lines.append("")

        # Captured Data Summary
        captured = agentic_data.get("captured_data")
        if captured:
            creds = captured.get("credentials", [])
            tokens = captured.get("tokens", [])
            cookies = captured.get("cookies", [])

            if creds or tokens or cookies:
                lines.extend([
                    "---",
                    "",
                    "## Captured Sensitive Data",
                    "",
                    f"- **Credentials:** {len(creds)}",
                    f"- **Tokens:** {len(tokens)}",
                    f"- **Session Cookies:** {len(cookies)}",
                    "",
                ])

                if creds:
                    lines.append("### Captured Credentials")
                    lines.append("")
                    lines.append("```")
                    for cred in creds[:5]:
                        # Redact actual passwords
                        username = cred.get("username", "N/A")
                        cred_type = cred.get("type", "N/A")
                        source = cred.get("source", "N/A")
                        lines.append(f"Type: {cred_type}, User: {username}, Source: {source}")
                    if len(creds) > 5:
                        lines.append(f"... and {len(creds) - 5} more")
                    lines.append("```")
                    lines.append("")

    # Footer
    lines.extend([
        "---",
        "",
        "*Report generated by VRAgent MITM Workbench*",
        "",
    ])

    return "\n".join(lines)


def generate_mitm_pdf_report(
    proxy_config: Dict,
    traffic_log: List[Dict],
    rules: List[Dict],
    analysis: Dict
) -> bytes:
    """Generate PDF report for MITM analysis."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        import re
        from io import BytesIO
    except ImportError:
        logger.error("reportlab not installed")
        return b"%PDF-1.4 placeholder"
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []
    writeup_body_style = ParagraphStyle(
        'WriteupBody', parent=styles['Normal'], fontSize=11, leading=15
    )
    writeup_heading_style = ParagraphStyle(
        'WriteupHeading', parent=styles['Heading3'], fontSize=14, spaceBefore=10, spaceAfter=6
    )
    writeup_subheading_style = ParagraphStyle(
        'WriteupSubHeading', parent=styles['Heading4'], fontSize=12, spaceBefore=8, spaceAfter=4
    )
    
    # Title
    title_style = ParagraphStyle(
        'Title', parent=styles['Heading1'], fontSize=24, spaceAfter=20, alignment=1
    )
    story.append(Paragraph("🔐 MITM Traffic Analysis Report", title_style))
    story.append(Spacer(1, 20))
    
    # Executive Summary
    story.append(Paragraph("Executive Summary", styles['Heading2']))
    
    summary_data = [
        ["Risk Level", analysis.get('risk_level', 'unknown').upper()],
        ["Risk Score", f"{analysis.get('risk_score', 0)}/100"],
        ["Traffic Analyzed", str(analysis.get('traffic_analyzed', 0))],
        ["Issues Found", str(len(analysis.get('findings', [])))],
    ]
    
    summary_table = Table(summary_data, colWidths=[2*inch, 3*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 20))
    
    # AI Analysis
    if analysis.get("ai_analysis"):
        story.append(Paragraph("AI Security Analysis", styles['Heading2']))
        # Clean text for PDF
        ai_text = analysis["ai_analysis"].replace("**", "").replace("*", "")
        story.append(Paragraph(ai_text, styles['Normal']))
        story.append(Spacer(1, 20))
    
    # Comprehensive AI Pentest Writeup
    if analysis.get("ai_writeup"):
        story.append(PageBreak())
        story.append(Paragraph("Comprehensive Penetration Test Report", styles['Heading2']))
        writeup_text = analysis["ai_writeup"]

        def _fmt_bold(text: str) -> str:
            return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)

        for line in writeup_text.splitlines():
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 6))
                continue

            if stripped.startswith("####"):
                story.append(Paragraph(stripped.lstrip("#").strip(), writeup_subheading_style))
                continue
            if stripped.startswith("###"):
                story.append(Paragraph(stripped.lstrip("#").strip(), writeup_heading_style))
                continue
            if stripped.startswith("##"):
                story.append(Paragraph(stripped.lstrip("#").strip(), styles['Heading2']))
                continue
            if stripped.startswith("#"):
                story.append(Paragraph(stripped.lstrip("#").strip(), styles['Heading1']))
                continue

            numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
            if numbered:
                story.append(Paragraph(_fmt_bold(numbered.group(2)), writeup_body_style, bulletText=f"{numbered.group(1)}."))
                continue

            if stripped.startswith("- ") or stripped.startswith("* "):
                story.append(Paragraph(_fmt_bold(stripped[2:]), writeup_body_style, bulletText="•"))
                continue

            story.append(Paragraph(_fmt_bold(stripped), writeup_body_style))
        story.append(Spacer(1, 20))
    
    # Attack Paths
    attack_paths = analysis.get("attack_paths", [])
    if attack_paths:
        story.append(Paragraph("Attack Paths", styles['Heading2']))
        for i, path in enumerate(attack_paths, 1):
            story.append(Paragraph(f"{i}. {path.get('name', 'Unknown')} (Risk: {path.get('risk', 'N/A')})", styles['Heading3']))
            story.append(Paragraph(path.get('description', 'N/A'), styles['Normal']))
            steps = path.get('steps', [])
            if steps:
                story.append(Paragraph("Steps:", styles['Normal']))
                for j, step in enumerate(steps, 1):
                    story.append(Paragraph(f"    {j}. {step}", styles['Normal']))
            story.append(Spacer(1, 10))
        story.append(Spacer(1, 10))
    
    # CVE References
    cve_refs = analysis.get("cve_references", [])
    if cve_refs:
        story.append(Paragraph("Related CVE References", styles['Heading2']))
        cve_data = [["CVE ID", "Severity", "Description"]]
        for cve in cve_refs[:10]:
            desc = cve.get('description', 'N/A')[:60] + "..." if len(cve.get('description', '')) > 60 else cve.get('description', 'N/A')
            cve_data.append([cve.get('id', 'N/A'), cve.get('severity', 'N/A'), desc])
        
        cve_table = Table(cve_data, colWidths=[1.2*inch, 0.8*inch, 4*inch])
        cve_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#333333')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 4),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(cve_table)
        story.append(Spacer(1, 20))
    
    # Exploit References
    exploit_refs = analysis.get("exploit_references", [])
    if exploit_refs:
        story.append(Paragraph("Exploit Database References", styles['Heading2']))
        for exploit in exploit_refs[:5]:
            story.append(Paragraph(f"• {exploit.get('title', 'Unknown')} ({exploit.get('platform', 'N/A')})", styles['Normal']))
            story.append(Paragraph(f"  Type: {exploit.get('type', 'N/A')} | Source: {exploit.get('source', 'N/A')}", styles['Normal']))
            story.append(Spacer(1, 6))
        story.append(Spacer(1, 10))
    
    # Findings
    findings = analysis.get("findings", [])
    if findings:
        story.append(Paragraph("Security Findings", styles['Heading2']))
        
        for i, f in enumerate(findings[:15], 1):
            story.append(Paragraph(f"{i}. {f.get('title', 'Unknown')}", styles['Heading3']))
            story.append(Paragraph(f"Severity: {f.get('severity', 'N/A').upper()} | Category: {f.get('category', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"Description: {f.get('description', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"Recommendation: {f.get('recommendation', 'N/A')}", styles['Normal']))
            story.append(Spacer(1, 10))
    
    # Recommendations
    recommendations = analysis.get("recommendations", [])
    if recommendations:
        story.append(Paragraph("Recommendations", styles['Heading2']))
        for i, rec in enumerate(recommendations, 1):
            story.append(Paragraph(f"{i}. {rec}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    story.append(Spacer(1, 20))
    story.append(Paragraph("Report generated by VRAgent MITM Workbench", 
        ParagraphStyle('footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey, alignment=1)))
    
    doc.build(story)
    return buffer.getvalue()


def generate_mitm_docx_report(
    proxy_config: Dict,
    traffic_log: List[Dict],
    rules: List[Dict],
    analysis: Dict
) -> bytes:
    """Generate DOCX report for MITM analysis."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        import re
    except ImportError:
        logger.error("python-docx not installed")
        return b"PK placeholder"
    
    doc = Document()
    
    # Title
    title = doc.add_heading("🔐 MITM Traffic Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ("Risk Level", analysis.get('risk_level', 'unknown').upper()),
        ("Risk Score", f"{analysis.get('risk_score', 0)}/100"),
        ("Traffic Analyzed", str(analysis.get('traffic_analyzed', 0))),
        ("Issues Found", str(len(analysis.get('findings', [])))),
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
        summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # AI Analysis
    if analysis.get("ai_analysis"):
        doc.add_heading("AI Security Analysis", level=1)
        doc.add_paragraph(analysis["ai_analysis"])
    
    # Comprehensive AI Pentest Writeup
    if analysis.get("ai_writeup"):
        doc.add_page_break()
        doc.add_heading("Comprehensive Penetration Test Report", level=1)
        writeup_text = analysis["ai_writeup"]
        def _add_bold_runs(paragraph, text: str):
            parts = text.split("**")
            for j, part in enumerate(parts):
                run = paragraph.add_run(part)
                if j % 2 == 1:
                    run.bold = True
                run.font.size = Pt(11)

        for line in writeup_text.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            if stripped.startswith("####"):
                doc.add_heading(stripped.lstrip("#").strip(), level=4)
                continue
            if stripped.startswith("###"):
                doc.add_heading(stripped.lstrip("#").strip(), level=3)
                continue
            if stripped.startswith("##"):
                doc.add_heading(stripped.lstrip("#").strip(), level=2)
                continue
            if stripped.startswith("#"):
                doc.add_heading(stripped.lstrip("#").strip(), level=1)
                continue

            numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
            if numbered:
                p = doc.add_paragraph(style='List Number')
                _add_bold_runs(p, numbered.group(2))
                continue

            if stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                _add_bold_runs(p, stripped[2:])
                continue

            p = doc.add_paragraph()
            _add_bold_runs(p, stripped)
    
    # Attack Paths
    attack_paths = analysis.get("attack_paths", [])
    if attack_paths:
        doc.add_heading("Attack Paths", level=1)
        for i, path in enumerate(attack_paths, 1):
            doc.add_heading(f"{i}. {path.get('name', 'Unknown')}", level=2)
            p = doc.add_paragraph()
            p.add_run("Risk Level: ").bold = True
            p.add_run(path.get('risk', 'N/A'))
            doc.add_paragraph(path.get('description', 'N/A'))
            steps = path.get('steps', [])
            if steps:
                p = doc.add_paragraph()
                p.add_run("Attack Steps:").bold = True
                for j, step in enumerate(steps, 1):
                    doc.add_paragraph(f"{j}. {step}", style='List Number')
            doc.add_paragraph()
    
    # CVE References
    cve_refs = analysis.get("cve_references", [])
    if cve_refs:
        doc.add_heading("Related CVE References", level=1)
        cve_table = doc.add_table(rows=1, cols=3)
        cve_table.style = 'Table Grid'
        header_cells = cve_table.rows[0].cells
        header_cells[0].text = "CVE ID"
        header_cells[1].text = "Severity"
        header_cells[2].text = "Description"
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        for cve in cve_refs[:10]:
            row = cve_table.add_row().cells
            row[0].text = cve.get('id', 'N/A')
            row[1].text = cve.get('severity', 'N/A')
            desc = cve.get('description', 'N/A')
            row[2].text = desc[:80] + "..." if len(desc) > 80 else desc
        doc.add_paragraph()
    
    # Exploit References
    exploit_refs = analysis.get("exploit_references", [])
    if exploit_refs:
        doc.add_heading("Exploit Database References", level=1)
        for exploit in exploit_refs[:5]:
            p = doc.add_paragraph()
            p.add_run(f"• {exploit.get('title', 'Unknown')}").bold = True
            p.add_run(f" ({exploit.get('platform', 'N/A')})")
            doc.add_paragraph(f"  Type: {exploit.get('type', 'N/A')} | Source: {exploit.get('source', 'N/A')}")
        doc.add_paragraph()
    
    # Findings
    findings = analysis.get("findings", [])
    if findings:
        doc.add_heading("Security Findings", level=1)
        
        for i, f in enumerate(findings[:15], 1):
            doc.add_heading(f"{i}. {f.get('title', 'Unknown')}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Severity: ").bold = True
            p.add_run(f"{f.get('severity', 'N/A').upper()} | ")
            p.add_run("Category: ").bold = True
            p.add_run(f.get('category', 'N/A'))
            
            p = doc.add_paragraph()
            p.add_run("Description: ").bold = True
            p.add_run(f.get('description', 'N/A'))
            
            p = doc.add_paragraph()
            p.add_run("Recommendation: ").bold = True
            p.add_run(f.get('recommendation', 'N/A'))
            
            doc.add_paragraph()
    
    # Recommendations
    recommendations = analysis.get("recommendations", [])
    if recommendations:
        doc.add_heading("Recommendations", level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Report generated by VRAgent MITM Workbench").italic = True
    
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_mitm_pcap(traffic_log: List[Dict]) -> bytes:
    """Generate a PCAP file from MITM traffic entries."""
    try:
        from scapy.all import IP, TCP, Raw, wrpcap
    except Exception as exc:
        raise RuntimeError("scapy not installed. Install with: pip install scapy") from exc

    packets = []

    def build_http_request(req: Dict[str, Any]) -> bytes:
        method = req.get("method", "GET")
        path = req.get("path", "/")
        headers = req.get("headers", {}) or {}
        if "Host" not in headers and req.get("host"):
            headers["Host"] = req["host"]
        lines = [f"{method} {path} HTTP/1.1"]
        for key, value in headers.items():
            lines.append(f"{key}: {value}")
        lines.append("")
        lines.append(req.get("body") or req.get("body_text") or "")
        return "\r\n".join(lines).encode("utf-8", errors="replace")

    def build_http_response(resp: Dict[str, Any]) -> bytes:
        status = resp.get("status_code", 200)
        status_text = resp.get("status_text") or resp.get("status_message") or "OK"
        headers = resp.get("headers", {}) or {}
        lines = [f"HTTP/1.1 {status} {status_text}"]
        for key, value in headers.items():
            lines.append(f"{key}: {value}")
        lines.append("")
        lines.append(resp.get("body") or resp.get("body_text") or "")
        return "\r\n".join(lines).encode("utf-8", errors="replace")

    for entry in traffic_log:
        request = entry.get("request", {}) or {}
        response = entry.get("response", {}) or {}
        client_ip = request.get("client_ip") or "127.0.0.1"
        client_port = request.get("client_port") or 0
        host = request.get("host") or "127.0.0.1"
        dest_ip = host if re.match(r"^\\d+\\.\\d+\\.\\d+\\.\\d+$", host) else "127.0.0.1"
        dest_port = request.get("port") or 80

        req_bytes = build_http_request(request)
        packets.append(IP(src=client_ip, dst=dest_ip) / TCP(sport=client_port, dport=dest_port) / Raw(load=req_bytes))

        if response:
            resp_bytes = build_http_response(response)
            packets.append(IP(src=dest_ip, dst=client_ip) / TCP(sport=dest_port, dport=client_port) / Raw(load=resp_bytes))

    buffer = BytesIO()
    wrpcap(buffer, packets)
    return buffer.getvalue()


# ============================================================================
# AI-Powered Natural Language Rule Creation
# ============================================================================

async def create_rule_from_natural_language(description: str) -> Dict:
    """
    Convert natural language description to a MITM interception rule.
    
    Examples:
    - "Block all requests to analytics domains"
    - "Add a 2 second delay to all API responses"
    - "Remove authentication headers from requests"
    - "Replace all error responses with success"
    """
    from ..core.config import settings
    
    if not settings.gemini_api_key:
        # Fallback to pattern matching for common requests
        return _parse_rule_with_patterns(description)
    
    try:
        from google import genai
        from google.genai import types
        
        client = genai.Client(api_key=settings.gemini_api_key)
        
        prompt = f"""You are a MITM (Man-in-the-Middle) proxy rule generator. Convert the user's natural language request into a JSON rule configuration.

USER REQUEST: "{description}"

Generate a JSON object with these fields (include only relevant fields):
{{
    "name": "Rule name (descriptive)",
    "enabled": true,
    "match_direction": "request" | "response" | "both",
    "match_host": "regex pattern or null",
    "match_path": "regex pattern or null",
    "match_method": "GET|POST|PUT|DELETE or null",
    "match_content_type": "content type pattern or null",
    "match_status_code": number or null,
    "action": "modify" | "drop" | "delay",
    "modify_headers": {{"Header-Name": "value"}} or null,
    "remove_headers": ["header1", "header2"] or null,
    "body_find_replace": {{"find": "replace"}} or null,
    "delay_ms": number (only if action is delay)
}}

Common patterns:
- For blocking: action="drop"
- For adding headers: action="modify", modify_headers={{...}}
- For removing headers: action="modify", remove_headers=[...]
- For modifying body: action="modify", body_find_replace={{...}}
- For delays: action="delay", delay_ms=number

Return ONLY valid JSON, no explanation."""

        response = client.models.generate_content(
            model=settings.gemini_model_id,
            contents=prompt,
            config=types.GenerateContentConfig(
                thinking_config={"thinking_level": "low"},
                max_output_tokens=500,
            )
        )
        
        if response and response.text:
            # Parse the JSON response
            import json
            text = response.text.strip()
            # Remove markdown code blocks if present
            if text.startswith("```"):
                text = text.split("```")[1]
                if text.startswith("json"):
                    text = text[4:]
            text = text.strip()
            
            rule = json.loads(text)
            rule["ai_generated"] = True
            rule["original_description"] = description
            
            return {
                "success": True,
                "rule": rule,
                "message": f"Created rule: {rule.get('name', 'Unnamed Rule')}"
            }
            
    except Exception as e:
        logger.warning(f"AI rule generation failed: {e}, falling back to pattern matching")
    
    # Fallback to pattern matching
    return _parse_rule_with_patterns(description)


def _parse_rule_with_patterns(description: str) -> Dict:
    """Fallback pattern-based rule parsing when AI is unavailable."""
    desc_lower = description.lower()
    rule = {
        "name": description[:50],
        "enabled": True,
        "match_direction": "both",
        "action": "modify"
    }
    
    # Detect blocking/dropping
    if any(word in desc_lower for word in ["block", "drop", "reject", "stop", "prevent"]):
        rule["action"] = "drop"
        
        # Common block targets
        if "analytics" in desc_lower:
            rule["match_host"] = "(google-analytics|googletagmanager|facebook|analytics|mixpanel|segment)"
            rule["name"] = "Block Analytics"
        elif "tracking" in desc_lower or "tracker" in desc_lower:
            rule["match_host"] = "(tracking|tracker|pixel|beacon)"
            rule["name"] = "Block Trackers"
        elif "ads" in desc_lower or "advert" in desc_lower:
            rule["match_host"] = "(doubleclick|adsense|adserver|ad\\.|ads\\.)"
            rule["name"] = "Block Ads"
    
    # Detect delays
    elif any(word in desc_lower for word in ["delay", "slow", "latency", "wait"]):
        rule["action"] = "delay"
        # Extract delay time
        import re
        time_match = re.search(r'(\d+)\s*(second|sec|s|ms|millisecond)', desc_lower)
        if time_match:
            value = int(time_match.group(1))
            unit = time_match.group(2)
            if unit.startswith("s") and not unit.startswith("ms"):
                value *= 1000  # Convert to ms
            rule["delay_ms"] = value
        else:
            rule["delay_ms"] = 2000
        rule["name"] = f"Add {rule['delay_ms']}ms Delay"
        rule["match_direction"] = "response"
    
    # Detect header removal
    elif "remove" in desc_lower and "header" in desc_lower:
        headers_to_remove = []
        if "auth" in desc_lower:
            headers_to_remove.extend(["Authorization", "X-Auth-Token", "X-API-Key"])
        if "security" in desc_lower:
            headers_to_remove.extend(["Content-Security-Policy", "X-Frame-Options", "X-XSS-Protection"])
        if "cors" in desc_lower:
            headers_to_remove.extend(["Access-Control-Allow-Origin", "Access-Control-Allow-Methods"])
        if "cookie" in desc_lower:
            headers_to_remove.extend(["Set-Cookie", "Cookie"])
        if "csrf" in desc_lower:
            headers_to_remove.extend(["X-CSRF-Token", "X-XSRF-Token"])
        
        if headers_to_remove:
            rule["remove_headers"] = headers_to_remove
            rule["name"] = f"Remove {', '.join(headers_to_remove[:2])}..."
            rule["match_direction"] = "request" if "request" in desc_lower else "response" if "response" in desc_lower else "both"
    
    # Detect header addition
    elif "add" in desc_lower and "header" in desc_lower:
        rule["modify_headers"] = {}
        if "debug" in desc_lower:
            rule["modify_headers"]["X-Debug"] = "true"
        if "admin" in desc_lower:
            rule["modify_headers"]["X-Admin"] = "true"
            rule["modify_headers"]["X-User-Role"] = "admin"
        if "cors" in desc_lower or "cross-origin" in desc_lower:
            rule["modify_headers"]["Access-Control-Allow-Origin"] = "*"
            rule["modify_headers"]["Access-Control-Allow-Methods"] = "GET, POST, PUT, DELETE, OPTIONS"
        rule["name"] = "Add Custom Headers"
        rule["match_direction"] = "request" if "request" in desc_lower else "response"
    
    # Detect body modification
    elif any(word in desc_lower for word in ["replace", "change", "modify body", "swap"]):
        rule["body_find_replace"] = {}
        if "error" in desc_lower and "success" in desc_lower:
            rule["body_find_replace"]["\"success\":false"] = "\"success\":true"
            rule["body_find_replace"]["\"error\":"] = "\"_hidden_error\":"
        elif "false" in desc_lower and "true" in desc_lower:
            rule["body_find_replace"]["false"] = "true"
        rule["name"] = "Modify Response Body"
        rule["match_direction"] = "response"
    
    return {
        "success": True,
        "rule": rule,
        "message": f"Created rule using pattern matching: {rule.get('name')}",
        "ai_generated": False
    }


# ============================================================================
# Real-time AI Traffic Suggestions
# ============================================================================

async def get_ai_traffic_suggestions(
    traffic_log: List[Dict],
    existing_rules: List[Dict],
    proxy_config: Dict
) -> Dict:
    """
    Analyze current traffic and suggest security tests or rules.

    Returns suggestions based on observed traffic patterns.
    """
    from ..core.config import settings

    generated_at = datetime.utcnow().isoformat() + "Z"

    methods: Dict[str, int] = {}
    paths: List[str] = []
    hosts: List[str] = []
    status_codes: Dict[int, int] = {}
    has_auth = False
    has_json = False
    has_cors = False
    has_cookies = False
    error_count = 0

    for entry in traffic_log[:250]:  # Analyze up to 250 entries for statistics
        req = entry.get("request", {})
        resp = entry.get("response", {}) or {}

        method = req.get("method", "GET")
        methods[method] = methods.get(method, 0) + 1

        path = req.get("path", "/")
        host = req.get("host", "")
        paths.append(path)
        if host:
            hosts.append(host)

        if resp:
            status = resp.get("status_code", 0)
            status_codes[status] = status_codes.get(status, 0) + 1
            if status >= 400:
                error_count += 1

        req_headers = {k.lower(): v for k, v in req.get("headers", {}).items()}
        resp_headers = {k.lower(): v for k, v in resp.get("headers", {}).items()} if resp else {}

        if "authorization" in req_headers or "x-api-key" in req_headers:
            has_auth = True
        if "application/json" in str(req_headers.get("content-type", "")) or            "application/json" in str(resp_headers.get("content-type", "")):
            has_json = True
        if "access-control-allow-origin" in resp_headers:
            has_cors = True
        if "cookie" in req_headers or "set-cookie" in resp_headers:
            has_cookies = True

    traffic_summary = {
        "total_requests": len(traffic_log),
        "unique_hosts": sorted(set(hosts))[:10],
        "unique_paths": sorted(set(paths))[:20],
        "auth_detected": has_auth,
        "json_apis": has_json,
        "has_cookies": has_cookies
    }

    if len(traffic_log) == 0:
        return {
            "suggestions": [{
                "id": "no-traffic",
                "title": "No traffic yet",
                "description": "Start sending traffic through the proxy to unlock AI suggestions.",
                "category": "learning",
                "priority": "low",
                "natural_language": "Show me how to capture traffic"
            }],
            "traffic_summary": traffic_summary,
            "generated_at": generated_at
        }

    suggestions: List[Dict[str, Any]] = []
    existing_rule_names = [r.get("name", "").lower() for r in existing_rules]

    if has_auth and "auth" not in str(existing_rule_names):
        suggestions.append({
            "id": "auth-bypass",
            "title": "Test authentication bypass",
            "description": "Authorization headers detected. Remove auth headers to test for bypass vulnerabilities.",
            "category": "security",
            "priority": "high",
            "rule": {
                "name": "Remove Auth Headers",
                "match_direction": "request",
                "action": "modify",
                "remove_headers": ["Authorization", "X-API-Key", "X-Auth-Token"]
            },
            "natural_language": "Remove Authorization and API key headers from all requests"
        })

    if has_json:
        suggestions.append({
            "id": "json-tamper",
            "title": "Test response tampering",
            "description": "JSON API detected. Modify response values to test client-side validation.",
            "category": "security",
            "priority": "medium",
            "rule": {
                "name": "Modify JSON Response",
                "match_direction": "response",
                "match_content_type": "application/json",
                "action": "modify",
                "body_find_replace": {
                    '"success":false': '"success":true',
                    '"authorized":false': '"authorized":true'
                }
            },
            "natural_language": "Change JSON response fields so success and authorized are true"
        })

    if has_cors:
        suggestions.append({
            "id": "cors-test",
            "title": "Test CORS policy",
            "description": "CORS headers detected. Try permissive headers to see what data is exposed.",
            "category": "security",
            "priority": "medium",
            "rule": {
                "name": "Modify CORS Headers",
                "match_direction": "response",
                "action": "modify",
                "modify_headers": {
                    "Access-Control-Allow-Origin": "*",
                    "Access-Control-Allow-Credentials": "true"
                }
            },
            "natural_language": "Set Access-Control-Allow-Origin to * on all responses"
        })

    if has_cookies:
        suggestions.append({
            "id": "cookie-test",
            "title": "Review cookie security",
            "description": "Cookies detected. Check for missing HttpOnly, Secure, or SameSite flags.",
            "category": "security",
            "priority": "high",
            "rule": {
                "name": "Expose Cookie Values",
                "match_direction": "response",
                "action": "modify",
                "remove_headers": ["Set-Cookie"]
            },
            "natural_language": "Remove Set-Cookie headers from responses"
        })

    if error_count > 0:
        suggestions.append({
            "id": "error-disclosure",
            "title": "Inspect error responses",
            "description": f"Found {error_count} error responses (4xx/5xx). Look for information disclosure.",
            "category": "debug",
            "priority": "medium",
            "rule": None,
            "natural_language": "Show me error responses with stack traces"
        })

    api_paths = [p for p in paths if "/api/" in p or "/v1/" in p or "/v2/" in p]
    if len(api_paths) > 5:
        suggestions.append({
            "id": "api-discovery",
            "title": "Explore API surface",
            "description": f"Found {len(set(api_paths))} unique API paths. Consider fuzzing for undocumented endpoints.",
            "category": "learning",
            "priority": "low",
            "rule": None,
            "natural_language": "List all API endpoints seen in traffic"
        })

    admin_paths = [p for p in paths if any(x in p.lower() for x in ["admin", "manage", "dashboard", "internal"])]
    if admin_paths:
        suggestions.append({
            "id": "admin-access",
            "title": "Test admin access control",
            "description": f"Potential admin paths detected: {', '.join(sorted(set(admin_paths))[:3])}.",
            "category": "security",
            "priority": "high",
            "rule": {
                "name": "Add Admin Headers",
                "match_direction": "request",
                "action": "modify",
                "modify_headers": {
                    "X-Admin": "true",
                    "X-User-Role": "administrator"
                }
            },
            "natural_language": "Add admin headers to requests for admin paths"
        })

    if "delay" not in str(existing_rule_names):
        suggestions.append({
            "id": "latency-test",
            "title": "Test timeout handling",
            "description": "Introduce latency to see how the client handles slow responses.",
            "category": "performance",
            "priority": "low",
            "rule": {
                "name": "Add 3s Delay",
                "match_direction": "response",
                "action": "delay",
                "delay_ms": 3000
            },
            "natural_language": "Add a 3 second delay to all responses"
        })

    if settings.gemini_api_key and len(traffic_log) >= 5:
        try:
            from google import genai
            from google.genai import types

            client = genai.Client(api_key=settings.gemini_api_key)

            traffic_sample = []
            for entry in traffic_log[:15]:
                req = entry.get("request", {})
                resp = entry.get("response", {})
                traffic_sample.append({
                    "method": req.get("method"),
                    "path": req.get("path"),
                    "status": resp.get("status_code") if resp else None,
                })

            prompt = f"""As a security expert, analyze this HTTP traffic and suggest 2-3 specific security tests.

TRAFFIC SAMPLE:
{json.dumps(traffic_sample, indent=2)}

EXISTING RULES: {len(existing_rules)} rules already configured

Provide suggestions in JSON format:
{{
    "ai_suggestions": [
        {{
            "title": "Test name",
            "description": "What to test and why",
            "test_type": "auth|injection|access_control|information_disclosure|other",
            "risk_level": "high|medium|low"
        }}
    ],
    "traffic_insight": "Brief observation about the traffic patterns (1 sentence)"
}}

Return ONLY JSON."""

            response = client.models.generate_content(
                model=settings.gemini_model_id,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config={"thinking_level": "medium"},
                    max_output_tokens=500,
                )
            )

            if response and response.text:
                text = response.text.strip()
                if text.startswith("```"):
                    text = text.split("```")[1]
                    if text.startswith("json"):
                        text = text[4:]
                text = text.strip()

                ai_data = json.loads(text)

                for idx, ai_sug in enumerate(ai_data.get("ai_suggestions", [])):
                    suggestions.insert(0, {
                        "id": f"ai-{idx}",
                        "title": ai_sug.get("title", "AI suggestion"),
                        "description": ai_sug.get("description", ""),
                        "category": "security",
                        "priority": ai_sug.get("risk_level", "medium"),
                        "rule": None,
                        "natural_language": ai_sug.get("title", "AI suggestion")
                    })

        except Exception as e:
            logger.warning(f"AI suggestions failed: {e}")

    return {
        "suggestions": suggestions[:8],
        "traffic_summary": traffic_summary,
        "generated_at": generated_at
    }


# ============================================================================
# AI-Powered MITM Analysis Features
# ============================================================================

@dataclass
class SensitiveDataMatch:
    """Represents a detected sensitive data item."""
    data_type: str  # pii, credential, api_key, token, financial, health
    field_name: str
    value_preview: str  # First/last chars only for safety
    confidence: float  # 0-1
    location: str  # request_header, request_body, response_header, response_body, url
    entry_id: str
    risk_level: str  # critical, high, medium, low
    recommendation: str


@dataclass
class InjectionPoint:
    """Represents a potential injection point in a request."""
    parameter_name: str
    parameter_value: str
    location: str  # query, body, header, path, cookie
    injection_types: List[str]  # sqli, xss, cmdi, xxe, ssti, etc.
    confidence: float
    entry_id: str
    reasoning: str
    suggested_payloads: List[str]


@dataclass
class TestCase:
    """Auto-generated test case."""
    id: str
    name: str
    description: str
    target_entry_id: str
    attack_type: str
    payloads: List[Dict]  # {position, original, payload, expected_indicator}
    risk_level: str
    prerequisites: List[str]


class MITMIntelligenceAnalyzer:
    """
    AI-powered intelligence analyzer for MITM traffic.
    
    Features:
    - Sensitive data detection (PII, credentials, tokens)
    - Injection point highlighting
    - Natural language traffic queries
    - Auto-generated test cases
    - Finding description generation
    """
    
    # Regex patterns for common sensitive data
    SENSITIVE_PATTERNS = {
        # Credentials
        "password": (r'(?i)(password|passwd|pwd|secret)["\s:=]+["\']?([^"\'&\s]{4,})', "credential", "critical"),
        "api_key": (r'(?i)(api[_-]?key|apikey)["\s:=]+["\']?([a-zA-Z0-9_\-]{16,})', "api_key", "critical"),
        "bearer_token": (r'(?i)bearer\s+([a-zA-Z0-9_\-\.]+)', "token", "critical"),
        "jwt": (r'eyJ[a-zA-Z0-9_-]+\.eyJ[a-zA-Z0-9_-]+\.[a-zA-Z0-9_-]+', "token", "high"),
        "basic_auth": (r'(?i)basic\s+([a-zA-Z0-9+/=]{10,})', "credential", "critical"),
        "aws_key": (r'AKIA[0-9A-Z]{16}', "api_key", "critical"),
        "github_token": (r'gh[ps]_[a-zA-Z0-9]{36}', "api_key", "critical"),
        "slack_token": (r'xox[baprs]-[0-9a-zA-Z]{10,}', "api_key", "high"),
        "private_key": (r'-----BEGIN\s+(RSA\s+)?PRIVATE\s+KEY-----', "credential", "critical"),
        
        # PII
        "email": (r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b', "pii", "medium"),
        "ssn": (r'\b\d{3}-\d{2}-\d{4}\b', "pii", "critical"),
        "phone": (r'\b(?:\+1)?[\s.-]?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b', "pii", "medium"),
        "credit_card": (r'\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13})\b', "financial", "critical"),
        "ip_address": (r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b', "pii", "low"),
        
        # Health
        "health_id": (r'(?i)(patient[_-]?id|medical[_-]?record|mrn)["\s:=]+["\']?([a-zA-Z0-9]{6,})', "health", "high"),
    }
    
    # Injection type signatures
    INJECTION_SIGNATURES = {
        "sqli": {
            "indicators": ["id", "user", "name", "search", "query", "filter", "sort", "order", "limit", "offset", "where", "select"],
            "risk_patterns": [r"^\d+$", r"^[a-zA-Z0-9_]+$"],  # Numeric IDs, simple strings
            "payloads": ["' OR '1'='1", "1; DROP TABLE users--", "' UNION SELECT NULL--", "1' AND '1'='1"]
        },
        "xss": {
            "indicators": ["name", "search", "q", "query", "message", "comment", "title", "content", "text", "value", "callback", "redirect", "url", "return"],
            "risk_patterns": [r"<", r">", r"script"],
            "payloads": ["<script>alert(1)</script>", "<img src=x onerror=alert(1)>", "javascript:alert(1)", "'-alert(1)-'"]
        },
        "cmdi": {
            "indicators": ["cmd", "command", "exec", "run", "shell", "ping", "host", "ip", "file", "path", "dir", "filename"],
            "risk_patterns": [r"^[\w\.\-\/]+$"],
            "payloads": ["; ls -la", "| cat /etc/passwd", "`whoami`", "$(id)", "&& dir"]
        },
        "xxe": {
            "indicators": [],  # Detected by content-type
            "content_types": ["application/xml", "text/xml"],
            "payloads": ["<!DOCTYPE foo [<!ENTITY xxe SYSTEM 'file:///etc/passwd'>]>"]
        },
        "ssti": {
            "indicators": ["template", "render", "view", "page", "name", "title"],
            "risk_patterns": [],
            "payloads": ["{{7*7}}", "${7*7}", "<%= 7*7 %>", "#{7*7}", "*{7*7}"]
        },
        "path_traversal": {
            "indicators": ["file", "path", "document", "doc", "pdf", "image", "img", "download", "read", "load"],
            "risk_patterns": [r"[\w\.\-]+\.\w+"],  # Filename-like
            "payloads": ["../../../etc/passwd", "....//....//etc/passwd", "..%2f..%2f..%2fetc/passwd"]
        },
        "idor": {
            "indicators": ["id", "user_id", "account", "profile", "order", "invoice", "document"],
            "risk_patterns": [r"^\d+$", r"^[a-f0-9\-]{36}$"],  # Numeric or UUID
            "payloads": []  # Dynamic - increment/decrement IDs
        }
    }
    
    def __init__(self):
        self.analysis_cache: Dict[str, Any] = {}
        
    def detect_sensitive_data(self, traffic_entries: List[Dict]) -> List[SensitiveDataMatch]:
        """
        Scan traffic for sensitive data exposure.
        
        Returns list of detected sensitive data with locations and recommendations.
        """
        matches = []
        
        for entry in traffic_entries:
            entry_id = entry.get("id", "unknown")
            
            # Check request
            request = entry.get("request", {})
            if request:
                # URL/Query params
                url = request.get("path", "")
                matches.extend(self._scan_text(url, "url", entry_id))
                
                # Headers
                for header, value in request.get("headers", {}).items():
                    if header.lower() in ["authorization", "x-api-key", "cookie", "x-auth-token"]:
                        matches.extend(self._scan_text(f"{header}: {value}", "request_header", entry_id))
                
                # Body
                body = request.get("body", "")
                if body:
                    if isinstance(body, dict):
                        body = json.dumps(body)
                    matches.extend(self._scan_text(str(body), "request_body", entry_id))
            
            # Check response
            response = entry.get("response", {})
            if response:
                # Headers
                for header, value in response.get("headers", {}).items():
                    if header.lower() in ["set-cookie", "x-auth-token", "authorization"]:
                        matches.extend(self._scan_text(f"{header}: {value}", "response_header", entry_id))
                
                # Body
                body = response.get("body", "")
                if body:
                    if isinstance(body, dict):
                        body = json.dumps(body)
                    matches.extend(self._scan_text(str(body), "response_body", entry_id))
        
        # Deduplicate by value preview
        seen = set()
        unique_matches = []
        for m in matches:
            key = f"{m.data_type}:{m.value_preview}:{m.location}"
            if key not in seen:
                seen.add(key)
                unique_matches.append(m)
        
        return sorted(unique_matches, key=lambda x: {"critical": 0, "high": 1, "medium": 2, "low": 3}.get(x.risk_level, 4))
    
    def _scan_text(self, text: str, location: str, entry_id: str) -> List[SensitiveDataMatch]:
        """Scan text for sensitive data patterns."""
        matches = []
        
        for name, (pattern, data_type, risk_level) in self.SENSITIVE_PATTERNS.items():
            try:
                for match in re.finditer(pattern, text):
                    value = match.group(0)
                    # Create safe preview (first 4 + last 4 chars)
                    if len(value) > 12:
                        preview = f"{value[:4]}...{value[-4:]}"
                    else:
                        preview = value[:4] + "****"
                    
                    matches.append(SensitiveDataMatch(
                        data_type=data_type,
                        field_name=name,
                        value_preview=preview,
                        confidence=0.9 if name in ["jwt", "aws_key", "github_token"] else 0.7,
                        location=location,
                        entry_id=entry_id,
                        risk_level=risk_level,
                        recommendation=self._get_recommendation(data_type, location)
                    ))
            except Exception:
                pass
        
        return matches
    
    def _get_recommendation(self, data_type: str, location: str) -> str:
        """Get recommendation based on data type and location."""
        recommendations = {
            ("credential", "response_body"): "Credentials should never be exposed in responses. Implement proper authentication tokens instead.",
            ("credential", "request_body"): "Consider using secure authentication methods like OAuth2 instead of sending credentials in requests.",
            ("api_key", "response_body"): "API keys should not be exposed in responses. Return only necessary data.",
            ("api_key", "url"): "API keys in URLs are logged in server logs. Use Authorization headers instead.",
            ("token", "response_body"): "Ensure tokens have appropriate expiration and are transmitted over HTTPS only.",
            ("pii", "response_body"): "Minimize PII exposure. Apply data masking or field-level encryption.",
            ("financial", "response_body"): "Financial data requires PCI-DSS compliance. Mask card numbers, encrypt at rest.",
            ("health", "response_body"): "Health data requires HIPAA compliance. Implement audit logging and access controls.",
        }
        return recommendations.get((data_type, location), f"Review {data_type} exposure in {location}. Apply least privilege principle.")
    
    def find_injection_points(self, traffic_entries: List[Dict]) -> List[InjectionPoint]:
        """
        Analyze traffic to identify potential injection points.
        
        Returns list of parameters that may be vulnerable to injection attacks.
        """
        injection_points = []
        
        for entry in traffic_entries:
            entry_id = entry.get("id", "unknown")
            request = entry.get("request", {})
            
            if not request:
                continue
            
            method = request.get("method", "GET")
            path = request.get("path", "")
            content_type = request.get("headers", {}).get("content-type", "").lower()
            
            # Parse query parameters
            if "?" in path:
                query_string = path.split("?", 1)[1]
                params = self._parse_query_string(query_string)
                for name, value in params.items():
                    injection_types = self._detect_injection_types(name, value, "query", content_type)
                    if injection_types:
                        injection_points.append(InjectionPoint(
                            parameter_name=name,
                            parameter_value=value[:50] if len(value) > 50 else value,
                            location="query",
                            injection_types=list(injection_types.keys()),
                            confidence=max(injection_types.values()),
                            entry_id=entry_id,
                            reasoning=self._build_reasoning(name, injection_types),
                            suggested_payloads=self._get_payloads(injection_types)
                        ))
            
            # Parse body parameters
            body = request.get("body", "")
            if body and method in ["POST", "PUT", "PATCH"]:
                body_params = self._parse_body(body, content_type)
                for name, value in body_params.items():
                    injection_types = self._detect_injection_types(name, str(value), "body", content_type)
                    if injection_types:
                        injection_points.append(InjectionPoint(
                            parameter_name=name,
                            parameter_value=str(value)[:50] if len(str(value)) > 50 else str(value),
                            location="body",
                            injection_types=list(injection_types.keys()),
                            confidence=max(injection_types.values()),
                            entry_id=entry_id,
                            reasoning=self._build_reasoning(name, injection_types),
                            suggested_payloads=self._get_payloads(injection_types)
                        ))
            
            # Check for XXE in XML content
            if "xml" in content_type:
                injection_points.append(InjectionPoint(
                    parameter_name="XML Body",
                    parameter_value="<xml content>",
                    location="body",
                    injection_types=["xxe"],
                    confidence=0.8,
                    entry_id=entry_id,
                    reasoning="XML content type detected - test for XML External Entity injection",
                    suggested_payloads=self.INJECTION_SIGNATURES["xxe"]["payloads"]
                ))
            
            # Check path parameters (e.g., /users/123)
            path_parts = path.split("?")[0].split("/")
            for i, part in enumerate(path_parts):
                if re.match(r"^\d+$", part) or re.match(r"^[a-f0-9\-]{36}$", part):
                    injection_points.append(InjectionPoint(
                        parameter_name=f"path_segment_{i}",
                        parameter_value=part,
                        location="path",
                        injection_types=["idor"],
                        confidence=0.75,
                        entry_id=entry_id,
                        reasoning=f"Numeric/UUID path segment detected at position {i} - possible IDOR",
                        suggested_payloads=[str(int(part) + 1) if part.isdigit() else part]
                    ))
        
        return sorted(injection_points, key=lambda x: x.confidence, reverse=True)
    
    def _parse_query_string(self, qs: str) -> Dict[str, str]:
        """Parse query string into dict."""
        params = {}
        for pair in qs.split("&"):
            if "=" in pair:
                key, value = pair.split("=", 1)
                params[key] = value
        return params
    
    def _parse_body(self, body: Any, content_type: str) -> Dict[str, Any]:
        """Parse request body into dict."""
        if isinstance(body, dict):
            return body
        
        try:
            if "json" in content_type:
                return json.loads(body) if isinstance(body, str) else body
            elif "x-www-form-urlencoded" in content_type:
                return self._parse_query_string(body)
        except (json.JSONDecodeError, ValueError, TypeError):
            pass
        return {}
    
    def _detect_injection_types(self, name: str, value: str, location: str, content_type: str) -> Dict[str, float]:
        """Detect possible injection types for a parameter."""
        results = {}
        name_lower = name.lower()
        
        for inj_type, config in self.INJECTION_SIGNATURES.items():
            confidence = 0.0
            
            # Check name indicators
            for indicator in config.get("indicators", []):
                if indicator in name_lower:
                    confidence = max(confidence, 0.7)
                    break
            
            # Check value patterns
            for pattern in config.get("risk_patterns", []):
                if re.match(pattern, value):
                    confidence = max(confidence, 0.6)
            
            # Check content type for XXE
            if inj_type == "xxe" and any(ct in content_type for ct in config.get("content_types", [])):
                confidence = 0.8
            
            if confidence >= 0.5:
                results[inj_type] = confidence
        
        return results
    
    def _build_reasoning(self, name: str, injection_types: Dict[str, float]) -> str:
        """Build human-readable reasoning for injection detection."""
        reasons = []
        for inj_type, confidence in injection_types.items():
            if inj_type == "sqli":
                reasons.append(f"Parameter '{name}' may be used in database queries")
            elif inj_type == "xss":
                reasons.append(f"Parameter '{name}' may be reflected in responses")
            elif inj_type == "cmdi":
                reasons.append(f"Parameter '{name}' may be used in system commands")
            elif inj_type == "idor":
                reasons.append(f"Parameter '{name}' appears to be an object reference")
            elif inj_type == "ssti":
                reasons.append(f"Parameter '{name}' may be used in template rendering")
        return "; ".join(reasons) if reasons else "Multiple injection vectors possible"
    
    def _get_payloads(self, injection_types: Dict[str, float]) -> List[str]:
        """Get suggested payloads for detected injection types."""
        payloads = []
        for inj_type in injection_types:
            payloads.extend(self.INJECTION_SIGNATURES.get(inj_type, {}).get("payloads", [])[:3])
        return payloads[:5]
    
    async def natural_language_query(self, traffic_entries: List[Dict], query: str) -> Dict:
        """
        Query traffic using natural language.
        
        Examples:
        - "Find all authentication requests"
        - "Show requests with user IDs"
        - "Find error responses"
        - "Show requests to admin endpoints"
        """
        from ..core.config import settings
        
        # First, try pattern-based matching for common queries
        results = self._pattern_based_query(traffic_entries, query)
        
        # If pattern matching found results, return them
        if results["matches"]:
            return results
        
        # Fall back to AI for complex queries
        if settings.gemini_api_key:
            try:
                from google import genai
                from google.genai import types
                
                client = genai.Client(api_key=settings.gemini_api_key)
                
                # Prepare traffic summary for AI
                traffic_summary = []
                for entry in traffic_entries[:50]:  # Limit for context
                    req = entry.get("request", {})
                    resp = entry.get("response", {})
                    traffic_summary.append({
                        "id": entry.get("id"),
                        "method": req.get("method"),
                        "path": req.get("path"),
                        "status": resp.get("status_code") if resp else None,
                        "content_type": req.get("headers", {}).get("content-type", ""),
                    })
                
                prompt = f"""Analyze this HTTP traffic and find entries matching the user's query.

USER QUERY: "{query}"

TRAFFIC (showing first 50 entries):
{json.dumps(traffic_summary, indent=2)}

Return a JSON response:
{{
    "matching_ids": ["id1", "id2"],  // IDs of matching entries
    "query_interpretation": "What I understood from the query",
    "filter_criteria": "The criteria used to filter"
}}

Return ONLY valid JSON."""

                response = client.models.generate_content(
                    model=settings.gemini_model_id,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        thinking_config={"thinking_level": "medium"},
                        max_output_tokens=500,
                    )
                )
                
                if response and response.text:
                    text = response.text.strip()
                    if text.startswith("```"):
                        text = text.split("```")[1]
                        if text.startswith("json"):
                            text = text[4:]
                    text = text.strip()
                    
                    ai_result = json.loads(text)
                    matching_ids = ai_result.get("matching_ids", [])
                    
                    matches = [e for e in traffic_entries if e.get("id") in matching_ids]
                    
                    return {
                        "query": query,
                        "interpretation": ai_result.get("query_interpretation", ""),
                        "filter_criteria": ai_result.get("filter_criteria", ""),
                        "matches": matches,
                        "total_matches": len(matches),
                        "ai_powered": True
                    }
            except Exception as e:
                logger.warning(f"AI query failed: {e}")
        
        return {
            "query": query,
            "interpretation": "Could not interpret query",
            "filter_criteria": "",
            "matches": [],
            "total_matches": 0,
            "ai_powered": False,
            "error": "No matches found for this query"
        }
    
    def _pattern_based_query(self, traffic_entries: List[Dict], query: str) -> Dict:
        """Pattern-based query matching for common queries."""
        query_lower = query.lower()
        matches = []
        interpretation = ""
        filter_criteria = ""
        
        # Authentication requests
        if any(word in query_lower for word in ["auth", "login", "signin", "sign in", "authenticate"]):
            interpretation = "Finding authentication-related requests"
            filter_criteria = "path contains /auth, /login, /signin, /token OR has Authorization header"
            for entry in traffic_entries:
                req = entry.get("request", {})
                path = req.get("path", "").lower()
                headers = req.get("headers", {})
                if any(p in path for p in ["/auth", "/login", "/signin", "/token", "/oauth"]) or \
                   "authorization" in [h.lower() for h in headers.keys()]:
                    matches.append(entry)
        
        # Error responses
        elif any(word in query_lower for word in ["error", "fail", "4xx", "5xx", "500", "404"]):
            interpretation = "Finding error responses"
            filter_criteria = "status code >= 400"
            for entry in traffic_entries:
                resp = entry.get("response", {})
                status = resp.get("status_code", 0)
                if status >= 400:
                    matches.append(entry)
        
        # Admin endpoints
        elif any(word in query_lower for word in ["admin", "management", "dashboard", "settings"]):
            interpretation = "Finding admin/management endpoints"
            filter_criteria = "path contains /admin, /manage, /dashboard, /settings"
            for entry in traffic_entries:
                req = entry.get("request", {})
                path = req.get("path", "").lower()
                if any(p in path for p in ["/admin", "/manage", "/dashboard", "/settings", "/config"]):
                    matches.append(entry)
        
        # User ID/account requests
        elif any(word in query_lower for word in ["user", "account", "profile", "id"]):
            interpretation = "Finding user/account related requests"
            filter_criteria = "path contains /user, /account, /profile OR has user-related parameters"
            for entry in traffic_entries:
                req = entry.get("request", {})
                path = req.get("path", "").lower()
                if any(p in path for p in ["/user", "/account", "/profile", "/member"]) or \
                   re.search(r"(user|account|profile)[_-]?id", path):
                    matches.append(entry)
        
        # POST requests
        elif "post" in query_lower:
            interpretation = "Finding POST requests"
            filter_criteria = "method = POST"
            for entry in traffic_entries:
                if entry.get("request", {}).get("method") == "POST":
                    matches.append(entry)
        
        # API endpoints
        elif "api" in query_lower:
            interpretation = "Finding API endpoints"
            filter_criteria = "path contains /api"
            for entry in traffic_entries:
                if "/api" in entry.get("request", {}).get("path", "").lower():
                    matches.append(entry)
        
        # JSON responses
        elif "json" in query_lower:
            interpretation = "Finding JSON responses"
            filter_criteria = "content-type contains application/json"
            for entry in traffic_entries:
                resp = entry.get("response", {})
                ct = resp.get("headers", {}).get("content-type", "")
                if "json" in ct.lower():
                    matches.append(entry)
        
        return {
            "query": query,
            "interpretation": interpretation,
            "filter_criteria": filter_criteria,
            "matches": matches,
            "total_matches": len(matches),
            "ai_powered": False
        }
    
    async def generate_test_cases(self, traffic_entries: List[Dict], entry_id: Optional[str] = None) -> List[TestCase]:
        """
        Auto-generate security test cases based on traffic analysis.
        
        If entry_id is provided, generate tests for that specific entry.
        Otherwise, analyze all traffic and prioritize.
        """
        test_cases = []
        
        # Filter to specific entry if provided
        entries_to_analyze = traffic_entries
        if entry_id:
            entries_to_analyze = [e for e in traffic_entries if e.get("id") == entry_id]
        
        # Find injection points first
        injection_points = self.find_injection_points(entries_to_analyze)
        
        # Generate test cases from injection points
        for i, ip in enumerate(injection_points[:20]):  # Limit to top 20
            for inj_type in ip.injection_types:
                test_id = f"tc_{inj_type}_{i}"
                payloads = []
                
                for payload in ip.suggested_payloads[:3]:
                    payloads.append({
                        "position": ip.location,
                        "parameter": ip.parameter_name,
                        "original": ip.parameter_value,
                        "payload": payload,
                        "expected_indicator": self._get_success_indicator(inj_type)
                    })
                
                test_cases.append(TestCase(
                    id=test_id,
                    name=f"{inj_type.upper()} test on {ip.parameter_name}",
                    description=f"Test {ip.parameter_name} in {ip.location} for {inj_type} vulnerability. {ip.reasoning}",
                    target_entry_id=ip.entry_id,
                    attack_type=inj_type,
                    payloads=payloads,
                    risk_level="high" if inj_type in ["sqli", "cmdi", "xxe"] else "medium",
                    prerequisites=["Valid session/auth if required"]
                ))
        
        # Add authentication bypass tests
        auth_entries = [e for e in entries_to_analyze if 
                       "authorization" in [h.lower() for h in e.get("request", {}).get("headers", {}).keys()]]
        
        if auth_entries:
            test_cases.append(TestCase(
                id="tc_auth_bypass_1",
                name="Authorization Header Removal",
                description="Test if endpoints are accessible without authentication by removing the Authorization header",
                target_entry_id=auth_entries[0].get("id"),
                attack_type="auth_bypass",
                payloads=[{"position": "header", "parameter": "Authorization", "original": "Bearer xxx", "payload": "", "expected_indicator": "200 OK or data returned"}],
                risk_level="high",
                prerequisites=[]
            ))
        
        return test_cases
    
    def _get_success_indicator(self, inj_type: str) -> str:
        """Get success indicator for injection type."""
        indicators = {
            "sqli": "SQL error message, different response, time delay",
            "xss": "Payload reflected unencoded in response",
            "cmdi": "Command output in response, time delay",
            "xxe": "File contents in response, DNS callback",
            "ssti": "Mathematical result (49) or error message",
            "path_traversal": "File contents in response",
            "idor": "Access to other user's data"
        }
        return indicators.get(inj_type, "Unexpected response or behavior")
    
    async def generate_finding_description(
        self,
        vulnerability_type: str,
        affected_endpoint: str,
        parameter: str,
        evidence: str,
        severity: str = "medium"
    ) -> Dict:
        """
        Generate a professional vulnerability finding description.
        
        Returns structured finding with description, impact, remediation.
        """
        from ..core.config import settings
        
        # Template-based generation for common vulns
        templates = {
            "sqli": {
                "title": f"SQL Injection in {parameter}",
                "description": f"A SQL injection vulnerability was identified in the '{parameter}' parameter of the {affected_endpoint} endpoint. The application appears to construct SQL queries using unsanitized user input, allowing an attacker to manipulate database queries.",
                "impact": "An attacker could extract sensitive data, modify or delete database records, bypass authentication, or potentially execute system commands depending on database configuration.",
                "remediation": "1. Use parameterized queries or prepared statements\n2. Implement input validation using allowlists\n3. Apply least privilege to database accounts\n4. Enable WAF rules for SQL injection",
                "references": ["OWASP SQL Injection", "CWE-89"]
            },
            "xss": {
                "title": f"Cross-Site Scripting (XSS) in {parameter}",
                "description": f"A cross-site scripting vulnerability was identified in the '{parameter}' parameter of the {affected_endpoint} endpoint. User-supplied input is reflected in the response without proper encoding.",
                "impact": "An attacker could execute arbitrary JavaScript in victims' browsers, steal session tokens, perform actions on behalf of users, or redirect users to malicious sites.",
                "remediation": "1. Encode all user output (HTML entity encoding)\n2. Implement Content-Security-Policy headers\n3. Use HTTPOnly and Secure flags on cookies\n4. Validate and sanitize input on the server side",
                "references": ["OWASP XSS", "CWE-79"]
            },
            "idor": {
                "title": f"Insecure Direct Object Reference in {parameter}",
                "description": f"An IDOR vulnerability was identified in the '{parameter}' parameter of the {affected_endpoint} endpoint. The application uses predictable identifiers without proper authorization checks.",
                "impact": "An attacker could access or modify other users' data, including sensitive personal information, financial records, or private documents.",
                "remediation": "1. Implement proper authorization checks for all object access\n2. Use indirect reference maps or UUIDs\n3. Verify user ownership before returning data\n4. Log and monitor access patterns",
                "references": ["OWASP IDOR", "CWE-639"]
            }
        }
        
        # Get template or use AI
        template = templates.get(vulnerability_type.lower().replace(" ", "_").replace("-", "_"))
        
        if template:
            return {
                "title": template["title"],
                "severity": severity,
                "description": template["description"],
                "impact": template["impact"],
                "remediation": template["remediation"],
                "evidence": evidence,
                "affected_endpoint": affected_endpoint,
                "affected_parameter": parameter,
                "references": template["references"],
                "cvss_estimate": {"high": "7.5-8.9", "medium": "4.0-6.9", "low": "1.0-3.9"}.get(severity, "4.0-6.9")
            }
        
        # Use AI for unknown vuln types
        if settings.gemini_api_key:
            try:
                from google import genai
                from google.genai import types
                
                client = genai.Client(api_key=settings.gemini_api_key)
                
                prompt = f"""Generate a professional security vulnerability finding description.

Vulnerability Type: {vulnerability_type}
Affected Endpoint: {affected_endpoint}
Affected Parameter: {parameter}
Severity: {severity}
Evidence: {evidence}

Return JSON:
{{
    "title": "Clear vulnerability title",
    "description": "Technical description of the vulnerability",
    "impact": "Business and security impact",
    "remediation": "Step-by-step fix recommendations",
    "references": ["Relevant standards/resources"]
}}

Return ONLY valid JSON."""

                response = client.models.generate_content(
                    model=settings.gemini_model_id,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        thinking_config={"thinking_level": "medium"},
                        max_output_tokens=800,
                    )
                )
                
                if response and response.text:
                    text = response.text.strip()
                    if text.startswith("```"):
                        text = text.split("```")[1]
                        if text.startswith("json"):
                            text = text[4:]
                    text = text.strip()
                    
                    ai_finding = json.loads(text)
                    ai_finding["severity"] = severity
                    ai_finding["evidence"] = evidence
                    ai_finding["affected_endpoint"] = affected_endpoint
                    ai_finding["affected_parameter"] = parameter
                    ai_finding["ai_generated"] = True
                    
                    return ai_finding
            except Exception as e:
                logger.warning(f"AI finding generation failed: {e}")
        
        # Fallback generic template
        return {
            "title": f"{vulnerability_type} Vulnerability",
            "severity": severity,
            "description": f"A {vulnerability_type} vulnerability was identified in the {parameter} parameter of {affected_endpoint}.",
            "impact": "The impact depends on the specific vulnerability type and application context.",
            "remediation": "Review the affected endpoint and implement appropriate security controls.",
            "evidence": evidence,
            "affected_endpoint": affected_endpoint,
            "affected_parameter": parameter,
            "references": ["OWASP Testing Guide"]
        }


# Global instance
mitm_intelligence = MITMIntelligenceAnalyzer()


# Convenience functions
async def analyze_traffic_sensitive_data(traffic_entries: List[Dict]) -> List[Dict]:
    """Analyze traffic for sensitive data exposure."""
    try:
        if not traffic_entries:
            return []
        matches = mitm_intelligence.detect_sensitive_data(traffic_entries)
        return [
            {
                "data_type": getattr(m, 'data_type', 'unknown'),
                "field_name": getattr(m, 'field_name', 'unknown'),
                "value_preview": getattr(m, 'value_preview', '****'),
                "confidence": getattr(m, 'confidence', 0.5),
                "location": getattr(m, 'location', 'unknown'),
                "entry_id": getattr(m, 'entry_id', 'unknown'),
                "risk_level": getattr(m, 'risk_level', 'medium'),
                "recommendation": getattr(m, 'recommendation', 'Review this data exposure.')
            }
        for m in matches
    ]
    except Exception as e:
        logger.warning(f"Sensitive data analysis failed: {e}")
        return []


async def analyze_traffic_injection_points(traffic_entries: List[Dict]) -> List[Dict]:
    """Analyze traffic for injection points."""
    try:
        if not traffic_entries:
            return []
        points = mitm_intelligence.find_injection_points(traffic_entries)
        return [
            {
                "parameter_name": getattr(p, 'parameter_name', 'unknown'),
                "parameter_value": getattr(p, 'parameter_value', ''),
                "location": getattr(p, 'location', 'unknown'),
                "injection_types": getattr(p, 'injection_types', []),
                "confidence": getattr(p, 'confidence', 0.5),
                "entry_id": getattr(p, 'entry_id', 'unknown'),
                "reasoning": getattr(p, 'reasoning', ''),
                "suggested_payloads": getattr(p, 'suggested_payloads', [])
            }
            for p in points
        ]
    except Exception as e:
        logger.warning(f"Injection point analysis failed: {e}")
        return []


async def query_traffic_natural_language(traffic_entries: List[Dict], query: str) -> Dict:
    """Query traffic using natural language."""
    try:
        if not query or not query.strip():
            return {
                "query": query or "",
                "interpretation": "Empty query provided",
                "filter_criteria": "",
                "matches": [],
                "total_matches": 0,
                "ai_powered": False
            }
        return await mitm_intelligence.natural_language_query(traffic_entries or [], query)
    except Exception as e:
        logger.warning(f"Natural language query failed: {e}")
        return {
            "query": query or "",
            "interpretation": "Query processing failed",
            "filter_criteria": "",
            "matches": [],
            "total_matches": 0,
            "ai_powered": False,
            "error": str(e)
        }


async def generate_security_test_cases(traffic_entries: List[Dict], entry_id: Optional[str] = None) -> List[Dict]:
    """Generate security test cases from traffic."""
    try:
        if not traffic_entries:
            return []
        test_cases = await mitm_intelligence.generate_test_cases(traffic_entries, entry_id)
        return [
            {
                "id": getattr(tc, 'id', f'tc_{i}'),
                "name": getattr(tc, 'name', 'Unknown Test'),
                "description": getattr(tc, 'description', ''),
                "target_entry_id": getattr(tc, 'target_entry_id', 'unknown'),
                "attack_type": getattr(tc, 'attack_type', 'unknown'),
                "payloads": getattr(tc, 'payloads', []),
                "risk_level": getattr(tc, 'risk_level', 'medium'),
                "prerequisites": getattr(tc, 'prerequisites', [])
            }
            for i, tc in enumerate(test_cases)
        ]
    except Exception as e:
        logger.warning(f"Test case generation failed: {e}")
        return []


async def generate_vulnerability_finding(
    vulnerability_type: str,
    affected_endpoint: str,
    parameter: str,
    evidence: str,
    severity: str = "medium"
) -> Dict:
    """Generate a professional vulnerability finding description."""
    try:
        # Validate inputs
        if not vulnerability_type:
            vulnerability_type = "Unknown"
        if not affected_endpoint:
            affected_endpoint = "Unknown endpoint"
        if not parameter:
            parameter = "Unknown parameter"
        if severity not in ["critical", "high", "medium", "low", "info"]:
            severity = "medium"
            
        return await mitm_intelligence.generate_finding_description(
            vulnerability_type, affected_endpoint, parameter, evidence or "", severity
        )
    except Exception as e:
        logger.warning(f"Finding generation failed: {e}")
        return {
            "title": f"{vulnerability_type or 'Unknown'} Vulnerability",
            "severity": severity or "medium",
            "description": f"Error generating finding description: {str(e)}",
            "impact": "Unable to determine impact.",
            "remediation": "Review the affected endpoint manually.",
            "evidence": evidence or "",
            "affected_endpoint": affected_endpoint or "",
            "affected_parameter": parameter or "",
            "references": []
        }
