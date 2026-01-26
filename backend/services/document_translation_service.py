"""
Document Translation Service

Translates documents using OCR (tesseract) and Gemini for literal translation.
Focuses on line-by-line preservation to reduce omissions and hallucinations.

Features:
- Parallel chunk translation (3-5x faster)
- Translation memory/cache (cost savings)
- Auto language detection
- Progress streaming callbacks
"""

import asyncio
import base64
import hashlib
import json
import logging
import os
import random
import re
import sqlite3
import statistics
import time
from collections import OrderedDict
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from backend.services.document_ai_service import document_ai_service
from backend.services.document_parser_service import DocumentParserService
from backend.services.tesseract_utils import configure_pytesseract

# Progress callback type: (current_step, total_steps, message)
ProgressCallback = Callable[[int, int, str], None]

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pytesseract
    from PIL import Image
    from pytesseract import Output
    _tesseract_cmd = configure_pytesseract()
    HAS_OCR = bool(_tesseract_cmd)
except ImportError:
    HAS_OCR = False
    pytesseract = None
    Image = None
    Output = None

logger = logging.getLogger(__name__)


class DocumentTranslationService:
    """Service for extracting and translating documents with strict coverage."""

    SUPPORTED_MIME_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        'application/vnd.openxmlformats-officedocument.presentationml.presentation': 'pptx',
        'application/vnd.ms-powerpoint': 'ppt',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'application/vnd.ms-excel': 'xls',
        'text/plain': 'txt',
        'text/markdown': 'md',
        'text/csv': 'csv',
        'application/json': 'json',
        'application/xml': 'xml',
        'text/xml': 'xml',
        'image/png': 'png',
        'image/jpeg': 'jpg',
        'image/jpg': 'jpg',
        'image/tiff': 'tiff',
        'image/bmp': 'bmp',
        'image/gif': 'gif',
    }

    # RTL languages require special handling
    RTL_LANGUAGES = {
        'arabic', 'persian', 'farsi', 'hebrew', 'urdu', 'pashto', 'sindhi',
        'ar', 'fa', 'he', 'ur', 'ps', 'sd'
    }

    # Font selection based on target language/script
    # Using Noto fonts for broad Unicode coverage
    # Font flags: bold=16 (bit 4), italic=2 (bit 1)
    FONT_MAP = {
        'cjk': '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        'cjk_bold': '/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc',
        'arabic': '/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf',
        'arabic_bold': '/usr/share/fonts/truetype/noto/NotoSansArabic-Bold.ttf',
        'hebrew': '/usr/share/fonts/truetype/noto/NotoSansHebrew-Regular.ttf',
        'hebrew_bold': '/usr/share/fonts/truetype/noto/NotoSansHebrew-Bold.ttf',
        'devanagari': '/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf',
        'devanagari_bold': '/usr/share/fonts/truetype/noto/NotoSansDevanagari-Bold.ttf',
        'thai': '/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf',
        'thai_bold': '/usr/share/fonts/truetype/noto/NotoSansThai-Bold.ttf',
        'default': '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        'default_bold': '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
        'default_italic': '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf',
        'default_bolditalic': '/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf',
    }

    # Language to script mapping
    LANGUAGE_SCRIPT_MAP = {
        'chinese': 'cjk', 'japanese': 'cjk', 'korean': 'cjk',
        'zh': 'cjk', 'ja': 'cjk', 'ko': 'cjk',
        'arabic': 'arabic', 'ar': 'arabic',
        'persian': 'arabic', 'farsi': 'arabic', 'fa': 'arabic',
        'urdu': 'arabic', 'ur': 'arabic',
        'hebrew': 'hebrew', 'he': 'hebrew',
        'hindi': 'devanagari', 'hi': 'devanagari',
        'thai': 'thai', 'th': 'thai',
    }

    MAX_LINES_PER_CHUNK = 40
    MAX_CHARS_PER_CHUNK = 2600
    MAX_OUTPUT_TOKENS = 2400
    MAX_RETRIES = 2
    PDF_MAX_FONT = 12
    PDF_MIN_FONT = 6
    PDF_LINE_HEIGHT = 1.25

    # Parallel processing settings
    MAX_PARALLEL_CHUNKS = 5  # Concurrent API calls limit

    # Translation cache settings
    CACHE_MAX_SIZE = 10000   # Max cached translations
    CACHE_TTL = 3600         # 1 hour TTL

    # Context overlap for better chunk boundary translations
    CONTEXT_OVERLAP_LINES = 3

    # Validation and hardening settings
    MAX_LENGTH_RATIO = 4.0      # Translation can't be >4x original length
    MIN_LENGTH_RATIO = 0.2      # Translation can't be <0.2x original length
    MAX_REPETITION_RATIO = 0.5  # Max 50% of output can be repeated phrases
    RETRY_BASE_DELAY = 1.0      # Base delay for exponential backoff
    RETRY_MAX_DELAY = 30.0      # Max retry delay

    # Persistent translation memory settings
    PERSISTENT_CACHE_DIR = "/app/storage/translation_cache"
    PERSISTENT_CACHE_DB = "translation_memory.db"

    # Multi-column detection settings
    COLUMN_GAP_THRESHOLD = 50     # Min horizontal gap between columns (points)
    COLUMN_CLUSTER_TOLERANCE = 30  # X-position tolerance for column clustering

    # Quality scoring settings
    MIN_BLEU_SCORE = 0.1  # Minimum acceptable BLEU for back-translation

    # Multi-page table stitching settings
    TABLE_PAGE_MARGIN = 40        # Distance from page edge for table continuation detection
    TABLE_COLUMN_TOLERANCE = 10   # X-position tolerance for matching table columns
    MIN_TABLE_ROWS = 2            # Minimum rows to consider for table stitching

    # Font extraction settings
    EXTRACTED_FONTS_DIR = "/app/storage/extracted_fonts"
    FONT_CACHE_ENABLED = True     # Enable/disable font extraction and caching

    # Hallucination prevention settings
    MAX_SENTENCE_COUNT_DRIFT = 0.3   # Allow 30% drift in sentence count
    ENTITY_PRESERVATION_THRESHOLD = 0.8  # 80% of entities must be preserved
    # Regex patterns for entity extraction
    ENTITY_PATTERNS = {
        'numbers': r'\b\d+(?:[.,]\d+)*\b',
        'dates': r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b',
        'times': r'\b\d{1,2}:\d{2}(?::\d{2})?(?:\s*[APap][Mm])?\b',
        'emails': r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
        'urls': r'https?://[^\s<>"]+|www\.[^\s<>"]+',
        'percentages': r'\b\d+(?:[.,]\d+)?\s*%',
        'currencies': r'[$€£¥₹]\s*\d+(?:[.,]\d+)*|\d+(?:[.,]\d+)*\s*[$€£¥₹]',
        'codes': r'\b[A-Z]{2,}[-_]?\d+\b|\b\d+[-_][A-Z]+\b',
    }

    def __init__(self):
        self.parser = DocumentParserService()
        self.client = document_ai_service.client
        self.model_name = document_ai_service.model_name
        self._has_pymupdf = HAS_PYMUPDF
        self._has_ocr = HAS_OCR
        # Translation memory cache (LRU) - in-memory
        self._translation_cache: OrderedDict = OrderedDict()
        self._cache_timestamps: Dict[str, float] = {}
        self._cache_hits = 0
        self._cache_misses = 0
        # Persistent translation memory (SQLite)
        self._init_persistent_cache()
        # Glossary for consistent terminology
        self._glossary: Dict[str, str] = {}
        # Checkpoint for resume from failure
        self._checkpoint_data: Dict[str, Any] = {}
        # Context overlap tracking
        self._last_chunk_context: List[Tuple[int, str]] = []

    def _init_persistent_cache(self) -> None:
        """Initialize SQLite-backed persistent translation memory."""
        try:
            os.makedirs(self.PERSISTENT_CACHE_DIR, exist_ok=True)
            db_path = os.path.join(self.PERSISTENT_CACHE_DIR, self.PERSISTENT_CACHE_DB)
            self._persist_db_path = db_path
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS translation_memory (
                    cache_key TEXT PRIMARY KEY,
                    source_text TEXT NOT NULL,
                    target_lang TEXT NOT NULL,
                    source_lang TEXT,
                    translated_text TEXT NOT NULL,
                    created_at REAL NOT NULL,
                    used_count INTEGER DEFAULT 1,
                    quality_score REAL
                )
            ''')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_tm_target_lang ON translation_memory(target_lang)')
            cursor.execute('CREATE INDEX IF NOT EXISTS idx_tm_created_at ON translation_memory(created_at)')
            conn.commit()
            conn.close()
            logger.info(f"Persistent translation memory initialized at {db_path}")
        except Exception as e:
            logger.warning(f"Failed to initialize persistent cache: {e}")
            self._persist_db_path = None

    def is_supported(self, mime_type: str) -> bool:
        return mime_type in self.SUPPORTED_MIME_TYPES

    # ===== Translation Cache Methods =====

    def _get_cache_key(self, text: str, target_lang: str, source_lang: Optional[str] = None) -> str:
        """Generate MD5 cache key from text and languages."""
        content = f"{text.strip()}|{target_lang.lower()}|{(source_lang or '').lower()}"
        return hashlib.md5(content.encode()).hexdigest()

    def _cache_get(self, key: str) -> Optional[str]:
        """Get from cache with TTL check (in-memory then persistent)."""
        # First check in-memory cache
        if key in self._translation_cache:
            timestamp = self._cache_timestamps.get(key, 0)
            if time.time() - timestamp <= self.CACHE_TTL:
                self._translation_cache.move_to_end(key)
                self._cache_hits += 1
                return self._translation_cache[key]
            else:
                del self._translation_cache[key]
                del self._cache_timestamps[key]

        # Check persistent cache
        result = self._persist_get(key)
        if result:
            # Populate in-memory cache
            self._cache_set(key, result, persist=False)
            self._cache_hits += 1
            self._persistent_cache_hits = getattr(self, '_persistent_cache_hits', 0) + 1
            return result

        self._cache_misses += 1
        return None

    def _cache_set(self, key: str, value: str, persist: bool = True, 
                   source_text: str = "", target_lang: str = "", source_lang: str = "") -> None:
        """Set cache with LRU eviction and optional persistence."""
        # In-memory cache
        if key in self._translation_cache:
            self._translation_cache.move_to_end(key)
        else:
            if len(self._translation_cache) >= self.CACHE_MAX_SIZE:
                self._translation_cache.popitem(last=False)
        self._translation_cache[key] = value
        self._cache_timestamps[key] = time.time()

        # Persistent cache
        if persist and self._persist_db_path:
            self._persist_set(key, source_text, target_lang, source_lang, value)

    def _persist_get(self, key: str) -> Optional[str]:
        """Get translation from persistent SQLite cache."""
        if not self._persist_db_path:
            return None
        try:
            conn = sqlite3.connect(self._persist_db_path)
            cursor = conn.cursor()
            cursor.execute(
                'SELECT translated_text FROM translation_memory WHERE cache_key = ?',
                (key,)
            )
            row = cursor.fetchone()
            if row:
                # Update usage count
                cursor.execute(
                    'UPDATE translation_memory SET used_count = used_count + 1 WHERE cache_key = ?',
                    (key,)
                )
                conn.commit()
            conn.close()
            return row[0] if row else None
        except Exception as e:
            logger.warning(f"Persistent cache read error: {e}")
            return None

    def _persist_set(self, key: str, source_text: str, target_lang: str, 
                     source_lang: str, translated_text: str, quality_score: float = None) -> None:
        """Store translation in persistent SQLite cache."""
        if not self._persist_db_path:
            return
        try:
            conn = sqlite3.connect(self._persist_db_path)
            cursor = conn.cursor()
            cursor.execute('''
                INSERT OR REPLACE INTO translation_memory 
                (cache_key, source_text, target_lang, source_lang, translated_text, created_at, used_count, quality_score)
                VALUES (?, ?, ?, ?, ?, ?, 
                    COALESCE((SELECT used_count FROM translation_memory WHERE cache_key = ?), 0) + 1,
                    ?)
            ''', (key, source_text, target_lang, source_lang or '', translated_text, time.time(), key, quality_score))
            conn.commit()
            conn.close()
        except Exception as e:
            logger.warning(f"Persistent cache write error: {e}")

    def get_persistent_cache_stats(self) -> Dict[str, Any]:
        """Return persistent translation memory statistics."""
        if not self._persist_db_path:
            return {"enabled": False}
        try:
            conn = sqlite3.connect(self._persist_db_path)
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) FROM translation_memory')
            total = cursor.fetchone()[0]
            cursor.execute('SELECT SUM(used_count) FROM translation_memory')
            total_uses = cursor.fetchone()[0] or 0
            cursor.execute('SELECT target_lang, COUNT(*) FROM translation_memory GROUP BY target_lang')
            by_lang = dict(cursor.fetchall())
            conn.close()
            return {
                "enabled": True,
                "total_entries": total,
                "total_uses": total_uses,
                "by_language": by_lang
            }
        except Exception as e:
            return {"enabled": True, "error": str(e)}

    def get_cache_stats(self) -> Dict[str, Any]:
        """Return cache statistics."""
        total = self._cache_hits + self._cache_misses
        return {
            "hits": self._cache_hits,
            "misses": self._cache_misses,
            "hit_rate": self._cache_hits / total if total > 0 else 0,
            "size": len(self._translation_cache),
            "max_size": self.CACHE_MAX_SIZE,
        }

    def clear_cache(self) -> None:
        """Clear translation cache."""
        self._translation_cache.clear()
        self._cache_timestamps.clear()
        self._cache_hits = 0
        self._cache_misses = 0

    def clear_persistent_cache(self, older_than_days: int = None) -> int:
        """Clear persistent translation memory, optionally only old entries."""
        if not self._persist_db_path:
            return 0
        try:
            conn = sqlite3.connect(self._persist_db_path)
            cursor = conn.cursor()
            if older_than_days:
                cutoff = time.time() - (older_than_days * 86400)
                cursor.execute('DELETE FROM translation_memory WHERE created_at < ?', (cutoff,))
            else:
                cursor.execute('DELETE FROM translation_memory')
            deleted = cursor.rowcount
            conn.commit()
            conn.close()
            return deleted
        except Exception as e:
            logger.warning(f"Failed to clear persistent cache: {e}")
            return 0

    # ===== Multi-Column Detection =====

    def _detect_columns(self, blocks: List[Tuple]) -> List[List[Tuple]]:
        """Detect multi-column layout by clustering text blocks by x-position.
        
        Returns list of columns, each containing blocks sorted top-to-bottom.
        """
        if not blocks:
            return [[]]

        # Extract x-positions (left edge of each block)
        x_positions = []
        for block in blocks:
            rect = block[1]  # (text, rect, format_meta) or (text, rect)
            x_positions.append(rect.x0)

        if not x_positions:
            return [blocks]

        # Cluster x-positions using simple gap detection
        sorted_x = sorted(set(x_positions))
        if len(sorted_x) < 2:
            return [blocks]

        # Find significant gaps between x-positions
        column_boundaries = [sorted_x[0] - 10]  # Start before first
        for i in range(1, len(sorted_x)):
            gap = sorted_x[i] - sorted_x[i-1]
            if gap > self.COLUMN_GAP_THRESHOLD:
                column_boundaries.append((sorted_x[i-1] + sorted_x[i]) / 2)
        column_boundaries.append(sorted_x[-1] + 1000)  # End after last

        # If only one column detected, return all blocks
        if len(column_boundaries) <= 2:
            return [sorted(blocks, key=lambda b: b[1].y0)]

        # Assign blocks to columns
        columns: List[List[Tuple]] = [[] for _ in range(len(column_boundaries) - 1)]
        for block in blocks:
            rect = block[1]
            x_center = (rect.x0 + rect.x1) / 2
            # Find which column this block belongs to
            for col_idx in range(len(column_boundaries) - 1):
                if column_boundaries[col_idx] <= x_center < column_boundaries[col_idx + 1]:
                    columns[col_idx].append(block)
                    break

        # Sort each column by y-position (top to bottom)
        for col in columns:
            col.sort(key=lambda b: b[1].y0)

        # Filter empty columns
        columns = [col for col in columns if col]

        return columns if columns else [blocks]

    def _sort_blocks_reading_order(self, blocks: List[Tuple]) -> List[Tuple]:
        """Sort text blocks in natural reading order (left-to-right columns, top-to-bottom)."""
        columns = self._detect_columns(blocks)

        # Flatten columns in reading order
        sorted_blocks = []
        for column in columns:
            sorted_blocks.extend(column)

        return sorted_blocks

    # ===== Quality Scoring =====

    def _compute_bleu_score(self, reference: str, candidate: str) -> float:
        """Compute simplified BLEU score between reference and candidate translations.
        
        Uses 1-4 gram precision with brevity penalty.
        """
        if not reference or not candidate:
            return 0.0

        ref_words = reference.lower().split()
        cand_words = candidate.lower().split()

        if not ref_words or not cand_words:
            return 0.0

        # Compute n-gram precisions
        precisions = []
        for n in range(1, min(5, len(cand_words) + 1)):
            ref_ngrams = self._get_ngrams(ref_words, n)
            cand_ngrams = self._get_ngrams(cand_words, n)

            if not cand_ngrams:
                break

            matches = sum(1 for ng in cand_ngrams if ng in ref_ngrams)
            precision = matches / len(cand_ngrams)
            precisions.append(precision)

        if not precisions:
            return 0.0

        # Geometric mean of precisions
        import math
        log_precisions = [math.log(p) if p > 0 else -float('inf') for p in precisions]
        if -float('inf') in log_precisions:
            return 0.0
        avg_log = sum(log_precisions) / len(log_precisions)
        geo_mean = math.exp(avg_log)

        # Brevity penalty
        bp = 1.0
        if len(cand_words) < len(ref_words):
            bp = math.exp(1 - len(ref_words) / len(cand_words))

        return bp * geo_mean

    def _get_ngrams(self, words: List[str], n: int) -> List[Tuple[str, ...]]:
        """Extract n-grams from word list."""
        return [tuple(words[i:i+n]) for i in range(len(words) - n + 1)]

    def _compute_similarity_score(self, text1: str, text2: str) -> float:
        """Compute word-overlap similarity between two texts."""
        if not text1 or not text2:
            return 0.0

        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())

        if not words1 or not words2:
            return 0.0

        intersection = len(words1 & words2)
        union = len(words1 | words2)

        return intersection / union if union > 0 else 0.0

    async def compute_translation_quality(self, original: str, translated: str,
                                           target_lang: str, source_lang: str = None) -> Dict[str, Any]:
        """Compute translation quality metrics including back-translation BLEU.
        
        Returns:
            - length_ratio: translation length / original length
            - similarity: word overlap between original and translation
            - back_translation_bleu: BLEU of back-translated text vs original
            - overall_score: weighted combination (0-1)
        """
        metrics = {
            "length_ratio": 0.0,
            "similarity": 0.0,
            "back_translation_bleu": None,
            "overall_score": 0.0
        }

        if not original or not translated:
            return metrics

        # Length ratio (normalized to 0-1 range where 1.0 is ideal)
        ratio = len(translated) / len(original)
        # Ideal ratio depends on languages, assume 0.8-1.5 is good
        if 0.8 <= ratio <= 1.5:
            length_score = 1.0
        elif ratio < 0.8:
            length_score = max(0, ratio / 0.8)
        else:
            length_score = max(0, 1 - (ratio - 1.5) / 2)
        metrics["length_ratio"] = round(ratio, 3)

        # Similarity (for related languages, some word overlap expected)
        metrics["similarity"] = round(self._compute_similarity_score(original, translated), 3)

        # Back-translation BLEU (translate back and compare to original)
        try:
            # Use source language detection if not provided
            back_lang = source_lang or "English"  # Fallback
            back_translation = await self._translate_text_gemini(
                translated, back_lang, source_lang=target_lang
            )
            if back_translation:
                bleu = self._compute_bleu_score(original, back_translation)
                metrics["back_translation_bleu"] = round(bleu, 3)
        except Exception as e:
            logger.debug(f"Back-translation failed: {e}")

        # Overall score (weighted)
        bleu_score = metrics["back_translation_bleu"] or 0
        overall = (
            0.3 * length_score +
            0.3 * min(1.0, metrics["similarity"] * 2) +  # Boost low similarity
            0.4 * bleu_score
        )
        metrics["overall_score"] = round(overall, 3)

        return metrics

    # ===== Translation Validation & Hardening =====

    def _extract_entities(self, text: str) -> Dict[str, set]:
        """Extract all preservable entities from text (numbers, dates, URLs, etc.).
        
        Safe extraction with error handling for malformed input.
        """
        entities: Dict[str, set] = {}

        if not text or not isinstance(text, str):
            return entities

        # Truncate extremely long text to avoid regex catastrophic backtracking
        safe_text = text[:50000] if len(text) > 50000 else text

        for entity_type, pattern in self.ENTITY_PATTERNS.items():
            try:
                # Use timeout-safe approach with limited iterations
                matches = set()
                for match in re.finditer(pattern, safe_text, re.IGNORECASE):
                    matches.add(match.group())
                    # Limit matches to prevent memory issues
                    if len(matches) >= 1000:
                        logger.debug(f"Entity extraction limit reached for {entity_type}")
                        break
                if matches:
                    entities[entity_type] = matches
            except re.error as e:
                logger.warning(f"Regex error extracting {entity_type}: {e}")
            except Exception as e:
                logger.warning(f"Unexpected error extracting {entity_type}: {e}")

        return entities

    def _normalize_entity(self, entity: str, entity_type: str) -> str:
        """Normalize entity for comparison (handle format variations)."""
        # Remove whitespace
        normalized = entity.strip()
        if entity_type == 'numbers':
            # Normalize decimal separators (1,234.56 vs 1.234,56)
            normalized = normalized.replace(' ', '')
        elif entity_type == 'urls':
            # Normalize URLs
            normalized = normalized.lower().rstrip('/')
        return normalized

    def _validate_entity_preservation(
        self, original: str, translated: str
    ) -> Tuple[bool, str, float]:
        """Check that critical entities are preserved in translation."""
        orig_entities = self._extract_entities(original)
        trans_entities = self._extract_entities(translated)

        if not orig_entities:
            return True, "", 1.0  # No entities to preserve

        total_entities = 0
        preserved_entities = 0
        missing_details: List[str] = []

        for entity_type, orig_set in orig_entities.items():
            trans_set = trans_entities.get(entity_type, set())

            # Normalize for comparison
            orig_normalized = {self._normalize_entity(e, entity_type) for e in orig_set}
            trans_normalized = {self._normalize_entity(e, entity_type) for e in trans_set}

            for entity in orig_normalized:
                total_entities += 1
                if entity in trans_normalized:
                    preserved_entities += 1
                else:
                    # Check if entity appears anywhere in translated text
                    if entity in translated or entity.lower() in translated.lower():
                        preserved_entities += 1
                    else:
                        missing_details.append(f"{entity_type}: {entity}")

        if total_entities == 0:
            return True, "", 1.0

        preservation_ratio = preserved_entities / total_entities
        if preservation_ratio < self.ENTITY_PRESERVATION_THRESHOLD:
            missing_str = ", ".join(missing_details[:5])
            if len(missing_details) > 5:
                missing_str += f" (+{len(missing_details) - 5} more)"
            return False, f"Entity loss detected ({preservation_ratio:.0%} preserved): {missing_str}", preservation_ratio

        return True, "", preservation_ratio

    def _count_sentences(self, text: str) -> int:
        """Count sentences in text using basic heuristics."""
        if not text:
            return 0
        # Split on sentence-ending punctuation
        # Handle common patterns: . ! ? and their combinations with quotes
        sentences = re.split(r'[.!?]+[\s\n]+|[.!?]+$', text.strip())
        # Filter out empty strings
        return len([s for s in sentences if s.strip()])

    def _validate_sentence_alignment(
        self, original: str, translated: str
    ) -> Tuple[bool, str, float]:
        """Check that translation has similar sentence count to original.
        
        SAFE: Returns valid result even on unexpected input.
        """
        try:
            orig_count = self._count_sentences(original)
            trans_count = self._count_sentences(translated)

            if orig_count == 0:
                return True, "", 1.0

            # Allow some drift due to language differences
            drift = abs(trans_count - orig_count) / orig_count
            if drift > self.MAX_SENTENCE_COUNT_DRIFT:
                return False, f"Sentence count mismatch (original: {orig_count}, translated: {trans_count}, drift: {drift:.0%})", 1.0 - drift

            return True, "", 1.0 - drift
        except Exception as e:
            logger.debug(f"Sentence alignment check failed safely: {e}")
            return True, "", 1.0  # Fail open - don't block translation

    def _validate_translation(self, original: str, translated: str) -> Tuple[bool, str]:
        """Validate translation quality to detect hallucinations and summarization.
        
        ROBUST: Never crashes - returns (True, "") on any unexpected error.
        This ensures translations proceed even with unusual content.
        """
        try:
            if not original or not translated:
                return True, ""

            # Safely get lengths
            try:
                orig_len = len(str(original).strip())
                trans_len = len(str(translated).strip())
            except Exception:
                return True, ""

            if orig_len == 0:
                return True, ""

            # 1. Length ratio check - translation shouldn't be wildly different
            try:
                ratio = trans_len / orig_len
                if ratio > self.MAX_LENGTH_RATIO:
                    return False, f"Translation too long ({ratio:.1f}x original, max {self.MAX_LENGTH_RATIO}x)"
                if ratio < self.MIN_LENGTH_RATIO:
                    return False, f"Translation too short ({ratio:.1f}x original, min {self.MIN_LENGTH_RATIO}x)"
            except ZeroDivisionError:
                pass

            # 2. Repetition detection - catch hallucinated repeated phrases
            try:
                words = str(translated).split()
                if len(words) > 10:
                    word_counts: Dict[str, int] = {}
                    for word in words:
                        word_lower = word.lower()
                        word_counts[word_lower] = word_counts.get(word_lower, 0) + 1

                    if word_counts:
                        max_repeat = max(word_counts.values())
                        repeat_ratio = max_repeat / len(words)
                        if repeat_ratio > self.MAX_REPETITION_RATIO:
                            most_repeated = max(word_counts, key=word_counts.get)
                            return False, f"Excessive repetition detected ('{most_repeated}' appears {max_repeat} times)"
            except Exception as e:
                logger.debug(f"Repetition check failed safely: {e}")

            # 3. Entity preservation check - numbers, dates, URLs must be preserved
            try:
                entity_valid, entity_reason, _ = self._validate_entity_preservation(original, translated)
                if not entity_valid:
                    return False, entity_reason
            except Exception as e:
                logger.debug(f"Entity preservation check failed safely: {e}")

            # 4. Sentence alignment check - detect summarization
            try:
                sentence_valid, sentence_reason, _ = self._validate_sentence_alignment(original, translated)
                if not sentence_valid:
                    # Log warning but don't reject - sentence counting is imprecise
                    logger.warning(f"Sentence alignment warning: {sentence_reason}")
            except Exception as e:
                logger.debug(f"Sentence alignment check failed safely: {e}")

            return True, ""

        except Exception as e:
            # Ultimate fallback - never crash validation
            logger.warning(f"Translation validation encountered unexpected error: {e}")
            return True, ""

    def _extract_page_images(self, page) -> List[Dict[str, Any]]:
        """Extract images from page to preserve during translation."""
        images: List[Dict[str, Any]] = []
        try:
            for img_idx, img_info in enumerate(page.get_images(full=True)):
                xref = img_info[0]
                try:
                    base_image = page.parent.extract_image(xref)
                    if base_image:
                        rects = page.get_image_rects(xref)
                        rect = rects[0] if rects else None
                        images.append({
                            "xref": xref,
                            "data": base_image["image"],
                            "ext": base_image.get("ext", "png"),
                            "rect": rect,
                            "colorspace": base_image.get("colorspace", 0),
                        })
                except Exception as e:
                    logger.debug(f"Could not extract image {xref}: {e}")
        except Exception as e:
            logger.warning(f"Failed to extract images from page: {e}")
        return images

    def _reinsert_images(self, page, images: List[Dict[str, Any]]) -> None:
        """Re-insert preserved images into page after translation."""
        for img in images:
            rect = img.get("rect")
            if rect and img.get("data"):
                try:
                    page.insert_image(rect, stream=img["data"])
                except Exception as e:
                    logger.warning(f"Failed to reinsert image: {e}")

    def _get_image_regions(self, page) -> List["fitz.Rect"]:
        """Get bounding rectangles of all images on a page."""
        regions = []
        try:
            for img_info in page.get_images(full=True):
                xref = img_info[0]
                rects = page.get_image_rects(xref)
                for r in rects:
                    if r and not r.is_empty:
                        regions.append(r)
        except Exception as e:
            logger.debug(f"Failed to get image regions: {e}")
        return regions

    def _rect_overlaps_images(self, rect: "fitz.Rect", image_regions: List["fitz.Rect"], threshold: float = 0.1) -> bool:
        """Check if a text rect significantly overlaps with any image region.
        
        Args:
            rect: Text block rectangle
            image_regions: List of image bounding rectangles
            threshold: Minimum overlap ratio (0-1) to consider as overlap
            
        Returns:
            True if text block overlaps significantly with any image
        """
        if not rect or rect.is_empty:
            return False
        rect_area = rect.width * rect.height
        if rect_area <= 0:
            return False
            
        for img_rect in image_regions:
            if not img_rect or img_rect.is_empty:
                continue
            # Calculate intersection area
            intersect = rect & img_rect  # fitz.Rect intersection operator
            if intersect and not intersect.is_empty:
                intersect_area = intersect.width * intersect.height
                overlap_ratio = intersect_area / rect_area
                if overlap_ratio >= threshold:
                    return True
        return False

    def _get_retry_delay(self, attempt: int) -> float:
        """Calculate exponential backoff delay with jitter."""
        base_delay = self.RETRY_BASE_DELAY * (2 ** (attempt - 1))
        jitter = random.uniform(0, 1)
        return min(base_delay + jitter, self.RETRY_MAX_DELAY)

    # ===== Quick Win #1: Header/Footer Detection =====
    
    # Font patterns indicating code (monospace fonts)
    CODE_FONT_PATTERNS = ("courier", "consolas", "mono", "menlo", "source code", "fira", "roboto mono", "inconsolata")
    
    # Text patterns indicating code (should not be translated)
    CODE_TEXT_PATTERNS = (
        "```", "def ", "class ", "function ", "function(", "import ", "from ", "#include",
        "int ", "void ", "public ", "private ", "return ", "if (", "for (", "while (",
        "const ", "let ", "var ", "=>", "async ", "await ", "try {", "catch (",
        "#!/", "<?php", "using ", "namespace ", "struct ", "#define", "typedef ",
    )

    def _detect_repeated_blocks(self, src_doc, min_occurrences: int = 3) -> set:
        """Find text blocks that repeat on multiple pages (headers/footers).
        
        Headers and footers typically appear on 3+ pages. Detecting them
        allows us to skip redundant translation calls and maintain consistency.
        
        Args:
            src_doc: PyMuPDF document
            min_occurrences: Minimum number of pages text must appear on
            
        Returns:
            Set of MD5 hashes of repeated text blocks
        """
        text_counts: Dict[str, int] = {}
        
        for page in src_doc:
            try:
                blocks = page.get_text("blocks")
                for block in blocks:
                    if len(block) >= 5 and block[6] == 0:  # text block (not image)
                        text = block[4].strip()
                        if len(text) < 5 or len(text) > 200:  # Skip very short/long
                            continue
                        # Normalize whitespace for consistent hashing
                        text_normalized = " ".join(text.split())
                        text_hash = hashlib.md5(text_normalized.encode()).hexdigest()
                        text_counts[text_hash] = text_counts.get(text_hash, 0) + 1
            except Exception as e:
                logger.debug(f"Error scanning page for headers/footers: {e}")
                
        repeated = {h for h, count in text_counts.items() if count >= min_occurrences}
        if repeated:
            logger.info(f"Detected {len(repeated)} repeated text blocks (headers/footers)")
        return repeated

    def _is_repeated_block(self, text: str, repeated_hashes: set) -> bool:
        """Check if text matches a detected header/footer pattern."""
        if not repeated_hashes or not text:
            return False
        text_normalized = " ".join(text.strip().split())
        text_hash = hashlib.md5(text_normalized.encode()).hexdigest()
        return text_hash in repeated_hashes

    # ===== Quick Win #2: Code Block Detection =====

    def _is_code_block(self, text: str, format_meta: Optional[Dict] = None) -> bool:
        """Detect if a text block is code (should not be translated).
        
        Code detection uses:
        1. Monospace font names (Courier, Consolas, etc.)
        2. Code syntax patterns (def, class, function, etc.)
        3. High symbol density ({}, [], ;, etc.)
        
        Args:
            text: Text content to analyze
            format_meta: Optional format metadata with font_name
            
        Returns:
            True if text appears to be code
        """
        if not text or len(text) < 3:
            return False
            
        # Check font name for monospace indicators
        if format_meta:
            font_name = format_meta.get("font_name", "").lower()
            if any(pattern in font_name for pattern in self.CODE_FONT_PATTERNS):
                return True
        
        # Check for code syntax patterns
        text_check = text.lower()
        if any(pattern.lower() in text_check for pattern in self.CODE_TEXT_PATTERNS):
            return True
            
        # Check for high symbol density (code has lots of {}[]();:=<>)
        if len(text) >= 10:
            code_symbols = "{}[]();:=<>|&!@#$%^*+-/\\"
            symbol_count = sum(1 for c in text if c in code_symbols)
            symbol_ratio = symbol_count / len(text)
            if symbol_ratio > 0.12:  # More than 12% symbols = likely code
                return True
                
        return False

    # ===== Quick Win #3: Hyperlink Preservation =====

    def _extract_page_links(self, page) -> List[Dict]:
        """Extract all hyperlinks from a page.
        
        Args:
            page: PyMuPDF page object
            
        Returns:
            List of link dictionaries (PyMuPDF link format)
        """
        try:
            links = page.get_links()
            return [link for link in links if link.get("uri") or link.get("page")]
        except Exception as e:
            logger.debug(f"Failed to extract links: {e}")
            return []

    def _reinsert_links(self, page, links: List[Dict]) -> int:
        """Reinsert hyperlinks into a page after text translation.
        
        Args:
            page: PyMuPDF page object
            links: List of link dictionaries to reinsert
            
        Returns:
            Number of links successfully reinserted
        """
        if not links:
            return 0
            
        count = 0
        for link in links:
            try:
                page.insert_link(link)
                count += 1
            except Exception as e:
                logger.debug(f"Failed to reinsert link: {e}")
        return count

    # ===== Image Text Translation (OCR within embedded images) =====

    def _ocr_image_with_boxes(
        self,
        pil_image: "Image.Image",
        ocr_lang: str = "eng+rus",
    ) -> List[Dict[str, Any]]:
        """OCR an image and return text regions with bounding boxes.
        
        Args:
            pil_image: PIL Image object
            ocr_lang: Tesseract language code(s)
            
        Returns:
            List of dicts with keys: text, x, y, w, h, conf
        """
        if not self._has_ocr or pytesseract is None:
            return []
        
        try:
            # Get word-level OCR data with positions
            data = pytesseract.image_to_data(pil_image, lang=ocr_lang, output_type=Output.DICT)
            
            regions = []
            n_boxes = len(data["text"])
            
            # Group words into lines for better translation context
            lines: Dict[Tuple[int, int, int], List[Dict]] = {}
            
            for i in range(n_boxes):
                text = data["text"][i].strip()
                conf = int(data["conf"][i]) if data["conf"][i] != "-1" else 0
                
                # Skip empty or low-confidence text
                if not text or conf < 30:
                    continue
                
                line_key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
                word_info = {
                    "text": text,
                    "x": data["left"][i],
                    "y": data["top"][i],
                    "w": data["width"][i],
                    "h": data["height"][i],
                    "conf": conf,
                }
                
                if line_key not in lines:
                    lines[line_key] = []
                lines[line_key].append(word_info)
            
            # Consolidate words into line regions
            for line_key, words in lines.items():
                if not words:
                    continue
                
                line_text = " ".join(w["text"] for w in words)
                x_min = min(w["x"] for w in words)
                y_min = min(w["y"] for w in words)
                x_max = max(w["x"] + w["w"] for w in words)
                y_max = max(w["y"] + w["h"] for w in words)
                avg_conf = sum(w["conf"] for w in words) / len(words)
                
                regions.append({
                    "text": line_text,
                    "x": x_min,
                    "y": y_min,
                    "w": x_max - x_min,
                    "h": y_max - y_min,
                    "conf": avg_conf,
                })
            
            return regions
            
        except Exception as e:
            logger.warning(f"OCR with boxes failed: {e}")
            return []

    def _get_dominant_color(
        self,
        pil_image: "Image.Image",
        x: int,
        y: int,
        w: int,
        h: int,
        sample_border: int = 3,
    ) -> Tuple[int, int, int]:
        """Get the dominant background color around a text region.
        
        Samples pixels from the border of the region to estimate background.
        """
        try:
            img_w, img_h = pil_image.size
            pixels = []
            
            # Sample from edges of the region (likely background)
            for dx in range(0, w, max(1, w // 10)):
                # Top edge
                px, py = x + dx, max(0, y - sample_border)
                if 0 <= px < img_w and 0 <= py < img_h:
                    pixels.append(pil_image.getpixel((px, py)))
                # Bottom edge
                px, py = x + dx, min(img_h - 1, y + h + sample_border)
                if 0 <= px < img_w and 0 <= py < img_h:
                    pixels.append(pil_image.getpixel((px, py)))
            
            for dy in range(0, h, max(1, h // 10)):
                # Left edge
                px, py = max(0, x - sample_border), y + dy
                if 0 <= px < img_w and 0 <= py < img_h:
                    pixels.append(pil_image.getpixel((px, py)))
                # Right edge
                px, py = min(img_w - 1, x + w + sample_border), y + dy
                if 0 <= px < img_w and 0 <= py < img_h:
                    pixels.append(pil_image.getpixel((px, py)))
            
            if not pixels:
                return (255, 255, 255)  # Default white
            
            # Handle grayscale images
            if isinstance(pixels[0], int):
                avg = sum(pixels) // len(pixels)
                return (avg, avg, avg)
            
            # Handle RGBA
            if len(pixels[0]) == 4:
                pixels = [(p[0], p[1], p[2]) for p in pixels]
            
            # Average color
            r = sum(p[0] for p in pixels) // len(pixels)
            g = sum(p[1] for p in pixels) // len(pixels)
            b = sum(p[2] for p in pixels) // len(pixels)
            return (r, g, b)
            
        except Exception:
            return (255, 255, 255)

    async def _translate_image_text(
        self,
        image_data: bytes,
        target_language: str,
        source_language: Optional[str] = None,
        ocr_lang: str = "eng+rus",
        min_text_length: int = 2,
    ) -> Tuple[bytes, int]:
        """Translate text within an image.
        
        Args:
            image_data: Original image bytes
            target_language: Target language for translation
            source_language: Source language (optional)
            ocr_lang: Tesseract language code(s)
            min_text_length: Minimum text length to translate
            
        Returns:
            Tuple of (modified image bytes, number of regions translated)
        """
        if not self._has_ocr or Image is None:
            return image_data, 0
        
        try:
            from PIL import ImageDraw, ImageFont
            from io import BytesIO
            
            # Load image
            pil_image = Image.open(BytesIO(image_data))
            if pil_image.mode not in ("RGB", "RGBA"):
                pil_image = pil_image.convert("RGB")
            
            # OCR to get text regions
            regions = self._ocr_image_with_boxes(pil_image, ocr_lang)
            
            if not regions:
                return image_data, 0
            
            # Filter regions with enough text
            regions = [r for r in regions if len(r["text"]) >= min_text_length]
            if not regions:
                return image_data, 0
            
            # Translate all text regions
            texts_to_translate = [r["text"] for r in regions]
            translated_texts = []
            
            for text in texts_to_translate:
                try:
                    translated = await self._translate_chunk(text, target_language, source_language)
                    translated_texts.append(translated if translated else text)
                except Exception:
                    translated_texts.append(text)
            
            # Create drawing context
            draw = ImageDraw.Draw(pil_image)
            
            # Try to load a font that supports target language
            font_path = self._get_font_for_language(target_language)
            
            translated_count = 0
            for region, translated in zip(regions, translated_texts):
                if translated == region["text"]:
                    continue  # Skip if translation is same as original
                
                x, y, w, h = region["x"], region["y"], region["w"], region["h"]
                
                # Get background color for this region
                bg_color = self._get_dominant_color(pil_image, x, y, w, h)
                
                # Cover original text with background color (with small padding)
                padding = 2
                draw.rectangle(
                    [x - padding, y - padding, x + w + padding, y + h + padding],
                    fill=bg_color
                )
                
                # Calculate font size to fit the region height
                target_font_size = max(8, int(h * 0.85))
                
                try:
                    font = ImageFont.truetype(font_path, target_font_size)
                except Exception:
                    font = ImageFont.load_default()
                
                # Determine text color (contrast with background)
                bg_luminance = (bg_color[0] * 0.299 + bg_color[1] * 0.587 + bg_color[2] * 0.114)
                text_color = (0, 0, 0) if bg_luminance > 128 else (255, 255, 255)
                
                # Draw translated text
                draw.text((x, y), translated, font=font, fill=text_color)
                translated_count += 1
            
            # Save modified image
            output = BytesIO()
            img_format = "PNG" if pil_image.mode == "RGBA" else "JPEG"
            pil_image.save(output, format=img_format, quality=95)
            
            return output.getvalue(), translated_count
            
        except Exception as e:
            logger.warning(f"Image text translation failed: {e}")
            return image_data, 0

    async def _translate_page_images(
        self,
        page_images: List[Dict[str, Any]],
        target_language: str,
        source_language: Optional[str] = None,
        ocr_lang: str = "eng+rus",
    ) -> Tuple[List[Dict[str, Any]], int]:
        """Translate text in all images on a page.
        
        Args:
            page_images: List of image dicts from _extract_page_images
            target_language: Target language
            source_language: Source language
            ocr_lang: OCR language code
            
        Returns:
            Tuple of (modified images list, total regions translated)
        """
        total_translated = 0
        modified_images = []
        
        for img in page_images:
            if not img.get("data"):
                modified_images.append(img)
                continue
            
            translated_data, count = await self._translate_image_text(
                img["data"],
                target_language,
                source_language,
                ocr_lang,
            )
            
            modified_images.append({
                **img,
                "data": translated_data,
            })
            total_translated += count
        
        return modified_images, total_translated

    # ===== Glossary/Terminology Support =====

    def set_glossary(self, glossary: Dict[str, str]) -> None:
        """Set terminology glossary for consistent translations."""
        self._glossary = {k.lower(): v for k, v in glossary.items()}
        logger.info(f"Glossary set with {len(self._glossary)} terms")

    def clear_glossary(self) -> None:
        """Clear the terminology glossary."""
        self._glossary.clear()

    def get_glossary(self) -> Dict[str, str]:
        """Get current glossary."""
        return self._glossary.copy()

    def _build_glossary_prompt(self) -> str:
        """Build glossary instruction for translation prompt."""
        if not self._glossary:
            return ""
        terms = "\n".join([f"'{k}' -> '{v}'" for k, v in list(self._glossary.items())[:50]])
        return f"\n\nIMPORTANT TERMINOLOGY (use these exact translations):\n{terms}\n"

    # ===== Checkpoint/Resume Support =====

    def save_checkpoint(self, checkpoint_path: str, data: Dict[str, Any]) -> None:
        """Save translation checkpoint for resume."""
        import json as json_module
        checkpoint = {
            "timestamp": time.time(),
            "completed_chunks": data.get("completed_chunks", []),
            "translated_text": data.get("translated_text", ""),
            "target_language": data.get("target_language"),
            "source_language": data.get("source_language"),
            "total_chunks": data.get("total_chunks", 0),
        }
        with open(checkpoint_path, 'w', encoding='utf-8') as f:
            json_module.dump(checkpoint, f, ensure_ascii=False, indent=2)
        logger.info(f"Checkpoint saved: {len(checkpoint['completed_chunks'])}/{checkpoint['total_chunks']} chunks")

    def load_checkpoint(self, checkpoint_path: str) -> Optional[Dict[str, Any]]:
        """Load translation checkpoint for resume."""
        import json as json_module
        try:
            with open(checkpoint_path, 'r', encoding='utf-8') as f:
                checkpoint = json_module.load(f)
            age_hours = (time.time() - checkpoint.get("timestamp", 0)) / 3600
            if age_hours > 24:
                logger.warning(f"Checkpoint is {age_hours:.1f} hours old, may be stale")
            logger.info(f"Checkpoint loaded: {len(checkpoint.get('completed_chunks', []))}/{checkpoint.get('total_chunks', 0)} chunks")
            return checkpoint
        except (FileNotFoundError, ValueError) as e:
            logger.warning(f"Could not load checkpoint: {e}")
            return None

    def delete_checkpoint(self, checkpoint_path: str) -> None:
        """Delete checkpoint file after successful completion."""
        from pathlib import Path
        try:
            Path(checkpoint_path).unlink(missing_ok=True)
            logger.info(f"Checkpoint deleted: {checkpoint_path}")
        except Exception as e:
            logger.warning(f"Could not delete checkpoint: {e}")

    # ===== Export Methods =====

    def export_to_docx(
        self,
        original: str,
        translated: str,
        output_path: str,
        side_by_side: bool = False,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Export translation to DOCX format with improved formatting.
        
        Features:
        - Professional table styling with borders
        - Alternating row colors for readability
        - Proper font sizes and spacing
        - Metadata section
        - Support for both side-by-side and translated-only views
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        
        # Set default font for document
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Add title with styling
        title = doc.add_heading('Document Translation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section with better formatting
        if metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            source_lang = metadata.get('source_language', 'Auto-detected')
            target_lang = metadata.get('target_language', 'Unknown')
            timestamp = time.strftime('%Y-%m-%d %H:%M:%S')
            
            run = meta_para.add_run(f"Source: {source_lang}  |  Target: {target_lang}  |  Generated: {timestamp}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Add quality metrics if available
            if metadata.get('quality_score'):
                quality_para = doc.add_paragraph()
                quality_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                quality_run = quality_para.add_run(f"Quality Score: {metadata.get('quality_score', 0):.2f}")
                quality_run.font.size = Pt(9)
                quality_run.font.color.rgb = RGBColor(0, 128, 0) if metadata.get('quality_score', 0) > 0.5 else RGBColor(200, 100, 0)
            
            doc.add_paragraph()  # Spacer

        def set_cell_shading(cell, color: str):
            """Set cell background color."""
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            cell._tc.get_or_add_tcPr().append(shading)
        
        def add_table_borders(table):
            """Add borders to all cells in table."""
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '666666')
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)

        if side_by_side:
            # Create professionally styled two-column table
            orig_lines = original.splitlines()
            trans_lines = translated.splitlines()
            max_lines = max(len(orig_lines), len(trans_lines))
            
            doc.add_heading('Side-by-Side Comparison', level=1)
            
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            
            # Style header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Original'
            hdr_cells[1].text = 'Translated'
            
            for cell in hdr_cells:
                set_cell_shading(cell, '4472C4')  # Blue header
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(11)

            # Add content rows with alternating colors
            for i in range(max_lines):
                row = table.add_row().cells
                row[0].text = orig_lines[i] if i < len(orig_lines) else ''
                row[1].text = trans_lines[i] if i < len(trans_lines) else ''
                
                # Alternating row colors
                bg_color = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
                for cell in row:
                    set_cell_shading(cell, bg_color)
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(4)
                        for run in para.runs:
                            run.font.size = Pt(10)
            
            # Add borders
            add_table_borders(table)
            
            # Set column widths (roughly equal)
            for row in table.rows:
                row.cells[0].width = Inches(3.5)
                row.cells[1].width = Inches(3.5)
                
        else:
            # Translated text only with proper paragraph formatting
            doc.add_heading('Translated Content', level=1)
            
            paragraphs = translated.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    para.paragraph_format.space_after = Pt(8)
                    para.paragraph_format.line_spacing = 1.15
                    
                    # Handle line breaks within paragraph
                    lines = para_text.strip().split('\n')
                    for idx, line in enumerate(lines):
                        run = para.add_run(line)
                        run.font.size = Pt(11)
                        if idx < len(lines) - 1:
                            para.add_run('\n')

        # Add footer with translation stats
        if metadata:
            doc.add_paragraph()  # Spacer
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            stats = []
            if metadata.get('page_count'):
                stats.append(f"Pages: {metadata['page_count']}")
            if metadata.get('chars_translated'):
                stats.append(f"Characters: {metadata['chars_translated']:,}")
            if metadata.get('tables_stitched'):
                stats.append(f"Tables: {metadata['tables_stitched']}")
            
            if stats:
                footer_run = footer_para.add_run(' | '.join(stats))
                footer_run.font.size = Pt(8)
                footer_run.font.color.rgb = RGBColor(150, 150, 150)

        doc.save(output_path)
        logger.info(f"DOCX exported to: {output_path}")
        return output_path

    def export_pdf_to_docx(
        self,
        pdf_path: str,
        output_path: Optional[str] = None,
        preserve_layout: bool = True,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Convert a translated PDF to DOCX format, preserving layout where possible.
        
        Args:
            pdf_path: Path to the PDF file
            output_path: Output DOCX path (defaults to pdf_path with .docx extension)
            preserve_layout: Try to preserve PDF layout in DOCX
            metadata: Optional metadata to include
            
        Returns:
            Path to the created DOCX file
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        if not self._has_pymupdf:
            raise ValueError("PyMuPDF required for PDF to DOCX conversion")
        
        if output_path is None:
            output_path = str(Path(pdf_path).with_suffix('.docx'))
        
        pdf_doc = fitz.open(pdf_path)
        doc = Document()
        
        # Set default style
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # Add title
        title = doc.add_heading('Translated Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata if provided
        if metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_run = meta_para.add_run(
                f"Source: {metadata.get('source_language', 'Unknown')} → "
                f"Target: {metadata.get('target_language', 'Unknown')}"
            )
            meta_run.italic = True
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()
        
        def add_table_borders(table):
            """Add borders to table."""
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '888888')
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
            if tbl.tblPr is None:
                tbl.insert(0, tblPr)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            
            # Add page separator for multi-page docs
            if page_num > 0:
                doc.add_page_break()
                page_header = doc.add_paragraph()
                page_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = page_header.add_run(f"— Page {page_num + 1} —")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(150, 150, 150)
                doc.add_paragraph()
            
            if preserve_layout:
                # Extract text blocks with position info
                blocks = page.get_text("dict")["blocks"]
                
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        for line in block.get("lines", []):
                            para = doc.add_paragraph()
                            
                            for span in line.get("spans", []):
                                text = span.get("text", "")
                                if not text.strip():
                                    continue
                                
                                run = para.add_run(text + " ")
                                
                                # Preserve font size (scale from PDF points)
                                font_size = span.get("size", 11)
                                run.font.size = Pt(min(max(font_size * 0.75, 8), 24))
                                
                                # Preserve bold/italic from font flags
                                flags = span.get("flags", 0)
                                run.bold = bool(flags & 2**4)  # Bold flag
                                run.italic = bool(flags & 2**1)  # Italic flag
                            
                            para.paragraph_format.space_after = Pt(4)
                    
                    elif block.get("type") == 1:  # Image block
                        # Add placeholder for images
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run("[Image]")
                        run.font.color.rgb = RGBColor(150, 150, 150)
                        run.italic = True
            else:
                # Simple text extraction
                text = page.get_text()
                for para_text in text.split('\n\n'):
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.paragraph_format.space_after = Pt(8)
            
            # Try to extract and convert tables
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    for pdf_table in tables.tables:
                        doc.add_paragraph()  # Spacer
                        
                        table_data = pdf_table.extract()
                        if table_data and len(table_data) > 0:
                            num_cols = max(len(row) for row in table_data)
                            docx_table = doc.add_table(rows=0, cols=num_cols)
                            docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            
                            for row_idx, row_data in enumerate(table_data):
                                row = docx_table.add_row()
                                for col_idx, cell_text in enumerate(row_data):
                                    if col_idx < num_cols:
                                        cell = row.cells[col_idx]
                                        cell.text = str(cell_text) if cell_text else ""
                                        
                                        # Style first row as header
                                        if row_idx == 0:
                                            for para in cell.paragraphs:
                                                for run in para.runs:
                                                    run.bold = True
                            
                            add_table_borders(docx_table)
                            doc.add_paragraph()  # Spacer after table
            except Exception as e:
                logger.warning(f"Could not extract tables from page {page_num + 1}: {e}")
        
        pdf_doc.close()
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(f"Converted from PDF | {len(pdf_doc)} pages | {time.strftime('%Y-%m-%d %H:%M')}")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        
        doc.save(output_path)
        logger.info(f"PDF converted to DOCX: {output_path}")
        return output_path

    def generate_side_by_side_pdf(
        self,
        original: str,
        translated: str,
        output_path: str,
        target_language: str = "english",
    ) -> str:
        """Generate side-by-side comparison PDF."""
        if not self._has_pymupdf:
            raise ValueError("PyMuPDF required for PDF generation")

        doc = fitz.open()
        page_width, page_height = 842, 595  # A4 landscape
        margin = 40
        col_width = (page_width - 3 * margin) / 2
        font_size = 10
        line_height = font_size * 1.4

        orig_lines = original.splitlines()
        trans_lines = translated.splitlines()
        max_lines = max(len(orig_lines), len(trans_lines))

        font_path = self._get_font_for_language(target_language)

        y = margin + 30
        page = doc.new_page(width=page_width, height=page_height)

        # Add header
        page.insert_text((margin, margin + 15), "Original", fontsize=12, fontname="helv")
        page.insert_text((margin + col_width + margin, margin + 15), "Translated", fontsize=12, fontname="helv")
        page.draw_line((margin, margin + 25), (page_width - margin, margin + 25))

        for i in range(max_lines):
            if y + line_height > page_height - margin:
                page = doc.new_page(width=page_width, height=page_height)
                y = margin + 30
                page.insert_text((margin, margin + 15), "Original", fontsize=12, fontname="helv")
                page.insert_text((margin + col_width + margin, margin + 15), "Translated", fontsize=12, fontname="helv")
                page.draw_line((margin, margin + 25), (page_width - margin, margin + 25))

            orig_text = orig_lines[i][:80] if i < len(orig_lines) else ""
            trans_text = trans_lines[i][:80] if i < len(trans_lines) else ""

            page.insert_text((margin, y), orig_text, fontsize=font_size, fontname="helv")

            try:
                page.insert_text(
                    (margin + col_width + margin, y),
                    trans_text,
                    fontsize=font_size,
                    fontfile=font_path,
                )
            except Exception:
                page.insert_text((margin + col_width + margin, y), trans_text, fontsize=font_size, fontname="helv")

            y += line_height

        doc.save(output_path)
        doc.close()
        logger.info(f"Side-by-side PDF exported to: {output_path}")
        return output_path

    # ===== Exact Font Embedding =====

    def _extract_document_fonts(self, src_doc: "fitz.Document") -> Dict[str, bytes]:
        """Extract embedded fonts from source PDF for exact font matching.
        
        Returns a dict mapping font name to font file bytes.
        """
        extracted_fonts: Dict[str, bytes] = {}
        
        if not self.FONT_CACHE_ENABLED:
            return extracted_fonts
        
        try:
            # Ensure fonts directory exists
            os.makedirs(self.EXTRACTED_FONTS_DIR, exist_ok=True)
            
            # Get all unique font references from the document
            for page_num in range(len(src_doc)):
                page = src_doc[page_num]
                fonts = page.get_fonts(full=True)
                
                for font_info in fonts:
                    xref = font_info[0]
                    font_name = font_info[3]  # Base font name
                    font_ext = font_info[4]    # Font extension type
                    
                    if not font_name or font_name in extracted_fonts:
                        continue
                    
                    # Skip fonts we already have fallbacks for
                    if any(skip in font_name.lower() for skip in ['arial', 'times', 'courier', 'helvetica']):
                        continue
                    
                    try:
                        # Extract the font data
                        font_data = src_doc.extract_font(xref)
                        if font_data and len(font_data) > 3:
                            font_buffer = font_data[3]  # Raw font bytes
                            if font_buffer and len(font_buffer) > 0:
                                extracted_fonts[font_name] = font_buffer
                                # Cache to disk for reuse
                                cache_path = os.path.join(
                                    self.EXTRACTED_FONTS_DIR, 
                                    f"{hashlib.md5(font_name.encode()).hexdigest()}.font"
                                )
                                with open(cache_path, 'wb') as f:
                                    f.write(font_buffer)
                                logger.debug(f"Extracted font: {font_name} ({len(font_buffer)} bytes)")
                    except Exception as e:
                        logger.debug(f"Could not extract font {font_name}: {e}")
                        
        except Exception as e:
            logger.warning(f"Font extraction failed: {e}")
            
        return extracted_fonts

    def _register_extracted_fonts(self, doc: "fitz.Document", 
                                  extracted_fonts: Dict[str, bytes]) -> Dict[str, str]:
        """Register extracted fonts with the output document.
        
        Returns a mapping from original font name to registered font name.
        """
        font_mapping: Dict[str, str] = {}
        
        for font_name, font_buffer in extracted_fonts.items():
            try:
                # Create a temporary file for the font
                font_path = os.path.join(
                    self.EXTRACTED_FONTS_DIR,
                    f"{hashlib.md5(font_name.encode()).hexdigest()}.font"
                )
                
                # Write font buffer if not already cached
                if not os.path.exists(font_path):
                    with open(font_path, 'wb') as f:
                        f.write(font_buffer)
                
                # PyMuPDF accepts font file path or buffer for inserting text
                # Store the path for later use with insert_text
                font_mapping[font_name] = font_path
                logger.debug(f"Registered font: {font_name} -> {font_path}")
                
            except Exception as e:
                logger.debug(f"Could not register font {font_name}: {e}")
                
        return font_mapping

    def _get_font_path_for_block(self, format_meta: Dict[str, Any],
                                  font_mapping: Dict[str, str],
                                  target_lang: str) -> Optional[str]:
        """Get the best font path for a text block.
        
        Tries to use the original font if available, falls back to language-appropriate font.
        """
        original_font = format_meta.get('font_name', '')
        
        # Try exact font match first
        if original_font and original_font in font_mapping:
            return font_mapping[original_font]
        
        # Try partial match (font families)
        for font_name, font_path in font_mapping.items():
            if original_font and (font_name.lower() in original_font.lower() or 
                                   original_font.lower() in font_name.lower()):
                return font_path
        
        # Fall back to language-appropriate font
        return None  # Will use default FONT_MAP

    # ===== Multi-Page Table Stitching =====

    def _detect_multi_page_tables(self, src_doc: "fitz.Document") -> List[Dict[str, Any]]:
        """Detect tables that span multiple pages.
        
        Returns list of table groups that should be merged before translation.
        """
        multi_page_tables: List[Dict[str, Any]] = []
        
        if not self._has_pymupdf:
            return multi_page_tables
        
        try:
            page_tables: List[List[Dict]] = []
            
            # Extract tables from all pages
            for page_num in range(len(src_doc)):
                page = src_doc[page_num]
                tables = self._extract_tables_from_page(page)
                page_tables.append(tables)
            
            # Find continuation patterns
            for page_num in range(len(src_doc) - 1):
                current_page_tables = page_tables[page_num]
                next_page_tables = page_tables[page_num + 1]
                
                page = src_doc[page_num]
                page_height = page.rect.height
                
                for curr_table in current_page_tables:
                    # Check if table extends to bottom of page
                    if not self._is_table_at_page_boundary(curr_table, page_height, 'bottom'):
                        continue
                    
                    # Find matching table at top of next page
                    for next_table in next_page_tables:
                        if not self._is_table_at_page_boundary(next_table, 0, 'top'):
                            continue
                        
                        # Check if column structure matches
                        if self._tables_have_matching_columns(curr_table, next_table):
                            multi_page_tables.append({
                                'start_page': page_num,
                                'end_page': page_num + 1,
                                'first_table': curr_table,
                                'second_table': next_table,
                                'merged': False
                            })
                            logger.debug(f"Detected multi-page table: pages {page_num}-{page_num+1}")
                            break
                            
        except Exception as e:
            logger.warning(f"Multi-page table detection failed: {e}")
            
        return multi_page_tables

    def _is_table_at_page_boundary(self, table: Dict[str, Any], 
                                    boundary_y: float, 
                                    position: str) -> bool:
        """Check if table is at page boundary (top or bottom)."""
        bbox = table.get('bbox', (0, 0, 0, 0))
        
        if position == 'bottom':
            # Table bottom edge is close to page bottom
            return abs(bbox[3] - boundary_y) < self.TABLE_PAGE_MARGIN
        elif position == 'top':
            # Table top edge is close to page top
            return bbox[1] < self.TABLE_PAGE_MARGIN
        
        return False

    def _tables_have_matching_columns(self, table1: Dict[str, Any], 
                                       table2: Dict[str, Any]) -> bool:
        """Check if two tables have matching column structure."""
        # Simple column count comparison
        col_count1 = table1.get('col_count', 0)
        col_count2 = table2.get('col_count', 0)
        
        if col_count1 == 0 or col_count2 == 0:
            return False
        
        if col_count1 != col_count2:
            return False
        
        # Check row structure consistency
        rows1 = table1.get('rows', [])
        rows2 = table2.get('rows', [])
        
        if not rows1 or not rows2:
            return False
        
        # Verify last row of first table and first row of second have same column count
        if len(rows1[-1]) != len(rows2[0]):
            return False
        
        return True

    def _merge_multi_page_table(self, table_group: Dict[str, Any]) -> Dict[str, Any]:
        """Merge a multi-page table into a single table structure."""
        first_table = table_group['first_table']
        second_table = table_group['second_table']
        
        merged_rows = first_table.get('rows', [])[:]
        
        # Add rows from second table (skip header if it's a duplicate)
        second_rows = second_table.get('rows', [])
        if second_rows:
            # Check if first row is a repeated header
            first_row_second = second_rows[0]
            first_row_first = first_table.get('rows', [[]])[0] if first_table.get('rows') else []
            
            # Skip header if it matches
            start_idx = 1 if first_row_second == first_row_first else 0
            merged_rows.extend(second_rows[start_idx:])
        
        return {
            'bbox': first_table.get('bbox'),
            'rows': merged_rows,
            'col_count': first_table.get('col_count', 0),
            'is_merged': True,
            'source_pages': [table_group['start_page'], table_group['end_page']]
        }

    async def _translate_merged_tables(
        self,
        multi_page_tables: List[Dict[str, Any]],
        target_lang: str,
        source_lang: Optional[str]
    ) -> Dict[int, Dict[str, Any]]:
        """Translate all merged multi-page tables.
        
        Returns a dict mapping start_page to translated merged table.
        """
        translated_tables: Dict[int, Dict[str, Any]] = {}
        
        for table_group in multi_page_tables:
            if table_group.get('merged'):
                continue
            
            merged_table = self._merge_multi_page_table(table_group)
            
            # Translate the merged table
            translated_table = await self._translate_table(
                merged_table, target_lang, source_lang
            )
            
            translated_tables[table_group['start_page']] = {
                'table': translated_table,
                'end_page': table_group['end_page'],
                'original_first': table_group['first_table'],
                'original_second': table_group['second_table']
            }
            
            table_group['merged'] = True
            
        return translated_tables

    # ===== Table Structure Preservation =====

    def _extract_tables_from_page(self, page: "fitz.Page") -> List[Dict[str, Any]]:
        """Extract table structures from PDF page."""
        tables = []
        try:
            table_finder = page.find_tables()
            for table in table_finder:
                table_data = {
                    "bbox": table.bbox,
                    "rows": [],
                    "col_count": len(table.header.cells) if table.header else 0,
                }
                for row in table.extract():
                    table_data["rows"].append([cell or "" for cell in row])
                if table_data["rows"]:
                    tables.append(table_data)
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")
        return tables

    async def _translate_table(
        self,
        table_data: Dict[str, Any],
        target_language: str,
        source_language: Optional[str],
    ) -> Dict[str, Any]:
        """Translate table cells while preserving structure."""
        translated_rows = []
        for row in table_data.get("rows", []):
            translated_row = []
            for cell in row:
                if cell and cell.strip() and not self._is_fixed_line(cell):
                    cache_key = self._get_cache_key(cell, target_language, source_language)
                    cached = self._cache_get(cache_key)
                    if cached:
                        translated_row.append(cached)
                    else:
                        result, _ = await self.translate_text(cell, target_language, source_language)
                        translated_row.append(result.strip())
                        self._cache_set(cache_key, result.strip(), persist=True,
                                       source_text=cell, target_lang=target_language, 
                                       source_lang=source_language or "")
                else:
                    translated_row.append(cell)
            translated_rows.append(translated_row)

        return {
            "bbox": table_data.get("bbox"),
            "rows": translated_rows,
            "col_count": table_data.get("col_count", 0),
        }

    def _render_table_to_page(
        self,
        page: "fitz.Page",
        table_data: Dict[str, Any],
        target_language: str,
    ) -> None:
        """Render translated table back to PDF page."""
        bbox = table_data.get("bbox")
        if not bbox:
            return
        rows = table_data.get("rows", [])
        if not rows:
            return

        x0, y0, x1, y1 = bbox
        row_height = (y1 - y0) / len(rows) if rows else 20
        col_count = table_data.get("col_count", 1) or 1
        col_width = (x1 - x0) / col_count
        font_path = self._get_font_for_language(target_language)
        font_size = min(10, row_height * 0.7)

        for row_idx, row in enumerate(rows):
            for col_idx, cell in enumerate(row[:col_count]):
                cell_x = x0 + col_idx * col_width + 2
                cell_y = y0 + row_idx * row_height + font_size + 2
                try:
                    page.insert_text(
                        (cell_x, cell_y),
                        str(cell)[:50],
                        fontsize=font_size,
                        fontfile=font_path,
                    )
                except Exception:
                    page.insert_text((cell_x, cell_y), str(cell)[:50], fontsize=font_size, fontname="helv")

    # ===== Language Detection =====

    async def detect_language(self, text: str) -> str:
        """Auto-detect source language using Gemini."""
        if not self.client:
            return "unknown"

        sample = text[:1000]
        prompt = f"""Detect the language of this text. Reply with ONLY the language name in lowercase English (e.g., 'russian', 'chinese', 'arabic', 'persian', 'english', 'spanish').

Text: {sample}

Language:"""

        try:
            from google.genai import types
            response = await asyncio.to_thread(
                self.client.models.generate_content,
                model=self.model_name,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(thinking_level="low"),
                    max_output_tokens=50,
                )
            )
            detected = response.text.strip().lower().split()[0].rstrip('.,;:') if response.text else "unknown"
            logger.info(f"Detected language: {detected}")
            return detected
        except Exception as e:
            logger.warning(f"Language detection failed: {e}")
            return "unknown"

    # ===== Translation with Cache =====

    async def _translate_chunk_with_cache(
        self,
        chunk: List[Tuple[int, str]],
        target_language: str,
        source_language: Optional[str] = None,
    ) -> Dict[int, str]:
        """Translate chunk with caching for individual lines."""
        result: Dict[int, str] = {}
        to_translate: List[Tuple[int, str]] = []

        # Check cache for each line
        for idx, text in chunk:
            cache_key = self._get_cache_key(text, target_language, source_language)
            cached = self._cache_get(cache_key)
            if cached is not None:
                result[idx] = cached
            else:
                to_translate.append((idx, text))

        # Translate uncached lines
        if to_translate:
            translated = await self._translate_chunk(to_translate, target_language, source_language)
            for idx, translation in translated.items():
                result[idx] = translation
                # Cache the result with persistent storage
                original_text = next((t for i, t in chunk if i == idx), None)
                if original_text:
                    cache_key = self._get_cache_key(original_text, target_language, source_language)
                    self._cache_set(cache_key, translation, persist=True,
                                   source_text=original_text, target_lang=target_language,
                                   source_lang=source_language or "")

        return result

    def extract_text_with_metadata(
        self,
        file_path: str,
        filename: str,
        mime_type: str,
        ocr_languages: Optional[str] = None,
    ) -> Tuple[str, Dict[str, Optional[str]]]:
        """Extract text and metadata using the offline parser."""
        try:
            content = Path(file_path).read_bytes()
        except Exception as e:
            raise ValueError(f"Failed to read file: {e}")

        encoded = base64.b64encode(content).decode("utf-8")
        parsed = self.parser.parse_document(
            filename=filename,
            content_base64=encoded,
            content_type=mime_type,
            ocr_languages=ocr_languages,
        )

        if not parsed.full_text or not parsed.full_text.strip():
            errors = "; ".join(parsed.parse_errors) if parsed.parse_errors else "No text extracted"
            raise ValueError(errors)

        metadata = {
            "ocr_used": bool(parsed.metadata.get("ocr_used")),
            "page_count": parsed.total_pages,
            "chars_extracted": parsed.total_chars,
            "document_type": parsed.document_type.value,
        }

        return parsed.full_text, metadata

    async def translate_pdf_preserve_layout(
        self,
        file_path: str,
        target_language: str,
        source_language: Optional[str] = None,
        ocr_languages: Optional[str] = None,
        progress_callback: Optional[ProgressCallback] = None,
        translate_images: bool = False,
    ) -> Tuple[str, Dict[str, Any]]:
        """Translate a PDF and preserve layout by replacing text blocks.
        
        Args:
            file_path: Path to source PDF
            target_language: Target language for translation
            source_language: Source language (auto-detected if None)
            ocr_languages: Tesseract language codes for scanned pages
            progress_callback: Progress callback function
            translate_images: If True, OCR and translate text within embedded images
        """
        if not self._has_pymupdf:
            raise ValueError("PyMuPDF not available for PDF translation")
        if not self.client:
            raise ValueError("Gemini API key not configured or google-genai not installed")

        src_doc = fitz.open(file_path)
        page_count = len(src_doc)
        out_doc = fitz.open()
        ocr_used = False
        chars_extracted = 0
        chars_translated = 0
        total_skipped_overlaps = 0
        image_text_regions_translated = 0
        
        # Progress tracking: total_steps = 3 (prep) + page_count (pages) + 1 (save)
        total_steps = page_count + 4
        current_step = 0
        
        # Progress: Step 1 - Font extraction
        if progress_callback:
            progress_callback(current_step, total_steps, "Extracting fonts from document...")
        
        # Large Effort #1: Extract all fonts from source document
        font_mapping = self._extract_document_fonts(src_doc)
        exact_fonts_used = 0
        logger.info(f"Extracted {len(font_mapping)} fonts from source document")
        current_step += 1
        
        # Progress: Step 2 - Multi-page table detection
        if progress_callback:
            progress_callback(current_step, total_steps, "Detecting multi-page tables...")
        
        # Large Effort #2: Detect multi-page tables (tables spanning page boundaries)
        multi_page_tables = self._detect_multi_page_tables(src_doc)
        tables_stitched = len(multi_page_tables)
        merged_table_translations: Dict[int, Dict[str, Any]] = {}
        merged_table_pages: set = set()  # Pages that have merged table parts
        
        if tables_stitched > 0:
            logger.info(f"Detected {tables_stitched} multi-page tables for stitching")
            if progress_callback:
                progress_callback(current_step, total_steps, f"Translating {tables_stitched} multi-page tables...")
            # Translate merged tables
            merged_table_translations = await self._translate_merged_tables(
                multi_page_tables, target_language, source_language
            )
            # Track which pages are part of merged tables
            for table_group in multi_page_tables:
                for pg in range(table_group['start_page'], table_group['end_page'] + 1):
                    merged_table_pages.add(pg)
        current_step += 1
        
        # Progress: Step 3 - Header/footer detection
        if progress_callback:
            progress_callback(current_step, total_steps, "Detecting headers and footers...")
        
        # Quick Win #1: Detect repeated blocks (headers/footers) across all pages
        repeated_hashes = self._detect_repeated_blocks(src_doc, min_occurrences=3)
        current_step += 1
        
        # Quick Win tracking counters
        headers_footers_skipped = 0
        code_blocks_skipped = 0
        links_preserved = 0
        
        # Medium effort: Multi-column detection
        multi_column_pages = 0

        for page_index in range(len(src_doc)):
            # Progress: Per-page update
            if progress_callback:
                progress_callback(
                    current_step + page_index, 
                    total_steps, 
                    f"Translating page {page_index + 1} of {page_count}..."
                )
            out_doc.insert_pdf(src_doc, from_page=page_index, to_page=page_index)
            out_page = out_doc[-1]
            page = src_doc[page_index]

            blocks = self._extract_text_blocks(page)
            if not blocks:
                ocr_blocks = self._extract_ocr_lines(page, ocr_languages)
                if ocr_blocks:
                    blocks = ocr_blocks
                    ocr_used = True

            if not blocks:
                continue

            # Get image regions to avoid covering text that overlaps images
            image_regions = self._get_image_regions(page)
            
            # Preserve images before covering text blocks (for fallback reinsertion)
            page_images = self._extract_page_images(out_page)
            
            # Quick Win #4: Extract hyperlinks for preservation
            page_links = self._extract_page_links(page)

            # Filter out text blocks that overlap with images
            # (they're likely captions or watermarks - leave them untouched)
            translatable_blocks = []
            skipped_image_overlaps = 0
            for block_id, block in enumerate(blocks):
                if len(block) == 3:
                    rect, text, format_meta = block
                else:
                    rect, text = block[:2]
                    format_meta = None
                
                if self._rect_overlaps_images(rect, image_regions, threshold=0.3):
                    # Skip text blocks significantly overlapping images
                    skipped_image_overlaps += 1
                    logger.debug(f"Skipping text block overlapping image: {text[:50]}...")
                    continue
                
                # Quick Win #1: Skip repeated headers/footers
                if self._is_repeated_block(text, repeated_hashes):
                    headers_footers_skipped += 1
                    logger.debug(f"Skipping repeated header/footer: {text[:50]}...")
                    continue
                
                # Quick Win #2: Skip code blocks (monospace font or code patterns)
                if self._is_code_block(text, format_meta):
                    code_blocks_skipped += 1
                    logger.debug(f"Skipping code block: {text[:50]}...")
                    continue
                    
                translatable_blocks.append((block_id, rect, text, format_meta))
            
            if skipped_image_overlaps > 0:
                logger.info(f"Page {page_index + 1}: Skipped {skipped_image_overlaps} text blocks overlapping images")
                total_skipped_overlaps += skipped_image_overlaps

            # Medium #1: Sort blocks in reading order (multi-column detection)
            if translatable_blocks:
                # Extract just (rect, text, format_meta) for sorting
                blocks_for_sorting = [(rect, text, fmt) for _, rect, text, fmt in translatable_blocks]
                sorted_blocks = self._sort_blocks_reading_order(blocks_for_sorting)
                
                # Detect column count for metadata
                columns = self._detect_columns(blocks_for_sorting)
                if len(columns) > 1:
                    logger.info(f"Page {page_index + 1}: Detected {len(columns)} columns")
                    multi_column_pages += 1
                
                # Rebuild translatable_blocks with sorted order
                translatable_blocks = [
                    (i, rect, text, fmt) for i, (rect, text, fmt) in enumerate(sorted_blocks)
                ]

            # Translate only non-overlapping blocks
            blocks_for_translation = [(rect, text, fmt) for _, rect, text, fmt in translatable_blocks]
            block_map = await self._translate_blocks(blocks_for_translation, target_language, source_language)
            
            for idx, (block_id, rect, text, format_meta) in enumerate(translatable_blocks):
                translated = block_map.get(idx, text)
                if not translated:
                    continue

                chars_extracted += len(text)
                chars_translated += len(translated)

                # Cover original text and insert translation with preserved formatting
                out_page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                
                # Large Effort #1: Try exact font matching
                original_font = format_meta.get("font_name", "") if format_meta else ""
                if font_mapping and original_font and original_font in font_mapping:
                    exact_fonts_used += 1
                
                self._insert_text_fit(out_page, rect, translated, target_language, format_meta, font_mapping)

            # ==== Table Translation and Rendering ====
            # Extract tables from current page
            page_tables = self._extract_tables_from_page(page)
            
            if page_tables:
                # Check if this page has merged multi-page table
                if page_index in merged_table_translations:
                    # This is the START of a multi-page table - render merged translation
                    merged_info = merged_table_translations[page_index]
                    translated_table = merged_info['table']
                    original_table = merged_info.get('original_first', page_tables[0] if page_tables else None)
                    
                    if original_table and 'rect' in original_table:
                        self._render_table_to_page(
                            out_page, translated_table, original_table['rect'], target_language, font_mapping
                        )
                        logger.info(f"Page {page_index + 1}: Rendered merged multi-page table")
                else:
                    # Regular page - translate and render tables normally
                    is_merged_continuation = any(
                        table_group['end_page'] == page_index and table_group['start_page'] != page_index
                        for table_group in multi_page_tables
                    )
                    
                    if not is_merged_continuation:
                        for table_data in page_tables:
                            if 'rows' in table_data and table_data['rows']:
                                translated_table = await self._translate_table(
                                    table_data, target_language, source_language
                                )
                                self._render_table_to_page(
                                    out_page, translated_table, table_data.get('rect', rect), 
                                    target_language, font_mapping
                                )
                    else:
                        logger.debug(f"Page {page_index + 1}: Skipping continuation part of merged table")

            # Translate text within images if enabled
            if translate_images and page_images:
                ocr_lang = ocr_languages or "eng+rus"
                page_images, img_regions = await self._translate_page_images(
                    page_images,
                    target_language,
                    source_language,
                    ocr_lang,
                )
                image_text_regions_translated += img_regions
                if img_regions > 0:
                    logger.info(f"Page {page_index + 1}: Translated {img_regions} text regions in images")

            # Re-insert images (translated if enabled, otherwise preserved originals)
            self._reinsert_images(out_page, page_images)
            
            # Quick Win #4: Re-insert hyperlinks
            page_links_restored = self._reinsert_links(out_page, page_links)
            links_preserved += page_links_restored
            if page_links_restored > 0:
                logger.debug(f"Page {page_index + 1}: Restored {page_links_restored} hyperlinks")

        # Progress: Saving file
        if progress_callback:
            progress_callback(total_steps - 1, total_steps, "Saving translated document...")
        
        output_path = str(Path(file_path).with_suffix("")) + "_translated.pdf"
        out_doc.save(output_path)
        out_doc.close()
        src_doc.close()
        
        # Progress: Complete!
        if progress_callback:
            progress_callback(total_steps, total_steps, "Translation complete!")

        metadata = {
            "ocr_used": ocr_used,
            "page_count": page_count,
            "chars_extracted": chars_extracted,
            "chars_translated": chars_translated,
            "blocks_skipped_image_overlap": total_skipped_overlaps,
            "image_text_regions_translated": image_text_regions_translated,
            # Quick Win metrics
            "headers_footers_skipped": headers_footers_skipped,
            "code_blocks_skipped": code_blocks_skipped,
            "links_preserved": links_preserved,
            # Medium effort metrics
            "multi_column_pages": multi_column_pages,
            "persistent_cache_hits": getattr(self, '_persistent_cache_hits', 0),
            # Large effort metrics
            "exact_fonts_used": exact_fonts_used,
            "fonts_extracted": len(font_mapping),
            "tables_stitched": tables_stitched,
        }
        return output_path, metadata

    async def translate_text(
        self,
        text: str,
        target_language: str,
        source_language: Optional[str] = None,
        progress_callback: Optional[ProgressCallback] = None,
    ) -> Tuple[str, Dict[str, Any]]:
        """Translate text line-by-line, ensuring full coverage."""
        if not self.client:
            raise ValueError("Gemini API key not configured or google-genai not installed")

        lines = text.splitlines()
        if not lines:
            return "", {"chunks": 0, "warnings": []}

        fixed_indices = {i for i, line in enumerate(lines) if self._is_fixed_line(line)}
        translate_items = [
            (i, lines[i]) for i in range(len(lines))
            if i not in fixed_indices and lines[i].strip()
        ]

        chunks = self._chunk_lines(translate_items)
        translated_lines = list(lines)
        warnings = []

        # Auto-detect language if not provided
        if not source_language:
            if progress_callback:
                progress_callback(0, len(chunks) + 1, "Detecting source language...")
            source_language = await self.detect_language(text)

        # Parallel translation with semaphore
        semaphore = asyncio.Semaphore(self.MAX_PARALLEL_CHUNKS)

        async def translate_with_progress(chunk_idx: int, chunk_data: List[Tuple[int, str]]) -> Tuple[int, Dict[int, str]]:
            async with semaphore:
                if progress_callback:
                    progress_callback(chunk_idx + 1, len(chunks) + 1, f"Translating chunk {chunk_idx + 1}/{len(chunks)}...")
                result = await self._translate_chunk_with_cache(chunk_data, target_language, source_language)
                return chunk_idx, result

        results = await asyncio.gather(*[
            translate_with_progress(i, chunk) for i, chunk in enumerate(chunks)
        ])

        # Apply results
        for chunk_idx, mapping in results:
            for idx, translation in mapping.items():
                translated_lines[idx] = translation
            chunk = chunks[chunk_idx]
            missing = [idx for idx, _ in chunk if idx not in mapping]
            if missing:
                warnings.append(f"Missing translations for lines: {missing[:5]}")

        if progress_callback:
            progress_callback(len(chunks) + 1, len(chunks) + 1, "Translation complete")

        translated_text = "\n".join(translated_lines)
        return translated_text, {"chunks": len(chunks), "warnings": warnings}

    def _is_fixed_line(self, line: str) -> bool:
        """Detect lines that should not be translated (page markers, separators)."""
        stripped = line.strip()
        if not stripped:
            return True
        if stripped.startswith("--- Page") or stripped.startswith("--- Slide"):
            return True
        if stripped.startswith("=== ") or stripped == "---":
            return True
        return False

    def _chunk_lines(self, items: List[Tuple[int, str]]) -> List[List[Tuple[int, str]]]:
        """Chunk line items to fit within token limits with context overlap."""
        chunks: List[List[Tuple[int, str]]] = []
        current: List[Tuple[int, str]] = []
        current_chars = 0
        prev_context: List[Tuple[int, str]] = []  # For context overlap

        for idx, text in items:
            line_len = len(text)
            if (
                current
                and (len(current) >= self.MAX_LINES_PER_CHUNK
                     or current_chars + line_len > self.MAX_CHARS_PER_CHUNK)
            ):
                chunks.append(current)
                # Save last N lines for context in next chunk's translation prompt
                prev_context = current[-self.CONTEXT_OVERLAP_LINES:] if len(current) >= self.CONTEXT_OVERLAP_LINES else current[:]
                current = []
                current_chars = 0

            current.append((idx, text))
            current_chars += line_len

        if current:
            chunks.append(current)

        # Store context for reference (available via _last_chunk_context)
        self._last_chunk_context = prev_context
        return chunks

    async def _translate_blocks(
        self,
        blocks: List[Tuple["fitz.Rect", str]],
        target_language: str,
        source_language: Optional[str],
    ) -> Dict[int, str]:
        """Translate blocks in parallel with caching."""
        items = []
        for idx, block in enumerate(blocks):
            # Handle both 2-tuple (legacy) and 3-tuple (with format_meta) blocks
            text = block[1] if len(block) >= 2 else ""
            safe_text = text.replace("\n", "\\n")
            items.append((idx, safe_text))

        chunks = self._chunk_lines(items)

        # Parallel translation with semaphore
        semaphore = asyncio.Semaphore(self.MAX_PARALLEL_CHUNKS)

        async def translate_chunk_limited(chunk):
            async with semaphore:
                return await self._translate_chunk_with_cache(chunk, target_language, source_language)

        results = await asyncio.gather(*[translate_chunk_limited(c) for c in chunks])

        mapping: Dict[int, str] = {}
        for result in results:
            for idx, translated in result.items():
                mapping[idx] = translated.replace("\\n", "\n")
        return mapping

    async def _translate_chunk(
        self,
        chunk: List[Tuple[int, str]],
        target_language: str,
        source_language: Optional[str] = None,
    ) -> Dict[int, str]:
        """Translate a chunk of lines and return mapping of index -> translation.
        
        ROBUST: Never raises exceptions - falls back to original text on failure.
        """
        if not chunk:
            return {}

        # Build fallback mapping (original text) in case of total failure
        fallback_mapping = {idx: text for idx, text in chunk}

        try:
            lines_text = "\n".join([f"{idx}|{text}" for idx, text in chunk])
        except Exception as e:
            logger.error(f"Failed to prepare chunk text: {e}")
            return fallback_mapping

        prompt = self._build_translation_prompt(lines_text, target_language, source_language)
        last_error = None

        for attempt in range(1, self.MAX_RETRIES + 1):
            try:
                from google.genai import types

                # API call with error handling
                try:
                    response = await asyncio.wait_for(
                        asyncio.to_thread(
                            self.client.models.generate_content,
                            model=self.model_name,
                            contents=prompt,
                            config=types.GenerateContentConfig(
                                thinking_config=types.ThinkingConfig(thinking_level="low"),
                                max_output_tokens=self.MAX_OUTPUT_TOKENS,
                            )
                        ),
                        timeout=60.0  # 60 second timeout per chunk
                    )
                except asyncio.TimeoutError:
                    logger.warning(f"Translation timeout on attempt {attempt}")
                    last_error = "API timeout"
                    if attempt < self.MAX_RETRIES:
                        await asyncio.sleep(self._get_retry_delay(attempt))
                        continue
                    break

                # Handle empty or missing response
                if not response or not hasattr(response, 'text') or not response.text:
                    logger.warning(f"Empty response from API on attempt {attempt}")
                    last_error = "Empty API response"
                    if attempt < self.MAX_RETRIES:
                        continue
                    break

                # Parse response with robust parser
                data = self._parse_json_response(response.text)
                mapping = self._normalize_translation_mapping(data)

                # Validate translations to detect hallucinations
                for idx, text in chunk:
                    if idx in mapping:
                        try:
                            valid, reason = self._validate_translation(text, mapping[idx])
                            if not valid:
                                logger.warning(f"Line {idx} validation failed: {reason}")
                                # Fall back to original text if validation fails
                                mapping[idx] = text
                        except Exception as e:
                            logger.debug(f"Validation error for line {idx}: {e}")
                            # Keep the translation if validation itself fails

                # Check for missing lines
                if self._has_missing_lines(chunk, mapping):
                    missing_count = sum(1 for idx, _ in chunk if idx not in mapping)
                    logger.info(f"Missing {missing_count} translations, attempt {attempt}")

                    if attempt < self.MAX_RETRIES:
                        prompt = self._build_retry_prompt(lines_text, target_language, source_language)
                        continue

                    # Fill in missing with original text
                    for idx, text in chunk:
                        if idx not in mapping:
                            mapping[idx] = text
                            logger.debug(f"Using original text for line {idx}")

                return mapping

            except ImportError as e:
                logger.error(f"Missing dependency for translation: {e}")
                return fallback_mapping

            except Exception as e:
                last_error = str(e)
                error_str = str(e).lower()
                
                # Check for rate limit errors (429) - need longer backoff
                is_rate_limit = (
                    '429' in error_str or 
                    'rate limit' in error_str or 
                    'quota' in error_str or
                    'too many requests' in error_str or
                    'resource_exhausted' in error_str
                )
                
                if is_rate_limit:
                    # Use much longer backoff for rate limits: 60s, 120s, 240s, max 300s
                    wait_time = min(60 * (2 ** (attempt - 1)), 300)
                    logger.warning(f"Rate limited! Waiting {wait_time}s before retry {attempt + 1}")
                    await asyncio.sleep(wait_time)
                    if attempt < self.MAX_RETRIES:
                        continue
                else:
                    logger.warning(f"Translation attempt {attempt} failed: {e}")
                    if attempt < self.MAX_RETRIES:
                        delay = self._get_retry_delay(attempt)
                        logger.info(f"Retrying translation in {delay:.1f}s...")
                        await asyncio.sleep(delay)

        # All retries exhausted - return fallback
        logger.error(f"All translation attempts failed. Last error: {last_error}. Using original text.")
        return fallback_mapping

    def _build_translation_prompt(
        self,
        lines_text: str,
        target_language: str,
        source_language: Optional[str],
    ) -> str:
        source_hint = f"Source language hint: {source_language}\n" if source_language else ""
        glossary_prompt = self._build_glossary_prompt()
        return f"""You are a LITERAL TRANSLATOR. Your ONLY job is to translate text from one language to another.

{source_hint}{glossary_prompt}## STRICT RULES - VIOLATION MEANS FAILURE:

1. **TRANSLATE LITERALLY** - Word-for-word translation preserving meaning
2. **NEVER SUMMARIZE** - If input has 10 sentences, output MUST have 10 sentences
3. **NEVER PARAPHRASE** - Don't "improve" or rephrase the original
4. **NEVER ADD CONTENT** - No explanations, notes, or commentary
5. **NEVER OMIT CONTENT** - Every word in input must appear in output
6. **PRESERVE ALL:**
   - Numbers exactly as they appear (123.45, 1,000, etc.)
   - Dates in their original format (2024-01-15, 15/01/2024)
   - Times (14:30, 2:30 PM)
   - URLs and emails unchanged
   - Codes and identifiers (ABC-123, ID_456)
   - Currency symbols and amounts ($100, €50.00)
   - Percentages (50%, 0.5%)
   - Punctuation and special characters
   - Line breaks marked as \\n

7. **OUTPUT FORMAT:** Return ONLY a JSON array:
   [{{{"id": <number>, "translation": "<translated_text>"}}}, ...]

## EXAMPLES OF VIOLATIONS (DO NOT DO THESE):
- Input: "The meeting is at 2:30 PM on 2024-01-15"
- WRONG: "There's a meeting scheduled" (summarized)
- WRONG: "The gathering occurs at half past two" (paraphrased)
- CORRECT: "La reunión es a las 2:30 PM el 2024-01-15"

## TRANSLATE THESE LINES:
{lines_text}
"""

    def _build_retry_prompt(
        self,
        lines_text: str,
        target_language: str,
        source_language: Optional[str],
    ) -> str:
        source_hint = f"Source language hint: {source_language}\n" if source_language else ""
        return f"""Retry translation. You missed some lines previously.

{source_hint}Translate EVERY line below to {target_language}. No omissions.

LINES:
{lines_text}

Return ONLY JSON array: [{{"id": <number>, "translation": "<text>"}}]
"""

    def _parse_json_response(self, response_text: str) -> object:
        """Parse JSON from AI response with multiple fallback strategies."""
        if not response_text:
            logger.warning("Empty response from translation API")
            return []

        text = response_text.strip()

        # Strategy 1: Strip markdown code blocks
        if text.startswith("```json"):
            text = text[7:]
        elif text.startswith("```"):
            text = text[3:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()

        # Strategy 2: Direct JSON parse
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

        # Strategy 3: Find JSON array in the response
        try:
            start = text.find('[')
            end = text.rfind(']')
            if start != -1 and end > start:
                json_substr = text[start:end + 1]
                return json.loads(json_substr)
        except json.JSONDecodeError:
            pass

        # Strategy 4: Try to fix common JSON issues
        try:
            # Fix unescaped quotes in values
            fixed = re.sub(r'(?<="translation":\s*")([^"]*?)(?<!\\)"(?=[^,}]*")', r'\1\\"', text)
            # Fix trailing commas
            fixed = re.sub(r',\s*([\]}])', r'\1', fixed)
            return json.loads(fixed)
        except (json.JSONDecodeError, re.error):
            pass

        # Strategy 5: Line-by-line extraction for severely malformed responses
        try:
            results = []
            # Match patterns like: {"id": 0, "translation": "text"}
            pattern = r'\{\s*["\']?id["\']?\s*:\s*(\d+)\s*,\s*["\']?translation["\']?\s*:\s*["\']([^"\'}]+)["\']\s*\}'
            for match in re.finditer(pattern, text, re.IGNORECASE):
                results.append({"id": int(match.group(1)), "translation": match.group(2)})
            if results:
                logger.info(f"Recovered {len(results)} translations via regex extraction")
                return results
        except Exception as e:
            logger.debug(f"Regex extraction failed: {e}")

        logger.error(f"All JSON parsing strategies failed. Response preview: {text[:200]}...")
        return []

    def _normalize_translation_mapping(self, data: object) -> Dict[int, str]:
        """Normalize parsed JSON to index->translation mapping with fallbacks."""
        mapping: Dict[int, str] = {}

        if data is None:
            return mapping

        # Handle list of translation objects (expected format)
        if isinstance(data, list):
            for item in data:
                try:
                    if not isinstance(item, dict):
                        continue

                    # Try multiple field names for ID
                    idx = item.get("id") or item.get("ID") or item.get("index") or item.get("line")
                    # Try multiple field names for translation
                    translation = (
                        item.get("translation") or
                        item.get("Translation") or
                        item.get("text") or
                        item.get("translated") or
                        item.get("result")
                    )

                    if idx is None or translation is None:
                        continue

                    # Sanitize the translation
                    idx_int = int(idx)
                    trans_str = str(translation).strip()

                    # Skip obviously broken translations
                    if not trans_str or trans_str in ('null', 'None', 'undefined'):
                        continue

                    mapping[idx_int] = trans_str
                except (ValueError, TypeError, AttributeError) as e:
                    logger.debug(f"Skipping malformed item: {item} - {e}")
                    continue

        # Handle dict with numeric keys (alternative format)
        elif isinstance(data, dict):
            for key, value in data.items():
                try:
                    idx = int(key)
                    trans = str(value).strip() if value else ""
                    if trans and trans not in ('null', 'None', 'undefined'):
                        mapping[idx] = trans
                except (ValueError, TypeError):
                    continue

        return mapping

    def _has_missing_lines(self, chunk: List[Tuple[int, str]], mapping: Dict[int, str]) -> bool:
        return any(idx not in mapping for idx, _ in chunk)

    def _extract_text_blocks(self, page: "fitz.Page") -> List[Tuple["fitz.Rect", str, Dict[str, Any]]]:
        """Extract text blocks from a PDF page with formatting metadata."""
        blocks = []
        
        # Get detailed text info with font properties
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:  # Skip images
                continue
            
            block_bbox = block.get("bbox", (0, 0, 0, 0))
            block_text_parts = []
            
            # Collect formatting info from spans
            font_sizes = []
            font_names = []
            font_colors = []
            font_flags = []
            
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        block_text_parts.append(text)
                        font_sizes.append(span.get("size", 12))
                        font_names.append(span.get("font", ""))
                        font_colors.append(span.get("color", 0))  # int RGB
                        font_flags.append(span.get("flags", 0))
            
            if not block_text_parts:
                continue
            
            full_text = " ".join(block_text_parts)
            rect = fitz.Rect(block_bbox)
            
            # Determine dominant formatting (most common)
            format_meta = {
                "font_size": max(set(font_sizes), key=font_sizes.count) if font_sizes else 12,
                "font_name": max(set(font_names), key=font_names.count) if font_names else "",
                "color": max(set(font_colors), key=font_colors.count) if font_colors else 0,
                "flags": max(set(font_flags), key=font_flags.count) if font_flags else 0,
                "is_bold": any(f & 16 for f in font_flags),  # Flag 16 = bold
                "is_italic": any(f & 2 for f in font_flags),  # Flag 2 = italic
            }
            
            blocks.append((rect, full_text, format_meta))
        
        # Fallback to simple extraction if dict method fails
        if not blocks:
            for block in page.get_text("blocks"):
                x0, y0, x1, y1, text, _, block_type = block
                if block_type != 0:
                    continue
                if not text or not text.strip():
                    continue
                rect = fitz.Rect(x0, y0, x1, y1)
                blocks.append((rect, text.strip(), {"font_size": 12, "color": 0, "is_bold": False, "is_italic": False}))
        
        return blocks

    def _extract_ocr_lines(
        self,
        page: "fitz.Page",
        ocr_languages: Optional[str],
    ) -> List[Tuple["fitz.Rect", str]]:
        """Extract OCR line blocks when a PDF page has no text layer."""
        if not self._has_ocr:
            return []
        try:
            pix = page.get_pixmap(dpi=300)
            mode = "RGB" if pix.alpha == 0 else "RGBA"
            image = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            if mode == "RGBA":
                image = image.convert("RGB")

            lang = ocr_languages or "eng"
            data = pytesseract.image_to_data(image, lang=lang, output_type=Output.DICT)

            lines = {}
            for i in range(len(data["text"])):
                text = data["text"][i].strip()
                if not text:
                    continue
                key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
                x, y, w, h = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
                if key not in lines:
                    lines[key] = {"words": [], "bbox": [x, y, x + w, y + h]}
                lines[key]["words"].append(text)
                bbox = lines[key]["bbox"]
                bbox[0] = min(bbox[0], x)
                bbox[1] = min(bbox[1], y)
                bbox[2] = max(bbox[2], x + w)
                bbox[3] = max(bbox[3], y + h)

            blocks = []
            for _, info in lines.items():
                text_line = " ".join(info["words"]).strip()
                if not text_line:
                    continue
                rect = fitz.Rect(*info["bbox"])
                blocks.append((rect, text_line))

            return blocks
        except Exception as e:
            logger.warning(f"OCR line extraction failed: {e}")
            return []

    def _get_font_for_language(self, target_language: str, is_bold: bool = False, is_italic: bool = False) -> str:
        """Get the appropriate font path for the target language and style.
        
        Args:
            target_language: Target language for translation
            is_bold: Whether to use bold variant
            is_italic: Whether to use italic variant
            
        Returns:
            Path to the appropriate font file
        """
        lang_lower = target_language.lower()
        script = self.LANGUAGE_SCRIPT_MAP.get(lang_lower, 'default')
        
        # Build font key with style variant
        font_key = script
        if is_bold and is_italic:
            # Try bold-italic first, fall back to bold, then regular
            if f"{script}_bolditalic" in self.FONT_MAP:
                font_key = f"{script}_bolditalic"
            elif f"{script}_bold" in self.FONT_MAP:
                font_key = f"{script}_bold"
        elif is_bold:
            if f"{script}_bold" in self.FONT_MAP:
                font_key = f"{script}_bold"
        elif is_italic:
            if f"{script}_italic" in self.FONT_MAP:
                font_key = f"{script}_italic"
        
        font_path = self.FONT_MAP.get(font_key, self.FONT_MAP.get(script, self.FONT_MAP['default']))
        
        # Verify font exists, fall back to regular if not
        if not Path(font_path).exists():
            font_path = self.FONT_MAP.get(script, self.FONT_MAP['default'])
            if not Path(font_path).exists():
                font_path = self.FONT_MAP['default']
        
        return font_path

    def _is_rtl_language(self, target_language: str) -> bool:
        """Check if target language is RTL."""
        return target_language.lower() in self.RTL_LANGUAGES

    def _insert_text_fit(
        self, 
        page: "fitz.Page", 
        rect: "fitz.Rect", 
        text: str, 
        target_language: str = "english",
        format_meta: Optional[Dict[str, Any]] = None,
        font_mapping: Optional[Dict[str, str]] = None
    ) -> None:
        """Insert text into a rectangle with proper Unicode font support and original formatting.
        
        Args:
            page: PyMuPDF page to insert text into
            rect: Rectangle bounding box for the text
            text: Text to insert
            target_language: Target language for font selection
            format_meta: Original formatting metadata (font_name, font_size, color, etc.)
            font_mapping: Optional mapping of original font names to extracted font paths
        """
        # Use original font size if available, otherwise calculate from rect
        if format_meta and format_meta.get("font_size"):
            font_size = format_meta["font_size"]
        else:
            font_size = min(self.PDF_MAX_FONT, rect.height * 0.7)
        font_size = max(font_size, self.PDF_MIN_FONT)
        font_size = min(font_size, self.PDF_MAX_FONT)

        # Detect bold/italic from font flags
        # Font flags: bit 0=superscript, bit 1=italic, bit 2=serifed, bit 3=script, bit 4=bold
        is_bold = False
        is_italic = False
        if format_meta and format_meta.get("font_flags"):
            flags = format_meta["font_flags"]
            is_bold = bool(flags & 16)   # bit 4
            is_italic = bool(flags & 2)  # bit 1

        # Try exact font matching first
        font_path = None
        original_font_name = format_meta.get("font_name", "") if format_meta else ""
        
        if font_mapping and original_font_name:
            font_path = self._get_font_path_for_block(format_meta, font_mapping, target_language)
        
        # Fall back to language-appropriate font if no exact match
        if not font_path:
            font_path = self._get_font_for_language(target_language, is_bold=is_bold, is_italic=is_italic)
        
        font_name = "translation_font"
        
        try:
            # Try to load custom font for better Unicode support
            if Path(font_path).exists():
                font = fitz.Font(fontfile=font_path)
                page.insert_font(fontname=font_name, fontbuffer=font.buffer)
            else:
                # Fallback to built-in font
                font_name = "helv"
        except Exception as e:
            logger.warning(f"Failed to load font {font_path}: {e}, using default")
            font_name = "helv"

        # RTL alignment
        align = 2 if self._is_rtl_language(target_language) else 0
        
        # Parse original color (int to RGB tuple)
        if format_meta and format_meta.get("color"):
            color_int = format_meta["color"]
            # Convert integer color to RGB tuple (0-1 range)
            r = ((color_int >> 16) & 0xFF) / 255.0
            g = ((color_int >> 8) & 0xFF) / 255.0
            b = (color_int & 0xFF) / 255.0
            color = (r, g, b)
        else:
            color = (0, 0, 0)  # Default black

        while font_size >= self.PDF_MIN_FONT:
            lines = self._wrap_text_simple(text, rect.width, font_size, font_name)
            total_height = len(lines) * font_size * self.PDF_LINE_HEIGHT
            if total_height <= rect.height:
                page.insert_textbox(
                    rect,
                    "\n".join(lines),
                    fontname=font_name,
                    fontsize=font_size,
                    color=color,
                    align=align,
                )
                return
            font_size -= 1

        # Final fallback: insert with min font size
        page.insert_textbox(
            rect,
            text,
            fontname=font_name,
            fontsize=self.PDF_MIN_FONT,
            color=color,
            align=align,
        )

    def _wrap_text_simple(self, text: str, max_width: float, font_size: float, font_name: str) -> List[str]:
        """Wrap text into lines that fit within max_width."""
        lines: List[str] = []
        for paragraph in text.split("\n"):
            words = paragraph.split()
            if not words:
                lines.append("")
                continue
            line = ""
            for word in words:
                candidate = f"{line} {word}".strip()
                try:
                    text_len = fitz.get_text_length(candidate, fontname=font_name, fontsize=font_size)
                except Exception:
                    # Estimate for non-standard fonts
                    text_len = len(candidate) * font_size * 0.5
                if text_len <= max_width:
                    line = candidate
                else:
                    if line:
                        lines.append(line)
                    line = word
            if line:
                lines.append(line)
        return lines

    # ===== Maintenance & Cleanup Methods =====

    def cleanup_font_cache(self, max_age_hours: int = 24) -> int:
        """Clean up old cached fonts to prevent disk space leak.
        
        Args:
            max_age_hours: Delete fonts older than this (default 24h)
            
        Returns:
            Number of files cleaned up
        """
        import time as time_module
        cleaned = 0
        try:
            font_dir = self.EXTRACTED_FONTS_DIR
            if os.path.exists(font_dir):
                cutoff = time_module.time() - (max_age_hours * 3600)
                for filename in os.listdir(font_dir):
                    filepath = os.path.join(font_dir, filename)
                    if os.path.isfile(filepath):
                        try:
                            if os.path.getmtime(filepath) < cutoff:
                                os.remove(filepath)
                                cleaned += 1
                        except OSError:
                            pass
                logger.info(f"Font cache cleanup: removed {cleaned} old font files")
        except Exception as e:
            logger.warning(f"Font cache cleanup failed: {e}")
        return cleaned

    # ===== Unified Entry Point =====

    async def translate_document(
        self,
        file_path: str,
        target_language: str,
        source_language: Optional[str] = None,
        output_path: Optional[str] = None,
        preserve_layout: bool = True,
        export_format: str = "pdf",
        translate_images: bool = False,
        glossary: Optional[Dict[str, str]] = None,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> Dict[str, Any]:
        """Unified document translation - main entry point for all file types.
        
        Args:
            file_path: Path to source document
            target_language: Target language for translation
            source_language: Source language (optional, auto-detected if not provided)
            output_path: Path for translated output
            preserve_layout: Whether to preserve document layout (PDF only)
            export_format: Output format - "pdf" or "docx"
            translate_images: Whether to OCR and translate text in images
            glossary: Optional term glossary for consistent translations
            progress_callback: Optional callback for progress updates (current, total, message)
            
        Returns:
            Dict with translation results including output_path, stats, etc.
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        # Clean up old cached fonts before starting (prevents memory leak)
        self.cleanup_font_cache(max_age_hours=48)
        
        result: Dict[str, Any] = {
            "source_path": file_path,
            "target_language": target_language,
            "source_language": source_language,
            "file_type": ext,
            "success": False,
        }
        
        try:
            if ext == '.pdf':
                pdf_result = await self.translate_pdf_preserve_layout(
                    file_path, 
                    target_language, 
                    source_language,
                    output_path, 
                    preserve_layout,
                    translate_images=translate_images,
                    glossary=glossary,
                    progress_callback=progress_callback
                )
                result.update(pdf_result)
                
                # Handle DOCX export if requested
                if export_format.lower() == 'docx' and result.get('output_path'):
                    pdf_output = result['output_path']
                    docx_path = pdf_output.rsplit('.', 1)[0] + '.docx'
                    try:
                        self.export_pdf_to_docx(pdf_output, docx_path, preserve_layout=True)
                        result['docx_path'] = docx_path
                        result['export_format'] = 'docx'
                    except Exception as e:
                        logger.warning(f"DOCX export failed: {e}")
                        result['docx_error'] = str(e)
                        
            elif ext in ['.docx', '.doc']:
                result.update(await self.translate_docx_preserve_formatting(
                    file_path, target_language, source_language, output_path
                ))
                
            elif ext in ['.pptx', '.ppt']:
                result.update(await self.translate_pptx(
                    file_path, target_language, source_language, output_path
                ))
                
            elif ext in ['.xlsx', '.xls']:
                result.update(await self.translate_xlsx(
                    file_path, target_language, source_language, output_path
                ))
                
            elif ext in ['.txt', '.md', '.rst']:
                text_result = await self.translate_text_file(
                    file_path, target_language, source_language
                )
                result.update(text_result)
                if output_path:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(text_result.get('translated_text', ''))
                    result['output_path'] = output_path
                    
            elif ext in ['.html', '.htm']:
                html_result = await self.translate_html(
                    file_path, target_language, source_language, output_path
                )
                result.update(html_result)
                
            else:
                result['error'] = f"Unsupported file type: {ext}"
                result['supported_types'] = ['.pdf', '.docx', '.doc', '.pptx', '.ppt', 
                                            '.xlsx', '.xls', '.txt', '.md', '.html']
                return result
                
            result['success'] = True
            
        except Exception as e:
            logger.error(f"Document translation failed: {e}", exc_info=True)
            result['error'] = str(e)
            result['success'] = False
            
        return result


document_translation_service = DocumentTranslationService()
