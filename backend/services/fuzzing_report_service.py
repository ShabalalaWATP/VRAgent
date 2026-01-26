"""
Fuzzing Campaign Report Service

Generates comprehensive AI-powered reports for completed fuzzing campaigns.
Supports export to Markdown, PDF, and Word formats.
"""

import asyncio
import hashlib
import io
import logging
import os
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class CrashFinding:
    """A crash finding with analysis."""
    crash_id: str
    crash_type: str
    exploitability: str
    confidence: float
    crash_address: Optional[str] = None
    stack_trace: List[str] = field(default_factory=list)
    root_cause: Optional[str] = None
    impact: Optional[str] = None
    recommendation: Optional[str] = None


@dataclass
class DecisionRecord:
    """A record of an AI decision."""
    decision_id: str
    timestamp: datetime
    decision_type: str
    reasoning: str
    outcome: Optional[str] = None
    was_effective: Optional[bool] = None


@dataclass
class ReportData:
    """Structured report data."""
    # Meta
    campaign_id: str
    binary_name: str
    binary_hash: str
    binary_type: str
    architecture: str

    # Timeline
    started_at: datetime
    completed_at: datetime
    duration: timedelta

    # Metrics
    total_executions: int
    executions_per_second: float
    final_coverage: float
    unique_crashes: int
    exploitable_crashes: int

    # AI Analysis
    executive_summary: str
    risk_rating: str  # Critical, High, Medium, Low, Informational
    key_findings: List[str]
    recommendations: List[str]

    # Detailed sections
    crash_findings: List[CrashFinding]
    decision_history: List[DecisionRecord]
    coverage_analysis: Dict[str, Any]
    strategy_effectiveness: Dict[str, Any]


# =============================================================================
# Report Generator
# =============================================================================

class FuzzingReportGenerator:
    """
    Generates comprehensive AI reports for fuzzing campaigns.
    """

    def __init__(self, ai_client: Optional[Any] = None):
        self.ai_client = ai_client

    async def generate_report(
        self,
        campaign_id: str,
        binary_name: str,
        binary_hash: str,
        binary_type: str,
        architecture: str,
        started_at: datetime,
        completed_at: datetime,
        total_executions: int,
        executions_per_second: float,
        final_coverage: float,
        unique_crashes: int,
        exploitable_crashes: int,
        crashes: List[Dict[str, Any]],
        decisions: List[Dict[str, Any]],
        coverage_data: Optional[Dict[str, Any]] = None,
    ) -> Tuple[ReportData, str]:
        """
        Generate a comprehensive report for a completed campaign.

        Returns:
            Tuple of (ReportData, markdown_report)
        """
        logger.info(f"Generating report for campaign {campaign_id}")

        duration = completed_at - started_at

        # Analyze crashes
        crash_findings = await self._analyze_crashes(crashes)

        # Analyze decisions
        decision_history = self._parse_decisions(decisions)

        # Calculate strategy effectiveness
        strategy_effectiveness = self._calculate_strategy_effectiveness(decisions, coverage_data)

        # Generate AI analysis
        executive_summary, risk_rating, key_findings, recommendations = await self._generate_ai_analysis(
            binary_name=binary_name,
            binary_type=binary_type,
            architecture=architecture,
            duration=duration,
            total_executions=total_executions,
            final_coverage=final_coverage,
            unique_crashes=unique_crashes,
            exploitable_crashes=exploitable_crashes,
            crash_findings=crash_findings,
            strategy_effectiveness=strategy_effectiveness,
        )

        # Build report data
        report_data = ReportData(
            campaign_id=campaign_id,
            binary_name=binary_name,
            binary_hash=binary_hash,
            binary_type=binary_type,
            architecture=architecture,
            started_at=started_at,
            completed_at=completed_at,
            duration=duration,
            total_executions=total_executions,
            executions_per_second=executions_per_second,
            final_coverage=final_coverage,
            unique_crashes=unique_crashes,
            exploitable_crashes=exploitable_crashes,
            executive_summary=executive_summary,
            risk_rating=risk_rating,
            key_findings=key_findings,
            recommendations=recommendations,
            crash_findings=crash_findings,
            decision_history=decision_history,
            coverage_analysis=coverage_data or {},
            strategy_effectiveness=strategy_effectiveness,
        )

        # Generate markdown report
        markdown_report = self._generate_markdown(report_data)

        return report_data, markdown_report

    async def _analyze_crashes(self, crashes: List[Dict[str, Any]]) -> List[CrashFinding]:
        """Analyze crash data and create findings."""
        findings = []

        for crash in crashes:
            finding = CrashFinding(
                crash_id=crash.get("crash_id") or crash.get("id") or "unknown",
                crash_type=crash.get("crash_type") or crash.get("type") or "unknown",
                exploitability=crash.get("exploitability") or "unknown",
                confidence=crash.get("confidence") or 0.0,
                crash_address=crash.get("crash_address"),
                stack_trace=crash.get("stack_trace") or crash.get("stack_frames") or [],
                root_cause=crash.get("root_cause"),
                impact=crash.get("impact"),
                recommendation=crash.get("recommendation"),
            )

            # Generate impact and recommendation if not present
            if not finding.impact:
                finding.impact = self._assess_crash_impact(finding)
            if not finding.recommendation:
                finding.recommendation = self._generate_crash_recommendation(finding)

            findings.append(finding)

        # Sort by exploitability
        exploitability_order = ["exploitable", "probably_exploitable", "probably_not_exploitable", "not_exploitable", "unknown"]
        findings.sort(key=lambda f: exploitability_order.index(f.exploitability.lower()) if f.exploitability.lower() in exploitability_order else 99)

        return findings

    def _assess_crash_impact(self, crash: CrashFinding) -> str:
        """Assess the security impact of a crash."""
        crash_type = crash.crash_type.lower()
        exploitability = crash.exploitability.lower()

        if exploitability == "exploitable":
            if "heap" in crash_type or "use-after-free" in crash_type:
                return "**CRITICAL**: This crash indicates a heap corruption vulnerability that could allow arbitrary code execution."
            elif "stack" in crash_type or "buffer" in crash_type:
                return "**CRITICAL**: This crash indicates a stack-based vulnerability that could allow code execution or control flow hijacking."
            else:
                return "**HIGH**: This crash is likely exploitable and could lead to arbitrary code execution."
        elif exploitability == "probably_exploitable":
            return "**HIGH**: This crash shows characteristics of an exploitable vulnerability. Further analysis recommended."
        elif exploitability == "probably_not_exploitable":
            return "**MEDIUM**: This crash may cause denial of service but is unlikely to be exploitable for code execution."
        else:
            return "**LOW**: This crash appears to be a non-security issue, but should be investigated for stability."

    def _generate_crash_recommendation(self, crash: CrashFinding) -> str:
        """Generate a recommendation for fixing a crash."""
        crash_type = crash.crash_type.lower()

        if "heap" in crash_type or "use-after-free" in crash_type:
            return "Review memory allocation patterns. Consider using smart pointers or RAII patterns. Enable AddressSanitizer for development builds."
        elif "stack" in crash_type or "buffer" in crash_type:
            return "Add bounds checking for all buffer operations. Use safe string functions (strncpy, snprintf). Enable stack canaries."
        elif "null" in crash_type or "nullptr" in crash_type:
            return "Add null pointer validation before dereferencing. Consider using optional types or assertions."
        elif "integer" in crash_type or "overflow" in crash_type:
            return "Add integer overflow checks. Use safe arithmetic functions. Consider using SafeInt or similar libraries."
        elif "divide" in crash_type or "division" in crash_type:
            return "Add zero-division checks before arithmetic operations."
        else:
            return "Investigate the crash root cause using a debugger. Consider adding input validation and error handling."

    def _parse_decisions(self, decisions: List[Dict[str, Any]]) -> List[DecisionRecord]:
        """Parse decision history into structured records."""
        records = []

        for d in decisions:
            timestamp = d.get("timestamp")
            if isinstance(timestamp, str):
                try:
                    timestamp = datetime.fromisoformat(timestamp.replace("Z", "+00:00"))
                except ValueError:
                    timestamp = datetime.utcnow()
            elif not isinstance(timestamp, datetime):
                timestamp = datetime.utcnow()

            records.append(DecisionRecord(
                decision_id=d.get("decision_id") or d.get("id") or "unknown",
                timestamp=timestamp,
                decision_type=d.get("decision_type") or d.get("type") or "unknown",
                reasoning=d.get("reasoning") or "",
                outcome=d.get("outcome"),
                was_effective=d.get("was_effective"),
            ))

        return records

    def _calculate_strategy_effectiveness(
        self,
        decisions: List[Dict[str, Any]],
        coverage_data: Optional[Dict[str, Any]],
    ) -> Dict[str, Any]:
        """Calculate effectiveness of different strategies used."""
        strategy_stats = {}

        for d in decisions:
            dtype = d.get("decision_type") or d.get("type") or "unknown"
            if dtype not in strategy_stats:
                strategy_stats[dtype] = {
                    "count": 0,
                    "effective": 0,
                    "ineffective": 0,
                }
            strategy_stats[dtype]["count"] += 1
            if d.get("was_effective"):
                strategy_stats[dtype]["effective"] += 1
            elif d.get("was_effective") is False:
                strategy_stats[dtype]["ineffective"] += 1

        # Calculate effectiveness rates
        for dtype, stats in strategy_stats.items():
            total_with_outcome = stats["effective"] + stats["ineffective"]
            if total_with_outcome > 0:
                stats["effectiveness_rate"] = stats["effective"] / total_with_outcome
            else:
                stats["effectiveness_rate"] = None

        return strategy_stats

    async def _generate_ai_analysis(
        self,
        binary_name: str,
        binary_type: str,
        architecture: str,
        duration: timedelta,
        total_executions: int,
        final_coverage: float,
        unique_crashes: int,
        exploitable_crashes: int,
        crash_findings: List[CrashFinding],
        strategy_effectiveness: Dict[str, Any],
    ) -> Tuple[str, str, List[str], List[str]]:
        """Generate AI-powered analysis sections."""

        # Calculate metrics for analysis
        hours = duration.total_seconds() / 3600
        exec_per_hour = total_executions / hours if hours > 0 else 0

        # Determine risk rating
        if exploitable_crashes >= 3:
            risk_rating = "Critical"
        elif exploitable_crashes >= 1:
            risk_rating = "High"
        elif unique_crashes >= 5:
            risk_rating = "Medium"
        elif unique_crashes >= 1:
            risk_rating = "Low"
        else:
            risk_rating = "Informational"

        # Try to use AI for enhanced analysis
        if self.ai_client:
            try:
                ai_result = await self._call_ai_for_analysis(
                    binary_name, binary_type, architecture, duration,
                    total_executions, final_coverage, unique_crashes,
                    exploitable_crashes, crash_findings, strategy_effectiveness
                )
                if ai_result:
                    return ai_result
            except Exception as e:
                logger.warning(f"AI analysis failed, using heuristic: {e}")

        # Fallback to heuristic analysis
        executive_summary = self._generate_heuristic_summary(
            binary_name, binary_type, architecture, duration,
            total_executions, final_coverage, unique_crashes, exploitable_crashes
        )

        key_findings = self._generate_heuristic_findings(
            final_coverage, unique_crashes, exploitable_crashes, crash_findings
        )

        recommendations = self._generate_heuristic_recommendations(
            final_coverage, unique_crashes, exploitable_crashes, crash_findings
        )

        return executive_summary, risk_rating, key_findings, recommendations

    async def _call_ai_for_analysis(
        self,
        binary_name: str,
        binary_type: str,
        architecture: str,
        duration: timedelta,
        total_executions: int,
        final_coverage: float,
        unique_crashes: int,
        exploitable_crashes: int,
        crash_findings: List[CrashFinding],
        strategy_effectiveness: Dict[str, Any],
    ) -> Optional[Tuple[str, str, List[str], List[str]]]:
        """Call AI model for enhanced analysis."""
        if not self.ai_client:
            return None

        # Build crash summary for prompt
        crash_summary = "\n".join([
            f"- {cf.crash_type} ({cf.exploitability}, {cf.confidence:.0%} confidence)"
            for cf in crash_findings[:10]
        ])

        hours = duration.total_seconds() / 3600

        prompt = f"""Analyze this fuzzing campaign and provide a security assessment:

**Binary:** {binary_name} ({binary_type}, {architecture})
**Duration:** {hours:.1f} hours
**Executions:** {total_executions:,} ({total_executions/hours:,.0f}/hour)
**Coverage:** {final_coverage:.1f}%
**Crashes Found:** {unique_crashes} total, {exploitable_crashes} exploitable

**Crash Summary:**
{crash_summary if crash_summary else "No crashes found"}

Provide:
1. A 2-3 sentence executive summary
2. Risk rating (Critical/High/Medium/Low/Informational)
3. 3-5 key findings as bullet points
4. 3-5 security recommendations as bullet points

Format your response as JSON:
{{
    "executive_summary": "...",
    "risk_rating": "...",
    "key_findings": ["...", "..."],
    "recommendations": ["...", "..."]
}}"""

        try:
            response = await self.ai_client.generate(prompt)
            if response and isinstance(response, dict):
                import json
                # Try to parse if it's a string
                if isinstance(response.get("content"), str):
                    try:
                        data = json.loads(response["content"])
                    except json.JSONDecodeError:
                        return None
                else:
                    data = response

                return (
                    data.get("executive_summary", ""),
                    data.get("risk_rating", "Medium"),
                    data.get("key_findings", []),
                    data.get("recommendations", []),
                )
        except Exception as e:
            logger.warning(f"AI analysis call failed: {e}")

        return None

    def _generate_heuristic_summary(
        self,
        binary_name: str,
        binary_type: str,
        architecture: str,
        duration: timedelta,
        total_executions: int,
        final_coverage: float,
        unique_crashes: int,
        exploitable_crashes: int,
    ) -> str:
        """Generate a heuristic executive summary."""
        hours = duration.total_seconds() / 3600

        if exploitable_crashes > 0:
            severity = "critical security vulnerabilities"
            action = "Immediate remediation is required."
        elif unique_crashes > 5:
            severity = "multiple stability issues and potential vulnerabilities"
            action = "Security review is recommended before deployment."
        elif unique_crashes > 0:
            severity = "minor issues"
            action = "Review and fix before production use."
        else:
            severity = "no crashes"
            action = "The binary appears stable under the tested conditions."

        return (
            f"Agentic fuzzing campaign completed for **{binary_name}** ({binary_type}/{architecture}) "
            f"after {hours:.1f} hours and {total_executions:,} executions. "
            f"Achieved {final_coverage:.1f}% code coverage. "
            f"Analysis identified {severity} with {unique_crashes} unique crashes "
            f"({exploitable_crashes} potentially exploitable). {action}"
        )

    def _generate_heuristic_findings(
        self,
        final_coverage: float,
        unique_crashes: int,
        exploitable_crashes: int,
        crash_findings: List[CrashFinding],
    ) -> List[str]:
        """Generate heuristic key findings."""
        findings = []

        # Coverage finding
        if final_coverage >= 80:
            findings.append(f"Excellent code coverage achieved ({final_coverage:.1f}%), indicating thorough testing.")
        elif final_coverage >= 60:
            findings.append(f"Good code coverage achieved ({final_coverage:.1f}%), but some paths remain unexplored.")
        elif final_coverage >= 40:
            findings.append(f"Moderate code coverage ({final_coverage:.1f}%). Significant portions of code were not exercised.")
        else:
            findings.append(f"Low code coverage ({final_coverage:.1f}%). Consider providing better seed inputs or increasing duration.")

        # Crash findings
        if exploitable_crashes > 0:
            findings.append(f"**{exploitable_crashes} exploitable crash(es)** detected requiring immediate attention.")

        if unique_crashes > exploitable_crashes:
            other_crashes = unique_crashes - exploitable_crashes
            findings.append(f"{other_crashes} additional crash(es) found that may indicate stability or security issues.")

        # Specific crash types
        crash_types = {}
        for cf in crash_findings:
            ct = cf.crash_type.lower()
            crash_types[ct] = crash_types.get(ct, 0) + 1

        for ct, count in sorted(crash_types.items(), key=lambda x: -x[1])[:3]:
            findings.append(f"Detected {count} {ct} issue(s).")

        if not crash_findings:
            findings.append("No crashes were detected during the fuzzing campaign.")

        return findings[:5]

    def _generate_heuristic_recommendations(
        self,
        final_coverage: float,
        unique_crashes: int,
        exploitable_crashes: int,
        crash_findings: List[CrashFinding],
    ) -> List[str]:
        """Generate heuristic recommendations."""
        recommendations = []

        if exploitable_crashes > 0:
            recommendations.append("**URGENT**: Investigate and fix all exploitable crashes before deployment.")
            recommendations.append("Enable memory safety tools (AddressSanitizer, Valgrind) in development builds.")

        if unique_crashes > 0:
            recommendations.append("Create unit tests to reproduce each crash for regression testing.")
            recommendations.append("Review input validation and boundary checking in affected code paths.")

        if final_coverage < 60:
            recommendations.append("Improve seed corpus with more diverse inputs to increase code coverage.")
            recommendations.append("Consider longer fuzzing duration or additional fuzzing strategies.")

        # Check for specific vulnerability patterns
        has_memory_issues = any("heap" in cf.crash_type.lower() or "buffer" in cf.crash_type.lower() for cf in crash_findings)
        if has_memory_issues:
            recommendations.append("Review memory management patterns. Consider using safe memory allocation practices.")

        if not recommendations:
            recommendations.append("Continue monitoring with regular fuzzing campaigns as part of CI/CD pipeline.")
            recommendations.append("Consider expanding test coverage with additional input formats and protocols.")

        return recommendations[:5]

    def _generate_markdown(self, report: ReportData) -> str:
        """Generate a formatted markdown report."""
        hours = report.duration.total_seconds() / 3600

        md = f"""# Agentic Fuzzing Campaign Report

**Campaign ID:** `{report.campaign_id}`
**Generated:** {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}

---

## Executive Summary

{report.executive_summary}

**Risk Rating:** {self._format_risk_badge(report.risk_rating)}

---

## Target Binary

| Property | Value |
|----------|-------|
| **Name** | {report.binary_name} |
| **Type** | {report.binary_type} |
| **Architecture** | {report.architecture} |
| **SHA-256** | `{report.binary_hash}` |

---

## Campaign Metrics

| Metric | Value |
|--------|-------|
| **Duration** | {hours:.1f} hours |
| **Total Executions** | {report.total_executions:,} |
| **Executions/Second** | {report.executions_per_second:,.0f} |
| **Code Coverage** | {report.final_coverage:.1f}% |
| **Unique Crashes** | {report.unique_crashes} |
| **Exploitable Crashes** | {report.exploitable_crashes} |
| **AI Decisions Made** | {len(report.decision_history)} |

---

## Key Findings

"""
        for i, finding in enumerate(report.key_findings, 1):
            md += f"{i}. {finding}\n"

        md += """
---

## Security Recommendations

"""
        for i, rec in enumerate(report.recommendations, 1):
            md += f"{i}. {rec}\n"

        if report.crash_findings:
            md += """
---

## Crash Analysis

"""
            for i, crash in enumerate(report.crash_findings, 1):
                exploitability_icon = self._get_exploitability_icon(crash.exploitability)
                md += f"""### Crash #{i}: {crash.crash_type}

{exploitability_icon} **Exploitability:** {crash.exploitability} ({crash.confidence:.0%} confidence)

**Crash ID:** `{crash.crash_id}`

"""
                if crash.crash_address:
                    md += f"**Address:** `{crash.crash_address}`\n\n"

                if crash.impact:
                    md += f"**Impact:** {crash.impact}\n\n"

                if crash.recommendation:
                    md += f"**Recommendation:** {crash.recommendation}\n\n"

                if crash.stack_trace:
                    md += "**Stack Trace:**\n```\n"
                    md += "\n".join(crash.stack_trace[:10])
                    if len(crash.stack_trace) > 10:
                        md += f"\n... and {len(crash.stack_trace) - 10} more frames"
                    md += "\n```\n\n"

        if report.decision_history:
            md += """
---

## AI Decision History

The agentic AI made the following strategic decisions during the campaign:

| Time | Decision | Reasoning |
|------|----------|-----------|
"""
            for decision in report.decision_history[:20]:
                elapsed = decision.timestamp - report.started_at
                elapsed_str = f"+{elapsed.total_seconds()/60:.0f}m"
                reasoning_short = decision.reasoning[:80] + "..." if len(decision.reasoning) > 80 else decision.reasoning
                md += f"| {elapsed_str} | {decision.decision_type} | {reasoning_short} |\n"

            if len(report.decision_history) > 20:
                md += f"\n*... and {len(report.decision_history) - 20} more decisions*\n"

        md += f"""
---

## Appendix

### Strategy Effectiveness

"""
        if report.strategy_effectiveness:
            md += "| Strategy | Count | Effectiveness |\n|----------|-------|---------------|\n"
            for strategy, stats in report.strategy_effectiveness.items():
                eff = f"{stats['effectiveness_rate']:.0%}" if stats.get('effectiveness_rate') is not None else "N/A"
                md += f"| {strategy} | {stats['count']} | {eff} |\n"
        else:
            md += "*No strategy effectiveness data available.*\n"

        md += f"""
---

*Report generated by VRAgent Agentic Binary Fuzzer*
*Campaign completed: {report.completed_at.strftime("%Y-%m-%d %H:%M UTC")}*
"""

        return md

    def _format_risk_badge(self, risk_rating: str) -> str:
        """Format risk rating with visual indicator."""
        badges = {
            "Critical": "🔴 **CRITICAL**",
            "High": "🟠 **HIGH**",
            "Medium": "🟡 **MEDIUM**",
            "Low": "🟢 **LOW**",
            "Informational": "🔵 **INFORMATIONAL**",
        }
        return badges.get(risk_rating, f"⚪ **{risk_rating.upper()}**")

    def _get_exploitability_icon(self, exploitability: str) -> str:
        """Get icon for exploitability level."""
        icons = {
            "exploitable": "🔴",
            "probably_exploitable": "🟠",
            "probably_not_exploitable": "🟡",
            "not_exploitable": "🟢",
            "unknown": "⚪",
        }
        return icons.get(exploitability.lower(), "⚪")


# =============================================================================
# Export Functions
# =============================================================================

def export_to_markdown(markdown_content: str) -> bytes:
    """Export report as Markdown file."""
    return markdown_content.encode("utf-8")


def export_to_pdf(markdown_content: str, title: str = "Fuzzing Campaign Report") -> bytes:
    """
    Export report as PDF.

    Uses markdown2 and weasyprint for conversion.
    Falls back to basic PDF if libraries not available.
    """
    try:
        import markdown2
        from weasyprint import HTML, CSS

        # Convert markdown to HTML
        html_content = markdown2.markdown(
            markdown_content,
            extras=["tables", "fenced-code-blocks", "header-ids"]
        )

        # Add CSS styling
        styled_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #3498db; color: white; }}
        tr:nth-child(even) {{ background-color: #f2f2f2; }}
        code {{ background-color: #f4f4f4; padding: 2px 6px; border-radius: 3px; font-family: 'Consolas', monospace; }}
        pre {{ background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        pre code {{ background-color: transparent; padding: 0; }}
        strong {{ color: #2c3e50; }}
        hr {{ border: none; border-top: 1px solid #eee; margin: 30px 0; }}
        .risk-critical {{ color: #e74c3c; font-weight: bold; }}
        .risk-high {{ color: #e67e22; font-weight: bold; }}
        .risk-medium {{ color: #f1c40f; font-weight: bold; }}
        .risk-low {{ color: #27ae60; font-weight: bold; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""

        # Generate PDF
        pdf_buffer = io.BytesIO()
        HTML(string=styled_html).write_pdf(pdf_buffer)
        return pdf_buffer.getvalue()

    except ImportError as e:
        logger.warning(f"PDF export dependencies not available: {e}")
        # Fallback: return a simple text-based PDF indicator
        return _create_fallback_pdf(markdown_content, title)


def _create_fallback_pdf(content: str, title: str) -> bytes:
    """Create a basic PDF without weasyprint."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
        )
        story.append(Paragraph(title, title_style))

        # Convert markdown to paragraphs (basic conversion)
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('---'):
                story.append(Spacer(1, 20))
            elif line.strip():
                # Basic formatting
                line = line.replace('**', '<b>').replace('**', '</b>')
                line = line.replace('`', '<code>').replace('`', '</code>')
                try:
                    story.append(Paragraph(line, styles['Normal']))
                except Exception:
                    story.append(Paragraph(line.replace('<', '&lt;').replace('>', '&gt;'), styles['Normal']))
            else:
                story.append(Spacer(1, 12))

        doc.build(story)
        return buffer.getvalue()

    except ImportError:
        logger.warning("reportlab not available, returning plain text as PDF")
        # Last resort: just return the markdown as bytes
        return f"PDF export not available. Please install weasyprint or reportlab.\n\n{content}".encode('utf-8')


def export_to_docx(markdown_content: str, title: str = "Fuzzing Campaign Report") -> bytes:
    """
    Export report as Word document.

    Uses python-docx for generation.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import re

        doc = Document()

        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Parse markdown and build document
        lines = markdown_content.split('\n')
        i = 0
        in_table = False
        table_data = []
        in_code_block = False
        code_lines = []

        while i < len(lines):
            line = lines[i]

            # Code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    if code_lines:
                        p = doc.add_paragraph()
                        p.style = 'Normal'
                        for cl in code_lines:
                            run = p.add_run(cl + '\n')
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Tables
            if '|' in line and not line.startswith('|--'):
                if line.startswith('|'):
                    cells = [c.strip() for c in line.split('|')[1:-1]]
                    if cells:
                        table_data.append(cells)
                i += 1
                continue
            elif table_data:
                # End of table, render it
                if table_data:
                    # Filter out separator rows
                    table_data = [row for row in table_data if not all(c.startswith('-') for c in row)]
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_text in enumerate(row_data):
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_text.replace('**', '').replace('`', '')
                                if row_idx == 0:
                                    cell.paragraphs[0].runs[0].bold = True
                        doc.add_paragraph()
                table_data = []

            # Headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('---'):
                # Horizontal rule - add some space
                doc.add_paragraph()
            elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                # Numbered list
                p = doc.add_paragraph(style='List Number')
                _add_formatted_text(p, line.split('.', 1)[1].strip())
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                # Bullet list
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_text(p, line.strip()[2:])
            elif line.strip():
                # Regular paragraph
                p = doc.add_paragraph()
                _add_formatted_text(p, line)

            i += 1

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError as e:
        logger.warning(f"python-docx not available: {e}")
        return f"DOCX export not available. Please install python-docx.\n\n{markdown_content}".encode('utf-8')


def _add_formatted_text(paragraph, text: str):
    """Add text with basic markdown formatting to a paragraph."""
    import re

    # Split by formatting markers
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|_[^_]+_)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
        elif (part.startswith('*') and part.endswith('*')) or (part.startswith('_') and part.endswith('_')):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
