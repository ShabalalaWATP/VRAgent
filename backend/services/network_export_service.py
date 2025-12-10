"""
Network Analysis Export Service for VRAgent.

Exports PCAP, Nmap, and SSL analysis reports to Markdown, PDF, and Word formats.
"""

import io
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from backend.core.logging import get_logger

logger = get_logger(__name__)


# ============================================================================
# SSL REPORT GENERATION
# ============================================================================

def generate_ssl_markdown_report(
    title: str,
    summary_data: Dict[str, Any],
    results_data: List[Dict[str, Any]],
    ai_report: Optional[Dict[str, Any]] = None,
) -> str:
    """Generate a Markdown report from SSL scan data with exploitation focus."""
    
    lines = []
    
    # Header
    lines.append(f"# 🔒 {title}")
    lines.append(f"\n**Analysis Type:** SSL/TLS Security Assessment")
    lines.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    
    # Risk Overview
    if ai_report and "structured_report" in ai_report:
        report = ai_report["structured_report"]
        risk_level = report.get("overall_risk_level", "Unknown")
        risk_score = report.get("risk_score", "N/A")
        
        lines.append("---")
        lines.append(f"\n## 🎯 Risk Assessment")
        lines.append(f"\n**Overall Risk Level:** {risk_level}")
        lines.append(f"**Risk Score:** {risk_score}/100")
        
        if report.get("executive_summary"):
            lines.append(f"\n{report['executive_summary']}")
        lines.append("")
    
    # Summary Statistics
    lines.append("---")
    lines.append("\n## 📊 Scan Summary")
    lines.append("")
    lines.append(f"- **Total Hosts Scanned:** {summary_data.get('total_hosts', 0)}")
    lines.append(f"- **Hosts with SSL/TLS:** {summary_data.get('hosts_with_ssl', 0)}")
    lines.append(f"- **Expired Certificates:** {summary_data.get('expired_certs', 0)}")
    lines.append(f"- **Self-Signed Certificates:** {summary_data.get('self_signed_certs', 0)}")
    lines.append(f"- **Hosts with Weak Protocols:** {summary_data.get('weak_protocols', 0)}")
    lines.append(f"- **Hosts with Weak Ciphers:** {summary_data.get('weak_ciphers', 0)}")
    lines.append("")
    
    # Vulnerability Summary
    if summary_data.get('total_vulnerabilities', 0) > 0:
        lines.append("### ⚠️ Vulnerability Statistics")
        lines.append(f"- **Total Vulnerabilities:** {summary_data.get('total_vulnerabilities', 0)}")
        lines.append(f"- **Critical Vulnerabilities:** {summary_data.get('critical_vulnerabilities', 0)}")
        lines.append(f"- **Exploitable Vulnerabilities:** {summary_data.get('exploitable_vulnerabilities', 0)}")
        lines.append(f"- **Certificate Chain Issues:** {summary_data.get('chain_issues', 0)}")
        lines.append("")
    
    # Vulnerability Details (grouped by severity)
    all_vulns = []
    for result in results_data:
        host = result.get('host', 'Unknown')
        port = result.get('port', 443)
        for vuln in result.get('vulnerabilities', []):
            vuln['target'] = f"{host}:{port}"
            all_vulns.append(vuln)
    
    if all_vulns:
        lines.append("---")
        lines.append("\n## 🔴 Detected Vulnerabilities")
        lines.append("")
        
        # Group by severity
        by_severity = {"critical": [], "high": [], "medium": [], "low": []}
        for v in all_vulns:
            sev = v.get("severity", "low").lower()
            if sev in by_severity:
                by_severity[sev].append(v)
        
        for severity in ["critical", "high", "medium", "low"]:
            vulns = by_severity[severity]
            if vulns:
                emoji = {"critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🟢"}[severity]
                lines.append(f"\n### {emoji} {severity.upper()} ({len(vulns)})")
                lines.append("")
                lines.append("| Vulnerability | CVE | CVSS | Target | Exploit Difficulty |")
                lines.append("|---------------|-----|------|--------|-------------------|")
                for v in vulns:
                    name = v.get('name', v.get('vuln_id', 'Unknown'))
                    cve = v.get('cve', 'N/A')
                    cvss = v.get('cvss', 'N/A')
                    target = v.get('target', 'Unknown')
                    exploit = v.get('exploit_difficulty', 'Unknown')
                    lines.append(f"| {name} | {cve} | {cvss} | {target} | {exploit} |")
                lines.append("")
                
                # Details for each vulnerability
                for v in vulns:
                    name = v.get('name', v.get('vuln_id', 'Unknown'))
                    lines.append(f"#### {name}")
                    lines.append(f"\n{v.get('description', 'No description available.')}")
                    if v.get('affected'):
                        lines.append(f"\n**Affected:** {v.get('affected')}")
                    if v.get('evidence'):
                        lines.append(f"\n**Evidence:** `{v.get('evidence')}`")
                    if v.get('is_exploitable'):
                        lines.append(f"\n⚠️ **This vulnerability is considered exploitable**")
                    lines.append("")
    
    # Certificate Chain Issues
    chain_issues = []
    for result in results_data:
        chain_info = result.get('chain_info', {})
        if chain_info and chain_info.get('chain_errors'):
            chain_issues.append({
                'target': f"{result.get('host', 'Unknown')}:{result.get('port', 443)}",
                'chain_info': chain_info
            })
    
    if chain_issues:
        lines.append("---")
        lines.append("\n## 🔗 Certificate Chain Issues")
        lines.append("")
        for issue in chain_issues:
            lines.append(f"### {issue['target']}")
            ci = issue['chain_info']
            lines.append(f"- **Chain Length:** {ci.get('chain_length', 'N/A')}")
            lines.append(f"- **Is Complete:** {'✅ Yes' if ci.get('is_complete') else '❌ No'}")
            lines.append(f"- **Is Trusted:** {'✅ Yes' if ci.get('is_trusted') else '❌ No'}")
            if ci.get('root_ca'):
                lines.append(f"- **Root CA:** {ci.get('root_ca')}")
            if ci.get('chain_errors'):
                lines.append("\n**Chain Errors:**")
                for err in ci['chain_errors']:
                    lines.append(f"- ⚠️ {err}")
            lines.append("")
    
    # Detailed Host Results
    lines.append("---")
    lines.append("\n## 🖥️ Host Details")
    lines.append("")
    
    for result in results_data:
        host = result.get('host', 'Unknown')
        port = result.get('port', 443)
        has_ssl = result.get('has_ssl', False)
        error = result.get('error')
        
        lines.append(f"### {host}:{port}")
        
        if error:
            lines.append(f"\n**Error:** {error}")
            lines.append("")
            continue
        
        if not has_ssl:
            lines.append("\n**Status:** No SSL/TLS detected")
            lines.append("")
            continue
        
        # Certificate Info
        cert = result.get('certificate', {})
        if cert:
            lines.append("\n**Certificate:**")
            lines.append(f"- **Subject:** {cert.get('subject', 'N/A')}")
            lines.append(f"- **Issuer:** {cert.get('issuer', 'N/A')}")
            lines.append(f"- **Valid From:** {cert.get('not_before', 'N/A')}")
            lines.append(f"- **Valid Until:** {cert.get('not_after', 'N/A')}")
            lines.append(f"- **Days Until Expiry:** {cert.get('days_until_expiry', 'N/A')}")
            lines.append(f"- **Is Expired:** {'⚠️ Yes' if cert.get('is_expired') else '✅ No'}")
            lines.append(f"- **Is Self-Signed:** {'⚠️ Yes' if cert.get('is_self_signed') else '✅ No'}")
            lines.append(f"- **Signature Algorithm:** {cert.get('signature_algorithm', 'N/A')}")
            lines.append(f"- **Key Type:** {cert.get('key_type', 'N/A')}")
            lines.append(f"- **Key Size:** {cert.get('key_size', 'N/A')} bits")
            if cert.get('san'):
                lines.append(f"- **SANs:** {', '.join(cert['san'][:10])}")
        
        # Supported Protocols
        protocols = result.get('supported_protocols', [])
        if protocols:
            lines.append(f"\n**Supported Protocols:** {', '.join(protocols)}")
            # Highlight weak protocols
            weak = [p for p in protocols if p in ['SSLv2', 'SSLv3', 'TLSv1.0', 'TLSv1.1']]
            if weak:
                lines.append(f"\n⚠️ **Weak Protocols Detected:** {', '.join(weak)}")
        
        # Cipher Suites
        ciphers = result.get('cipher_suites', [])
        if ciphers:
            lines.append(f"\n**Cipher Suites ({len(ciphers)} total):**")
            for cipher in ciphers[:10]:
                lines.append(f"- `{cipher}`")
            if len(ciphers) > 10:
                lines.append(f"- ... and {len(ciphers) - 10} more")
        
        # Security Findings
        findings = result.get('findings', [])
        if findings:
            lines.append(f"\n**Security Findings ({len(findings)}):**")
            for f in findings[:10]:
                sev = f.get('severity', 'info').upper()
                title = f.get('title', 'Unknown')
                lines.append(f"- [{sev}] {title}")
            if len(findings) > 10:
                lines.append(f"- ... and {len(findings) - 10} more")
        
        lines.append("")
    
    # AI Exploitation Analysis
    if ai_report and "structured_report" in ai_report:
        report = ai_report["structured_report"]
        
        # Exploitation Scenarios
        if report.get("exploitation_scenarios"):
            lines.append("---")
            lines.append("\n## 💀 Exploitation Scenarios")
            lines.append("")
            lines.append("> ⚠️ **OFFENSIVE SECURITY ANALYSIS** - For authorized penetration testing only")
            lines.append("")
            
            for i, scenario in enumerate(report["exploitation_scenarios"], 1):
                lines.append(f"### Scenario {i}: {scenario.get('attack_name', 'Unknown Attack')}")
                lines.append(f"\n**Target Vulnerability:** {scenario.get('target_vulnerability', 'N/A')}")
                lines.append(f"\n**Exploit Difficulty:** {scenario.get('exploit_difficulty', 'N/A')}")
                lines.append(f"\n**Required Tools:** {scenario.get('required_tools', 'N/A')}")
                lines.append(f"\n**Potential Impact:** {scenario.get('potential_impact', 'N/A')}")
                
                if scenario.get('attack_steps'):
                    lines.append("\n**Attack Steps:**")
                    for step in scenario['attack_steps']:
                        lines.append(f"1. {step}")
                
                if scenario.get('example_commands'):
                    lines.append("\n**Example Commands:**")
                    lines.append("```bash")
                    for cmd in scenario['example_commands']:
                        lines.append(cmd)
                    lines.append("```")
                
                if scenario.get('indicators_of_compromise'):
                    lines.append("\n**Indicators of Compromise:**")
                    for ioc in scenario['indicators_of_compromise']:
                        lines.append(f"- {ioc}")
                
                lines.append("")
        
        # High-Value Targets
        if report.get("high_value_targets"):
            lines.append("---")
            lines.append("\n## 🎯 High-Value Targets")
            lines.append("")
            lines.append("| Host | Risk Level | Weaknesses | Attack Priority |")
            lines.append("|------|------------|------------|-----------------|")
            for target in report["high_value_targets"]:
                host = target.get('host', 'Unknown')
                risk = target.get('risk_level', 'N/A')
                weaknesses = ', '.join(target.get('weaknesses', [])[:3])
                priority = target.get('attack_priority', 'N/A')
                lines.append(f"| {host} | {risk} | {weaknesses} | {priority} |")
            lines.append("")
            
            for target in report["high_value_targets"]:
                lines.append(f"### {target.get('host', 'Unknown')}")
                lines.append(f"\n**Why High Value:** {target.get('why_high_value', 'N/A')}")
                lines.append(f"\n**Recommended Attack Vector:** {target.get('recommended_attack_vector', 'N/A')}")
                lines.append("")
        
        # Lateral Movement
        if report.get("lateral_movement_opportunities"):
            lines.append("---")
            lines.append("\n## 🔀 Lateral Movement Opportunities")
            lines.append("")
            for opp in report["lateral_movement_opportunities"]:
                lines.append(f"- **{opp.get('from_host', 'N/A')} → {opp.get('to_host', 'N/A')}:** {opp.get('technique', 'N/A')}")
            lines.append("")
        
        # Tool Recommendations
        if report.get("tool_recommendations"):
            lines.append("---")
            lines.append("\n## 🛠️ Recommended Tools")
            lines.append("")
            lines.append("| Tool | Purpose | Target |")
            lines.append("|------|---------|--------|")
            for tool in report["tool_recommendations"]:
                name = tool.get('tool_name', 'Unknown')
                purpose = tool.get('purpose', 'N/A')
                target = tool.get('target_host', 'All')
                lines.append(f"| {name} | {purpose} | {target} |")
            lines.append("")
    
    # Footer
    lines.append("---")
    lines.append(f"\n*Report generated by VRAgent SSL/TLS Security Scanner*")
    lines.append(f"*{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
    lines.append("\n> ⚠️ This report is intended for authorized security testing only.")
    lines.append("> Unauthorized access to computer systems is illegal.")
    
    return "\n".join(lines)


def generate_ssl_pdf_report(markdown_content: str) -> bytes:
    """Generate PDF from SSL markdown content."""
    # Use the generic PDF generator which handles the markdown format
    return generate_pdf_report(markdown_content)


def generate_ssl_docx_report(markdown_content: str) -> bytes:
    """Generate Word document from SSL markdown content."""
    # Use the generic Word generator which handles the markdown format
    return generate_docx_report(markdown_content)


def generate_markdown_report(
    analysis_type: str,
    title: str,
    summary_data: Dict[str, Any],
    findings_data: list,
    ai_report: Optional[Dict[str, Any]] = None,
) -> str:
    """Generate a Markdown report from network analysis data."""
    
    lines = []
    
    # Header
    lines.append(f"# {title}")
    lines.append(f"\n**Analysis Type:** {analysis_type.upper()}")
    lines.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    
    # Risk Overview (from AI report if available)
    if ai_report and "structured_report" in ai_report:
        report = ai_report["structured_report"]
        risk_level = report.get("risk_level", "Unknown")
        risk_score = report.get("risk_score", "N/A")
        
        lines.append("---")
        lines.append(f"\n## 🎯 Risk Assessment")
        lines.append(f"\n**Risk Level:** {risk_level}")
        lines.append(f"**Risk Score:** {risk_score}/100")
        lines.append("")
    
    # Executive Summary
    if ai_report and "structured_report" in ai_report:
        report = ai_report["structured_report"]
        if report.get("executive_summary"):
            lines.append("---")
            lines.append("\n## 📋 Executive Summary")
            lines.append(f"\n{report['executive_summary']}")
            lines.append("")
    
    # Summary Statistics
    lines.append("---")
    lines.append("\n## 📊 Summary Statistics")
    lines.append("")
    
    if analysis_type == "pcap":
        lines.append(f"- **Total Packets:** {summary_data.get('total_packets', 'N/A'):,}")
        lines.append(f"- **Duration:** {summary_data.get('duration_seconds', 'N/A')} seconds")
        lines.append(f"- **Security Findings:** {summary_data.get('potential_issues', 0)}")
        
        if summary_data.get("protocols"):
            lines.append("\n### Protocol Distribution")
            lines.append("| Protocol | Count |")
            lines.append("|----------|-------|")
            for proto, count in sorted(summary_data["protocols"].items(), key=lambda x: x[1], reverse=True)[:15]:
                lines.append(f"| {proto} | {count:,} |")
        
        if summary_data.get("top_talkers"):
            lines.append("\n### Top Communicating Hosts")
            lines.append("| IP Address | Packets | Bytes |")
            lines.append("|------------|---------|-------|")
            for host in summary_data["top_talkers"][:10]:
                lines.append(f"| {host['ip']} | {host['packets']:,} | {host['bytes']:,} |")
        
        if summary_data.get("dns_queries"):
            lines.append("\n### DNS Queries")
            for query in summary_data["dns_queries"][:30]:
                lines.append(f"- `{query}`")
        
        if summary_data.get("http_hosts"):
            lines.append("\n### HTTP Hosts")
            for host in summary_data["http_hosts"][:30]:
                lines.append(f"- `{host}`")
    
    elif analysis_type == "nmap":
        lines.append(f"- **Total Hosts:** {summary_data.get('total_hosts', 'N/A')}")
        lines.append(f"- **Hosts Up:** {summary_data.get('hosts_up', 'N/A')}")
        lines.append(f"- **Open Ports:** {summary_data.get('open_ports', 'N/A')}")
        lines.append(f"- **Scan Type:** {summary_data.get('scan_type', 'N/A')}")
        lines.append(f"- **Scan Command:** `{summary_data.get('command', 'N/A')}`")
        
        if summary_data.get("services_detected"):
            lines.append("\n### Services Detected")
            lines.append("| Service | Count |")
            lines.append("|---------|-------|")
            for svc, count in sorted(summary_data["services_detected"].items(), key=lambda x: x[1], reverse=True)[:15]:
                lines.append(f"| {svc} | {count} |")
    
    # Security Findings
    if findings_data:
        lines.append("")
        lines.append("---")
        lines.append("\n## ⚠️ Security Findings")
        lines.append("")
        
        # Group by severity
        by_severity = {"critical": [], "high": [], "medium": [], "low": [], "info": []}
        for f in findings_data:
            sev = f.get("severity", "info").lower()
            if sev in by_severity:
                by_severity[sev].append(f)
            else:
                by_severity["info"].append(f)
        
        for severity in ["critical", "high", "medium", "low", "info"]:
            findings = by_severity[severity]
            if findings:
                emoji = {"critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🟢", "info": "🔵"}[severity]
                lines.append(f"\n### {emoji} {severity.upper()} ({len(findings)})")
                lines.append("")
                for f in findings:
                    lines.append(f"#### {f.get('title', 'Unknown Finding')}")
                    lines.append(f"\n{f.get('description', 'No description')}")
                    if f.get("source_ip") or f.get("host"):
                        lines.append(f"\n- **Source/Host:** {f.get('source_ip') or f.get('host')}")
                    if f.get("dest_ip"):
                        lines.append(f"- **Destination:** {f.get('dest_ip')}")
                    if f.get("port"):
                        lines.append(f"- **Port:** {f.get('port')}")
                    if f.get("evidence"):
                        lines.append(f"- **Evidence:** `{f.get('evidence')[:200]}...`")
                    lines.append("")
    
    # AI Analysis Details
    if ai_report and "structured_report" in ai_report:
        report = ai_report["structured_report"]
        
        # Key Findings from AI
        if report.get("key_findings"):
            lines.append("---")
            lines.append("\n## 🔍 Key Findings (AI Analysis)")
            lines.append("")
            for i, finding in enumerate(report["key_findings"], 1):
                lines.append(f"### {i}. {finding.get('title', 'Finding')}")
                lines.append(f"\n**Severity:** {finding.get('severity', 'Unknown')}")
                lines.append(f"\n{finding.get('description', '')}")
                if finding.get("evidence"):
                    lines.append(f"\n**Evidence:** {finding.get('evidence')}")
                if finding.get("recommendation"):
                    lines.append(f"\n**Recommendation:** {finding.get('recommendation')}")
                lines.append("")
        
        # Credential Exposure (PCAP)
        if report.get("credential_exposure") and report["credential_exposure"].get("severity") != "None":
            cred = report["credential_exposure"]
            lines.append("---")
            lines.append("\n## 🔐 Credential Exposure")
            lines.append(f"\n**Severity:** {cred.get('severity', 'Unknown')}")
            lines.append(f"\n{cred.get('summary', '')}")
            
            if cred.get("exposed_credentials"):
                lines.append("\n| Type | Service | Source | Destination |")
                lines.append("|------|---------|--------|-------------|")
                for c in cred["exposed_credentials"]:
                    lines.append(f"| {c.get('type', '')} | {c.get('service', '')} | {c.get('source_ip', '')} | {c.get('dest_ip', '')} |")
            
            if cred.get("immediate_actions"):
                lines.append("\n**Immediate Actions:**")
                for action in cred["immediate_actions"]:
                    lines.append(f"- ⚡ {action}")
            lines.append("")
        
        # High Risk Hosts (Nmap)
        if report.get("high_risk_hosts"):
            lines.append("---")
            lines.append("\n## 🖥️ High-Risk Hosts")
            lines.append("")
            for host in report["high_risk_hosts"]:
                lines.append(f"### {host.get('ip', 'Unknown')} ({host.get('hostname', 'no hostname')})")
                lines.append(f"\n- **Risk Level:** {host.get('risk_level', 'Unknown')}")
                lines.append(f"- **Open Ports:** {host.get('open_ports_count', 'N/A')}")
                lines.append(f"- **OS:** {host.get('os', 'Unknown')}")
                lines.append(f"- **Concerns:** {host.get('concerns', '')}")
                if host.get("priority_actions"):
                    lines.append("\n**Priority Actions:**")
                    for action in host["priority_actions"]:
                        lines.append(f"- {action}")
                lines.append("")
        
        # Attack Vectors/Indicators
        if report.get("attack_vectors"):
            lines.append("---")
            lines.append("\n## 🎯 Attack Vectors")
            lines.append("")
            for vector in report["attack_vectors"]:
                lines.append(f"### {vector.get('vector', 'Unknown Vector')}")
                lines.append(f"\n**Severity:** {vector.get('severity', 'Unknown')} | **Likelihood:** {vector.get('likelihood', 'Unknown')}")
                lines.append(f"\n{vector.get('description', '')}")
                lines.append(f"\n**Potential Impact:** {vector.get('potential_impact', 'Unknown')}")
                lines.append("")
        
        # IOCs
        if report.get("indicators_of_compromise"):
            lines.append("---")
            lines.append("\n## 🚨 Indicators of Compromise")
            lines.append("")
            lines.append("| Type | Value | Threat Level | Context |")
            lines.append("|------|-------|--------------|---------|")
            for ioc in report["indicators_of_compromise"]:
                lines.append(f"| {ioc.get('type', '')} | `{ioc.get('value', '')}` | {ioc.get('threat_level', '')} | {ioc.get('context', '')} |")
            lines.append("")
        
        # Recommendations
        if report.get("recommendations"):
            lines.append("---")
            lines.append("\n## ✅ Recommendations")
            lines.append("")
            
            # Group by priority
            by_priority = {"immediate": [], "high": [], "medium": [], "low": []}
            for rec in report["recommendations"]:
                priority = rec.get("priority", "medium").lower()
                if priority in by_priority:
                    by_priority[priority].append(rec)
                else:
                    by_priority["medium"].append(rec)
            
            for priority in ["immediate", "high", "medium", "low"]:
                recs = by_priority[priority]
                if recs:
                    lines.append(f"\n### {priority.upper()} Priority")
                    lines.append("")
                    for rec in recs:
                        lines.append(f"- **[{rec.get('category', 'General')}]** {rec.get('action', '')}")
                        if rec.get("rationale"):
                            lines.append(f"  - *Rationale:* {rec.get('rationale')}")
                        if rec.get("effort"):
                            lines.append(f"  - *Effort:* {rec.get('effort')}")
                    lines.append("")
    
    # Footer
    lines.append("---")
    lines.append(f"\n*Report generated by VRAgent Network Analysis*")
    lines.append(f"*{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
    
    return "\n".join(lines)


def generate_pdf_report(markdown_content: str) -> bytes:
    """Generate PDF from markdown content."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        import re
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='CustomTitle', fontSize=24, spaceAfter=30, textColor=colors.HexColor("#1e40af")))
        styles.add(ParagraphStyle(name='CustomH2', fontSize=16, spaceAfter=12, spaceBefore=20, textColor=colors.HexColor("#1e40af")))
        styles.add(ParagraphStyle(name='CustomH3', fontSize=14, spaceAfter=10, spaceBefore=15, textColor=colors.HexColor("#374151")))
        styles.add(ParagraphStyle(name='CustomBody', fontSize=11, spaceAfter=8, leading=14))
        styles.add(ParagraphStyle(name='CustomCode', fontSize=9, fontName='Courier', backColor=colors.HexColor("#f3f4f6"), spaceAfter=8))
        
        story = []
        
        # Parse markdown and convert to reportlab elements
        lines = markdown_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Skip empty lines
            if not line.strip():
                story.append(Spacer(1, 6))
                i += 1
                continue
            
            # Headers
            if line.startswith('# '):
                story.append(Paragraph(line[2:], styles['CustomTitle']))
            elif line.startswith('## '):
                text = re.sub(r'[🎯📋📊⚠️🔍🔐🖥️🚨✅]', '', line[3:])
                story.append(Paragraph(text.strip(), styles['CustomH2']))
            elif line.startswith('### '):
                text = re.sub(r'[🔴🟠🟡🟢🔵⚡]', '', line[4:])
                story.append(Paragraph(text.strip(), styles['CustomH3']))
            elif line.startswith('#### '):
                story.append(Paragraph(f"<b>{line[5:]}</b>", styles['CustomBody']))
            # Horizontal rule
            elif line.startswith('---'):
                story.append(Spacer(1, 20))
            # Bold text
            elif line.startswith('**') and line.endswith('**'):
                story.append(Paragraph(f"<b>{line[2:-2]}</b>", styles['CustomBody']))
            # List items
            elif line.startswith('- '):
                # Convert markdown bold **text** to HTML <b>text</b>
                text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', line[2:])
                text = text.replace('`', '')
                story.append(Paragraph(f"• {text}", styles['CustomBody']))
            # Table detection
            elif line.startswith('|') and i + 1 < len(lines) and lines[i + 1].startswith('|'):
                # Parse table
                table_data = []
                while i < len(lines) and lines[i].startswith('|'):
                    if '---' not in lines[i]:  # Skip separator line
                        row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                        table_data.append(row)
                    i += 1
                
                if table_data:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1e40af")),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 10))
                continue
            # Code blocks
            elif line.startswith('`') and line.endswith('`'):
                story.append(Paragraph(line[1:-1], styles['CustomCode']))
            # Regular text
            else:
                # Convert markdown bold **text** to HTML <b>text</b>
                text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', line)
                text = text.replace('`', '')
                if text.strip():
                    story.append(Paragraph(text, styles['CustomBody']))
            
            i += 1
        
        doc.build(story)
        return buffer.getvalue()
        
    except ImportError:
        logger.error("reportlab not installed for PDF generation")
        raise RuntimeError("PDF export requires reportlab. Install with: pip install reportlab")


def generate_docx_report(markdown_content: str) -> bytes:
    """Generate Word document from markdown content."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            if not line.strip():
                i += 1
                continue
            
            # Headers
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=0)
                p.runs[0].font.color.rgb = RGBColor(30, 64, 175)
            elif line.startswith('## '):
                text = re.sub(r'[🎯📋📊⚠️🔍🔐🖥️🚨✅]', '', line[3:]).strip()
                p = doc.add_heading(text, level=1)
            elif line.startswith('### '):
                text = re.sub(r'[🔴🟠🟡🟢🔵⚡]', '', line[4:]).strip()
                p = doc.add_heading(text, level=2)
            elif line.startswith('#### '):
                p = doc.add_paragraph()
                run = p.add_run(line[5:])
                run.bold = True
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('|') and i + 1 < len(lines) and lines[i + 1].startswith('|'):
                # Parse table
                table_data = []
                while i < len(lines) and lines[i].startswith('|'):
                    if '---' not in lines[i]:
                        row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                        table_data.append(row)
                    i += 1
                
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for row_idx, row in enumerate(table_data):
                        for col_idx, cell in enumerate(row):
                            table.rows[row_idx].cells[col_idx].text = cell.replace('`', '')
                            if row_idx == 0:
                                table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
                continue
            elif line.strip():
                clean_line = line.replace('**', '').replace('`', '')
                doc.add_paragraph(clean_line)
            
            i += 1
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
        
    except ImportError:
        logger.error("python-docx not installed for Word generation")
        raise RuntimeError("Word export requires python-docx. Install with: pip install python-docx")
