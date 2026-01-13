"""
Fuzzer Report Export Service

Generates professional Markdown, PDF, and Word documents from agentic fuzzer reports.
Includes proper formatting with bold headings, bullet points, tables, and styling.
"""

from io import BytesIO
from typing import Dict, Any, List, Optional
from datetime import datetime
import json

from backend.core.logging import get_logger
from backend.models.models import AgenticFuzzerReport

logger = get_logger(__name__)


# =============================================================================
# MARKDOWN GENERATION
# =============================================================================

def generate_fuzzer_report_markdown(report: AgenticFuzzerReport) -> str:
    """
    Generate a professional Markdown report from an agentic fuzzer report.
    
    Args:
        report: The AgenticFuzzerReport database model
        
    Returns:
        Formatted Markdown string
    """
    md = []
    
    # Title and metadata
    md.append(f"# 🔐 Security Assessment Report")
    md.append("")
    md.append(f"**Target:** {report.target_url}")
    md.append(f"**Session ID:** `{report.session_id}`")
    md.append(f"**Scan Profile:** {report.scan_profile or 'Default'}")
    md.append(f"**Generated:** {report.completed_at.strftime('%Y-%m-%d %H:%M:%S UTC') if report.completed_at else 'N/A'}")
    md.append("")
    
    # Duration
    if report.duration_seconds:
        mins = int(report.duration_seconds // 60)
        secs = int(report.duration_seconds % 60)
        md.append(f"**Duration:** {mins}m {secs}s")
    md.append("")
    
    # Horizontal rule
    md.append("---")
    md.append("")
    
    # Executive Summary
    md.append("## 📋 Executive Summary")
    md.append("")
    
    if report.executive_summary:
        md.append(report.executive_summary)
    elif report.ai_report and isinstance(report.ai_report, dict):
        summary = report.ai_report.get("executive_summary", "")
        if summary:
            md.append(summary)
        else:
            md.append("*No executive summary available.*")
    else:
        md.append("*No executive summary available.*")
    md.append("")
    
    # Findings Summary Box
    md.append("---")
    md.append("")
    md.append("## 🎯 Findings Overview")
    md.append("")
    
    total_findings = (
        (report.findings_critical or 0) + 
        (report.findings_high or 0) + 
        (report.findings_medium or 0) + 
        (report.findings_low or 0) + 
        (report.findings_info or 0)
    )
    
    md.append(f"**Total Findings:** {total_findings}")
    md.append("")
    md.append("| Severity | Count |")
    md.append("|----------|-------|")
    md.append(f"| 🔴 Critical | {report.findings_critical or 0} |")
    md.append(f"| 🟠 High | {report.findings_high or 0} |")
    md.append(f"| 🟡 Medium | {report.findings_medium or 0} |")
    md.append(f"| 🔵 Low | {report.findings_low or 0} |")
    md.append(f"| ⚪ Info | {report.findings_info or 0} |")
    md.append("")
    
    # Test Statistics
    md.append("### 📊 Test Statistics")
    md.append("")
    md.append(f"- **Total Iterations:** {report.total_iterations or 0}")
    md.append(f"- **Total Requests:** {report.total_requests or 0}")
    md.append(f"- **Duplicates Filtered:** {report.duplicates_filtered or 0}")
    md.append("")
    
    # Detailed Findings
    md.append("---")
    md.append("")
    md.append("## 🔍 Detailed Findings")
    md.append("")
    
    findings = report.findings or []
    if not findings:
        md.append("*No vulnerabilities discovered.*")
        md.append("")
    else:
        # Group by severity
        severity_order = ["critical", "high", "medium", "low", "info"]
        findings_by_severity: Dict[str, List] = {s: [] for s in severity_order}
        
        for finding in findings:
            sev = finding.get("severity", "info").lower()
            if sev in findings_by_severity:
                findings_by_severity[sev].append(finding)
            else:
                findings_by_severity["info"].append(finding)
        
        severity_icons = {
            "critical": "🔴",
            "high": "🟠",
            "medium": "🟡",
            "low": "🔵",
            "info": "⚪"
        }
        
        finding_num = 0
        for severity in severity_order:
            sev_findings = findings_by_severity[severity]
            if not sev_findings:
                continue
                
            md.append(f"### {severity_icons[severity]} {severity.upper()} Severity ({len(sev_findings)})")
            md.append("")
            
            for finding in sev_findings:
                finding_num += 1
                md.append(f"#### Finding #{finding_num}: {finding.get('title', 'Untitled')}")
                md.append("")
                
                # Technique
                if finding.get("technique"):
                    md.append(f"**Technique:** `{finding['technique']}`")
                    md.append("")
                
                # Description
                if finding.get("description"):
                    md.append(f"**Description:**")
                    md.append(finding["description"])
                    md.append("")
                
                # Endpoint
                if finding.get("endpoint"):
                    md.append(f"**Affected Endpoint:** `{finding['endpoint']}`")
                
                # Parameter
                if finding.get("parameter"):
                    md.append(f"**Parameter:** `{finding['parameter']}`")
                md.append("")
                
                # CVSS Score
                if finding.get("cvss_score"):
                    md.append(f"**CVSS Score:** {finding['cvss_score']}")
                    if finding.get("cvss_vector"):
                        md.append(f"**CVSS Vector:** `{finding['cvss_vector']}`")
                    md.append("")
                
                # CWE
                if finding.get("cwe_id"):
                    md.append(f"**CWE:** [{finding['cwe_id']}](https://cwe.mitre.org/data/definitions/{finding['cwe_id'].replace('CWE-', '')}.html)")
                    md.append("")
                
                # Payload
                if finding.get("payload"):
                    md.append("**Payload:**")
                    md.append("```")
                    md.append(finding["payload"])
                    md.append("```")
                    md.append("")
                
                # Evidence
                evidence = finding.get("evidence", [])
                if evidence:
                    md.append("**Evidence:**")
                    for ev in evidence[:5]:  # Limit to 5 items
                        md.append(f"- {ev}")
                    md.append("")
                
                # Proof of Concept
                if finding.get("proof_of_concept"):
                    md.append("**Proof of Concept:**")
                    md.append("```")
                    md.append(finding["proof_of_concept"])
                    md.append("```")
                    md.append("")
                
                # Recommendation
                if finding.get("recommendation"):
                    md.append("**Recommendation:**")
                    md.append(f"> {finding['recommendation']}")
                    md.append("")
                
                md.append("---")
                md.append("")
    
    # AI Analysis (from ai_report)
    if report.ai_report and isinstance(report.ai_report, dict):
        md.append("## 🤖 AI Security Analysis")
        md.append("")
        
        # Critical findings analysis
        if report.ai_report.get("critical_findings_analysis"):
            md.append("### Critical Findings Analysis")
            md.append("")
            analysis = report.ai_report["critical_findings_analysis"]
            if isinstance(analysis, str):
                md.append(analysis)
            elif isinstance(analysis, list):
                for item in analysis:
                    if isinstance(item, dict):
                        md.append(f"- **{item.get('title', 'Finding')}**: {item.get('analysis', '')}")
                    else:
                        md.append(f"- {item}")
            md.append("")
        
        # Risk assessment
        if report.ai_report.get("risk_assessment"):
            md.append("### Risk Assessment")
            md.append("")
            risk = report.ai_report["risk_assessment"]
            if isinstance(risk, str):
                md.append(risk)
            elif isinstance(risk, dict):
                if risk.get("overall_risk"):
                    md.append(f"**Overall Risk Level:** {risk['overall_risk']}")
                if risk.get("details"):
                    md.append(risk["details"])
            md.append("")
        
        # False positive assessment
        if report.ai_report.get("false_positive_assessment"):
            md.append("### False Positive Assessment")
            md.append("")
            fp = report.ai_report["false_positive_assessment"]
            if isinstance(fp, str):
                md.append(fp)
            elif isinstance(fp, dict):
                for key, value in fp.items():
                    md.append(f"- **{key}**: {value}")
            md.append("")
        
        # Remediation priorities
        if report.ai_report.get("remediation_priorities"):
            md.append("### Remediation Priorities")
            md.append("")
            remediation = report.ai_report["remediation_priorities"]
            if isinstance(remediation, str):
                md.append(remediation)
            elif isinstance(remediation, list):
                for i, item in enumerate(remediation, 1):
                    if isinstance(item, dict):
                        md.append(f"{i}. **{item.get('priority', f'Priority {i}')}**: {item.get('action', item.get('description', ''))}")
                    else:
                        md.append(f"{i}. {item}")
            md.append("")
        
        # Additional recommendations
        if report.ai_report.get("additional_testing_recommendations"):
            md.append("### Additional Testing Recommendations")
            md.append("")
            recs = report.ai_report["additional_testing_recommendations"]
            if isinstance(recs, str):
                md.append(recs)
            elif isinstance(recs, list):
                for rec in recs:
                    md.append(f"- {rec}")
            md.append("")
        
        # Compliance implications
        if report.ai_report.get("compliance_implications"):
            md.append("### Compliance Implications")
            md.append("")
            compliance = report.ai_report["compliance_implications"]
            if isinstance(compliance, str):
                md.append(compliance)
            elif isinstance(compliance, dict):
                for standard, implication in compliance.items():
                    md.append(f"- **{standard}**: {implication}")
            elif isinstance(compliance, list):
                for item in compliance:
                    md.append(f"- {item}")
            md.append("")
    
    # Correlation Analysis (Attack Chains)
    if report.correlation_analysis:
        md.append("---")
        md.append("")
        md.append("## 🔗 Correlation Analysis")
        md.append("")
        
        attack_chains = report.correlation_analysis.get("attack_chains", [])
        if attack_chains:
            md.append("### Attack Chains Discovered")
            md.append("")
            for chain in attack_chains:
                name = chain.get("name", "Unnamed Chain")
                md.append(f"#### {name}")
                md.append("")
                
                steps = chain.get("steps", [])
                if steps:
                    md.append("**Attack Path:**")
                    md.append(f"`{' → '.join(steps)}`")
                    md.append("")
                
                if chain.get("impact"):
                    md.append(f"**Impact:** {chain['impact']}")
                if chain.get("confidence"):
                    conf = chain["confidence"]
                    if isinstance(conf, float):
                        md.append(f"**Confidence:** {conf:.0%}")
                    else:
                        md.append(f"**Confidence:** {conf}")
                md.append("")
        
        root_causes = report.correlation_analysis.get("root_causes", [])
        if root_causes:
            md.append("### Root Cause Analysis")
            md.append("")
            for rc in root_causes:
                cause = rc.get("cause", "Unknown")
                affected = rc.get("affected_findings", [])
                md.append(f"- **{cause}**: Affects {len(affected)} finding(s)")
            md.append("")
    
    # Techniques Used
    if report.techniques_used:
        md.append("---")
        md.append("")
        md.append("## 🛠️ Techniques Used")
        md.append("")
        
        techniques = report.techniques_used
        if isinstance(techniques, list):
            # Group by category if possible
            for tech in techniques:
                if isinstance(tech, dict):
                    md.append(f"- **{tech.get('name', tech.get('id', 'Unknown'))}**: {tech.get('description', '')}")
                else:
                    md.append(f"- `{tech}`")
        md.append("")
    
    # Crawl Results
    if report.crawl_results:
        md.append("---")
        md.append("")
        md.append("## 🕸️ Crawl & Discovery Results")
        md.append("")
        
        stats = report.crawl_results.get("statistics", {})
        if stats:
            md.append("### Discovery Statistics")
            md.append("")
            md.append(f"- **URLs Crawled:** {stats.get('total_urls_crawled', 0)}")
            md.append(f"- **Parameters Found:** {stats.get('total_parameters', 0)}")
            md.append(f"- **Forms Found:** {stats.get('total_forms', 0)}")
            md.append("")
        
        auth_endpoints = report.crawl_results.get("auth_endpoints", [])
        if auth_endpoints:
            md.append("### Authentication Endpoints Discovered")
            md.append("")
            for ep in auth_endpoints[:10]:
                if isinstance(ep, dict):
                    md.append(f"- `{ep.get('url', ep)}`")
                else:
                    md.append(f"- `{ep}`")
            md.append("")
        
        api_endpoints = report.crawl_results.get("api_endpoints", [])
        if api_endpoints:
            md.append("### API Endpoints Discovered")
            md.append("")
            for ep in api_endpoints[:15]:
                if isinstance(ep, dict):
                    method = ep.get("method", "GET")
                    url = ep.get("url", ep)
                    md.append(f"- `{method}` `{url}`")
                else:
                    md.append(f"- `{ep}`")
            md.append("")
    
    # Engine Statistics
    if report.engine_stats:
        md.append("---")
        md.append("")
        md.append("## ⚙️ Engine Statistics")
        md.append("")
        
        for engine_name, stats in report.engine_stats.items():
            if not stats or (isinstance(stats, dict) and not stats.get("available", True)):
                continue
            
            md.append(f"### {engine_name.replace('_', ' ').title()}")
            md.append("")
            
            if isinstance(stats, dict):
                for key, value in stats.items():
                    if key != "available":
                        md.append(f"- **{key.replace('_', ' ').title()}:** {value}")
            md.append("")
    
    # Footer
    md.append("---")
    md.append("")
    md.append("*Report generated by VRAgent Agentic Fuzzer*")
    md.append(f"*Session: {report.session_id}*")
    
    return "\n".join(md)


# =============================================================================
# PDF GENERATION
# =============================================================================

def generate_fuzzer_report_pdf(report: AgenticFuzzerReport) -> BytesIO:
    """
    Generate a professional PDF report from an agentic fuzzer report.
    
    Args:
        report: The AgenticFuzzerReport database model
        
    Returns:
        BytesIO containing the PDF
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, ListFlowable, ListItem, HRFlowable
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    except ImportError:
        logger.error("ReportLab not installed. Cannot generate PDF.")
        raise ImportError("ReportLab is required for PDF generation. Install with: pip install reportlab")
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=letter,
        rightMargin=72, 
        leftMargin=72,
        topMargin=72, 
        bottomMargin=72
    )
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(
        name='ReportTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        textColor=colors.HexColor('#1a237e'),
        alignment=TA_CENTER
    ))
    
    styles.add(ParagraphStyle(
        name='SectionHeader',
        parent=styles['Heading2'],
        fontSize=16,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.HexColor('#1565c0'),
        borderPadding=5
    ))
    
    styles.add(ParagraphStyle(
        name='SubSection',
        parent=styles['Heading3'],
        fontSize=13,
        spaceBefore=15,
        spaceAfter=8,
        textColor=colors.HexColor('#424242')
    ))
    
    styles.add(ParagraphStyle(
        name='BodyText',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=8,
        alignment=TA_JUSTIFY
    ))
    
    styles.add(ParagraphStyle(
        name='FindingTitle',
        parent=styles['Heading4'],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6,
        textColor=colors.HexColor('#d32f2f')
    ))
    
    styles.add(ParagraphStyle(
        name='Code',
        parent=styles['Code'],
        fontSize=8,
        backColor=colors.HexColor('#f5f5f5'),
        borderPadding=5,
        leftIndent=10
    ))
    
    story = []
    
    # Title
    story.append(Paragraph("🔐 Security Assessment Report", styles['ReportTitle']))
    story.append(Spacer(1, 12))
    
    # Metadata table
    meta_data = [
        ["Target:", report.target_url],
        ["Session ID:", report.session_id],
        ["Scan Profile:", report.scan_profile or "Default"],
        ["Generated:", report.completed_at.strftime('%Y-%m-%d %H:%M:%S UTC') if report.completed_at else "N/A"],
    ]
    
    if report.duration_seconds:
        mins = int(report.duration_seconds // 60)
        secs = int(report.duration_seconds % 60)
        meta_data.append(["Duration:", f"{mins}m {secs}s"])
    
    meta_table = Table(meta_data, colWidths=[1.5*inch, 5*inch])
    meta_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 20))
    
    # HR
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')))
    story.append(Spacer(1, 15))
    
    # Executive Summary
    story.append(Paragraph("📋 Executive Summary", styles['SectionHeader']))
    
    summary_text = report.executive_summary
    if not summary_text and report.ai_report:
        summary_text = report.ai_report.get("executive_summary", "")
    
    if summary_text:
        story.append(Paragraph(str(summary_text), styles['BodyText']))
    else:
        story.append(Paragraph("<i>No executive summary available.</i>", styles['BodyText']))
    
    story.append(Spacer(1, 15))
    
    # Findings Summary Table
    story.append(Paragraph("🎯 Findings Overview", styles['SectionHeader']))
    
    total_findings = (
        (report.findings_critical or 0) + 
        (report.findings_high or 0) + 
        (report.findings_medium or 0) + 
        (report.findings_low or 0) + 
        (report.findings_info or 0)
    )
    
    findings_data = [
        ["Severity", "Count"],
        ["🔴 Critical", str(report.findings_critical or 0)],
        ["🟠 High", str(report.findings_high or 0)],
        ["🟡 Medium", str(report.findings_medium or 0)],
        ["🔵 Low", str(report.findings_low or 0)],
        ["⚪ Info", str(report.findings_info or 0)],
        ["Total", str(total_findings)],
    ]
    
    findings_table = Table(findings_data, colWidths=[2*inch, 1.5*inch])
    findings_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1565c0')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (1, 0), (1, -1), 'CENTER'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e0e0e0')),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#f5f5f5')),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
    ]))
    story.append(findings_table)
    story.append(Spacer(1, 15))
    
    # Test Statistics
    story.append(Paragraph("📊 Test Statistics", styles['SubSection']))
    stats_text = f"""
    <b>Total Iterations:</b> {report.total_iterations or 0}<br/>
    <b>Total Requests:</b> {report.total_requests or 0}<br/>
    <b>Duplicates Filtered:</b> {report.duplicates_filtered or 0}
    """
    story.append(Paragraph(stats_text, styles['BodyText']))
    
    # Page break before detailed findings
    story.append(PageBreak())
    
    # Detailed Findings
    story.append(Paragraph("🔍 Detailed Findings", styles['SectionHeader']))
    
    findings = report.findings or []
    if not findings:
        story.append(Paragraph("<i>No vulnerabilities discovered.</i>", styles['BodyText']))
    else:
        severity_order = ["critical", "high", "medium", "low", "info"]
        severity_colors = {
            "critical": colors.HexColor('#d32f2f'),
            "high": colors.HexColor('#ff5722'),
            "medium": colors.HexColor('#ff9800'),
            "low": colors.HexColor('#2196f3'),
            "info": colors.HexColor('#9e9e9e')
        }
        
        findings_by_severity: Dict[str, List] = {s: [] for s in severity_order}
        for finding in findings:
            sev = finding.get("severity", "info").lower()
            if sev in findings_by_severity:
                findings_by_severity[sev].append(finding)
            else:
                findings_by_severity["info"].append(finding)
        
        finding_num = 0
        for severity in severity_order:
            sev_findings = findings_by_severity[severity]
            if not sev_findings:
                continue
            
            # Section for this severity
            sev_style = ParagraphStyle(
                name=f'{severity}Header',
                parent=styles['SubSection'],
                textColor=severity_colors[severity]
            )
            story.append(Paragraph(f"{severity.upper()} Severity ({len(sev_findings)})", sev_style))
            
            for finding in sev_findings:
                finding_num += 1
                
                # Finding title
                story.append(Paragraph(
                    f"<b>Finding #{finding_num}:</b> {finding.get('title', 'Untitled')}",
                    styles['FindingTitle']
                ))
                
                # Details
                details = []
                if finding.get("technique"):
                    details.append(f"<b>Technique:</b> {finding['technique']}")
                if finding.get("endpoint"):
                    details.append(f"<b>Endpoint:</b> {finding['endpoint']}")
                if finding.get("parameter"):
                    details.append(f"<b>Parameter:</b> {finding['parameter']}")
                if finding.get("cvss_score"):
                    details.append(f"<b>CVSS Score:</b> {finding['cvss_score']}")
                if finding.get("cwe_id"):
                    details.append(f"<b>CWE:</b> {finding['cwe_id']}")
                
                if details:
                    story.append(Paragraph("<br/>".join(details), styles['BodyText']))
                
                if finding.get("description"):
                    story.append(Paragraph(f"<b>Description:</b> {finding['description']}", styles['BodyText']))
                
                if finding.get("payload"):
                    story.append(Paragraph("<b>Payload:</b>", styles['BodyText']))
                    story.append(Paragraph(finding["payload"][:500], styles['Code']))
                
                if finding.get("recommendation"):
                    story.append(Paragraph(f"<b>Recommendation:</b> {finding['recommendation']}", styles['BodyText']))
                
                story.append(Spacer(1, 10))
                story.append(HRFlowable(width="80%", thickness=0.5, color=colors.HexColor('#e0e0e0')))
                story.append(Spacer(1, 10))
    
    # AI Analysis
    if report.ai_report and isinstance(report.ai_report, dict):
        story.append(PageBreak())
        story.append(Paragraph("🤖 AI Security Analysis", styles['SectionHeader']))
        
        if report.ai_report.get("risk_assessment"):
            story.append(Paragraph("Risk Assessment", styles['SubSection']))
            risk = report.ai_report["risk_assessment"]
            if isinstance(risk, str):
                story.append(Paragraph(risk, styles['BodyText']))
            elif isinstance(risk, dict) and risk.get("details"):
                story.append(Paragraph(str(risk["details"]), styles['BodyText']))
        
        if report.ai_report.get("remediation_priorities"):
            story.append(Paragraph("Remediation Priorities", styles['SubSection']))
            remediation = report.ai_report["remediation_priorities"]
            if isinstance(remediation, str):
                story.append(Paragraph(remediation, styles['BodyText']))
            elif isinstance(remediation, list):
                for i, item in enumerate(remediation[:10], 1):
                    if isinstance(item, dict):
                        text = f"{i}. <b>{item.get('priority', '')}:</b> {item.get('action', item.get('description', ''))}"
                    else:
                        text = f"{i}. {item}"
                    story.append(Paragraph(text, styles['BodyText']))
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e0e0e0')))
    story.append(Paragraph(
        f"<i>Report generated by VRAgent Agentic Fuzzer | Session: {report.session_id}</i>",
        ParagraphStyle(name='Footer', parent=styles['Normal'], fontSize=8, textColor=colors.gray, alignment=TA_CENTER)
    ))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer


# =============================================================================
# WORD (DOCX) GENERATION
# =============================================================================

def generate_fuzzer_report_docx(report: AgenticFuzzerReport) -> BytesIO:
    """
    Generate a professional Word document from an agentic fuzzer report.
    
    Args:
        report: The AgenticFuzzerReport database model
        
    Returns:
        BytesIO containing the DOCX
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.error("python-docx not installed. Cannot generate Word document.")
        raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('🔐 Security Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata paragraph
    meta = doc.add_paragraph()
    meta.add_run('Target: ').bold = True
    meta.add_run(f'{report.target_url}\n')
    meta.add_run('Session ID: ').bold = True
    meta.add_run(f'{report.session_id}\n')
    meta.add_run('Scan Profile: ').bold = True
    meta.add_run(f'{report.scan_profile or "Default"}\n')
    meta.add_run('Generated: ').bold = True
    meta.add_run(f'{report.completed_at.strftime("%Y-%m-%d %H:%M:%S UTC") if report.completed_at else "N/A"}\n')
    
    if report.duration_seconds:
        mins = int(report.duration_seconds // 60)
        secs = int(report.duration_seconds % 60)
        meta.add_run('Duration: ').bold = True
        meta.add_run(f'{mins}m {secs}s')
    
    doc.add_paragraph()  # Spacer
    
    # Executive Summary
    doc.add_heading('📋 Executive Summary', level=1)
    
    summary_text = report.executive_summary
    if not summary_text and report.ai_report:
        summary_text = report.ai_report.get("executive_summary", "")
    
    if summary_text:
        doc.add_paragraph(str(summary_text))
    else:
        doc.add_paragraph('No executive summary available.', style='Intense Quote')
    
    # Findings Overview
    doc.add_heading('🎯 Findings Overview', level=1)
    
    total_findings = (
        (report.findings_critical or 0) + 
        (report.findings_high or 0) + 
        (report.findings_medium or 0) + 
        (report.findings_low or 0) + 
        (report.findings_info or 0)
    )
    
    # Findings table
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Severity'
    header_cells[1].text = 'Count'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    severity_data = [
        ('🔴 Critical', report.findings_critical or 0),
        ('🟠 High', report.findings_high or 0),
        ('🟡 Medium', report.findings_medium or 0),
        ('🔵 Low', report.findings_low or 0),
        ('⚪ Info', report.findings_info or 0),
        ('Total', total_findings),
    ]
    
    for i, (severity, count) in enumerate(severity_data, 1):
        table.rows[i].cells[0].text = severity
        table.rows[i].cells[1].text = str(count)
        if i == len(severity_data):  # Bold the total row
            for cell in table.rows[i].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph()  # Spacer
    
    # Test Statistics
    doc.add_heading('📊 Test Statistics', level=2)
    stats_para = doc.add_paragraph()
    stats_para.add_run('Total Iterations: ').bold = True
    stats_para.add_run(f'{report.total_iterations or 0}\n')
    stats_para.add_run('Total Requests: ').bold = True
    stats_para.add_run(f'{report.total_requests or 0}\n')
    stats_para.add_run('Duplicates Filtered: ').bold = True
    stats_para.add_run(f'{report.duplicates_filtered or 0}')
    
    # Detailed Findings
    doc.add_page_break()
    doc.add_heading('🔍 Detailed Findings', level=1)
    
    findings = report.findings or []
    if not findings:
        doc.add_paragraph('No vulnerabilities discovered.', style='Intense Quote')
    else:
        severity_order = ["critical", "high", "medium", "low", "info"]
        severity_icons = {
            "critical": "🔴",
            "high": "🟠",
            "medium": "🟡",
            "low": "🔵",
            "info": "⚪"
        }
        
        findings_by_severity: Dict[str, List] = {s: [] for s in severity_order}
        for finding in findings:
            sev = finding.get("severity", "info").lower()
            if sev in findings_by_severity:
                findings_by_severity[sev].append(finding)
            else:
                findings_by_severity["info"].append(finding)
        
        finding_num = 0
        for severity in severity_order:
            sev_findings = findings_by_severity[severity]
            if not sev_findings:
                continue
            
            doc.add_heading(f'{severity_icons[severity]} {severity.upper()} Severity ({len(sev_findings)})', level=2)
            
            for finding in sev_findings:
                finding_num += 1
                
                # Finding heading
                h = doc.add_heading(f'Finding #{finding_num}: {finding.get("title", "Untitled")}', level=3)
                
                # Details
                details = doc.add_paragraph()
                
                if finding.get("technique"):
                    details.add_run('Technique: ').bold = True
                    details.add_run(f'{finding["technique"]}\n')
                
                if finding.get("endpoint"):
                    details.add_run('Endpoint: ').bold = True
                    details.add_run(f'{finding["endpoint"]}\n')
                
                if finding.get("parameter"):
                    details.add_run('Parameter: ').bold = True
                    details.add_run(f'{finding["parameter"]}\n')
                
                if finding.get("cvss_score"):
                    details.add_run('CVSS Score: ').bold = True
                    details.add_run(f'{finding["cvss_score"]}\n')
                
                if finding.get("cwe_id"):
                    details.add_run('CWE: ').bold = True
                    details.add_run(f'{finding["cwe_id"]}\n')
                
                if finding.get("description"):
                    desc = doc.add_paragraph()
                    desc.add_run('Description: ').bold = True
                    desc.add_run(finding["description"])
                
                if finding.get("payload"):
                    doc.add_paragraph().add_run('Payload:').bold = True
                    # Add as code block style
                    code_para = doc.add_paragraph(finding["payload"][:500])
                    code_para.style = 'Quote'
                
                if finding.get("recommendation"):
                    rec = doc.add_paragraph()
                    rec.add_run('Recommendation: ').bold = True
                    rec.add_run(finding["recommendation"])
                
                doc.add_paragraph()  # Spacer
    
    # AI Analysis
    if report.ai_report and isinstance(report.ai_report, dict):
        doc.add_page_break()
        doc.add_heading('🤖 AI Security Analysis', level=1)
        
        if report.ai_report.get("risk_assessment"):
            doc.add_heading('Risk Assessment', level=2)
            risk = report.ai_report["risk_assessment"]
            if isinstance(risk, str):
                doc.add_paragraph(risk)
            elif isinstance(risk, dict) and risk.get("details"):
                doc.add_paragraph(str(risk["details"]))
        
        if report.ai_report.get("remediation_priorities"):
            doc.add_heading('Remediation Priorities', level=2)
            remediation = report.ai_report["remediation_priorities"]
            if isinstance(remediation, str):
                doc.add_paragraph(remediation)
            elif isinstance(remediation, list):
                for i, item in enumerate(remediation[:10], 1):
                    p = doc.add_paragraph(style='List Number')
                    if isinstance(item, dict):
                        p.add_run(f"{item.get('priority', '')}: ").bold = True
                        p.add_run(item.get('action', item.get('description', '')))
                    else:
                        p.add_run(str(item))
    
    # Correlation Analysis
    if report.correlation_analysis:
        attack_chains = report.correlation_analysis.get("attack_chains", [])
        if attack_chains:
            doc.add_heading('🔗 Attack Chains Discovered', level=1)
            for chain in attack_chains:
                name = chain.get("name", "Unnamed Chain")
                doc.add_heading(name, level=2)
                
                steps = chain.get("steps", [])
                if steps:
                    p = doc.add_paragraph()
                    p.add_run('Attack Path: ').bold = True
                    p.add_run(' → '.join(steps))
                
                if chain.get("impact"):
                    p = doc.add_paragraph()
                    p.add_run('Impact: ').bold = True
                    p.add_run(chain["impact"])
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f'Report generated by VRAgent Agentic Fuzzer | Session: {report.session_id}').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def export_fuzzer_report(report: AgenticFuzzerReport, format: str) -> tuple[BytesIO, str, str]:
    """
    Export a fuzzer report to the specified format.
    
    Args:
        report: The AgenticFuzzerReport database model
        format: Export format ('markdown', 'pdf', 'docx')
        
    Returns:
        Tuple of (content_buffer, filename, content_type)
    """
    safe_name = report.session_id.replace("/", "_").replace("\\", "_")
    
    if format.lower() == "markdown" or format.lower() == "md":
        content = generate_fuzzer_report_markdown(report)
        buffer = BytesIO(content.encode('utf-8'))
        return buffer, f"fuzzer_report_{safe_name}.md", "text/markdown; charset=utf-8"
    
    elif format.lower() == "pdf":
        buffer = generate_fuzzer_report_pdf(report)
        return buffer, f"fuzzer_report_{safe_name}.pdf", "application/pdf"
    
    elif format.lower() == "docx" or format.lower() == "word":
        buffer = generate_fuzzer_report_docx(report)
        return buffer, f"fuzzer_report_{safe_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    else:
        raise ValueError(f"Unsupported format: {format}. Use 'markdown', 'pdf', or 'docx'.")
